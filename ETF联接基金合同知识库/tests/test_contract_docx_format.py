import io
import unittest
import xml.etree.ElementTree as ET
import zipfile

from docx import Document
from docx.oxml.ns import qn

from app import ContractEngine


def _sample_contract_text():
    return "\n".join(
        [
            "测试ETF联接基金基金合同",
            "基金管理人：南方基金管理股份有限公司",
            "基金托管人：测试银行股份有限公司",
            "目 录",
            "第一部分  前言\t1",
            "第一部分  前言",
            "一、订立本基金合同的目的、依据和原则",
            "1、订立本基金合同的目的是保护投资人合法权益，明确基金合同当事人的权利义务，规范基金运作。",
        ]
    )


class ContractDocxFormatTests(unittest.TestCase):
    def _docx_bytes(self):
        return ContractEngine().build_docx(_sample_contract_text())

    def _xml_map(self, docx_bytes):
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            return {
                name: zf.read(name).decode("utf-8", errors="ignore")
                for name in zf.namelist()
                if name in {"word/document.xml", "word/settings.xml"}
                or name.startswith("word/header")
                or name.startswith("word/footer")
            }

    def test_contract_docx_uses_reference_section_grid_and_cover_spacing(self):
        doc = Document(io.BytesIO(self._docx_bytes()))
        line_pitches = [
            [
                grid.get(qn("w:linePitch"))
                for grid in section._sectPr.findall(qn("w:docGrid"))
            ]
            for section in doc.sections
        ]
        title_index = next(
            idx for idx, paragraph in enumerate(doc.paragraphs)
            if paragraph.text.strip().endswith("基金合同（草案）")
        )
        manager_index = next(
            idx for idx, paragraph in enumerate(doc.paragraphs)
            if paragraph.text.strip().startswith("基金管理人：")
        )
        blank_count = sum(
            1
            for paragraph in doc.paragraphs[title_index + 1 : manager_index]
            if not paragraph.text.strip()
        )

        self.assertEqual([["312"], ["312"]], line_pitches)
        self.assertEqual(6, blank_count)

    def test_contract_docx_body_paragraph_uses_farming_reference_base_style(self):
        doc = Document(io.BytesIO(self._docx_bytes()))
        paragraph = next(
            p for p in doc.paragraphs
            if "订立本基金合同的目的是保护投资人合法权益" in p.text
        )
        pPr = paragraph._p.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        indent = pPr.find(qn("w:ind"))
        run_properties = paragraph.runs[0]._r.get_or_add_rPr()
        size = run_properties.find(qn("w:sz"))

        self.assertEqual("360", spacing.get(qn("w:line")))
        self.assertEqual("auto", spacing.get(qn("w:lineRule")))
        self.assertEqual("0", spacing.get(qn("w:after")))
        self.assertEqual("480", indent.get(qn("w:firstLine")))
        self.assertIsNone(indent.get(qn("w:firstLineChars")))
        self.assertEqual("24", size.get(qn("w:val")))

    def test_contract_docx_follows_etf_header_footer_and_toc_behavior(self):
        xml_map = self._xml_map(self._docx_bytes())
        settings_xml = xml_map["word/settings.xml"]
        header_xml = "".join(v for k, v in xml_map.items() if k.startswith("word/header"))
        footer_xml = "".join(v for k, v in xml_map.items() if k.startswith("word/footer"))
        document_root = ET.fromstring(xml_map["word/document.xml"])
        settings_root = ET.fromstring(settings_xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        pg_num_types = document_root.findall(".//w:pgNumType", ns)
        default_tab_stop = settings_root.find(".//w:defaultTabStop", ns)

        self.assertIn('TOC \\\\o "1-1" \\\\h \\\\z \\\\u', xml_map["word/document.xml"])
        self.assertIn("updateFields", settings_xml)
        self.assertIn("基金合同（草案）", header_xml)
        self.assertIn("PAGE", footer_xml)
        self.assertTrue(any(node.get(qn("w:start")) == "1" for node in pg_num_types))
        self.assertIsNotNone(default_tab_stop)
        self.assertEqual("720", default_tab_stop.get(qn("w:val")))


if __name__ == "__main__":
    unittest.main()
