import io
import unittest
import zipfile

from docx import Document
from docx.oxml.ns import qn

from app import ProspectusEngine, _reinforce_prospectus_normal_text_format


def _sample_prospectus_text():
    return "\n".join(
        [
            "南方测试交易型开放式指数证券投资基金发起式联接基金招募说明书",
            "基金管理人：南方基金管理股份有限公司",
            "基金托管人：招商银行股份有限公司",
            "重要提示",
            "这是重要提示正文段落。",
            "目录",
            "一、绪言",
            "二、释义",
            "一、绪言",
            "这是绪言正文第一段。",
            "二、释义",
            "这是释义正文第一段。",
        ]
    )


class ProspectusDocxFormatTests(unittest.TestCase):
    def _docx_bytes(self):
        return ProspectusEngine().build_docx(_sample_prospectus_text())

    def _xml_map(self, docx_bytes):
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            return {
                name: zf.read(name).decode("utf-8", errors="ignore")
                for name in zf.namelist()
                if name in {"word/document.xml", "word/settings.xml", "word/styles.xml"}
                or name.startswith("word/header")
                or name.startswith("word/footer")
            }

    def test_prospectus_docx_uses_reference_page_setup_and_header(self):
        docx_bytes = self._docx_bytes()
        doc = Document(io.BytesIO(docx_bytes))
        xml_map = self._xml_map(docx_bytes)
        header_xml = "".join(value for key, value in xml_map.items() if key.startswith("word/header"))

        self.assertEqual(2, len(doc.sections))
        self.assertIn('w:defaultTabStop w:val="420"', xml_map["word/settings.xml"])
        self.assertIn("《南方测试交易型开放式指数证券投资基金发起式联接基金招募说明书（草案）》", header_xml)
        self.assertNotIn("基金合同", header_xml)

    def test_prospectus_docx_uses_reference_cover_body_and_notice_styles(self):
        doc = Document(io.BytesIO(self._docx_bytes()))
        title = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip().endswith("招募说明书"))
        manager = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith("基金管理人："))
        notice_title = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "重要提示")
        notice_body = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "这是重要提示正文段落。")
        body = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "这是绪言正文第一段。")

        self.assertEqual("44", title.runs[0]._r.get_or_add_rPr().find(qn("w:sz")).get(qn("w:val")))
        self.assertEqual("30", manager.runs[0]._r.get_or_add_rPr().find(qn("w:sz")).get(qn("w:val")))
        self.assertTrue(manager.runs[0].bold)
        self.assertEqual("21", notice_title.runs[0]._r.get_or_add_rPr().find(qn("w:sz")).get(qn("w:val")))
        self.assertTrue(notice_title.runs[0].bold)

        for paragraph in (notice_body, body):
            ppr = paragraph._p.get_or_add_pPr()
            spacing = ppr.find(qn("w:spacing"))
            indent = ppr.find(qn("w:ind"))
            snap_to_grid = ppr.find(qn("w:snapToGrid"))
            run_properties = paragraph.runs[0]._r.get_or_add_rPr()

            self.assertEqual("360", spacing.get(qn("w:line")))
            self.assertEqual("auto", spacing.get(qn("w:lineRule")))
            self.assertEqual("420", indent.get(qn("w:firstLine")))
            self.assertEqual("200", indent.get(qn("w:firstLineChars")))
            self.assertEqual("0", snap_to_grid.get(qn("w:val")))
            self.assertEqual("21", run_properties.find(qn("w:sz")).get(qn("w:val")))

    def test_reinforces_normal_text_format_after_word_strips_direct_xml(self):
        raw_doc = Document()
        notice_title = raw_doc.add_paragraph()
        notice_title.alignment = 1
        notice_title.add_run("重要提示").bold = True
        notice_body = raw_doc.add_paragraph("提示正文")
        source = io.BytesIO()
        raw_doc.save(source)

        reinforced = Document(io.BytesIO(_reinforce_prospectus_normal_text_format(source.getvalue())))
        reinforced_title = next(p for p in reinforced.paragraphs if p.text == "重要提示")
        reinforced_body = next(p for p in reinforced.paragraphs if p.text == "提示正文")
        body_spacing = reinforced_body._p.get_or_add_pPr().find(qn("w:spacing"))

        self.assertEqual("21", reinforced_title.runs[0]._r.get_or_add_rPr().find(qn("w:sz")).get(qn("w:val")))
        self.assertEqual("21", reinforced_body.runs[0]._r.get_or_add_rPr().find(qn("w:sz")).get(qn("w:val")))
        self.assertEqual("360", body_spacing.get(qn("w:line")))
        self.assertEqual("auto", body_spacing.get(qn("w:lineRule")))

    def test_custody_summary_internal_party_headings_do_not_enter_toc(self):
        text = "\n".join(
            [
                "南方测试交易型开放式指数证券投资基金发起式联接基金招募说明书",
                "基金管理人：南方基金管理股份有限公司",
                "基金托管人：招商银行股份有限公司",
                "目录",
                "一、绪言",
                "二、释义",
                "三、基金管理人",
                "四、基金托管人",
                "二十、基金托管协议的内容摘要",
                "二十一、基金份额持有人服务",
                "二十二、备查文件",
                "绪言",
                "这是绪言正文。",
                "二十、基金托管协议的内容摘要",
                "一、基金托管协议当事人",
                "基金管理人",
                "名称：南方基金管理股份有限公司",
                "基金托管人",
                "名称：招商银行股份有限公司",
                "二、基金托管人对基金管理人的业务监督和核查",
                "监督正文。",
                "二十一、基金份额持有人服务",
                "服务正文。",
                "二十二、备查文件",
                "一、备查文件清单",
            ]
        )

        doc = Document(io.BytesIO(ProspectusEngine().build_docx(text)))
        summary_seen = False
        internal_party_styles = []
        for paragraph in doc.paragraphs:
            stripped = paragraph.text.strip()
            if stripped == "基金托管协议的内容摘要":
                summary_seen = True
                continue
            if stripped == "备查文件":
                break
            if summary_seen and stripped in {"基金管理人", "基金托管人"}:
                internal_party_styles.append(paragraph.style.name)

        self.assertEqual(["Normal", "Normal"], internal_party_styles)

    def test_render_model_keeps_custody_summary_numbered_headings_inside_chapter(self):
        text = "\n".join(
            [
                "南方测试交易型开放式指数证券投资基金发起式联接基金招募说明书",
                "基金管理人：南方基金管理股份有限公司",
                "基金托管人：招商银行股份有限公司",
                "目录",
                "一、绪言",
                "二十、基金托管协议的内容摘要",
                "二十一、基金份额持有人服务",
                "一、绪言",
                "这是绪言正文。",
                "二十、基金托管协议的内容摘要",
                "一、基金托管协议当事人",
                "当事人正文。",
                "二、基金托管人对基金管理人的业务监督和核查",
                "监督正文。",
                "二十一、基金份额持有人服务",
                "服务正文。",
            ]
        )

        model = ProspectusEngine()._build_render_model_from_text(text)
        chapter_titles = [chapter["title"] for chapter in model["chapters"]]
        summary = next(chapter for chapter in model["chapters"] if chapter["title"] == "基金托管协议的内容摘要")

        self.assertEqual(["绪言", "基金托管协议的内容摘要", "基金份额持有人服务"], chapter_titles)
        self.assertEqual(
            [
                "一、基金托管协议当事人",
                "当事人正文。",
                "二、基金托管人对基金管理人的业务监督和核查",
                "监督正文。",
            ],
            [block["text"] for block in summary["blocks"]],
        )


if __name__ == "__main__":
    unittest.main()
