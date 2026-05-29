from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from app import _docx_document_to_text, _extract_custody_agreement_summary_from_docx, _extract_custody_agreement_summary_from_text  # noqa: E402


def _append_parenthesized_chinese_numbering(document, num_id="99"):
    numbering = document.part.numbering_part.element
    abstract_id = "99"
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (
        ("w:start", "1"),
        ("w:numFmt", "japaneseCounting"),
        ("w:lvlText", "（%1）"),
    ):
        child = OxmlElement(tag)
        child.set(qn("w:val"), value)
        level.append(child)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering.append(num)


def _set_paragraph_numbering(paragraph, num_id="99", ilvl="0"):
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    paragraph._p.get_or_add_pPr().append(num_pr)


def _add_tracked_change_paragraph(document, change_type: str, text: str):
    paragraph = document.add_paragraph()
    change = OxmlElement(f"w:{change_type}")
    change.set(qn("w:id"), "1")
    change.set(qn("w:author"), "审核工作台")
    change.set(qn("w:date"), "2026-05-18T00:00:00Z")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:delText" if change_type == "del" else "w:t")
    text_node.text = text
    run.append(text_node)
    change.append(run)
    paragraph._p.append(change)
    return paragraph


def test_custody_summary_extracts_non_contiguous_source_sections():
    text = "\n".join(
        [
            "一、基金托管协议当事人",
            "当事人正文。",
            "二、基金托管协议的依据、目的和原则",
            "依据正文不应进入当事人摘要。",
            "三、基金托管人对基金管理人的业务监督和核查",
            "监督核查正文。",
            "四、基金管理人对基金托管人的业务核查",
            "业务核查正文。",
            "五、基金财产的保管",
            "保管正文。",
            "六、指令的发送、确认及执行",
            "指令正文不应进入摘要。",
            "八、基金资产净值计算和会计核算",
            "净值正文。",
            "十二、基金份额持有人名册的登记与保管",
            "名册正文。",
            "十六、基金托管协议的变更、终止与基金财产的清算",
            "变更清算正文。",
            "十八、争议解决方式",
            "争议正文。",
            "本协议受中国法律管辖。",
            "十九、托管协议的效力",
            "效力正文不应进入摘要。",
        ]
    )

    result = _extract_custody_agreement_summary_from_text(text)

    assert result["source_type"] == "custody_agreement"
    assert result["missing_sections"] == []
    assert "一、基金托管协议当事人\n当事人正文。" in result["summary_text"]
    assert "依据正文不应进入当事人摘要" not in result["summary_text"]
    assert "三、基金管理人对基金托管人的业务监督和核查\n业务核查正文。" in result["summary_text"]
    assert "五、基金资产净值计算、估值和会计核算\n净值正文。" in result["summary_text"]
    assert "六、基金份额持有人名册的保管\n名册正文。" in result["summary_text"]
    assert "七、争议解决方式\n争议正文。\n本协议受中国法律管辖。" in result["summary_text"]
    assert "八、托管协议的变更、终止与基金财产的清算\n变更清算正文。" in result["summary_text"]


def test_custody_summary_matches_complete_source_chapters_by_title_not_number():
    text = "\n".join(
        [
            "第一章 前言",
            "前言正文不应进入摘要。",
            "第二章 基金托管协议当事人",
            "当事人整章第一段。",
            "当事人整章第二段。",
            "第三章 基金托管人对基金管理人的监督和核查",
            "托管人监督整章。",
            "第四章 基金管理人对基金托管人的业务核查",
            "管理人核查整章。",
            "第五章 资产保管",
            "保管整章。",
            "第六章 指令的发送、确认及执行",
            "指令正文不应进入摘要。",
            "第八章 基金资产净值计算、估值与会计核算",
            "估值核算整章。",
            "第十二章 基金份额持有人名册登记与保管",
            "名册整章。",
            "第十七章 托管协议的变更、终止和基金财产的清算",
            "变更终止清算整章。",
            "第十八章 适用法律和争议解决方式",
            "争议解决整章。",
            "第十九章 托管协议的效力",
            "效力正文不应进入摘要。",
        ]
    )

    result = _extract_custody_agreement_summary_from_text(text)

    assert result["source_type"] == "custody_agreement"
    assert result["missing_sections"] == []
    assert "一、基金托管协议当事人\n当事人整章第一段。\n当事人整章第二段。" in result["summary_text"]
    assert "二、基金托管人对基金管理人的业务监督和核查\n托管人监督整章。" in result["summary_text"]
    assert "三、基金管理人对基金托管人的业务监督和核查\n管理人核查整章。" in result["summary_text"]
    assert "四、基金财产的保管\n保管整章。" in result["summary_text"]
    assert "五、基金资产净值计算、估值和会计核算\n估值核算整章。" in result["summary_text"]
    assert "六、基金份额持有人名册的保管\n名册整章。" in result["summary_text"]
    assert "七、争议解决方式\n争议解决整章。" in result["summary_text"]
    assert "八、托管协议的变更、终止与基金财产的清算\n变更终止清算整章。" in result["summary_text"]
    assert "前言正文不应进入摘要" not in result["summary_text"]
    assert "指令正文不应进入摘要" not in result["summary_text"]
    assert "效力正文不应进入摘要" not in result["summary_text"]


def test_custody_summary_matches_source_chapters_by_keywords_when_titles_differ():
    text = "\n".join(
        [
            "第一章 基金管理人与基金托管人",
            "当事人关键词正文。",
            "第二章 托管人对管理人的投资监督",
            "托管人监督关键词正文。",
            "第三章 管理人对托管人的履职核查",
            "管理人核查关键词正文。",
            "第四章 资产保管安排",
            "保管关键词正文。",
            "第五章 估值和会计核算",
            "估值核算关键词正文。",
            "第六章 名册登记保管",
            "名册关键词正文。",
            "第七章 法律适用及争议处理",
            "争议关键词正文。",
            "第八章 协议变更终止及财产清算",
            "清算关键词正文。",
        ]
    )

    result = _extract_custody_agreement_summary_from_text(text)

    assert result["missing_sections"] == []
    assert "一、基金托管协议当事人\n当事人关键词正文。" in result["summary_text"]
    assert "二、基金托管人对基金管理人的业务监督和核查\n托管人监督关键词正文。" in result["summary_text"]
    assert "三、基金管理人对基金托管人的业务监督和核查\n管理人核查关键词正文。" in result["summary_text"]
    assert "四、基金财产的保管\n保管关键词正文。" in result["summary_text"]
    assert "五、基金资产净值计算、估值和会计核算\n估值核算关键词正文。" in result["summary_text"]
    assert "六、基金份额持有人名册的保管\n名册关键词正文。" in result["summary_text"]
    assert "七、争议解决方式\n争议关键词正文。" in result["summary_text"]
    assert "八、托管协议的变更、终止与基金财产的清算\n清算关键词正文。" in result["summary_text"]


def test_custody_summary_preserves_word_auto_numbered_party_subheadings():
    doc = Document()
    _append_parenthesized_chinese_numbering(doc)
    doc.add_paragraph("一、基金托管协议当事人")
    manager = doc.add_paragraph("基金管理人")
    _set_paragraph_numbering(manager)
    doc.add_paragraph("名称：南方基金管理股份有限公司")
    custodian = doc.add_paragraph("基金托管人")
    _set_paragraph_numbering(custodian)
    doc.add_paragraph("名称：中信证券股份有限公司")
    doc.add_paragraph("二、基金托管人对基金管理人的业务监督和核查")
    doc.add_paragraph("监督正文。")

    plain_text = _docx_document_to_text(doc)
    assert "（一）基金管理人" in plain_text
    assert "（二）基金托管人" in plain_text

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    result = _extract_custody_agreement_summary_from_docx(buf, "托管协议.docx")

    assert "一、基金托管协议当事人\n（一）基金管理人" in result["summary_text"]
    assert "（二）基金托管人\n名称：中信证券股份有限公司" in result["summary_text"]


def test_custody_summary_docx_accepts_tracked_insertions_and_ignores_deletions():
    doc = Document()
    _add_tracked_change_paragraph(doc, "del", "第一章 被删除章节")
    doc.add_paragraph("被删除章节正文不应出现。")
    for heading, body in [
        ("第一章 基金托管协议当事人", "当事人修订正文。"),
        ("第二章 基金托管人对基金管理人的监督和核查", "托管人监督修订正文。"),
        ("第三章 基金管理人对基金托管人的业务核查", "管理人核查修订正文。"),
        ("第四章 基金财产保管", "保管修订正文。"),
        ("第五章 指令的发送、确认及执行", "指令正文不应进入摘要。"),
        ("第六章 基金资产净值计算和会计核算", "净值修订正文。"),
        ("第七章 基金份额持有人名册的保管", "名册修订正文。"),
        ("第八章 适用法律与争议解决方式", "争议修订正文。"),
        ("第九章 托管协议的变更、终止与基金财产的清算", "清算修订正文。"),
    ]:
        _add_tracked_change_paragraph(doc, "ins", heading)
        doc.add_paragraph(body)

    plain_text = _docx_document_to_text(doc)
    assert "第一章 基金托管协议当事人" in plain_text
    assert "第一章 被删除章节" not in plain_text

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    result = _extract_custody_agreement_summary_from_docx(buf, "修订版托管协议.docx")

    assert result["missing_sections"] == []
    assert "一、基金托管协议当事人\n当事人修订正文。" in result["summary_text"]
    assert "二、基金托管人对基金管理人的业务监督和核查\n托管人监督修订正文。" in result["summary_text"]
    assert "三、基金管理人对基金托管人的业务监督和核查\n管理人核查修订正文。" in result["summary_text"]
    assert "五、基金资产净值计算、估值和会计核算\n净值修订正文。" in result["summary_text"]
    assert "七、争议解决方式\n争议修订正文。" in result["summary_text"]
    assert "八、托管协议的变更、终止与基金财产的清算\n清算修订正文。" in result["summary_text"]
    assert "指令正文不应进入摘要" not in result["summary_text"]
    assert "被删除章节" not in result["summary_text"]
