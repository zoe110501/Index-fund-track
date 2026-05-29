"""
ETF联接基金合同知识库 Web 应用
Flask 后端服务器 + 合同生成引擎
"""

import copy
import base64
import difflib
import html
import io
import json
import logging
import os
import re
import secrets
import threading
import time
import unicodedata
import webbrowser
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from flask import Flask, Response, abort, after_this_request, jsonify, render_template, request, send_file, stream_with_context
from werkzeug.exceptions import BadRequest, HTTPException, RequestEntityTooLarge
from jinja2 import Environment
from review_builtin_rules import (
    BUILTIN_LINKED_FUND_CROSS_RULES,
    BUILTIN_LINKED_FUND_SUMMARY_RULES,
    CROSS_RULE_FIELDS,
    SUMMARY_RULE_FIELDS,
)

BASE_DIR = Path(__file__).parent
TEMPLATE_MD = BASE_DIR / "01_通用模板.md"
SCHEMA_JSON = BASE_DIR / "02_变量定义表.json"
DIFF_TABLE_MD = BASE_DIR / "03_差异条款匹配表.md"
CLAUSES_JSON = BASE_DIR / "04_差异条款原文库.json"
ENTRY_TABLE_MD = BASE_DIR / "05_要素录入表.md"
PROSPECTUS_TEMPLATE_MD = BASE_DIR / "07_招募说明书模板.md"
PROSPECTUS_CLAUSES_JSON = BASE_DIR / "08_招募说明书差异条款库.json"
ALLOWED_SUFFIXES = {".md", ".json"}
EDITABLE_FILE_NAMES = {
    "01_通用模板.md",
    "02_变量定义表.json",
    "04_差异条款原文库.json",
    "07_招募说明书模板.md",
    "08_招募说明书差异条款库.json",
}
MAX_DOCX_ENTRIES = 2500
MAX_DOCX_UNCOMPRESSED_BYTES = 80 * 1024 * 1024
REVIEW_JOB_TTL_SECONDS = 2 * 60 * 60


def _editor_group_for_filename(filename: str) -> str:
    if "模板" in filename:
        return "模板"
    if "变量定义" in filename:
        return "变量库"
    if "条款" in filename:
        return "条款库"
    return "规则库"


CLAUSE_LIBRARY_SOURCE_CANDIDATES = {
    "contract": ("04_差异条款原文库.json",),
    "prospectus": ("08_招募说明书差异条款库.json",),
}
CLAUSE_LIBRARY_DOCUMENT_LABELS = {
    "contract": "基金合同",
    "prospectus": "招募说明书",
}
CLAUSE_LIBRARY_METADATA_FIELDS = {
    "applicable_to",
    "condition",
    "description",
    "legacy_source",
    "location",
    "number_label",
    "preservation_rule",
    "source_sample",
    "title",
}
CLAUSE_LIBRARY_FIELD_LABELS = {
    "text": "正文",
    "investment_scope_text": "投资范围",
    "component_stock_strategy": "成份股策略",
    "risk_return_characteristics": "风险收益特征",
}
CLAUSE_LIBRARY_SCENE_LABELS = {
    "DEFAULT": "通用",
    "STANDARD": "标准",
    "HK_CONNECT": "港股通",
    "NON_HK_CONNECT": "非港股通",
    "SSE": "上交所",
    "SZSE": "深交所",
    "SSE_HK": "上交所港股通",
    "SZSE_HK": "深交所港股通",
    "SSE_CROSS": "上交所跨市场",
    "SSE_SINGLE": "上交所单市场",
    "SZSE_CROSS": "深交所跨市场",
    "SZSE_SINGLE": "深交所单市场",
    "INITIATING": "发起式基金",
}


def _resolve_base_dir_data_path(*names: str) -> Path:
    for name in names:
        candidate = BASE_DIR / name
        if candidate.exists():
            return candidate
    return BASE_DIR / names[0]


def _clause_library_sources() -> list[dict]:
    return [
        {
            "document_type": document_type,
            "document_label": CLAUSE_LIBRARY_DOCUMENT_LABELS[document_type],
            "path": _resolve_base_dir_data_path(*names),
        }
        for document_type, names in CLAUSE_LIBRARY_SOURCE_CANDIDATES.items()
    ]


def _encode_clause_path_id(document_type: str, path_parts: list[str]) -> str:
    payload = {"document_type": document_type, "path": path_parts}
    raw = json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
    return base64.urlsafe_b64encode(raw).decode("ascii").rstrip("=")


def _decode_clause_path_id(path_id: str) -> tuple[str, list[str]]:
    padded = str(path_id or "") + "=" * (-len(str(path_id or "")) % 4)
    try:
        payload = json.loads(base64.urlsafe_b64decode(padded.encode("ascii")).decode("utf-8"))
    except Exception as exc:
        raise ValueError("无效的条款路径") from exc
    document_type = str(payload.get("document_type") or "")
    path_parts = payload.get("path")
    if document_type not in CLAUSE_LIBRARY_SOURCE_CANDIDATES or not isinstance(path_parts, list):
        raise ValueError("无效的条款路径")
    return document_type, [str(part) for part in path_parts]


def _clause_library_source_for(document_type: str) -> dict:
    for source in _clause_library_sources():
        if source["document_type"] == document_type:
            return source
    raise ValueError("未知条款库文档类型")


def _clause_scene_label(scene_key: str) -> str:
    return CLAUSE_LIBRARY_SCENE_LABELS.get(scene_key, scene_key or "通用")


def _clause_field_label(field_key: str) -> str:
    if field_key in CLAUSE_LIBRARY_FIELD_LABELS:
        return CLAUSE_LIBRARY_FIELD_LABELS[field_key]
    return str(field_key or "").replace("_", " ").strip() or "正文"


def _clause_chapter_hint(*values: str) -> str:
    text = " ".join(str(value or "") for value in values)
    match = re.search(r"(第[一二三四五六七八九十百零\d]+(?:部分|章)[^，。；；/、]*)", text)
    return match.group(1).strip() if match else "未归类"


def _clause_source_note(*nodes: dict) -> str:
    parts = []
    for node in nodes:
        if not isinstance(node, dict):
            continue
        if node.get("source_sample"):
            parts.append(f"来源：{node.get('source_sample')}")
        legacy_source = node.get("legacy_source")
        if isinstance(legacy_source, dict):
            clause = legacy_source.get("clause") or ""
            variant = legacy_source.get("variant") or ""
            parts.append(f"引用：{clause}/{variant}".rstrip("/"))
        if node.get("preservation_rule"):
            parts.append(str(node.get("preservation_rule")))
    return "；".join(dict.fromkeys(part for part in parts if part))


def _is_editable_clause_field(field_key: str, value) -> bool:
    return isinstance(value, str) and field_key not in CLAUSE_LIBRARY_METADATA_FIELDS


def _clause_entry_title(clause_key: str, clause: dict, field_key: str, target_node: dict) -> str:
    description = str((clause or {}).get("description") or "").strip()
    child_title = str((target_node or {}).get("title") or "").strip()
    field_label = _clause_field_label(field_key)
    if child_title and field_key == "text":
        return child_title
    if description and field_key == "text":
        return description
    if description:
        return f"{description}｜{field_label}"
    return f"{clause_key}｜{field_label}"


def _build_clause_library_entry(
    *,
    source: dict,
    clause_key: str,
    clause: dict,
    path_parts: list[str],
    field_key: str,
    content: str,
    scene_key: str,
    target_node: dict | None = None,
    readonly: bool = False,
    readonly_reason: str = "",
) -> dict:
    target_node = target_node if isinstance(target_node, dict) else {}
    description = str((clause or {}).get("description") or "").strip()
    location = str((clause or {}).get("location") or target_node.get("location") or "").strip()
    condition = str(target_node.get("condition") or (clause or {}).get("condition") or "").strip()
    applicable_to = target_node.get("applicable_to") or (clause or {}).get("applicable_to") or []
    if isinstance(applicable_to, list) and applicable_to:
        condition = f"{condition}；适用：{'、'.join(str(item) for item in applicable_to)}".strip("；")
    applicability = "；".join(part for part in (_clause_scene_label(scene_key), condition) if part)
    return {
        "path_id": _encode_clause_path_id(source["document_type"], path_parts),
        "raw_path": ".".join(path_parts),
        "document_type": source["document_type"],
        "document_label": source["document_label"],
        "file_name": source["path"].name,
        "clause_key": clause_key,
        "title": _clause_entry_title(clause_key, clause, field_key, target_node),
        "chapter": _clause_chapter_hint(location, description, condition),
        "location": location,
        "scene_key": scene_key,
        "scene_label": _clause_scene_label(scene_key),
        "field_key": field_key,
        "field_label": _clause_field_label(field_key),
        "condition": condition,
        "applicability": applicability,
        "source_note": _clause_source_note(clause, target_node),
        "content": str(content or ""),
        "readonly": bool(readonly),
        "readonly_reason": readonly_reason,
        "search_text": " ".join(
            str(part or "")
            for part in (
                source["document_label"],
                clause_key,
                description,
                location,
                scene_key,
                _clause_scene_label(scene_key),
                applicability,
                field_key,
                _clause_field_label(field_key),
                condition,
                content,
            )
        ),
    }


def _append_clause_entries_for_node(entries: list[dict], source: dict, clause_key: str, clause: dict) -> None:
    def append_fields(target_node: dict, base_path: list[str], scene_key: str) -> bool:
        appended = False
        for field_key, value in target_node.items():
            if _is_editable_clause_field(field_key, value):
                entries.append(
                    _build_clause_library_entry(
                        source=source,
                        clause_key=clause_key,
                        clause=clause,
                        path_parts=[*base_path, field_key],
                        field_key=field_key,
                        content=value,
                        scene_key=scene_key,
                        target_node=target_node,
                    )
                )
                appended = True
        return appended

    variants = clause.get("variants")
    if isinstance(variants, dict):
        for variant_key, variant_node in variants.items():
            if isinstance(variant_node, dict):
                if not append_fields(variant_node, ["clauses", clause_key, "variants", str(variant_key)], str(variant_key)):
                    entries.append(
                        _build_clause_library_entry(
                            source=source,
                            clause_key=clause_key,
                            clause=clause,
                            path_parts=["clauses", clause_key, "variants", str(variant_key)],
                            field_key="",
                            content="",
                            scene_key=str(variant_key),
                            target_node=variant_node,
                            readonly=True,
                            readonly_reason="复杂结构，只读展示，暂不支持直接编辑",
                        )
                    )
            else:
                entries.append(
                    _build_clause_library_entry(
                        source=source,
                        clause_key=clause_key,
                        clause=clause,
                        path_parts=["clauses", clause_key, "variants", str(variant_key)],
                        field_key="",
                        content="",
                        scene_key=str(variant_key),
                        readonly=True,
                        readonly_reason="列表或非文本结构，只读展示",
                    )
                )
        return

    child_clauses = clause.get("clauses")
    if isinstance(child_clauses, dict):
        for child_key, child_node in child_clauses.items():
            if isinstance(child_node, dict):
                if not append_fields(child_node, ["clauses", clause_key, "clauses", str(child_key)], str(child_key)):
                    entries.append(
                        _build_clause_library_entry(
                            source=source,
                            clause_key=clause_key,
                            clause=clause,
                            path_parts=["clauses", clause_key, "clauses", str(child_key)],
                            field_key="",
                            content="",
                            scene_key=str(child_key),
                            target_node=child_node,
                            readonly=True,
                            readonly_reason="子条款没有可编辑正文",
                        )
                    )
        return

    if isinstance(clause.get("text"), str):
        entries.append(
            _build_clause_library_entry(
                source=source,
                clause_key=clause_key,
                clause=clause,
                path_parts=["clauses", clause_key, "text"],
                field_key="text",
                content=clause.get("text", ""),
                scene_key="DEFAULT",
                target_node=clause,
            )
        )
        return

    entries.append(
        _build_clause_library_entry(
            source=source,
            clause_key=clause_key,
            clause=clause,
            path_parts=["clauses", clause_key],
            field_key="",
            content="",
            scene_key="DEFAULT",
            target_node=clause,
            readonly=True,
            readonly_reason="目录或复杂配置，只读展示",
        )
    )


def _build_clause_library_catalog() -> dict:
    entries = []
    sources = []
    for source in _clause_library_sources():
        sources.append({
            "document_type": source["document_type"],
            "document_label": source["document_label"],
            "file_name": source["path"].name,
        })
        if not source["path"].exists():
            continue
        data = json.loads(source["path"].read_text(encoding="utf-8"))
        clauses = data.get("clauses", {}) if isinstance(data, dict) else {}
        if not isinstance(clauses, dict):
            continue
        for clause_key, clause in clauses.items():
            if isinstance(clause, dict):
                _append_clause_entries_for_node(entries, source, str(clause_key), clause)
            else:
                entries.append(
                    _build_clause_library_entry(
                        source=source,
                        clause_key=str(clause_key),
                        clause={},
                        path_parts=["clauses", str(clause_key)],
                        field_key="",
                        content="",
                        scene_key="DEFAULT",
                        readonly=True,
                        readonly_reason="非对象结构，只读展示",
                    )
                )
    entries.sort(key=lambda item: (item["document_type"], item["chapter"], item["clause_key"], item["scene_key"], item["field_key"]))
    return {"success": True, "sources": sources, "entries": entries}


def _save_clause_library_field(path_id: str, content: str) -> dict:
    document_type, path_parts = _decode_clause_path_id(path_id)
    if len(path_parts) < 3 or path_parts[0] != "clauses":
        raise ValueError("只能保存条款正文路径")
    field_key = path_parts[-1]
    source = _clause_library_source_for(document_type)
    if not source["path"].is_file():
        raise FileNotFoundError(f"条款库文件不存在：{source['path'].name}")
    data = json.loads(source["path"].read_text(encoding="utf-8"))
    target = data
    for part in path_parts[:-1]:
        if not isinstance(target, dict) or part not in target:
            raise ValueError("条款路径不存在")
        target = target[part]
    if not isinstance(target, dict) or not _is_editable_clause_field(field_key, target.get(field_key)):
        raise ValueError("该条款字段不可编辑")
    target[field_key] = str(content or "")

    tmp_path = source["path"].with_name(f".{source['path'].name}.tmp")
    try:
        tmp_path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        os.replace(tmp_path, source["path"])
    finally:
        if tmp_path.exists():
            _safe_unlink(tmp_path)
    _reload_generation_engines()
    return {
        "success": True,
        "message": "条款正文已保存",
        "entry": _build_clause_library_entry(
            source=source,
            clause_key=path_parts[1],
            clause=data.get("clauses", {}).get(path_parts[1], {}),
            path_parts=path_parts,
            field_key=field_key,
            content=str(content or ""),
            scene_key=path_parts[-2] if len(path_parts) >= 4 else "DEFAULT",
            target_node=target,
        ),
    }


LINKED_CONTRACT_FORMAT_TEMPLATE_FILENAME = "1、南方中证全指农牧渔交易型开放式指数证券投资基金发起式联接基金基金合同（草案）.docx"
LINKED_PROSPECTUS_FORMAT_TEMPLATE_FILENAME = "2、南方中证全指农牧渔交易型开放式指数证券投资基金发起式联接基金招募说明书（草案）.docx"


def _linked_legal_template_dir(env_var_name: str) -> Path:
    configured = str(os.environ.get(env_var_name) or "").strip()
    if configured:
        return Path(configured)
    packaged = Path(__file__).resolve().parent / "packaged_assets" / "legal_templates"
    if packaged.exists():
        return packaged
    return Path.home() / "Desktop" / "联接基金法律文件"


def _require_linked_format_template(filename: str, env_var_name: str, label: str) -> Path:
    template_dir = _linked_legal_template_dir(env_var_name)
    template_docx = template_dir / filename
    if template_docx.exists():
        return template_docx
    raise FileNotFoundError(
        f"缺少{label} Word 导出所需的固定格式底稿：{template_docx}"
    )


DEFAULT_DISPUTE_RESOLUTION_VENUE = "SZ_SCIA"
DISPUTE_RESOLUTION_VENUES = {
    "SZ_SCIA": {
        "label": "深圳-深圳国际仲裁院",
        "institution": "深圳国际仲裁院",
        "location": "深圳市",
        "contract_clause": (
            "各方当事人同意，因《基金合同》而产生的或与《基金合同》有关的一切争议，"
            "如经友好协商未能解决的，任何一方均有权将争议提交深圳国际仲裁院，"
            "按照深圳国际仲裁院届时有效的仲裁规则进行仲裁。仲裁地点为深圳市，"
            "仲裁裁决是终局的，对当事人均有约束力。除非仲裁裁决另有规定，"
            "仲裁费、律师费由败诉方承担。"
        ),
    },
    "BJ_CIETAC": {
        "label": "北京-中国国际经济贸易仲裁委员会",
        "institution": "中国国际经济贸易仲裁委员会",
        "location": "北京市",
        "contract_clause": (
            "各方当事人同意，因《基金合同》而产生的或与《基金合同》有关的一切争议，"
            "如经友好协商未能解决的，任何一方均有权将争议提交中国国际经济贸易仲裁委员会，"
            "按照中国国际经济贸易仲裁委员会届时有效的仲裁规则进行仲裁。仲裁地点为北京市，"
            "仲裁裁决是终局的，对当事人均有约束力。除非仲裁裁决另有规定，"
            "仲裁费、律师费由败诉方承担。"
        ),
    },
}
DISPUTE_RESOLUTION_VENUE_ALIASES = {
    "": DEFAULT_DISPUTE_RESOLUTION_VENUE,
    "SZ": "SZ_SCIA",
    "SZ_SCIA": "SZ_SCIA",
    "SHENZHEN_SCIA": "SZ_SCIA",
    "深圳": "SZ_SCIA",
    "深圳市": "SZ_SCIA",
    "深圳国际仲裁院": "SZ_SCIA",
    "深圳-深圳国际仲裁院": "SZ_SCIA",
    "BJ": "BJ_CIETAC",
    "BJ_CIETAC": "BJ_CIETAC",
    "BEIJING_CIETAC": "BJ_CIETAC",
    "北京": "BJ_CIETAC",
    "北京市": "BJ_CIETAC",
    "中国国际经济贸易仲裁委员会": "BJ_CIETAC",
    "北京-中国国际经济贸易仲裁委员会": "BJ_CIETAC",
}

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB 上传限制

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
logger = logging.getLogger(__name__)


# A4 页面布局参数（单位：Twips，1 Twips = 1/20 pt = 1/1440 inch）
PAGE_LAYOUT = dict(
    page_width=11906,      # 210mm
    page_height=16838,     # 297mm
    top_margin=1440,       # 2.54cm
    bottom_margin=1440,    # 2.54cm
    left_margin=1800,      # 3.17cm
    right_margin=1800,     # 3.17cm
    header_distance=851,   # 1.50cm
    footer_distance=992,   # 1.75cm
)


def _as_bool(value) -> bool:
    """将各种形式的布尔值转换为 Python bool"""
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in ("true", "1", "yes", "是", "on")
    return bool(value)


def _safe_unlink(path):
    """安全删除临时文件，失败时仅记录日志"""
    try:
        os.unlink(path)
    except Exception:
        logger.debug("Failed to clean temp file: %s", path)


def _same_origin_request() -> bool:
    """Best-effort browser CSRF guard for this local-only tool."""
    origin = request.headers.get("Origin")
    if origin:
        allowed = {request.host_url.rstrip("/")}
        if origin.rstrip("/") not in allowed:
            return False
    sec_fetch_site = request.headers.get("Sec-Fetch-Site", "").lower()
    if sec_fetch_site in {"cross-site", "same-site"}:
        return False
    return True


@app.before_request
def reject_cross_origin_writes():
    if request.method in {"POST", "PUT", "PATCH", "DELETE"} and not _same_origin_request():
        return jsonify({"success": False, "error": "跨站写入请求已被拒绝"}), 403


def _validate_docx_upload(file_obj, filename: str = "") -> io.BytesIO:
    if filename and not filename.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx 文件")
    raw = file_obj.read()
    if len(raw) > app.config["MAX_CONTENT_LENGTH"]:
        raise ValueError("文件过大，请上传 16MB 以内的文件")
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_DOCX_ENTRIES:
                raise ValueError("DOCX 文件结构过于复杂，请精简后再上传")
            total_uncompressed = sum(info.file_size for info in infos)
            if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValueError("DOCX 解压后体积过大，请精简后再上传")
            if total_uncompressed and len(raw) and total_uncompressed / max(len(raw), 1) > 200:
                raise ValueError("DOCX 压缩比异常，请确认文件安全后再上传")
            for info in infos:
                if info.flag_bits & 0x1:
                    raise ValueError("不支持加密 DOCX 文件")
            if "word/document.xml" not in {info.filename for info in infos}:
                raise ValueError("DOCX 文件缺少正文内容")
    except zipfile.BadZipFile as exc:
        raise ValueError("DOCX 文件格式无效") from exc
    stream = io.BytesIO(raw)
    stream.name = filename or "uploaded.docx"
    return stream


def _sanitize_review_html(value) -> str:
    """Escape arbitrary HTML while preserving diff-only tags generated by the server."""
    raw = str(value or "")
    raw = re.sub(r"(?is)<script\b[^>]*>.*?</script>", "", raw)
    raw = re.sub(r"(?is)<style\b[^>]*>.*?</style>", "", raw)
    raw = re.sub(r"(?is)<(?!/?(?:del|ins|mark)\b)[^>]*>", "", raw)
    raw = re.sub(r"(?is)<(del|ins|mark)\b[^>]*>", r"<\1>", raw)
    escaped = html.escape(raw, quote=False)
    for tag in ("del", "ins", "mark"):
        escaped = re.sub(rf"&lt;{tag}&gt;", f"<{tag}>", escaped, flags=re.IGNORECASE)
        escaped = re.sub(rf"&lt;/{tag}&gt;", f"</{tag}>", escaped, flags=re.IGNORECASE)
    return escaped


def _sanitize_review_markup_fields(value):
    if isinstance(value, list):
        return [_sanitize_review_markup_fields(item) for item in value]
    if isinstance(value, dict):
        sanitized = {}
        for key, child in value.items():
            if isinstance(child, str) and (key == "html" or key.endswith("_html")):
                sanitized[key] = _sanitize_review_html(child)
            else:
                sanitized[key] = _sanitize_review_markup_fields(child)
        return sanitized
    return value


def _resolve_dispute_resolution_venue(raw_value: str | None) -> tuple[str, dict[str, str]]:
    normalized = str(raw_value or "").strip()
    key = DISPUTE_RESOLUTION_VENUE_ALIASES.get(normalized)
    if key is None:
        key = DISPUTE_RESOLUTION_VENUE_ALIASES.get(normalized.upper(), DEFAULT_DISPUTE_RESOLUTION_VENUE)
    return key, DISPUTE_RESOLUTION_VENUES[key]


def _normalize_fee_payment_clause_text(text: str, fee_name: str) -> str:
    """Accept either a full fee sentence or only the payment tail expected by the template."""
    normalized = str(text or "").strip()
    if not normalized:
        return ""
    normalized = re.sub(
        rf"^{re.escape(fee_name)}每日(?:计提|计算)[^。]*?按月支付[。，]\s*",
        "",
        normalized,
        count=1,
    ).strip()
    return normalized


def _resolve_payment_clause_value(value, variants, aliases=None, fee_name=""):
    """Return (enum_key, rendered_text) for either a legacy enum or verbatim clause text."""
    raw = "" if value is None else str(value).strip()
    if not raw:
        return "", ""
    normalized_key = (aliases or {}).get(raw, raw)
    if normalized_key in variants:
        return normalized_key, variants[normalized_key]["text"]
    if raw in variants:
        return raw, variants[raw]["text"]
    return "", _normalize_fee_payment_clause_text(raw, fee_name)


_LAYOUT_BLANK_RE = re.compile(r"[\s\u00a0\u3000\u200b\u200c\u200d\ufeff]+")


def _normalize_layout_text(value: str) -> str:
    return (
        str(value or "")
        .replace("\u00a0", " ")
        .replace("\u3000", " ")
        .replace("\u200b", "")
        .replace("\u200c", "")
        .replace("\u200d", "")
        .replace("\ufeff", "")
    )


def _is_layout_blank_line(value: str) -> bool:
    return not _LAYOUT_BLANK_RE.sub("", str(value or ""))


def _is_markdown_table_separator_line(value: str) -> bool:
    text = str(value or "").strip()
    if "|" not in text:
        return False
    cells = [cell.strip() for cell in text.strip("|").split("|")]
    return len(cells) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _clean_docx_cell_text(value: str) -> str:
    return re.sub(r"\s+", " ", _normalize_layout_text(value).replace("\r", "\n")).strip()


def _int_to_cn_numeral(value: int) -> str:
    digits = "零一二三四五六七八九"
    try:
        number = int(value)
    except Exception:
        return str(value)
    if number <= 0:
        return str(value)
    if number < 10:
        return digits[number]
    if number < 20:
        return "十" + (digits[number % 10] if number % 10 else "")
    if number < 100:
        return digits[number // 10] + "十" + (digits[number % 10] if number % 10 else "")
    if number < 1000:
        hundreds = digits[number // 100] + "百"
        rest = number % 100
        if not rest:
            return hundreds
        if rest < 10:
            return hundreds + "零" + digits[rest]
        return hundreds + _int_to_cn_numeral(rest)
    return str(value)


def _alpha_sequence(value: int, upper=False) -> str:
    number = max(int(value), 1)
    chars = []
    while number:
        number -= 1
        chars.append(chr(ord("A" if upper else "a") + number % 26))
        number //= 26
    return "".join(reversed(chars))


def _roman_sequence(value: int, upper=False) -> str:
    number = max(int(value), 1)
    pairs = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    result = []
    for arabic, roman in pairs:
        while number >= arabic:
            result.append(roman)
            number -= arabic
    text = "".join(result)
    return text if upper else text.lower()


def _build_docx_numbering_lookup(doc) -> dict[tuple[str, int], dict]:
    from docx.oxml.ns import qn

    numbering_part = getattr(getattr(doc, "part", None), "numbering_part", None)
    root = getattr(numbering_part, "element", None)
    if root is None:
        return {}

    abstract_levels = {}
    for abstract in root.findall(qn("w:abstractNum")):
        abstract_id = abstract.get(qn("w:abstractNumId"))
        levels = {}
        for level in abstract.findall(qn("w:lvl")):
            ilvl = int(level.get(qn("w:ilvl")) or 0)
            fmt = level.find(qn("w:numFmt"))
            text = level.find(qn("w:lvlText"))
            start = level.find(qn("w:start"))
            levels[ilvl] = {
                "format": fmt.get(qn("w:val")) if fmt is not None else "decimal",
                "level_text": text.get(qn("w:val")) if text is not None else f"%{ilvl + 1}.",
                "start": int(start.get(qn("w:val")) or 1) if start is not None else 1,
            }
        abstract_levels[abstract_id] = levels

    lookup = {}
    for num in root.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        abstract_ref = num.find(qn("w:abstractNumId"))
        abstract_id = abstract_ref.get(qn("w:val")) if abstract_ref is not None else ""
        for ilvl, level_info in (abstract_levels.get(abstract_id) or {}).items():
            lookup[(str(num_id), ilvl)] = dict(level_info)
    return lookup


def _paragraph_numbering_key(paragraph):
    from docx.oxml.ns import qn

    def read_num_pr(num_pr):
        if num_pr is None:
            return None
        num_id = num_pr.find(qn("w:numId"))
        ilvl = num_pr.find(qn("w:ilvl"))
        if num_id is None:
            return None
        return str(num_id.get(qn("w:val"))), int(ilvl.get(qn("w:val")) if ilvl is not None else 0)

    ppr = paragraph._p.pPr
    direct = read_num_pr(ppr.numPr if ppr is not None else None)
    if direct:
        return direct

    style_element = getattr(getattr(paragraph, "style", None), "_element", None)
    style_ppr = getattr(style_element, "pPr", None)
    return read_num_pr(style_ppr.numPr if style_ppr is not None else None)


def _docx_text_local_name(tag: str) -> str:
    return str(tag or "").rsplit("}", 1)[-1]


def _docx_xml_text_accepting_revisions(node) -> str:
    """Read visible text as if tracked insertions are accepted and deletions are rejected."""
    parts = []

    def walk(current):
        name = _docx_text_local_name(getattr(current, "tag", ""))
        if name in {"del", "moveFrom"}:
            return
        if name == "t" and current.text:
            parts.append(str(current.text))
        elif name == "tab":
            parts.append("\t")
        elif name in {"br", "cr"}:
            parts.append("\n")
        for child in list(current):
            walk(child)

    walk(node)
    return _normalize_layout_text("".join(parts)).strip()


def _docx_paragraph_visible_text(paragraph) -> str:
    return _docx_xml_text_accepting_revisions(paragraph._p)


def _format_docx_numbering_value(fmt: str, value: int) -> str:
    fmt = str(fmt or "").strip()
    if fmt in {"decimal", "decimalZero"}:
        return str(value).zfill(2) if fmt == "decimalZero" else str(value)
    if fmt in {"japaneseCounting", "chineseCounting", "chineseCountingThousand", "ideographDigital"}:
        return _int_to_cn_numeral(value)
    if fmt == "lowerLetter":
        return _alpha_sequence(value, upper=False)
    if fmt == "upperLetter":
        return _alpha_sequence(value, upper=True)
    if fmt == "lowerRoman":
        return _roman_sequence(value, upper=False)
    if fmt == "upperRoman":
        return _roman_sequence(value, upper=True)
    return str(value)


def _docx_paragraph_text_with_numbering(paragraph, numbering_lookup, counters) -> str:
    text = _docx_paragraph_visible_text(paragraph)
    key = _paragraph_numbering_key(paragraph)
    if not key or key not in numbering_lookup:
        return text

    num_id, ilvl = key
    for counter_key in list(counters):
        if counter_key[0] == num_id and counter_key[1] > ilvl:
            counters.pop(counter_key, None)
    info = numbering_lookup[key]
    current = counters.get(key, int(info.get("start", 1)) - 1) + 1
    counters[key] = current

    rendered = _format_docx_numbering_value(info.get("format"), current)
    label = re.sub(r"%\d+", rendered, str(info.get("level_text") or ""))
    if not label or _compact_heading_text(text).startswith(_compact_heading_text(label)):
        return text
    return f"{label}{text}"


def _docx_document_to_text(doc) -> str:
    lines = []
    numbering_lookup = _build_docx_numbering_lookup(doc)
    numbering_counters = {}
    for block_type, block in _iter_docx_blocks(doc):
        if block_type == "paragraph":
            text = _docx_paragraph_text_with_numbering(block, numbering_lookup, numbering_counters)
            if not _is_layout_blank_line(text):
                lines.append(text)
        elif block_type == "table":
            lines.extend(_docx_table_to_markdown_lines(block))
    return "\n".join(lines)


def _docx_to_text(file_obj) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_obj)
    return _docx_document_to_text(doc)


def _read_uploaded_text(file_obj, filename):
    name = filename or ""
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    if ext == "docx":
        return _docx_to_text(_validate_docx_upload(file_obj, name))
    if ext in ("txt", "md"):
        raw = file_obj.read()
        if len(raw) > 20 * 1024 * 1024:
            raise ValueError("文件过大，请上传 20MB 以内的文件")
        for enc in ("utf-8", "gbk", "utf-16"):
            try:
                return raw.decode(enc)
            except Exception:
                logger.debug("Encoding %s failed for uploaded file %s", enc, name)
        raise ValueError("文件编码无法识别，请另存为 UTF-8 后再上传")
    raise ValueError("仅支持 .docx、.txt、.md 格式")


def _compact_heading_text(text):
    return re.sub(r"\s+", "", text or "")


def _strip_heading_number(text):
    text = _compact_heading_text(text)
    return re.sub(r"^(第[一二三四五六七八九十百]+[章节部分]|[一二三四五六七八九十]+[、.．])", "", text)


def _is_likely_toc_line(line):
    stripped = line.strip()
    return bool("\t" in stripped and re.search(r"\d+\s*$", stripped))


def _line_matches_heading_alias(line, aliases):
    stripped = line.strip()
    if not stripped or _is_likely_toc_line(stripped):
        return False
    compact = _compact_heading_text(stripped)
    if len(compact) > 120:
        return False
    without_number = _strip_heading_number(stripped)
    alias_compacts = [_compact_heading_text(alias) for alias in aliases]
    return any(without_number == alias or compact == alias for alias in alias_compacts)


def _is_chinese_top_heading(line):
    stripped = line.strip()
    if not stripped or _is_likely_toc_line(stripped):
        return False
    compact = _compact_heading_text(stripped)
    return len(compact) <= 120 and bool(
        re.match(r"^(?:[一二三四五六七八九十百]+、|第[一二三四五六七八九十百千]+[章节部分])", compact)
    )


def _next_chinese_top_heading(lines, start):
    for index in range(start, len(lines)):
        if _is_chinese_top_heading(lines[index]):
            return index
    return len(lines)


def _custody_summary_label_variants(value: str) -> list[str]:
    text = unicodedata.normalize("NFKC", str(value or "")).strip()
    if not text:
        return []

    parts = [part.strip() for part in re.split(r"[+/／]", text) if part and part.strip()]
    variants = []
    for part in parts or [text]:
        candidates = [part]
        candidates.append(re.sub(r"^(?:摘要)?[一二三四五六七八九十百千]+[、．]\s*", "", part).strip())
        candidates.append(re.sub(r"^(?:摘要)?[（(][一二三四五六七八九十百千]+[）)]\s*", "", part).strip())
        candidates.append(re.sub(r"^(?:摘要)?第[一二三四五六七八九十百千]+[章节部分]\s*", "", part).strip())
        candidates.append(re.sub(r"^(?:摘要)?\d+[.、．]\s*", "", part).strip())
        candidates.append(re.sub(r"^(?:摘要)?[（(]\d+[）)]\s*", "", part).strip())
        for candidate in candidates:
            candidate = candidate.strip(" /")
            if candidate and candidate not in variants:
                variants.append(candidate)
    return variants


def _custody_summary_label_ordinal(value: str):
    text = unicodedata.normalize("NFKC", str(value or "")).strip()
    if not text:
        return None

    patterns = (
        r"^(?:摘要)?([一二三四五六七八九十百千]+)[、．]",
        r"^(?:摘要)?[（(]([一二三四五六七八九十百千]+)[）)]",
        r"^(?:摘要)?第([一二三四五六七八九十百千]+)[章节部分]",
        r"^(?:摘要)?(\d+)[.、．]",
        r"^(?:摘要)?[（(](\d+)[）)]",
    )
    for pattern in patterns:
        match = re.match(pattern, text)
        if not match:
            continue
        ordinal = match.group(1)
        return int(ordinal) if ordinal.isdigit() else _cn_numeral_to_int(ordinal)
    return None


def _custody_summary_heading_key(value: str) -> str:
    text = unicodedata.normalize("NFKC", str(value or ""))
    text = re.sub(r"^(?:摘要)?第[一二三四五六七八九十百千]+[章节部分]\s*", "", text)
    text = re.sub(r"^(?:摘要)?[一二三四五六七八九十百千]+[、．]\s*", "", text)
    text = re.sub(r"^(?:摘要)?[（(][一二三四五六七八九十百千]+[）)]\s*", "", text)
    text = re.sub(r"^(?:摘要)?\d+[.、．]\s*", "", text)
    text = re.sub(r"^(?:摘要)?[（(]\d+[）)]\s*", "", text)
    text = re.sub(r"\([^)]*\)", "", text)
    text = re.sub(r"[（(][^）)]*[）)]", "", text)
    text = re.sub(r"[\s的与及和、，,：:；;/-]", "", text)
    replacements = (
        ("基金托管协议", "托管协议"),
        ("基金托管人", "托管人"),
        ("基金管理人", "管理人"),
        ("业务监督核查", "监督核查"),
        ("业务核查", "核查"),
        ("基金托管协议变更终止基金财产清算", "托管协议变更终止基金财产清算"),
    )
    for source, target in replacements:
        text = text.replace(source, target)
    return text


def _custody_summary_heading_direction(value: str) -> str:
    key = _custody_summary_heading_key(value)
    if "托管人对管理人" in key:
        return "custodian_to_manager"
    if "管理人对托管人" in key:
        return "manager_to_custodian"
    return ""


def _custody_summary_labels_equivalent(left: str, right: str) -> bool:
    for left_variant in _custody_summary_label_variants(left):
        left_key = _custody_summary_heading_key(left_variant)
        left_ordinal = _custody_summary_label_ordinal(left_variant)
        if not left_key:
            continue
        for right_variant in _custody_summary_label_variants(right):
            right_key = _custody_summary_heading_key(right_variant)
            if not right_key:
                continue
            if left_key == right_key:
                return True
    return False


def _split_custody_agreement_sections(text, doc=None):
    """拆出托管协议真实大章，逻辑对齐 ETF 合同知识库的托管协议导入。"""
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    if not text.strip():
        return []

    def heading_style_level(style_name: str):
        style_name = str(style_name or "").strip()
        if not style_name:
            return None
        match = re.match(r"^(?:Heading|标题)\s*(\d+)$", style_name, flags=re.IGNORECASE)
        return int(match.group(1)) if match else None

    def is_toc_line(line: str) -> bool:
        stripped = (line or "").strip()
        return bool(stripped and "\t" in stripped and re.search(r"\t+\d+\s*$", stripped))

    def toc_line_heading(line: str) -> str:
        stripped = (line or "").strip()
        if not stripped or not is_toc_line(stripped):
            return ""
        return re.sub(r"\t+\d+\s*$", "", stripped).strip()

    def line_is_plain_heading(line: str) -> bool:
        stripped = (line or "").strip()
        if not stripped or is_toc_line(stripped):
            return False
        if len(stripped) > 80:
            return False
        if re.search(r"[。！？；]", stripped):
            return False
        return True

    def unique_ordered(values):
        result = []
        for value in values:
            if value and value not in result:
                result.append(value)
        return result

    lines = text.split("\n")
    line_offsets = []
    offset = 0
    for line in lines:
        line_offsets.append((offset, line))
        offset += len(line) + 1

    def build_sections_from_line_matches(matches):
        deduped_matches = []
        seen = set()
        for line_index, pos, heading in matches:
            key = (line_index, heading)
            if key in seen:
                continue
            seen.add(key)
            deduped_matches.append((line_index, pos, heading))

        sections = []
        for i, (line_index, pos, heading) in enumerate(deduped_matches):
            line_text = lines[line_index]
            start = pos + len(line_text)
            end = deduped_matches[i + 1][1] if i + 1 < len(deduped_matches) else len(text)
            content = text[start:end].strip()
            if heading and content:
                sections.append({"heading": heading.strip(), "content": content})
        return sections

    def find_heading_line_matches(headings):
        matches = []
        search_start = 0
        for heading in headings:
            found = None
            for line_index in range(search_start, len(line_offsets)):
                pos, line = line_offsets[line_index]
                stripped = line.strip()
                if not line_is_plain_heading(stripped):
                    continue
                if _custody_summary_labels_equivalent(stripped, heading):
                    found = (line_index, pos, stripped)
                    break
            if not found:
                continue
            matches.append(found)
            search_start = found[0] + 1
        return matches

    def find_regex_heading_line_matches():
        matches = []
        re_top = re.compile(
            r"^((?:[一二三四五六七八九十百千]+、|第[一二三四五六七八九十百千]+[章节部分]\s*)[^\n]+)",
            re.MULTILINE,
        )
        for match in re_top.finditer(text):
            heading = match.group(1).strip()
            if is_toc_line(heading):
                continue
            line_index = text[:match.start()].count("\n")
            matches.append((line_index, match.start(), heading))
        return matches

    styled_headings = []
    if doc:
        styled_entries = []
        for paragraph in list(getattr(doc, "paragraphs", None) or []):
            ptext = (getattr(paragraph, "text", "") or "").strip()
            if not ptext:
                continue
            level = heading_style_level(getattr(getattr(paragraph, "style", None), "name", "") or "")
            if level is None:
                continue
            styled_entries.append({"text": ptext, "heading_level": level})
        if styled_entries:
            min_level = min(entry["heading_level"] for entry in styled_entries)
            styled_headings = [entry["text"] for entry in styled_entries if entry["heading_level"] == min_level]

    toc_headings = [toc_line_heading(line) for _, line in line_offsets]
    candidate_headings = unique_ordered([*styled_headings, *toc_headings])
    heading_matches = find_heading_line_matches(candidate_headings)
    regex_matches = find_regex_heading_line_matches()

    if heading_matches:
        # 修订版 DOCX 的标题可能藏在 w:ins 里，python-docx 的样式标题列表会漏掉。
        # 当按“接受修订”文本扫描能找到更多顶级章时，优先用可见文本拆章。
        if len(regex_matches) > len(heading_matches):
            return build_sections_from_line_matches(regex_matches)
        return build_sections_from_line_matches(heading_matches)

    return build_sections_from_line_matches(regex_matches)


def _clean_custody_agreement_sections(sections):
    by_key = {}
    ordered_keys = []
    for section in sections or []:
        heading = str((section or {}).get("heading") or "").strip()
        content = str((section or {}).get("content") or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not heading or not content:
            continue
        key = _custody_summary_heading_key(heading)
        if not key:
            continue
        if key not in by_key:
            ordered_keys.append(key)
            by_key[key] = {"heading": heading, "content": content}
            continue
        if len(content) > len(by_key[key].get("content") or ""):
            by_key[key] = {"heading": heading, "content": content}
    return [by_key[key] for key in ordered_keys]


def _match_custody_agreement_summary_section(sections, spec):
    labels = [spec.get("title", ""), *(spec.get("aliases") or [])]
    candidates = []
    for label in labels:
        for variant in _custody_summary_label_variants(label):
            if variant and variant not in candidates:
                candidates.append(variant)
    best = {"section": None, "match_method": "", "score": 0.0}
    method_rank = {"exact": 5, "normalized": 4, "keywords": 3, "soft_heading": 2, "": 0}
    target_direction = _custody_summary_heading_direction(spec.get("title", ""))

    def update(section, match_method, score):
        current_rank = method_rank.get(best["match_method"], 0)
        new_rank = method_rank.get(match_method, 0)
        if new_rank > current_rank or (new_rank == current_rank and score > best["score"]):
            best["section"] = section
            best["match_method"] = match_method
            best["score"] = score

    def keyword_match_score(heading_key: str) -> float:
        for forbidden in spec.get("forbidden_keywords") or []:
            forbidden_key = _custody_summary_heading_key(forbidden)
            if forbidden_key and forbidden_key in heading_key:
                return 0.0
        keyword_sets = spec.get("keyword_sets") or []
        best_score = 0.0
        for keyword_set in keyword_sets:
            score = 0.0
            matched_groups = 0
            for group in keyword_set:
                alternatives = group if isinstance(group, (list, tuple, set)) else [group]
                matched = ""
                for alternative in alternatives:
                    alt_key = _custody_summary_heading_key(alternative)
                    if alt_key and alt_key in heading_key:
                        matched = alt_key
                        break
                if not matched:
                    matched_groups = 0
                    break
                matched_groups += 1
                score += len(matched)
            if matched_groups:
                best_score = max(best_score, score + matched_groups * 100)
        return best_score

    for candidate in candidates:
        candidate_key = _custody_summary_heading_key(candidate)
        if not candidate_key:
            continue
        for section in sections or []:
            heading = str((section or {}).get("heading") or "").strip()
            heading_key = _custody_summary_heading_key(heading)
            section_direction = _custody_summary_heading_direction(heading)
            if not heading_key:
                continue
            if target_direction and section_direction and section_direction != target_direction:
                continue
            if candidate_key == heading_key:
                update(section, "exact", float(len(candidate_key)))
                continue
            if candidate_key in heading_key or heading_key in candidate_key:
                update(section, "normalized", float(min(len(candidate_key), len(heading_key))))
                continue
            keyword_score = keyword_match_score(heading_key)
            if keyword_score:
                update(section, "keywords", keyword_score)
                continue
            if min(len(candidate_key), len(heading_key)) < 4:
                continue
            ratio = difflib.SequenceMatcher(None, candidate_key, heading_key, autojunk=False).ratio()
            if ratio >= 0.55:
                update(section, "soft_heading", ratio)
    return best


def _extract_chapter_by_title(text, title, next_titles):
    lines = text.splitlines()
    positions = []
    title_compact = _compact_heading_text(title)
    for index, line in enumerate(lines):
        stripped = line.strip()
        if _is_likely_toc_line(stripped):
            continue
        compact = _compact_heading_text(stripped)
        if compact == title_compact or _strip_heading_number(stripped) == title_compact:
            positions.append(index)
    if not positions:
        return ""
    start = positions[-1]
    next_positions = []
    for next_title in next_titles:
        next_compact = _compact_heading_text(next_title)
        for index in range(start + 1, len(lines)):
            stripped = lines[index].strip()
            if _is_likely_toc_line(stripped):
                continue
            if _compact_heading_text(stripped) == next_compact or _strip_heading_number(stripped) == next_compact:
                next_positions.append(index)
                break
    end = min(next_positions) if next_positions else len(lines)
    return "\n".join(lines[start + 1 : end]).strip()


CUSTODY_AGREEMENT_SUMMARY_SECTION_SPECS = [
    {
        "title": "一、基金托管协议当事人",
        "aliases": ["基金托管协议当事人", "托管协议当事人"],
        "keyword_sets": [
            [["当事人"]],
            [["管理人"], ["托管人"]],
        ],
        "forbidden_keywords": ["托管人对管理人", "管理人对托管人", "监督", "核查"],
    },
    {
        "title": "二、基金托管人对基金管理人的业务监督和核查",
        "aliases": [
            "基金托管人对基金管理人的业务监督和核查",
            "基金托管人对基金管理人的监督和核查",
            "基金托管人对基金管理人的业务监督",
        ],
        "keyword_sets": [
            [["托管人对管理人"], ["监督", "核查", "监督核查"]],
        ],
    },
    {
        "title": "三、基金管理人对基金托管人的业务监督和核查",
        "aliases": [
            "基金管理人对基金托管人的业务监督和核查",
            "基金管理人对基金托管人的业务核查",
            "基金管理人对基金托管人的监督和核查",
            "基金管理人对基金托管人的核查",
        ],
        "keyword_sets": [
            [["管理人对托管人"], ["监督", "核查", "监督核查"]],
        ],
    },
    {
        "title": "四、基金财产的保管",
        "aliases": ["基金财产的保管", "基金资产的保管", "资产保管", "基金财产保管"],
        "keyword_sets": [
            [["基金财产", "基金资产", "资产"], ["保管"]],
        ],
    },
    {
        "title": "五、基金资产净值计算、估值和会计核算",
        "aliases": [
            "基金资产净值计算、估值和会计核算",
            "基金资产净值计算、估值与会计核算",
            "基金资产净值计算和会计核算",
            "基金资产净值计算与会计核算",
            "基金资产估值和会计核算",
        ],
        "keyword_sets": [
            [["净值", "估值"], ["计算", "会计", "核算"]],
        ],
    },
    {
        "title": "六、基金份额持有人名册的保管",
        "aliases": ["基金份额持有人名册的保管", "基金份额持有人名册的登记与保管"],
        "keyword_sets": [
            [["持有人名册", "名册"], ["保管", "登记"]],
        ],
    },
    {
        "title": "七、争议解决方式",
        "aliases": ["争议解决方式", "适用法律与争议解决方式", "适用法律和争议解决方式", "法律适用和争议解决方式"],
        "keyword_sets": [
            [["争议"], ["解决", "处理", "适用法律", "法律适用"]],
        ],
    },
    {
        "title": "八、托管协议的变更、终止与基金财产的清算",
        "aliases": [
            "托管协议的变更、终止与基金财产的清算",
            "托管协议的变更与终止",
            "托管协议的变更、终止",
        ],
        "keyword_sets": [
            [["托管协议", "协议"], ["变更"], ["终止"], ["清算"]],
            [["变更"], ["终止"], ["财产清算", "清算"]],
        ],
    },
]


def _custody_agreement_summary_placeholder_block(title):
    return f"{title}\n{{{{{title}}}}}"


def _default_custody_agreement_summary_text():
    return "\n\n".join(
        _custody_agreement_summary_placeholder_block(spec["title"])
        for spec in CUSTODY_AGREEMENT_SUMMARY_SECTION_SPECS
    )


def _build_custody_agreement_summary_from_sections(sections):
    sections = _clean_custody_agreement_sections(sections)
    remaining_sections = list(sections)
    parts = []
    matched = []
    missing = []

    for spec in CUSTODY_AGREEMENT_SUMMARY_SECTION_SPECS:
        match = _match_custody_agreement_summary_section(remaining_sections, spec)
        matched_section = match.get("section")
        if not matched_section:
            missing.append(spec["title"])
            parts.append(_custody_agreement_summary_placeholder_block(spec["title"]))
            continue
        content = str(matched_section.get("content") or "").strip()
        parts.append("\n".join([spec["title"], content]).strip())
        matched.append(spec["title"])
        remaining_sections = [section for section in remaining_sections if section is not matched_section]

    return {
        "summary_text": "\n\n".join(parts).strip(),
        "matched_sections": matched,
        "missing_sections": missing,
        "source_type": "custody_agreement",
    }


def _extract_custody_agreement_summary_from_text(text, doc=None):
    prospectus_chapter = _extract_chapter_by_title(
        text,
        "基金托管协议的内容摘要",
        ["基金份额持有人服务", "其它应披露事项", "招募说明书存放及其查阅方式", "备查文件"],
    )
    if prospectus_chapter:
        return {
            "summary_text": prospectus_chapter,
            "matched_sections": ["基金托管协议的内容摘要"],
            "missing_sections": [],
            "source_type": "prospectus_chapter",
        }

    sections = _split_custody_agreement_sections(text, doc=doc)
    if sections:
        return _build_custody_agreement_summary_from_sections(sections)

    raw_lines = [line.rstrip() for line in text.splitlines()]
    lines = [line for line in raw_lines if line.strip()]
    starts = []
    for spec in CUSTODY_AGREEMENT_SUMMARY_SECTION_SPECS:
        candidates = [
            index
            for index in range(len(lines))
            if _line_matches_heading_alias(lines[index], spec["aliases"])
        ]
        top_heading_candidates = [
            index for index in candidates if _is_chinese_top_heading(lines[index])
        ]
        start = (top_heading_candidates or candidates or [None])[0]
        starts.append(start)

    parts = []
    matched = []
    missing = []
    for spec, start in zip(CUSTODY_AGREEMENT_SUMMARY_SECTION_SPECS, starts):
        if start is None:
            missing.append(spec["title"])
            parts.append(_custody_agreement_summary_placeholder_block(spec["title"]))
            continue
        end = _next_chinese_top_heading(lines, start + 1)
        body_lines = lines[start + 1 : end]
        block = "\n".join([spec["title"]] + body_lines).strip()
        if block:
            parts.append(block)
            matched.append(spec["title"])

    return {
        "summary_text": "\n\n".join(parts).strip(),
        "matched_sections": matched,
        "missing_sections": missing,
        "source_type": "custody_agreement",
    }


def _extract_custody_agreement_summary_from_docx(file_obj, filename=""):
    from docx import Document as DocxDocument

    stream = _validate_docx_upload(file_obj, filename or "uploaded.docx")
    doc = DocxDocument(stream)
    text = _docx_document_to_text(doc)
    return _extract_custody_agreement_summary_from_text(text, doc=doc)


@app.after_request
def add_no_cache_headers(response):
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


@app.errorhandler(HTTPException)
def handle_api_http_exception(exc):
    if request.path.startswith("/api/"):
        message = "上传文件过大，请控制在 16MB 以内" if isinstance(exc, RequestEntityTooLarge) else exc.description
        return jsonify({"success": False, "error": message}), exc.code or 500
    return exc


@app.errorhandler(Exception)
def handle_api_unexpected_exception(exc):
    if request.path.startswith("/api/"):
        logger.exception("Unhandled API error")
        return jsonify({"success": False, "error": "服务内部错误，请查看服务器日志"}), 500
    raise exc


def _clear_xml_children(node):
    for child in list(node):
        node.remove(child)


def _reset_header_footer_part(part):
    try:
        part.is_linked_to_previous = False
    except Exception as e:
        logger.warning("Header/footer link reset failed: %s", e)
    _clear_xml_children(part._element)
    paragraph = part.add_paragraph()
    _clear_xml_children(paragraph._p)
    return paragraph


def _prune_unreferenced_docx_header_footer_parts(docx_bytes: bytes) -> bytes:
    import posixpath
    import xml.etree.ElementTree as ET

    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    def _local_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as src:
        parts = {name: src.read(name) for name in src.namelist()}

    document_root = ET.fromstring(parts["word/document.xml"])
    rels_root = ET.fromstring(parts["word/_rels/document.xml.rels"])
    content_root = ET.fromstring(parts["[Content_Types].xml"])

    used_rel_ids = set()
    for sect_pr in document_root.iter(f"{{{word_ns}}}sectPr"):
        for child in sect_pr:
            if _local_name(child.tag) not in {"headerReference", "footerReference"}:
                continue
            rel_id = child.get(f"{{{office_rel_ns}}}id")
            if rel_id:
                used_rel_ids.add(rel_id)

    used_targets = set()
    for rel in list(rels_root):
        rel_type = rel.get("Type", "")
        is_header_footer = rel_type.endswith("/header") or rel_type.endswith("/footer")
        if not is_header_footer:
            continue
        rel_id = rel.get("Id")
        target = rel.get("Target", "")
        normalized_target = posixpath.normpath(posixpath.join("word", target)).replace("\\", "/")
        if rel_id in used_rel_ids:
            used_targets.add(normalized_target)
        else:
            rels_root.remove(rel)

    for override in list(content_root):
        if _local_name(override.tag) != "Override":
            continue
        part_name = override.get("PartName", "").lstrip("/")
        if not (part_name.startswith("word/header") or part_name.startswith("word/footer")):
            continue
        if part_name not in used_targets:
            content_root.remove(override)

    for part_name in list(parts):
        if part_name.startswith("word/header") or part_name.startswith("word/footer"):
            if part_name not in used_targets:
                parts.pop(part_name, None)

    parts["word/_rels/document.xml.rels"] = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
    parts["[Content_Types].xml"] = ET.tostring(content_root, encoding="utf-8", xml_declaration=True)

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for name, data in parts.items():
            dst.writestr(name, data)
    return output.getvalue()


def _align_cover_section_references_with_template(docx_bytes: bytes, template_docx) -> bytes:
    import posixpath
    import xml.etree.ElementTree as ET

    if not template_docx or not Path(template_docx).exists():
        return docx_bytes

    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    package_rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"

    def _local_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    def _first_cover_sectpr(document_root):
        body = document_root.find(f"{{{word_ns}}}body")
        if body is None:
            return None
        for child in body:
            if _local_name(child.tag) != "p":
                continue
            p_pr = child.find(f"{{{word_ns}}}pPr")
            if p_pr is None:
                continue
            sect_pr = p_pr.find(f"{{{word_ns}}}sectPr")
            if sect_pr is not None:
                return sect_pr
        return None

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as src:
        parts = {name: src.read(name) for name in src.namelist()}
    with zipfile.ZipFile(template_docx) as src:
        template_parts = {name: src.read(name) for name in src.namelist()}

    document_root = ET.fromstring(parts["word/document.xml"])
    rels_root = ET.fromstring(parts["word/_rels/document.xml.rels"])
    content_root = ET.fromstring(parts["[Content_Types].xml"])
    template_document_root = ET.fromstring(template_parts["word/document.xml"])
    template_rels_root = ET.fromstring(template_parts["word/_rels/document.xml.rels"])
    template_content_root = ET.fromstring(template_parts["[Content_Types].xml"])

    cover_sect_pr = _first_cover_sectpr(document_root)
    template_cover_sect_pr = _first_cover_sectpr(template_document_root)
    if cover_sect_pr is None or template_cover_sect_pr is None:
        return docx_bytes

    for child in list(cover_sect_pr):
        if _local_name(child.tag) in {"headerReference", "footerReference"}:
            cover_sect_pr.remove(child)

    existing_rel_ids = {
        rel.get("Id")
        for rel in rels_root
        if _local_name(rel.tag) == "Relationship" and rel.get("Id")
    }

    def _next_rel_id(preferred: str) -> str:
        if preferred not in existing_rel_ids:
            existing_rel_ids.add(preferred)
            return preferred
        index = 1
        while True:
            candidate = f"rId{100 + index}"
            if candidate not in existing_rel_ids:
                existing_rel_ids.add(candidate)
                return candidate
            index += 1

    rel_map = {
        rel.get("Id"): rel
        for rel in template_rels_root
        if _local_name(rel.tag) == "Relationship"
    }
    template_overrides = {
        override.get("PartName", ""): override
        for override in template_content_root
        if _local_name(override.tag) == "Override"
    }

    insert_index = 0
    for ref in list(template_cover_sect_pr):
        local_name = _local_name(ref.tag)
        if local_name not in {"headerReference", "footerReference"}:
            continue
        template_rel_id = ref.get(f"{{{office_rel_ns}}}id")
        template_rel = rel_map.get(template_rel_id)
        if template_rel is None:
            continue
        target = template_rel.get("Target", "")
        full_target = posixpath.normpath(posixpath.join("word", target)).replace("\\", "/")
        if full_target not in template_parts:
            continue
        new_rel_id = _next_rel_id(template_rel_id or "rId200")
        new_rel = ET.Element(f"{{{package_rel_ns}}}Relationship")
        for key, value in template_rel.attrib.items():
            new_rel.set(key, new_rel_id if key == "Id" else value)
        rels_root.append(new_rel)
        parts[full_target] = template_parts[full_target]

        part_name = "/" + full_target
        if not any(
            _local_name(node.tag) == "Override" and node.get("PartName") == part_name
            for node in content_root
        ):
            template_override = template_overrides.get(part_name)
            if template_override is not None:
                content_root.append(ET.fromstring(ET.tostring(template_override, encoding="utf-8")))

        new_ref = ET.fromstring(ET.tostring(ref, encoding="utf-8"))
        new_ref.set(f"{{{office_rel_ns}}}id", new_rel_id)
        cover_sect_pr.insert(insert_index, new_ref)
        insert_index += 1

    parts["word/document.xml"] = ET.tostring(document_root, encoding="utf-8", xml_declaration=True)
    parts["word/_rels/document.xml.rels"] = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
    parts["[Content_Types].xml"] = ET.tostring(content_root, encoding="utf-8", xml_declaration=True)

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for name, data in parts.items():
            dst.writestr(name, data)
    return output.getvalue()


def _append_run_properties(run, OxmlElement, qn, ascii_font=None, eastasia_font=None, size=None, bold=False):
    rPr = OxmlElement("w:rPr")
    if ascii_font or eastasia_font:
        rFonts = OxmlElement("w:rFonts")
        if ascii_font:
            rFonts.set(qn("w:ascii"), ascii_font)
            rFonts.set(qn("w:hAnsi"), ascii_font)
        if eastasia_font:
            rFonts.set(qn("w:eastAsia"), eastasia_font)
        rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size))
        rPr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(size))
        rPr.append(sz_cs)
    run.append(rPr)


def _append_page_field(paragraph, OxmlElement, qn, ascii_font="Times New Roman", eastasia_font="Times New Roman", size=18):
    fld_begin = OxmlElement("w:r")
    _append_run_properties(fld_begin, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    fld_begin.append(fld_char_begin)

    instr_run = OxmlElement("w:r")
    _append_run_properties(instr_run, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = "PAGE   \\* MERGEFORMAT"
    instr_run.append(instr_text)

    fld_sep = OxmlElement("w:r")
    _append_run_properties(fld_sep, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_sep.append(fld_char_sep)

    fld_text = OxmlElement("w:r")
    _append_run_properties(fld_text, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_text_run = OxmlElement("w:t")
    fld_text_run.text = "1"
    fld_text.append(fld_text_run)

    fld_end = OxmlElement("w:r")
    _append_run_properties(fld_end, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    fld_end.append(fld_char_end)

    paragraph._p.append(fld_begin)
    paragraph._p.append(instr_run)
    paragraph._p.append(fld_sep)
    paragraph._p.append(fld_text)
    paragraph._p.append(fld_end)


def _set_section_page_numbers(section, OxmlElement, qn, start=None):
    processed_parts = set()
    footers = [section.footer, section.first_page_footer]
    if hasattr(section, "even_page_footer"):
        footers.append(section.even_page_footer)
    for footer in footers:
        try:
            footer.is_linked_to_previous = False
        except Exception:
            pass
        footer_key = id(footer._element)
        if footer_key in processed_parts:
            continue
        processed_parts.add(footer_key)

        paragraph = _reset_header_footer_part(footer)
        pPr = paragraph._p.get_or_add_pPr()
        _clear_xml_children(pPr)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        _append_page_field(paragraph, OxmlElement, qn, ascii_font="Times New Roman", eastasia_font="Times New Roman", size=18)

    sectPr = section._sectPr
    for node in list(sectPr.findall(qn("w:pgNumType"))):
        sectPr.remove(node)
    if start is not None:
        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:start"), str(start))
        sectPr.append(pg_num_type)


def _clear_header_footer_part(part):
    _reset_header_footer_part(part)


def _remove_section_refs(section, qn):
    sectPr = section._sectPr
    for tag in ("w:headerReference", "w:footerReference"):
        for node in list(sectPr.findall(qn(tag))):
            sectPr.remove(node)


def _set_cover_section_no_page_number(section, OxmlElement, qn):
    _remove_section_refs(section, qn)
    sectPr = section._sectPr
    for tag in ("w:pgNumType", "w:titlePg"):
        for node in list(sectPr.findall(qn(tag))):
            sectPr.remove(node)


def _finalize_doc_page_numbers(doc, OxmlElement, qn, body_start_index=1):
    sections = list(doc.sections)
    if not sections:
        return

    body_start_index = max(1, min(body_start_index, len(sections)))
    for section in sections[:body_start_index]:
        _set_cover_section_no_page_number(section, OxmlElement, qn)

    body_sections = sections[body_start_index:] or sections[-1:]
    for idx, section in enumerate(body_sections):
        _set_section_page_numbers(section, OxmlElement, qn, start=1 if idx == 0 else None)


def _set_update_fields_on_open(document, OxmlElement, qn):
    settings = document.settings.element
    for node in list(settings.findall(qn("w:updateFields"))):
        settings.remove(node)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def _update_docx_fields_with_word(docx_bytes: bytes, toc_entries: list[str] | None = None) -> bytes:
    """Update TOC/page fields with local Microsoft Word when available."""
    if os.environ.get("DISABLE_WORD_FIELD_UPDATE") == "1" or os.environ.get("PYTEST_CURRENT_TEST"):
        return docx_bytes
    if os.name != "nt":
        return docx_bytes
    import tempfile

    tmp_path = ""
    word = None
    document = None
    try:
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="field_update_")
        tmp.write(docx_bytes)
        tmp.close()
        tmp_path = tmp.name

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(tmp_path, ReadOnly=False, AddToRecentFiles=False)
        for toc in document.TablesOfContents:
            toc.Update()
        document.Fields.Update()
        for section in document.Sections:
            for header in section.Headers:
                header.Range.Fields.Update()
            for footer in section.Footers:
                footer.Range.Fields.Update()
        document.Save()
        document.Close(False)
        document = None
        return Path(tmp_path).read_bytes()
    except Exception:
        logger.exception("Failed to update DOCX fields with Microsoft Word")
        return docx_bytes
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            import pythoncom

            pythoncom.CoUninitialize()
        except Exception:
            pass
        if tmp_path:
            _safe_unlink(tmp_path)


def _mark_docx_fields_for_update(docx_bytes: bytes) -> bytes:
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    replacements: dict[str, bytes | str] = {}
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        related_parts = [
            info.filename
            for info in archive.infolist()
            if info.filename.startswith(("word/header", "word/footer")) and info.filename.endswith(".xml")
        ]
    for part_name in ["word/document.xml", *related_parts]:
        xml_text = _read_docx_xml_part(docx_bytes, part_name)
        if not xml_text:
            continue
        try:
            root = ET.fromstring(xml_text)
        except ET.ParseError:
            continue
        changed = False
        for node in root.iter():
            if _docx_local_name(node.tag) == "fldChar":
                node.set(f"{{{_DOCX_W_NS}}}dirty", "true")
                changed = True
        if changed:
            replacements[part_name] = _serialize_docx_document_xml(ET, root, xml_text)

    settings_xml = _read_docx_xml_part(docx_bytes, "word/settings.xml")
    if settings_xml:
        try:
            settings_root = ET.fromstring(settings_xml)
        except ET.ParseError:
            settings_root = ET.Element(f"{{{_DOCX_W_NS}}}settings")
    else:
        settings_root = ET.Element(f"{{{_DOCX_W_NS}}}settings")
    for node in list(settings_root):
        if _docx_local_name(node.tag) == "updateFields":
            settings_root.remove(node)
    update_fields = ET.Element(f"{{{_DOCX_W_NS}}}updateFields")
    update_fields.set(f"{{{_DOCX_W_NS}}}val", "true")
    settings_root.append(update_fields)
    replacements["word/settings.xml"] = ET.tostring(settings_root, encoding="utf-8", xml_declaration=True)
    return _repack_docx_parts(docx_bytes, replacements)


def _finalize_review_docx_for_export(docx_bytes: bytes) -> bytes:
    marked = _mark_docx_fields_for_update(docx_bytes)
    return _update_docx_fields_with_word(marked)


def _reinforce_prospectus_normal_text_format(docx_bytes: bytes) -> bytes:
    """Word may drop explicit Normal text formatting after field updates; restore it where needed."""
    try:
        from lxml import etree
    except Exception:
        return docx_bytes

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    try:
        source = io.BytesIO(docx_bytes)
        target = io.BytesIO()
        changed = False
        with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                payload = zin.read(info.filename)
                if info.filename != "word/document.xml":
                    zout.writestr(info, payload)
                    continue

                parser = etree.XMLParser(remove_blank_text=False, recover=True)
                root = etree.fromstring(payload, parser)
                for paragraph in root.xpath(".//w:p", namespaces=ns):
                    p_style = paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=ns)
                    if p_style and p_style[0] != "Normal":
                        continue
                    p_pr = paragraph.find("w:pPr", namespaces=ns)
                    if p_pr is None:
                        p_pr = etree.Element(f"{{{ns['w']}}}pPr")
                        paragraph.insert(0, p_pr)
                    is_centered = paragraph.xpath("./w:pPr/w:jc[@w:val='center']", namespaces=ns)
                    has_text = bool(paragraph.xpath("./w:r/w:t[normalize-space(.) != '']", namespaces=ns))
                    if has_text and not is_centered:
                        spacing = p_pr.find("w:spacing", namespaces=ns)
                        if spacing is None:
                            spacing = etree.Element(f"{{{ns['w']}}}spacing")
                            ind = p_pr.find("w:ind", namespaces=ns)
                            if ind is not None:
                                p_pr.insert(list(p_pr).index(ind), spacing)
                            else:
                                p_pr.append(spacing)
                            changed = True
                        if spacing.get(f"{{{ns['w']}}}line") != "360":
                            spacing.set(f"{{{ns['w']}}}line", "360")
                            changed = True
                        if spacing.get(f"{{{ns['w']}}}lineRule") != "auto":
                            spacing.set(f"{{{ns['w']}}}lineRule", "auto")
                            changed = True
                    for run in paragraph.xpath("./w:r", namespaces=ns):
                        text_nodes = run.xpath("./w:t", namespaces=ns)
                        if not any((node.text or "").strip() for node in text_nodes):
                            continue
                        r_pr = run.find("w:rPr", namespaces=ns)
                        if r_pr is None:
                            r_pr = etree.Element(f"{{{ns['w']}}}rPr")
                            run.insert(0, r_pr)
                        for tag in ("sz", "szCs"):
                            if r_pr.find(f"w:{tag}", namespaces=ns) is not None:
                                continue
                            node = etree.Element(f"{{{ns['w']}}}{tag}")
                            node.set(f"{{{ns['w']}}}val", "21")
                            r_pr.append(node)
                            changed = True
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                zout.writestr(info, payload)
        return target.getvalue() if changed else docx_bytes
    except Exception:
        logger.exception("Failed to reinforce prospectus normal text format")
        return docx_bytes


def _normalize_prospectus_toc_result(docx_bytes: bytes, toc_entries: list[str]) -> bytes:
    import xml.etree.ElementTree as ET

    if not toc_entries:
        return docx_bytes

    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": word_ns}
    w = f"{{{word_ns}}}"

    def paragraph_text(paragraph):
        parts = []
        for node in paragraph.iter():
            if node.tag == f"{w}t":
                parts.append(node.text or "")
            elif node.tag == f"{w}tab":
                parts.append("\t")
        return "".join(parts)

    def paragraph_style(paragraph):
        pstyle = paragraph.find("./w:pPr/w:pStyle", ns)
        return pstyle.get(f"{w}val") if pstyle is not None else ""

    chapter_prefixes = []
    for entry in toc_entries:
        match = re.match(r"^([一二三四五六七八九十百]+)、\s*(.+)$", str(entry or "").strip())
        if match:
            chapter_prefixes.append(match.group(1))

    if not chapter_prefixes:
        return docx_bytes

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as src:
        parts = {name: src.read(name) for name in src.namelist()}
        infos = {info.filename: info for info in src.infolist()}

    document_xml = parts.get("word/document.xml")
    if not document_xml:
        return docx_bytes
    root = ET.fromstring(document_xml)
    body = root.find("w:body", ns)
    if body is None:
        return docx_bytes

    parent_map = {child: parent for parent in root.iter() for child in parent}
    toc_index = 0
    changed = False
    for paragraph in list(root.findall(".//w:p", ns)):
        style = paragraph_style(paragraph)
        text = paragraph_text(paragraph).strip()
        if style in {"TOC1", "11"}:
            parent = parent_map.get(paragraph)
            if parent is not None:
                parent.remove(paragraph)
            else:
                body.remove(paragraph)
            changed = True
            continue
        if style not in {"TOC2", "25"}:
            continue
        if toc_index >= len(chapter_prefixes):
            continue
        if re.match(r"^[一二三四五六七八九十百]+、", text):
            toc_index += 1
            continue
        first_text = paragraph.find(".//w:t", ns)
        if first_text is not None:
            first_text.text = f"{chapter_prefixes[toc_index]}、\t{first_text.text or ''}"
            toc_index += 1
            changed = True

    if not changed:
        return docx_bytes

    parts["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
        for name, data in parts.items():
            info = infos.get(name)
            if info is not None:
                dst.writestr(info, data)
            else:
                dst.writestr(name, data)
    return out.getvalue()


def _align_prospectus_docx_format_with_reference(docx_bytes: bytes, template_docx) -> bytes:
    import copy
    import xml.etree.ElementTree as ET

    if not template_docx or not Path(template_docx).exists():
        return docx_bytes

    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": word_ns}
    w = f"{{{word_ns}}}"
    style_ids_to_copy = {"11", "21", "25", "34", "36", "HD", "a9", "ab", "af5"}
    paragraph_style_map = {
        "TOC2": "25",
        "33": "34",
    }

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as src:
        parts = {name: src.read(name) for name in src.namelist()}
        infos = {info.filename: info for info in src.infolist()}
    with zipfile.ZipFile(template_docx, "r") as src:
        template_parts = {name: src.read(name) for name in src.namelist()}

    changed = False

    document_xml = parts.get("word/document.xml")
    template_document_xml = template_parts.get("word/document.xml")
    if document_xml and template_document_xml:
        root = ET.fromstring(document_xml)
        template_root = ET.fromstring(template_document_xml)

        for p_style in root.findall(".//w:pPr/w:pStyle", ns):
            current = p_style.get(f"{w}val")
            replacement = paragraph_style_map.get(current)
            if replacement:
                p_style.set(f"{w}val", replacement)
                changed = True

        tables = root.findall(".//w:tbl", ns)
        template_tables = template_root.findall(".//w:tbl", ns)
        for table, template_table in zip(tables, template_tables):
            table_pr = table.find("./w:tblPr", ns)
            template_table_pr = template_table.find("./w:tblPr", ns)
            if template_table_pr is not None:
                if table_pr is not None:
                    table.remove(table_pr)
                table.insert(0, copy.deepcopy(template_table_pr))
                changed = True

            table_grid = table.find("./w:tblGrid", ns)
            template_table_grid = template_table.find("./w:tblGrid", ns)
            if template_table_grid is not None:
                if table_grid is not None:
                    table.remove(table_grid)
                insert_at = 1 if table.find("./w:tblPr", ns) is not None else 0
                table.insert(insert_at, copy.deepcopy(template_table_grid))
                changed = True

            rows = table.findall("./w:tr", ns)
            template_rows = template_table.findall("./w:tr", ns)
            for row, template_row in zip(rows, template_rows):
                row_pr = row.find("./w:trPr", ns)
                template_row_pr = template_row.find("./w:trPr", ns)
                if template_row_pr is not None:
                    if row_pr is not None:
                        row.remove(row_pr)
                    row.insert(0, copy.deepcopy(template_row_pr))
                    changed = True
                cells = row.findall("./w:tc", ns)
                template_cells = template_row.findall("./w:tc", ns)
                for cell, template_cell in zip(cells, template_cells):
                    cell_pr = cell.find("./w:tcPr", ns)
                    template_cell_pr = template_cell.find("./w:tcPr", ns)
                    if template_cell_pr is not None:
                        if cell_pr is not None:
                            cell.remove(cell_pr)
                        cell.insert(0, copy.deepcopy(template_cell_pr))
                        changed = True

        parts["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    styles_xml = parts.get("word/styles.xml")
    template_styles_xml = template_parts.get("word/styles.xml")
    if styles_xml and template_styles_xml:
        styles_root = ET.fromstring(styles_xml)
        template_styles_root = ET.fromstring(template_styles_xml)
        existing = {
            style.get(f"{w}styleId"): style
            for style in styles_root.findall("w:style", ns)
            if style.get(f"{w}styleId")
        }
        for style_id in style_ids_to_copy:
            template_style = template_styles_root.find(f".//w:style[@w:styleId='{style_id}']", ns)
            if template_style is None:
                continue
            current_style = existing.get(style_id)
            if current_style is not None:
                index = list(styles_root).index(current_style)
                styles_root.remove(current_style)
                styles_root.insert(index, copy.deepcopy(template_style))
            else:
                styles_root.append(copy.deepcopy(template_style))
            changed = True
        parts["word/styles.xml"] = ET.tostring(styles_root, encoding="utf-8", xml_declaration=True)

    if not changed:
        return docx_bytes

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
        for name, data in parts.items():
            info = infos.get(name)
            if info is not None:
                dst.writestr(info, data)
            else:
                dst.writestr(name, data)
    return out.getvalue()


def _normalize_prospectus_header_footer_structure(docx_bytes: bytes, template_docx) -> bytes:
    import posixpath
    import xml.etree.ElementTree as ET

    if not template_docx or not Path(template_docx).exists():
        return docx_bytes

    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    package_rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    content_type_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    w = f"{{{word_ns}}}"
    r = f"{{{office_rel_ns}}}"
    pr = f"{{{package_rel_ns}}}"
    ct = f"{{{content_type_ns}}}"

    def local_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    def rel_targets(rels_root):
        return {
            rel.get("Id"): posixpath.normpath(posixpath.join("word", rel.get("Target", ""))).replace("\\", "/")
            for rel in rels_root
            if local_name(rel.tag) == "Relationship"
        }

    def ref_part_from_sect(sect_pr, kind: str, ref_type: str, rel_target_map, parts):
        for child in sect_pr:
            if local_name(child.tag) != f"{kind}Reference":
                continue
            if child.get(f"{w}type") != ref_type:
                continue
            target = rel_target_map.get(child.get(f"{r}id"))
            if target and target in parts:
                return parts[target]
        return None

    def blank_part(kind: str) -> bytes:
        root_name = "hdr" if kind == "header" else "ftr"
        return (
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:{root_name} xmlns:w="{word_ns}"><w:p/></w:{root_name}>'
        ).encode("utf-8")

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as src:
        parts = {name: src.read(name) for name in src.namelist()}
        infos = {info.filename: info for info in src.infolist()}
    with zipfile.ZipFile(template_docx, "r") as src:
        template_parts = {name: src.read(name) for name in src.namelist()}

    document_root = ET.fromstring(parts["word/document.xml"])
    rels_root = ET.fromstring(parts["word/_rels/document.xml.rels"])
    content_root = ET.fromstring(parts["[Content_Types].xml"])
    sect_prs = list(document_root.iter(f"{w}sectPr"))
    if len(sect_prs) < 2:
        return docx_bytes

    rel_target_map = rel_targets(rels_root)
    cover_sect = sect_prs[0]
    body_sect = sect_prs[-1]

    cover_footer = (
        ref_part_from_sect(cover_sect, "footer", "default", rel_target_map, parts)
        or template_parts.get("word/footer1.xml")
        or blank_part("footer")
    )
    body_header_default = (
        ref_part_from_sect(body_sect, "header", "default", rel_target_map, parts)
        or ref_part_from_sect(body_sect, "header", "first", rel_target_map, parts)
        or blank_part("header")
    )
    body_header_first = (
        ref_part_from_sect(body_sect, "header", "first", rel_target_map, parts)
        or body_header_default
    )
    body_footer_default = (
        ref_part_from_sect(body_sect, "footer", "default", rel_target_map, parts)
        or ref_part_from_sect(body_sect, "footer", "first", rel_target_map, parts)
        or blank_part("footer")
    )
    body_footer_first = (
        ref_part_from_sect(body_sect, "footer", "first", rel_target_map, parts)
        or body_footer_default
    )

    desired_parts = {
        "word/footer1.xml": cover_footer,
        "word/header1.xml": body_header_default,
        "word/header2.xml": body_header_default,
        "word/footer2.xml": body_footer_default,
        "word/footer3.xml": body_footer_default,
        "word/header3.xml": body_header_first,
        "word/footer4.xml": body_footer_first,
    }

    for child in list(cover_sect):
        if local_name(child.tag) in {"headerReference", "footerReference"}:
            cover_sect.remove(child)
    for child in list(body_sect):
        if local_name(child.tag) in {"headerReference", "footerReference"}:
            body_sect.remove(child)

    def add_reference(sect_pr, kind: str, ref_type: str, rel_id: str, insert_at: int):
        ref = ET.Element(f"{w}{kind}Reference")
        ref.set(f"{w}type", ref_type)
        ref.set(f"{r}id", rel_id)
        sect_pr.insert(insert_at, ref)

    add_reference(cover_sect, "footer", "default", "rId8", 0)
    for idx, (kind, ref_type, rel_id) in enumerate(
        [
            ("header", "even", "rId9"),
            ("header", "default", "rId10"),
            ("footer", "even", "rId11"),
            ("footer", "default", "rId12"),
            ("header", "first", "rId13"),
            ("footer", "first", "rId14"),
        ]
    ):
        add_reference(body_sect, kind, ref_type, rel_id, idx)

    for rel in list(rels_root):
        rel_type = rel.get("Type", "")
        if rel_type.endswith("/header") or rel_type.endswith("/footer"):
            rels_root.remove(rel)

    rel_specs = [
        ("rId8", "footer", "footer1.xml"),
        ("rId9", "header", "header1.xml"),
        ("rId10", "header", "header2.xml"),
        ("rId11", "footer", "footer2.xml"),
        ("rId12", "footer", "footer3.xml"),
        ("rId13", "header", "header3.xml"),
        ("rId14", "footer", "footer4.xml"),
    ]
    for rel_id, kind, target in rel_specs:
        rel = ET.Element(f"{pr}Relationship")
        rel.set("Id", rel_id)
        rel.set("Type", f"http://schemas.openxmlformats.org/officeDocument/2006/relationships/{kind}")
        rel.set("Target", target)
        rels_root.append(rel)

    for name in list(parts):
        if name.startswith("word/header") or name.startswith("word/footer"):
            parts.pop(name, None)
    parts.update(desired_parts)

    for override in list(content_root):
        if local_name(override.tag) != "Override":
            continue
        part_name = override.get("PartName", "").lstrip("/")
        if part_name.startswith("word/header") or part_name.startswith("word/footer"):
            content_root.remove(override)
    for part_name in desired_parts:
        override = ET.Element(f"{ct}Override")
        override.set("PartName", f"/{part_name}")
        kind = "header" if part_name.startswith("word/header") else "footer"
        override.set(
            "ContentType",
            f"application/vnd.openxmlformats-officedocument.wordprocessingml.{kind}+xml",
        )
        content_root.append(override)

    parts["word/document.xml"] = ET.tostring(document_root, encoding="utf-8", xml_declaration=True)
    parts["word/_rels/document.xml.rels"] = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
    parts["[Content_Types].xml"] = ET.tostring(content_root, encoding="utf-8", xml_declaration=True)

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
        for name, data in parts.items():
            info = infos.get(name)
            if info is not None:
                dst.writestr(info, data)
            else:
                dst.writestr(name, data)
    return out.getvalue()


def _append_word_toc_field(
    paragraph,
    OxmlElement,
    qn,
    *,
    levels: str,
    placeholder: str = "右键更新目录",
    ascii_font: str = "Times New Roman",
    eastasia_font: str = "宋体",
    size: int = 24,
    use_outline_levels: bool = True,
    style_map: str | None = None,
):
    if style_map:
        instruction = f' TOC \\\\t "{style_map}" \\\\h \\\\z'
    else:
        instruction = f' TOC \\\\o "{levels}" \\\\h \\\\z'
    if use_outline_levels and not style_map:
        instruction += " \\\\u"
    instruction += " "

    begin = OxmlElement("w:r")
    _append_run_properties(begin, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    begin.append(fld_begin)
    paragraph._p.append(begin)

    instr_run = OxmlElement("w:r")
    _append_run_properties(instr_run, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = instruction
    instr_run.append(instr_text)
    paragraph._p.append(instr_run)

    sep = OxmlElement("w:r")
    _append_run_properties(sep, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep.append(fld_sep)
    paragraph._p.append(sep)

    if placeholder:
        text_run = OxmlElement("w:r")
        _append_run_properties(text_run, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
        text_node = OxmlElement("w:t")
        text_node.text = placeholder
        text_run.append(text_node)
        paragraph._p.append(text_run)

    end = OxmlElement("w:r")
    _append_run_properties(end, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    paragraph._p.append(end)


def _ensure_word_toc_styles(
    document,
    *,
    heading_eastasia: str = "宋体",
    heading_ascii: str = "宋体",
    heading_size_half_pt: int = 28,
    heading_bold: bool = False,
    entry_eastasia: str = "宋体",
    entry_ascii: str = "Times New Roman",
    entry_size_half_pt: int = 24,
    max_level: int = 3,
):
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def ensure_style(style_name: str, base_style_name: str | None = None):
        styles = document.styles
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        if base_style_name:
            try:
                style.base_style = styles[base_style_name]
            except KeyError:
                pass
        return style

    def set_style_font(style, *, eastasia: str, ascii_font: str, size_half_pt: int, bold: bool = False):
        style.font.name = ascii_font
        style.font.size = Pt(size_half_pt / 2)
        style.font.bold = bold
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), ascii_font)
        rfonts.set(qn("w:hAnsi"), ascii_font)
        rfonts.set(qn("w:eastAsia"), eastasia)
        for tag in ("w:sz", "w:szCs"):
            node = rpr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                rpr.append(node)
            node.set(qn("w:val"), str(size_half_pt))
        for tag in ("w:b", "w:bCs"):
            node = rpr.find(qn(tag))
            if bold:
                if node is None:
                    rpr.append(OxmlElement(tag))
            elif node is not None:
                rpr.remove(node)

    toc_heading = ensure_style("TOC Heading", "Normal")
    set_style_font(
        toc_heading,
        eastasia=heading_eastasia,
        ascii_font=heading_ascii,
        size_half_pt=heading_size_half_pt,
        bold=heading_bold,
    )
    toc_heading.paragraph_format.first_line_indent = Pt(0)

    for level in range(1, max_level + 1):
        toc_style = ensure_style(f"TOC {level}", "Normal")
        set_style_font(
            toc_style,
            eastasia=entry_eastasia,
            ascii_font=entry_ascii,
            size_half_pt=entry_size_half_pt,
            bold=False,
        )
        toc_style.paragraph_format.first_line_indent = Pt(0)


# ═══════════════════════════════════════════════════════════════════════════════
#  ContractEngine — 6步处理管线
# ═══════════════════════════════════════════════════════════════════════════════
class ContractEngine:
    _CONTRACT_SUMMARY_PLACEHOLDER_RE = re.compile(r"\{(CONTRACT_SUMMARY_[A-Z0-9_]*)\}")
    CONTRACT_SUMMARY_SPECS = (
        {
            "placeholder": "CONTRACT_SUMMARY_PARTIES",
            "fragments": (
                {
                    "locator": "第七部分 基金合同当事人及权利义务 / （二）基金管理人的权利与义务",
                    "heading": "（一）基金管理人的权利与义务",
                },
                {
                    "locator": "第七部分 基金合同当事人及权利义务 / （二）基金托管人的权利与义务",
                    "heading": "（二）基金托管人的权利与义务",
                },
                {
                    "locator": "第七部分 基金合同当事人及权利义务 / 三、基金份额持有人",
                    "heading": "（三）基金份额持有人",
                },
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_MEETING",
            "fragments": (
                {"locator": "第八部分 基金份额持有人大会", "take": "intro", "stop_before": "一、召开事由"},
                {"locator": "第八部分 基金份额持有人大会 / 一、召开事由", "heading": "（一）召开事由"},
                {"locator": "第八部分 基金份额持有人大会 / 二、会议召集人及召集方式", "heading": "（二）会议召集人及召集方式"},
                {
                    "locator": "第八部分 基金份额持有人大会 / 三、召开基金份额持有人大会的通知时间、通知内容、通知方式",
                    "heading": "（三）召开基金份额持有人大会的通知时间、通知内容、通知方式",
                },
                {"locator": "第八部分 基金份额持有人大会 / 四、基金份额持有人出席会议的方式", "heading": "（四）基金份额持有人出席会议的方式"},
                {"locator": "第八部分 基金份额持有人大会 / 五、议事内容与程序", "heading": "（五）议事内容与程序"},
                {"locator": "第八部分 基金份额持有人大会 / 六、表决", "heading": "（六）表决"},
                {"locator": "第八部分 基金份额持有人大会 / 七、计票", "heading": "（七）计票"},
                {"locator": "第八部分 基金份额持有人大会 / 八、生效与公告", "heading": "（八）生效与公告"},
                {
                    "locator": "第八部分 基金份额持有人大会 / 九、实施侧袋机制期间基金份额持有人大会的特殊约定",
                    "heading": "（九）实施侧袋机制期间基金份额持有人大会的特殊约定",
                },
                {
                    "locator": "第八部分 基金份额持有人大会 / 十、本部分关于基金份额持有人大会召开事由、召开条件、议事程序、表决条件等规定，凡是直接引用法律法规或监管规则的部分，如将来法律法规或监管规则修改导致相关内容被取消或变更的，基金管理人与基金托管人根据新颁布的法律法规或监管规则协商一致并提前公告后，可直接对本部分内容进行修改和调整，无需召开基金份额持有人大会审议。",
                    "heading": "（十）本部分关于基金份额持有人大会召开事由、召开条件、议事程序、表决条件等规定，凡是直接引用法律法规或监管规则的部分，如将来法律法规或监管规则修改导致相关内容被取消或变更的，基金管理人与基金托管人根据新颁布的法律法规或监管规则协商一致并提前公告后，可直接对本部分内容进行修改和调整，无需召开基金份额持有人大会审议。",
                },
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_DISTRIBUTION",
            "fragments": (
                {"locator": "第十六部分 基金的收益与分配 / 三、基金收益分配原则", "heading": "（一）基金收益分配原则"},
                {"locator": "第十六部分 基金的收益与分配 / 四、收益分配方案", "heading": "（二）收益分配方案"},
                {
                    "locator": "第十六部分 基金的收益与分配 / 五、收益分配方案的确定、公告与实施",
                    "heading": "（三）收益分配方案的确定、公告与实施",
                },
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_FEES",
            "fragments": (
                {"locator": "第十五部分 基金费用与税收 / 一、基金费用的种类", "heading": "（一）基金费用的种类"},
                {
                    "locator": "第十五部分 基金费用与税收 / 二、基金费用计提方法、计提标准和支付方式",
                    "heading": "（二）基金费用计提方法、计提标准和支付方式",
                },
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_INVESTMENT",
            "fragments": (
                {"locator": "第十二部分 基金的投资 / 二、投资范围", "heading": "（一）投资范围"},
                {"locator": "第十二部分 基金的投资 / 四、投资限制", "heading": "（二）投资限制"},
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_VALUATION",
            "fragments": (
                {"locator": "第十四部分 基金资产估值 / 五、估值程序"},
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_CHANGE_TERMINATION",
            "fragments": (
                {"locator": "第十九部分 基金合同的变更、终止与基金财产的清算 / 一、《基金合同》的变更", "heading": "（一）《基金合同》的变更"},
                {"locator": "第十九部分 基金合同的变更、终止与基金财产的清算 / 二、《基金合同》的终止事由", "heading": "（二）《基金合同》的终止事由"},
                {"locator": "第十九部分 基金合同的变更、终止与基金财产的清算 / 三、基金财产的清算", "heading": "（三）基金财产的清算"},
                {"locator": "第十九部分 基金合同的变更、终止与基金财产的清算 / 四、清算费用", "heading": "（四）清算费用"},
                {
                    "locator": "第十九部分 基金合同的变更、终止与基金财产的清算 / 五、基金财产清算剩余资产的分配",
                    "heading": "（五）基金财产清算剩余资产的分配",
                },
                {"locator": "第十九部分 基金合同的变更、终止与基金财产的清算 / 六、基金财产清算的公告", "heading": "（六）基金财产清算的公告"},
                {
                    "locator": "第十九部分 基金合同的变更、终止与基金财产的清算 / 七、基金财产清算账册及文件的保存",
                    "heading": "（七）基金财产清算账册及文件的保存",
                },
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_DISPUTE",
            "fragments": (
                {"locator": "第二十一部分 争议的处理和适用的法律"},
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_EFFECT",
            "fragments": (
                {"locator": "第二十二部分 基金合同的效力"},
            ),
        },
    )

    def __init__(self):
        with open(CLAUSES_JSON, encoding="utf-8") as f:
            self.clauses = json.load(f)["clauses"]

    def _contract_clause_text(self, clause_key: str, variant_key: str, field: str = "text") -> str:
        return str(
            self.clauses.get(clause_key, {})
            .get("variants", {})
            .get(variant_key, {})
            .get(field, "")
            or ""
        )

    def _initiating_definition_terms(self) -> list[str]:
        terms = (
            self.clauses.get("INITIATING_FUND_DEFINITIONS", {})
            .get("variants", {})
            .get("DEFAULT", {})
            .get("terms", [])
        )
        return [str(term or "").strip() for term in terms if str(term or "").strip()]

    # ── Step 1: 派生变量 ─────────────────────────────────────────────────────
    def _derive_variables(self, v: dict) -> dict:
        v = _merge_schema_variable_defaults(v)

        # HAS_HK_CONNECT from INDEX_NAME or direct setting (keyword detection has final say)
        hk_detected = "港股通" in v.get("INDEX_NAME", "") or "港股通" in v.get("FUND_NAME", "")
        v["HAS_HK_CONNECT"] = hk_detected if hk_detected else v.get("HAS_HK_CONNECT", False)

        # IS_INITIATING_FUND from FUND_NAME
        init_detected = "发起式" in v.get("FUND_NAME", "")
        v["IS_INITIATING_FUND"] = init_detected if init_detected else v.get("IS_INITIATING_FUND", False)

        # Special product flags from fund/index names as backend fallback.
        product_name_text = f"{v.get('INDEX_NAME', '')} {v.get('FUND_NAME', '')}"
        chuangye_detected = "创业板" in product_name_text
        kechuang_detected = "科创" in product_name_text
        if chuangye_detected:
            v["IS_CHUANGYE_FUND"] = True
        if kechuang_detected:
            v["IS_KECHUANG_FUND"] = True

        # 中文布尔 → Python bool
        condition_bool_keys = (
            "HAS_HK_CONNECT",
            "IS_INITIATING_FUND",
            "CUSTODIAN_HAS_OFFICE_ADDRESS",
            "IS_CHUANGYE_FUND",
            "IS_KECHUANG_FUND",
            "INCLUDE_COVER_DATE",
            "INDEX_MAY_CHANGE",
            "TARGET_ETF_FULL_REPLICATION",
            "CONSIDERATION_MAINLY_BASKET",
            "SHOW_FUTURES_SETTLEMENT_ACCOUNT",
            "HAS_CUSTODIAN_SUBSIDIARY_CLAUSE",
            "HK_DAILY_QUOTA_CONTROL",
            "CHUANGYE_RISK_RETURN_FEATURE",
            "KECHUANG_RISK_DISCLOSURE",
            "HK_RISK_DISCLOSURE",
            "SPECIAL_ETF_SUBSCRIPTION_REDEMPTION",
        )
        for key in condition_bool_keys:
            if key in v:
                v[key] = _as_bool(v[key])

        # 托管人有无办公地址 — 从04.json CUSTODIAN_INFO_KNOWN_VALUES动态读取
        custodian = v.get("CUSTODIAN_NAME", "")
        custodians_data = self.clauses.get("CUSTODIAN_INFO_KNOWN_VALUES", {}).get("custodians", {})
        custodian_info = custodians_data.get(custodian, {})
        v.setdefault(
            "CUSTODIAN_HAS_OFFICE_ADDRESS",
            custodian_info.get("has_office_address", False),
        )
        v.setdefault("CUSTODIAN_TYPE", custodian_info.get("custodian_type") or ("证券公司" if "证券" in custodian else "商业银行"))
        v.setdefault("HAS_CUSTODIAN_SUBSIDIARY_CLAUSE", custodian_info.get("has_special_transfer", False))

        # C类销售服务费默认值
        v.setdefault("SALES_SERVICE_FEE_RATE", "0.20")
        v.setdefault("INCLUDE_COVER_DATE", True)
        v.setdefault("TARGET_ETF_FULL_REPLICATION", True)
        v.setdefault("CONSIDERATION_MAINLY_BASKET", not v.get("HAS_HK_CONNECT"))
        v.setdefault("SHOW_FUTURES_SETTLEMENT_ACCOUNT", True)
        v.setdefault("HK_DAILY_QUOTA_CONTROL", v.get("HAS_HK_CONNECT", False))
        v.setdefault("CHUANGYE_RISK_RETURN_FEATURE", v.get("IS_CHUANGYE_FUND", False))
        v.setdefault("KECHUANG_RISK_DISCLOSURE", v.get("IS_KECHUANG_FUND", False))
        v.setdefault("HK_RISK_DISCLOSURE", v.get("HAS_HK_CONNECT", False))
        v.setdefault("INDEX_WEBSITE", self._infer_index_website(v.get("INDEX_COMPILER", ""), v.get("INDEX_NAME", "")))
        dispute_key, dispute_config = _resolve_dispute_resolution_venue(
            v.get("DISPUTE_RESOLUTION_VENUE") or v.get("DISPUTE_RESOLUTION_PLACE")
        )
        v["DISPUTE_RESOLUTION_VENUE"] = dispute_key
        v["DISPUTE_RESOLUTION_LABEL"] = dispute_config["label"]
        v["DISPUTE_ARBITRATION_INSTITUTION"] = dispute_config["institution"]
        v["DISPUTE_ARBITRATION_LOCATION"] = dispute_config["location"]
        v["DISPUTE_RESOLUTION_CLAUSE"] = dispute_config["contract_clause"]

        dist_periods = {"MONTHLY": "月", "QUARTERLY": "季度", "SEMI_ANNUAL": "半年"}
        v["DISTRIBUTION_REVIEW_PERIOD"] = dist_periods.get(v.get("DISTRIBUTION_FREQ", "QUARTERLY"), "季度")

        is_initiating = _as_bool(v.get("IS_INITIATING_FUND"))
        is_hk = _as_bool(v.get("HAS_HK_CONNECT"))
        is_chuangye = _as_bool(v.get("IS_CHUANGYE_FUND"))
        is_kechuang = _as_bool(v.get("IS_KECHUANG_FUND"))
        market_scope = v.get("MARKET_SCOPE") or ("跨市场" if is_hk else "跨市场")
        single_exchange = v.get("SINGLE_MARKET_EXCHANGE", "")

        # Special product flags are mutually exclusive in the templates. Normalize conflicts
        # before deriving conditions so generation cannot mix HK/ChiNext/STAR clauses.
        if is_hk:
            is_chuangye = False
            is_kechuang = False
            market_scope = "跨市场"
            single_exchange = ""
        elif is_chuangye:
            is_kechuang = False
            market_scope = "单市场"
            single_exchange = "深交所"
        elif is_kechuang:
            market_scope = "单市场"
            single_exchange = "上交所"

        v["HAS_HK_CONNECT"] = is_hk
        v["IS_CHUANGYE_FUND"] = is_chuangye
        v["IS_KECHUANG_FUND"] = is_kechuang
        v["MARKET_SCOPE"] = market_scope
        v["SINGLE_MARKET_EXCHANGE"] = single_exchange

        custodian_type = v.get("CUSTODIAN_TYPE", "")
        is_cross_market = market_scope == "跨市场"
        is_bank_custodian = custodian_type == "商业银行"
        is_securities_custodian = custodian_type == "证券公司"
        is_standard = not is_initiating

        v.update({
            "COND_INITIATING_FUND": is_initiating,
            "COND_STANDARD_FUND": is_standard,
            "COND_HK_CONNECT": is_hk,
            "COND_NON_HK_CONNECT": not is_hk,
            "COND_CHUANGYE": is_chuangye,
            "COND_KECHUANG": is_kechuang,
            "COND_CROSS_MARKET": is_cross_market,
            "COND_CROSS_MARKET_NON_HK": is_cross_market and not is_hk,
            "COND_SINGLE_MARKET_SH_NON_SPECIAL": (
                market_scope == "单市场" and single_exchange == "上交所" and not is_hk and not is_chuangye and not is_kechuang
            ),
            "COND_SINGLE_MARKET_SZ_NON_SPECIAL": (
                market_scope == "单市场" and single_exchange == "深交所" and not is_hk and not is_chuangye and not is_kechuang
            ),
            "COND_BANK_CUSTODIAN": is_bank_custodian,
            "COND_SECURITIES_CUSTODIAN": is_securities_custodian,
            "COND_CUSTODIAN_SUBSIDIARY_CLAUSE": _as_bool(v.get("HAS_CUSTODIAN_SUBSIDIARY_CLAUSE")),
            "COND_NOT_CUSTODIAN_SUBSIDIARY_CLAUSE": not _as_bool(v.get("HAS_CUSTODIAN_SUBSIDIARY_CLAUSE")),
        })
        v.update({
            "COND_INITIATING_BANK_CROSS_NON_HK": (
                is_initiating and is_bank_custodian and is_cross_market and not is_hk and not is_chuangye and not is_kechuang
            ),
            "COND_INITIATING_SECURITIES_CUSTODIAN": is_initiating and is_securities_custodian,
            # Historical template names: these branches contain special product clauses, not only standard-fund clauses.
            "COND_STANDARD_CHUANGYE": is_chuangye and market_scope == "单市场" and single_exchange == "深交所" and not is_kechuang and not is_hk,
            "COND_STANDARD_KECHUANG": is_kechuang and market_scope == "单市场" and single_exchange == "上交所" and not is_chuangye and not is_hk,
            # Historical template name: this branch contains HK Connect clauses, not only standard-fund clauses.
            "COND_STANDARD_HK_CONNECT": is_cross_market and is_hk,
            "COND_STANDARD_NON_CHUANGYE_NON_KECHUANG_NON_HK": is_standard and not is_chuangye and not is_kechuang and not is_hk,
        })
        v["COND_STANDARD_CHUANGYE_OR_KECHUANG"] = v["COND_STANDARD_CHUANGYE"] or v["COND_STANDARD_KECHUANG"]
        v["COND_NON_HK_NON_CHUANGYE_NON_KECHUANG"] = not is_hk and not is_chuangye and not is_kechuang
        v["COND_NOT_STANDARD_CHUANGYE"] = not v["COND_STANDARD_CHUANGYE"]
        v["COND_NOT_KECHUANG_NOT_HK"] = not is_kechuang and not is_hk
        v["COND_INITIATING_SECURITIES_OR_STANDARD_CHUANGYE"] = v["COND_INITIATING_SECURITIES_CUSTODIAN"] or v["COND_STANDARD_CHUANGYE"]

        return v

    @staticmethod
    def _infer_index_website(index_compiler: str, index_name: str) -> str:
        compiler = index_compiler or ""
        name = index_name or ""
        if "深圳证券信息" in compiler or any(k in name for k in ("深证", "国证", "巨潮")):
            return "www.cnindex.com.cn"
        if "中证" in compiler or "中证" in name or "上证" in name or "科创" in name:
            return "www.csindex.com.cn"
        return ""

    # ── Step 2: 注入差异条款原文 ─────────────────────────────────────────────
    def _inject_clause_texts(self, v: dict) -> dict:
        v = dict(v)

        # DISTRIBUTION_FREQ_CLAUSE
        df = v.get("DISTRIBUTION_FREQ", "QUARTERLY")
        df_variants = self.clauses["DISTRIBUTION_FREQ_CLAUSE"]["variants"]
        v["DISTRIBUTION_FREQ_CLAUSE"] = df_variants.get(df, df_variants["QUARTERLY"])["text"]

        # MGMT_FEE_PAYMENT_METHOD text (accepts either a legacy enum or verbatim clause text)
        method_aliases = {"AUTO_V2": "AUTO", "AUTO_CHECK": "AUTO"}
        mfpm_raw = v.get("MGMT_FEE_PAYMENT_METHOD") or "CONSULT"
        mgmt_variants = self.clauses["MGMT_FEE_PAYMENT"]["variants"]
        mfpm_enum, mfpm_text = _resolve_payment_clause_value(
            mfpm_raw,
            mgmt_variants,
            method_aliases,
            "基金管理费",
        )
        # Store original enum/raw values for custody lookup.
        v["_MGMT_FEE_PAYMENT_ENUM"] = mfpm_enum
        v["_MGMT_FEE_PAYMENT_RAW"] = "" if mfpm_enum else mfpm_text
        v["MGMT_FEE_PAYMENT_METHOD"] = mfpm_text or mgmt_variants["CONSULT"]["text"]
        v["MGMT_FEE_PAYMENT"] = v["MGMT_FEE_PAYMENT_METHOD"]  # mirror for business text overrides

        # CUSTODY_FEE_PAYMENT_METHOD text
        # Alias mapping for UI/backward-compatible values.
        _CUSTODY_ALIAS = {"AUTO_V2": "AUTO", "AUTO_CHECK": "AUTO"}
        cfpm_raw = v.get("CUSTODY_FEE_PAYMENT_METHOD") or ""
        if not cfpm_raw:
            # Derive from custodian data in 04.json
            cust_info2 = self.clauses.get("CUSTODIAN_INFO_KNOWN_VALUES", {}).get("custodians", {}).get(v.get("CUSTODIAN_NAME", ""), {})
            cfpm_raw = (
                cust_info2.get("custody_fee_payment")
                or v.get("_MGMT_FEE_PAYMENT_ENUM")
                or v.get("_MGMT_FEE_PAYMENT_RAW")
                or "CONSULT"
            )
        cust_variants = self.clauses["CUSTODY_FEE_PAYMENT"]["variants"]
        _, cfpm_text = _resolve_payment_clause_value(
            cfpm_raw,
            cust_variants,
            _CUSTODY_ALIAS,
            "基金托管费",
        )
        v["CUSTODY_FEE_PAYMENT_METHOD"] = cfpm_text or cust_variants["CONSULT"]["text"]
        v["CUSTODY_FEE_PAYMENT"] = v["CUSTODY_FEE_PAYMENT_METHOD"]  # mirror for business text overrides

        # DISPUTE_RESOLUTION_CLAUSE
        dispute_key = v.get("DISPUTE_RESOLUTION_VENUE", DEFAULT_DISPUTE_RESOLUTION_VENUE)
        dispute_variants = self.clauses.get("DISPUTE_RESOLUTION_CLAUSE", {}).get("variants", {})
        if dispute_variants:
            v["DISPUTE_RESOLUTION_CLAUSE"] = dispute_variants.get(
                dispute_key,
                dispute_variants.get(DEFAULT_DISPUTE_RESOLUTION_VENUE, {}),
            ).get("text", v.get("DISPUTE_RESOLUTION_CLAUSE", ""))

        # FUND_PROFIT_DEF (for HK_CONNECT)
        fp_variants = self.clauses["FUND_PROFIT_DEF"]["variants"]
        if v.get("HAS_HK_CONNECT"):
            v["FUND_PROFIT_DEF"] = fp_variants["HK_CONNECT"]["text"]
        else:
            v["FUND_PROFIT_DEF"] = fp_variants["STANDARD"]["text"]

        v["INVESTMENT_SCOPE_TEXT"] = self._default_investment_scope_text(v)
        v["COMPONENT_STOCK_STRATEGY_TEXT"] = self._component_stock_strategy_text(v)
        v["RISK_RETURN_CHARACTERISTICS_TEXT"] = self._risk_return_characteristics_text(v)

        return v

    def _process_conditionals(self, text: str, v: dict) -> str:
        # We process from innermost outward using a stack-based approach
        max_passes = 10
        for _ in range(max_passes):
            new_text = self._single_pass_conditionals(text, v)
            if new_text == text:
                break
            text = new_text
        return text

    def _single_pass_conditionals(self, text: str, v: dict) -> str:
        # Match innermost {{IF ...}}...{{ENDIF}} (no nested IF inside)
        pattern_if = re.compile(
            r'\{\{IF\s+(\w+)\}\}((?:(?!\{\{IF)[\s\S])*?)\{\{ENDIF\}\}',
            re.DOTALL
        )
        pattern_if_not = re.compile(
            r'\{\{IF[_\s]+NOT\s+(\w+)\}\}((?:(?!\{\{IF)[\s\S])*?)\{\{ENDIF\}\}',
            re.DOTALL
        )

        def replace_if(m):
            var_name = m.group(1)
            content = m.group(2)
            val = v.get(var_name, False)
            if isinstance(val, str):
                val = val.lower() in ("true", "1", "yes", "是")
            return content if val else ""

        def replace_if_not(m):
            var_name = m.group(1)
            content = m.group(2)
            val = v.get(var_name, False)
            if isinstance(val, str):
                val = val.lower() in ("true", "1", "yes", "是")
            return "" if val else content

        text = pattern_if.sub(replace_if, text)
        text = pattern_if_not.sub(replace_if_not, text)
        return text

    # ── Step 5: 替换占位符 ───────────────────────────────────────────────────
    def _replace_placeholders(self, text: str, v: dict) -> str:
        def replacer(m):
            key = m.group(1)
            val = v.get(key)
            if val is None:
                return m.group(0)  # 保留未定义的占位符
            if isinstance(val, bool):
                return "是" if val else "否"
            return str(val)

        return re.sub(r"\{([A-Z_][A-Z0-9_]*)\}", replacer, text)

    def _definition_tail_from_variant(self, variant_condition: str, include_initiating_terms: bool) -> str:
        template_text = TEMPLATE_MD.read_text(encoding="utf-8")
        pattern = re.compile(
            rf"\{{\{{IF {re.escape(variant_condition)}\}}\}}([\s\S]*?)\{{\{{ENDIF\}}\}}"
        )
        variant_body = None
        for match in pattern.finditer(template_text):
            body = match.group(1)
            if "18、" in body and "基金合同当事人" in body:
                variant_body = body
                break
        if variant_body is None:
            return ""

        lines = []
        for line in variant_body.strip("\n").splitlines():
            stripped = line.strip()
            if "投资人、投资者：" in stripped and not include_initiating_terms:
                line = re.sub(r"、?发起资金提供方、?", "、", line)
                line = line.replace("合格境外投资者、以及", "合格境外投资者以及")
            if not include_initiating_terms and stripped.startswith(
                (
                    "64、发起式基金：",
                    "65、发起式基金：",
                    "65、发起资金：",
                    "66、发起资金：",
                    "66、发起资金提供方：",
                    "67、发起资金提供方：",
                )
            ):
                continue
            lines.append(line)
        return "\n".join(lines).strip()

    def _ensure_complete_definition_section(self, text: str, v: dict) -> str:
        if not v.get("COND_STANDARD_NON_CHUANGYE_NON_KECHUANG_NON_HK"):
            return text

        lines = text.splitlines()
        start = None
        for index, line in enumerate(lines):
            if "第二部分" in line and "释义" in line and any(
                "在本基金合同中" in candidate for candidate in lines[index : index + 8]
            ):
                start = index
                break
        if start is None:
            return text

        end = next(
            (
                index
                for index in range(start + 1, len(lines))
                if "第三部分" in lines[index]
            ),
            None,
        )
        if end is None:
            return text

        definition_count = sum(
            1 for line in lines[start:end] if re.match(r"\s*\d+、", line)
        )
        if definition_count > 17:
            return text

        marker = next(
            (
                index
                for index in range(start, end)
                if "以上释义中涉及法律法规" in lines[index]
            ),
            None,
        )
        if marker is None:
            return text

        variant = (
            "COND_INITIATING_BANK_CROSS_NON_HK"
            if v.get("CUSTODIAN_TYPE") == "商业银行"
            else "COND_INITIATING_SECURITIES_CUSTODIAN"
        )
        tail = self._definition_tail_from_variant(
            variant, include_initiating_terms=bool(v.get("IS_INITIATING_FUND"))
        )
        if not tail:
            return text
        tail = self._process_conditionals(tail, v)
        tail = self._replace_placeholders(tail, v)

        new_lines = lines[:marker] + [""] + tail.splitlines() + [""] + lines[marker:]
        return "\n".join(new_lines)

    def _ensure_initiating_definition_terms(self, text: str, v: dict) -> str:
        if not v.get("IS_INITIATING_FUND"):
            return text

        lines = text.splitlines()
        start = None
        for index, line in enumerate(lines):
            if "第二部分" in line and "释义" in line and any(
                "在本基金合同中" in candidate for candidate in lines[index : index + 8]
            ):
                start = index
                break
        if start is None:
            return text

        end = next(
            (
                index
                for index in range(start + 1, len(lines))
                if "第三部分" in lines[index]
            ),
            None,
        )
        if end is None:
            return text

        for index in range(start, end):
            if "投资人、投资者：" in lines[index] and "发起资金提供方" not in lines[index]:
                lines[index] = lines[index].replace(
                    "合格境外投资者以及法律法规",
                    "合格境外投资者、发起资金提供方以及法律法规",
                )
                break

        section = "\n".join(lines[start:end])
        if "发起式基金：指符合《运作办法》" in section:
            return "\n".join(lines)

        numbers = [
            int(match.group(1))
            for line in lines[start:end]
            for match in [re.match(r"\s*(\d+)、", line)]
            if match
        ]
        next_number = (max(numbers) if numbers else 0) + 1
        initiating_terms = [
            f"{next_number + offset}、{term}"
            for offset, term in enumerate(self._initiating_definition_terms())
        ]
        if not initiating_terms:
            return "\n".join(lines)
        marker = next(
            (
                index
                for index in range(start, end)
                if "以上释义中涉及法律法规" in lines[index]
            ),
            end,
        )
        return "\n".join(lines[:marker] + [""] + initiating_terms + [""] + lines[marker:])

    def _default_investment_scope_text(self, v: dict) -> str:
        if v.get("HAS_HK_CONNECT"):
            variant = "HK_CONNECT"
        elif v.get("IS_KECHUANG_FUND"):
            variant = "KECHUANG"
        else:
            variant = "STANDARD"
        return self._contract_clause_text("INVESTMENT_SCOPE_TEXT", variant)

    def _ensure_investment_scope_sections(self, text: str, v: dict) -> str:
        scope_text = self._default_investment_scope_text(v)
        lines = text.splitlines()

        def insert_after_heading(section_start_match, heading_text, block_end_match, section_end_match):
            nonlocal lines
            starts = [i for i, line in enumerate(lines) if section_start_match(line)]
            for start in starts:
                end = next(
                    (i for i in range(start + 1, len(lines)) if section_end_match(lines[i])),
                    len(lines),
                )
                heading = next(
                    (i for i in range(start, end) if heading_text in lines[i]),
                    None,
                )
                if heading is None:
                    continue
                block_end = next(
                    (i for i in range(heading + 1, end) if block_end_match(lines[i])),
                    end,
                )
                if any("本基金主要投资于目标ETF基金份额、标的指数成份股、备选成份股" in line for line in lines[heading:block_end]):
                    return
                lines = lines[: heading + 1] + ["", scope_text] + lines[heading + 1 :]
                return

        insert_after_heading(
            lambda line: "第十二部分" in line and "基金的投资" in line,
            "二、投资范围",
            lambda line: line.strip().startswith("三、投资策略"),
            lambda line: line.startswith("第十三部分"),
        )
        insert_after_heading(
            lambda line: "五、基金财产的投资范围和投资限制" in line,
            "（一）投资范围",
            lambda line: line.strip().startswith("（二）投资限制"),
            lambda line: line.startswith("六、基金资产净值的计算方法和公告方式"),
        )
        return "\n".join(lines)

    def _component_stock_strategy_text(self, v: dict) -> str:
        variant = "HK_CONNECT" if v.get("HAS_HK_CONNECT") else "STANDARD"
        return self._contract_clause_text("COMPONENT_STOCK_STRATEGY", variant)

    def _risk_return_characteristics_text(self, v: dict) -> str:
        if v.get("HAS_HK_CONNECT"):
            variant = "HK_CONNECT"
        elif v.get("IS_CHUANGYE_FUND"):
            variant = "CHUANGYE"
        elif v.get("IS_KECHUANG_FUND"):
            variant = "KECHUANG"
        else:
            variant = "BASE"
        return self._contract_clause_text("RISK_RETURN_CHARACTERISTICS", variant)

    def _ensure_investment_strategy_and_risk_sections(self, text: str, v: dict) -> str:
        lines = text.splitlines()

        start = next(
            (i for i, line in enumerate(lines) if "第十二部分" in line and "基金的投资" in line and "\t" not in line),
            None,
        )
        if start is None:
            return text
        end = next((i for i in range(start + 1, len(lines)) if lines[i].startswith("第十三部分")), len(lines))

        strategy_heading = next(
            (i for i in range(start, end) if "成份股、备选成份股投资策略" in lines[i]),
            None,
        )
        if strategy_heading is not None:
            next_strategy = next(
                (i for i in range(strategy_heading + 1, end) if re.match(r"^\d+、", lines[i].strip())),
                end,
            )
            if not any("本基金对成份股、备选成份股的投资目的是" in line or "本基金可以通过投资成份股、备选成份股" in line for line in lines[strategy_heading:next_strategy]):
                lines = (
                    lines[: strategy_heading + 1]
                    + ["", self._component_stock_strategy_text(v)]
                    + lines[strategy_heading + 1 :]
                )
                inserted = 2
                end += inserted

        section_has_risk = any(
            "本基金的目标ETF为股票型基金" in line
            for line in lines[start:end]
        )
        if not section_has_risk:
            insert_before = next(
                (i for i in range(start, end) if "基金管理人代表基金行使股东或债权人权利" in lines[i]),
                None,
            )
            if insert_before is not None:
                lines = (
                    lines[:insert_before]
                    + ["六、风险收益特征", "", self._risk_return_characteristics_text(v), ""]
                    + lines[insert_before:]
                )

        return "\n".join(lines)

    # ── Step 5b: 重排阿拉伯序号（修复条件块删除项目后的跳跃）──────────────
    def _renumber_sequences(self, text: str) -> str:
        """
        修复条件块删除项目后阿拉伯序号列表的跳跃。
        规则：
          - 只处理行首 `数字、` 格式
          - num == 1：新列表开始，重置计数器，不修改
          - num > expected_next（且 num > 1）：检测到跳跃，将该行及后续连续项重排
          - 遇到中文序号标题（一、二、…）或章节标题（第X部分）时重置计数器
        """
        lines = text.split("\n")
        RE_NUM = re.compile(r"^(\d+)(、)")
        RE_RESET = re.compile(
            r"^(?:[一二三四五六七八九十百]+、|第[一二三四五六七八九十百]+部分)"
        )
        last_num = None
        result = []
        for line in lines:
            # 章节/中文序号 → 重置
            if RE_RESET.match(line.strip()):
                last_num = None
                result.append(line)
                continue
            m = RE_NUM.match(line)
            if m:
                num = int(m.group(1))
                sep = m.group(2)
                if num == 1:
                    # 新列表起点，不修改，重置计数器
                    last_num = 1
                elif last_num is not None and num > last_num + 1:
                    # 条件跳跃：把该项编号改为期望值
                    expected = last_num + 1
                    line = re.sub(r"^\d+、", f"{expected}、", line, count=1)
                    last_num = expected
                else:
                    last_num = num
            result.append(line)
        return "\n".join(result)

    # ── Step 6: 清理 ─────────────────────────────────────────────────────────
    def _cleanup(self, text: str) -> str:
        # Remove markdown blockquote header lines (> ...)
        lines = text.split("\n")
        clean = []
        skip_header = True
        for line in lines:
            # Skip the preamble before the actual contract starts
            if skip_header:
                if line.strip().startswith("{FUND_NAME}") or (
                    "基金合同" in line and "模板说明" not in line and not line.startswith(">")
                    and not line.startswith("#") and not line.startswith("**") and not line.startswith("---")
                ):
                    skip_header = False
                    clean.append(line)
                # Skip preamble lines
                continue
            # Remove comment/annotation lines
            if line.strip().startswith(">") or line.strip().startswith("**条件变量") or line.strip().startswith("**差异条款"):
                continue
            clean.append(line)

        text = "\n".join(clean)

        # Collapse 3+ consecutive blank lines into 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        return text

    @staticmethod
    def _normalize_contract_text(text: str) -> str:
        return text.replace("具体办理时间为指", "具体办理时间为")

    @staticmethod
    def _extract_contract_summary_block_body(section_text: str, label: str) -> str:
        blocks = _split_review_blocks(section_text)
        label_text = _normalize_review_text(label)
        if not blocks or not label_text:
            return ""

        exact_index = next(
            (
                index
                for index, block in enumerate(blocks)
                if _normalize_review_text(block.get("heading", "")) == label_text
            ),
            None,
        )
        if exact_index is not None:
            expanded = _expand_review_block_with_descendants(blocks, exact_index)
            return (expanded.get("body") or expanded.get("text") or "").strip()

        fallback_indexes = []
        for index, block in enumerate(blocks):
            heading = block.get("heading", "")
            if _review_labels_match(heading, label) or _review_soft_heading_match(heading, label):
                fallback_indexes.append((len(_normalize_review_text(heading)), index))
        if fallback_indexes:
            _, best_index = max(fallback_indexes)
            expanded = _expand_review_block_with_descendants(blocks, best_index)
            return (expanded.get("body") or expanded.get("text") or "").strip()
        return ""

    @staticmethod
    def _extract_contract_summary_intro_before(section_text: str, stop_label: str) -> str:
        blocks = _split_review_blocks(section_text)
        stop_label_text = _normalize_review_text(stop_label)
        if not blocks or not section_text or not stop_label_text:
            return ""

        target_index = next(
            (
                index
                for index, block in enumerate(blocks)
                if _normalize_review_text(block.get("heading", "")) == stop_label_text
            ),
            None,
        )
        if target_index is None:
            fallback_indexes = []
            for index, block in enumerate(blocks):
                heading = block.get("heading", "")
                if _review_labels_match(heading, stop_label) or _review_soft_heading_match(heading, stop_label):
                    fallback_indexes.append((len(_normalize_review_text(heading)), index))
            if fallback_indexes:
                _, target_index = max(fallback_indexes)
        if target_index is None:
            return ""

        block_text = blocks[target_index].get("text") or ""
        if block_text and block_text in section_text:
            return section_text.split(block_text, 1)[0].strip()
        return ""

    def _extract_contract_summary_source_text(self, target: dict, locator: str) -> str:
        segments = _split_review_locator(_normalize_review_locator(locator))
        sub_labels = segments[1:] if segments and ("第" in segments[0] or "部分" in segments[0]) else segments
        section = target.get("section") or {}
        if sub_labels and section.get("content"):
            expanded_body = self._extract_contract_summary_block_body(section.get("content") or "", sub_labels[-1])
            if expanded_body:
                return expanded_body
        return (target.get("body_text") or target.get("text") or "").strip()

    @staticmethod
    def _contract_summary_heading_marker(heading: str) -> str:
        match = re.match(r"^\s*([一二三四五六七八九十百千]+)[、.．]", str(heading or ""))
        return match.group(1) if match else ""

    @staticmethod
    def _contract_summary_parenthesized_heading_marker(heading: str) -> str:
        match = re.match(r"^\s*[（(]([一二三四五六七八九十百千]+)[）)]", str(heading or ""))
        return match.group(1) if match else ""

    @staticmethod
    def _rewrite_contract_summary_internal_references(source_text: str, locator: str, output_heading: str | None) -> str:
        if not source_text or not output_heading:
            return source_text
        segments = _split_review_locator(_normalize_review_locator(locator))
        if not segments:
            return source_text
        source_heading = segments[-1]
        if not source_heading or source_heading == output_heading:
            return source_text

        rewritten = source_text
        for open_quote, close_quote in (("“", "”"), ("‘", "’"), ('"', '"')):
            rewritten = rewritten.replace(
                f"{open_quote}{source_heading}{close_quote}",
                f"{open_quote}{output_heading}{close_quote}",
            )

        source_marker = ContractEngine._contract_summary_heading_marker(source_heading)
        output_marker = ContractEngine._contract_summary_parenthesized_heading_marker(output_heading)
        if source_marker and output_marker and source_marker == output_marker:
            rewritten = re.sub(
                rf"第{re.escape(source_marker)}条",
                f"第（{output_marker}）条",
                rewritten,
            )
        if segments and "基金份额持有人大会" in segments[0]:
            rewritten = re.sub(
                r"(下列|本部分|上述|前述)第([一二三四五六七八九十百千]+)条",
                r"\1第（\2）条",
                rewritten,
            )
        return rewritten

    @staticmethod
    def _normalize_contract_summary_numbered_heading_references(text: str) -> str:
        if not text:
            return text

        text = re.sub(
            r"([“‘\"])([一二三四五六七八九十百千]+)[、.．]([^”’\"\n]+)([”’\"])",
            lambda m: f"{m.group(1)}（{m.group(2)}）{m.group(3)}{m.group(4)}",
            text,
        )
        return re.sub(
            r"([“‘\"])[（(]([一二三四五六七八九十百千]+)[）)][、.．]([^”’\"\n]+)([”’\"])",
            lambda m: f"{m.group(1)}（{m.group(2)}）{m.group(3)}{m.group(4)}",
            text,
        )

    def _locate_contract_summary_source(self, sections: list[dict], locator: str, placeholder: str) -> dict:
        target = _locate_review_rule_target(sections, locator)
        if not target.get("matched"):
            reason = target.get("missing_reason") or "未匹配"
            raise ValueError(f"基金合同摘要占位符 {placeholder} 无法定位正文来源：{locator}（{reason}）")
        source_text = self._extract_contract_summary_source_text(target, locator)
        if not source_text:
            raise ValueError(f"基金合同摘要占位符 {placeholder} 的正文来源为空：{locator}")
        target = dict(target)
        target["summary_source_text"] = source_text
        return target

    def _render_contract_summary_fragment(self, sections: list[dict], placeholder: str, fragment: dict) -> str:
        locator = str(fragment.get("locator") or "").strip()
        target = self._locate_contract_summary_source(sections, locator, placeholder)
        source_text = str(target.get("summary_source_text") or "").strip()
        if fragment.get("take") == "intro":
            return self._extract_contract_summary_intro_before(
                source_text,
                str(fragment.get("stop_before") or ""),
            )

        heading = fragment.get("heading")
        source_text = self._rewrite_contract_summary_internal_references(source_text, locator, heading)
        if heading:
            source_heading = (_split_review_locator(_normalize_review_locator(locator)) or [""])[-1]
            if _normalize_review_text(source_text) == _normalize_review_text(source_heading):
                return str(heading).strip()
            return f"{heading}\n\n{source_text}".strip()
        return source_text

    def _render_contract_summary_spec(self, sections: list[dict], spec: dict) -> str:
        placeholder = spec.get("placeholder") or ""
        rendered_fragments = [
            self._render_contract_summary_fragment(sections, placeholder, fragment)
            for fragment in spec.get("fragments", ())
        ]
        rendered_fragments = [fragment for fragment in rendered_fragments if fragment.strip()]
        if not rendered_fragments:
            raise ValueError(f"基金合同摘要占位符 {placeholder} 未生成任何内容")
        rendered = "\n\n".join(rendered_fragments).strip()
        for fragment in spec.get("fragments", ()):
            rendered = self._rewrite_contract_summary_internal_references(
                rendered,
                str(fragment.get("locator") or ""),
                fragment.get("heading"),
            )
        rendered = self._rewrite_contract_summary_internal_references(
            rendered,
            str(spec.get("locator") or ""),
            spec.get("heading"),
        )
        return self._normalize_contract_summary_numbered_heading_references(rendered)

    def _replace_contract_summary_placeholders(self, text: str) -> str:
        placeholders = set(self._CONTRACT_SUMMARY_PLACEHOLDER_RE.findall(text or ""))
        if not placeholders:
            return text

        specs_by_placeholder = {
            spec["placeholder"]: spec
            for spec in self.CONTRACT_SUMMARY_SPECS
        }
        unknown = sorted(placeholders - set(specs_by_placeholder))
        if unknown:
            raise ValueError(f"未知基金合同摘要占位符：{', '.join(unknown)}")

        sections = _split_contract_sections(text)
        rendered_by_placeholder = {
            placeholder: self._render_contract_summary_spec(sections, specs_by_placeholder[placeholder])
            for placeholder in placeholders
        }

        def replace_summary_placeholder(match):
            return rendered_by_placeholder[match.group(1)]

        rendered = self._CONTRACT_SUMMARY_PLACEHOLDER_RE.sub(replace_summary_placeholder, text)
        remaining = sorted(set(self._CONTRACT_SUMMARY_PLACEHOLDER_RE.findall(rendered)))
        if remaining:
            raise ValueError(f"基金合同摘要仍残留占位符：{', '.join(remaining)}")
        return rendered

    # ── 主方法：生成合同 ─────────────────────────────────────────────────────
    def generate(self, form_data: dict) -> str:
        # Step 1
        v = self._derive_variables(form_data)
        # Step 2
        v = self._inject_clause_texts(v)
        # Step 5: read template
        template_text = TEMPLATE_MD.read_text(encoding="utf-8")
        # Step 6
        text = self._process_conditionals(template_text, v)
        # Step 7
        text = self._replace_placeholders(text, v)
        text = self._ensure_complete_definition_section(text, v)
        text = self._ensure_initiating_definition_terms(text, v)
        text = self._ensure_investment_scope_sections(text, v)
        text = self._ensure_investment_strategy_and_risk_sections(text, v)
        # Step 8: 修复条件删除导致的序号跳跃
        text = self._renumber_sequences(text)
        # Step 9
        text = self._cleanup(text)
        text = self._normalize_contract_text(text)
        text = self._replace_contract_summary_placeholders(text)
        text = self._normalize_contract_text(text)
        return text

    # ── 标准化中文标点 ────────────────────────────────────────────────────────
    @staticmethod
    def _to_chinese_punct(text: str) -> str:
        """将合同中常见的半角/英文标点替换为中文全角标点（数字间的小数点除外）。"""
        import re as _re
        # 半角逗号/句号（非小数点）→ 中文全角（谨慎，仅ASCII区间）
        # 此函数仅做最保守替换，避免破坏数字、代码等内容
        replacements = [
            (r'(?<!\d)\.(?!\d)', '。'),  # 非数字间的句点 → 句号
        ]
        for pat, rep in replacements:
            text = _re.sub(pat, rep, text)
        return text

    # ── Word (.docx) 导出 ────────────────────────────────────────────────────
    def build_docx(self, contract_text: str) -> bytes:
        """
        将合同纯文本转换为格式化 Word 文档，与参考合同格式完全对齐。
        - 封面：独立 section + vAlign=center，彻底消除内容溢出第二页的问题
        - 签署页：左对齐 + 精确空行间距（来自参考文档XML）
        - 章节标题：Times New Roman 15pt 加粗，pageBreakBefore
        - 正文：宋体 12pt（中文），Times New Roman 12pt（英文/数字），两端对齐，首行缩进2字，1.5倍行距
        """
        def _normalize_contract_cover_text(text: str) -> str:
            lines = str(text or "").split("\n")
            toc_index = next((idx for idx, line in enumerate(lines) if re.match(r"^目\s*录$", line.strip())), None)
            cover_end = toc_index if toc_index is not None else len(lines)
            cover_lines = lines[:cover_end]
            tail_lines = lines[cover_end:]
            normalized_cover = []
            title_done = False
            for raw in cover_lines:
                stripped = raw.strip()
                if stripped and re.fullmatch(r"\d{4}年\d{1,2}月\d{1,2}日", stripped):
                    continue
                if stripped and not title_done and "基金合同" in stripped:
                    stripped = re.sub(r"基金合同(?:（草案）)?$", "基金合同（草案）", stripped)
                    title_done = True
                    normalized_cover.append(stripped)
                    continue
                normalized_cover.append(raw)
            return "\n".join([*normalized_cover, *tail_lines])

        contract_text = _normalize_contract_cover_text(contract_text)
        from pathlib import Path
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Twips

        # ── 1. 文档与页面设置 ──────────────────────────────────────────────
        template_docx = _require_linked_format_template(
            LINKED_CONTRACT_FORMAT_TEMPLATE_FILENAME,
            "CONTRACT_TEMPLATE_DIR",
            "基金合同",
        )
        doc = Document(str(template_docx))

        def _collect_signature_format_refs(document):
            refs = {}
            in_signature_block = False
            for paragraph in document.paragraphs:
                text_value = (paragraph.text or "").strip()
                if "无正文" in text_value and ("签字页" in text_value or "签署页" in text_value):
                    key = "title"
                    in_signature_block = True
                elif not in_signature_block:
                    continue
                elif text_value.startswith("基金管理人"):
                    key = "manager"
                elif text_value.startswith("基金托管人"):
                    key = "custodian"
                elif text_value.startswith("法定代表人"):
                    key = "representative"
                elif text_value.startswith("签订地"):
                    key = "place"
                elif text_value.startswith("签订日"):
                    key = "date"
                else:
                    continue
                first_run = next((run for run in paragraph.runs if run.text), None)
                refs.setdefault(key, {
                    "pPr": copy.deepcopy(paragraph._p.pPr) if paragraph._p.pPr is not None else None,
                    "rPr": copy.deepcopy(first_run._r.rPr) if first_run is not None and first_run._r.rPr is not None else None,
                })
            return refs

        signature_format_refs = _collect_signature_format_refs(doc)

        body = doc._element.body
        cover_sectpr_template = None
        for child in body:
            if child.tag != qn("w:p"):
                continue
            p_pr = child.find(qn("w:pPr"))
            if p_pr is None:
                continue
            sect_pr = p_pr.find(qn("w:sectPr"))
            if sect_pr is not None:
                cover_sectpr_template = copy.deepcopy(sect_pr)
                break
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
        sec = doc.sections[0]
        sec.page_width        = Twips(PAGE_LAYOUT["page_width"])
        sec.page_height       = Twips(PAGE_LAYOUT["page_height"])
        sec.top_margin        = Twips(PAGE_LAYOUT["top_margin"])
        sec.bottom_margin     = Twips(PAGE_LAYOUT["bottom_margin"])
        sec.left_margin       = Twips(PAGE_LAYOUT["left_margin"])
        sec.right_margin      = Twips(PAGE_LAYOUT["right_margin"])
        sec.header_distance   = Twips(PAGE_LAYOUT["header_distance"])
        sec.footer_distance   = Twips(PAGE_LAYOUT["footer_distance"])
        sectPr = sec._sectPr
        for node in list(sectPr.findall(qn("w:docGrid"))):
            sectPr.remove(node)
        docGrid = OxmlElement("w:docGrid")
        docGrid.set(qn("w:type"), "lines")
        docGrid.set(qn("w:linePitch"), "312")
        sectPr.append(docGrid)
        settings = doc.settings.element
        default_tab_stop = settings.find(qn("w:defaultTabStop"))
        if default_tab_stop is None:
            default_tab_stop = OxmlElement("w:defaultTabStop")
            settings.insert(0, default_tab_stop)
        default_tab_stop.set(qn("w:val"), "720")

        # ── 2. Normal 样式：中文宋体/英数 Times New Roman 12pt，两端对齐，孤行控制关闭 ──
        normal_style = doc.styles["Normal"]
        normal_rpr = normal_style.element.get_or_add_rPr()
        for old in normal_rpr.findall(qn("w:rFonts")):
            normal_rpr.remove(old)
        for old in normal_rpr.findall(qn("w:sz")):
            normal_rpr.remove(old)
        for old in normal_rpr.findall(qn("w:szCs")):
            normal_rpr.remove(old)
        rFonts_n = OxmlElement("w:rFonts")
        rFonts_n.set(qn("w:ascii"),    "Times New Roman")
        rFonts_n.set(qn("w:hAnsi"),    "Times New Roman")
        rFonts_n.set(qn("w:eastAsia"), "宋体")
        normal_rpr.insert(0, rFonts_n)
        sz_n = OxmlElement("w:sz"); sz_n.set(qn("w:val"), "24"); normal_rpr.append(sz_n)
        normal_ppr = normal_style.element.get_or_add_pPr()
        _clear_xml_children(normal_ppr)
        wc = OxmlElement("w:widowControl"); wc.set(qn("w:val"), "0"); normal_ppr.insert(0, wc)
        jc_n = OxmlElement("w:jc"); jc_n.set(qn("w:val"), "both"); normal_ppr.append(jc_n)
        sp_n = OxmlElement("w:spacing")
        sp_n.set(qn("w:line"), "360"); sp_n.set(qn("w:lineRule"), "auto")
        sp_n.set(qn("w:before"), "0"); sp_n.set(qn("w:after"), "0")
        normal_ppr.append(sp_n)
        _ensure_word_toc_styles(
            doc,
            heading_eastasia="宋体",
            heading_ascii="宋体",
            heading_size_half_pt=28,
            heading_bold=False,
            entry_eastasia="宋体",
            entry_ascii="Times New Roman",
            entry_size_half_pt=24,
            max_level=1,
        )

        # ── 3. XML 辅助函数 ───────────────────────────────────────────────
        def _set_para(p, jc=None, line=None, line_rule="auto",
                      before=None, after=None,
                      first_line=None, first_line_chars=None, left_ind=None,
                      keep_lines=False, page_break_before=False,
                      snap_to_grid=None):
            pPr = p._p.get_or_add_pPr()
            if keep_lines:
                pPr.append(OxmlElement("w:keepLines"))
            if page_break_before:
                pPr.append(OxmlElement("w:pageBreakBefore"))
            if snap_to_grid is not None:
                sg = OxmlElement("w:snapToGrid")
                sg.set(qn("w:val"), "1" if snap_to_grid else "0")
                pPr.append(sg)
            if before is not None or after is not None or line is not None:
                sp = OxmlElement("w:spacing")
                if before is not None: sp.set(qn("w:before"), str(before))
                if after  is not None: sp.set(qn("w:after"),  str(after))
                if line   is not None:
                    sp.set(qn("w:line"),     str(line))
                    sp.set(qn("w:lineRule"), line_rule)
                pPr.append(sp)
            if first_line is not None or first_line_chars is not None or left_ind is not None:
                ind = OxmlElement("w:ind")
                if left_ind          is not None: ind.set(qn("w:left"),           str(left_ind))
                if first_line_chars  is not None: ind.set(qn("w:firstLineChars"), str(first_line_chars))
                if first_line        is not None: ind.set(qn("w:firstLine"),      str(first_line))
                pPr.append(ind)
            if jc is not None:
                jc_el = OxmlElement("w:jc"); jc_el.set(qn("w:val"), jc); pPr.append(jc_el)

        def _set_run(r, ascii_font=None, eastasia_font=None, hint=None,
                     sz=None, sz_cs=None, bold=False, bcs=False, color=None):
            rPr = r._r.get_or_add_rPr()
            if ascii_font or eastasia_font or hint:
                rF = OxmlElement("w:rFonts")
                if ascii_font:    rF.set(qn("w:ascii"), ascii_font); rF.set(qn("w:hAnsi"), ascii_font)
                if eastasia_font: rF.set(qn("w:eastAsia"), eastasia_font)
                if hint:          rF.set(qn("w:hint"), hint)
                rPr.insert(0, rF)
            if bold: rPr.append(OxmlElement("w:b"))
            if bcs:  rPr.append(OxmlElement("w:bCs"))
            if color:
                col = OxmlElement("w:color"); col.set(qn("w:val"), color); rPr.append(col)
            if sz is not None:
                s = OxmlElement("w:sz"); s.set(qn("w:val"), str(sz)); rPr.append(s)
            if sz_cs is not None:
                sc = OxmlElement("w:szCs"); sc.set(qn("w:val"), str(sz_cs)); rPr.append(sc)

        def _apply_signature_format(paragraph, run, format_key: str, page_break_before=False):
            refs = signature_format_refs.get(format_key) or {}
            ref_ppr = refs.get("pPr")
            if ref_ppr is not None:
                current_ppr = paragraph._p.pPr
                if current_ppr is not None:
                    paragraph._p.remove(current_ppr)
                paragraph._p.insert(0, copy.deepcopy(ref_ppr))
            if page_break_before:
                pPr = paragraph._p.get_or_add_pPr()
                if pPr.find(qn("w:pageBreakBefore")) is None:
                    pPr.append(OxmlElement("w:pageBreakBefore"))
            ref_rpr = refs.get("rPr")
            if ref_rpr is not None:
                current_rpr = run._r.rPr
                if current_rpr is not None:
                    run._r.remove(current_rpr)
                run._r.insert(0, copy.deepcopy(ref_rpr))

        def _signing_empty(n=1):
            """签署页专用空行：左对齐，snapToGrid=0，宋体，bCs，与参考文档一致。"""
            for _ in range(n):
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                sg = OxmlElement("w:snapToGrid"); sg.set(qn("w:val"), "0"); pPr.append(sg)
                sp = OxmlElement("w:spacing")
                sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
                sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
                pPr.append(sp)
                jc_el = OxmlElement("w:jc"); jc_el.set(qn("w:val"), "left"); pPr.append(jc_el)
                mrPr = OxmlElement("w:rPr")
                rF2 = OxmlElement("w:rFonts")
                rF2.set(qn("w:ascii"), "Times New Roman"); rF2.set(qn("w:hAnsi"), "Times New Roman")
                mrPr.append(rF2)
                mrPr.append(OxmlElement("w:bCs"))
                szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "21"); mrPr.append(szCs)
                pPr.append(mrPr)

        def _update_contract_header_text(header_text: str, sections):
            if not sections:
                return
            for section in sections:
                for header in (section.header, section.first_page_header):
                    paragraph = _reset_header_footer_part(header)
                    pPr = paragraph._p.get_or_add_pPr()
                    _clear_xml_children(pPr)
                    jc = OxmlElement("w:jc")
                    jc.set(qn("w:val"), "right")
                    pPr.append(jc)
                    p_bdr = OxmlElement("w:pBdr")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:color"), "auto")
                    bottom.set(qn("w:sz"), "6")
                    bottom.set(qn("w:space"), "1")
                    p_bdr.append(bottom)
                    pPr.append(p_bdr)
                    run = paragraph.add_run(header_text)
                    _set_run(run, ascii_font="Times New Roman", eastasia_font="宋体", sz=18, sz_cs=18)

        def _set_section_doc_grid(section, line_pitch: str):
            sect_pr = section._sectPr
            for node in list(sect_pr.findall(qn("w:docGrid"))):
                sect_pr.remove(node)
            doc_grid = OxmlElement("w:docGrid")
            doc_grid.set(qn("w:type"), "lines")
            doc_grid.set(qn("w:linePitch"), str(line_pitch))
            sect_pr.append(doc_grid)

        def _cover_section_break():
            """
            在封面最后一段的 pPr 中嵌入 sectPr（封面 section 定义）。
            使用 vAlign=center 确保内容永远不溢出，并通过 nextPage 分隔目录。
            """
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            # 封面 section 的属性
            cover_sp = copy.deepcopy(cover_sectpr_template) if cover_sectpr_template is not None else OxmlElement("w:sectPr")
            for node in list(cover_sp.findall(qn("w:type"))):
                cover_sp.remove(node)
            # 分节类型：下一页
            t = OxmlElement("w:type"); t.set(qn("w:val"), "nextPage"); cover_sp.append(t)
            # 页面尺寸（A4）
            for node in list(cover_sp.findall(qn("w:pgSz"))):
                cover_sp.remove(node)
            pgSz = OxmlElement("w:pgSz")
            pgSz.set(qn("w:w"), "11906"); pgSz.set(qn("w:h"), "16838")
            cover_sp.append(pgSz)
            # 页边距（与正文一致）
            for node in list(cover_sp.findall(qn("w:pgMar"))):
                cover_sp.remove(node)
            pgMar = OxmlElement("w:pgMar")
            pgMar.set(qn("w:top"),    "1440"); pgMar.set(qn("w:right"),  "1800")
            pgMar.set(qn("w:bottom"), "1440"); pgMar.set(qn("w:left"),   "1800")
            pgMar.set(qn("w:header"), "851");  pgMar.set(qn("w:footer"), "992")
            pgMar.set(qn("w:gutter"), "0")
            cover_sp.append(pgMar)
            # 文档网格
            for node in list(cover_sp.findall(qn("w:docGrid"))):
                cover_sp.remove(node)
            dg = OxmlElement("w:docGrid")
            dg.set(qn("w:type"), "lines"); dg.set(qn("w:linePitch"), "312")
            cover_sp.append(dg)
            # 垂直居中——确保内容不管多少行都不溢出第二页
            for node in list(cover_sp.findall(qn("w:vAlign"))):
                cover_sp.remove(node)
            vAlign = OxmlElement("w:vAlign"); vAlign.set(qn("w:val"), "center")
            cover_sp.append(vAlign)
            pPr.append(cover_sp)

        # ── 4. 正则分类 ──────────────────────────────────────────────────
        RE_TOC_ENTRY = re.compile(r"^第[一二三四五六七八九十百]+部分\s+.+\t\d+$")
        RE_PART_HEAD = re.compile(r"^第[一二三四五六七八九十百]+部分\s+\S")
        RE_UPPER_NUM = re.compile(r"^[一二三四五六七八九十百]+、")
        RE_PAREN_CN  = re.compile(r"^（[一二三四五六七八九十百]+）")
        RE_ARAB_DOT  = re.compile(r"^\d+、")
        RE_PAREN_NUM = re.compile(r"^（\d+）")

        # ── 书签计数器与锚点辅助 ────────────────────────────────────────────────
        _bm_id = [0]  # 可变单元格，允许在内层代码中递增

        def _part_anchor(heading_s: str) -> str:
            """从"第X部分..."中提取书签锚名，例如 'part_二十六'"""
            m = re.match(r'^第([一二三四五六七八九十百]+)部分', heading_s)
            return f"part_{m.group(1)}" if m else ""

        # ── 5. 主处理循环 ─────────────────────────────────────────────────
        lines = contract_text.split("\n")
        phase          = "cover"
        cover_idx      = 0      # 非空封面行计数
        signing_started = False
        signing_idx    = 0      # 签署页内容行计数
        pending_signing_date = None
        signing_representative_count = 0
        current_part   = 0      # 当前所在部分编号（用于判断前言是否加空行）

        def _is_signing_page_start(text_value: str) -> bool:
            return "无正文" in text_value and ("签署页" in text_value or "签字页" in text_value)

        def _add_signing_line(text_value: str):
            nonlocal signing_representative_count
            p = doc.add_paragraph()
            _set_para(p, snap_to_grid=False, jc="left", line=360, before=0, after=0)
            r = p.add_run(text_value)
            _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                     hint="eastAsia", sz=24, bcs=True)
            if text_value.startswith("基金管理人"):
                _apply_signature_format(p, r, "manager")
            elif text_value.startswith("基金托管人"):
                _apply_signature_format(p, r, "custodian")
            elif text_value.startswith("法定代表人"):
                _apply_signature_format(p, r, "representative")
            elif text_value.startswith("签订地"):
                _apply_signature_format(p, r, "place")
            elif text_value.startswith("签订日"):
                _apply_signature_format(p, r, "date")
            if "盖章" in text_value:
                _signing_empty(5)
            elif "签字/章" in text_value or "签字或盖章" in text_value or "签名" in text_value:
                signing_representative_count += 1
                _signing_empty(6 if signing_representative_count >= 2 else 5)

        for raw in lines:
            s = raw.strip()

            # 所有空行跳过（签署页内的空行由代码主动添加）
            if not s:
                continue

            # ════ 封面 → 目录 ════
            if phase == "cover" and re.match(r"^目\s*录$", s):
                phase = "toc"
                # 封面 section break（vAlign=center, nextPage）
                _cover_section_break()
                p = doc.add_paragraph(style="TOC Heading")
                _set_para(p, jc="center", line=360, before=0, after=0)
                r = p.add_run("目    录")
                _set_run(r, sz=28, sz_cs=28, bcs=True)
                toc_field = doc.add_paragraph()
                _set_para(toc_field, line=360, before=0, after=0)
                _append_word_toc_field(
                    toc_field,
                    OxmlElement,
                    qn,
                    levels="1-1",
                    placeholder="右键更新目录",
                    ascii_font="Times New Roman",
                    eastasia_font="宋体",
                    size=24,
                )
                continue

            # ════ 封面内容：仅4行，vAlign=center 负责垂直定位 ════
            if phase == "cover":
                p = doc.add_paragraph()
                if cover_idx == 0:
                    # 合同标题：sz=48(24pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=48, sz_cs=48, bold=True, bcs=True)
                    # 标题与管理人之间加6个空行，和ETF合同知识库保持一致
                    for _ in range(6):
                        ep = doc.add_paragraph()
                        _set_para(ep, jc="center", line=360, before=0, after=0)
                elif s.startswith("基金管理人") or s.startswith("基金托管人"):
                    # 管理人/托管人：sz=36(18pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=36, sz_cs=36, bold=True, bcs=True)
                else:
                    # 日期：sz=36(18pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=36, sz_cs=36, bold=True, bcs=True)
                cover_idx += 1
                continue

            # ════ 目录条目 ════
            if phase == "toc":
                if RE_TOC_ENTRY.match(s):
                    continue
                else:
                    phase = "body"

            # ════ 签署页检测 ════
            if phase == "body" and not signing_started and _is_signing_page_start(s):
                signing_started = True
                p = doc.add_paragraph()
                _set_para(p, snap_to_grid=False, jc="left", line=360, before=0, after=0,
                          page_break_before=True)   # 签署页强制分页
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                _apply_signature_format(p, r, "title", page_break_before=True)
                _signing_empty(4)   # 签署页标题后 4 个空行（参考文档）
                signing_idx = 1
                continue

            # ════ 签署页内容行 ════
            if phase == "body" and signing_started:
                if s.startswith("签订日"):
                    pending_signing_date = s
                    signing_idx += 1
                    continue
                _add_signing_line(s)
                if s.startswith("签订地") and pending_signing_date:
                    _add_signing_line(pending_signing_date)
                    pending_signing_date = None
                signing_idx += 1
                continue

            # ════ 部分标题（第X部分…） ════
            if phase == "body" and RE_PART_HEAD.match(s):
                current_part += 1
                p = doc.add_paragraph(style="Heading 1")
                _set_para(p, jc="center", line=360, before=0, after=0,
                          keep_lines=True, page_break_before=True)
                # 添加书签，供目录超链接跳转
                anchor = _part_anchor(s)
                if anchor:
                    _bm_id[0] += 1
                    bm_start = OxmlElement("w:bookmarkStart")
                    bm_start.set(qn("w:id"), str(_bm_id[0]))
                    bm_start.set(qn("w:name"), anchor)
                    bm_end = OxmlElement("w:bookmarkEnd")
                    bm_end.set(qn("w:id"), str(_bm_id[0]))
                    p._p.append(bm_start)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         sz=30, sz_cs=30, bold=True, bcs=False, color="auto")
                if anchor:
                    p._p.append(bm_end)
                continue

            # ════ 一级子标题（一、二、…） ════
            if phase == "body" and RE_UPPER_NUM.match(s):
                # 第一部分（前言）子标题间不加空行，其余部分各子标题前空一行
                if current_part != 1:
                    ep = doc.add_paragraph()
                    _set_para(ep, line=360, before=0, after=0)
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 二级子标题（（一）（二）…） ════
            if phase == "body" and RE_PAREN_CN.match(s):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 数字条款（1、2、… 或（1）（2）…） ════
            if phase == "body" and (RE_ARAB_DOT.match(s) or RE_PAREN_NUM.match(s)):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 普通正文 ════
            if phase in ("body", "toc"):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)

        # ── 6. 序列化 ──────────────────────────────────────────────────────
        body_sections = list(doc.sections[1:]) if len(doc.sections) > 1 else list(doc.sections)
        for section in body_sections:
            _set_section_doc_grid(section, "312")
        _update_contract_header_text("基金合同（草案）", body_sections)
        _finalize_doc_page_numbers(doc, OxmlElement, qn)
        _set_update_fields_on_open(doc, OxmlElement, qn)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

# ═══════════════════════════════════════════════════════════════════════════════
#  ProspectusEngine — 招募说明书生成引擎（8步管线）
# ═══════════════════════════════════════════════════════════════════════════════
class ProspectusEngine:
    def __init__(self, contract_engine=None):
        with open(PROSPECTUS_CLAUSES_JSON, encoding="utf-8") as f:
            data = json.load(f)
            self.clauses = data["clauses"]
        self.contract_engine = contract_engine or ContractEngine()

    # ── Step 1: 变量推导 ──────────────────────────────────────────────────
    def _derive_variables(self, v):
        v = self.contract_engine._derive_variables(v)
        v["SALES_SERVICE_FEE_PAYMENT_METHOD"] = self._sales_service_fee_payment_method_text(v)
        v.setdefault("PRODUCT_TYPE", "ETF联接基金")
        v.setdefault("REDEMPTION_FEE_STYLE", "区分个人与非个人")
        v.setdefault("NON_PERSONAL_REDEMPTION_FEE_STYLE", "长期递减")
        return v

    def _sales_service_fee_payment_method_text(self, v):
        custodian_info = (
            self.contract_engine.clauses
            .get("CUSTODIAN_INFO_KNOWN_VALUES", {})
            .get("custodians", {})
            .get(v.get("CUSTODIAN_NAME", ""), {})
        )
        raw_method = (
            v.get("MGMT_FEE_PAYMENT_METHOD")
            or custodian_info.get("mgmt_fee_payment")
            or custodian_info.get("mgmt_fee_payment_method")
            or custodian_info.get("payment_method")
            or "CONSULT"
        )
        variants = self.contract_engine.clauses["MGMT_FEE_PAYMENT"]["variants"]
        _, method_text = _resolve_payment_clause_value(
            raw_method,
            variants,
            {"AUTO_V2": "AUTO", "AUTO_CHECK": "AUTO"},
            "基金管理费",
        )
        if not method_text:
            method_text = variants["CONSULT"]["text"].strip()
        return self._normalize_sales_service_fee_payment_text(method_text)

    @staticmethod
    def _normalize_sales_service_fee_payment_text(text):
        text = re.sub(r"^基金管理费每日(?:计提|计算)[^。]*按月支付[。，]\s*", "", text.strip())
        text = text.replace("基金管理费", "销售服务费")
        text = text.replace("支付给基金管理人", "进行资金支付")
        text = text.replace("向基金管理人支付", "进行资金支付")
        return text

    @staticmethod
    def _first_nonempty(v, *keys, default=""):
        for key in keys:
            value = v.get(key)
            if value not in (None, ""):
                return value
        return default

    @staticmethod
    def _strip_leading_exact_heading(text, heading):
        lines = str(text or "").splitlines()
        while lines and not lines[0].strip():
            lines.pop(0)
        if lines and lines[0].strip() == heading:
            lines = lines[1:]
            while lines and not lines[0].strip():
                lines.pop(0)
        return "\n".join(lines).strip()

    def _build_jinja_context(self, v):
        fund_name = v.get("FUND_NAME", "")
        contract_name = self._first_nonempty(v, "CONTRACT_NAME", "FUND_CONTRACT_NAME", default=f"{fund_name}基金合同")
        custody_agreement_name = self._first_nonempty(
            v, "CUSTODY_AGREEMENT_NAME", "TRUSTEE_AGREEMENT_NAME", default=f"{fund_name}托管协议"
        )
        prospectus_name = self._first_nonempty(
            v, "PROSPECTUS_NAME", default=f"{fund_name}招募说明书"
        )
        product_summary_name = self._first_nonempty(
            v, "PRODUCT_SUMMARY_NAME", default=f"{fund_name}基金产品资料概要"
        )
        sale_notice_name = self._first_nonempty(
            v, "FUND_OFFERING_NOTICE_NAME", default=f"{fund_name}基金份额发售公告"
        )
        registration_doc_no = self._first_nonempty(
            v, "REGISTRATION_APPROVAL_NO", default="证监许可〔202X〕XXXX号"
        )
        registration_year = self._first_nonempty(
            v, "REGISTRATION_APPROVAL_YEAR", default="202X"
        )
        registration_number = self._first_nonempty(
            v, "REGISTRATION_APPROVAL_NUMBER", default="XXXX"
        )
        manager_name = self._first_nonempty(v, "MANAGER_NAME", "FUND_MANAGER", default="南方基金管理股份有限公司")
        target_etf_short_name = self._first_nonempty(v, "TARGET_ETF_SHORT_NAME", default=v.get("TARGET_ETF_NAME", ""))
        fund_manager_resume = self._first_nonempty(
            v,
            "FUND_MANAGER_RESUME",
            "FUND_MANAGER_PROFILE",
            default="{{基金经理姓名及简历全文}}",
        )

        hk = bool(v.get("HAS_HK_CONNECT"))
        initiating = bool(v.get("IS_INITIATING_FUND"))
        chuangye = bool(v.get("IS_CHUANGYE_FUND"))
        kechuang = bool(v.get("IS_KECHUANG_FUND"))
        market_scope = v.get("MARKET_SCOPE", "跨市场")
        single_exchange = v.get("SINGLE_MARKET_EXCHANGE", "")
        custodian_type = v.get("CUSTODIAN_TYPE", "")
        custodian_name = v.get("CUSTODIAN_NAME", "")
        subscription_fee_style = v.get("SUBSCRIPTION_FEE_STYLE", v.get("认申购费率模式", ""))
        redemption_fee_style = v.get("REDEMPTION_FEE_STYLE", v.get("赎回费率模式", ""))
        non_personal_redemption_style = v.get("NON_PERSONAL_REDEMPTION_FEE_STYLE", v.get("非个人赎回费率模式", ""))

        values = {
            "基金名称": fund_name,
            "基金简称": v.get("FUND_SHORT_NAME", ""),
            "目标ETF名称": v.get("TARGET_ETF_NAME", ""),
            "目标ETF简称": target_etf_short_name,
            "标的指数": v.get("INDEX_NAME", ""),
            "基金管理人名称": manager_name,
            "基金托管人名称": custodian_name,
            "托管协议名称": custody_agreement_name,
            "基金合同名称": contract_name,
            "招募说明书名称": prospectus_name,
            "基金产品资料概要名称": product_summary_name,
            "基金份额发售公告名称": sale_notice_name,
            "注册批文日期": self._first_nonempty(v, "REGISTRATION_APPROVAL_DATE", "REG_APPROVAL_DATE", default="202X年X月X日"),
            "注册批文文号": registration_doc_no,
            "注册批文年份": registration_year,
            "注册批文编号": registration_number,
            "封面日期": self._first_nonempty(v, "COVER_DATE", "CONTRACT_DATE", default="{{封面日期}}"),
            "指数编制机构名称": self._first_nonempty(v, "INDEX_COMPILER", default="{{指数编制机构名称}}"),
            "指数编制机构网址": self._first_nonempty(v, "INDEX_WEBSITE", default="{{指数编制机构网址}}"),
            "标的指数简介": self._first_nonempty(v, "INDEX_DESCRIPTION", default="{{标的指数简介}}"),
            "基金管理人信息版本": self._first_nonempty(v, "MANAGER_INFO_VERSION", default="{{基金管理人信息版本}}"),
            "基金托管人信息版本": self._first_nonempty(v, "CUSTODIAN_INFO_VERSION", default="{{基金托管人信息版本}}"),
            "基金经理姓名": self._first_nonempty(v, "FUND_MANAGER_NAME", default="{{基金经理姓名}}"),
            "基金经理简历全文": fund_manager_resume,
            "基金经理姓名及简历全文": fund_manager_resume,
            "相关服务机构全文": self._first_nonempty(v, "SERVICE_ORGANIZATIONS_TEXT", default="{{相关服务机构全文}}"),
            "其它应披露事项全文": self._first_nonempty(v, "OTHER_DISCLOSURE_TEXT", default="{{其它应披露事项全文}}"),
            "托管协议内容摘要全文": self._first_nonempty(
                v,
                "CUSTODY_AGREEMENT_SUMMARY_TEXT",
                default=_default_custody_agreement_summary_text(),
            ),
            "会计师事务所简介": self._strip_leading_exact_heading(
                self._first_nonempty(v, "ACCOUNTING_FIRM_PROFILE", default="{{会计师事务所简介}}"),
                "四、审计基金财产的会计师事务所",
            ),
            "保留占位符，前端填写": self._first_nonempty(
                v, "CUSTODIAN_PROSPECTUS_TEXT", default="{{保留占位符，前端填写}}"
            ),
            "风险揭示全文_按产品特征补充": self._first_nonempty(v, "RISK_DISCLOSURE_FULL_TEXT", default="{{风险揭示全文_按产品特征补充}}"),
        }

        conditions = {
            "条件_草案": bool(v.get("IS_DRAFT", v.get("是否草案", False))),
            "条件_正式稿": not bool(v.get("IS_DRAFT", v.get("是否草案", False))),
            "条件_发起式基金": initiating,
            "条件_普通基金": not initiating,
            "条件_托管人为商业银行": custodian_type == "商业银行",
            "条件_托管人为证券公司": custodian_type == "证券公司",
            "条件_托管人为工商银行": custodian_name == "中国工商银行股份有限公司",
            "条件_托管人为招商银行": custodian_name == "招商银行股份有限公司",
            "条件_单市场": market_scope == "单市场",
            "条件_跨市场": market_scope == "跨市场",
            "条件_单市场上交所": market_scope == "单市场" and single_exchange == "上交所",
            "条件_单市场深交所": market_scope == "单市场" and single_exchange == "深交所",
            "条件_创业板基金": chuangye,
            "条件_科创板基金": kechuang,
            "条件_港股通基金": hk,
            "条件_指数编制机构为中证指数": v.get("INDEX_COMPILER") == "中证指数有限公司",
            "条件_指数编制机构为深证信息": v.get("INDEX_COMPILER") == "深圳证券信息有限公司",
            "条件_低费率且销售服务费可返还": subscription_fee_style == "低费率直销免0.3且C类销售服务费可返还",
            "条件_普通股票指数费率": subscription_fee_style == "普通股票指数费率",
            "条件_赎回费率区分个人与非个人": redemption_fee_style == "区分个人与非个人",
            "条件_赎回费率不区分投资人类型": redemption_fee_style == "不区分投资人类型",
            "条件_非个人赎回费率长期递减": non_personal_redemption_style == "长期递减",
            "条件_C类销售服务费返还": bool(v.get("IS_C_SALES_SERVICE_FEE_RETURN", v.get("是否C类销售服务费返还", False))),
            "条件_管理人披露监事会成员": bool(v.get("DISCLOSE_SUPERVISORY_BOARD", v.get("是否披露监事会成员", False))),
            "条件_招募含港股通风险": hk,
            "条件_招募含科创板风险": kechuang,
            "条件_招募含创业板风险": chuangye,
            "风险总述口径": v.get("RISK_SUMMARY_STYLE", "市场风险"),
        }

        aliases = {}
        context = {}
        for idx, (name, value) in enumerate({**values, **conditions}.items()):
            alias = f"JINJA_VAR_{idx}"
            aliases[name] = alias
            context[alias] = value
        return context, aliases

    @staticmethod
    def _strip_jinja_reference_preamble(text):
        marker = re.search(r"(?m)^#?\s*(?:\{\{\s*基金名称\s*\}\}|\{FUND_NAME\})招募说明书", text)
        if marker:
            return text[marker.start():]
        return text

    @staticmethod
    def _strip_prospectus_template_notes(text):
        text = ProspectusEngine._strip_jinja_reference_preamble(text)
        cleaned = []
        in_fence = False
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith("```"):
                in_fence = not in_fence
                continue
            if in_fence:
                continue
            if stripped.startswith(">"):
                continue
            cleaned.append(line)
        return "\n".join(cleaned)

    def _apply_prospectus_template_values(self, v):
        v = dict(v)
        context, aliases = self._build_jinja_context(v)
        key_map = {
            "基金名称": "FUND_NAME",
            "基金简称": "FUND_SHORT_NAME",
            "目标ETF名称": "TARGET_ETF_NAME",
            "目标ETF简称": "TARGET_ETF_SHORT_NAME",
            "标的指数": "INDEX_NAME",
            "基金管理人名称": "MANAGER_NAME",
            "基金托管人名称": "CUSTODIAN_NAME",
            "托管协议名称": "CUSTODY_AGREEMENT_NAME",
            "基金合同名称": "CONTRACT_NAME",
            "招募说明书名称": "PROSPECTUS_NAME",
            "基金产品资料概要名称": "PRODUCT_SUMMARY_NAME",
            "基金份额发售公告名称": "FUND_OFFERING_NOTICE_NAME",
            "注册批文日期": "REGISTRATION_APPROVAL_DATE",
            "注册批文文号": "REGISTRATION_APPROVAL_NO",
            "注册批文年份": "REGISTRATION_APPROVAL_YEAR",
            "注册批文编号": "REGISTRATION_APPROVAL_NUMBER",
            "封面日期": "COVER_DATE",
            "指数编制机构名称": "INDEX_COMPILER",
            "指数编制机构网址": "INDEX_WEBSITE",
            "标的指数简介": "INDEX_DESCRIPTION",
            "基金管理人信息版本": "MANAGER_INFO_VERSION",
            "基金托管人信息版本": "CUSTODIAN_INFO_VERSION",
            "基金经理姓名": "FUND_MANAGER_NAME",
            "基金经理简历全文": "FUND_MANAGER_RESUME",
            "基金经理姓名及简历全文": "FUND_MANAGER_RESUME",
            "相关服务机构全文": "SERVICE_ORGANIZATIONS_TEXT",
            "其它应披露事项全文": "OTHER_DISCLOSURE_TEXT",
            "托管协议内容摘要全文": "CUSTODY_AGREEMENT_SUMMARY_TEXT",
            "会计师事务所简介": "ACCOUNTING_FIRM_PROFILE",
            "保留占位符，前端填写": "CUSTODIAN_PROSPECTUS_TEXT",
            "风险揭示全文_按产品特征补充": "RISK_DISCLOSURE_FULL_TEXT",
        }
        for name, key in key_map.items():
            alias = aliases.get(name)
            if not alias:
                continue
            value = context.get(alias)
            if isinstance(value, str) and value.strip().startswith("{{"):
                continue
            if v.get(key) in (None, ""):
                v[key] = value
        if v.get("ACCOUNTING_FIRM_PROFILE"):
            v["ACCOUNTING_FIRM_PROFILE"] = self._strip_leading_exact_heading(
                v.get("ACCOUNTING_FIRM_PROFILE"),
                "四、审计基金财产的会计师事务所",
            )

        hk = bool(v.get("HAS_HK_CONNECT"))
        initiating = bool(v.get("IS_INITIATING_FUND"))
        chuangye = bool(v.get("IS_CHUANGYE_FUND"))
        kechuang = bool(v.get("IS_KECHUANG_FUND"))
        market_scope = v.get("MARKET_SCOPE", "跨市场")
        subscription_fee_style = v.get("SUBSCRIPTION_FEE_STYLE", v.get("认购费率模式", ""))
        non_personal_redemption_style = v.get("NON_PERSONAL_REDEMPTION_FEE_STYLE", v.get("非个人赎回费率模式", ""))
        risk_summary_style = v.get("RISK_SUMMARY_STYLE", "市场风险")

        v["COND_RISK_SUMMARY_MARKET"] = risk_summary_style == "市场风险"
        v["COND_RISK_SUMMARY_SYSTEMIC"] = risk_summary_style == "系统性风险和非系统性风险"
        v["COND_LOW_FEE_C_RETURN"] = subscription_fee_style == "低费率直销免0.3且C类销售服务费可返还"
        v["COND_NON_PERSONAL_REDEMPTION_LONG_DECREASE"] = non_personal_redemption_style == "长期递减"
        v["COND_PROSPECTUS_RISK_HK"] = hk
        v["COND_PROSPECTUS_RISK_KECHUANG"] = (not hk) and kechuang
        v["COND_PROSPECTUS_RISK_CHUANGYE"] = (not hk) and (not kechuang) and chuangye
        v["COND_PROSPECTUS_RISK_NON_SPECIAL"] = (
            (not hk) and (not kechuang) and (not chuangye)
        )
        v["COND_PROSPECTUS_RISK_FALLBACK"] = not (
            v["COND_PROSPECTUS_RISK_HK"]
            or v["COND_PROSPECTUS_RISK_KECHUANG"]
            or v["COND_PROSPECTUS_RISK_CHUANGYE"]
            or v["COND_PROSPECTUS_RISK_NON_SPECIAL"]
        )
        if v["COND_PROSPECTUS_RISK_FALLBACK"]:
            if not v.get("RISK_DISCLOSURE_FULL_TEXT"):
                v["RISK_DISCLOSURE_FULL_TEXT"] = self._default_prospectus_risk_disclosure(v)
            import logging
            logging.getLogger(__name__).warning(
                "风险揭示章节使用了FALLBACK简化内容: hk=%s, kechuang=%s, chuangye=%s, initiating=%s, market_scope=%s",
                hk, kechuang, chuangye, initiating, market_scope,
            )
        return v

    @staticmethod
    def _default_prospectus_risk_disclosure(v):
        market_risk = ProductSummaryEngine._build_default_market_risk_summary()
        specific_risk = ProductSummaryEngine._build_default_specific_risk_summary(v)
        return (
            "本基金的目标ETF为股票型基金，一般而言，其长期平均风险与预期收益高于混合型基金、"
            "债券型基金与货币市场基金。本基金为ETF联接基金，通过投资于目标ETF跟踪标的指数表现，"
            "具有与标的指数以及标的指数所代表的市场相似的风险收益特征。\n\n"
            "本基金面临的主要风险有市场风险、管理风险、流动性风险、本基金特有风险、其他风险、"
            "基金风险评价不一致的风险等。\n\n"
            f"一、市场风险\n\n{market_risk}\n\n"
            "二、管理风险\n\n"
            "在基金管理运作过程中，基金管理人的知识、经验、判断、决策、技能等，"
            "会影响其对信息的占有和对经济形势、证券价格走势的判断，从而影响基金收益水平。\n\n"
            "三、流动性风险\n\n"
            "本基金为开放式基金，基金管理人有义务接受投资者的申购和赎回。"
            "如出现较大数额赎回申请或所投资资产流动性受限，基金资产可能难以及时变现。\n\n"
            f"四、本基金特有的风险\n\n{specific_risk}\n\n"
            "五、其他风险\n\n"
            "本基金可能面临技术风险、操作风险、不可抗力风险以及法律法规、监管政策变化等其他风险。\n\n"
            "六、基金风险评价不一致的风险\n\n"
            "基金管理人、销售机构或其他评价机构对本基金的风险评级可能存在差异，"
            "投资者应结合自身风险承受能力审慎作出投资决策。"
        )

    def _render_jinja_template(self, text, v):
        text = self._strip_jinja_reference_preamble(text)
        context, aliases = self._build_jinja_context(v)
        protected = {}

        def protect_literal(value):
            token = f"__PROSPECTUS_LITERAL_{len(protected)}__"
            protected[token] = value
            return token

        def replace_variable(m):
            inner = m.group(1).strip()
            if inner in aliases:
                return "{{ " + aliases[inner] + " }}"
            if "引用基金合同" in inner:
                return protect_literal(m.group(0))
            return protect_literal(m.group(0))

        def replace_block(m):
            body = m.group(1)
            for name, alias in sorted(aliases.items(), key=lambda item: len(item[0]), reverse=True):
                body = body.replace(name, alias)
            return "{% " + body.strip() + " %}"

        text = re.sub(r"\{\{\s*(.*?)\s*\}\}", replace_variable, text, flags=re.DOTALL)
        text = re.sub(r"\{%-?\s*(.*?)\s*-?%\}", replace_block, text, flags=re.DOTALL)
        rendered = Environment(autoescape=False, keep_trailing_newline=True).from_string(text).render(context)
        for token, literal in protected.items():
            rendered = rendered.replace(token, literal)
        return rendered

    @staticmethod
    def _normalize_reference_text(text):
        return re.sub(r"\s+", "", str(text or "")).strip("。；;：:")

    @staticmethod
    def _heading_level(line):
        stripped = line.strip()
        if re.match(r"^第[一二三四五六七八九十百]+部分", stripped):
            return 0
        if re.match(r"^[一二三四五六七八九十百]+、", stripped):
            return 1
        if re.match(r"^（[一二三四五六七八九十百]+）", stripped):
            return 2
        if re.match(r"^\d+、", stripped):
            return 3
        if re.match(r"^（\d+）", stripped):
            return 4
        if re.match(r"^\d+[）)]", stripped):
            return 4
        return None

    def _extract_contract_section(self, contract_text, section_title):
        target = self._normalize_reference_text(section_title)
        lines = contract_text.splitlines()
        start = None
        for i, line in enumerate(lines):
            if "\t" in line:
                continue
            if line.strip().startswith("第") and "部分" in line and self._normalize_reference_text(line) == target:
                start = i
                break
        if start is None:
            return ""
        end = next(
            (
                i
                for i in range(start + 1, len(lines))
                if "\t" not in lines[i] and re.match(r"^第[一二三四五六七八九十百]+部分", lines[i].strip())
            ),
            len(lines),
        )
        return _strip_contract_signing_page_text("\n".join(lines[start + 1 : end]).strip())

    def _extract_contract_block(self, text, heading_query):
        query = self._normalize_reference_text(heading_query).strip("“”\"")
        if not query:
            return ""
        lines = text.splitlines()

        start = next(
            (
                i
                for i, line in enumerate(lines)
                if query in self._normalize_reference_text(line)
                and self._heading_level(line) is not None
            ),
            None,
        )
        if start is None:
            for line in lines:
                if query and query in self._normalize_reference_text(line):
                    return line.strip()
            return ""
        level = self._heading_level(lines[start])
        end = next(
            (
                i
                for i in range(start + 1, len(lines))
                if self._heading_level(lines[i]) is not None and self._heading_level(lines[i]) <= level
            ),
            len(lines),
        )
        return "\n".join(lines[start + 1 : end]).strip()

    def _extract_contract_reference(self, contract_text, path):
        path = re.sub(r"\s+", " ", path).strip("。 ")
        parts = [part.strip() for part in path.split(">") if part.strip()]
        if not parts:
            return ""

        current = self._extract_contract_section(contract_text, parts[0])
        if not current:
            return ""
        if len(parts) == 1:
            return current

        for part in parts[1:]:
            items = [item.strip() for item in re.split(r"[；;]", part) if item.strip()]
            if len(items) > 1:
                blocks = [self._extract_contract_block(current, item) for item in items]
                blocks = [block for block in blocks if block]
                return "\n\n".join(blocks).strip()
            block = self._extract_contract_block(current, items[0] if items else part)
            if block:
                current = block
        return current.strip()

    def _definition_reference_for_prospectus(self, text, v):
        prospectus_name = self._first_nonempty(
            v,
            "PROSPECTUS_NAME",
            default=f"{v.get('FUND_NAME', '')}招募说明书",
        )
        is_initiating = bool(v.get("IS_INITIATING_FUND"))
        has_hk_connect = bool(v.get("HAS_HK_CONNECT"))

        lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if "在本基金合同中" in stripped and "下列词语或简称具有如下含义" in stripped:
                line = "在本招募说明书中，除非文意另有所指，下列词语或简称具有如下含义："
            elif re.match(r"^\d+、基金合同或本基金合同：", stripped):
                line = line.replace("基金合同或本基金合同", "基金合同", 1)
            elif re.match(r"^\d+、招募说明书：", stripped):
                number = stripped.split("、", 1)[0]
                line = f"{number}、招募说明书或本招募说明书：指《{prospectus_name}》及其更新"

            if not is_initiating and re.match(r"^\d+、(发起式基金|发起资金|发起资金提供方)：", stripped):
                continue
            if not has_hk_connect and (
                "内地与香港股票市场交易互联互通机制" in stripped
                or re.match(r"^\d+、港股通：", stripped)
            ):
                continue
            lines.append(line)

        return "\n".join(lines).strip()

    def _extract_prefixed_contract_reference(self, contract_text, path, prefix):
        referenced = self._extract_contract_reference(contract_text, path)
        prefix = re.sub(r"\s+", "", prefix or "")
        if not referenced or not prefix:
            return referenced

        tokens = [token for token in prefix.split("、") if token]
        if not tokens:
            return referenced

        blocks = []
        for token in tokens:
            block = self._extract_contract_block(referenced, f"{token}、")
            if block:
                blocks.append(block)
        return "\n\n".join(blocks).strip() or referenced

    def _render_reference_markers(self, text, contract_text=None, v=None):
        def replace_reference(match):
            path = re.sub(r"\s+", " ", match.group("path").strip())
            prefix = re.sub(r"\s+", "", match.group("prefix") or "")
            if contract_text:
                referenced = (
                    self._extract_prefixed_contract_reference(contract_text, path, prefix)
                    if prefix
                    else self._extract_contract_reference(contract_text, path)
                )
                if referenced:
                    if not prefix and self._normalize_reference_text(path) == self._normalize_reference_text("第二部分 释义"):
                        referenced = self._definition_reference_for_prospectus(referenced, v or {})
                    return referenced
            return f"本部分{prefix + '相关内容' if prefix else '相关内容'}引用《基金合同》{path}。"

        return re.sub(
            r"\{\{\s*(?P<prefix>.*?)引用基金合同\s*[：:]\s*(?P<path>.*?)\s*\}\}",
            replace_reference,
            text,
            flags=re.DOTALL,
        )

    @staticmethod
    def _remove_subsection_by_heading(text, chapter_title, next_chapter_title, subsection_heading, body_contains=None):
        chapter_start = text.rfind(chapter_title)
        if chapter_start == -1:
            return text
        chapter_end = text.find(next_chapter_title, chapter_start + len(chapter_title))
        if chapter_end == -1:
            chapter_end = len(text)

        chapter = text[chapter_start:chapter_end]
        heading_pattern = re.escape(subsection_heading)
        pattern = re.compile(
            rf"\n*{heading_pattern}\s*\n[\s\S]*?(?=\n[一二三四五六七八九十百]+、|\Z)"
        )

        def replace_match(match):
            block = match.group(0)
            if body_contains and body_contains not in block:
                return block
            return "\n"

        chapter = pattern.sub(replace_match, chapter)
        return text[:chapter_start] + chapter + text[chapter_end:]

    @staticmethod
    def _prospectus_clause_library() -> dict:
        with open(PROSPECTUS_CLAUSES_JSON, encoding="utf-8") as f:
            return json.load(f).get("clauses", {})

    @staticmethod
    def _prospectus_replacements(category: str | None = None) -> list[dict]:
        replacements = ProspectusEngine._prospectus_clause_library().get(
            "PROSPECTUS_REFERENCE_REPLACEMENTS",
            {},
        ).get("replacements", [])
        if category is None:
            return [item for item in replacements if isinstance(item, dict)]
        return [
            item
            for item in replacements
            if isinstance(item, dict) and item.get("category") == category
        ]

    @staticmethod
    def _replace_product_summary_disclosure_wording(text):
        replacement = (
            ProspectusEngine._prospectus_clause_library()
            .get("PRODUCT_SUMMARY_DISCLOSURE_WORDING", {})
            .get("variants", {})
            .get("DEFAULT", {})
        )
        old = str(replacement.get("old") or "")
        new = str(replacement.get("new") or "")
        return text.replace(old, new) if old and new else text

    @staticmethod
    def _replace_side_pocket_reference_wording(text):
        for item in ProspectusEngine._prospectus_replacements("side_pocket"):
            old = str(item.get("old") or "")
            new = str(item.get("new") or "")
            if old and new:
                text = text.replace(old, new)
        return text

    @staticmethod
    def _relocate_sales_service_fee_details(text):
        chapter_title = "十三、基金的费用与税收"
        next_chapter_title = "十四、基金的会计与审计"
        chapter_start = text.rfind(chapter_title)
        if chapter_start == -1:
            return text
        chapter_end = text.find(next_chapter_title, chapter_start + len(chapter_title))
        if chapter_end == -1:
            return text

        chapter = text[chapter_start:chapter_end]
        detail_heading = "C类基金份额的销售服务费具体收取/返还方式："
        detail_heading_pos = chapter.rfind(detail_heading)
        if detail_heading_pos == -1:
            return text

        prelude = "招募说明书需额外保留以下“C类基金份额的销售服务费具体收取/返还方式”："
        detail_remove_start = detail_heading_pos
        prelude_pos = chapter.rfind(prelude, 0, detail_heading_pos)
        if prelude_pos != -1:
            between = chapter[prelude_pos + len(prelude):detail_heading_pos].strip()
            if not between:
                detail_remove_start = prelude_pos

        detail_block = chapter[detail_heading_pos:].strip()
        chapter_without_detail = chapter[:detail_remove_start].rstrip() + "\n\n"
        anchors = [
            "C类基金份额的销售服务费每日计算，具体收取/返还方式详见《招募说明书》。",
            "C类基金份额的销售服务费每日计提，具体收取/返还方式详见《招募说明书》。",
        ]
        for anchor in anchors:
            if anchor in chapter_without_detail:
                chapter_without_detail = chapter_without_detail.replace(anchor, detail_block, 1)
                return text[:chapter_start] + chapter_without_detail + text[chapter_end:]

        fallback = re.search(r"\n上述“(?:一、|（一）、)基金费用的种类”中第\d+-\d+项费用", chapter_without_detail)
        if fallback:
            insert_at = fallback.start()
            chapter_without_detail = (
                chapter_without_detail[:insert_at].rstrip()
                + "\n\n"
                + detail_block
                + "\n\n"
                + chapter_without_detail[insert_at:].lstrip()
            )
        return text[:chapter_start] + chapter_without_detail + text[chapter_end:]

    @staticmethod
    def _apply_prospectus_reference_overrides(text):
        text = text.replace("本基金合同", "基金合同")
        for item in ProspectusEngine._prospectus_replacements("reference_override"):
            old = str(item.get("old") or "")
            new = str(item.get("new") or "")
            if old and new:
                text = text.replace(old, new)
        text = ProspectusEngine._remove_subsection_by_heading(
            text,
            "十一、基金资产估值",
            "十二、基金的收益与分配",
            "十、实施侧袋机制期间的基金资产估值",
            body_contains="应根据本部分的约定对主袋账户资产进行估值",
        )
        text = ProspectusEngine._replace_side_pocket_reference_wording(text)
        text = ProspectusEngine._replace_product_summary_disclosure_wording(text)
        text = ProspectusEngine._relocate_sales_service_fee_details(text)
        return text

    @staticmethod
    def _cleanup_prospectus(text):
        lines = []
        for line in text.split("\n"):
            stripped = line.strip()
            if stripped.startswith(">"):
                continue
            if stripped.startswith("招募说明书需额外保留"):
                continue
            if stripped.startswith("生成招募说明书时"):
                continue
            if re.match(r"^\d+\.\s*(篇首句|“基金合同|“招募说明书|若)", stripped):
                continue
            if stripped in ("---", "```"):
                continue
            if stripped.startswith("|") and stripped.endswith("|"):
                if re.fullmatch(r"[|\s:\-]+", stripped):
                    continue
                cells = [cell.strip() for cell in stripped.strip("|").split("|")]
                lines.append("    ".join(cell for cell in cells if cell))
                continue
            heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if heading:
                lines.append(heading.group(2).strip())
                continue
            lines.append(line)
        text = "\n".join(lines)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    # ── 条件求值 ──────────────────────────────────────────────────────────
    def _eval_condition(self, v, condition_str):
        if not condition_str:
            return True
        key, val = condition_str.split("=", 1)
        key = key.strip()
        val = val.strip().lower()
        actual = v.get(key, False)
        if isinstance(actual, bool):
            return actual == (val in ("true", "1", "yes"))
        return str(actual).lower() == val

    # ── Step 2: 注入招募专属差异条款（双重匹配） ─────────────────────────────
    @staticmethod
    def _prospectus_clause_scene(key, v):
        if key in {"RISK_DISCLOSURE_HK", "VALUATION_TIMING_TEXT"}:
            return "HK_CONNECT" if v.get("HAS_HK_CONNECT") else "STANDARD"
        if key in {"IMPORTANT_NOTICE_TEXT", "CHAPTER6_INTRO_TEXT"}:
            return "INITIATING" if v.get("IS_INITIATING_FUND") else "STANDARD"
        return "DEFAULT"
    def _inject_prospectus_clauses(self, v):
        v = dict(v)
        for key, clause_data in self.clauses.items():
            if key == "PROSPECTUS_VARIANTS":
                continue
            variants = clause_data.get("variants", {})
            scene = self._prospectus_clause_scene(key, v)

            # 第一轮：条件匹配（解析 "KEY=VALUE" 格式）
            matched_text = ""
            for variant_name, variant_data in variants.items():
                condition = variant_data.get("condition", "")
                if condition and self._eval_condition(v, condition):
                    matched_text = variant_data.get("text", "")
                    break

            # 第二轮：场景名匹配（variant name = scene name）
            if not matched_text:
                if scene in variants:
                    matched_text = variants[scene].get("text", "")
                elif "DEFAULT" in variants:
                    matched_text = variants["DEFAULT"].get("text", "")

            v[key] = matched_text
        return v

    def generate(self, form_data):
        v = self._derive_variables(form_data)
        v = self._inject_prospectus_clauses(v)
        v = self._apply_prospectus_template_values(v)
        contract_text = self.contract_engine.generate(form_data)
        text = PROSPECTUS_TEMPLATE_MD.read_text(encoding="utf-8")
        text = self._strip_prospectus_template_notes(text)
        text = self.contract_engine._process_conditionals(text, v)
        text = self.contract_engine._replace_placeholders(text, v)
        text = self._render_reference_markers(text, contract_text, v)
        text = self.contract_engine._renumber_sequences(text)
        text = self._apply_prospectus_reference_overrides(text)
        text = self._cleanup_prospectus(text)
        return text

    # ── generate_bundle ───────────────────────────────────────────────────
    def generate_bundle(self, form_data):
        text = self.generate(form_data)
        v = self._derive_variables(form_data)
        v = self._inject_prospectus_clauses(v)
        return {"text": text, "variables": v}

    # ── DOCX 导出 ─────────────────────────────────────────────────────────
    @staticmethod
    def _format_prospectus_draft_header_text(title_text):
        header = str(title_text or "招募说明书").strip().strip("《》")
        if "招募说明书" not in header:
            header = f"{header}招募说明书"
        header = re.sub(r"招募说明书(?:（草案）)?$", "招募说明书（草案）", header)
        if not header.endswith("（草案）"):
            header = header.replace("招募说明书", "招募说明书（草案）", 1)
        return f"《{header}》"

    @staticmethod
    def _is_prospectus_heading_line(text):
        return bool(re.match(r"^[一二三四五六七八九十百]+、", (text or "").strip()))

    def _parse_docx_prospectus_text(self, text):
        lines = [
            line.strip()
            for line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").splitlines()
            if line.strip()
        ]
        notice_index = next((idx for idx, line in enumerate(lines) if line == "重要提示"), None)
        toc_index = next((idx for idx, line in enumerate(lines) if line == "目录"), None)
        cover_end_candidates = [idx for idx in (notice_index, toc_index) if idx is not None]
        cover_end = min(cover_end_candidates) if cover_end_candidates else len(lines)
        cover_lines = lines[:cover_end] or ["招募说明书"]

        important_notice_blocks = []
        if notice_index is not None:
            notice_end = toc_index if toc_index is not None and toc_index > notice_index else len(lines)
            important_notice_blocks = lines[notice_index + 1 : notice_end]

        toc_entries = []
        body_lines = []
        if toc_index is not None:
            first_toc_entry = ""
            body_started = False
            for line in lines[toc_index + 1 :]:
                normalized_line = re.sub(r"[\t ]+\d+\s*$", "", line).strip()
                if not body_started:
                    if first_toc_entry and normalized_line == first_toc_entry:
                        body_started = True
                        body_lines.append(normalized_line)
                        continue
                    if self._is_prospectus_heading_line(normalized_line):
                        if not first_toc_entry:
                            first_toc_entry = normalized_line
                        toc_entries.append(normalized_line)
                        continue
                    if toc_entries:
                        body_started = True
                if body_started:
                    body_lines.append(line)
        else:
            body_start = (notice_index + 1) if notice_index is not None else cover_end
            body_lines = lines[body_start:]

        if not body_lines and toc_entries:
            body_lines = toc_entries

        return {
            "cover_lines": cover_lines,
            "important_notice_blocks": important_notice_blocks,
            "toc_entries": toc_entries,
            "body_lines": body_lines,
        }

    def _build_render_model_from_text(self, text):
        parsed = self._parse_docx_prospectus_text(text)

        def split_heading(line):
            match = re.match(r"^([一二三四五六七八九十百]+)、\s*(.*)$", str(line or "").strip())
            if not match:
                return "", str(line or "").strip()
            return match.group(1), match.group(2).strip()

        toc_entries = []
        for entry in parsed.get("toc_entries") or []:
            chapter_cn, title = split_heading(entry)
            toc_entries.append({"chapter_cn": chapter_cn, "title": title})

        chapter_heading_entries = []
        for entry in toc_entries:
            chapter_cn = str(entry.get("chapter_cn") or "").strip()
            title = str(entry.get("title") or "").strip()
            if chapter_cn and title:
                chapter_heading_entries.append({
                    "chapter_cn": chapter_cn,
                    "title": title,
                    "full": f"{chapter_cn}、{title}",
                })

        next_chapter_index = 0

        def match_body_chapter(line):
            nonlocal next_chapter_index
            stripped = str(line or "").strip()
            if chapter_heading_entries:
                for idx in range(next_chapter_index, len(chapter_heading_entries)):
                    entry = chapter_heading_entries[idx]
                    if stripped == entry["full"]:
                        next_chapter_index = idx + 1
                        return entry["chapter_cn"], entry["title"]
                if next_chapter_index < len(chapter_heading_entries):
                    entry = chapter_heading_entries[next_chapter_index]
                    if stripped == entry["title"]:
                        next_chapter_index += 1
                        return entry["chapter_cn"], entry["title"]
                return "", ""
            if self._is_prospectus_heading_line(stripped):
                return split_heading(stripped)
            return "", ""

        chapters = []
        current_chapter = None
        for line in parsed.get("body_lines") or []:
            chapter_cn, title = match_body_chapter(line)
            if chapter_cn or title:
                current_chapter = {"chapter_cn": chapter_cn, "title": title, "blocks": []}
                chapters.append(current_chapter)
                continue
            if current_chapter is None:
                current_chapter = {"chapter_cn": "", "title": "", "blocks": []}
                chapters.append(current_chapter)
            current_chapter["blocks"].append({"type": "paragraph", "text": line})

        if not toc_entries:
            toc_entries = [
                {"chapter_cn": chapter.get("chapter_cn", ""), "title": chapter.get("title", "")}
                for chapter in chapters
                if chapter.get("title") or chapter.get("chapter_cn")
            ]

        return {
            "cover_lines": parsed.get("cover_lines") or [],
            "important_notice_title": "重要提示",
            "important_notice_blocks": parsed.get("important_notice_blocks") or [],
            "toc_title": "目录",
            "toc_entries": toc_entries,
            "chapters": chapters,
        }

    @staticmethod
    def _prospectus_template_dir():
        return _linked_legal_template_dir("PROSPECTUS_TEMPLATE_DIR")

    @classmethod
    def _format_template_prospectus_docx(cls, v=None):
        template_docx = cls._prospectus_template_dir() / LINKED_PROSPECTUS_FORMAT_TEMPLATE_FILENAME
        return template_docx if template_docx.exists() else None

    @classmethod
    def _require_format_template_prospectus_docx(cls, v=None):
        template_docx = cls._format_template_prospectus_docx(v)
        if template_docx and Path(template_docx).exists():
            return template_docx
        template_dir = cls._prospectus_template_dir()
        raise FileNotFoundError(
            f"缺少招募说明书 Word 导出所需的参考 DOCX 资产，请检查参考目录：{template_dir}"
        )

    def build_docx(self, text, v=None):
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, Twips
        import copy

        model = self._parse_docx_prospectus_text(text)
        template_docx = self._require_format_template_prospectus_docx(v)
        doc = Document(str(template_docx))

        def format_lookup_key(text_value):
            return re.sub(r"\s+", " ", str(text_value or "").strip())

        def build_reference_format_lookup(docx_path):
            if not docx_path:
                return {}
            reference_doc = Document(str(docx_path))
            lookup = {}
            for ref_para in reference_doc.paragraphs:
                key = format_lookup_key(ref_para.text)
                if not key or key in lookup:
                    continue
                first_run = next((run for run in ref_para.runs if run.text), None)
                lookup[key] = {
                    "pPr": copy.deepcopy(ref_para._p.pPr) if ref_para._p.pPr is not None else None,
                    "rPr": copy.deepcopy(first_run._r.rPr) if first_run is not None and first_run._r.rPr is not None else None,
                }
            return lookup

        reference_format_lookup = build_reference_format_lookup(template_docx)

        def build_reference_table_formats(docx_path):
            if not docx_path:
                return []
            reference_doc = Document(str(docx_path))
            table_formats = []
            for ref_table in reference_doc.tables:
                ref_tbl = ref_table._tbl
                table_formats.append({
                    "tblPr": copy.deepcopy(ref_tbl.tblPr) if ref_tbl.tblPr is not None else None,
                    "tblGrid": copy.deepcopy(ref_tbl.find(qn("w:tblGrid"))),
                    "rows": [
                        {
                            "trPr": copy.deepcopy(ref_row._tr.trPr) if ref_row._tr.trPr is not None else None,
                            "cells": [
                                copy.deepcopy(ref_cell._tc.tcPr) if ref_cell._tc.tcPr is not None else None
                                for ref_cell in ref_row.cells
                            ],
                        }
                        for ref_row in ref_table.rows
                    ],
                })
            return table_formats

        reference_table_formats = build_reference_table_formats(template_docx)
        reference_table_index = 0

        def replace_child(element, child_tag, replacement, insert_at):
            existing = element.find(qn(child_tag))
            if existing is not None:
                element.remove(existing)
            if replacement is not None:
                element.insert(insert_at, copy.deepcopy(replacement))

        def apply_reference_table_format(table):
            nonlocal reference_table_index
            if reference_table_index >= len(reference_table_formats):
                return
            table_format = reference_table_formats[reference_table_index]
            reference_table_index += 1

            tbl = table._tbl
            replace_child(tbl, "w:tblPr", table_format.get("tblPr"), 0)
            replace_child(tbl, "w:tblGrid", table_format.get("tblGrid"), 1 if tbl.tblPr is not None else 0)

            row_formats = table_format.get("rows") or []
            for row, row_format in zip(table.rows, row_formats):
                replace_child(row._tr, "w:trPr", row_format.get("trPr"), 0)
                for cell, cell_pr in zip(row.cells, row_format.get("cells") or []):
                    replace_child(cell._tc, "w:tcPr", cell_pr, 0)

        def clear_body_keep_section(document):
            body = document._element.body
            for child in list(body):
                if child.tag != qn("w:sectPr"):
                    body.remove(child)

        def set_default_tab_stop(document, value="420"):
            settings = document.settings.element
            default_tab_stop = settings.find(qn("w:defaultTabStop"))
            if default_tab_stop is None:
                default_tab_stop = OxmlElement("w:defaultTabStop")
                settings.insert(0, default_tab_stop)
            default_tab_stop.set(qn("w:val"), str(value))

        def apply_reference_page_setup(section):
            section.page_width = Twips(PAGE_LAYOUT["page_width"])
            section.page_height = Twips(PAGE_LAYOUT["page_height"])
            section.top_margin = Twips(PAGE_LAYOUT["top_margin"])
            section.bottom_margin = Twips(PAGE_LAYOUT["bottom_margin"])
            section.left_margin = Twips(PAGE_LAYOUT["left_margin"])
            section.right_margin = Twips(PAGE_LAYOUT["right_margin"])
            section.header_distance = Twips(PAGE_LAYOUT["header_distance"])
            section.footer_distance = Twips(PAGE_LAYOUT["footer_distance"])
            sect_pr = section._sectPr
            for node in list(sect_pr.findall(qn("w:docGrid"))):
                sect_pr.remove(node)
            doc_grid = OxmlElement("w:docGrid")
            doc_grid.set(qn("w:type"), "lines")
            doc_grid.set(qn("w:linePitch"), "312")
            sect_pr.append(doc_grid)

        def normalize_normal_style(document):
            normal_style = document.styles["Normal"]
            normal_style.font.name = "Times New Roman"
            normal_style.font.size = Pt(10.5)
            normal_style.paragraph_format.line_spacing = 1.5
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(0)
            rpr = normal_style._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")
            rfonts.set(qn("w:eastAsia"), "宋体")
            for tag in ("w:sz", "w:szCs"):
                node = rpr.find(qn(tag))
                if node is None:
                    node = OxmlElement(tag)
                    rpr.append(node)
                node.set(qn("w:val"), "21")

        def set_run_font(run, eastasia="宋体", ascii_font="Times New Roman", size_half_pt=21, bold=False):
            run.font.name = ascii_font
            run.font.size = Pt(size_half_pt / 2)
            run.bold = bold
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), ascii_font)
            rfonts.set(qn("w:hAnsi"), ascii_font)
            rfonts.set(qn("w:eastAsia"), eastasia)
            for tag in ("w:sz", "w:szCs"):
                node = rpr.find(qn(tag))
                if node is None:
                    node = OxmlElement(tag)
                    rpr.append(node)
                node.set(qn("w:val"), str(size_half_pt))
            for tag in ("w:b", "w:bCs"):
                node = rpr.find(qn(tag))
                if bold:
                    if node is None:
                        rpr.append(OxmlElement(tag))
                elif node is not None:
                    rpr.remove(node)

        def set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False):
            paragraph.alignment = align
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.first_line_indent = Pt(21 if first_line else 0)
            ppr = paragraph._p.get_or_add_pPr()
            snap_to_grid = ppr.find(qn("w:snapToGrid"))
            if snap_to_grid is None:
                snap_to_grid = OxmlElement("w:snapToGrid")
                spacing = ppr.find(qn("w:spacing"))
                children = list(ppr)
                insert_at = children.index(spacing) if spacing is not None else len(children)
                ppr.insert(insert_at, snap_to_grid)
            snap_to_grid.set(qn("w:val"), "0")
            ind = ppr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                ppr.append(ind)
            for attr in ("w:firstLine", "w:firstLineChars", "w:hanging", "w:hangingChars"):
                attr_name = qn(attr)
                if ind.get(attr_name) is not None:
                    del ind.attrib[attr_name]
            if first_line:
                ind.set(qn("w:firstLine"), "420")
                ind.set(qn("w:firstLineChars"), "200")
            else:
                ind.set(qn("w:firstLine"), "0")

        def add_paragraph(text_value="", *, style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, eastasia="宋体", ascii_font="Times New Roman", size_half_pt=21, bold=False, apply_run_font=True):
            paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            set_paragraph_format(paragraph, align=align, first_line=first_line)
            if text_value:
                run = paragraph.add_run(text_value)
                if apply_run_font:
                    set_run_font(run, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt, bold=bold)
            return paragraph

        def reference_format_uses_heading_style(fmt):
            p_pr = (fmt or {}).get("pPr")
            if p_pr is None:
                return False
            p_style = p_pr.find(qn("w:pStyle"))
            if p_style is None:
                return False
            return str(p_style.get(qn("w:val")) or "") in {"Heading1", "Heading2", "Heading3", "1", "2", "3", "21"}

        def apply_reference_format(paragraph, run, text_value, *, allow_heading_style=False):
            fmt = reference_format_lookup.get(format_lookup_key(text_value))
            if not fmt:
                return False
            if not allow_heading_style and reference_format_uses_heading_style(fmt):
                return False
            if fmt.get("pPr") is not None:
                existing_ppr = paragraph._p.pPr
                if existing_ppr is not None:
                    paragraph._p.remove(existing_ppr)
                paragraph._p.insert(0, copy.deepcopy(fmt["pPr"]))
            if run is not None and fmt.get("rPr") is not None:
                existing_rpr = run._r.rPr
                if existing_rpr is not None:
                    run._r.remove(existing_rpr)
                run._r.insert(0, copy.deepcopy(fmt["rPr"]))
            return True

        def add_reference_paragraph(text_value):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(str(text_value or ""))
            if apply_reference_format(paragraph, run, text_value):
                return paragraph
            set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True)
            set_run_font(run, eastasia="宋体", ascii_font="Times New Roman", size_half_pt=21)
            return paragraph

        def set_run_font_family_only(run, *, eastasia="宋体", ascii_font="宋体", cs_font="Times New Roman"):
            run.font.name = ascii_font
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), ascii_font)
            rfonts.set(qn("w:eastAsia"), eastasia)
            rfonts.set(qn("w:hAnsi"), ascii_font)
            rfonts.set(qn("w:cs"), cs_font)

        def set_paragraph_mark_font_family(paragraph, *, eastasia="宋体", ascii_font="宋体", cs_font="Times New Roman"):
            ppr = paragraph._p.get_or_add_pPr()
            ppr_rpr = ppr.find(qn("w:rPr"))
            if ppr_rpr is None:
                ppr_rpr = OxmlElement("w:rPr")
                ppr.append(ppr_rpr)
            rfonts = ppr_rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                ppr_rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), ascii_font)
            rfonts.set(qn("w:eastAsia"), eastasia)
            rfonts.set(qn("w:hAnsi"), ascii_font)
            rfonts.set(qn("w:cs"), cs_font)

        def add_chapter_heading(heading_text):
            paragraph = doc.add_paragraph(style="Heading 2")
            run = paragraph.add_run(str(heading_text or ""))
            if apply_reference_format(paragraph, run, heading_text, allow_heading_style=True):
                return paragraph
            set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
            set_paragraph_mark_font_family(paragraph)
            set_run_font_family_only(run)
            return paragraph

        def add_cover_line(text_value, *, title=False):
            if title:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ppr = paragraph._p.get_or_add_pPr()
                for node in list(ppr.findall(qn("w:spacing"))):
                    ppr.remove(node)
                ind = ppr.find(qn("w:ind"))
                if ind is None:
                    ind = OxmlElement("w:ind")
                    ppr.append(ind)
                for attr in ("w:left", "w:right", "w:hanging", "w:firstLineChars", "w:hangingChars"):
                    attr_name = qn(attr)
                    if ind.get(attr_name) is not None:
                        del ind.attrib[attr_name]
                ind.set(qn("w:firstLine"), "0")
                run = paragraph.add_run(text_value)
                run.font.name = "宋体"
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.rFonts
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.append(rfonts)
                rfonts.set(qn("w:ascii"), "宋体")
                rfonts.set(qn("w:hAnsi"), "宋体")
                rfonts.set(qn("w:eastAsia"), "宋体")
                for tag in ("w:sz", "w:szCs"):
                    node = rpr.find(qn(tag))
                    if node is None:
                        node = OxmlElement(tag)
                        rpr.append(node)
                    node.set(qn("w:val"), "44")
                for tag in ("w:b", "w:bCs"):
                    if rpr.find(qn(tag)) is None:
                        rpr.append(OxmlElement(tag))
                return paragraph
            return add_paragraph(
                text_value,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                first_line=False,
                eastasia="宋体",
                ascii_font="宋体",
                size_half_pt=30,
                bold=True,
            )

        def split_prospectus_heading(text_value):
            stripped = str(text_value or "").strip()
            match = re.match(r"^([一二三四五六七八九十百]+)、\s*(.+)$", stripped)
            if not match:
                return "", stripped
            return match.group(1), match.group(2).strip()

        chapter_heading_entries = []
        chapter_heading_full_by_line = {}
        for toc_entry in model.get("toc_entries") or []:
            if isinstance(toc_entry, dict):
                chapter_cn = str(toc_entry.get("chapter_cn") or "").strip()
                chapter_title = str(toc_entry.get("title") or "").strip()
                normalized_entry = f"{chapter_cn}、{chapter_title}" if chapter_cn and chapter_title else chapter_title
            else:
                normalized_entry = str(toc_entry or "").strip()
                chapter_cn, chapter_title = split_prospectus_heading(normalized_entry)
            if chapter_cn and chapter_title:
                chapter_heading_entries.append({
                    "full": normalized_entry,
                    "title": chapter_title,
                })
                chapter_heading_full_by_line[normalized_entry] = chapter_title

        body_chapter_count = 0
        last_body_kind = None
        next_chapter_index = 0

        def is_internal_section_heading(text_value):
            stripped = str(text_value or "").strip()
            if stripped in chapter_heading_full_by_line:
                return False
            return bool(re.match(r"^[一二三四五六七八九十百]+、\s*\S+", stripped))

        def match_body_chapter_title(stripped):
            nonlocal next_chapter_index
            if not chapter_heading_entries:
                if self._is_prospectus_heading_line(stripped):
                    return split_prospectus_heading(stripped)[1]
                return ""

            for idx in range(next_chapter_index, len(chapter_heading_entries)):
                entry = chapter_heading_entries[idx]
                if stripped == entry["full"]:
                    next_chapter_index = idx + 1
                    return entry["title"]

            if next_chapter_index < len(chapter_heading_entries):
                entry = chapter_heading_entries[next_chapter_index]
                if stripped == entry["title"]:
                    next_chapter_index += 1
                    return entry["title"]
            return ""

        def add_body_line(text_value):
            nonlocal body_chapter_count, last_body_kind
            stripped = str(text_value or "").strip()
            chapter_title = match_body_chapter_title(stripped)
            if chapter_title:
                heading_text = chapter_title
                if body_chapter_count:
                    doc.add_page_break()
                paragraph = add_chapter_heading(heading_text)
                add_paragraph("", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False, apply_run_font=False)
                body_chapter_count += 1
                last_body_kind = "chapter_heading"
                return paragraph
            if is_internal_section_heading(stripped) and last_body_kind == "content":
                add_paragraph("", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False, apply_run_font=False)
            paragraph = add_reference_paragraph(text_value)
            last_body_kind = "internal_heading" if is_internal_section_heading(stripped) else "content"
            return paragraph

        def is_markdown_table_row(text_value):
            stripped = str(text_value or "").strip()
            return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2

        def parse_markdown_table_row(text_value):
            return [cell.strip() for cell in str(text_value or "").strip().strip("|").split("|")]

        def parse_aligned_table_row(text_value):
            stripped = str(text_value or "").strip()
            if not stripped:
                return None
            cells = [cell.strip() for cell in re.split(r"\s{2,}", stripped) if cell.strip()]
            if len(cells) < 2:
                return None
            if len(cells) > 4:
                return None
            return cells

        def looks_like_fee_table_row(cells, in_table=False):
            if not cells:
                return False
            joined = " ".join(cells)
            if any(token in joined for token in ("认购金额", "申购金额", "持有期限", "认购费率", "申购费率", "赎回费率")):
                return True
            if in_table and re.search(r"[MN][＜≤≥>]|万|元|%|费率|每笔", joined):
                return True
            return False

        def is_markdown_table_separator(cells):
            if not cells:
                return False
            return all(re.fullmatch(r":?-{3,}:?", re.sub(r"\s+", "", cell or "")) for cell in cells)

        def set_cell_shading(cell, fill="D9D9D9"):
            tc_pr = cell._element.get_or_add_tcPr()
            existing = tc_pr.find(qn("w:shd"))
            if existing is not None:
                tc_pr.remove(existing)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            shd.set(qn("w:val"), "clear")
            tc_pr.append(shd)

        def add_word_table(rows):
            normalized_rows = [
                [str(cell or "").strip() for cell in row]
                for row in rows
                if any(str(cell or "").strip() for cell in row)
            ]
            if not normalized_rows:
                return None
            col_count = max(len(row) for row in normalized_rows)
            table = doc.add_table(rows=len(normalized_rows), cols=col_count, style="Table Grid")
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for row_idx, row_data in enumerate(normalized_rows):
                for col_idx in range(col_count):
                    cell = table.cell(row_idx, col_idx)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.clear()
                    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
                    run = paragraph.add_run(row_data[col_idx] if col_idx < len(row_data) else "")
                    set_run_font(run, eastasia="宋体", ascii_font="Times New Roman", size_half_pt=21, bold=row_idx == 0)
                    if row_idx == 0:
                        set_cell_shading(cell)
            apply_reference_table_format(table)
            return table

        def add_body_lines(lines):
            table_rows = []

            def flush_table():
                nonlocal table_rows, last_body_kind
                if table_rows:
                    add_word_table(table_rows)
                    table_rows = []
                    last_body_kind = "content"

            for line in lines:
                if is_markdown_table_row(line):
                    cells = parse_markdown_table_row(line)
                    if not is_markdown_table_separator(cells):
                        table_rows.append(cells)
                    continue
                aligned_cells = parse_aligned_table_row(line)
                if aligned_cells and looks_like_fee_table_row(aligned_cells, in_table=bool(table_rows)):
                    table_rows.append(aligned_cells)
                    continue
                flush_table()
                add_body_line(line)
            flush_table()

        def update_header_text(header_text, sections):
            for section in sections:
                headers = [section.header, section.first_page_header]
                if hasattr(section, "even_page_header"):
                    headers.append(section.even_page_header)
                for header in headers:
                    try:
                        header.is_linked_to_previous = False
                    except Exception:
                        pass
                    paragraph = _reset_header_footer_part(header)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ppr = paragraph._p.get_or_add_pPr()
                    existing = ppr.find(qn("w:pBdr"))
                    if existing is not None:
                        ppr.remove(existing)
                    pbdr = OxmlElement("w:pBdr")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:color"), "auto")
                    bottom.set(qn("w:sz"), "6")
                    bottom.set(qn("w:space"), "1")
                    pbdr.append(bottom)
                    ppr.append(pbdr)
                    run = paragraph.add_run(header_text)
                    set_run_font(run, eastasia="宋体", ascii_font="Times New Roman", size_half_pt=18)

        clear_body_keep_section(doc)
        set_default_tab_stop(doc)
        normalize_normal_style(doc)
        _set_update_fields_on_open(doc, OxmlElement, qn)
        _ensure_word_toc_styles(
            doc,
            heading_eastasia="宋体",
            heading_ascii="宋体",
            heading_size_half_pt=28,
            heading_bold=False,
            entry_eastasia="宋体",
            entry_ascii="Times New Roman",
            entry_size_half_pt=21,
            max_level=3,
        )

        cover_lines = model["cover_lines"]
        for _ in range(5):
            add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        add_cover_line(cover_lines[0], title=True)
        for _ in range(10):
            add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        for cover_line in cover_lines[1:]:
            add_cover_line(cover_line)

        add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        doc.add_page_break()

        if model["important_notice_blocks"]:
            add_paragraph(
                "重要提示",
                align=WD_ALIGN_PARAGRAPH.CENTER,
                first_line=False,
                eastasia="宋体",
                ascii_font="宋体",
                size_half_pt=21,
                bold=True,
            )
            for block in model["important_notice_blocks"]:
                add_paragraph(block, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True)
            doc.add_page_break()

        add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        add_paragraph(
            "目录",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line=False,
            eastasia="宋体",
            ascii_font="宋体",
            size_half_pt=28,
            bold=True,
        )
        toc_field_paragraph = doc.add_paragraph()
        set_paragraph_format(toc_field_paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
        _append_word_toc_field(
            toc_field_paragraph,
            OxmlElement,
            qn,
            levels="2-2",
            placeholder="右键更新目录",
            ascii_font="Times New Roman",
            eastasia_font="宋体",
            size=21,
            use_outline_levels=False,
        )
        doc.add_section(WD_SECTION.NEW_PAGE)

        add_body_lines(model["body_lines"])

        for section in doc.sections:
            apply_reference_page_setup(section)
        update_header_text(self._format_prospectus_draft_header_text(cover_lines[0]), list(doc.sections[1:]))
        _finalize_doc_page_numbers(doc, OxmlElement, qn, body_start_index=1)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        def finalize_footer_parts(docx_bytes: bytes) -> bytes:
            source = io.BytesIO(docx_bytes)
            target = io.BytesIO()
            with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zout:
                canonical_footer = None
                for info in zin.infolist():
                    if info.filename.startswith("word/footer"):
                        footer_xml = zin.read(info.filename)
                        if b"PAGE" in footer_xml:
                            canonical_footer = footer_xml
                            break

                for info in zin.infolist():
                    payload = zin.read(info.filename)
                    if info.filename == "word/document.xml":
                        document_xml = payload.decode("utf-8")

                        def strip_cover_refs(match):
                            sect_xml = match.group(0)
                            return re.sub(r"<w:(?:headerReference|footerReference)\b[^>]*/>", "", sect_xml)

                        document_xml = re.sub(
                            r"<w:sectPr\b[\s\S]*?</w:sectPr>",
                            strip_cover_refs,
                            document_xml,
                            count=1,
                        )
                        payload = document_xml.encode("utf-8")
                    elif info.filename.startswith("word/footer") and canonical_footer is not None and b"PAGE" not in payload:
                        payload = canonical_footer
                    zout.writestr(info, payload)
            return target.getvalue()

        docx_bytes = buf.read()
        docx_bytes = _update_docx_fields_with_word(docx_bytes, model.get("toc_entries") or [])
        docx_bytes = _reinforce_prospectus_normal_text_format(docx_bytes)
        return docx_bytes


# ═══════════════════════════════════════════════════════════════════════════════
#  ProductSummaryEngine — ETF联接基金产品资料概要
# ═══════════════════════════════════════════════════════════════════════════════

class ProductSummaryEngine:
    """生成 ETF联接基金 A/C 份额产品资料概要（6部分，含pipe-format表格）"""

    TEMPLATE_BASE_FUND_NAME = "南方中证全指农牧渔交易型开放式指数证券投资基金发起式联接基金"
    TEMPLATE_TARGET_ETF_NAME = "南方中证全指农牧渔交易型开放式指数证券投资基金"
    TEMPLATE_SHORT_NAME_BASE = "南方中证全指农牧渔ETF发起联接"
    TEMPLATE_INDEX_NAME = "中证全指农牧渔指数"
    TEMPLATE_CUSTODIAN_NAME = "中国工商银行股份有限公司"
    TEMPLATE_FUND_MANAGER_NAME = "杨恺宁"
    TEMPLATE_FILENAMES = {
        "A": "4、南方中证全指农牧渔交易型开放式指数证券投资基金发起式联接基金（A类份额）基金产品资料概要.docx",
        "C": "5、南方中证全指农牧渔交易型开放式指数证券投资基金联接基金（C类份额）基金产品资料概要.docx",
    }

    def __init__(self, prospectus_engine=None):
        self.prospectus_engine = prospectus_engine or ProspectusEngine()

    @staticmethod
    def _template_dir():
        configured = os.environ.get("PRODUCT_SUMMARY_TEMPLATE_DIR")
        if configured:
            return Path(configured)
        packaged = Path(__file__).resolve().parent / "packaged_assets" / "product_summary_templates"
        if packaged.exists():
            return packaged
        return Path.home() / "Desktop" / "联接基金法律文件"

    @classmethod
    def _template_path(cls, share_class):
        filename = cls.TEMPLATE_FILENAMES.get(str(share_class or "").upper())
        if not filename:
            return None
        return cls._template_dir() / filename

    @staticmethod
    def _docx_xml_escape(value):
        return html.escape(str(value or ""), quote=False)

    def _template_text_replacements(self, render_model):
        v = render_model.get("variables", {}) or {}
        share_class = str(render_model.get("share_class") or "").upper()
        fund_name = v.get("FUND_NAME") or re.sub(r"（[AC]类份额）$", "", render_model.get("fund_name", ""))
        target_etf = v.get("TARGET_ETF_NAME") or fund_name.replace("发起式联接基金", "").replace("联接基金", "")
        short_name_base = v.get("FUND_SHORT_NAME") or self.TEMPLATE_SHORT_NAME_BASE
        short_name = self._share_short_name(v, share_class)
        index_name = v.get("INDEX_NAME", self.TEMPLATE_INDEX_NAME)
        benchmark = v.get("BENCHMARK") or f"{index_name}收益率×95%+活期存款基准利率×5%"
        fund_code = self._share_code(v, share_class)
        manager_name = v.get("MANAGER_NAME", v.get("FUND_MANAGER", ""))
        custodian_name = v.get("CUSTODIAN_NAME", "")
        fund_manager_name = v.get("FUND_MANAGER_NAME", "详见招募说明书")

        replacements = [
            (f"{self.TEMPLATE_BASE_FUND_NAME}（{share_class}类份额）基金产品资料概要", f"{fund_name}（{share_class}类份额）基金产品资料概要"),
            (self.TEMPLATE_BASE_FUND_NAME, fund_name),
            (self.TEMPLATE_TARGET_ETF_NAME, target_etf),
            (f"{self.TEMPLATE_SHORT_NAME_BASE}{share_class}", short_name),
            (self.TEMPLATE_SHORT_NAME_BASE, short_name_base),
            (f"{self.TEMPLATE_INDEX_NAME}收益率×95%+活期存款基准利率×5%", benchmark),
            (self.TEMPLATE_INDEX_NAME, index_name),
            (self.TEMPLATE_CUSTODIAN_NAME, custodian_name),
            (self.TEMPLATE_FUND_MANAGER_NAME, fund_manager_name),
            ("南方基金管理股份有限公司", manager_name),
            ("XXXXX", fund_code),
            ("0.15%", f"{v.get('MGMT_FEE_RATE', '0.50')}%"),
            ("0.05%", f"{v.get('CUSTODY_FEE_RATE', '0.10')}%"),
            ("0.20%", f"{v.get('SALES_SERVICE_FEE_RATE', '0.40')}%"),
        ]
        return sorted(
            [(old, new) for old, new in replacements if old and new is not None],
            key=lambda item: len(item[0]),
            reverse=True,
        )

    @staticmethod
    def _template_table_row_replacements(render_model):
        replacements = {}
        target_labels = {"其他", "投资范围", "风险收益特征", "其他费用", "管理费", "托管费", "销售服务费"}

        def collect_rows(rows):
            for row in rows or []:
                if not row:
                    continue
                label = str(row[0] or "").strip()
                if label in target_labels:
                    replacements[label] = [str(cell or "") for cell in row[1:]]

        for section in render_model.get("sections", []):
            if section.get("type") == "table":
                collect_rows(section.get("rows"))
            for subsection in section.get("subsections", []):
                if subsection.get("type") == "table":
                    collect_rows(subsection.get("rows"))
        return replacements

    @staticmethod
    def _template_risk_disclosure_text(render_model):
        for section in render_model.get("sections", []):
            if section.get("title") != "四、风险揭示与重要提示":
                continue
            for subsection in section.get("subsections", []):
                if subsection.get("title") == "（一）风险揭示":
                    return str(subsection.get("content") or "")
        return ""

    @classmethod
    def _template_paragraphs(cls, share_class="A"):
        try:
            from docx import Document

            template_path = cls._template_path(share_class)
            if not template_path or not template_path.exists():
                return []
            return [p.text.strip() for p in Document(template_path).paragraphs if p.text.strip()]
        except Exception as e:
            logger.warning("Failed to read template paragraphs: %s", e)
            return []

    @classmethod
    def _template_section_text(cls, share_class, start_heading, end_heading=None):
        paragraphs = cls._template_paragraphs(share_class)
        start = next((idx for idx, text in enumerate(paragraphs) if text == start_heading), None)
        if start is None:
            return ""
        if end_heading:
            end = next((idx for idx, text in enumerate(paragraphs[start + 1 :], start + 1) if text == end_heading), None)
        else:
            end = len(paragraphs)
        if end is None or end <= start + 1:
            return ""
        return "\n".join(paragraphs[start + 1 : end]).strip()

    @classmethod
    def _template_risk_tail_text(cls, share_class="A"):
        paragraphs = cls._template_paragraphs(share_class)

        start = next((idx for idx, text in enumerate(paragraphs) if text == "三）管理风险"), None)
        end = next((idx for idx, text in enumerate(paragraphs) if "重要提示" in text and "（二）" in text), None)
        if start is None or end is None or end <= start:
            return ""
        return "\n".join(paragraphs[start:end]).strip()

    @classmethod
    def _apply_template_plain_replacements(cls, text, v, share_class, arbitration_notice=""):
        fund_name = v.get("FUND_NAME", "")
        target_etf = v.get("TARGET_ETF_NAME") or fund_name.replace("发起式联接基金", "").replace("联接基金", "")
        short_name_base = v.get("FUND_SHORT_NAME") or cls.TEMPLATE_SHORT_NAME_BASE
        index_name = v.get("INDEX_NAME", cls.TEMPLATE_INDEX_NAME)
        benchmark = v.get("BENCHMARK") or f"{index_name}收益率×95%+活期存款基准利率×5%"
        fund_code = ProductSummaryEngine._share_code(v, share_class)
        short_name = ProductSummaryEngine._share_short_name(v, share_class)
        manager_name = v.get("MANAGER_NAME", v.get("FUND_MANAGER", ""))
        custodian_name = v.get("CUSTODIAN_NAME", "")
        fund_manager_name = v.get("FUND_MANAGER_NAME", "详见招募说明书")

        replacements = [
            (f"{cls.TEMPLATE_BASE_FUND_NAME}（{share_class}类份额）基金产品资料概要", f"{fund_name}（{share_class}类份额）基金产品资料概要"),
            (cls.TEMPLATE_BASE_FUND_NAME, fund_name),
            (cls.TEMPLATE_TARGET_ETF_NAME, target_etf),
            (f"{cls.TEMPLATE_SHORT_NAME_BASE}{share_class}", short_name),
            (cls.TEMPLATE_SHORT_NAME_BASE, short_name_base),
            (f"{cls.TEMPLATE_INDEX_NAME}收益率×95%+活期存款基准利率×5%", benchmark),
            (cls.TEMPLATE_INDEX_NAME, index_name),
            (cls.TEMPLATE_CUSTODIAN_NAME, custodian_name),
            (cls.TEMPLATE_FUND_MANAGER_NAME, fund_manager_name),
            ("南方基金管理股份有限公司", manager_name),
            ("XXXXX", fund_code),
            ("0.15%", f"{v.get('MGMT_FEE_RATE', '0.50')}%"),
            ("0.05%", f"{v.get('CUSTODY_FEE_RATE', '0.10')}%"),
            ("0.20%", f"{v.get('SALES_SERVICE_FEE_RATE', '0.40')}%"),
        ]
        for old, new in sorted(replacements, key=lambda item: len(item[0]), reverse=True):
            if old and new is not None:
                text = text.replace(old, str(new))
        if arbitration_notice:
            lines = [
                arbitration_notice
                if ("仲裁地点为" in line and "仲裁" in line)
                else line
                for line in text.splitlines()
            ]
            text = "\n".join(lines)
        return text

    def _build_docx_from_reference_template(self, render_model):
        import zipfile
        from lxml import etree

        share_class = str(render_model.get("share_class") or "").upper()
        template_path = self._template_path(share_class)
        if not template_path or not template_path.exists():
            return None

        replacements = self._template_text_replacements(render_model)
        table_row_replacements = self._template_table_row_replacements(render_model)
        risk_disclosure_text = self._template_risk_disclosure_text(render_model)

        def replace_in_word_xml(data):
            parser = etree.XMLParser(remove_blank_text=False, recover=True)
            root = etree.fromstring(data, parser)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            xml_space = "{http://www.w3.org/XML/1998/namespace}space"
            changed = False

            def element_text(element):
                return "".join(node.text or "" for node in element.xpath(".//w:t", namespaces=ns)).strip()

            def set_paragraph_text(paragraph, value):
                text_nodes = paragraph.xpath(".//w:t", namespaces=ns)
                if not text_nodes:
                    return False
                text_nodes[0].text = str(value or "")
                text_nodes[0].set(xml_space, "preserve")
                for node in text_nodes[1:]:
                    node.text = ""
                if str(value or "").startswith("证券市场价格受到经济因素"):
                    p_pr = paragraph.find(f"{{{ns['w']}}}pPr")
                    if p_pr is None:
                        p_pr = etree.Element(f"{{{ns['w']}}}pPr")
                        paragraph.insert(0, p_pr)
                    ind = p_pr.find(f"{{{ns['w']}}}ind")
                    if ind is None:
                        ind = etree.Element(f"{{{ns['w']}}}ind")
                        p_pr.append(ind)
                    ind.set(f"{{{ns['w']}}}firstLine", "420")
                return True

            def set_cell_text(cell, value):
                paragraphs = cell.xpath("./w:p", namespaces=ns)
                if not paragraphs:
                    return False
                lines = [line.strip() for line in str(value or "").splitlines() if line.strip()]
                if not lines:
                    lines = [""]
                while len(paragraphs) < len(lines):
                    new_paragraph = copy.deepcopy(paragraphs[-1])
                    set_paragraph_text(new_paragraph, "")
                    cell.append(new_paragraph)
                    paragraphs.append(new_paragraph)
                changed_cell = False
                for idx, line in enumerate(lines):
                    changed_cell = set_paragraph_text(paragraphs[idx], line) or changed_cell
                for paragraph in paragraphs[len(lines):]:
                    parent = paragraph.getparent()
                    if parent is not None:
                        parent.remove(paragraph)
                        changed_cell = True
                return changed_cell

            def remove_paragraph(paragraph):
                parent = paragraph.getparent()
                if parent is not None:
                    parent.remove(paragraph)
                    return True
                return False

            def insert_paragraph_after(anchor):
                parent = anchor.getparent()
                if parent is None:
                    return None
                new_paragraph = copy.deepcopy(anchor)
                set_paragraph_text(new_paragraph, "")
                parent.insert(parent.index(anchor) + 1, new_paragraph)
                return new_paragraph

            def replace_risk_disclosure_block():
                if not risk_disclosure_text:
                    return False
                paragraphs = root.xpath(".//w:p", namespaces=ns)
                paragraph_texts = [element_text(paragraph) for paragraph in paragraphs]
                start_idx = next(
                    (
                        idx
                        for idx, text in enumerate(paragraph_texts)
                        if "风险揭示" in text and "（一）" in text
                    ),
                    None,
                )
                if start_idx is None:
                    return False
                end_idx = next(
                    (
                        idx
                        for idx in range(start_idx + 1, len(paragraphs))
                        if "重要提示" in paragraph_texts[idx] and "（二）" in paragraph_texts[idx]
                    ),
                    None,
                )
                if end_idx is None or end_idx <= start_idx + 1:
                    return False
                risk_lines = [line.strip() for line in risk_disclosure_text.splitlines() if line.strip()]
                block_paragraphs = paragraphs[start_idx + 1 : end_idx]
                if not block_paragraphs:
                    return False
                while len(block_paragraphs) < len(risk_lines):
                    new_paragraph = insert_paragraph_after(block_paragraphs[-1])
                    if new_paragraph is None:
                        break
                    block_paragraphs.append(new_paragraph)
                for idx, line in enumerate(risk_lines):
                    set_paragraph_text(block_paragraphs[idx], line)
                for paragraph in block_paragraphs[len(risk_lines):]:
                    remove_paragraph(paragraph)
                return True

            def replace_arbitration_notice():
                arbitration_notice = str(render_model.get("arbitration_notice") or "").strip()
                if not arbitration_notice:
                    return False
                changed_notice = False
                for paragraph in root.xpath(".//w:p", namespaces=ns):
                    text = element_text(paragraph)
                    if "仲裁地点为" in text and "仲裁" in text:
                        changed_notice = set_paragraph_text(paragraph, arbitration_notice) or changed_notice
                return changed_notice

            changed = replace_risk_disclosure_block() or changed
            changed = replace_arbitration_notice() or changed

            for row in root.xpath(".//w:tbl/w:tr", namespaces=ns):
                cells = row.xpath("./w:tc", namespaces=ns)
                if not cells:
                    continue
                label = element_text(cells[0])
                values = table_row_replacements.get(label)
                if values is None:
                    continue
                for offset, value in enumerate(values, start=1):
                    if offset >= len(cells):
                        break
                    changed = set_cell_text(cells[offset], value) or changed

            for paragraph in root.xpath(".//w:p", namespaces=ns):
                text_nodes = paragraph.xpath(".//w:t", namespaces=ns)
                if not text_nodes:
                    continue
                original = "".join(node.text or "" for node in text_nodes)
                updated = original
                for old, new in replacements:
                    updated = updated.replace(old, new)
                if updated == original:
                    continue
                text_nodes[0].text = updated
                text_nodes[0].set(xml_space, "preserve")
                for node in text_nodes[1:]:
                    node.text = ""
                changed = True

            if not changed:
                return data
            return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

        buf = io.BytesIO()
        with zipfile.ZipFile(template_path, "r") as source, zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                data = source.read(info.filename)
                if info.filename == "word/document.xml" or re.match(r"word/(header|footer)\d+\.xml$", info.filename):
                    data = replace_in_word_xml(data)
                elif info.filename.endswith(".xml"):
                    try:
                        xml = data.decode("utf-8")
                    except UnicodeDecodeError:
                        xml = ""
                    if xml:
                        updated = xml
                        for old, new in replacements:
                            updated = updated.replace(self._docx_xml_escape(old), self._docx_xml_escape(new))
                        if updated != xml:
                            data = updated.encode("utf-8")
                target.writestr(info, data)
        buf.seek(0)
        return buf.read()

    # ── 变量推导 ──────────────────────────────────────────────────────────
    def _derive_variables(self, form_data):
        v = _merge_schema_variable_defaults(form_data)
        v.setdefault("FUND_TYPE", "股票型")
        v.setdefault("OPERATION_MODE", "普通开放式")
        if not v.get("BENCHMARK"):
            v["BENCHMARK"] = f"{v.get('INDEX_NAME','')}收益率×95%+活期存款基准利率×5%"
        v.setdefault("RISK_LEVEL", "中高风险")
        # 费率默认值
        v.setdefault("MGMT_FEE_RATE", "0.50")
        v.setdefault("CUSTODY_FEE_RATE", "0.10")
        v.setdefault("SALES_SERVICE_FEE_RATE", "0.40")
        v.setdefault("TARGET_ETF_RATIO_MIN", "90")
        v["HAS_HK_CONNECT"] = _as_bool(v.get("HAS_HK_CONNECT", False))
        v["IS_INITIATING_FUND"] = _as_bool(v.get("IS_INITIATING_FUND", False))
        return v

    @staticmethod
    def _share_short_name(v, share_class):
        base = (v.get("FUND_SHORT_NAME") or v.get("FUND_NAME") or "").strip()
        if not base:
            return f"{share_class}类"
        base = re.sub(r"\s*A/C\s*$", "", base, flags=re.IGNORECASE)
        base = re.sub(r"\s*[AC]类?\s*$", "", base, flags=re.IGNORECASE)
        return f"{base}{share_class}"

    @staticmethod
    def _share_code(v, share_class):
        return str(v.get(f"FUND_CODE_{share_class}") or v.get(f"{share_class}_FUND_CODE") or v.get("FUND_CODE") or "")

    @staticmethod
    def _subsection_between(text, start_marker, end_markers):
        heading_pattern = re.compile(rf"(?m)^[ \t　]*{re.escape(start_marker)}[ \t　]*$")
        matches = list(heading_pattern.finditer(text))
        if matches:
            start = matches[-1].end()
        else:
            pos = text.find(start_marker)
            if pos == -1:
                return ""
            start = pos + len(start_marker)
        end_candidates = []
        for marker in end_markers:
            marker_pattern = re.compile(rf"(?m)^[ \t　]*{re.escape(marker)}[ \t　]*$")
            marker_match = marker_pattern.search(text, start)
            if marker_match:
                end_candidates.append(marker_match.start())
                continue
            pos = text.find(marker, start)
            if pos != -1:
                end_candidates.append(pos)
        end = min(end_candidates) if end_candidates else len(text)
        return text[start:end].strip()

    @staticmethod
    def _subsection_between_any(text, start_markers, end_markers):
        source = text or ""
        for marker in start_markers:
            block = ProductSummaryEngine._subsection_between(source, marker, end_markers)
            if block:
                return block
        return ""

    @staticmethod
    def _normalize_extracted_block(text):
        if not text:
            return ""
        return re.sub(r"\n{2,}", "\n", "\n".join(line.strip() for line in text.splitlines() if line.strip())).strip()

    @staticmethod
    def _extract_risk_disclosure_chapter(prospectus_text):
        return ProductSummaryEngine._subsection_between_any(
            prospectus_text or "",
            (
                "第十五部分  风险揭示",
                "第十五部分 风险揭示",
                "十五、风险揭示",
                "十七、风险揭示",
            ),
            [
                "第十六部分",
                "十六、基金合同的变更、终止和基金财产的清算",
                "十八、基金合同的变更、终止和基金财产的清算",
                "十八、基金合同",
            ],
        )

    @staticmethod
    def _extract_investment_range(prospectus_text):
        if not prospectus_text:
            return ""
        investment = ProductSummaryEngine._subsection_between_any(
            prospectus_text,
            (
                "第八部分  基金的投资",
                "第八部分 基金的投资",
                "八、基金的投资",
                "九、基金的投资",
                "第十二部分  基金的投资",
                "第十二部分 基金的投资",
            ),
            ["第九部分", "九、基金的财产", "十、基金的财产", "第十三部分"],
        )
        source = investment or prospectus_text
        scope = ProductSummaryEngine._subsection_between(
            source,
            "二、投资范围",
            ["三、投资策略"],
        )
        return ProductSummaryEngine._normalize_extracted_block(scope)

    @staticmethod
    def _extract_risk_return_characteristics(prospectus_text):
        if not prospectus_text:
            return ""
        investment = ProductSummaryEngine._subsection_between_any(
            prospectus_text,
            (
                "第八部分  基金的投资",
                "第八部分 基金的投资",
                "八、基金的投资",
                "九、基金的投资",
                "第十二部分  基金的投资",
                "第十二部分 基金的投资",
            ),
            ["第九部分", "九、基金的财产", "十、基金的财产", "第十三部分"],
        )
        source = investment or prospectus_text
        risk = ProductSummaryEngine._subsection_between(
            source,
            "六、风险收益特征",
            ["七、基金管理人代表基金行使股东或债权人权利的处理原则及方法", "七、", "基金的财产"],
        )
        return ProductSummaryEngine._normalize_extracted_block(risk)

    @staticmethod
    def _extract_other_operation_fees(prospectus_text):
        fee_chapter = ProductSummaryEngine._subsection_between_any(
            prospectus_text,
            (
                "第十一部分  费用与税收",
                "第十一部分  基金费用与税收",
                "第十一部分  基金的费用与税收",
                "第十一部分 费用与税收",
                "第十一部分 基金费用与税收",
                "第十一部分 基金的费用与税收",
                "十一、费用与税收",
                "十一、基金费用与税收",
                "十一、基金的费用与税收",
                "十三、基金的费用与税收",
            ),
            ["第十二部分", "十二、基金的收益与分配", "十四、基金的会计与审计"],
        )
        fee_types = ProductSummaryEngine._subsection_between(
            fee_chapter,
            "一、基金费用的种类",
            ["二、基金费用计提方法、计提标准和支付方式"],
        )
        items = []
        for raw in fee_types.splitlines():
            line = raw.strip()
            if not line:
                continue
            line = re.sub(r"^\d+、\s*", "", line)
            line = line.rstrip("；;。")
            if line in (
                "基金管理人的管理费",
                "基金托管人的托管费",
                "从C类基金份额的基金财产中计提的销售服务费",
            ):
                continue
            if line:
                items.append(line)
        if not items:
            items = [
                "《基金合同》生效后与基金相关的信息披露费用",
                "《基金合同》生效后与基金相关的会计师费、律师费、公证费、诉讼费和仲裁费",
                "基金份额持有人大会费用",
                "基金的证券/期货/期权交易费用",
                "基金的银行汇划费用",
                "基金相关账户的开户及维护费用",
                "因参与融资及转融通证券出借业务而产生的各项合理费用",
                "基金投资目标ETF的相关费用（包括但不限于目标ETF的交易费用、申赎费用等）",
                "按照国家有关规定和《基金合同》约定，可以在基金财产中列支的其他费用",
            ]
        return "；".join(items) + "。"

    @staticmethod
    def _build_default_market_risk_summary():
        return (
            "证券市场价格受到经济因素、政治因素、投资心理和交易制度等各种因素的影响，"
            "导致基金收益水平变化，产生风险，主要包括："
            "1、政策风险。"
            "2、经济周期风险。"
            "3、利率风险。"
            "4、上市公司经营风险。"
            "5、购买力风险。"
            "6、信用风险。"
            "7、债券收益率曲线变动风险。"
            "8、再投资风险。"
        )

    @staticmethod
    def _clean_risk_heading_title(line):
        line = line.strip()
        line = re.sub(r"^[（(]\d+[）)]\s*", "", line)
        line = re.sub(r"^\d+、\s*", "", line)
        line = re.split(r"[。；;：:]", line, maxsplit=1)[0].strip()
        return line if "风险" in line else ""

    @staticmethod
    def _extract_numbered_risk_items(section_text, include_title=None):
        items = []
        by_title = {}
        current = None
        include_title = include_title or (lambda _title: True)

        for raw in section_text.splitlines():
            line = raw.strip()
            if not line:
                continue
            number_match = re.match(r"^\d+、", line)
            sub_match = re.match(r"^（\d+）", line)
            if number_match:
                title = ProductSummaryEngine._clean_risk_heading_title(line)
                if title and include_title(title):
                    current = {"title": title, "subheads": []}
                    items.append(current)
                    by_title[title] = current
                else:
                    current = None
                continue
            if sub_match and current is not None:
                sub_title = ProductSummaryEngine._clean_risk_heading_title(line)
                if sub_title:
                    current["subheads"].append(sub_title)

        return items, by_title

    @staticmethod
    def _build_market_risk_summary(prospectus_text):
        fixed_intro = (
            "证券市场价格受到经济因素、政治因素、投资心理和交易制度等各种因素的影响，"
            "导致基金收益水平变化，产生风险，主要包括："
        )
        risk_chapter = ProductSummaryEngine._extract_risk_disclosure_chapter(prospectus_text)
        market_section = ProductSummaryEngine._subsection_between(
            risk_chapter,
            "一、市场风险",
            ["二、管理风险", "二、"],
        )
        product_specific_keywords = (
            "标的指数",
            "目标ETF",
            "跟踪",
            "业绩比较基准",
            "存托凭证",
            "指数编制机构",
            "成份股",
            "股指期货",
            "股票期权",
            "资产支持证券",
            "流通受限证券",
            "融资",
            "转融通",
            "基金合同终止",
            "港股通",
            "创业板",
            "科创板",
        )
        items, _by_title = ProductSummaryEngine._extract_numbered_risk_items(
            market_section,
            include_title=lambda title: not any(keyword in title for keyword in product_specific_keywords),
        )

        if not items:
            return ProductSummaryEngine._build_default_market_risk_summary()

        parts = [fixed_intro]
        for idx, item in enumerate(items, start=1):
            title = item["title"].rstrip("。")
            subheads = item["subheads"]
            if subheads:
                compact_subheads = "".join(f"（{sub_idx}）{subhead.rstrip('。')}" for sub_idx, subhead in enumerate(subheads, start=1))
                parts.append(f"{idx}、{title}：{compact_subheads}。")
            else:
                parts.append(f"{idx}、{title}。")
        return "".join(parts)

    @staticmethod
    def _build_default_specific_risk_summary(v=None):
        v = v or {}
        parts = [
            "1、标的指数的风险：标的指数编制方法或调整可能导致指数表现与总体市场表现存在差异，并增加基金跟踪成本和跟踪误差。",
            "2、标的指数波动的风险：标的指数成份股价格受多种因素影响发生波动，可能导致基金收益水平变化。",
            "3、跟踪偏离风险/跟踪误差控制未达约定目标的风险：本基金跟踪标的指数时可能因成份股调整、交易成本、申赎、停牌、费用等因素产生跟踪偏离或跟踪误差。",
            "4、投资于目标ETF带来的风险：本基金主要投资目标ETF，目标ETF的管理、操作、二级市场价格波动及折溢价等风险可能影响本基金收益。",
            "5、基金投资组合收益率与业绩比较基准收益率偏离的风险：本基金实际投资组合与业绩比较基准的构成和权重可能不同，存在收益率偏离风险。",
            "6、投资于存托凭证的风险：本基金可投资中国存托凭证，可能面临价格大幅波动、基础证券发行机制及境内外制度差异等风险。",
            "7、指数编制机构停止服务的风险：指数编制机构可能停止标的指数的管理和维护，基金可能面临转换运作方式、合并或终止基金合同等风险。",
            "8、成份股停牌、摘牌的风险：标的指数成份股可能停牌、摘牌或调整，影响基金及时交易成份股并扩大跟踪偏离和跟踪误差。",
            "9、投资股指期货的风险：本基金可投资股指期货，可能面临价格波动、流动性、基差、保证金、杠杆、信用及操作等风险。",
            "10、投资股票期权的风险：本基金可投资股票期权，期权价格受标的资产价格、波动率、到期时间和利率等因素影响，可能增加基金净值波动。",
            "11、本基金投资资产支持证券的风险：本基金可投资资产支持证券，可能面临信用、利率、流动性、提前偿付、法律和操作等风险。",
            "12、本基金投资流通受限证券的风险：本基金可投资流通受限证券，可能面临无法及时变现、折价出售及大额赎回下的流动性压力。",
            "13、本基金参与转融通证券出借业务的风险：本基金可参与转融通证券出借业务，可能面临出借证券无法及时处置、对手方履约、流动性及操作等风险。",
        ]
        if v.get("IS_INITIATING_FUND"):
            parts.append(
                "14、基金合同终止的风险：本基金为发起式基金，基金合同生效满三年后的对应日若基金资产净值低于2亿元，基金合同将自动终止。"
            )
        return "\n".join(parts)

    @staticmethod
    def _extract_specific_risk_text(prospectus_text):
        risk_chapter = ProductSummaryEngine._extract_risk_disclosure_chapter(prospectus_text)
        specific = ProductSummaryEngine._subsection_between(
            risk_chapter,
            "四、本基金特有的风险",
            ["五、其他风险", "五、基金风险评价", "五、"],
        )
        return ProductSummaryEngine._normalize_extracted_block(specific)

    @staticmethod
    def _extract_numbered_risk_blocks(section_text):
        blocks = []
        current = None

        for raw in (section_text or "").splitlines():
            line = raw.strip()
            if not line:
                continue
            if re.match(r"^\d+、", line):
                title = ProductSummaryEngine._clean_risk_heading_title(line) or re.sub(r"^\d+、\s*", "", line).strip()
                current = {"title": title.rstrip("。"), "lines": []}
                blocks.append(current)
                tail = re.sub(r"^\d+、\s*", "", line).strip()
                if "：" in tail:
                    current["lines"].append(tail.split("：", 1)[1].strip())
                elif ":" in tail:
                    current["lines"].append(tail.split(":", 1)[1].strip())
                continue
            if current is not None:
                current["lines"].append(line)

        return blocks

    @staticmethod
    def _first_risk_sentence(lines, max_len=120):
        text_parts = []
        for line in lines or []:
            if re.match(r"^（\d+）", line):
                continue
            if line.startswith("除上述风险外"):
                continue
            text_parts.append(line)
            if re.search(r"[。；;]$", line):
                break

        text = re.sub(r"\s+", "", "".join(text_parts))
        text = re.sub(r"包括但不限于以下风险[:：]?$", "存在相关风险", text)
        match = re.search(r"(.{1,%d}?[。；;])" % max_len, text)
        if match:
            return match.group(1).rstrip("；;")
        if len(text) > max_len:
            return text[:max_len].rstrip("，、；;：:") + "。"
        return text

    @staticmethod
    def _specific_risk_summary_for_title(title, lines):
        rules = [
            (("标的指数的风险",), "标的指数编制方法或调整可能导致指数表现与总体市场表现存在差异，并增加基金跟踪成本和跟踪误差。"),
            (("标的指数波动",), "标的指数成份股价格受多种因素影响发生波动，可能导致基金收益水平变化。"),
            (("跟踪偏离",), "本基金跟踪标的指数时可能因成份股调整、交易成本、申赎、停牌、费用等因素产生跟踪偏离或跟踪误差。"),
            (("目标ETF",), "本基金主要投资目标ETF，目标ETF的管理、操作、二级市场价格波动及折溢价等风险可能影响本基金收益。"),
            (("业绩比较基准",), "本基金实际投资组合与业绩比较基准的构成和权重可能不同，存在收益率偏离风险。"),
            (("存托凭证",), "本基金可投资中国存托凭证，可能面临价格大幅波动、基础证券发行机制及境内外制度差异等风险。"),
            (("指数编制机构停止服务",), "指数编制机构可能停止标的指数的管理和维护，基金可能面临转换运作方式、合并或终止基金合同等风险。"),
            (("成份股停牌",), "标的指数成份股可能停牌、摘牌或调整，影响基金及时交易成份股并扩大跟踪偏离和跟踪误差。"),
            (("股指期货",), "本基金可投资股指期货，可能面临价格波动、流动性、基差、保证金、杠杆、信用及操作等风险。"),
            (("股票期权",), "本基金可投资股票期权，期权价格受标的资产价格、波动率、到期时间和利率等因素影响，可能增加基金净值波动。"),
            (("资产支持证券",), "本基金可投资资产支持证券，可能面临信用、利率、流动性、提前偿付、法律和操作等风险。"),
            (("流通受限证券",), "本基金可投资流通受限证券，可能面临无法及时变现、折价出售及大额赎回下的流动性压力。"),
            (("转融通证券出借",), "本基金可参与转融通证券出借业务，可能面临出借证券无法及时处置、对手方履约、流动性及操作等风险。"),
            (("基金合同终止",), "本基金为发起式基金，基金合同生效满三年后的对应日若基金资产净值低于2亿元，基金合同将自动终止。"),
        ]
        for keywords, summary in rules:
            if all(keyword in title for keyword in keywords):
                return summary
        return ProductSummaryEngine._first_risk_sentence(lines)

    @staticmethod
    def _build_specific_risk_summary(prospectus_text):
        return ProductSummaryEngine._extract_specific_risk_text(prospectus_text)

    @staticmethod
    def _extract_arbitration_notice(prospectus_text):
        for raw in (prospectus_text or "").splitlines():
            line = raw.strip()
            if "仲裁地点为" in line and "仲裁" in line and "《基金合同》" in line:
                return line
        return ""

    # ── 表格1：产品概况 (9行×4列) ─────────────────────────────────────────
    def _build_overview_table(self, v, share_class):
        fund_manager_start = v.get("FUND_MANAGER_START_DATE", "")
        fund_manager_license_date = v.get("FUND_MANAGER_SECURITIES_DATE", v.get("FUND_MANAGER_START_WORK_DATE", "XXXX年X月X日"))
        rows = [
            ["基金简称", self._share_short_name(v, share_class), "基金代码", self._share_code(v, share_class)],
            ["基金管理人", v.get("MANAGER_NAME", v.get("FUND_MANAGER", "")),
             "基金托管人", v.get("CUSTODIAN_NAME", "")],
            ["基金合同生效日", v.get("CONTRACT_EFFECTIVE_DATE", ""), "", ""],
            ["基金类型", v.get("FUND_TYPE", ""), "交易币种", "人民币"],
            ["运作方式", v.get("OPERATION_MODE", ""), "", ""],
            ["开放频率", "每个开放日", "", ""],
            ["基金经理", v.get("FUND_MANAGER_NAME", "详见招募说明书"), "开始担任本基金基金经理的日期", fund_manager_start],
            ["基金经理", v.get("FUND_MANAGER_NAME", "详见招募说明书"), "证券从业日期", fund_manager_license_date],
        ]
        other_text = ""
        if v.get("IS_INITIATING_FUND"):
            other_text = (
                "《基金合同》生效之日起3年后的对应日，若基金资产净值低于2亿元，基金合同自动终止，"
                "且不得通过召开基金份额持有人大会延续基金合同期限。法律法规或中国证监会另有规定的，从其规定。"
            )
        rows.append(["其他", other_text, "", ""])
        return rows

    # ── 表格2：投资目标与策略 (5行×2列) ──────────────────────────────────
    def _build_investment_table(self, v, risk_return_text, prospectus_text):
        index_name = v.get("INDEX_NAME", "标的指数")
        extracted_scope = self._extract_investment_range(prospectus_text)

        if not extracted_scope:
            extracted_scope = (
                "本基金主要投资于目标ETF基金份额、标的指数成份股、备选成份股（含存托凭证）。"
                "为更好地实现投资目标，本基金可少量投资于非成份股（包含主板、科创板、创业板及其他经中国证监会核准或注册发行的股票、存托凭证）、"
                "金融衍生品（股指期货、股票期权等）、债券（包括国债、金融债、企业债、公司债、政府机构债券、地方政府债券、次级债、可转换债券、"
                "可交换债券、央行票据、中期票据、短期融资券、超短期融资券等）、资产支持证券、债券回购、银行存款（包括协议存款、定期存款及其他银行存款）、"
                "同业存单、货币市场工具以及法律法规或中国证监会允许基金投资的其他金融工具（但须符合中国证监会的相关规定）。\n"
                "本基金根据相关规定可参与融资、转融通证券出借业务。\n"
                "如法律法规或监管机构以后允许基金投资其他品种，基金管理人在履行适当程序后，可以将其纳入投资范围。\n"
                f"基金的投资组合比例为：在建仓完成后，本基金投资于目标ETF的比例不低于基金资产净值的{v.get('TARGET_ETF_RATIO_MIN','90')}%。"
                "每个交易日日终在扣除金融衍生品合约需缴纳的交易保证金后，现金或到期日在一年以内的政府债券的投资比例不低于基金资产净值的5%，"
                "其中现金不包括结算备付金、存出保证金、应收申购款等。\n"
                "如法律法规或监管机构变更投资品种的投资比例限制，基金管理人在履行适当程序后，可以调整上述投资品种的投资比例。"
            )
        index_sentence = f"本基金的标的指数为{index_name}。"
        inv_range = extracted_scope
        if "本基金的标的指数为" not in inv_range:
            inv_range = f"{inv_range}\n{index_sentence}" if inv_range else index_sentence

        inv_strategy = (
            "本基金为被动式指数基金，主要投资于目标ETF。本基金并不参与目标ETF的管理。"
            "在正常市场情况下，本基金力争净值增长率与业绩比较基准之间的日均跟踪偏离度不超过0.35%，年跟踪误差不超过4%。"
            "如因指数编制规则调整或其他因素导致跟踪偏离度和跟踪误差超过上述范围，基金管理人应采取合理措施避免跟踪偏离度、跟踪误差进一步扩大。"
            "具体投资策略包括：1、资产配置策略；2、目标ETF投资策略；3、成份股、备选成份股投资策略；4、金融衍生品投资策略；"
            "5、债券投资策略；6、可转换债券及可交换债券投资策略；7、资产支持证券投资策略；8、融资及转融通证券出借业务投资策略；9、存托凭证投资策略。"
        )

        return [
            ["投资目标",
             "本基金通过投资于目标ETF，紧密跟踪标的指数，追求与业绩比较基准相似的回报。"],
            ["投资范围", inv_range],
            ["投资策略", inv_strategy],
            ["业绩比较基准", v.get("BENCHMARK", "")],
            ["风险收益特征", risk_return_text],
        ]

    # ── 表格3：基金销售相关费用 ───────────────────────────────────────────
    def _build_sales_fee_table(self, share_class):
        redemption_rows = [
            ["赎回费", "N＜7天", "1.5%", "个人投资者"],
            ["赎回费", "7天≤N", "0", "个人投资者"],
            ["赎回费", "N＜7天", "1.5%", "个人投资者以外的其他投资者"],
            ["赎回费", "7天≤N＜30天", "1.0%", "个人投资者以外的其他投资者"],
            ["赎回费", "30天≤N＜180天", "0.5%", "个人投资者以外的其他投资者"],
            ["赎回费", "180天≤N", "0", "个人投资者以外的其他投资者"],
        ]
        if share_class == "C":
            return [["费用类型", "份额（S）或金额（M）/持有期限（N）", "收费方式/费率", "备注"]] + redemption_rows
        return [
            ["费用类型", "份额（S）或金额（M）/持有期限（N）", "收费方式/费率", "备注"],
            ["认购费", "M＜100万元", "1.00%", ""],
            ["认购费", "100万元≤M＜300万元", "0.60%", ""],
            ["认购费", "300万元≤M＜500万元", "0.30%", ""],
            ["认购费", "M≥500万元", "每笔1,000元", ""],
            ["申购费（前端）", "M＜100万元", "1.20%", ""],
            ["申购费（前端）", "100万元≤M＜300万元", "0.80%", ""],
            ["申购费（前端）", "300万元≤M＜500万元", "0.40%", ""],
            ["申购费（前端）", "M≥500万元", "每笔1,000元", ""],
        ] + redemption_rows

    # ── 表格4：基金运作相关费用 ───────────────────────────────────────────
    def _build_operation_fee_table(self, v, share_class, other_fees_text):
        mgmt_rate = v.get("MGMT_FEE_RATE", "0.50")
        custody_rate = v.get("CUSTODY_FEE_RATE", "0.10")
        svc_rate = v.get("SALES_SERVICE_FEE_RATE", "0.40")
        rows = [
            ["费用类别", "收费方式/年费率或金额", "收取方"],
            ["管理费",
             f"{mgmt_rate}%",
             "基金管理人和销售机构"],
            ["托管费",
             f"{custody_rate}%",
             "基金托管人"],
        ]
        if share_class == "C":
            rows.append(
                ["销售服务费",
                 f"{svc_rate}%",
                 "代销机构"]
            )
        rows.append(["其他费用", other_fees_text, other_fees_text])
        return rows

    # ── 风险揭示 ──────────────────────────────────────────────────────────
    def _build_risk_section(self, v, prospectus_text, share_class="A"):
        lines = [
            "本基金不提供任何保证。投资者可能损失本金。",
            "投资有风险，投资者购买基金时应认真阅读本基金的《招募说明书》等销售文件。",
        ]
        specific_risk = self._build_specific_risk_summary(prospectus_text)
        if not specific_risk:
            specific_risk = self._build_default_specific_risk_summary(v)
        lines.extend(["一）本基金特有的风险", specific_risk])
        lines.extend(["二）市场风险", self._build_market_risk_summary(prospectus_text)])
        risk_tail = self._template_risk_tail_text(share_class)
        if risk_tail:
            lines.append(risk_tail)
        return "\n".join(lines)

    # ── 重要提示 ──────────────────────────────────────────────────────────
    def _build_important_notice(self, v, share_class="A", arbitration_notice=""):
        template_text = self._template_section_text(share_class, "（二） 重要提示", "五、其他资料查询方式")
        if template_text:
            return self._apply_template_plain_replacements(template_text, v, share_class, arbitration_notice)

        is_init = v.get("IS_INITIATING_FUND", False)
        lines = [
            "本基金经中国证监会注册，但中国证监会对本基金募集的注册，并不表明其对本基金的投资价值和市场前景做出实质性判断或保证，也不表明投资于本基金没有风险。",
            "投资有风险，投资者在投资本基金前应认真阅读本基金的招募说明书、基金合同和基金产品资料概要等信息披露文件。",
            "基金管理人依照恪尽职守、诚实信用、谨慎勤勉的原则管理和运用基金财产，但不保证基金一定盈利，也不保证最低收益。",
        ]
        if is_init:
            lines.append(
                "本基金为发起式基金，在基金合同生效之日起三年后的对应日，若基金资产净值低于2亿元，"
                "基金合同自动终止，且不得通过召开基金份额持有人大会延续基金合同期限。投资者将面临基金合同可能终止的不确定性风险。"
            )
        return "\n".join(lines)

    def _build_template_text_section(self, v, share_class, start_heading, end_heading, fallback_text, arbitration_notice=""):
        template_text = self._template_section_text(share_class, start_heading, end_heading)
        if not template_text:
            return fallback_text
        return self._apply_template_plain_replacements(template_text, v, share_class, arbitration_notice)

    # ── 组装全部章节 ──────────────────────────────────────────────────────
    def _build_sections(self, v, share_class, prospectus_text):
        risk_return_text = self._extract_risk_return_characteristics(prospectus_text) or (
            "本基金的目标ETF为股票型基金，一般而言，其长期平均风险与预期收益高于混合型基金、债券型基金与货币市场基金。"
            "本基金为ETF联接基金，通过投资于目标ETF跟踪标的指数表现，具有与标的指数以及标的指数所代表的市场相似的风险收益特征。"
        )
        other_fees_text = self._extract_other_operation_fees(prospectus_text)
        overview = self._build_overview_table(v, share_class)
        investment = self._build_investment_table(v, risk_return_text, prospectus_text)
        sales_fee = self._build_sales_fee_table(share_class)
        operation_fee = self._build_operation_fee_table(v, share_class, other_fees_text)
        risk_text = self._build_risk_section(v, prospectus_text, share_class)
        arbitration_notice = self._extract_arbitration_notice(prospectus_text)
        notice_text = self._build_important_notice(v, share_class, arbitration_notice)
        query_text = self._build_template_text_section(
            v,
            share_class,
            "五、其他资料查询方式",
            "六、其他情况说明",
            (
                "以下资料详见基金管理人网站[www.nffund.com][客服电话：400-889-8899]：\n"
                f"1、《{v.get('FUND_NAME', '')}基金合同》、《{v.get('FUND_NAME', '')}托管协议》、《{v.get('FUND_NAME', '')}招募说明书》\n"
                "2、定期报告，包括基金季度报告、中期报告和年度报告\n"
                "3、基金份额净值\n"
                "4、基金销售机构及联系方式\n"
                "5、其他重要资料"
            ),
            arbitration_notice,
        )
        other_note_text = self._build_template_text_section(
            v,
            share_class,
            "六、其他情况说明",
            None,
            "暂无。",
            arbitration_notice,
        )

        base_fund_name = v.get("FUND_NAME", "")
        fund_name = f"{base_fund_name}（{share_class}类份额）"

        # 组装 render_model
        render_model = {
            "fund_name": fund_name,
            "arbitration_notice": arbitration_notice,
            "sections": [
                {"title": "一、产品概况", "type": "table", "rows": overview},
                {"title": "二、基金投资与净值表现",
                 "type": "mixed",
                 "subsections": [
                     {"title": "（一）投资目标与投资策略", "type": "table", "rows": investment},
                     {"title": "（二）投资组合资产配置图表/区域配置图表",
                      "type": "text",
                      "content": "无"},
                 ]},
                {"title": "三、投资本基金涉及的费用",
                 "type": "mixed",
                 "subsections": [
                     {"title": "（一）基金销售相关费用", "type": "table", "rows": sales_fee},
                     {"title": "（二）基金运作相关费用", "type": "table", "rows": operation_fee,
                      "note": "注：本基金交易证券、基金等产生的费用和税负，按实际发生额从基金资产扣除。\n本基金基金财产中投资于目标ETF的部分不收取管理费。本基金基金财产中投资于目标ETF的部分不收取托管费。"},
                 ]},
                {"title": "四、风险揭示与重要提示",
                 "type": "mixed",
                 "subsections": [
                     {"title": "（一）风险揭示", "type": "text", "content": risk_text},
                     {"title": "（二）重要提示", "type": "text", "content": notice_text},
                 ]},
                {"title": "五、其他资料查询方式",
                 "type": "text",
                 "content": query_text},
                {"title": "六、其他情况说明",
                 "type": "text",
                 "content": other_note_text},
            ]
        }
        return render_model

    # ── render_model → 纯文本 ─────────────────────────────────────────────
    def _render_model_to_text(self, render_model):
        """将 render_model 转换为 pipe-format 纯文本"""
        lines = []
        fund_name = render_model.get("fund_name", "")
        lines.append(f"# {fund_name}基金产品资料概要")
        lines.append("")
        lines.append(f"基金管理人：{render_model.get('fund_manager','详见招募说明书')}")
        lines.append(f"基金托管人：{render_model.get('custodian_name','详见招募说明书')}")
        lines.append("")

        for i, sec in enumerate(render_model.get("sections", [])):
            if i > 0:
                lines.append("")
            lines.append(f"## {sec['title']}")
            lines.append("")

            if sec["type"] == "table":
                self._render_table(lines, sec["rows"])
            elif sec["type"] == "text":
                lines.append(sec.get("content", ""))
            elif sec["type"] == "mixed":
                for sub_idx, sub in enumerate(sec.get("subsections", [])):
                    if sub_idx > 0:
                        lines.append("")
                    if sub.get("title"):
                        lines.append(f"### {sub['title']}")
                        lines.append("")
                    if sub["type"] == "table":
                        self._render_table(lines, sub["rows"])
                        if sub.get("note"):
                            lines.append("")
                            lines.append(sub["note"])
                    elif sub["type"] == "text":
                        lines.append(sub.get("content", ""))

        return "\n".join(lines)

    def _render_table(self, lines, rows):
        """将二维列表渲染为 pipe-format 表格"""
        if not rows:
            return
        col_count = max(len(r) for r in rows)

        def render_row(row):
            cells = []
            for c in range(col_count):
                cell = row[c] if c < len(row) else ""
                cells.append(str(cell) if cell is not None else "")
            return "| " + " | ".join(cells) + " |"

        lines.append(render_row(rows[0]))
        lines.append("| " + " | ".join(["---"] * col_count) + " |")
        for row in rows[1:]:
            lines.append(render_row(row))

    # ── generate / generate_bundle ────────────────────────────────────────
    def generate(self, form_data):
        return self.generate_bundle(form_data)["text"]

    def generate_bundle(self, form_data):
        v = self._derive_variables(form_data)
        prospectus_text = str(form_data.get("PROSPECTUS_TEXT") or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not prospectus_text:
            prospectus_text = self.prospectus_engine.generate(v)
        render_models = []
        for share_class in ("A", "C"):
            render_model = self._build_sections(v, share_class, prospectus_text)
            render_model["share_class"] = share_class
            render_model["fund_manager"] = v.get("MANAGER_NAME", v.get("FUND_MANAGER", ""))
            render_model["custodian_name"] = v.get("CUSTODIAN_NAME", "")
            render_model["variables"] = dict(v)
            render_models.append(render_model)
        text = "\n\n---\n\n".join(self._render_model_to_text(model) for model in render_models)
        return {
            "text": text,
            "render_models": render_models,
            "render_model": render_models[0],
            "variables": v,
        }

    @staticmethod
    def _safe_filename(name):
        return re.sub(r'[\\/:*?"<>|]', "_", name)

    def build_docx_bundle(self, render_models):
        import zipfile

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for model in render_models:
                share_class = model.get("share_class", "")
                filename = self._safe_filename(f"{model.get('fund_name', 'ETF联接基金')}基金产品资料概要.docx")
                if share_class and f"{share_class}类份额" not in filename:
                    filename = self._safe_filename(
                        f"{model.get('fund_name', 'ETF联接基金')}（{share_class}类份额）基金产品资料概要.docx"
                    )
                archive.writestr(filename, self.build_docx(model))
        buf.seek(0)
        return buf.read()

    # ── DOCX 导出 ─────────────────────────────────────────────────────────
    def build_docx(self, render_model, form_data=None):
        """从 render_model 生成产品资料概要 DOCX"""
        template_docx = self._build_docx_from_reference_template(render_model)
        if template_docx is not None:
            return template_docx

        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Twips, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        sec = doc.sections[0]
        sec.page_width = Twips(PAGE_LAYOUT["page_width"])
        sec.page_height = Twips(PAGE_LAYOUT["page_height"])
        sec.top_margin = Twips(PAGE_LAYOUT["top_margin"])
        sec.bottom_margin = Twips(PAGE_LAYOUT["bottom_margin"])
        sec.left_margin = Twips(PAGE_LAYOUT["left_margin"])
        sec.right_margin = Twips(PAGE_LAYOUT["right_margin"])

        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{render_model.get('fund_name','')}基金产品资料概要")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 基本信息行
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info.add_run(
            f"基金管理人：{render_model.get('fund_manager','')}    基金托管人：{render_model.get('custodian_name','')}"
        )
        info_run.font.size = Pt(9)

        doc.add_paragraph()  # 空行

        for sec in render_model.get("sections", []):
            # 章节标题
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            r = h.add_run(sec["title"])
            r.bold = True
            r.font.size = Pt(12)
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            if sec["type"] == "table":
                self._add_docx_table(doc, sec["rows"])
            elif sec["type"] == "text":
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                run = p.add_run(sec.get("content", ""))
                run.font.size = Pt(10.5)
            elif sec["type"] == "mixed":
                for sub in sec.get("subsections", []):
                    if sub.get("title"):
                        sh = doc.add_paragraph()
                        sh.paragraph_format.space_before = Pt(8)
                        sr = sh.add_run(sub["title"])
                        sr.bold = True
                        sr.font.size = Pt(10.5)
                    if sub["type"] == "table":
                        self._add_docx_table(doc, sub["rows"])
                        if sub.get("note"):
                            np = doc.add_paragraph()
                            nr = np.add_run(sub["note"])
                            nr.font.size = Pt(9)
                            nr.italic = True
                    elif sub["type"] == "text":
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Cm(0.74)
                        run = p.add_run(sub.get("content", ""))
                        run.font.size = Pt(10.5)

        # 保存
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def _add_docx_table(self, doc, rows):
        """向文档添加格式化表格"""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn

        if not rows:
            return
        col_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=col_count, style='Table Grid')
        table.alignment = 1  # center

        for r_idx, row_data in enumerate(rows):
            for c_idx in range(col_count):
                cell = table.cell(r_idx, c_idx)
                text = str(row_data[c_idx]) if c_idx < len(row_data) else ""
                # 清除默认段落
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(text)
                run.font.size = Pt(9)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if r_idx == 0:
                    run.bold = True
                    # 表头灰色背景
                    shading = cell._element.get_or_add_tcPr()
                    shd = shading.makeelement(qn('w:shd'), {
                        qn('w:fill'): 'D9D9D9',
                        qn('w:val'): 'clear',
                    })
                    shading.append(shd)

        doc.add_paragraph()  # 表后空行


# ═══════════════════════════════════════════════════════════════════════════════
#  Flask Routes
# ═══════════════════════════════════════════════════════════════════════════════

generation_engine_lock = threading.Lock()

def _reload_generation_engines() -> None:
    global engine, prospectus_engine, product_summary_engine
    with generation_engine_lock:
        engine = ContractEngine()
        prospectus_engine = ProspectusEngine()
        product_summary_engine = ProductSummaryEngine()


def _generation_quality_check(text: str, doc_type: str = "") -> dict:
    """Post-generation sanity checks for conditional-block joins."""
    issues = []
    lines = str(text or "").splitlines()

    def add_issue(severity, code, message, line_no=None, snippet="", extra=None):
        issue = {
            "severity": severity,
            "code": code,
            "message": message,
            "line": line_no,
            "snippet": (snippet or "").strip()[:180],
        }
        if extra:
            issue.update(extra)
        issues.append(issue)

    unresolved_patterns = [
        ("UNRESOLVED_CONDITION", r"\{\{\s*(?:IF|IF_NOT|IF\s+NOT|ENDIF|ELSE)\b[^}]*\}\}", "存在未解析条件语句标记"),
        ("UNRESOLVED_TEMPLATE", r"\{\{\s*[^}]+?\s*\}\}", "存在未解析模板占位符"),
        ("UNRESOLVED_PLACEHOLDER", r"\{[A-Z][A-Z0-9_]{2,}\}", "存在未替换的大写变量占位符"),
        ("UNRESOLVED_ALIAS", r"JINJA_VAR_\d+|__PROSPECTUS_LITERAL_\d+__", "存在模板渲染中间变量残留"),
        ("BAD_CHAR", r"�", "存在疑似编码异常字符"),
    ]
    for code, pattern, message in unresolved_patterns:
        for match in re.finditer(pattern, text or ""):
            line_no = (text or "")[:match.start()].count("\n") + 1
            add_issue("error", code, message, line_no, match.group(0))
            if sum(1 for issue in issues if issue["code"] == code) >= 5:
                break

    normalized_seen = {}
    previous_nonblank = None
    previous_heading = None
    previous_number = None
    previous_number_line = None
    current_duplicate_scope = ""
    reset_heading_pattern = re.compile(r"^(?:第[一二三四五六七八九十百]+部分|[一二三四五六七八九十百]+、|##+|###)")
    number_pattern = re.compile(r"^\s*(\d+)、")

    for idx, raw in enumerate(lines, start=1):
        line = raw.strip()
        if not line:
            previous_nonblank = None
            previous_number = None
            previous_number_line = None
            continue

        compact = re.sub(r"\s+", "", line)
        if previous_nonblank and compact == previous_nonblank["compact"] and len(compact) >= 6:
            add_issue("warning", "ADJACENT_DUPLICATE_TEXT", "相邻文本重复，可能是条件块拼接重复", idx, line)
        previous_nonblank = {"compact": compact, "line": idx}

        if reset_heading_pattern.match(line):
            previous_number = None
            previous_number_line = None

        heading_like = bool(re.match(r"^(?:第[一二三四五六七八九十百]+部分|[一二三四五六七八九十百]+、|（[一二三四五六七八九十百]+）|##+)", line))
        if heading_like:
            current_duplicate_scope = compact
            if previous_heading and previous_heading["compact"] == compact:
                add_issue("warning", "ADJACENT_DUPLICATE_HEADING", "相邻标题重复，可能是条件块保留了两份标题", idx, line)
            previous_heading = {"compact": compact, "line": idx}
        else:
            previous_heading = None

        number_match = number_pattern.match(line)
        if number_match:
            current_number = int(number_match.group(1))
            if previous_number is not None and current_number == previous_number:
                add_issue(
                    "warning",
                    "REPEATED_NUMBER",
                    "同级编号连续重复，可能是条件块删除后编号错乱",
                    idx,
                    f"上一编号行 {previous_number_line}；当前：{line}",
                )
            if previous_number is not None and current_number > previous_number + 1 and current_number <= 20:
                add_issue(
                    "warning",
                    "SKIPPED_NUMBER",
                    "同级编号疑似跳号，请核对条件块附近文本",
                    idx,
                    f"上一编号 {previous_number}，当前编号 {current_number}：{line}",
                )
            previous_number = current_number
            previous_number_line = idx

        formula_like = bool(re.match(r"^[A-Z]\u4e3a", compact))
        if len(compact) >= 36 and not formula_like:
            first_seen = normalized_seen.get(compact)
            if (
                first_seen
                and idx - first_seen["line"] <= 12
                and first_seen.get("scope") == current_duplicate_scope
            ):
                add_issue(
                    "warning",
                    "NEAR_DUPLICATE_PARAGRAPH",
                    "同一小节内短距离出现重复段落，可能是条件块文本重复",
                    idx,
                    line,
                    {"related_line": first_seen["line"]},
                )
            normalized_seen.setdefault(compact, {"line": idx, "scope": current_duplicate_scope})

    # Focus on conditional-join scars: empty headings or punctuation stranded before headings.
    for idx, raw in enumerate(lines[:-1], start=1):
        line = raw.strip()
        next_line = lines[idx].strip()
        if line.endswith(("：", "、", "（")) and re.match(r"^(?:第[一二三四五六七八九十百]+部分|[一二三四五六七八九十百]+、|##+)", next_line):
            add_issue("warning", "DANGLING_PUNCTUATION", "标题前一行疑似断裂，请核对条件语句附近文本", idx, line)

    issue_counts = {
        "error": sum(1 for issue in issues if issue["severity"] == "error"),
        "warning": sum(1 for issue in issues if issue["severity"] == "warning"),
    }
    return {
        "doc_type": doc_type,
        "ok": issue_counts["error"] == 0,
        "issue_counts": issue_counts,
        "issues": issues[:50],
    }


def _quality_error_response(text: str, doc_type: str):
    quality = _generation_quality_check(text, doc_type)
    if not quality["ok"]:
        return jsonify({
            "success": False,
            "error": f"{doc_type}生成后检查发现未解析条件或占位符，请先修复后导出。",
            "quality_check": quality,
        }), 400
    return None

# 全局引擎实例 — 线程安全前提：引擎方法在请求处理中只读 self 属性（clauses/schema），
# 不修改 self 状态。若未来需要在请求中修改引擎状态，必须改为每请求创建实例或加锁。
engine = ContractEngine()
prospectus_engine = ProspectusEngine()
product_summary_engine = ProductSummaryEngine()


_VARIABLE_CONDITION_TYPES = {"boolean", "condition", "condition_boolean"}
_VARIABLE_CONDITION_PREFIXES = ("IS_", "HAS_", "USE_", "NOT_", "INCLUDE_", "SHOW_")
_VARIABLE_OPTION_MISSING = object()


def _schema_json_error(message: str, status: int = 400):
    return jsonify({"ok": False, "success": False, "error": message}), status


def _load_schema_json() -> dict:
    with open(SCHEMA_JSON, encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        return {"groups": {}}
    data.setdefault("groups", {})
    return data


def _save_schema_json(data: dict) -> None:
    data["last_updated"] = datetime.now(timezone.utc).date().isoformat()
    with open(SCHEMA_JSON, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
        f.write("\n")


def _schema_variable_key(item: dict) -> str:
    return str((item or {}).get("key") or (item or {}).get("name") or "").strip()


def _schema_variable_label(item: dict) -> str:
    return str(
        (item or {}).get("label")
        or (item or {}).get("ui_entry")
        or (item or {}).get("applies_to")
        or _schema_variable_key(item)
    ).strip()


def _schema_variable_group_label(group: dict, group_key: str) -> str:
    return str((group or {}).get("label") or (group or {}).get("description") or group_key).strip()


def _schema_variable_is_condition(item: dict) -> bool:
    key = _schema_variable_key(item)
    field_type = str((item or {}).get("type") or "").strip()
    return field_type in _VARIABLE_CONDITION_TYPES or key.startswith(_VARIABLE_CONDITION_PREFIXES)


def _schema_variable_default(item: dict):
    if "default" in (item or {}):
        return item.get("default")
    return False if _schema_variable_is_condition(item) else ""


def _schema_variable_item_payload(item: dict, group_key: str, group: dict) -> dict:
    key = _schema_variable_key(item)
    options = (item or {}).get("options")
    return {
        "key": key,
        "label": _schema_variable_label(item),
        "type": str((item or {}).get("type") or "string"),
        "default": _schema_variable_default(item),
        "group": group_key,
        "group_label": _schema_variable_group_label(group, group_key),
        "required": bool((item or {}).get("required", False)),
        "usage": str((item or {}).get("usage") or (item or {}).get("applies_to") or ""),
        "source_field": str((item or {}).get("source_field") or ""),
        "selection_rule": str((item or {}).get("selection_rule") or ""),
        "options": options,
        "options_json": "" if options in (None, "") else json.dumps(options, ensure_ascii=False, indent=2),
        "is_condition": _schema_variable_is_condition(item),
    }


def _schema_variable_group_options(data: dict) -> list[dict]:
    return [
        {"key": group_key, "label": _schema_variable_group_label(group, group_key)}
        for group_key, group in (data.get("groups") or {}).items()
    ]


def _flatten_schema_variables(data: dict) -> list[dict]:
    items = []
    for group_key, group in (data.get("groups") or {}).items():
        for item in group.get("variables") or []:
            if isinstance(item, dict) and _schema_variable_key(item):
                items.append(_schema_variable_item_payload(item, group_key, group))
    return items


def _find_schema_variable(data: dict, key: str):
    for group_key, group in (data.get("groups") or {}).items():
        variables = group.get("variables") or []
        for index, item in enumerate(variables):
            if isinstance(item, dict) and _schema_variable_key(item) == key:
                return group_key, group, index, item
    return None, None, None, None


def _parse_schema_variable_boolean(value) -> bool:
    if isinstance(value, bool):
        return value
    return str(value or "").strip().lower() in {"1", "true", "yes", "y", "on", "是", "启用", "开启"}


def _schema_variable_generation_default(item: dict):
    if "default" not in (item or {}):
        return _VARIABLE_OPTION_MISSING
    default = item.get("default")
    if _schema_variable_is_condition(item):
        return _parse_schema_variable_boolean(default)
    return default


def _merge_schema_variable_defaults(values: dict | None) -> dict:
    merged = dict(values or {})
    try:
        data = _load_schema_json()
    except Exception:
        return merged
    for group in (data.get("groups") or {}).values():
        for item in group.get("variables") or []:
            if not isinstance(item, dict):
                continue
            key = _schema_variable_key(item)
            if not key or key in merged:
                continue
            default = _schema_variable_generation_default(item)
            if default is _VARIABLE_OPTION_MISSING:
                continue
            merged[key] = default
    return merged


def _parse_schema_variable_options(payload: dict):
    if "options_json" not in payload and "options" not in payload:
        return _VARIABLE_OPTION_MISSING
    raw = payload.get("options_json", payload.get("options"))
    if raw in (None, ""):
        return None
    if isinstance(raw, (dict, list)):
        return raw
    return json.loads(str(raw))


def _normalize_schema_variable_payload(payload: dict, existing_key: str | None = None) -> dict:
    payload = payload or {}
    key = str(payload.get("key") or payload.get("name") or existing_key or "").strip()
    if not key:
        raise ValueError("变量名不能为空")
    if not key.replace("_", "").isalnum():
        raise ValueError("变量名只能包含字母、数字和下划线")
    field_type = str(payload.get("type") or "string").strip() or "string"
    if not field_type.replace("_", "").isalnum():
        raise ValueError("变量类型只能包含字母、数字和下划线")
    try:
        options = _parse_schema_variable_options(payload)
    except json.JSONDecodeError as exc:
        raise ValueError(f"选项 JSON 格式错误：{exc}") from exc
    default = payload.get("default")
    if field_type in _VARIABLE_CONDITION_TYPES:
        default = _parse_schema_variable_boolean(default)
    elif default is None:
        default = ""
    return {
        "key": key,
        "label": str(payload.get("label") or payload.get("ui_entry") or key).strip(),
        "type": field_type,
        "default": default,
        "group": str(payload.get("group") or "custom").strip() or "custom",
        "group_label": str(payload.get("group_label") or payload.get("group") or "自定义变量").strip(),
        "required": _parse_schema_variable_boolean(payload.get("required")),
        "usage": str(payload.get("usage") or payload.get("applies_to") or "").strip(),
        "source_field": str(payload.get("source_field") or "").strip(),
        "selection_rule": str(payload.get("selection_rule") or "").strip(),
        "options": options,
    }


def _schema_prefers_name_key(data: dict) -> bool:
    key_count = 0
    name_count = 0
    for group in (data.get("groups") or {}).values():
        for item in group.get("variables") or []:
            key_count += int("key" in item)
            name_count += int("name" in item)
    return name_count > key_count


def _ensure_schema_variable_group(data: dict, group_key: str, group_label: str) -> dict:
    groups = data.setdefault("groups", {})
    group = groups.setdefault(group_key, {"label": group_label or group_key, "variables": []})
    group.setdefault("variables", [])
    if group_label and not (group.get("label") or group.get("description")):
        group["label"] = group_label
    return group


def _apply_schema_variable_update(item: dict, normalized: dict, *, key_field: str) -> dict:
    item[key_field] = normalized["key"]
    if key_field == "key" and "name" in item:
        item.pop("name", None)
    if key_field == "name" and "key" in item:
        item.pop("key", None)
    if "ui_entry" in item and "label" not in item:
        item["ui_entry"] = normalized["label"]
    else:
        item["label"] = normalized["label"]
    item["type"] = normalized["type"]
    item["default"] = normalized["default"]
    item["required"] = normalized["required"]
    if "applies_to" in item and "usage" not in item:
        item["applies_to"] = normalized["usage"]
    elif normalized["usage"] or "usage" in item:
        item["usage"] = normalized["usage"]
    if normalized["source_field"] or "source_field" in item:
        item["source_field"] = normalized["source_field"]
    if normalized["selection_rule"] or "selection_rule" in item:
        item["selection_rule"] = normalized["selection_rule"]
    if normalized["options"] is not _VARIABLE_OPTION_MISSING:
        if normalized["options"] is None:
            item.pop("options", None)
        else:
            item["options"] = normalized["options"]
    return item


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/favicon.ico")
def favicon():
    return Response(status=204)


@app.route("/api/schema")
def api_schema():
    with open(SCHEMA_JSON, encoding="utf-8") as f:
        return jsonify(json.load(f))


@app.get("/api/variables")
def api_variables():
    data = _load_schema_json()
    items = _flatten_schema_variables(data)
    return jsonify(
        {
            "source_template": str(SCHEMA_JSON),
            "count": len(items),
            "condition_count": sum(1 for item in items if item.get("is_condition")),
            "group_options": _schema_variable_group_options(data),
            "items": items,
        }
    )


@app.post("/api/variables")
def api_create_variable():
    try:
        payload = request.get_json(silent=True) or {}
        normalized = _normalize_schema_variable_payload(payload)
        data = _load_schema_json()
        if _find_schema_variable(data, normalized["key"])[3] is not None:
            return _schema_json_error("变量名已存在")
        group = _ensure_schema_variable_group(data, normalized["group"], normalized["group_label"])
        key_field = "name" if _schema_prefers_name_key(data) else "key"
        item = _apply_schema_variable_update({}, normalized, key_field=key_field)
        group["variables"].append(item)
        _save_schema_json(data)
        return jsonify({"ok": True, "success": True, "id": normalized["key"], "item": _schema_variable_item_payload(item, normalized["group"], group)})
    except ValueError as exc:
        return _schema_json_error(str(exc))


@app.put("/api/variables/<path:key>")
def api_update_variable(key: str):
    try:
        payload = request.get_json(silent=True) or {}
        normalized = _normalize_schema_variable_payload(payload, existing_key=key)
        data = _load_schema_json()
        old_group_key, old_group, index, item = _find_schema_variable(data, key)
        if item is None:
            return _schema_json_error("未找到该变量", 404)
        _duplicate_group, _duplicate_group_obj, _duplicate_index, duplicate_item = _find_schema_variable(data, normalized["key"])
        if normalized["key"] != key and duplicate_item is not None:
            return _schema_json_error("新变量名已存在")
        key_field = "name" if "name" in item and "key" not in item else "key"
        updated = _apply_schema_variable_update(dict(item), normalized, key_field=key_field)
        target_group = _ensure_schema_variable_group(data, normalized["group"], normalized["group_label"])
        if normalized["group"] != old_group_key:
            old_group["variables"].pop(index)
            target_group["variables"].append(updated)
        else:
            old_group["variables"][index] = updated
        _save_schema_json(data)
        return jsonify({"ok": True, "success": True, "id": normalized["key"], "item": _schema_variable_item_payload(updated, normalized["group"], target_group)})
    except ValueError as exc:
        return _schema_json_error(str(exc))


@app.delete("/api/variables/<path:key>")
def api_delete_variable(key: str):
    data = _load_schema_json()
    group_key, group, index, item = _find_schema_variable(data, key)
    if item is None:
        return _schema_json_error("未找到该变量", 404)
    removed = group["variables"].pop(index)
    _save_schema_json(data)
    return jsonify({"ok": True, "success": True, "id": key, "item": _schema_variable_item_payload(removed, group_key, group)})


@app.route("/api/custodians")
def api_custodians():
    """返回已知托管人的自动填充数据"""
    with open(CLAUSES_JSON, encoding="utf-8") as f:
        data = json.load(f)
    custodians = data["clauses"]["CUSTODIAN_INFO_KNOWN_VALUES"]["custodians"]
    return jsonify(custodians)


@app.route("/api/diff_table")
def api_diff_table():
    """返回差异条款匹配表（Markdown原文）"""
    content = DIFF_TABLE_MD.read_text(encoding="utf-8")
    return jsonify({"content": content})


@app.route("/api/extract_text", methods=["POST"])
def api_extract_text():
    """从上传的 .docx / .txt / .md 文件提取纯文本，用于合同对比。"""
    if "file" not in request.files:
        return jsonify({"error": "未收到文件"}), 400
    f = request.files["file"]
    name = f.filename or ""
    try:
        text = _read_uploaded_text(f, name)
    except Exception as e:
        logger.exception("Failed to read uploaded file")
        return jsonify({"error": f"无法读取文件：{e}"}), 400

    return jsonify({"filename": name, "text": text,
                    "lines": len(text.splitlines()),
                    "size":  len(text)})


@app.route("/api/extract_custody_agreement_summary", methods=["POST"])
def api_extract_custody_agreement_summary():
    if "file" not in request.files:
        return jsonify({"success": False, "error": "未收到托管协议文件"}), 400
    f = request.files["file"]
    name = f.filename or ""
    try:
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        if ext == "docx":
            result = _extract_custody_agreement_summary_from_docx(f, name)
        else:
            text = _read_uploaded_text(f, name)
            result = _extract_custody_agreement_summary_from_text(text)
    except Exception as e:
        logger.exception("Failed to parse custody agreement")
        return jsonify({"success": False, "error": f"无法解析托管协议：{e}"}), 400

    if not result["summary_text"]:
        return jsonify({
            "success": False,
            "error": "未匹配到托管协议摘要章节，请确认文件包含托管协议正文或招募说明书的托管协议摘要章",
            "filename": name,
            "matched_sections": result["matched_sections"],
            "missing_sections": result["missing_sections"],
        }), 400

    return jsonify({
        "success": True,
        "filename": name,
        "summary_text": result["summary_text"],
        "matched_sections": result["matched_sections"],
        "missing_sections": result["missing_sections"],
        "source_type": result["source_type"],
        "size": len(result["summary_text"]),
    })


@app.route("/api/compare", methods=["POST"])
def api_compare():
    """
    合同对比分析：接收两段合同文本，返回逐行差异（含字符级高亮HTML）。
    响应格式：
      summary: { total_a, total_b, changed_lines, similarity, diff_groups }
      hunks:   每个 hunk 为 { type: "lines"|"skip", lines?, count?, preview? }
               每行: { tag, a, b, a_html, b_html, num_a, num_b }
    """
    import difflib
    import html as html_mod

    d = request.get_json(force=True)
    text1  = d.get("text1",  "")
    text2  = d.get("text2",  "")
    label1 = d.get("label1", "合同A")
    label2 = d.get("label2", "合同B")
    CTX    = 3   # 每个差异块上下各保留 3 行上下文

    def prepare_compare_lines(text: str):
        items = []
        raw_lines = text.splitlines()
        for line_no, raw in enumerate(raw_lines, start=1):
            if _is_layout_blank_line(raw) or _is_markdown_table_separator_line(raw):
                continue
            stripped = _normalize_layout_text(raw).strip()
            if _is_layout_blank_line(stripped) or _is_markdown_table_separator_line(stripped):
                continue
            # Ignore layout-only whitespace differences while keeping original line numbers.
            key = re.sub(r"\s+", " ", stripped)
            items.append({"text": stripped, "key": key, "num": line_no})
        return raw_lines, items

    raw_lines1, items1 = prepare_compare_lines(text1)
    raw_lines2, items2 = prepare_compare_lines(text2)
    lines1 = [item["text"] for item in items1]
    lines2 = [item["text"] for item in items2]
    keys1 = [item["key"] for item in items1]
    keys2 = [item["key"] for item in items2]

    if not lines1 or not lines2:
        return jsonify({"error": "两份合同文本去除空行后均不能为空"}), 400

    # 字符级差异 HTML（对单行 a, b 比较）
    def char_diff(a: str, b: str):
        m = difflib.SequenceMatcher(None, a, b, autojunk=False)
        ra, rb = [], []
        for tag, i1, i2, j1, j2 in m.get_opcodes():
            ea = html_mod.escape(a[i1:i2])
            eb = html_mod.escape(b[j1:j2])
            if tag == "equal":
                ra.append(ea); rb.append(eb)
            elif tag == "replace":
                ra.append(f'<del>{ea}</del>'); rb.append(f'<ins>{eb}</ins>')
            elif tag == "delete":
                ra.append(f'<del>{ea}</del>')
            elif tag == "insert":
                rb.append(f'<ins>{eb}</ins>')
        return "".join(ra), "".join(rb)

    # 逐行对比，构建 result_lines 列表
    sm = difflib.SequenceMatcher(None, keys1, keys2, autojunk=False)
    result_lines = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for k in range(i2 - i1):
                t = html_mod.escape(lines1[i1 + k])
                result_lines.append({"tag": "equal",
                                     "a": t, "b": t,
                                     "a_html": t, "b_html": t,
                                     "num_a": items1[i1+k]["num"], "num_b": items2[j1+k]["num"]})
        elif tag == "replace":
            len_a, len_b = i2 - i1, j2 - j1
            for k in range(max(len_a, len_b)):
                la = lines1[i1+k] if k < len_a else None
                lb = lines2[j1+k] if k < len_b else None
                na = items1[i1+k]["num"] if la is not None else None
                nb = items2[j1+k]["num"] if lb is not None else None
                if la is not None and lb is not None:
                    ah, bh = char_diff(la, lb)
                    result_lines.append({"tag": "replace",
                                         "a": html_mod.escape(la), "b": html_mod.escape(lb),
                                         "a_html": ah, "b_html": bh,
                                         "num_a": na, "num_b": nb})
                elif la is not None:
                    ea = html_mod.escape(la)
                    result_lines.append({"tag": "delete",
                                         "a": ea, "b": "",
                                         "a_html": f"<del>{ea}</del>", "b_html": "",
                                         "num_a": na, "num_b": None})
                else:
                    eb = html_mod.escape(lb)
                    result_lines.append({"tag": "insert",
                                         "a": "", "b": eb,
                                         "a_html": "", "b_html": f"<ins>{eb}</ins>",
                                         "num_a": None, "num_b": nb})
        elif tag == "delete":
            for k in range(i2 - i1):
                ea = html_mod.escape(lines1[i1+k])
                result_lines.append({"tag": "delete",
                                     "a": ea, "b": "",
                                     "a_html": f"<del>{ea}</del>", "b_html": "",
                                     "num_a": items1[i1+k]["num"], "num_b": None})
        elif tag == "insert":
            for k in range(j2 - j1):
                eb = html_mod.escape(lines2[j1+k])
                result_lines.append({"tag": "insert",
                                     "a": "", "b": eb,
                                     "a_html": "", "b_html": f"<ins>{eb}</ins>",
                                     "num_a": None, "num_b": items2[j1+k]["num"]})

    # 压缩连续的 equal 区域（超过 CTX*2+1 行则折叠中间部分）
    hunks = []
    idx, n = 0, len(result_lines)
    while idx < n:
        r = result_lines[idx]
        if r["tag"] == "equal":
            j = idx
            while j < n and result_lines[j]["tag"] == "equal":
                j += 1
            count = j - idx
            if count > CTX * 2 + 1:
                hunks.append({"type": "lines", "lines": result_lines[idx: idx+CTX]})
                hunks.append({"type": "skip",
                              "count": count - CTX*2,
                              "preview": result_lines[idx+CTX]["a"][:60]})
                hunks.append({"type": "lines", "lines": result_lines[j-CTX: j]})
            else:
                hunks.append({"type": "lines", "lines": result_lines[idx:j]})
            idx = j
        else:
            j = idx
            while j < n and result_lines[j]["tag"] != "equal":
                j += 1
            hunks.append({"type": "lines", "lines": result_lines[idx:j], "is_diff": True})
            idx = j

    # 统计
    changed = sum(1 for r in result_lines if r["tag"] != "equal")
    total   = len(result_lines)
    sim     = round((total - changed) / total * 100, 1) if total > 0 else 100.0
    diff_groups = sum(1 for h in hunks if h.get("is_diff"))

    return jsonify({
        "label1": label1, "label2": label2,
        "summary": {
            "total_a": len(lines1), "total_b": len(lines2),
            "raw_total_a": len(raw_lines1), "raw_total_b": len(raw_lines2),
            "ignored_blank_a": len(raw_lines1) - len(lines1),
            "ignored_blank_b": len(raw_lines2) - len(lines2),
            "changed_lines": changed, "similarity": sim,
            "diff_groups": diff_groups
        },
        "hunks": hunks
    })


@app.route("/api/generate", methods=["POST"])
def api_generate():
    """接收表单数据，返回合同全文"""
    form_data = request.get_json(force=True)
    try:
        contract_text = engine.generate(form_data)
        quality = _generation_quality_check(contract_text, "基金合同")
        return jsonify({"success": True, "text": contract_text, "quality_check": quality})
    except Exception as e:
        logger.exception("Generate contract failed")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export", methods=["POST"])
def api_export():
    """返回 .txt 文件下载"""
    form_data = request.get_json(force=True)
    try:
        contract_text = engine.generate(form_data)
        quality_response = _quality_error_response(contract_text, "基金合同")
        if quality_response:
            return quality_response
        fund_name = form_data.get("FUND_NAME", "ETF联接基金合同")
        # Write to a temp file
        import tempfile
        tmp = tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", encoding="utf-8",
            delete=False, prefix="contract_"
        )
        tmp.write(contract_text)
        tmp.close()
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        tmp_path = tmp.name
        @after_this_request
        def _cleanup(resp):
            _safe_unlink(tmp_path)
            return resp
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}基金合同.txt",
            mimetype="text/plain",
        )
    except Exception as e:
        logger.exception("Export contract txt failed")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export_docx", methods=["POST"])
def api_export_docx():
    """返回格式化 .docx 文件下载"""
    form_data = request.get_json(force=True)
    try:
        contract_text = engine.generate(form_data)
        quality_response = _quality_error_response(contract_text, "基金合同")
        if quality_response:
            return quality_response
        docx_bytes = engine.build_docx(contract_text)
        fund_name = form_data.get("FUND_NAME", "ETF联接基金合同")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile, os as _os
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="contract_")
        tmp.write(docx_bytes)
        tmp.close()
        tmp_path = tmp.name
        @after_this_request
        def _cleanup(resp):
            _safe_unlink(tmp_path)
            return resp
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}基金合同.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        logger.exception("Export contract docx failed")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/files")
def api_files():
    """列出所有可编辑文件"""
    files = []
    for name in sorted(EDITABLE_FILE_NAMES):
        p = BASE_DIR / name
        if not p.is_file():
            continue
        stat = p.stat()
        files.append({
            "name": p.name,
            "group": _editor_group_for_filename(p.name),
            "size": stat.st_size,
            "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
        })
    return jsonify(files)


@app.route("/api/files/<filename>", methods=["GET", "POST"])
def api_file(filename):
    """读取或保存指定文件（防路径遍历）"""
    # Security: forbid path separators in filename
    if os.sep in filename or "/" in filename or "\\" in filename:
        logger.warning("Path traversal blocked (separator in filename): %s", filename)
        abort(403)
    target = (BASE_DIR / filename).resolve()
    # Security: only allow direct children of BASE_DIR with allowed suffixes
    if (
        target.parent.resolve() != BASE_DIR.resolve()
        or target.suffix not in ALLOWED_SUFFIXES
        or target.name not in EDITABLE_FILE_NAMES
    ):
        logger.warning("Blocked non-editable file access: %s -> %s", filename, target)
        abort(403)
    if not target.is_file():
        abort(404)

    if request.method == "GET":
        content = target.read_text(encoding="utf-8")
        return jsonify({"name": filename, "content": content})

    # POST: save
    data = request.get_json(silent=True)
    if not isinstance(data, dict) or not isinstance(data.get("content"), str):
        return jsonify({"success": False, "error": "请求体必须是包含 content 字符串字段的 JSON 对象"}), 400
    content = data["content"]
    if target.suffix == ".json":
        try:
            json.loads(content)
        except json.JSONDecodeError as exc:
            return jsonify({"success": False, "error": f"JSON 格式错误：第 {exc.lineno} 行第 {exc.colno} 列"}), 400
    backup_dir = BASE_DIR / "backups"
    backup_dir.mkdir(exist_ok=True)
    backup_name = f"{target.name}.bak_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    backup_path = backup_dir / backup_name
    backup_path.write_text(target.read_text(encoding="utf-8"), encoding="utf-8")
    tmp_path = target.with_name(f".{target.name}.tmp")
    try:
        tmp_path.write_text(content, encoding="utf-8")
        os.replace(tmp_path, target)
    finally:
        if tmp_path.exists():
            _safe_unlink(tmp_path)
    # Reload engine after file changes (picks up fresh overrides and templates)
    try:
        _reload_generation_engines()
    except Exception as exc:
        target.write_text(backup_path.read_text(encoding="utf-8"), encoding="utf-8")
        logger.exception("Reload after editable file save failed")
        return jsonify({"success": False, "error": f"保存后重载失败，已回滚：{exc}"}), 400
    return jsonify({"success": True, "message": f"已保存 {filename}"})


@app.route("/api/clause_library")
def api_clause_library():
    """返回面向业务维护的条款库扁平条目。"""
    return jsonify(_build_clause_library_catalog())


@app.route("/api/clause_library", methods=["POST"])
def api_clause_library_save():
    """保存单个条款正文字段，不接受整文件覆盖。"""
    data = request.get_json(silent=True)
    if not isinstance(data, dict):
        return jsonify({"success": False, "error": "请求体必须是 JSON 对象"}), 400
    try:
        result = _save_clause_library_field(
            str(data.get("path_id") or ""),
            str(data.get("content") or ""),
        )
    except FileNotFoundError as exc:
        return jsonify({"success": False, "error": str(exc)}), 404
    except (ValueError, json.JSONDecodeError) as exc:
        return jsonify({"success": False, "error": str(exc)}), 400
    except Exception as exc:
        logger.exception("Save clause library field failed")
        return jsonify({"success": False, "error": f"保存条款正文失败：{exc}"}), 500
    return jsonify(result)


@app.route("/api/preview_clause", methods=["POST"])
def api_preview_clause():
    """根据当前表单值，实时返回关键差异条款预览文本"""
    form_data = request.get_json(force=True)
    try:
        v = engine._derive_variables(form_data)
        v = engine._inject_clause_texts(v)
        preview = {
            "MGMT_FEE_PAYMENT_METHOD": v.get("MGMT_FEE_PAYMENT_METHOD", ""),
            "CUSTODY_FEE_PAYMENT_METHOD": v.get("CUSTODY_FEE_PAYMENT_METHOD", ""),
            "DISPUTE_RESOLUTION_LABEL": v.get("DISPUTE_RESOLUTION_LABEL", ""),
            "DISPUTE_RESOLUTION_CLAUSE": v.get("DISPUTE_RESOLUTION_CLAUSE", ""),
            "HAS_HK_CONNECT": v.get("HAS_HK_CONNECT", False),
            "IS_INITIATING_FUND": v.get("IS_INITIATING_FUND", False),
            "FUND_PROFIT_DEF": v.get("FUND_PROFIT_DEF", ""),
            "DISTRIBUTION_FREQ_CLAUSE": v.get("DISTRIBUTION_FREQ_CLAUSE", ""),
            "SALES_SERVICE_FEE_RATE": v.get("SALES_SERVICE_FEE_RATE", "0.20"),
        }
        return jsonify(preview)
    except Exception as e:
        logger.exception("Preview clause failed")
        return jsonify({"error": str(e)}), 400


@app.route("/api/generate_prospectus", methods=["POST"])
def api_generate_prospectus():
    form_data = request.get_json(force=True)
    try:
        text = prospectus_engine.generate(form_data)
        quality = _generation_quality_check(text, "招募说明书")
        return jsonify({"success": True, "text": text, "quality_check": quality})
    except Exception as e:
        logger.exception("Generate prospectus failed")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export_prospectus_docx", methods=["POST"])
def api_export_prospectus_docx():
    form_data = request.get_json(force=True)
    try:
        text = prospectus_engine.generate(form_data)
        quality_response = _quality_error_response(text, "招募说明书")
        if quality_response:
            return quality_response
        docx_bytes = prospectus_engine.build_docx(text, form_data)
        fund_name = form_data.get("FUND_NAME", "ETF联接基金招募说明书")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="prospectus_")
        tmp.write(docx_bytes)
        tmp.close()
        tmp_path = tmp.name
        @after_this_request
        def _cleanup(resp):
            _safe_unlink(tmp_path)
            return resp
        return send_file(
            tmp.name, as_attachment=True,
            download_name=f"{safe_name}招募说明书.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        logger.exception("Export prospectus docx failed")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export_prospectus_txt", methods=["POST"])
def api_export_prospectus_txt():
    form_data = request.get_json(force=True)
    try:
        text = prospectus_engine.generate(form_data)
        quality_response = _quality_error_response(text, "招募说明书")
        if quality_response:
            return quality_response
        fund_name = form_data.get("FUND_NAME", "ETF联接基金招募说明书")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".txt", delete=False, prefix="prospectus_")
        tmp.write(text.encode("utf-8"))
        tmp.close()
        tmp_path = tmp.name
        @after_this_request
        def _cleanup(resp):
            _safe_unlink(tmp_path)
            return resp
        return send_file(
            tmp.name, as_attachment=True,
            download_name=f"{safe_name}招募说明书.txt",
            mimetype="text/plain",
        )
    except Exception as e:
        logger.exception("Export prospectus txt failed")
        return jsonify({"success": False, "error": str(e)}), 400


# ── 产品资料概要生成 API ────────────────────────────────────────────────────
@app.route("/api/generate_product_summary", methods=["POST"])
def api_generate_product_summary():
    form_data = request.get_json(force=True)
    try:
        bundle = product_summary_engine.generate_bundle(form_data)
        quality = _generation_quality_check(bundle["text"], "产品资料概要")
        return jsonify({"success": True, **bundle, "quality_check": quality})
    except Exception as e:
        logger.exception("Generate product summary failed")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export_product_summary_docx", methods=["POST"])
def api_export_product_summary_docx():
    form_data = request.get_json(force=True)
    try:
        bundle = product_summary_engine.generate_bundle(form_data)
        quality_response = _quality_error_response(bundle["text"], "产品资料概要")
        if quality_response:
            return quality_response
        zip_bytes = product_summary_engine.build_docx_bundle(bundle["render_models"])
        fund_name = form_data.get("FUND_NAME", "ETF联接基金产品资料概要")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".zip", delete=False, prefix="product_summary_")
        tmp.write(zip_bytes)
        tmp.close()
        tmp_path = tmp.name
        @after_this_request
        def _cleanup(resp):
            _safe_unlink(tmp_path)
            return resp
        return send_file(
            tmp.name, as_attachment=True,
            download_name=f"{safe_name}产品资料概要.zip",
            mimetype="application/zip",
        )
    except Exception as e:
        logger.exception("Export product summary docx bundle failed")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export_product_summary_txt", methods=["POST"])
def api_export_product_summary_txt():
    form_data = request.get_json(force=True)
    try:
        text = product_summary_engine.generate(form_data)
        quality_response = _quality_error_response(text, "产品资料概要")
        if quality_response:
            return quality_response
        fund_name = form_data.get("FUND_NAME", "ETF联接基金产品资料概要")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".txt", delete=False, prefix="product_summary_")
        tmp.write(text.encode("utf-8"))
        tmp.close()
        tmp_path = tmp.name
        @after_this_request
        def _cleanup(resp):
            _safe_unlink(tmp_path)
            return resp
        return send_file(
            tmp.name, as_attachment=True,
            download_name=f"{safe_name}产品资料概要.txt",
            mimetype="text/plain",
        )
    except Exception as e:
        logger.exception("Export product summary txt failed")
        return jsonify({"success": False, "error": str(e)}), 400


# ═══════════════════════════════════════════════════════════════════════════════
#  复核系统 (Review System)
# ═══════════════════════════════════════════════════════════════════════════════

# ── 复核常量 ──────────────────────────────────────────────────────────────────
SUMMARY_PROBLEM_SIMILARITY_THRESHOLD = 95.0
SUMMARY_WARNING_SIMILARITY_THRESHOLD = 80.0
REVIEW_SOFT_HEADING_MATCH_MIN_KEY_LENGTH = 16
REVIEW_SOFT_HEADING_MATCH_RATIO = 0.88
_REVIEW_NUM_PREFIX_RE = re.compile(r'^[\d]+[\.\、\)）]\s*')
_REVIEW_CHAPTER_PREFIX_RE = re.compile(
    r'^(第([一二三四五六七八九十百千\d]+)(?:部分|章|节|条))\s*(.*)'
)


# ── 中文数字转换 ──────────────────────────────────────────────────────────────
def _cn_numeral_to_int(cn: str):
    if not cn:
        return None
    if cn.isdigit():
        return int(cn)
    digits = {"零": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
              "六": 6, "七": 7, "八": 8, "九": 9}
    units = {"十": 10, "百": 100, "千": 1000}
    total = 0
    current = 0
    for ch in cn:
        if ch in digits:
            current = digits[ch]
            continue
        if ch in units:
            if current == 0:
                current = 1
            total += current * units[ch]
            current = 0
            continue
        return None
    return total + current if current else (total if total > 0 else None)


# ── 文本规范化 ────────────────────────────────────────────────────────────────
def _normalize_review_text(text):
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", str(text))
    text = re.sub(r'\s+', '', text)
    return text


def _normalize_review_locator(locator):
    if not locator:
        return ""
    return unicodedata.normalize("NFKC", str(locator)).strip()


def _normalize_contract_prospectus_compare_text(text):
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", str(text))
    text = re.sub(r'\s+', '', text)
    text = re.sub(r'[（(][^）)]*[）)]', '', text)
    return text


def _normalize_review_doc_title_key(text):
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", str(text))
    text = re.sub(r'[\s的与及和、，,：:；;/-]', '', text)
    return text


# ── 标题匹配 ──────────────────────────────────────────────────────────────────
def _review_heading_key(value: str) -> str:
    text = unicodedata.normalize("NFKC", str(value or ""))
    text = text.replace("基金合同的基本情况", "基金的基本情况")
    text = text.replace("基金合同基本情况", "基金基本情况")
    text = re.sub(r'\([^)]*\)', '', text)
    text = re.sub(r'（[^）]*）', '', text)
    return re.sub(r'[\s的与及和、，,：:；;/-]', '', text)


def _review_label_variants(chapter_name):
    if not chapter_name:
        return []
    name = str(chapter_name).strip()
    variants = [name]
    m = _REVIEW_CHAPTER_PREFIX_RE.match(name)
    if m:
        ordinal_str = m.group(2)
        ordinal_num = _cn_numeral_to_int(ordinal_str)
        title = m.group(3).strip()
        if title:
            variants.append(title)
        if ordinal_num is not None:
            prefix = str(ordinal_num)
            variants.append(f"{prefix}、{title}" if title else prefix)
            variants.append(f"{prefix}.{title}" if title else prefix)
    return [v for v in variants if v]


def _review_label_ordinal(label):
    if not label:
        return None
    label = str(label).strip()
    m = _REVIEW_CHAPTER_PREFIX_RE.match(label)
    if m:
        return _cn_numeral_to_int(m.group(2))
    m2 = re.match(r'^(\d+)[\.\、]', label)
    if m2:
        return int(m2.group(1))
    return None


def _review_labels_match(a, b):
    if not a or not b:
        return False
    a_ord = _review_label_ordinal(a)
    b_ord = _review_label_ordinal(b)
    if a_ord is not None and b_ord is not None and a_ord != b_ord:
        return False
    key_a = _review_heading_key(a)
    key_b = _review_heading_key(b)
    if len(key_a) >= 4 and len(key_b) >= 4 and (key_a in key_b or key_b in key_a):
        return True
    return key_a == key_b


def _review_soft_heading_key(value):
    if not value:
        return ""
    text = unicodedata.normalize("NFKC", str(value))
    text = re.sub(r'[\s、，,：:；;/]', '', text)
    return text


def _review_soft_heading_match(a, b):
    if not a or not b:
        return False
    key_a = _review_soft_heading_key(a)
    key_b = _review_soft_heading_key(b)
    if len(key_a) < REVIEW_SOFT_HEADING_MATCH_MIN_KEY_LENGTH or \
       len(key_b) < REVIEW_SOFT_HEADING_MATCH_MIN_KEY_LENGTH:
        return False
    if key_a == key_b:
        return True
    ratio = difflib.SequenceMatcher(None, key_a, key_b, autojunk=False).ratio()
    return ratio >= REVIEW_SOFT_HEADING_MATCH_RATIO


def _match_review_section(sections, chapter_name):
    candidates = _review_label_variants(chapter_name)
    if not candidates:
        return {"section": None, "match_method": "", "score": 0.0}
    best = {"section": None, "match_method": "", "score": 0.0}
    method_rank = {"exact": 4, "normalized": 3, "chapter_title": 2,
                   "soft_heading": 1, "": 0}

    def update(section, match_method, score):
        current_rank = method_rank.get(best["match_method"], 0)
        new_rank = method_rank.get(match_method, 0)
        if new_rank > current_rank or (new_rank == current_rank and score > best["score"]):
            best["section"] = section
            best["match_method"] = match_method
            best["score"] = score

    for sec in sections:
        h = sec.get("heading", "")
        for candidate in candidates:
            if h.strip() == candidate:
                update(sec, "exact", 1.0)
                continue
            if _review_heading_key(h) == _review_heading_key(candidate):
                update(sec, "normalized", 0.95)
                continue
            mm = _REVIEW_CHAPTER_PREFIX_RE.match(candidate)
            if mm and mm.group(3).strip():
                title = mm.group(3).strip()
                if title in h:
                    update(sec, "chapter_title", 0.85)
                    continue
            if _review_soft_heading_match(h, candidate):
                score = difflib.SequenceMatcher(
                    None, _review_soft_heading_key(h),
                    _review_soft_heading_key(candidate), autojunk=False).ratio()
                update(sec, "soft_heading", score)
    return best


def _find_review_section(sections, heading):
    result = _match_review_section(sections, heading)
    return result["section"]


# ── 块操作 ────────────────────────────────────────────────────────────────────
def _split_review_blocks(text):
    if not text:
        return []
    RE_HEADING = re.compile(
        r'^((?:第[一二三四五六七八九十百\d]+(?:部分|章|节|条))'
        r'|(?:[一二三四五六七八九十]+)[、．.]'
        r'|(?:（[一二三四五六七八九十\d]+）)'
        r'|(?:\d+[\.\、](?!\d))'
        r'|(?:（\d+）))'
        r'\s*\S', re.MULTILINE)
    RE_HEADING = re.compile(
        r'^((?:第[一二三四五六七八九十百\d]+(?:部分|章|节|条)'
        r'|[一二三四五六七八九十]+[、．.]'
        r'|（[一二三四五六七八九十\d]+）'
        r'|\d+[\.、](?!\d)'
        r'|（\d+）)\s*\S[^\n]*)',
        re.MULTILINE)
    blocks = []
    matches = list(RE_HEADING.finditer(text))
    for i, m in enumerate(matches):
        heading = m.group(1).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        blocks.append({
            "heading": heading,
            "body": body,
            "text": text[m.start():end].strip(),
            "start": m.start(),
        })
    if not blocks and text.strip():
        blocks.append({"heading": "", "body": text.strip(), "start": 0})
    return blocks


def _split_summary_top_blocks(text):
    if not text:
        return []
    re_heading = re.compile(r'^([一二三四五六七八九十]+[、．.]\s*\S[^\n]*)', re.MULTILINE)
    blocks = []
    matches = list(re_heading.finditer(text))
    for i, match in enumerate(matches):
        heading = match.group(1).strip()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        blocks.append({
            "heading": heading,
            "body": text[start:end].strip(),
            "start": match.start(),
        })
    if not blocks:
        return _split_review_blocks(text)
    return blocks


def _review_block_heading_style(line):
    if not line:
        return None
    line = line.strip()
    if re.match(r'^第[一二三四五六七八九十百\d]+部分', line):
        return "part"
    if re.match(r'^[一二三四五六七八九十]+[、．.]', line):
        return "cn_numbered"
    if re.match(r'^（\d+）', line):
        return "digit_paren"
    if re.match(r'^（[一二三四五六七八九十]+）', line):
        return "cn_paren"
    if re.match(r'^\d+[\.\、]', line):
        return "digit_numbered"
    return None


def _expand_review_block_with_descendants(blocks, start_index):
    if start_index < 0 or start_index >= len(blocks):
        return {"text": "", "body": ""}

    block = blocks[start_index]
    base_style = _review_block_heading_style(block.get("heading"))
    descendant_styles = {
        "cn_numbered": {"cn_paren", "digit_numbered", "digit_paren"},
        "cn_paren": {"digit_numbered", "digit_paren"},
        "digit_numbered": {"digit_paren"},
    }.get(base_style, set())
    if not descendant_styles:
        return {
            "text": block.get("text") or "",
            "body": block.get("body") or "",
        }

    text_parts = [block.get("text") or ""]
    body_parts = [part for part in [block.get("body") or ""] if str(part).strip()]
    for idx in range(start_index + 1, len(blocks)):
        current = blocks[idx]
        current_style = _review_block_heading_style(current.get("heading"))
        if current_style == base_style:
            break
        if current_style not in descendant_styles:
            break
        current_text = current.get("text") or ""
        text_parts.append(current_text)
        body_parts.append(current_text)

    return {
        "text": "\n".join(part for part in text_parts if str(part).strip()).strip(),
        "body": "\n".join(part for part in body_parts if str(part).strip()).strip(),
    }


def _match_review_block_by_label(section, label):
    if not section or not label:
        return None
    if isinstance(section, str):
        content = section
    else:
        content = section.get("content", "") or section.get("body", "")
    blocks = _split_review_blocks(content)
    label_key = _review_heading_key(label)
    for block in blocks:
        if _review_heading_key(block["heading"]) == label_key:
            return block
    for block in blocks:
        if _review_soft_heading_match(block["heading"], label):
            return block
    return None


def _find_review_block_by_label(section, label):
    return _match_review_block_by_label(section, label)


def _locate_review_subheading(section, subheading):
    if not section or not subheading:
        return None
    content = section.get("content", "") or section.get("body", "")
    blocks = _split_review_blocks(content)
    for block in blocks:
        h = block.get("heading", "")
        if subheading.strip() in h or _review_heading_key(h) == _review_heading_key(subheading):
            return block
    return None


def _split_review_locator(locator):
    if not locator:
        return []
    return [p.strip() for p in locator.split("/") if p.strip()]


def _format_review_locator_path(parts):
    if not parts:
        return ""
    return " → ".join(parts)


def _format_review_target_locator(target_info):
    if not target_info:
        return ""
    parts = []
    if target_info.get("section_heading"):
        parts.append(target_info["section_heading"])
    if target_info.get("target_heading"):
        parts.append(target_info["target_heading"])
    return " → ".join(parts) if parts else ""


def _locate_review_rule_target(sections, locator: str, section_hint: str = ""):
    result = {
        "section": None, "section_heading": "", "target_heading": "",
        "text": "", "body_text": "", "matched": False, "match_method": "",
        "missing_reason": "", "locator_parts": [],
    }
    if not locator or _is_missing_review_locator(locator):
        result["missing_reason"] = "无定位器或标记为无对应"
        return result

    parts = _split_review_locator(locator)
    result["locator_parts"] = parts
    if not parts:
        result["missing_reason"] = "定位器为空"
        return result

    chapter_name = parts[0]
    section_match = _match_review_section(sections, chapter_name)
    section = section_match["section"]

    if not section and section_hint:
        section_match = _match_review_section(sections, section_hint)
        section = section_match["section"]

    if not section:
        result["missing_reason"] = f'未找到章节: {chapter_name}'
        return result

    result["section"] = section
    result["section_heading"] = section.get("heading", "")
    result["matched"] = True
    result["match_method"] = section_match["match_method"]

    content = section.get("content", "") or section.get("body", "")
    result["body_text"] = content

    if len(parts) == 1:
        result["text"] = content
        result["target_heading"] = result["section_heading"]
    else:
        sub_locator = parts[1]
        blocks = _split_review_blocks(content)
        block = _find_review_block_by_label(section, sub_locator)
        if block:
            block_index = next((idx for idx, item in enumerate(blocks) if item == block), -1)
            expanded_block = _expand_review_block_with_descendants(blocks, block_index)
            result["target_heading"] = block["heading"]
            result["text"] = (expanded_block.get("text") or block.get("text") or block.get("body") or "").strip()
            result["body_text"] = (expanded_block.get("body") or block.get("body") or result["text"]).strip()
        else:
            escaped = re.escape(sub_locator)
            m = re.search(escaped, content)
            if m:
                result["target_heading"] = sub_locator
                result["text"] = content[m.start():]
                result["body_text"] = result["text"]
            else:
                result["text"] = content

    return result


# ── 章节元数据 ────────────────────────────────────────────────────────────────
def _extract_review_chapter_meta(heading):
    if not heading:
        return {"ordinal": None, "ordinal_num": None, "title": "", "prefix": ""}
    heading = str(heading).strip()
    m = _REVIEW_CHAPTER_PREFIX_RE.match(heading)
    if m:
        return {
            "ordinal": m.group(2),
            "ordinal_num": _cn_numeral_to_int(m.group(2)),
            "title": m.group(3).strip(),
            "prefix": m.group(1),
        }
    return {"ordinal": None, "ordinal_num": None, "title": heading, "prefix": ""}


def _is_missing_review_locator(locator):
    if not locator:
        return True
    return str(locator).startswith("无直接对应") or str(locator).strip() == ""


# ── 比较工具 ──────────────────────────────────────────────────────────────────
def _review_similarity_severity(similarity):
    try:
        v = float(similarity)
    except (TypeError, ValueError):
        v = 0.0
    if v >= SUMMARY_PROBLEM_SIMILARITY_THRESHOLD:
        return "info"
    if v >= SUMMARY_WARNING_SIMILARITY_THRESHOLD:
        return "warning"
    return "error"


def _summary_cross_severity(similarity):
    return _review_similarity_severity(similarity)


def _cross_check_result_is_problem(result):
    if not result:
        return False
    return result.get("consistency_result", "") in ("error", "warning", "fail")


def _summary_cross_result_is_problem(result):
    if not result:
        return False
    return result.get("severity", "") in ("error", "warning")


def _strip_review_number_prefix(text):
    if not text:
        return ""
    return _REVIEW_NUM_PREFIX_RE.sub('', text.strip()).strip()


def _truncate_review_excerpt(text, max_len=200):
    if not text:
        return ""
    if len(text) <= max_len:
        return text
    return text[:max_len] + "…"


def _collect_nonblank_compare_lines(lines, *, normalize_line=None):
    if normalize_line is None:
        return [
            _normalize_layout_text(line).strip()
            for line in (lines or [])
            if not _is_layout_blank_line(line) and not _is_markdown_table_separator_line(line)
        ]

    rows = []
    for lineno, raw_line in enumerate(str(lines or "").splitlines(), 1):
        if _is_layout_blank_line(raw_line) or _is_markdown_table_separator_line(raw_line):
            continue
        normalized = normalize_line(raw_line)
        if normalized is None:
            continue
        normalized = str(normalized)
        if _is_layout_blank_line(normalized) or _is_markdown_table_separator_line(normalized):
            continue
        rows.append({
            "line": raw_line,
            "compare": normalized,
            "lineno": lineno,
        })
    return rows


def _build_review_excerpt_pair(contract_lines, prospectus_lines, context_lines=2):
    contract_lines = contract_lines or []
    prospectus_lines = prospectus_lines or []
    sm = difflib.SequenceMatcher(None, contract_lines, prospectus_lines, autojunk=False)
    diffs = {"contract": [], "prospectus": []}
    window = context_lines * 2
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            span = i2 - i1
            if span > window:
                for k in range(context_lines):
                    if i1 + k < i2:
                        diffs["contract"].append({"type": "equal", "text": contract_lines[i1 + k]})
                        diffs["prospectus"].append({"type": "equal", "text": prospectus_lines[j1 + k]})
                skip_count = span - window
                if skip_count > 0:
                    diffs["contract"].append({"type": "skip", "count": skip_count})
                    diffs["prospectus"].append({"type": "skip", "count": skip_count})
                for k in range(context_lines):
                    idx = i2 - context_lines + k
                    if idx >= i1 + context_lines and idx < i2:
                        diffs["contract"].append({"type": "equal", "text": contract_lines[idx]})
                        diffs["prospectus"].append({"type": "equal", "text": prospectus_lines[j2 - context_lines + k]})
            else:
                for k in range(span):
                    diffs["contract"].append({"type": "equal", "text": contract_lines[i1 + k]})
                    diffs["prospectus"].append({"type": "equal", "text": prospectus_lines[j1 + k]})
        elif tag == "replace":
            for k in range(i2 - i1):
                diffs["contract"].append({"type": "replace", "text": contract_lines[i1 + k]})
            for k in range(j2 - j1):
                diffs["prospectus"].append({"type": "replace", "text": prospectus_lines[j1 + k]})
        elif tag == "delete":
            for k in range(i2 - i1):
                diffs["contract"].append({"type": "delete", "text": contract_lines[i1 + k]})
            for k in range(j2 - j1):
                diffs["prospectus"].append({"type": "insert", "text": ""})
        elif tag == "insert":
            for k in range(i2 - i1):
                diffs["contract"].append({"type": "delete", "text": ""})
            for k in range(j2 - j1):
                diffs["prospectus"].append({"type": "insert", "text": prospectus_lines[j1 + k]})
    return diffs


def _prepare_review_compare_lines(text):
    items = []
    for line_no, raw in enumerate(str(text or "").splitlines(), start=1):
        if _is_layout_blank_line(raw) or _is_markdown_table_separator_line(raw):
            continue
        stripped = _normalize_layout_text(raw).strip()
        if _is_layout_blank_line(stripped) or _is_markdown_table_separator_line(stripped):
            continue
        key = re.sub(r"\s+", " ", stripped)
        items.append({"text": stripped, "key": key, "num": line_no})
    return items


def _review_char_diff_html(left: str, right: str):
    matcher = difflib.SequenceMatcher(None, left or "", right or "", autojunk=False)
    left_html = []
    right_html = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        left_part = html.escape((left or "")[i1:i2], quote=False)
        right_part = html.escape((right or "")[j1:j2], quote=False)
        if tag == "equal":
            left_html.append(left_part)
            right_html.append(right_part)
        elif tag == "replace":
            left_html.append(f"<del>{left_part}</del>")
            right_html.append(f"<ins>{right_part}</ins>")
        elif tag == "delete":
            left_html.append(f"<del>{left_part}</del>")
        elif tag == "insert":
            right_html.append(f"<ins>{right_part}</ins>")
    return "".join(left_html), "".join(right_html)


def _build_review_line_hunks(contract_text, prospectus_text, context_lines=2):
    contract_items = _prepare_review_compare_lines(contract_text)
    prospectus_items = _prepare_review_compare_lines(prospectus_text)
    contract_keys = [item["key"] for item in contract_items]
    prospectus_keys = [item["key"] for item in prospectus_items]
    rows = []

    matcher = difflib.SequenceMatcher(None, contract_keys, prospectus_keys, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for k in range(i2 - i1):
                c_item = contract_items[i1 + k]
                p_item = prospectus_items[j1 + k]
                escaped = html.escape(c_item["text"], quote=False)
                rows.append({
                    "tag": "equal",
                    "contract_ln": c_item["num"],
                    "prospectus_ln": p_item["num"],
                    "contract_text": c_item["text"],
                    "prospectus_text": p_item["text"],
                    "contract_html": escaped,
                    "prospectus_html": escaped,
                })
            continue

        if tag == "replace":
            span = max(i2 - i1, j2 - j1)
            for k in range(span):
                c_item = contract_items[i1 + k] if i1 + k < i2 else None
                p_item = prospectus_items[j1 + k] if j1 + k < j2 else None
                c_text = c_item["text"] if c_item else ""
                p_text = p_item["text"] if p_item else ""
                c_html, p_html = _review_char_diff_html(c_text, p_text)
                rows.append({
                    "tag": "replace",
                    "contract_ln": c_item["num"] if c_item else None,
                    "prospectus_ln": p_item["num"] if p_item else None,
                    "contract_text": c_text,
                    "prospectus_text": p_text,
                    "contract_html": c_html,
                    "prospectus_html": p_html,
                })
            continue

        if tag == "delete":
            for k in range(i1, i2):
                c_item = contract_items[k]
                escaped = html.escape(c_item["text"], quote=False)
                rows.append({
                    "tag": "delete",
                    "contract_ln": c_item["num"],
                    "prospectus_ln": None,
                    "contract_text": c_item["text"],
                    "prospectus_text": "",
                    "contract_html": f"<del>{escaped}</del>",
                    "prospectus_html": "",
                })
            continue

        if tag == "insert":
            for k in range(j1, j2):
                p_item = prospectus_items[k]
                escaped = html.escape(p_item["text"], quote=False)
                rows.append({
                    "tag": "insert",
                    "contract_ln": None,
                    "prospectus_ln": p_item["num"],
                    "contract_text": "",
                    "prospectus_text": p_item["text"],
                    "contract_html": "",
                    "prospectus_html": f"<ins>{escaped}</ins>",
                })

    if not rows:
        return []

    compacted = []
    idx = 0
    while idx < len(rows):
        row = rows[idx]
        if row["tag"] != "equal":
            compacted.append(row)
            idx += 1
            continue

        end = idx
        while end < len(rows) and rows[end]["tag"] == "equal":
            end += 1
        equal_count = end - idx
        if equal_count > context_lines * 2 + 1:
            compacted.extend(rows[idx:idx + context_lines])
            compacted.append({"tag": "skip", "count": equal_count - context_lines * 2})
            compacted.extend(rows[end - context_lines:end])
        else:
            compacted.extend(rows[idx:end])
        idx = end
    return compacted


def _normalize_review_rate_value(value):
    if value is None:
        return ""
    try:
        return str(float(str(value).strip().rstrip("%")))
    except ValueError:
        return re.sub(r"\s+", "", str(value))


def _extract_first_review_rate(text, label_patterns):
    text = str(text or "")
    for label_pattern in label_patterns:
        pattern = rf'{label_pattern}[\s\S]{{0,120}}?(\d+(?:\.\d+)?)\s*%[\s\S]{{0,30}}?(?:年费率|年费|计提)'
        match = re.search(pattern, text)
        if match:
            raw = match.group(1)
            return {"display": f"{raw}%", "compare": _normalize_review_rate_value(raw)}
    return None


def _extract_first_review_text_fact(text, patterns):
    text = str(text or "")
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            value = re.split(r'[；;。,\n]', match.group(1).strip(), maxsplit=1)[0].strip()
            value = re.sub(r'\s+', '', value)
            if value:
                return {"display": value, "compare": _review_heading_key(value)}
    return None


def _extract_review_key_facts(text):
    facts = {}
    rate_specs = [
        ("management_fee_rate", "管理费率", [r"(?:基金管理人的)?管理费"]),
        ("custody_fee_rate", "托管费率", [r"(?:基金托管人的)?托管费"]),
        ("sales_service_fee_rate", "销售服务费率", [r"销售服务费"]),
    ]
    for key, label, patterns in rate_specs:
        value = _extract_first_review_rate(text, patterns)
        if value:
            facts[key] = {"label": label, **value}

    text_specs = [
        ("fund_manager", "基金管理人", [r"基金管理人[：:]\s*([^\n]{4,80})"]),
        ("fund_custodian", "基金托管人", [r"基金托管人[：:]\s*([^\n]{4,80})"]),
        ("fund_type", "基金类型", [r"基金类型[：:]\s*([^\n]{2,80})"]),
        ("operation_mode", "运作方式", [r"运作方式[：:]\s*([^\n]{2,80})"]),
        ("target_index", "标的指数", [
            r"标的指数[：:]\s*([^\n，。；]{2,80})",
            r"标的指数为([^\n，。；]{2,80})",
        ]),
    ]
    for key, label, patterns in text_specs:
        value = _extract_first_review_text_fact(text, patterns)
        if value:
            facts[key] = {"label": label, **value}
    return facts


def _compare_review_key_facts(contract_text, prospectus_text):
    contract_facts = _extract_review_key_facts(contract_text)
    prospectus_facts = _extract_review_key_facts(prospectus_text)
    mismatches = []
    for key, contract_fact in contract_facts.items():
        prospectus_fact = prospectus_facts.get(key)
        if not prospectus_fact:
            continue
        if contract_fact.get("compare") == prospectus_fact.get("compare"):
            continue
        mismatches.append({
            "key": key,
            "label": contract_fact.get("label", key),
            "contract": contract_fact.get("display", ""),
            "prospectus": prospectus_fact.get("display", ""),
        })
    return mismatches


def _review_fact_mismatch_message(mismatches):
    if not mismatches:
        return ""
    parts = [
        f'{item["label"]}：合同为{item["contract"]}，招募为{item["prospectus"]}'
        for item in mismatches
    ]
    return "关键要素不一致：" + "；".join(parts)


def _finalize_review_issue(issue):
    item = _sanitize_review_markup_fields(dict(issue or {}))
    detail = item.get("description") or item.get("detail") or item.get("message") or ""
    item["detail"] = detail
    item["description"] = detail
    severity = item.get("severity") or "info"
    item["severity"] = severity
    item["is_problem"] = bool(item.get("is_problem", severity in ("error", "warning")))
    if "line" not in item and item.get("location"):
        match = re.search(r'第(\d+)行', str(item.get("location")))
        if match:
            item["line"] = int(match.group(1))
    return item


def _finalize_review_issues(issues):
    return [_finalize_review_issue(issue) for issue in (issues or [])]


# ── 摘要辅助 ──────────────────────────────────────────────────────────────────
def _parse_summary_rule_path(path):
    if not path:
        return ("", "")
    path = str(path).strip()
    m = re.match(r'^([一二三四五六七八九十]+)[、．.](.+)', path)
    if m:
        return (m.group(1), m.group(2).strip())
    return ("", path)


def _summary_group_path_from_heading(heading):
    if not heading:
        return ""
    heading = str(heading).strip()
    m = re.match(r'^([一二三四五六七八九十]+)[、．.]', heading)
    if m:
        return m.group(1)
    return heading


def _summary_heading_key(heading):
    if not heading:
        return ""
    text = unicodedata.normalize("NFKC", str(heading))
    text = re.sub(r'[\s、，,：:；;/\(\)（）]', '', text)
    return text


def _summary_heading_ordinal(heading):
    if not heading:
        return None
    heading = str(heading).strip()
    m = re.match(r'^([一二三四五六七八九十]+)', heading)
    if m:
        return _cn_numeral_to_int(m.group(1))
    m2 = re.match(r'^(\d+)', heading)
    if m2:
        return int(m2.group(1))
    return None


def _enrich_summary_rule_entry(entry):
    if not entry:
        return entry
    entry = dict(entry)
    entry["_ordinal"] = _summary_heading_ordinal(entry.get("summary_pos", ""))
    entry["_key"] = _summary_heading_key(entry.get("summary_pos", ""))
    return entry


def _find_matching_summary_rule(summary_pos, summary_rules):
    if not summary_pos or not summary_rules:
        return None
    key = _summary_heading_key(summary_pos)
    ordinal = _summary_heading_ordinal(summary_pos)
    for rule in summary_rules:
        enriched = _enrich_summary_rule_entry(rule)
        if ordinal is not None and enriched["_ordinal"] == ordinal:
            return enriched
        if key and enriched["_key"] == key:
            return enriched
    for rule in summary_rules:
        enriched = _enrich_summary_rule_entry(rule)
        if key and enriched["_key"] and (key in enriched["_key"] or enriched["_key"] in key):
            return enriched
    return None


def _find_summary_block_by_rule(blocks, summary_pos):
    if not blocks or not summary_pos:
        return None
    target_key = _summary_heading_key(summary_pos)
    target_ordinal = _summary_heading_ordinal(summary_pos)

    for block in blocks:
        heading = block.get("heading", "")
        if target_key and _summary_heading_key(heading) == target_key:
            return block
        if _review_soft_heading_match(heading, summary_pos):
            return block

    if target_ordinal is not None:
        for block in blocks:
            if _summary_heading_ordinal(block.get("heading", "")) == target_ordinal:
                return block

    return None


def _review_cross_section_optional(chapter_name):
    if not chapter_name:
        return True
    return "招募独有" in str(chapter_name) or "说明性" in str(chapter_name)


# ── 基金名称推断 ──────────────────────────────────────────────────────────────
def _clean_review_fund_name(value):
    name = Path(str(value or "")).stem.strip()
    if not name:
        return ""
    name = re.sub(r'^\d+\s*[、.．]\s*', '', name)
    name = re.sub(r'[（(]\s*草案\s*[）)]', '', name)
    name = re.sub(r'(基金合同|招募说明书|托管协议|法律意见书).*$','', name).strip()
    name = re.sub(r'[：:].*$', '', name).strip()
    return name if "基金" in name else ""


def _infer_review_fund_name(text, filename=""):
    filename_name = _clean_review_fund_name(filename)
    if filename_name:
        return filename_name

    text = str(text or "")
    nonempty_lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in nonempty_lines[:20]:
        cleaned = _clean_review_fund_name(line)
        if cleaned:
            return cleaned

    patterns = [
        r'基金名称[：:]\s*([^\n]{2,80})',
        r'基金或本基金[：:]\s*指\s*([^\n，。；]{4,100}基金)',
        r'([^\s，。；：“”"《》]{4,100}联接基金)',
        r'([^\s，。；：“”"《》]{4,100}证券投资基金)',
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            candidate = _clean_review_fund_name(m.group(1))
            if candidate and not re.match(r'^\d+\s*[、.．]', candidate):
                return candidate
    return ""


def _fallback_review_filename(context):
    ctx = context or {}
    fn = ctx.get("contract_filename", "") or ctx.get("filename", "") or "未命名"
    return fn.rsplit(".", 1)[0] if "." in fn else fn


# ── DOCX 解析 ─────────────────────────────────────────────────────────────────
def _iter_docx_blocks(parent):
    from docx.document import Document as DocumentObject
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
        parent_obj = parent
    else:
        parent_elm = parent._tc
        parent_obj = parent
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, parent_obj))
        elif isinstance(child, CT_Tbl):
            yield ("table", Table(child, parent_obj))


def _docx_table_to_markdown_lines(table):
    def cell_visible_text(cell):
        cell_lines = []
        for block_type, block in _iter_docx_blocks(cell):
            if block_type == "paragraph":
                text = _docx_paragraph_visible_text(block)
                if not _is_layout_blank_line(text):
                    cell_lines.append(text)
            elif block_type == "table":
                cell_lines.extend(_docx_table_to_markdown_lines(block))
        return _clean_docx_cell_text("\n".join(cell_lines))

    rows = []
    for row in table.rows:
        cells = [cell_visible_text(cell) for cell in row.cells]
        if any(not _is_layout_blank_line(cell) for cell in cells):
            rows.append(cells)
    return [f"|{'|'.join(cells)}|" for cells in rows]


def _extract_review_docx_text(doc):
    lines = []
    for block_type, block in _iter_docx_blocks(doc):
        if block_type == "paragraph":
            text = _normalize_layout_text(block.text).strip()
            if not _is_layout_blank_line(text):
                lines.append(text)
        elif block_type == "table":
            lines.extend(_docx_table_to_markdown_lines(block))
    return "\n".join(lines)


def _detect_review_blank_pages(doc):
    blank_pages = []
    page_texts = []
    current_text = []
    from docx.oxml.ns import qn
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        has_page_break = False
        if pPr is not None:
            for child in pPr.iterchildren():
                if child.tag == qn('w:pageBreakBefore'):
                    has_page_break = True
                    break
        if has_page_break and current_text:
            page_texts.append(" ".join(current_text))
            current_text = []
        current_text.append(para.text)
    if current_text:
        page_texts.append(" ".join(current_text))
    for i, text in enumerate(page_texts):
        if not text.strip():
            blank_pages.append(i + 1)
    return blank_pages


def _extract_review_doc_metadata(doc):
    meta = {}
    try:
        props = doc.core_properties
        if props.title:
            meta["title"] = props.title
        if props.author:
            meta["author"] = props.author
    except Exception as e:
        logger.debug("Docx core properties unavailable: %s", e)
    return meta


def _collect_review_doc_header_entries(doc):
    entries = []
    for i, para in enumerate(doc.paragraphs):
        if i >= 10:
            break
        text = para.text.strip()
        if text:
            entries.append(text)
    return entries


def _strip_contract_signing_page_text(text):
    if not text:
        return ""
    markers = ["签署页", "基金管理人：", "基金托管人：", "本页无正文"]
    cut_pos = len(text)
    for marker in markers:
        pos = text.rfind(marker)
        if pos > 0 and pos > len(text) * 0.8:
            marker_start = pos
            while marker_start > 0 and text[marker_start - 1] in " \t\r\n（(":
                marker_start -= 1
            cut_pos = min(cut_pos, marker_start)
    return text[:cut_pos].strip()


# ── Review Store ──────────────────────────────────────────────────────────────
_review_store = {}
_revision_workbench_store = {}
_review_lock = threading.Lock()


def _ensure_review_store_shape():
    if not isinstance(_review_store.get("jobs"), dict):
        _review_store["jobs"] = {}
    _review_store.setdefault("latest_job_id", None)


def _prune_review_jobs(now=None):
    now = now or time.time()
    jobs = _review_store.get("jobs", {})
    expired = [
        job_id for job_id, data in jobs.items()
        if now - data.get("_created_at", now) > REVIEW_JOB_TTL_SECONDS
    ]
    for job_id in expired:
        jobs.pop(job_id, None)
    if _review_store.get("latest_job_id") in expired:
        _review_store["latest_job_id"] = next(reversed(jobs), None) if jobs else None


def _store_review_job(data: dict) -> str:
    job_id = secrets.token_urlsafe(18)
    data = dict(data)
    data["_created_at"] = time.time()
    with _review_lock:
        _ensure_review_store_shape()
        _prune_review_jobs(data["_created_at"])
        _review_store["jobs"][job_id] = data
        _review_store["latest_job_id"] = job_id
        # Compatibility for older scripts that inspect the last uploaded data.
        _review_store["data"] = data
    return job_id


def _request_review_job_id(body=None) -> str:
    if body is None and request.method in {"POST", "PUT", "PATCH"}:
        body = request.get_json(silent=True) or {}
    body = body or {}
    return (
        request.args.get("job_id")
        or request.headers.get("X-Review-Job-Id")
        or body.get("job_id")
        or body.get("review_job_id")
        or ""
    )


def _get_review_data(body=None):
    job_id = _request_review_job_id(body)
    with _review_lock:
        _ensure_review_store_shape()
        _prune_review_jobs()
        if job_id:
            return _review_store["jobs"].get(job_id, {}), job_id
        latest_job_id = _review_store.get("latest_job_id")
        if latest_job_id and latest_job_id in _review_store["jobs"]:
            return _review_store["jobs"][latest_job_id], latest_job_id
        return _review_store.get("data", {}), ""


# ── 章节拆分 ──────────────────────────────────────────────────────────────────
def _split_contract_sections(text):
    if not text:
        return []
    RE_PART = re.compile(r'^(第[一二三四五六七八九十百]+部分\s*\S[^\n]*)', re.MULTILINE)
    RE_TOC_LINE = re.compile(r'\t\d+\s*')
    part_iter = [m for m in RE_PART.finditer(text)
                 if not RE_TOC_LINE.search(m.group(1))]
    sections = []
    for i, m in enumerate(part_iter):
        heading = m.group(1).split('\t')[0].strip()
        start = m.end()
        end = part_iter[i + 1].start() if i + 1 < len(part_iter) else len(text)
        content = _strip_contract_signing_page_text(text[start:end])
        sections.append({"heading": heading, "content": content})
    return sections


def _split_prospectus_sections(text, doc=None):
    if not text:
        return []

    def is_toc_line(line):
        stripped = (line or "").strip()
        return bool(stripped and "\t" in stripped and re.search(r'\t+\d+\s*$', stripped))

    def toc_heading(line):
        if not is_toc_line(line):
            return ""
        heading = re.sub(r'\t+\d+\s*$', '', line.strip()).strip()
        heading = re.sub(r'^[一二三四五六七八九十百]+、\s*', '', heading.replace("\t", " ")).strip()
        return heading

    def line_is_plain_heading(line):
        stripped = (line or "").strip()
        if not stripped or is_toc_line(stripped):
            return False
        if len(stripped) > 80:
            return False
        if re.search(r'[。！？；]', stripped):
            return False
        return True

    def labels_equivalent(left, right):
        for left_variant in _review_label_variants(left):
            left_key = _review_heading_key(left_variant)
            left_ordinal = _review_label_ordinal(left_variant)
            if not left_key:
                continue
            for right_variant in _review_label_variants(right):
                right_key = _review_heading_key(right_variant)
                right_ordinal = _review_label_ordinal(right_variant)
                if not right_key:
                    continue
                if left_ordinal is not None and right_ordinal is not None and left_ordinal != right_ordinal:
                    continue
                if left_key == right_key:
                    return True
        return False

    lines = text.split("\n")
    offsets = []
    offset = 0
    for line in lines:
        offsets.append((offset, line))
        offset += len(line) + 1

    toc_headings = []
    for _, line in offsets:
        heading = toc_heading(line)
        if heading and heading not in toc_headings:
            toc_headings.append(heading)

    matches = []
    search_start = 0
    for heading in toc_headings:
        for line_index in range(search_start, len(offsets)):
            pos, line = offsets[line_index]
            stripped = line.strip()
            if not line_is_plain_heading(stripped):
                continue
            if labels_equivalent(stripped, heading):
                matches.append((line_index, pos, stripped))
                search_start = line_index + 1
                break

    if not matches:
        re_part = re.compile(r'^(?:##\s+)?(第[一二三四五六七八九十百]+部分\s*\S[^\n]*)', re.MULTILINE)
        matches = [(text[:m.start()].count("\n"), m.start(), m.group(1).strip()) for m in re_part.finditer(text)]

    if not matches:
        re_top = re.compile(r'^([一二三四五六七八九十百]+、[^\n]+)', re.MULTILINE)
        for m in re_top.finditer(text):
            heading = m.group(1).strip()
            if is_toc_line(heading):
                continue
            matches.append((text[:m.start()].count("\n"), m.start(), heading))

    sections = []
    seen = set()
    deduped = []
    for line_index, pos, heading in matches:
        key = (line_index, heading)
        if key not in seen:
            seen.add(key)
            deduped.append((line_index, pos, heading))

    for i, (line_index, pos, heading) in enumerate(deduped):
        line_text = lines[line_index]
        start = pos + len(line_text)
        end = deduped[i + 1][1] if i + 1 < len(deduped) else len(text)
        sections.append({"heading": heading, "content": text[start:end].strip()})
    return sections


# ── 规则构建 ──────────────────────────────────────────────────────────────────
def _build_review_cross_rules():
    return list(BUILTIN_LINKED_FUND_CROSS_RULES)


def _load_review_rules(fund_type="联接基金"):
    cross_rules = list(BUILTIN_LINKED_FUND_CROSS_RULES)
    summary_rules = [_enrich_summary_rule_entry(r) for r in BUILTIN_LINKED_FUND_SUMMARY_RULES]
    return {
        "cross": [],
        "summary": list(summary_rules),
        "chapter_rules": list(cross_rules),
        "chapter_level": list(cross_rules),
        "detail_rules": [],
        "summary_rules": list(summary_rules),
        "key_diffs": [],
        "key_diff_rules": [],
        "workbook_meta": {"fund_type": fund_type or "联接基金"},
    }


def _truncate_ai_review_text(text, limit=4000):
    text = str(text or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "\n……（以下内容已截断）"


def _build_ai_cross_review_items(contract_sections, prospectus_sections, review_rules, fund_type="联接基金"):
    items = []
    seen = set()
    cross_rules = review_rules.get("chapter_rules") or review_rules.get("chapter_level") or _build_review_cross_rules()

    for rule in cross_rules:
        contract_target = _locate_review_rule_target(
            contract_sections,
            rule.get("contract_locator") or rule.get("contract_chapter") or "",
            rule.get("contract_section_name") or rule.get("contract_chapter") or "",
        )
        prospectus_target = _locate_review_rule_target(
            prospectus_sections,
            rule.get("prospectus_locator") or rule.get("prospectus_chapter") or "",
            rule.get("prospectus_section_name") or rule.get("prospectus_chapter") or "",
        )
        if not contract_target.get("matched") or not prospectus_target.get("matched"):
            continue

        contract_locator = _format_review_target_locator(contract_target)
        prospectus_locator = _format_review_target_locator(prospectus_target)
        dedupe_key = (contract_locator, prospectus_locator)
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)

        contract_text = (contract_target.get("body_text") or contract_target.get("text") or "").strip()
        prospectus_text = (prospectus_target.get("body_text") or prospectus_target.get("text") or "").strip()
        if not contract_text or not prospectus_text:
            continue

        items.append({
            "label": f"合同↔招募: {contract_locator} ↔ {prospectus_locator}",
            "contract_locator": contract_locator,
            "prospectus_locator": prospectus_locator,
            "contract_text": _truncate_ai_review_text(contract_text),
            "prospectus_text": _truncate_ai_review_text(prospectus_text),
            "rule": rule,
        })
    return items


def _build_ai_summary_review_items(contract_sections, review_rules):
    contract_summary = _find_contract_summary_section(contract_sections)
    if not contract_summary:
        return []

    summary_rules = review_rules.get("summary_rules") or review_rules.get("summary") or []
    summary_blocks = _split_summary_top_blocks(contract_summary.get("content", ""))
    items = []
    seen = set()

    for block in summary_blocks:
        summary_heading = block.get("heading", "")
        summary_text = (block.get("body") or "").strip()
        if not summary_text:
            continue
        rule = _find_matching_summary_rule(summary_heading, summary_rules)
        if not rule:
            continue
        method_text = _normalize_review_text(rule.get("method", ""))
        status_text = _normalize_review_text(rule.get("status", ""))
        if "未纳入摘要" in method_text or "未进入摘要" in method_text or "未纳入摘要" in status_text:
            continue

        target = _locate_review_rule_target(contract_sections, rule.get("contract_pos", ""))
        if not target.get("matched"):
            continue

        contract_locator = _format_review_target_locator(target)
        if (summary_heading, contract_locator) in seen:
            continue
        seen.add((summary_heading, contract_locator))

        contract_text = (target.get("body_text") or target.get("text") or "").strip()
        if not contract_text:
            continue
        items.append({
            "label": f"正文↔摘要: {contract_locator}",
            "summary_heading": summary_heading,
            "summary_text": _truncate_ai_review_text(summary_text),
            "contract_locator": contract_locator,
            "contract_text": _truncate_ai_review_text(contract_text),
            "rule": rule,
        })
    return items


# ── 摘要章节查找 ──────────────────────────────────────────────────────────────
def _find_contract_summary_section(sections):
    for sec in sections:
        h = sec.get("heading", "")
        if "基金合同的内容摘要" in h or "基金合同内容摘要" in h:
            return sec
    for sec in sections:
        h = sec.get("heading", "")
        if "摘要" in h and "合同" in h:
            return sec
    return None


# ── 审核工作台：DOCX 修订阅读模型 ───────────────────────────────────────────────
_DOCX_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_DOCX_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_DOCX_NS = {"w": _DOCX_W_NS}

_REVISION_SCOPE_LABELS = {
    "contract_body": "基金合同正文",
    "contract_summary": "基金合同摘要",
    "prospectus_body": "招募说明书",
    "prospectus_contract_summary": "招募说明书内基金合同摘要",
    "custody_agreement": "托管协议",
    "product_summary": "产品资料概要",
}

_REVISION_IMPACT_TARGETS = {
    "contract_body": ["contract_summary", "prospectus_body", "prospectus_contract_summary", "product_summary"],
    "contract_summary": ["prospectus_contract_summary", "product_summary"],
    "prospectus_body": ["product_summary"],
    "prospectus_contract_summary": ["product_summary"],
    "custody_agreement": ["prospectus_body"],
    "product_summary": [],
}


def _docx_local_name(tag: str) -> str:
    return str(tag or "").rsplit("}", 1)[-1]


def _docx_attr(node, name: str) -> str:
    return str(node.attrib.get(f"{{{_DOCX_W_NS}}}{name}") or node.attrib.get(name) or "").strip()


def _docx_xml_text(node) -> str:
    parts = []
    for child in node.iter():
        name = _docx_local_name(child.tag)
        if name in {"t", "delText"} and child.text:
            parts.append(str(child.text))
        elif name == "tab":
            parts.append("\t")
        elif name in {"br", "cr"}:
            parts.append("\n")
    return "".join(parts).strip()


def _read_docx_xml_part(docx_bytes: bytes, part_name: str) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
            with zf.open(part_name) as part:
                return part.read().decode("utf-8")
    except (KeyError, zipfile.BadZipFile, UnicodeDecodeError):
        return ""


def _load_docx_comments(docx_bytes: bytes) -> dict[str, dict]:
    xml_text = _read_docx_xml_part(docx_bytes, "word/comments.xml")
    if not xml_text:
        return {}
    import xml.etree.ElementTree as ET

    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return {}

    comments = {}
    for comment in root.findall(".//w:comment", _DOCX_NS):
        comment_id = _docx_attr(comment, "id")
        if comment_id:
            comments[comment_id] = {
                "author": _docx_attr(comment, "author"),
                "date": _docx_attr(comment, "date"),
                "comment_text": _docx_xml_text(comment),
            }
    return comments


def _revision_base_document_kind(document_kind: str) -> str:
    kind = str(document_kind or "").strip()
    if kind.startswith("product_summary"):
        return "product_summary"
    return kind


def _revision_default_scope(document_kind: str) -> str:
    base_kind = _revision_base_document_kind(document_kind)
    if base_kind == "contract":
        return "contract_body"
    if base_kind == "prospectus":
        return "prospectus_body"
    if base_kind == "product_summary":
        return "product_summary"
    if base_kind == "custody_agreement":
        return "custody_agreement"
    return base_kind or "unknown"


def _revision_scope_for_paragraph(document_kind: str, paragraph_text: str, current_scope: str) -> str:
    base_kind = _revision_base_document_kind(document_kind)
    compact = re.sub(r"\s+", "", _normalize_review_text(paragraph_text))
    if base_kind == "contract" and re.fullmatch(r"(?:第二十四部分)?基金合同(?:的)?内容摘要", compact):
        return "contract_summary"
    if base_kind == "prospectus":
        appendix_prefix = r"(?:附件[一二三四五六七八九十]+[、:：]?)?"
        numbered_prefix = r"(?:[一二三四五六七八九十]+[、:：])?"
        if re.fullmatch(rf"{appendix_prefix}{numbered_prefix}基金合同(?:的)?内容摘要", compact):
            return "prospectus_contract_summary"
        if re.fullmatch(rf"{appendix_prefix}{numbered_prefix}(?:基金)?托管协议(?:的)?内容摘要", compact):
            return "prospectus_body"
    return current_scope


def _revision_impact_targets(scope: str) -> list[dict]:
    return [
        {"key": target, "label": _REVISION_SCOPE_LABELS.get(target, target)}
        for target in _REVISION_IMPACT_TARGETS.get(scope, [])
    ]


def _revision_label_path(scope: str) -> str:
    label = _REVISION_SCOPE_LABELS.get(scope, scope)
    targets = _revision_impact_targets(scope)
    return " → ".join([label, *[target["label"] for target in targets]])


def _build_revision_reader_revision(
    *,
    revision_id: str,
    document_kind: str,
    scope: str,
    filename: str,
    revision_type: str,
    block_id: str,
    context: str,
    text: str = "",
    old_text: str = "",
    new_text: str = "",
    author: str = "",
    date: str = "",
    comment_text: str = "",
) -> dict:
    label = _REVISION_SCOPE_LABELS.get(scope, scope or document_kind)
    return {
        "id": revision_id,
        "document_kind": document_kind,
        "document_scope": scope,
        "document_label": label,
        "source_file": filename,
        "revision_type": revision_type,
        "text": text,
        "old_text": old_text,
        "new_text": new_text,
        "comment_text": comment_text,
        "block_id": block_id,
        "context": context,
        "author": author,
        "date": date,
        "decision": "accepted",
        "impact_targets": _revision_impact_targets(scope),
        "impact_path": _revision_label_path(scope),
    }


def _append_revision_text_run(
    runs: list[dict],
    *,
    text: str,
    mark_type: str = "normal",
    revision_type: str = "",
    author: str = "",
    date: str = "",
    comment_id: str = "",
    comment_text: str = "",
) -> None:
    if text:
        runs.append({
            "text": text,
            "mark_type": mark_type,
            "revision_type": revision_type,
            "author": author,
            "date": date,
            "comment_id": comment_id,
            "comment_text": comment_text,
        })


def _append_revision_plain_text_run(runs: list[dict], text: str, active_comments: list[str], comments: dict[str, dict]) -> None:
    if not text:
        return
    comment_id = active_comments[-1] if active_comments else ""
    if comment_id and comment_id in comments:
        comment = comments[comment_id]
        _append_revision_text_run(
            runs,
            text=text,
            mark_type="comment",
            revision_type="comment",
            author=comment.get("author", ""),
            date=comment.get("date", ""),
            comment_id=comment_id,
            comment_text=comment.get("comment_text", ""),
        )
    else:
        _append_revision_text_run(runs, text=text)


def _append_revision_inline_runs(node, runs: list[dict], active_comments: list[str], comments: dict[str, dict]) -> None:
    name = _docx_local_name(node.tag)
    if name in {"ins", "del", "moveTo", "moveFrom"}:
        revision_text = _docx_xml_text(node)
        if revision_text:
            revision_type = "insert" if name in {"ins", "moveTo"} else "delete"
            _append_revision_text_run(
                runs,
                text=revision_text,
                mark_type=revision_type,
                revision_type=revision_type,
                author=_docx_attr(node, "author"),
                date=_docx_attr(node, "date"),
            )
        return
    if name == "commentRangeStart":
        comment_id = _docx_attr(node, "id")
        if comment_id:
            active_comments.append(comment_id)
        return
    if name == "commentRangeEnd":
        comment_id = _docx_attr(node, "id")
        if comment_id in active_comments:
            active_comments[:] = [cid for cid in active_comments if cid != comment_id]
        elif active_comments:
            active_comments.pop()
        return
    if name in {"t", "delText"}:
        _append_revision_plain_text_run(runs, str(node.text or ""), active_comments, comments)
        return
    if name == "tab":
        _append_revision_plain_text_run(runs, "\t", active_comments, comments)
        return
    if name in {"br", "cr"}:
        _append_revision_plain_text_run(runs, "\n", active_comments, comments)
        return
    for child in list(node):
        _append_revision_inline_runs(child, runs, active_comments, comments)


def _parse_revision_paragraph_runs(paragraph, comments: dict[str, dict]) -> list[dict]:
    runs: list[dict] = []
    active_comments: list[str] = []
    for child in list(paragraph):
        name = _docx_local_name(child.tag)
        if name in {"commentRangeStart", "commentRangeEnd"}:
            _append_revision_inline_runs(child, runs, active_comments, comments)
            continue
        if name in {"ins", "del", "moveTo", "moveFrom", "r", "hyperlink", "sdt"}:
            _append_revision_inline_runs(child, runs, active_comments, comments)
    return runs


def _finalize_revision_runs(
    *,
    runs: list[dict],
    document_kind: str,
    scope: str,
    filename: str,
    block_id: str,
    context: str,
    revision_counter: int,
) -> tuple[list[dict], list[dict], int]:
    revisions: list[dict] = []
    finalized = [dict(run) for run in runs]
    index = 0
    comment_revision_ids: dict[str, str] = {}
    while index < len(finalized):
        run = finalized[index]
        mark_type = run.get("mark_type")
        next_run = finalized[index + 1] if index + 1 < len(finalized) else None
        next_type = (next_run or {}).get("mark_type")
        is_modify_pair = {mark_type, next_type} == {"delete", "insert"}

        if is_modify_pair:
            revision_counter += 1
            revision_id = f"{document_kind}-reader-{revision_counter}"
            from_run = run if mark_type == "delete" else next_run
            to_run = run if mark_type == "insert" else next_run
            run["mark_type"] = "modify_from" if mark_type == "delete" else "modify_to"
            next_run["mark_type"] = "modify_from" if next_type == "delete" else "modify_to"
            run["revision_id"] = revision_id
            next_run["revision_id"] = revision_id
            revisions.append(_build_revision_reader_revision(
                revision_id=revision_id,
                document_kind=document_kind,
                scope=scope,
                filename=filename,
                revision_type="modify",
                block_id=block_id,
                context=context,
                text=f"{from_run.get('text', '')} → {to_run.get('text', '')}",
                old_text=from_run.get("text", ""),
                new_text=to_run.get("text", ""),
                author=to_run.get("author") or from_run.get("author", ""),
                date=to_run.get("date") or from_run.get("date", ""),
            ))
            index += 2
            continue

        if mark_type in {"insert", "delete"}:
            revision_counter += 1
            revision_id = f"{document_kind}-reader-{revision_counter}"
            run["revision_id"] = revision_id
            revisions.append(_build_revision_reader_revision(
                revision_id=revision_id,
                document_kind=document_kind,
                scope=scope,
                filename=filename,
                revision_type=mark_type,
                block_id=block_id,
                context=context,
                text=run.get("text", ""),
                author=run.get("author", ""),
                date=run.get("date", ""),
            ))
        elif mark_type == "comment":
            comment_id = run.get("comment_id") or f"{block_id}-comment-{index}"
            revision_id = comment_revision_ids.get(comment_id)
            if not revision_id:
                revision_counter += 1
                revision_id = f"{document_kind}-reader-{revision_counter}"
                comment_revision_ids[comment_id] = revision_id
                revisions.append(_build_revision_reader_revision(
                    revision_id=revision_id,
                    document_kind=document_kind,
                    scope=scope,
                    filename=filename,
                    revision_type="comment",
                    block_id=block_id,
                    context=context,
                    text=run.get("text", ""),
                    author=run.get("author", ""),
                    date=run.get("date", ""),
                    comment_text=run.get("comment_text", ""),
                ))
            run["revision_id"] = revision_id
        index += 1

    for run in finalized:
        for key in ("revision_type", "comment_id"):
            run.pop(key, None)
    return finalized, revisions, revision_counter


def _build_revision_paragraph_block(*, paragraph, document_kind: str, filename: str, block_id: str, scope: str, comments: dict[str, dict], revision_counter: int):
    runs = _parse_revision_paragraph_runs(paragraph, comments)
    text = "".join(run.get("text", "") for run in runs).strip()
    next_scope = _revision_scope_for_paragraph(document_kind, text, scope) if text else scope
    finalized_runs, revisions, revision_counter = _finalize_revision_runs(
        runs=runs,
        document_kind=document_kind,
        scope=next_scope,
        filename=filename,
        block_id=block_id,
        context=text,
        revision_counter=revision_counter,
    )
    block = {
        "id": block_id,
        "type": "paragraph",
        "document_scope": next_scope,
        "document_label": _REVISION_SCOPE_LABELS.get(next_scope, next_scope),
        "text": text,
        "runs": finalized_runs,
        "revision_ids": [revision["id"] for revision in revisions],
    }
    return block, revisions, revision_counter, next_scope


def _build_revision_table_block(*, table, document_kind: str, filename: str, block_id: str, scope: str, comments: dict[str, dict], revision_counter: int):
    rows = []
    revisions: list[dict] = []
    current_scope = scope
    cell_texts: list[str] = []
    revision_ids: list[str] = []
    for row_index, row in enumerate(table.findall(".//w:tr", _DOCX_NS), start=1):
        cells = []
        for cell_index, cell in enumerate(row.findall("./w:tc", _DOCX_NS), start=1):
            cell_blocks = []
            for paragraph_index, paragraph in enumerate(cell.findall("./w:p", _DOCX_NS), start=1):
                child_block_id = f"{block_id}-r{row_index}c{cell_index}p{paragraph_index}"
                paragraph_block, paragraph_revisions, revision_counter, current_scope = _build_revision_paragraph_block(
                    paragraph=paragraph,
                    document_kind=document_kind,
                    filename=filename,
                    block_id=child_block_id,
                    scope=current_scope,
                    comments=comments,
                    revision_counter=revision_counter,
                )
                cell_blocks.append(paragraph_block)
                revisions.extend(paragraph_revisions)
                revision_ids.extend(paragraph_block["revision_ids"])
                if paragraph_block.get("text"):
                    cell_texts.append(paragraph_block["text"])
            cells.append({"blocks": cell_blocks, "text": "\n".join(block.get("text", "") for block in cell_blocks if block.get("text"))})
        rows.append(cells)
    block = {
        "id": block_id,
        "type": "table",
        "document_scope": current_scope,
        "document_label": _REVISION_SCOPE_LABELS.get(current_scope, current_scope),
        "text": "\n".join(cell_texts),
        "rows": rows,
        "revision_ids": revision_ids,
    }
    return block, revisions, revision_counter, current_scope


def _count_revision_items(items: list[dict]) -> dict:
    counts = {scope: 0 for scope in _REVISION_SCOPE_LABELS}
    for item in items or []:
        scope = str((item or {}).get("document_scope") or "").strip()
        if scope:
            counts[scope] = counts.get(scope, 0) + 1
    counts["total"] = sum(counts.values())
    return counts


def _build_revision_document_view(docx_source, *, document_kind: str, filename: str = "") -> dict:
    if hasattr(docx_source, "read"):
        current_pos = None
        try:
            current_pos = docx_source.tell()
        except Exception:
            current_pos = None
        docx_bytes = docx_source.read()
        if current_pos is not None:
            try:
                docx_source.seek(current_pos)
            except Exception:
                pass
    else:
        docx_bytes = bytes(docx_source or b"")

    import xml.etree.ElementTree as ET

    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        return {"document_kind": document_kind, "filename": filename, "blocks": [], "revisions": [], "revision_counts": _count_revision_items([])}
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return {"document_kind": document_kind, "filename": filename, "blocks": [], "revisions": [], "revision_counts": _count_revision_items([])}

    body = root.find("w:body", _DOCX_NS)
    comments = _load_docx_comments(docx_bytes)
    blocks = []
    revisions: list[dict] = []
    scope = _revision_default_scope(document_kind)
    revision_counter = 0
    block_counter = 0
    for child in list(body) if body is not None else []:
        name = _docx_local_name(child.tag)
        if name not in {"p", "tbl"}:
            continue
        block_counter += 1
        block_id = f"{document_kind}-block-{block_counter}"
        if name == "p":
            block, block_revisions, revision_counter, scope = _build_revision_paragraph_block(
                paragraph=child,
                document_kind=document_kind,
                filename=filename,
                block_id=block_id,
                scope=scope,
                comments=comments,
                revision_counter=revision_counter,
            )
        else:
            block, block_revisions, revision_counter, scope = _build_revision_table_block(
                table=child,
                document_kind=document_kind,
                filename=filename,
                block_id=block_id,
                scope=scope,
                comments=comments,
                revision_counter=revision_counter,
            )
        if block.get("text") or block.get("type") == "table":
            blocks.append(block)
        revisions.extend(block_revisions)
    return {
        "document_kind": document_kind,
        "filename": filename,
        "blocks": blocks,
        "revisions": revisions,
        "revision_counts": _count_revision_items(revisions),
    }


# ═══════════════════════════════════════════════════════════════════════════════
#  复核 API 路由
# ═══════════════════════════════════════════════════════════════════════════════

def _normalize_review_export_decision(value: str) -> str:
    raw = str(value or "").strip().lower()
    if raw in {"accept", "accepted", "接受", "已接受"}:
        return "accept"
    if raw in {"reject", "rejected", "refuse", "拒绝", "已拒绝"}:
        return "reject"
    return "pending"


def _register_docx_xml_namespaces(ET) -> None:
    namespaces = {
        "w": _DOCX_W_NS,
        "r": _DOCX_R_NS,
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
        "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    }
    for prefix, uri in namespaces.items():
        try:
            ET.register_namespace(prefix, uri)
        except ValueError:
            pass


def _copy_docx_element(ET, node):
    return ET.fromstring(ET.tostring(node, encoding="utf-8"))


def _docx_root_start_tag(xml_text: str) -> str:
    match = re.search(r"<[^!?][^>]*\bdocument\b[^>]*>", str(xml_text or ""))
    return match.group(0) if match else ""


def _serialize_docx_document_xml(ET, root, original_xml_text: str | bytes = "") -> bytes:
    rendered = ET.tostring(root, encoding="utf-8", xml_declaration=True).decode("utf-8")
    original = bytes(original_xml_text).decode("utf-8", errors="ignore") if isinstance(original_xml_text, (bytes, bytearray)) else str(original_xml_text or "")
    original_start = _docx_root_start_tag(original)
    rendered_start = _docx_root_start_tag(rendered)
    if not original_start or not rendered_start:
        return rendered.encode("utf-8")

    missing_declarations = []
    for declaration in re.findall(r'xmlns(?::[\w.\-]+)?="[^"]+"', original_start):
        declaration_name = declaration.split("=", 1)[0]
        if f"{declaration_name}=" not in rendered_start:
            missing_declarations.append(declaration)
    if not missing_declarations:
        return rendered.encode("utf-8")

    insert_at = rendered.find(">", rendered.find(rendered_start))
    if insert_at < 0:
        return rendered.encode("utf-8")
    rendered = f"{rendered[:insert_at]} {' '.join(missing_declarations)}{rendered[insert_at:]}"
    return rendered.encode("utf-8")


def _convert_deleted_run_to_plain_text(run) -> None:
    for child in run.iter():
        if _docx_local_name(child.tag) == "delText":
            child.tag = f"{{{_DOCX_W_NS}}}t"


def _deleted_revision_to_plain_runs(ET, revision_node) -> list:
    runs = []
    for run in revision_node.findall(".//w:r", _DOCX_NS):
        copied = _copy_docx_element(ET, run)
        _convert_deleted_run_to_plain_text(copied)
        runs.append(copied)
    if runs:
        return runs

    text = _docx_xml_text(revision_node)
    if not text:
        return []
    run = ET.Element(f"{{{_DOCX_W_NS}}}r")
    text_node = ET.SubElement(run, f"{{{_DOCX_W_NS}}}t")
    text_node.text = text
    return [run]


def _deleted_revision_to_plain_text_nodes(ET, revision_node) -> list:
    text = _docx_xml_text(revision_node)
    if not text:
        return []
    text_node = ET.Element(f"{{{_DOCX_W_NS}}}t")
    text_node.text = text
    return [text_node]


def _replace_docx_child(parent, child, replacements: list) -> None:
    children = list(parent)
    try:
        index = children.index(child)
    except ValueError:
        return
    parent.remove(child)
    for offset, replacement in enumerate(replacements or []):
        parent.insert(index + offset, replacement)


def _remove_docx_child(parent, child) -> None:
    try:
        parent.remove(child)
    except ValueError:
        pass


def _docx_revision_token_from_child(child, active_comments: list[str], comments: dict[str, dict]) -> dict | None:
    name = _docx_local_name(child.tag)
    if name in {"ins", "del", "moveTo", "moveFrom"}:
        revision_text = _docx_xml_text(child)
        if not revision_text:
            return None
        revision_type = "insert" if name in {"ins", "moveTo"} else "delete"
        return {
            "mark_type": revision_type,
            "revision_type": revision_type,
            "text": revision_text,
            "element": child,
            "author": _docx_attr(child, "author"),
            "date": _docx_attr(child, "date"),
        }
    if name in {"r", "hyperlink", "sdt"}:
        text = _docx_xml_text(child)
        if not text:
            return None
        comment_id = active_comments[-1] if active_comments else ""
        if comment_id and comment_id in comments:
            return {
                "mark_type": "comment",
                "revision_type": "comment",
                "text": text,
                "element": child,
                "comment_id": comment_id,
                "comment_text": comments[comment_id].get("comment_text", ""),
                "author": comments[comment_id].get("author", ""),
                "date": comments[comment_id].get("date", ""),
            }
        return {"mark_type": "normal", "text": text, "element": child}
    return None


def _append_review_export_text_token(
    tokens: list[dict],
    *,
    text: str,
    element,
    parent,
    active_comments: list[str],
    comments: dict[str, dict],
) -> None:
    if not text:
        return
    comment_id = active_comments[-1] if active_comments else ""
    if comment_id and comment_id in comments:
        tokens.append({
            "mark_type": "comment",
            "revision_type": "comment",
            "text": text,
            "element": element,
            "parent": parent,
            "comment_id": comment_id,
            "comment_text": comments[comment_id].get("comment_text", ""),
            "author": comments[comment_id].get("author", ""),
            "date": comments[comment_id].get("date", ""),
        })
        return
    tokens.append({"mark_type": "normal", "text": text, "element": element, "parent": parent})


def _append_review_export_tokens_from_node(
    node,
    parent,
    tokens: list[dict],
    active_comments: list[str],
    comments: dict[str, dict],
) -> None:
    name = _docx_local_name(node.tag)
    if name in {"ins", "del", "moveTo", "moveFrom"}:
        revision_text = _docx_xml_text(node)
        if revision_text:
            revision_type = "insert" if name in {"ins", "moveTo"} else "delete"
            tokens.append({
                "mark_type": revision_type,
                "revision_type": revision_type,
                "text": revision_text,
                "element": node,
                "parent": parent,
                "author": _docx_attr(node, "author"),
                "date": _docx_attr(node, "date"),
            })
        return
    if name == "commentRangeStart":
        comment_id = _docx_attr(node, "id")
        if comment_id:
            active_comments.append(comment_id)
        return
    if name == "commentRangeEnd":
        comment_id = _docx_attr(node, "id")
        if comment_id in active_comments:
            active_comments[:] = [cid for cid in active_comments if cid != comment_id]
        elif active_comments:
            active_comments.pop()
        return
    if name in {"t", "delText"}:
        _append_review_export_text_token(
            tokens,
            text=str(node.text or ""),
            element=node,
            parent=parent,
            active_comments=active_comments,
            comments=comments,
        )
        return
    if name == "tab":
        _append_review_export_text_token(
            tokens,
            text="\t",
            element=node,
            parent=parent,
            active_comments=active_comments,
            comments=comments,
        )
        return
    if name in {"br", "cr"}:
        _append_review_export_text_token(
            tokens,
            text="\n",
            element=node,
            parent=parent,
            active_comments=active_comments,
            comments=comments,
        )
        return

    for child in list(node):
        _append_review_export_tokens_from_node(child, node, tokens, active_comments, comments)


def _review_export_paragraph_tokens(paragraph, comments: dict[str, dict]) -> list[dict]:
    tokens: list[dict] = []
    active_comments: list[str] = []
    for child in list(paragraph):
        _append_review_export_tokens_from_node(child, paragraph, tokens, active_comments, comments)
    return tokens


def _apply_rejected_docx_revision(parent, token: dict, ET) -> None:
    mark_type = token.get("mark_type")
    element = token.get("element")
    token_parent = token.get("parent")
    if token_parent is not None:
        parent = token_parent
    if element is None:
        return
    if mark_type == "insert":
        _remove_docx_child(parent, element)
    elif mark_type == "delete":
        if _docx_local_name(getattr(parent, "tag", "")) == "r":
            replacements = _deleted_revision_to_plain_text_nodes(ET, element)
        else:
            replacements = _deleted_revision_to_plain_runs(ET, element)
        _replace_docx_child(parent, element, replacements)


def _strip_docx_comment_markers(root, comment_ids: set[str]) -> None:
    if not comment_ids:
        return
    for parent in root.iter():
        for child in list(parent):
            name = _docx_local_name(child.tag)
            if name in {"commentRangeStart", "commentRangeEnd"} and _docx_attr(child, "id") in comment_ids:
                parent.remove(child)
                continue
            if name == "r":
                references = child.findall(".//w:commentReference", _DOCX_NS)
                if any(_docx_attr(ref, "id") in comment_ids for ref in references):
                    parent.remove(child)


def _copy_docx_run_properties(ET, run):
    if run is None:
        return None
    r_pr = run.find("w:rPr", _DOCX_NS)
    return _copy_docx_element(ET, r_pr) if r_pr is not None else None


DEFAULT_REVIEW_AUTHOR = "审核工作台"


def _normalize_review_author(value) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", str(value or ""))
    text = re.sub(r"\s+", " ", text).strip()
    return (text[:80].strip() or DEFAULT_REVIEW_AUTHOR)


def _normalize_review_revision_edits(revision_edits: dict | None) -> dict[str, dict[str, str]]:
    if not isinstance(revision_edits, dict):
        return {}
    normalized: dict[str, dict[str, str]] = {}
    for revision_id, payload in revision_edits.items():
        clean_id = str(revision_id or "").strip()
        if not clean_id:
            continue
        text = payload.get("text", "") if isinstance(payload, dict) else payload
        clean_text = str(text or "").strip()
        if clean_text:
            normalized[clean_id] = {"text": clean_text}
            if isinstance(payload, dict) and str(payload.get("mode") or "").strip().lower() == "paragraph":
                normalized[clean_id]["mode"] = "paragraph"
    return normalized


def _review_revision_edit_text(revision_edits: dict[str, dict[str, str]], revision_id: str) -> str:
    payload = revision_edits.get(revision_id) or {}
    return str(payload.get("text") or "").strip()


def _review_revision_edit_mode(revision_edits: dict[str, dict[str, str]], revision_id: str) -> str:
    payload = revision_edits.get(revision_id) or {}
    return str(payload.get("mode") or "").strip().lower()


def _docx_original_text_from_node(node) -> str:
    def iter_text(current):
        name = _docx_local_name(current.tag)
        if name in {"ins", "moveTo"}:
            return
        if name in {"del", "moveFrom"}:
            text = _docx_xml_text(current)
            if text:
                yield text
            return
        if name in {"t", "delText"} and current.text:
            yield str(current.text)
        elif name == "tab":
            yield "\t"
        elif name in {"br", "cr"}:
            yield "\n"
        for child in list(current):
            yield from iter_text(child)

    return "".join(iter_text(node)).strip()


def _copy_docx_revision_run_properties(ET, revision_node):
    if revision_node is None:
        return None
    for run in revision_node.findall(".//w:r", _DOCX_NS):
        run_properties = _copy_docx_run_properties(ET, run)
        if run_properties is not None:
            return run_properties
    return None


def _append_docx_revision_text_runs(ET, container, text: str, *, tag_name: str, run_properties=None) -> None:
    lines = str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    for index, line in enumerate(lines):
        if index:
            break_run = ET.SubElement(container, f"{{{_DOCX_W_NS}}}r")
            if run_properties is not None:
                break_run.append(_copy_docx_element(ET, run_properties))
            ET.SubElement(break_run, f"{{{_DOCX_W_NS}}}br")
        if line or len(lines) == 1:
            run = ET.SubElement(container, f"{{{_DOCX_W_NS}}}r")
            if run_properties is not None:
                run.append(_copy_docx_element(ET, run_properties))
            run.append(_docx_revision_text_element(ET, tag_name, line))


def _replace_docx_revision_text(ET, revision_node, text: str, *, review_author: str | None = None) -> bool:
    if revision_node is None:
        return False
    if review_author:
        revision_node.set(f"{{{_DOCX_W_NS}}}author", _normalize_review_author(review_author))
        revision_node.set(f"{{{_DOCX_W_NS}}}date", datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"))
    tag_name = "delText" if _docx_local_name(revision_node.tag) in {"del", "moveFrom"} else "t"
    run_properties = _copy_docx_revision_run_properties(ET, revision_node)
    for child in list(revision_node):
        revision_node.remove(child)
    _append_docx_revision_text_runs(ET, revision_node, text, tag_name=tag_name, run_properties=run_properties)
    return True


def _insert_docx_revision_after_token(ET, root, token: dict, text: str, *, review_author: str = DEFAULT_REVIEW_AUTHOR) -> bool:
    parent = token.get("parent")
    element = token.get("element")
    if parent is None or element is None:
        return False
    children = list(parent)
    try:
        index = children.index(element)
    except ValueError:
        return False

    change = ET.Element(f"{{{_DOCX_W_NS}}}ins")
    change.set(f"{{{_DOCX_W_NS}}}id", str(_next_docx_change_id(root)))
    change.set(f"{{{_DOCX_W_NS}}}author", _normalize_review_author(review_author))
    change.set(
        f"{{{_DOCX_W_NS}}}date",
        datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
    )
    _append_docx_revision_text_runs(
        ET,
        change,
        text,
        tag_name="t",
        run_properties=_copy_docx_revision_run_properties(ET, element),
    )
    parent.insert(index + 1, change)
    return True


def _docx_revision_change_with_text_runs(ET, *, text: str, change_type: str, change_id: int, author: str, date: str, run_properties=None):
    change = ET.Element(f"{{{_DOCX_W_NS}}}{change_type}")
    change.set(f"{{{_DOCX_W_NS}}}id", str(change_id))
    change.set(f"{{{_DOCX_W_NS}}}author", author)
    change.set(f"{{{_DOCX_W_NS}}}date", date)
    _append_docx_revision_text_runs(
        ET,
        change,
        text,
        tag_name="delText" if change_type == "del" else "t",
        run_properties=run_properties,
    )
    return change


def _replace_paragraph_with_full_redline(ET, paragraph, old_text: str, new_text: str, *, start_change_id: int, author: str, date: str) -> int:
    p_pr = _copy_docx_paragraph_properties(ET, paragraph)
    run_properties = _copy_docx_first_run_properties(ET, paragraph)
    for child in list(paragraph):
        paragraph.remove(child)
    if p_pr is not None:
        paragraph.append(p_pr)

    next_id = start_change_id
    if old_text:
        paragraph.append(_docx_revision_change_with_text_runs(
            ET,
            text=old_text,
            change_type="del",
            change_id=next_id,
            author=author,
            date=date,
            run_properties=run_properties,
        ))
        next_id += 1
    if new_text:
        paragraph.append(_docx_revision_change_with_text_runs(
            ET,
            text=new_text,
            change_type="ins",
            change_id=next_id,
            author=author,
            date=date,
            run_properties=run_properties,
        ))
        next_id += 1
    return next_id


_REVIEW_NUMBER_PREFIX_PATTERN = re.compile(
    r"^(\s*(?:[（(]\s*(?:[0-9０-９]+|[一二三四五六七八九十百千]+)\s*[）)]|(?:[0-9０-９]+|[一二三四五六七八九十百千]+)[、.．]))(.*)$",
    re.S,
)


def _split_review_number_prefix(text: str) -> tuple[str, str]:
    match = _REVIEW_NUMBER_PREFIX_PATTERN.match(str(text or ""))
    if not match:
        return "", str(text or "")
    return match.group(1), match.group(2)


def _review_lines_need_number_prefix_redline(old_line: str, new_line: str) -> bool:
    old_prefix, old_body = _split_review_number_prefix(old_line)
    new_prefix, new_body = _split_review_number_prefix(new_line)
    if not old_prefix or not new_prefix or old_prefix == new_prefix:
        return False
    return _normalize_contract_prospectus_compare_text(old_body) == _normalize_contract_prospectus_compare_text(new_body)


def _review_lines_have_number_prefix_redlines(old_lines: list[str], new_lines: list[str]) -> bool:
    return any(
        _review_lines_need_number_prefix_redline(old_line, new_line)
        for old_line, new_line in zip(old_lines, new_lines)
    )


def _docx_first_text_node(run, tag_name: str = "t"):
    if run is None or _docx_local_name(getattr(run, "tag", "")) != "r":
        return None
    return run.find(f".//w:{tag_name}", _DOCX_NS)


def _normalize_docx_leading_number_revision_marks(root, ET, *, review_author: str = DEFAULT_REVIEW_AUTHOR) -> None:
    number_text_pattern = re.compile(r"^\s*(?:[0-9０-９]+|[一二三四五六七八九十百千]+)\s*$")
    opening_pattern = re.compile(r"^\s*[（(]\s*$")
    for paragraph in root.findall(".//w:p", _DOCX_NS):
        changed = True
        while changed:
            changed = False
            children = list(paragraph)
            for index in range(0, max(len(children) - 3, 0)):
                open_run, first_change, second_change, close_run = children[index:index + 4]
                if _docx_local_name(open_run.tag) != "r" or _docx_local_name(close_run.tag) != "r":
                    continue
                first_name = _docx_local_name(first_change.tag)
                second_name = _docx_local_name(second_change.tag)
                if {first_name, second_name} != {"del", "ins"}:
                    continue
                delete_node = first_change if first_name == "del" else second_change
                insert_node = first_change if first_name == "ins" else second_change
                old_number = _docx_xml_text(delete_node).strip()
                new_number = _docx_xml_text(insert_node).strip()
                if not number_text_pattern.fullmatch(old_number) or not number_text_pattern.fullmatch(new_number):
                    continue
                open_text_node = _docx_first_text_node(open_run)
                close_text_node = _docx_first_text_node(close_run)
                open_text = str(open_text_node.text or "") if open_text_node is not None else ""
                close_text = str(close_text_node.text or "") if close_text_node is not None else ""
                if not opening_pattern.fullmatch(open_text) or close_text[:1] not in {"）", ")"}:
                    continue

                open_char = "（" if "（" in open_text else "("
                close_char = close_text[0]
                old_prefix = f"{open_char}{old_number}{close_char}"
                new_prefix = f"{open_char}{new_number}{close_char}"
                run_properties = _copy_docx_run_properties(ET, open_run) or _copy_docx_run_properties(ET, close_run)
                delete_id = _docx_attr(delete_node, "id") or str(_next_docx_change_id(root))
                if _docx_attr(insert_node, "id"):
                    insert_id = _docx_attr(insert_node, "id")
                elif str(delete_id).isdigit():
                    insert_id = str(int(str(delete_id)) + 1)
                else:
                    insert_id = str(_next_docx_change_id(root))
                delete_change = _docx_inline_change(
                    ET,
                    text=old_prefix,
                    change_type="del",
                    change_id=delete_id,
                    author=_docx_attr(delete_node, "author") or _normalize_review_author(review_author),
                    date=_docx_attr(delete_node, "date") or datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
                    run_properties=run_properties,
                )
                insert_change = _docx_inline_change(
                    ET,
                    text=new_prefix,
                    change_type="ins",
                    change_id=insert_id,
                    author=_docx_attr(insert_node, "author") or _normalize_review_author(review_author),
                    date=_docx_attr(insert_node, "date") or datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
                    run_properties=run_properties,
                )

                close_text_node.text = close_text[1:]
                insert_at = index
                for node in (open_run, first_change, second_change):
                    try:
                        paragraph.remove(node)
                    except ValueError:
                        pass
                if not close_text_node.text:
                    try:
                        paragraph.remove(close_run)
                    except ValueError:
                        pass
                paragraph.insert(insert_at, insert_change)
                paragraph.insert(insert_at, delete_change)
                changed = True
                break


def _remove_comments_from_comments_xml(comments_xml: bytes | str, comment_ids: set[str]) -> bytes | str:
    if not comment_ids:
        return comments_xml
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    is_bytes = isinstance(comments_xml, (bytes, bytearray))
    raw = bytes(comments_xml).decode("utf-8", errors="ignore") if is_bytes else str(comments_xml or "")
    try:
        root = ET.fromstring(raw)
    except ET.ParseError:
        for comment_id in comment_ids:
            raw = re.sub(
                rf"<w:comment\b[^>]*\bw:id=\"{re.escape(comment_id)}\"[\s\S]*?</w:comment>",
                "",
                raw,
            )
        return raw.encode("utf-8") if is_bytes else raw

    for parent in root.iter():
        for child in list(parent):
            if _docx_local_name(child.tag) == "comment" and _docx_attr(child, "id") in comment_ids:
                parent.remove(child)
    rendered = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    return rendered if is_bytes else rendered.decode("utf-8")


def _repack_docx_parts(docx_bytes: bytes, replacements: dict[str, bytes | str], *, removals: set[str] | None = None) -> bytes:
    output = io.BytesIO()
    removals = removals or set()
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as src, zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        written = set()
        for item in src.infolist():
            if item.filename in removals:
                continue
            data = replacements.get(item.filename)
            if data is None:
                data = src.read(item.filename)
            if isinstance(data, str):
                data = data.encode("utf-8")
            dst.writestr(item, data)
            written.add(item.filename)
        for name, data in replacements.items():
            if name in written or name in removals:
                continue
            if isinstance(data, str):
                data = data.encode("utf-8")
            dst.writestr(name, data)
    return output.getvalue()


def _review_export_revision_actions(tokens: list[dict], document_kind: str, revision_counter: int, decisions: dict, revision_edits: dict) -> tuple[list[dict], int]:
    actions: list[dict] = []
    index = 0
    while index < len(tokens):
        token = tokens[index]
        mark_type = token.get("mark_type")
        next_token = tokens[index + 1] if index + 1 < len(tokens) else None
        next_type = (next_token or {}).get("mark_type")
        is_modify_pair = {mark_type, next_type} == {"delete", "insert"}

        if is_modify_pair:
            revision_counter += 1
            revision_id = f"{document_kind}-reader-{revision_counter}"
            actions.append({
                "revision_id": revision_id,
                "kind": "modify",
                "decision": _normalize_review_export_decision(decisions.get(revision_id)),
                "token": token,
                "next_token": next_token,
                "edit_text": _review_revision_edit_text(revision_edits, revision_id),
                "edit_mode": _review_revision_edit_mode(revision_edits, revision_id),
            })
            index += 2
            continue

        if mark_type in {"insert", "delete"}:
            revision_counter += 1
            revision_id = f"{document_kind}-reader-{revision_counter}"
            actions.append({
                "revision_id": revision_id,
                "kind": mark_type,
                "decision": _normalize_review_export_decision(decisions.get(revision_id)),
                "token": token,
                "edit_text": _review_revision_edit_text(revision_edits, revision_id),
                "edit_mode": _review_revision_edit_mode(revision_edits, revision_id),
            })
        elif mark_type == "comment":
            comment_id = token.get("comment_id") or ""
            revision_id = token.get("_revision_id")
            if not revision_id:
                revision_counter += 1
                revision_id = f"{document_kind}-reader-{revision_counter}"
                for peer in tokens:
                    if peer.get("mark_type") == "comment" and peer.get("comment_id") == comment_id:
                        peer["_revision_id"] = revision_id
                actions.append({
                    "revision_id": revision_id,
                    "kind": "comment",
                    "decision": _normalize_review_export_decision(decisions.get(revision_id)),
                    "token": token,
                    "comment_id": comment_id,
                })
        index += 1
    return actions, revision_counter


def _apply_docx_paragraph_revision_edit(ET, root, paragraph, action: dict, *, review_author: str = DEFAULT_REVIEW_AUTHOR) -> bool:
    edit_text = str(action.get("edit_text") or "").strip()
    if not edit_text:
        return False
    token = action.get("token") or {}
    old_text = _docx_original_text_from_node(paragraph)
    start_change_id = _next_docx_change_id(root)
    _replace_paragraph_with_full_redline(
        ET,
        paragraph,
        old_text,
        edit_text,
        start_change_id=start_change_id,
        author=_normalize_review_author(review_author),
        date=datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
    )
    return True


def _apply_review_revision_decisions_to_docx(
    docx_bytes: bytes,
    *,
    document_kind: str,
    filename: str = "",
    decisions: dict | None = None,
    revision_edits: dict | None = None,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    """Apply review decisions while keeping accepted changes as Word redlines."""
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    decisions = decisions or {}
    revision_edits = _normalize_review_revision_edits(revision_edits)
    review_author = _normalize_review_author(review_author)
    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        raise ValueError("DOCX 缺少 word/document.xml，无法生成红线终稿")

    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError as exc:
        raise ValueError("DOCX 正文 XML 无法解析，无法生成红线终稿") from exc

    comments = _load_docx_comments(docx_bytes)
    revision_counter = 0
    accepted_count = 0
    rejected_count = 0
    pending_count = 0
    rejected_comment_ids: set[str] = set()
    applied_ids: list[str] = []
    edited_revision_ids: list[str] = []

    for paragraph in root.findall(".//w:p", _DOCX_NS):
        tokens = _review_export_paragraph_tokens(paragraph, comments)
        actions, revision_counter = _review_export_revision_actions(tokens, document_kind, revision_counter, decisions, revision_edits)
        if not actions:
            continue

        for action in actions:
            revision_id = action["revision_id"]
            decision = action["decision"]
            applied_ids.append(revision_id)
            if decision == "reject":
                rejected_count += 1
                if action["kind"] == "comment" and action.get("comment_id"):
                    rejected_comment_ids.add(action["comment_id"])
            elif decision == "pending":
                pending_count += 1
            else:
                accepted_count += 1

        paragraph_edit = next(
            (
                action
                for action in actions
                if action["kind"] in {"insert", "delete", "modify"}
                and action["decision"] == "accept"
                and action.get("edit_text")
                and action.get("edit_mode") == "paragraph"
            ),
            None,
        )
        if paragraph_edit:
            if _apply_docx_paragraph_revision_edit(ET, root, paragraph, paragraph_edit, review_author=review_author):
                edited_revision_ids.append(paragraph_edit["revision_id"])
            continue

        for action in actions:
            decision = action["decision"]
            kind = action["kind"]
            token = action.get("token")
            next_token = action.get("next_token")
            revision_id = action["revision_id"]

            if decision == "reject":
                if kind == "modify":
                    for part in (token, next_token):
                        if part and part.get("mark_type") == "delete":
                            _apply_rejected_docx_revision(paragraph, part, ET)
                    for part in (token, next_token):
                        if part and part.get("mark_type") == "insert":
                            _apply_rejected_docx_revision(paragraph, part, ET)
                elif kind in {"insert", "delete"}:
                    _apply_rejected_docx_revision(paragraph, token, ET)
                continue
            if decision != "accept":
                continue

            edit_text = str(action.get("edit_text") or "").strip()
            if not edit_text:
                continue
            if kind == "modify":
                insert_part = token if token and token.get("mark_type") == "insert" else next_token if next_token and next_token.get("mark_type") == "insert" else None
                if insert_part and _replace_docx_revision_text(ET, insert_part.get("element"), edit_text, review_author=review_author):
                    edited_revision_ids.append(revision_id)
            elif kind == "insert" and _replace_docx_revision_text(ET, token.get("element"), edit_text, review_author=review_author):
                edited_revision_ids.append(revision_id)
            elif kind == "delete" and _insert_docx_revision_after_token(ET, root, token, edit_text, review_author=review_author):
                edited_revision_ids.append(revision_id)

    _strip_docx_comment_markers(root, rejected_comment_ids)
    _normalize_docx_leading_number_revision_marks(root, ET, review_author=review_author)
    replacements: dict[str, bytes | str] = {
        "word/document.xml": _serialize_docx_document_xml(ET, root, xml_text)
    }
    comments_xml = _read_docx_xml_part(docx_bytes, "word/comments.xml")
    if comments_xml and rejected_comment_ids:
        replacements["word/comments.xml"] = _remove_comments_from_comments_xml(comments_xml, rejected_comment_ids)

    final_bytes = _repack_docx_parts(docx_bytes, replacements)
    return final_bytes, {
        "document_kind": document_kind,
        "filename": filename,
        "accepted_count": accepted_count,
        "rejected_count": rejected_count,
        "pending_count": pending_count,
        "edited_count": len(edited_revision_ids),
        "applied_revision_ids": applied_ids,
        "edited_revision_ids": edited_revision_ids,
        "rejected_comment_ids": sorted(rejected_comment_ids),
    }


def _docx_effective_text_from_bytes(docx_bytes: bytes) -> str:
    import xml.etree.ElementTree as ET

    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        return ""
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return ""

    def iter_effective_text(node, excluded: bool = False):
        name = _docx_local_name(node.tag)
        if name in {"del", "moveFrom"}:
            excluded = True
        if not excluded:
            if name == "t" and node.text:
                yield str(node.text)
            elif name == "tab":
                yield "\t"
            elif name in {"br", "cr"}:
                yield "\n"
        for child in list(node):
            yield from iter_effective_text(child, excluded)

    lines = []
    for paragraph in root.findall(".//w:p", _DOCX_NS):
        text = "".join(iter_effective_text(paragraph)).strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def _build_review_decision_csv(
    decisions: dict,
    documents: list[dict],
    revision_edits: dict | None = None,
    applied_edit_ids: set[str] | None = None,
) -> str:
    import csv

    revision_edits = _normalize_review_revision_edits(revision_edits)
    applied_edit_ids = applied_edit_ids or set()
    revisions_by_id = {
        revision.get("id"): revision
        for document in documents or []
        for revision in (document or {}).get("revisions", [])
        if revision.get("id")
    }
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["revision_id", "decision", "document", "scope", "type", "text", "edited_text", "edit_applied"])
    for revision_id in sorted(revisions_by_id):
        revision = revisions_by_id[revision_id]
        edited_text = _review_revision_edit_text(revision_edits, revision_id)
        writer.writerow([
            revision_id,
            _normalize_review_export_decision(decisions.get(revision_id)),
            revision.get("source_file", ""),
            revision.get("document_label", ""),
            revision.get("revision_type", ""),
            revision.get("text") or revision.get("new_text") or revision.get("old_text") or "",
            edited_text,
            "yes" if revision_id in applied_edit_ids else "no",
        ])
    return output.getvalue()


def _render_review_contract_summary_from_text(contract_text: str) -> tuple[str, dict]:
    template_text = TEMPLATE_MD.read_text(encoding="utf-8")
    template_summary = _find_contract_summary_section(_split_contract_sections(template_text))
    if not template_summary:
        raise ValueError("未在基金合同模板中找到第二十四部分摘要结构")
    scaffold = "\n".join([
        str(template_summary.get("heading") or "第二十四部分  基金合同内容摘要"),
        str(template_summary.get("content") or "").strip(),
    ]).strip()
    scaffold = engine._process_conditionals(scaffold, {"CONTRACT_SUMMARY_LEGACY_DISABLED": False})
    merged = f"{contract_text.rstrip()}\n\n{scaffold}"
    rendered = engine._replace_contract_summary_placeholders(merged)
    rendered_sections = _split_contract_sections(rendered)
    rendered_summary_candidates = [
        section
        for section in rendered_sections
        if _review_docx_is_fund_contract_summary_heading(section.get("heading") or "")
        and len(_strip_contract_signing_page_text(section.get("content") or "")) > 10
    ]
    rendered_summary = rendered_summary_candidates[-1] if rendered_summary_candidates else _find_contract_summary_section(rendered_sections)
    if not rendered_summary:
        raise ValueError("基金合同摘要重算后无法定位摘要章节")
    summary_text = "\n".join([
        str(rendered_summary.get("heading") or "").strip(),
        str(rendered_summary.get("content") or "").strip(),
    ]).strip()
    return summary_text, {
        "status": "generated",
        "message": "已根据应用审核决定后的基金合同正文重算基金合同摘要",
        "characters": len(summary_text),
    }


def _effective_docx_text_from_node(node) -> str:
    def iter_text(current, excluded: bool = False):
        name = _docx_local_name(current.tag)
        if name in {"del", "moveFrom"}:
            excluded = True
        if not excluded:
            if name == "t" and current.text:
                yield str(current.text)
            elif name == "tab":
                yield "\t"
            elif name in {"br", "cr"}:
                yield "\n"
        for child in list(current):
            yield from iter_text(child, excluded)

    return "".join(iter_text(node)).strip()


def _docx_revision_text_element(ET, tag_name: str, text: str):
    text_node = ET.Element(f"{{{_DOCX_W_NS}}}{tag_name}")
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    return text_node


def _next_docx_change_id(root) -> int:
    max_id = 0
    for node in root.iter():
        node_id = _docx_attr(node, "id")
        if node_id.isdigit():
            max_id = max(max_id, int(node_id))
    return max_id + 1


def _tracked_review_paragraph(ET, *, text: str, change_type: str, change_id: int, author: str, date: str):
    paragraph = ET.Element(f"{{{_DOCX_W_NS}}}p")
    change = ET.SubElement(paragraph, f"{{{_DOCX_W_NS}}}{change_type}")
    change.set(f"{{{_DOCX_W_NS}}}id", str(change_id))
    change.set(f"{{{_DOCX_W_NS}}}author", author)
    change.set(f"{{{_DOCX_W_NS}}}date", date)
    run = ET.SubElement(change, f"{{{_DOCX_W_NS}}}r")
    text_tag = "delText" if change_type == "del" else "t"
    run.append(_docx_revision_text_element(ET, text_tag, text))
    return paragraph


def _review_docx_summary_heading_compact(text: str) -> str:
    compact = re.sub(r"\s+", "", text or "")
    replacements = {
        "基金合同的内容摘要": "基金合同内容摘要",
        "基金托管协议的内容摘要": "基金托管协议内容摘要",
        "托管协议的内容摘要": "托管协议内容摘要",
    }
    for old, new in replacements.items():
        compact = compact.replace(old, new)
    return compact


def _review_docx_heading_has_toc_page_number(compact_heading: str) -> bool:
    return bool(re.search(r"[0-9０-９]+$", compact_heading or ""))


def _review_docx_is_fund_contract_summary_heading(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    if _review_docx_heading_has_toc_page_number(compact):
        return False
    return "基金合同内容摘要" in compact and "托管协议" not in compact


def _strip_matching_heading_from_summary(summary_text: str, heading_text: str) -> str:
    lines = [line for line in str(summary_text or "").splitlines()]
    if not lines:
        return ""
    heading_compact = _review_docx_summary_heading_compact(heading_text)
    first_compact = _review_docx_summary_heading_compact(lines[0])
    if heading_compact and first_compact == heading_compact:
        return "\n".join(lines[1:]).strip()
    if "基金合同内容摘要" in first_compact:
        return "\n".join(lines[1:]).strip()
    return "\n".join(lines).strip()


def _docx_text_is_contract_signing_page_start(text: str) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if not compact:
        return False
    if "签署页" in compact and ("基金合同" in compact or "无正文" in compact):
        return True
    if "本页" in compact and "基金合同" in compact and ("无正文" in compact or "签署页" in compact):
        return True
    return False


def _replace_docx_section_content_with_redline(
    docx_bytes: bytes,
    *,
    is_heading,
    is_stop_heading,
    new_text: str,
    report_label: str,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：DOCX 缺少正文 XML"}
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：正文 XML 无法解析"}
    _normalize_docx_leading_number_revision_marks(root, ET, review_author=review_author)

    body = root.find("w:body", _DOCX_NS)
    if body is None:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：未找到正文 body"}

    children = list(body)
    heading_index = None
    heading_text = ""
    for index, child in enumerate(children):
        if _docx_local_name(child.tag) != "p":
            continue
        text = _effective_docx_text_from_node(child)
        if is_heading(text):
            heading_index = index
            heading_text = text
            break
    if heading_index is None:
        return docx_bytes, {"status": "not_found", "message": f"{report_label}：未定位到目标摘要标题，未自动替换"}

    end_index = len(children)
    for index in range(heading_index + 1, len(children)):
        child = children[index]
        if _docx_local_name(child.tag) == "sectPr":
            end_index = index
            break
        if _docx_local_name(child.tag) == "p":
            if child.find("w:pPr/w:sectPr", _DOCX_NS) is not None:
                end_index = index
                break
            paragraph_text = _effective_docx_text_from_node(child)
            if is_stop_heading(paragraph_text) or _docx_text_is_contract_signing_page_start(paragraph_text):
                end_index = index
                break

    old_blocks = children[heading_index + 1:end_index]
    old_units = []
    for child in old_blocks:
        if _docx_local_name(child.tag) != "p":
            continue
        line = _effective_docx_text_from_node(child).strip()
        if not line:
            continue
        old_units.append({
            "node": child,
            "line": line,
            "compare": _normalize_contract_prospectus_compare_text(_strip_review_number_prefix(line)),
        })
    old_text = "\n".join(unit["line"] for unit in old_units).strip()
    replacement_text = _strip_matching_heading_from_summary(new_text, heading_text)
    if not replacement_text:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：新摘要为空，未替换"}
    if re.sub(r"\s+", "", old_text) == re.sub(r"\s+", "", replacement_text):
        return docx_bytes, {"status": "unchanged", "message": f"{report_label}：内容已一致，未新增系统红线"}

    for child in old_blocks:
        try:
            body.remove(child)
        except ValueError:
            pass

    next_id = _next_docx_change_id(root)
    now = datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
    author = _normalize_review_author(review_author)
    replacements = []
    new_rows = _review_sync_line_rows(replacement_text)
    new_lines = [row["line"].strip() for row in new_rows]
    old_compare = [unit["compare"] for unit in old_units]
    new_compare = [row["compare"] for row in new_rows]
    old_nodes = [unit["node"] for unit in old_units]
    old_lines = [unit["line"] for unit in old_units]
    matcher = difflib.SequenceMatcher(None, old_compare, new_compare, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(i2 - i1):
                paragraph = old_nodes[i1 + offset]
                old_line = old_lines[i1 + offset]
                new_line = new_lines[j1 + offset]
                if _review_lines_need_number_prefix_redline(old_line, new_line):
                    next_id = _replace_paragraph_with_minimal_redline(
                        ET,
                        paragraph,
                        old_line,
                        new_line,
                        start_change_id=next_id,
                        author=author,
                        date=now,
                    )
                replacements.append(paragraph)
            continue
        if tag == "replace" and (i2 - i1) == (j2 - j1):
            for offset in range(i2 - i1):
                paragraph = old_nodes[i1 + offset]
                next_id = _replace_paragraph_with_minimal_redline(
                    ET,
                    paragraph,
                    old_lines[i1 + offset],
                    new_lines[j1 + offset],
                    start_change_id=next_id,
                    author=author,
                    date=now,
                )
                replacements.append(paragraph)
            continue
        template = old_nodes[i1] if i1 < len(old_nodes) else (old_nodes[i1 - 1] if i1 > 0 else None)
        for old_line in old_lines[i1:i2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=old_line, change_type="del", change_id=next_id, author=author, date=now))
            next_id += 1
        for new_line in new_lines[j1:j2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=new_line, change_type="ins", change_id=next_id, author=author, date=now))
            next_id += 1

    insert_index = heading_index + 1
    for offset, replacement in enumerate(replacements):
        body.insert(insert_index + offset, replacement)

    final_bytes = _repack_docx_parts(
        docx_bytes,
        {"word/document.xml": _serialize_docx_document_xml(ET, root, xml_text)},
    )
    return final_bytes, {
        "status": "redlined",
        "message": f"{report_label}：已以红线方式替换摘要内容",
        "old_characters": len(old_text),
        "new_characters": len(replacement_text),
    }


def _contract_summary_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    return compact == "第二十四部分基金合同内容摘要"


def _contract_summary_stop_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    return bool(re.match(r"^第[一二三四五六七八九十百]+部分", compact or "") and compact != "第二十四部分基金合同内容摘要")


def _prospectus_contract_summary_heading_match(text: str) -> bool:
    return _review_docx_is_fund_contract_summary_heading(text)


def _prospectus_contract_summary_stop_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    if _review_docx_heading_has_toc_page_number(compact):
        return False
    return "基金托管协议内容摘要" in compact or "托管协议内容摘要" in compact


def _prospectus_custody_summary_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    if _review_docx_heading_has_toc_page_number(compact):
        return False
    return "基金托管协议内容摘要" in compact or "托管协议内容摘要" in compact


def _prospectus_custody_summary_stop_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    if not compact or _review_docx_heading_has_toc_page_number(compact):
        return False
    if _prospectus_custody_summary_heading_match(text):
        return False
    stop_keywords = (
        "基金份额持有人服务",
        "其他应披露事项",
        "其它应披露事项",
        "招募说明书存放",
        "备查文件",
    )
    return any(keyword in compact for keyword in stop_keywords)


def _sync_review_contract_summary_redlines(contract_docx: bytes, prospectus_docx: bytes, *, review_author: str = DEFAULT_REVIEW_AUTHOR) -> tuple[bytes, bytes, dict]:
    contract_text = _docx_effective_text_from_bytes(contract_docx)
    try:
        summary_text, summary_report = _render_review_contract_summary_from_text(contract_text)
    except Exception as exc:
        return contract_docx, prospectus_docx, {
            "contract_summary": {"status": "skipped", "message": "未能自动重算基金合同摘要", "error": str(exc)},
            "prospectus_contract_summary": {"status": "skipped", "message": "因合同摘要未重算，未同步招募说明书内合同摘要"},
        }

    contract_docx, contract_apply_report = _replace_docx_section_content_with_redline(
        contract_docx,
        is_heading=_contract_summary_heading_match,
        is_stop_heading=_contract_summary_stop_heading_match,
        new_text=summary_text,
        report_label="基金合同第 24 部分摘要",
        review_author=review_author,
    )
    prospectus_docx, prospectus_apply_report = _replace_docx_section_content_with_redline(
        prospectus_docx,
        is_heading=_prospectus_contract_summary_heading_match,
        is_stop_heading=_prospectus_contract_summary_stop_heading_match,
        new_text=summary_text,
        report_label="招募说明书内基金合同摘要",
        review_author=review_author,
    )
    return contract_docx, prospectus_docx, {
        "contract_summary": {**summary_report, "apply": contract_apply_report},
        "prospectus_contract_summary": prospectus_apply_report,
    }


def _sync_review_custody_summary_redlines(
    custody_docx: bytes,
    prospectus_docx: bytes,
    *,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    custody_text = _docx_effective_text_from_bytes(custody_docx)
    summary_report = _extract_custody_agreement_summary_from_text(custody_text)
    summary_text = str(summary_report.get("summary_text") or "").strip()
    if not summary_text:
        return prospectus_docx, {
            "status": "skipped",
            "message": "未能从托管协议提取可同步的内容摘要",
            "summary": summary_report,
            "apply": {"status": "skipped", "message": "未同步招募说明书托管协议摘要"},
        }

    prospectus_docx, apply_report = _replace_docx_section_content_with_redline(
        prospectus_docx,
        is_heading=_prospectus_custody_summary_heading_match,
        is_stop_heading=_prospectus_custody_summary_stop_heading_match,
        new_text=summary_text,
        report_label="招募说明书内托管协议摘要",
        review_author=review_author,
    )
    return prospectus_docx, {
        "status": apply_report.get("status") or "skipped",
        "message": apply_report.get("message") or "已尝试同步招募说明书托管协议摘要",
        "summary": summary_report,
        "apply": apply_report,
    }


def _review_cross_rule_is_direct_sync_candidate(rule: dict) -> tuple[bool, str]:
    relation = _normalize_review_text((rule or {}).get("relation"))
    consistency = _normalize_review_text((rule or {}).get("consistency"))
    combined = " ".join([relation, consistency, _normalize_review_text((rule or {}).get("detail"))])
    contract_locator = _normalize_review_text((rule or {}).get("contract_locator") or (rule or {}).get("contract_chapter"))
    prospectus_locator = _normalize_review_text((rule or {}).get("prospectus_locator") or (rule or {}).get("prospectus_chapter"))
    locator_pair = f"{contract_locator} {prospectus_locator}"
    is_whole_chapter_rule = (
        not any(separator in contract_locator for separator in ("/", "／"))
        and not any(separator in prospectus_locator for separator in ("/", "／"))
    )
    is_whole_distribution_chapter = (
        "基金的收益与分配" in locator_pair
        and is_whole_chapter_rule
    )
    if is_whole_distribution_chapter:
        return False, "收益分配章节含招募细化条款，不做整章自动替换"
    if any(keyword in combined for keyword in ("无直接对应", "招募独有", "合同独有", "无同名独立条款", "无对应独立章节")):
        return False, "无直接对应或单方独有"
    if any(keyword in relation for keyword in ("招募细化", "细化", "部分对应", "说明性对应")):
        return False, "招募细化或部分对应规则不自动替换"
    if is_whole_chapter_rule:
        return False, "整章对应需经过招募说明书后处理，不做整章自动替换"
    if "直接对应" not in relation and "完全一致" not in consistency:
        return False, "非直接对应规则"
    if consistency and not any(keyword in consistency for keyword in ("完全一致", "基本一致")):
        return False, "一致性口径不是完全一致或基本一致"
    return True, ""


def _review_sync_minimal_change_guard(source_text: str, target_text: str) -> tuple[bool, str]:
    source_rows = _review_sync_line_rows(source_text)
    target_rows = _review_sync_line_rows(target_text)
    source_compare = [row["compare"] for row in source_rows if row.get("compare")]
    target_compare = [row["compare"] for row in target_rows if row.get("compare")]
    if not source_compare or not target_compare:
        return False, "同步源或目标为空"
    if len(source_compare) != len(target_compare):
        if len(target_compare) > len(source_compare):
            return False, "招募说明书存在额外内容，未自动删除"
        return False, "合同来源内容存在新增段落，需人工确认后同步"
    source_joined = "\n".join(source_compare)
    target_joined = "\n".join(target_compare)
    similarity = difflib.SequenceMatcher(None, source_joined, target_joined, autojunk=False).ratio()
    if similarity < 0.72:
        return False, f"文本差异较大（相似度 {similarity:.0%}），未自动替换"
    return True, ""


def _review_sync_line_rows(text: str) -> list[dict]:
    return _collect_nonblank_compare_lines(
        text,
        normalize_line=lambda line: _normalize_contract_prospectus_compare_text(_strip_review_number_prefix(line)),
    )


def _fallback_review_rule_target_from_text(full_text: str, locator: str, section_hint: str = "") -> dict:
    parts = _split_review_locator(locator)
    labels = [part for part in [*(reversed(parts)), section_hint, locator] if _normalize_review_text(part)]
    seen = set()
    for label in labels:
        normalized = _normalize_review_text(label)
        if normalized in seen:
            continue
        seen.add(normalized)
        block = _find_review_block_by_label(full_text, label)
        if not block:
            continue
        heading = _normalize_review_text(block.get("heading"))
        body_text = (block.get("body") or block.get("text") or "").strip()
        return {
            "section": None,
            "section_heading": "",
            "target_heading": heading,
            "text": (block.get("text") or "").strip(),
            "body_text": body_text,
            "matched": bool(body_text),
            "match_method": "full_text_block",
            "missing_reason": "" if body_text else "body_empty",
            "locator_parts": [heading] if heading else [],
        }
    return {
        "section": None,
        "section_heading": "",
        "target_heading": locator,
        "text": "",
        "body_text": "",
        "matched": False,
        "match_method": "",
        "missing_reason": "fallback_block_missing",
        "locator_parts": [],
    }


def _docx_body_paragraphs(root) -> list[dict]:
    body = root.find("w:body", _DOCX_NS)
    if body is None:
        return []
    paragraphs = []
    for child in list(body):
        if _docx_local_name(child.tag) != "p":
            continue
        text = _effective_docx_text_from_node(child)
        if text:
            paragraphs.append({"node": child, "text": text, "compare": _normalize_contract_prospectus_compare_text(_strip_review_number_prefix(text))})
    return paragraphs


def _find_docx_paragraph_span_by_text(root, target_text: str) -> tuple[list, str]:
    target_rows = _review_sync_line_rows(target_text)
    target_compare = [row["compare"] for row in target_rows if row.get("compare")]
    if not target_compare:
        return [], "target_empty"

    paragraphs = _docx_body_paragraphs(root)
    paragraph_compare = [row["compare"] for row in paragraphs]
    target_len = len(target_compare)
    for start in range(0, max(len(paragraph_compare) - target_len + 1, 0)):
        if paragraph_compare[start:start + target_len] == target_compare:
            return [row["node"] for row in paragraphs[start:start + target_len]], ""

    if target_len == 1:
        target = target_compare[0]
        for row in paragraphs:
            if row["compare"] == target:
                return [row["node"]], ""
    return [], "target_not_found_in_docx"


def _copy_docx_paragraph_properties(ET, paragraph):
    p_pr = paragraph.find("w:pPr", _DOCX_NS)
    return _copy_docx_element(ET, p_pr) if p_pr is not None else None


def _copy_docx_first_run_properties(ET, paragraph):
    if paragraph is None:
        return None
    for run in paragraph.findall(".//w:r", _DOCX_NS):
        r_pr = run.find("w:rPr", _DOCX_NS)
        if r_pr is not None:
            return _copy_docx_element(ET, r_pr)
    return None


def _docx_plain_text_run(ET, text: str, run_properties=None):
    run = ET.Element(f"{{{_DOCX_W_NS}}}r")
    if run_properties is not None:
        run.append(_copy_docx_element(ET, run_properties))
    run.append(_docx_revision_text_element(ET, "t", text))
    return run


def _docx_inline_change(ET, *, text: str, change_type: str, change_id: int, author: str, date: str, run_properties=None):
    change = ET.Element(f"{{{_DOCX_W_NS}}}{change_type}")
    change.set(f"{{{_DOCX_W_NS}}}id", str(change_id))
    change.set(f"{{{_DOCX_W_NS}}}author", author)
    change.set(f"{{{_DOCX_W_NS}}}date", date)
    run = ET.SubElement(change, f"{{{_DOCX_W_NS}}}r")
    if run_properties is not None:
        run.append(_copy_docx_element(ET, run_properties))
    run.append(_docx_revision_text_element(ET, "delText" if change_type == "del" else "t", text))
    return change


def _review_sync_diff_tokens(text: str) -> list[str]:
    return re.findall(r"\d+(?:\.\d+)?%?|[A-Za-z]+(?:[-_][A-Za-z0-9]+)*|[\u4e00-\u9fff]+|[^\u4e00-\u9fffA-Za-z\d]+", str(text or ""))


def _append_minimal_redline_text_runs(ET, paragraph, old_text: str, new_text: str, *, next_id: int, author: str, date: str, run_properties=None) -> int:
    old_tokens = _review_sync_diff_tokens(old_text)
    new_tokens = _review_sync_diff_tokens(new_text)
    matcher = difflib.SequenceMatcher(None, old_tokens, new_tokens, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        old_part = "".join(old_tokens[i1:i2])
        new_part = "".join(new_tokens[j1:j2])
        if tag == "equal":
            if old_part:
                paragraph.append(_docx_plain_text_run(ET, old_part, run_properties=run_properties))
            continue
        if tag in {"replace", "delete"} and old_part:
            paragraph.append(_docx_inline_change(ET, text=old_part, change_type="del", change_id=next_id, author=author, date=date, run_properties=run_properties))
            next_id += 1
        if tag in {"replace", "insert"} and new_part:
            paragraph.append(_docx_inline_change(ET, text=new_part, change_type="ins", change_id=next_id, author=author, date=date, run_properties=run_properties))
            next_id += 1
    return next_id


def _replace_paragraph_with_minimal_redline(ET, paragraph, old_text: str, new_text: str, *, start_change_id: int, author: str, date: str) -> int:
    p_pr = _copy_docx_paragraph_properties(ET, paragraph)
    run_properties = _copy_docx_first_run_properties(ET, paragraph)
    for child in list(paragraph):
        paragraph.remove(child)
    if p_pr is not None:
        paragraph.append(p_pr)

    next_id = start_change_id
    old_prefix, old_body = _split_review_number_prefix(old_text)
    new_prefix, new_body = _split_review_number_prefix(new_text)
    if old_prefix and new_prefix and old_prefix != new_prefix:
        paragraph.append(_docx_inline_change(ET, text=old_prefix, change_type="del", change_id=next_id, author=author, date=date, run_properties=run_properties))
        next_id += 1
        paragraph.append(_docx_inline_change(ET, text=new_prefix, change_type="ins", change_id=next_id, author=author, date=date, run_properties=run_properties))
        next_id += 1
        return _append_minimal_redline_text_runs(
            ET,
            paragraph,
            old_body,
            new_body,
            next_id=next_id,
            author=author,
            date=date,
            run_properties=run_properties,
        )
    return _append_minimal_redline_text_runs(
        ET,
        paragraph,
        old_text,
        new_text,
        next_id=next_id,
        author=author,
        date=date,
        run_properties=run_properties,
    )


def _tracked_review_paragraph_from_template(ET, *, template=None, text: str, change_type: str, change_id: int, author: str, date: str):
    paragraph = ET.Element(f"{{{_DOCX_W_NS}}}p")
    p_pr = _copy_docx_paragraph_properties(ET, template) if template is not None else None
    run_properties = _copy_docx_first_run_properties(ET, template)
    if p_pr is not None:
        paragraph.append(p_pr)
    paragraph.append(_docx_inline_change(ET, text=text, change_type=change_type, change_id=change_id, author=author, date=date, run_properties=run_properties))
    return paragraph


def _redline_docx_paragraph_span_to_text(
    docx_bytes: bytes,
    target_text: str,
    replacement_text: str,
    *,
    report_label: str,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：DOCX 缺少正文 XML"}
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：正文 XML 无法解析"}
    body = root.find("w:body", _DOCX_NS)
    if body is None:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：未找到正文 body"}

    span_nodes, missing_reason = _find_docx_paragraph_span_by_text(root, target_text)
    if not span_nodes:
        return docx_bytes, {"status": "not_found", "message": f"{report_label}：未定位到招募正文段落，未自动替换", "reason": missing_reason}

    old_rows = _review_sync_line_rows(target_text)
    new_rows = _review_sync_line_rows(replacement_text)
    old_lines = [row["line"].strip() for row in old_rows]
    new_lines = [row["line"].strip() for row in new_rows]
    old_compare = [row["compare"] for row in old_rows]
    new_compare = [row["compare"] for row in new_rows]
    if old_compare == new_compare and not _review_lines_have_number_prefix_redlines(old_lines, new_lines):
        return docx_bytes, {"status": "unchanged", "message": f"{report_label}：内容已一致，未新增系统红线"}

    children = list(body)
    try:
        start_index = children.index(span_nodes[0])
    except ValueError:
        return docx_bytes, {"status": "not_found", "message": f"{report_label}：目标段落位置失效，未自动替换"}
    span_set = set(span_nodes)
    for node in span_nodes:
        try:
            body.remove(node)
        except ValueError:
            pass

    now = datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
    author = _normalize_review_author(review_author)
    next_id = _next_docx_change_id(root)
    replacements = []
    matcher = difflib.SequenceMatcher(None, old_compare, new_compare, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(i2 - i1):
                paragraph = span_nodes[i1 + offset]
                old_line = old_lines[i1 + offset]
                new_line = new_lines[j1 + offset]
                if _review_lines_need_number_prefix_redline(old_line, new_line):
                    next_id = _replace_paragraph_with_minimal_redline(
                        ET,
                        paragraph,
                        old_line,
                        new_line,
                        start_change_id=next_id,
                        author=author,
                        date=now,
                    )
                replacements.append(paragraph)
            continue
        if tag == "replace" and (i2 - i1) == (j2 - j1):
            for offset in range(i2 - i1):
                paragraph = span_nodes[i1 + offset]
                next_id = _replace_paragraph_with_minimal_redline(
                    ET,
                    paragraph,
                    old_lines[i1 + offset],
                    new_lines[j1 + offset],
                    start_change_id=next_id,
                    author=author,
                    date=now,
                )
                replacements.append(paragraph)
            continue
        template = span_nodes[i1] if i1 < len(span_nodes) else (span_nodes[i1 - 1] if i1 > 0 else None)
        for old_line in old_lines[i1:i2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=old_line, change_type="del", change_id=next_id, author=author, date=now))
            next_id += 1
        for new_line in new_lines[j1:j2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=new_line, change_type="ins", change_id=next_id, author=author, date=now))
            next_id += 1

    for offset, node in enumerate(replacements):
        body.insert(start_index + offset, node)

    final_bytes = _repack_docx_parts(
        docx_bytes,
        {"word/document.xml": _serialize_docx_document_xml(ET, root, xml_text)},
    )
    return final_bytes, {
        "status": "redlined",
        "message": f"{report_label}：已按最小文本差异写入红线",
        "old_lines": len(old_lines),
        "new_lines": len(new_lines),
        "changed_paragraphs": max(len(old_lines), len(new_lines)),
    }


def _redline_docx_paragraph_nodes_in_parent(
    ET,
    root,
    parent,
    span_nodes: list,
    replacement_text: str,
    *,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bool, dict]:
    span_nodes = [node for node in span_nodes or [] if _docx_local_name(getattr(node, "tag", "")) == "p"]
    new_rows = _review_sync_line_rows(replacement_text)
    new_lines = [row["line"].strip() for row in new_rows]
    new_compare = [row["compare"] for row in new_rows]
    if not span_nodes or not new_compare:
        return False, {"status": "skipped", "reason": "empty_paragraphs_or_replacement"}

    old_lines = [_effective_docx_text_from_node(node).strip() for node in span_nodes]
    old_lines = [line for line in old_lines if line]
    old_compare = [
        _normalize_contract_prospectus_compare_text(_strip_review_number_prefix(line))
        for line in old_lines
    ]
    if old_compare == new_compare and not _review_lines_have_number_prefix_redlines(old_lines, new_lines):
        return False, {"status": "unchanged"}

    children = list(parent)
    try:
        start_index = children.index(span_nodes[0])
    except ValueError:
        return False, {"status": "not_found", "reason": "paragraph_position_lost"}

    for node in span_nodes:
        _remove_docx_child(parent, node)

    now = datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
    author = _normalize_review_author(review_author)
    next_id = _next_docx_change_id(root)
    replacements = []
    matcher = difflib.SequenceMatcher(None, old_compare, new_compare, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(i2 - i1):
                paragraph = span_nodes[i1 + offset]
                old_line = old_lines[i1 + offset]
                new_line = new_lines[j1 + offset]
                if _review_lines_need_number_prefix_redline(old_line, new_line):
                    next_id = _replace_paragraph_with_minimal_redline(
                        ET,
                        paragraph,
                        old_line,
                        new_line,
                        start_change_id=next_id,
                        author=author,
                        date=now,
                    )
                replacements.append(paragraph)
            continue
        if tag == "replace" and (i2 - i1) == (j2 - j1):
            for offset in range(i2 - i1):
                paragraph = span_nodes[i1 + offset]
                next_id = _replace_paragraph_with_minimal_redline(
                    ET,
                    paragraph,
                    old_lines[i1 + offset],
                    new_lines[j1 + offset],
                    start_change_id=next_id,
                    author=author,
                    date=now,
                )
                replacements.append(paragraph)
            continue
        template = span_nodes[i1] if i1 < len(span_nodes) else (span_nodes[i1 - 1] if i1 > 0 else None)
        for old_line in old_lines[i1:i2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=old_line, change_type="del", change_id=next_id, author=author, date=now))
            next_id += 1
        for new_line in new_lines[j1:j2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=new_line, change_type="ins", change_id=next_id, author=author, date=now))
            next_id += 1

    for offset, node in enumerate(replacements):
        parent.insert(start_index + offset, node)
    return True, {
        "status": "redlined",
        "old_lines": len(old_lines),
        "new_lines": len(new_lines),
        "changed_paragraphs": max(len(old_lines), len(new_lines)),
    }


def _detect_product_summary_share_class(filename: str = "", text: str = "") -> str:
    combined = f"{filename}\n{text}"
    if re.search(r"(?:C类|（C|C类份额)", combined, flags=re.IGNORECASE):
        return "C"
    if re.search(r"(?:A类|（A|A类份额)", combined, flags=re.IGNORECASE):
        return "A"
    return "A"


def _product_summary_model_for_share(render_models: list[dict], share_class: str) -> dict:
    share_class = (share_class or "A").upper()
    for model in render_models or []:
        if str(model.get("share_class") or "").upper() == share_class:
            return model
    return (render_models or [{}])[0]


def _product_summary_prospectus_driven_table_values(render_model: dict) -> dict[str, list[str]]:
    target_labels = {"投资范围", "风险收益特征", "其他费用"}
    rows_by_label: dict[str, list[str]] = {}

    def collect(rows):
        for row in rows or []:
            if not row:
                continue
            label = str(row[0] or "").strip()
            if label in target_labels:
                rows_by_label[label] = [str(cell or "") for cell in row[1:]]

    for section in (render_model or {}).get("sections", []):
        if section.get("type") == "table":
            collect(section.get("rows"))
        for subsection in section.get("subsections", []):
            if subsection.get("type") == "table":
                collect(subsection.get("rows"))
    return rows_by_label


def _product_summary_model_risk_disclosure_text(render_model: dict) -> str:
    for section in (render_model or {}).get("sections", []):
        if section.get("title") != "四、风险揭示与重要提示":
            continue
        for subsection in section.get("subsections", []):
            if subsection.get("title") == "（一）风险揭示":
                return str(subsection.get("content") or "").strip()
    return ""


def _redline_product_summary_table_rows_from_model(
    docx_bytes: bytes,
    render_model: dict,
    *,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    import xml.etree.ElementTree as ET

    expected_rows = _product_summary_prospectus_driven_table_values(render_model)
    report = {
        "status": "completed",
        "checked_fields": sorted(expected_rows),
        "updated_fields": [],
        "unchanged_fields": [],
        "skipped": [],
    }
    if not expected_rows:
        report["status"] = "skipped"
        report["message"] = "未生成可同步的产品资料概要字段"
        return docx_bytes, report

    _register_docx_xml_namespaces(ET)
    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        report["status"] = "skipped"
        report["message"] = "产品资料概要 DOCX 缺少正文 XML"
        return docx_bytes, report
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        report["status"] = "skipped"
        report["message"] = "产品资料概要正文 XML 无法解析"
        return docx_bytes, report

    matched_labels: set[str] = set()
    changed = False
    for row in root.findall(".//w:tr", _DOCX_NS):
        cells = row.findall("./w:tc", _DOCX_NS)
        if len(cells) < 2:
            continue
        label = _effective_docx_text_from_node(cells[0]).strip()
        if label not in expected_rows:
            continue
        matched_labels.add(label)
        expected_values = expected_rows[label]
        row_changed = False
        for offset, expected_value in enumerate(expected_values, start=1):
            if offset >= len(cells):
                report["skipped"].append({"field": label, "reason": "target_cell_missing", "cell_index": offset + 1})
                continue
            paragraphs = cells[offset].findall("./w:p", _DOCX_NS)
            if not paragraphs:
                paragraph = ET.Element(f"{{{_DOCX_W_NS}}}p")
                cells[offset].append(paragraph)
                paragraphs = [paragraph]
            cell_changed, cell_report = _redline_docx_paragraph_nodes_in_parent(
                ET,
                root,
                cells[offset],
                paragraphs,
                expected_value,
                review_author=review_author,
            )
            row_changed = row_changed or cell_changed
            changed = changed or cell_changed
            if cell_report.get("status") not in {"redlined", "unchanged"}:
                report["skipped"].append({"field": label, **cell_report})
        if row_changed:
            report["updated_fields"].append(label)
        else:
            report["unchanged_fields"].append(label)

    for label in sorted(set(expected_rows) - matched_labels):
        report["skipped"].append({"field": label, "reason": "row_not_found"})

    if changed:
        report["status"] = "redlined"
        final_bytes = _repack_docx_parts(
            docx_bytes,
            {"word/document.xml": _serialize_docx_document_xml(ET, root, xml_text)},
        )
        return final_bytes, report
    report["status"] = "unchanged"
    return docx_bytes, report


def _sync_review_product_summary_from_prospectus_redlines(
    product_summary_docx: bytes,
    prospectus_docx: bytes,
    *,
    form_data: dict | None = None,
    filename: str = "",
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    prospectus_text = _docx_effective_text_from_bytes(prospectus_docx)
    if not prospectus_text:
        return product_summary_docx, {
            "status": "skipped",
            "message": "未能从最终招募说明书提取文本，产品资料概要未同步",
        }
    try:
        current_product_text = _docx_effective_text_from_bytes(product_summary_docx)
        share_class = _detect_product_summary_share_class(filename, current_product_text)
        bundle = product_summary_engine.generate_bundle({**(form_data or {}), "PROSPECTUS_TEXT": prospectus_text})
        render_model = _product_summary_model_for_share(bundle.get("render_models") or [], share_class)
        current_docx, table_report = _redline_product_summary_table_rows_from_model(
            product_summary_docx,
            render_model,
            review_author=review_author,
        )

        risk_text = _product_summary_model_risk_disclosure_text(render_model)
        risk_report = {"status": "skipped", "message": "未生成风险揭示同步文本"}
        if risk_text:
            current_docx, risk_report = _replace_docx_section_content_with_redline(
                current_docx,
                is_heading=lambda value: _normalize_review_text(value) == _normalize_review_text("（一）风险揭示"),
                is_stop_heading=lambda value: _normalize_review_text(value) in {
                    _normalize_review_text("（二）重要提示"),
                    _normalize_review_text("五、其他资料查询方式"),
                    _normalize_review_text("六、其他情况说明"),
                },
                new_text=risk_text,
                report_label="产品资料概要风险揭示",
                review_author=review_author,
            )

        updated_fields = list(table_report.get("updated_fields") or [])
        if risk_report.get("status") == "redlined":
            updated_fields.append("风险揭示")
        status = "synced_from_prospectus" if updated_fields else "unchanged"
        return current_docx, {
            "status": status,
            "message": "已根据应用审核决定后的招募说明书同步产品资料概要红线" if updated_fields else "产品资料概要与最终招募说明书派生内容已一致",
            "share_class": share_class,
            "updated_fields": updated_fields,
            "table_sync": table_report,
            "risk_sync": risk_report,
        }
    except Exception as exc:
        logger.exception("Failed to sync product summary from final prospectus")
        return product_summary_docx, {
            "status": "skipped",
            "message": "产品资料概要根据最终招募说明书同步失败",
            "error": str(exc),
        }


_LINKED_PROSPECTUS_TEMPLATE_SYNC_RULES = (
    ("第二部分 释义", "第一部分 释义", "definition_reference"),
    ("第十一部分 基金的投资", "第八部分 基金的投资", "reference_overrides"),
    ("第十二部分 基金的财产", "第九部分 基金的财产", "reference_overrides"),
    ("第十四部分 基金资产估值", "第十部分 基金资产估值", "reference_overrides"),
    ("第十五部分 基金费用与税收", "第十一部分 费用与税收", "reference_overrides"),
    ("第十六部分 基金的收益与分配", "第十二部分 基金的收益与分配", "reference_overrides"),
    ("第十七部分 基金的会计与审计", "第十三部分 基金的会计与审计", "reference_overrides"),
    ("第十八部分 基金的信息披露", "第十四部分 基金的信息披露", "reference_overrides"),
    ("第二十部分 基金合同的变更、终止与基金财产的清算", "第十六部分 基金合同的变更、终止与基金财产的清算", "reference_overrides"),
)


def _build_review_prospectus_template_sync_rules() -> list[dict]:
    rules = []
    for contract_locator, prospectus_locator, postprocess in _LINKED_PROSPECTUS_TEMPLATE_SYNC_RULES:
        rules.append({
            "source": "prospectus_template",
            "sheet_name": "招募说明书模板",
            "contract_chapter": contract_locator,
            "prospectus_chapter": prospectus_locator,
            "contract_locator": contract_locator,
            "prospectus_locator": prospectus_locator,
            "contract_section_name": contract_locator,
            "prospectus_section_name": prospectus_locator,
            "relation": "模板取自基金合同",
            "consistency": "模板来源同步",
            "detail": "招募说明书模板中的引用基金合同内容，导出时先执行招募说明书后处理再写入红线。",
            "postprocess": postprocess,
        })
    return rules


def _review_rule_is_prospectus_template_source(rule: dict) -> bool:
    return (rule or {}).get("source") == "prospectus_template"


def _postprocess_review_prospectus_template_source(text: str, rule: dict) -> str:
    mode = (rule or {}).get("postprocess") or ""
    source = (text or "").strip()
    if not source:
        return ""
    if mode == "definition_reference":
        return prospectus_engine._definition_reference_for_prospectus(source, {})
    if mode == "reference_overrides":
        source = ProspectusEngine._apply_prospectus_reference_overrides(source)
        return ProspectusEngine._cleanup_prospectus(source)
    return source


def _sync_review_direct_contract_sources_redlines(
    contract_docx: bytes,
    prospectus_docx: bytes,
    *,
    rules: list[dict] | None = None,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    contract_text = _docx_effective_text_from_bytes(contract_docx)
    prospectus_text = _docx_effective_text_from_bytes(prospectus_docx)
    contract_sections = _split_contract_sections(contract_text)
    prospectus_sections = _split_prospectus_sections(prospectus_text)
    candidate_rules = list(rules) if rules is not None else (
        _build_review_cross_rules()
        + _build_review_prospectus_template_sync_rules()
    )

    report = {
        "status": "completed",
        "message": "已检查招募说明书正文中可确定为直接对应的合同来源内容",
        "checked_count": 0,
        "updated_count": 0,
        "unchanged_count": 0,
        "template_checked_count": 0,
        "template_updated_count": 0,
        "template_unchanged_count": 0,
        "skipped": [],
        "updated": [],
    }
    current_prospectus = prospectus_docx

    for rule in candidate_rules:
        is_template_source = _review_rule_is_prospectus_template_source(rule)
        ok, reason = (True, "") if is_template_source else _review_cross_rule_is_direct_sync_candidate(rule)
        label = f"{rule.get('contract_locator') or rule.get('contract_chapter') or ''} → {rule.get('prospectus_locator') or rule.get('prospectus_chapter') or ''}".strip()
        if not ok:
            report["skipped"].append({"label": label, "reason": reason})
            continue
        report["checked_count"] += 1
        if is_template_source:
            report["template_checked_count"] += 1

        prospectus_text = _docx_effective_text_from_bytes(current_prospectus)
        prospectus_sections = _split_prospectus_sections(prospectus_text)
        contract_target = _locate_review_rule_target(
            contract_sections,
            rule.get("contract_locator") or rule.get("contract_chapter") or "",
            rule.get("contract_section_name") or rule.get("contract_chapter") or "",
        )
        if not contract_target.get("matched"):
            contract_target = _fallback_review_rule_target_from_text(
                contract_text,
                rule.get("contract_locator") or rule.get("contract_chapter") or "",
                rule.get("contract_section_name") or rule.get("contract_chapter") or "",
            )
        prospectus_target = _locate_review_rule_target(
            prospectus_sections,
            rule.get("prospectus_locator") or rule.get("prospectus_chapter") or "",
            rule.get("prospectus_section_name") or rule.get("prospectus_chapter") or "",
        )
        if not prospectus_target.get("matched"):
            prospectus_target = _fallback_review_rule_target_from_text(
                prospectus_text,
                rule.get("prospectus_locator") or rule.get("prospectus_chapter") or "",
                rule.get("prospectus_section_name") or rule.get("prospectus_chapter") or "",
            )
        if not contract_target.get("matched") or not prospectus_target.get("matched"):
            report["skipped"].append({
                "label": label,
                "reason": "locator_not_matched",
                "contract_matched": bool(contract_target.get("matched")),
                "prospectus_matched": bool(prospectus_target.get("matched")),
            })
            continue

        contract_source = (contract_target.get("body_text") or contract_target.get("text") or "").strip()
        prospectus_target_text = (prospectus_target.get("body_text") or prospectus_target.get("text") or "").strip()
        if not contract_source or not prospectus_target_text:
            report["skipped"].append({"label": label, "reason": "empty_source_or_target", "source": rule.get("source", "")})
            continue
        replacement_text = (
            _postprocess_review_prospectus_template_source(contract_source, rule)
            if is_template_source
            else contract_source
        )
        if not replacement_text:
            report["skipped"].append({"label": label, "reason": "empty_postprocessed_source", "source": rule.get("source", "")})
            continue
        minimal_ok, minimal_reason = _review_sync_minimal_change_guard(replacement_text, prospectus_target_text)
        if not minimal_ok:
            report["skipped"].append({"label": label, "reason": minimal_reason, "source": rule.get("source", "")})
            continue

        current_prospectus, apply_report = _redline_docx_paragraph_span_to_text(
            current_prospectus,
            prospectus_target_text,
            replacement_text,
            report_label=label or "招募说明书合同来源正文",
            review_author=review_author,
        )
        if apply_report.get("status") == "redlined":
            report["updated_count"] += 1
            if is_template_source:
                report["template_updated_count"] += 1
            report["updated"].append({"label": label, "source": rule.get("source", ""), "postprocess": rule.get("postprocess", ""), **apply_report})
        elif apply_report.get("status") == "unchanged":
            report["unchanged_count"] += 1
            if is_template_source:
                report["template_unchanged_count"] += 1
        else:
            report["skipped"].append({"label": label, "reason": apply_report.get("reason") or apply_report.get("message", ""), "source": rule.get("source", "")})

    return current_prospectus, report


def _build_revision_workbench_sync_report(
    *,
    contract_bytes: bytes,
    prospectus_bytes: bytes,
    product_summary_bytes: bytes | None,
    form_data: dict,
    document_reports: dict,
    summary_sync_report: dict | None = None,
    direct_sync_report: dict | None = None,
    product_summary_sync_report: dict | None = None,
    custody_sync_report: dict | None = None,
    product_summary_outputs: list[dict] | None = None,
) -> dict:
    summary_report = (summary_sync_report or {}).get("contract_summary") or {
        "status": "skipped",
        "message": "未能自动重算基金合同摘要",
    }

    prospectus_report = {
        "status": "report_only",
        "message": "已保留应用审核决定后的招募说明书红线；无法确定为直接对应的合同来源片段未做猜测替换",
    }
    if summary_sync_report and summary_sync_report.get("prospectus_contract_summary"):
        prospectus_report["contract_summary_apply"] = summary_sync_report["prospectus_contract_summary"]
    if summary_report.get("status") == "generated":
        prospectus_report["contract_summary_source"] = "已生成新的基金合同摘要，可用于招募说明书内合同摘要同步"
    if direct_sync_report:
        prospectus_report["direct_contract_sources"] = direct_sync_report
    if custody_sync_report:
        prospectus_report["custody_summary_apply"] = custody_sync_report

    product_summary_report = {
        "status": "skipped",
        "message": "产品资料概要未上传，未生成该终稿",
    }
    product_summaries = []
    if product_summary_outputs:
        for output in product_summary_outputs:
            item_report = dict(output.get("sync_report") or {})
            item_report["filename"] = output.get("filename", "")
            item_report["output_name"] = output.get("output_name", "")
            item_report["share_class"] = output.get("share_class", "")
            product_summaries.append(item_report)
        product_summary_report = product_summaries[0] if len(product_summaries) == 1 else {
            "status": "batch_processed",
            "message": f"已处理 {len(product_summaries)} 份产品资料概要",
            "items": product_summaries,
        }
        if form_data:
            product_summary_report["form_data_keys"] = sorted(str(key) for key in form_data.keys())[:30]
    elif product_summary_bytes:
        product_summary_report = {
            "status": "kept_redline",
            "message": "已输出应用审核决定后的产品资料概要红线；本版未在缺少完整表单数据时强制重算",
        }
        if product_summary_sync_report:
            product_summary_report = product_summary_sync_report
        if form_data:
            product_summary_report["form_data_keys"] = sorted(str(key) for key in form_data.keys())[:30]

    return {
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
        "contract_summary": summary_report,
        "prospectus": prospectus_report,
        "custody_agreement": custody_sync_report or {"status": "skipped", "message": "托管协议未上传，未同步招募说明书托管协议摘要"},
        "product_summary": product_summary_report,
        "product_summaries": product_summaries,
        "documents": document_reports,
    }


def _revision_product_summary_share_class(filename: str, docx_bytes: bytes) -> str:
    text = _docx_effective_text_from_bytes(docx_bytes)
    share = _detect_product_summary_share_class(filename, text)
    return share if share in {"A", "C"} else ""


def _revision_product_summary_document_kind(filename: str, docx_bytes: bytes, index: int, total: int, used: set[str]) -> tuple[str, str]:
    share = _revision_product_summary_share_class(filename, docx_bytes)
    if total <= 1:
        return "product_summary", share
    suffix = share or str(index)
    candidate = f"product_summary_{suffix}"
    if candidate in used:
        candidate = f"product_summary_{suffix}_{index}"
    used.add(candidate)
    return candidate, share


def _revision_product_summary_output_name(entry: dict, index: int, total: int) -> str:
    share = str(entry.get("share_class") or "").strip().upper()
    if total <= 1 and not share:
        return "产品资料概要.docx"
    if total <= 1:
        return "产品资料概要.docx"
    if share in {"A", "C"}:
        return f"产品资料概要-{share}类.docx"
    return f"产品资料概要-{index}.docx"


@app.route("/api/review/upload", methods=["POST"])
def api_review_upload():
    try:
        contract_file = request.files.get("contract_docx") or request.files.get("contract")
        prospectus_file = request.files.get("prospectus_docx") or request.files.get("prospectus")
        if not contract_file:
            return jsonify({"success": False, "error": "未上传合同文件"}), 400

        from docx import Document

        contract_doc = Document(_validate_docx_upload(contract_file, contract_file.filename or "合同.docx"))
        contract_text = _extract_review_docx_text(contract_doc)

        contract_filename = contract_file.filename or "合同.docx"
        fund_name = _infer_review_fund_name(contract_text, contract_filename) or \
                    _fallback_review_filename({"contract_filename": contract_filename})
        contract_sections = _split_contract_sections(contract_text)

        data = {
            "contract_text": contract_text,
            "contract_filename": contract_filename,
            "contract_sections": contract_sections,
            "fund_name": fund_name,
            "prospectus_text": None,
            "prospectus_filename": None,
            "prospectus_sections": [],
        }

        if prospectus_file:
            prospectus_doc = Document(_validate_docx_upload(prospectus_file, prospectus_file.filename or "招募说明书.docx"))
            prospectus_text = _extract_review_docx_text(prospectus_doc)

            data["prospectus_text"] = prospectus_text
            data["prospectus_filename"] = prospectus_file.filename or "招募说明书.docx"
            data["prospectus_sections"] = _split_prospectus_sections(prospectus_text, prospectus_doc)

        job_id = _store_review_job(data)

        return jsonify({
            "success": True,
            "job_id": job_id,
            "review_job_id": job_id,
            "fund_name": fund_name,
            "contract_filename": contract_filename,
            "contract_section_count": len(contract_sections),
            "contract_sections": len(contract_sections),
            "prospectus_filename": data["prospectus_filename"],
            "prospectus_section_count": len(data["prospectus_sections"])
            if data["prospectus_text"] else 0,
            "prospectus_sections": len(data["prospectus_sections"])
            if data["prospectus_text"] else 0,
        })
    except Exception as e:
        logger.exception("Review upload failed")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/revision_workbench/upload", methods=["POST"])
def api_revision_workbench_upload():
    try:
        file_specs = []
        for document_kind, field_names in [
            ("contract", ("contract_revision_docx", "contract_docx", "contract")),
            ("prospectus", ("prospectus_revision_docx", "prospectus_docx", "prospectus")),
            ("custody_agreement", ("custody_agreement_revision_docx", "custody_agreement_docx", "custody_agreement", "custody")),
        ]:
            file_obj = next((request.files.get(field_name) for field_name in field_names if request.files.get(field_name)), None)
            if file_obj:
                file_specs.append((document_kind, file_obj))

        product_summary_files = []
        for field_name in ("product_summary_revision_docx", "product_summary_docx", "product_summary"):
            product_summary_files.extend([file_obj for file_obj in request.files.getlist(field_name) if file_obj])
        documents = []
        revision_items = []
        uploaded_originals = {}
        product_summary_entries = []

        for document_kind, file_obj in file_specs:
            if not file_obj:
                continue
            filename = file_obj.filename or f"{document_kind}.docx"
            stream = _validate_docx_upload(file_obj, filename)
            raw_bytes = stream.getvalue()
            document = _build_revision_document_view(raw_bytes, document_kind=document_kind, filename=filename)
            documents.append(document)
            revision_items.extend(document.get("revisions") or [])
            uploaded_originals[document_kind] = {"filename": filename, "bytes": raw_bytes}

        used_product_kinds = set()
        for index, file_obj in enumerate(product_summary_files, start=1):
            filename = file_obj.filename or f"product_summary_{index}.docx"
            stream = _validate_docx_upload(file_obj, filename)
            raw_bytes = stream.getvalue()
            document_kind, share_class = _revision_product_summary_document_kind(
                filename,
                raw_bytes,
                index,
                len(product_summary_files),
                used_product_kinds,
            )
            document = _build_revision_document_view(raw_bytes, document_kind=document_kind, filename=filename)
            documents.append(document)
            revision_items.extend(document.get("revisions") or [])
            uploaded_originals[document_kind] = {"filename": filename, "bytes": raw_bytes}
            product_summary_entries.append({
                "document_kind": document_kind,
                "filename": filename,
                "bytes": raw_bytes,
                "share_class": share_class,
            })

        if not documents:
            return jsonify({"success": False, "error": "请至少上传一份 DOCX 文件"}), 400

        revision_counts = _count_revision_items(revision_items)
        data = {
            "documents": documents,
            "revision_items": revision_items,
            "revision_counts": revision_counts,
            "revision_total": len(revision_items),
            "uploaded_originals": uploaded_originals,
            "file_bytes": {key: value.get("bytes", b"") for key, value in uploaded_originals.items()},
            "filenames": {key: value.get("filename", "") for key, value in uploaded_originals.items()},
            "product_summary_entries": product_summary_entries,
            "created_at": datetime.now().isoformat(timespec="seconds"),
        }
        _revision_workbench_store["data"] = data

        return jsonify({
            "success": True,
            "documents": documents,
            "revision_items": revision_items,
            "revision_counts": revision_counts,
            "revision_total": len(revision_items),
            "scope_labels": _REVISION_SCOPE_LABELS,
            "impact_targets": _REVISION_IMPACT_TARGETS,
        })
    except Exception as e:
        logger.exception("Revision workbench upload failed")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/revision_workbench/export_final", methods=["POST"])
def api_revision_workbench_export_final():
    """根据审核工作台决策导出红线终稿 ZIP，不覆盖原文件。"""
    data = _revision_workbench_store.get("data") or {}
    file_bytes = data.get("file_bytes") or {key: value.get("bytes") for key, value in (data.get("uploaded_originals") or {}).items()}
    if not file_bytes:
        return jsonify({"error": "请先上传修订版 DOCX 文件"}), 400
    if not file_bytes.get("contract") or not file_bytes.get("prospectus"):
        return jsonify({"error": "生成红线终稿至少需要上传基金合同和招募说明书"}), 400

    payload = request.get_json(force=True, silent=True) or {}
    decisions = payload.get("decisions") or {}
    revision_edits = _normalize_review_revision_edits(payload.get("revision_edits") or {})
    review_author = _normalize_review_author(payload.get("review_author"))
    form_data = payload.get("form_data") or {}
    session_operations = payload.get("session_operations") or []
    documents = data.get("documents") or []
    revisions = [
        revision
        for document in documents
        for revision in (document or {}).get("revisions", [])
        if revision.get("id")
    ]
    pending_ids = [
        revision["id"]
        for revision in revisions
        if _normalize_review_export_decision(decisions.get(revision["id"])) == "pending"
    ]
    if pending_ids:
        return jsonify({
            "error": f"仍有 {len(pending_ids)} 条待定修订，请先全部标记为接受或拒绝后再生成红线终稿",
            "pending_count": len(pending_ids),
            "pending_ids": pending_ids[:50],
        }), 400

    filenames = data.get("filenames") or {key: value.get("filename", "") for key, value in (data.get("uploaded_originals") or {}).items()}
    document_reports = {}
    try:
        contract_docx, contract_report = _apply_review_revision_decisions_to_docx(
            file_bytes["contract"],
            document_kind="contract",
            filename=filenames.get("contract", ""),
            decisions=decisions,
            revision_edits=revision_edits,
            review_author=review_author,
        )
        prospectus_docx, prospectus_report = _apply_review_revision_decisions_to_docx(
            file_bytes["prospectus"],
            document_kind="prospectus",
            filename=filenames.get("prospectus", ""),
            decisions=decisions,
            revision_edits=revision_edits,
            review_author=review_author,
        )
        document_reports["contract"] = contract_report
        document_reports["prospectus"] = prospectus_report
        contract_docx, prospectus_docx, summary_sync_report = _sync_review_contract_summary_redlines(
            contract_docx,
            prospectus_docx,
            review_author=review_author,
        )
        prospectus_docx, direct_sync_report = _sync_review_direct_contract_sources_redlines(
            contract_docx,
            prospectus_docx,
            review_author=review_author,
        )

        custody_docx = None
        custody_sync_report = None
        if file_bytes.get("custody_agreement"):
            custody_docx, custody_report = _apply_review_revision_decisions_to_docx(
                file_bytes["custody_agreement"],
                document_kind="custody_agreement",
                filename=filenames.get("custody_agreement", ""),
                decisions=decisions,
                revision_edits=revision_edits,
                review_author=review_author,
            )
            document_reports["custody_agreement"] = custody_report
            prospectus_docx, custody_sync_report = _sync_review_custody_summary_redlines(
                custody_docx,
                prospectus_docx,
                review_author=review_author,
            )

        product_summary_outputs = []
        product_summary_entries = list(data.get("product_summary_entries") or [])
        if not product_summary_entries and file_bytes.get("product_summary"):
            product_summary_entries = [{
                "document_kind": "product_summary",
                "filename": filenames.get("product_summary", ""),
                "bytes": file_bytes["product_summary"],
                "share_class": _revision_product_summary_share_class(filenames.get("product_summary", ""), file_bytes["product_summary"]),
            }]
        for index, entry in enumerate(product_summary_entries, start=1):
            document_kind = entry.get("document_kind") or "product_summary"
            filename = entry.get("filename") or filenames.get(document_kind, "")
            product_summary_docx, product_summary_report = _apply_review_revision_decisions_to_docx(
                entry.get("bytes") or file_bytes.get(document_kind, b""),
                document_kind=document_kind,
                filename=filename,
                decisions=decisions,
                revision_edits=revision_edits,
                review_author=review_author,
            )
            product_summary_report["document_kind"] = document_kind
            document_reports[document_kind] = product_summary_report
            product_summary_docx, product_summary_sync_report = _sync_review_product_summary_from_prospectus_redlines(
                product_summary_docx,
                prospectus_docx,
                form_data=form_data,
                filename=filename,
                review_author=review_author,
            )
            product_summary_outputs.append({
                "document_kind": document_kind,
                "filename": filename,
                "share_class": entry.get("share_class") or _revision_product_summary_share_class(filename, product_summary_docx),
                "bytes": product_summary_docx,
                "sync_report": product_summary_sync_report,
            })

        for index, output_item in enumerate(product_summary_outputs, start=1):
            output_item["output_name"] = _revision_product_summary_output_name(output_item, index, len(product_summary_outputs))

        product_summary_docx = product_summary_outputs[0]["bytes"] if len(product_summary_outputs) == 1 else None

        sync_report = _build_revision_workbench_sync_report(
            contract_bytes=contract_docx,
            prospectus_bytes=prospectus_docx,
            product_summary_bytes=product_summary_docx,
            form_data=form_data,
            document_reports=document_reports,
            summary_sync_report=summary_sync_report,
            direct_sync_report=direct_sync_report,
            product_summary_sync_report=product_summary_outputs[0]["sync_report"] if len(product_summary_outputs) == 1 else None,
            custody_sync_report=custody_sync_report,
            product_summary_outputs=product_summary_outputs,
        )
        sync_report["review_author"] = review_author
        applied_edit_ids = {
            revision_id
            for report in document_reports.values()
            for revision_id in (report.get("edited_revision_ids") or [])
        }

        decision_record = {
            "generated_at": sync_report["generated_at"],
            "review_author": review_author,
            "decisions": decisions,
            "revision_edits": revision_edits,
            "edited_count": len(applied_edit_ids),
            "applied_edit_ids": sorted(applied_edit_ids),
            "session_operations": session_operations,
            "documents": [
                {
                    "document_kind": (document or {}).get("document_kind", ""),
                    "filename": (document or {}).get("filename", ""),
                    "revision_count": len((document or {}).get("revisions", []) or []),
                }
                for document in documents
            ],
        }
        csv_text = _build_review_decision_csv(decisions, documents, revision_edits, applied_edit_ids)
        contract_docx = _finalize_review_docx_for_export(contract_docx)
        prospectus_docx = _finalize_review_docx_for_export(prospectus_docx)
        if custody_docx:
            custody_docx = _finalize_review_docx_for_export(custody_docx)
        for output_item in product_summary_outputs:
            output_item["bytes"] = _finalize_review_docx_for_export(output_item["bytes"])

        output = io.BytesIO()
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("基金合同.docx", contract_docx)
            zf.writestr("招募说明书.docx", prospectus_docx)
            if custody_docx:
                zf.writestr("托管协议.docx", custody_docx)
            for output_item in product_summary_outputs:
                zf.writestr(output_item.get("output_name") or "产品资料概要.docx", output_item["bytes"])
            zf.writestr("审核决策记录.json", json.dumps(decision_record, ensure_ascii=False, indent=2).encode("utf-8"))
            zf.writestr("审核决策记录.csv", "\ufeff" + csv_text)
            zf.writestr("同步报告.json", json.dumps(sync_report, ensure_ascii=False, indent=2).encode("utf-8"))
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name="审核工作台红线终稿.zip",
            mimetype="application/zip",
        )
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": f"生成红线终稿失败：{exc}"}), 500

@app.route("/api/review/get_text", methods=["GET"])
def api_review_get_text():
    data, job_id = _get_review_data()
    if not data or not data.get("contract_text"):
        return jsonify({"success": False, "error": "未上传复核文件"}), 404
    return jsonify({
        "success": True,
        "job_id": job_id,
        "contract_text": data["contract_text"],
        "fund_name": data.get("fund_name", ""),
        "contract_filename": data.get("contract_filename", ""),
        "prospectus_available": bool(data.get("prospectus_text")),
    })


@app.route("/api/review/cross_check", methods=["POST"])
def api_review_cross_check():
    req = request.get_json(silent=True) or {}
    data, job_id = _get_review_data(req)
    if not data or not data.get("contract_text"):
        return jsonify({"success": False, "error": "请先上传合同文件"}), 400

    contract_sections = data.get("contract_sections", [])
    prospectus_sections = data.get("prospectus_sections", [])
    has_prospectus = bool(data.get("prospectus_text"))
    rules = _build_review_cross_rules()
    results = []

    for rule in rules:
        contract_chapter = rule.get("contract_chapter", "")
        prospectus_chapter = rule.get("prospectus_chapter", "")
        contract_locator = rule.get("contract_locator", "")
        prospectus_locator = rule.get("prospectus_locator", "")
        relation = rule.get("relation", "")
        expected_consistency = rule.get("consistency", "")
        detail = rule.get("detail", "")
        suggestion = rule.get("suggestion", "")

        contract_target = _locate_review_rule_target(contract_sections, contract_locator)
        prospectus_target = None
        if has_prospectus:
            prospectus_target = _locate_review_rule_target(prospectus_sections, prospectus_locator)

        contract_missing = not contract_target.get("matched")
        prospectus_missing = has_prospectus and not (
            prospectus_target and prospectus_target.get("matched"))

        # Determine comparison mode
        if "完全一致" in expected_consistency:
            compare_mode = "strict"
        elif "基本对应" in expected_consistency or "招募细化" in relation or "招募补充" in detail:
            compare_mode = "descriptive"
        else:
            compare_mode = "descriptive"

        similarity = None
        diffs = None
        hunks = None
        consistency_result = "pass"
        result_detail = detail

        if _is_missing_review_locator(contract_locator):
            consistency_result = "pass"
            result_detail = detail or "合同无此章节，为招募独有内容"
        elif contract_missing:
            if "风险揭示" in contract_chapter:
                consistency_result = "pass"
                result_detail = detail or "合同通常不设独立风险揭示章节，相关风险提示在前言及招募说明书中展开"
            else:
                consistency_result = "error"
                result_detail = f'合同缺少章节: {contract_chapter}'
        elif prospectus_missing and has_prospectus:
            if _is_missing_review_locator(prospectus_locator):
                consistency_result = "pass"
            else:
                consistency_result = "warning"
                result_detail = f'招募说明书缺少章节: {prospectus_chapter}'
        elif has_prospectus and contract_target.get("text") and \
             prospectus_target and prospectus_target.get("text"):
            c_text = contract_target["text"]
            p_text = prospectus_target["text"]
            fact_mismatches = _compare_review_key_facts(c_text, p_text)

            if fact_mismatches:
                consistency_result = "error"
                result_detail = _review_fact_mismatch_message(fact_mismatches)
                norm_c = _normalize_contract_prospectus_compare_text(c_text)
                norm_p = _normalize_contract_prospectus_compare_text(p_text)
                similarity = difflib.SequenceMatcher(
                    None, norm_c, norm_p, autojunk=False).ratio() * 100
                c_lines = [l for l in c_text.split('\n') if l.strip()]
                p_lines = [l for l in p_text.split('\n') if l.strip()]
                diffs = _build_review_excerpt_pair(c_lines, p_lines, context_lines=2)
                hunks = _build_review_line_hunks(c_text, p_text, context_lines=2)
            elif compare_mode == "strict":
                norm_c = _normalize_contract_prospectus_compare_text(c_text)
                norm_p = _normalize_contract_prospectus_compare_text(p_text)
                similarity = difflib.SequenceMatcher(
                    None, norm_c, norm_p, autojunk=False).ratio() * 100
                if similarity > 98.0:
                    consistency_result = "pass"
                elif similarity > 90.0:
                    consistency_result = "warning"
                else:
                    consistency_result = "error"
            elif compare_mode == "normalized":
                norm_c = _normalize_contract_prospectus_compare_text(c_text)
                norm_p = _normalize_contract_prospectus_compare_text(p_text)
                similarity = difflib.SequenceMatcher(
                    None, norm_c, norm_p, autojunk=False).ratio() * 100
                if similarity > 95.0:
                    consistency_result = "pass"
                elif similarity > 80.0:
                    consistency_result = "warning"
                else:
                    consistency_result = "error"
            else:
                consistency_result = "pass"
                similarity = 100.0

            if compare_mode in ("strict", "normalized"):
                c_lines = [l for l in c_text.split('\n') if l.strip()]
                p_lines = [l for l in p_text.split('\n') if l.strip()]
                diffs = _build_review_excerpt_pair(c_lines, p_lines, context_lines=2)
                if consistency_result in ("error", "warning"):
                    hunks = _build_review_line_hunks(c_text, p_text, context_lines=2)
        elif not has_prospectus:
            consistency_result = "pass"
            result_detail = "未上传招募说明书，跳过比对"

        severity = "info"
        if consistency_result == "error":
            severity = "error"
        elif consistency_result == "warning":
            severity = "warning"

        contract_display_locator = _format_review_target_locator(contract_target) or contract_locator
        prospectus_display_locator = (
            _format_review_target_locator(prospectus_target) if prospectus_target else ""
        ) or prospectus_locator
        is_problem = severity in ("error", "warning")

        results.append({
            "rule_source": rule.get("source", ""),
            "rule": rule,
            "contract_chapter": contract_chapter,
            "prospectus_chapter": prospectus_chapter,
            "contract_section_name": rule.get("contract_section_name", ""),
            "prospectus_section_name": rule.get("prospectus_section_name", ""),
            "relation": relation,
            "consistency_result": consistency_result,
            "status": consistency_result,
            "severity": severity,
            "detail": result_detail,
            "message": result_detail,
            "is_problem": is_problem,
            "similarity": round(similarity, 1) if similarity is not None else None,
            "suggestion": suggestion,
            "contract_matched": contract_target.get("matched", False),
            "contract_match_method": contract_target.get("match_method", ""),
            "prospectus_matched": prospectus_target.get("matched", False)
            if prospectus_target else False,
            "contract_locator": contract_display_locator,
            "prospectus_locator": prospectus_display_locator,
            "contract_anchor_text": contract_target.get("target_heading", ""),
            "prospectus_anchor_text": prospectus_target.get("target_heading", "")
            if prospectus_target else "",
            "contract_context_excerpt": _truncate_review_excerpt(contract_target.get("text", ""), 800),
            "prospectus_context_excerpt": _truncate_review_excerpt(prospectus_target.get("text", ""), 800)
            if prospectus_target else "",
            "diffs": diffs,
            "hunks": hunks or [],
        })

    errors = sum(1 for r in results if r["severity"] == "error")
    warnings = sum(1 for r in results if r["severity"] == "warning")

    return jsonify({
        "success": True,
        "results": results,
        "total": len(results),
        "errors": errors,
        "warnings": warnings,
        "has_prospectus": has_prospectus,
    })


@app.route("/api/review/format_check", methods=["POST"])
def api_review_format_check():
    req = request.get_json(silent=True) or {}
    data, job_id = _get_review_data(req)
    if not data or not data.get("contract_text"):
        return jsonify({"success": False, "error": "请先上传合同文件"}), 400

    raw_contract_text = data["contract_text"]
    format_contract_text = _strip_contract_signing_page_text(raw_contract_text)
    contract_sections = data.get("contract_sections") or []
    if contract_sections:
        contract_text = "\n".join(
            f"{section.get('heading', '')}\n{section.get('content', '')}"
            for section in contract_sections
        )
    else:
        contract_text = format_contract_text
    contract_text = _strip_contract_signing_page_text(contract_text)
    issues = []

    def _review_line_position(text, char_index):
        text = str(text or "")
        line_no = text[:char_index].count("\n") + 1
        line_start = text.rfind("\n", 0, char_index) + 1
        return line_no, char_index - line_start + 1

    def _highlight_line_column(line, column):
        line = str(line or "")
        if not column or column < 1 or column > len(line):
            return html.escape(line, quote=False)
        index = column - 1
        return (
            html.escape(line[:index], quote=False)
            + "<mark>"
            + html.escape(line[index:index + 1], quote=False)
            + "</mark>"
            + html.escape(line[index + 1:], quote=False)
        )

    def _review_line_location(text, line_no, column=None):
        try:
            line_no = int(line_no or 0)
        except (TypeError, ValueError):
            return None
        lines = str(text or "").splitlines()
        if line_no < 1 or line_no > len(lines):
            return None
        line = lines[line_no - 1]
        location = {"ln": line_no, "text": line}
        if column:
            location["html"] = _highlight_line_column(line, column)
        return location

    def _nearby_context_line_numbers(lines, line_no, radius=1, nonblank_radius=25, nonblank_each_side=2):
        line_count = len(lines)
        selected = set(range(max(1, line_no - radius), min(line_count, line_no + radius) + 1))

        before = []
        for ln in range(line_no - 1, max(0, line_no - nonblank_radius) - 1, -1):
            if lines[ln - 1].strip():
                before.append(ln)
                if len(before) >= nonblank_each_side:
                    break

        after = []
        for ln in range(line_no + 1, min(line_count, line_no + nonblank_radius) + 1):
            if lines[ln - 1].strip():
                after.append(ln)
                if len(after) >= nonblank_each_side:
                    break

        selected.update(before)
        selected.update(after)
        return sorted(ln for ln in selected if 1 <= ln <= line_count)

    def _nearest_review_heading(lines, line_no, max_scan=80):
        heading_pattern = re.compile(
            r"^(?:第[一二三四五六七八九十百]+部分|[一二三四五六七八九十百]+、|"
            r"[（(][一二三四五六七八九十百0-9]+[）)]|\d+、)"
        )
        for ln in range(line_no - 1, max(0, line_no - max_scan) - 1, -1):
            stripped = lines[ln - 1].strip()
            if stripped and len(stripped) <= 120 and heading_pattern.match(stripped):
                return {"ln": ln, "text": stripped}
        return None

    def _attach_issue_context(issue, text, line_no=None, column=None, radius=1):
        if not line_no:
            line_no = issue.get("line")
        location = _review_line_location(text, line_no, column)
        if not location:
            return issue
        line_no = location["ln"]
        issue["line"] = line_no
        if column:
            issue["column"] = column
        issue.setdefault("text", location["text"].strip())
        lines = str(text or "").splitlines()
        context = []
        for ln in _nearby_context_line_numbers(lines, line_no, radius=radius):
            ctx = _review_line_location(text, ln, column if ln == line_no else None)
            if ctx:
                ctx["current"] = ln == line_no
                context.append(ctx)
        issue["context"] = context
        nearest_heading = _nearest_review_heading(lines, line_no)
        if nearest_heading:
            issue["section"] = nearest_heading
        return issue

    quality_type_map = {
        "UNRESOLVED_CONDITION": "unresolved_condition_tag",
        "UNRESOLVED_TEMPLATE": "unresolved_template_tag",
        "UNRESOLVED_PLACEHOLDER": "unresolved_placeholder",
        "UNRESOLVED_ALIAS": "unresolved_render_alias",
        "BAD_CHAR": "bad_character",
        "ADJACENT_DUPLICATE_TEXT": "adjacent_duplicate_text",
        "ADJACENT_DUPLICATE_HEADING": "adjacent_duplicate_heading",
        "NEAR_DUPLICATE_PARAGRAPH": "near_duplicate_paragraph",
        "REPEATED_NUMBER": "repeated_number",
        "SKIPPED_NUMBER": "skipped_number",
        "DANGLING_PUNCTUATION": "dangling_punctuation",
    }
    quality = _generation_quality_check(format_contract_text, "基金合同复核")
    for quality_issue in quality.get("issues", []):
        code = quality_issue.get("code", "")
        issue = {
            "type": quality_type_map.get(code, code.lower() if code else "quality_issue"),
            "severity": quality_issue.get("severity", "warning"),
            "detail": quality_issue.get("message", "条件语句附近文本疑似异常"),
            "location": f'第{quality_issue.get("line")}行' if quality_issue.get("line") else "",
            "line": quality_issue.get("line"),
            "text": quality_issue.get("snippet", ""),
            "suggestion": "请核对条件语句、模板变量或相邻文本是否被错误保留、重复拼接。",
        }
        _attach_issue_context(issue, format_contract_text)
        related_line = quality_issue.get("related_line")
        if related_line:
            issue["locations"] = [
                loc for loc in (
                    _review_line_location(format_contract_text, related_line),
                    _review_line_location(format_contract_text, quality_issue.get("line")),
                ) if loc
            ]
        issues.append(issue)

    def is_toc_or_heading_spacing_line(line):
        stripped = (line or "").strip()
        if not stripped:
            return True
        if "\t" in stripped and re.search(r'\t+\d+\s*$', stripped):
            return True
        if re.sub(r'\s+', '', stripped) == "目录":
            return True
        if re.match(r'^第[一二三四五六七八九十百]+部分\s+\S+', stripped):
            return True
        return False

    # 1. Heading sequence continuity
    RE_PART = re.compile(r'^第([一二三四五六七八九十百]+)部分', re.MULTILINE)
    parts = [(m.group(1), _cn_numeral_to_int(m.group(1))) for m in RE_PART.finditer(contract_text)]
    prev_num = 0
    for cn, num in parts:
        if num is None:
            continue
        if prev_num > 0 and num != prev_num + 1:
            issues.append({
                "type": "heading_sequence",
                "severity": "warning",
                "detail": f'章节编号跳跃: 第{prev_num}部分 → 第{num}部分',
                "location": f'第{cn}部分',
            })
        prev_num = num

    # 2. Chinese numbered list continuity
    RE_CN_NUM = re.compile(r'^([一二三四五六七八九十]+)[、．.]', re.MULTILINE)
    cn_nums = [(m.group(1), _cn_numeral_to_int(m.group(1)))
               for m in RE_CN_NUM.finditer(contract_text)]
    prev_cn = 0
    seen = set()
    for cn, num in cn_nums:
        if num is None or num in seen:
            continue
        seen.add(num)
        if prev_cn > 0 and num != prev_cn + 1:
            issues.append({
                "type": "cn_numbering_sequence",
                "severity": "info",
                "detail": f'中文编号跳跃: {prev_cn} → {num}',
                "location": f'{cn}、',
            })
        prev_cn = num

    # 3. Bracket pairing
    bracket_map = {"（": "）", "(": ")", "《": "》", "【": "】", "「": "」"}
    bracket_stack = []
    for i, ch in enumerate(contract_text):
        if ch in bracket_map:
            bracket_stack.append((ch, i))
        elif ch in bracket_map.values():
            if ch == "）":
                line_start = contract_text.rfind("\n", 0, i) + 1
                prefix = contract_text[line_start:i].strip()
                if re.match(r'^(?:\d+|[一二三四五六七八九十百]+)$', prefix):
                    continue
            if bracket_stack:
                open_ch, open_pos = bracket_stack.pop()
                expected_close = bracket_map.get(open_ch)
                if expected_close != ch:
                    line_num, column = _review_line_position(contract_text, i)
                    open_line, open_column = _review_line_position(contract_text, open_pos)
                    issue = {
                        "type": "bracket_mismatch",
                        "severity": "error",
                        "detail": f'括号不匹配: 第{open_line}行第{open_column}列的 {open_ch} 期望 {expected_close}，但第{line_num}行第{column}列为 {ch}',
                        "location": f'第{line_num}行第{column}列',
                        "char": ch,
                        "expected": expected_close,
                        "opening_char": open_ch,
                        "opening_line": open_line,
                        "opening_column": open_column,
                        "suggestion": "请检查高亮位置与对应开括号是否为同一种括号，或是否有括号被误删。",
                    }
                    _attach_issue_context(issue, contract_text, line_num, column)
                    opening_loc = _review_line_location(contract_text, open_line, open_column)
                    if opening_loc:
                        issue["locations"] = [
                            loc for loc in (
                                opening_loc,
                                _review_line_location(contract_text, line_num, column),
                            ) if loc
                        ]
                    issues.append(issue)
            else:
                line_num, column = _review_line_position(contract_text, i)
                issue = {
                    "type": "extra_close_bracket",
                    "severity": "warning",
                    "detail": f'多余的闭合括号: {ch}（第{line_num}行第{column}列）',
                    "location": f'第{line_num}行第{column}列',
                    "char": ch,
                    "suggestion": "请检查高亮位置前是否缺少匹配的开括号；若这是合法编号，请改成成对括号格式。",
                }
                _attach_issue_context(issue, contract_text, line_num, column)
                issues.append(issue)
    for open_ch, pos in bracket_stack:
        line_num, column = _review_line_position(contract_text, pos)
        issue = {
            "type": "unclosed_bracket",
            "severity": "warning",
            "detail": f'未闭合的括号: {open_ch}（第{line_num}行第{column}列）',
            "location": f'第{line_num}行第{column}列',
            "char": open_ch,
            "suggestion": "请检查高亮位置之后是否缺少对应的闭合括号。",
        }
        _attach_issue_context(issue, contract_text, line_num, column)
        issues.append(issue)

    # 4. English punctuation in Chinese context
    RE_ENG_PUNCT = re.compile(r'[一-鿿][,;\.][一-鿿]')
    for m in RE_ENG_PUNCT.finditer(contract_text):
        line_num, column = _review_line_position(contract_text, m.start() + 1)
        if line_num not in {i.get("line") for i in issues if i["type"] == "english_punctuation"}:
            issue = {
                "type": "english_punctuation",
                "severity": "info",
                "detail": f'中文字符间使用了英文标点: "{m.group()}"',
                "location": f'第{line_num}行第{column}列',
            }
            _attach_issue_context(issue, contract_text, line_num, column)
            issues.append(issue)

    # 5. Extra spaces between Chinese characters
    RE_CN_SPACE = re.compile(r'[一-鿿]\s+[一-鿿]')
    for line_num, line in enumerate(contract_text.splitlines(), 1):
        if is_toc_or_heading_spacing_line(line):
            continue
        if RE_CN_SPACE.search(line):
            column = RE_CN_SPACE.search(line).start() + 2
            issue = {
                "type": "extra_spaces",
                "severity": "info",
                "detail": "中文字符间有多余空格",
                "location": f'第{line_num}行第{column}列',
            }
            _attach_issue_context(issue, contract_text, line_num, column)
            issues.append(issue)

    # 6. Cross-reference validation
    RE_XREF = re.compile(r'详见第([一二三四五六七八九十百]+)部分')
    valid_nums = {num for _, num in parts if num is not None}
    for m in RE_XREF.finditer(contract_text):
        ref_num = _cn_numeral_to_int(m.group(1))
        if ref_num is not None and ref_num not in valid_nums:
            line_num, column = _review_line_position(contract_text, m.start())
            issue = {
                "type": "invalid_cross_reference",
                "severity": "warning",
                "detail": f'交叉引用指向不存在的章节: 第{m.group(1)}部分',
                "location": f'第{line_num}行第{column}列',
            }
            _attach_issue_context(issue, contract_text, line_num, column)
            issues.append(issue)

    issues = _finalize_review_issues(issues)
    problem_issues = [i for i in issues if i["severity"] in ("error", "warning")]
    info_issues = [i for i in issues if i["severity"] == "info"]
    error_count = sum(1 for i in problem_issues if i["severity"] == "error")
    warning_count = sum(1 for i in problem_issues if i["severity"] == "warning")
    info_count = len(info_issues)

    return jsonify({
        "success": True,
        "issues": problem_issues,
        "info_issues": info_issues,
        "total": len(problem_issues),
        "errors": error_count,
        "warnings": warning_count,
        "infos": info_count,
    })


@app.route("/api/review/consistency_check", methods=["POST"])
def api_review_consistency_check():
    req = request.get_json(silent=True) or {}
    data, job_id = _get_review_data(req)
    if not data or not data.get("contract_text"):
        return jsonify({"success": False, "error": "请先上传合同文件"}), 400

    contract_text = data["contract_text"]
    issues = []

    def clean_party_name(name):
        name = str(name or "").strip()
        name = re.sub(r'^(?:指|由)', '', name).strip()
        name = re.sub(r'[（(]法人盖章[）)]', '', name).strip()
        name = re.split(r'[；;。]', name, maxsplit=1)[0].strip()
        return name

    def collect_party_names(label):
        pattern = rf'^(?:\d+、)?{label}[：:]\s*([^\n]+)$'
        names = set()
        for line in contract_text.splitlines():
            if "新任" in line or "临时" in line:
                continue
            match = re.search(pattern, line.strip())
            if match:
                name = clean_party_name(match.group(1))
                if name and len(name) >= 4 and "管理人" not in name.replace(label, "") and "托管人" not in name.replace(label, ""):
                    names.add(name)
        return names

    # 1. Fund name consistency
    canonical_fund_name = data.get("fund_name") or _infer_review_fund_name(contract_text, data.get("contract_filename", ""))
    suspicious_fund_names = set()
    fund_name_pat = r'([^\s，。；：“”"《》]{4,100}联接基金)'
    for m in re.finditer(fund_name_pat, contract_text):
        candidate = _clean_review_fund_name(m.group(1))
        if "证券投资基金" not in candidate:
            continue
        if candidate and canonical_fund_name and canonical_fund_name not in candidate and candidate not in canonical_fund_name:
            suspicious_fund_names.add(candidate)
    if suspicious_fund_names:
        issues.append({
            "type": "fund_name_inconsistency",
            "severity": "error",
            "detail": f'基金名称不一致: {", ".join(sorted(suspicious_fund_names))}',
        })

    # 2. Fund manager name consistency
    managers = collect_party_names("基金管理人")
    if len(managers) > 1:
        issues.append({
            "type": "manager_name_inconsistency",
            "severity": "error",
            "detail": f'基金管理人名称不一致: {", ".join(sorted(managers))}',
        })

    # 3. Custodian name consistency
    custodians = collect_party_names("基金托管人")
    if len(custodians) > 1:
        issues.append({
            "type": "custodian_name_inconsistency",
            "severity": "error",
            "detail": f'基金托管人名称不一致: {", ".join(sorted(custodians))}',
        })

    # 4. Management fee rate consistency
    mgmt_rate_pat = r'管理费.*?(\d+\.\d+)%.*?年费'
    mgmt_rates = set()
    for m in re.finditer(mgmt_rate_pat, contract_text):
        mgmt_rates.add(m.group(1))
    if len(mgmt_rates) > 1:
        issues.append({
            "type": "mgmt_fee_rate_inconsistency",
            "severity": "warning",
            "detail": f'管理费率数值不一致: {", ".join(sorted(mgmt_rates))}%',
        })

    # 5. Custody fee rate consistency
    custody_rate_pat = r'托管费.*?(\d+\.\d+)%.*?年费'
    custody_rates = set()
    for m in re.finditer(custody_rate_pat, contract_text):
        custody_rates.add(m.group(1))
    if len(custody_rates) > 1:
        issues.append({
            "type": "custody_fee_rate_inconsistency",
            "severity": "warning",
            "detail": f'托管费率数值不一致: {", ".join(sorted(custody_rates))}%',
        })

    issues = _finalize_review_issues(issues)
    error_count = sum(1 for i in issues if i["severity"] == "error")
    warning_count = sum(1 for i in issues if i["severity"] == "warning")

    return jsonify({
        "success": True,
        "issues": issues,
        "total": len(issues),
        "errors": error_count,
        "warnings": warning_count,
    })


@app.route("/api/review/ai_check", methods=["POST"])
def api_review_ai_check():
    req = request.get_json(silent=True) or {}
    data, job_id = _get_review_data(req)
    if not data or not data.get("contract_text"):
        return jsonify({"success": False, "error": "\u8bf7\u5148\u4e0a\u4f20\u5408\u540c\u6587\u4ef6"}), 400

    fund_type = req.get("fund_type") or "联接基金"
    if not _as_bool(req.get("allow_external_ai")):
        return jsonify({
            "success": False,
            "error": "AI智能审查会将合同/招募文本发送至外部 DeepSeek 服务，请明确勾选外发确认后再运行。",
        }), 403

    def event(payload):
        return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    @stream_with_context
    def generate():
        api_key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
        if not api_key:
            yield event({
                "type": "error",
                "message": "\u672a\u914d\u7f6eDEEPSEEK_API_KEY\u73af\u5883\u53d8\u91cf\uff0cAI\u667a\u80fd\u5ba1\u67e5\u5df2\u8df3\u8fc7\u3002",
            })
            yield event({"type": "done"})
            return

        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com", timeout=60)
        except ImportError:
            yield event({"type": "error", "message": "未安装openai库"})
            yield event({"type": "done"})
            return

        rules = _load_review_rules(fund_type)
        contract_sections = data.get("contract_sections", [])
        prospectus_sections = data.get("prospectus_sections", [])

        def call_ai(prompt, chunk_label):
            issues = []
            try:
                response = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.1,
                    max_tokens=3000,
                )
                resp_text = response.choices[0].message.content.strip()
                if resp_text.startswith("```"):
                    resp_text = re.sub(r"^```\w*\n?", "", resp_text)
                    resp_text = re.sub(r"\n?```$", "", resp_text)
                result = json.loads(resp_text)
                for issue in result.get("issues", []):
                    issue["chapter"] = chunk_label
                    issue["source"] = "AI"
                    issues.append(issue)
            except Exception as exc:
                issues.append({
                    "chapter": chunk_label,
                    "type": "error",
                    "description": f"AI检查出错: {exc}",
                    "source": "AI",
                })
            return issues

        if data.get("prospectus_text"):
            checked = 0
            for item in _build_ai_cross_review_items(contract_sections, prospectus_sections, rules, fund_type):
                if checked >= 8:
                    break
                rule = item["rule"]
                known_diff = str(rule.get("expected_diff") or rule.get("diff") or "").strip()
                prompt = f"""你是基金法律文件复核专家。请根据勾稽规则审查以下基金合同与招募说明书的对应章节。

【勾稽规则】
章节: 合同「{item['contract_locator']}」↔ 招募「{item['prospectus_locator']}」
预期对应关系: {rule.get('relation') or '按勾稽规则对应'}
检查要点: {rule.get('suggestion') or rule.get('detail') or rule.get('consistency') or ''}
已知合理差异（不需要标记为问题）: {known_diff}

【基金合同原文】
{item['contract_text']}

【招募说明书原文】
{item['prospectus_text']}

请重点检查实质内容差异、数值费率日期不一致、权利义务遗漏或表述矛盾。
以JSON格式返回：
{{"issues": [{{"location": "问题所在原文", "type": "内容差异/数据不一致/条款遗漏/表述矛盾", "description": "具体问题描述", "suggestion": "修改建议"}}]}}
没有问题则返回 {{"issues": []}}。只返回JSON。"""
                label = item["label"]
                yield event({"type": "progress", "message": f"正在检查: {label}"})
                for issue in call_ai(prompt, label):
                    yield event({"type": "issue", "data": issue})
                checked += 1

        checked = 0
        for item in _build_ai_summary_review_items(contract_sections, rules):
            if checked >= 5:
                break
            prompt = f"""你是基金法律文件复核专家。请审查基金合同正文与合同内容摘要的对应关系。

【规则】
摘要位置: {item['summary_heading']}
对应正文: {item['contract_locator']}
收录方式: {item['rule'].get('method') or ''}
参考说明: {item['rule'].get('detail') or item['rule'].get('status') or ''}

【合同正文原文】
{item['contract_text']}

【合同摘要原文】
{item['summary_text']}

请检查摘要是否准确摘录正文核心条款，是否存在数值、条件、比例不一致，是否遗漏关键条款。
以JSON格式返回：
{{"issues": [{{"location": "问题所在原文", "type": "摘要不一致/关键遗漏/数据错误", "description": "具体问题描述", "suggestion": "修改建议"}}]}}
没有问题则返回 {{"issues": []}}。只返回JSON。"""
            label = item["label"]
            yield event({"type": "progress", "message": f"正在检查: {label}"})
            for issue in call_ai(prompt, label):
                yield event({"type": "issue", "data": issue})
            checked += 1

        key_chapters = ["前言", "基金的基本情况", "基金份额的发售", "违约责任", "基金合同的效力"]
        checked = 0
        for keyword in key_chapters:
            if checked >= 3:
                break
            sec = _find_review_section(contract_sections, keyword)
            if not sec or len(sec.get("content", "")) < 50:
                continue
            prompt = f"""请审查以下基金法律文件章节，检查错别字、语病、术语误用、明显数据错误或前后矛盾。

章节: {sec.get('heading', '')}
文本:
{_truncate_ai_review_text(sec.get('content', ''), 4000)}

以JSON格式返回：
{{"issues": [{{"location": "问题所在原文", "type": "错别字/语病/术语误用/数据错误/逻辑矛盾", "description": "问题描述", "suggestion": "修改建议"}}]}}
没有问题则返回 {{"issues": []}}。只返回JSON。"""
            label = sec.get("heading", keyword)
            yield event({"type": "progress", "message": f"正在检查: {label}"})
            for issue in call_ai(prompt, label):
                yield event({"type": "issue", "data": issue})
            checked += 1

        yield event({"type": "done"})

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/api/review/summary_cross_check", methods=["POST"])
def api_review_summary_cross_check():
    req = request.get_json(silent=True) or {}
    data, job_id = _get_review_data(req)
    if not data or not data.get("contract_text"):
        return jsonify({"success": False, "error": "请先上传合同文件"}), 400

    contract_sections = data.get("contract_sections", [])
    contract_summary = _find_contract_summary_section(contract_sections)
    if not contract_summary:
        issue = {
            "summary_pos": "基金合同的内容摘要",
            "contract_pos": "第二十四部分 基金合同的内容摘要",
            "status": "未找到",
            "method": "章节存在性检查",
            "contract_summary_heading": "",
            "prospectus_summary_heading": "",
            "contract_found": False,
            "contract_body_found": False,
            "prospectus_found": False,
            "similarity": None,
            "severity": "error",
            "is_problem": True,
            "message": "合同中未找到基金合同内容摘要章节",
            "suggestion": "请确认合同模板已生成“基金合同的内容摘要”章节，或检查章节标题是否被改写。",
            "diffs": None,
            "hunks": [],
        }
        return jsonify({
            "success": True,
            "results": [issue],
            "total": 1,
            "errors": 1,
            "warnings": 0,
            "problem_count": 1,
            "is_problem": True,
            "message": issue["message"],
        })

    prospectus_summary = None
    has_prospectus = bool(data.get("prospectus_text"))
    if data.get("prospectus_sections"):
        for sec in data["prospectus_sections"]:
            h = sec.get("heading", "")
            if "基金合同的内容摘要" in h or "基金合同内容摘要" in h:
                prospectus_summary = sec
                break

    if not has_prospectus:
        return jsonify({
            "success": True,
            "results": [],
            "total": 0,
            "errors": 0,
            "warnings": 0,
            "problem_count": 0,
            "is_problem": False,
            "message": "未上传招募说明书，已跳过合同摘要↔招募摘要对比",
        })

    summary_rules = list(BUILTIN_LINKED_FUND_SUMMARY_RULES)
    results = []
    contract_blocks = _split_summary_top_blocks(contract_summary.get("content", ""))
    p_blocks = _split_summary_top_blocks(prospectus_summary.get("content", "")) if prospectus_summary else []

    for rule in summary_rules:
        summary_pos = rule.get("summary_pos", "")
        contract_pos = rule.get("contract_pos", "")
        status = rule.get("status", "")
        method = rule.get("method", "")

        contract_block = _find_summary_block_by_rule(contract_blocks, summary_pos)
        contract_text = contract_block.get("body", "") if contract_block else ""

        matched_section = _match_review_section(contract_sections, contract_pos)
        contract_body_text = ""
        if matched_section["section"]:
            contract_body_text = matched_section["section"].get("content", "")

        prospectus_text = ""
        prospectus_block = None
        if prospectus_summary:
            prospectus_block = _find_summary_block_by_rule(p_blocks, summary_pos)
            prospectus_text = prospectus_block.get("body", "") if prospectus_block else ""

        similarity = None
        diff_pair = None
        hunks = []
        message = ""
        suggestion = "请核对合同摘要与招募说明书摘要口径，必要时同步修正。"
        if contract_text and prospectus_text:
            norm_c = _normalize_review_text(contract_text)
            norm_p = _normalize_review_text(prospectus_text)
            similarity = difflib.SequenceMatcher(
                None, norm_c, norm_p, autojunk=False).ratio() * 100
            c_lines = [l for l in contract_text.split('\n') if l.strip()]
            p_lines = [l for l in prospectus_text.split('\n') if l.strip()]
            if c_lines or p_lines:
                diff_pair = _build_review_excerpt_pair(c_lines, p_lines, context_lines=2)
                hunks = _build_review_line_hunks(contract_text, prospectus_text, context_lines=2)
            fact_mismatches = _compare_review_key_facts(contract_text, prospectus_text)
            if fact_mismatches:
                severity = "error"
                message = _review_fact_mismatch_message(fact_mismatches)
            else:
                severity = _summary_cross_severity(similarity)
                if severity in ("error", "warning"):
                    message = f"合同摘要与招募摘要存在差异，相似度 {round(similarity, 1)}%"
                else:
                    message = "合同摘要与招募摘要一致"
        elif not contract_text:
            severity = "error"
            message = f"合同摘要缺少“{summary_pos}”"
            suggestion = "请补齐合同摘要对应条目。"
        elif not prospectus_text:
            severity = "error"
            message = f"招募说明书摘要缺少“{summary_pos}”"
            suggestion = "请补齐招募说明书中的基金合同内容摘要对应条目。"
        else:
            severity = "error"
            message = f"摘要条目“{summary_pos}”未能定位"

        is_problem = severity in ("error", "warning")

        results.append({
            "summary_pos": summary_pos,
            "contract_pos": contract_pos,
            "status": status,
            "method": method,
            "contract_summary_heading": contract_block.get("heading", "") if contract_block else "",
            "prospectus_summary_heading": prospectus_block.get("heading", "") if prospectus_block else "",
            "contract_found": bool(contract_text),
            "contract_body_found": bool(contract_body_text),
            "prospectus_found": bool(prospectus_text),
            "similarity": round(similarity, 1) if similarity is not None else None,
            "severity": severity,
            "is_problem": is_problem,
            "message": message,
            "detail": message,
            "suggestion": suggestion,
            "contract_summary_text": _truncate_review_excerpt(contract_text, 1000),
            "prospectus_summary_text": _truncate_review_excerpt(prospectus_text, 1000),
            "diffs": hunks,
            "diff_pair": diff_pair,
            "hunks": hunks,
        })

    errors = sum(1 for r in results if r["severity"] == "error")
    warnings = sum(1 for r in results if r["severity"] == "warning")
    problem_count = errors + warnings
    problem_results = [r for r in results if r.get("is_problem")]
    worst_similarity = None
    similarities = [r.get("similarity") for r in results if r.get("similarity") is not None]
    if similarities:
        worst_similarity = min(similarities)

    return jsonify({
        "success": True,
        "results": results,
        "total": len(results),
        "errors": errors,
        "warnings": warnings,
        "problem_count": problem_count,
        "is_problem": problem_count > 0,
        "similarity": worst_similarity,
        "changed_lines": sum(
            1 for r in problem_results for h in (r.get("hunks") or [])
            if h.get("tag") not in ("equal", "skip")
        ),
        "message": f"发现 {problem_count} 个摘要对比问题" if problem_count else "摘要对比通过",
    })


@app.route("/api/review/report", methods=["POST"])
def api_review_report():
    body = request.get_json(silent=True) or {}
    cross_results = _finalize_review_issues(body.get("cross_results") or body.get("cross_check") or [])
    format_issues = _finalize_review_issues(body.get("format_issues", []))
    consistency_issues = _finalize_review_issues(body.get("consistency_issues", []))
    summary_results = body.get("summary_results") or []
    summary_cross_check = body.get("summary_cross_check") or {}
    if summary_cross_check and summary_cross_check.get("results"):
        summary_results = list(summary_results) + list(summary_cross_check.get("results") or [])
    elif summary_cross_check and summary_cross_check.get("is_problem"):
        summary_results = list(summary_results) + [{
            "summary_pos": summary_cross_check.get("contract_summary_heading", "合同摘要↔招募摘要"),
            "severity": summary_cross_check.get("severity") or (
                "warning" if float(summary_cross_check.get("similarity") or 0) >= 80 else "error"
            ),
            "detail": summary_cross_check.get("message")
            or f"合同摘要与招募摘要存在差异，相似度 {summary_cross_check.get('similarity', '')}%",
            "suggestion": summary_cross_check.get("suggestion", "请核对合同摘要与招募说明书摘要口径。"),
        }]
    summary_results = _finalize_review_issues(summary_results)
    review_data, job_id = _get_review_data(body)
    fund_name = body.get("fund_name") or review_data.get("fund_name", "")

    cross_errors = sum(1 for r in cross_results if r.get("severity") == "error")
    cross_warnings = sum(1 for r in cross_results if r.get("severity") == "warning")
    format_errors = sum(1 for i in format_issues if i.get("severity") == "error")
    format_warnings = sum(1 for i in format_issues if i.get("severity") == "warning")
    consistency_errors = sum(1 for i in consistency_issues if i.get("severity") == "error")
    consistency_warnings = sum(1 for i in consistency_issues if i.get("severity") == "warning")
    summary_errors = sum(1 for r in summary_results if r.get("severity") == "error")
    summary_warnings = sum(1 for r in summary_results if r.get("severity") == "warning")

    total_errors = cross_errors + format_errors + consistency_errors + summary_errors
    total_warnings = cross_warnings + format_warnings + consistency_warnings + summary_warnings

    if total_errors > 0:
        overall = "不通过"
    elif total_warnings > 5:
        overall = "需复核"
    else:
        overall = "通过"

    report_html = render_template(
        "review_report.html",
        fund_name=fund_name or "基金合同",
        overall=overall,
        total_errors=total_errors,
        total_warnings=total_warnings,
        cross_results=cross_results,
        cross_errors=cross_errors,
        cross_warnings=cross_warnings,
        cross_total=len(cross_results),
        format_issues=format_issues,
        format_errors=format_errors,
        format_warnings=format_warnings,
        format_total=len(format_issues),
        consistency_issues=consistency_issues,
        consistency_errors=consistency_errors,
        consistency_warnings=consistency_warnings,
        consistency_total=len(consistency_issues),
        summary_results=summary_results,
        summary_errors=summary_errors,
        summary_warnings=summary_warnings,
        summary_total=len(summary_results),
        generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    )

    return report_html, 200, {"Content-Type": "text/html; charset=utf-8"}


# ═══════════════════════════════════════════════════════════════════════════════
#  启动
# ═══════════════════════════════════════════════════════════════════════════════
def open_browser():
    time.sleep(1.2)
    webbrowser.open("http://127.0.0.1:5000")


if __name__ == "__main__":
    t = threading.Thread(target=open_browser, daemon=True)
    t.start()
    app.run(debug=False, port=5000, use_reloader=False)
