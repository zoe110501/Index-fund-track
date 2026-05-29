import gc
import io
import inspect
import json
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
import unittest
import warnings
import zipfile
import importlib
from unittest import mock
from copy import deepcopy
from pathlib import Path
from types import SimpleNamespace

import app
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

AUTOMOTIVE_ETF_CONTRACT_DOCX = Path(
    r"C:\Users\12534\Downloads\南方中证全指汽车交易型开放式指数证券投资基金基金合同.docx"
)
AUTOMOTIVE_ETF_PROSPECTUS_DOCX = Path(
    r"C:\Users\12534\Downloads\南方中证全指汽车交易型开放式指数证券投资基金招募说明书.docx"
)


class ProspectusEngineSourceAuditTests(unittest.TestCase):
    def test_app_source_has_single_prospectus_engine_definition_without_monkey_patches(self):
        app_source = Path(app.__file__).read_text(encoding="utf-8")
        self.assertEqual(app_source.count("class ProspectusEngine:"), 1)
        self.assertNotIn("ProspectusEngine.__init__ =", app_source)
        self.assertNotIn("ProspectusEngine.build_docx_prospectus =", app_source)

    def test_builtin_review_rules_are_kept_in_dedicated_module(self):
        rules_module = importlib.import_module("review_builtin_rules")
        self.assertGreaterEqual(len(rules_module.BUILTIN_ETF_CROSS_RULES), 20)
        self.assertEqual(9, len(rules_module.BUILTIN_ETF_SUMMARY_RULES))
        self.assertEqual(app.BUILTIN_ETF_CROSS_RULES, rules_module.BUILTIN_ETF_CROSS_RULES)
        self.assertEqual(app.BUILTIN_ETF_SUMMARY_RULES, rules_module.BUILTIN_ETF_SUMMARY_RULES)

    def test_build_docx_prospectus_signature_accepts_optional_form_data(self):
        signature = inspect.signature(app.ProspectusEngine.build_docx_prospectus)
        params = list(signature.parameters.values())

        self.assertEqual([param.name for param in params[:2]], ["self", "prospectus_text"])
        self.assertEqual(params[2].name, "form_data")
        self.assertIsNone(params[2].default)


class ContractEngineRegressionTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.engine = app.ContractEngine()

    def build_form(self, **overrides):
        data = {
            "FUND_NAME": "测试上交所科创板交易型开放式指数证券投资基金",
            "FUND_SHORT_NAME": "测试科创ETF",
            "INDEX_NAME": "上证测试指数",
            "INDEX_CODE": "000001",
            "INDEX_COMPILER": "中证指数有限公司",
            "CONTRACT_DATE": "2026年1月",
            "EXCHANGE": "SSE",
            "MARKET_TYPE": "KECHUANG",
            "CUSTODIAN_NAME": "中信建投证券股份有限公司",
            "DISTRIBUTION_FREQ": "QUARTERLY",
            "MGMT_FEE_PAYMENT_METHOD": "CONSULT",
            "CUSTODY_FEE_PAYMENT_METHOD": "CONSULT",
            "MGMT_FEE_RATE": "0.15",
            "CUSTODY_FEE_RATE": "0.05",
        }
        data.update(overrides)
        return data

    def extract_preface(self, text):
        return text.rsplit("第一部分  前言", 1)[1].split("第二部分  释义", 1)[0]

    def extract_contract_summary(self, text):
        match = re.search(
            r"第二十六部分\s+基金合同内容摘要\s*(?P<body>.*?)(?:\n（本页为|$)",
            text,
            flags=re.S,
        )
        self.assertIsNotNone(match, "未找到基金合同内容摘要")
        return match.group("body").strip()

    def test_schema_defaults_are_available_to_generation_templates(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            template_path = tmp_path / "template.md"
            schema_path = tmp_path / "schema.json"
            base_template = app.TEMPLATE_MD.read_text(encoding="utf-8")
            template_path.write_text(
                base_template
                + "\n\n"
                + (
                    "schema variable: {CUSTOM_NOTICE}\n"
                    "{{IF SHOW_CUSTOM_NOTICE}}schema condition enabled{{ENDIF}}\n"
                    "{{IF_NOT SHOW_CUSTOM_NOTICE}}schema condition disabled{{ENDIF}}\n"
                ),
                encoding="utf-8",
            )
            schema_path.write_text(
                json.dumps(
                    {
                        "version": "test",
                        "groups": {
                            "custom_generation": {
                                "label": "Custom generation",
                                "variables": [
                                    {
                                        "key": "CUSTOM_NOTICE",
                                        "type": "string",
                                        "default": "default-from-schema",
                                    },
                                    {
                                        "key": "SHOW_CUSTOM_NOTICE",
                                        "type": "boolean",
                                        "default": True,
                                    },
                                ],
                            }
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            with (
                mock.patch.object(app, "TEMPLATE_MD", template_path),
                mock.patch.object(app, "SCHEMA_JSON", schema_path),
            ):
                contract = app.ContractEngine().generate(self.build_form())
                overridden = app.ContractEngine().generate(
                    self.build_form(CUSTOM_NOTICE="manual-value", SHOW_CUSTOM_NOTICE=False)
                )

        self.assertIn("schema variable: default-from-schema", contract)
        self.assertIn("schema condition enabled", contract)
        self.assertNotIn("schema condition disabled", contract)
        self.assertIn("schema variable: manual-value", overridden)
        self.assertNotIn("schema condition enabled", overridden)
        self.assertIn("schema condition disabled", overridden)

    def extract_summary_group(self, summary, heading):
        match = re.search(
            rf"^{re.escape(heading)}\s*(?P<body>.*?)(?=^[一二三四五六七八九十百]+、|\Z)",
            summary,
            flags=re.S | re.M,
        )
        self.assertIsNotNone(match, f"未找到摘要小节：{heading}")
        return match.group("body").strip()

    def write_summary_placeholder_template(self, path):
        path.write_text(
            """{FUND_NAME}基金合同

第九部分  基金合同当事人及权利义务
一、基金管理人
（一）基金管理人简况
管理人简况不应进入摘要。
（二）基金管理人的权利与义务
管理人自动来源正文。
二、基金托管人
（一）基金托管人简况
托管人简况不应进入摘要。
（二）基金托管人的权利与义务
托管人自动来源正文。
三、基金份额持有人
持有人自动来源正文。

第十部分  基金份额持有人大会
一、召开事由
召开事由自动来源正文。
二、会议召集人及召集方式
召集方式自动来源正文。
三、召开基金份额持有人大会的通知时间、通知内容、通知方式
通知方式自动来源正文。
四、基金份额持有人出席会议的方式
出席方式自动来源正文。
五、议事内容与程序
议事程序自动来源正文。
六、表决
表决规则自动来源正文。
七、计票
计票规则自动来源正文。
八、生效与公告
生效公告自动来源正文。
九、本部分关于基金份额持有人大会召开事由、召开条件、议事程序、表决条件等规定，凡是直接引用法律法规或监管规则的部分，如将来法律法规或监管规则修改导致相关内容被取消或变更的，基金管理人与基金托管人协商一致并提前公告后，可直接对本部分内容进行修改和调整，无需召开基金份额持有人大会审议。
监管调整自动来源正文。

第十四部分  基金的投资
一、投资目标
投资目标不应进入摘要。
二、投资范围
投资范围自动来源正文。
三、投资策略
投资策略不应进入摘要。
四、投资限制
投资限制自动来源正文。

第十六部分  基金资产估值
一、估值日
估值日不应进入摘要。
五、估值程序
估值程序自动来源正文。

第十七部分  基金费用与税收
一、基金费用的种类
费用种类自动来源正文。
二、基金费用计提方法、计提标准和支付方式
管理费自动来源正文。
三、不列入基金费用的项目
不列入基金费用不应进入摘要。
四、基金税收
基金税收不应进入摘要。

第十八部分  基金的收益与分配
一、基金利润的构成
利润构成不应进入摘要。
三、基金收益分配原则
收益原则自动来源正文。
四、收益分配方案
收益方案自动来源正文。
五、收益分配方案的确定、公告与实施
公告实施自动来源正文。
六、基金收益分配中发生的费用
分配费用自动来源正文。

第二十一部分  基金合同的变更、终止与基金财产的清算
一、《基金合同》的变更
合同变更自动来源正文。
二、《基金合同》的终止事由
终止事由自动来源正文。
三、基金财产的清算
清算程序自动来源正文。
四、清算费用
清算费用自动来源正文。
五、基金财产清算剩余资产的分配
剩余资产分配自动来源正文。
六、基金财产清算的公告
清算公告自动来源正文。
七、基金财产清算账册及文件的保存
清算文件保存自动来源正文。

第二十三部分  争议的处理和适用的法律
争议解决自动来源正文。

第二十四部分  基金合同的效力
合同效力自动来源正文。

第二十六部分  基金合同内容摘要
一、基金合同当事人的权利、义务
{CONTRACT_SUMMARY_PARTIES}
二、基金份额持有人大会召集、议事及表决的程序和规则
{CONTRACT_SUMMARY_MEETING}
三、基金收益分配原则、执行方式
{CONTRACT_SUMMARY_DISTRIBUTION}
四、基金费用与税收
{CONTRACT_SUMMARY_FEES}
五、基金财产的投资范围和投资限制
{CONTRACT_SUMMARY_INVESTMENT}
六、基金资产净值的计算方法和公告方式
{CONTRACT_SUMMARY_VALUATION}
七、基金合同的变更、终止与基金财产的清算
{CONTRACT_SUMMARY_CHANGE_TERMINATION}
八、争议解决方式
{CONTRACT_SUMMARY_DISPUTE}
九、基金合同的效力
{CONTRACT_SUMMARY_EFFECT}

（本页为《{FUND_NAME}基金合同》签署页，无正文）
""",
            encoding="utf-8",
        )

    def test_contract_part8_sec6_keeps_contract_specific_wording(self):
        expected_fee = "4、投资人在申购或赎回基金份额时，申购赎回代理券商可按照一定的标准收取佣金，其中包含证券交易所、登记机构等收取的相关费用，具体规定请参见招募说明书及基金产品资料概要。"
        forms = [
            self.build_form(),
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT"),
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"),
        ]
        for form in forms:
            with self.subTest(exchange=form["EXCHANGE"], market_type=form["MARKET_TYPE"]):
                text = self.engine.generate(form)
                exchange_name = "上海证券交易所" if form["EXCHANGE"] == "SSE" else "深圳证券交易所"
                self.assertIn(f"3、申购赎回清单由基金管理人编制。T日的申购赎回清单在当日{exchange_name}开市前公告。申购赎回清单的内容与格式示例参见招募说明书。", text)
                self.assertIn(expected_fee, text)
                self.assertNotIn("申购赎回代理券商可按不超过0.2%的标准收取佣金", text)

    def test_contract_summary_placeholders_are_rendered_from_generated_body(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "contract_template.md"
            self.write_summary_placeholder_template(template_path)
            with mock.patch.object(app, "TEMPLATE_MD", template_path):
                text = self.engine.generate(self.build_form(FUND_NAME="测试基金"))

        summary = self.extract_contract_summary(text)
        self.assertNotIn("{CONTRACT_SUMMARY_", summary)
        self.assertIn("管理人自动来源正文。", summary)
        self.assertIn("托管人自动来源正文。", summary)
        self.assertIn("持有人自动来源正文。", summary)
        self.assertIn("争议解决自动来源正文。", summary)
        self.assertIn("合同效力自动来源正文。", summary)
        self.assertNotIn("管理人简况不应进入摘要。", summary)
        self.assertNotIn("托管人简况不应进入摘要。", summary)

        fees = self.extract_summary_group(summary, "四、基金费用与税收")
        self.assertIn("费用种类自动来源正文。", fees)
        self.assertIn("管理费自动来源正文。", fees)
        self.assertNotIn("不列入基金费用不应进入摘要。", fees)
        self.assertNotIn("基金税收不应进入摘要。", fees)

        investment = self.extract_summary_group(summary, "五、基金财产的投资范围和投资限制")
        self.assertIn("投资范围自动来源正文。", investment)
        self.assertIn("投资限制自动来源正文。", investment)
        self.assertNotIn("投资目标不应进入摘要。", investment)
        self.assertNotIn("投资策略不应进入摘要。", investment)

        distribution = self.extract_summary_group(summary, "三、基金收益分配原则、执行方式")
        self.assertIn("收益原则自动来源正文。", distribution)
        self.assertNotIn("利润构成不应进入摘要。", distribution)

    def test_contract_summary_keeps_current_scope_and_dynamic_values(self):
        text = self.engine.generate(
            self.build_form(
                MGMT_FEE_RATE="0.33",
                CUSTODY_FEE_RATE="0.07",
                DISPUTE_RESOLUTION_VENUE="BJ_CIETAC",
            )
        )

        summary = self.extract_contract_summary(text)
        self.assertNotIn("{CONTRACT_SUMMARY_", summary)
        self.assertIn("0.33%", summary)
        self.assertIn("0.07%", summary)
        self.assertIn("仲裁地点为北京市。", summary)

        parties = self.extract_summary_group(summary, "一、基金合同当事人的权利、义务")
        self.assertNotIn("基金托管人简况", parties)
        self.assertNotIn("基金托管资格批文及文号", parties)

        fees = self.extract_summary_group(summary, "四、基金费用与税收")
        self.assertNotIn("不列入基金费用的项目", fees)
        self.assertNotIn("基金税收", fees)

        investment = self.extract_summary_group(summary, "五、基金财产的投资范围和投资限制")
        self.assertNotIn("（一）投资目标", investment)
        self.assertNotIn("（三）投资策略", investment)

    def test_contract_summary_rewrites_meeting_internal_number_references(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "contract_template.md"
            self.write_summary_placeholder_template(template_path)
            text = template_path.read_text(encoding="utf-8")
            text = text.replace(
                "议事程序自动来源正文。",
                (
                    "2、议事程序\n"
                    "（1）现场开会\n"
                    "在现场开会的方式下，首先由大会主持人按照下列第七条规定程序确定和公布监票人，"
                    "然后由大会主持人宣读提案，经讨论后进行表决，并形成大会决议。"
                ),
            )
            template_path.write_text(text, encoding="utf-8")
            with mock.patch.object(app, "TEMPLATE_MD", template_path):
                contract_text = self.engine.generate(self.build_form(FUND_NAME="测试基金"))

        summary = self.extract_contract_summary(contract_text)
        meeting = self.extract_summary_group(summary, "二、基金份额持有人大会召集、议事及表决的程序和规则")
        self.assertIn("第（七）条规定程序", meeting)
        self.assertNotIn("第七条规定程序", meeting)

    def test_contract_summary_rewrites_quoted_top_level_heading_references(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "contract_template.md"
            self.write_summary_placeholder_template(template_path)
            text = template_path.read_text(encoding="utf-8")
            text = text.replace(
                "管理费自动来源正文。",
                "上述“一、基金费用的种类”中第4-12项费用，根据有关法规及相应协议规定，按费用实际支出金额列入当期费用。",
            )
            template_path.write_text(text, encoding="utf-8")
            with mock.patch.object(app, "TEMPLATE_MD", template_path):
                contract_text = self.engine.generate(self.build_form(FUND_NAME="测试基金"))

        summary = self.extract_contract_summary(contract_text)
        fees = self.extract_summary_group(summary, "四、基金费用与税收")
        self.assertIn("上述“（一）基金费用的种类”中第4-12项费用", fees)
        self.assertNotIn("上述“一、基金费用的种类”", fees)

    def test_contract_sse_exchange_defaults_working_day_type_when_missing(self):
        text = self.engine.generate(self.build_form())

        self.assertIn("工作日：指上海证券交易所的正常交易日", text)
        self.assertNotIn("工作日：指深圳证券交易所的正常交易日", text)

    def test_contract_sse_kechuang_defaults_stock_subscription_true(self):
        text = self.engine.generate(self.build_form())

        self.assertIn("基金管理人可选择网上现金认购、网下现金认购和网下股票认购3种方式发售本基金", text)
        self.assertIn("2、募集期间认购资金与股票", text)
        self.assertIn("基金募集金额（含网下股票认购所募集的股票市值）不少于2亿元人民币", text)

    def test_contract_subscription_method_uses_selected_exchange_name(self):
        text = self.engine.generate(self.build_form(HAS_STOCK_SUBSCRIPTION=True))

        self.assertIn("上海证券交易所网上系统以现金进行的认购", text)
        self.assertNotIn("深圳证券交易所网上系统以现金进行的认购", text)
        self.assertNotIn("上海证券交易所证券交易所网上系统", text)

    def test_contract_dispute_resolution_venue_can_use_beijing_cietac(self):
        text = self.engine.generate(self.build_form(DISPUTE_RESOLUTION_VENUE="BJ_CIETAC"))

        self.assertIn("将争议提交中国国际经济贸易仲裁委员会", text)
        self.assertIn("仲裁地点为北京市。", text)
        self.assertNotIn("将争议提交深圳国际仲裁院", text)
        self.assertNotIn("{DISPUTE_RESOLUTION_CLAUSE}", text)

    def test_contract_generated_clause_texts_are_loaded_from_clause_library(self):
        engine = app.ContractEngine(business_text_overrides={})
        for clause_key in [
            "NON_COMPONENT_SCOPE_INTRO",
            "DEFINITION_TAIL",
            "PREFACE_RISK_DISCLOSURE_TAIL",
            "PURCHASE_REDEMPTION_CONSIDERATION_TEXTS",
            "VALUATION_DISCLOSURE_TEXTS",
            "LISTING_TRADING_TEXTS",
            "PURCHASE_REDEMPTION_SECTION_TEXTS",
            "HK_PROXY_VOTING_TEXT",
        ]:
            self.assertIn(clause_key, engine.clauses)

        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT", HAS_HK_CONNECT=True, INDEX_NAME="港股通测试指数")
        variables = engine._inject_clause_texts(engine._derive_variables(form))
        library_section4 = engine.clauses["PURCHASE_REDEMPTION_SECTION_TEXTS"]["variants"]["SSE_HK"]["part8_section4_text"]

        self.assertEqual(variables["PART8_SECTION4_TEXT"], library_section4)

    def test_contract_kechuang_preface_includes_kechuang_risk_disclosure(self):
        text = self.engine.generate(self.build_form())
        preface = self.extract_preface(text)

        self.assertIn("本基金投资于科创板股票，会面临科创板机制下因投资标的、市场制度以及交易规则等差异带来的特有风险", preface)


    def test_contract_business_text_override_file_replaces_listing_section_2_body_for_sse_hk(self):
        original_base_dir = app.BASE_DIR
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                override_text = "这是来自外部覆盖文件的上市交易正文。"
                (base_dir / "09_业务正文覆盖.json").write_text(
                    json.dumps(
                        {
                            "contract": {
                                "LISTING_SECTION2_BODY": {
                                    "SSE_HK": override_text,
                                }
                            },
                            "prospectus": {},
                            "meta": {"version": "1.0"},
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )

                with mock.patch.object(app, "BASE_DIR", base_dir):
                    engine = app.ContractEngine()
                    variables = engine._inject_clause_texts(
                        engine._derive_variables(
                            self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT")
                        )
                    )

                self.assertEqual(override_text, variables["LISTING_SECTION2_BODY"])
        finally:
            app.BASE_DIR = original_base_dir

    def test_contract_business_text_override_file_matches_condition_specific_text_then_falls_back(self):
        original_base_dir = app.BASE_DIR
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                base_text = "base contract body"
                exact_text = "exact contract body"
                (base_dir / "09_业务正文覆盖.json").write_text(
                    json.dumps(
                        {
                            "contract": {
                                "LISTING_SECTION2_BODY": {
                                    "DEFAULT": {
                                        "ALL": {
                                            "ALL": base_text,
                                        },
                                        "ETF": {
                                            "A_SHARE": exact_text,
                                        },
                                    }
                                }
                            },
                            "prospectus": {},
                            "meta": {"version": "2.0"},
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )

                with mock.patch.object(app, "BASE_DIR", base_dir):
                    engine = app.ContractEngine()
                    exact_variables = engine._inject_clause_texts(
                        engine._derive_variables(
                            self.build_form(MARKET_TYPE="A_SHARE", PRODUCT_TYPE="ETF")
                        )
                    )
                    fallback_variables = engine._inject_clause_texts(
                        engine._derive_variables(
                            self.build_form(MARKET_TYPE="KECHUANG", PRODUCT_TYPE="ETF")
                        )
                    )

                self.assertEqual(exact_text, exact_variables["LISTING_SECTION2_BODY"])
                self.assertEqual(base_text, fallback_variables["LISTING_SECTION2_BODY"])
        finally:
            app.BASE_DIR = original_base_dir

    def test_contract_business_text_override_file_matches_exchange_specific_text_then_falls_back(self):
        base_text = "base by exchange"
        sse_text = "sse exchange body"
        engine = app.ContractEngine(
            business_text_overrides={
                "contract": {
                    "LISTING_ITEM3_DEF": {
                        "DEFAULT": {
                            "ALL": {
                                "ALL": {
                                    "ALL": base_text,
                                    "SSE": sse_text,
                                }
                            }
                        }
                    }
                },
                "prospectus": {},
                "meta": {"version": "2.0"},
            }
        )

        sse_variables = engine._inject_clause_texts(
            engine._derive_variables(
                self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")
            )
        )
        szse_variables = engine._inject_clause_texts(
            engine._derive_variables(
                self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE")
            )
        )

        self.assertEqual(sse_text, sse_variables["LISTING_ITEM3_DEF"])
        self.assertEqual(base_text, szse_variables["LISTING_ITEM3_DEF"])


class ProspectusEngineRoutingTests(unittest.TestCase):
    REFERENCE_FUND_NAME = "南方中证全指红利质量交易型开放式指数证券投资基金"
    PROSPECTUS_PRODUCT_SUMMARY_ITEM4 = (
        "4、基金产品资料概要是基金招募说明书的摘要文件，用于向投资者提供简明的基金概要信息。"
        "《基金合同》生效后，基金产品资料概要的信息发生重大变更的，基金管理人应当在三个工作日内，"
        "更新基金产品资料概要，并登载在规定网站及基金销售机构网站或营业网点；基金产品资料概要其他"
        "信息发生变更的，基金管理人至少每年更新一次。基金终止运作的，基金管理人不再更新基金产品资料概要。"
        "基金产品资料概要的正文应当包括产品概况、基金投资与净值表现、投资本基金涉及的费用、风险揭示"
        "与重要提示等中国证监会规定的披露事项，相关内容不得与基金合同、招募说明书有实质性差异。"
    )
    CHAPTER_TITLES = {
        "一": "绪言",
        "二": "释义",
        "三": "基金管理人",
        "四": "基金托管人",
        "五": "相关服务机构",
        "六": "基金的募集",
        "七": "基金合同的生效",
        "八": "基金份额折算与变更登记",
        "九": "基金份额的上市交易",
        "十": "基金份额的申购赎回",
        "十一": "基金的投资",
        "十二": "基金的财产",
        "十三": "基金资产估值",
        "十四": "基金的收益与分配",
        "十五": "基金的费用与税收",
        "十六": "基金的会计与审计",
        "十七": "基金的信息披露",
        "十八": "风险揭示",
        "十九": "基金合同的变更、终止和基金财产的清算",
        "二十": "基金合同的内容摘要",
        "二十一": "基金托管协议的内容摘要",
        "二十二": "基金份额持有人服务",
        "二十三": "其他应披露事项",
        "二十四": "招募说明书存放及其查阅方式",
        "二十五": "备查文件",
    }

    @classmethod
    def setUpClass(cls):
        cls.engine = app.ProspectusEngine()

    def build_form(self, **overrides):
        data = {
            "FUND_NAME": "测试交易型开放式指数证券投资基金",
            "INDEX_NAME": "中证测试指数",
            "INDEX_DESCRIPTION": "中证测试指数选取测试样本证券作为指数样本，以反映测试市场的整体表现。",
            "INDEX_WEBSITE": "www.example-index.com",
            "INDEX_COMPILER": "测试指数公司",
            "IMPORTANT_NOTICE_STYLE": "AUTO",
            "EXCHANGE": "SSE",
            "MARKET_TYPE": "A_SHARE",
            "CUSTODIAN_NAME": "招商银行股份有限公司",
            "CONTRACT_DATE": "2026年1月1日",
            "FUND_MANAGER_NAME": "张三",
            "FUND_MANAGER_BIO": "张三，硕士，2020年起任基金经理。",
            "MIN_SUB_UNIT": "250,000份",
        }
        data.update(overrides)
        return data

    def chapter_heading(self, chapter_cn):
        return f"{chapter_cn}、{self.CHAPTER_TITLES[chapter_cn]}"

    def extract_chapter(self, text, chapter_cn):
        pattern = re.compile(rf"^第{chapter_cn}章[^\n]*", re.MULTILINE)
        matches = list(pattern.finditer(text))
        if matches:
            start = matches[-1].start()
            next_pattern = re.compile(r"^第[一二三四五六七八九十百]+章[^\n]*", re.MULTILINE)
            nxt = next_pattern.search(text, matches[-1].end())
            end = nxt.start() if nxt else len(text)
            return text[start:end].strip()

        lines = text.splitlines()
        heading = self.chapter_heading(chapter_cn)
        headings = {self.chapter_heading(cn) for cn in self.CHAPTER_TITLES}
        matches = [idx for idx, line in enumerate(lines) if line.strip() == heading]
        self.assertTrue(matches, f"missing chapter {chapter_cn}")
        start_idx = matches[-1]

        end_idx = len(lines)
        for idx in range(start_idx + 1, len(lines)):
            if lines[idx].strip() in headings:
                end_idx = idx
                break
        return "\n".join(lines[start_idx:end_idx]).strip()

    def extract_toc(self, text):
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        toc_idx = lines.index("目录")
        first_heading = self.chapter_heading("一")
        headings = []
        body_start = None
        for idx in range(toc_idx + 1, len(lines)):
            line = lines[idx]
            if line == first_heading and headings:
                body_start = idx
                break
            headings.append(line)
        self.assertIsNotNone(body_start, "missing formatted body start after TOC")
        return "\n".join(headings)

    def extract_preface_before_toc(self, text):
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        toc_idx = lines.index("\u76ee\u5f55")
        return "\n".join(lines[:toc_idx])

    def extract_important_notice(self, text):
        preface = self.extract_preface_before_toc(text)
        lines = [line.strip() for line in preface.splitlines() if line.strip()]
        notice_idx = lines.index("\u91cd\u8981\u63d0\u793a")
        return "\n".join(lines[notice_idx:])

    def test_prospectus_business_text_override_file_replaces_subscription_account_clause_for_sse_hk(self):
        original_base_dir = app.BASE_DIR
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                override_text = "这是来自外部覆盖文件的开户条款。"
                (base_dir / "09_业务正文覆盖.json").write_text(
                    json.dumps(
                        {
                            "contract": {},
                            "prospectus": {
                                "SUBSCRIPTION_ACCOUNT_CLAUSE": {
                                    "SSE_HK": override_text,
                                }
                            },
                            "meta": {"version": "1.0"},
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )

                with mock.patch.object(app, "BASE_DIR", base_dir):
                    prospectus_engine = app.ProspectusEngine()
                    variables = prospectus_engine._inject_clause_texts(
                        self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT")
                    )

                self.assertEqual(override_text, variables["SUBSCRIPTION_ACCOUNT_CLAUSE"])
        finally:
            app.BASE_DIR = original_base_dir

    def test_prospectus_business_text_override_file_replaces_index_info_source_clause(self):
        original_base_dir = app.BASE_DIR
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                override_text = "\u8fd9\u662f\u6765\u81ea\u5916\u90e8\u8986\u76d6\u6587\u4ef6\u7684\u6307\u6570\u6765\u6e90\u6761\u6b3e\u3002"
                (base_dir / "09_\u4e1a\u52a1\u6b63\u6587\u8986\u76d6.json").write_text(
                    json.dumps(
                        {
                            "contract": {},
                            "prospectus": {
                                "INDEX_INFO_SOURCE_CLAUSE": {
                                    "DEFAULT": override_text,
                                }
                            },
                            "meta": {"version": "1.0"},
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )

                with mock.patch.object(app, "BASE_DIR", base_dir):
                    prospectus_engine = app.ProspectusEngine()
                    variables = prospectus_engine._inject_clause_texts(
                        self.build_form(
                            EXCHANGE="SSE",
                            MARKET_TYPE="HK_CONNECT",
                            INDEX_COMPILER="\u4e2d\u8bc1\u6307\u6570\u6709\u9650\u516c\u53f8",
                            INDEX_WEBSITE="https://example.com/index",
                        )
                    )

                self.assertEqual(override_text, variables["INDEX_INFO_SOURCE_CLAUSE"])
        finally:
            app.BASE_DIR = original_base_dir

    def test_prospectus_business_text_override_file_replaces_important_notice_index_source_sentence(self):
        original_base_dir = app.BASE_DIR
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                override_template = "自定义指数说明请见{INDEX_COMPILER}官网：{INDEX_WEBSITE}。"
                (base_dir / "09_业务正文覆盖.json").write_text(
                    json.dumps(
                        {
                            "contract": {},
                            "prospectus": {
                                "IMPORTANT_NOTICE_INDEX_SOURCE_SENTENCE": {
                                    "DEFAULT": override_template,
                                }
                            },
                            "meta": {"version": "1.0"},
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )

                with mock.patch.object(app, "BASE_DIR", base_dir):
                    prospectus_engine = app.ProspectusEngine()
                    with mock.patch.object(
                        prospectus_engine,
                        "_load_reference_fixed_content",
                        return_value={
                            "important_notice": (
                                "重要提示\n"
                                "原始提示。\n"
                                "本基金是跟踪标的指数的交易型开放式基金，本基金标的指数为旧指数。\n"
                                "本基金被动跟踪标的指数“旧指数”，因此，本基金的业绩表现与标的指数的表现密切相关。"
                            )
                        },
                    ):
                        variables = prospectus_engine._inject_clause_texts(
                            self.build_form(
                                INDEX_NAME="测试指数",
                                INDEX_COMPILER="中证指数有限公司",
                                INDEX_WEBSITE="https://example.com/index",
                            )
                        )
                        blocks = prospectus_engine._build_important_notice_blocks(variables)

                self.assertIn(
                    "自定义指数说明请见中证指数有限公司官网：https://example.com/index。",
                    blocks,
                )
        finally:
            app.BASE_DIR = original_base_dir

    def test_prospectus_business_text_override_file_replaces_chapter_ten_sec5_default_body(self):
        original_base_dir = app.BASE_DIR
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                override_template = (
                    "五、申购和赎回的数额限制\n"
                    "自定义的最小申购赎回单位为{MIN_SUB_UNIT}。"
                )
                (base_dir / "09_业务正文覆盖.json").write_text(
                    json.dumps(
                        {
                            "contract": {},
                            "prospectus": {
                                "CHAPTER10_SEC5_DEFAULT_BODY": {
                                    "DEFAULT": override_template,
                                }
                            },
                            "meta": {"version": "1.0"},
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )

                with mock.patch.object(app, "BASE_DIR", base_dir):
                    prospectus_engine = app.ProspectusEngine()
                    variables = prospectus_engine._inject_clause_texts(
                        self.build_form(MIN_SUB_UNIT="250万份")
                    )
                    section = prospectus_engine._build_chapter_ten_sec5({}, variables)

                self.assertEqual(
                    "五、申购和赎回的数额限制\n自定义的最小申购赎回单位为250万份。",
                    section,
                )
        finally:
            app.BASE_DIR = original_base_dir

    def test_prospectus_generated_clause_texts_are_loaded_from_clause_library(self):
        engine = app.ProspectusEngine()
        self.assertIn("PRODUCT_SUMMARY_DISCLOSURE_ITEM", engine.pro_clauses)
        self.assertIn("DISTRIBUTION_CONDITIONS_SECTION", engine.pro_clauses)
        self.assertEqual(
            engine._prospectus_product_summary_item_four_text(),
            engine.pro_clauses["PRODUCT_SUMMARY_DISCLOSURE_ITEM"]["variants"]["DEFAULT"]["text"],
        )

    def extract_section(self, chapter_text, section_cn):
        pattern = re.compile(rf"^{section_cn}、[^\n]*", re.MULTILINE)
        match = pattern.search(chapter_text)
        self.assertIsNotNone(match, f"missing section {section_cn}")
        next_pattern = re.compile(r"^[一二三四五六七八九十百]+、[^\n]*", re.MULTILINE)
        nxt = next_pattern.search(chapter_text, match.end())
        end = nxt.start() if nxt else len(chapter_text)
        return chapter_text[match.start():end].strip()

    def extract_numbered_item(self, section_text, item_no):
        pattern = re.compile(rf"^{item_no}、[^\n]*", re.MULTILINE)
        match = pattern.search(section_text)
        self.assertIsNotNone(match, f"missing item {item_no}")
        next_pattern = re.compile(r"^\d+、[^\n]*", re.MULTILINE)
        nxt = next_pattern.search(section_text, match.end())
        end = nxt.start() if nxt else len(section_text)
        return section_text[match.start():end].strip()

    def _docx_xml(self, docx_bytes):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(docx_bytes)
            tmp.flush()
            with zipfile.ZipFile(tmp.name) as zf:
                return {
                    name: zf.read(name).decode("utf-8", errors="ignore")
                    for name in zf.namelist()
                    if name in {"word/document.xml", "word/settings.xml", "word/styles.xml", "word/numbering.xml"}
                    or name.startswith("word/footer")
                    or name.startswith("word/header")
                }

    def _docx_document(self, docx_bytes):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(docx_bytes)
            tmp.flush()
            return Document(tmp.name)

    def _docx_footer_xml_parts(self, docx_bytes):
        xml_map = self._docx_xml(docx_bytes)
        return {
            name: xml
            for name, xml in xml_map.items()
            if name.startswith("word/footer")
        }

    def test_mark_docx_fields_preserves_settings_ignorable_namespaces(self):
        source = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=["第一部分  前言", "正文"],
        ).getvalue()
        settings_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:settings xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
            'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
            'xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" '
            'mc:Ignorable="w14 w15 w16se"><w:zoom w:percent="100"/></w:settings>'
        )
        prepared = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(source), "r") as src, zipfile.ZipFile(prepared, "w", zipfile.ZIP_DEFLATED) as dst:
            for item in src.infolist():
                data = settings_xml.encode("utf-8") if item.filename == "word/settings.xml" else src.read(item.filename)
                dst.writestr(item, data)

        marked = app._mark_docx_fields_for_update(prepared.getvalue())
        with zipfile.ZipFile(io.BytesIO(marked), "r") as zf:
            updated_settings = zf.read("word/settings.xml").decode("utf-8")

        root_start = updated_settings[updated_settings.find("<w:settings"):updated_settings.find(">", updated_settings.find("<w:settings"))]
        self.assertIn('mc:Ignorable="w14 w15 w16se"', root_start)
        self.assertIn('xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"', root_start)
        self.assertIn('xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"', root_start)
        self.assertIn('xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"', root_start)
        self.assertIn("<w:updateFields", updated_settings)

    def _consume_response(self, response, *, as_text=False, parser=None):
        try:
            if parser is not None:
                return parser(response)
            return response.get_data(as_text=as_text)
        finally:
            response.close()

    def _assert_no_resource_warning(self, callback):
        with warnings.catch_warnings(record=True) as seen:
            warnings.simplefilter("always", ResourceWarning)
            result = callback()
            gc.collect()
        resource_warnings = [w for w in seen if issubclass(w.category, ResourceWarning)]
        self.assertEqual([], resource_warnings)
        return result

    def _build_review_docx_buffer(self, *, title_lines, body_lines, header_text="", extra_page_breaks=0):
        from docx import Document as DocxDocument
        from docx.enum.text import WD_BREAK

        doc = DocxDocument()
        if header_text:
            header = doc.sections[0].header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.text = header_text

        for line in title_lines:
            doc.add_paragraph(line)
        for line in body_lines:
            doc.add_paragraph(line)
        for _ in range(extra_page_breaks):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _build_tracked_custodian_docx_buffer(self):
        doc = Document()
        doc.add_paragraph("南方测试ETF托管协议")

        def add_change_paragraph(change_type, text):
            paragraph = doc.add_paragraph()
            change = OxmlElement(f"w:{change_type}")
            change.set(qn("w:id"), "1")
            change.set(qn("w:author"), "律所律师")
            change.set(qn("w:date"), "2026-05-18T00:00:00Z")
            run = OxmlElement("w:r")
            text_node = OxmlElement("w:delText" if change_type == "del" else "w:t")
            text_node.text = text
            run.append(text_node)
            change.append(run)
            paragraph._p.append(change)
            return paragraph

        add_change_paragraph("del", "一、被删除章节")
        doc.add_paragraph("被删除章节正文不应作为托管协议摘要章节。")
        add_change_paragraph("ins", "一、托管协议当事人")
        doc.add_paragraph("当事人修订正文。")
        add_change_paragraph("ins", "二、基金托管人对基金管理人的业务监督和核查")
        doc.add_paragraph("监督核查修订正文。")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _rewrite_docx_xml_buffer(self, buffer, replacements, extra_files=None):
        source = io.BytesIO(buffer.getvalue())
        target = io.BytesIO()
        with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    text = data.decode("utf-8")
                    for old, new in replacements:
                        self.assertIn(old, text)
                        text = text.replace(old, new, 1)
                    data = text.encode("utf-8")
                zout.writestr(item, data)
            for name, content in (extra_files or {}).items():
                zout.writestr(name, content.encode("utf-8"))
        target.seek(0)
        return target

    def _build_tracked_review_docx_buffer(self):
        base = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第一部分  前言",
                "合同正文段落",
                "第二十六部分  基金合同内容摘要",
                "摘要段落",
            ],
        )
        comments_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="9" w:author="律所律师" w:date="2026-05-16T08:00:00Z">
    <w:p><w:r><w:t>摘要口径需同步基金合同正文。</w:t></w:r></w:p>
  </w:comment>
</w:comments>"""
        return self._rewrite_docx_xml_buffer(
            base,
            [
                (
                    "<w:t>合同正文段落</w:t>",
                    '<w:t>合同正文段落</w:t><w:ins w:id="1" w:author="律所律师" w:date="2026-05-16T08:00:00Z"><w:r><w:t>新增正文意见</w:t></w:r></w:ins>',
                ),
                (
                    "<w:t>摘要段落</w:t>",
                    '<w:t>摘要段落</w:t><w:del w:id="2" w:author="律所律师" w:date="2026-05-16T08:05:00Z"><w:r><w:delText>删除摘要旧口径</w:delText></w:r></w:del><w:commentRangeStart w:id="9"/><w:r><w:t>摘要需确认文字</w:t></w:r><w:commentRangeEnd w:id="9"/><w:r><w:commentReference w:id="9"/></w:r>',
                ),
            ],
            {"word/comments.xml": comments_xml},
        )

    def _build_tracked_prospectus_docx_buffer(self):
        base = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "一、绪言",
                "招募正文段落",
                "基金合同内容摘要",
                "招募摘要段落",
            ],
        )
        return self._rewrite_docx_xml_buffer(
            base,
            [
                (
                    "<w:t>招募正文段落</w:t>",
                    '<w:t>招募正文段落</w:t><w:ins w:id="3" w:author="律所律师" w:date="2026-05-16T08:10:00Z"><w:r><w:t>新增招募意见</w:t></w:r></w:ins>',
                ),
                (
                    "<w:t>招募摘要段落</w:t>",
                    '<w:t>招募摘要段落</w:t><w:del w:id="4" w:author="律所律师" w:date="2026-05-16T08:15:00Z"><w:r><w:delText>删除招募摘要旧口径</w:delText></w:r></w:del>',
                ),
            ],
        )

    def _build_review_reader_docx_buffer(self):
        base = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第一部分  前言",
                "普通上下文段落",
                "费率表述段落",
                "第二十六部分  基金合同内容摘要",
                "摘要批注段落",
            ],
        )
        comments_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="11" w:author="律所律师" w:date="2026-05-16T09:00:00Z">
    <w:p><w:r><w:t>这里需要确认摘要是否同步。</w:t></w:r></w:p>
  </w:comment>
</w:comments>"""
        return self._rewrite_docx_xml_buffer(
            base,
            [
                (
                    "<w:t>普通上下文段落</w:t>",
                    '<w:t>普通上下文段落</w:t><w:ins w:id="7" w:author="律所律师" w:date="2026-05-16T09:05:00Z"><w:r><w:t>新增披露文字</w:t></w:r></w:ins>',
                ),
                (
                    "<w:t>费率表述段落</w:t>",
                    '<w:t>费率表述段落</w:t><w:del w:id="8" w:author="律所律师" w:date="2026-05-16T09:10:00Z"><w:r><w:delText>0.15%</w:delText></w:r></w:del><w:ins w:id="9" w:author="律所律师" w:date="2026-05-16T09:11:00Z"><w:r><w:t>0.20%</w:t></w:r></w:ins>',
                ),
                (
                    "<w:t>摘要批注段落</w:t>",
                    '<w:commentRangeStart w:id="11"/><w:r><w:t>摘要批注段落</w:t></w:r><w:commentRangeEnd w:id="11"/><w:r><w:commentReference w:id="11"/></w:r>',
                ),
            ],
            {"word/comments.xml": comments_xml},
        )

    def _build_review_reader_table_docx_buffer(self):
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("测试基金基金合同")
        doc.add_paragraph("第一部分  前言")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "项目"
        table.cell(0, 1).text = "表格费率段落"
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return self._rewrite_docx_xml_buffer(
            buffer,
            [
                (
                    "<w:t>表格费率段落</w:t>",
                    '<w:t>表格费率段落</w:t><w:del w:id="12" w:author="律所律师" w:date="2026-05-16T09:20:00Z"><w:r><w:delText>0.10%</w:delText></w:r></w:del><w:ins w:id="13" w:author="律所律师" w:date="2026-05-16T09:21:00Z"><w:r><w:t>0.12%</w:t></w:r></w:ins>',
                ),
            ],
        )

    def _build_review_rules_workbook(self, *, linked=False):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "说明" if linked else "总览"
        if linked:
            ws["A1"] = "项目"
            ws["B1"] = "内容"
            ws["A2"] = "基金类型"
            ws["B2"] = "ETF联接基金（发起式）"
            ws["A3"] = "重点结论"
            ws["B3"] = "招募说明书对合同存在细化披露。"
        else:
            ws["A1"] = "特别说明"
            ws["B1"] = "目录标题与正文章节标题可能不一致。"
            ws["A2"] = "本基金的几个关键判断"
            ws["A3"] = "1. 招募说明书与基金合同存在选择性差异。"

        if linked:
            ws = wb.create_sheet("招募-合同_总表")
            ws.append(["序号", "招募说明书章节", "对应基金合同章节", "关系判定", "勾稽结论"])
            ws.append([1, "基金募集", "第四部分基金份额的发售", "部分对应（招募细化）", "招募补充费率、公式和份额类别。"])

            ws = wb.create_sheet("招募-合同_明细")
            ws.append(["序号", "招募一级章节", "招募条款", "对应基金合同位置", "勾稽关系", "一致性判断", "主要差异/补充点", "备注"])
            ws.append([1, "基金募集", "一、募集期", "第四部分基金份额的发售—一、1、发售时间", "直接对应", "完全一致", "募集期最长不超过3个月", ""])

            ws = wb.create_sheet("正文-摘要_总表")
            ws.append(["序号", "基金合同正文章节", "对应摘要位置", "纳入情况", "勾稽结论"])
            ws.append([1, "第十六部分基金的收益与分配", "三、基金收益分配原则、执行方式", "部分纳入", "摘要仅保留执行性条款。"])

            ws = wb.create_sheet("正文-摘要_明细")
            ws.append(["序号", "基金合同正文位置", "摘要位置", "纳入方式", "一致性判断", "主要差异/省略点", "备注"])
            ws.append([1, "第十六部分—三、基金收益分配原则", "三、（一）基金收益分配原则", "摘录/转录", "完全一致", "仅编号层级调整", ""])

            ws = wb.create_sheet("重点差异")
            ws.append(["序号", "差异类别", "位置", "差异说明", "处理建议"])
            ws.append([1, "招募对合同的细化", "招募—基金募集", "招募补充A/C类份额与费率公式。", "审核募集安排时应优先核对招募参数。"])
        else:
            ws = wb.create_sheet("招募-合同(章节级)")
            ws.append(["序号", "招募说明书位置", "对应基金合同位置", "关系类型", "一致性判断", "文本相似度(参考)", "详细勾稽说明"])
            ws.append([1, "基金份额的申购赎回", "第八部分基金份额的申购与赎回", "直接对应（招募细化）", "基本一致", "0.66", "招募补充程序、公式与清单格式。"])

            ws = wb.create_sheet("招募-合同(详细)")
            ws.append(["序号", "招募章节", "招募内容点", "对应基金合同位置", "关系类型", "一致性判断", "文本相似度(参考)", "详细说明"])
            ws.append([1, "基金份额的申购赎回", "申购和赎回的对价、费用及其用途", "第八部分基金份额的申购与赎回", "直接对应（招募细化）", "基本一致", "0.60", "招募补充基金份额净值计算公式和代理券商佣金限额。"])

            ws = wb.create_sheet("正文-摘要(章节级)")
            ws.append(["序号", "基金合同正文章节", "基金合同摘要位置", "收录状态", "收录方式", "文本相似度(参考)", "详细说明"])
            ws.append([1, "第十八部分基金的收益与分配", "摘要三、基金收益分配原则、执行方式", "部分收录", "选择性摘录", "0.86", "摘要仅保留收益分配执行性条款。"])

            ws = wb.create_sheet("正文-摘要(详细)")
            ws.append(["序号", "基金合同摘要位置", "对应基金合同正文位置", "收录/未收录情况", "收录方式", "文本相似度(参考)", "详细说明"])
            ws.append([1, "一、摘要条款", "第三部分 目标章节", "已收录", "选择性摘录", "0.55", "摘要仅摘录本章的核心执行性条款。"])

            ws = wb.create_sheet("参考规则")
            ws.append(["基金类型", "合同", "招募", "对应关系", "提示词", "内容完全一致", "内容有差异"])
            ws.append(["ETF", "基金份额的上市交易", "基金份额的上市交易", "招募有补充", "重点核对IOPV公式", "一、基金份额的上市；二、基金份额的上市交易", "四、基金份额参考净值（IOPV）的计算与公告"])

        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.close()
        wb.save(tmp.name)
        return tmp.name

    def _build_summary_detail_locator_workbook(self, contract_locator, summary_pos="一、摘要条款"):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "总览"
        ws["A1"] = "特别说明"
        ws["B1"] = "摘要复核优先按明细规则定位到小标题。"

        ws = wb.create_sheet("正文-摘要(详细)")
        ws.append(["序号", "基金合同摘要位置", "对应基金合同正文位置", "收录/未收录情况", "收录方式", "文本相似度(参考)", "详细说明"])
        ws.append([1, summary_pos, contract_locator, "已收录", "选择性摘录", "0.73", "摘要应按Excel明细规则定位到指定小标题。"])

        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.close()
        wb.save(tmp.name)
        return tmp.name

    def _build_summary_chapter_locator_workbook(self, contract_chapter, summary_pos):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "总览"
        ws["A1"] = "特别说明"
        ws["B1"] = "章节级规则命中后仍需继续锁定章内小标题。"

        ws = wb.create_sheet("正文-摘要(章节级)")
        ws.append(["序号", "基金合同正文章节", "基金合同摘要位置", "收录状态", "收录方式", "文本相似度(参考)", "详细说明"])
        ws.append([1, contract_chapter, summary_pos, "部分收录", "选择性摘录", "0.81", "先锁定章节，再按摘要标题继续定位小标题。"])

        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.close()
        wb.save(tmp.name)
        return tmp.name

    def _build_summary_rules_workbook(self, *, detail_rows=None, chapter_rows=None):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "总览"
        ws["A1"] = "特别说明"
        ws["B1"] = "摘要复核测试专用规则表。"

        if chapter_rows:
            ws = wb.create_sheet("正文-摘要(章节级)")
            ws.append(["序号", "基金合同正文章节", "基金合同摘要位置", "收录状态", "收录方式", "文本相似度(参考)", "详细说明"])
            for index, row in enumerate(chapter_rows, 1):
                ws.append([
                    index,
                    row.get("contract_pos", ""),
                    row.get("summary_pos", ""),
                    row.get("status", "部分收录"),
                    row.get("method", "选择性摘录"),
                    row.get("similarity", "0.80"),
                    row.get("detail", ""),
                ])

        if detail_rows:
            ws = wb.create_sheet("正文-摘要(详细)")
            ws.append(["序号", "基金合同摘要位置", "对应基金合同正文位置", "收录/未收录情况", "收录方式", "文本相似度(参考)", "详细说明"])
            for index, row in enumerate(detail_rows, 1):
                ws.append([
                    index,
                    row.get("summary_pos", ""),
                    row.get("contract_pos", ""),
                    row.get("status", "已收录"),
                    row.get("method", "近全文摘录"),
                    row.get("similarity", "0.95"),
                    row.get("detail", ""),
                ])

        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.close()
        wb.save(tmp.name)
        return tmp.name

    def _build_cross_check_workbook(self, rows):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "规则"
        ws.append(["基金类型", "合同", "招募", "对应关系", "提示语", "完全一致内容", "差异内容"])
        for row in rows:
            ws.append(row)

        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.close()
        wb.save(tmp.name)
        return tmp.name

    def _run_cross_check_with_review_rules(self, contract_sections, prospectus_sections, review_rules, *, fund_type="ETF"):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        original_load_review_rules = app._load_review_rules
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": contract_sections,
            "prospectus_sections": prospectus_sections,
        }

        try:
            app._load_review_rules = lambda _fund_type=fund_type: review_rules
            response = client.post("/api/review/cross_check", json={"fund_type": fund_type})
            self.assertEqual(response.status_code, 200)
            return self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            app._load_review_rules = original_load_review_rules
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

    def _with_review_rules(self, path):
        class _Ctx:
            def __init__(self, outer, workbook_path):
                self.outer = outer
                self.workbook_path = workbook_path
                self.original = app._review_store.get("rules_xlsx_path")

            def __enter__(self):
                app._review_store["rules_xlsx_path"] = self.workbook_path
                return self.workbook_path

            def __exit__(self, exc_type, exc, tb):
                if self.original is None:
                    app._review_store.pop("rules_xlsx_path", None)
                else:
                    app._review_store["rules_xlsx_path"] = self.original
                Path(self.workbook_path).unlink(missing_ok=True)

        return _Ctx(self, path)

    def _build_etf_cross_check_detail_workbook(self):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "总览"
        ws["A1"] = "基金类型"
        ws["B1"] = "ETF"
        ws["A2"] = "特别说明"
        ws["B2"] = "合同摘要不得并入签署页；申购赎回与上市交易章节按明细规则优先复核。"

        ws = wb.create_sheet("招募-合同(章节级)")
        ws.append(["序号", "招募说明书位置", "对应基金合同位置", "关系类型", "一致性判断", "文本相似度(参考)", "详细勾稽说明"])
        ws.append([1, "基金份额的上市交易", "第七部分基金份额的上市交易", "直接对应（招募细化）", "基本一致", "0.91", "除IOPV外其余条款基本一致。"])
        ws.append([2, "基金份额的申购赎回", "第八部分基金份额的申购与赎回", "直接对应（招募细化）", "基本一致", "0.66", "合同提供制度框架；招募细化程序、单位、对价和清单格式。"])

        ws = wb.create_sheet("招募-合同(详细)")
        ws.append(["序号", "招募章节", "招募内容点", "对应基金合同位置", "关系类型", "一致性判断", "文本相似度(参考)", "详细说明"])
        ws.append([1, "基金份额的上市交易", "一、基金份额的上市", "第七部分基金份额的上市交易 / 一、基金份额的上市", "直接对应", "完全一致", "1", "对应标题一致，招募说明书与基金合同对应条款实质一致。"])
        ws.append([2, "基金份额的上市交易", "二、基金份额的上市交易", "第七部分基金份额的上市交易 / 二、基金份额的上市交易", "直接对应", "完全一致", "1", "对应标题一致，招募说明书与基金合同对应条款实质一致。"])
        ws.append([3, "基金份额的上市交易", "三、终止上市交易", "第七部分基金份额的上市交易 / 三、终止上市交易", "直接对应", "完全一致", "1", "对应标题一致，招募说明书与基金合同对应条款实质一致。"])
        ws.append([4, "基金份额的上市交易", "四、基金份额参考净值（IOPV）的计算与公告", "第七部分基金份额的上市交易 / 四、基金份额参考净值（IOPV）的计算与公告", "直接对应（招募细化）", "部分对应", "0.6393", "合同仅原则性规定基金份额参考净值计算方法详见招募说明书；招募补充IOPV计算公式、保留位数和调整口径。"])
        ws.append([5, "基金份额的上市交易", "五、在不违反法律法规及不损害基金份额持有人利益的前提下，本基金可以申请在其他证券交易所（含境外证券交易所）上市交易，而无需召开基金份额持有人大会审议。", "第七部分基金份额的上市交易 / 五、在不违反法律法规及不损害基金份额持有人利益的前提下，本基金可以申请在其他证券交易所（含境外证券交易所）上市交易，而无需召开基金份额持有人大会审议。", "直接对应", "完全一致", "1", "对应标题一致，招募说明书与基金合同对应条款实质一致。"])
        ws.append([6, "基金份额的上市交易", "六、相关法律法规、中国证监会及上海证券交易所对基金上市交易的规则等相关规定内容进行调整的，基金合同相应予以修改，并按照新规定执行，且此项修改无须召开基金份额持有人大会。", "第七部分基金份额的上市交易 / 六、相关法律法规、中国证监会及上海证券交易所对基金上市交易的规则等相关规定内容进行调整的，本基金合同相应予以修改，并按照新规定执行，且此项修改无须召开基金份额持有人大会。", "直接对应", "基本一致", "0.9941", "与合同基本一致，仅“基金合同/本基金合同”等个别自指表述存在细微差异。"])
        ws.append([7, "基金份额的上市交易", "七、若上海证券交易所、中国证券登记结算有限责任公司增加了基金上市交易的新功能，基金管理人可以在履行适当的程序后增加相应功能。", "第七部分基金份额的上市交易 / 七、若上海证券交易所、中国证券登记结算有限责任公司增加了基金上市交易的新功能，基金管理人可以在履行适当的程序后增加相应功能。", "直接对应", "完全一致", "1", "对应标题一致，招募说明书与基金合同对应条款实质一致。"])
        ws.append([8, "基金份额的申购赎回", "一、申购和赎回场所", "第八部分基金份额的申购与赎回 / 一、申购和赎回场所", "直接对应", "完全一致", "1", "对应条款一致。"])
        ws.append([9, "基金份额的申购赎回", "二、申购和赎回的开放日及时间", "第八部分基金份额的申购与赎回 / 二、申购和赎回的开放日及时间", "直接对应", "基本一致", "0.9989", "对应条款基本一致，文字表述仅有细微差异。"])
        ws.append([10, "基金份额的申购赎回", "三、申购和赎回的原则", "第八部分基金份额的申购与赎回 / 三、申购与赎回的原则", "直接对应", "基本一致", "0.9965", "对应条款基本一致，“申购和赎回/申购与赎回”仅为连词差异。"])
        ws.append([11, "基金份额的申购赎回", "四、申购和赎回的程序", "第八部分基金份额的申购与赎回 / 四、申购与赎回的程序", "直接对应（招募细化）", "部分对应", "0.7812", "合同规定确认、清算交收与登记的原则框架，并明确具体在招募说明书中列示；招募进一步写明T日/T+1/T+2清算流程、交收节奏和投资人当日可卖出安排。"])
        ws.append([12, "基金份额的申购赎回", "五、申购和赎回的数额限制", "第八部分基金份额的申购与赎回 / 五、申购和赎回的数额限制", "直接对应（招募细化）", "基本一致", "0.9037", "合同仅要求最小申购赎回单位由管理人确定并在招募说明书披露；招募明确当前最小申购赎回单位为100万份，并补充上限及规模控制提示。"])
        ws.append([13, "基金份额的申购赎回", "六、申购、赎回的对价、费用及其用途", "第八部分基金份额的申购与赎回 / 六、申购和赎回的对价、费用及其用途", "直接对应（招募细化）", "部分对应", "0.9133", "合同原则规定对价和费用安排；招募增加基金份额净值计算公式、T日/T+1公告口径和代理券商佣金上限等细节。"])
        ws.append([14, "基金份额的申购赎回", "七、申购赎回清单的内容与格式", "基金合同无同名独立条款（系第八部分第六项和第四项的细化展开）", "招募细化", "无直接对应", "", "申购赎回清单的构成、现金替代分类、计算公式和格式示例等为招募说明书对合同规则的操作化细化，合同正文未单独成章。"])
        ws.append([15, "基金份额的申购赎回", "八、拒绝或暂停申购的情形", "第八部分基金份额的申购与赎回 / 七、拒绝或暂停申购的情形", "直接对应", "基本一致", "0.9989", "对应条款基本一致。"])
        ws.append([16, "基金份额的申购赎回", "九、暂停赎回或延缓支付赎回对价的情形", "第八部分基金份额的申购与赎回 / 八、暂停赎回或延缓支付赎回对价的情形", "直接对应", "基本一致", "0.9987", "对应条款基本一致。"])
        ws.append([17, "基金份额的申购赎回", "十、其他申购赎回方式", "第八部分基金份额的申购与赎回 / 九、其他申购赎回方式", "直接对应", "基本一致", "0.9969", "对应条款基本一致。"])
        ws.append([18, "基金份额的申购赎回", "十一、基金的非交易过户", "第八部分基金份额的申购与赎回 / 十、基金的非交易过户", "直接对应", "基本一致", "0.9984", "对应条款基本一致。"])
        ws.append([19, "基金份额的申购赎回", "十二、基金份额的冻结和解冻", "第八部分基金份额的申购与赎回 / 十一、基金份额的冻结和解冻", "直接对应", "基本一致", "0.9936", "对应条款基本一致。"])
        ws.append([20, "基金份额的申购赎回", "十三、基金份额的转让", "第八部分基金份额的申购与赎回 / 十二、基金份额的转让", "直接对应", "基本一致", "0.9932", "对应条款基本一致。"])
        ws.append([21, "基金份额的申购赎回", "十四、其他业务", "第八部分基金份额的申购与赎回 / 十三、其他业务", "直接对应", "基本一致", "0.9912", "对应条款基本一致。"])
        ws.append([22, "基金份额的申购赎回", "十五、基金合同生效后，若上海证券交易所和中国证券登记结算有限责任公司针对交易型开放式证券投资基金推出新的清算交收与登记模式并引入新的申购、赎回方式，本基金管理人有权调整本基金的清算交收与登记模式及申购、赎回方式，或新增本基金的清算交收与登记模式并引入新的申购、赎回方式，届时将发布公告予以披露并在本基金招募说明书及其更新中予以更新，无须召开基金份额持有人大会审议。", "第八部分基金份额的申购与赎回 / 十四、本基金合同生效后，若上海证券交易所和中国证券登记结算有限责任公司针对交易型开放式证券投资基金推出新的清算交收与登记模式并引入新的申购、赎回方式，本基金管理人有权调整本基金的清算交收与登记模式及申购、赎回方式，或新增本基金的清算交收与登记模式并引入新的申购、赎回方式，届时将发布公告予以披露并在本基金招募说明书及其更新中予以更新，无须召开基金份额持有人大会审议。", "直接对应", "基本一致", "0.9918", "对应条款基本一致。"])

        ws = wb.create_sheet("参考规则")
        ws.append(["基金类型", "合同", "招募", "对应关系", "提示词", "内容完全一致", "内容有差异"])
        ws.append(["ETF", "基金份额的上市交易", "基金份额的上市交易", "招募有补充", "重点核对IOPV公式", "一、基金份额的上市；二、基金份额的上市交易；三、终止上市交易；五、在其他证券交易所上市交易；七、新功能扩展", "四、基金份额参考净值（IOPV）的计算与公告"])
        ws.append(["ETF", "基金份额的申购与赎回", "基金份额的申购赎回", "招募有补充", "重点核对程序、最小单位、对价费用和申购赎回清单", "一、申购和赎回场所；二、申购和赎回的开放日及时间；三、申购与赎回的原则；八、拒绝或暂停申购的情形；九、暂停赎回或延缓支付赎回对价的情形；十、其他申购赎回方式；十一、基金的非交易过户；十二、基金份额的冻结和解冻；十三、基金份额的转让；十四、其他业务；十五、清算交收与登记模式调整", "四、申购和赎回的程序；五、申购和赎回的数额限制；六、申购、赎回的对价、费用及其用途；七、申购赎回清单的内容与格式"])

        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.close()
        wb.save(tmp.name)
        return tmp.name

    def test_open_browser_helper_is_defined(self):
        self.assertTrue(callable(app.open_browser))

    def test_market_scope_defaults_to_single_for_kechuang(self):
        derived = self.engine._derive_variables(self.build_form(MARKET_TYPE="KECHUANG"))
        self.assertEqual(derived["MARKET_SCOPE"], "SINGLE_MARKET")

    def test_market_scope_defaults_to_cross_for_standard_a_share(self):
        derived = self.engine._derive_variables(self.build_form(MARKET_TYPE="A_SHARE"))
        self.assertEqual(derived["MARKET_SCOPE"], "CROSS_MARKET")

    def test_market_scope_override_is_preserved(self):
        derived = self.engine._derive_variables(self.build_form(MARKET_TYPE="A_SHARE", MARKET_SCOPE="SINGLE_MARKET"))
        self.assertEqual(derived["MARKET_SCOPE"], "SINGLE_MARKET")

    def test_single_market_types_force_single_market_scope_even_when_form_posts_cross_market(self):
        kechuang = self.engine._derive_variables(
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="KECHUANG", MARKET_SCOPE="CROSS_MARKET")
        )
        chuangye = self.engine._derive_variables(
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="CHUANGYE", MARKET_SCOPE="CROSS_MARKET")
        )

        self.assertEqual("SINGLE_MARKET", kechuang["MARKET_SCOPE"])
        self.assertEqual("SINGLE_MARKET", chuangye["MARKET_SCOPE"])

    def test_stock_subscription_defaults_true_for_all_non_hk_etfs(self):
        sse_cross = self.engine._derive_variables(
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="CROSS_MARKET")
        )
        sse_single = self.engine._derive_variables(
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="SINGLE_MARKET")
        )
        szse_single = self.engine._derive_variables(
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="CHUANGYE", MARKET_SCOPE="SINGLE_MARKET")
        )
        hk = self.engine._derive_variables(
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT")
        )

        self.assertTrue(sse_cross["HAS_STOCK_SUBSCRIPTION"])
        self.assertTrue(sse_single["HAS_STOCK_SUBSCRIPTION"])
        self.assertTrue(szse_single["HAS_STOCK_SUBSCRIPTION"])
        self.assertFalse(hk["HAS_STOCK_SUBSCRIPTION"])

    def test_product_type_defaults_to_etf(self):
        derived = self.engine._derive_variables(self.build_form())
        self.assertEqual(derived["PRODUCT_TYPE"], "ETF")

    def test_derive_variables_auto_fills_csindex_provider_and_fund_names_from_index_name(self):
        derived = self.engine._derive_variables(
            self.build_form(
                FUND_NAME="",
                FUND_SHORT_NAME="",
                INDEX_NAME="上证科创板人工智能指数",
                INDEX_COMPILER="",
                INDEX_WEBSITE="",
            )
        )
        self.assertEqual("中证指数有限公司", derived["INDEX_COMPILER"])
        self.assertEqual("https://www.csindex.com.cn", derived["INDEX_WEBSITE"])
        self.assertEqual("https://www.csindex.com.cn", derived["INDEX_COMPILER_web"])
        self.assertEqual("南方上证科创板人工智能交易型开放式指数证券投资基金", derived["FUND_NAME"])
        self.assertEqual("南方上证科创板人工智能ETF", derived["FUND_SHORT_NAME"])

    def test_derive_variables_prefers_cnindex_provider_for_chuangyeban_keywords(self):
        derived = self.engine._derive_variables(
            self.build_form(
                FUND_NAME="",
                FUND_SHORT_NAME="",
                INDEX_NAME="中证创业板成长指数",
                INDEX_COMPILER="",
                INDEX_WEBSITE="",
            )
        )
        self.assertEqual("深圳证券信息有限公司", derived["INDEX_COMPILER"])
        self.assertEqual("http://www.cnindex.com.cn", derived["INDEX_WEBSITE"])
        self.assertEqual("南方中证创业板成长交易型开放式指数证券投资基金", derived["FUND_NAME"])
        self.assertEqual("南方中证创业板成长ETF", derived["FUND_SHORT_NAME"])

    def test_derive_variables_maps_hengsheng_index_to_cnindex_rule(self):
        derived = self.engine._derive_variables(
            self.build_form(
                FUND_NAME="",
                FUND_SHORT_NAME="",
                INDEX_NAME="恒生科技指数",
                INDEX_COMPILER="",
                INDEX_WEBSITE="",
            )
        )
        self.assertEqual("深圳证券信息有限公司", derived["INDEX_COMPILER"])
        self.assertEqual("http://www.cnindex.com.cn", derived["INDEX_WEBSITE"])
        self.assertEqual("南方恒生科技交易型开放式指数证券投资基金", derived["FUND_NAME"])
        self.assertEqual("南方恒生科技ETF", derived["FUND_SHORT_NAME"])

    def test_variant_key_uses_six_way_matrix(self):
        self.assertEqual(
            self.engine._get_prospectus_variant_key(self.build_form(EXCHANGE="SSE", MARKET_TYPE="KECHUANG")),
            "SSE_SINGLE",
        )
        self.assertEqual(
            self.engine._get_prospectus_variant_key(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")),
            "SSE_CROSS",
        )
        self.assertEqual(
            self.engine._get_prospectus_variant_key(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="CHUANGYE")),
            "SZSE_SINGLE",
        )
        self.assertEqual(
            self.engine._get_prospectus_variant_key(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE")),
            "SZSE_CROSS",
        )
        self.assertEqual(
            self.engine._get_prospectus_variant_key(self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT")),
            "SSE_HK",
        )

    def test_variant_key_ignores_cross_market_scope_for_single_market_types(self):
        self.assertEqual(
            "SSE_SINGLE",
            self.engine._get_prospectus_variant_key(
                self.build_form(EXCHANGE="SSE", MARKET_TYPE="KECHUANG", MARKET_SCOPE="CROSS_MARKET")
            ),
        )
        self.assertEqual(
            "SZSE_SINGLE",
            self.engine._get_prospectus_variant_key(
                self.build_form(EXCHANGE="SZSE", MARKET_TYPE="CHUANGYE", MARKET_SCOPE="CROSS_MARKET")
            ),
        )
        self.assertEqual(
            self.engine._get_prospectus_variant_key(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT")),
            "SZSE_HK",
        )

    def test_hk_connect_derives_exchange_specific_variant_and_business_rules_type(self):
        derived_sse = app.engine._derive_variables(
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT", DISTRIBUTION_FREQ="MONTHLY")
        )
        derived_sz = app.engine._derive_variables(
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT", DISTRIBUTION_FREQ="QUARTERLY")
        )

        self.assertEqual("SSE_HK", derived_sse["HK_CONNECT_EXCHANGE_VARIANT"])
        self.assertEqual("SSE", derived_sse["BUSINESS_RULES_TYPE"])
        self.assertEqual("SZSE_HK", derived_sz["HK_CONNECT_EXCHANGE_VARIANT"])
        self.assertEqual("CSDC", derived_sz["BUSINESS_RULES_TYPE"])

    def test_reference_prospectus_resolution_prefers_repo_manual_docx_over_legacy(self):
        temp_path = Path("D:/codex/_tmp_reference_resolution")
        if temp_path.exists():
            shutil.rmtree(temp_path)
        temp_path.mkdir(parents=True)
        repo_docx = temp_path / "repo.docx"
        legacy_docx = temp_path / "legacy.docx"
        repo_docx.write_bytes(b"repo")
        legacy_docx.write_bytes(b"legacy")

        original_repo_map = dict(app.REFERENCE_PROSPECTUS_DOCX_MAP)
        original_legacy_map = dict(app.LEGACY_REFERENCE_PROSPECTUS_DOCX_MAP)
        original_env = os.environ.get("ETF_REFERENCE_PROSPECTUS_DIR")
        try:
            os.environ.pop("ETF_REFERENCE_PROSPECTUS_DIR", None)
            app.REFERENCE_PROSPECTUS_DOCX_MAP = {**original_repo_map, "SSE_CROSS": repo_docx}
            app.LEGACY_REFERENCE_PROSPECTUS_DOCX_MAP = {**original_legacy_map, "SSE_CROSS": legacy_docx}
            self.assertEqual(repo_docx, app._resolve_reference_prospectus_docx("SSE_CROSS"))
        finally:
            app.REFERENCE_PROSPECTUS_DOCX_MAP = original_repo_map
            app.LEGACY_REFERENCE_PROSPECTUS_DOCX_MAP = original_legacy_map
            if original_env is None:
                os.environ.pop("ETF_REFERENCE_PROSPECTUS_DIR", None)
            else:
                os.environ["ETF_REFERENCE_PROSPECTUS_DIR"] = original_env
            if temp_path.exists():
                shutil.rmtree(temp_path)

    def test_chapter_four_is_placeholder_only(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "四")
        self.assertEqual(chapter.splitlines()[1:], ["【托管人情况待填写】"])

    def test_chapter_twenty_one_adds_placeholder_under_each_title(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "\u4e8c\u5341\u4e00")
        self.assertIn("一、基金托管协议当事人", chapter)
        self.assertIn("二、基金托管人对基金管理人的业务监督和核查", chapter)
        self.assertIn("八、基金托管协议的变更、终止与基金财产的清算", chapter)
        self.assertIn("【待填写】", chapter)

    def test_basket_is_derived_from_min_sub_unit_in_wan_units(self):
        variables = self.engine._inject_clause_texts(
            self.engine._derive_variables(self.build_form(MIN_SUB_UNIT="250,000份"))
        )
        self.assertEqual("25", variables["BASKET"])

    def test_generated_prospectus_keeps_basket_style_for_non_wan_min_sub_unit(self):
        text = self.engine.generate(self.build_form(MIN_SUB_UNIT="250,000份"))
        sec5 = self.extract_section(self.extract_chapter(text, "十"), "五")
        risk = self.extract_chapter(text, "十八")

        self.assertIn("25万份", sec5)
        self.assertNotIn("250,000份", sec5)
        self.assertIn("25万份", risk)
        self.assertNotIn("250,000份", risk)

    def test_chapter_ten_uses_frontend_min_sub_unit_value(self):
        text = self.engine.generate(self.build_form(MIN_SUB_UNIT="3,000,000份"))
        chapter = self.extract_chapter(text, "十")
        sec5 = self.extract_section(chapter, "五")
        self.assertIn("300万份", sec5)
        self.assertNotIn("3,000,000份", sec5)

    def test_chapter_ten_sec5_keeps_reference_sentence_and_only_updates_min_sub_unit_number(self):
        text = self.engine.generate(self.build_form(MIN_SUB_UNIT="2,000,000份"))
        chapter = self.extract_chapter(text, "十")
        sec5 = self.extract_section(chapter, "五")
        self.assertIn(
            "1、投资人申购、赎回的基金份额需为最小申购赎回单位的整数倍。目前，本基金最小申购赎回单位为200万份",
            sec5,
        )
        self.assertNotIn("需为最小申购赎回单位为", sec5)
        self.assertNotIn("本基金目前最小申购赎回单位为", sec5)

    def test_chapter_ten_sec5_uses_frontend_min_sub_unit_in_szse_a_share_branch(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE", MIN_SUB_UNIT="2,000,000份"))
        chapter = self.extract_chapter(text, "十")
        sec5 = self.extract_section(chapter, "五")
        self.assertIn("投资人申购、赎回的基金份额需为最小申购赎回单位的整数倍。目前，本基金最小申购赎回单位为200万份", sec5)
        self.assertNotIn("需为最小申购赎回单位为", sec5)
        self.assertNotIn("150万份", sec5)

    def test_important_notice_normalizes_szse_index_compiler_name(self):
        text = self.engine.generate(
            self.build_form(
                EXCHANGE="SZSE",
                MARKET_TYPE="CHUANGYE",
                INDEX_NAME="创业板中盘200指数",
                INDEX_DESCRIPTION="创业板中盘200指数反映创业板市值规模中等、流动性较好的200家公司股价变化情况。",
                INDEX_COMPILER="深证信息有限公司",
                INDEX_WEBSITE="www.cnindex.com.cn",
            )
        )
        notice = self.extract_important_notice(text)
        self.assertIn("深圳证券信息有限公司网站", notice)
        self.assertNotIn("深证信息有限公司网站", notice)

    def test_chapter_ten_sec2_removes_extra_zhi_character_from_open_time_sentence(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"))
        chapter = self.extract_chapter(text, "十")
        sec2 = self.extract_section(chapter, "二")
        self.assertIn("具体办理时间为深圳证券交易所的正常交易日的交易时间", sec2)
        self.assertNotIn("具体办理时间为指深圳证券交易所的正常交易日的交易时间", sec2)

    def test_chapter_ten_sec6_updates_fee_wording_and_removes_deleted_sentences(self):
        expected = "4、投资人在申购或赎回基金份额时，申购赎回代理券商可按不超过0.2%的标准收取佣金，其中包含证券交易所、登记机构等收取的相关费用。"
        forms = [
            self.build_form(),
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT"),
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"),
        ]
        for form in forms:
            with self.subTest(exchange=form["EXCHANGE"], market_type=form["MARKET_TYPE"]):
                text = self.engine.generate(form)
                chapter = self.extract_chapter(text, "十")
                sec6 = self.extract_section(chapter, "六")
                exchange_name = "上海证券交易所" if form["EXCHANGE"] == "SSE" else "深圳证券交易所"
                self.assertIn(f"3、申购赎回清单由基金管理人编制。T日的申购赎回清单在当日{exchange_name}开市前公告。", sec6)
                self.assertIn(expected, sec6)
                self.assertNotIn("申购赎回清单的内容与格式示例参见招募说明书", sec6)
                self.assertNotIn("具体规定请参见招募说明书及基金产品资料概要", sec6)
                self.assertNotIn("按照一定的标准收取佣金", sec6)

    def test_chapter_ten_sec10_removes_ben_from_other_redemption_methods_clause(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "十")
        self.assertIn("十、其他申购赎回方式", chapter)
        self.assertIn("6、基金管理人指定的代理机构可依据基金合同开展其他服务，双方需签订书面委托代理协议。", chapter)
        self.assertNotIn("依据本基金合同开展其他服务", chapter)

    def test_risk_chapter_updates_min_sub_unit_in_liquidity_risk_sentence(self):
        text = self.engine.generate(self.build_form(MIN_SUB_UNIT="2,000,000份"))
        risk = self.extract_chapter(text, "十八")
        self.assertIn(
            "①本基金最小申购赎回单位设置较高（目前为200万份），中小投资者只能在二级市场上按交易价格卖出基金份额。",
            risk,
        )

    def test_chapter_ten_uses_default_min_sub_unit_when_missing(self):
        form = self.build_form()
        form.pop("MIN_SUB_UNIT")
        text = self.engine.generate(form)
        chapter = self.extract_chapter(text, "十")
        sec5 = self.extract_section(chapter, "五")
        self.assertIn("100万份", sec5)
        self.assertNotIn("1,000,000份", sec5)
        self.assertNotIn("（即100万份）", sec5)

        sh_text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        sz_text = self.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"))
        sh_ch6 = self.extract_section(self.extract_chapter(sh_text, "六"), "七")
        sz_ch6 = self.extract_section(self.extract_chapter(sz_text, "六"), "七")
        self.assertNotEqual(sh_ch6, sz_ch6)
        sh_ch9 = self.extract_chapter(sh_text, "九")
        sz_ch9 = self.extract_chapter(sz_text, "九")
        self.assertNotEqual(sh_ch9, sz_ch9)
        sh_ch10 = self.extract_section(self.extract_chapter(sh_text, "十"), "四")
        sz_ch10 = self.extract_section(self.extract_chapter(sz_text, "十"), "四")
        self.assertNotEqual(sh_ch10, sz_ch10)

    def test_risk_disclosure_uses_market_specific_text(self):
        hk_text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT"))
        kechuang_text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="KECHUANG"))
        chuangye_text = self.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="CHUANGYE"))
        a_share_text = self.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"))

        self.assertIn("港股通股票", self.extract_chapter(hk_text, "十八"))
        self.assertIn("科创板股票", self.extract_chapter(kechuang_text, "十八"))
        self.assertIn("创业板股票", self.extract_chapter(chuangye_text, "十八"))
        self.assertIn("本基金属于股票型基金", self.extract_chapter(a_share_text, "十八"))

    def test_important_notice_prefers_manual_reference_subscription_account_clause(self):
        form = self.build_form(
            EXCHANGE="SSE",
            MARKET_TYPE="HK_CONNECT",
            IMPORTANT_NOTICE_STYLE="DIVIDEND_QUALITY",
            INDEX_NAME="中证测试指数",
        )
        reference_notice = self.engine._load_reference_fixed_content(form).get("important_notice", "")
        expected_line = next(
            line.strip()
            for line in reference_notice.splitlines()
            if "基金账户只能进行基金的现金认购和二级市场交易" in line
        )

        notice = self.extract_important_notice(self.engine.generate(form))

        self.assertIn(expected_line, notice)
        self.assertNotIn("{SUBSCRIPTION_ACCOUNT_CLAUSE}", notice)

    def test_chapter_nine_prefers_manual_reference_body(self):
        form = self.build_form(
            EXCHANGE="SZSE",
            MARKET_TYPE="HK_CONNECT",
            IMPORTANT_NOTICE_STYLE="CHUANGYE_200",
            INDEX_NAME="国证测试指数",
        )
        reference_body = self.engine._load_reference_fixed_content(form).get("九", {}).get("body", "")
        self.assertTrue(reference_body)

        expected = self.engine._normalize_top_level_section_spacing(
            self.engine._ensure_chapter_nine_body(
                self.engine._replace_reference_fund_name(reference_body, form)
            )
        )
        actual = self.engine._chapter_body_text(self.engine.generate(form), "九")

        self.assertEqual(expected, actual)

    def test_prospectus_template_marks_chapter_six_section_four_as_manual_reference(self):
        template = (app.BASE_DIR / "02_招募说明书模板.md").read_text(encoding="utf-8")

        self.assertIn("四、发售方式", template)
        self.assertIn("【按当前ETF类型对应的manual参考招募说明书生成】", template)
        self.assertNotIn("{METHOD_SUBSCRIBE_DERIVED_FROM_CONTRACT}", template)

    def test_prospectus_template_marks_chapter_six_section_seven_as_manual_reference(self):
        template = (app.BASE_DIR / "02_招募说明书模板.md").read_text(encoding="utf-8")

        self.assertIn("七、认购开户", template)
        self.assertIn("【按当前ETF类型对应的manual参考招募说明书生成】", template)
        self.assertNotIn("{SUB_ACCOUNT_OPENING_CLAUSE}", template)

    def test_prospectus_template_marks_chapter_nine_sections_as_manual_reference(self):
        template = (app.BASE_DIR / "02_招募说明书模板.md").read_text(encoding="utf-8")

        self.assertIn("基金份额的上市交易", template)
        self.assertNotIn("{LISTING_DERIVED_FROM_CONTRACT_SEC7_1}", template)
        self.assertNotIn("{LISTING_DERIVED_FROM_CONTRACT_SEC7_2}", template)
        self.assertNotIn("{LISTING_DERIVED_FROM_CONTRACT_SEC7_3}", template)
        self.assertNotIn("{LISTING_DERIVED_FROM_CONTRACT_SEC7_5}", template)
        self.assertNotIn("{LISTING_DERIVED_FROM_CONTRACT_SEC7_6}", template)
        self.assertNotIn("{LISTING_DERIVED_FROM_CONTRACT_SEC7_7}", template)

    def test_risk_chapter_prefers_manual_reference_body(self):
        cases = (
            {
                "EXCHANGE": "SSE",
                "MARKET_TYPE": "KECHUANG",
                "IMPORTANT_NOTICE_STYLE": "DIVIDEND_QUALITY",
                "INDEX_NAME": "中证测试指数",
            },
            {
                "EXCHANGE": "SZSE",
                "MARKET_TYPE": "CHUANGYE",
                "IMPORTANT_NOTICE_STYLE": "CHUANGYE_200",
                "INDEX_NAME": "国证测试指数",
            },
            {
                "EXCHANGE": "SSE",
                "MARKET_TYPE": "HK_CONNECT",
                "IMPORTANT_NOTICE_STYLE": "DIVIDEND_QUALITY",
                "INDEX_NAME": "中证测试指数",
            },
        )

        for overrides in cases:
            with self.subTest(**overrides):
                form = self.build_form(**overrides)
                variables = self.engine._inject_clause_texts(dict(form))
                reference_body = self.engine._load_reference_fixed_content(form).get("十八", {}).get("body", "")
                self.assertTrue(reference_body)

                expected = self.engine._normalize_top_level_section_spacing(
                    self.engine._normalize_prospectus_risk_chapter(
                        self.engine._replace_reference_fund_name(reference_body, variables),
                        variables,
                    )
                )
                actual = self.engine._chapter_body_text(self.engine.generate(form), "十八")

                self.assertEqual(expected, actual)

    def test_definitions_are_normalized_for_prospectus_wording(self):
        form = self.build_form()
        variables = self.engine._extract_contract_sections(form)
        text = self.engine.generate(form)
        chapter = self.extract_chapter(text, "二")
        chapter_body = chapter.splitlines()[1:]
        while chapter_body and not chapter_body[0].strip():
            chapter_body.pop(0)
        chapter_text = "\n".join(chapter_body).strip()
        self.assertNotEqual(chapter_text, variables["CONTRACT_DEFS_TEXT"].strip())
        self.assertNotIn("在本基金合同中", chapter)
        self.assertIn("在本招募说明书中", chapter)
        self.assertIn("4、基金合同：", chapter)
        self.assertIn("6、招募说明书或本招募说明书：", chapter)
        self.assertNotRegex(chapter, r"\d+、不可抗力：指本基金合同当事人")
        self.assertRegex(chapter, r"\d+、不可抗力：指基金合同当事人")

    def test_generated_prospectus_replaces_ben_jijinhetong_outside_contract_summary(self):
        text = self.engine.generate(self.build_form())
        contract_summary = self.engine._chapter_body_text(text, "二十")
        outside_contract_summary = text.replace(contract_summary, "", 1) if contract_summary else text

        self.assertNotIn("本基金合同", outside_contract_summary)
        self.assertIn("4、基金合同：", self.extract_chapter(text, "二"))

    def test_contract_term_fallback_skips_prospectus_contract_summary_chapter(self):
        text = (
            "第十九章 基金合同的变更、终止和基金财产的清算\n"
            "基金管理人根据本基金合同履行信息披露义务。\n"
            "第二十章 基金合同的内容摘要\n"
            "本基金合同自基金份额持有人大会审议通过后生效。\n"
            "第二十一章 基金托管协议的内容摘要\n"
            "基金管理人依据本基金合同开展其他服务。\n"
        )

        normalized = self.engine._normalize_prospectus_contract_references_outside_contract_summary(text)
        chapter_19 = self.engine._chapter_body_text(normalized, "十九")
        chapter_20 = self.engine._chapter_body_text(normalized, "二十")
        chapter_21 = self.engine._chapter_body_text(normalized, "二十一")

        self.assertIn("根据基金合同履行信息披露义务", chapter_19)
        self.assertIn("本基金合同自基金份额持有人大会审议通过后生效", chapter_20)
        self.assertIn("依据基金合同开展其他服务", chapter_21)
        self.assertNotIn("本基金合同", f"{chapter_19}\n{chapter_21}")

    def test_chapter_seventeen_section_five_rewrites_product_summary_item_four(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "十七")
        section = self.extract_section(chapter, "五")

        item4 = self.extract_numbered_item(section, 4)
        self.assertIn(self.PROSPECTUS_PRODUCT_SUMMARY_ITEM4, item4)
        self.assertIn("相关内容不得与基金合同、招募说明书有实质性差异", item4)
        self.assertNotIn("基金产品资料概要的内容及编制等具体要求，按照招募说明书相关规定执行", item4)

    def test_chapter_seventeen_section_five_keeps_disclosure_items_after_product_summary_rewrite(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "十七")
        section = self.extract_section(chapter, "五")

        for heading in (
            "（二）基金份额发售公告",
            "（三）《基金合同》生效公告",
            "（四）基金净值信息",
            "（五）基金份额折算日公告、基金份额折算结果公告",
            "（六）基金份额上市交易公告书",
            "（七）申购赎回清单",
            "（八）基金定期报告，包括基金年度报告、基金中期报告和基金季度报告",
            "（九）临时报告",
        ):
            with self.subTest(heading=heading):
                self.assertIn(heading, section)

    def test_generated_prospectus_drops_template_metadata_block(self):
        text = self.engine.generate(self.build_form())
        self.assertTrue(text.startswith("测试交易型开放式指数证券投资基金招募说明书"))
        self.assertNotIn("VALUATION_TIMING_CLAUSE", text)
        self.assertNotIn("模板说明", text)
        self.assertNotIn("条件变量引用说明", text)
        self.assertNotIn("\n---\n", text)

    def test_generated_prospectus_uses_reference_style_toc_and_body_titles(self):
        text = self.engine.generate(self.build_form())
        toc = self.extract_toc(text)

        self.assertIn("一、绪言", toc)
        self.assertIn("四、基金托管人", toc)
        self.assertIn("二十一、基金托管协议的内容摘要", toc)
        self.assertNotIn("一、基金托管协议当事人", toc)
        self.assertNotIn("【托管人情况待填写】", toc)
        self.assertNotIn("第一章  绪言", text)
        self.assertIn("\n一、绪言\n", text)

    def test_chapter_one_reuses_reference_preface_body_with_only_fund_name_changed(self):
        form = self.build_form(FUND_NAME="测试红利质量交易型开放式指数证券投资基金")
        text = self.engine.generate(form)
        chapter = self.extract_chapter(text, "一")
        chapter_body = "\n".join(chapter.splitlines()[1:]).strip()
        reference_body = self.engine._load_reference_fixed_content(form)["一"]["body"]
        reference_name = self.engine._get_reference_prospectus_docx(form).stem.replace("招募说明书", "").strip()

        def normalize(body):
            normalized = body.replace(reference_name, "{FUND}").replace(form["FUND_NAME"], "{FUND}")
            normalized = re.sub(r"《[^》]*基金合同》", "《{FUND}基金合同》", normalized)
            return normalized

        self.assertEqual(normalize(reference_body), normalize(chapter_body))

    def test_chapter_three_uses_runtime_fund_manager_values(self):
        text = self.engine.generate(self.build_form(FUND_MANAGER_NAME="张三", FUND_MANAGER_BIO="张三简介"))
        chapter = self.extract_chapter(text, "三")

        self.assertIn("本基金的基金经理为张三。", chapter)
        self.assertIn("张三简介", chapter)
        self.assertNotIn("[待填写：基金经理姓名]", chapter)

    def test_contract_definition_of_bank_regulator_is_conditional_on_bank_custodian(self):
        bank_contract = app.engine.generate(self.build_form(CUSTODIAN_NAME="招商银行股份有限公司"))
        non_bank_contract = app.engine.generate(self.build_form(CUSTODIAN_NAME="中信证券股份有限公司"))

        self.assertIn("19、银行业监督管理机构：指中国人民银行和/或国家金融监督管理总局", bank_contract)
        self.assertNotIn("银行业监督管理机构：指中国人民银行和/或国家金融监督管理总局", non_bank_contract)
        self.assertIn("19、基金合同当事人：指受基金合同约束", non_bank_contract)
        self.assertNotIn("20、基金合同当事人：指受基金合同约束", non_bank_contract)

    def test_chapter_five_uses_reference_sections_with_lawyer_and_audit_overrides(self):
        form = self.build_form()
        text = self.engine.generate(form)
        chapter = self.extract_chapter(text, "五")
        reference = self.engine._load_reference_fixed_content(form)

        self.assertEqual(reference["五"]["sections"]["一"], self.extract_section(chapter, "一"))
        self.assertEqual(reference["五"]["sections"]["二"], self.extract_section(chapter, "二"))
        self.assertIn("经办律师：丁媛、李晓露", self.extract_section(chapter, "三"))
        self.assertEqual("四、审计基金财产的会计师事务所\n【待填写】", self.extract_section(chapter, "四"))

    def test_generate_render_model_keeps_chapter_titles_unumbered(self):
        model = self.engine.generate_render_model(self.build_form())

        self.assertEqual(model["chapters"][0]["title"], "绪言")
        self.assertEqual(model["chapters"][1]["title"], "释义")
        self.assertNotIn("、", model["chapters"][0]["title"])
        self.assertEqual(model["toc_title"], "目录")
        self.assertEqual(model["toc_titles"][0], "绪言")

    def test_generate_render_model_keeps_chapter_order_without_false_section_match(self):
        model = self.engine.generate_render_model(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))

        self.assertEqual(25, len(model["chapters"]))
        self.assertEqual(
            [("九", "基金份额的上市交易"), ("十", "基金份额的申购赎回"), ("十一", "基金的投资")],
            [(chapter["chapter_cn"], chapter["title"]) for chapter in model["chapters"][8:11]],
        )
        self.assertNotIn("二、基金份额的上市交易", [chapter["display_title"] for chapter in model["chapters"]])

    def test_generate_prospectus_api_includes_render_model(self):
        client = app.app.test_client()
        response = client.post("/api/generate_prospectus", json=self.build_form())

        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        self.assertTrue(payload["success"])
        self.assertIn("render_model", payload)
        self.assertEqual(payload["render_model"]["chapters"][0]["title"], "绪言")
        self.assertIn("一、绪言", payload["text"])

    def test_generate_product_summary_api_requires_generated_prospectus_text(self):
        client = app.app.test_client()

        response = client.post("/api/generate_product_summary", json=self.build_form())

        self.assertEqual(response.status_code, 400)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        self.assertFalse(payload["success"])
        self.assertIn("招募说明书", payload["error"])

    def test_generate_product_summary_api_returns_six_sections_and_four_tables(self):
        client = app.app.test_client()
        form = self.build_form()
        prospectus_bundle = self.engine.generate_bundle(form)

        response = client.post(
            "/api/generate_product_summary",
            json={
                **form,
                "PROSPECTUS_TEXT": prospectus_bundle["text"],
                "PROSPECTUS_RENDER_MODEL": prospectus_bundle["render_model"],
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        self.assertTrue(payload["success"])
        self.assertEqual(
            [
                "一、产品概况",
                "二、基金投资与净值表现",
                "三、投资本基金涉及的费用",
                "四、风险揭示与重要提示",
                "五、其他资料查询方式",
                "六、其他情况说明",
            ],
            [section["title"] for section in payload["render_model"]["sections"]],
        )
        table_blocks = [
            block
            for section in payload["render_model"]["sections"]
            for block in section.get("blocks", [])
            if isinstance(block, dict) and block.get("type") == "table"
        ]
        self.assertEqual(4, len(table_blocks))
        self.assertEqual(
            ["基金简称", "测试交易型开放式指数证券投资基金", "基金代码", "-"],
            table_blocks[0]["rows"][0],
        )
        self.assertIn("本概要提供本基金的重要信息，是招募说明书的一部分。", payload["text"])

    def test_generate_product_summary_distinguishes_sse_and_szse_templates(self):
        sse_bundle = app.product_summary_engine.generate_bundle(
            {
                **self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"),
                "PROSPECTUS_TEXT": "placeholder",
            }
        )
        szse_bundle = app.product_summary_engine.generate_bundle(
            {
                **self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"),
                "PROSPECTUS_TEXT": "placeholder",
            }
        )

        sse_table = next(
            block
            for block in sse_bundle["render_model"]["sections"][1]["blocks"]
            if isinstance(block, dict) and block.get("type") == "table"
        )
        szse_table = next(
            block
            for block in szse_bundle["render_model"]["sections"][1]["blocks"]
            if isinstance(block, dict) and block.get("type") == "table"
        )
        sse_benchmark = sse_table["rows"][3][1]
        szse_benchmark = szse_table["rows"][3][1]
        self.assertEqual("中证测试指数收益率", sse_benchmark)
        self.assertEqual("本基金的业绩比较基准为标的指数收益率。本基金标的指数为中证测试指数。", szse_benchmark)

    def test_generate_product_summary_includes_sample_note_line(self):
        client = app.app.test_client()
        form = self.build_form()
        prospectus_bundle = self.engine.generate_bundle(form)

        response = client.post(
            "/api/generate_product_summary",
            json={
                **form,
                "PROSPECTUS_TEXT": prospectus_bundle["text"],
                "PROSPECTUS_RENDER_MODEL": prospectus_bundle["render_model"],
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        self.assertIn("注:详见《测试交易型开放式指数证券投资基金招募说明书》“基金的投资”", payload["text"])

    def test_export_product_summary_docx_renders_native_tables(self):
        client = app.app.test_client()
        form = self.build_form()
        prospectus_bundle = self.engine.generate_bundle(form)

        response = client.post(
            "/api/export_product_summary_docx",
            json={
                **form,
                "PROSPECTUS_TEXT": prospectus_bundle["text"],
                "PROSPECTUS_RENDER_MODEL": prospectus_bundle["render_model"],
            },
        )

        self.assertEqual(response.status_code, 200)
        docx_bytes = response.data
        doc = self._docx_document(docx_bytes)
        self.assertEqual(4, len(doc.tables))
        self.assertIn("基金产品资料概要", "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip()))
        first_row = [cell.text.strip() for cell in doc.tables[0].rows[0].cells]
        self.assertEqual(["基金简称", "测试交易型开放式指数证券投资基金", "基金代码", "-"], first_row)

    def test_export_product_summary_docx_formats_cover_notice_lines_like_sample(self):
        client = app.app.test_client()
        form = self.build_form()
        prospectus_bundle = self.engine.generate_bundle(form)

        response = client.post(
            "/api/export_product_summary_docx",
            json={
                **form,
                "PROSPECTUS_TEXT": prospectus_bundle["text"],
                "PROSPECTUS_RENDER_MODEL": prospectus_bundle["render_model"],
            },
        )

        self.assertEqual(response.status_code, 200)
        doc = self._docx_document(response.data)
        notice_paragraph = next(
            paragraph
            for paragraph in doc.paragraphs
            if "本概要提供本基金的重要信息，是招募说明书的一部分。" in paragraph.text
        )
        notice_run = notice_paragraph.runs[0]
        self.assertTrue(bool(notice_run.bold))
        self.assertAlmostEqual(15.5, notice_run.font.size.pt, places=1)

    def test_export_product_summary_docx_uses_sample_style_ids(self):
        client = app.app.test_client()
        form = self.build_form()
        prospectus_bundle = self.engine.generate_bundle(form)

        response = client.post(
            "/api/export_product_summary_docx",
            json={
                **form,
                "PROSPECTUS_TEXT": prospectus_bundle["text"],
                "PROSPECTUS_RENDER_MODEL": prospectus_bundle["render_model"],
            },
        )

        self.assertEqual(response.status_code, 200)
        xml_map = self._docx_xml(response.data)
        document_xml = xml_map["word/document.xml"]
        self.assertIn('w:pStyle w:val="-1"', document_xml)
        self.assertIn('w:pStyle w:val="-2"', document_xml)
        self.assertIn('w:pStyle w:val="-7"', document_xml)
        self.assertIn('w:tblStyle w:val="-noheader"', document_xml)

    def test_product_summary_template_is_fixed_to_petrochemical_sample(self):
        template_path = app._resolve_product_summary_template_docx()

        self.assertIn("石油天然气", template_path.name)

    def test_generate_product_summary_uses_explicit_prospectus_sections_for_risk_return_and_other_fees(self):
        form = self.build_form(
            MGMT_FEE_RATE="0.30",
            CUSTODY_FEE_RATE="0.10",
        )
        custom_render_model = {
            "chapters": [
                {
                    "title": "基金的投资",
                    "blocks": [
                        {"type": "paragraph", "text": "六、风险收益特征"},
                        {"type": "paragraph", "text": "这里是自定义风险收益特征。"},
                    ],
                },
                {
                    "title": "基金的费用与税收",
                    "blocks": [
                        {"type": "paragraph", "text": "一、基金费用的种类"},
                        {"type": "paragraph", "text": "1、基金管理人的管理费；"},
                        {"type": "paragraph", "text": "2、基金托管人的托管费；"},
                        {"type": "paragraph", "text": "3、审计费；"},
                        {"type": "paragraph", "text": "4、上市费；"},
                    ],
                },
            ]
        }

        bundle = app.product_summary_engine.generate_bundle(
            {
                **form,
                "PROSPECTUS_TEXT": "placeholder",
                "PROSPECTUS_RENDER_MODEL": custom_render_model,
            }
        )

        investment_table = next(
            block
            for block in bundle["render_model"]["sections"][1]["blocks"]
            if isinstance(block, dict) and block.get("type") == "table"
        )
        fee_table = next(
            block
            for block in bundle["render_model"]["sections"][2]["blocks"]
            if isinstance(block, dict) and block.get("type") == "table" and block["rows"][0][0] == "费用类别"
        )
        self.assertEqual("这里是自定义风险收益特征。", investment_table["rows"][4][1])
        self.assertEqual("0.30%", fee_table["rows"][1][1])
        self.assertEqual("0.10%", fee_table["rows"][2][1])
        self.assertEqual("审计费；上市费。", fee_table["rows"][4][1])

    def test_export_product_summary_docx_replaces_template_content_in_place(self):
        client = app.app.test_client()
        form = self.build_form(
            FUND_NAME="自定义石油ETF基金",
            FUND_SHORT_NAME="自定义石油ETF",
            CUSTODIAN_NAME="自定义托管银行股份有限公司",
            FUND_MANAGER_NAME="李四",
            FUND_MANAGER_START_DATE="2026年1月1日",
            FUND_MANAGER_SECURITY_DATE="2018年1月1日",
            MGMT_FEE_RATE="0.30",
            CUSTODY_FEE_RATE="0.10",
            INDEX_NAME="自定义指数",
            EXCHANGE="SZSE",
            ON_EXCHANGE_SHORT_NAME="自定义ETF南方",
        )
        custom_render_model = {
            "chapters": [
                {
                    "title": "基金的投资",
                    "blocks": [
                        {"type": "paragraph", "text": "六、风险收益特征"},
                        {"type": "paragraph", "text": "这里是导出风险收益特征。"},
                    ],
                },
                {
                    "title": "基金的费用与税收",
                    "blocks": [
                        {"type": "paragraph", "text": "一、基金费用的种类"},
                        {"type": "paragraph", "text": "1、基金管理人的管理费；"},
                        {"type": "paragraph", "text": "2、基金托管人的托管费；"},
                        {"type": "paragraph", "text": "3、审计费；"},
                        {"type": "paragraph", "text": "4、上市费；"},
                    ],
                },
                {
                    "title": "风险揭示",
                    "blocks": [
                        {"type": "paragraph", "text": "一、市场风险"},
                        {"type": "paragraph", "text": "1、政策风险。"},
                        {"type": "paragraph", "text": "2、经济周期风险。"},
                        {"type": "paragraph", "text": "9、投资股指期货的风险"},
                        {"type": "paragraph", "text": "（1）市场风险"},
                        {"type": "paragraph", "text": "（2）流动性风险"},
                        {"type": "paragraph", "text": "（3）基差风险"},
                        {"type": "paragraph", "text": "13、本基金参与转融通证券出借业务的风险"},
                        {"type": "paragraph", "text": "（1）流动性风险"},
                        {"type": "paragraph", "text": "（2）信用风险"},
                        {"type": "paragraph", "text": "（3）市场风险"},
                        {"type": "paragraph", "text": "四、本基金特有的风险"},
                        {"type": "paragraph", "text": "1、这里是导出的特有风险。"},
                        {"type": "paragraph", "text": "2、这里是第二条导出的特有风险。"},
                    ],
                },
            ]
        }

        response = client.post(
            "/api/export_product_summary_docx",
            json={
                **form,
                "PROSPECTUS_TEXT": "placeholder",
                "PROSPECTUS_RENDER_MODEL": custom_render_model,
            },
        )

        self.assertEqual(response.status_code, 200)
        doc = self._docx_document(response.data)
        self.assertEqual("自定义石油ETF基金基金产品资料概要", doc.paragraphs[0].text.strip())
        self.assertEqual("编制日期：202X年X月X日", doc.paragraphs[2].text.strip())
        self.assertEqual("送出日期：202X年X月X日", doc.paragraphs[3].text.strip())
        self.assertEqual(
            "注:详见《自定义石油ETF基金招募说明书》“基金的投资”部分。",
            doc.paragraphs[9].text.strip(),
        )
        self.assertEqual("自定义石油ETF", doc.tables[0].rows[0].cells[1].text.strip())
        self.assertEqual("自定义托管银行股份有限公司", doc.tables[0].rows[1].cells[3].text.strip())
        self.assertEqual("深圳证券交易所", doc.tables[0].rows[2].cells[3].text.strip())
        self.assertEqual("李四", doc.tables[0].rows[7].cells[1].text.strip())
        self.assertEqual("202X年X月X日", doc.tables[0].rows[8].cells[3].text.strip())
        self.assertEqual("场内简称：XXXETF南方", doc.tables[0].rows[9].cells[1].text.strip())
        self.assertEqual("这里是导出风险收益特征。", doc.tables[1].rows[4].cells[1].text.strip())
        self.assertEqual("0.30%", doc.tables[3].rows[1].cells[1].text.strip())
        self.assertEqual("0.10%", doc.tables[3].rows[2].cells[1].text.strip())
        self.assertEqual("审计费；上市费。", doc.tables[3].rows[4].cells[1].text.strip())
        all_paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        self.assertIn("一）本基金特有的风险", all_paragraphs)
        self.assertIn("1、这里是导出的特有风险。", all_paragraphs)
        self.assertIn("2、这里是第二条导出的特有风险。", all_paragraphs)
        market_risk_index = all_paragraphs.index("二）市场风险")
        self.assertEqual(
            "证券市场价格受到经济因素、政治因素、投资心理和交易制度等各种因素的影响，导致基金收益水平变化，产生风险，主要包括："
            "1、政策风险。2、经济周期风险。9、投资股指期货的风险：（1）市场风险（2）流动性风险（3）基差风险。"
            "13、本基金参与转融通证券出借业务的风险：（1）流动性风险（2）信用风险（3）市场风险。",
            all_paragraphs[market_risk_index + 1],
        )

    def test_generate_product_summary_risk_section_reuses_special_risk_body(self):
        form = self.build_form(FUND_NAME="风险测试基金")
        custom_render_model = {
            "chapters": [
                {
                    "title": "风险揭示",
                    "blocks": [
                        {"type": "paragraph", "text": "一、市场风险"},
                        {"type": "paragraph", "text": "1、政策风险。"},
                        {"type": "paragraph", "text": "2、利率风险。"},
                        {"type": "paragraph", "text": "9、投资股指期货的风险"},
                        {"type": "paragraph", "text": "（1）市场风险"},
                        {"type": "paragraph", "text": "（2）流动性风险"},
                        {"type": "paragraph", "text": "13、本基金参与转融通证券出借业务的风险"},
                        {"type": "paragraph", "text": "（1）流动性风险"},
                        {"type": "paragraph", "text": "（2）信用风险"},
                        {"type": "paragraph", "text": "四、本基金特有的风险"},
                        {"type": "paragraph", "text": "1、这里是不应进入概要的特有风险正文。"},
                        {"type": "paragraph", "text": "2、这里也是不应进入概要的特有风险正文。"},
                    ],
                },
            ]
        }

        bundle = app.product_summary_engine.generate_bundle(
            {
                **form,
                "PROSPECTUS_TEXT": "placeholder",
                "PROSPECTUS_RENDER_MODEL": custom_render_model,
            }
        )

        risk_section = bundle["text"].split("四、风险揭示与重要提示", 1)[1].split("五、其他资料查询方式", 1)[0]

        self.assertIn("一）本基金特有的风险", risk_section)
        self.assertIn("1、这里是不应进入概要的特有风险正文。", risk_section)
        self.assertIn("2、这里也是不应进入概要的特有风险正文。", risk_section)
        self.assertIn("二）市场风险", risk_section)
        self.assertIn(
            "证券市场价格受到经济因素、政治因素、投资心理和交易制度等各种因素的影响，导致基金收益水平变化，产生风险，主要包括："
            "1、政策风险。2、利率风险。9、投资股指期货的风险：（1）市场风险（2）流动性风险。"
            "13、本基金参与转融通证券出借业务的风险：（1）流动性风险（2）信用风险。",
            risk_section,
        )

    def test_generate_product_summary_important_notice_matches_template_sample(self):
        form = self.build_form(FUND_NAME="重要提示测试基金", CSRC_APPROVAL_NO="2026年4月1日证监许可〔2026〕100号")
        bundle = app.product_summary_engine.generate_bundle(
            {
                **form,
                "PROSPECTUS_TEXT": "placeholder",
                "PROSPECTUS_RENDER_MODEL": {"chapters": []},
            }
        )

        important_notice = bundle["text"].split("（二） 重要提示", 1)[1].split("五、其他资料查询方式", 1)[0]

        self.assertIn(
            "重要提示测试基金（以下简称“本基金”）经中国证监会2026年4月1日证监许可〔2026〕100号文注册募集。",
            important_notice,
        )
        self.assertIn(
            "与本基金/基金合同相关的争议解决方式为仲裁。因《基金合同》而产生的或与《基金合同》有关的一切争议，如经友好协商未能解决的，任何一方均有权将争议提交深圳国际仲裁院，按照深圳国际仲裁院届时有效的仲裁规则进行仲裁，仲裁地点为深圳市。",
            important_notice,
        )
        self.assertIn(
            "基金产品资料概要信息发生重大变更的，基金管理人将在三个工作日内更新，其他信息发生变更的，基金管理人每年更新一次。",
            important_notice,
        )

    def test_generate_product_summary_important_notice_uses_beijing_cietac(self):
        form = self.build_form(
            FUND_NAME="重要提示测试基金",
            CSRC_APPROVAL_NO="2026年4月1日证监许可〔2026〕100号",
            DISPUTE_RESOLUTION_VENUE="BJ_CIETAC",
        )
        bundle = app.product_summary_engine.generate_bundle(
            {
                **form,
                "PROSPECTUS_TEXT": "placeholder",
                "PROSPECTUS_RENDER_MODEL": {"chapters": []},
            }
        )

        important_notice = bundle["text"].split("（二） 重要提示", 1)[1].split("五、其他资料查询方式", 1)[0]

        self.assertIn("将争议提交中国国际经济贸易仲裁委员会", important_notice)
        self.assertIn("仲裁地点为北京市。", important_notice)
        self.assertNotIn("将争议提交深圳国际仲裁院", important_notice)

    def test_generate_product_summary_market_risk_uses_sse_single_manual_headings(self):
        manual_path = app.REFERENCE_PROSPECTUS_REPO_DIR / "SSE_SINGLE.txt"
        prospectus_text = manual_path.read_text(encoding="utf-8")
        bundle = app.product_summary_engine.generate_bundle(
            {
                **self.build_form(EXCHANGE="SSE", MARKET_TYPE="KECHUANG"),
                "PROSPECTUS_TEXT": prospectus_text,
            }
        )

        risk_section = bundle["text"].split("四、风险揭示与重要提示", 1)[1].split("五、其他资料查询方式", 1)[0]

        self.assertIn("二）市场风险", risk_section)
        self.assertIn(app.product_summary_engine._fixed_market_risk_intro(), risk_section)
        self.assertIn("9、投资科创板股票的风险。", risk_section)
        self.assertIn("10、投资股指期货的风险：（1）市场风险（2）流动性风险（3）基差风险", risk_section)
        self.assertIn("14、本基金参与转融通证券出借业务的风险：（1）流动性风险（2）信用风险（3）市场风险（4）操作风险。", risk_section)

    def test_export_product_summary_docx_updates_header_fund_name(self):
        client = app.app.test_client()
        form = self.build_form(FUND_NAME="页眉测试基金")
        response = client.post(
            "/api/export_product_summary_docx",
            json={
                **form,
                "PROSPECTUS_TEXT": "placeholder",
                "PROSPECTUS_RENDER_MODEL": {"chapters": []},
            },
        )

        self.assertEqual(response.status_code, 200)
        doc = self._docx_document(response.data)
        self.assertEqual("页眉测试基金基金产品资料概要", doc.sections[0].header.paragraphs[0].text.strip())

    def test_index_template_includes_product_summary_panel_and_compare_section(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn('data-panel="product-summary"', html)
        self.assertIn('id="product-summary-panel"', html)
        self.assertIn("生成产品资料概要", html)
        self.assertIn("产品资料概要对比", html)
        self.assertIn("loadGeneratedProductSummaryAsA()", html)
        self.assertIn("runProductSummaryCompare()", html)

    def test_prospectus_custodian_summary_upload_status_and_clear(self):
        client = app.app.test_client()
        original_store = getattr(app, "_prospectus_store", None)
        if isinstance(original_store, dict):
            original_store = deepcopy(original_store)

        upload_buffer = self._build_review_docx_buffer(
            title_lines=["南方中证全指红利质量ETF发起联接基金托管协议"],
            body_lines=[
                "一、托管协议当事人",
                "这里是托管协议当事人正文。",
                "二、基金托管人对基金管理人的业务监督和核查",
                "这里是监督和核查正文。",
            ],
        )

        try:
            upload_response = client.post(
                "/api/prospectus/custodian_summary/upload",
                data={"file": (upload_buffer, "sample-custodian.docx")},
                content_type="multipart/form-data",
            )
            self.assertEqual(upload_response.status_code, 200)
            upload_payload = self._consume_response(upload_response, parser=lambda resp: resp.get_json())
            self.assertTrue(upload_payload["success"])
            self.assertEqual("sample-custodian.docx", upload_payload["filename"])
            self.assertGreaterEqual(upload_payload["section_count"], 2)

            status_response = client.get("/api/prospectus/custodian_summary/status")
            self.assertEqual(status_response.status_code, 200)
            status_payload = self._consume_response(status_response, parser=lambda resp: resp.get_json())
            self.assertTrue(status_payload["uploaded"])
            self.assertEqual("sample-custodian.docx", status_payload["filename"])
            self.assertGreaterEqual(status_payload["section_count"], 2)

            clear_response = client.post("/api/prospectus/custodian_summary/clear", json={})
            self.assertEqual(clear_response.status_code, 200)
            clear_payload = self._consume_response(clear_response, parser=lambda resp: resp.get_json())
            self.assertTrue(clear_payload["success"])

            status_after_clear = client.get("/api/prospectus/custodian_summary/status")
            self.assertEqual(status_after_clear.status_code, 200)
            status_after_clear_payload = self._consume_response(status_after_clear, parser=lambda resp: resp.get_json())
            self.assertFalse(status_after_clear_payload["uploaded"])
        finally:
            if original_store is None:
                if hasattr(app, "_prospectus_store"):
                    app._prospectus_store.clear()
            else:
                app._prospectus_store.clear()
                app._prospectus_store.update(original_store)

    def test_custodian_summary_upload_extracts_tracked_revision_docx(self):
        upload_buffer = self._build_tracked_custodian_docx_buffer()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(upload_buffer.getvalue())
            tmp_path = tmp.name
        try:
            text, sections = app._extract_custodian_summary_sections_from_docx(tmp_path)
        finally:
            os.unlink(tmp_path)

        headings = [section["heading"] for section in sections]
        self.assertIn("一、托管协议当事人", headings)
        self.assertIn("二、基金托管人对基金管理人的业务监督和核查", headings)
        self.assertIn("当事人修订正文。", text)
        self.assertIn("监督核查修订正文。", text)
        self.assertNotIn("一、被删除章节", text)

    def test_extract_review_revision_items_labels_contract_summary_and_impact_chain(self):
        items = app._extract_review_docx_revision_items(
            self._build_tracked_review_docx_buffer().getvalue(),
            document_kind="contract",
            filename="测试基金合同-律所修订.docx",
        )

        inserted = next(item for item in items if item["revision_type"] == "insert")
        summary_delete = next(item for item in items if item["revision_type"] == "delete")
        summary_comment = next(item for item in items if item["revision_type"] == "comment")

        self.assertEqual("contract_body", inserted["document_scope"])
        self.assertEqual("基金合同正文", inserted["document_label"])
        self.assertEqual("新增正文意见", inserted["text"])
        inserted_targets = [target["label"] for target in inserted["impact_targets"]]
        self.assertIn("基金合同摘要", inserted_targets)
        self.assertIn("招募说明书", inserted_targets)
        self.assertIn("招募说明书内基金合同摘要", inserted_targets)
        self.assertIn("产品资料概要", inserted_targets)

        self.assertEqual("contract_summary", summary_delete["document_scope"])
        self.assertEqual("基金合同摘要", summary_delete["document_label"])
        self.assertEqual("删除摘要旧口径", summary_delete["text"])
        self.assertIn("招募说明书内基金合同摘要", [target["label"] for target in summary_delete["impact_targets"]])

        self.assertEqual("contract_summary", summary_comment["document_scope"])
        self.assertEqual("摘要口径需同步基金合同正文。", summary_comment["comment_text"])
        self.assertEqual("pending", summary_comment["decision"])

    def test_extract_review_revision_items_marks_prospectus_to_product_summary_impact(self):
        items = app._extract_review_docx_revision_items(
            self._build_tracked_prospectus_docx_buffer().getvalue(),
            document_kind="prospectus",
            filename="测试招募说明书-律所修订.docx",
        )

        body_item = next(item for item in items if item["revision_type"] == "insert")
        summary_item = next(item for item in items if item["revision_type"] == "delete")

        self.assertEqual("prospectus_body", body_item["document_scope"])
        self.assertEqual("招募说明书", body_item["document_label"])
        self.assertIn("产品资料概要", [target["label"] for target in body_item["impact_targets"]])
        self.assertEqual("prospectus_contract_summary", summary_item["document_scope"])
        self.assertEqual("招募说明书内基金合同摘要", summary_item["document_label"])
        self.assertIn("产品资料概要", [target["label"] for target in summary_item["impact_targets"]])

    def test_extract_review_revision_items_handles_appendix_contract_summary_boundary(self):
        base = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "一、绪言",
                "招募正文段落",
                "附件一：基金合同内容摘要",
                "招募合同摘要段落",
                "附件二：托管协议内容摘要",
                "托管协议摘要段落",
            ],
        )
        docx_buffer = self._rewrite_docx_xml_buffer(
            base,
            [
                (
                    "<w:t>招募合同摘要段落</w:t>",
                    '<w:t>招募合同摘要段落</w:t><w:del w:id="5" w:author="律所律师" w:date="2026-05-16T08:20:00Z"><w:r><w:delText>删除招募合同摘要旧口径</w:delText></w:r></w:del>',
                ),
                (
                    "<w:t>托管协议摘要段落</w:t>",
                    '<w:t>托管协议摘要段落</w:t><w:ins w:id="6" w:author="律所律师" w:date="2026-05-16T08:25:00Z"><w:r><w:t>新增托管协议摘要意见</w:t></w:r></w:ins>',
                ),
            ],
        )

        items = app._extract_review_docx_revision_items(
            docx_buffer.getvalue(),
            document_kind="prospectus",
            filename="测试招募说明书-律所修订.docx",
        )

        contract_summary_item = next(item for item in items if item["revision_type"] == "delete")
        custodian_summary_item = next(item for item in items if item["revision_type"] == "insert")

        self.assertEqual("prospectus_contract_summary", contract_summary_item["document_scope"])
        self.assertEqual("prospectus_body", custodian_summary_item["document_scope"])

    def test_build_review_revision_document_view_preserves_inline_revision_flow(self):
        document = app._build_review_revision_document_view(
            self._build_review_reader_docx_buffer().getvalue(),
            document_kind="contract",
            filename="测试基金合同-律所修订.docx",
        )

        self.assertEqual("contract", document["document_kind"])
        self.assertEqual("测试基金合同-律所修订.docx", document["filename"])
        paragraph_blocks = [block for block in document["blocks"] if block["type"] == "paragraph"]
        self.assertGreaterEqual(len(paragraph_blocks), 5)

        insert_block = next(block for block in paragraph_blocks if "普通上下文段落" in block["text"])
        self.assertEqual("contract_body", insert_block["document_scope"])
        self.assertEqual(["normal", "insert"], [run["mark_type"] for run in insert_block["runs"]])
        self.assertEqual("新增披露文字", insert_block["runs"][1]["text"])

        modify_block = next(block for block in paragraph_blocks if "费率表述段落" in block["text"])
        modify_runs = [run for run in modify_block["runs"] if run["mark_type"].startswith("modify")]
        self.assertEqual(["modify_from", "modify_to"], [run["mark_type"] for run in modify_runs])
        self.assertEqual(["0.15%", "0.20%"], [run["text"] for run in modify_runs])
        self.assertEqual(modify_runs[0]["revision_id"], modify_runs[1]["revision_id"])
        modify_revision = next(revision for revision in document["revisions"] if revision["revision_type"] == "modify")
        self.assertEqual("0.15%", modify_revision["old_text"])
        self.assertEqual("0.20%", modify_revision["new_text"])

        comment_block = next(block for block in paragraph_blocks if "摘要批注段落" in block["text"])
        self.assertEqual("contract_summary", comment_block["document_scope"])
        self.assertTrue(any(run["mark_type"] == "comment" for run in comment_block["runs"]))
        comment_revision = next(revision for revision in document["revisions"] if revision["revision_type"] == "comment")
        self.assertEqual("这里需要确认摘要是否同步。", comment_revision["comment_text"])

    def test_build_review_revision_document_view_keeps_table_revisions_in_flow(self):
        document = app._build_review_revision_document_view(
            self._build_review_reader_table_docx_buffer().getvalue(),
            document_kind="contract",
            filename="测试基金合同-律所修订.docx",
        )

        table_block = next(block for block in document["blocks"] if block["type"] == "table")
        cell_blocks = [
            child_block
            for row in table_block["rows"]
            for cell in row
            for child_block in cell["blocks"]
        ]
        table_runs = [run for block in cell_blocks for run in block["runs"]]

        self.assertIn("表格费率段落", table_block["text"])
        self.assertTrue(any(run["mark_type"] == "modify_from" for run in table_runs))
        self.assertTrue(any(run["mark_type"] == "modify_to" for run in table_runs))
        self.assertTrue(any(revision["revision_type"] == "modify" for revision in document["revisions"]))

    def test_revision_workbench_upload_returns_revision_items_and_summary_bucket(self):
        client = app.app.test_client()
        original_data = app._revision_workbench_store.get("data")

        try:
            response = client.post(
                "/api/revision_workbench/upload",
                data={
                    "contract": (self._build_tracked_review_docx_buffer(), "测试基金合同-律所修订.docx"),
                    "prospectus": (self._build_tracked_prospectus_docx_buffer(), "测试招募说明书-律所修订.docx"),
                },
                content_type="multipart/form-data",
            )

            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
            self.assertGreaterEqual(payload["revision_total"], 5)
            self.assertGreaterEqual(payload["revision_counts"]["contract_body"], 1)
            self.assertGreaterEqual(payload["revision_counts"]["contract_summary"], 2)
            self.assertGreaterEqual(payload["revision_counts"]["prospectus_body"], 1)
            self.assertGreaterEqual(payload["revision_counts"]["prospectus_contract_summary"], 1)

            stored_items = app._revision_workbench_store["data"]["revision_items"]
            self.assertEqual(payload["revision_total"], len(stored_items))
            self.assertTrue(any(item["document_scope"] == "contract_summary" for item in stored_items))

            self.assertIn("documents", payload)
            contract_doc = next(doc for doc in payload["documents"] if doc["document_kind"] == "contract")
            self.assertIn("blocks", contract_doc)
            self.assertIn("revisions", contract_doc)
            self.assertIn("revision_counts", contract_doc)
            self.assertTrue(any(block["type"] == "paragraph" for block in contract_doc["blocks"]))
            self.assertTrue(any(run["mark_type"] == "insert" for block in contract_doc["blocks"] for run in block.get("runs", [])))
        finally:
            if original_data is None:
                app._revision_workbench_store.pop("data", None)
            else:
                app._revision_workbench_store["data"] = original_data

    def test_apply_review_revision_decisions_accepts_and_rejects_docx_redlines(self):
        source = self._build_review_reader_docx_buffer().getvalue()
        document = app._build_review_revision_document_view(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
        )
        insert_revision = next(item for item in document["revisions"] if item["revision_type"] == "insert")
        modify_revision = next(item for item in document["revisions"] if item["revision_type"] == "modify")
        comment_revision = next(item for item in document["revisions"] if item["revision_type"] == "comment")

        final_bytes, report = app._apply_review_revision_decisions_to_docx(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
            decisions={
                insert_revision["id"]: "accept",
                modify_revision["id"]: "reject",
                comment_revision["id"]: "reject",
            },
        )

        xml_map = self._docx_xml(final_bytes)
        document_xml = xml_map["word/document.xml"]
        self.assertIn("<w:ins", document_xml)
        self.assertIn("0.15%", document_xml)
        self.assertNotIn("0.20%", document_xml)
        self.assertNotIn("commentRangeStart", document_xml)
        self.assertNotIn("commentReference", document_xml)
        with zipfile.ZipFile(io.BytesIO(final_bytes), "r") as zf:
            comments_xml = zf.read("word/comments.xml").decode("utf-8", errors="ignore")
        self.assertNotIn('w:id="11"', comments_xml)
        self.assertEqual("contract", report["document_kind"])
        self.assertEqual(1, report["accepted_count"])
        self.assertEqual(2, report["rejected_count"])

    def test_apply_review_revision_decisions_applies_edit_to_modify_redline(self):
        source = self._build_review_reader_docx_buffer().getvalue()
        document = app._build_review_revision_document_view(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
        )
        modify_revision = next(item for item in document["revisions"] if item["revision_type"] == "modify")

        final_bytes, report = app._apply_review_revision_decisions_to_docx(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
            decisions={modify_revision["id"]: "accept"},
            revision_edits={modify_revision["id"]: {"text": "0.33%"}},
        )

        document_xml = self._docx_xml(final_bytes)["word/document.xml"]
        self.assertIn("0.15%", document_xml)
        self.assertIn("0.33%", document_xml)
        self.assertNotIn("0.20%", document_xml)
        self.assertEqual(1, report["edited_count"])
        self.assertEqual([modify_revision["id"]], report["edited_revision_ids"])

    def test_apply_review_revision_decisions_applies_paragraph_edit_to_modify_redline(self):
        source = self._build_review_reader_docx_buffer().getvalue()
        document = app._build_review_revision_document_view(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
        )
        modify_revision = next(item for item in document["revisions"] if item["revision_type"] == "modify")

        final_bytes, report = app._apply_review_revision_decisions_to_docx(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
            decisions={modify_revision["id"]: "accept"},
            revision_edits={modify_revision["id"]: {"text": "费率表述段落最终整段人工编辑。", "mode": "paragraph"}},
        )

        document_xml = self._docx_xml(final_bytes)["word/document.xml"]
        self.assertIn("0.15%", document_xml)
        self.assertIn("<w:del", document_xml)
        self.assertIn("<w:ins", document_xml)
        self.assertIn("费率表述段落最终整段人工编辑。", document_xml)
        self.assertNotIn("0.20%", document_xml)
        self.assertEqual(1, report["edited_count"])
        self.assertEqual([modify_revision["id"]], report["edited_revision_ids"])

    def test_apply_review_revision_decisions_applies_edit_to_insert_redline(self):
        source = self._build_review_reader_docx_buffer().getvalue()
        document = app._build_review_revision_document_view(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
        )
        insert_revision = next(item for item in document["revisions"] if item["revision_type"] == "insert")

        final_bytes, report = app._apply_review_revision_decisions_to_docx(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
            decisions={insert_revision["id"]: "accept"},
            revision_edits={insert_revision["id"]: {"text": "Manual inserted text"}},
            review_author="产品部审核",
        )

        document_xml = self._docx_xml(final_bytes)["word/document.xml"]
        self.assertIn("<w:ins", document_xml)
        self.assertIn('w:author="产品部审核"', document_xml)
        self.assertIn("Manual inserted text", document_xml)
        self.assertNotIn("鏂板鎶湶鏂囧瓧", document_xml)
        self.assertEqual(1, report["edited_count"])

    def test_apply_review_revision_decisions_applies_edit_to_delete_as_insert_redline(self):
        source = self._build_tracked_review_docx_buffer().getvalue()
        document = app._build_review_revision_document_view(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
        )
        delete_revision = next(item for item in document["revisions"] if item["revision_type"] == "delete")

        final_bytes, report = app._apply_review_revision_decisions_to_docx(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
            decisions={delete_revision["id"]: "accept"},
            revision_edits={delete_revision["id"]: {"text": "Replacement for deleted text"}},
        )

        document_xml = self._docx_xml(final_bytes)["word/document.xml"]
        self.assertIn("<w:del", document_xml)
        self.assertIn("<w:ins", document_xml)
        self.assertIn("Replacement for deleted text", document_xml)
        self.assertEqual(1, report["edited_count"])

    def test_apply_review_revision_decisions_ignores_edit_when_rejected(self):
        source = self._build_review_reader_docx_buffer().getvalue()
        document = app._build_review_revision_document_view(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
        )
        modify_revision = next(item for item in document["revisions"] if item["revision_type"] == "modify")

        final_bytes, report = app._apply_review_revision_decisions_to_docx(
            source,
            document_kind="contract",
            filename="contract-redline.docx",
            decisions={modify_revision["id"]: "reject"},
            revision_edits={modify_revision["id"]: {"text": "Rejected edit text"}},
        )

        document_xml = self._docx_xml(final_bytes)["word/document.xml"]
        self.assertIn("0.15%", document_xml)
        self.assertNotIn("0.20%", document_xml)
        self.assertNotIn("Rejected edit text", document_xml)
        self.assertEqual(0, report["edited_count"])

    def test_revision_workbench_export_final_blocks_pending_decisions(self):
        client = app.app.test_client()
        original_data = app._revision_workbench_store.get("data")

        try:
            upload_response = client.post(
                "/api/revision_workbench/upload",
                data={
                    "contract": (self._build_tracked_review_docx_buffer(), "contract.docx"),
                    "prospectus": (self._build_tracked_prospectus_docx_buffer(), "prospectus.docx"),
                },
                content_type="multipart/form-data",
            )
            self.assertEqual(upload_response.status_code, 200)
            upload_response.close()

            response = client.post(
                "/api/revision_workbench/export_final",
                json={"decisions": {}, "form_data": self.build_form()},
            )

            self.assertEqual(response.status_code, 400)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
            self.assertIn("待定", payload["error"])
        finally:
            if original_data is None:
                app._revision_workbench_store.pop("data", None)
            else:
                app._revision_workbench_store["data"] = original_data

    def test_revision_workbench_upload_accepts_custody_agreement_docx(self):
        client = app.app.test_client()
        original_data = app._revision_workbench_store.get("data")

        try:
            response = client.post(
                "/api/revision_workbench/upload",
                data={
                    "custody_agreement": (self._build_tracked_custodian_docx_buffer(), "托管协议.docx"),
                },
                content_type="multipart/form-data",
            )

            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
            self.assertIn("custody_agreement", payload["revision_counts"])
            self.assertGreaterEqual(payload["revision_counts"]["custody_agreement"], 2)
            self.assertIn("托管协议", payload["scope_labels"]["custody_agreement"])
            self.assertTrue(any(document["document_kind"] == "custody_agreement" for document in payload["documents"]))
        finally:
            if original_data is None:
                app._revision_workbench_store.pop("data", None)
            else:
                app._revision_workbench_store["data"] = original_data

    def test_revision_workbench_export_final_returns_zip_without_product_summary(self):
        client = app.app.test_client()
        original_data = app._revision_workbench_store.get("data")

        try:
            upload_response = client.post(
                "/api/revision_workbench/upload",
                data={
                    "contract": (self._build_tracked_review_docx_buffer(), "contract.docx"),
                    "prospectus": (self._build_tracked_prospectus_docx_buffer(), "prospectus.docx"),
                },
                content_type="multipart/form-data",
            )
            self.assertEqual(upload_response.status_code, 200)
            payload = self._consume_response(upload_response, parser=lambda resp: resp.get_json())
            decisions = {
                revision["id"]: "accept"
                for document in payload["documents"]
                for revision in document["revisions"]
            }
            first_revision_id = next(
                revision["id"]
                for document in payload["documents"]
                for revision in document["revisions"]
                if revision["revision_type"] in {"insert", "modify", "delete"}
            )

            response = client.post(
                "/api/revision_workbench/export_final",
                json={
                    "decisions": decisions,
                    "revision_edits": {first_revision_id: {"text": "Manual export edit"}},
                    "review_author": "产品部审核",
                    "form_data": self.build_form(),
                    "session_operations": [{"actionLabel": "批量接受", "affectedCount": len(decisions)}],
                },
            )

            self.assertEqual(response.status_code, 200)
            self.assertEqual(response.mimetype, "application/zip")
            zip_bytes = self._consume_response(response)
            with zipfile.ZipFile(io.BytesIO(zip_bytes), "r") as zf:
                names = set(zf.namelist())
                self.assertIn("基金合同.docx", names)
                self.assertIn("招募说明书.docx", names)
                self.assertIn("审核决策记录.json", names)
                self.assertIn("审核决策记录.csv", names)
                self.assertIn("同步报告.json", names)
                self.assertNotIn("产品资料概要.docx", names)
                json_payloads = [json.loads(zf.read(name).decode("utf-8")) for name in names if name.endswith(".json")]
                decision_record = next(item for item in json_payloads if "decisions" in item)
                decision_csv = next(zf.read(name).decode("utf-8-sig") for name in names if name.endswith(".csv"))
                report = json.loads(zf.read("同步报告.json").decode("utf-8"))
                contract_xml = zipfile.ZipFile(io.BytesIO(zf.read("基金合同.docx")), "r").read("word/document.xml").decode("utf-8", errors="ignore")
            self.assertIn("产品资料概要未上传", report["product_summary"]["message"])
            self.assertEqual("产品部审核", decision_record["review_author"])
            self.assertEqual(1, decision_record["edited_count"])
            self.assertEqual("Manual export edit", decision_record["revision_edits"][first_revision_id]["text"])
            self.assertIn("edited_text", decision_csv.splitlines()[0])
            self.assertIn("Manual export edit", decision_csv)
            self.assertIn("<w:ins", contract_xml)
            root_start = contract_xml[contract_xml.find("<w:document"):contract_xml.find(">", contract_xml.find("<w:document"))]
            ignorable = re.search(r'mc:Ignorable="([^"]+)"', root_start)
            self.assertIsNotNone(ignorable)
            for prefix in ignorable.group(1).split():
                self.assertIn(f"xmlns:{prefix}=", root_start)
        finally:
            if original_data is None:
                app._revision_workbench_store.pop("data", None)
            else:
                app._revision_workbench_store["data"] = original_data

    def test_sync_review_custody_summary_redlines_prospectus_summary(self):
        custody = self._build_review_docx_buffer(
            title_lines=["测试基金托管协议"],
            body_lines=[
                "一、基金托管协议当事人",
                "托管协议当事人最终正文。",
                "二、基金托管人对基金管理人的业务监督和核查",
                "监督核查最终正文。",
            ],
        ).getvalue()
        prospectus = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "二十一、基金托管协议的内容摘要",
                "旧托管协议摘要正文。",
                "二十二、基金份额持有人服务",
            ],
        ).getvalue()

        updated, report = app._sync_review_custody_summary_redlines(custody, prospectus)

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual("custody_agreement", report["summary"]["source_type"])
        self.assertEqual("redlined", report["apply"]["status"])
        self.assertIn("旧托管协议摘要正文", document_xml)
        self.assertIn("托管协议当事人最终正文", document_xml)
        self.assertIn("<w:del", document_xml)
        self.assertIn("<w:ins", document_xml)

    def test_finalize_review_docx_for_export_marks_update_fields_on_open(self):
        source = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=["目录", "一、绪言"],
        ).getvalue()

        with mock.patch.dict(os.environ, {"DISABLE_WORD_FIELD_UPDATE": "1"}):
            finalized = app._finalize_review_docx_for_export(source)

        settings_xml = self._docx_xml(finalized)["word/settings.xml"]
        self.assertIn("updateFields", settings_xml)
        self.assertIn('w:val="true"', settings_xml)

    def test_sync_review_product_summary_from_final_prospectus_updates_table_rows(self):
        prospectus = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "基金的投资",
                "六、风险收益特征",
                "最终风险收益特征。",
                "基金的费用与税收",
                "一、基金费用的种类",
                "1、基金管理人的管理费；",
                "2、基金托管人的托管费；",
                "3、审计费；",
                "4、上市费；",
            ],
        ).getvalue()
        doc = Document()
        doc.add_paragraph("产品资料概要")
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "项目"
        table.rows[0].cells[1].text = "内容"
        table.rows[1].cells[0].text = "风险收益特征"
        table.rows[1].cells[1].text = "旧风险收益特征。"
        table.rows[2].cells[0].text = "其他费用"
        table.rows[2].cells[1].text = "旧其他费用。"
        product_buffer = io.BytesIO()
        doc.save(product_buffer)

        updated, report = app._sync_review_product_summary_from_prospectus_redlines(
            product_buffer.getvalue(),
            prospectus,
            form_data=self.build_form(
                FUND_NAME="测试交易型开放式指数证券投资基金",
                INDEX_NAME="测试指数",
                EXCHANGE="SSE",
                MARKET_TYPE="A_SHARE",
                PROSPECTUS_RENDER_MODEL={
                    "chapters": [
                        {
                            "title": "基金的投资",
                            "blocks": [
                                {"type": "paragraph", "text": "六、风险收益特征"},
                                {"type": "paragraph", "text": "旧模型风险收益特征。"},
                            ],
                        },
                        {
                            "title": "基金的费用与税收",
                            "blocks": [
                                {"type": "paragraph", "text": "一、基金费用的种类"},
                                {"type": "paragraph", "text": "1、旧模型费用；"},
                            ],
                        },
                    ],
                },
            ),
            filename="product_summary.docx",
        )

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual("synced_from_prospectus", report["status"])
        self.assertIn("旧风险收益特征", document_xml)
        self.assertIn("最终风险收益特征", document_xml)
        self.assertIn("旧其他费用", document_xml)
        self.assertIn("审计费；上市费", document_xml)
        self.assertNotIn("旧模型风险收益特征", document_xml)
        self.assertNotIn("旧模型费用", document_xml)
        self.assertIn("<w:del", document_xml)
        self.assertIn("<w:ins", document_xml)

    def test_sync_review_product_summary_from_final_prospectus_updates_investment_range(self):
        prospectus = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "基金的投资",
                "二、投资范围",
                "最终投资范围包含转融通证券出借业务。",
                "三、主要投资策略",
                "策略正文。",
                "六、风险收益特征",
                "最终风险收益特征。",
                "基金的费用与税收",
                "一、基金费用的种类",
                "1、基金管理人的管理费；",
                "2、基金托管人的托管费；",
                "3、审计费；",
            ],
        ).getvalue()
        doc = Document()
        doc.add_paragraph("产品资料概要")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "项目"
        table.rows[0].cells[1].text = "内容"
        table.rows[1].cells[0].text = "投资范围"
        table.rows[1].cells[1].text = "旧投资范围。"
        product_buffer = io.BytesIO()
        doc.save(product_buffer)

        updated, report = app._sync_review_product_summary_from_prospectus_redlines(
            product_buffer.getvalue(),
            prospectus,
            form_data=self.build_form(
                FUND_NAME="测试交易型开放式指数证券投资基金",
                INDEX_NAME="测试指数",
                EXCHANGE="SSE",
                MARKET_TYPE="A_SHARE",
                PROSPECTUS_RENDER_MODEL={
                    "chapters": [
                        {
                            "title": "基金的投资",
                            "blocks": [
                                {"type": "paragraph", "text": "二、投资范围"},
                                {"type": "paragraph", "text": "旧模型投资范围。"},
                            ],
                        },
                    ],
                },
            ),
            filename="product_summary.docx",
        )

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual("synced_from_prospectus", report["status"])
        self.assertIn("投资范围", report["updated_fields"])
        self.assertIn("旧投资范围", document_xml)
        self.assertIn("最终投资范围包含转融通证券出借业务", document_xml)
        self.assertNotIn("旧模型投资范围", document_xml)
        self.assertIn("<w:del", document_xml)
        self.assertIn("<w:ins", document_xml)

    def test_render_review_contract_summary_uses_rerendered_summary_not_existing_summary(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "contract-template.md"
            template_path.write_text(
                "第二十六部分  基金合同内容摘要\n{CONTRACT_SUMMARY_FEES}",
                encoding="utf-8",
            )
            stale_contract_text = (
                "第一部分  前言\n"
                "合同正文内容\n"
                "第二十六部分  基金合同内容摘要\n"
                "旧摘要内容不应该返回且不应误判一致"
            )
            fake_engine = SimpleNamespace(
                _replace_contract_summary_placeholders=lambda text: text.replace(
                    "{CONTRACT_SUMMARY_FEES}",
                    "新摘要内容应该返回并用于同步判断",
                )
            )

            with mock.patch.object(app, "TEMPLATE_MD", template_path), mock.patch.object(app, "engine", fake_engine):
                summary_text, report = app._render_review_contract_summary_from_text(stale_contract_text)

        self.assertEqual("generated", report["status"])
        self.assertIn("新摘要内容应该返回并用于同步判断", summary_text)
        self.assertNotIn("旧摘要内容不应该返回且不应误判一致", summary_text)

    def test_prospectus_contract_summary_heading_match_accepts_de_particle(self):
        self.assertTrue(app._prospectus_contract_summary_heading_match("二十、基金合同的内容摘要"))
        self.assertTrue(app._prospectus_contract_summary_stop_heading_match("二十一、基金托管协议的内容摘要"))
        self.assertFalse(app._prospectus_contract_summary_heading_match("二十一、基金托管协议的内容摘要"))
        self.assertFalse(app._prospectus_contract_summary_heading_match("二十、基金合同的内容摘要88"))

    def test_prospectus_contract_summary_redline_skips_toc_heading(self):
        source = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "目录",
                "二十、基金合同的内容摘要88",
                "二十一、基金托管协议的内容摘要99",
                "二十、基金合同的内容摘要",
                "旧招募摘要内容需要删除",
                "二十一、基金托管协议的内容摘要",
                "托管协议摘要内容保留",
            ],
        ).getvalue()

        updated, report = app._replace_docx_section_content_with_redline(
            source,
            is_heading=app._prospectus_contract_summary_heading_match,
            is_stop_heading=app._prospectus_contract_summary_stop_heading_match,
            new_text=(
                "第二十六部分  基金合同内容摘要\n"
                "新合同摘要内容需要写入"
            ),
            report_label="招募说明书内基金合同摘要",
        )

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual("redlined", report["status"])
        self.assertGreater(report["old_characters"], 0)
        self.assertIn("<w:delText>旧招募摘要内容需要删除</w:delText>", document_xml)
        self.assertIn("托管协议摘要内容保留", document_xml)

    def test_contract_summary_redline_preserves_signing_page_and_uses_minimal_diff(self):
        source = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第二十六部分  基金合同内容摘要",
                "一、基金费用",
                "本基金管理费率为0.15%，托管费率为0.05%。",
                "（本页为《测试基金基金合同》签署页，无正文）",
                "基金管理人：测试管理人（盖章）",
                "基金托管人：测试托管人（盖章）",
            ],
        ).getvalue()

        updated, report = app._replace_docx_section_content_with_redline(
            source,
            is_heading=app._contract_summary_heading_match,
            is_stop_heading=app._contract_summary_stop_heading_match,
            new_text=(
                "第二十六部分  基金合同内容摘要\n"
                "一、基金费用\n"
                "本基金管理费率为0.20%，托管费率为0.05%。"
            ),
            report_label="基金合同第 26 部分摘要",
        )

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual("redlined", report["status"])
        self.assertIn("签署页", document_xml)
        self.assertIn("基金管理人：测试管理人", document_xml)
        self.assertNotIn("<w:delText>（本页为《测试基金基金合同》签署页，无正文）</w:delText>", document_xml)
        self.assertNotIn("<w:delText>基金管理人：测试管理人（盖章）</w:delText>", document_xml)
        self.assertNotIn("<w:delText>本基金管理费率为0.15%，托管费率为0.05%。</w:delText>", document_xml)
        self.assertIn("<w:delText>0.15%</w:delText>", document_xml)
        self.assertIn("<w:t>0.20%</w:t>", document_xml)
        self._docx_document(updated)

    def test_contract_summary_minimal_redline_preserves_run_formatting(self):
        doc = Document()
        doc.add_paragraph("测试基金基金合同")
        doc.add_paragraph("第二十六部分  基金合同内容摘要")
        doc.add_paragraph("一、基金费用")
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("本基金管理费率为0.15%，托管费率为0.05%。")
        run.bold = True
        doc.add_paragraph("（本页为《测试基金基金合同》签署页，无正文）")
        buffer = io.BytesIO()
        doc.save(buffer)

        updated, report = app._replace_docx_section_content_with_redline(
            buffer.getvalue(),
            is_heading=app._contract_summary_heading_match,
            is_stop_heading=app._contract_summary_stop_heading_match,
            new_text=(
                "第二十六部分  基金合同内容摘要\n"
                "一、基金费用\n"
                "本基金管理费率为0.20%，托管费率为0.05%。"
            ),
            report_label="基金合同第 26 部分摘要",
        )

        self.assertEqual("redlined", report["status"])
        document_xml = self._docx_xml(updated)["word/document.xml"]
        import xml.etree.ElementTree as ET

        root = ET.fromstring(document_xml)
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        delete_run = root.find(".//w:del/w:r", namespace)
        insert_run = root.find(".//w:ins/w:r", namespace)
        self.assertIsNotNone(delete_run)
        self.assertIsNotNone(insert_run)
        self.assertIsNotNone(delete_run.find("w:rPr/w:b", namespace))
        self.assertIsNotNone(insert_run.find("w:rPr/w:b", namespace))

    def test_minimal_redline_uses_whole_number_prefix_for_renumbering(self):
        doc = Document()
        paragraph = doc.add_paragraph("（16）本基金投资存托凭证。")
        buffer = io.BytesIO()
        doc.save(buffer)
        document_xml = self._docx_xml(buffer.getvalue())["word/document.xml"]

        import xml.etree.ElementTree as ET

        app._register_docx_xml_namespaces(ET)
        root = ET.fromstring(document_xml)
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraph_node = root.find(".//w:p", namespace)
        self.assertIsNotNone(paragraph_node)

        app._replace_paragraph_with_minimal_redline(
            ET,
            paragraph_node,
            "（16）本基金投资存托凭证。",
            "（15）本基金投资存托凭证。",
            start_change_id=1,
            author="审核工作台",
            date="2026-05-16T00:00:00Z",
        )

        updated_xml = ET.tostring(root, encoding="unicode")
        self.assertIn("<w:delText>（16）</w:delText>", updated_xml)
        self.assertIn("<w:t>（15）</w:t>", updated_xml)
        self.assertNotIn("<w:t>（</w:t>", updated_xml)

    def test_apply_review_decisions_normalizes_original_leading_number_revisions(self):
        base = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=["（16）本基金投资存托凭证。"],
        )
        source = self._rewrite_docx_xml_buffer(
            base,
            [
                (
                    "<w:r><w:t>（16）本基金投资存托凭证。</w:t></w:r>",
                    (
                        "<w:r><w:t>（</w:t></w:r>"
                        '<w:del w:id="1" w:author="Llinks" w:date="2026-05-09T10:57:00Z">'
                        "<w:r><w:delText>16</w:delText></w:r>"
                        "</w:del>"
                        '<w:ins w:id="2" w:author="Llinks" w:date="2026-05-09T10:57:00Z">'
                        "<w:r><w:t>15</w:t></w:r>"
                        "</w:ins>"
                        "<w:r><w:t>）本基金投资存托凭证。</w:t></w:r>"
                    ),
                )
            ],
        ).getvalue()

        updated, report = app._apply_review_revision_decisions_to_docx(
            source,
            document_kind="prospectus",
            filename="prospectus.docx",
            decisions={"prospectus-reader-1": "accept"},
        )

        self.assertEqual(1, report["accepted_count"])
        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertIn("<w:delText>（16）</w:delText>", document_xml)
        self.assertIn("<w:t>（15）</w:t>", document_xml)
        self.assertNotIn("<w:t>（</w:t><w:del", document_xml)

    def test_contract_summary_redline_renumbers_equal_following_items(self):
        source = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第二十六部分  基金合同内容摘要",
                "一、基金合同当事人的权利、义务",
                "（8）返还在基金交易过程中因任何原因获得的不当得利；",
                "（9）法律法规及中国证监会规定的和《基金合同》约定的其他义务。",
                "第二十七部分  其他事项",
            ],
        ).getvalue()

        updated, report = app._replace_docx_section_content_with_redline(
            source,
            is_heading=app._contract_summary_heading_match,
            is_stop_heading=app._contract_summary_stop_heading_match,
            new_text=(
                "第二十六部分  基金合同内容摘要\n"
                "一、基金合同当事人的权利、义务\n"
                "（8）返还在基金交易过程中因任何原因获得的不当得利；\n"
                "（9）发起资金提供方持有认购的基金份额不少于3年；\n"
                "（10）法律法规及中国证监会规定的和《基金合同》约定的其他义务。"
            ),
            report_label="基金合同第 26 部分摘要",
            review_author="产品部审核",
        )

        self.assertEqual("redlined", report["status"])
        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertIn("发起资金提供方持有认购的基金份额不少于3年", document_xml)
        self.assertIn("<w:delText>（9）</w:delText>", document_xml)
        self.assertIn("<w:t>（10）</w:t>", document_xml)

    def test_contract_summary_redline_preserves_section_break_before_signing_page(self):
        from docx.enum.section import WD_SECTION

        doc = Document()
        doc.add_paragraph("测试基金基金合同")
        doc.add_paragraph("第二十六部分  基金合同内容摘要")
        doc.add_paragraph("一、基金费用")
        doc.add_paragraph("本基金管理费率为0.15%，托管费率为0.05%。")
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_paragraph("（本页为《测试基金基金合同》签署页，无正文）")
        doc.add_paragraph("基金管理人：测试管理人（盖章）")
        buffer = io.BytesIO()
        doc.save(buffer)
        source_xml = self._docx_xml(buffer.getvalue())["word/document.xml"]

        updated, report = app._replace_docx_section_content_with_redline(
            buffer.getvalue(),
            is_heading=app._contract_summary_heading_match,
            is_stop_heading=app._contract_summary_stop_heading_match,
            new_text=(
                "第二十六部分  基金合同内容摘要\n"
                "一、基金费用\n"
                "本基金管理费率为0.20%，托管费率为0.05%。"
            ),
            report_label="基金合同第 26 部分摘要",
        )

        self.assertEqual("redlined", report["status"])
        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual(source_xml.count("<w:sectPr"), document_xml.count("<w:sectPr"))
        self.assertIn("签署页", document_xml)
        self._docx_document(updated)

    def test_contract_summary_redline_preserves_ignorable_namespace_declarations(self):
        source = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第二十六部分  基金合同内容摘要",
                "一、基金费用",
                "本基金管理费率为0.15%，托管费率为0.05%。",
            ],
        )
        source = source.getvalue()

        updated, report = app._replace_docx_section_content_with_redline(
            source,
            is_heading=app._contract_summary_heading_match,
            is_stop_heading=app._contract_summary_stop_heading_match,
            new_text=(
                "第二十六部分  基金合同内容摘要\n"
                "一、基金费用\n"
                "本基金管理费率为0.20%，托管费率为0.05%。"
            ),
            report_label="基金合同第 26 部分摘要",
        )

        self.assertEqual("redlined", report["status"])
        document_xml = self._docx_xml(updated)["word/document.xml"]
        root_start = document_xml[document_xml.find("<w:document"):document_xml.find(">", document_xml.find("<w:document"))]
        self.assertIn('mc:Ignorable="w14 wp14"', root_start)
        self.assertIn('xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"', root_start)

    def test_sync_review_direct_contract_sources_redlines_minimal_prospectus_body_change(self):
        contract = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第十七部分  基金费用与税收",
                "一、基金费用的种类",
                "本基金管理费率为0.20%，托管费率为0.05%。",
            ],
        ).getvalue()
        prospectus = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "十七、基金费用与税收",
                "一、基金费用的种类",
                "本基金管理费率为0.15%，托管费率为0.05%。",
            ],
        ).getvalue()
        rules = [
            {
                "contract_locator": "第十七部分 基金费用与税收 / 一、基金费用的种类",
                "prospectus_locator": "十七、基金费用与税收 / 一、基金费用的种类",
                "contract_section_name": "第十七部分 基金费用与税收",
                "prospectus_section_name": "十七、基金费用与税收",
                "relation": "直接对应",
                "consistency": "完全一致",
            }
        ]

        updated, report = app._sync_review_direct_contract_sources_redlines(
            contract,
            prospectus,
            rules=rules,
        )

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual(1, report["updated_count"])
        self.assertIn("<w:del", document_xml)
        self.assertIn("<w:ins", document_xml)
        self.assertIn("<w:delText>0.15%</w:delText>", document_xml)
        self.assertIn("<w:t>0.20%</w:t>", document_xml)
        self.assertNotIn("<w:delText>本基金管理费率为0.15%，托管费率为0.05%。</w:delText>", document_xml)
        self.assertIn("<w:t>本基金管理费率为</w:t>", document_xml)
        self.assertIn("<w:t>，托管费率为0.05%。</w:t>", document_xml)

    def test_sync_review_template_source_uses_prospectus_postprocess_before_redline(self):
        contract = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第十四部分  基金的投资",
                "一、投资目标",
                "在本基金合同中，投资目标更新为0.20%。",
            ],
        ).getvalue()
        prospectus = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "十一、基金的投资",
                "一、投资目标",
                "在本招募说明书中，投资目标更新为0.10%。",
            ],
        ).getvalue()

        updated, report = app._sync_review_direct_contract_sources_redlines(contract, prospectus)

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual(1, report["updated_count"])
        self.assertEqual("prospectus_template", report["updated"][0]["source"])
        self.assertIn("<w:delText>0.10%</w:delText>", document_xml)
        self.assertIn("<w:t>0.20%</w:t>", document_xml)
        self.assertIn("在本招募说明书中", document_xml)
        self.assertNotIn("在本基金合同中，投资目标更新为0.20%。", document_xml)

    def test_sync_review_direct_contract_sources_skips_prospectus_expansion_rules(self):
        contract = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第八部分  基金份额的申购与赎回",
                "四、申购与赎回的程序",
                "合同原则性程序。",
            ],
        ).getvalue()
        prospectus = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "八、基金份额的申购赎回",
                "四、申购和赎回的程序",
                "招募说明书细化程序。",
            ],
        ).getvalue()
        rules = [
            {
                "contract_locator": "第八部分 基金份额的申购与赎回 / 四、申购与赎回的程序",
                "prospectus_locator": "八、基金份额的申购赎回 / 四、申购和赎回的程序",
                "contract_section_name": "第八部分 基金份额的申购与赎回",
                "prospectus_section_name": "八、基金份额的申购赎回",
                "relation": "直接对应（招募细化）",
                "consistency": "部分对应",
            }
        ]

        updated, report = app._sync_review_direct_contract_sources_redlines(
            contract,
            prospectus,
            rules=rules,
        )

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual(0, report["updated_count"])
        self.assertIn("招募细化", report["skipped"][0]["reason"])
        self.assertNotIn("<w:del", document_xml)
        self.assertNotIn("<w:ins", document_xml)

    def test_sync_review_direct_contract_sources_keeps_prospectus_distribution_conditions(self):
        contract = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第十八部分  基金的收益与分配",
                "一、基金利润的构成",
                "基金利润指基金利息收入、投资收益和其他收入扣除相关费用后的余额。",
                "三、基金收益分配原则",
                "1、每一基金份额享有同等分配权；",
            ],
        ).getvalue()
        prospectus = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "十四、基金的收益与分配",
                "一、基金利润的构成",
                "基金利润指基金利息收入、投资收益和其他收入扣除相关费用后的余额。",
                "三、基金收益分配原则",
                "4、每一基金份额享有同等分配权；",
                "四、收益分配条件",
                "1、本基金在基金收益评价日对基金相对业绩比较基准的超额收益率以及基金的可供分配利润进行评价。",
            ],
        ).getvalue()

        rules = [
            {
                "contract_locator": "第十八部分 基金的收益与分配",
                "prospectus_locator": "十四、基金的收益与分配",
                "contract_section_name": "第十八部分 基金的收益与分配",
                "prospectus_section_name": "十四、基金的收益与分配",
                "relation": "直接对应",
                "consistency": "基本一致",
            }
        ]

        ok, reason = app._review_cross_rule_is_direct_sync_candidate(rules[0])
        self.assertFalse(ok, reason)

        updated, report = app._sync_review_direct_contract_sources_redlines(contract, prospectus, rules=rules)

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual(0, report["updated_count"])
        self.assertIn("四、收益分配条件", document_xml)
        self.assertNotIn("<w:delText>四、收益分配条件</w:delText>", document_xml)
        self.assertGreaterEqual(len(report["skipped"]), 1)

    def test_sync_review_direct_contract_sources_skips_whole_chapter_replacements(self):
        contract = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第十七部分  基金费用与税收",
                "一、基金费用的种类",
                "本基金管理费率为0.20%，托管费率为0.05%。",
            ],
        ).getvalue()
        prospectus = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "十五、基金的费用与税收",
                "一、基金费用的种类",
                "本基金管理费率为0.15%，托管费率为0.05%。",
                "招募说明书补充披露销售服务、申购赎回及信息披露相关费用安排。",
            ],
        ).getvalue()
        rules = [
            {
                "contract_locator": "第十七部分 基金费用与税收",
                "prospectus_locator": "十五、基金的费用与税收",
                "contract_section_name": "第十七部分 基金费用与税收",
                "prospectus_section_name": "十五、基金的费用与税收",
                "relation": "直接对应",
                "consistency": "基本一致",
            }
        ]

        updated, report = app._sync_review_direct_contract_sources_redlines(contract, prospectus, rules=rules)

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual(0, report["updated_count"])
        self.assertIn("招募说明书补充披露销售服务", document_xml)
        self.assertNotIn("<w:delText>招募说明书补充披露销售服务、申购赎回及信息披露相关费用安排。</w:delText>", document_xml)
        self.assertIn("不做整章自动替换", report["skipped"][0]["reason"])

    def test_sync_review_direct_contract_sources_skips_subsection_with_prospectus_extra_content(self):
        contract = self._build_review_docx_buffer(
            title_lines=["测试基金基金合同"],
            body_lines=[
                "第十七部分  基金费用与税收",
                "一、基金费用的种类",
                "本基金管理费率为0.20%，托管费率为0.05%。",
            ],
        ).getvalue()
        prospectus = self._build_review_docx_buffer(
            title_lines=["测试基金招募说明书"],
            body_lines=[
                "十五、基金的费用与税收",
                "一、基金费用的种类",
                "本基金管理费率为0.15%，托管费率为0.05%。",
                "招募说明书补充披露销售服务、申购赎回及信息披露相关费用安排。",
            ],
        ).getvalue()
        rules = [
            {
                "contract_locator": "第十七部分 基金费用与税收 / 一、基金费用的种类",
                "prospectus_locator": "十五、基金的费用与税收 / 一、基金费用的种类",
                "contract_section_name": "第十七部分 基金费用与税收",
                "prospectus_section_name": "十五、基金的费用与税收",
                "relation": "直接对应",
                "consistency": "基本一致",
            }
        ]

        updated, report = app._sync_review_direct_contract_sources_redlines(contract, prospectus, rules=rules)

        document_xml = self._docx_xml(updated)["word/document.xml"]
        self.assertEqual(0, report["updated_count"])
        self.assertIn("招募说明书补充披露销售服务", document_xml)
        self.assertNotIn("<w:delText>招募说明书补充披露销售服务、申购赎回及信息披露相关费用安排。</w:delText>", document_xml)
        self.assertIn("招募说明书存在额外内容", report["skipped"][0]["reason"])

    def test_builtin_cross_rules_do_not_auto_sync_whole_chapters(self):
        unsafe = []
        for rule in app.BUILTIN_ETF_CROSS_RULES:
            contract_locator = rule.get("contract_locator") or rule.get("contract_chapter") or ""
            prospectus_locator = rule.get("prospectus_locator") or rule.get("prospectus_chapter") or ""
            if "/" in contract_locator or "／" in contract_locator or "/" in prospectus_locator or "／" in prospectus_locator:
                continue
            ok, reason = app._review_cross_rule_is_direct_sync_candidate(rule)
            if ok:
                unsafe.append((contract_locator, prospectus_locator, reason))

        self.assertEqual([], unsafe)

    def test_generate_prospectus_chapter_twenty_one_uses_uploaded_custodian_summary_sections(self):
        client = app.app.test_client()
        original_store = getattr(app, "_prospectus_store", None)
        if isinstance(original_store, dict):
            original_store = deepcopy(original_store)

        upload_buffer = self._build_review_docx_buffer(
            title_lines=["南方中证全指红利质量ETF发起联接基金托管协议"],
            body_lines=[
                "一、托管协议当事人",
                "托管协议当事人正文应导入第21章。",
                "二、基金托管人对基金管理人的业务监督和核查",
                "监督和核查正文应导入第21章。",
            ],
        )

        try:
            upload_response = client.post(
                "/api/prospectus/custodian_summary/upload",
                data={"file": (upload_buffer, "sample-custodian.docx")},
                content_type="multipart/form-data",
            )
            self.assertEqual(upload_response.status_code, 200)
            upload_response.close()

            response = client.post("/api/generate_prospectus", json=self.build_form())

            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
            chapter = self.extract_chapter(payload["text"], "二十一")
            self.assertIn("一、基金托管协议当事人", chapter)
            self.assertIn("托管协议当事人正文应导入第21章。", chapter)
            self.assertIn("二、基金托管人对基金管理人的业务监督和核查", chapter)
            self.assertIn("监督和核查正文应导入第21章。", chapter)
            self.assertIn("【待填写】", chapter)
        finally:
            if original_store is None:
                if hasattr(app, "_prospectus_store"):
                    app._prospectus_store.clear()
            else:
                app._prospectus_store.clear()
                app._prospectus_store.update(original_store)

    def test_generate_prospectus_custodian_summary_keeps_supervision_direction_distinct(self):
        client = app.app.test_client()
        original_store = getattr(app, "_prospectus_store", None)
        if isinstance(original_store, dict):
            original_store = deepcopy(original_store)

        upload_buffer = self._build_review_docx_buffer(
            title_lines=["南方中证全指红利质量ETF发起联接基金托管协议"],
            body_lines=[
                "三、基金管理人对基金托管人的业务核查",
                "这里只能落到第三节，不能串到第二节。",
            ],
        )

        try:
            upload_response = client.post(
                "/api/prospectus/custodian_summary/upload",
                data={"file": (upload_buffer, "sample-custodian.docx")},
                content_type="multipart/form-data",
            )
            self.assertEqual(upload_response.status_code, 200)
            upload_response.close()

            response = client.post("/api/generate_prospectus", json=self.build_form())

            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
            chapter = self.extract_chapter(payload["text"], "二十一")
            sec2 = self.extract_section(chapter, "二")
            sec3 = self.extract_section(chapter, "三")
            self.assertIn("【待填写】", sec2)
            self.assertIn("这里只能落到第三节，不能串到第二节。", sec3)
        finally:
            if original_store is None:
                if hasattr(app, "_prospectus_store"):
                    app._prospectus_store.clear()
            else:
                app._prospectus_store.clear()
                app._prospectus_store.update(original_store)

    def test_generate_render_model_keeps_subscription_fee_table_as_structured_block(self):
        model = self.engine.generate_render_model(self.build_form(IMPORTANT_NOTICE_STYLE="DIVIDEND_QUALITY"))
        chapter_six = model["chapters"][5]

        fee_tables = [
            block
            for block in chapter_six["blocks"]
            if isinstance(block, dict)
            and block.get("type") == "table"
            and block.get("table_kind") == "subscription_fee"
        ]

        self.assertEqual(1, len(fee_tables))
        self.assertEqual(
            [["认购份额（S）", "认购费率"], ["S＜100万份", "0.30%"], ["S≥100万份", "每笔500元"]],
            fee_tables[0]["rows"],
        )

    def test_generate_render_model_keeps_section7_reference_tables_as_structured_blocks(self):
        model = self.engine.generate_render_model(self.build_form(IMPORTANT_NOTICE_STYLE="DIVIDEND_QUALITY"))
        chapter_ten = model["chapters"][9]

        section7_tables = [
            block
            for block in chapter_ten["blocks"]
            if isinstance(block, dict)
            and block.get("type") == "table"
            and block.get("table_kind") == "purchase_redemption_list_format"
        ]

        self.assertEqual(["基本信息", "T-1日信息内容", "T日信息内容", "成份股信息内容"], [block.get("title") for block in section7_tables])
        self.assertEqual(["最新公告日期", "现金差额（单位:元）", "预估现金差额（单位:元）", "证券代码"], [block["rows"][0][0] for block in section7_tables])
        self.assertTrue(all(block.get("source") == "reference_docx" for block in section7_tables))

    def test_schema_api_includes_basket_variable_definition(self):
        client = app.app.test_client()
        response = client.get("/api/schema")

        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        variables = payload["groups"]["chapter_10_purchase_redemption"]["variables"]
        keys = {item["key"] for item in variables}
        self.assertIn("BASKET", keys)

    def test_schema_api_includes_user_facing_prospectus_direct_input_variables(self):
        client = app.app.test_client()
        response = client.get("/api/schema")

        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        variables = payload["groups"]["prospectus_direct_inputs"]["variables"]
        keys = {item["key"] for item in variables}

        self.assertTrue({
            "CSRC_APPROVAL_NO",
            "FUND_MANAGER_NAME",
            "FUND_MANAGER_SEX",
            "FUND_MANAGER_BIO",
            "CUSTODIAN_INTRO",
            "ACCOUNTANT",
        }.issubset(keys))

    def test_variables_endpoint_updates_schema_variable_definitions(self):
        schema = {
            "version": "test",
            "groups": {
                "basic": {
                    "description": "基本变量",
                    "variables": [
                        {
                            "key": "OLD_FLAG",
                            "type": "boolean",
                            "ui_entry": "旧开关",
                            "default": False,
                            "applies_to": "旧用途",
                        }
                    ],
                }
            },
        }
        client = app.app.test_client()
        with tempfile.TemporaryDirectory() as tmpdir:
            schema_path = Path(tmpdir) / "schema.json"
            schema_path.write_text(json.dumps(schema, ensure_ascii=False), encoding="utf-8")
            with mock.patch.object(app, "SCHEMA_JSON", schema_path):
                list_response = client.get("/api/variables")
                self.assertEqual(list_response.status_code, 200)
                list_payload = self._consume_response(list_response, parser=lambda resp: resp.get_json())
                self.assertEqual(1, list_payload["count"])
                self.assertEqual(1, list_payload["condition_count"])
                self.assertEqual("OLD_FLAG", list_payload["items"][0]["key"])

                update_response = client.put(
                    "/api/variables/OLD_FLAG",
                    json={
                        "key": "NEW_FLAG",
                        "label": "新开关",
                        "type": "boolean",
                        "default": True,
                        "group": "basic",
                        "usage": "新用途",
                    },
                )
                self.assertEqual(update_response.status_code, 200)

                create_response = client.post(
                    "/api/variables",
                    json={
                        "key": "CUSTOM_TEXT",
                        "label": "自定义文本",
                        "type": "textarea",
                        "default": "默认文本",
                        "group": "custom",
                        "group_label": "自定义变量",
                    },
                )
                self.assertEqual(create_response.status_code, 200)

                delete_response = client.delete("/api/variables/CUSTOM_TEXT")
                self.assertEqual(delete_response.status_code, 200)

            saved = json.loads(schema_path.read_text(encoding="utf-8"))
            variables = saved["groups"]["basic"]["variables"]
            self.assertEqual("NEW_FLAG", variables[0]["key"])
            self.assertEqual("新开关", variables[0]["ui_entry"])
            self.assertTrue(variables[0]["default"])
            self.assertEqual("新用途", variables[0]["applies_to"])
            self.assertEqual([], saved["groups"]["custom"]["variables"])

    def test_index_template_includes_variable_library_panel(self):
        index_html = Path(app.__file__).with_name("templates").joinpath("index.html").read_text(encoding="utf-8")

        self.assertIn('data-panel="variables"', index_html)
        self.assertIn('id="variable-library-rows"', index_html)
        self.assertIn('/api/variables', index_html)
        self.assertIn('openVariableEditor', index_html)
        self.assertIn('id="variable-type-hint"', index_html)
        self.assertIn('VARIABLE_TYPE_HINTS', index_html)
        self.assertIn('applyVariableTypeHints', index_html)
        self.assertIn('SHOW_CUSTOM_NOTICE', index_html)
        self.assertIn('id="variable-template-snippet"', index_html)
        self.assertIn('canonicalVariableType', index_html)
        self.assertIn('buildVariableTemplateSuggestion', index_html)
        self.assertNotIn('option value="condition_boolean"', index_html)
        self.assertNotIn('option value="direct_input"', index_html)
        self.assertNotIn('option value="direct_input_enum"', index_html)
        self.assertNotIn('option value="derived_clause_placeholder"', index_html)

    def test_generate_prospectus_api_rejects_ambiguous_notice_style_in_auto_mode(self):
        client = app.app.test_client()
        response = client.post(
            "/api/generate_prospectus",
            json=self.build_form(
                INDEX_NAME="中证创业板双创指数",
                INDEX_DESCRIPTION="双创指数反映相关上市公司证券的整体表现。",
            ),
        )

        self.assertEqual(response.status_code, 400)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        self.assertFalse(payload["success"])
        self.assertIn("重要提示", payload["error"])

    def test_summary_cross_check_prefers_contract_summary_over_custodian_summary(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        contract_summary_content = (
            "一、合同摘要第一项\n"
            "这是基金合同摘要的核心内容，包含基金份额、申购赎回、投资范围、收益分配、费用承担等关键安排，"
            "用于验证合同摘要比对时应当选择基金合同对应章节而不是其他摘要章节。\n"
            "二、合同摘要第二项\n"
            "这是另一段合同摘要内容，继续补充基金合同中的重要约定、风险安排、信息披露和争议解决机制。"
        )
        custody_summary_content = (
            "一、托管协议摘要第一项\n"
            "这里是完全不同的托管协议内容，主要描述托管人职责、账户管理、资金清算、监督机制和交接流程，"
            "故意与基金合同摘要保持明显差异，以便准确识别误选托管协议摘要的问题。\n"
            "二、托管协议摘要第二项\n"
            "这里继续描述托管协议安排，包括估值复核、会计核算、资料保管、信息提供以及异常情况处理。"
        )
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {
                    "heading": "第二十六部分 基金合同内容摘要",
                    "content": contract_summary_content,
                }
            ],
            "prospectus_sections": [
                {
                    "heading": "二十、基金合同的内容摘要",
                    "content": contract_summary_content,
                },
                {
                    "heading": "二十一、基金托管协议的内容摘要",
                    "content": custody_summary_content,
                },
            ],
        }

        try:
            response = client.post("/api/review/summary_cross_check", json={})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        self.assertEqual(payload["prospectus_summary_heading"], "二十、基金合同的内容摘要")
        self.assertEqual(payload["contract_summary_heading"], "第二十六部分 基金合同内容摘要")
        self.assertEqual(payload["similarity"], 100.0)

    def test_summary_check_excludes_signing_page_from_effect_section_source_text(self):
        client = app.app.test_client()
        effect_content = (
            "基金合同自基金份额持有人大会审议通过并经监管机构履行程序后生效。\n"
            "基金份额持有人、基金管理人和基金托管人自基金合同生效之日起享有相应权利并承担相应义务。\n"
            "本条款用于验证基金合同的效力章节在摘要核对时不应夹带签署页文字。\n"
            "（本页无正文，为《测试基金合同》签署页）\n"
            "基金管理人：测试管理人（盖章）\n"
            "基金托管人：测试托管人（盖章）"
        )
        summary_content = (
            "一、基金合同的效力\n"
            "基金合同自基金份额持有人大会审议通过并经监管机构履行程序后生效。\n"
            "基金份额持有人、基金管理人和基金托管人自基金合同生效之日起享有相应权利并承担相应义务。\n"
            "本条款用于验证基金合同的效力章节在摘要核对时不应夹带签署页文字。"
        )
        contract_text = (
            "第一部分 前言\n"
            "这是前言部分的说明文字，用于充实测试文本长度，避免被章节过滤逻辑误判。\n"
            "第七部分 基金合同的生效\n"
            f"{effect_content}\n"
            "第二十六部分 基金合同内容摘要\n"
            f"{summary_content}\n"
        )

        response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        self.assertEqual(payload["total_subsections"], 1)
        result = payload["results"][0]

        self.assertEqual(result["matched_section"], "第七部分 基金合同的生效")
        self.assertNotIn("本页无正文", result["matched_section_text"])
        self.assertNotIn("签署页", result["matched_section_text"])
        self.assertNotIn("（盖章）", result["matched_section_text"])

    def test_summary_check_returns_original_source_texts_for_diff_review(self):
        client = app.app.test_client()
        body_content = (
            "基金合同约定基金管理人应当依法履行审慎勤勉义务，并及时办理信息披露、估值和收益分配事项。\n"
            "基金托管人应当按照法律法规和基金合同约定履行监督职责，确保基金财产安全。\n"
            "本段用于验证摘要核对结果中需要回传原始正文和摘要文本，方便直接核看原文。\n"
            "基金管理人与基金托管人还应当按照约定保存记录、接受监督并配合信息披露义务履行。"
        )
        summary_content = (
            "一、管理人与托管人义务\n"
            "基金合同约定基金管理人应当依法履行审慎勤勉义务，并及时办理信息披露、估值和收益分配事项。\n"
            "基金托管人应当按照法律法规和基金合同约定履行监督职责，确保基金财产安全。\n"
            "基金管理人与基金托管人还应当按照约定保存记录、接受监督并配合信息披露义务履行。"
        )
        contract_text = (
            "第一部分 前言\n"
            "这是前言部分的说明文字，用于充实测试文本长度，避免被章节过滤逻辑误判。\n"
            "第十二部分 基金管理人与基金托管人\n"
            f"{body_content}\n"
            "第二十六部分 基金合同内容摘要\n"
            f"{summary_content}\n"
        )

        response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        result = payload["results"][0]

        self.assertEqual(result["matched_section_text"], body_content)
        self.assertEqual(result["summary_text"], summary_content.split("\n", 1)[1])

    def test_summary_check_uses_uploaded_contract_text_when_request_omits_text(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        contract_text = (
            "第一部分 前言\n"
            "这是前言部分的说明文字，用于充实测试文本长度，避免被章节过滤逻辑误判。\n"
            "第七部分 基金合同的生效\n"
            "基金合同自基金份额持有人大会审议通过并经监管机构履行程序后生效。\n"
            "基金份额持有人、基金管理人和基金托管人自基金合同生效之日起享有相应权利并承担相应义务。\n"
            "第二十六部分 基金合同内容摘要\n"
            "一、基金合同的效力\n"
            "基金合同自基金份额持有人大会审议通过并经监管机构履行程序后生效。\n"
            "基金份额持有人、基金管理人和基金托管人自基金合同生效之日起享有相应权利并承担相应义务。\n"
        )
        app._review_store["data"] = {"contract_text": contract_text}

        try:
            response = client.post("/api/summary_check", json={})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        self.assertEqual(payload["total_subsections"], 1)
        self.assertEqual(payload["results"][0]["matched_section"], "第七部分 基金合同的生效")

    def test_summary_cross_check_treats_self_reference_wording_as_literal_difference(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        contract_summary_content = (
            "一、基金合同的效力\n"
            "本基金合同自基金份额持有人大会审议通过后生效，并对基金合同当事人具有法律约束力。\n"
            "本基金合同的变更、终止与清算安排按照法律法规和基金合同约定执行。\n"
            "本段用于验证摘要交叉比对时，本基金合同与基金合同的口径差异不应被视为实质性差异。"
        )
        prospectus_summary_content = (
            "一、基金合同的效力\n"
            "基金合同自基金份额持有人大会审议通过后生效，并对基金合同当事人具有法律约束力。\n"
            "基金合同的变更、终止与清算安排按照法律法规和基金合同约定执行。\n"
            "本段用于验证摘要交叉比对时，基金合同与本基金合同的口径差异不应被视为实质性差异。"
        )
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {
                    "heading": "第二十六部分 基金合同内容摘要",
                    "content": contract_summary_content,
                }
            ],
            "prospectus_sections": [
                {
                    "heading": "二十、基金合同的内容摘要",
                    "content": prospectus_summary_content,
                }
            ],
        }

        try:
            response = client.post("/api/review/summary_cross_check", json={})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        self.assertTrue(payload["is_problem"])
        self.assertLess(payload["similarity"], 100.0)
        self.assertGreater(payload["changed_lines"], 0)

    def test_summary_cross_check_strips_signing_page_from_contract_summary(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        contract_summary_content = (
            "一、基金合同的效力\n"
            "基金合同自基金份额持有人大会审议通过后生效，并对基金合同当事人具有法律约束力。\n"
            "本段用于验证合同摘要与签署页之间必须强制分离，摘要交叉比对不能把签署页并入摘要正文。\n"
            "（本页无正文，为《测试基金合同》签署页）\n"
            "基金管理人：测试管理人（盖章）\n"
            "基金托管人：测试托管人（盖章）"
        )
        prospectus_summary_content = (
            "一、基金合同的效力\n"
            "基金合同自基金份额持有人大会审议通过后生效，并对基金合同当事人具有法律约束力。\n"
            "本段用于验证合同摘要与签署页之间必须强制分离，摘要交叉比对不能把签署页并入摘要正文。"
        )
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {
                    "heading": "第二十六部分 基金合同内容摘要",
                    "content": contract_summary_content,
                }
            ],
            "prospectus_sections": [
                {
                    "heading": "二十、基金合同的内容摘要",
                    "content": prospectus_summary_content,
                }
            ],
        }

        try:
            response = client.post("/api/review/summary_cross_check", json={})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        self.assertEqual(payload["similarity"], 100.0)
        self.assertEqual(payload["changed_lines"], 0)
        contract_html = "\n".join(hunk["contract_html"] for hunk in payload["hunks"])
        self.assertNotIn("签署页", contract_html)
        self.assertNotIn("盖章", contract_html)

    def test_summary_check_sets_is_problem_for_matched_and_unmatched_rows(self):
        client = app.app.test_client()
        workbook = self._build_summary_chapter_locator_workbook(
            "第十八部分 基金的收益与分配",
            "三、基金收益分配原则、执行方式",
        )
        matched_body = (
            "基金收益分配应当依照合同约定的时间、登记与执行要求推进，并在披露口径上保持一致。"
        )
        contract_text = "\n".join(
            [
                "第一部分 前言",
                "前言部分用于填充测试文本。",
                "第十八部分 基金的收益与分配",
                "（三）基金收益分配原则及执行方式",
                matched_body,
                "第二十六部分 基金合同内容摘要",
                "三、基金收益分配原则、执行方式",
                matched_body,
                "四、完全无关的摘要小标题",
                "这一段摘要文本与任何正文章节都没有对应关系，用于验证 is_problem 标记。",
            ]
        )

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        self.assertEqual(payload["total_subsections"], 2)
        self.assertFalse(payload["results"][0]["is_problem"])
        self.assertTrue(payload["results"][1]["is_problem"])

    def test_summary_check_normalizes_internal_clause_reference_style_for_full_inclusion_rule(self):
        client = app.app.test_client()
        workbook = self._build_summary_rules_workbook(
            chapter_rows=[
                {
                    "contract_pos": "第十部分 基金份额持有人大会",
                    "summary_pos": "二、基金份额持有人大会召集、议事及表决的程序和规则",
                    "status": "已收录",
                    "method": "近全文摘录",
                    "detail": "摘要应与正文保持一致，仅允许内部编号口径差异。",
                }
            ]
        )
        chapter_body = (
            "2、在法律法规规定和《基金合同》约定的范围内且对基金份额持有人利益无实质性不利影响的前提下，"
            "以下情况可由基金管理人和基金托管人协商后修改，不需召开基金份额持有人大会：\n"
            "(1) 法律法规要求增加的基金费用的收取；\n"
            "(2) 在法律法规和《基金合同》规定的范围内调整本基金的申购费率、变更收费方式；\n"
            "2、议事程序\n"
            "(1) 现场开会\n"
            "在现场开会的方式下，首先由大会主持人按照下列第七条规定程序确定和公布监票人，然后由大会主持人宣读提案，经讨论后进行表决，并形成大会决议。"
        )
        summary_body = (
            "2、在法律法规规定和《基金合同》约定的范围内且对基金份额持有人利益无实质性不利影响的前提下，"
            "以下情况可由基金管理人和基金托管人协商后修改，不需召开基金份额持有人大会：\n"
            "(1) 法律法规要求增加的基金费用的收取；\n"
            "(2) 在法律法规和《基金合同》规定的范围内调整本基金的申购费率、变更收费方式；\n"
            "2、议事程序\n"
            "(1) 现场开会\n"
            "在现场开会的方式下，首先由大会主持人按照下列第（七）条规定程序确定和公布监票人，然后由大会主持人宣读提案，经讨论后进行表决，并形成大会决议。"
        )
        contract_text = "\n".join(
            [
                "第一部分 前言",
                "这是前言部分的说明文字，用于充实测试文本长度，避免章节过滤逻辑误判，同时确保摘要定位稳定。",
                "第十部分 基金份额持有人大会",
                chapter_body,
                "第二十六部分 基金合同内容摘要",
                "二、基金份额持有人大会召集、议事及表决的程序和规则",
                summary_body,
            ]
        )

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result = payload["results"][0]
        self.assertFalse(result["is_problem"])
        self.assertEqual("info", result["severity"])
        self.assertEqual(0, result["changed_lines"])

    def test_summary_check_flags_substantive_difference_for_full_inclusion_rule(self):
        client = app.app.test_client()
        workbook = self._build_summary_rules_workbook(
            chapter_rows=[
                {
                    "contract_pos": "第十部分 基金份额持有人大会",
                    "summary_pos": "二、基金份额持有人大会召集、议事及表决的程序和规则",
                    "status": "已收录",
                    "method": "近全文摘录",
                    "detail": "摘要应与正文保持一致，不应新增实质性收费调整内容。",
                }
            ]
        )
        chapter_body = (
            "2、在法律法规规定和《基金合同》约定的范围内且对基金份额持有人利益无实质性不利影响的前提下，"
            "以下情况可由基金管理人和基金托管人协商后修改，不需召开基金份额持有人大会：\n"
            "(1) 法律法规要求增加的基金费用的收取；\n"
            "(2) 在法律法规和《基金合同》规定的范围内调整本基金的申购费率、变更收费方式；\n"
            "(3) 基金管理人、相关证券交易所、登记机构、销售机构调整有关基金认购、申购、赎回、交易、非交易过户、收益分配、信息披露等业务的规则。"
        )
        summary_body = (
            "2、在法律法规规定和《基金合同》约定的范围内且对基金份额持有人利益无实质性不利影响的前提下，"
            "以下情况可由基金管理人和基金托管人协商后修改，不需召开基金份额持有人大会：\n"
            "(1) 法律法规要求增加的基金费用的收取；\n"
            "(2) 在法律法规和《基金合同》规定的范围内调整本基金的申购费率、调低赎回费率、变更收费方式；\n"
            "(3) 基金管理人、相关证券交易所、登记机构、销售机构调整有关基金认购、申购、赎回、交易、非交易过户、收益分配、信息披露等业务的规则。"
        )
        contract_text = "\n".join(
            [
                "第一部分 前言",
                "这是前言部分的说明文字，用于充实测试文本长度，避免章节过滤逻辑误判，同时确保摘要定位稳定。",
                "第十部分 基金份额持有人大会",
                chapter_body,
                "第二十六部分 基金合同内容摘要",
                "二、基金份额持有人大会召集、议事及表决的程序和规则",
                summary_body,
            ]
        )

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result = payload["results"][0]
        self.assertTrue(result["is_problem"])
        self.assertEqual("error", result["severity"])
        self.assertGreater(result["changed_lines"], 0)

    def test_summary_cross_check_returns_problem_flag_locators_and_clean_texts(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        contract_summary_content = (
            "一、基金合同的效力\n"
            "基金合同自基金份额持有人大会审议通过后生效，并对基金合同当事人具有法律约束力。"
        )
        prospectus_summary_content = (
            "一、基金合同的效力\n"
            "基金合同自基金份额持有人大会审议通过后生效，并对基金合同当事人具有法律约束力。"
        )
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {
                    "heading": "第二十六部分 基金合同内容摘要",
                    "content": contract_summary_content,
                }
            ],
            "prospectus_sections": [
                {
                    "heading": "二十、基金合同的内容摘要",
                    "content": prospectus_summary_content,
                }
            ],
        }

        try:
            response = client.post("/api/review/summary_cross_check", json={})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        self.assertFalse(payload["is_problem"])
        self.assertEqual("第二十六部分 基金合同内容摘要", payload["contract_summary_locator"])
        self.assertEqual("二十、基金合同的内容摘要", payload["prospectus_summary_locator"])
        self.assertEqual(contract_summary_content, payload["contract_summary_text"])
        self.assertEqual(prospectus_summary_content, payload["prospectus_summary_text"])

    def test_review_cross_check_sets_is_problem_for_missing_and_pass_rows(self):
        pass_payload = self._run_cross_check_with_review_rules(
            [
                {
                    "heading": "第十六部分 基金资产估值",
                    "content": "2、基金管理人应每个估值日对基金资产估值，但基金管理人根据法律法规或基金合同的规定暂停估值时除外。",
                }
            ],
            [
                {
                    "heading": "基金资产估值",
                    "content": "2、基金管理人应每个估值日对基金资产估值，但基金管理人根据法律法规或基金合同的规定暂停估值时除外。",
                }
            ],
            {
                "chapter_rules": [
                    {
                        "source": "chapter_rules",
                        "contract": "第十六部分 基金资产估值",
                        "prospectus": "基金资产估值",
                        "relation": "直接对应",
                        "consistency": "完全一致",
                        "detail": "基金合同与招募说明书存在文件自指口径差异。",
                        "sheet_name": "测试",
                    }
                ],
                "detail_rules": [],
                "key_diff_rules": [],
                "chapter_level": [],
                "key_diffs": [],
            },
        )
        self.assertFalse(pass_payload["results"][0]["is_problem"])

        missing_payload = self._run_cross_check_with_review_rules(
            [
                {
                    "heading": "第八部分 基金份额的申购与赎回",
                    "content": "三、申购与赎回的原则\n申购和赎回遵循相同的基本要求。",
                }
            ],
            [
                {
                    "heading": "基金份额的申购赎回",
                    "content": "七、申购赎回清单的内容与格式\n招募说明书仅保留了清单内容，缺少程序段落。",
                }
            ],
            {
                "chapter_rules": [
                    {
                        "source": "chapter_rules",
                        "contract": "第八部分 基金份额的申购与赎回",
                        "prospectus": "基金份额的申购赎回",
                        "relation": "直接对应",
                        "consistency": "完全一致",
                        "detail": "",
                        "sheet_name": "测试",
                    }
                ],
                "detail_rules": [
                    {
                        "source": "detail_rules",
                        "contract": "第八部分 基金份额的申购与赎回 / 四、申购和赎回的程序",
                        "prospectus_point": "基金份额的申购赎回 / 四、申购和赎回的程序",
                        "relation": "直接对应",
                        "consistency": "完全一致",
                        "detail": "",
                        "sheet_name": "测试",
                    }
                ],
                "key_diff_rules": [],
                "chapter_level": [],
                "key_diffs": [],
            },
        )
        self.assertTrue(missing_payload["results"][0]["is_problem"])

    def test_build_ai_cross_review_items_use_precise_locators_and_middle_windowing(self):
        long_contract = "A" * 7000 + "CONTRACT-MIDDLE" + "B" * 7000
        long_prospectus = "C" * 7000 + "PROSPECTUS-MIDDLE" + "D" * 7000
        items = app._build_ai_cross_review_items(
            [
                {
                    "heading": "第八部分 基金份额的申购与赎回",
                    "content": "四、申购和赎回的程序\n" + long_contract,
                }
            ],
            [
                {
                    "heading": "基金份额的申购赎回",
                    "content": "四、申购和赎回的程序\n" + long_prospectus,
                }
            ],
            {
                "chapter_rules": [
                    {
                        "source": "chapter_rules",
                        "contract": "第八部分 基金份额的申购与赎回 / 四、申购和赎回的程序",
                        "prospectus": "基金份额的申购赎回 / 四、申购和赎回的程序",
                        "relation": "直接对应",
                        "consistency": "完全一致",
                        "detail": "",
                        "sheet_name": "测试",
                    }
                ],
                "detail_rules": [],
                "key_diff_rules": [],
                "chapter_level": [],
                "key_diffs": [],
            },
            fund_type="ETF",
        )

        self.assertEqual(1, len(items))
        item = items[0]
        self.assertIn("第八部分 基金份额的申购与赎回 / 四、申购和赎回的程序", item["label"])
        self.assertIn("基金份额的申购赎回 / 四、申购和赎回的程序", item["label"])
        self.assertIn("已省略中间", item["contract_text"])
        self.assertIn("已省略中间", item["prospectus_text"])
        self.assertNotIn("CONTRACT-MIDDLE", item["contract_text"])
        self.assertNotIn("PROSPECTUS-MIDDLE", item["prospectus_text"])

    def test_build_ai_summary_review_items_reuse_summary_locator_targets(self):
        matched_body = (
            "基金收益分配的执行应当按照合同约定的时间、登记与披露安排推进，并保持口径一致。"
        )
        items = app._build_ai_summary_review_items(
            [
                {
                    "heading": "第十八部分 基金的收益与分配",
                    "content": "（三）基金收益分配原则及执行方式\n" + matched_body,
                },
                {
                    "heading": "第二十六部分 基金合同内容摘要",
                    "content": "三、基金收益分配原则、执行方式\n" + matched_body,
                },
            ],
            {
                "summary_rules": [
                    {
                        "summary_pos": "三、基金收益分配原则、执行方式",
                        "contract_pos": "第十八部分 基金的收益与分配 / （三）基金收益分配原则及执行方式",
                        "method": "摘录/转录",
                        "detail": "优先按 Excel locator 定位到小标题。",
                        "rule_level": "detail",
                    }
                ]
            },
        )

        self.assertEqual(1, len(items))
        self.assertIn(
            "第十八部分 基金的收益与分配 / （三）基金收益分配原则及执行方式",
            items[0]["label"],
        )
        self.assertEqual(matched_body, items[0]["summary_text"])
        self.assertTrue(items[0]["contract_text"].startswith(matched_body))

    def test_load_review_rules_parses_etf_workbook_meta_detail_and_reference_rules(self):
        workbook = self._build_review_rules_workbook(linked=False)
        with self._with_review_rules(workbook):
            rules = app._load_review_rules("ETF")

        self.assertIn("workbook_meta", rules)
        self.assertIn("chapter_rules", rules)
        self.assertIn("detail_rules", rules)
        self.assertIn("summary_rules", rules)
        self.assertIn("key_diff_rules", rules)
        self.assertEqual(rules["workbook_meta"]["special_notes"][0], "目录标题与正文章节标题可能不一致。")
        self.assertEqual(rules["detail_rules"][0]["detail"], "招募补充基金份额净值计算公式和代理券商佣金限额。")
        self.assertEqual(rules["key_diff_rules"][0]["description"], "四、基金份额参考净值（IOPV）的计算与公告")

    def test_load_review_rules_parses_linked_fund_summary_orientation(self):
        workbook = self._build_review_rules_workbook(linked=True)
        with self._with_review_rules(workbook):
            rules = app._load_review_rules("\u8054\u63a5\u57fa\u91d1")

        self.assertEqual(rules["workbook_meta"]["fund_type"], "ETF联接基金（发起式）")
        self.assertEqual(rules["summary_rules"][0]["contract_pos"], "第十六部分 / 三、基金收益分配原则")
        self.assertEqual(rules["summary_rules"][0]["summary_pos"], "三、（一）基金收益分配原则")
        self.assertEqual(rules["summary_rules"][0]["detail"], "仅编号层级调整")

    def test_find_first_existing_review_workbook_prefers_fund_name_match(self):
        tmp_path = Path("D:/codex/_tmp_review_workbook_match")
        if tmp_path.exists():
            shutil.rmtree(tmp_path)
        tmp_path.mkdir(parents=True)
        etf_path = tmp_path / "\u5357\u65b9\u4e2d\u8bc1\u5168\u6307\u7ea2\u5229\u8d28\u91cfETF_\u52fe\u7a3d\u5173\u7cfb\u6574\u7406.xlsx"
        linked_path = tmp_path / "\u5357\u65b9\u4e2d\u8bc1\u901a\u7528\u822a\u7a7a\u4e3b\u9898ETF\u53d1\u8d77\u5f0f\u8054\u63a5\u57fa\u91d1_\u52fe\u7a3d\u5173\u7cfb\u6574\u7406.xlsx"
        etf_path.touch()
        linked_path.touch()

        original_candidates = app.REVIEW_XLSX_CANDIDATES
        original_data = app._review_store.get("data")
        app.REVIEW_XLSX_CANDIDATES = [etf_path, linked_path]
        app._review_store["data"] = {
            "fund_name": "\u5357\u65b9\u4e2d\u8bc1\u901a\u7528\u822a\u7a7a\u4e3b\u9898ETF\u53d1\u8d77\u5f0f\u8054\u63a5\u57fa\u91d1"
        }

        try:
            selected = app._find_first_existing_review_workbook("ETF")
        finally:
            app.REVIEW_XLSX_CANDIDATES = original_candidates
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data
            if tmp_path.exists():
                shutil.rmtree(tmp_path)

        self.assertEqual(str(linked_path), selected)

    def test_summary_check_prefers_excel_rule_match_over_similarity(self):
        client = app.app.test_client()
        workbook = self._build_review_rules_workbook(linked=False)
        summary_line = (
            "\u8fd9\u6bb5\u6458\u8981\u6761\u6b3e\u4e13\u95e8\u7528\u4e8e\u8bef\u5bfc\u7eaf\u76f8\u4f3c\u5ea6\u5339\u914d\uff0c"
            "\u540c\u65f6\u8865\u5145\u4e86\u4e00\u4e9b\u6267\u884c\u6027\u5b89\u6392\u3001\u62ab\u9732\u4e49\u52a1\u548c\u9009\u62e9\u6027\u6458\u5f55\u8bf4\u660e\uff0c"
            "\u7528\u4e8e\u9a8c\u8bc1\u89c4\u5219\u4f18\u5148\u65f6\u4e0d\u5e94\u88ab\u7eaf\u76f8\u4f3c\u5ea6\u7ed3\u679c\u5e26\u504f\uff0c"
            "\u540c\u65f6\u8865\u5145\u8bf4\u660e\u8fd9\u4e9b\u5185\u5bb9\u53ea\u4f5c\u4e3a\u6458\u8981\u4fdd\u7559\u7684\u6838\u5fc3\u6267\u884c\u6027\u6761\u6b3e\uff0c"
            "\u800c\u4e0d\u662f\u5bf9\u539f\u6587\u7ae0\u8282\u7684\u5168\u6587\u6444\u5f55\u3002"
        )
        contract_text = "\n".join([
            "\u7b2c\u4e00\u90e8\u5206 \u524d\u8a00",
            "\u8fd9\u662f\u524d\u8a00\uff0c\u7528\u4e8e\u62c9\u957f\u6d4b\u8bd5\u6587\u672c\u5e76\u907f\u514d\u88ab\u7ae0\u8282\u8fc7\u6ee4\u903b\u8f91\u5ffd\u7565\u3002",
            "\u7b2c\u4e8c\u90e8\u5206 \u9ad8\u76f8\u4f3c\u7ae0\u8282",
            summary_line,
            "\u7b2c\u4e09\u90e8\u5206 \u76ee\u6807\u7ae0\u8282",
            "\u8fd9\u662f\u6b63\u786e\u7684\u6b63\u6587\u6765\u6e90\uff0c\u53ea\u4fdd\u7559\u4e86\u6838\u5fc3\u6267\u884c\u6027\u6761\u6b3e\uff0c"
            "\u4e0e\u6458\u8981\u5c5e\u4e8e\u9009\u62e9\u6027\u6458\u5f55\u5173\u7cfb\uff0c\u8fd8\u989d\u5916\u8bf4\u660e\u4e86\u5b9e\u65bd\u7a0b\u5e8f\u3001"
            "\u62ab\u9732\u8981\u6c42\u548c\u4e0d\u7eb3\u5165\u6458\u8981\u7684\u8fb9\u754c\uff0c\u7528\u4e8e\u9a8c\u8bc1 Excel \u89c4\u5219\u53ef\u4ee5\u4f18\u5148\u6307\u5411\u6b64\u7ae0\u3002",
            "\u7b2c\u4e8c\u5341\u516d\u90e8\u5206 \u57fa\u91d1\u5408\u540c\u5185\u5bb9\u6458\u8981",
            "\u4e00\u3001\u6458\u8981\u6761\u6b3e",
            summary_line,
        ])

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result = payload["results"][0]
        self.assertEqual(result["matched_section"], "\u7b2c\u4e09\u90e8\u5206 \u76ee\u6807\u7ae0\u8282")
        self.assertEqual(result["rule_source"], "summary_rules")
        self.assertEqual(result["expected_status"], "\u5df2\u6536\u5f55")
        self.assertIn("\u9009\u62e9\u6027\u6458\u5f55", result["expected_method"])

    def test_summary_check_returns_locator_fields_for_excel_detail_rule(self):
        client = app.app.test_client()
        workbook = self._build_summary_detail_locator_workbook(
            "第八部分 基金份额的申购与赎回 / 四、申购和赎回的程序"
        )
        subsection_body = (
            "本小标题用于验证摘要复核会优先命中Excel指定的小标题，而不是只停留在整章层面。"
            "正文同时保留程序确认、清算交收、登记处理与投资者当日卖出安排等执行性描述，以便与摘要内容形成稳定匹配。"
            "这里额外补充复核、披露、交收节奏和登记确认等描述，用于确保摘要章节长度稳定超过过滤阈值。"
        )
        contract_text = "\n".join([
            "第一部分 前言",
            "这是前言部分的说明文字，用于充实测试文本长度，避免被章节过滤逻辑误判，同时确保摘要定位不会误回退到前言章节。",
            "第八部分 基金份额的申购与赎回",
            "一、申购和赎回场所",
            "投资人可按照约定办理申购和赎回，相关安排由基金管理人另行披露。",
            "四、申购和赎回的程序",
            subsection_body,
            "五、申购和赎回的数额限制",
            "基金管理人可以根据业务安排调整申购赎回数量限制，并按规定公告。",
            "第二十六部分 基金合同内容摘要",
            "一、摘要条款",
            subsection_body,
        ])

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result = payload["results"][0]
        self.assertEqual(result["matched_section"], "第八部分 基金份额的申购与赎回")
        self.assertEqual(result["matched_target_heading"], "四、申购和赎回的程序")
        self.assertEqual(result["matched_locator"], "第八部分 基金份额的申购与赎回 / 四、申购和赎回的程序")
        self.assertEqual(result["rule_locator"], "第八部分 基金份额的申购与赎回 / 四、申购和赎回的程序")
        self.assertEqual(result["rule_level"], "detail")
        self.assertTrue(result["locator_matched"])
        self.assertTrue(result["matched_section_text"].startswith("四、申购和赎回的程序"))

    def test_summary_check_marks_missing_when_excel_detail_subheading_not_found(self):
        client = app.app.test_client()
        workbook = self._build_summary_detail_locator_workbook(
            "第八部分 基金份额的申购与赎回 / 六、未出现的小标题"
        )
        shared_body = (
            "这段摘要条款与正文整体存在较高文字重合，用于验证系统在Excel指定小标题缺失时不能退回整章冒充命中。"
            "即使整章相似度较高，也应明确提示未找到指定小标题，并把规则定位路径保留下来供人工复核。"
            "这里额外补充流程、披露与登记说明，用于确保摘要章节长度稳定超过过滤阈值，避免测试因长度过滤提前失败。"
        )
        contract_text = "\n".join([
            "第一部分 前言",
            "这是前言部分的说明文字，用于充实测试文本长度，避免被章节过滤逻辑误判，同时确保摘要定位不会误回退到前言章节。",
            "第八部分 基金份额的申购与赎回",
            "一、申购和赎回场所",
            "投资人可按照约定办理申购和赎回，相关安排由基金管理人另行披露。",
            "四、申购和赎回的程序",
            shared_body,
            "五、申购和赎回的数额限制",
            "基金管理人可以根据业务安排调整申购赎回数量限制，并按规定公告。",
            "第二十六部分 基金合同内容摘要",
            "一、摘要条款",
            shared_body,
        ])

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result = payload["results"][0]
        self.assertEqual(result["rule_source"], "summary_rules")
        self.assertEqual(result["rule_level"], "detail")
        self.assertFalse(result["locator_matched"])
        self.assertEqual(result["severity"], "warning")
        self.assertEqual(result["matched_locator"], "")
        self.assertEqual(result["matched_section_text"], "")
        self.assertEqual(result["rule_locator"], "第八部分 基金份额的申购与赎回 / 六、未出现的小标题")

    def test_summary_check_uses_chapter_rule_then_normalized_subheading_locator(self):
        client = app.app.test_client()
        workbook = self._build_summary_chapter_locator_workbook(
            "第十八部分 基金的收益与分配",
            "三、基金收益分配原则、执行方式",
        )
        subsection_body = (
            "基金收益分配遵循约定比例、现金分红安排与登记确认规则，并保留必要的执行性表述。"
            "本段用于验证章节级规则命中后，系统仍会继续在章内找到对应小标题，而不是把整章文本直接当作命中结果。"
            "这里额外补充收益分配登记、披露和执行安排说明，用于确保摘要章节长度稳定超过过滤阈值。"
        )
        contract_text = "\n".join([
            "第一部分 前言",
            "这是前言部分的说明文字，用于充实测试文本长度，避免被章节过滤逻辑误判，同时确保摘要定位不会误回退到前言章节。",
            "第十八部分 基金的收益与分配",
            "（一）收益分配原则",
            "基金收益分配原则由基金管理人根据合同约定执行，并保持披露口径一致。",
            "（三）基金收益分配原则及执行方式",
            subsection_body,
            "（四）收益分配方案",
            "基金管理人应当按照规定披露收益分配方案并组织实施。",
            "第二十六部分 基金合同内容摘要",
            "三、基金收益分配原则、执行方式",
            subsection_body,
        ])

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result = payload["results"][0]
        self.assertEqual(result["rule_level"], "chapter")
        self.assertTrue(result["locator_matched"])
        self.assertEqual(result["matched_section"], "第十八部分 基金的收益与分配")
        self.assertEqual(result["matched_target_heading"], "（三）基金收益分配原则及执行方式")
        self.assertEqual(result["matched_locator"], "第十八部分 基金的收益与分配 / （三）基金收益分配原则及执行方式")
        self.assertTrue(result["matched_section_text"].startswith("（三）基金收益分配原则及执行方式"))

    def test_summary_check_hides_low_confidence_similarity_fallback_matches(self):
        client = app.app.test_client()
        original_rules_path = app._review_store.get("rules_xlsx_path")
        app._review_store["rules_xlsx_path"] = "__missing_review_rules__.xlsx"
        contract_text = "\n".join([
            "第一部分 前言",
            "这里是基金合同的前言部分，用于说明订立目的和法律依据。",
            "第十部分 基金份额持有人大会",
            "本部分详细说明基金份额持有人大会的召集、议事和表决程序。",
            "第二十六部分 基金合同内容摘要",
            "一、完全无关的摘要标题",
            "本摘要内容主要讨论基金费用支付安排，与持有人大会程序并无对应关系。",
        ])

        try:
            response = client.post("/api/summary_check", json={"text": contract_text})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_rules_path is None:
                app._review_store.pop("rules_xlsx_path", None)
            else:
                app._review_store["rules_xlsx_path"] = original_rules_path

        result = payload["results"][0]
        self.assertFalse(result["locator_matched"])
        self.assertIsNone(result["matched_section"])
        self.assertEqual(result["matched_locator"], "")
        self.assertEqual(result["matched_section_text"], "")

    def test_summary_check_groups_nested_summary_details_and_prefers_title_match_when_rule_ordinal_differs(self):
        client = app.app.test_client()
        workbook = self._build_summary_rules_workbook(
            chapter_rows=[
                {
                    "contract_pos": "第七部分基金合同当事人及权利义务",
                    "summary_pos": "摘要一、基金合同当事人的权利、义务",
                    "status": "已收录",
                    "method": "近全文摘录",
                    "similarity": "0.95",
                    "detail": "摘要一对应基金合同当事人的权利义务章节。",
                }
            ],
            detail_rows=[
                {
                    "summary_pos": "摘要一 / （一）基金管理人的权利与义务",
                    "contract_pos": "第七部分—一、基金管理人 / （二）基金管理人的权利与义务",
                    "status": "已收录",
                    "method": "近全文摘录",
                    "similarity": "0.99",
                    "detail": "基金管理人的权利义务应精确定位到对应小标题。",
                },
                {
                    "summary_pos": "摘要一 / （二）基金托管人的权利与义务",
                    "contract_pos": "第七部分—二、基金托管人 / （二）基金托管人的权利与义务",
                    "status": "已收录",
                    "method": "近全文摘录",
                    "similarity": "0.99",
                    "detail": "基金托管人的权利义务应精确定位到对应小标题。",
                },
            ],
        )
        manager_body = (
            "基金管理人依法享有投资管理、信息披露安排和相关执行性权利，并承担诚实信用、谨慎勤勉义务。"
            "本段用于验证系统能够在摘要主项下继续锁定到基金管理人的权利义务小标题。"
        )
        trustee_body = (
            "基金托管人依法享有监督、复核和账户管理等权利，并承担安全保管基金财产和监督运作义务。"
            "本段用于验证系统能够在摘要主项下继续锁定到基金托管人的权利义务小标题。"
        )
        contract_text = "\n".join([
            "第一部分 前言",
            "前言部分用于充实测试文本长度，避免摘要定位退回到无关章节。",
            "第九部分 基金合同当事人及权利义务",
            "一、基金管理人",
            "（二）基金管理人的权利与义务",
            manager_body,
            "二、基金托管人",
            "（二）基金托管人的权利与义务",
            trustee_body,
            "第二十六部分 基金合同内容摘要",
            "一、基金合同当事人的权利、义务",
            "（一）基金管理人的权利与义务",
            manager_body,
            "（二）基金托管人的权利与义务",
            trustee_body,
        ])

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        self.assertEqual(payload["total_subsections"], 1)
        result = payload["results"][0]
        self.assertEqual(result["detail_count"], 2)
        self.assertEqual(result["problem_count"], 0)
        self.assertEqual(result["pass_count"], 2)
        self.assertFalse(result["is_problem"])
        self.assertEqual(len(result["detail_results"]), 2)

        manager_detail = next(item for item in result["detail_results"] if "基金管理人" in item["summary_heading"])
        trustee_detail = next(item for item in result["detail_results"] if "基金托管人" in item["summary_heading"])

        self.assertEqual(manager_detail["summary_path"], "摘要一 / （一）基金管理人的权利与义务")
        self.assertEqual(manager_detail["matched_section"], "第九部分 基金合同当事人及权利义务")
        self.assertEqual(manager_detail["matched_target_heading"], "（二）基金管理人的权利与义务")
        self.assertEqual(manager_detail["matched_locator"], "第九部分 基金合同当事人及权利义务 / 一、基金管理人 / （二）基金管理人的权利与义务")
        self.assertTrue(manager_detail["locator_matched"])
        self.assertEqual(trustee_detail["matched_target_heading"], "（二）基金托管人的权利与义务")

    def test_summary_check_limits_selected_summary_detail_scope_for_configured_groups(self):
        client = app.app.test_client()
        workbook = self._build_summary_rules_workbook(
            detail_rows=[
                {
                    "summary_pos": "摘要一 / （一）基金管理人的权利与义务",
                    "contract_pos": "第七部分 基金合同当事人的权利、义务 / 一、基金管理人 / （二）基金管理人的权利与义务",
                },
                {
                    "summary_pos": "摘要一 / （二）基金托管人的权利与义务",
                    "contract_pos": "第七部分 基金合同当事人的权利、义务 / 二、基金托管人 / （二）基金托管人的权利与义务",
                },
                {
                    "summary_pos": "摘要一 / （三）基金份额持有人的权利与义务",
                    "contract_pos": "第七部分 基金合同当事人的权利、义务 / 三、基金份额持有人 / （二）基金份额持有人的权利与义务",
                },
                {
                    "summary_pos": "摘要四 / 一、基金费用的种类",
                    "contract_pos": "第十五部分 基金的费用与税收 / 一、基金费用的种类",
                },
                {
                    "summary_pos": "摘要四 / 二、基金费用计提方法、计提标准和支付方式",
                    "contract_pos": "第十五部分 基金的费用与税收 / 二、基金费用计提方法、计提标准和支付方式",
                },
                {
                    "summary_pos": "摘要四 / 三、与基金运作有关的税收",
                    "contract_pos": "第十五部分 基金的费用与税收 / 三、与基金运作有关的税收",
                },
                {
                    "summary_pos": "摘要五 / （一）投资范围",
                    "contract_pos": "第十一部分 基金的投资 / 一、投资范围",
                },
                {
                    "summary_pos": "摘要五 / （二）投资限制",
                    "contract_pos": "第十一部分 基金的投资 / 二、投资限制",
                },
                {
                    "summary_pos": "摘要五 / （三）投资策略",
                    "contract_pos": "第十一部分 基金的投资 / 三、投资策略",
                },
                {
                    "summary_pos": "摘要六 / （五）估值程序",
                    "contract_pos": "第十三部分 基金资产估值 / 五、估值程序",
                },
                {
                    "summary_pos": "摘要六 / （六）暂停估值的情形",
                    "contract_pos": "第十三部分 基金资产估值 / 六、暂停估值的情形",
                },
            ]
        )
        manager_body = (
            "基金管理人依法享有投资管理、信息披露安排和相关执行性权利，并承担诚实信用、谨慎勤勉义务。"
            "本段用于验证摘要复核会保留管理人权利义务这一项的正文定位与文本比对。"
        )
        trustee_body = (
            "基金托管人依法享有监督、复核和账户管理等权利，并承担安全保管基金财产和监督运作义务。"
            "本段用于验证摘要复核会保留托管人权利义务这一项的正文定位与文本比对。"
        )
        holder_body = (
            "基金份额持有人享有收益分配、表决和查询等权利，并承担遵守合同约定和法律法规的义务。"
            "本段故意保留在摘要和规则中，用于验证本次复核不会再把持有人权利义务纳入摘要一致性明细。"
        )
        fee_types_body = (
            "基金费用的种类包括管理费、托管费、信息披露费以及基金合同约定的其他费用。"
            "本段用于验证基金费用章节只保留费用种类这一项参与摘要一致性复核。"
        )
        fee_calc_body = (
            "基金费用计提方法、计提标准和支付方式应按合同约定执行，并履行必要的信息披露义务。"
            "本段用于验证基金费用章节只保留计提方法与支付方式这一项参与摘要一致性复核。"
        )
        fee_tax_body = (
            "与基金运作有关的税收按照国家税收法律法规执行。"
            "本段用于验证税收说明即便在规则表中存在，也不应再出现在摘要一致性复核明细中。"
        )
        invest_scope_body = (
            "投资范围包括标的指数成份证券、备选成份证券以及法律法规允许基金投资的其他金融工具。"
            "本段用于验证投资章节只保留投资范围这一项参与摘要一致性复核。"
        )
        invest_limit_body = (
            "投资限制包括比例限制、集中度限制以及法律法规和监管规则要求的其他限制。"
            "本段用于验证投资章节只保留投资限制这一项参与摘要一致性复核。"
        )
        invest_strategy_body = (
            "投资策略包括完全复制、抽样复制和组合优化等安排。"
            "本段用于验证投资策略不再纳入基金合同摘要一致性的明细对比。"
        )
        valuation_process_body = (
            "估值程序包括估值对象确认、估值方法选取、结果复核以及对外披露安排。"
            "本段用于验证资产净值章节只保留估值程序这一项参与摘要一致性复核。"
        )
        suspend_value_body = (
            "暂停估值的情形包括证券市场异常波动、不可抗力以及监管规则规定的其他情形。"
            "本段用于验证暂停估值情形不再纳入基金合同摘要一致性复核明细。"
        )
        contract_text = "\n".join([
            "第一部分 前言",
            "前言部分用于充实测试文本长度，避免摘要定位退回到无关章节，并确保复核逻辑使用真实正文来源。",
            "第七部分 基金合同当事人的权利、义务",
            "一、基金管理人",
            "（二）基金管理人的权利与义务",
            manager_body,
            "二、基金托管人",
            "（二）基金托管人的权利与义务",
            trustee_body,
            "三、基金份额持有人",
            "（二）基金份额持有人的权利与义务",
            holder_body,
            "第十五部分 基金的费用与税收",
            "一、基金费用的种类",
            fee_types_body,
            "二、基金费用计提方法、计提标准和支付方式",
            fee_calc_body,
            "三、与基金运作有关的税收",
            fee_tax_body,
            "第十一部分 基金的投资",
            "一、投资范围",
            invest_scope_body,
            "二、投资限制",
            invest_limit_body,
            "三、投资策略",
            invest_strategy_body,
            "第十三部分 基金资产估值",
            "五、估值程序",
            valuation_process_body,
            "六、暂停估值的情形",
            suspend_value_body,
            "第二十六部分 基金合同内容摘要",
            "一、基金合同当事人的权利、义务",
            "（一）基金管理人的权利与义务",
            manager_body,
            "（二）基金托管人的权利与义务",
            trustee_body,
            "（三）基金份额持有人的权利与义务",
            holder_body,
            "四、基金费用与税收",
            "一、基金费用的种类",
            fee_types_body,
            "二、基金费用计提方法、计提标准和支付方式",
            fee_calc_body,
            "三、与基金运作有关的税收",
            fee_tax_body,
            "五、基金财产的投资范围和投资限制",
            "（一）投资范围",
            invest_scope_body,
            "（二）投资限制",
            invest_limit_body,
            "（三）投资策略",
            invest_strategy_body,
            "六、基金资产净值的计算方法和公告方式",
            "（一）基金资产净值",
            "本段用于验证资产净值小项不会再参与摘要一致性复核。",
            "（五）估值程序",
            valuation_process_body,
            "（六）暂停估值的情形",
            suspend_value_body,
        ])

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result_map = {item["summary_heading"]: item for item in payload["results"]}
        self.assertEqual(4, payload["total_subsections"])

        party_result = result_map["一、基金合同当事人的权利、义务"]
        self.assertEqual(2, party_result["detail_count"])
        self.assertEqual(
            ["（一）基金管理人的权利与义务", "（二）基金托管人的权利与义务"],
            [item["summary_heading"] for item in party_result["detail_results"]],
        )

        fee_result = result_map["四、基金费用与税收"]
        self.assertEqual(2, fee_result["detail_count"])
        self.assertEqual(
            ["一、基金费用的种类", "二、基金费用计提方法、计提标准和支付方式"],
            [item["summary_heading"] for item in fee_result["detail_results"]],
        )

        invest_result = result_map["五、基金财产的投资范围和投资限制"]
        self.assertEqual(2, invest_result["detail_count"])
        self.assertEqual(
            ["（一）投资范围", "（二）投资限制"],
            [item["summary_heading"] for item in invest_result["detail_results"]],
        )

        valuation_result = result_map["六、基金资产净值的计算方法和公告方式"]
        self.assertEqual(1, valuation_result["detail_count"])
        self.assertEqual(
            ["（五）估值程序"],
            [item["summary_heading"] for item in valuation_result["detail_results"]],
        )

    def test_build_ai_summary_review_items_respects_selected_summary_detail_scope(self):
        manager_body = (
            "基金管理人依法享有投资管理、信息披露安排和相关执行性权利，并承担诚实信用、谨慎勤勉义务。"
            "本段用于验证文件复核中的摘要一致性明细只保留管理人和托管人两项。"
        )
        trustee_body = (
            "基金托管人依法享有监督、复核和账户管理等权利，并承担安全保管基金财产和监督运作义务。"
            "本段用于验证文件复核中的摘要一致性明细只保留管理人和托管人两项。"
        )
        holder_body = (
            "基金份额持有人享有收益分配、表决和查询等权利，并承担遵守合同约定和法律法规的义务。"
            "本段用于验证即便规则表仍保留持有人项，文件复核也不会继续把它生成为摘要一致性明细。"
        )
        items = app._build_ai_summary_review_items(
            [
                {
                    "heading": "第七部分 基金合同当事人的权利、义务",
                    "content": "\n".join([
                        "一、基金管理人",
                        "（二）基金管理人的权利与义务",
                        manager_body,
                        "二、基金托管人",
                        "（二）基金托管人的权利与义务",
                        trustee_body,
                        "三、基金份额持有人",
                        "（二）基金份额持有人的权利与义务",
                        holder_body,
                    ]),
                },
                {
                    "heading": "第二十六部分 基金合同内容摘要",
                    "content": "\n".join([
                        "一、基金合同当事人的权利、义务",
                        "（一）基金管理人的权利与义务",
                        manager_body,
                        "（二）基金托管人的权利与义务",
                        trustee_body,
                        "（三）基金份额持有人的权利与义务",
                        holder_body,
                    ]),
                },
            ],
            {
                "summary_rules": [
                    {
                        "summary_pos": "摘要一 / （一）基金管理人的权利与义务",
                        "contract_pos": "第七部分 基金合同当事人的权利、义务 / 一、基金管理人 / （二）基金管理人的权利与义务",
                        "method": "近全文摘录",
                        "status": "已收录",
                        "rule_level": "detail",
                    },
                    {
                        "summary_pos": "摘要一 / （二）基金托管人的权利与义务",
                        "contract_pos": "第七部分 基金合同当事人的权利、义务 / 二、基金托管人 / （二）基金托管人的权利与义务",
                        "method": "近全文摘录",
                        "status": "已收录",
                        "rule_level": "detail",
                    },
                    {
                        "summary_pos": "摘要一 / （三）基金份额持有人的权利与义务",
                        "contract_pos": "第七部分 基金合同当事人的权利、义务 / 三、基金份额持有人 / （二）基金份额持有人的权利与义务",
                        "method": "近全文摘录",
                        "status": "已收录",
                        "rule_level": "detail",
                    },
                ]
            },
        )

        self.assertEqual(2, len(items))
        self.assertEqual(
            ["（一）基金管理人的权利与义务", "（二）基金托管人的权利与义务"],
            [item["summary_heading"] for item in items],
        )
        self.assertTrue(all("基金份额持有人" not in item["label"] for item in items))

    def test_summary_check_maps_section_six_group_content_to_valuation_procedure(self):
        client = app.app.test_client()
        workbook = self._build_summary_rules_workbook(
            chapter_rows=[
                {
                    "contract_pos": "第十六部分 基金资产估值",
                    "summary_pos": "六、基金资产净值的计算方法和公告方式",
                    "status": "部分收录",
                    "method": "选择性摘录",
                }
            ]
        )
        valuation_body = "\n".join([
            "1、基金份额净值是按照每个估值日闭市后，基金资产净值除以当日基金份额的余额数量计算，精确到0.0001元。",
            "基金管理人应每个估值日计算基金资产净值及基金份额净值，并按规定披露。",
            "2、基金管理人应每个估值日对基金资产估值，但基金管理人根据法律法规或本基金合同的规定暂停估值时除外。",
            "基金管理人每个估值日对基金资产估值后，将基金份额净值结果发送基金托管人，经基金托管人复核无误后，由基金管理人按规定对外公布。",
        ])
        contract_text = "\n".join([
            "第一部分 前言",
            "前言用于补足测试样本文本长度。",
            "第十六部分 基金资产估值",
            "五、估值程序",
            valuation_body,
            "第二十六部分 基金合同内容摘要",
            "六、基金资产净值的计算方法和公告方式",
            valuation_body,
        ])

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result = payload["results"][0]
        self.assertEqual(1, result["detail_count"])
        detail = result["detail_results"][0]
        self.assertEqual("五、估值程序", detail["summary_heading"])
        self.assertEqual("第十六部分 基金资产估值 / 五、估值程序", detail["matched_locator"])
        self.assertEqual(100.0, detail["similarity"])
        self.assertFalse(detail["is_problem"])

    def test_summary_check_detail_results_include_excerpt_fields_for_problem_items(self):
        client = app.app.test_client()
        workbook = self._build_summary_rules_workbook(
            detail_rows=[
                {
                    "summary_pos": "摘要三 / （一）基金收益分配原则",
                    "contract_pos": "第十八部分—三、基金收益分配原则",
                    "status": "已收录",
                    "method": "选择性摘录",
                    "similarity": "0.95",
                    "detail": "收益分配原则应按规则定位到正文对应小标题。",
                }
            ]
        )
        contract_body = (
            "基金收益分配遵循约定比例、登记确认和现金分红要求，管理人应及时披露收益分配安排。"
            "本段用于验证差异片段会返回正文附近原文。"
        )
        summary_body = (
            "基金收益分配遵循约定比例、登记确认和现金分红要求，管理人应立即披露收益分配安排。"
            "本段用于验证差异片段会返回摘要附近原文。"
        )
        contract_text = "\n".join([
            "第一部分 前言",
            "这里是前言，用于避免测试文本过短。",
            "第十八部分 基金的收益与分配",
            "三、基金收益分配原则",
            contract_body,
            "第二十六部分 基金合同内容摘要",
            "三、基金收益分配原则、执行方式",
            "（一）基金收益分配原则",
            summary_body,
        ])

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result = payload["results"][0]["detail_results"][0]
        self.assertTrue(result["is_problem"])
        self.assertIn("及时披露收益分配安排", result["contract_anchor_text"])
        self.assertIn("立即披露收益分配安排", result["summary_anchor_text"])
        self.assertIn("收益分配安排", result["contract_context_excerpt"])
        self.assertIn("收益分配安排", result["summary_context_excerpt"])

    def test_summary_check_ignores_blank_lines_in_diff(self):
        client = app.app.test_client()
        workbook = self._build_summary_rules_workbook(
            detail_rows=[
                {
                    "summary_pos": "摘要四 / 一、基金费用的种类",
                    "contract_pos": "第十五部分 基金的费用与税收 / 一、基金费用的种类",
                    "status": "已收录",
                    "method": "全文摘录",
                }
            ]
        )
        body = (
            "基金费用的种类包括管理费、托管费、信息披露费以及基金合同约定的其他费用。"
            "本段用于验证摘要一致性在仅存在空行差异时应视为一致。"
        )
        contract_text = "\n".join([
            "第一部分 前言",
            "前言用于补足测试样本文本长度。",
            "第十五部分 基金的费用与税收",
            "一、基金费用的种类",
            body,
            "第二十六部分 基金合同内容摘要",
            "四、基金费用与税收",
            "一、基金费用的种类",
            "",
            body,
            "",
        ])

        with self._with_review_rules(workbook):
            response = client.post("/api/summary_check", json={"text": contract_text})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        result = payload["results"][0]["detail_results"][0]
        self.assertEqual(0, result["changed_lines"])
        self.assertEqual(100.0, result["similarity"])
        self.assertFalse(result["is_problem"])

    def test_review_cross_check_ignores_blank_lines(self):
        payload = self._run_cross_check_with_review_rules(
            [
                {
                    "heading": "第八部分 基金份额的申购与赎回",
                    "content": "\n".join([
                        "四、申购和赎回的程序",
                        "投资人应按规定提交申请。",
                        "",
                        "登记机构确认后办理。",
                    ]),
                }
            ],
            [
                {
                    "heading": "基金份额的申购赎回",
                    "content": "\n".join([
                        "四、申购和赎回的程序",
                        "投资人应按规定提交申请。",
                        "登记机构确认后办理。",
                    ]),
                }
            ],
            {
                "chapter_rules": [
                    {
                        "contract": "第八部分 基金份额的申购与赎回 / 四、申购和赎回的程序",
                        "prospectus": "基金份额的申购赎回 / 四、申购和赎回的程序",
                        "relation": "直接对应",
                        "consistency": "完全一致",
                        "detail": "",
                        "sheet_name": "测试",
                        "rule_level": "chapter",
                    }
                ],
                "detail_rules": [],
                "key_diff_rules": [],
                "chapter_level": [],
                "key_diffs": [],
            },
        )

        result = payload["results"][0]
        self.assertEqual("pass", result["status"])
        self.assertFalse(result["is_problem"])
        self.assertEqual(100.0, result["similarity"])

    def test_summary_cross_check_ignores_blank_lines(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {
                    "heading": "第二十六部分 基金合同内容摘要",
                    "content": "\n".join([
                        "四、基金费用与税收",
                        "一、基金费用的种类",
                        "基金费用包括管理费。",
                        "",
                        "基金费用包括托管费。",
                    ]),
                }
            ],
            "prospectus_sections": [
                {
                    "heading": "二十、基金合同的内容摘要",
                    "content": "\n".join([
                        "四、基金费用与税收",
                        "一、基金费用的种类",
                        "基金费用包括管理费。",
                        "基金费用包括托管费。",
                    ]),
                }
            ],
        }

        try:
            response = client.post("/api/review/summary_cross_check", json={})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        self.assertEqual(0, payload["changed_lines"])
        self.assertEqual(100.0, payload["similarity"])
        self.assertFalse(payload["is_problem"])

    def test_review_cross_check_returns_excerpt_fields_for_problem_rows(self):
        payload = self._run_cross_check_with_review_rules(
            [
                {
                    "heading": "第八部分 基金份额的申购与赎回",
                    "content": "四、申购和赎回的程序\n基金合同仅规定确认、清算交收与登记的原则框架，投资人T日申请后按约定办理。",
                }
            ],
            [
                {
                    "heading": "基金份额的申购赎回",
                    "content": "四、申购和赎回的程序\n招募说明书补充T日、T+1和T+2的清算流程，并明确投资人当日可卖出安排。",
                }
            ],
            {
                "chapter_rules": [
                    {
                        "source": "chapter_rules",
                        "contract": "第八部分 基金份额的申购与赎回",
                        "prospectus": "基金份额的申购赎回",
                        "relation": "直接对应",
                        "consistency": "完全一致",
                        "detail": "对应条款应当保持一致。",
                        "sheet_name": "测试",
                    }
                ],
                "detail_rules": [],
                "key_diff_rules": [],
                "chapter_level": [],
                "key_diffs": [],
            },
        )

        result = payload["results"][0]
        self.assertTrue(result["is_problem"])
        self.assertIn("确认、清算交收与登记", result["contract_anchor_text"])
        self.assertIn("T日、T+1和T+2", result["prospectus_anchor_text"])
        self.assertIn("投资人", result["contract_context_excerpt"])
        self.assertIn("投资人", result["prospectus_context_excerpt"])

    def test_summary_cross_check_returns_excerpt_fields_for_problem_rows(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {
                    "heading": "第二十六部分 基金合同内容摘要",
                    "content": (
                        "一、基金合同的效力\n"
                        "基金合同自基金份额持有人大会审议通过后生效，并对基金合同当事人具有法律约束力。\n"
                        "基金合同的效力条款用于验证合同摘要差异片段。"
                    ),
                }
            ],
            "prospectus_sections": [
                {
                    "heading": "二十、基金合同的内容摘要",
                    "content": (
                        "一、基金合同的效力\n"
                        "基金合同自基金份额持有人大会审议通过后生效，并对基金合同当事人具有持续法律约束力。\n"
                        "基金合同的效力条款用于验证招募摘要差异片段。"
                    ),
                }
            ],
        }

        try:
            response = client.post("/api/review/summary_cross_check", json={})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        self.assertTrue(payload["is_problem"])
        self.assertIn("法律约束力", payload["contract_anchor_text"])
        self.assertIn("持续法律约束力", payload["prospectus_anchor_text"])
        self.assertIn("效力条款", payload["contract_context_excerpt"])
        self.assertIn("效力条款", payload["prospectus_context_excerpt"])

    def test_review_cross_check_returns_rule_annotations_from_excel(self):
        client = app.app.test_client()
        workbook = self._build_review_rules_workbook(linked=False)
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {"heading": "\u7b2c\u516b\u90e8\u5206 \u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u4e0e\u8d4e\u56de", "content": "\u5408\u540c\u53ea\u89c4\u5b9a\u7533\u8d2d\u8d4e\u56de\u7684\u5236\u5ea6\u6846\u67b6\u3002"},
            ],
            "prospectus_sections": [
                {"heading": "\u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u8d4e\u56de", "content": "\u62db\u52df\u8bf4\u660e\u4e66\u989d\u5916\u8865\u5145\u4e86\u7a0b\u5e8f\u3001\u516c\u5f0f\u4e0e\u6e05\u5355\u683c\u5f0f\u3002"},
            ],
        }
        try:
            with self._with_review_rules(workbook):
                response = client.post("/api/review/cross_check", json={"fund_type": "ETF"})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        result = payload["results"][0]
        self.assertEqual(result["rule_source"], "detail_rules")
        self.assertEqual(result["sheet_name"], "招募-合同(详细)")
        self.assertEqual(result["expected_relation"], "直接对应（招募细化）")
        self.assertIn("招募补充基金份额净值计算公式和代理券商佣金限额。", result["expected_diff"])

    def test_locate_review_rule_target_ignores_page_markers_in_locator(self):
        sections = [
            {"heading": "绪言", "content": "本招募说明书依据基金合同编写。"},
        ]

        result = app._locate_review_rule_target(sections, "绪言（P5）", "绪言（P5）")

        self.assertTrue(result["matched"])
        self.assertEqual(result["section_heading"], "绪言")
        self.assertEqual(result["target_heading"], "绪言")

    def test_strip_contract_signing_page_text_trims_draft_and_truncated_signing_markers(self):
        cases = [
            (
                "\u6458\u8981\u6b63\u6587\u7b2c\u4e00\u6bb5\u3002\n\u672c\u9875\u4e3a\u300a\u5357\u65b9\u6d4b\u8bd5\u57fa\u91d1\u5408\u540c\uff08\u8349\u6848\uff09\u300b\u7b7e\u7f72\u9875\n\u57fa\u91d1\u7ba1\u7406\u4eba\uff1a\u6d4b\u8bd5\u516c\u53f8\uff08\u76d6\u7ae0\uff09",
                "\u6458\u8981\u6b63\u6587\u7b2c\u4e00\u6bb5\u3002",
            ),
            (
                "\u6458\u8981\u6b63\u6587\u7b2c\u4e8c\u6bb5\u3002\n\u672c\u9875\u4e3a\u300a\u5357\u65b9\u6d4b\u8bd5\u57fa\u91d1\u5408\u540c\uff08\u8349\u6848\uff09\u300b\n\u57fa\u91d1\u6258\u7ba1\u4eba\uff1a\u6d4b\u8bd5\u94f6\u884c\uff08\u7b7e\u5b57\u6216\u76d6\u7ae0\uff09",
                "\u6458\u8981\u6b63\u6587\u7b2c\u4e8c\u6bb5\u3002",
            ),
            (
                "\u6458\u8981\u6b63\u6587\u7b2c\u4e09\u6bb5\u3002\n\u7b7e\u7f72\u9875\uff0c\u65e0\u6b63\u6587\n\u57fa\u91d1\u7ba1\u7406\u4eba\uff1a\u6d4b\u8bd5\u516c\u53f8\uff08\u76d6\u7ae0\uff09",
                "\u6458\u8981\u6b63\u6587\u7b2c\u4e09\u6bb5\u3002",
            ),
        ]

        for text, expected in cases:
            with self.subTest(text=text):
                self.assertEqual(expected, app._strip_contract_signing_page_text(text))

    def test_split_prospectus_sections_skips_toc_and_accepts_heading2_aliases(self):
        fake_doc = SimpleNamespace(
            paragraphs=[
                SimpleNamespace(text="\u7eea\u8a00", style=SimpleNamespace(name="Heading 2")),
                SimpleNamespace(text="\u57fa\u91d1\u7684\u52df\u96c6", style=SimpleNamespace(name="\u6807\u9898 2")),
            ]
        )
        text = "\n".join([
            "\u76ee\u5f55",
            "\u7eea\u8a00\t1",
            "\u57fa\u91d1\u7684\u52df\u96c6\t6",
            "\u7eea\u8a00",
            "\u672c\u7ae0\u4e3a\u6b63\u6587\u7eea\u8a00\u3002",
            "\u57fa\u91d1\u7684\u52df\u96c6",
            "\u672c\u7ae0\u4e3a\u57fa\u91d1\u52df\u96c6\u6b63\u6587\u3002",
        ])

        sections = app._split_prospectus_sections(text, fake_doc)

        self.assertEqual(
            ["\u7eea\u8a00", "\u57fa\u91d1\u7684\u52df\u96c6"],
            [section["heading"] for section in sections],
        )
        self.assertEqual("\u672c\u7ae0\u4e3a\u6b63\u6587\u7eea\u8a00\u3002", sections[0]["content"])
        self.assertEqual("\u672c\u7ae0\u4e3a\u57fa\u91d1\u52df\u96c6\u6b63\u6587\u3002", sections[1]["content"])

    def test_split_prospectus_sections_detects_plain_body_headings_without_style(self):
        text = "\n".join([
            "\u76ee\u5f55",
            "\u7eea\u8a00\t1",
            "\u91ca\u4e49\t2",
            "\u7eea\u8a00",
            "\u8fd9\u91cc\u662f\u7eea\u8a00\u6b63\u6587\u3002",
            "\u91ca\u4e49",
            "\u8fd9\u91cc\u662f\u91ca\u4e49\u6b63\u6587\u3002",
        ])

        sections = app._split_prospectus_sections(text)

        self.assertEqual(
            ["\u7eea\u8a00", "\u91ca\u4e49"],
            [section["heading"] for section in sections],
        )

    def test_split_prospectus_sections_uses_toc_variants_to_match_plain_body_headings(self):
        fake_doc = SimpleNamespace(
            paragraphs=[
                SimpleNamespace(text="\u7eea\u8a00", style=SimpleNamespace(name="Heading 2")),
                SimpleNamespace(text="\u91ca\u4e49", style=SimpleNamespace(name="Heading 2")),
            ]
        )
        text = "\n".join([
            "\u76ee\u5f55",
            "\u7eea\u8a00\t1",
            "\u91ca\u4e49\t4",
            "\u4e8c\u5341\u3001\u57fa\u91d1\u5408\u540c\u7684\u5185\u5bb9\u6458\u8981\t82",
            "\u4e8c\u5341\u4e00\u3001\u57fa\u91d1\u6258\u7ba1\u534f\u8bae\u7684\u5185\u5bb9\u6458\u8981\t84",
            "\u7eea\u8a00",
            "\u8fd9\u91cc\u662f\u7eea\u8a00\u6b63\u6587\u3002",
            "\u91ca\u4e49",
            "\u8fd9\u91cc\u662f\u91ca\u4e49\u6b63\u6587\u3002",
            "\u57fa\u91d1\u5408\u540c\u7684\u5185\u5bb9\u6458\u8981",
            "\u8fd9\u91cc\u662f\u5408\u540c\u6458\u8981\u6b63\u6587\u3002",
            "\u57fa\u91d1\u6258\u7ba1\u534f\u8bae\u7684\u5185\u5bb9\u6458\u8981",
            "\u8fd9\u91cc\u662f\u6258\u7ba1\u534f\u8bae\u6458\u8981\u6b63\u6587\u3002",
        ])

        sections = app._split_prospectus_sections(text, fake_doc)

        self.assertEqual(
            [
                "\u7eea\u8a00",
                "\u91ca\u4e49",
                "\u57fa\u91d1\u5408\u540c\u7684\u5185\u5bb9\u6458\u8981",
                "\u57fa\u91d1\u6258\u7ba1\u534f\u8bae\u7684\u5185\u5bb9\u6458\u8981",
            ],
            [section["heading"] for section in sections],
        )
        self.assertEqual("\u8fd9\u91cc\u662f\u5408\u540c\u6458\u8981\u6b63\u6587\u3002", sections[2]["content"])
        self.assertEqual("\u8fd9\u91cc\u662f\u6258\u7ba1\u534f\u8bae\u6458\u8981\u6b63\u6587\u3002", sections[3]["content"])

    def test_split_prospectus_sections_prefers_heading2_positions_over_fuzzy_body_matches(self):
        fake_doc = SimpleNamespace(
            paragraphs=[
                SimpleNamespace(text="基金的收益与分配", style=SimpleNamespace(name="Heading 2")),
                SimpleNamespace(text="基金合同的内容摘要", style=SimpleNamespace(name="Heading 2")),
                SimpleNamespace(text="基金托管协议的内容摘要", style=SimpleNamespace(name="Heading 2")),
            ]
        )
        text = "\n".join(
            [
                "目录",
                "十四、基金的收益与分配\t61",
                "二十、基金合同的内容摘要\t82",
                "二十一、基金托管协议的内容摘要\t101",
                "基金的收益与分配",
                "这里是收益与分配正文。",
                "基金合同的内容摘要",
                "一、基金合同当事人的权利、义务",
                "（四）基金收益分配中发生的费用",
                "基金收益分配时所发生的银行转账或其他手续费用由投资人自行承担。",
                "四、基金费用与税收",
                "（一）基金费用的种类",
                "五、基金财产的投资范围和投资限制",
                "（二）投资限制",
                "1、组合限制",
                "基金的投资组合应遵循以下限制：",
                "（1）本基金投资于标的指数成份股、备选成份股的资产比例不低于基金资产净值的90%。",
                "九、基金合同的效力",
                "《基金合同》自生效之日起对各方当事人具有同等法律约束力。",
                "基金托管协议的内容摘要",
                "这里是托管协议摘要正文。",
            ]
        )

        sections = app._split_prospectus_sections(text, fake_doc)

        self.assertEqual(
            ["基金的收益与分配", "基金合同的内容摘要", "基金托管协议的内容摘要"],
            [section["heading"] for section in sections],
        )
        self.assertIn("（四）基金收益分配中发生的费用", sections[1]["content"])
        self.assertIn("四、基金费用与税收", sections[1]["content"])
        self.assertIn("基金的投资组合应遵循以下限制：", sections[1]["content"])
        self.assertIn("九、基金合同的效力", sections[1]["content"])
        self.assertNotIn("基金托管协议的内容摘要", sections[1]["content"])

    @unittest.skipUnless(
        AUTOMOTIVE_ETF_CONTRACT_DOCX.exists() and AUTOMOTIVE_ETF_PROSPECTUS_DOCX.exists(),
        "需要本地汽车ETF样本文档",
    )
    def test_split_prospectus_sections_keeps_automotive_etf_summary_as_single_chapter(self):
        contract_doc = Document(str(AUTOMOTIVE_ETF_CONTRACT_DOCX))
        prospectus_doc = Document(str(AUTOMOTIVE_ETF_PROSPECTUS_DOCX))
        contract_text = "\n".join((p.text or "").strip() for p in contract_doc.paragraphs if (p.text or "").strip())
        prospectus_text = "\n".join((p.text or "").strip() for p in prospectus_doc.paragraphs if (p.text or "").strip())

        contract_sections = app._split_contract_sections(contract_text)
        prospectus_sections = app._split_prospectus_sections(prospectus_text, prospectus_doc)
        prospectus_summary = app._find_contract_summary_section(prospectus_sections)

        self.assertEqual(26, len(contract_sections))
        self.assertEqual(25, len(prospectus_sections))
        self.assertIsNotNone(prospectus_summary)
        self.assertEqual("基金合同的内容摘要", prospectus_summary["heading"])
        self.assertIn("四、基金费用与税收", prospectus_summary["content"])
        self.assertIn("五、基金财产的投资范围和投资限制", prospectus_summary["content"])
        self.assertIn("九、基金合同的效力", prospectus_summary["content"])
        self.assertNotIn("基金托管协议的内容摘要", prospectus_summary["content"])

    @unittest.skipUnless(
        AUTOMOTIVE_ETF_CONTRACT_DOCX.exists() and AUTOMOTIVE_ETF_PROSPECTUS_DOCX.exists(),
        "需要本地汽车ETF样本文档",
    )
    def test_summary_cross_check_uses_full_automotive_etf_summary_text(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        contract_doc = Document(str(AUTOMOTIVE_ETF_CONTRACT_DOCX))
        prospectus_doc = Document(str(AUTOMOTIVE_ETF_PROSPECTUS_DOCX))
        contract_text = "\n".join((p.text or "").strip() for p in contract_doc.paragraphs if (p.text or "").strip())
        prospectus_text = "\n".join((p.text or "").strip() for p in prospectus_doc.paragraphs if (p.text or "").strip())

        app._review_store["data"] = {
            "prospectus_text": prospectus_text,
            "contract_sections": app._split_contract_sections(contract_text),
            "prospectus_sections": app._split_prospectus_sections(prospectus_text, prospectus_doc),
        }

        try:
            response = client.post("/api/review/summary_cross_check", json={})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        self.assertIn("四、基金费用与税收", payload["prospectus_summary_text"])
        self.assertIn("九、基金合同的效力", payload["prospectus_summary_text"])
        self.assertGreater(payload["similarity"], 95.0)

    @unittest.skipUnless(
        AUTOMOTIVE_ETF_CONTRACT_DOCX.exists() and AUTOMOTIVE_ETF_PROSPECTUS_DOCX.exists(),
        "需要本地汽车ETF样本文档",
    )
    def test_review_upload_then_summary_check_is_all_green_for_automotive_etf_samples(self):
        client = app.app.test_client()

        with open(AUTOMOTIVE_ETF_CONTRACT_DOCX, "rb") as contract_file, open(AUTOMOTIVE_ETF_PROSPECTUS_DOCX, "rb") as prospectus_file:
            upload = client.post(
                "/api/review/upload",
                data={
                    "contract": (contract_file, AUTOMOTIVE_ETF_CONTRACT_DOCX.name),
                    "prospectus": (prospectus_file, AUTOMOTIVE_ETF_PROSPECTUS_DOCX.name),
                },
                content_type="multipart/form-data",
            )

        self.assertEqual(upload.status_code, 200)
        response = client.post("/api/summary_check", json={})
        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())

        self.assertEqual(9, payload["total_subsections"])
        self.assertTrue(all(not item["is_problem"] for item in payload["results"]))
        self.assertTrue(all((item.get("severity") or "info") == "info" for item in payload["results"]))

    @unittest.skipUnless(
        AUTOMOTIVE_ETF_CONTRACT_DOCX.exists() and AUTOMOTIVE_ETF_PROSPECTUS_DOCX.exists(),
        "需要本地汽车ETF样本文档",
    )
    def test_review_upload_then_all_review_checks_are_green_for_automotive_etf_samples(self):
        client = app.app.test_client()

        with open(AUTOMOTIVE_ETF_CONTRACT_DOCX, "rb") as contract_file, open(AUTOMOTIVE_ETF_PROSPECTUS_DOCX, "rb") as prospectus_file:
            upload = client.post(
                "/api/review/upload",
                data={
                    "contract": (contract_file, AUTOMOTIVE_ETF_CONTRACT_DOCX.name),
                    "prospectus": (prospectus_file, AUTOMOTIVE_ETF_PROSPECTUS_DOCX.name),
                },
                content_type="multipart/form-data",
            )

        self.assertEqual(upload.status_code, 200)

        cross_response = client.post("/api/review/cross_check", json={})
        self.assertEqual(cross_response.status_code, 200)
        cross_payload = self._consume_response(cross_response, parser=lambda resp: resp.get_json())
        self.assertGreater(cross_payload["total_rules"], 0)
        self.assertTrue(all(not item.get("is_problem") for item in cross_payload["results"]))

        summary_cross_response = client.post("/api/review/summary_cross_check", json={})
        self.assertEqual(summary_cross_response.status_code, 200)
        summary_cross_payload = self._consume_response(summary_cross_response, parser=lambda resp: resp.get_json())
        self.assertFalse(summary_cross_payload["is_problem"])

        summary_response = client.post("/api/summary_check", json={})
        self.assertEqual(summary_response.status_code, 200)
        summary_payload = self._consume_response(summary_response, parser=lambda resp: resp.get_json())
        self.assertEqual(9, summary_payload["total_subsections"])
        self.assertTrue(all(not item["is_problem"] for item in summary_payload["results"]))

    def test_locate_review_rule_target_reports_normalized_match_method_for_alias_heading(self):
        sections = [
            {"heading": "\u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u8d4e\u56de", "content": "\u62db\u52df\u7ae0\u8282\u5185\u5bb9\u3002"},
        ]

        result = app._locate_review_rule_target(
            sections,
            "\u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u4e0e\u8d4e\u56de",
            "\u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u4e0e\u8d4e\u56de",
        )

        self.assertTrue(result["matched"])
        self.assertEqual("normalized", result["match_method"])
        self.assertEqual("\u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u8d4e\u56de", result["section_heading"])

    def test_review_cross_check_reports_subheading_missing_instead_of_missing_section(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        original_load_review_rules = app._load_review_rules
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {
                    "heading": "\u7b2c\u56db\u90e8\u5206 \u57fa\u91d1\u4efd\u989d\u7684\u53d1\u552e",
                    "content": "\u5408\u540c\u7ae0\u8282\u6b63\u6587\u3002",
                }
            ],
            "prospectus_sections": [
                {
                    "heading": "\u57fa\u91d1\u7684\u52df\u96c6",
                    "content": "\u4e00\u3001\u52df\u96c6\u5b89\u6392\n\u4ec5\u6709\u4e00\u7ea7\u5c0f\u8282\u3002",
                }
            ],
        }

        try:
            app._load_review_rules = lambda _fund_type="ETF": {
                "chapter_rules": [
                    {
                        "source": "chapter_rules",
                        "contract": "\u7b2c\u56db\u90e8\u5206 \u57fa\u91d1\u4efd\u989d\u7684\u53d1\u552e",
                        "prospectus": "\u57fa\u91d1\u7684\u52df\u96c6 / \u4e8c\u3001\u53d1\u552e\u65b9\u5f0f",
                        "relation": "\u5b8c\u5168\u4e00\u81f4",
                        "consistency": "\u5b8c\u5168\u4e00\u81f4",
                        "detail": "",
                        "sheet_name": "\u6d4b\u8bd5",
                    }
                ],
                "detail_rules": [],
                "key_diff_rules": [],
                "chapter_level": [],
                "key_diffs": [],
            }
            response = client.post("/api/review/cross_check", json={"fund_type": "ETF"})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            app._load_review_rules = original_load_review_rules
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        result = payload["results"][0]
        self.assertEqual("missing", result["status"])
        self.assertEqual("subheading_missing", result["prospectus_missing_reason"])
        self.assertEqual("\u7ae0\u8282\u5df2\u627e\u5230\uff0c\u4f46\u62db\u52df\u5b50\u6761\u6b3e\u672a\u627e\u5230", result["message"])
        self.assertEqual("exact", result["prospectus_match_method"])

    def test_review_cross_check_treats_self_reference_differences_as_normalized_match(self):
        payload = self._run_cross_check_with_review_rules(
            [
                {
                    "heading": "\u7b2c\u5341\u516d\u90e8\u5206 \u57fa\u91d1\u8d44\u4ea7\u4f30\u503c",
                    "content": "2\u3001\u57fa\u91d1\u7ba1\u7406\u4eba\u5e94\u6bcf\u4e2a\u4f30\u503c\u65e5\u5bf9\u57fa\u91d1\u8d44\u4ea7\u4f30\u503c\uff0c\u4f46\u57fa\u91d1\u7ba1\u7406\u4eba\u6839\u636e\u6cd5\u5f8b\u6cd5\u89c4\u6216\u672c\u57fa\u91d1\u5408\u540c\u7684\u89c4\u5b9a\u6682\u505c\u4f30\u503c\u65f6\u9664\u5916\u3002",
                }
            ],
            [
                {
                    "heading": "\u57fa\u91d1\u8d44\u4ea7\u4f30\u503c",
                    "content": "2\u3001\u57fa\u91d1\u7ba1\u7406\u4eba\u5e94\u6bcf\u4e2a\u4f30\u503c\u65e5\u5bf9\u57fa\u91d1\u8d44\u4ea7\u4f30\u503c\uff0c\u4f46\u57fa\u91d1\u7ba1\u7406\u4eba\u6839\u636e\u6cd5\u5f8b\u6cd5\u89c4\u6216\u57fa\u91d1\u5408\u540c\u7684\u89c4\u5b9a\u6682\u505c\u4f30\u503c\u65f6\u9664\u5916\u3002",
                }
            ],
            {
                "chapter_rules": [
                    {
                        "source": "chapter_rules",
                        "contract": "\u7b2c\u5341\u516d\u90e8\u5206 \u57fa\u91d1\u8d44\u4ea7\u4f30\u503c",
                        "prospectus": "\u57fa\u91d1\u8d44\u4ea7\u4f30\u503c",
                        "relation": "\u76f4\u63a5\u5bf9\u5e94",
                        "consistency": "\u5b8c\u5168\u4e00\u81f4",
                        "detail": "\u57fa\u91d1\u5408\u540c\u4e0e\u62db\u52df\u8bf4\u660e\u4e66\u5b58\u5728\u6587\u4ef6\u81ea\u6307\u53e3\u5f84\u5dee\u5f02\u3002",
                        "sheet_name": "\u6d4b\u8bd5",
                    }
                ],
                "detail_rules": [],
                "key_diff_rules": [],
                "chapter_level": [],
                "key_diffs": [],
            },
        )

        result = payload["results"][0]
        self.assertEqual("pass", result["status"])
        self.assertEqual("info", result["severity"])

    def test_review_cross_check_compares_subheading_body_not_only_heading_text(self):
        payload = self._run_cross_check_with_review_rules(
            [
                {
                    "heading": "\u7b2c\u516b\u90e8\u5206 \u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u4e0e\u8d4e\u56de",
                    "content": "\u4e09\u3001\u7533\u8d2d\u4e0e\u8d4e\u56de\u7684\u539f\u5219\n\u7533\u8d2d\u548c\u8d4e\u56de\u9075\u5faa\u76f8\u540c\u7684\u57fa\u672c\u8981\u6c42\u3002",
                }
            ],
            [
                {
                    "heading": "\u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u8d4e\u56de",
                    "content": "\u4e09\u3001\u7533\u8d2d\u548c\u8d4e\u56de\u7684\u539f\u5219\n\u7533\u8d2d\u548c\u8d4e\u56de\u9075\u5faa\u76f8\u540c\u7684\u57fa\u672c\u8981\u6c42\u3002",
                }
            ],
            {
                "chapter_rules": [
                    {
                        "source": "chapter_rules",
                        "contract": "\u7b2c\u516b\u90e8\u5206 \u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u4e0e\u8d4e\u56de",
                        "prospectus": "\u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u8d4e\u56de",
                        "relation": "\u76f4\u63a5\u5bf9\u5e94",
                        "consistency": "\u57fa\u672c\u4e00\u81f4",
                        "detail": "",
                        "sheet_name": "\u6d4b\u8bd5",
                    }
                ],
                "detail_rules": [
                    {
                        "source": "detail_rules",
                        "contract": "\u7b2c\u516b\u90e8\u5206 \u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u4e0e\u8d4e\u56de / \u4e09\u3001\u7533\u8d2d\u4e0e\u8d4e\u56de\u7684\u539f\u5219",
                        "prospectus": "\u57fa\u91d1\u4efd\u989d\u7684\u7533\u8d2d\u8d4e\u56de",
                        "prospectus_point": "\u4e09\u3001\u7533\u8d2d\u548c\u8d4e\u56de\u7684\u539f\u5219",
                        "relation": "\u76f4\u63a5\u5bf9\u5e94",
                        "consistency": "\u57fa\u672c\u4e00\u81f4",
                        "detail": "\u201c\u7533\u8d2d\u548c\u8d4e\u56de/\u7533\u8d2d\u4e0e\u8d4e\u56de\u201d\u4ec5\u4e3a\u8fde\u8bcd\u5dee\u5f02\u3002",
                        "sheet_name": "\u6d4b\u8bd5",
                    }
                ],
                "key_diff_rules": [],
                "chapter_level": [],
                "key_diffs": [],
            },
        )

        result = payload["results"][0]
        self.assertEqual("pass", result["status"])
        self.assertEqual("info", result["severity"])
        self.assertGreaterEqual(result["similarity"], 95.0)

    def test_review_chapter_prefers_detail_results_skips_composite_rules(self):
        chapter_rule = {
            "contract": "第一部分前言 + 第八部分基金份额的申购与赎回 + 第十四部分基金的投资 / 五、六",
            "prospectus": "风险揭示 / 四、本基金特有的风险",
        }
        matched_detail_rules = [
            {
                "prospectus": "基金份额的申购赎回",
                "prospectus_point": "一、申购和赎回场所",
                "contract": "第八部分基金份额的申购与赎回 / 一、申购和赎回场所",
            }
        ]

        self.assertFalse(app._review_chapter_prefers_detail_results(chapter_rule, matched_detail_rules))

    def test_find_matching_detail_rules_ignores_composite_contract_locators_for_single_chapter_rule(self):
        chapter_rule = {
            "contract": "第七部分基金份额的上市交易",
            "prospectus": "基金份额的上市交易",
        }
        detail_rules = [
            {
                "prospectus": "风险揭示",
                "prospectus_point": "四、本基金特有的风险",
                "contract": "第一部分前言 + 第七部分基金份额的上市交易 + 第八部分基金份额的申购与赎回 + 第十四部分基金的投资 / 五、六",
            }
        ]

        self.assertEqual(app._find_matching_detail_rules(detail_rules, chapter_rule), [])

    def test_locate_review_rule_target_soft_matches_long_subheading_with_minor_wording_difference(self):
        sections = [
            {
                "heading": "第七部分 基金份额的上市交易",
                "content": (
                    "六、相关法律法规、中国证监会及深圳证券交易所对基金上市交易的规则等相关规定内容进行调整的，"
                    "基金合同相应予以修改，并按照新规定执行，且此项修改无须召开基金份额持有人大会。\n"
                    "正文说明。"
                ),
            }
        ]

        result = app._locate_review_rule_target(
            sections,
            "第七部分基金份额的上市交易 / 六、相关法律法规、中国证监会及上海证券交易所对基金上市交易的规则等相关规定内容进行调整的，本基金合同相应予以修改，并按照新规定执行，且此项修改无须召开基金份额持有人大会。",
            "基金份额的上市交易",
        )

        self.assertTrue(result["matched"])
        self.assertEqual(result["section_heading"], "第七部分 基金份额的上市交易")
        self.assertEqual(
            result["target_heading"],
            "六、相关法律法规、中国证监会及深圳证券交易所对基金上市交易的规则等相关规定内容进行调整的，基金合同相应予以修改，并按照新规定执行，且此项修改无须召开基金份额持有人大会。",
        )

    def test_review_cross_check_prefers_detail_rules_for_subscription_redemption(self):
        client = app.app.test_client()
        workbook = self._build_etf_cross_check_detail_workbook()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {
                    "heading": "第八部分 基金份额的申购与赎回",
                    "content": (
                        "一、申购和赎回场所\n"
                        "投资人可在交易场所办理申购和赎回。\n"
                        "二、申购和赎回的开放日及时间\n"
                        "开放日安排遵循交易所规则。\n"
                        "三、申购与赎回的原则\n"
                        "申购与赎回遵循份额对应和实物申赎原则。\n"
                        "四、申购与赎回的程序\n"
                        "基金合同仅规定确认、清算交收与登记的原则框架，具体在招募说明书中列示。\n"
                        "五、申购和赎回的数额限制\n"
                        "最小申购赎回单位由基金管理人确定并在招募说明书中公告。\n"
                        "六、申购和赎回的对价、费用及其用途\n"
                        "申购对价、赎回对价以及相关费用安排按照基金合同原则执行。\n"
                        "七、拒绝或暂停申购的情形\n"
                        "发生规定情形时可拒绝或暂停申购。\n"
                        "八、暂停赎回或延缓支付赎回对价的情形\n"
                        "发生规定情形时可暂停赎回或延缓支付赎回对价。\n"
                        "九、其他申购赎回方式\n"
                        "基金管理人可在条件成熟后开通其他申购赎回方式。\n"
                        "十、基金的非交易过户\n"
                        "基金登记机构可办理非交易过户。\n"
                        "十一、基金份额的冻结和解冻\n"
                        "基金份额可依法冻结和解冻。\n"
                        "十二、基金份额的转让\n"
                        "基金份额转让按规定办理。\n"
                        "十三、其他业务\n"
                        "基金管理人可依法办理其他业务。\n"
                        "十四、本基金合同生效后，若上海证券交易所和中国证券登记结算有限责任公司针对交易型开放式证券投资基金推出新的清算交收与登记模式并引入新的申购、赎回方式，本基金管理人有权调整本基金的清算交收与登记模式及申购、赎回方式。\n"
                        "届时将发布公告予以披露并在本基金招募说明书及其更新中予以更新。"
                    ),
                },
            ],
            "prospectus_sections": [
                {
                    "heading": "基金份额的申购赎回",
                    "content": (
                        "一、申购和赎回场所\n"
                        "投资人可在交易场所办理申购和赎回。\n"
                        "二、申购和赎回的开放日及时间\n"
                        "开放日安排遵循交易所规则并在公告中列明。\n"
                        "三、申购和赎回的原则\n"
                        "申购和赎回遵循份额对应和实物申赎原则。\n"
                        "四、申购和赎回的程序\n"
                        "招募说明书进一步写明T日、T+1和T+2的清算流程、交收节奏和投资人当日可卖出安排。\n"
                        "五、申购和赎回的数额限制\n"
                        "当前最小申购赎回单位为100万份，并补充上限及规模控制提示。\n"
                        "六、申购、赎回的对价、费用及其用途\n"
                        "招募说明书进一步列明基金份额净值计算公式、T日/T+1公告口径以及代理券商佣金上限不超过0.2%。\n"
                        "七、申购赎回清单的内容与格式\n"
                        "本节列明申购赎回清单构成、现金替代分类、计算公式与格式示例。\n"
                        "八、拒绝或暂停申购的情形\n"
                        "发生规定情形时可拒绝或暂停申购。\n"
                        "九、暂停赎回或延缓支付赎回对价的情形\n"
                        "发生规定情形时可暂停赎回或延缓支付赎回对价。\n"
                        "十、其他申购赎回方式\n"
                        "基金管理人可在条件成熟后开通其他申购赎回方式。\n"
                        "十一、基金的非交易过户\n"
                        "基金登记机构可办理非交易过户。\n"
                        "十二、基金份额的冻结和解冻\n"
                        "基金份额可依法冻结和解冻。\n"
                        "十三、基金份额的转让\n"
                        "基金份额转让按规定办理。\n"
                        "十四、其他业务\n"
                        "基金管理人可依法办理其他业务。\n"
                        "十五、基金合同生效后，若上海证券交易所和中国证券登记结算有限责任公司针对交易型开放式证券投资基金推出新的清算交收与登记模式并引入新的申购、赎回方式，本基金管理人有权调整本基金的清算交收与登记模式及申购、赎回方式。\n"
                        "届时将发布公告予以披露并在本基金招募说明书及其更新中予以更新。"
                    ),
                },
            ],
        }

        try:
            with self._with_review_rules(workbook):
                response = client.post("/api/review/cross_check", json={"fund_type": "ETF"})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        detail_results = [result for result in payload["results"] if result["rule_source"] == "detail_rules"]
        self.assertGreaterEqual(len(detail_results), 10)

        program_rule = next(result for result in detail_results if result.get("prospectus_target") == "四、申购和赎回的程序")
        self.assertEqual(program_rule["expected_relation"], "直接对应（招募细化）")
        self.assertEqual(program_rule["rule"]["consistency"], "部分对应")
        self.assertNotEqual(program_rule["status"], "fail")
        self.assertIn("T日/T+1/T+2", program_rule["expected_diff"])
        self.assertEqual(program_rule["contract_locator"], "第八部分 基金份额的申购与赎回 / 四、申购与赎回的程序")
        self.assertEqual(program_rule["prospectus_locator"], "基金份额的申购赎回 / 四、申购和赎回的程序")

        checklist_rule = next(result for result in detail_results if result.get("prospectus_target") == "七、申购赎回清单的内容与格式")
        self.assertEqual(checklist_rule["expected_relation"], "招募细化")
        self.assertIn("无直接对应", checklist_rule["rule"]["consistency"])
        self.assertEqual(checklist_rule["status"], "pass")
        self.assertIn("无直接对应", checklist_rule["message"])

    def test_review_cross_check_prefers_detail_rules_for_listing_and_trading(self):
        client = app.app.test_client()
        workbook = self._build_etf_cross_check_detail_workbook()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "prospectus_text": "mock prospectus",
            "contract_sections": [
                {
                    "heading": "第七部分 基金份额的上市交易",
                    "content": (
                        "一、基金份额的上市\n"
                        "基金份额可依法上市。\n"
                        "二、基金份额的上市交易\n"
                        "基金份额上市后可以依法交易。\n"
                        "三、终止上市交易\n"
                        "符合条件时可以终止上市交易。\n"
                        "四、基金份额参考净值（IOPV）的计算与公告\n"
                        "基金合同仅原则性规定基金份额参考净值的计算方法详见招募说明书。\n"
                        "五、在不违反法律法规及不损害基金份额持有人利益的前提下，本基金可以申请在其他证券交易所（含境外证券交易所）上市交易，而无需召开基金份额持有人大会审议。\n"
                        "本条款与招募说明书一致。\n"
                        "六、相关法律法规、中国证监会及上海证券交易所对基金上市交易的规则等相关规定内容进行调整的，本基金合同相应予以修改，并按照新规定执行，且此项修改无须召开基金份额持有人大会。\n"
                        "七、若上海证券交易所、中国证券登记结算有限责任公司增加了基金上市交易的新功能，基金管理人可以在履行适当的程序后增加相应功能。"
                    ),
                },
            ],
            "prospectus_sections": [
                {
                    "heading": "基金份额的上市交易",
                    "content": (
                        "一、基金份额的上市\n"
                        "基金份额可依法上市。\n"
                        "二、基金份额的上市交易\n"
                        "基金份额上市后可以依法交易。\n"
                        "三、终止上市交易\n"
                        "符合条件时可以终止上市交易。\n"
                        "四、基金份额参考净值（IOPV）的计算与公告\n"
                        "招募说明书进一步补充IOPV计算公式、保留位数和调整口径。\n"
                        "五、在不违反法律法规及不损害基金份额持有人利益的前提下，本基金可以申请在其他证券交易所（含境外证券交易所）上市交易，而无需召开基金份额持有人大会审议。\n"
                        "本条款与基金合同一致。\n"
                        "六、相关法律法规、中国证监会及上海证券交易所对基金上市交易的规则等相关规定内容进行调整的，基金合同相应予以修改，并按照新规定执行，且此项修改无须召开基金份额持有人大会。\n"
                        "七、若上海证券交易所、中国证券登记结算有限责任公司增加了基金上市交易的新功能，基金管理人可以在履行适当的程序后增加相应功能。"
                    ),
                },
            ],
        }

        try:
            with self._with_review_rules(workbook):
                response = client.post("/api/review/cross_check", json={"fund_type": "ETF"})
            self.assertEqual(response.status_code, 200)
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        detail_results = [result for result in payload["results"] if result["rule_source"] == "detail_rules"]
        iopv_rule = next(result for result in detail_results if result.get("prospectus_target") == "四、基金份额参考净值（IOPV）的计算与公告")
        self.assertEqual(iopv_rule["expected_relation"], "直接对应（招募细化）")
        self.assertEqual(iopv_rule["rule"]["consistency"], "部分对应")
        self.assertNotEqual(iopv_rule["status"], "fail")
        self.assertIn("IOPV", iopv_rule["expected_diff"])

        adjustment_rule = next(result for result in detail_results if result.get("prospectus_target", "").startswith("六、相关法律法规"))
        self.assertEqual(adjustment_rule["rule"]["consistency"], "基本一致")
        self.assertEqual(adjustment_rule["status"], "pass")

    def test_response_helper_closes_exported_file_handles_without_resource_warning(self):
        client = app.app.test_client()
        text = self._assert_no_resource_warning(
            lambda: self._consume_response(
                client.post("/api/export_prospectus_txt", json=self.build_form()),
                as_text=True,
            )
        )

        self.assertIn("一、绪言", text)
    def test_generated_prospectus_puts_variant_specific_important_notice_before_toc(self):
        standard_text = self.engine.generate(
            self.build_form(
                EXCHANGE="SSE",
                MARKET_TYPE="A_SHARE",
                INDEX_NAME="中证红利质量指数",
                INDEX_DESCRIPTION="中证红利质量指数选取50只高股息且盈利质量较好的上市公司证券作为样本。",
                INDEX_COMPILER="中证指数有限公司",
                INDEX_WEBSITE="www.csindex.com.cn",
            )
        )
        kechuang_text = self.engine.generate(
            self.build_form(
                EXCHANGE="SSE",
                MARKET_TYPE="KECHUANG",
                INDEX_NAME="科创人工智能指数",
                INDEX_DESCRIPTION="科创人工智能指数反映科创板人工智能主题上市公司证券的整体表现。",
                INDEX_COMPILER="中证指数有限公司",
                INDEX_WEBSITE="www.csindex.com.cn",
            )
        )
        chuangye_text = self.engine.generate(
            self.build_form(
                EXCHANGE="SZSE",
                MARKET_TYPE="CHUANGYE",
                INDEX_NAME="创业板中盘200指数",
                INDEX_DESCRIPTION="创业板中盘200指数反映创业板市值规模中等、流动性较好的200家公司股价变化情况。",
                INDEX_COMPILER="深圳证券信息有限公司",
                INDEX_WEBSITE="www.cnindex.com.cn",
            )
        )
        hk_text = self.engine.generate(
            self.build_form(
                EXCHANGE="SSE",
                MARKET_TYPE="HK_CONNECT",
                INDEX_NAME="港股通互联网指数",
                INDEX_DESCRIPTION="港股通互联网指数反映港股通范围内互联网主题上市公司证券的整体表现。",
                INDEX_COMPILER="中证指数有限公司",
                INDEX_WEBSITE="www.csindex.com.cn",
                IMPORTANT_NOTICE_STYLE="DIVIDEND_QUALITY",
            )
        )

        standard_preface = self.extract_preface_before_toc(standard_text)
        kechuang_preface = self.extract_preface_before_toc(kechuang_text)
        chuangye_preface = self.extract_preface_before_toc(chuangye_text)
        hk_preface = self.extract_preface_before_toc(hk_text)

        self.assertIn("\u91cd\u8981\u63d0\u793a", standard_preface)
        self.assertIn("\u91cd\u8981\u63d0\u793a", kechuang_preface)
        self.assertIn("\u91cd\u8981\u63d0\u793a", chuangye_preface)
        self.assertIn("\u91cd\u8981\u63d0\u793a", hk_preface)
        self.assertIn(
            "\u672c\u57fa\u91d1\u7ecf\u4e2d\u56fd\u8bc1\u76d1\u4f1a202X\u5e74X\u6708X\u65e5\u8bc1\u76d1\u8bb8\u53ef\u3014202X\u3015XXX\u53f7\u6587\u6ce8\u518c\u52df\u96c6\u3002",
            standard_preface,
        )
        self.assertIn("\u4e2d\u8bc1\u7ea2\u5229\u8d28\u91cf\u6307\u6570", standard_preface)
        self.assertIn("\u79d1\u521b\u4eba\u5de5\u667a\u80fd\u6307\u6570", kechuang_preface)
        self.assertIn("\u521b\u4e1a\u677f\u4e2d\u76d8200\u6307\u6570", chuangye_preface)
        self.assertIn("\u6e2f\u80a1\u901a\u4e92\u8054\u7f51\u6307\u6570", hk_preface)
        self.assertIn(
            "\u6709\u5173\u6807\u7684\u6307\u6570\u5177\u4f53\u7f16\u5236\u65b9\u6848\u53ca\u6210\u4efd\u80a1\u4fe1\u606f\u8be6\u89c1\u4e2d\u8bc1\u6307\u6570\u6709\u9650\u516c\u53f8\u7f51\u7ad9\uff0c\u7f51\u5740\uff1awww.csindex.com.cn\u3002",
            standard_preface,
        )
        self.assertIn(
            "\u6709\u5173\u6807\u7684\u6307\u6570\u5177\u4f53\u7f16\u5236\u65b9\u6848\u53ca\u6210\u4efd\u80a1\u4fe1\u606f\u8be6\u89c1\u6df1\u5733\u8bc1\u5238\u4fe1\u606f\u6709\u9650\u516c\u53f8\u7f51\u7ad9\uff0c\u7f51\u5740\uff1awww.cnindex.com.cn\u3002",
            chuangye_preface,
        )
        self.assertNotIn("\u53ca\u5176\u672a\u6765\u53ef\u80fd\u53d1\u751f\u7684\u53d8\u66f4", standard_preface)
        self.assertNotIn("\u53ca\u5176\u672a\u6765\u53ef\u80fd\u53d1\u751f\u7684\u53d8\u66f4", kechuang_preface)
        self.assertNotIn("\u53ca\u5176\u672a\u6765\u53ef\u80fd\u53d1\u751f\u7684\u53d8\u66f4", chuangye_preface)
        self.assertNotIn("\u53ca\u5176\u672a\u6765\u53ef\u80fd\u53d1\u751f\u7684\u53d8\u66f4", hk_preface)

    def test_generate_prospectus_requires_index_description(self):
        with self.assertRaisesRegex(ValueError, "INDEX_DESCRIPTION"):
            self.engine.generate(self.build_form(INDEX_DESCRIPTION=""))

    def test_generate_prospectus_requires_manual_notice_style_for_ambiguous_index_name(self):
        with self.assertRaisesRegex(ValueError, "重要提示"):
            self.engine.generate(
                self.build_form(
                    INDEX_NAME="中证创业板双创指数",
                    INDEX_DESCRIPTION="双创指数反映相关上市公司证券的整体表现。",
                )
            )

    def test_generate_prospectus_accepts_manual_notice_style_override_for_ambiguous_index_name(self):
        text = self.engine.generate(
            self.build_form(
                INDEX_NAME="中证创业板双创指数",
                INDEX_DESCRIPTION="双创指数反映相关上市公司证券的整体表现。",
                IMPORTANT_NOTICE_STYLE="CHUANGYE_200",
            )
        )

        notice = self.extract_important_notice(text)
        self.assertNotIn("及其未来可能发生的变更", notice)
        self.assertIn("本基金标的指数为中证创业板双创指数。双创指数反映相关上市公司证券的整体表现。", notice)
        self.assertNotIn("及其未来可能发生的变更", text)

    def test_important_notice_uses_name_keywords_instead_of_market_defaults(self):
        text = self.engine.generate(
            self.build_form(
                EXCHANGE="SZSE",
                MARKET_TYPE="A_SHARE",
                INDEX_NAME="国证红利成长指数",
                INDEX_DESCRIPTION="国证红利成长指数反映相关样本证券的整体表现。",
                INDEX_COMPILER="深圳证券信息有限公司",
                INDEX_WEBSITE="www.cnindex.com.cn",
            )
        )

        notice = self.extract_important_notice(text)
        self.assertIn("本基金标的指数为国证红利成长指数。", notice)
        self.assertIn("本基金被动跟踪标的指数“国证红利成长指数”，因此，本基金的业绩表现与标的指数的表现密切相关。", notice)
        self.assertNotIn("及其未来可能发生的变更", text)

    def test_important_notice_uses_sse_single_market_account_clause(self):
        text = self.engine.generate(
            self.build_form(
                EXCHANGE="SSE",
                MARKET_TYPE="A_SHARE",
                MARKET_SCOPE="SINGLE_MARKET",
            )
        )

        notice = self.extract_important_notice(text)
        self.assertIn("投资人投资本基金时需具有上海证券交易所A股账户或基金账户。", notice)
        self.assertIn("如投资人需要参与网下股票认购或基金的申购、赎回，则应开立上海证券交易所A股账户。", notice)
        self.assertNotIn("投资人认购本基金时需具有证券账户", notice)

    def test_important_notice_uses_sse_cross_market_account_clause(self):
        text = self.engine.generate(
            self.build_form(
                EXCHANGE="SSE",
                MARKET_TYPE="A_SHARE",
                MARKET_SCOPE="CROSS_MARKET",
            )
        )

        notice = self.extract_important_notice(text)
        self.assertIn("投资人投资本基金时需具有上海证券交易所A股账户或基金账户。", notice)
        self.assertIn("如投资人需要使用标的指数成份股中的上海证券交易所上市股票参与网下股票认购或基金的申购、赎回，则应开立上海证券交易所A股账户；", notice)
        self.assertIn("如投资人需要使用标的指数成份股中的深圳证券交易所上市股票参与网下股票认购，则还应开立深圳证券交易所A股账户。", notice)

    def test_important_notice_uses_szse_cross_market_account_clause(self):
        text = self.engine.generate(
            self.build_form(
                EXCHANGE="SZSE",
                MARKET_TYPE="A_SHARE",
                MARKET_SCOPE="CROSS_MARKET",
            )
        )

        notice = self.extract_important_notice(text)
        self.assertIn("投资人投资本基金时需具有深圳证券交易所A股账户或基金账户。", notice)
        self.assertIn("如投资人需要使用标的指数成份股中的深圳证券交易所上市股票参与基金的申购、赎回，则应开立深圳证券交易所A股账户。", notice)

    def test_important_notice_uses_sse_hk_account_clause(self):
        text = self.engine.generate(
            self.build_form(
                EXCHANGE="SSE",
                MARKET_TYPE="HK_CONNECT",
                INDEX_NAME="中证港股通互联网指数",
                INDEX_DESCRIPTION="中证港股通互联网指数反映港股通范围内互联网主题上市公司证券的整体表现。",
                INDEX_COMPILER="中证指数有限公司",
                INDEX_WEBSITE="www.csindex.com.cn",
            )
        )

        notice = self.extract_important_notice(text)
        self.assertIn("投资人投资本基金时需具有上海证券交易所A股账户或基金账户。", notice)
        self.assertIn("如投资人需要参与基金的申购、赎回，则应开立并使用上海证券交易所A股账户。", notice)

    def test_reused_chapters_drop_contract_wording(self):
        text = self.engine.generate(self.build_form())
        invest = self.extract_chapter(text, "十一")
        valuation = self.extract_chapter(text, "十三")
        fees = self.extract_chapter(text, "十五")
        self.assertNotIn("本基金合同", invest)
        self.assertNotIn("本基金合同", valuation)
        self.assertIn("《基金合同》", fees)

    def test_reused_chapter_normalizer_fixes_side_pocket_reference(self):
        normalized = self.engine._normalize_reused_prospectus_chapter(
            "实施侧袋机制期间的投资运作安排，详见招募说明书的规定。基金管理人根据本基金合同履行信息披露义务。"
        )
        self.assertIn("详见招募说明书“侧袋机制”部分的规定", normalized)
        self.assertNotIn("详见招募说明书的规定", normalized)
        self.assertNotIn("本基金合同", normalized)

    def test_risk_chapter_normalizer_uses_frontend_min_sub_unit_value(self):
        normalized = self.engine._normalize_prospectus_risk_chapter(
            "投资人按原1,000,000份申购并持有的基金份额，可能无法按照新的1,000,000份全部赎回。",
            self.build_form(MIN_SUB_UNIT="3,000,000份"),
        )
        self.assertIn("300万份", normalized)
        self.assertNotIn("1,000,000份", normalized)

    def test_etf_chapter_ten_keeps_prospectus_only_sections(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "十")
        self.assertIn("一、申购和赎回场所", chapter)
        self.assertIn("四、申购和赎回的程序", chapter)
        self.assertIn("五、申购和赎回的数额限制", chapter)
        self.assertIn("对价、费用及其用途", chapter)
        self.assertIn("七、申购赎回清单的内容与格式", chapter)

    def test_etf_chapter_ten_sec5_formats_min_sub_unit_without_parenthetical_numeric_style(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "\u5341")
        sec5 = self.extract_section(chapter, "\u4e94")
        self.assertIn("\u5f53\u63a5\u53d7\u7533\u8d2d\u7533\u8bf7\u5bf9\u5b58\u91cf\u57fa\u91d1\u4efd\u989d\u6301\u6709\u4eba\u5229\u76ca\u6784\u6210\u6f5c\u5728\u91cd\u5927\u4e0d\u5229\u5f71\u54cd\u65f6", sec5)
        self.assertIn("25万份", sec5)
        self.assertNotIn("250,000份", sec5)
        self.assertNotIn("1,000,000份", sec5)
        self.assertNotIn("（即", sec5)
        self.assertNotIn("|\u9879\u76ee|\u5185\u5bb9|", sec5)
        self.assertNotIn("|---|---|", sec5)

    def test_chapter_ten_includes_venue_section_before_open_day_section(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        chapter = self.extract_chapter(text, "十")
        self.assertEqual(1, chapter.count("一、申购和赎回场所"))
        self.assertLess(chapter.index("一、申购和赎回场所"), chapter.index("二、申购和赎回的开放日及时间"))

    def test_etf_chapter_ten_sec7_keeps_formula_and_markdown_tables_for_cross_market(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        chapter = self.extract_chapter(text, "\u5341")
        sec7 = self.extract_section(chapter, "\u4e03")
        self.assertIn("七、申购赎回清单的内容与格式", sec7)
        self.assertIn("T日申购赎回清单的格式举例如下：", sec7)
        self.assertIn("基本信息\nT-1日信息内容\nT日信息内容\n成份股信息内容。", sec7)
        self.assertIn("说明：此表仅为示意", sec7)
        self.assertIn("现金替代比例的计算公式为：", sec7)

    def test_market_sh_standard_purchase_redemption_list_clause_keeps_sample_body(self):
        clause = self.engine._resolve_prospectus_clause_text(
            "PURCHASE_REDEMPTION_LIST_CLAUSE",
            "MARKET_SH_STANDARD",
        )

        self.assertIn("现金替代比例的计算公式为：", clause)
        self.assertIn(
            "现金替代比例（%）＝Σ（第i只替代证券的数量×该证券参考价格）×100%／（申购基金份额×参考基金份额净值）",
            clause,
        )
        self.assertIn(
            "对于沪市成份证券，现金替代的类型可以设为：\"禁止\"、\"允许\"和\"必须\"。对于非沪市成份证券，可以设为：\"允许\"和\"必须\"。",
            clause,
        )
        self.assertIn(
            "T日申购赎回清单中公告T日预估现金差额。T日预估现金差额＝T-1日最小申购赎回单位的基金资产净值－（申购赎回清单中必须现金替代的固定替代金额＋申购赎回清单中可以现金替代成份证券的数量与该证券调整后T日开盘参考价相乘之和+申购赎回清单中禁止现金替代成份证券的数量与该证券调整后T日开盘参考价相乘之和）。",
            clause,
        )

    def test_etf_chapter_ten_sec7_matches_reference_for_all_variants(self):
        variant_forms = {
            "SSE_CROSS": self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="CROSS_MARKET"),
            "SSE_SINGLE": self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="SINGLE_MARKET"),
            "SZSE_SINGLE": self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="SINGLE_MARKET"),
            "SZSE_CROSS": self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="CROSS_MARKET"),
            "SSE_HK": self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT", HAS_HK_CONNECT=True),
            "SZSE_HK": self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT", HAS_HK_CONNECT=True),
        }

        for variant_key, form in variant_forms.items():
            with self.subTest(variant_key=variant_key):
                generated = self.engine.generate(form)
                chapter = self.extract_chapter(generated, "十")
                sec7 = self.extract_section(chapter, "七")
                ref = self.engine._load_reference_fixed_content(form)
                expected = self.engine._restore_missing_prospectus_blocks(
                    self.engine._reference_chapter_ten_sec7(ref, form)
                )
                self.assertEqual(expected, sec7)

    def test_etf_chapter_ten_sections_four_and_eight_to_thirteen_follow_manual_reference(self):
        variant_forms = {
            "SSE_CROSS": self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="CROSS_MARKET"),
            "SSE_SINGLE": self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="SINGLE_MARKET"),
            "SZSE_SINGLE": self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="SINGLE_MARKET"),
            "SZSE_CROSS": self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="CROSS_MARKET"),
            "SSE_HK": self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT", HAS_HK_CONNECT=True),
            "SZSE_HK": self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT", HAS_HK_CONNECT=True),
        }

        for variant_key, form in variant_forms.items():
            with self.subTest(variant_key=variant_key):
                ref = self.engine._load_reference_fixed_content(form)
                ref_sections = ref.get("十", {}).get("sections", {})
                self.assertTrue(ref_sections, f"{variant_key} is missing manual chapter ten reference content")

                generated = self.engine.generate(form)
                chapter = self.extract_chapter(generated, "十")
                generated_sections = self.engine._split_top_sections("\n".join(chapter.splitlines()[1:]))

                for sec_cn in ("四", "八", "九", "十", "十一", "十二", "十三"):
                    manual_section = str(ref_sections.get(sec_cn) or "").strip()
                    actual_section = str(generated_sections.get(sec_cn) or "").strip()
                    if manual_section:
                        expected_section = self.engine._ensure_subsection_heading(
                            self.engine._normalize_prospectus_chapter_ten_text(
                                self.engine._replace_manual_reference_text(manual_section, form)
                            ),
                            sec_cn,
                        )
                        self.assertEqual(expected_section, actual_section)
                    else:
                        self.assertEqual("", actual_section, f"{variant_key} should omit section {sec_cn} when manual lacks it")

    def test_etf_chapter_ten_sec7_splits_format_note_to_standalone_line(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        chapter = self.extract_chapter(text, "\u5341")
        sec7 = self.extract_section(chapter, "\u4e03")

        self.assertIn("成份股信息内容。\n说明：此表仅为示意", sec7)

    def test_etf_chapter_ten_sec7_keeps_body_paragraphs_when_inserting_format_block(self):
        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")
        form["PROSPECTUS_CH10_SEC7_BODY"] = "\n".join([
            "七、申购赎回清单的内容与格式",
            "T日申购赎回清单的格式举例如下：",
            "基本信息",
            "本基金管理人将在T日申购赎回清单公告后及时披露相关信息。",
            "投资者应结合申购赎回清单内容办理申购、赎回申请。",
            "说明：此表仅为示意",
        ])
        form["PROSPECTUS_CH10_SEC7_FORMAT_BLOCK"] = "\n".join([
            "基本信息",
            "T-1日信息内容",
            "T日信息内容",
            "成份股信息内容。",
        ])

        sec7 = self.engine._build_chapter_ten_sec7(form)

        self.assertIn("T日申购赎回清单的格式举例如下：", sec7)
        self.assertIn("基本信息", sec7)
        self.assertIn("T-1日信息内容", sec7)
        self.assertIn("T日信息内容", sec7)
        self.assertIn("成份股信息内容。", sec7)
        self.assertIn("本基金管理人将在T日申购赎回清单公告后及时披露相关信息。", sec7)
        self.assertIn("投资者应结合申购赎回清单内容办理申购、赎回申请。", sec7)

    def test_etf_chapter_ten_sec7_indents_format_block_titles_on_separate_lines(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        chapter = self.extract_chapter(text, "十")
        sec7 = self.extract_section(chapter, "七")
        self.assertIn("基本信息\nT-1日信息内容\nT日信息内容\n成份股信息内容。", sec7)
        self.assertNotRegex(sec7, r"\n[ \t\u3000]+基本信息\n")

    def test_etf_chapter_ten_sec7_keeps_hk_variant_branch(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))
        chapter = self.extract_chapter(text, "十")
        sec7 = self.extract_section(chapter, "七")
        self.assertIn("基金合同生效后，本基金的申购赎回采用全现金替代", sec7)
        self.assertIn("T日申购赎回清单的格式举例如下：", sec7)
        self.assertIn("基本信息\nT-1日信息内容\nT日信息内容\n成份股信息内容。", sec7)
        self.assertIn("现金替代保证金率", sec7)
        self.assertNotIn("现金替代比例（%）＝", sec7)

    def test_contract_hk_connect_uses_exchange_specific_nav_term(self):
        sh_text = app.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT"))
        sz_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))

        self.assertIn("基金份额净值：指估值日基金资产净值除以估值日基金份额总数", sh_text)
        self.assertIn("基金份额净值：指计算日基金资产净值除以计算日基金份额总数", sz_text)

    def test_contract_hk_connect_uses_exchange_specific_purchase_redemption_principle(self):
        sh_text = app.engine.generate(
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT", DISTRIBUTION_FREQ="MONTHLY")
        )
        sz_text = app.engine.generate(
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT", DISTRIBUTION_FREQ="QUARTERLY")
        )

        self.assertIn(
            "2、本基金的申购对价、赎回对价包括组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价。",
            sh_text,
        )
        self.assertIn(
            "2、本基金的申购对价、赎回对价包括现金替代、现金差额及其他对价。",
            sz_text,
        )
        self.assertNotIn(
            "2、本基金的申购对价、赎回对价包括组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价。",
            sz_text,
        )

    def test_contract_hk_connect_uses_exchange_specific_purchase_redemption_consideration_definition(self):
        sh_text = app.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT"))
        sz_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))

        self.assertIn(
            "2、申购对价是指投资人申购基金份额时应交付的组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价。",
            sh_text,
        )
        self.assertIn(
            "2、申购对价是指投资人申购基金份额时应交付的现金替代、现金差额及其他对价。",
            sz_text,
        )
        self.assertNotIn(
            "2、申购对价是指投资人申购基金份额时应交付的组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价。",
            sz_text,
        )

    def test_chapter_six_reuses_reference_stock_subscription_formula(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", HAS_STOCK_SUBSCRIPTION=True))
        chapter = self.extract_chapter(text, "\u516d")
        sec11 = self.extract_section(chapter, "\u5341\u4e00")
        self.assertIn("认购份额的计算公式", sec11)
        self.assertIn("第i只股票在网下股票认购期最后一日的均价", sec11)
        self.assertIn("\u9664\u606f\u4e14\u9001\u80a1", sec11)
        self.assertIn("\u9664\u606f\u4e14\u914d\u80a1", sec11)

    def test_contract_hk_connect_inserts_preamble_hk_risk_clause(self):
        hk_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))
        standard_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"))

        self.assertIn("本基金资产投资于港股，会面临港股通机制下因投资环境、投资标的、市场制度以及交易规则等差异带来的特有风险", hk_text)
        self.assertIn("汇率风险（汇率波动可能对基金的投资收益造成损失）", hk_text)
        self.assertNotIn("本基金资产投资于港股，会面临港股通机制下因投资环境、投资标的、市场制度以及交易规则等差异带来的特有风险", standard_text)

    def test_contract_hk_connect_inserts_hk_connect_definitions(self):
        hk_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))
        standard_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"))

        self.assertIn("内地与香港股票市场交易互联互通机制：指上海证券交易所、深圳证券交易所分别和香港联合交易所有限公司", hk_text)
        self.assertIn("港股通：指内地投资者委托内地证券公司", hk_text)
        self.assertIn("流动性受限资产", hk_text)
        self.assertNotIn("内地与香港股票市场交易互联互通机制：指上海证券交易所、深圳证券交易所分别和香港联合交易所有限公司", standard_text)

    def test_contract_hk_connect_updates_glossary_purchase_redemption_and_profit_definitions(self):
        sh_text = app.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT"))
        sz_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))
        standard_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"))

        self.assertIn("49、申购对价：指投资人申购基金份额时，按基金合同和招募说明书规定应交付的组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价", sh_text)
        self.assertIn("50、赎回对价：指基金份额持有人赎回基金份额时，基金管理人按基金合同和招募说明书规定应交付给赎回人的组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价", sh_text)
        self.assertIn("49、申购对价：指投资人申购基金份额时，按基金合同和招募说明书规定应交付的现金替代、现金差额及其他对价", sz_text)
        self.assertIn("50、赎回对价：指基金份额持有人赎回基金份额时，基金管理人按基金合同和招募说明书规定应交付给赎回人的现金替代、现金差额及其他对价", sz_text)
        self.assertIn("57、基金利润：指基金利息收入、汇兑损益、投资收益、公允价值变动收益和其他收入扣除相关费用后的余额", sh_text)
        self.assertIn("57、基金利润：指基金利息收入、汇兑损益、投资收益、公允价值变动收益和其他收入扣除相关费用后的余额", sz_text)
        self.assertIn("57、基金利润：指基金利息收入、投资收益、公允价值变动收益和其他收入扣除相关费用后的余额", standard_text)

    def test_contract_hk_connect_rewrites_subscription_and_record_sections(self):
        sh_text = app.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT"))
        sz_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))

        self.assertIn("基金管理人可选择网上现金认购和网下现金认购2种方式发售本基金。", sh_text)
        self.assertIn("基金管理人可选择网上现金认购和网下现金认购2种方式发售本基金。", sz_text)
        self.assertNotIn("网下股票认购3种方式发售本基金", sh_text)
        self.assertIn("2、募集期间认购资金的处理方式", sh_text)
        self.assertIn("2、募集期间认购资金的处理方式", sz_text)
        self.assertNotIn("2、募集期间认购资金与股票", sh_text)
        self.assertNotIn("基金募集金额（）不少于2亿元人民币", sh_text)
        self.assertIn("基金募集金额不少于2亿元人民币且基金认购人数不少于200人的条件下", sh_text)
        self.assertIn("基金募集金额不少于2亿元人民币且基金认购人数不少于200人的条件下", sz_text)
        self.assertNotIn("网下股票认购募集的股票按照交易所和登记机构的规则和流程予以冻结", sh_text)
        self.assertNotIn("对于基金募集期间网下股票认购所募集的股票", sh_text)
        self.assertIn("有效认购款项在募集期间产生的利息将折算为基金份额归基金份额持有人所有", sz_text)

    def test_contract_hk_connect_uses_exchange_specific_listing_sections(self):
        sh_text = app.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT"))
        sz_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))

        self.assertIn("3、符合上海证券交易所规定的其他条件。", sh_text)
        self.assertIn("本基金基金份额在上海证券交易所的上市交易需遵照《上海证券交易所交易规则》", sh_text)
        self.assertIn("三、终止上市交易", sh_text)
        self.assertIn("基金管理人或基金管理人委托的其他机构可以在相关证券交易所开市后根据申购赎回清单、汇率数据和组合证券内各只证券的实时成交数据，计算基金份额参考净值（IOPV）", sh_text)
        self.assertIn("3、《深圳证券交易所证券投资基金上市规则》规定的其他条件。", sz_text)
        self.assertIn("本基金基金份额在深圳证券交易所的上市交易，应遵照《深圳证券交易所交易规则》", sz_text)
        self.assertIn("三、停复牌、暂停上市、恢复上市及终止上市交易", sz_text)
        self.assertIn("并将计算结果向深圳证券交易所发送，由深圳证券交易所对外发布", sz_text)

    def test_contract_hk_connect_uses_exchange_specific_purchase_redemption_sections(self):
        sh_text = app.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="HK_CONNECT"))
        sz_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))

        self.assertIn("目前本基金的申购赎回采用全现金替代模式，申购对价、赎回对价包括现金替代、现金差额及其他对价。未来在证券交易所和登记机构系统允许的情况下，本基金可采用实物申购赎回模式", sh_text)
        self.assertIn("具体办理时间为上海证券交易所的正常交易日的交易时间（若该交易日非港股通交易日，则本基金不开放申购和赎回）", sh_text)
        self.assertIn("如遇港股通暂停交易或交收、登记公司系统故障", sh_text)
        self.assertIn("10、因港股通每日额度等原因需要控制基金申购规模", sh_text)
        self.assertIn("9、法律法规规定、中国证监会或上海证券交易所认定的其他情形。", sh_text)
        self.assertIn("未来在证券交易所和登记结算机构系统允许的情况下，本基金可采用实物申购赎回模式", sh_text)

        self.assertIn("本基金的开放日为深圳证券交易所的交易日（若该交易日非港股通交易日，则本基金不开放申购和赎回）", sz_text)
        self.assertIn("2、申购对价是指投资人申购基金份额时应交付的现金替代、现金差额及其他对价。", sz_text)
        self.assertIn("9、因港股通每日额度等原因需要控制基金申购规模", sz_text)
        self.assertIn("8、法律法规、深圳证券交易所规定或中国证监会认定的其他情形。", sz_text)
        self.assertIn("7、基金管理人指定的代理机构可依据本基金合同开展其他服务", sz_text)

    def test_contract_hk_connect_adds_proxy_voting_section_and_removes_stock_refund_duty(self):
        hk_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))

        self.assertIn("八、港股通股票投资的代理投票", hk_text)
        self.assertIn("中国证券登记结算有限责任公司以自己的名义，通过香港中央结算有限公司行使对该股票发行人的权利", hk_text)
        self.assertIn("（24）基金管理人在募集期间未能达到基金的备案条件，《基金合同》不能生效，基金管理人承担全部募集费用，将已募集资金并加计银行同期活期存款利息在基金募集期结束后30日内退还基金认购人；", hk_text)
        self.assertNotIn("投资者以股票认购的，相关股票的解冻按照《业务规则》的规定处理；", hk_text)

    def test_contract_hk_connect_adds_valuation_clause_for_hk_connect_assets(self):
        hk_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))
        standard_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"))

        self.assertIn("5、港股通投资持有外币证券资产估值涉及到的主要货币对人民币汇率，以估值日中国人民银行或其授权机构公布的人民币汇率中间价为准。", hk_text)
        self.assertIn("基金管理人每个估值日对基金资产估值后，将基金资产净值、基金份额净值结果发送基金托管人", hk_text)
        self.assertNotIn("港股通投资持有外币证券资产估值涉及到的主要货币对人民币汇率", standard_text)

    def test_contract_hk_connect_adds_disclosure_for_hk_connect_stock_investment(self):
        hk_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))
        standard_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"))

        self.assertIn("（十六）基金投资港股通股票的信息披露", hk_text)
        self.assertIn("本基金在季度报告、中期报告、年度报告等定期报告和招募说明书（更新）等文件中披露参与港股通交易的相关情况。", hk_text)
        self.assertIn("（十七）基金投资资产支持证券的信息披露", hk_text)
        self.assertIn("（十六）基金投资资产支持证券的信息披露", standard_text)

    def test_contract_hk_connect_updates_fee_ranges_summary_and_payment_wording(self):
        hk_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT"))
        standard_text = app.engine.generate(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"))

        self.assertIn("上述“（一）基金费用的种类”中第3－12项费用", hk_text)
        self.assertIn("上述“（一）基金费用的种类”中第3－11项费用", standard_text)
        self.assertEqual(2, hk_text.count("因投资港股通股票而产生的各项合理费用"))
        self.assertEqual(0, standard_text.count("因投资港股通股票而产生的各项合理费用"))
        self.assertIn("缴纳基金认购款项、应付申购对价及法律法规和《基金合同》所规定的费用；", hk_text)
        self.assertNotIn("缴纳基金认购款项或股票、应付申购对价及法律法规和《基金合同》所规定的费用；", hk_text)

    def test_chapter_six_includes_subscription_fee_table(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        chapter = self.extract_chapter(text, "\u516d")
        sec8 = self.extract_section(chapter, "\u516b")

        self.assertIn("|认购份额（S）|认购费率|", sec8)
        self.assertIn("|S＜100万份|0.30%|", sec8)
        self.assertIn("|S≥100万份|每笔500元|", sec8)

    def test_chapter_six_removes_subscription_account_placeholder(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        chapter = self.extract_chapter(text, "\u516d")
        self.assertNotIn("{SUBSCRIPTION_ACCOUNT_CLAUSE}", chapter)

    def test_chapter_six_keeps_required_reference_sections(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", HAS_STOCK_SUBSCRIPTION=True))
        chapter = self.extract_chapter(text, "六")

        for section_cn, heading in (
            ("一", "募集期"),
            ("二", "发售对象"),
            ("三", "首次募集规模上限"),
            ("五", "募集场所"),
            ("六", "基金份额发售面值和认购价格"),
            ("九", "网上现金认购"),
            ("十", "网下现金认购"),
        ):
            if section_cn == "六":
                self.assertIn(f"{section_cn}、{heading}", chapter)
            else:
                section = self.extract_section(chapter, section_cn)
                self.assertIn(f"{section_cn}、{heading}", section)

    def test_chapter_six_sections_from_eight_onward_prefer_variant_reference(self):
        variables = self.engine._inject_clause_texts(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT", HAS_STOCK_SUBSCRIPTION=True))
        variables["ONLINE_CASH_SUBSCRIPTION_CLAUSE"] = "九、网上现金认购\n条款九"
        variables["OFFLINE_CASH_SUBSCRIPTION_CLAUSE"] = "十、网下现金认购\n条款十"
        variables["OFFLINE_STOCK_SUBSCRIPTION_CLAUSE"] = "十一、网下股票认购\n条款十一"
        variables["PROSPECTUS_CH6_SEC12"] = "十二、认购的确认\n条款十二"
        variables["PROSPECTUS_CH6_SEC13"] = "十三、特别事项\n条款十三"
        ref = {
            "六": {
                "sections": {
                    "八": "八、认购费用\n参考八",
                    "九": "九、网上现金认购\n参考九",
                    "十": "十、网下现金认购\n参考十",
                    "十一": "十一、网下股票认购\n6、认购份额的计算公式：",
                    "十二": "十二、认购的确认\n参考十二",
                    "十三": "十三、特别事项\n参考十三",
                }
            }
        }

        chapter = self.engine._build_chapter_six_body(variables, ref)

        self.assertEqual("八、认购费用\n参考八", self.extract_section(chapter, "八"))
        self.assertEqual("九、网上现金认购\n参考九", self.extract_section(chapter, "九"))
        self.assertEqual("十、网下现金认购\n参考十", self.extract_section(chapter, "十"))
        self.assertIn("投资者的认购份额＝Σ", self.extract_section(chapter, "十一"))
        self.assertEqual("十二、认购的确认\n参考十二", self.extract_section(chapter, "十二"))
        self.assertEqual("十三、特别事项\n参考十三", self.extract_section(chapter, "十三"))

    def test_chapter_six_section_four_prefers_variant_reference(self):
        variables = self.engine._inject_clause_texts(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT", HAS_STOCK_SUBSCRIPTION=True))
        variables["METHOD_SUBSCRIBE_DERIVED_FROM_CONTRACT"] = "四、发售方式\n合同条款四"
        variables["PROSPECTUS_CH6_SEC4"] = "四、发售方式\n模板条款四"
        ref = {
            "六": {
                "sections": {
                    "四": "四、发售方式\n参考四",
                }
            }
        }

        chapter = self.engine._build_chapter_six_body(variables, ref)
        self.assertEqual("四、发售方式\n参考四", self.extract_section(chapter, "四"))

    def test_chapter_six_section_seven_prefers_manual_reference(self):
        variables = self.engine._inject_clause_texts(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT", HAS_STOCK_SUBSCRIPTION=True))
        variables["SUB_ACCOUNT_OPENING_CLAUSE"] = "七、认购开户\n合同条款七"
        variables["PROSPECTUS_CH6_SEC7"] = "七、认购开户\n模板条款七"
        ref = {
            "六": {
                "sections": {
                    "七": "七、认购开户\n参考七",
                }
            }
        }

        chapter = self.engine._build_chapter_six_body(variables, ref)
        self.assertEqual("七、认购开户\n参考七", self.extract_section(chapter, "七"))

    def test_chapter_six_without_stock_subscription_stops_at_section_twelve(self):
        variables = self.engine._inject_clause_texts(self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT", HAS_STOCK_SUBSCRIPTION=False))
        variables["PROSPECTUS_CH6_SEC13"] = "十三、特别事项\n不应出现"
        ref = {
            "六": {
                "sections": {
                    "八": "八、认购费用\n参考八",
                    "九": "九、网上现金认购\n参考九",
                    "十": "十、网下现金认购\n参考十",
                    "十一": "十一、认购份额的计算\n参考十一",
                    "十二": "十二、认购的确认\n参考十二",
                    "十三": "十三、特别事项\n参考十三",
                }
            }
        }

        chapter = self.engine._build_chapter_six_body(variables, ref)

        self.assertEqual("十二、认购的确认\n参考十二", self.extract_section(chapter, "十二"))
        self.assertNotIn("十三、特别事项", chapter)

    def test_chapter_six_sse_stock_subscription_does_not_invent_section_thirteen_when_reference_omits_it(self):
        variables = self.engine._inject_clause_texts(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", HAS_STOCK_SUBSCRIPTION=True))
        variables["PROSPECTUS_CH6_SEC13"] = "十三、特别事项\n不应出现"
        ref = {
            "六": {
                "sections": {
                    "八": "八、认购费用\n参考八",
                    "九": "九、网上现金认购\n参考九",
                    "十": "十、网下现金认购\n参考十",
                    "十一": "十一、网下股票认购\n6、认购份额的计算公式：",
                    "十二": "十二、募集期间的资金、股票与费用\n参考十二",
                }
            }
        }

        chapter = self.engine._build_chapter_six_body(variables, ref)

        self.assertEqual("十二、募集期间的资金、股票与费用\n参考十二", self.extract_section(chapter, "十二"))
        self.assertNotIn("十三、特别事项", chapter)

    def test_reference_fixed_content_falls_back_to_manual_txt_for_chapter_six_sections(self):
        ref = self.engine._load_reference_fixed_content(
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="SINGLE_MARKET", HAS_STOCK_SUBSCRIPTION=True)
        )
        sections = ref["六"]["sections"]

        self.assertIn("认购费用由投资人承担，不高于0.3%", sections["八"])
        self.assertIn("认购佣金由发售代理机构向投资人收取", sections["九"])
        self.assertIn("认购金额＝认购价格×认购份额", sections["十"])
        self.assertIn("网下股票认购", sections["十一"])
        self.assertIn("募集资金利息与募集股票权益的处理方式", sections["十二"])
        self.assertIn("募集期间的资金、股票与费用", sections["十三"])

    def test_generated_chapter_six_uses_manual_sections_eight_to_thirteen(self):
        text = self.engine.generate(
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="SINGLE_MARKET", HAS_STOCK_SUBSCRIPTION=True)
        )
        chapter = self.extract_chapter(text, "六")

        self.assertIn("认购费用由投资人承担，不高于0.3%", self.extract_section(chapter, "八"))
        sec9 = self.extract_section(chapter, "九")
        sec10 = self.extract_section(chapter, "十")
        sec11 = self.extract_section(chapter, "十一")
        sec12 = self.extract_section(chapter, "十二")
        sec13 = self.extract_section(chapter, "十三")

        self.assertIn("认购佣金由发售代理机构向投资人收取", sec9)
        self.assertNotIn("0.08%", sec9)
        self.assertNotIn("0.8元", sec9)
        self.assertIn("认购金额＝认购价格×认购份额", sec10)
        self.assertNotIn("0.8%", sec10)
        self.assertNotIn("100,800元", sec10)
        self.assertIn("网下股票认购", sec11)
        self.assertIn("募集资金利息与募集股票权益的处理方式", sec12)
        self.assertIn("募集期间的资金、股票与费用", sec13)

    def test_generated_sse_single_chapter_six_uses_manual_fee_examples(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="KECHUANG"))
        chapter = self.extract_chapter(text, "六")

        sec8 = self.extract_section(chapter, "八")
        sec9 = self.extract_section(chapter, "九")
        sec10 = self.extract_section(chapter, "十")
        sec11 = self.extract_section(chapter, "十一")
        sec12 = self.extract_section(chapter, "十二")

        self.assertIn("认购费用由投资人承担，不高于0.3%", sec8)
        self.assertIn("S＜100万份", sec8)
        self.assertNotIn("0.8%", sec8)
        self.assertIn("1,003元", sec9)
        self.assertNotIn("1,008元", sec9)
        self.assertIn("认购金额＝认购价格×认购份额", sec10)
        self.assertIn("100,000元", sec10)
        self.assertNotIn("100,800元", sec10)
        self.assertIn("网下股票认购", sec11)
        self.assertIn("募集期间的资金、股票与费用", sec12)
        self.assertNotIn("十三、", chapter)

    def test_generated_chapter_six_section_four_uses_manual_reference(self):
        text = self.engine.generate(
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE", MARKET_SCOPE="SINGLE_MARKET", HAS_STOCK_SUBSCRIPTION=True)
        )
        chapter = self.extract_chapter(text, "六")
        sec4 = self.extract_section(chapter, "四")

        self.assertIn("基金管理人可选择网上现金认购、网下现金认购和网下股票认购3种方式发售本基金", sec4)
        self.assertNotIn("本基金将自", sec4)

    def test_generated_chapter_six_without_stock_subscription_uses_manual_sections_only_through_twelve(self):
        text = self.engine.generate(
            self.build_form(EXCHANGE="SZSE", MARKET_TYPE="HK_CONNECT", HAS_STOCK_SUBSCRIPTION=False)
        )
        chapter = self.extract_chapter(text, "六")

        self.assertIn("认购费用由投资人承担，不高于0.3%", self.extract_section(chapter, "八"))
        self.assertIn("认购佣金由发售代理机构向投资人收取", self.extract_section(chapter, "九"))
        self.assertIn("认购金额＝认购价格×认购份额", self.extract_section(chapter, "十"))
        self.assertIn("募集资金利息的处理方式", self.extract_section(chapter, "十一"))
        self.assertIn("募集期间的资金与费用", self.extract_section(chapter, "十二"))
        self.assertNotIn("十三、", chapter)

    def test_chapter_eight_starts_with_reference_intro_sentence(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "八")
        chapter_body = "\n".join(chapter.splitlines()[1:]).strip()

        self.assertTrue(
            chapter_body.startswith("基金合同生效后，为提高交易便利，本基金可以进行基金份额折算。"),
            chapter_body[:120],
        )

    def test_chapter_nine_uses_correct_title_and_first_subsection(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        chapter = self.extract_chapter(text, "九")
        sec1 = self.extract_section(chapter, "一")

        self.assertTrue(chapter.startswith("九、基金份额的上市交易"))
        self.assertTrue(sec1.startswith("一、基金份额的上市"))
        self.assertNotIn("一、基金份额的上市交易", sec1)

    def test_chapter_ten_section_one_prefers_contract_clause_over_manual_reference(self):
        variables = self.engine._inject_clause_texts(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        variables["CONTRACT_PART8_SEC1"] = "一、申购和赎回场所\n合同第一节"
        ref = {
            "十": {
                "sections": {
                    "一": "一、申购和赎回场所\n参考第一节",
                }
            }
        }

        chapter = self.engine._build_chapter_ten_body(variables, ref)

        self.assertEqual("一、申购和赎回场所\n合同第一节", self.extract_section(chapter, "一"))

    def test_generated_prospectus_leaves_blank_line_between_top_level_sections(self):
        text = self.engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", HAS_STOCK_SUBSCRIPTION=True))

        self.assertIn("一、募集期\n", text)
        self.assertIn("\n\n二、发售对象\n", text)
        self.assertIn("\n\n三、首次募集规模上限\n", text)

    def test_contract_summary_excludes_signing_page_content(self):
        variables = self.engine._extract_contract_sections(self.build_form())
        summary = variables["CONTRACT_SUMMARY_TEXT"]
        self.assertNotIn("签署页", summary)
        self.assertNotIn("本页无正文", summary)
        self.assertNotIn("（盖章）", summary)
        self.assertNotIn("（签字或盖章）", summary)
        self.assertNotIn("（签名）", summary)

    def test_etf_chapter_fourteen_keeps_distribution_conditions(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "十四")
        self.assertIn("四、收益分配条件", chapter)

    def test_chapter_fourteen_section_four_matches_canonical_reference(self):
        text = self.engine.generate(self.build_form())
        chapter = self.extract_chapter(text, "\u5341\u56db")
        actual_section = self.extract_section(chapter, "\u56db")
        reference = self.engine._load_reference_fixed_content(self.build_form())
        expected_section = self.extract_section(reference["\u5341\u56db"]["body"], "\u56db")

        self.assertEqual(expected_section, actual_section)

    def test_chapters_twenty_two_to_twenty_five_match_canonical_reference(self):
        text = self.engine.generate(self.build_form())
        reference = self.engine._load_reference_fixed_content(self.build_form())
        reference_name = self.engine._get_reference_prospectus_docx(self.build_form()).stem.replace("招募说明书", "").strip()
        target_name = self.build_form()["FUND_NAME"]

        for chapter_cn in ("\u4e8c\u5341\u4e8c", "\u4e8c\u5341\u4e09", "\u4e8c\u5341\u56db", "\u4e8c\u5341\u4e94"):
            with self.subTest(chapter_cn=chapter_cn):
                chapter = self.extract_chapter(text, chapter_cn)
                actual_body = "\n".join(chapter.splitlines()[1:]).strip()
                expected_body = reference[chapter_cn]["body"].strip()
                def normalize(body):
                    lines = []
                    for line in body.strip().splitlines():
                        stripped = line.strip()
                        if not stripped:
                            continue
                        stripped = re.sub(r"^\d+、", "", stripped)
                        stripped = stripped.replace(reference_name, "{FUND}").replace(target_name, "{FUND}")
                        stripped = re.sub(r"《[^》]*基金合同》", "《{FUND}基金合同》", stripped)
                        stripped = re.sub(r"《[^》]*托管协议》", "《{FUND}托管协议》", stripped)
                        stripped = re.sub(r"《[^》]*登记结算服务协议》", "《{FUND}登记结算服务协议》", stripped)
                        lines.append(stripped)
                    return "\n".join(lines)
                self.assertEqual(normalize(expected_body), normalize(actual_body))

    def test_generated_prospectus_toc_and_tail_stop_at_chapter_twenty_five(self):
        text = self.engine.generate(self.build_form())
        toc = self.extract_toc(text)
        toc_lines = [line.strip() for line in toc.splitlines() if line.strip()]
        model = self.engine.generate_render_model(self.build_form())
        tail_text = text[text.rfind("\u4e8c\u5341\u4e94\u3001\u5907\u67e5\u6587\u4ef6"):]

        self.assertEqual("\u4e8c\u5341\u4e94\u3001\u5907\u67e5\u6587\u4ef6", toc_lines[-1])
        self.assertEqual(25, len(model["chapters"]))
        self.assertEqual("\u4e8c\u5341\u4e94", model["chapters"][-1]["chapter_cn"])
        self.assertEqual("\u5907\u67e5\u6587\u4ef6", model["chapters"][-1]["title"])
        self.assertNotIn("\n\u4e09\u3001\u57fa\u91d1\u7ba1\u7406\u4eba\n", tail_text)
        self.assertNotIn("\n\u56db\u3001\u57fa\u91d1\u6258\u7ba1\u4eba\n", tail_text)

    def test_contract_docx_has_body_page_numbers_starting_at_one(self):
        contract_text = app.engine.generate(self.build_form())
        xml_map = self._docx_xml(app.engine.build_docx(contract_text))
        document_xml = xml_map["word/document.xml"]
        footer_xml = "".join(v for k, v in xml_map.items() if k.startswith("word/footer"))

        self.assertIn('w:start="1"', document_xml)
        self.assertIn("<w:titlePg", document_xml)
        self.assertIn("PAGE", footer_xml)
        self.assertIn('w:jc w:val="center"', footer_xml)

    def test_contract_docx_uses_word_native_toc_field(self):
        contract_text = app.engine.generate(self.build_form())
        xml_map = self._docx_xml(app.engine.build_docx(contract_text))
        document_xml = xml_map["word/document.xml"]
        settings_xml = xml_map["word/settings.xml"]

        self.assertIn('TOC \\\\o "1-1" \\\\h \\\\z \\\\u', document_xml)
        self.assertNotIn("PAGEREF", document_xml)
        self.assertIn("updateFields", settings_xml)

    def test_contract_docx_uses_reference_section_grid(self):
        from docx.oxml.ns import qn

        contract_text = app.engine.generate(self.build_form())
        doc = self._docx_document(app.engine.build_docx(contract_text))
        line_pitches = [
            [
                grid.get(qn("w:linePitch"))
                for grid in section._sectPr.findall(qn("w:docGrid"))
            ]
            for section in doc.sections
        ]

        self.assertEqual([["312"], ["312"]], line_pitches)

    def test_contract_docx_body_paragraph_uses_farming_reference_base_style(self):
        from docx.oxml.ns import qn

        contract_text = app.engine.generate(self.build_form())
        doc = self._docx_document(app.engine.build_docx(contract_text))
        paragraph = next(
            p for p in doc.paragraphs
            if "订立本基金合同的目的是保护投资人合法权益" in p.text
        )
        pPr = paragraph._p.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        indent = pPr.find(qn("w:ind"))
        run_properties = paragraph.runs[0]._r.get_or_add_rPr()
        size = run_properties.find(qn("w:sz"))

        self.assertEqual("360", spacing.get(qn("w:line")))
        self.assertEqual("auto", spacing.get(qn("w:lineRule")))
        self.assertEqual("0", spacing.get(qn("w:after")))
        self.assertEqual("480", indent.get(qn("w:firstLine")))
        self.assertIsNone(indent.get(qn("w:firstLineChars")))
        self.assertEqual("24", size.get(qn("w:val")))

    def test_contract_docx_cover_keeps_six_blank_lines_before_manager_line(self):
        contract_text = app.engine.generate(self.build_form())
        doc = self._docx_document(app.engine.build_docx(contract_text))
        title_index = next(
            idx for idx, paragraph in enumerate(doc.paragraphs)
            if paragraph.text.strip().endswith("基金合同（草案）")
        )
        manager_index = next(
            idx for idx, paragraph in enumerate(doc.paragraphs)
            if paragraph.text.strip().startswith("基金管理人：")
        )

        blank_count = sum(1 for paragraph in doc.paragraphs[title_index + 1:manager_index] if not paragraph.text.strip())
        self.assertEqual(6, blank_count)

    def test_contract_docx_adds_body_header_title(self):
        import xml.etree.ElementTree as ET

        contract_text = app.engine.generate(self.build_form())
        xml_map = self._docx_xml(app.engine.build_docx(contract_text))
        header_xml = "".join(v for k, v in xml_map.items() if k.startswith("word/header"))
        document_root = ET.fromstring(xml_map["word/document.xml"])
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        sect_prs = document_root.findall(".//w:sectPr", ns)

        self.assertGreaterEqual(len(sect_prs), 2)
        first_ref_tags = [child.tag.rsplit("}", 1)[-1] for child in list(sect_prs[0])]
        self.assertNotIn("headerReference", first_ref_tags)
        self.assertIn("基金合同（草案）", header_xml)
        self.assertIn("right", header_xml)

    def test_contract_docx_uses_reference_default_tab_stop(self):
        contract_text = app.engine.generate(self.build_form())
        xml_map = self._docx_xml(app.engine.build_docx(contract_text))

        self.assertIn('w:defaultTabStop w:val="720"', xml_map["word/settings.xml"])

    def test_contract_docx_cover_uses_draft_title_without_date_line(self):
        contract_text = app.engine.generate(self.build_form())
        doc = self._docx_document(app.engine.build_docx(contract_text))
        nonempty = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]

        self.assertEqual("测试交易型开放式指数证券投资基金基金合同（草案）", nonempty[0])
        self.assertNotIn("2026年1月1日", nonempty[:6])

    def test_contract_open_day_time_sentence_does_not_include_extra_zhi_character(self):
        contract_text = app.engine.generate(
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", WORKING_DAY_TYPE="BOTH")
        )

        self.assertIn(
            "投资人在开放日办理基金份额的申购和赎回，具体办理时间为上海证券交易所、深圳证券交易所的正常交易日的交易时间",
            contract_text,
        )
        self.assertNotIn(
            "投资人在开放日办理基金份额的申购和赎回，具体办理时间为指上海证券交易所、深圳证券交易所的正常交易日的交易时间",
            contract_text,
        )

    def test_prospectus_docx_has_body_page_numbers_starting_at_one(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        xml_map = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text))
        document_xml = xml_map["word/document.xml"]
        footer_xml = "".join(v for k, v in xml_map.items() if k.startswith("word/footer"))
        header_xml = "".join(v for k, v in xml_map.items() if k.startswith("word/header"))
        settings_xml = xml_map["word/settings.xml"]

        self.assertGreaterEqual(document_xml.count("<w:sectPr"), 2)
        self.assertIn('w:start="1"', document_xml)
        self.assertIn("<w:titlePg", document_xml)
        self.assertIn('w:top="1440"', document_xml)
        self.assertIn('w:bottom="1440"', document_xml)
        self.assertIn('w:left="1800"', document_xml)
        self.assertIn('w:right="1800"', document_xml)
        self.assertIn('w:header="851"', document_xml)
        self.assertIn('w:footer="992"', document_xml)
        self.assertIn("PAGE", footer_xml)
        self.assertIn('w:jc w:val="center"', footer_xml)
        self.assertIn("\u6d4b\u8bd5\u4ea4\u6613\u578b\u5f00\u653e\u5f0f\u6307\u6570\u8bc1\u5238\u6295\u8d44\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66", header_xml)
        self.assertIn('w:jc w:val="right"', header_xml)
        self.assertIn('w:pBdr', xml_map["word/styles.xml"])
        self.assertIn('w:val="single"', header_xml)
        self.assertIn('w:val="18"', footer_xml)
        self.assertIn("updateFields", settings_xml)

    def test_prospectus_docx_footer_parts_do_not_keep_template_page_number_paragraphs(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        footer_parts = self._docx_footer_xml_parts(app.prospectus_engine.build_docx_prospectus(prospectus_text))

        self.assertTrue(footer_parts)
        for name, footer_xml in footer_parts.items():
            self.assertEqual(1, footer_xml.count("PAGE"), name)
            self.assertLessEqual(1, footer_xml.count('w:fldCharType="begin"'), name)

    def test_prospectus_docx_formats_important_notice_and_chapter_titles(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        xml_map = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text))
        document_xml = xml_map["word/document.xml"]
        styles_xml = xml_map["word/styles.xml"]

        self.assertRegex(document_xml, r"w:jc w:val=\"center\".{0,500}\u91cd\u8981\u63d0\u793a")
        self.assertIn("<w:t>绪言</w:t>", document_xml)
        self.assertRegex(document_xml, r'w:pStyle w:val="[^"]+"')
        self.assertNotIn("<w:t>一、绪言</w:t>", document_xml)
        self.assertIn('w:eastAsia="黑体"', styles_xml)
        self.assertIn('w:sz w:val="32"', styles_xml)

    def test_prospectus_docx_renders_section7_markdown_tables_as_word_tables(self):
        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")
        prospectus_text = app.prospectus_engine.generate(form)
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))

        self.assertGreaterEqual(len(doc.tables), 5)
        table_headers = [table.rows[0].cells[0].text for table in doc.tables]
        self.assertIn("认购份额（S）", table_headers)
        self.assertIn("最新公告日期", table_headers)
        self.assertIn("现金差额（单位:元）", table_headers)
        self.assertIn("预估现金差额（单位:元）", table_headers)
        self.assertIn("证券代码", table_headers)

    def test_prospectus_docx_keeps_subscription_fee_formula_and_section7_note(self):
        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", HAS_STOCK_SUBSCRIPTION=True)
        prospectus_text = app.prospectus_engine.generate(form)
        docx_bytes = app.prospectus_engine.build_docx_prospectus(prospectus_text, form)
        doc = self._docx_document(docx_bytes)
        paragraph_text = "\n".join(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
        table_text = "\n".join(
            "\n".join(cell.text.strip() for cell in row.cells)
            for table in doc.tables
            for row in table.rows
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(docx_bytes)
            temp_path = tmp.name
        try:
            extracted_text = app._extract_docx_text_with_tables(temp_path)
        finally:
            Path(temp_path).unlink(missing_ok=True)

        self.assertIn("认购份额（S）", table_text)
        self.assertIn("认购费率", table_text)
        self.assertIn(app.ProspectusEngine._stock_subscription_formula_text(), extracted_text)
        self.assertIn(app.ProspectusEngine._cash_substitution_ratio_formula_label(), extracted_text)
        self.assertIn("说明：此表仅为示意", extracted_text)

    def test_prospectus_docx_cover_spacing_and_body_indentation(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        docx_bytes = app.prospectus_engine.build_docx_prospectus(prospectus_text)
        document_xml = self._docx_xml(docx_bytes)["word/document.xml"]
        doc = self._docx_document(docx_bytes)

        cover_title_text = "测试交易型开放式指数证券投资基金招募说明书"
        cover_title = f"<w:t>{cover_title_text}</w:t>"
        manager_label = "<w:t>基金管理人：</w:t>"
        important_notice = "<w:t>重要提示</w:t>"
        self.assertIn(cover_title, document_xml)
        self.assertIn(manager_label, document_xml)
        self.assertIn(important_notice, document_xml)

        cover_title_index = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text == cover_title_text)
        self.assertGreaterEqual(cover_title_index, 3)

        cover_start = document_xml.index(cover_title)
        manager_start = document_xml.index(manager_label)
        notice_start = document_xml.index(important_notice)
        cover_slice = document_xml[cover_start:manager_start]
        transition_slice = document_xml[manager_start:notice_start]
        notice_slice = document_xml[notice_start:document_xml.index("<w:t>目录</w:t>")]

        self.assertGreaterEqual(cover_slice.count("</w:pPr></w:p>"), 6)
        self.assertIn("<w:sectPr>", transition_slice)
        self.assertIn('w:firstLine="420"', notice_slice)
        self.assertIn('w:firstLineChars="200"', notice_slice)
        self.assertNotIn('<w:ind w:firstLineChars="200"', cover_slice)

    def test_prospectus_docx_places_toc_on_its_own_page_before_body(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        document_xml = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text))["word/document.xml"]

        toc_start = document_xml.index("<w:t>目录</w:t>")
        body_start = document_xml.index("<w:t>绪言</w:t>")
        toc_slice = document_xml[toc_start:body_start]

        self.assertIn("<w:sectPr>", toc_slice)

    def test_prospectus_docx_starts_each_chapter_on_new_page(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        xml_map = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text))
        document_xml = xml_map["word/document.xml"]
        model = app.prospectus_engine.generate_render_model(self.build_form())

        # Cover page + body preface share the first two sections; every later chapter should add one more section.
        self.assertGreaterEqual(document_xml.count("<w:sectPr"), len(model["chapters"]) + 1)
    def test_prospectus_docx_uses_word_native_toc_field(self):
        form = self.build_form()
        prospectus_text = app.prospectus_engine.generate(form)
        xml_map = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))
        document_xml = xml_map["word/document.xml"]
        settings_xml = xml_map["word/settings.xml"]

        self.assertIn('TOC \\\\o "2-2" \\\\h \\\\z \\\\u', document_xml)
        self.assertIn("updateFields", settings_xml)
        self.assertNotIn("PAGEREF", document_xml)
        self.assertIn("目录", document_xml)
        self.assertIn("<w:t>绪言</w:t>", document_xml)

    def test_prospectus_docx_exports_cash_substitution_and_stock_subscription_as_native_math(self):
        form = self.build_form(IMPORTANT_NOTICE_STYLE="DIVIDEND_QUALITY", HAS_STOCK_SUBSCRIPTION=True)
        prospectus_text = app.prospectus_engine.generate(form)
        document_xml = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))["word/document.xml"]

        self.assertGreaterEqual(document_xml.count("m:oMath"), 2)
        self.assertIn("现金替代比例", document_xml)
        self.assertIn("投资者的认购份额", document_xml)

    def test_generated_prospectus_restores_cash_substitution_formula_text_when_reference_sec7_omits_it(self):
        form = self.build_form(
            EXCHANGE="SSE",
            MARKET_TYPE="A_SHARE",
            IMPORTANT_NOTICE_STYLE="DIVIDEND_QUALITY",
            HAS_STOCK_SUBSCRIPTION=True,
        )

        prospectus_text = app.prospectus_engine.generate(form)

        self.assertIn(app.ProspectusEngine._cash_substitution_ratio_formula_label(), prospectus_text)
        self.assertIn(app.ProspectusEngine._cash_substitution_ratio_formula_text(), prospectus_text)

    def test_prospectus_docx_native_formula_xml_uses_nary_math_without_hidden_text_runs(self):
        form = self.build_form(
            EXCHANGE="SSE",
            MARKET_TYPE="A_SHARE",
            IMPORTANT_NOTICE_STYLE="DIVIDEND_QUALITY",
            HAS_STOCK_SUBSCRIPTION=True,
        )
        prospectus_text = app.prospectus_engine.generate(form)
        document_xml = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))["word/document.xml"]

        self.assertGreaterEqual(document_xml.count("<m:nary>"), 2)
        self.assertGreaterEqual(document_xml.count("<m:sub>"), 2)
        self.assertGreaterEqual(document_xml.count("<m:sup>"), 2)
        self.assertNotIn("w:vanish", document_xml)

    def test_prospectus_docx_gives_subscription_fee_table_visible_borders(self):
        form = self.build_form(IMPORTANT_NOTICE_STYLE="DIVIDEND_QUALITY")
        prospectus_text = app.prospectus_engine.generate(form)
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))

        fee_table = next(table for table in doc.tables if table.rows and table.rows[0].cells[0].text == "认购份额（S）")

        self.assertIn("tblBorders", fee_table._tbl.xml)
        self.assertIn('w:val="single"', fee_table._tbl.xml)

    def test_prospectus_docx_toc_styles_are_defined_for_word_generated_entries(self):
        form = self.build_form()
        prospectus_text = app.prospectus_engine.generate(form)
        xml_map = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))
        document_xml = xml_map["word/document.xml"]
        styles_xml = xml_map["word/styles.xml"]

        self.assertRegex(document_xml, r"w:sz w:val=\"28\".{0,400}目录")
        self.assertRegex(document_xml, r"<w:instrText[^>]*> TOC \\\\o &quot;2-2&quot; \\\\h \\\\z \\\\u </w:instrText>|<w:instrText[^>]*> TOC \\\\o \"2-2\" \\\\h \\\\z \\\\u </w:instrText>")
        self.assertIn('w:styleId="TOCHeading"', styles_xml)
        self.assertIn('w:styleId="TOC1"', styles_xml)

    def test_prospectus_docx_heading_style_uses_heiti_without_auto_numbering(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        styles_xml = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text))["word/styles.xml"]
        style_blocks = re.findall(r"<w:style\b[\s\S]*?</w:style>", styles_xml)
        matching_blocks = [
            block for block in style_blocks
            if 'w:eastAsia="黑体"' in block and 'w:ascii="Arial"' in block and 'w:sz w:val="32"' in block
        ]

        self.assertTrue(matching_blocks)
        self.assertTrue(any("<w:numPr>" not in block for block in matching_blocks))

    def test_prospectus_docx_adds_blank_line_after_chapter_heading_and_keeps_left_alignment(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text))

        heading_index = next(
            i for i, paragraph in enumerate(doc.paragraphs)
            if paragraph.style.name == "Heading 2" and "绪言" in paragraph.text
        )
        heading = doc.paragraphs[heading_index]
        spacer = doc.paragraphs[heading_index + 1]
        body = doc.paragraphs[heading_index + 2]

        self.assertEqual("", spacer.text)
        self.assertTrue(body.text.strip())
        self.assertEqual(0, heading.paragraph_format.first_line_indent)
        self.assertEqual(0, spacer.paragraph_format.first_line_indent)
        self.assertEqual("一、\t绪言", heading.text)
        self.assertIn("<w:tab/>", heading._p.xml)
        self.assertIn("w:tabs", heading._p.xml)
        self.assertIn('w:pos="840"', heading._p.xml)
        self.assertIn('w:firstLine="420"', body._p.xml)

    def test_prospectus_docx_section7_reference_titles_follow_body_indentation(self):
        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")
        prospectus_text = app.prospectus_engine.generate(form)
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))

        target_titles = ["基本信息", "T-1日信息内容", "T日信息内容", "成份股信息内容"]
        paragraphs = {paragraph.text.strip(): paragraph for paragraph in doc.paragraphs if paragraph.text.strip() in target_titles}

        for title in target_titles:
            self.assertIn(title, paragraphs)
            self.assertIn('w:firstLine="420"', paragraphs[title]._p.xml)

    def test_prospectus_docx_adds_blank_line_before_top_level_section_heading(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", HAS_STOCK_SUBSCRIPTION=True))
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text))

        sec2_index = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text == "二、发售对象")
        spacer = doc.paragraphs[sec2_index - 1]

        self.assertEqual("", spacer.text)

    def test_prospectus_docx_toc_excludes_placeholder_and_chapter_twenty_one_subitems(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        document_xml = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text))["word/document.xml"]

        pre_body_xml = document_xml[:document_xml.index("\u7eea\u8a00")]
        self.assertNotIn("【托管人情况待填写】", pre_body_xml)
        self.assertNotIn("基金托管协议当事人", pre_body_xml)
        self.assertNotIn("基金管理人对基金托管人的业务核查", pre_body_xml)

    def test_prospectus_docx_keeps_numbered_paragraph_indentation(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE"))
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text, self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")))

        targets = {
            "\u4e03\u3001\u7533\u8d2d\u8d4e\u56de\u6e05\u5355\u7684\u5185\u5bb9\u4e0e\u683c\u5f0f",
            "1\u3001\u7533\u8d2d\u8d4e\u56de\u6e05\u5355\u7684\u5185\u5bb9",
            "\uff081\uff09\u7ec4\u5408\u8bc1\u5238\u76f8\u5173\u5185\u5bb9",
        }
        found = {}
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text in targets:
                found[text] = paragraph

        self.assertEqual(targets, set(found))
        for paragraph in found.values():
            self.assertIn('w:firstLine="420"', paragraph._p.xml)

    def test_prospectus_docx_body_uses_reference_style_chapter_headings(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        xml_map = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text))
        document_xml = xml_map["word/document.xml"]
        numbering_xml = xml_map["word/numbering.xml"]

        self.assertIn(">绪言<", document_xml)
        self.assertIn(">基金托管协议的内容摘要<", document_xml)
        self.assertIn("一、基金托管协议当事人", document_xml)
        self.assertIn('w:abstractNumId w:val="1"', numbering_xml)
        self.assertIn("upperLetter", numbering_xml)

    def test_prospectus_docx_reuses_reference_section7_tables_for_variant(self):
        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")
        prospectus_text = app.prospectus_engine.generate(form)
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))

        self.assertEqual(5, len(doc.tables))
        section7_tables = [table for table in doc.tables if table.rows and table.rows[0].cells[0].text != "认购份额（S）"]
        self.assertEqual([4, 3, 15, 2], [len(table.rows) for table in section7_tables[:3]] + [len(section7_tables[3].rows)])
        self.assertEqual([2, 2, 2, 8], [len(table.columns) for table in section7_tables])
        self.assertEqual(8, len(section7_tables[3].rows[0].cells))
        self.assertIn("\u8bc1\u5238\u4ee3\u7801", section7_tables[3].rows[0].cells[0].text)

    def test_prospectus_docx_applies_half_point_borders_to_every_table(self):
        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")
        prospectus_text = app.prospectus_engine.generate(form)
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))

        border_pattern = re.compile(
            r"<w:(?:top|left|bottom|right|insideH|insideV|start|end)\b[^>]*w:val=\"single\"[^>]*w:sz=\"(\d+)\""
        )

        self.assertTrue(doc.tables)
        for table in doc.tables:
            first_cell_text = table.rows[0].cells[0].text if table.rows and table.rows[0].cells else ""
            self.assertIn("<w:tblBorders>", table._tbl.xml, first_cell_text)
            self.assertEqual({"4"}, set(border_pattern.findall(table._tbl.xml)), first_cell_text)

    def test_prospectus_docx_replaces_reference_fund_name_in_szse_cross_section7_tables(self):
        form = self.build_form(
            EXCHANGE="SZSE",
            MARKET_TYPE="A_SHARE",
            MARKET_SCOPE="CROSS_MARKET",
            FUND_NAME="Test ETF Fund",
        )
        engine = app.ProspectusEngine()
        reference_tables = engine._load_reference_section7_table_xmls(form)
        self.assertTrue(reference_tables)

        block = engine._section7_reference_table_block(reference_tables[0], form)
        fund_name_row = next(row for row in block["rows"] if row[0].strip().rstrip("：:") == "基金名称")

        self.assertEqual(form["FUND_NAME"], fund_name_row[1])
        if block.get("table_xml"):
            self.assertIn(form["FUND_NAME"], block["table_xml"])
            self.assertNotIn("南方中证电池主题交易型开放式指数证券投资基金", block["table_xml"])

    def test_prospectus_docx_inserts_reference_section7_titles_before_tables(self):
        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")
        prospectus_text = app.prospectus_engine.generate(form)
        document_xml = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))["word/document.xml"]

        title_and_header_pairs = [
            ("\u57fa\u672c\u4fe1\u606f", "\u6700\u65b0\u516c\u544a\u65e5\u671f"),
            ("T-1\u65e5\u4fe1\u606f\u5185\u5bb9", "\u73b0\u91d1\u5dee\u989d\uff08\u5355\u4f4d:\u5143\uff09"),
            ("T\u65e5\u4fe1\u606f\u5185\u5bb9", "\u9884\u4f30\u73b0\u91d1\u5dee\u989d\uff08\u5355\u4f4d:\u5143\uff09"),
            ("\u6210\u4efd\u80a1\u4fe1\u606f\u5185\u5bb9", "\u8bc1\u5238\u4ee3\u7801"),
        ]

        section7_start = document_xml.index("\u7533\u8d2d\u8d4e\u56de\u6e05\u5355\u7684\u683c\u5f0f\u4e3e\u4f8b\u5982\u4e0b\uff1a")
        previous_title_index = -1
        previous_header_index = -1
        for title_text, header_text in title_and_header_pairs:
            title_index = document_xml.index(title_text, section7_start)
            header_index = document_xml.index(header_text, section7_start)
            self.assertLess(title_index, header_index)
            self.assertGreater(title_index, previous_title_index)
            self.assertGreater(header_index, previous_header_index)
            previous_title_index = title_index
            previous_header_index = header_index

    def test_prospectus_docx_body_numbered_paragraphs_are_not_bold_but_titles_and_table_headers_stay_bold(self):
        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")
        prospectus_text = app.prospectus_engine.generate(form)
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))

        body_targets = {
            "\u4e03\u3001\u7533\u8d2d\u8d4e\u56de\u6e05\u5355\u7684\u5185\u5bb9\u4e0e\u683c\u5f0f",
            "1\u3001\u7533\u8d2d\u8d4e\u56de\u6e05\u5355\u7684\u5185\u5bb9",
            "\uff081\uff09\u7ec4\u5408\u8bc1\u5238\u76f8\u5173\u5185\u5bb9",
        }
        found = {}
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text in body_targets:
                found[text] = paragraph

        self.assertEqual(body_targets, set(found))
        for paragraph in found.values():
            visible_runs = [run for run in paragraph.runs if run.text.strip()]
            self.assertTrue(visible_runs)
            self.assertTrue(all(run.bold is not True for run in visible_runs), paragraph.text)

        chapter_heading = next(
            paragraph for paragraph in doc.paragraphs
            if paragraph.style.name == "Heading 2" and "\u7eea\u8a00" in paragraph.text
        )
        self.assertTrue(any(run.bold for run in chapter_heading.runs if run.text.strip()))

        header_cell_xml = doc.tables[-4].rows[0].cells[0]._tc.xml
        self.assertIn("w:b", header_cell_xml)

    def test_prospectus_docx_still_inserts_section7_table_when_reference_title_is_missing(self):
        form = self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE")
        prospectus_text = app.prospectus_engine.generate(form)
        original_loader = app.prospectus_engine._load_reference_section7_table_xmls
        reference_entries = original_loader(form)
        self.assertTrue(reference_entries)

        def fake_loader(_):
            first_entry = dict(reference_entries[0])
            first_entry["title"] = None
            return [first_entry]

        app.prospectus_engine._load_reference_section7_table_xmls = fake_loader
        try:
            doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))
        finally:
            app.prospectus_engine._load_reference_section7_table_xmls = original_loader

        self.assertIn("最新公告日期", doc.tables[-1].rows[0].cells[0].text)
        section7_paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        self.assertEqual(0, section7_paragraphs.count("基本信息"))

    def test_prospectus_docx_keeps_stock_subscription_formula(self):
        prospectus_text = app.prospectus_engine.generate(
            self.build_form(EXCHANGE="SSE", MARKET_TYPE="A_SHARE", HAS_STOCK_SUBSCRIPTION=True)
        )
        document_xml = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text))["word/document.xml"]

        self.assertIn("认购份额的计算公式", document_xml)
        self.assertIn("第i只股票在网下股票认购期最后一日的均价", document_xml)
        self.assertIn("除息且送股", document_xml)
        self.assertIn("除息且配股", document_xml)

    def test_prospectus_docx_body_tail_does_not_append_manager_or_custodian_chapters(self):
        prospectus_text = app.prospectus_engine.generate(self.build_form())
        document_xml = self._docx_xml(app.prospectus_engine.build_docx_prospectus(prospectus_text))["word/document.xml"]

        tail_xml = document_xml[document_xml.rfind("\u5907\u67e5\u6587\u4ef6"):]
        self.assertNotIn(">\u57fa\u91d1\u7ba1\u7406\u4eba<", tail_xml)
        self.assertNotIn(">\u57fa\u91d1\u6258\u7ba1\u4eba<", tail_xml)
    def test_prospectus_export_validation_detects_metadata_leak(self):
        report = self.engine.validate_exportable_text(
            "VALUATION_TIMING_CLAUSE：见 06_招募说明书差异条款库.json\n---\n正式正文"
        )
        self.assertFalse(report["ok"])
        self.assertEqual(report["error_type"], "template_metadata_leaked")
        self.assertIn("VALUATION_TIMING_CLAUSE", "\n".join(report["matches"]))

    def test_prospectus_export_validation_allows_unresolved_placeholders(self):
        report = self.engine.validate_exportable_text(
            "测试招募说明书\n【待填写】\n[待填写：律师事务所名称]\n{CUSTODIAN_NAME}\n{{PLACEHOLDER}}"
        )
        self.assertTrue(report["ok"])
        self.assertEqual(report["matches"], [])

    def test_export_prospectus_docx_allows_incomplete_text(self):
        client = app.app.test_client()
        response = client.post("/api/export_prospectus_docx", json=self.build_form())
        self.assertEqual(response.status_code, 200)
        self.assertIn("application/vnd.openxmlformats-officedocument.wordprocessingml.document", response.content_type)
        response.close()

    def test_export_prospectus_txt_allows_incomplete_text(self):
        client = app.app.test_client()
        response = client.post("/api/export_prospectus_txt", json=self.build_form())
        self.assertEqual(response.status_code, 200)
        self.assertIn("text/plain", response.content_type)
        self.assertIn("\u4e00\u3001\u7eea\u8a00", self._consume_response(response, as_text=True))

    def test_export_prospectus_txt_rewrites_product_summary_item_four_and_removes_ben_jijinhetong(self):
        client = app.app.test_client()
        response = client.post("/api/export_prospectus_txt", json=self.build_form())

        self.assertEqual(response.status_code, 200)
        txt_text = self._consume_response(response, as_text=True)
        contract_summary = self.engine._chapter_body_text(txt_text, "二十")
        outside_contract_summary = txt_text.replace(contract_summary, "", 1) if contract_summary else txt_text
        self.assertNotIn("本基金合同", outside_contract_summary)
        self.assertIn("基金产品资料概要是基金招募说明书的摘要文件", txt_text)
        self.assertIn("相关内容不得与基金合同、招募说明书有实质性差异", txt_text)
        self.assertNotIn("基金产品资料概要的内容及编制等具体要求，按照招募说明书相关规定执行", txt_text)

    def test_export_prospectus_docx_rewrites_product_summary_item_four_and_removes_ben_jijinhetong(self):
        client = app.app.test_client()
        response = client.post("/api/export_prospectus_docx", json=self.build_form())

        self.assertEqual(response.status_code, 200)
        docx_bytes = self._consume_response(response)
        document_xml = self._docx_xml(docx_bytes)["word/document.xml"]
        document_text = "\n".join(paragraph.text.strip() for paragraph in self._docx_document(docx_bytes).paragraphs if paragraph.text.strip())
        outside_contract_summary = re.sub(
            r"(?s)二十、\s*基金合同的内容摘要.*?(?=二十一、\s*基金托管协议的内容摘要)",
            "",
            document_text,
            count=1,
        )
        self.assertNotEqual(document_text, outside_contract_summary)
        self.assertNotIn("本基金合同", outside_contract_summary)
        self.assertIn("基金产品资料概要是基金招募说明书的摘要文件", document_xml)
        self.assertIn("相关内容不得与基金合同、招募说明书有实质性差异", document_xml)
        self.assertNotIn("基金产品资料概要的内容及编制等具体要求，按照招募说明书相关规定执行", document_xml)

    def test_export_prospectus_docx_prefers_posted_preview_text_over_regeneration(self):
        client = app.app.test_client()
        form = self.build_form()
        prospectus_text = app.prospectus_engine.generate(form)
        original_generate = app.prospectus_engine.generate

        def fail_generate(_form_data):
            raise AssertionError("route should reuse posted prospectus text")

        app.prospectus_engine.generate = fail_generate
        try:
            response = client.post(
                "/api/export_prospectus_docx",
                json={**form, "PROSPECTUS_TEXT": prospectus_text},
            )
        finally:
            app.prospectus_engine.generate = original_generate

        self.assertEqual(response.status_code, 200)
        document_xml = self._docx_xml(self._consume_response(response))["word/document.xml"]
        self.assertIn("\u516b\u3001\u62d2\u7edd\u6216\u6682\u505c\u7533\u8d2d\u7684\u60c5\u5f62", document_xml)
        self.assertIn("\u4e5d\u3001\u6682\u505c\u8d4e\u56de\u6216\u5ef6\u7f13\u652f\u4ed8\u8d4e\u56de\u5bf9\u4ef7\u7684\u60c5\u5f62", document_xml)
        self.assertIn("\u5341\u3001\u5176\u4ed6\u7533\u8d2d\u8d4e\u56de\u65b9\u5f0f", document_xml)

    def test_export_prospectus_docx_returns_missing_reference_asset_error_when_required_docx_unavailable(self):
        client = app.app.test_client()
        missing_dir = Path(tempfile.mkdtemp())
        missing_variant = missing_dir / "SZSE_CROSS.docx"
        missing_template = missing_dir / "SSE_CROSS.docx"
        original_get_reference = app.prospectus_engine._get_reference_prospectus_docx
        original_get_template = app.prospectus_engine._get_format_template_prospectus_docx

        app.prospectus_engine._get_reference_prospectus_docx = lambda *_args, **_kwargs: missing_variant
        app.prospectus_engine._get_format_template_prospectus_docx = lambda: missing_template
        try:
            response = client.post(
                "/api/export_prospectus_docx",
                json=self.build_form(EXCHANGE="SZSE", MARKET_TYPE="A_SHARE"),
            )
        finally:
            app.prospectus_engine._get_reference_prospectus_docx = original_get_reference
            app.prospectus_engine._get_format_template_prospectus_docx = original_get_template

        self.assertEqual(response.status_code, 400)
        payload = response.get_json()
        self.assertEqual("missing_reference_assets", payload["error_type"])
        self.assertIn("SZSE_CROSS.docx", "\n".join(payload.get("matches", [])))
        self.assertIn("SSE_CROSS.docx", "\n".join(payload.get("matches", [])))

    def test_prospectus_exports_keep_dynamic_chinese_fields_readable(self):
        client = app.app.test_client()
        form = self.build_form(
            FUND_NAME="测试交易型开放式指数证券投资基金",
            CUSTODIAN_NAME="招商银行股份有限公司",
            MIN_SUB_UNIT="123,456份",
        )
        normalized_min_sub_unit = "12.3456万份"

        txt_response = client.post("/api/export_prospectus_txt", json=form)
        self.assertEqual(txt_response.status_code, 200)
        txt_text = self._consume_response(txt_response, as_text=True)
        self.assertIn(form["FUND_NAME"], txt_text)
        self.assertIn(form["CUSTODIAN_NAME"], txt_text)
        self.assertIn(normalized_min_sub_unit, txt_text)
        self.assertNotIn(form["MIN_SUB_UNIT"], txt_text)
        self.assertNotIn("æµ", txt_text)
        self.assertNotIn("æ", txt_text)

        docx_response = client.post("/api/export_prospectus_docx", json=form)
        self.assertEqual(docx_response.status_code, 200)
        document_xml = self._docx_xml(self._consume_response(docx_response))["word/document.xml"]
        self.assertIn(form["FUND_NAME"], document_xml)
        self.assertIn(form["CUSTODIAN_NAME"], document_xml)
        self.assertIn(normalized_min_sub_unit, document_xml)
        self.assertNotIn(form["MIN_SUB_UNIT"], document_xml)
        self.assertNotIn("æµ", document_xml)
        self.assertNotIn("æ", document_xml)

    def test_export_prospectus_txt_prefers_posted_preview_text_over_regeneration(self):
        client = app.app.test_client()
        form = self.build_form()
        prospectus_text = app.prospectus_engine.generate(form)
        original_generate = app.prospectus_engine.generate

        def fail_generate(_form_data):
            raise AssertionError("route should reuse posted prospectus text")

        app.prospectus_engine.generate = fail_generate
        try:
            response = client.post(
                "/api/export_prospectus_txt",
                json={**form, "PROSPECTUS_TEXT": prospectus_text},
            )
        finally:
            app.prospectus_engine.generate = original_generate

        self.assertEqual(response.status_code, 200)
        txt_text = self._consume_response(response, as_text=True)
        self.assertIn("\u516b\u3001\u62d2\u7edd\u6216\u6682\u505c\u7533\u8d2d\u7684\u60c5\u5f62", txt_text)
        self.assertIn("\u4e5d\u3001\u6682\u505c\u8d4e\u56de\u6216\u5ef6\u7f13\u652f\u4ed8\u8d4e\u56de\u5bf9\u4ef7\u7684\u60c5\u5f62", txt_text)
        self.assertIn("\u5341\u3001\u5176\u4ed6\u7533\u8d2d\u8d4e\u56de\u65b9\u5f0f", txt_text)

    def test_index_template_exposes_dedicated_prospectus_sidebar_entry(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")
        self.assertIn('data-panel="prospectus"', html)
        self.assertIn('id="prospectus-panel"', html)
        self.assertIn("招募说明书", html)
        self.assertIn('id="INDEX_DESCRIPTION"', html)
        self.assertIn('id="IMPORTANT_NOTICE_STYLE"', html)
        self.assertIn('id="important-notice-style-group"', html)

    def test_index_template_exposes_prospectus_compare_section(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")
        self.assertIn('id="prospectus-compare-panel"', html)
        self.assertIn('id="btn-use-generated-prospectus"', html)
        self.assertIn('id="btn-compare-prospectus"', html)
        self.assertIn("招募说明书对比", html)

    def test_index_template_exposes_business_text_editor_panel(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")
        self.assertIn('data-panel="business-texts"', html)
        self.assertIn('id="business-texts-panel"', html)
        self.assertIn("/api/business_texts", html)
        self.assertIn("saveBusinessTextOverride", html)
        self.assertIn('id="business-text-scene"', html)
        self.assertIn('id="business-text-product-type"', html)
        self.assertIn('id="business-text-market-type"', html)
        self.assertIn('id="business-text-exchange"', html)
        self.assertIn('id="business-text-processing-card"', html)
        self.assertIn("business-text-list-scroll", html)

    def test_index_template_exposes_form_history_controls(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="form-history-panel"', html)
        self.assertIn('id="form-history-list"', html)
        self.assertIn("FORM_AUTODRAFT_STORAGE_KEY", html)
        self.assertIn("FORM_HISTORY_STORAGE_KEY", html)
        self.assertIn("saveCurrentFormHistory", html)
        self.assertIn("restoreSelectedFormHistory", html)
        self.assertIn("deleteSelectedFormHistory", html)
        self.assertIn("clearFormHistory", html)
        self.assertIn("applyFormData", html)

    def test_collect_form_data_includes_prospectus_notice_fields(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")
        self.assertIn("'INDEX_DESCRIPTION'", html)
        self.assertIn("'IMPORTANT_NOTICE_STYLE'", html)
        self.assertIn("syncImportantNoticeStyleControl", html)
        self.assertIn("deriveIndexDescriptionFromIndexName", html)
        self.assertIn("INDEX_DESCRIPTION: deriveIndexDescriptionFromIndexName", html)
        self.assertIn("PROSPECTUS_TEXT: lastProspectusText || ''", html)

    def test_review_templates_render_locator_paths(self):
        index_html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")
        report_html = Path(app.BASE_DIR / "templates" / "review_report.html").read_text(encoding="utf-8")

        self.assertIn("matched_locator", index_html)
        self.assertIn("contract_locator", index_html)
        self.assertIn("prospectus_locator", index_html)
        self.assertIn("matched_locator", report_html)
        self.assertIn("contract_locator", report_html)
        self.assertIn("prospectus_locator", report_html)

    def test_review_index_template_uses_auto_selected_rules_workbook(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertNotIn('id="review-rules-file"', html)
        self.assertNotIn("fd.append('rules_xlsx'", html)

    def test_review_page_tracks_result_signature_for_report_export(self):
        index_html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn("buildReviewRunSignature", index_html)
        self.assertIn("reviewRunSignatureMatchesCurrentInputs", index_html)

    def test_review_templates_include_problem_only_toggle_controls(self):
        index_html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")
        report_html = Path(app.BASE_DIR / "templates" / "review_report.html").read_text(encoding="utf-8")

        self.assertIn("仅看问题", index_html)
        self.assertIn("显示全部", index_html)
        self.assertIn("仅看问题", report_html)
        self.assertIn("显示全部", report_html)

    def test_review_template_includes_law_firm_revision_workbench(self):
        index_html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")
        review_panel = index_html.split('<div id="review-panel"', 1)[1].split('<!-- /review-panel -->', 1)[0]

        self.assertIn('data-panel="revision-workbench"', index_html)
        nav_html = index_html.split('<nav id="sidebar">', 1)[1].split('</nav>', 1)[0]
        self.assertLess(
            nav_html.index('data-panel="product-summary"'),
            nav_html.index('data-panel="revision-workbench"'),
        )
        self.assertLess(
            nav_html.index('data-panel="revision-workbench"'),
            nav_html.index('data-panel="editor"'),
        )
        self.assertIn('id="revision-workbench-panel"', index_html)
        self.assertNotIn("律所修订痕迹 · 四件套同步影响", index_html)
        self.assertIn("headerBadge.style.display = currentTitle[1] ? 'inline-block' : 'none'", index_html)
        self.assertIn("startRevisionWorkbench", index_html)
        self.assertIn('id="chk-revisions"', index_html)
        self.assertIn('id="rv-revision-detail"', index_html)
        self.assertIn('id="revision-reader-layout"', index_html)
        self.assertIn('id="revision-file-nav"', index_html)
        self.assertIn('class="revision-file-strip"', index_html)
        self.assertIn('class="revision-file-nav-strip"', index_html)
        self.assertIn('id="revision-scope-filter"', index_html)
        self.assertIn('class="revision-scope-strip"', index_html)
        self.assertIn("flex-wrap:nowrap", index_html)
        self.assertIn('id="revision-document-viewer"', index_html)
        self.assertIn('id="revision-inspector-panel"', index_html)
        self.assertIn('id="btn-revision-prev"', index_html)
        self.assertIn('id="btn-revision-next"', index_html)
        self.assertIn('id="revision-left-toggle"', index_html)
        self.assertIn('id="revision-right-toggle"', index_html)
        self.assertIn('id="revision-left-main-toggle"', index_html)
        self.assertIn('id="revision-right-main-toggle"', index_html)
        self.assertIn("revision-reader-header-actions", index_html)
        self.assertIn("rightMainToggle", index_html)
        self.assertIn(".revision-reader-main{grid-column:1", index_html)
        self.assertIn("#revision-inspector-panel{grid-column:2", index_html)
        self.assertIn(".revision-reader-content{display:grid;grid-template-columns:minmax(240px,24%) minmax(0,1fr)", index_html)
        self.assertIn("grid-template-rows:minmax(0,1fr)", index_html)
        self.assertIn("height:auto;overflow:hidden;background:#fbfcfe", index_html)
        self.assertIn("revision-reader-nav-panel", index_html)
        self.assertIn("visibility:hidden;pointer-events:none", index_html)
        self.assertIn("align-items:stretch;height:calc(100vh + 20px);min-height:760px;max-height:none", index_html)
        self.assertIn("min-height:0;height:auto;max-height:none", index_html)
        self.assertIn("overscroll-behavior:contain", index_html)
        self.assertIn('id="rv-revision-body" style="padding:10px;flex:1;min-height:0;overflow:auto;overscroll-behavior:contain"', index_html)
        self.assertIn("data-nav-revision-id", index_html)
        self.assertIn("_scrollRevisionIntoDocumentViewer", index_html)
        self.assertIn("_scrollRevisionNavIntoView", index_html)
        self.assertIn("container.scrollTo", index_html)
        self.assertIn("previousScrollBehavior", index_html)
        self.assertIn('id="revision-workbench-topbar"', index_html)
        self.assertIn('id="revision-session-actions"', index_html)
        self.assertNotIn('id="revision-settings"', index_html)
        self.assertIn('id="revision-display-mode"', index_html)
        self.assertIn('id="revision-document-range"', index_html)
        self.assertIn('value="context">修订上下文', index_html)
        self.assertIn("let _revisionDocumentRange = 'context';", index_html)
        self.assertIn("setRevisionDocumentRange", index_html)
        self.assertIn("_blocksForRevisionDocumentViewer", index_html)
        self.assertIn("当前仅显示修订上下文", index_html)
        self.assertIn('id="revision-type-filter"', index_html)
        self.assertIn('id="revision-status-filter"', index_html)
        self.assertIn('id="revision-search-input"', index_html)
        self.assertIn('id="revision-batch-bar"', index_html)
        self.assertIn('id="revision-select-current"', index_html)
        self.assertIn('id="revision-final-note"', index_html)
        self.assertIn("不覆盖原文件", index_html)
        self.assertIn("生成红线终稿", index_html)
        self.assertIn('id="btn-export-final-redline"', index_html)
        self.assertIn('id="revision-final-modal"', index_html)
        self.assertIn('id="btn-batch-accept"', index_html)
        self.assertIn('id="btn-batch-reject"', index_html)
        self.assertIn('id="btn-batch-pending"', index_html)
        self.assertIn('id="revision-batch-modal"', index_html)
        self.assertIn('id="revision-undo-toast"', index_html)
        self.assertIn('id="revision-operation-drawer"', index_html)
        self.assertIn('id="revision-sync-impact-toggle"', index_html)
        self.assertIn("renderRevisionReader", index_html)
        self.assertIn("selectRevisionDocument", index_html)
        self.assertIn("selectReaderRevision", index_html)
        self.assertIn("selectAdjacentReaderRevision", index_html)
        self.assertIn('id="btn-revision-select-all"', index_html)
        self.assertIn('id="btn-revision-clear-selection"', index_html)
        self.assertIn("revision-nav-select-toolbar", index_html)
        self.assertIn("toggleRevisionPane", index_html)
        self.assertIn("showRevisionBatchConfirm", index_html)
        self.assertIn("applyRevisionBatchAction", index_html)
        self.assertIn("exportRevisionFinalRedline", index_html)
        self.assertIn("undoRevisionOperation", index_html)
        self.assertIn("toggleRevisionOperationDrawer", index_html)
        self.assertIn("exportRevisionOperations", index_html)
        self.assertIn("toggleRevisionSyncImpact", index_html)
        self.assertIn("_initializeRevisionDecisions", index_html)
        self.assertIn("revision-decision-checkbox", index_html)
        self.assertIn("let reviewRevisionEdits = {}", index_html)
        self.assertIn("let reviewRevisionAuthor", index_html)
        self.assertIn("revision-author-input", index_html)
        self.assertIn("review_author", index_html)
        self.assertIn('id="revision-custody-file"', index_html)
        self.assertIn("custody_agreement", index_html)
        self.assertIn("revision-edit-box", index_html)
        self.assertIn("revision.context", index_html)
        self.assertIn("mode: 'paragraph'", index_html)
        self.assertIn("revision-edit-paragraph", index_html)
        self.assertIn("???????", index_html)
        self.assertIn("????", index_html)
        self.assertIn("revision_edits", index_html)
        self.assertIn("???", index_html)
        self.assertIn("默认接受", index_html)
        self.assertIn("reviewRevisionDecisions", index_html)
        self.assertIn("基金合同摘要", index_html)
        self.assertIn("产品资料概要", index_html)
        self.assertNotIn('id="revision-reader-layout"', review_panel)
        self.assertNotIn('id="chk-revisions"', review_panel)
        self.assertNotIn('id="review-product-summary-file"', review_panel)

    def test_review_report_renders_format_issue_source_context_and_text(self):
        client = app.app.test_client()

        response = client.post(
            "/api/review/report",
            json={
                "format_issues": [
                    {
                        "type": "\u6807\u70b9",
                        "severity": "warning",
                        "line": 13,
                        "description": "\u6807\u70b9\u7528\u6cd5\u4e0d\u4e00\u81f4",
                        "text": "\u539f\u6587\u95ee\u9898\u7247\u6bb5",
                        "context": [
                            {"ln": 12, "text": "\u4e0a\u6587\u7b2c\u4e00\u884c", "html": "\u4e0a\u6587\u7b2c\u4e00\u884c", "current": False},
                            {"ln": 13, "text": "\u539f\u6587\u95ee\u9898\u7247\u6bb5", "html": "<mark>\u539f\u6587</mark>\u95ee\u9898\u7247\u6bb5", "current": True},
                            {"ln": 14, "text": "\u4e0b\u6587\u7b2c\u4e00\u884c", "html": "\u4e0b\u6587\u7b2c\u4e00\u884c", "current": False},
                        ],
                        "suggestion": "\u7edf\u4e00\u4fee\u6539",
                    }
                ]
            },
        )
        html = self._consume_response(response, as_text=True)

        self.assertIn("\u95ee\u9898\u539f\u6587", html)
        self.assertIn("\u539f\u6587\u95ee\u9898\u7247\u6bb5", html)
        self.assertIn("<mark>\u539f\u6587</mark>\u95ee\u9898\u7247\u6bb5", html)
        self.assertIn("\u4e0a\u6587\u7b2c\u4e00\u884c", html)
        self.assertIn("\u4e0b\u6587\u7b2c\u4e00\u884c", html)

    def test_review_report_renders_format_issue_locations_when_context_is_absent(self):
        client = app.app.test_client()

        response = client.post(
            "/api/review/report",
            json={
                "format_issues": [
                    {
                        "type": "\u4e00\u81f4\u6027",
                        "severity": "warning",
                        "description": "\u540c\u4e00\u95ee\u9898\u51fa\u73b0 5 \u5904",
                        "locations": [
                            {"ln": 42, "text": "\u7b2c\u4e00\u5904\u539f\u6587", "html": "<mark>\u7b2c\u4e00\u5904</mark>\u539f\u6587"},
                            {"ln": 73, "text": "\u7b2c\u4e8c\u5904\u539f\u6587", "html": "<mark>\u7b2c\u4e8c\u5904</mark>\u539f\u6587"},
                        ],
                        "suggestion": "\u8bf7\u9010\u5904\u6838\u5bf9",
                    }
                ]
            },
        )
        html = self._consume_response(response, as_text=True)

        self.assertIn("<mark>\u7b2c\u4e00\u5904</mark>\u539f\u6587", html)
        self.assertIn("<mark>\u7b2c\u4e8c\u5904</mark>\u539f\u6587", html)

    def test_format_check_ignores_numbered_items_with_single_right_parenthesis_style(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "contract_text": "\n".join(
                [
                    "（12）本基金参与股指期货交易，应当符合下列投资限制：",
                    "1）本基金在任何交易日日终，持有的买入股指期货合约价值，不得超过基金资产净值的10%；",
                    "2）本基金在任何交易日日终，持有的买入股指期货合约价值与有价证券市值之和，不得超过基金资产净值的100%；",
                    "3）本基金在任何交易日日终，持有的卖出股指期货合约价值不得超过基金持有的股票总市值的20%；",
                    "（13）本基金参与股票期权交易，应当符合下列投资限制：",
                ]
            )
        }
        try:
            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        bracket_issue_types = [issue["type"] for issue in payload["issues"] if issue.get("type") == "bracket_mismatch"]
        self.assertEqual([], bracket_issue_types)

    def test_format_check_still_reports_true_unmatched_fullwidth_parenthesis(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "contract_text": "\n".join(
                [
                    "第一部分 前言",
                    "这是一个真正未配对的示例）需要继续提示。",
                ]
            )
        }
        try:
            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        bracket_issue = next(issue for issue in payload["issues"] if issue.get("type") == "bracket_mismatch")
        self.assertIn("符号不配对", bracket_issue["description"])
        self.assertEqual(2, bracket_issue["line"])

    def test_format_check_reports_digit_parenthesis_sequence_gap(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "contract_text": "\n".join(
                [
                    "一、投资限制",
                    "1）第一项",
                    "3）第三项",
                ]
            )
        }
        try:
            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        heading_issue = next(issue for issue in payload["issues"] if issue.get("type") == "heading_sequence")
        self.assertIn("数字序号", heading_issue["description"])
        self.assertEqual(3, heading_issue["line"])

    def test_format_check_accepts_digit_sequence_followed_by_book_title_mark(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "contract_text": "\n".join(
                [
                    "二、释义",
                    "16、《流动性风险管理规定》：指中国证监会2017年8月31日颁布、同年10月1日实施的《公开募集开放式证券投资基金流动性风险管理规定》及颁布机关对其不时做出的修订",
                    "17、《指数基金指引》：指中国证监会2021年1月22日颁布、同年2月1日实施的《公开募集证券投资基金运作指引第3号——指数基金指引》及颁布机关对其不时做出的修订",
                    "18、中国证监会：指中国证券监督管理委员会",
                    "19、银行业监督管理机构：指中国人民银行和/或国家金融监督管理总局",
                    "20、基金合同当事人：指受基金合同约束，根据基金合同享有权利并承担义务的法律主体，包括基金管理人、基金托管人和基金份额持有人",
                ]
            )
        }
        try:
            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        heading_issues = [issue for issue in payload["issues"] if issue.get("type") == "heading_sequence"]
        self.assertEqual([], heading_issues)

    def test_format_check_accepts_digit_sequence_followed_by_ascii_letters(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "contract_text": "\n".join(
                [
                    "二、释义",
                    "26、交易型开放式证券投资基金：指《中国证券登记结算有限责任公司关于交易所交易型开放式证券投资基金登记结算业务实施细则》定义的交易型开放式证券投资基金，简称ETF",
                    "27、ETF联接基金：指将其绝大部分基金财产投资于跟踪同一标的指数的ETF，紧密跟踪标的指数表现，追求跟踪偏离度和跟踪误差最小化，采用开放式运作方式的基金",
                    "28、基金销售业务：指基金管理人或销售机构宣传推介基金，发售基金份额，办理基金份额的申购、赎回等业务",
                ]
            )
        }
        try:
            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        heading_issues = [issue for issue in payload["issues"] if issue.get("type") == "heading_sequence"]
        self.assertEqual([], heading_issues)

    def test_format_check_tracks_numeric_sequence_by_style_level(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "contract_text": "\n".join(
                [
                    "一、总则",
                    "1、第一项",
                    "2、第二项",
                    "1）说明事项",
                    "3、第三项",
                ]
            )
        }
        try:
            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        heading_issues = [issue for issue in payload["issues"] if issue.get("type") == "heading_sequence"]
        self.assertEqual([], heading_issues)

    def test_format_check_reports_single_extra_space_between_chinese_characters(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "contract_text": "\n".join(
                [
                    "第一部分 前言",
                    "基金 管理人应当勤勉尽责。",
                ]
            )
        }
        try:
            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        extra_space_issue = next(issue for issue in payload["issues"] if issue.get("type") == "extra_space")
        self.assertEqual(2, extra_space_issue["line"])
        self.assertIn("空格", extra_space_issue["description"])

    def test_format_check_reports_blank_pages_from_uploaded_contract_docx(self):
        client = app.app.test_client()
        contract_buffer = self._build_review_docx_buffer(
            title_lines=["南方测试成长交易型开放式指数证券投资基金基金合同"],
            body_lines=["第一部分 前言", "前言正文"],
            extra_page_breaks=2,
        )
        try:
            upload_response = client.post(
                "/api/review/upload",
                data={"contract": (contract_buffer, "南方测试成长交易型开放式指数证券投资基金基金合同.docx")},
                content_type="multipart/form-data",
            )
            self.assertEqual(upload_response.status_code, 200)
            upload_response.close()

            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            app._review_store.pop("data", None)

        blank_page_issue = next(issue for issue in payload["issues"] if issue.get("type") == "blank_page")
        self.assertIn("空白页", blank_page_issue["description"])
        self.assertIn("第2页", blank_page_issue["locations"][0]["text"])

    def test_format_check_reports_header_title_mismatch_from_uploaded_contract_docx(self):
        client = app.app.test_client()
        contract_buffer = self._build_review_docx_buffer(
            title_lines=["南方测试成长交易型开放式指数证券投资基金基金合同"],
            body_lines=["第一部分 前言", "前言正文"],
            header_text="错误的页眉标题",
        )
        try:
            upload_response = client.post(
                "/api/review/upload",
                data={"contract": (contract_buffer, "南方测试成长交易型开放式指数证券投资基金基金合同.docx")},
                content_type="multipart/form-data",
            )
            self.assertEqual(upload_response.status_code, 200)
            upload_response.close()

            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            app._review_store.pop("data", None)

        header_issue = next(issue for issue in payload["issues"] if issue.get("type") == "header_title_mismatch")
        self.assertIn("页眉", header_issue["description"])
        self.assertIn("错误的页眉标题", header_issue["locations"][0]["text"])

    def test_format_check_ignores_header_suffix_noise_when_matching_uploaded_contract_docx(self):
        client = app.app.test_client()
        contract_buffer = self._build_review_docx_buffer(
            title_lines=["南方测试成长交易型开放式指数证券投资基金基金合同"],
            body_lines=["第一部分 前言", "前言正文"],
            header_text="南方测试成长交易型开放式指数证券投资基金基金合同",
        )
        try:
            upload_response = client.post(
                "/api/review/upload",
                data={"contract": (contract_buffer, "南方测试成长交易型开放式指数证券投资基金基金合同（草案）-0323.docx")},
                content_type="multipart/form-data",
            )
            self.assertEqual(upload_response.status_code, 200)
            upload_response.close()

            response = client.post("/api/review/format_check", json={})
            payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        finally:
            app._review_store.pop("data", None)

        header_issue_types = [issue["type"] for issue in payload["issues"] if issue.get("type") == "header_title_mismatch"]
        self.assertEqual([], header_issue_types)

    def test_review_index_template_renders_format_issue_locations(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn("it.locations", html)
        self.assertIn("_renderLocs(it.locations)", html)

    def test_compare_template_supports_generated_and_uploaded_sources_on_both_sides(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="file-input-a"', html)
        self.assertIn('onclick="loadGeneratedAsB()"', html)
        self.assertIn('id="prospectus-file-input-a"', html)
        self.assertIn('onclick="loadGeneratedProspectusAsB()"', html)

    def test_prospectus_template_includes_custodian_summary_upload_controls(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="prospectus-custodian-summary-file"', html)
        self.assertIn('id="prospectus-custodian-summary-status"', html)
        self.assertIn('onclick="triggerProspectusCustodianSummaryUpload()"', html)
        self.assertIn('onclick="clearProspectusCustodianSummary()"', html)

    def test_compare_template_uses_clean_copy_and_single_helper_definitions(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertGreaterEqual(html.count("点击选择或拖放文件"), 4)
        self.assertGreaterEqual(html.count("或直接粘贴文本："), 4)
        self.assertEqual(1, html.count("function loadGeneratedAsA()"))
        self.assertEqual(1, html.count("function loadGeneratedAsB()"))
        self.assertEqual(1, html.count("function loadGeneratedProspectusAsA()"))
        self.assertEqual(1, html.count("function loadGeneratedProspectusAsB()"))
        self.assertNotIn("鐐瑰嚮閫夋嫨鎴栨嫋鏀炬枃浠", html)
        self.assertNotIn("馃摠", html)

    def test_compare_template_supports_expanding_skipped_equal_lines(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn("function toggleDiffSkipRows(", html)
        self.assertIn("function appendDiffRows(", html)
        self.assertNotIn('showToast(\'展开功能需要重新加载完整文本对比，请取消勾选"仅显示差异行"查看全部内容\')', html)
        self.assertNotIn('showToast(\'当前版本请取消“仅显示差异行”查看完整上下文\')', html)

    def test_generated_summary_check_reuses_review_summary_card_renderer(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn("detailItems.map(item => _renderSummaryReviewItemCard(item)).join('')", html)
        self.assertGreaterEqual(html.count("_renderSummaryReviewItemCard("), 2)

    def test_compare_api_skip_hunks_include_hidden_equal_lines_for_expand(self):
        client = app.app.test_client()
        left_lines = [f"相同前文{i}" for i in range(1, 16)] + ["左侧差异"] + [f"相同后文{i}" for i in range(1, 16)]
        right_lines = [f"相同前文{i}" for i in range(1, 16)] + ["右侧差异"] + [f"相同后文{i}" for i in range(1, 16)]

        response = client.post(
            "/api/compare",
            json={
                "text1": "\n".join(left_lines),
                "text2": "\n".join(right_lines),
                "label1": "A",
                "label2": "B",
            },
        )
        payload = response.get_json()

        self.assertEqual(200, response.status_code)
        skip_hunks = [hunk for hunk in payload["hunks"] if hunk.get("type") == "skip"]
        self.assertTrue(skip_hunks)
        for hunk in skip_hunks:
            self.assertEqual(hunk["count"], len(hunk.get("lines") or []))
            self.assertTrue(all(line.get("tag") == "equal" for line in hunk.get("lines") or []))

    def test_review_diff_panels_default_to_expanded_complete_diff(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn('class="sc-accordion-item sc-open"', html)
        self.assertIn('class="sc-accordion-header sc-open"', html)
        self.assertIn('class="sc-accordion-body sc-open"', html)
        self.assertIn('<details style="margin:8px 0 10px"><summary style="cursor:pointer;font-size:12px;color:var(--accent);font-weight:600">查看完整原文对照</summary>', html)
        self.assertIn('<details style="margin:8px 0 10px"><summary style="cursor:pointer;font-size:12px;color:var(--accent);font-weight:600">查看章节原文</summary>', html)
        self.assertIn('<details open><summary style="cursor:pointer;font-size:12px;color:var(--accent);font-weight:600">查看完整差异', html)
        self.assertIn('<details open><summary style="cursor:pointer;font-size:12px;color:var(--accent);font-weight:600">查看差异原文', html)

    def test_review_report_prefers_stored_fund_name_and_marks_manual_review_rows(self):
        client = app.app.test_client()
        original_data = app._review_store.get("data")
        app._review_store["data"] = {
            "fund_name": "南方港股通测试基金",
            "contract_filename": "合同.docx",
            "prospectus_filename": "招募说明书.docx",
        }

        try:
            response = client.post(
                "/api/review/report",
                json={
                    "summary_check": [
                        {
                            "summary_heading": "二、基金份额持有人大会召集、议事及表决的程序和规则",
                            "matched_section": None,
                            "matched_locator": "",
                            "locator_matched": False,
                            "content_match_score": 6.1,
                            "similarity": 0.0,
                            "changed_lines": 125,
                            "total_lines": 125,
                            "rule_locator": "",
                            "rule_level": "",
                        }
                    ]
                },
            )
            html = self._consume_response(response, as_text=True)
        finally:
            if original_data is None:
                app._review_store.pop("data", None)
            else:
                app._review_store["data"] = original_data

        self.assertIn("基金: 南方港股通测试基金", html)
        self.assertIn("待人工复核", html)

    def test_extract_text_api_reads_docx_tables_in_order(self):
        client = app.app.test_client()

        from docx import Document as DocxDocument
        import io

        doc = DocxDocument()
        doc.add_paragraph("\u7b2c\u4e00\u6bb5")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        doc.add_paragraph("\u7b2c\u4e8c\u6bb5")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = client.post(
            "/api/extract_text",
            data={"file": (buffer, "sample.docx")},
            content_type="multipart/form-data",
        )

        self.assertEqual(response.status_code, 200)
        payload = self._consume_response(response, parser=lambda resp: resp.get_json())
        text = payload["text"]
        self.assertIn("\u7b2c\u4e00\u6bb5", text)
        self.assertIn("|A|B|", text)
        self.assertIn("|1|2|", text)
        self.assertIn("\u7b2c\u4e8c\u6bb5", text)
        self.assertLess(text.index("\u7b2c\u4e00\u6bb5"), text.index("|A|B|"))
        self.assertLess(text.index("|1|2|"), text.index("\u7b2c\u4e8c\u6bb5"))

    def test_clean_prospectus_docx_does_not_include_placeholder_markers(self):
        clean_text = "\n".join([
            "测试交易型开放式指数证券投资基金招募说明书",
            "基金管理人：南方基金管理股份有限公司",
            "基金托管人：招商银行股份有限公司",
            "2026年1月1日",
            "目录",
            "一、绪言",
            "一、绪言",
            "本招募说明书用于测试导出。",
        ])
        xml_map = self._docx_xml(app.prospectus_engine.build_docx_prospectus(clean_text))
        document_xml = xml_map["word/document.xml"]

        self.assertNotIn("VALUATION_TIMING_CLAUSE", document_xml)
        self.assertNotIn("【待填写】", document_xml)
        self.assertNotIn("[待填写", document_xml)
        self.assertNotIn("{CUSTODIAN_NAME}", document_xml)

    def test_build_docx_prospectus_keeps_body_when_toc_has_page_numbers(self):
        clean_text = "\n".join([
            "\u6d4b\u8bd5\u4ea4\u6613\u578b\u5f00\u653e\u5f0f\u6307\u6570\u8bc1\u5238\u6295\u8d44\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66",
            "\u57fa\u91d1\u7ba1\u7406\u4eba\uff1a\u5357\u65b9\u57fa\u91d1\u7ba1\u7406\u80a1\u4efd\u6709\u9650\u516c\u53f8",
            "\u57fa\u91d1\u6258\u7ba1\u4eba\uff1a\u62db\u5546\u94f6\u884c\u80a1\u4efd\u6709\u9650\u516c\u53f8",
            "\u91cd\u8981\u63d0\u793a",
            "\u8fd9\u662f\u91cd\u8981\u63d0\u793a\u6bb5\u843d\u3002",
            "\u76ee\u5f55",
            "\u4e00\u3001\u7eea\u8a00\t4",
            "\u4e8c\u3001\u91ca\u4e49\t5",
            "\u7eea\u8a00",
            "\u8fd9\u662f\u7eea\u8a00\u6b63\u6587\u7b2c\u4e00\u6bb5\u3002",
            "\u91ca\u4e49",
            "\u8fd9\u662f\u91ca\u4e49\u6b63\u6587\u7b2c\u4e00\u6bb5\u3002",
        ])

        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(clean_text))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        self.assertIn("\u8fd9\u662f\u7eea\u8a00\u6b63\u6587\u7b2c\u4e00\u6bb5\u3002", lines)
        self.assertIn("\u8fd9\u662f\u91ca\u4e49\u6b63\u6587\u7b2c\u4e00\u6bb5\u3002", lines)

    def test_build_docx_prospectus_preserves_generated_body_lines_after_normalization(self):
        form = self.build_form()
        prospectus_text = app.prospectus_engine.generate(form)
        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(prospectus_text, form))

        def relaxed(text):
            normalized = unicodedata.normalize("NFKC", text or "")
            return "".join(
                ch
                for ch in normalized
                if not unicodedata.category(ch).startswith(("P", "Z", "C"))
            )

        stock_formula_text = app.ProspectusEngine._stock_subscription_formula_text()
        cash_formula_text = app.ProspectusEngine._cash_substitution_ratio_formula_text()
        source_lines = [
            line.strip()
            for line in prospectus_text.splitlines()
            if line.strip()
            and not line.strip().startswith("|")
            and line.strip() != stock_formula_text
            and line.strip() != cash_formula_text
        ]
        exported_lines = {
            relaxed(paragraph.text.strip())
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        }
        missing = [line for line in source_lines if relaxed(line) not in exported_lines]

        self.assertEqual([], missing[:10], f"Missing normalized lines: {missing[:10]}")

    def test_build_docx_prospectus_uses_reference_prospectus_page_and_header_format(self):
        clean_text = "\n".join([
            "测试交易型开放式指数证券投资基金招募说明书",
            "基金管理人：南方基金管理股份有限公司",
            "基金托管人：招商银行股份有限公司",
            "重要提示",
            "这是重要提示正文段落。",
            "目录",
            "一、绪言",
            "二、释义",
            "一、绪言",
            "这是绪言正文第一段。",
            "二、释义",
            "这是释义正文第一段。",
        ])

        docx_bytes = app.prospectus_engine.build_docx_prospectus(clean_text)
        doc = self._docx_document(docx_bytes)
        xml_map = self._docx_xml(docx_bytes)
        header_xml = "".join(value for key, value in xml_map.items() if key.startswith("word/header"))

        self.assertEqual(2, len(doc.sections))
        self.assertIn('w:defaultTabStop w:val="420"', xml_map["word/settings.xml"])
        self.assertIn("《测试交易型开放式指数证券投资基金招募说明书（草案）》", header_xml)

    def test_build_docx_prospectus_uses_reference_prospectus_cover_body_and_notice_styles(self):
        clean_text = "\n".join([
            "测试交易型开放式指数证券投资基金招募说明书",
            "基金管理人：南方基金管理股份有限公司",
            "基金托管人：招商银行股份有限公司",
            "重要提示",
            "这是重要提示正文段落。",
            "目录",
            "一、绪言",
            "一、绪言",
            "这是绪言正文第一段。",
        ])

        doc = self._docx_document(app.prospectus_engine.build_docx_prospectus(clean_text))
        title = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip().endswith("招募说明书"))
        manager = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith("基金管理人："))
        notice_title = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "重要提示")
        notice_body = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "这是重要提示正文段落。")
        body = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "这是绪言正文第一段。")

        self.assertEqual("44", title.runs[0]._r.get_or_add_rPr().find(qn("w:sz")).get(qn("w:val")))
        self.assertEqual("30", manager.runs[0]._r.get_or_add_rPr().find(qn("w:sz")).get(qn("w:val")))
        self.assertTrue(manager.runs[0].bold)
        self.assertEqual("21", notice_title.runs[0]._r.get_or_add_rPr().find(qn("w:sz")).get(qn("w:val")))
        self.assertTrue(notice_title.runs[0].bold)

        for paragraph in (notice_body, body):
            ppr = paragraph._p.get_or_add_pPr()
            spacing = ppr.find(qn("w:spacing"))
            indent = ppr.find(qn("w:ind"))
            snap_to_grid = ppr.find(qn("w:snapToGrid"))
            run_properties = paragraph.runs[0]._r.get_or_add_rPr()

            self.assertEqual("360", spacing.get(qn("w:line")))
            self.assertEqual("auto", spacing.get(qn("w:lineRule")))
            self.assertEqual("420", indent.get(qn("w:firstLine")))
            self.assertEqual("200", indent.get(qn("w:firstLineChars")))
            self.assertEqual("0", snap_to_grid.get(qn("w:val")))
            self.assertEqual("21", run_properties.find(qn("w:sz")).get(qn("w:val")))

    def test_extract_text_round_trips_exported_prospectus_text_for_compare(self):
        client = app.app.test_client()
        form = self.build_form(
            EXCHANGE="SSE",
            MARKET_TYPE="A_SHARE",
            IMPORTANT_NOTICE_STYLE="DIVIDEND_QUALITY",
            HAS_STOCK_SUBSCRIPTION=True,
        )
        prospectus_text = app.prospectus_engine.generate(form)

        export_response = client.post(
            "/api/export_prospectus_docx",
            json={**form, "PROSPECTUS_TEXT": prospectus_text},
        )
        self.assertEqual(export_response.status_code, 200)
        docx_bytes = self._consume_response(export_response)

        upload = io.BytesIO(docx_bytes)
        upload.name = "prospectus.docx"
        extract_response = client.post(
            "/api/extract_text",
            data={"file": (upload, "prospectus.docx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(extract_response.status_code, 200)
        payload = self._consume_response(extract_response, parser=lambda resp: resp.get_json())

        expected_lines = [row["line"] for row in app._collect_nonblank_compare_lines(prospectus_text)]
        extracted_lines = [row["line"] for row in app._collect_nonblank_compare_lines(payload["text"])]
        self.assertEqual(expected_lines, extracted_lines)


class CustodianDirectoryTests(unittest.TestCase):
    EXPECTED_DIRECTORY_NAMES = [
        "中国工商银行股份有限公司",
        "中国农业银行股份有限公司",
        "中国银行股份有限公司",
        "中国建设银行股份有限公司",
        "交通银行股份有限公司",
        "华夏银行股份有限公司",
        "中国光大银行股份有限公司",
        "招商银行股份有限公司",
        "中信银行股份有限公司",
        "中国民生银行股份有限公司",
        "兴业银行股份有限公司",
        "上海浦东发展银行股份有限公司",
        "北京银行股份有限公司",
        "平安银行股份有限公司",
        "广发银行股份有限公司",
        "中国邮政储蓄银行股份有限公司",
        "上海银行股份有限公司",
        "国信证券股份有限公司",
        "招商证券股份有限公司",
        "广发证券股份有限公司",
        "国泰海通证券股份有限公司",
        "江苏银行股份有限公司",
        "中国银河证券股份有限公司",
        "中信证券股份有限公司",
        "兴业证券股份有限公司",
        "中信建投证券股份有限公司",
        "中泰证券股份有限公司",
        "国投证券股份有限公司",
        "东方证券股份有限公司",
        "申万宏源证券有限公司",
        "长城证券股份有限公司",
        "长江证券股份有限公司",
        "浙商证券股份有限公司",
        "东方财富证券股份有限公司",
    ]

    @classmethod
    def setUpClass(cls):
        cls.client = app.app.test_client()
        cls.engine = app.ProspectusEngine()

    def test_api_custodians_returns_official_directory_with_new_payment_fields(self):
        response = self.client.get("/api/custodians")
        self.assertEqual(200, response.status_code)
        payload = response.get_json()
        response.close()

        self.assertEqual(34, len(payload))
        self.assertEqual(self.EXPECTED_DIRECTORY_NAMES, list(payload.keys()))

        zhongtai = payload["中泰证券股份有限公司"]
        self.assertEqual(27, zhongtai["directory_order"])
        self.assertEqual("山东", zhongtai["jurisdiction"])
        self.assertTrue(zhongtai["address"])
        self.assertTrue(zhongtai["legal_rep"])
        self.assertEqual("", zhongtai["mgmt_fee_payment_method"])
        self.assertEqual("", zhongtai["custody_fee_payment_method"])
        self.assertIn("协商一致", zhongtai["mgmt_fee_payment_text"])
        self.assertIn("协商一致", zhongtai["custody_fee_payment_text"])

        cmb = payload["招商银行股份有限公司"]
        self.assertEqual(8, cmb["directory_order"])
        self.assertEqual("深圳", cmb["jurisdiction"])
        self.assertEqual("CONSULT", cmb["payment_method"])
        self.assertIn("协商一致", cmb["mgmt_fee_payment_text"])
        self.assertIn("协商一致", cmb["custody_fee_payment_text"])

    def test_external_supplemented_custodians_expose_contract_brief_fields(self):
        response = self.client.get("/api/custodians")
        self.assertEqual(200, response.status_code)
        payload = response.get_json()
        response.close()

        supplemented_names = [
            "中国工商银行股份有限公司",
            "中国农业银行股份有限公司",
            "中国银行股份有限公司",
            "中国建设银行股份有限公司",
            "华夏银行股份有限公司",
            "中国光大银行股份有限公司",
            "中国民生银行股份有限公司",
            "北京银行股份有限公司",
            "平安银行股份有限公司",
            "广发银行股份有限公司",
            "中国邮政储蓄银行股份有限公司",
            "上海银行股份有限公司",
            "国信证券股份有限公司",
            "招商证券股份有限公司",
            "广发证券股份有限公司",
            "中国银河证券股份有限公司",
            "中信银行股份有限公司",
            "兴业银行股份有限公司",
            "上海浦东发展银行股份有限公司",
            "江苏银行股份有限公司",
            "国投证券股份有限公司",
            "兴业证券股份有限公司",
            "中泰证券股份有限公司",
            "东方证券股份有限公司",
            "申万宏源证券有限公司",
            "长城证券股份有限公司",
            "长江证券股份有限公司",
            "浙商证券股份有限公司",
            "东方财富证券股份有限公司",
        ]

        for name in supplemented_names:
            info = payload[name]
            self.assertTrue(info["address"], name)
            self.assertTrue(info["legal_rep"], name)
            self.assertTrue(info["established"], name)
            self.assertTrue(info["approval_no"], name)
            self.assertTrue(info["registered_capital"], name)
            self.assertTrue(info["custody_license"], name)

        expected_payment_keywords = {
            "中国工商银行股份有限公司": "协商一致",
            "中国农业银行股份有限公司": "协商一致",
            "中国银行股份有限公司": "划付指令",
            "中国建设银行股份有限公司": "自动",
            "华夏银行股份有限公司": "协商一致",
            "中国光大银行股份有限公司": "协商一致",
            "中国民生银行股份有限公司": "协商一致",
            "北京银行股份有限公司": "协商一致",
            "平安银行股份有限公司": "自动付费授权函",
            "广发银行股份有限公司": "划款指令",
            "中国邮政储蓄银行股份有限公司": "无需再出具资金划拨指令",
            "上海银行股份有限公司": "次月前3个工作日内",
            "国信证券股份有限公司": "资金划拨指令",
            "招商证券股份有限公司": "划付指令",
            "广发证券股份有限公司": "协商一致",
            "中国银河证券股份有限公司": "指定的账户路径",
            "中信银行股份有限公司": "协商一致",
            "兴业银行股份有限公司": "协商一致",
            "上海浦东发展银行股份有限公司": "协商一致",
            "江苏银行股份有限公司": "自动于次月首日起5个工作日内",
            "国投证券股份有限公司": "无需再出具资金划拨指令",
            "兴业证券股份有限公司": "划付指令",
            "中泰证券股份有限公司": "协商一致",
            "东方证券股份有限公司": "资金划拨指令",
            "申万宏源证券有限公司": "指定的账户路径",
            "长城证券股份有限公司": "划款指令",
            "长江证券股份有限公司": "基金财产中一次性",
            "浙商证券股份有限公司": "资金划拨指令",
            "东方财富证券股份有限公司": "无需再出具资金划拨指令",
        }

        for name, expected_keyword in expected_payment_keywords.items():
            info = payload[name]
            self.assertIn(expected_keyword, info["mgmt_fee_payment_text"], name)
            self.assertIn(expected_keyword, info["custody_fee_payment_text"], name)
            self.assertTrue(info["payment_method_source"], name)

    def test_inject_clause_texts_supports_verbatim_payment_clauses_and_defaults_legacy_enums(self):
        blank_result = self.engine._inject_clause_texts(
            {
                "MGMT_FEE_PAYMENT_METHOD": "",
                "CUSTODY_FEE_PAYMENT_METHOD": "",
            }
        )
        self.assertEqual("", blank_result["MGMT_FEE_PAYMENT_METHOD"])
        self.assertEqual("", blank_result["CUSTODY_FEE_PAYMENT_METHOD"])

        verbatim_result = self.engine._inject_clause_texts(
            {
                "MGMT_FEE_PAYMENT_METHOD": "经基金管理人与基金托管人核对一致后，基金托管人于次月前5个工作日内支付给基金管理人。",
                "CUSTODY_FEE_PAYMENT_METHOD": "经基金管理人与基金托管人核对一致后，基金托管人于次月前5个工作日内支取。",
            }
        )
        self.assertEqual(
            "经基金管理人与基金托管人核对一致后，基金托管人于次月前5个工作日内支付给基金管理人。",
            verbatim_result["MGMT_FEE_PAYMENT_METHOD"],
        )
        self.assertEqual(
            "经基金管理人与基金托管人核对一致后，基金托管人于次月前5个工作日内支取。",
            verbatim_result["CUSTODY_FEE_PAYMENT_METHOD"],
        )

        default_result = self.engine._inject_clause_texts({})
        self.assertIn("协商一致", default_result["MGMT_FEE_PAYMENT_METHOD"])
        self.assertIn("协商一致", default_result["CUSTODY_FEE_PAYMENT_METHOD"])

    def test_index_template_builds_custodian_quick_options_from_api_and_shows_pending_notice(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn("renderCustodianQuickOptions()", html)
        self.assertIn("该机构仅完成名录收录，合同简况及支付方式待补录", html)
        self.assertIn("<textarea id=\"MGMT_FEE_PAYMENT_METHOD\"", html)
        self.assertIn("<textarea id=\"CUSTODY_FEE_PAYMENT_METHOD\"", html)
        self.assertIn("<input id=\"INDEX_WEBSITE\"", html)
        self.assertIn("autoFillIndexLinkedFields", html)
        self.assertIn("INDEX_COMPILER','INDEX_WEBSITE'", html)
        self.assertIn("id=\"CSRC_APPROVAL_NO\"", html)
        self.assertIn("id=\"FUND_MANAGER_NAME\"", html)
        self.assertIn("id=\"FUND_MANAGER_SEX\"", html)
        self.assertIn("id=\"FUND_MANAGER_BIO\"", html)
        self.assertIn("id=\"CUSTODIAN_INTRO\"", html)
        self.assertIn("id=\"ACCOUNTANT\"", html)
        self.assertNotIn('<option value="招商银行股份有限公司">招商银行</option>', html)

    def test_index_template_exposes_accountant_firm_dropdown_profiles_and_manual_edit(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="ACCOUNTANT_CHOICE"', html)
        self.assertIn("ACCOUNTANT_PROFILES", html)
        self.assertIn("onAccountantChoice", html)
        self.assertIn("德勤华永会计师事务所（特殊普通合伙）", html)
        self.assertIn("安永华明会计师事务所（特殊普通合伙）", html)
        self.assertIn("容诚会计师事务所（特殊普通合伙）", html)
        self.assertIn("经办注册会计师：汪芳、陈思雨", html)
        self.assertIn("签章注册会计师：高鹤、邓雯", html)
        self.assertIn("经办注册会计师：陈熹、成磊", html)
        self.assertIn('value="OTHER"', html)
        self.assertIn("手动编辑", html)

    def test_index_template_apply_auto_derived_value_updates_restored_auto_values(self):
        html_path = Path(app.BASE_DIR / "templates" / "index.html")
        script = f"""
const fs = require('fs');
const html = fs.readFileSync({json.dumps(str(html_path))}, 'utf8');
const start = html.indexOf('function applyAutoDerivedValue(id, value, options = {{}}) {{');
const end = html.indexOf('function autoFillIndexLinkedFields()', start);
if (start === -1 || end === -1) {{
  throw new Error('applyAutoDerivedValue not found');
}}
const fnSource = html.slice(start, end);
const input = {{ value: '南方上证科创板人工智能ETF', dataset: {{}} }};
global.document = {{ getElementById() {{ return input; }} }};
eval(fnSource);
applyAutoDerivedValue('FUND_SHORT_NAME', '南方上证科创板人工智能ETF');
applyAutoDerivedValue('FUND_SHORT_NAME', '南方上证科创板机器人ETF');
if (input.value !== '南方上证科创板机器人ETF') {{
  throw new Error(`unexpected value: ${{input.value}}`);
}}
"""
        subprocess.run(["node", "-e", script], check=True, capture_output=True, text=True)

    def test_index_template_auto_fill_index_linked_fields_refreshes_fund_names(self):
        html_path = Path(app.BASE_DIR / "templates" / "index.html")
        script = f"""
const fs = require('fs');
const html = fs.readFileSync({json.dumps(str(html_path))}, 'utf8');
const start = html.indexOf('function inferIndexProvider(indexName) {{');
const end = html.indexOf('// ═══════════════════════════════════════════════════════════════', start);
if (start === -1 || end === -1) {{
  throw new Error('index auto-fill functions not found');
}}
const fnSource = html.slice(start, end);
const nodes = {{
  INDEX_NAME: {{ value: '中证机器人指数', dataset: {{}} }},
  FUND_NAME: {{ value: '旧基金全名', dataset: {{}} }},
  FUND_SHORT_NAME: {{ value: '旧基金简称', dataset: {{}} }},
  INDEX_COMPILER: {{ value: '', dataset: {{}} }},
  INDEX_WEBSITE: {{ value: '', dataset: {{}} }},
  INDEX_DESCRIPTION: {{ value: '', dataset: {{}} }},
}};
global.document = {{
  getElementById(id) {{
    return nodes[id];
  }}
}};
eval(fnSource);
autoFillIndexLinkedFields();
if (nodes.FUND_NAME.value !== '南方中证机器人交易型开放式指数证券投资基金') {{
  throw new Error(`unexpected fund name: ${{nodes.FUND_NAME.value}}`);
}}
if (nodes.FUND_SHORT_NAME.value !== '南方中证机器人ETF') {{
  throw new Error(`unexpected fund short name: ${{nodes.FUND_SHORT_NAME.value}}`);
}}
"""
        subprocess.run(["node", "-e", script], check=True, capture_output=True, text=True)

    def test_index_template_exposes_readiness_and_demo_fill_controls(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="btn-fill-demo"', html)
        self.assertIn('id="contract-readiness-panel"', html)
        self.assertIn('id="contract-readiness-actions"', html)
        self.assertIn('id="contract-generate-hint"', html)
        self.assertIn('id="prospectus-readiness-panel"', html)
        self.assertIn('id="prospectus-readiness-actions"', html)
        self.assertIn('id="prospectus-export-hint"', html)
        self.assertIn("fillDemoData()", html)
        self.assertIn("renderReadinessPanels()", html)

    def test_index_template_includes_mobile_responsive_breakpoints_for_sidebar_and_forms(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn("@media (max-width: 960px)", html)
        self.assertIn("@media (max-width: 720px)", html)
        self.assertIn("#sidebar {", html)
        self.assertIn(".steps-bar {", html)
        self.assertIn(".form-with-preview {", html)
        self.assertIn(".editor-layout {", html)

    def test_entry_table_covers_new_user_facing_template_variables(self):
        entry_table = Path(app.BASE_DIR / "06_要素录入表.md").read_text(encoding="utf-8")

        for key in (
            "CSRC_APPROVAL_NO",
            "FUND_MANAGER_NAME",
            "FUND_MANAGER_SEX",
            "FUND_MANAGER_BIO",
            "CUSTODIAN_INTRO",
            "ACCOUNTANT",
        ):
            self.assertIn(key, entry_table)


class ProspectusMaterialsLibraryTests(unittest.TestCase):
    def _material_file_paths(self):
        materials_dir = Path(app.BASE_DIR / "prospectus_materials")
        return {
            "overview": materials_dir / "00_素材库说明.md",
            "reference_index": materials_dir / "01_参考素材索引.json",
            "wording_map": materials_dir / "02_表述来源映射表.md",
            "function_map": materials_dir / "03_函数映射表.json",
            "shared_dependencies": materials_dir / "04_共享依赖清单.md",
            "howto": materials_dir / "05_修改操作指引.md",
        }

    def test_api_files_lists_only_generation_files_for_template_editor(self):
        client = app.app.test_client()
        with tempfile.TemporaryDirectory() as tmp_dir:
            base_dir = Path(tmp_dir)
            (base_dir / "02_招募说明书模板.md").write_text("root template", encoding="utf-8")
            (base_dir / "09_业务正文覆盖.json").write_text("{}", encoding="utf-8")
            (base_dir / "prospectus_materials").mkdir()
            (base_dir / "prospectus_materials" / "00_素材库说明.md").write_text("materials", encoding="utf-8")
            (base_dir / "other_notes").mkdir()
            (base_dir / "other_notes" / "ignore.md").write_text("ignore", encoding="utf-8")

            with mock.patch.object(app, "BASE_DIR", base_dir):
                response = client.get("/api/files")

            self.assertEqual(200, response.status_code)
            payload = response.get_json()
            response.close()

        paths = {item.get("path") for item in payload}
        self.assertIn("02_招募说明书模板.md", paths)
        self.assertNotIn("09_业务正文覆盖.json", paths)
        self.assertNotIn("prospectus_materials/00_素材库说明.md", paths)
        self.assertNotIn("other_notes/ignore.md", paths)
        self.assertTrue(all("path" in item for item in payload))

    def test_api_file_blocks_non_generation_paths_and_rejects_traversal(self):
        client = app.app.test_client()
        with tempfile.TemporaryDirectory() as tmp_dir:
            base_dir = Path(tmp_dir)
            editable_file = base_dir / "02_招募说明书模板.md"
            editable_file.write_text("before save", encoding="utf-8")
            material_dir = base_dir / "prospectus_materials"
            material_dir.mkdir()
            material_file = material_dir / "00_素材库说明.md"
            material_file.write_text("before save", encoding="utf-8")

            with mock.patch.object(app, "BASE_DIR", base_dir):
                get_response = client.get("/api/files/02_招募说明书模板.md")
                self.assertEqual(200, get_response.status_code)
                self.assertEqual("before save", get_response.get_json()["content"])
                get_response.close()

                save_response = client.post(
                    "/api/files/02_招募说明书模板.md",
                    json={"content": "after save"},
                )
                self.assertEqual(200, save_response.status_code)
                self.assertTrue(save_response.get_json()["success"])
                save_response.close()

                material_response = client.get("/api/files/prospectus_materials/00_素材库说明.md")
                self.assertEqual(403, material_response.status_code)
                material_response.close()

                traversal_response = client.get("/api/files/%2E%2E%2Fsecret.md")
                self.assertEqual(403, traversal_response.status_code)
                traversal_response.close()

            self.assertEqual("after save", editable_file.read_text(encoding="utf-8"))
            self.assertEqual("before save", material_file.read_text(encoding="utf-8"))

    def test_api_clause_library_flattens_business_edit_entries_and_readonly_items(self):
        client = app.app.test_client()

        response = client.get("/api/clause_library")

        self.assertEqual(200, response.status_code)
        payload = response.get_json()
        response.close()

        self.assertTrue(payload["success"])
        self.assertIn("entries", payload)

        editable = next(
            item
            for item in payload["entries"]
            if item["document_type"] == "contract"
            and item["clause_key"] == "LISTING_TRADING_TEXTS"
            and item["field_key"] == "LISTING_SECTION2_BODY"
            and item["scene_key"] == "SSE_HK"
        )
        self.assertEqual("基金合同", editable["document_label"])
        self.assertFalse(editable["readonly"])
        self.assertIn("上市交易", editable["title"])
        self.assertIn("上海证券交易所", editable["content"])
        self.assertIn("path_id", editable)
        self.assertIn("raw_path", editable)
        self.assertIn("applicability", editable)
        self.assertIn("上交所港股通", editable["applicability"])
        self.assertIn(editable["applicability"], editable["search_text"])

        readonly = next(
            item
            for item in payload["entries"]
            if item["document_type"] == "prospectus"
            and item["clause_key"] == "PROSPECTUS_VARIANTS"
        )
        self.assertTrue(readonly["readonly"])
        self.assertIn("只读", readonly["readonly_reason"])
        self.assertIn("招募说明书", readonly["document_label"])

    def test_api_clause_library_save_updates_single_text_field_and_reloads_engine(self):
        client = app.app.test_client()
        original_base_dir = app.BASE_DIR
        original_engine = app.engine
        original_prospectus_engine = app.prospectus_engine
        original_product_summary_engine = app.product_summary_engine
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                contract_data = {
                    "version": "test",
                    "clauses": {
                        "LISTING_TRADING_TEXTS": {
                            "description": "上市交易测试条款",
                            "variants": {
                                "SSE_HK": {
                                    "condition": "测试条件",
                                    "LISTING_SECTION2_BODY": "保存前正文",
                                }
                            },
                        }
                    },
                }
                prospectus_data = {"version": "test", "clauses": {}}
                (base_dir / "04_差异条款原文库.json").write_text(
                    json.dumps(contract_data, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
                (base_dir / "06_招募说明书差异条款库.json").write_text(
                    json.dumps(prospectus_data, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )

                with mock.patch.object(app, "BASE_DIR", base_dir):
                    list_response = client.get("/api/clause_library")
                    self.assertEqual(200, list_response.status_code)
                    entry = next(
                        item
                        for item in list_response.get_json()["entries"]
                        if item["field_key"] == "LISTING_SECTION2_BODY"
                    )
                    list_response.close()

                    save_response = client.post(
                        "/api/clause_library",
                        json={"path_id": entry["path_id"], "content": "保存后正文"},
                    )
                    self.assertEqual(200, save_response.status_code)
                    self.assertTrue(save_response.get_json()["success"])
                    save_response.close()

                saved = json.loads((base_dir / "04_差异条款原文库.json").read_text(encoding="utf-8"))
                self.assertEqual(
                    "保存后正文",
                    saved["clauses"]["LISTING_TRADING_TEXTS"]["variants"]["SSE_HK"]["LISTING_SECTION2_BODY"],
                )
                self.assertEqual(
                    "测试条件",
                    saved["clauses"]["LISTING_TRADING_TEXTS"]["variants"]["SSE_HK"]["condition"],
                )
        finally:
            app.BASE_DIR = original_base_dir
            app.engine = original_engine
            app.prospectus_engine = original_prospectus_engine
            app.product_summary_engine = original_product_summary_engine

    def test_api_business_texts_returns_default_listing_text_without_override_file(self):
        client = app.app.test_client()
        original_base_dir = app.BASE_DIR
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                with mock.patch.object(app, "BASE_DIR", base_dir):
                    response = client.get("/api/business_texts")

                self.assertEqual(200, response.status_code)
                payload = response.get_json()
                response.close()

            self.assertTrue(payload["success"])
            listing_section2 = payload["groups"]["contract"]["LISTING_SECTION2_BODY"]
            self.assertIn("SSE_HK", listing_section2["variants"])
            self.assertIn(
                "上海证券交易所",
                listing_section2["variants"]["SSE_HK"]["default"],
            )
            self.assertEqual(
                "",
                listing_section2["variants"]["SSE_HK"]["override"],
            )
            prospectus_subscription = payload["groups"]["prospectus"]["SUBSCRIPTION_ACCOUNT_CLAUSE"]
            self.assertIn("SSE_HK", prospectus_subscription["variants"])
            self.assertIn(
                "上海证券交易所A股账户",
                prospectus_subscription["variants"]["SSE_HK"]["default"],
            )
            chapter_ten_sec5 = payload["groups"]["prospectus"]["CHAPTER10_SEC5_DEFAULT_BODY"]
            self.assertIn("DEFAULT", chapter_ten_sec5["variants"])
            self.assertIn(
                "申购和赎回的数额限制",
                chapter_ten_sec5["variants"]["DEFAULT"]["default"],
            )
        finally:
            app.BASE_DIR = original_base_dir

    def test_api_business_texts_save_persists_override_and_reloads_engine(self):
        client = app.app.test_client()
        original_base_dir = app.BASE_DIR
        original_engine = app.engine
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                override_path = base_dir / "09_业务正文覆盖.json"
                override_text = "通过接口保存后的上市交易正文。"

                with mock.patch.object(app, "BASE_DIR", base_dir):
                    response = client.post(
                        "/api/business_texts",
                        json={
                            "group": "contract",
                            "key": "LISTING_SECTION2_BODY",
                            "variant": "SSE_HK",
                            "product_type": "ETF",
                            "market_type": "HK_CONNECT",
                            "content": override_text,
                        },
                    )

                self.assertEqual(200, response.status_code)
                payload = response.get_json()
                response.close()

                self.assertTrue(payload["success"])
                self.assertTrue(override_path.exists())

                saved = json.loads(override_path.read_text(encoding="utf-8"))
                self.assertEqual(
                    override_text,
                    saved["contract"]["LISTING_SECTION2_BODY"]["SSE_HK"]["ETF"]["HK_CONNECT"]["ALL"],
                )

                form = {
                    "FUND_NAME": "测试交易型开放式指数证券投资基金",
                    "FUND_SHORT_NAME": "测试ETF",
                    "INDEX_NAME": "中证测试指数",
                    "INDEX_CODE": "000001",
                    "INDEX_COMPILER": "中证指数有限公司",
                    "CONTRACT_DATE": "2026年4月24日",
                    "EXCHANGE": "SSE",
                    "MARKET_TYPE": "HK_CONNECT",
                    "CUSTODIAN_NAME": "招商银行股份有限公司",
                    "DISTRIBUTION_FREQ": "QUARTERLY",
                    "MGMT_FEE_PAYMENT_METHOD": "CONSULT",
                    "CUSTODY_FEE_PAYMENT_METHOD": "CONSULT",
                    "MGMT_FEE_RATE": "0.15",
                    "CUSTODY_FEE_RATE": "0.05",
                }
                variables = app.engine._inject_clause_texts(
                    app.engine._derive_variables(form)
                )
                self.assertEqual(override_text, variables["LISTING_SECTION2_BODY"])
        finally:
            app.BASE_DIR = original_base_dir
            app.engine = original_engine

    def test_api_business_texts_returns_condition_options_and_processing_metadata(self):
        client = app.app.test_client()
        original_base_dir = app.BASE_DIR
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                with mock.patch.object(app, "BASE_DIR", base_dir):
                    response = client.get("/api/business_texts")

                self.assertEqual(200, response.status_code)
                payload = response.get_json()
                response.close()

            self.assertIn("condition_options", payload)
            self.assertIn("ETF", payload["condition_options"]["product_types"])
            self.assertIn("A_SHARE", payload["condition_options"]["market_types"])
            self.assertIn("SSE", payload["condition_options"].get("exchanges", []))

            listing_section2 = payload["groups"]["contract"]["LISTING_SECTION2_BODY"]
            self.assertIn("scenes", listing_section2)
            self.assertIn("supports_conditions", listing_section2)
            self.assertIn("processing_tags", listing_section2)
            self.assertIn("processing_description", listing_section2)

            listing_item3 = payload["groups"]["contract"]["LISTING_ITEM3_DEF"]
            self.assertIn("上市交易所", listing_item3["processing_description"])
            self.assertIn(
                "SSE",
                listing_item3["scenes"]["DEFAULT"]["matrix"]["ALL"]["ALL"],
            )

            chapter_ten_sec5 = payload["groups"]["prospectus"]["CHAPTER10_SEC5_DEFAULT_BODY"]
            self.assertIn("DEFAULT", chapter_ten_sec5["scenes"])
            self.assertTrue(chapter_ten_sec5["processing_description"])
            self.assertIn("章节拼装/规范化", chapter_ten_sec5["processing_tags"])

        finally:
            app.BASE_DIR = original_base_dir

    def test_api_business_texts_accepts_old_request_shape_but_saves_condition_matrix(self):
        client = app.app.test_client()
        original_base_dir = app.BASE_DIR
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                base_dir = Path(tmp_dir)
                override_path = base_dir / "09_业务正文覆盖.json"

                with mock.patch.object(app, "BASE_DIR", base_dir):
                    response = client.post(
                        "/api/business_texts",
                        json={
                            "group": "prospectus",
                            "key": "INDEX_INFO_SOURCE_CLAUSE",
                            "variant": "DEFAULT",
                            "content": "legacy request payload",
                        },
                    )

                self.assertEqual(200, response.status_code)
                payload = response.get_json()
                response.close()
                self.assertTrue(payload["success"])

                saved = json.loads(override_path.read_text(encoding="utf-8"))
                self.assertEqual(
                    "legacy request payload",
                    saved["prospectus"]["INDEX_INFO_SOURCE_CLAUSE"]["DEFAULT"]["ALL"]["ALL"]["ALL"],
                )

        finally:
            app.BASE_DIR = original_base_dir

    def test_prospectus_material_library_files_cover_reference_assets_and_core_functions(self):
        material_paths = self._material_file_paths()
        for path in material_paths.values():
            self.assertTrue(path.exists(), f"missing material file: {path}")

        reference_index = json.loads(material_paths["reference_index"].read_text(encoding="utf-8"))
        function_map = json.loads(material_paths["function_map"].read_text(encoding="utf-8"))

        self.assertIn("items", reference_index)
        self.assertIn("functions", function_map)

        variant_keys = {item["variant_key"] for item in reference_index["items"]}
        self.assertEqual(
            {"SSE_CROSS", "SSE_SINGLE", "SSE_HK", "SZSE_CROSS", "SZSE_SINGLE", "SZSE_HK"},
            variant_keys,
        )
        source_types = {item["source_type"] for item in reference_index["items"]}
        self.assertTrue({"reference_docx", "reference_txt"}.issubset(source_types))

        function_names = {item["function_name"] for item in function_map["functions"]}
        for required_name in (
            "_reference_prospectus_docx_candidates",
            "_resolve_reference_prospectus_docx",
            "_load_reference_fixed_content_from_txt",
            "_load_reference_fixed_content",
            "_replace_manual_reference_text",
            "_replace_reference_fund_name",
            "_apply_reference_fixed_content",
            "_apply_canonical_reference_overrides",
            "_build_display_prospectus_text",
            "generate_bundle",
            "_ensure_docx_reference_assets",
            "build_docx_prospectus",
        ):
            self.assertIn(required_name, function_names)

    def test_index_template_groups_generation_files_and_supports_clause_editor(self):
        html = Path(app.BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")

        self.assertIn("f.path || f.name", html)
        self.assertIn("groupFilesForEditor", html)
        self.assertIn("变量库", html)
        self.assertIn("条款库", html)
        self.assertIn("模板", html)
        self.assertIn("clause-structured-editor", html)
        self.assertIn("clause-business-search", html)
        self.assertIn("loadClauseBusinessLibrary", html)
        self.assertIn("/api/clause_library", html)
        self.assertIn("encodeURIComponent(currentFile)", html)


if __name__ == "__main__":
    unittest.main()
