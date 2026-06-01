"""
ETF基金合同知识库 Web 应用
Flask 后端服务器 + 合同生成引擎
"""

import base64
import io
import json
import logging
import os
import re
import sys
import threading
import time
import unicodedata
import webbrowser
import zipfile
import xml.sax.saxutils
import difflib
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from flask import Flask, abort, jsonify, render_template, request, send_file
from packaging_support import (
    LEGACY_PRODUCT_SUMMARY_TEMPLATE_CANDIDATES,
    LEGACY_REFERENCE_PROSPECTUS_SOURCE_MAP,
    LEGACY_REVIEW_XLSX_CANDIDATES,
    LEGACY_RULES_XLSX_CANDIDATES,
    PACKAGED_ASSETS_RELATIVE_DIR,
    PRODUCT_SUMMARY_TEMPLATE_FILENAME,
    REVIEW_WORKBOOK_FILENAMES,
    RULES_XLSX_FILENAME,
    compute_app_root,
)
from review_builtin_rules import BUILTIN_ETF_CROSS_RULES, BUILTIN_ETF_SUMMARY_RULES

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
logger = logging.getLogger(__name__)

# ── 路径常量 ────────────────────────────────────────────────────────────────
SOURCE_DIR = Path(__file__).resolve().parent
BASE_DIR = compute_app_root(__file__, frozen=bool(getattr(sys, "frozen", False)))
PACKAGED_ASSETS_DIR = BASE_DIR / PACKAGED_ASSETS_RELATIVE_DIR
TEMPLATES_DIR = BASE_DIR / "templates"
BUSINESS_TEXT_OVERRIDES_FILENAME = "09_业务正文覆盖.json"


def _resolve_repo_data_path(*names: str) -> Path:
    for name in names:
        candidate = BASE_DIR / name
        if candidate.exists():
            return candidate
    return BASE_DIR / names[0]


TEMPLATE_MD = _resolve_repo_data_path("01_通用模板.md", "01_基金合同模板.md")
SCHEMA_JSON = _resolve_repo_data_path("02_变量定义表.json", "03_变量定义表.json")
DIFF_TABLE_MD = _resolve_repo_data_path("03_差异条款匹配表.md", "04_基金合同差异条款匹配表.md")
CLAUSES_JSON = _resolve_repo_data_path("04_差异条款原文库.json", "07_基金合同差异条款原文库.json")
ENTRY_TABLE_MD = _resolve_repo_data_path("05_要素录入表.md", "06_要素录入表.md")
PROSPECTUS_TEMPLATE_MD = _resolve_repo_data_path("01_招募说明书模板.md", "02_招募说明书模板.md")
PROSPECTUS_CLAUSES_JSON = _resolve_repo_data_path("06_招募说明书差异条款库.json", "08_招募说明书差异条款库.json")
REFERENCE_PROSPECTUS_REPO_DIR = BASE_DIR / "manual_regression_20260310_201555"
PACKAGED_REFERENCE_PROSPECTUS_DIR = PACKAGED_ASSETS_DIR / "reference_prospectus"
REFERENCE_PROSPECTUS_DOCX_MAP = {
    "SSE_CROSS": REFERENCE_PROSPECTUS_REPO_DIR / "SSE_CROSS.docx",
    "SSE_SINGLE": REFERENCE_PROSPECTUS_REPO_DIR / "SSE_SINGLE.docx",
    "SSE_HK": REFERENCE_PROSPECTUS_REPO_DIR / "SSE_HK.docx",
    "SZSE_CROSS": REFERENCE_PROSPECTUS_REPO_DIR / "SZSE_CROSS.docx",
    "SZSE_SINGLE": REFERENCE_PROSPECTUS_REPO_DIR / "SZSE_SINGLE.docx",
    "SZSE_HK": REFERENCE_PROSPECTUS_REPO_DIR / "SZSE_HK.docx",
}
PACKAGED_REFERENCE_PROSPECTUS_DOCX_MAP = {
    key: PACKAGED_REFERENCE_PROSPECTUS_DIR / f"{key}.docx" for key in REFERENCE_PROSPECTUS_DOCX_MAP
}
LEGACY_REFERENCE_PROSPECTUS_DOCX_MAP = {
    key: sources[0] for key, sources in LEGACY_REFERENCE_PROSPECTUS_SOURCE_MAP.items()
}
PRODUCT_SUMMARY_TEMPLATE_DOCX = PACKAGED_ASSETS_DIR / "product_summary" / PRODUCT_SUMMARY_TEMPLATE_FILENAME
LEGACY_PRODUCT_SUMMARY_TEMPLATE_DOCX_CANDIDATES = list(LEGACY_PRODUCT_SUMMARY_TEMPLATE_CANDIDATES)
_REFERENCE_PROSPECTUS_DOCX_QUALITY_CACHE: dict[str, tuple[int, int, int, int]] = {}

ALLOWED_SUFFIXES = {".md", ".json"}


EDITOR_GENERATION_FILE_CANDIDATE_NAMES = [
    ("01_通用模板.md", "01_基金合同模板.md"),
    ("01_招募说明书模板.md", "02_招募说明书模板.md"),
    ("02_变量定义表.json", "03_变量定义表.json"),
    ("04_差异条款原文库.json", "07_基金合同差异条款原文库.json"),
    ("06_招募说明书差异条款库.json", "08_招募说明书差异条款库.json"),
]


def _editable_generation_file_paths() -> list[Path]:
    candidates: list[Path] = []
    for names in EDITOR_GENERATION_FILE_CANDIDATE_NAMES:
        selected = None
        for name in names:
            candidate = BASE_DIR / name
            if candidate.exists():
                selected = candidate
                break
        candidates.append(selected or (BASE_DIR / names[0]))
    seen = set()
    paths: list[Path] = []
    for candidate in candidates:
        resolved = candidate.resolve()
        if resolved in seen:
            continue
        seen.add(resolved)
        paths.append(candidate)
    return paths

BUSINESS_TEXT_PRODUCT_TYPE_OPTIONS = [
    ("ALL", "全部产品类型"),
    ("ETF", "ETF"),
]
BUSINESS_TEXT_MARKET_TYPE_OPTIONS = [
    ("ALL", "全部市场类型"),
    ("CHUANGYE", "创业板"),
    ("KECHUANG", "科创板"),
    ("A_SHARE", "普通A股"),
    ("HK_CONNECT", "港股通"),
]
BUSINESS_TEXT_EXCHANGE_OPTIONS = [
    ("ALL", "全部上市交易所"),
    ("SSE", "上海证券交易所"),
    ("SZSE", "深圳证券交易所"),
]
DEFAULT_DISPUTE_RESOLUTION_VENUE = "SZ_SCIA"
DISPUTE_RESOLUTION_VENUES = {
    "SZ_SCIA": {
        "label": "深圳-深圳国际仲裁院",
        "institution": "深圳国际仲裁院",
        "location": "深圳市",
        "contract_clause": (
            "各方当事人同意，因《基金合同》而产生的或与《基金合同》有关的一切争议，"
            "如经友好协商未能解决的，任何一方均有权将争议提交深圳国际仲裁院，"
            "按照深圳国际仲裁院届时有效的仲裁规则进行仲裁。仲裁地点为深圳市。"
            "仲裁裁决是终局的，对当事人均有约束力。除非仲裁裁决另有规定，"
            "仲裁费、律师费由败诉方承担。"
        ),
        "product_summary_sentence": (
            "与本基金/基金合同相关的争议解决方式为仲裁。因《基金合同》而产生的或与《基金合同》有关的一切争议，"
            "如经友好协商未能解决的，任何一方均有权将争议提交深圳国际仲裁院，"
            "按照深圳国际仲裁院届时有效的仲裁规则进行仲裁，仲裁地点为深圳市。"
        ),
    },
    "BJ_CIETAC": {
        "label": "北京-中国国际经济贸易仲裁委员会",
        "institution": "中国国际经济贸易仲裁委员会",
        "location": "北京市",
        "contract_clause": (
            "各方当事人同意，因《基金合同》而产生的或与《基金合同》有关的一切争议，"
            "如经友好协商未能解决的，任何一方均有权将争议提交中国国际经济贸易仲裁委员会，"
            "按照中国国际经济贸易仲裁委员会届时有效的仲裁规则进行仲裁。仲裁地点为北京市。"
            "仲裁裁决是终局的，对当事人均有约束力。除非仲裁裁决另有规定，"
            "仲裁费、律师费由败诉方承担。"
        ),
        "product_summary_sentence": (
            "与本基金/基金合同相关的争议解决方式为仲裁。因《基金合同》而产生的或与《基金合同》有关的一切争议，"
            "如经友好协商未能解决的，任何一方均有权将争议提交中国国际经济贸易仲裁委员会，"
            "按照中国国际经济贸易仲裁委员会届时有效的仲裁规则进行仲裁，仲裁地点为北京市。"
        ),
    },
}
DISPUTE_RESOLUTION_VENUE_ALIASES = {
    "": DEFAULT_DISPUTE_RESOLUTION_VENUE,
    "SZ": "SZ_SCIA",
    "SZ_SCIA": "SZ_SCIA",
    "SHENZHEN_SCIA": "SZ_SCIA",
    "SHENZHEN_SCIAC": "SZ_SCIA",
    "深圳": "SZ_SCIA",
    "深圳市": "SZ_SCIA",
    "深圳国际仲裁院": "SZ_SCIA",
    "深圳-深圳国际仲裁院": "SZ_SCIA",
    "BJ": "BJ_CIETAC",
    "BJ_CIETAC": "BJ_CIETAC",
    "BEIJING_CIETAC": "BJ_CIETAC",
    "北京": "BJ_CIETAC",
    "北京市": "BJ_CIETAC",
    "中国国际经济贸易仲裁委员会": "BJ_CIETAC",
    "北京-中国国际经济贸易仲裁委员会": "BJ_CIETAC",
}


def _resolve_dispute_resolution_venue(raw_value: str | None) -> tuple[str, dict[str, str]]:
    normalized = str(raw_value or "").strip()
    key = DISPUTE_RESOLUTION_VENUE_ALIASES.get(normalized)
    if key is None:
        key = DISPUTE_RESOLUTION_VENUE_ALIASES.get(normalized.upper(), DEFAULT_DISPUTE_RESOLUTION_VENUE)
    return key, DISPUTE_RESOLUTION_VENUES[key]
BUSINESS_TEXT_SCENE_LABELS = {
    "DEFAULT": "通用场景",
    "HK_CONNECT": "港股通通用",
    "SSE_HK": "上交所港股通",
    "SZSE_HK": "深交所港股通",
    "SSE_CROSS": "上交所跨市场",
    "SSE_SINGLE": "上交所单市场",
    "SZSE_CROSS": "深交所跨市场",
    "SZSE_SINGLE": "深交所单市场",
}
DEFAULT_DIRECT_PROCESSING_DESCRIPTION = "保存后按场景、产品类型、市场类型和上市交易所条件匹配直接使用，不再做额外处理。"
DEFAULT_HK_EXCHANGE_PROCESSING_DESCRIPTION = (
    "默认文案会先按上市交易所场景派生；如填写覆盖文案，系统会再按场景、产品类型、市场类型和上市交易所条件直接使用，不再替换正文。"
)

CONTRACT_BUSINESS_TEXT_SPECS = {
    "DEFINITION_TAIL": {"variant_mode": "hk_flag"},
    "PREFACE_RISK_DISCLOSURE_TAIL": {"variant_mode": "hk_flag"},
    "PURCHASE_CONSIDERATION_GLOSSARY_DEF": {"variant_mode": "hk_exchange"},
    "REDEMPTION_CONSIDERATION_GLOSSARY_DEF": {"variant_mode": "hk_exchange"},
    "FUND_SHARE_NAV_DEF": {"variant_mode": "hk_exchange"},
    "PURCHASE_REDEMPTION_PRINCIPLE_DEF": {"variant_mode": "hk_exchange"},
    "PURCHASE_REDEMPTION_CONSIDERATION_DEF": {"variant_mode": "hk_exchange"},
    "VALUATION_METHODS_TAIL_TEXT": {"variant_mode": "hk_flag"},
    "VALUATION_PROCEDURE_TEXT": {"variant_mode": "hk_flag"},
    "SPECIAL_DISCLOSURE_TEXT": {"variant_mode": "hk_flag"},
    "LISTING_ITEM3_DEF": {"variant_mode": "hk_exchange"},
    "LISTING_SECTION2_BODY": {"variant_mode": "hk_exchange"},
    "LISTING_SECTION3_TEXT": {"variant_mode": "hk_exchange"},
    "LISTING_SECTION4_TEXT": {"variant_mode": "hk_exchange"},
    "PART8_INTRO_TEXT": {"variant_mode": "hk_exchange"},
    "OPEN_DAY_CLAUSE": {"variant_mode": "hk_exchange"},
    "OPEN_DAY_ADJUSTMENT_CLAUSE": {"variant_mode": "hk_exchange"},
    "PART8_SECTION4_TEXT": {"variant_mode": "hk_exchange"},
    "PART8_SECTION6_TEXT": {"variant_mode": "hk_exchange"},
    "PART8_SECTION7_TEXT": {"variant_mode": "hk_exchange"},
    "PART8_SECTION8_TEXT": {"variant_mode": "hk_exchange"},
    "PART8_SECTION9_TEXT": {"variant_mode": "hk_exchange"},
    "HK_PROXY_VOTING_TEXT": {"variant_mode": "hk_flag"},
}

PROSPECTUS_BUSINESS_TEXT_SPECS = {
    "SUBSCRIPTION_ACCOUNT_CLAUSE": {
        "variant_mode": "prospectus_variant",
        "processing_tags": ["直接使用"],
        "processing_description": "保存后按场景、产品类型和市场类型匹配，直接写入招募说明书正文。",
    },
    "INDEX_INFO_SOURCE_CLAUSE": {
        "variant_mode": "DEFAULT",
        "processing_tags": ["直接使用"],
        "processing_description": "保存后按条件匹配，直接作为指数信息来源条款使用。",
    },
    "IMPORTANT_NOTICE_APPROVAL_SENTENCE": {
        "variant_mode": "DEFAULT",
        "render_placeholders": True,
        "processing_tags": ["占位符替换", "章节拼装/规范化"],
        "processing_description": "保存后会替换 {CSRC_APPROVAL_NO}，并在“重要提示”组装阶段写入正文。",
    },
    "IMPORTANT_NOTICE_INDEX_SOURCE_SENTENCE": {
        "variant_mode": "DEFAULT",
        "render_placeholders": True,
        "processing_tags": ["占位符替换", "章节拼装/规范化"],
        "processing_description": "保存后会替换 {INDEX_COMPILER}、{INDEX_WEBSITE}，并在“重要提示”组装阶段写入正文。",
    },
    "CHAPTER6_INTRO_REGISTRATION_SENTENCE": {
        "variant_mode": "DEFAULT",
        "render_placeholders": True,
        "processing_tags": ["占位符替换", "章节拼装/规范化"],
        "processing_description": "保存后会替换 {CSRC_APPROVAL_NO}，并在招募说明书第六章拼装时写入正文。",
    },
    "CHAPTER6_PRODUCT_TYPE_SENTENCE": {
        "variant_mode": "DEFAULT",
        "processing_tags": ["章节拼装/规范化"],
        "processing_description": "保存后会在招募说明书第六章拼装时写入正文。",
    },
    "CHAPTER6_SEC8_DEFAULT_BODY": {
        "variant_mode": "DEFAULT",
        "processing_tags": ["章节拼装/规范化"],
        "processing_description": "保存后会作为第六章第八节默认文案，并继续经过表格/格式规范化处理。",
    },
    "CHAPTER10_SEC5_DEFAULT_BODY": {
        "variant_mode": "DEFAULT",
        "render_placeholders": True,
        "processing_tags": ["占位符替换", "章节拼装/规范化"],
        "processing_description": "保存后会替换 {MIN_SUB_UNIT}，并在第十章第五节拼装时继续做条目规范化。",
    },
    "CHAPTER21_PLACEHOLDER_TEXT": {
        "variant_mode": "DEFAULT",
        "processing_tags": ["章节拼装/规范化"],
        "processing_description": "保存后会作为第二十一章缺省占位文本，在章节拼装阶段写入。",
    },
}

BUSINESS_TEXT_GROUP_SPECS = {
    "contract": CONTRACT_BUSINESS_TEXT_SPECS,
    "prospectus": PROSPECTUS_BUSINESS_TEXT_SPECS,
}


def _business_text_override_path() -> Path:
    return BASE_DIR / BUSINESS_TEXT_OVERRIDES_FILENAME


def _empty_business_text_overrides() -> dict:
    return {
        "contract": {},
        "prospectus": {},
        "meta": {"version": "3.0"},
    }


def _normalize_business_text_overrides(data: dict | None) -> dict:
    normalized = _empty_business_text_overrides()
    if not isinstance(data, dict):
        return normalized

    for group in ("contract", "prospectus"):
        raw_group = data.get(group, {})
        if isinstance(raw_group, dict):
            group_result = {}
            for key, key_data in raw_group.items():
                if not isinstance(key_data, dict):
                    continue
                normalized_key = {}
                for scene, scene_data in key_data.items():
                    if isinstance(scene_data, str):
                        normalized_key[str(scene)] = {"ALL": {"ALL": {"ALL": str(scene_data)}}}
                        continue
                    if not isinstance(scene_data, dict):
                        continue
                    product_map = {}
                    for product_type, market_map in scene_data.items():
                        if isinstance(market_map, str):
                            product_map[str(product_type)] = {"ALL": {"ALL": str(market_map)}}
                            continue
                        if not isinstance(market_map, dict):
                            continue
                        normalized_market_map = {}
                        for market_type, content in market_map.items():
                            if isinstance(content, str):
                                normalized_market_map[str(market_type)] = {"ALL": str(content)}
                                continue
                            if not isinstance(content, dict):
                                continue
                            exchange_map = {}
                            for exchange, exchange_content in content.items():
                                if exchange_content is None:
                                    continue
                                exchange_map[str(exchange)] = str(exchange_content)
                            if exchange_map:
                                normalized_market_map[str(market_type)] = exchange_map
                        if normalized_market_map:
                            product_map[str(product_type)] = normalized_market_map
                    if product_map:
                        normalized_key[str(scene)] = product_map
                if normalized_key:
                    group_result[str(key)] = normalized_key
            normalized[group] = group_result

    raw_meta = data.get("meta", {})
    if isinstance(raw_meta, dict):
        normalized["meta"].update(raw_meta)

    return normalized


def _load_business_text_overrides() -> dict:
    path = _business_text_override_path()
    if not path.exists():
        return _empty_business_text_overrides()

    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    return _normalize_business_text_overrides(data)


def _save_business_text_overrides(data: dict) -> None:
    path = _business_text_override_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(_normalize_business_text_overrides(data), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def _business_text_product_type(v: dict) -> str:
    product_type = str((v or {}).get("PRODUCT_TYPE", "") or "").strip().upper()
    return product_type or "ETF"


def _business_text_market_type(v: dict) -> str:
    market_type = str((v or {}).get("MARKET_TYPE", "") or "").strip().upper()
    return market_type or "A_SHARE"


def _business_text_exchange(v: dict) -> str:
    exchange = str((v or {}).get("EXCHANGE", "") or "").strip().upper()
    return exchange if exchange in {"SSE", "SZSE"} else "ALL"


def _business_text_scene_label(scene: str) -> str:
    return BUSINESS_TEXT_SCENE_LABELS.get(scene, scene or "通用场景")


def _business_text_product_type_values() -> list[str]:
    return [value for value, _ in BUSINESS_TEXT_PRODUCT_TYPE_OPTIONS]


def _business_text_market_type_values() -> list[str]:
    return [value for value, _ in BUSINESS_TEXT_MARKET_TYPE_OPTIONS]


def _business_text_exchange_values() -> list[str]:
    return [value for value, _ in BUSINESS_TEXT_EXCHANGE_OPTIONS]


def _business_text_exact_override(
    overrides: dict,
    group: str,
    key: str,
    scene: str,
    product_type: str = "ALL",
    market_type: str = "ALL",
    exchange: str = "ALL",
) -> str | None:
    group_data = overrides.get(group, {})
    if not isinstance(group_data, dict):
        return None
    key_data = group_data.get(key, {})
    if not isinstance(key_data, dict):
        return None
    scene_data = key_data.get(scene, {})
    if not isinstance(scene_data, dict):
        return None
    product_data = scene_data.get(product_type, {})
    if not isinstance(product_data, dict):
        return None
    market_data = product_data.get(market_type, {})
    if not isinstance(market_data, dict):
        return None
    value = market_data.get(exchange)
    return None if value is None else str(value)


def _get_business_text_override(
    overrides: dict,
    group: str,
    key: str,
    scene: str,
    product_type: str = "ALL",
    market_type: str = "ALL",
    exchange: str = "ALL",
) -> str | None:
    for candidate_product_type, candidate_market_type, candidate_exchange in (
        (product_type, market_type, exchange),
        (product_type, market_type, "ALL"),
        (product_type, "ALL", exchange),
        ("ALL", market_type, exchange),
        (product_type, "ALL", "ALL"),
        ("ALL", market_type, "ALL"),
        ("ALL", "ALL", exchange),
        ("ALL", "ALL", "ALL"),
    ):
        value = _business_text_exact_override(
            overrides,
            group,
            key,
            scene,
            candidate_product_type,
            candidate_market_type,
            candidate_exchange,
        )
        if value is not None:
            return value
    return None


def _set_business_text_override(
    overrides: dict,
    group: str,
    key: str,
    scene: str,
    content: str,
    product_type: str = "ALL",
    market_type: str = "ALL",
    exchange: str = "ALL",
) -> dict:
    updated = _normalize_business_text_overrides(deepcopy(overrides))
    updated.setdefault(group, {})
    updated[group].setdefault(key, {})
    updated[group][key].setdefault(scene, {})
    updated[group][key][scene].setdefault(product_type, {})
    updated[group][key][scene][product_type].setdefault(market_type, {})
    updated[group][key][scene][product_type][market_type][exchange] = str(content or "")
    return updated


def _delete_business_text_override(
    overrides: dict,
    group: str,
    key: str,
    scene: str,
    product_type: str = "ALL",
    market_type: str = "ALL",
    exchange: str = "ALL",
) -> dict:
    updated = _normalize_business_text_overrides(deepcopy(overrides))
    scene_map = updated.get(group, {}).get(key, {}).get(scene, {})
    if isinstance(scene_map, dict):
        product_map = scene_map.get(product_type, {})
        if isinstance(product_map, dict):
            market_map = product_map.get(market_type, {})
            if isinstance(market_map, dict):
                market_map.pop(exchange, None)
                if not market_map:
                    product_map.pop(market_type, None)
            if not product_map:
                scene_map.pop(product_type, None)
        if not scene_map:
            updated[group].get(key, {}).pop(scene, None)
        if not updated[group].get(key, {}):
            updated[group].pop(key, None)
    return updated


def _contract_business_text_variant(v: dict, variant_mode: str) -> str:
    if variant_mode == "hk_exchange":
        hk_variant = str(v.get("HK_CONNECT_EXCHANGE_VARIANT", "") or "").strip()
        return hk_variant if hk_variant in {"SSE_HK", "SZSE_HK"} else "DEFAULT"
    if variant_mode == "hk_flag":
        return "HK_CONNECT" if v.get("HAS_HK_CONNECT") else "DEFAULT"
    return "DEFAULT"


def _contract_business_text_variant_values(variant_mode: str) -> list[str]:
    if variant_mode == "hk_exchange":
        return ["SSE_HK", "SZSE_HK", "DEFAULT"]
    if variant_mode == "hk_flag":
        return ["HK_CONNECT", "DEFAULT"]
    return ["DEFAULT"]


def _prospectus_variant_key(v: dict) -> str:
    exchange = str((v or {}).get("EXCHANGE", "") or "").strip().upper()
    market_type = str((v or {}).get("MARKET_TYPE", "") or "").strip().upper()
    market_scope = str((v or {}).get("MARKET_SCOPE", "") or "").strip().upper()

    if market_type in {"KECHUANG", "CHUANGYE"}:
        market_scope = "SINGLE_MARKET"
    elif market_scope not in {"SINGLE_MARKET", "CROSS_MARKET"}:
        market_scope = "CROSS_MARKET"

    if market_type == "HK_CONNECT":
        return "SSE_HK" if exchange == "SSE" else "SZSE_HK"
    if exchange == "SSE":
        return "SSE_SINGLE" if market_scope == "SINGLE_MARKET" else "SSE_CROSS"
    return "SZSE_SINGLE" if market_scope == "SINGLE_MARKET" else "SZSE_CROSS"


def _prospectus_business_text_variant(v: dict, variant_mode: str) -> str:
    if variant_mode == "prospectus_variant":
        return _prospectus_variant_key(v)
    return "DEFAULT"


def _business_text_variant_values(group: str, variant_mode: str) -> list[str]:
    if group == "contract":
        return _contract_business_text_variant_values(variant_mode)
    if group == "prospectus" and variant_mode == "prospectus_variant":
        return ["SSE_CROSS", "SSE_SINGLE", "SSE_HK", "SZSE_CROSS", "SZSE_SINGLE", "SZSE_HK"]
    return ["DEFAULT"]


def _business_text_processing_tags(spec: dict) -> list[str]:
    tags = spec.get("processing_tags")
    if isinstance(tags, list) and tags:
        return [str(tag) for tag in tags]
    if str(spec.get("variant_mode") or "") == "hk_exchange":
        return ["直接使用", "默认值按上市交易所派生"]
    return ["直接使用"]


def _business_text_processing_description(spec: dict) -> str:
    description = str(spec.get("processing_description") or "").strip()
    if description:
        return description
    if str(spec.get("variant_mode") or "") == "hk_exchange":
        return DEFAULT_HK_EXCHANGE_PROCESSING_DESCRIPTION
    return DEFAULT_DIRECT_PROCESSING_DESCRIPTION


def _business_text_condition_options() -> dict:
    return {
        "product_types": _business_text_product_type_values(),
        "product_type_labels": {value: label for value, label in BUSINESS_TEXT_PRODUCT_TYPE_OPTIONS},
        "market_types": _business_text_market_type_values(),
        "market_type_labels": {value: label for value, label in BUSINESS_TEXT_MARKET_TYPE_OPTIONS},
        "exchanges": _business_text_exchange_values(),
        "exchange_labels": {value: label for value, label in BUSINESS_TEXT_EXCHANGE_OPTIONS},
    }


CLAUSE_LIBRARY_SOURCE_CANDIDATES = {
    "contract": ("04_差异条款原文库.json", "07_基金合同差异条款原文库.json"),
    "prospectus": ("06_招募说明书差异条款库.json", "08_招募说明书差异条款库.json"),
}
CLAUSE_LIBRARY_DOCUMENT_LABELS = {
    "contract": "基金合同",
    "prospectus": "招募说明书",
}
CLAUSE_LIBRARY_METADATA_FIELDS = {
    "applicable_to",
    "condition",
    "description",
    "legacy_source",
    "location",
    "number_label",
    "preservation_rule",
    "source_sample",
    "title",
}
CLAUSE_LIBRARY_FIELD_LABELS = {
    "text": "正文",
    "LISTING_ITEM3_DEF": "上市释义",
    "LISTING_SECTION2_BODY": "上市交易正文",
    "LISTING_SECTION3_TEXT": "参考净值正文",
    "LISTING_SECTION4_TEXT": "停复牌终止上市正文",
    "VALUATION_METHODS_TAIL_TEXT": "估值方法补充",
    "VALUATION_PROCEDURE_TEXT": "估值程序",
    "SPECIAL_DISCLOSURE_TEXT": "特殊信息披露",
    "PART8_INTRO_TEXT": "申购赎回引言",
    "OPEN_DAY_CLAUSE": "开放日定义",
    "OPEN_DAY_ADJUSTMENT_CLAUSE": "开放日调整",
    "PART8_SECTION4_TEXT": "申购赎回程序",
    "PART8_SECTION6_TEXT": "申购赎回对价费用",
    "PART8_SECTION7_TEXT": "申购赎回清单",
    "PART8_SECTION8_TEXT": "拒绝或暂停申购",
    "PART8_SECTION9_TEXT": "暂停赎回",
}
CLAUSE_LIBRARY_SCENE_LABELS = {
    "DEFAULT": "通用",
    "STANDARD": "标准",
    "HK_CONNECT": "港股通",
    "NON_HK_CONNECT": "非港股通",
    "SSE": "上交所",
    "SZSE": "深交所",
    "SSE_HK": "上交所港股通",
    "SZSE_HK": "深交所港股通",
    "SSE_CROSS": "上交所跨市场",
    "SSE_SINGLE": "上交所单市场",
    "SZSE_CROSS": "深交所跨市场",
    "SZSE_SINGLE": "深交所单市场",
    "MARKET_SH_STANDARD": "上交所非港股通",
    "MARKET_SZ_STANDARD": "深交所非港股通",
    "MARKET_SH_HK_CONNECT": "上交所港股通",
    "MARKET_SZ_HK_CONNECT": "深交所港股通",
}


def _resolve_base_dir_data_path(*names: str) -> Path:
    for name in names:
        candidate = BASE_DIR / name
        if candidate.exists():
            return candidate
    return BASE_DIR / names[0]


def _clause_library_sources() -> list[dict]:
    return [
        {
            "document_type": document_type,
            "document_label": CLAUSE_LIBRARY_DOCUMENT_LABELS[document_type],
            "path": _resolve_base_dir_data_path(*names),
        }
        for document_type, names in CLAUSE_LIBRARY_SOURCE_CANDIDATES.items()
    ]


def _encode_clause_path_id(document_type: str, path_parts: list[str]) -> str:
    payload = {"document_type": document_type, "path": path_parts}
    raw = json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
    return base64.urlsafe_b64encode(raw).decode("ascii").rstrip("=")


def _decode_clause_path_id(path_id: str) -> tuple[str, list[str]]:
    padded = str(path_id or "") + "=" * (-len(str(path_id or "")) % 4)
    try:
        payload = json.loads(base64.urlsafe_b64decode(padded.encode("ascii")).decode("utf-8"))
    except Exception as exc:
        raise ValueError("无效的条款路径") from exc
    document_type = str(payload.get("document_type") or "")
    path_parts = payload.get("path")
    if document_type not in CLAUSE_LIBRARY_SOURCE_CANDIDATES or not isinstance(path_parts, list):
        raise ValueError("无效的条款路径")
    return document_type, [str(part) for part in path_parts]


def _clause_library_source_for(document_type: str) -> dict:
    for source in _clause_library_sources():
        if source["document_type"] == document_type:
            return source
    raise ValueError("未知条款库文档类型")


def _clause_scene_label(scene_key: str) -> str:
    return CLAUSE_LIBRARY_SCENE_LABELS.get(scene_key, _business_text_scene_label(scene_key))


def _clause_field_label(field_key: str) -> str:
    if field_key in CLAUSE_LIBRARY_FIELD_LABELS:
        return CLAUSE_LIBRARY_FIELD_LABELS[field_key]
    label = str(field_key or "").strip()
    if label == "text":
        return "正文"
    return label.replace("_", " ").strip() or "正文"


def _clause_chapter_hint(*values: str) -> str:
    text = " ".join(str(value or "") for value in values)
    match = re.search(r"(第[一二三四五六七八九十百零\d]+(?:部分|章)[^，。；；/、]*)", text)
    if match:
        return match.group(1).strip()
    return "未归类"


def _clause_source_note(*nodes: dict) -> str:
    parts = []
    for node in nodes:
        if not isinstance(node, dict):
            continue
        if node.get("source_sample"):
            parts.append(f"来源：{node.get('source_sample')}")
        legacy_source = node.get("legacy_source")
        if isinstance(legacy_source, dict):
            clause = legacy_source.get("clause") or ""
            variant = legacy_source.get("variant") or ""
            parts.append(f"引用：{clause}/{variant}".rstrip("/"))
        if node.get("preservation_rule"):
            parts.append(str(node.get("preservation_rule")))
    return "；".join(dict.fromkeys(part for part in parts if part))


def _is_editable_clause_field(field_key: str, value) -> bool:
    return isinstance(value, str) and field_key not in CLAUSE_LIBRARY_METADATA_FIELDS


def _clause_entry_title(clause_key: str, clause: dict, field_key: str, target_node: dict) -> str:
    description = str((clause or {}).get("description") or "").strip()
    child_title = str((target_node or {}).get("title") or "").strip()
    field_label = _clause_field_label(field_key)
    if child_title and field_key == "text":
        return child_title
    if description and field_key == "text":
        return description
    if description:
        return f"{description}｜{field_label}"
    return f"{clause_key}｜{field_label}"


def _build_clause_library_entry(
    *,
    source: dict,
    clause_key: str,
    clause: dict,
    path_parts: list[str],
    field_key: str,
    content: str,
    scene_key: str,
    target_node: dict | None = None,
    readonly: bool = False,
    readonly_reason: str = "",
) -> dict:
    target_node = target_node if isinstance(target_node, dict) else {}
    description = str((clause or {}).get("description") or "").strip()
    location = str((clause or {}).get("location") or target_node.get("location") or "").strip()
    condition = str(target_node.get("condition") or (clause or {}).get("condition") or "").strip()
    applicable_to = target_node.get("applicable_to") or (clause or {}).get("applicable_to") or []
    if isinstance(applicable_to, list) and applicable_to:
        condition = f"{condition}；适用：{'、'.join(str(item) for item in applicable_to)}".strip("；")
    applicability = "；".join(part for part in (_clause_scene_label(scene_key), condition) if part)
    path_id = _encode_clause_path_id(source["document_type"], path_parts)
    raw_path = ".".join(path_parts)
    return {
        "path_id": path_id,
        "raw_path": raw_path,
        "document_type": source["document_type"],
        "document_label": source["document_label"],
        "file_name": source["path"].name,
        "clause_key": clause_key,
        "title": _clause_entry_title(clause_key, clause, field_key, target_node),
        "chapter": _clause_chapter_hint(location, description, condition),
        "location": location,
        "scene_key": scene_key,
        "scene_label": _clause_scene_label(scene_key),
        "field_key": field_key,
        "field_label": _clause_field_label(field_key),
        "condition": condition,
        "applicability": applicability,
        "source_note": _clause_source_note(clause, target_node),
        "content": str(content or ""),
        "readonly": bool(readonly),
        "readonly_reason": readonly_reason,
        "search_text": " ".join(
            str(part or "")
            for part in (
                source["document_label"],
                clause_key,
                description,
                location,
                scene_key,
                _clause_scene_label(scene_key),
                applicability,
                field_key,
                _clause_field_label(field_key),
                condition,
                content,
            )
        ),
    }


def _append_clause_entries_for_node(entries: list[dict], source: dict, clause_key: str, clause: dict) -> None:
    def append_fields(target_node: dict, base_path: list[str], scene_key: str) -> bool:
        appended = False
        for field_key, value in target_node.items():
            if _is_editable_clause_field(field_key, value):
                entries.append(
                    _build_clause_library_entry(
                        source=source,
                        clause_key=clause_key,
                        clause=clause,
                        path_parts=[*base_path, field_key],
                        field_key=field_key,
                        content=value,
                        scene_key=scene_key,
                        target_node=target_node,
                    )
                )
                appended = True
        return appended

    variants = clause.get("variants")
    if isinstance(variants, dict):
        for variant_key, variant_node in variants.items():
            if isinstance(variant_node, dict):
                if not append_fields(variant_node, ["clauses", clause_key, "variants", str(variant_key)], str(variant_key)):
                    entries.append(
                        _build_clause_library_entry(
                            source=source,
                            clause_key=clause_key,
                            clause=clause,
                            path_parts=["clauses", clause_key, "variants", str(variant_key)],
                            field_key="",
                            content="",
                            scene_key=str(variant_key),
                            target_node=variant_node,
                            readonly=True,
                            readonly_reason="复杂结构，只读展示，暂不支持直接编辑",
                        )
                    )
            else:
                entries.append(
                    _build_clause_library_entry(
                        source=source,
                        clause_key=clause_key,
                        clause=clause,
                        path_parts=["clauses", clause_key, "variants", str(variant_key)],
                        field_key="",
                        content="",
                        scene_key=str(variant_key),
                        readonly=True,
                        readonly_reason="列表或非文本结构，只读展示",
                    )
                )
        return

    child_clauses = clause.get("clauses")
    if isinstance(child_clauses, dict):
        for child_key, child_node in child_clauses.items():
            if isinstance(child_node, dict):
                if not append_fields(child_node, ["clauses", clause_key, "clauses", str(child_key)], str(child_key)):
                    entries.append(
                        _build_clause_library_entry(
                            source=source,
                            clause_key=clause_key,
                            clause=clause,
                            path_parts=["clauses", clause_key, "clauses", str(child_key)],
                            field_key="",
                            content="",
                            scene_key=str(child_key),
                            target_node=child_node,
                            readonly=True,
                            readonly_reason="子条款没有可编辑正文",
                        )
                    )
        return

    if isinstance(clause.get("text"), str):
        entries.append(
            _build_clause_library_entry(
                source=source,
                clause_key=clause_key,
                clause=clause,
                path_parts=["clauses", clause_key, "text"],
                field_key="text",
                content=clause.get("text", ""),
                scene_key="DEFAULT",
                target_node=clause,
            )
        )
        return

    entries.append(
        _build_clause_library_entry(
            source=source,
            clause_key=clause_key,
            clause=clause,
            path_parts=["clauses", clause_key],
            field_key="",
            content="",
            scene_key="DEFAULT",
            target_node=clause,
            readonly=True,
            readonly_reason="目录或复杂配置，只读展示",
        )
    )


def _build_clause_library_catalog() -> dict:
    entries = []
    sources = []
    for source in _clause_library_sources():
        sources.append({
            "document_type": source["document_type"],
            "document_label": source["document_label"],
            "file_name": source["path"].name,
        })
        if not source["path"].exists():
            continue
        data = json.loads(source["path"].read_text(encoding="utf-8"))
        clauses = data.get("clauses", {}) if isinstance(data, dict) else {}
        if not isinstance(clauses, dict):
            continue
        for clause_key, clause in clauses.items():
            if isinstance(clause, dict):
                _append_clause_entries_for_node(entries, source, str(clause_key), clause)
            else:
                entries.append(
                    _build_clause_library_entry(
                        source=source,
                        clause_key=str(clause_key),
                        clause={},
                        path_parts=["clauses", str(clause_key)],
                        field_key="",
                        content="",
                        scene_key="DEFAULT",
                        readonly=True,
                        readonly_reason="非对象结构，只读展示",
                    )
                )
    entries.sort(key=lambda item: (item["document_type"], item["chapter"], item["clause_key"], item["scene_key"], item["field_key"]))
    return {"success": True, "sources": sources, "entries": entries}


def _save_clause_library_field(path_id: str, content: str) -> dict:
    document_type, path_parts = _decode_clause_path_id(path_id)
    if len(path_parts) < 3 or path_parts[0] != "clauses":
        raise ValueError("只能保存条款正文路径")
    field_key = path_parts[-1]
    source = _clause_library_source_for(document_type)
    if not source["path"].is_file():
        raise FileNotFoundError(f"条款库文件不存在：{source['path'].name}")
    data = json.loads(source["path"].read_text(encoding="utf-8"))
    target = data
    for part in path_parts[:-1]:
        if not isinstance(target, dict) or part not in target:
            raise ValueError("条款路径不存在")
        target = target[part]
    if not isinstance(target, dict) or not _is_editable_clause_field(field_key, target.get(field_key)):
        raise ValueError("该条款字段不可编辑")
    target[field_key] = str(content or "")

    tmp_path = source["path"].with_name(f".{source['path'].name}.tmp")
    try:
        tmp_path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        os.replace(tmp_path, source["path"])
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass
    _reload_generation_engines()
    return {
        "success": True,
        "message": "条款正文已保存",
        "entry": _build_clause_library_entry(
            source=source,
            clause_key=path_parts[1],
            clause=data.get("clauses", {}).get(path_parts[1], {}),
            path_parts=path_parts,
            field_key=field_key,
            content=str(content or ""),
            scene_key=path_parts[-2] if len(path_parts) >= 4 else "DEFAULT",
            target_node=target,
        ),
    }


def _editor_group_for_relative_path(relative_path: str) -> str:
    normalized = str(relative_path or "").replace("\\", "/").strip("/")
    target = (BASE_DIR / normalized).resolve()
    filename = Path(normalized).name
    if target in {TEMPLATE_MD.resolve(), PROSPECTUS_TEMPLATE_MD.resolve()}:
        return "模板"
    if target == SCHEMA_JSON.resolve():
        return "变量库"
    if target in {CLAUSES_JSON.resolve(), PROSPECTUS_CLAUSES_JSON.resolve()}:
        return "条款库"
    if "模板" in filename:
        return "模板"
    if "变量定义" in filename:
        return "变量库"
    if "条款" in filename:
        return "条款库"
    return "规则库"


def _iter_editable_knowledge_files() -> list[dict]:
    def _build_entry(path_obj: Path) -> dict:
        stat = path_obj.stat()
        relative_path = path_obj.relative_to(BASE_DIR).as_posix()
        return {
            "name": path_obj.name,
            "path": relative_path,
            "group": _editor_group_for_relative_path(relative_path),
            "size": stat.st_size,
            "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
        }

    files = []
    for path_obj in _editable_generation_file_paths():
        if path_obj.is_file() and path_obj.suffix in ALLOWED_SUFFIXES:
            files.append(_build_entry(path_obj))
    return sorted(files, key=lambda item: (item["group"], item["name"]))


def _resolve_editable_knowledge_file(relative_path: str) -> tuple[Path, str]:
    normalized = str(relative_path or "").replace("\\", "/").strip()
    normalized = normalized.lstrip("/")
    rel_path = Path(normalized)

    if not normalized or rel_path.is_absolute() or ".." in rel_path.parts:
        abort(403)
    if rel_path.suffix not in ALLOWED_SUFFIXES:
        abort(403)

    target = (BASE_DIR / rel_path).resolve()

    if BASE_DIR.resolve() not in target.parents:
        abort(403)
    allowed_targets = {
        path_obj.resolve()
        for path_obj in _editable_generation_file_paths()
        if path_obj.is_file()
    }
    if target not in allowed_targets:
        abort(403)
    if not target.is_file():
        abort(404)

    return target, rel_path.as_posix()


def _reference_prospectus_docx_candidates(variant_key: str) -> list[Path]:
    variant = variant_key if variant_key in REFERENCE_PROSPECTUS_DOCX_MAP else "SSE_CROSS"
    candidates = []
    env_dir = str(os.getenv("ETF_REFERENCE_PROSPECTUS_DIR") or "").strip()
    if env_dir:
        candidates.append(Path(env_dir) / f"{variant}.docx")
    candidates.append(REFERENCE_PROSPECTUS_DOCX_MAP[variant])
    packaged_candidate = PACKAGED_REFERENCE_PROSPECTUS_DOCX_MAP[variant]
    if packaged_candidate not in candidates:
        candidates.append(packaged_candidate)
    for legacy_candidate in LEGACY_REFERENCE_PROSPECTUS_SOURCE_MAP.get(
        variant, LEGACY_REFERENCE_PROSPECTUS_SOURCE_MAP["SSE_CROSS"]
    ):
        if legacy_candidate not in candidates:
            candidates.append(legacy_candidate)
    return candidates


def _score_reference_prospectus_docx(candidate: Path) -> tuple[int, int, int, int]:
    cache_key = str(candidate)
    cached = _REFERENCE_PROSPECTUS_DOCX_QUALITY_CACHE.get(cache_key)
    if cached is not None:
        return cached

    if not candidate.exists():
        score = (-1, -1, -1, -1)
        _REFERENCE_PROSPECTUS_DOCX_QUALITY_CACHE[cache_key] = score
        return score

    try:
        from docx import Document

        doc = Document(str(candidate))
    except Exception:
        score = (-1, -1, -1, -1)
        _REFERENCE_PROSPECTUS_DOCX_QUALITY_CACHE[cache_key] = score
        return score

    heading_styles = {"heading 1", "heading 2", "标题 1", "标题1", "标题 2", "标题2"}
    heading_re = re.compile(r"^(第[一二三四五六七八九十百]+章\s+.+|[一二三四五六七八九十百]+、.+)$")
    intro_markers = {
        "申购赎回清单的格式举例如下：",
        "T日申购赎回清单的格式举例如下：",
    }

    heading_count = 0
    intro_count = 0
    toc_count = 0
    for paragraph in doc.paragraphs:
        text_value = (paragraph.text or "").strip()
        if not text_value:
            continue
        style_name = ""
        try:
            style_name = (paragraph.style.name or "").strip().lower()
        except Exception:
            style_name = ""
        if style_name in heading_styles or heading_re.match(text_value):
            heading_count += 1
        if text_value == "目录":
            toc_count += 1
        if text_value in intro_markers:
            intro_count += 1

    score = (intro_count, len(doc.tables), heading_count, toc_count)
    _REFERENCE_PROSPECTUS_DOCX_QUALITY_CACHE[cache_key] = score
    return score


def _resolve_reference_prospectus_docx(variant_key: str) -> Path:
    candidates = _reference_prospectus_docx_candidates(variant_key)
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def _product_summary_template_docx_candidates() -> list[Path]:
    candidates = []
    env_file = str(os.getenv("ETF_PRODUCT_SUMMARY_TEMPLATE_DOCX") or "").strip()
    if env_file:
        candidates.append(Path(env_file))
    candidates.append(PRODUCT_SUMMARY_TEMPLATE_DOCX)
    candidates.extend(LEGACY_PRODUCT_SUMMARY_TEMPLATE_DOCX_CANDIDATES)
    return candidates


def _resolve_product_summary_template_docx() -> Path:
    candidates = _product_summary_template_docx_candidates()
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


class MissingProspectusReferenceAssetsError(RuntimeError):
    def __init__(self, missing_assets: list[str]):
        self.missing_assets = missing_assets
        super().__init__("缺少招募说明书导出所需参考资产")


def _resolve_payment_clause_value(value, variants):
    """Return (enum_key, rendered_text) for either a legacy enum or verbatim clause text."""
    raw = "" if value is None else str(value).strip()
    if not raw:
        return "", ""
    if raw in variants:
        return raw, variants[raw]["text"]
    return "", raw


def _iter_docx_blocks(parent):
    from docx.document import Document as DocumentObject
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
        parent_obj = parent
    else:
        parent_elm = parent._tc
        parent_obj = parent

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, parent_obj))
        elif isinstance(child, CT_Tbl):
            yield ("table", Table(child, parent_obj))


_LAYOUT_BLANK_RE = re.compile(r"[\s\u00a0\u3000\u200b\u200c\u200d\ufeff]+")


def _normalize_layout_text(value: str) -> str:
    return (
        str(value or "")
        .replace("\u00a0", " ")
        .replace("\u3000", " ")
        .replace("\u200b", "")
        .replace("\u200c", "")
        .replace("\u200d", "")
        .replace("\ufeff", "")
    )


def _is_layout_blank_line(value: str) -> bool:
    return not _LAYOUT_BLANK_RE.sub("", str(value or ""))


def _is_markdown_table_separator_line(value: str) -> bool:
    text = str(value or "").strip()
    if "|" not in text:
        return False
    cells = [cell.strip() for cell in text.strip("|").split("|")]
    return len(cells) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _clean_docx_cell_text(value: str) -> str:
    return re.sub(r"\s+", " ", _normalize_layout_text(value).replace("\r", "\n")).strip()


def _docx_text_local_name(tag: str) -> str:
    return str(tag or "").rsplit("}", 1)[-1]


def _docx_paragraph_visible_text_with_math(paragraph) -> str:
    """Read DOCX text as if tracked insertions are accepted and deletions are rejected."""
    from docx.oxml.ns import qn

    parts = []

    def walk(node):
        name = _docx_text_local_name(getattr(node, "tag", ""))
        if name in {"del", "moveFrom"}:
            return
        if name == "r":
            rpr = node.find(qn("w:rPr"))
            if rpr is not None and rpr.find(qn("w:vanish")) is not None:
                return
        if name == "t" and node.text:
            parts.append(str(node.text))
            return
        if name == "tab":
            parts.append("\t")
            return
        if name in {"br", "cr"}:
            parts.append("\n")
            return
        if name in {"oMath", "oMathPara"}:
            math_text = _extract_math_plain_text_from_node(node)
            if math_text:
                parts.append(math_text)
            return
        for child in list(node):
            walk(child)

    walk(paragraph._p)
    return _normalize_extracted_docx_line("".join(parts))


def _docx_table_to_markdown_lines(table) -> list[str]:
    def cell_visible_text(cell):
        cell_lines = []
        for block_type, block in _iter_docx_blocks(cell):
            if block_type == "paragraph":
                text = _paragraph_text_with_math(block)
                if not _is_layout_blank_line(text):
                    cell_lines.append(text)
            elif block_type == "table":
                cell_lines.extend(_docx_table_to_markdown_lines(block))
        return _clean_docx_cell_text("\n".join(cell_lines))

    rows = []
    for row in table.rows:
        values = [cell_visible_text(cell) for cell in row.cells]
        if any(not _is_layout_blank_line(value) for value in values):
            rows.append(values)
    if not rows:
        return []
    return [f"|{'|'.join(values)}|" for values in rows]


def _normalize_extracted_docx_line(text: str) -> str:
    normalized = _normalize_layout_text(text).replace("\t", "")
    normalized = re.sub(r"^([一二三四五六七八九十百千]+)、\s+", r"\1、", normalized)
    if normalized == "成份股信息内容":
        normalized = "成份股信息内容。"
    return normalized.strip()


def _extract_math_plain_text_from_node(node) -> str:
    raw_text = "".join(
        str(child.text or "")
        for child in node.iter()
        if str(getattr(child, "tag", "")).endswith("}t") and child.text
    ).strip()
    if not raw_text:
        return ""
    if "投资者的认购份额" in raw_text:
        return ProspectusEngine._stock_subscription_formula_text()
    if "现金替代比例" in raw_text:
        return ProspectusEngine._cash_substitution_ratio_formula_text()
    return raw_text


def _paragraph_text_with_math(paragraph) -> str:
    return _docx_paragraph_visible_text_with_math(paragraph)


def _extract_docx_text_with_tables(docx_path: str) -> str:
    from docx import Document

    doc = Document(docx_path)
    section7_titles = set(ProspectusEngine._section7_format_titles())
    is_generated_prospectus = any(str(paragraph.text or "").strip().endswith("招募说明书") for paragraph in doc.paragraphs)
    toc_lines = [
        _normalize_extracted_docx_line(paragraph.text)
        for paragraph in doc.paragraphs
        if getattr(getattr(paragraph, "style", None), "name", "") == "Heading 2" and paragraph.text.strip()
    ]
    lines = []
    last_text = ""
    for block_type, block in _iter_docx_blocks(doc):
        if block_type == "paragraph":
            text = _paragraph_text_with_math(block)
            if text:
                if text == "右键更新目录":
                    continue
                if text == "目录" and toc_lines:
                    lines.append(text)
                    lines.extend(toc_lines)
                    continue
                lines.append(text)
                last_text = text
            continue

        if last_text.rstrip("。") in section7_titles:
            continue
        table_lines = _docx_table_to_markdown_lines(block)
        if table_lines:
            if is_generated_prospectus and len(table_lines) >= 1:
                header_cells = [cell.strip() for cell in table_lines[0].strip("|").split("|")]
                separator = f"|{'|'.join('---' for _ in header_cells)}|"
                if separator not in table_lines:
                    table_lines = [table_lines[0], separator, *table_lines[1:]]
            lines.extend(table_lines)

    return "\n".join(lines)


def _empty_custodian_summary_state() -> dict:
    return {
        "filename": "",
        "text": "",
        "sections": [],
    }


_prospectus_store = {
    "custodian_summary": _empty_custodian_summary_state(),
}


def _get_custodian_summary_state() -> dict:
    state = _prospectus_store.get("custodian_summary")
    if not isinstance(state, dict):
        state = _empty_custodian_summary_state()
        _prospectus_store["custodian_summary"] = state
    return state


def _set_custodian_summary_state(filename: str, text: str, sections: list[dict]) -> dict:
    state = {
        "filename": str(filename or "").strip(),
        "text": str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip(),
        "sections": list(sections or []),
    }
    _prospectus_store["custodian_summary"] = state
    return state


def _clear_custodian_summary_state() -> dict:
    state = _empty_custodian_summary_state()
    _prospectus_store["custodian_summary"] = state
    return state


def _serialize_custodian_summary_state(state: dict | None = None) -> dict:
    current = state or _get_custodian_summary_state()
    sections = list(current.get("sections") or [])
    filename = str(current.get("filename") or "").strip()
    return {
        "uploaded": bool(filename and sections),
        "filename": filename,
        "section_count": len(sections),
        "section_titles": [str(section.get("heading") or "").strip() for section in sections[:8]],
    }


def _clean_custodian_summary_sections(sections: list[dict]) -> list[dict]:
    cleaned = []
    seen = set()
    for section in sections or []:
        heading = str((section or {}).get("heading") or "").strip()
        content = str((section or {}).get("content") or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not heading or not content:
            continue
        heading_key = _review_heading_key(heading)
        if not heading_key or heading_key in seen:
            continue
        seen.add(heading_key)
        cleaned.append({
            "heading": heading,
            "content": content,
        })
    return cleaned


def _extract_custodian_summary_sections_from_docx(docx_path: str) -> tuple[str, list[dict]]:
    from docx import Document

    doc = Document(docx_path)
    text = _extract_docx_text_with_tables(docx_path)
    sections = _clean_custodian_summary_sections(_split_prospectus_sections(text, doc=doc))
    return text, sections


def _normalize_review_doc_title_key(value: str) -> str:
    text = unicodedata.normalize("NFKC", str(value or "")).strip()
    if not text:
        return ""

    if text.lower().endswith((".doc", ".docx")):
        text = Path(text).stem

    text = re.sub(
        r"[（(][^）)]*(?:草案|修订|更新|版本|版|稿|llinks|\d{4}|\d{1,2}[月日号._/-]?)[^）)]*[）)]",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"[-_—–－ ]+(?:草案|修订稿|修订版|更新版|llinks)$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"[-_—–－ ]+\d{3,8}(?:[_-]?[A-Za-z]+)?$", "", text)
    text = re.sub(r"(基金合同|招募说明书|托管协议|合同|说明书)$", "", text).strip()
    text = re.sub(r"[《》“”\"'`·•]", "", text)
    return _review_heading_key(text)


def _collect_review_doc_header_entries(doc) -> list[dict]:
    entries = []
    seen = set()

    for section_index, section in enumerate(getattr(doc, "sections", []), start=1):
        for kind, header in (("default", section.header), ("first_page", section.first_page_header)):
            lines = []
            for paragraph in getattr(header, "paragraphs", []):
                text = re.sub(r"\s+", " ", (paragraph.text or "").strip())
                if text:
                    lines.append(text)
            header_text = " ".join(lines).strip()
            header_key = _review_heading_key(header_text)
            if not header_key or header_key in seen:
                continue
            seen.add(header_key)
            entries.append(
                {
                    "section": section_index,
                    "kind": kind,
                    "text": header_text,
                }
            )

    return entries


def _paragraph_review_segments(paragraph) -> list[dict]:
    from docx.oxml.ns import qn

    segments = []
    current = []
    p = paragraph._p
    ppr = getattr(p, "pPr", None)
    if ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None:
        segments.append({"kind": "page_break"})

    def flush_text():
        text = re.sub(r"\s+", " ", "".join(current)).strip()
        if text:
            segments.append({"kind": "text", "text": text})
        current.clear()

    for node in p.iter():
        tag = node.tag
        if tag == qn("w:t") and node.text:
            current.append(node.text)
        elif tag == qn("w:tab"):
            current.append("\t")
        elif tag == qn("w:cr"):
            current.append("\n")
        elif tag == qn("w:br"):
            if node.get(qn("w:type")) == "page":
                flush_text()
                segments.append({"kind": "page_break"})
            else:
                current.append("\n")
        elif tag == qn("w:lastRenderedPageBreak"):
            flush_text()
            segments.append({"kind": "page_break"})

    flush_text()

    if ppr is not None:
        sect_pr = ppr.find(qn("w:sectPr"))
        if sect_pr is not None:
            sect_type = sect_pr.find(qn("w:type"))
            sect_type_val = sect_type.get(qn("w:val")) if sect_type is not None else "nextPage"
            if sect_type_val == "nextPage":
                segments.append({"kind": "page_break"})

    return segments


def _detect_review_blank_pages(doc) -> list[int]:
    pages = [[]]
    current_page = pages[0]

    for block_type, block in _iter_docx_blocks(doc):
        if block_type == "table":
            table_lines = _docx_table_to_markdown_lines(block)
            if table_lines:
                current_page.extend(line for line in table_lines if str(line or "").strip())
            continue

        for segment in _paragraph_review_segments(block):
            if segment["kind"] == "text":
                current_page.append(segment["text"])
                continue
            pages.append([])
            current_page = pages[-1]

    blank_pages = []
    for page_number, snippets in enumerate(pages, start=1):
        if any(re.sub(r"[\s\u00a0]+", "", str(snippet or "")) for snippet in snippets):
            continue
        blank_pages.append(page_number)
    return blank_pages


def _extract_review_doc_metadata(doc, *, filename: str = "", title_reference: str = "") -> dict:
    return {
        "filename": str(filename or ""),
        "title_reference": str(title_reference or ""),
        "header_entries": _collect_review_doc_header_entries(doc),
        "blank_pages": _detect_review_blank_pages(doc),
    }


def _clear_xml_children(node):
    for child in list(node):
        node.remove(child)


def _reset_header_footer_part(part):
    try:
        part.is_linked_to_previous = False
    except Exception:
        pass
    _clear_xml_children(part._element)
    paragraph = part.add_paragraph()
    _clear_xml_children(paragraph._p)
    return paragraph


def _prune_unreferenced_docx_header_footer_parts(docx_bytes: bytes) -> bytes:
    import posixpath
    import xml.etree.ElementTree as ET

    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    content_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    def _local_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as src:
        parts = {name: src.read(name) for name in src.namelist()}

    document_root = ET.fromstring(parts["word/document.xml"])
    rels_root = ET.fromstring(parts["word/_rels/document.xml.rels"])
    content_root = ET.fromstring(parts["[Content_Types].xml"])

    used_rel_ids = set()
    for sect_pr in document_root.iter(f"{{{word_ns}}}sectPr"):
        for child in sect_pr:
            if _local_name(child.tag) not in {"headerReference", "footerReference"}:
                continue
            rel_id = child.get(f"{{{office_rel_ns}}}id")
            if rel_id:
                used_rel_ids.add(rel_id)

    used_targets = set()
    for rel in list(rels_root):
        rel_type = rel.get("Type", "")
        is_header_footer = rel_type.endswith("/header") or rel_type.endswith("/footer")
        if not is_header_footer:
            continue
        rel_id = rel.get("Id")
        target = rel.get("Target", "")
        normalized_target = posixpath.normpath(posixpath.join("word", target)).replace("\\", "/")
        if rel_id in used_rel_ids:
            used_targets.add(normalized_target)
        else:
            rels_root.remove(rel)

    for override in list(content_root):
        if _local_name(override.tag) != "Override":
            continue
        part_name = override.get("PartName", "").lstrip("/")
        if not (part_name.startswith("word/header") or part_name.startswith("word/footer")):
            continue
        if part_name not in used_targets:
            content_root.remove(override)

    for part_name in list(parts):
        if part_name.startswith("word/header") or part_name.startswith("word/footer"):
            if part_name not in used_targets:
                parts.pop(part_name, None)

    parts["word/_rels/document.xml.rels"] = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
    parts["[Content_Types].xml"] = ET.tostring(content_root, encoding="utf-8", xml_declaration=True)

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for name, data in parts.items():
            dst.writestr(name, data)
    return output.getvalue()


def _align_cover_section_references_with_template(docx_bytes: bytes, template_docx: Path | None) -> bytes:
    import posixpath
    import xml.etree.ElementTree as ET

    if not template_docx or not template_docx.exists():
        return docx_bytes

    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    package_rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"

    def _local_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    def _first_cover_sectpr(document_root):
        body = document_root.find(f"{{{word_ns}}}body")
        if body is None:
            return None
        for child in body:
            if _local_name(child.tag) != "p":
                continue
            p_pr = child.find(f"{{{word_ns}}}pPr")
            if p_pr is None:
                continue
            sect_pr = p_pr.find(f"{{{word_ns}}}sectPr")
            if sect_pr is not None:
                return sect_pr
        return None

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as src:
        parts = {name: src.read(name) for name in src.namelist()}
    with zipfile.ZipFile(template_docx) as src:
        template_parts = {name: src.read(name) for name in src.namelist()}

    document_root = ET.fromstring(parts["word/document.xml"])
    rels_root = ET.fromstring(parts["word/_rels/document.xml.rels"])
    content_root = ET.fromstring(parts["[Content_Types].xml"])
    template_document_root = ET.fromstring(template_parts["word/document.xml"])
    template_rels_root = ET.fromstring(template_parts["word/_rels/document.xml.rels"])
    template_content_root = ET.fromstring(template_parts["[Content_Types].xml"])

    cover_sect_pr = _first_cover_sectpr(document_root)
    template_cover_sect_pr = _first_cover_sectpr(template_document_root)
    if cover_sect_pr is None or template_cover_sect_pr is None:
        return docx_bytes

    for child in list(cover_sect_pr):
        if _local_name(child.tag) in {"headerReference", "footerReference"}:
            cover_sect_pr.remove(child)

    existing_rel_ids = {
        rel.get("Id")
        for rel in rels_root
        if _local_name(rel.tag) == "Relationship" and rel.get("Id")
    }

    def _next_rel_id(preferred: str) -> str:
        if preferred not in existing_rel_ids:
            existing_rel_ids.add(preferred)
            return preferred
        index = 1
        while True:
            candidate = f"rId{100 + index}"
            if candidate not in existing_rel_ids:
                existing_rel_ids.add(candidate)
                return candidate
            index += 1

    rel_map = {
        rel.get("Id"): rel
        for rel in template_rels_root
        if _local_name(rel.tag) == "Relationship"
    }
    template_overrides = {
        override.get("PartName", ""): override
        for override in template_content_root
        if _local_name(override.tag) == "Override"
    }

    insert_index = 0
    for ref in list(template_cover_sect_pr):
        local_name = _local_name(ref.tag)
        if local_name not in {"headerReference", "footerReference"}:
            continue
        template_rel_id = ref.get(f"{{{office_rel_ns}}}id")
        template_rel = rel_map.get(template_rel_id)
        if template_rel is None:
            continue
        target = template_rel.get("Target", "")
        full_target = posixpath.normpath(posixpath.join("word", target)).replace("\\", "/")
        if full_target not in template_parts:
            continue
        new_rel_id = _next_rel_id(template_rel_id or "rId200")
        new_rel = ET.Element(f"{{{package_rel_ns}}}Relationship")
        for key, value in template_rel.attrib.items():
            new_rel.set(key, new_rel_id if key == "Id" else value)
        rels_root.append(new_rel)
        parts[full_target] = template_parts[full_target]

        part_name = "/" + full_target
        if not any(
            _local_name(node.tag) == "Override" and node.get("PartName") == part_name
            for node in content_root
        ):
            template_override = template_overrides.get(part_name)
            if template_override is not None:
                content_root.append(ET.fromstring(ET.tostring(template_override, encoding="utf-8")))

        new_ref = ET.fromstring(ET.tostring(ref, encoding="utf-8"))
        new_ref.set(f"{{{office_rel_ns}}}id", new_rel_id)
        cover_sect_pr.insert(insert_index, new_ref)
        insert_index += 1

    parts["word/document.xml"] = ET.tostring(document_root, encoding="utf-8", xml_declaration=True)
    parts["word/_rels/document.xml.rels"] = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
    parts["[Content_Types].xml"] = ET.tostring(content_root, encoding="utf-8", xml_declaration=True)

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for name, data in parts.items():
            dst.writestr(name, data)
    return output.getvalue()


def _append_run_properties(run, OxmlElement, qn, ascii_font=None, eastasia_font=None, size=None, bold=False):
    rPr = OxmlElement("w:rPr")
    if ascii_font or eastasia_font:
        rFonts = OxmlElement("w:rFonts")
        if ascii_font:
            rFonts.set(qn("w:ascii"), ascii_font)
            rFonts.set(qn("w:hAnsi"), ascii_font)
        if eastasia_font:
            rFonts.set(qn("w:eastAsia"), eastasia_font)
        rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size))
        rPr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(size))
        rPr.append(sz_cs)
    run.append(rPr)


def _append_page_field(paragraph, OxmlElement, qn, ascii_font="Times New Roman", eastasia_font="Times New Roman", size=18):
    fld_begin = OxmlElement("w:r")
    _append_run_properties(fld_begin, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    fld_begin.append(fld_char_begin)

    instr_run = OxmlElement("w:r")
    _append_run_properties(instr_run, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = "PAGE   \\* MERGEFORMAT"
    instr_run.append(instr_text)

    fld_sep = OxmlElement("w:r")
    _append_run_properties(fld_sep, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_sep.append(fld_char_sep)

    fld_text = OxmlElement("w:r")
    _append_run_properties(fld_text, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_text_run = OxmlElement("w:t")
    fld_text_run.text = "1"
    fld_text.append(fld_text_run)

    fld_end = OxmlElement("w:r")
    _append_run_properties(fld_end, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    fld_end.append(fld_char_end)

    paragraph._p.append(fld_begin)
    paragraph._p.append(instr_run)
    paragraph._p.append(fld_sep)
    paragraph._p.append(fld_text)
    paragraph._p.append(fld_end)


def _set_section_page_numbers(section, OxmlElement, qn, start=None):
    processed_parts = set()
    for footer in (section.footer, section.first_page_footer):
        footer_key = id(footer._element)
        if footer_key in processed_parts:
            continue
        processed_parts.add(footer_key)

        paragraph = _reset_header_footer_part(footer)
        pPr = paragraph._p.get_or_add_pPr()
        _clear_xml_children(pPr)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        _append_page_field(paragraph, OxmlElement, qn, ascii_font="Times New Roman", eastasia_font="Times New Roman", size=18)

    sectPr = section._sectPr
    for node in list(sectPr.findall(qn("w:pgNumType"))):
        sectPr.remove(node)
    if start is not None:
        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:start"), str(start))
        sectPr.append(pg_num_type)


def _clear_header_footer_part(part):
    _reset_header_footer_part(part)


def _remove_section_refs(section, qn):
    sectPr = section._sectPr
    for tag in ("w:headerReference", "w:footerReference"):
        for node in list(sectPr.findall(qn(tag))):
            sectPr.remove(node)


def _set_cover_section_no_page_number(section, OxmlElement, qn):
    section.different_first_page_header_footer = True
    for part in (section.header, section.first_page_header):
        _clear_header_footer_part(part)
    for footer in (section.footer, section.first_page_footer):
        paragraph = _reset_header_footer_part(footer)
        pPr = paragraph._p.get_or_add_pPr()
        _clear_xml_children(pPr)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        _append_page_field(paragraph, OxmlElement, qn, ascii_font="Times New Roman", eastasia_font="Times New Roman", size=18)
    _remove_section_refs(section, qn)


def _finalize_doc_page_numbers(doc, OxmlElement, qn, body_start_index=1):
    sections = list(doc.sections)
    if not sections:
        return

    body_start_index = max(1, min(body_start_index, len(sections)))
    for section in sections[:body_start_index]:
        _set_cover_section_no_page_number(section, OxmlElement, qn)

    body_sections = sections[body_start_index:] or sections[-1:]
    for idx, section in enumerate(body_sections):
        _set_section_page_numbers(section, OxmlElement, qn, start=1 if idx == 0 else None)


def _set_update_fields_on_open(document, OxmlElement, qn):
    settings = document.settings.element
    for node in list(settings.findall(qn("w:updateFields"))):
        settings.remove(node)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def _update_docx_fields_with_word(docx_bytes: bytes) -> bytes:
    """Update TOC/page fields with local Microsoft Word when available."""
    if os.environ.get("DISABLE_WORD_FIELD_UPDATE") == "1" or os.environ.get("PYTEST_CURRENT_TEST") or os.name != "nt":
        return docx_bytes
    import tempfile

    tmp_path = ""
    word = None
    document = None
    try:
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="field_update_")
        tmp.write(docx_bytes)
        tmp.close()
        tmp_path = tmp.name

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(tmp_path, ReadOnly=False, AddToRecentFiles=False)
        for toc in document.TablesOfContents:
            toc.Update()
        document.Fields.Update()
        for section in document.Sections:
            for header in section.Headers:
                header.Range.Fields.Update()
            for footer in section.Footers:
                footer.Range.Fields.Update()
        document.Save()
        document.Close(False)
        document = None
        return Path(tmp_path).read_bytes()
    except Exception:
        logger.exception("Failed to update DOCX fields with Microsoft Word")
        return docx_bytes
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            import pythoncom

            pythoncom.CoUninitialize()
        except Exception:
            pass
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _mark_docx_fields_for_update(docx_bytes: bytes) -> bytes:
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    replacements: dict[str, bytes | str] = {}
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        related_parts = [
            info.filename
            for info in archive.infolist()
            if info.filename.startswith(("word/header", "word/footer")) and info.filename.endswith(".xml")
        ]
    for part_name in ["word/document.xml", *related_parts]:
        xml_text = _read_docx_xml_part(docx_bytes, part_name)
        if not xml_text:
            continue
        try:
            root = ET.fromstring(xml_text)
        except ET.ParseError:
            continue
        changed = False
        for node in root.iter():
            if _docx_local_name(node.tag) == "fldChar":
                node.set(f"{{{_DOCX_W_NS}}}dirty", "true")
                changed = True
        if changed:
            replacements[part_name] = _serialize_docx_document_xml(ET, root, xml_text)

    settings_xml = _read_docx_xml_part(docx_bytes, "word/settings.xml")
    if settings_xml:
        try:
            settings_root = ET.fromstring(settings_xml)
        except ET.ParseError:
            settings_root = ET.Element(f"{{{_DOCX_W_NS}}}settings")
    else:
        settings_root = ET.Element(f"{{{_DOCX_W_NS}}}settings")
    for node in list(settings_root):
        if _docx_local_name(node.tag) == "updateFields":
            settings_root.remove(node)
    update_fields = ET.Element(f"{{{_DOCX_W_NS}}}updateFields")
    update_fields.set(f"{{{_DOCX_W_NS}}}val", "true")
    settings_root.append(update_fields)
    replacements["word/settings.xml"] = _serialize_docx_document_xml(ET, settings_root, settings_xml or "")
    return _repack_docx_parts(docx_bytes, replacements)


def _finalize_review_docx_for_export(docx_bytes: bytes) -> bytes:
    marked = _mark_docx_fields_for_update(docx_bytes)
    return _update_docx_fields_with_word(marked)


def _append_word_toc_field(
    paragraph,
    OxmlElement,
    qn,
    *,
    levels: str,
    placeholder: str = "右键更新目录",
    ascii_font: str = "Times New Roman",
    eastasia_font: str = "宋体",
    size: int = 24,
):
    instruction = f' TOC \\\\o "{levels}" \\\\h \\\\z \\\\u '

    begin = OxmlElement("w:r")
    _append_run_properties(
        begin,
        OxmlElement,
        qn,
        ascii_font=ascii_font,
        eastasia_font=eastasia_font,
        size=size,
    )
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    begin.append(fld_begin)
    paragraph._p.append(begin)

    instr_run = OxmlElement("w:r")
    _append_run_properties(
        instr_run,
        OxmlElement,
        qn,
        ascii_font=ascii_font,
        eastasia_font=eastasia_font,
        size=size,
    )
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = instruction
    instr_run.append(instr_text)
    paragraph._p.append(instr_run)

    sep = OxmlElement("w:r")
    _append_run_properties(
        sep,
        OxmlElement,
        qn,
        ascii_font=ascii_font,
        eastasia_font=eastasia_font,
        size=size,
    )
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep.append(fld_sep)
    paragraph._p.append(sep)

    if placeholder:
        text_run = OxmlElement("w:r")
        _append_run_properties(
            text_run,
            OxmlElement,
            qn,
            ascii_font=ascii_font,
            eastasia_font=eastasia_font,
            size=size,
        )
        text_node = OxmlElement("w:t")
        text_node.text = placeholder
        text_run.append(text_node)
        paragraph._p.append(text_run)

    end = OxmlElement("w:r")
    _append_run_properties(
        end,
        OxmlElement,
        qn,
        ascii_font=ascii_font,
        eastasia_font=eastasia_font,
        size=size,
    )
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    paragraph._p.append(end)


def _ensure_word_toc_styles(
    document,
    *,
    heading_eastasia: str = "宋体",
    heading_ascii: str = "宋体",
    heading_size_half_pt: int = 28,
    heading_bold: bool = False,
    entry_eastasia: str = "宋体",
    entry_ascii: str = "Times New Roman",
    entry_size_half_pt: int = 24,
    max_level: int = 3,
):
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def ensure_style(style_name: str, base_style_name: str | None = None):
        styles = document.styles
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        if base_style_name:
            try:
                style.base_style = styles[base_style_name]
            except KeyError:
                pass
        return style

    def set_style_font(style, *, eastasia: str, ascii_font: str, size_half_pt: int, bold: bool = False):
        style.font.name = ascii_font
        style.font.size = Pt(size_half_pt / 2)
        style.font.bold = bold
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), ascii_font)
        rfonts.set(qn("w:hAnsi"), ascii_font)
        rfonts.set(qn("w:eastAsia"), eastasia)
        for tag in ("w:sz", "w:szCs"):
            node = rpr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                rpr.append(node)
            node.set(qn("w:val"), str(size_half_pt))
        for tag in ("w:b", "w:bCs"):
            node = rpr.find(qn(tag))
            if bold:
                if node is None:
                    rpr.append(OxmlElement(tag))
            elif node is not None:
                rpr.remove(node)

    toc_heading = ensure_style("TOC Heading", "Normal")
    set_style_font(
        toc_heading,
        eastasia=heading_eastasia,
        ascii_font=heading_ascii,
        size_half_pt=heading_size_half_pt,
        bold=heading_bold,
    )
    toc_heading.paragraph_format.first_line_indent = Pt(0)

    for level in range(1, max_level + 1):
        toc_style = ensure_style(f"TOC {level}", "Normal")
        set_style_font(
            toc_style,
            eastasia=entry_eastasia,
            ascii_font=entry_ascii,
            size_half_pt=entry_size_half_pt,
            bold=False,
        )
        toc_style.paragraph_format.first_line_indent = Pt(0)


def _strip_contract_signing_page_text(text: str) -> str:
    text = (text or "").strip()
    if not text:
        return ""
    positions = []
    signing_patterns = [
        r"（?\s*本页无正文[^）\n]*）?",
        r"（?\s*签署页[，,、\s]*(?:无正文)?[^）\n]*）?",
        r"（?\s*本页为《[^》\n]{0,120}?基金合同(?:（草案）)?》(?:签署页)?[^）\n]*）?",
        r"《[^》\n]{0,120}?基金合同(?:（草案）)?》签署页",
    ]
    for pattern in signing_patterns:
        for match in re.finditer(pattern, text):
            positions.append(match.start())

    signatory_line_patterns = [
        r"^\s*(?:基金管理人|基金托管人|法定代表人|授权代表).*(?:盖章|签字或盖章|签名)\s*$",
        r"^\s*签署日期[:：]?\s*.*$",
    ]
    offset = 0
    for line in text.splitlines(keepends=True):
        stripped = line.strip()
        if stripped:
            for pattern in signatory_line_patterns:
                if re.search(pattern, stripped):
                    positions.append(offset + max(line.find(stripped), 0))
                    break
        offset += len(line)

    if positions:
        text = text[:min(positions)].rstrip()
    return text


def _normalize_contract_prospectus_compare_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", text or "")
    normalized = re.sub(r"《\s*基金合同\s*》", "基金合同", normalized)
    normalized = re.sub(r"《\s*招募说明书\s*》", "招募说明书", normalized)
    normalized = re.sub(r"(?<=\d)\.\s+(?=\d)", ".", normalized)
    replacements = [
        ("在本基金合同中", "在本文中"),
        ("在本招募说明书中", "在本文中"),
        ("在基金合同中", "在本文中"),
        ("在招募说明书中", "在本文中"),
        ("基金合同或本基金合同", "基金合同"),
        ("招募说明书或本招募说明书", "招募说明书"),
        ("本基金合同当事人", "基金合同当事人"),
        ("本招募说明书当事人", "招募说明书当事人"),
        ("对本基金合同的任何有效修订和补充", "对基金合同的任何有效修订和补充"),
        ("本基金合同", "基金合同"),
        ("本招募说明书", "招募说明书"),
    ]
    for old, new in replacements:
        normalized = normalized.replace(old, new)
    return re.sub(r"[ \t\u3000]+", "", normalized)


app = Flask(__name__, template_folder=str(TEMPLATES_DIR))


# ═══════════════════════════════════════════════════════════════════════════════
#  ContractEngine — 6步处理管线
# ═══════════════════════════════════════════════════════════════════════════════
class ContractEngine:
    _CONTRACT_SUMMARY_PLACEHOLDER_RE = re.compile(r"\{(CONTRACT_SUMMARY_[A-Z0-9_]*)\}")
    CONTRACT_SUMMARY_SPECS = (
        {
            "placeholder": "CONTRACT_SUMMARY_PARTIES",
            "fragments": (
                {
                    "locator": "第九部分 基金合同当事人及权利义务 / 一、基金管理人 / （二）基金管理人的权利与义务",
                    "heading": "（一）基金管理人的权利与义务",
                },
                {
                    "locator": "第九部分 基金合同当事人及权利义务 / 二、基金托管人 / （二）基金托管人的权利与义务",
                    "heading": "（二）基金托管人的权利与义务",
                },
                {
                    "locator": "第九部分 基金合同当事人及权利义务 / 三、基金份额持有人",
                    "heading": "（三）基金份额持有人",
                },
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_MEETING",
            "fragments": (
                {"locator": "第十部分 基金份额持有人大会", "take": "intro", "stop_before": "一、召开事由"},
                {"locator": "第十部分 基金份额持有人大会 / 一、召开事由", "heading": "（一）召开事由"},
                {"locator": "第十部分 基金份额持有人大会 / 二、会议召集人及召集方式", "heading": "（二）会议召集人及召集方式"},
                {
                    "locator": "第十部分 基金份额持有人大会 / 三、召开基金份额持有人大会的通知时间、通知内容、通知方式",
                    "heading": "（三）召开基金份额持有人大会的通知时间、通知内容、通知方式",
                },
                {"locator": "第十部分 基金份额持有人大会 / 四、基金份额持有人出席会议的方式", "heading": "（四）基金份额持有人出席会议的方式"},
                {"locator": "第十部分 基金份额持有人大会 / 五、议事内容与程序", "heading": "（五）议事内容与程序"},
                {"locator": "第十部分 基金份额持有人大会 / 六、表决", "heading": "（六）表决"},
                {"locator": "第十部分 基金份额持有人大会 / 七、计票", "heading": "（七）计票"},
                {"locator": "第十部分 基金份额持有人大会 / 八、生效与公告", "heading": "（八）生效与公告"},
                {
                    "locator": "第十部分 基金份额持有人大会 / 九、本部分关于基金份额持有人大会召开事由、召开条件、议事程序、表决条件等规定，凡是直接引用法律法规或监管规则的部分，如将来法律法规或监管规则修改导致相关内容被取消或变更的，基金管理人与基金托管人协商一致并提前公告后，可直接对本部分内容进行修改和调整，无需召开基金份额持有人大会审议。",
                    "heading": "（九）本部分关于基金份额持有人大会召开事由、召开条件、议事程序、表决条件等规定，凡是直接引用法律法规或监管规则的部分，如将来法律法规或监管规则修改导致相关内容被取消或变更的，基金管理人与基金托管人根据新颁布的法律法规或监管规则协商一致并提前公告后，可直接对本部分内容进行修改和调整，无需召开基金份额持有人大会审议。",
                },
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_DISTRIBUTION",
            "fragments": (
                {"locator": "第十八部分 基金的收益与分配 / 三、基金收益分配原则", "heading": "（一）基金收益分配原则"},
                {"locator": "第十八部分 基金的收益与分配 / 四、收益分配方案", "heading": "（二）收益分配方案"},
                {
                    "locator": "第十八部分 基金的收益与分配 / 五、收益分配方案的确定、公告与实施",
                    "heading": "（三）收益分配方案的确定、公告与实施",
                },
                {"locator": "第十八部分 基金的收益与分配 / 六、基金收益分配中发生的费用", "heading": "（四）基金收益分配中发生的费用"},
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_FEES",
            "fragments": (
                {"locator": "第十七部分 基金费用与税收 / 一、基金费用的种类", "heading": "（一）基金费用的种类"},
                {
                    "locator": "第十七部分 基金费用与税收 / 二、基金费用计提方法、计提标准和支付方式",
                    "heading": "（二）基金费用计提方法、计提标准和支付方式",
                },
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_INVESTMENT",
            "fragments": (
                {"locator": "第十四部分 基金的投资 / 二、投资范围", "heading": "（一）投资范围"},
                {"locator": "第十四部分 基金的投资 / 四、投资限制", "heading": "（二）投资限制"},
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_VALUATION",
            "fragments": (
                {"locator": "第十六部分 基金资产估值 / 五、估值程序"},
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_CHANGE_TERMINATION",
            "fragments": (
                {"locator": "第二十一部分 基金合同的变更、终止与基金财产的清算 / 一、《基金合同》的变更", "heading": "（一）《基金合同》的变更"},
                {"locator": "第二十一部分 基金合同的变更、终止与基金财产的清算 / 二、《基金合同》的终止事由", "heading": "（二）《基金合同》的终止事由"},
                {"locator": "第二十一部分 基金合同的变更、终止与基金财产的清算 / 三、基金财产的清算", "heading": "（三）基金财产的清算"},
                {"locator": "第二十一部分 基金合同的变更、终止与基金财产的清算 / 四、清算费用", "heading": "（四）清算费用"},
                {
                    "locator": "第二十一部分 基金合同的变更、终止与基金财产的清算 / 五、基金财产清算剩余资产的分配",
                    "heading": "（五）基金财产清算剩余资产的分配",
                },
                {"locator": "第二十一部分 基金合同的变更、终止与基金财产的清算 / 六、基金财产清算的公告", "heading": "（六）基金财产清算的公告"},
                {
                    "locator": "第二十一部分 基金合同的变更、终止与基金财产的清算 / 七、基金财产清算账册及文件的保存",
                    "heading": "（七）基金财产清算账册及文件的保存",
                },
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_DISPUTE",
            "fragments": (
                {"locator": "第二十三部分 争议的处理和适用的法律"},
            ),
        },
        {
            "placeholder": "CONTRACT_SUMMARY_EFFECT",
            "fragments": (
                {"locator": "第二十四部分 基金合同的效力"},
            ),
        },
    )

    def __init__(self, business_text_overrides: dict | None = None):
        with open(CLAUSES_JSON, encoding="utf-8") as f:
            self.clauses = json.load(f)["clauses"]
        self.business_text_overrides = (
            _normalize_business_text_overrides(business_text_overrides)
            if business_text_overrides is not None
            else _load_business_text_overrides()
        )

    @staticmethod
    def _render_contract_clause_template(text: str, v: dict) -> str:
        replacements = {
            "EXCHANGE_NAME_CN": v.get("EXCHANGE_NAME_CN", ""),
            "WORKING_DAY_DEF": v.get("WORKING_DAY_DEF", ""),
            "PURCHASE_REDEMPTION_CONSIDERATION_DEF": v.get("PURCHASE_REDEMPTION_CONSIDERATION_DEF", ""),
        }
        rendered = str(text or "")
        for key, value in replacements.items():
            rendered = rendered.replace("{" + key + "}", str(value or ""))
        return rendered

    def _contract_clause_field(self, clause_key: str, variant_key: str, field: str, v: dict) -> str:
        variant = self.clauses[clause_key]["variants"][variant_key]
        return self._render_contract_clause_template(str(variant.get(field, "") or ""), v)

    def _apply_migrated_contract_clause_texts(self, v: dict) -> dict:
        has_hk_connect = bool(v.get("HAS_HK_CONNECT"))
        hk_variant = v.get("HK_CONNECT_EXCHANGE_VARIANT", "")
        hk_exchange_variant = hk_variant if hk_variant in {"SSE_HK", "SZSE_HK"} else "STANDARD"

        intro_variant = "HK_CONNECT" if has_hk_connect else "STANDARD"
        v["NON_COMPONENT_SCOPE_INTRO"] = self._contract_clause_field("NON_COMPONENT_SCOPE_INTRO", intro_variant, "text", v)

        definition_variant = "HK_CONNECT" if has_hk_connect else "STANDARD"
        v["DEFINITION_TAIL"] = self._contract_clause_field("DEFINITION_TAIL", definition_variant, "text", v)

        if has_hk_connect:
            preface_variant = "HK_CONNECT"
        elif v.get("IS_KECHUANG"):
            preface_variant = "KECHUANG"
        else:
            preface_variant = "STANDARD"
        v["PREFACE_RISK_DISCLOSURE_TAIL"] = self._contract_clause_field(
            "PREFACE_RISK_DISCLOSURE_TAIL",
            preface_variant,
            "text",
            v,
        )

        for field in (
            "PURCHASE_CONSIDERATION_GLOSSARY_DEF",
            "REDEMPTION_CONSIDERATION_GLOSSARY_DEF",
            "FUND_SHARE_NAV_DEF",
            "PURCHASE_REDEMPTION_PRINCIPLE_DEF",
            "PURCHASE_REDEMPTION_CONSIDERATION_DEF",
        ):
            v[field] = self._contract_clause_field("PURCHASE_REDEMPTION_CONSIDERATION_TEXTS", hk_exchange_variant, field, v)
        v["FUND_PROFIT_GLOSSARY_DEF"] = f"57、基金利润：{v['FUND_PROFIT_DEF']}"

        valuation_variant = "HK_CONNECT" if has_hk_connect else "STANDARD"
        for field in ("VALUATION_METHODS_TAIL_TEXT", "VALUATION_PROCEDURE_TEXT", "SPECIAL_DISCLOSURE_TEXT"):
            v[field] = self._contract_clause_field("VALUATION_DISCLOSURE_TEXTS", valuation_variant, field, v)

        for field in ("LISTING_ITEM3_DEF", "LISTING_SECTION2_BODY", "LISTING_SECTION3_TEXT", "LISTING_SECTION4_TEXT"):
            v[field] = self._contract_clause_field("LISTING_TRADING_TEXTS", hk_exchange_variant, field, v)

        section_fields = {
            "PART8_INTRO_TEXT": "part8_intro_text",
            "OPEN_DAY_CLAUSE": "open_day_clause",
            "OPEN_DAY_ADJUSTMENT_CLAUSE": "open_day_adjustment_clause",
            "PART8_SECTION4_TEXT": "part8_section4_text",
            "PART8_SECTION6_TEXT": "part8_section6_text",
            "PART8_SECTION7_TEXT": "part8_section7_text",
            "PART8_SECTION8_TEXT": "part8_section8_text",
            "PART8_SECTION9_TEXT": "part8_section9_text",
        }
        for target_key, field in section_fields.items():
            v[target_key] = self._contract_clause_field("PURCHASE_REDEMPTION_SECTION_TEXTS", hk_exchange_variant, field, v)

        proxy_variant = "HK_CONNECT" if has_hk_connect else "STANDARD"
        v["HK_PROXY_VOTING_TEXT"] = self._contract_clause_field("HK_PROXY_VOTING_TEXT", proxy_variant, "text", v)
        return v

    def _apply_business_text_overrides(self, v: dict) -> dict:
        result = dict(v)
        for key, spec in CONTRACT_BUSINESS_TEXT_SPECS.items():
            if key not in result:
                continue
            scene = _contract_business_text_variant(result, spec.get("variant_mode", "DEFAULT"))
            override = _get_business_text_override(
                self.business_text_overrides,
                "contract",
                key,
                scene,
                _business_text_product_type(result),
                _business_text_market_type(result),
                _business_text_exchange(result),
            )
            if override is not None:
                result[key] = override
        return result

    @staticmethod
    def _cn_numeral_to_int(cn: str):
        """
        Convert common Chinese numerals to int.
        Supports values used in headings like: 一, 十, 十一, 二十六, 一百零二.
        Returns None when conversion fails.
        """
        if not cn:
            return None
        if cn.isdigit():
            return int(cn)

        digits = {
            "零": 0, "一": 1, "二": 2, "三": 3, "四": 4,
            "五": 5, "六": 6, "七": 7, "八": 8, "九": 9,
        }
        units = {"十": 10, "百": 100, "千": 1000}

        total = 0
        current = 0
        for ch in cn:
            if ch in digits:
                current = digits[ch]
                continue
            if ch in units:
                if current == 0:
                    current = 1
                total += current * units[ch]
                current = 0
                continue
            return None
        total += current
        return total if total > 0 else None

    @staticmethod
    def _infer_index_provider(index_name: str) -> tuple[str, str]:
        name = str(index_name or "").strip()
        if not name:
            return "", ""

        cnindex_keywords = ("创业板", "深证", "国证", "恒生")
        if any(keyword in name for keyword in cnindex_keywords):
            return "深圳证券信息有限公司", "http://www.cnindex.com.cn"

        csindex_keywords = ("中证", "上证", "科创板")
        if any(keyword in name for keyword in csindex_keywords):
            return "中证指数有限公司", "https://www.csindex.com.cn"

        return "", ""

    @staticmethod
    def _strip_index_suffix(index_name: str) -> str:
        return str(index_name or "").strip().replace("指数", "")

    @classmethod
    def _derive_fund_names_from_index_name(cls, index_name: str) -> tuple[str, str]:
        stem = cls._strip_index_suffix(index_name)
        if not stem:
            return "", ""
        return (
            f"南方{stem}交易型开放式指数证券投资基金",
            f"南方{stem}ETF",
        )

    # ── Step 1: 派生变量 ─────────────────────────────────────────────────────
    def _derive_variables(self, v: dict) -> dict:
        v = _merge_schema_variable_defaults(v)

        index_name = str(v.get("INDEX_NAME", "") or "").strip()
        inferred_index_compiler, inferred_index_website = self._infer_index_provider(index_name)
        if not str(v.get("INDEX_COMPILER", "") or "").strip() and inferred_index_compiler:
            v["INDEX_COMPILER"] = inferred_index_compiler
        if not str(v.get("INDEX_WEBSITE", "") or "").strip() and inferred_index_website:
            v["INDEX_WEBSITE"] = inferred_index_website

        index_website = str(v.get("INDEX_WEBSITE", "") or "").strip()
        legacy_index_website = str(v.get("INDEX_COMPILER_web", "") or "").strip()
        if index_website and not legacy_index_website:
            v["INDEX_COMPILER_web"] = index_website
        elif legacy_index_website and not index_website:
            v["INDEX_WEBSITE"] = legacy_index_website

        fund_name, fund_short_name = self._derive_fund_names_from_index_name(index_name)
        if not str(v.get("FUND_NAME", "") or "").strip() and fund_name:
            v["FUND_NAME"] = fund_name
        if not str(v.get("FUND_SHORT_NAME", "") or "").strip() and fund_short_name:
            v["FUND_SHORT_NAME"] = fund_short_name

        # 交易所 → 中文名
        exchange = v.get("EXCHANGE", "")
        if exchange == "SZSE":
            v.setdefault("EXCHANGE_NAME_CN", "深圳证券交易所")
        elif exchange == "SSE":
            v.setdefault("EXCHANGE_NAME_CN", "上海证券交易所")
        v["IS_SZSE"] = (exchange == "SZSE")
        v["IS_SSE"] = (exchange == "SSE")

        dispute_key, dispute_config = _resolve_dispute_resolution_venue(
            v.get("DISPUTE_RESOLUTION_VENUE") or v.get("DISPUTE_RESOLUTION_PLACE")
        )
        v["DISPUTE_RESOLUTION_VENUE"] = dispute_key
        v["DISPUTE_RESOLUTION_LABEL"] = dispute_config["label"]
        v["DISPUTE_ARBITRATION_INSTITUTION"] = dispute_config["institution"]
        v["DISPUTE_ARBITRATION_LOCATION"] = dispute_config["location"]
        v["DISPUTE_RESOLUTION_CLAUSE"] = dispute_config["contract_clause"]
        v["PRODUCT_SUMMARY_DISPUTE_RESOLUTION_TEXT"] = dispute_config["product_summary_sentence"]

        if not str(v.get("WORKING_DAY_TYPE", "") or "").strip():
            if exchange == "SZSE":
                v["WORKING_DAY_TYPE"] = "SZSE_ONLY"
            elif exchange == "SSE":
                v["WORKING_DAY_TYPE"] = "SSE_ONLY"

        # 市场类型 → HAS_HK_CONNECT / 港股通子场景 / 跟踪误差默认值 / 市场标志
        market = v.get("MARKET_TYPE", "")
        hk_connect_variant = ""
        if market == "HK_CONNECT":
            v["HAS_HK_CONNECT"] = True
            if exchange == "SSE":
                hk_connect_variant = "SSE_HK"
            elif exchange == "SZSE":
                hk_connect_variant = "SZSE_HK"
            v.setdefault("TRACKING_ERROR_DAILY", 0.35)
            v.setdefault("TRACKING_ERROR_ANNUAL", 4)
        else:
            v.setdefault("HAS_HK_CONNECT", False)
            v.setdefault("TRACKING_ERROR_DAILY", 0.2)
            v.setdefault("TRACKING_ERROR_ANNUAL", 2)
        v["HK_CONNECT_EXCHANGE_VARIANT"] = hk_connect_variant
        v["IS_CHUANGYE"] = (market == "CHUANGYE")
        v["IS_KECHUANG"] = (market == "KECHUANG")

        if not str(v.get("HAS_STOCK_SUBSCRIPTION", "") or "").strip():
            v["HAS_STOCK_SUBSCRIPTION"] = (market != "HK_CONNECT")

        if not str(v.get("BUSINESS_RULES_TYPE", "") or "").strip():
            if hk_connect_variant == "SZSE_HK":
                v["BUSINESS_RULES_TYPE"] = "CSDC"
            elif exchange == "SZSE":
                v["BUSINESS_RULES_TYPE"] = "SZSE"
            elif exchange == "SSE":
                v["BUSINESS_RULES_TYPE"] = "SSE"

        # 中文布尔 → Python bool
        for key in ("HAS_HK_CONNECT", "HAS_AML", "HAS_CUSTODIAN_TRANSFER_SPECIAL",
                    "HAS_STOCK_SUBSCRIPTION", "HAS_CDR", "CUSTODIAN_HAS_OFFICE_ADDRESS"):
            if key in v:
                raw = v[key]
                if isinstance(raw, str):
                    v[key] = raw.lower() in ("true", "1", "yes", "是")

        # 托管人有无办公地址
        custodian = v.get("CUSTODIAN_NAME", "")
        v["CUSTODIAN_IS_BANK"] = "银行" in str(custodian or "")
        custodians_with_office = {"交通银行股份有限公司", "中信证券股份有限公司"}
        v.setdefault(
            "CUSTODIAN_HAS_OFFICE_ADDRESS",
            custodian in custodians_with_office,
        )

        return v

    # ── Step 2: 注入差异条款原文 ─────────────────────────────────────────────
    def _inject_clause_texts(self, v: dict) -> dict:
        v = dict(v)

        # WORKING_DAY_DEF
        wdt = v.get("WORKING_DAY_TYPE", "SZSE_ONLY")
        wday_variants = self.clauses["WORKING_DAY_DEF"]["variants"]
        v["WORKING_DAY_DEF"] = wday_variants.get(wdt, wday_variants["SZSE_ONLY"])["text"]

        # BUSINESS_RULES_DEF
        brt = v.get("BUSINESS_RULES_TYPE", "SZSE")
        br_variants = self.clauses["BUSINESS_RULES_DEF"]["variants"]
        v["BUSINESS_RULES_DEF"] = br_variants.get(brt, br_variants["SZSE"])["text"]

        # DISPUTE_RESOLUTION_CLAUSE
        dispute_key = v.get("DISPUTE_RESOLUTION_VENUE", DEFAULT_DISPUTE_RESOLUTION_VENUE)
        dispute_variants = self.clauses.get("DISPUTE_RESOLUTION_CLAUSE", {}).get("variants", {})
        if dispute_variants:
            fallback_key = DEFAULT_DISPUTE_RESOLUTION_VENUE
            v["DISPUTE_RESOLUTION_CLAUSE"] = dispute_variants.get(
                dispute_key, dispute_variants.get(fallback_key, {})
            ).get("text", v.get("DISPUTE_RESOLUTION_CLAUSE", ""))

        # NON_COMPONENT_SCOPE
        mt = v.get("MARKET_TYPE", "CHUANGYE")
        nc_variants = self.clauses["NON_COMPONENT_SCOPE"]["variants"]
        # Map MARKET_TYPE value to variant key
        mt_to_key = {
            "CHUANGYE": "CHUANGYE",
            "KECHUANG": "KECHUANG",
            "A_SHARE": "A_SHARE",
            "HK_CONNECT": "HK_CONNECT",
        }
        nc_key = mt_to_key.get(mt, "CHUANGYE")
        v["NON_COMPONENT_SCOPE"] = nc_variants.get(nc_key, nc_variants["CHUANGYE"])["text"]

        # NON_COMPONENT_SCOPE_INTRO
        if mt == "HK_CONNECT":
            v["NON_COMPONENT_SCOPE_INTRO"] = (
                "包括内地与香港股票市场交易互联互通机制允许买卖的规定范围内的"
                "香港联合交易所上市的股票（简称\u201c港股通股票\u201d）、存托凭证，下同"
            )
        else:
            v["NON_COMPONENT_SCOPE_INTRO"] = "含存托凭证"

        # DISTRIBUTION_FREQ_CLAUSE
        df = v.get("DISTRIBUTION_FREQ", "MONTHLY")
        df_variants = self.clauses["DISTRIBUTION_FREQ_CLAUSE"]["variants"]
        v["DISTRIBUTION_FREQ_CLAUSE"] = df_variants.get(df, df_variants["MONTHLY"])["text"]

        # MGMT_FEE_PAYMENT_METHOD text (accepts either a legacy enum or verbatim clause text)
        if "MGMT_FEE_PAYMENT_METHOD" in v:
            mfpm = v.get("MGMT_FEE_PAYMENT_METHOD") or ""
        else:
            mfpm = "CONSULT"
        mgmt_variants = self.clauses["MGMT_FEE_PAYMENT"]["variants"]
        mfpm_enum, mfpm_text = _resolve_payment_clause_value(mfpm, mgmt_variants)
        # Store original enum/raw values for custody fallback.
        v["_MGMT_FEE_PAYMENT_ENUM"] = mfpm_enum
        v["_MGMT_FEE_PAYMENT_RAW"] = "" if mfpm_enum else mfpm_text
        v["MGMT_FEE_PAYMENT_METHOD"] = mfpm_text

        # CUSTODY_FEE_PAYMENT_METHOD text
        if "CUSTODY_FEE_PAYMENT_METHOD" in v:
            cfpm = v.get("CUSTODY_FEE_PAYMENT_METHOD") or ""
        else:
            cfpm = v.get("_MGMT_FEE_PAYMENT_ENUM") or v.get("_MGMT_FEE_PAYMENT_RAW") or "CONSULT"
        cust_variants = self.clauses["CUSTODY_FEE_PAYMENT"]["variants"]
        _, cfpm_text = _resolve_payment_clause_value(cfpm, cust_variants)
        v["CUSTODY_FEE_PAYMENT_METHOD"] = cfpm_text

        # FUND_PROFIT_DEF (for HK_CONNECT)
        fp_variants = self.clauses["FUND_PROFIT_DEF"]["variants"]
        if v.get("HAS_HK_CONNECT"):
            v["FUND_PROFIT_DEF"] = fp_variants["HK_CONNECT"]["text"]
        else:
            v["FUND_PROFIT_DEF"] = fp_variants["STANDARD"]["text"]

        hk_definitions = self.clauses["HK_CONNECT_DEFINITIONS"]["clauses"]
        liquidity_def = (
            "流动性受限资产：指由于法律法规、监管、合同或操作障碍等原因无法以合理价格予以变现的资产，"
            "包括但不限于到期日在10个交易日以上的逆回购与银行定期存款（含协议约定有条件提前支取的银行存款）、"
            "停牌股票、流通受限的新股及非公开发行股票、资产支持证券、因发行人债务违约无法进行转让或交易的债券等，"
            "但法律法规或中国证监会另有规定的，从其规定"
        )
        force_majeure_def = "不可抗力：指本基金合同当事人不能预见、不能避免且不能克服的客观事件"
        if v.get("HAS_HK_CONNECT"):
            v["DEFINITION_TAIL"] = (
                f"64、内地与香港股票市场交易互联互通机制：{hk_definitions['clause_mutual_market']['text']}\n"
                f"65、港股通：{hk_definitions['clause_hk_connect']['text']}\n"
                f"66、{liquidity_def}\n"
                f"67、{force_majeure_def}"
            )
            v["PREFACE_RISK_DISCLOSURE_TAIL"] = (
                "七、本基金资产投资于港股，会面临港股通机制下因投资环境、投资标的、市场制度以及交易规则等差异带来的特有风险，"
                "包括港股市场股价波动较大的风险（港股市场实行T+0回转交易，且对个股不设涨跌幅限制，港股股价可能表现出比A股更为剧烈的股价波动）、"
                "汇率风险（汇率波动可能对基金的投资收益造成损失）、港股通机制下交易日不连贯可能带来的风险"
                "（在内地开市香港休市的情形下，港股通不能正常交易，港股不能及时卖出，可能带来一定的流动性风险）等。\n"
                "八、本基金为指数基金，投资者投资于本基金面临跟踪误差控制未达约定目标、指数编制机构停止服务、成份股停牌、摘牌等潜在风险，详见本基金招募说明书。\n"
                "九、本基金投资范围包括国内依法发行上市的存托凭证（“中国存托凭证”），将面临中国存托凭证价格大幅波动甚至出现较大亏损的风险、与中国存托凭证发行机制相关的风险等。"
                "但基金资产并非必然参与存托凭证的投资，基金可根据投资策略需要或不同市场环境的变化，选择是否采用存托凭证投资策略。\n"
                "十、本基金投资范围包括股指期货、股票期权等金融衍生品，可能给本基金带来额外风险。"
                "投资股指期货的风险包括但不限于杠杆风险、保证金风险、期货价格与基金投资品种价格的相关度降低带来的风险等；"
                "投资股票期权的风险包括但不限于市场风险、流动性风险、交易对手信用风险、操作风险、保证金风险等；由此可能增加本基金净值的波动性。"
            )
        else:
            v["DEFINITION_TAIL"] = (
                f"64、{liquidity_def}\n"
                f"65、{force_majeure_def}"
            )
            preface_market_risk = ""
            if v.get("IS_KECHUANG"):
                preface_market_risk = (
                    "本基金投资于科创板股票，会面临科创板机制下因投资标的、市场制度以及交易规则等差异带来的特有风险，"
                    "包括股价波动风险、流动性风险、退市风险和投资集中风险等。"
                )
            v["PREFACE_RISK_DISCLOSURE_TAIL"] = (
                "七、本基金为指数基金，投资者投资于本基金面临跟踪误差控制未达约定目标、指数编制机构停止服务、成份股停牌、摘牌等潜在风险，详见本基金招募说明书。"
                f"{preface_market_risk}\n"
                "八、本基金投资范围包括国内依法发行上市的存托凭证（“中国存托凭证”），将面临中国存托凭证价格大幅波动甚至出现较大亏损的风险、与中国存托凭证发行机制相关的风险等。"
                "但基金资产并非必然参与存托凭证的投资，基金可根据投资策略需要或不同市场环境的变化，选择是否采用存托凭证投资策略。\n"
                "九、本基金投资范围包括股指期货、股票期权等金融衍生品，可能给本基金带来额外风险。"
                "投资股指期货的风险包括但不限于杠杆风险、保证金风险、期货价格与基金投资品种价格的相关度降低带来的风险等；"
                "投资股票期权的风险包括但不限于市场风险、流动性风险、交易对手信用风险、操作风险、保证金风险等；由此可能增加本基金净值的波动性。"
            )

        hk_variant = v.get("HK_CONNECT_EXCHANGE_VARIANT", "")
        exch_cn = v.get("EXCHANGE_NAME_CN", "深圳证券交易所")
        if hk_variant == "SSE_HK":
            v["PURCHASE_CONSIDERATION_GLOSSARY_DEF"] = (
                "49、申购对价：指投资人申购基金份额时，按基金合同和招募说明书规定应交付的组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价"
            )
            v["REDEMPTION_CONSIDERATION_GLOSSARY_DEF"] = (
                "50、赎回对价：指基金份额持有人赎回基金份额时，基金管理人按基金合同和招募说明书规定应交付给赎回人的组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价"
            )
            v["FUND_SHARE_NAV_DEF"] = "60、基金份额净值：指估值日基金资产净值除以估值日基金份额总数"
            v["PURCHASE_REDEMPTION_PRINCIPLE_DEF"] = (
                "2、本基金的申购对价、赎回对价包括组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价。"
            )
            v["PURCHASE_REDEMPTION_CONSIDERATION_DEF"] = (
                "2、申购对价是指投资人申购基金份额时应交付的组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价。"
                "赎回对价是指基金份额持有人赎回基金份额时，基金管理人应交付的组合证券、现金替代、现金差额及其他对价或现金替代、现金差额及其他对价。"
                "申购对价、赎回对价根据申购赎回清单和投资人申购、赎回的基金份额数额确定。"
            )
        elif hk_variant == "SZSE_HK":
            v["PURCHASE_CONSIDERATION_GLOSSARY_DEF"] = (
                "49、申购对价：指投资人申购基金份额时，按基金合同和招募说明书规定应交付的现金替代、现金差额及其他对价"
            )
            v["REDEMPTION_CONSIDERATION_GLOSSARY_DEF"] = (
                "50、赎回对价：指基金份额持有人赎回基金份额时，基金管理人按基金合同和招募说明书规定应交付给赎回人的现金替代、现金差额及其他对价"
            )
            v["FUND_SHARE_NAV_DEF"] = "60、基金份额净值：指计算日基金资产净值除以计算日基金份额总数"
            v["PURCHASE_REDEMPTION_PRINCIPLE_DEF"] = (
                "2、本基金的申购对价、赎回对价包括现金替代、现金差额及其他对价。"
            )
            v["PURCHASE_REDEMPTION_CONSIDERATION_DEF"] = (
                "2、申购对价是指投资人申购基金份额时应交付的现金替代、现金差额及其他对价。"
                "赎回对价是指基金份额持有人赎回基金份额时，基金管理人应交付的现金替代、现金差额及其他对价。"
                "申购对价、赎回对价根据申购赎回清单和投资人申购、赎回的基金份额数额确定。"
            )
        else:
            v["PURCHASE_CONSIDERATION_GLOSSARY_DEF"] = (
                "49、申购对价：指投资人申购基金份额时，按基金合同和招募说明书规定应交付的组合证券、现金替代、现金差额和/或其他对价"
            )
            v["REDEMPTION_CONSIDERATION_GLOSSARY_DEF"] = (
                "50、赎回对价：指基金份额持有人赎回基金份额时，基金管理人按基金合同和招募说明书规定应交付给赎回人的组合证券、现金替代、现金差额和/或其他对价"
            )
            v["FUND_SHARE_NAV_DEF"] = "60、基金份额净值：指计算日基金资产净值除以计算日基金份额总数"
            v["PURCHASE_REDEMPTION_PRINCIPLE_DEF"] = (
                "2、本基金的申购对价、赎回对价包括组合证券、现金替代、现金差额和/或其他对价。"
            )
            v["PURCHASE_REDEMPTION_CONSIDERATION_DEF"] = (
                "2、申购对价是指投资人申购基金份额时应交付的组合证券、现金替代、现金差额和/或其他对价。"
                "赎回对价是指基金份额持有人赎回基金份额时，基金管理人应交付的组合证券、现金替代、现金差额和/或其他对价。"
                "申购对价、赎回对价根据申购赎回清单和投资人申购、赎回的基金份额数额确定。"
            )
        v["FUND_PROFIT_GLOSSARY_DEF"] = f"57、基金利润：{v['FUND_PROFIT_DEF']}"

        # SUBSCRIPTION_METHOD_TEXT
        sub_variants = self.clauses["SUBSCRIPTION_METHOD"]["variants"]
        if v.get("HAS_STOCK_SUBSCRIPTION"):
            sub_text = sub_variants["THREE_TYPES"]["text"]
        else:
            sub_text = sub_variants["TWO_TYPES"]["text"]
        # Replace {EXCHANGE_NAME_CN} inside subscription text
        sub_text = sub_text.replace("{EXCHANGE_NAME_CN}", v.get("EXCHANGE_NAME_CN", "深圳证券交易所"))
        v["SUBSCRIPTION_METHOD_TEXT"] = sub_text

        # SUBSCRIPTION_FUNDS_AND_STOCKS_TEXT
        sub_funds_variants = self.clauses["SUBSCRIPTION_FUNDS_AND_STOCKS"]["variants"]
        sub_funds_key = f"{v.get('EXCHANGE', 'SZSE')}_{'WITH_STOCK' if v.get('HAS_STOCK_SUBSCRIPTION') else 'NO_STOCK'}"
        v["SUBSCRIPTION_FUNDS_AND_STOCKS_TEXT"] = sub_funds_variants.get(
            sub_funds_key,
            sub_funds_variants["SZSE_WITH_STOCK"],
        )["text"]

        v = self._apply_migrated_contract_clause_texts(v)
        return self._apply_business_text_overrides(v)

        v["VALUATION_METHODS_TAIL_TEXT"] = (
            "4、本基金投资股指期货合约，一般以估值当日结算价进行估值，估值当日无结算价的，且最近交易日后经济环境未发生重大变化的，采用最近交易日结算价估值。本基金投资股票期权合约，根据相关法律法规以及监管部门的规定估值。法律法规另有规定的，从其规定。\n"
            "5、基金参与融资、转融通证券出借业务的，应参照相关法律法规和中国证券投资基金业协会发布的相关规定进行估值，确保估值的公允性。\n"
            "6、本基金投资存托凭证的估值核算依照境内上市交易的股票执行。\n"
            "7、对于发行人已破产、发行人未能按时足额偿付本金或利息，或者有其它可靠信息表明本金或利息无法按时足额偿付的债券投资品种，第三方估值基准服务机构可在提供推荐价格的同时提供价格区间作为公允价值的参考范围以及公允价值存在重大不确定性的相关提示。基金管理人在与基金托管人协商一致后，可采用价格区间中的数据作为该债券投资品种的公允价值。\n"
            "8、如有确凿证据表明按上述方法进行估值不能客观反映其公允价值的，基金管理人可根据具体情况与基金托管人商定后，按最能反映公允价值的价格估值。\n"
            "9、相关法律法规以及监管部门有强制规定的，从其规定。如有新增事项，按国家最新规定估值。"
        )
        v["VALUATION_PROCEDURE_TEXT"] = (
            "1、基金份额净值是按照每个估值日闭市后，基金资产净值除以当日基金份额的余额数量计算，精确到0.0001元，小数点后第5位四舍五入。基金管理人可以设立大额赎回情形下的净值精度应急调整机制。国家另有规定的，从其规定。\n"
            "基金管理人应每个估值日计算基金资产净值及基金份额净值，并按规定披露。\n"
            "2、基金管理人应每个估值日对基金资产估值。但基金管理人根据法律法规或本基金合同的规定暂停估值时除外。基金管理人每个估值日对基金资产估值后，将基金份额净值结果发送基金托管人，经基金托管人复核无误后，由基金管理人按规定对外公布。"
        )
        v["SPECIAL_DISCLOSURE_TEXT"] = (
            "（十二）基金投资股指期货的信息披露\n"
            "本基金在季度报告、中期报告、年度报告等定期报告和招募说明书（更新）等文件中披露股指期货交易情况，包括交易政策、持仓情况、损益情况、风险指标等，并充分揭示股指期货交易对基金总体风险的影响以及是否符合既定的投资政策和投资目标等。\n"
            "（十三）基金投资股票期权的信息披露\n"
            "本基金在定期信息披露文件中披露参与股票期权交易的有关情况，包括投资政策、持仓情况、损益情况、风险指标、估值方法等，并充分揭示股票期权交易对基金总体风险的影响等。\n"
            "（十四）基金参与融资业务的信息披露\n"
            "本基金在季度报告、中期报告、年度报告等定期报告和招募说明书（更新）等文件中披露参与融资交易情况，包括投资策略、业务开展情况、损益情况、风险及管理情况等。\n"
            "（十五）基金参与转融通证券出借业务的信息披露\n"
            "本基金在季度报告、中期报告、年度报告等定期报告和招募说明书（更新）等文件中披露参与转融通证券出借交易情况，包括投资策略、业务开展情况、损益情况、风险及其管理情况等，并就报告期内本基金参与转融通证券出借业务发生的重大关联交易事项做详细说明。\n"
            "（十六）基金投资资产支持证券的信息披露\n"
            "本基金在中期报告、年度报告中披露其持有的资产支持证券总额、资产支持证券市值占基金净资产的比例和报告期内所有的资产支持证券明细。本基金在季度报告中披露其持有的资产支持证券总额、资产支持证券市值占基金净资产的比例和报告期末按市值占基金净资产比例大小排序的前10名资产支持证券明细。\n"
            "（十七）清算报告\n"
            "基金合同终止的，基金管理人应当组织基金财产清算小组对基金财产进行清算并作出清算报告。基金财产清算小组应当将清算报告登载在规定网站上，并将清算报告提示性公告登载在规定报刊上。\n"
            "（十八）本基金投资存托凭证的信息披露依照境内上市交易的股票执行。\n"
            "（十九）中国证监会规定的其他信息。"
        )

        if v.get("HAS_HK_CONNECT"):
            v["VALUATION_METHODS_TAIL_TEXT"] = (
                "4、本基金投资股指期货合约，一般以估值当日结算价进行估值，估值当日无结算价的，且最近交易日后经济环境未发生重大变化的，采用最近交易日结算价估值。本基金投资股票期权合约，根据相关法律法规以及监管部门的规定估值。法律法规另有规定的，从其规定。\n"
                "5、港股通投资持有外币证券资产估值涉及到的主要货币对人民币汇率，以估值日中国人民银行或其授权机构公布的人民币汇率中间价为准。\n"
                "6、基金参与融资、转融通证券出借业务的，应参照相关法律法规和中国证券投资基金业协会发布的相关规定进行估值，确保估值的公允性。\n"
                "7、如有确凿证据表明按原有方法进行估值不能客观反映上述资产或负债公允价值的，基金管理人可根据具体情况与基金托管人商定后，按最能反映公允价值的方法估值。\n"
                "8、本基金投资存托凭证的估值核算依照境内上市交易的股票执行。\n"
                "9、相关法律法规以及监管部门、自律规则另有规定的，从其规定。如有新增事项，按国家最新规定估值。"
            )
            v["VALUATION_PROCEDURE_TEXT"] = (
                "1、基金份额净值是按照每个估值日闭市后，基金资产净值除以当日基金份额的余额数量计算，精确到0.0001元，小数点后第5位四舍五入。基金管理人可以设立大额赎回情形下的净值精度应急调整机制。国家另有规定的，从其规定。\n"
                "基金管理人应每个估值日计算基金资产净值及基金份额净值，并按规定披露。\n"
                "2、基金管理人应每个估值日对基金资产估值，但基金管理人根据法律法规或本基金合同的规定暂停估值时除外。基金管理人每个估值日对基金资产估值后，将基金资产净值、基金份额净值结果发送基金托管人，经基金托管人复核无误后，由基金管理人按规定对外公布。"
            )
            v["SPECIAL_DISCLOSURE_TEXT"] = (
                "（十二）基金投资股指期货的信息披露\n"
                "本基金在季度报告、中期报告、年度报告等定期报告和招募说明书（更新）等文件中披露股指期货交易情况，包括交易政策、持仓情况、损益情况、风险指标等，并充分揭示股指期货交易对基金总体风险的影响以及是否符合既定的投资政策和投资目标等。\n"
                "（十三）基金投资股票期权的信息披露\n"
                "本基金在定期信息披露文件中披露参与股票期权交易的有关情况，包括投资政策、持仓情况、损益情况、风险指标、估值方法等，并充分揭示股票期权交易对基金总体风险的影响等。\n"
                "（十四）基金参与融资业务的信息披露\n"
                "本基金在季度报告、中期报告、年度报告等定期报告和招募说明书（更新）等文件中披露参与融资交易情况，包括投资策略、业务开展情况、损益情况、风险及管理情况等。\n"
                "（十五）基金参与转融通证券出借业务的信息披露\n"
                "本基金在季度报告、中期报告、年度报告等定期报告和招募说明书（更新）等文件中披露参与转融通证券出借交易情况，包括投资策略、业务开展情况、损益情况、风险及其管理情况等，并就报告期内本基金参与转融通证券出借业务发生的重大关联交易事项做详细说明。\n"
                "（十六）基金投资港股通股票的信息披露\n"
                "本基金在季度报告、中期报告、年度报告等定期报告和招募说明书（更新）等文件中披露参与港股通交易的相关情况。\n"
                "（十七）基金投资资产支持证券的信息披露\n"
                "本基金在中期报告、年度报告中披露其持有的资产支持证券总额、资产支持证券市值占基金净资产的比例和报告期内所有的资产支持证券明细。本基金在季度报告中披露其持有的资产支持证券总额、资产支持证券市值占基金净资产的比例和报告期末按市值占基金净资产比例大小排序的前10名资产支持证券明细。\n"
                "（十八）清算报告\n"
                "基金合同终止的，基金管理人应当组织基金财产清算小组对基金财产进行清算并作出清算报告。基金财产清算小组应当将清算报告登载在规定网站上，并将清算报告提示性公告登载在规定报刊上。\n"
                "（十九）本基金投资存托凭证的信息披露依照境内上市交易的股票执行。\n"
                "（二十）中国证监会规定的其他信息。"
            )

        if hk_variant == "SSE_HK":
            v["LISTING_ITEM3_DEF"] = "3、符合上海证券交易所规定的其他条件。"
            v["LISTING_SECTION2_BODY"] = (
                "本基金基金份额在上海证券交易所的上市交易需遵照《上海证券交易所交易规则》、"
                "《上海证券交易所证券投资基金上市规则》、《上海证券交易所交易型开放式指数基金业务实施细则》及其他有关规定。"
            )
            v["LISTING_SECTION3_TEXT"] = (
                "三、终止上市交易\n"
                "基金份额上市交易后，有下列情形之一的，上海证券交易所可终止基金份额的上市交易：\n"
                "1、不再具备本部分第一条规定的上市条件；\n"
                "2、基金合同终止；\n"
                "3、基金份额持有人大会决定终止上市；\n"
                "4、基金合同约定的终止上市的其他情形；\n"
                "5、上海证券交易所认为应当终止上市的其他情形。\n"
                "若因上述 1、3、4、5 项原因使本基金不再具备上市条件而被上海证券交易所终止上市的，本基金可由交易型开放式基金变更为跟踪标的指数的普通开放式基金或上市开放式基金（LOF），且因上述 1、4、5 项之一情形终止上市的，本基金变更为跟踪标的指数的普通开放式基金或上市开放式基金（LOF）无需召开基金份额持有人大会。届时，基金管理人可变更本基金的登记机构、相应调整申购赎回业务规则、提前制定基金终止上市后场内份额的处理规则并公告，同时，基金管理人可按照《信息披露办法》的规定，公告变更后的基金合同及招募说明书。若届时本基金管理人已有以该指数作为标的指数的基金，则本基金将本着维护投资者合法权益的原则，履行适当的程序后与该指数基金合并。具体情况见基金管理人届时公告。"
            )
            v["LISTING_SECTION4_TEXT"] = (
                "四、基金份额参考净值（IOPV）的计算与公告\n"
                "基金管理人在每一交易日开市前公告当日的申购赎回清单，基金管理人或基金管理人委托的其他机构可以在相关证券交易所开市后根据申购赎回清单、汇率数据和组合证券内各只证券的实时成交数据，计算基金份额参考净值（IOPV），并将计算结果向上海证券交易所发送，由上海证券交易所对外发布，仅供投资者交易、申购、赎回基金份额时参考。基金份额参考净值的计算方法详见招募说明书。\n"
                "上海证券交易所和基金管理人可以调整基金份额参考净值的计算方法，并予以公告。"
            )
        elif hk_variant == "SZSE_HK":
            v["LISTING_ITEM3_DEF"] = "3、《深圳证券交易所证券投资基金上市规则》规定的其他条件。"
            v["LISTING_SECTION2_BODY"] = (
                "本基金基金份额在深圳证券交易所的上市交易，应遵照《深圳证券交易所交易规则》、"
                "《深圳证券交易所证券投资基金上市规则》、《深圳证券交易所证券投资基金交易和申购赎回实施细则》等有关规定。"
            )
            v["LISTING_SECTION3_TEXT"] = (
                "三、停复牌、暂停上市、恢复上市及终止上市交易\n"
                f"基金份额在深圳证券交易所上市后，如遇停复牌、暂停上市、恢复上市或终止上市的情形，按照深圳证券交易所的相关规定执行。当基金发生深圳证券交易所规定的因不再具备上市条件而应当终止上市的情形时，在法律法规和监管机构允许的情况下，本基金可变更为跟踪标的指数的普通开放式基金或上市开放式基金（LOF），无需召开基金份额持有人大会。届时，基金管理人可变更本基金的登记机构、相应调整申购赎回业务规则、提前制定基金终止上市后场内份额的处理规则并公告，同时，基金管理人可按照《信息披露办法》的规定，公告变更后的基金合同及招募说明书。若届时本基金管理人已有以该指数作为标的指数的基金，则本基金将本着维护投资者合法权益的原则，履行适当的程序后与该指数基金合并。具体情况见基金管理人届时公告。"
            )
            v["LISTING_SECTION4_TEXT"] = (
                "四、基金份额参考净值（IOPV）的计算与公告\n"
                "基金管理人在每一交易日开市前公告当日的申购赎回清单，基金管理人或基金管理人委托的其他机构可以在相关证券交易所开市后根据申购赎回清单、汇率数据和组合证券内各只证券的实时成交数据，计算基金份额参考净值（IOPV），并将计算结果向深圳证券交易所发送，由深圳证券交易所对外发布，仅供投资者交易、申购、赎回基金份额时参考。基金份额参考净值的计算方法详见招募说明书。\n"
                "深圳证券交易所和基金管理人可以调整基金份额参考净值的计算方法，并予以公告。"
            )
        else:
            v["LISTING_ITEM3_DEF"] = f"3、《{exch_cn}证券投资基金上市规则》规定的其他条件。"
            v["LISTING_SECTION2_BODY"] = (
                f"本基金基金份额在{exch_cn}的上市交易，应遵照《{exch_cn}交易规则》、《{exch_cn}证券投资基金上市规则》、"
                f"《{exch_cn}证券投资基金交易和申购赎回实施细则》等有关规定。"
            )
            v["LISTING_SECTION3_TEXT"] = (
                "三、停复牌、暂停上市、恢复上市及终止上市交易\n"
                f"基金份额在{exch_cn}上市后，如遇停复牌、暂停上市、恢复上市或终止上市的情形，按照{exch_cn}的相关规定执行。当基金发生{exch_cn}规定的因不再具备上市条件而应当终止上市的情形时，在法律法规和监管机构允许的情况下，本基金可变更为跟踪标的指数的普通开放式基金或上市开放式基金（LOF），无需召开基金份额持有人大会。届时，基金管理人可变更本基金的登记机构、相应调整申购赎回业务规则、提前制定基金终止上市后场内份额的处理规则并公告，同时，基金管理人可按照《信息披露办法》的规定，公告变更后的基金合同及招募说明书。若届时本基金管理人已有以该指数作为标的指数的基金，则本基金将本着维护投资者合法权益的原则，履行适当的程序后与该指数基金合并或选取其他合适的指数作为标的指数。具体情况见基金管理人届时公告。"
            )
            v["LISTING_SECTION4_TEXT"] = "四、基金份额参考净值（IOPV）的计算与公告详见招募说明书。"

        standard_part8_section4 = (
            "四、申购与赎回的程序\n"
            "1、申购和赎回的申请方式\n"
            "投资人须按申购赎回代理券商规定的手续，在开放日的开放时间提出申购、赎回的申请。投资人申购本基金时，须根据申购赎回清单备足申购对价；投资人提交赎回申请时，必须持有足够的基金份额余额和现金。否则所提交的申购、赎回申请无效。\n"
            "2、申购和赎回申请的确认\n"
            "本基金申购和赎回的确认适用《业务规则》的规定，具体在招募说明书中列示。如投资者未能提供符合要求的申购对价，则申购申请失败。如投资者持有的符合要求的基金份额不足或未能根据要求准备足额的现金，或基金投资组合内不具备足额的符合要求的赎回对价，或投资人提交的赎回申请超过基金管理人设定的当日净赎回份额上限、当日累计赎回份额上限、单个账户当日净赎回份额上限或单个账户当日累计赎回份额上限，则赎回申请失败。\n"
            "申购赎回代理券商受理申购、赎回申请并不代表该申购、赎回申请一定成功。申购、赎回的确认以登记机构的确认结果为准。对于申购、赎回申请的确认情况，投资人应及时查询并妥善行使合法权利。\n"
            "3、申购和赎回的清算交收与登记\n"
            "本基金申购和赎回过程中涉及的申购赎回对价和基金份额的交收适用《业务规则》的规定，具体在招募说明书中列示。\n"
            "如果登记机构和基金管理人在清算交收时发现不能正常履约的情形，则依据《业务规则》和参与各方相关协议及其不时修订的有关规定进行处理。\n"
            "投资人应按照本基金合同的约定和申购赎回代理券商的规定按时足额支付应付的现金差额、现金替代和现金替代退补款。因投资人原因导致现金差额、现金替代和现金替代退补款未能按时足额交收的，基金管理人有权为基金的利益向该投资人追偿并要求其承担由此导致的其他基金份额持有人或基金资产的损失。\n"
            f"4、如{exch_cn}、中国证券登记结算有限责任公司修改或更新上述规则并适用于本基金的，则按照新的规则执行。基金管理人在不损害基金份额持有人权益、并不违背交易所和登记机构相关规则的情况下可更改上述程序。基金管理人最迟须于新规则开始日前按照《信息披露办法》的有关规定在规定媒介公告。"
        )
        standard_part8_section6 = (
            "六、申购和赎回的对价、费用及其用途\n"
            "1、本基金份额净值的计算，保留到小数点后4位，小数点后第5位四舍五入，由此产生的收益或损失由基金财产承担。T日的基金份额净值在当天收市后计算，并在T+1日内公告。遇特殊情况，经履行适当程序，可以适当延迟计算或公告。\n"
            f"{v['PURCHASE_REDEMPTION_CONSIDERATION_DEF']}\n"
            f"3、申购赎回清单由基金管理人编制。T日的申购赎回清单在当日{exch_cn}开市前公告。申购赎回清单的内容与格式示例参见招募说明书。\n"
            "4、投资人在申购或赎回基金份额时，申购赎回代理券商可按照一定的标准收取佣金，其中包含证券交易所、登记机构等收取的相关费用，具体规定请参见招募说明书及基金产品资料概要。\n"
            "5、若市场情况发生变化，或相关业务规则发生变化，基金管理人可以在不违反相关法律法规且对基金份额持有人利益无实质性不利影响的情况下，对基金份额净值、申购赎回清单的计算和公告时间进行调整并提前公告。"
        )
        standard_part8_section7 = (
            "七、拒绝或暂停申购的情形\n"
            "发生下列情况时，基金管理人可拒绝或暂停接受投资人的申购申请：\n"
            "1、因不可抗力导致基金无法正常运作。\n"
            "2、发生基金合同规定的暂停基金资产估值情况时，基金管理人可暂停接受投资人的申购申请。\n"
            "3、本基金进行交易的主要证券/期货交易所交易时间非正常停市，可能影响本基金投资运作，或导致基金管理人无法计算当日基金资产净值。\n"
            "4、接受某笔或某些申购申请可能会影响或损害现有基金份额持有人利益或对存量基金份额持有人利益构成潜在重大不利影响时。\n"
            "5、基金资产规模过大，使基金管理人无法找到合适的投资品种，或其他可能对基金业绩产生负面影响，或发生其他损害现有基金份额持有人利益的情形。\n"
            "6、当前一估值日基金资产净值50%以上的资产出现无可参考的活跃市场价格且采用估值技术仍导致公允价值存在重大不确定性时，经与基金托管人协商确认后，基金管理人应当暂停接受基金申购申请。\n"
            "7、相关证券/期货交易所、申购赎回代理券商、登记机构等因异常情况无法办理申购，或者指数编制单位、相关证券/期货交易所等因异常情况使申购赎回清单无法编制或编制不当。上述异常情况指基金管理人无法预见并不可控制的情形，包括但不限于系统故障、网络故障、通讯故障、电力故障、数据错误等。\n"
            "8、基金管理人开市前因异常情况无法公布申购赎回清单，或基金管理人在开市后发现申购赎回清单编制错误或基金份额参考净值计算错误。\n"
            "9、当日申购申请达到基金管理人设定的申购份额上限时。\n"
            "10、本基金进行交易的主要证券/期货市场或本基金的资产组合中的重要部分发生暂停交易或其他重大事件，继续接受申购可能会影响或损害其他基金份额持有人利益时。\n"
            f"11、法律法规、{exch_cn}规定或中国证监会认定的其他情形。\n"
            "发生上述除第4项和第9项以外的暂停申购情形之一且基金管理人决定暂停接受申购申请时，基金管理人应当根据有关规定在规定媒介上刊登暂停申购公告。如果投资人的申购申请被拒绝，被拒绝的申购对价将退还给投资人。在暂停申购的情况消除时，基金管理人应及时恢复申购业务的办理。"
        )
        standard_part8_section8 = (
            "八、暂停赎回或延缓支付赎回对价的情形\n"
            "发生下列情形时，基金管理人可暂停接受投资人的赎回申请或延缓支付赎回对价：\n"
            "1、因不可抗力导致基金管理人不能支付赎回对价。\n"
            "2、发生基金合同规定的暂停基金资产估值情况时，基金管理人可暂停接受投资人的赎回申请或延缓支付赎回对价。\n"
            "3、本基金进行交易的主要证券/期货交易所交易时间非正常停市，可能影响本基金投资运作，或导致基金管理人无法计算当日基金资产净值。\n"
            "4、发生继续接受赎回申请将损害现有基金份额持有人利益的情形时，基金管理人可暂停接受基金份额持有人的赎回申请。\n"
            "5、当前一估值日基金资产净值50%以上的资产出现无可参考的活跃市场价格且采用估值技术仍导致公允价值存在重大不确定性时，经与基金托管人协商确认后，基金管理人应当延缓支付赎回对价或暂停接受基金赎回申请。\n"
            "6、基金管理人开市前因异常情况无法公布申购赎回清单，或基金管理人在开市后发现申购赎回清单编制错误或基金份额参考净值计算错误。\n"
            "7、当日赎回申请超过基金管理人根据市场情况设置的当日净赎回份额上限、当日累计赎回份额上限、单个账户当日净赎回份额上限或单个账户当日累计赎回份额上限。\n"
            "8、本基金进行交易的主要证券/期货市场或本基金的资产组合中的重要部分发生暂停交易或其他重大事件，继续接受赎回可能会影响或损害其他基金份额持有人利益时。\n"
            f"9、法律法规、{exch_cn}规定或中国证监会认定的其他情形。\n"
            "发生上述第4项和第7项以外的暂停赎回情形之一且基金管理人决定暂停赎回申请或延缓支付赎回对价时，基金管理人应当根据有关规定在规定媒介上刊登暂停赎回公告。在暂停赎回的情况消除时，基金管理人应及时恢复赎回业务的办理并公告。"
        )
        standard_part8_section9 = (
            "九、其他申购赎回方式\n"
            "1、若基金管理人推出以本基金为目标ETF的联接基金（可能由基金管理人另行募集或由基金管理人已管理的其他证券投资基金转型而形成），本基金可根据实际情况需要向本基金的联接基金开通特殊申购，不收取申购费用。\n"
            "2、在不违反法律法规且对基金份额持有人利益无实质性不利影响的情况下， 基金管理人可以根据具体情况履行适当程序后开通本基金的场外申购赎回等业务，无需召开基金份额持有人大会。场外申购赎回的具体办理方式等相关事项届时将另行公告。\n"
            "3、基金管理人可以在不违反法律法规规定且对基金份额持有人利益无实质性不利影响的情况下，调整基金申购赎回方式或申购赎回对价组成，并提前公告。\n"
            "4、在条件允许时，基金管理人可开放集合申购。在不损害基金份额持有人利益的前提下，基金管理人有权制定集合申购业务的相关规则，集合申购业务的相关规则在开始执行前将予以公告。\n"
            "5、在对基金份额持有人利益无实质性不利影响的情况下，基金管理人也可采取其他合理的申购方式，并于新的申购方式开始执行前予以公告。\n"
            "6、基金管理人指定的代理机构可依据基金合同开展其他服务，双方需签订书面委托代理协议。"
        )

        v["PART8_INTRO_TEXT"] = ""
        v["OPEN_DAY_CLAUSE"] = (
            f"投资人在开放日办理基金份额的申购和赎回，具体办理时间为{v['WORKING_DAY_DEF']}的交易时间，但基金管理人根据法律法规、中国证监会的要求或本基金合同的规定公告暂停申购、赎回时除外。"
        )
        v["OPEN_DAY_ADJUSTMENT_CLAUSE"] = (
            "基金合同生效后，若出现新的证券/期货交易市场、证券/期货交易所交易时间变更或其他特殊情况，基金管理人将视情况对前述开放日及开放时间进行相应的调整，但应在实施日前依照《信息披露办法》的有关规定在规定媒介上公告。"
        )
        v["PART8_SECTION4_TEXT"] = standard_part8_section4
        v["PART8_SECTION6_TEXT"] = standard_part8_section6
        v["PART8_SECTION7_TEXT"] = standard_part8_section7
        v["PART8_SECTION8_TEXT"] = standard_part8_section8
        v["PART8_SECTION9_TEXT"] = standard_part8_section9
        v["HK_PROXY_VOTING_TEXT"] = ""

        if hk_variant == "SSE_HK":
            hk_intro = (
                "目前本基金的申购赎回采用全现金替代模式，申购对价、赎回对价包括现金替代、现金差额及其他对价。未来在证券交易所和登记机构系统允许的情况下，本基金可采用实物申购赎回模式，申购对价、赎回对价包括组合证券、现金替代、现金差额及其他对价。"
            )
            v["PART8_INTRO_TEXT"] = hk_intro
            v["OPEN_DAY_CLAUSE"] = (
                "投资人在开放日办理基金份额的申购和赎回，具体办理时间为上海证券交易所的正常交易日的交易时间（若该交易日非港股通交易日，则本基金不开放申购和赎回），但基金管理人根据法律法规、中国证监会的要求或本基金合同的规定公告暂停申购、赎回时除外。"
            )
            v["OPEN_DAY_ADJUSTMENT_CLAUSE"] = (
                "基金合同生效后，若出现新的证券/期货交易市场、证券/期货交易所交易时间、港股通交易规则变更或其他特殊情况，基金管理人将视情况对前述开放日及开放时间进行相应的调整，但应在实施日前依照《信息披露办法》的有关规定在规定媒介上公告。"
            )
            v["PART8_SECTION4_TEXT"] = (
                "四、申购与赎回的程序\n"
                f"{hk_intro}\n"
                "1、申购和赎回的申请方式\n"
                "投资人须按申购赎回代理券商规定的手续，在开放日的开放时间提出申购、赎回的申请。投资人申购本基金时，须根据申购赎回清单备足申购对价；投资人提交赎回申请时，必须持有足够的基金份额余额和现金。否则所提交的申购、赎回申请无效。\n"
                "2、申购和赎回申请的确认\n"
                "本基金申购和赎回的确认适用《业务规则》的规定，具体在招募说明书中列示。如投资者未能提供符合要求的申购对价，则申购申请失败。如投资者持有的符合要求的基金份额不足或未能根据要求准备足额的现金，或基金投资组合内不具备足额的符合要求的赎回对价，或投资人提交的赎回申请超过基金管理人设定的当日净赎回份额上限、当日累计赎回份额上限、单个账户当日净赎回份额上限或单个账户当日累计赎回份额上限，则赎回申请失败。\n"
                "申购赎回代理券商对申购、赎回申请的受理并不代表该申请一定成功，而仅代表申购赎回代理券商确实接收到该申请。申购、赎回的确认以登记机构的确认结果为准。对于申购、赎回申请的确认情况，投资人应及时查询并妥善行使合法权利。如投资者未能提供符合要求的申购对价，则申购申请失败。申购款项本金将退回投资人账户，基金管理人、基金托管人和销售机构等不承担由此产生的利息等任何损失。\n"
                "3、申购和赎回的清算交收与登记\n"
                "本基金申购赎回过程中涉及的基金份额、组合证券（如适用）、现金替代、现金差额及其他对价的清算交收适用《业务规则》和参与各方相关协议及其不时修订的有关规定。具体规则在招募说明书中说明。\n"
                "如果登记机构和基金管理人在清算交收时发现不能正常履约的情形，则依据《业务规则》和参与各方相关协议及其不时修订的有关规定进行处理。\n"
                "投资人应按照本基金合同的约定和申购赎回代理券商的规定按时足额支付应付的现金差额、现金替代和现金替代退补款。因投资人原因导致现金差额、现金替代和现金替代退补款未能按时足额交收的，基金管理人有权为基金的利益向该投资人追偿并要求其承担由此导致的其他基金份额持有人或基金资产的损失。\n"
                "如遇港股通暂停交易或交收、登记公司系统故障、交易所或交易市场数据传输延迟、通讯系统故障、银行数据交换系统故障、港股通交易系统或港股通资金交收规则限制或其它非基金管理人及基金托管人所能控制的因素影响业务处理流程，则赎回对价的支付时间可相应顺延。在发生本基金合同载明的其他暂停赎回或延缓支付赎回对价的情形时，赎回对价的支付办法参照本基金合同有关条款处理。\n"
                "4、如上海证券交易所、中国证券登记结算有限责任公司修改或更新上述规则并适用于本基金的，则按照新的规则执行，并在招募说明书中进行更新。基金管理人、上海证券交易所和登记机构可在法律法规允许的范围内，对基金份额持有人利益不存在实质不利影响的前提下，对上述规则进行调整，并在开始实施前按照《信息披露办法》的有关规定在规定媒介上予以公告。"
            )
            v["PART8_SECTION6_TEXT"] = (
                "六、申购和赎回的对价、费用及其用途\n"
                "1、本基金份额净值的计算，保留到小数点后4位，小数点后第5位四舍五入，由此产生的收益或损失由基金财产承担。T日的基金份额净值在当天收市后计算，并在T+1日内公告。遇特殊情况，经履行适当程序，可以适当延迟计算或公告。未来，若市场情况发生变化，或实际情况需要，经中国证监会允许，本基金可相应调整基金净值计算和公告时间或频率并依照《信息披露办法》的有关规定提前公告。\n"
                f"{v['PURCHASE_REDEMPTION_CONSIDERATION_DEF']}\n"
                "3、申购赎回清单由基金管理人编制。T日的申购赎回清单在当日上海证券交易所开市前公告。申购赎回清单的内容与格式示例参见招募说明书。\n"
                "4、投资人在申购或赎回基金份额时，申购赎回代理券商可按照一定的标准收取佣金，其中包含证券交易所、登记机构等收取的相关费用，具体规定请参见招募说明书及基金产品资料概要。\n"
                "5、若市场情况发生变化，或相关业务规则发生变化，基金管理人可以在不违反相关法律法规且对基金份额持有人利益无实质性不利影响的情况下，对申购对价/赎回对价组成、基金份额净值、申购赎回清单的计算和公告时间进行调整并提前公告。"
            )
            v["PART8_SECTION7_TEXT"] = (
                "七、拒绝或暂停申购的情形\n"
                "发生下列情况时，基金管理人可拒绝或暂停接受投资人的申购申请：\n"
                "1、因不可抗力导致基金无法正常运作。\n"
                "2、发生基金合同规定的暂停基金资产估值情况时，基金管理人可暂停接受投资人的申购申请。\n"
                "3、本基金进行交易的主要证券/期货交易所交易时间非正常停市或港股通临时停市，可能影响本基金投资运作，或导致基金管理人无法计算当日基金资产净值。\n"
                "4、接受某笔或某些申购申请可能会影响或损害现有基金份额持有人利益时。\n"
                "5、基金资产规模过大，使基金管理人无法找到合适的投资品种，或其他可能对基金业绩产生负面影响，或发生其他损害现有基金份额持有人利益的情形。\n"
                "6、当前一估值日基金资产净值50%以上的资产出现无可参考的活跃市场价格且采用估值技术仍导致公允价值存在重大不确定性时，经与基金托管人协商确认后，基金管理人应当暂停接受基金申购申请。\n"
                "7、相关证券/期货交易所、申购赎回代理券商、登记机构等因异常情况无法办理申购，或者指数编制单位、相关证券/期货交易所等因异常情况使申购赎回清单无法编制或编制不当。上述异常情况指基金管理人无法预见并不可控制的情形，包括但不限于系统故障、网络故障、通讯故障、电力故障、数据错误等。\n"
                "8、基金管理人开市前因异常情况无法公布申购赎回清单，或基金管理人在开市后发现申购赎回清单编制错误或基金份额参考净值计算错误。\n"
                "9、基金管理人可根据市场情况在申购赎回清单中设置申购上限，当一笔新的申购申请被确认成功，会使本基金当日申购超过申购赎回清单中规定的申购上限时，该笔申购申请将被拒绝。\n"
                "10、因港股通每日额度等原因需要控制基金申购规模（基金管理人可根据市场情况进行调整）。\n"
                "11、本基金进行交易的主要证券/期货市场或本基金的资产组合中的重要部分发生暂停交易或其他重大事件，继续接受申购可能会影响或损害其他基金份额持有人利益时。\n"
                "12、法律法规规定、中国证监会或上海证券交易所认定的其他情形。\n"
                "发生上述除第4项和第9项以外的暂停申购情形之一且基金管理人决定暂停接受申购申请时，基金管理人应当根据有关规定在规定媒介上刊登暂停申购公告。如果投资人的申购申请被拒绝，被拒绝的申购对价将退还给投资人。在暂停申购的情况消除时，基金管理人应及时恢复申购业务的办理。"
            )
            v["PART8_SECTION8_TEXT"] = (
                "八、暂停赎回或延缓支付赎回对价的情形\n"
                "发生下列情形时，基金管理人可暂停接受投资人的赎回申请或延缓支付赎回对价：\n"
                "1、因不可抗力导致基金管理人不能支付赎回对价。\n"
                "2、发生基金合同规定的暂停基金资产估值情况时，基金管理人可暂停接受投资人的赎回申请或延缓支付赎回对价。 \n"
                "3、本基金进行交易的主要证券/期货交易所交易时间非正常停市或港股通临时停市，可能影响本基金投资运作，或导致基金管理人无法计算当日基金资产净值。\n"
                "4、发生继续接受赎回申请将损害现有基金份额持有人利益的情形时，基金管理人可暂停接受基金份额持有人的赎回申请。\n"
                "5、当前一估值日基金资产净值50%以上的资产出现无可参考的活跃市场价格且采用估值技术仍导致公允价值存在重大不确定性时，经与基金托管人协商确认后，基金管理人应当暂停接受基金赎回申请或延缓支付赎回对价。\n"
                "6、基金管理人开市前因异常情况无法公布申购赎回清单，或基金管理人在开市后发现申购赎回清单编制错误或基金份额参考净值计算错误。\n"
                "7、当日赎回申请超过基金管理人根据市场情况设置的当日净赎回份额上限、当日累计赎回份额上限、单个账户当日净赎回份额上限或单个账户当日累计赎回份额上限。\n"
                "8、本基金进行交易的主要证券/期货市场或本基金的资产组合中的重要部分发生暂停交易或其他重大事件，继续接受赎回可能会影响或损害其他基金份额持有人利益时。\n"
                "9、法律法规规定、中国证监会或上海证券交易所认定的其他情形。\n"
                "发生上述第4项和第7项以外的暂停赎回情形之一且基金管理人决定暂停赎回申请或延缓支付赎回对价时，基金管理人应当根据有关规定在规定媒介上刊登暂停赎回公告。在暂停赎回的情况消除时，基金管理人应及时恢复赎回业务的办理并公告。"
            )
            v["PART8_SECTION9_TEXT"] = (
                "九、其他申购赎回方式\n"
                "1、若基金管理人推出以本基金为目标ETF的联接基金（可能由基金管理人另行募集或由基金管理人已管理的其他证券投资基金转型而形成），本基金可根据实际情况需要向本基金的联接基金开通特殊申购，不收取申购费用。\n"
                "2、在不违反法律法规且对基金份额持有人利益无实质性不利影响的情况下， 基金管理人可以根据具体情况履行适当程序后开通本基金的场外申购赎回等业务，无需召开基金份额持有人大会。场外申购赎回的具体办理方式等相关事项届时将另行公告。\n"
                "3、基金管理人可以在不违反法律法规规定且对持有人利益无实质性不利影响的情况下，调整基金申购赎回方式或申购赎回对价组成，并提前公告。\n"
                "4、在条件允许时，基金管理人可开放集合申购。在不损害基金份额持有人利益的前提下，基金管理人有权制定集合申购业务的相关规则，集合申购业务的相关规则在开始执行前将予以公告。\n"
                "5、在对基金份额持有人利益无实质性不利影响的情况下，基金管理人也可采取其他合理的申购方式，并于新的申购方式开始执行前予以公告。\n"
                "6、未来在证券交易所和登记结算机构系统允许的情况下，本基金可采用实物申购赎回模式，申购对价、赎回对价包括组合证券、现金替代、现金差额及其他对价，具体的规则和程序由基金管理人在开通前予以公告，不需召开基金份额持有人大会。\n"
                "7、基金管理人指定的代理机构可依据本基金合同开展其他服务，双方需签订书面委托代理协议。"
            )
        elif hk_variant == "SZSE_HK":
            hk_intro = (
                "目前本基金的申购赎回采用全现金替代模式，申购对价、赎回对价包括现金替代、现金差额及其他对价。未来在证券交易所和登记机构系统允许的情况下，本基金可采用实物申购赎回模式，申购对价、赎回对价包括组合证券、现金替代、现金差额及其他对价。"
            )
            v["PART8_INTRO_TEXT"] = hk_intro
            v["OPEN_DAY_CLAUSE"] = (
                "投资人在开放日办理基金份额的申购和赎回，本基金的开放日为深圳证券交易所的交易日（若该交易日非港股通交易日，则本基金不开放申购和赎回），具体办理时间为深圳证券交易所的正常交易日的交易时间，但基金管理人根据法律法规、中国证监会的要求或本基金合同的规定公告暂停申购、赎回时除外。"
            )
            v["OPEN_DAY_ADJUSTMENT_CLAUSE"] = (
                "基金合同生效后，若出现新的证券/期货交易市场、证券/期货交易所交易时间变更、港股通交易规则变更或其他特殊情况，基金管理人将视情况对前述开放日及开放时间进行相应的调整，但应在实施日前依照《信息披露办法》的有关规定在规定媒介上公告。"
            )
            v["PART8_SECTION4_TEXT"] = (
                "四、申购与赎回的程序\n"
                f"{hk_intro}\n"
                "1、申购和赎回的申请方式\n"
                "投资人须按申购赎回代理券商规定的手续，在开放日的开放时间提出申购、赎回的申请。投资人申购本基金时，须根据申购赎回清单备足申购对价；投资人提交赎回申请时，必须持有足够的基金份额余额和现金。否则所提交的申购、赎回申请无效。\n"
                "2、申购和赎回申请的确认\n"
                "本基金申购和赎回的确认适用《业务规则》的规定，具体在招募说明书中列示。如投资者未能提供符合要求的申购对价，则申购申请失败。如投资者持有的符合要求的基金份额不足或未能根据要求准备足额的现金，或基金投资组合内不具备足额的符合要求的赎回对价，或投资人提交的赎回申请超过基金管理人设定的当日净赎回份额上限、当日累计赎回份额上限、单个账户当日净赎回份额上限或单个账户当日累计赎回份额上限，则赎回申请失败。\n"
                "申购赎回代理券商受理申购、赎回申请并不代表该申购、赎回申请一定成功。申购、赎回的确认以登记机构的确认结果为准。对于申购、赎回申请的确认情况，投资人应及时查询并妥善行使合法权利。\n"
                "3、申购和赎回的清算交收与登记\n"
                "本基金申购和赎回过程中涉及的申购赎回对价和基金份额的交收适用《业务规则》的规定，具体在招募说明书中列示。\n"
                "如果登记机构和基金管理人在清算交收时发现不能正常履约的情形，则依据《业务规则》和参与各方相关协议及其不时修订的有关规定进行处理。\n"
                "投资人应按照本基金合同的约定和申购赎回代理券商的规定按时足额支付应付的现金差额、现金替代和现金替代退补款。因投资人原因导致现金差额、现金替代和现金替代退补款未能按时足额交收的，基金管理人有权为基金的利益向该投资人追偿并要求其承担由此导致的其他基金份额持有人或基金资产的损失。\n"
                "如遇港股通暂停交易或交收、登记公司系统故障、交易所或交易市场数据传输延迟、通讯系统故障、银行数据交换系统故障、港股通交易系统或港股通资金交收规则限制或其它非基金管理人及基金托管人所能控制的因素影响业务处理流程，则赎回对价的支付时间可相应顺延。在发生本基金合同载明的其他暂停赎回或延缓支付赎回对价的情形时，赎回对价的支付办法参照本基金合同有关条款处理。\n"
                "4、如深圳证券交易所、中国证券登记结算有限责任公司修改或更新上述规则并适用于本基金的，则按照新的规则执行。基金管理人在不损害基金份额持有人权益、并不违背交易所和登记机构相关规则的情况下可更改上述程序。基金管理人最迟须于新规则开始日前按照《信息披露办法》的有关规定在规定媒介公告。"
            )
            v["PART8_SECTION6_TEXT"] = (
                "六、申购和赎回的对价、费用及其用途\n"
                "1、本基金份额净值的计算，保留到小数点后4位，小数点后第5位四舍五入，由此产生的收益或损失由基金财产承担。T日的基金份额净值在当天收市后计算，并在T+1日内公告。遇特殊情况，经履行适当程序，可以适当延迟计算或公告。\n"
                f"{v['PURCHASE_REDEMPTION_CONSIDERATION_DEF']}\n"
                "3、申购赎回清单由基金管理人编制。T日的申购赎回清单在当日深圳证券交易所开市前公告。申购赎回清单的内容与格式示例参见招募说明书。\n"
                "4、投资人在申购或赎回基金份额时，申购赎回代理券商可按照一定的标准收取佣金，其中包含证券交易所、登记机构等收取的相关费用，具体规定请参见招募说明书及基金产品资料概要。\n"
                "5、若市场情况发生变化，或相关业务规则发生变化，基金管理人可以在不违反相关法律法规且对基金份额持有人利益无实质性不利影响的情况下，对申购对价/赎回对价组成、基金份额净值、申购赎回清单的计算和公告时间进行调整并提前公告。"
            )
            v["PART8_SECTION7_TEXT"] = (
                "七、拒绝或暂停申购的情形\n"
                "发生下列情况时，基金管理人可拒绝或暂停接受投资人的申购申请：\n"
                "1、因不可抗力导致基金无法正常运作。\n"
                "2、发生基金合同规定的暂停基金资产估值情况时，基金管理人可暂停接受投资人的申购申请。\n"
                "3、本基金进行交易的主要证券/期货交易所交易时间非正常停市或港股通临时停市，可能影响本基金投资运作，或导致基金管理人无法计算当日基金资产净值。\n"
                "4、接受某笔或某些申购申请可能会影响或损害现有基金份额持有人利益或对存量基金份额持有人利益构成潜在重大不利影响时。\n"
                "5、基金资产规模过大，使基金管理人无法找到合适的投资品种，或其他可能对基金业绩产生负面影响，或发生其他损害现有基金份额持有人利益的情形。\n"
                "6、当前一估值日基金资产净值50%以上的资产出现无可参考的活跃市场价格且采用估值技术仍导致公允价值存在重大不确定性时，经与基金托管人协商确认后，基金管理人应当暂停接受基金申购申请。\n"
                "7、当日申购申请达到基金管理人设定的申购份额上限时。\n"
                "8、本基金进行交易的主要证券/期货市场或本基金的资产组合中的重要部分发生暂停交易或其他重大事件，继续接受申购可能会影响或损害其他基金份额持有人利益时。\n"
                "9、因港股通每日额度等原因需要控制基金申购规模（基金管理人可根据市场情况进行调整）。\n"
                "10、法律法规、深圳证券交易所规定或中国证监会认定的其他情形。\n"
                "发生上述除第4项和第7项以外的暂停申购情形之一且基金管理人决定暂停接受申购申请时，基金管理人应当根据有关规定在规定媒介上刊登暂停申购公告。如果投资人的申购申请被拒绝，被拒绝的申购对价将退还给投资人。在暂停申购的情况消除时，基金管理人应及时恢复申购业务的办理。"
            )
            v["PART8_SECTION8_TEXT"] = (
                "八、暂停赎回或延缓支付赎回对价的情形\n"
                "发生下列情形时，基金管理人可暂停接受投资人的赎回申请或延缓支付赎回对价：\n"
                "1、因不可抗力导致基金管理人不能支付赎回对价。\n"
                "2、发生基金合同规定的暂停基金资产估值情况时，基金管理人可暂停接受投资人的赎回申请或延缓支付赎回对价。\n"
                "3、本基金进行交易的主要证券/期货交易所交易时间非正常停市或港股通临时停市，可能影响本基金投资运作，或导致基金管理人无法计算当日基金资产净值。\n"
                "4、发生继续接受赎回申请将损害现有基金份额持有人利益的情形时，基金管理人可暂停接受基金份额持有人的赎回申请。\n"
                "5、当前一估值日基金资产净值50%以上的资产出现无可参考的活跃市场价格且采用估值技术仍导致公允价值存在重大不确定性时，经与基金托管人协商确认后，基金管理人应当延缓支付赎回对价或暂停接受基金赎回申请。\n"
                "6、当日赎回申请超过基金管理人根据市场情况设置的当日净赎回份额上限、当日累计赎回份额上限、单个账户当日净赎回份额上限或单个账户当日累计赎回份额上限。\n"
                "7、本基金进行交易的主要证券/期货市场或本基金的资产组合中的重要部分发生暂停交易或其他重大事件，继续接受赎回可能会影响或损害其他基金份额持有人利益时。\n"
                "8、法律法规、深圳证券交易所规定或中国证监会认定的其他情形。\n"
                "发生上述第4项和第6项以外的暂停赎回情形之一且基金管理人决定暂停赎回申请或延缓支付赎回对价时，基金管理人应当根据有关规定在规定媒介上刊登暂停赎回公告。在暂停赎回的情况消除时，基金管理人应及时恢复赎回业务的办理并公告。"
            )
            v["PART8_SECTION9_TEXT"] = (
                "九、其他申购赎回方式\n"
                "1、若基金管理人推出以本基金为目标ETF的联接基金（可能由基金管理人另行募集或由基金管理人已管理的其他证券投资基金转型而形成），本基金可根据实际情况需要向本基金的联接基金开通特殊申购，不收取申购费用。\n"
                "2、在不违反法律法规且对基金份额持有人利益无实质性不利影响的情况下， 基金管理人可以根据具体情况履行适当程序后开通本基金的场外申购赎回等业务，无需召开基金份额持有人大会。场外申购赎回的具体办理方式等相关事项届时将另行公告。\n"
                "3、基金管理人可以在不违反法律法规规定且对持有人利益无实质性不利影响的情况下，调整基金申购赎回方式或申购赎回对价组成，并提前公告。\n"
                "4、在条件允许时，基金管理人可开放集合申购。在不损害基金份额持有人利益的前提下，基金管理人有权制定集合申购业务的相关规则，集合申购业务的相关规则在开始执行前将予以公告。\n"
                "5、在对基金份额持有人利益无实质性不利影响的情况下，基金管理人也可采取其他合理的申购方式，并于新的申购方式开始执行前予以公告。\n"
                "6、未来在证券交易所和登记结算机构系统允许的情况下，本基金可采用实物申购赎回模式，申购对价、赎回对价包括组合证券、现金替代、现金差额及其他对价，具体的规则和程序由基金管理人在开通前予以公告，不需召开基金份额持有人大会。\n"
                "7、基金管理人指定的代理机构可依据本基金合同开展其他服务，双方需签订书面委托代理协议。"
            )

        if v.get("HAS_HK_CONNECT"):
            v["HK_PROXY_VOTING_TEXT"] = (
                "八、港股通股票投资的代理投票\n"
                "本基金通过港股通买入的股票记录在中国证券登记结算有限责任公司在香港中央结算有限公司开立的证券账户。中国证券登记结算有限责任公司以自己的名义，通过香港中央结算有限公司行使对该股票发行人的权利。中国证券登记结算有限责任公司行使对该股票发行人的权利，将通过证券公司事先征求包括基金管理人在内的内地投资者的意见，并按照其意见办理。法律法规另有规定的，从其规定。"
            )

        v = self._apply_migrated_contract_clause_texts(v)
        return self._apply_business_text_overrides(v)

    # ── Step 4: 处理条件块（支持嵌套） ──────────────────────────────────────
    def _process_conditionals(self, text: str, v: dict) -> str:
        # We process from innermost outward using a stack-based approach
        max_passes = 10
        for _ in range(max_passes):
            new_text = self._single_pass_conditionals(text, v)
            if new_text == text:
                break
            text = new_text
        return text

    def _single_pass_conditionals(self, text: str, v: dict) -> str:
        # Match innermost {{IF ...}}...{{ENDIF}} (no nested IF inside)
        pattern_if = re.compile(
            r'\{\{IF\s+(\w+)\}\}((?:(?!\{\{IF)[\s\S])*?)\{\{ENDIF\}\}',
            re.DOTALL
        )
        pattern_if_not = re.compile(
            r'\{\{IF_NOT\s+(\w+)\}\}((?:(?!\{\{IF)[\s\S])*?)\{\{ENDIF\}\}',
            re.DOTALL
        )

        def replace_if(m):
            var_name = m.group(1)
            content = m.group(2)
            val = v.get(var_name, False)
            if isinstance(val, str):
                val = val.lower() in ("true", "1", "yes", "是")
            return content if val else ""

        def replace_if_not(m):
            var_name = m.group(1)
            content = m.group(2)
            val = v.get(var_name, False)
            if isinstance(val, str):
                val = val.lower() in ("true", "1", "yes", "是")
            return "" if val else content

        text = pattern_if.sub(replace_if, text)
        text = pattern_if_not.sub(replace_if_not, text)
        return text

    # ── Step 5: 替换占位符 ───────────────────────────────────────────────────
    def _replace_placeholders(self, text: str, v: dict) -> str:
        def replacer(m):
            key = m.group(1)
            val = v.get(key)
            if val is None:
                return m.group(0)  # 保留未定义的占位符
            if isinstance(val, bool):
                return "是" if val else "否"
            return str(val)

        return re.sub(r"\{([A-Z_][A-Z0-9_]*)\}", replacer, text)

    # ── Step 6: 清理 ─────────────────────────────────────────────────────────
    @staticmethod
    def _is_prospectus_toc_placeholder_line(line: str) -> bool:
        return bool(re.search(r'[\u3010\[](?:\u5f85\u586b\u5199|\u5f85\u8865\u5145)', line or ""))

    @staticmethod
    def _parse_prospectus_chapter_heading(line: str):
        match = re.match("^\u7b2c([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e]+)\u7ae0\\s+(.+?)(?:[\\t ]+\\d+)?$", (line or "").strip())
        if not match:
            return None
        return match.group(1), match.group(2).strip()

    @staticmethod
    def _format_prospectus_reference_heading(chapter_cn: str, title: str) -> str:
        return f"{chapter_cn}\u3001{(title or '').strip()}"

    @staticmethod
    def _parse_prospectus_toc_entry(line: str):
        match = re.match(
            r"^([一二三四五六七八九十百]+)、\s*(.+?)(?:[\t ]+\d+)?$",
            (line or "").strip(),
        )
        if not match:
            return None
        return match.group(1), match.group(2).strip()

    def _extract_prospectus_toc_entries(self, text: str) -> list[dict]:
        lines = text.splitlines()
        entries = []
        in_toc = False
        started = False

        for idx, raw in enumerate(lines):
            stripped = raw.strip()
            if not in_toc:
                if stripped == "目录":
                    in_toc = True
                continue

            if not stripped:
                continue

            parsed = self._parse_prospectus_toc_entry(stripped)
            if parsed:
                chapter_cn, title = parsed
                entries.append(
                    {
                        "chapter_cn": chapter_cn,
                        "title": title,
                        "display_title": self._format_prospectus_reference_heading(chapter_cn, title),
                        "line_index": idx,
                    }
                )
                started = True
                continue

            if started:
                break

        return entries

    def _format_reference_style_prospectus_with_plain_titles(self, text: str, toc_entries: list[dict]) -> str:
        lines = text.split("\n")
        formatted = []
        phase = "cover"
        body_index = 0

        for raw in lines:
            stripped = raw.strip()
            if not stripped:
                if phase == "body" and formatted and formatted[-1] != "":
                    formatted.append("")
                continue

            if phase == "cover":
                if stripped == "目录":
                    formatted.append("目录")
                    phase = "toc"
                    continue
                formatted.append(stripped)
                continue

            if phase == "toc":
                if self._is_prospectus_toc_placeholder_line(stripped):
                    continue

                parsed_toc = self._parse_prospectus_toc_entry(stripped)
                if parsed_toc:
                    formatted.append(self._format_prospectus_reference_heading(*parsed_toc))
                    continue

                if body_index < len(toc_entries) and stripped == toc_entries[body_index]["title"]:
                    formatted.append(toc_entries[body_index]["display_title"])
                    body_index += 1
                    phase = "body"
                    continue

                parsed_heading = self._parse_prospectus_chapter_heading(stripped)
                if parsed_heading:
                    formatted.append(self._format_prospectus_reference_heading(*parsed_heading))
                    phase = "body"
                    body_index = min(len(toc_entries), body_index + 1)
                continue

            if body_index < len(toc_entries) and stripped == toc_entries[body_index]["title"]:
                formatted.append(toc_entries[body_index]["display_title"])
                body_index += 1
                continue

            parsed_heading = self._parse_prospectus_chapter_heading(stripped)
            if parsed_heading:
                formatted.append(self._format_prospectus_reference_heading(*parsed_heading))
            else:
                formatted.append(raw.rstrip())

        text = "\n".join(formatted)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _format_reference_style_prospectus(self, text: str) -> str:
        toc_entries = self._extract_prospectus_toc_entries(text)
        if toc_entries:
            return self._format_reference_style_prospectus_with_plain_titles(text, toc_entries)

        lines = text.split("\n")
        formatted = []
        phase = "cover"
        toc_headings = []

        for raw in lines:
            stripped = raw.strip()
            if not stripped:
                if phase == "body" and formatted and formatted[-1] != "":
                    formatted.append("")
                continue

            if phase == "cover":
                if re.match(r"^\u76ee\s*\u5f55$", stripped):
                    formatted.append("\u76ee\u5f55")
                    phase = "toc"
                    continue
                formatted.append(stripped)
                continue

            parsed = self._parse_prospectus_chapter_heading(stripped)
            if phase == "toc":
                if self._is_prospectus_toc_placeholder_line(stripped):
                    continue
                if parsed:
                    heading = self._format_prospectus_reference_heading(*parsed)
                    if toc_headings and heading == toc_headings[0]:
                        phase = "body"
                        formatted.append(heading)
                    elif heading not in toc_headings:
                        toc_headings.append(heading)
                        formatted.append(heading)
                    continue
                if re.match(r"^[???????????]+?", stripped):
                    continue
                continue

            if parsed:
                formatted.append(self._format_prospectus_reference_heading(*parsed))
            else:
                formatted.append(raw.rstrip())

        text = "\n".join(formatted)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _inject_important_notice_before_toc(self, text: str, v: dict) -> str:
        ref = self._load_reference_fixed_content(v)
        notice = (ref.get("important_notice") or "").strip()
        if not notice or "重要提示" in text:
            return text

        lines = text.splitlines()
        toc_idx = next((idx for idx, line in enumerate(lines) if line.strip() == "目录"), None)
        if toc_idx is None:
            return text

        cover = "\n".join(lines[:toc_idx]).strip()
        tail = "\n".join(lines[toc_idx:]).strip()
        return "\n\n".join(part for part in (cover, notice, tail) if part).strip()

    def _cleanup(self, text: str) -> str:
        # Remove markdown blockquote header lines (> ...)
        lines = text.split("\n")
        clean = []
        skip_header = True
        for line in lines:
            # Skip the preamble before the actual contract starts
            if skip_header:
                if line.strip().startswith("{FUND_NAME}") or (
                    "基金合同" in line and "模板说明" not in line and not line.startswith(">")
                    and not line.startswith("#") and not line.startswith("**") and not line.startswith("---")
                ):
                    skip_header = False
                    clean.append(line)
                # Skip preamble lines
                continue
            # Remove comment/annotation lines
            if line.strip().startswith(">") or line.strip().startswith("**条件变量") or line.strip().startswith("**差异条款"):
                continue
            clean.append(line)

        text = "\n".join(clean)

        # Collapse 3+ consecutive blank lines into 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        return text

    @staticmethod
    def _normalize_contract_text(text: str) -> str:
        return text.replace("具体办理时间为指", "具体办理时间为")

    @staticmethod
    def _extract_contract_summary_block_body(section_text: str, label: str) -> str:
        blocks = _split_review_blocks(section_text)
        label_text = _normalize_review_text(label)
        if not blocks or not label_text:
            return ""

        exact_index = next(
            (
                index
                for index, block in enumerate(blocks)
                if _normalize_review_text(block.get("heading", "")) == label_text
            ),
            None,
        )
        if exact_index is not None:
            expanded = _expand_review_block_with_descendants(blocks, exact_index)
            return (expanded.get("body") or expanded.get("text") or "").strip()

        fallback_indexes = []
        for index, block in enumerate(blocks):
            heading = block.get("heading", "")
            if (
                _review_labels_match(heading, label_text)
                or _review_soft_heading_match(heading, label_text)
            ):
                fallback_indexes.append((len(_normalize_review_text(heading)), index))
        if fallback_indexes:
            _, best_index = max(fallback_indexes)
            expanded = _expand_review_block_with_descendants(blocks, best_index)
            return (expanded.get("body") or expanded.get("text") or "").strip()
        return ""

    @staticmethod
    def _extract_contract_summary_intro_before(section_text: str, stop_label: str) -> str:
        blocks = _split_review_blocks(section_text)
        stop_label_text = _normalize_review_text(stop_label)
        if not blocks or not section_text or not stop_label_text:
            return ""

        exact_index = next(
            (
                index
                for index, block in enumerate(blocks)
                if _normalize_review_text(block.get("heading", "")) == stop_label_text
            ),
            None,
        )
        fallback_indexes = []
        if exact_index is None:
            for index, block in enumerate(blocks):
                heading = block.get("heading", "")
                if _review_labels_match(heading, stop_label_text) or _review_soft_heading_match(heading, stop_label_text):
                    fallback_indexes.append((len(_normalize_review_text(heading)), index))

        target_index = exact_index
        if target_index is None and fallback_indexes:
            _, target_index = max(fallback_indexes)
        if target_index is None:
            return ""

        block_text = blocks[target_index].get("text") or ""
        if block_text and block_text in section_text:
            return section_text.split(block_text, 1)[0].strip()
        return ""

    def _extract_contract_summary_source_text(self, target: dict, locator: str) -> str:
        segments = _split_review_locator(_normalize_review_locator(locator))
        if segments and ("第" in segments[0] or "部分" in segments[0] or "章" in segments[0]):
            sub_labels = segments[1:]
        else:
            sub_labels = segments

        section = target.get("section") or {}
        if sub_labels and section.get("content"):
            expanded_body = self._extract_contract_summary_block_body(section.get("content") or "", sub_labels[-1])
            if expanded_body:
                return expanded_body
        return (target.get("body_text") or target.get("text") or "").strip()

    @staticmethod
    def _rewrite_contract_summary_internal_references(source_text: str, locator: str, output_heading: str | None) -> str:
        if not source_text or not output_heading:
            return source_text
        segments = _split_review_locator(_normalize_review_locator(locator))
        if not segments:
            return source_text
        source_heading = segments[-1]
        if not source_heading or source_heading == output_heading:
            return source_text

        rewritten = source_text
        for open_quote, close_quote in (("“", "”"), ("‘", "’"), ('"', '"')):
            rewritten = rewritten.replace(
                f"{open_quote}{source_heading}{close_quote}",
                f"{open_quote}{output_heading}{close_quote}",
            )

        source_marker = ContractEngine._contract_summary_heading_marker(source_heading)
        output_marker = ContractEngine._contract_summary_parenthesized_heading_marker(output_heading)
        if source_marker and output_marker and source_marker == output_marker:
            rewritten = re.sub(
                rf"第{re.escape(source_marker)}条",
                f"第（{output_marker}）条",
                rewritten,
            )
        return rewritten

    @staticmethod
    def _normalize_contract_summary_numbered_heading_references(text: str) -> str:
        if not text:
            return text

        text = re.sub(
            r"([“‘\"])([一二三四五六七八九十百千]+)[、.．]([^”’\"\n]+)([”’\"])",
            lambda m: f"{m.group(1)}（{m.group(2)}）{m.group(3)}{m.group(4)}",
            text,
        )
        return re.sub(
            r"([“‘\"])[（(]([一二三四五六七八九十百千]+)[）)][、.．]([^”’\"\n]+)([”’\"])",
            lambda m: f"{m.group(1)}（{m.group(2)}）{m.group(3)}{m.group(4)}",
            text,
        )

    @staticmethod
    def _contract_summary_heading_marker(heading: str) -> str:
        match = re.match(r"^\s*([一二三四五六七八九十百千]+)[、.．]", str(heading or ""))
        return match.group(1) if match else ""

    @staticmethod
    def _contract_summary_parenthesized_heading_marker(heading: str) -> str:
        match = re.match(r"^\s*[（(]([一二三四五六七八九十百千]+)[）)]", str(heading or ""))
        return match.group(1) if match else ""

    def _locate_contract_summary_source(self, sections: list[dict], locator: str, placeholder: str) -> dict:
        target = _locate_review_rule_target(sections, locator)
        if not target.get("matched"):
            reason = target.get("missing_reason") or "未匹配"
            raise ValueError(f"基金合同摘要占位符 {placeholder} 无法定位正文来源：{locator}（{reason}）")
        source_text = self._extract_contract_summary_source_text(target, locator)
        if not source_text:
            raise ValueError(f"基金合同摘要占位符 {placeholder} 的正文来源为空：{locator}")
        target = dict(target)
        target["summary_source_text"] = source_text
        return target

    def _render_contract_summary_fragment(self, sections: list[dict], placeholder: str, fragment: dict) -> str:
        locator = str(fragment.get("locator") or "").strip()
        target = self._locate_contract_summary_source(sections, locator, placeholder)
        source_text = str(target.get("summary_source_text") or "").strip()
        if fragment.get("take") == "intro":
            return self._extract_contract_summary_intro_before(
                source_text,
                str(fragment.get("stop_before") or ""),
            )

        heading = fragment.get("heading")
        source_text = self._rewrite_contract_summary_internal_references(source_text, locator, heading)
        if heading:
            segments = _split_review_locator(_normalize_review_locator(locator))
            source_headings = [segments[-1] if segments else "", target.get("target_heading") or ""]
            if (
                any(
                    candidate
                    and _normalize_review_text(source_text) == _normalize_review_text(candidate)
                    for candidate in source_headings
                )
                and "\n" not in source_text
            ):
                source_text = ""
            if not source_text:
                return str(heading).strip()
            return f"{heading}\n{source_text}".strip()
        return source_text

    def _render_contract_summary_spec(self, sections: list[dict], spec: dict) -> str:
        placeholder = str(spec.get("placeholder") or "").strip()
        parts = [
            self._render_contract_summary_fragment(sections, placeholder, fragment)
            for fragment in spec.get("fragments", ())
        ]
        rendered = "\n".join(part for part in parts if part.strip()).strip()
        for fragment in spec.get("fragments", ()):
            rendered = self._rewrite_contract_summary_internal_references(
                rendered,
                str(fragment.get("locator") or ""),
                fragment.get("heading"),
            )
        rendered = self._normalize_contract_summary_numbered_heading_references(rendered)
        if not rendered:
            raise ValueError(f"基金合同摘要占位符 {placeholder} 未生成任何内容")
        return rendered

    def _replace_contract_summary_placeholders(self, text: str) -> str:
        placeholders = set(self._CONTRACT_SUMMARY_PLACEHOLDER_RE.findall(text or ""))
        if not placeholders:
            return text

        specs_by_placeholder = {
            spec["placeholder"]: spec
            for spec in self.CONTRACT_SUMMARY_SPECS
        }
        unknown_placeholders = sorted(placeholders - set(specs_by_placeholder))
        if unknown_placeholders:
            raise ValueError(f"基金合同摘要存在未配置的占位符：{', '.join(unknown_placeholders)}")

        sections = _split_contract_sections(text)
        replacements = {
            placeholder: self._render_contract_summary_spec(sections, specs_by_placeholder[placeholder])
            for placeholder in sorted(placeholders)
        }

        def replace_summary_placeholder(match):
            return replacements[match.group(1)]

        rendered = self._CONTRACT_SUMMARY_PLACEHOLDER_RE.sub(replace_summary_placeholder, text)
        remaining = sorted(set(self._CONTRACT_SUMMARY_PLACEHOLDER_RE.findall(rendered)))
        if remaining:
            raise ValueError(f"基金合同摘要占位符未完全替换：{', '.join(remaining)}")
        return rendered

    # ── 主方法：生成合同 ─────────────────────────────────────────────────────
    def generate(self, form_data: dict) -> str:
        # Step 1
        v = self._derive_variables(form_data)
        # Step 2
        v = self._inject_clause_texts(v)
        # Step 3: read template
        template_text = TEMPLATE_MD.read_text(encoding="utf-8")
        # Step 4
        text = self._process_conditionals(template_text, v)
        # Step 5
        text = self._replace_placeholders(text, v)
        # Step 5b: 修复条件删除导致的序号跳跃
        text = self._renumber_sequences(text)
        # Step 6
        text = self._cleanup(text)
        text = self._normalize_contract_text(text)
        text = self._replace_contract_summary_placeholders(text)
        text = self._normalize_contract_text(text)
        return text

    # ── Step 5b: 重排阿拉伯序号（修复条件块删除项目后的跳跃）──────────────
    def _renumber_sequences(self, text: str) -> str:
        """
        修复条件块删除项目后阿拉伯序号列表的跳跃。
        规则：
          - 只处理行首 `数字、` 格式
          - num == 1：新列表开始，重置计数器，不修改
          - num > expected_next（且 num > 1）：检测到跳跃，将该行及后续连续项重排
          - 遇到中文序号标题（一、二、…）或章节标题（第X部分）时重置计数器
        """
        lines = text.split("\n")
        RE_NUM = re.compile(r"^(\d+)(、)")
        RE_RESET = re.compile(
            r"^(?:[一二三四五六七八九十百]+、|第[一二三四五六七八九十百]+部分)"
        )
        last_num = None
        result = []
        for line in lines:
            if RE_RESET.match(line.strip()):
                last_num = None
                result.append(line)
                continue
            m = RE_NUM.match(line)
            if m:
                num = int(m.group(1))
                if num == 1:
                    last_num = 1
                elif last_num is not None and num > last_num + 1:
                    expected = last_num + 1
                    line = re.sub(r"^\d+、", f"{expected}、", line, count=1)
                    last_num = expected
                else:
                    last_num = num
            result.append(line)
        return "\n".join(result)

    # ── 中文标点规范化 ───────────────────────────────────────────────────────
    @staticmethod
    def _to_chinese_punct(text: str) -> str:
        """将文本中的半角标点转为全角中文标点（保留数字中的逗号与小数点）。"""
        text = text.replace('(', '（').replace(')', '）')
        text = text.replace('[', '【').replace(']', '】')
        text = text.replace(';', '；')
        # 逗号仅在两侧均非数字时替换，保留千位分隔符如"1,000"
        text = re.sub(r'(?<!\d),(?!\d)', '，', text)
        return text

    # ── Word (.docx) 导出 ────────────────────────────────────────────────────
    def build_docx(self, contract_text: str) -> bytes:
        """
        将合同纯文本转换为格式化 Word 文档，与参考合同格式完全对齐。
        - 封面：独立 section + vAlign=center，彻底消除内容溢出第二页的问题
        - 签署页：左对齐 + 精确空行间距（来自参考文档XML）
        - 章节标题：Times New Roman 15pt 加粗，pageBreakBefore
        - 正文：宋体 12pt，两端对齐，首行缩进2字，1.5倍行距
        - 数字/英文：Times New Roman；中文：宋体
        """
        # 标准化中文标点
        contract_text = self._to_chinese_punct(contract_text)

        def _normalize_contract_cover_text(text: str) -> str:
            lines = str(text or "").split("\n")
            toc_index = next((idx for idx, line in enumerate(lines) if re.match(r"^目\s*录$", line.strip())), None)
            cover_end = toc_index if toc_index is not None else len(lines)
            cover_lines = lines[:cover_end]
            tail_lines = lines[cover_end:]
            normalized_cover = []
            title_done = False
            for raw in cover_lines:
                stripped = raw.strip()
                if stripped and re.fullmatch(r"\d{4}年\d{1,2}月\d{1,2}日", stripped):
                    continue
                if stripped and not title_done and "基金合同" in stripped:
                    stripped = re.sub(r"基金合同(?:（草案）)?$", "基金合同（草案）", stripped)
                    title_done = True
                    normalized_cover.append(stripped)
                    continue
                normalized_cover.append(raw)
            return "\n".join([*normalized_cover, *tail_lines])

        contract_text = _normalize_contract_cover_text(contract_text)
        from pathlib import Path
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Twips

        # ── 1. 文档与页面设置 ──────────────────────────────────────────────
        packaged_contract_template_dir = BASE_DIR / "packaged_assets" / "contract_templates"
        template_candidates = []
        if packaged_contract_template_dir.is_dir():
            template_candidates = [
                p for p in packaged_contract_template_dir.iterdir()
                if p.suffix.lower() == ".docx" and "基金合同" in p.name
            ]
        template_docx = next((candidate for candidate in template_candidates if candidate.exists()), None)
        doc = Document(str(template_docx)) if template_docx else Document()

        body = doc._element.body
        cover_sectpr_template = None
        for child in body:
            if child.tag != qn("w:p"):
                continue
            p_pr = child.find(qn("w:pPr"))
            if p_pr is None:
                continue
            sect_pr = p_pr.find(qn("w:sectPr"))
            if sect_pr is not None:
                cover_sectpr_template = deepcopy(sect_pr)
                break
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
        sec = doc.sections[0]
        sec.page_width        = Twips(11906)
        sec.page_height       = Twips(16838)
        sec.top_margin        = Twips(1440)
        sec.bottom_margin     = Twips(1440)
        sec.left_margin       = Twips(1800)
        sec.right_margin      = Twips(1800)
        sec.header_distance   = Twips(851)
        sec.footer_distance   = Twips(992)
        sectPr = sec._sectPr
        for node in list(sectPr.findall(qn("w:docGrid"))):
            sectPr.remove(node)
        docGrid = OxmlElement("w:docGrid")
        docGrid.set(qn("w:type"), "lines")
        docGrid.set(qn("w:linePitch"), "312")
        sectPr.append(docGrid)
        settings = doc.settings.element
        default_tab_stop = settings.find(qn("w:defaultTabStop"))
        if default_tab_stop is None:
            default_tab_stop = OxmlElement("w:defaultTabStop")
            settings.insert(0, default_tab_stop)
        default_tab_stop.set(qn("w:val"), "720")

        # ── 2. Normal 样式：中文宋体/英数 Times New Roman 12pt，两端对齐，孤行控制关闭 ──
        normal_style = doc.styles["Normal"]
        normal_rpr = normal_style.element.get_or_add_rPr()
        for old in normal_rpr.findall(qn("w:rFonts")):
            normal_rpr.remove(old)
        for old in normal_rpr.findall(qn("w:sz")):
            normal_rpr.remove(old)
        for old in normal_rpr.findall(qn("w:szCs")):
            normal_rpr.remove(old)
        rFonts_n = OxmlElement("w:rFonts")
        rFonts_n.set(qn("w:ascii"),    "Times New Roman")
        rFonts_n.set(qn("w:hAnsi"),    "Times New Roman")
        rFonts_n.set(qn("w:eastAsia"), "宋体")
        normal_rpr.insert(0, rFonts_n)
        sz_n = OxmlElement("w:sz"); sz_n.set(qn("w:val"), "24"); normal_rpr.append(sz_n)
        normal_ppr = normal_style.element.get_or_add_pPr()
        _clear_xml_children(normal_ppr)
        wc = OxmlElement("w:widowControl"); wc.set(qn("w:val"), "0"); normal_ppr.insert(0, wc)
        jc_n = OxmlElement("w:jc"); jc_n.set(qn("w:val"), "both"); normal_ppr.append(jc_n)
        sp_n = OxmlElement("w:spacing")
        sp_n.set(qn("w:line"), "360"); sp_n.set(qn("w:lineRule"), "auto")
        sp_n.set(qn("w:before"), "0"); sp_n.set(qn("w:after"), "0")
        normal_ppr.append(sp_n)
        _ensure_word_toc_styles(
            doc,
            heading_eastasia="宋体",
            heading_ascii="宋体",
            heading_size_half_pt=28,
            heading_bold=False,
            entry_eastasia="宋体",
            entry_ascii="Times New Roman",
            entry_size_half_pt=24,
            max_level=1,
        )

        # ── 3. XML 辅助函数 ───────────────────────────────────────────────
        def _set_para(p, jc=None, line=None, line_rule="auto",
                      before=None, after=None,
                      first_line=None, first_line_chars=None, left_ind=None,
                      keep_lines=False, page_break_before=False,
                      snap_to_grid=None):
            pPr = p._p.get_or_add_pPr()
            if keep_lines:
                pPr.append(OxmlElement("w:keepLines"))
            if page_break_before:
                pPr.append(OxmlElement("w:pageBreakBefore"))
            if snap_to_grid is not None:
                sg = OxmlElement("w:snapToGrid")
                sg.set(qn("w:val"), "1" if snap_to_grid else "0")
                pPr.append(sg)
            if before is not None or after is not None or line is not None:
                sp = OxmlElement("w:spacing")
                if before is not None: sp.set(qn("w:before"), str(before))
                if after  is not None: sp.set(qn("w:after"),  str(after))
                if line   is not None:
                    sp.set(qn("w:line"),     str(line))
                    sp.set(qn("w:lineRule"), line_rule)
                pPr.append(sp)
            if first_line is not None or first_line_chars is not None or left_ind is not None:
                ind = OxmlElement("w:ind")
                if left_ind          is not None: ind.set(qn("w:left"),           str(left_ind))
                if first_line_chars  is not None: ind.set(qn("w:firstLineChars"), str(first_line_chars))
                if first_line        is not None: ind.set(qn("w:firstLine"),      str(first_line))
                pPr.append(ind)
            if jc is not None:
                jc_el = OxmlElement("w:jc"); jc_el.set(qn("w:val"), jc); pPr.append(jc_el)

        def _set_run(r, ascii_font=None, eastasia_font=None, hint=None,
                     sz=None, sz_cs=None, bold=False, bcs=False, color=None):
            rPr = r._r.get_or_add_rPr()
            if ascii_font or eastasia_font or hint:
                rF = OxmlElement("w:rFonts")
                if ascii_font:    rF.set(qn("w:ascii"), ascii_font); rF.set(qn("w:hAnsi"), ascii_font)
                if eastasia_font: rF.set(qn("w:eastAsia"), eastasia_font)
                if hint:          rF.set(qn("w:hint"), hint)
                rPr.insert(0, rF)
            if bold: rPr.append(OxmlElement("w:b"))
            if bcs:  rPr.append(OxmlElement("w:bCs"))
            if color:
                col = OxmlElement("w:color"); col.set(qn("w:val"), color); rPr.append(col)
            if sz is not None:
                s = OxmlElement("w:sz"); s.set(qn("w:val"), str(sz)); rPr.append(s)
            if sz_cs is not None:
                sc = OxmlElement("w:szCs"); sc.set(qn("w:val"), str(sz_cs)); rPr.append(sc)

        def _signing_empty(n=1):
            """签署页专用空行：左对齐，snapToGrid=0，宋体，bCs，与参考文档一致。"""
            for _ in range(n):
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                sg = OxmlElement("w:snapToGrid"); sg.set(qn("w:val"), "0"); pPr.append(sg)
                sp = OxmlElement("w:spacing")
                sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
                sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
                pPr.append(sp)
                jc_el = OxmlElement("w:jc"); jc_el.set(qn("w:val"), "left"); pPr.append(jc_el)
                mrPr = OxmlElement("w:rPr")
                rF2 = OxmlElement("w:rFonts")
                rF2.set(qn("w:ascii"), "Times New Roman"); rF2.set(qn("w:hAnsi"), "Times New Roman")
                mrPr.append(rF2)
                mrPr.append(OxmlElement("w:bCs"))
                szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "21"); mrPr.append(szCs)
                pPr.append(mrPr)

        def _update_contract_header_text(header_text: str, sections):
            if not sections:
                return
            for section in sections:
                for header in (section.header, section.first_page_header):
                    paragraph = _reset_header_footer_part(header)
                    pPr = paragraph._p.get_or_add_pPr()
                    _clear_xml_children(pPr)
                    jc = OxmlElement("w:jc")
                    jc.set(qn("w:val"), "right")
                    pPr.append(jc)
                    p_bdr = OxmlElement("w:pBdr")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:color"), "auto")
                    bottom.set(qn("w:sz"), "6")
                    bottom.set(qn("w:space"), "1")
                    p_bdr.append(bottom)
                    pPr.append(p_bdr)
                    run = paragraph.add_run(header_text)
                    _set_run(run, ascii_font="Times New Roman", eastasia_font="宋体", sz=18, sz_cs=18)

        def _set_section_doc_grid(section, line_pitch: str):
            sect_pr = section._sectPr
            for node in list(sect_pr.findall(qn("w:docGrid"))):
                sect_pr.remove(node)
            doc_grid = OxmlElement("w:docGrid")
            doc_grid.set(qn("w:type"), "lines")
            doc_grid.set(qn("w:linePitch"), str(line_pitch))
            sect_pr.append(doc_grid)

        def _cover_section_break():
            """
            在封面最后一段的 pPr 中嵌入 sectPr（封面 section 定义）。
            使用 vAlign=center 确保内容永远不溢出，并通过 nextPage 分隔目录。
            """
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            cover_sp = deepcopy(cover_sectpr_template) if cover_sectpr_template is not None else OxmlElement("w:sectPr")
            for node in list(cover_sp.findall(qn("w:type"))):
                cover_sp.remove(node)
            t = OxmlElement("w:type"); t.set(qn("w:val"), "nextPage"); cover_sp.insert(0, t)
            for node in list(cover_sp.findall(qn("w:pgSz"))):
                cover_sp.remove(node)
            pgSz = OxmlElement("w:pgSz")
            pgSz.set(qn("w:w"), "11906"); pgSz.set(qn("w:h"), "16838")
            cover_sp.append(pgSz)
            for node in list(cover_sp.findall(qn("w:pgMar"))):
                cover_sp.remove(node)
            pgMar = OxmlElement("w:pgMar")
            pgMar.set(qn("w:top"),    "1440"); pgMar.set(qn("w:right"),  "1800")
            pgMar.set(qn("w:bottom"), "1440"); pgMar.set(qn("w:left"),   "1800")
            pgMar.set(qn("w:header"), "851");  pgMar.set(qn("w:footer"), "992")
            pgMar.set(qn("w:gutter"), "0")
            cover_sp.append(pgMar)
            for node in list(cover_sp.findall(qn("w:docGrid"))):
                cover_sp.remove(node)
            dg = OxmlElement("w:docGrid")
            dg.set(qn("w:type"), "lines"); dg.set(qn("w:linePitch"), "312")
            cover_sp.append(dg)
            for node in list(cover_sp.findall(qn("w:vAlign"))):
                cover_sp.remove(node)
            vAlign = OxmlElement("w:vAlign"); vAlign.set(qn("w:val"), "center")
            cover_sp.append(vAlign)
            pPr.append(cover_sp)

        # ── 4. 正则分类 ──────────────────────────────────────────────────
        RE_TOC_ENTRY = re.compile(r"^(第[一二三四五六七八九十百]+部分\s+.+?)[\t ]+(\d+)$")
        RE_PART_HEAD = re.compile(r"^第[一二三四五六七八九十百]+部分\s+\S")
        RE_UPPER_NUM = re.compile(r"^[一二三四五六七八九十百]+、")
        RE_PAREN_CN  = re.compile(r"^（[一二三四五六七八九十百]+）")
        RE_ARAB_DOT  = re.compile(r"^\d+、")
        RE_PAREN_NUM = re.compile(r"^（\d+）")
        RE_MD_TABLE_ROW = re.compile(r"^\|.*\|\s*$")
        RE_MD_TABLE_DIVIDER = re.compile(r"^\|\s*:?-{3,}:?(?:\|\s*:?-{3,}:?)+\|\s*$")

        # ── 书签计数器与锚点辅助 ────────────────────────────────────────────────
        _bm_id = [0]  # 可变单元格，允许在内层代码中递增

        def _part_anchor(heading_s: str) -> str:
            """从"第X部分..."中提取 ASCII 书签锚名，例如 'part_26'。"""
            m = re.match(r'^第([一二三四五六七八九十百]+)部分', heading_s)
            if not m:
                return ""
            part_no = self._cn_numeral_to_int(m.group(1))
            return f"part_{part_no}" if part_no is not None else ""

        def _append_xml_run(paragraph, text=None, *, tab=False, ascii_font="Times New Roman",
                            eastasia_font="宋体", size=24, hyperlink_style=False):
            run_element = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            if hyperlink_style:
                r_style = OxmlElement("w:rStyle")
                r_style.set(qn("w:val"), "Hyperlink")
                r_pr.append(r_style)
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), ascii_font)
            r_fonts.set(qn("w:hAnsi"), ascii_font)
            r_fonts.set(qn("w:eastAsia"), eastasia_font)
            r_pr.append(r_fonts)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(size))
            r_pr.append(sz)
            sz_cs = OxmlElement("w:szCs")
            sz_cs.set(qn("w:val"), str(size))
            r_pr.append(sz_cs)
            run_element.append(r_pr)
            if tab:
                run_element.append(OxmlElement("w:tab"))
            elif text is not None:
                text_node = OxmlElement("w:t")
                text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                text_node.text = text
                run_element.append(text_node)
            paragraph._p.append(run_element)
            return run_element

        def _append_field(paragraph, instruction: str, *, placeholder="1",
                          ascii_font="Times New Roman", eastasia_font="宋体", size=24):
            begin = OxmlElement("w:r")
            _append_run_properties(begin, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            fld_begin.set(qn("w:dirty"), "true")
            begin.append(fld_begin)
            paragraph._p.append(begin)

            instr_run = OxmlElement("w:r")
            _append_run_properties(instr_run, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
            instr_text = OxmlElement("w:instrText")
            instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            instr_text.text = instruction
            instr_run.append(instr_text)
            paragraph._p.append(instr_run)

            sep = OxmlElement("w:r")
            _append_run_properties(sep, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            sep.append(fld_sep)
            paragraph._p.append(sep)

            _append_xml_run(
                paragraph,
                placeholder,
                ascii_font=ascii_font,
                eastasia_font=eastasia_font,
                size=size,
            )

            end = OxmlElement("w:r")
            _append_run_properties(end, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            end.append(fld_end)
            paragraph._p.append(end)

        # ── 5. 主处理循环 ─────────────────────────────────────────────────
        lines = contract_text.split("\n")
        phase          = "cover"
        cover_idx      = 0      # 非空封面行计数
        signing_started = False
        signing_idx    = 0      # 签署页内容行计数
        current_part   = 0      # 当前所在部分编号（用于判断前言是否加空行）

        for raw in lines:
            s = raw.strip()

            # 所有空行跳过（签署页内的空行由代码主动添加）
            if not s:
                continue

            # ════ 封面 → 目录 ════
            if phase == "cover" and re.match(r"^目\s*录$", s):
                phase = "toc"
                # 封面 section break（vAlign=center, nextPage）
                _cover_section_break()
                p = doc.add_paragraph(style="TOC Heading")
                _set_para(p, jc="center", line=360, before=0, after=0)
                r = p.add_run("目    录")
                _set_run(r, sz=28, sz_cs=28, bcs=True)
                toc_field = doc.add_paragraph()
                _set_para(toc_field, line=360, before=0, after=0)
                _append_word_toc_field(
                    toc_field,
                    OxmlElement,
                    qn,
                    levels="1-1",
                    placeholder="右键更新目录",
                    ascii_font="Times New Roman",
                    eastasia_font="宋体",
                    size=24,
                )
                continue

            # ════ 封面内容：仅4行，vAlign=center 负责垂直定位 ════
            if phase == "cover":
                p = doc.add_paragraph()
                if cover_idx == 0:
                    # 合同标题：sz=48(24pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=48, sz_cs=48, bold=True, bcs=True)
                    # 标题与管理人之间加6个空行，和参考合同封面版式保持一致
                    for _ in range(6):
                        ep = doc.add_paragraph()
                        _set_para(ep, jc="center", line=360, before=0, after=0)
                elif s.startswith("基金管理人") or s.startswith("基金托管人"):
                    # 管理人/托管人：sz=36(18pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=36, sz_cs=36, bold=True, bcs=True)
                else:
                    # 日期：sz=36(18pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=36, sz_cs=36, bold=True, bcs=True)
                cover_idx += 1
                continue

            # ════ 目录条目 ════
            if phase == "toc":
                m_toc = RE_TOC_ENTRY.match(s)
                if m_toc:
                    continue
                else:
                    phase = "body"

            # ════ 签署页检测 ════
            if phase == "body" and not signing_started and "签署页" in s and "无正文" in s:
                signing_started = True
                p = doc.add_paragraph()
                _set_para(p, snap_to_grid=False, jc="left", line=360, before=0, after=0,
                          page_break_before=True)   # 签署页强制分页
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                _signing_empty(2)   # 签署页标题后 2 个空行（参考文档）
                signing_idx = 1
                continue

            # ════ 签署页内容行 ════
            if phase == "body" and signing_started:
                p = doc.add_paragraph()
                is_seal = "（盖章）" in s
                _set_para(p, snap_to_grid=False, jc="left", line=360, before=0, after=0)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                # 每一内容行后补充空行（对应参考文档签名留白）
                if "（盖章）" in s:
                    _signing_empty(4)   # 盖章行后 4 个空行（管理人/托管人签名区）
                elif "（签字或盖章）：" in s or "（签名）" in s:
                    _signing_empty(4)   # 法定代表人签名行后 4 个空行
                # 签订地点/日期行后不加空行
                signing_idx += 1
                continue

            # ════ 部分标题（第X部分…） ════
            if phase == "body" and RE_PART_HEAD.match(s):
                current_part += 1
                p = doc.add_paragraph(style="Heading 1")
                _set_para(p, jc="center", line=360, before=0, after=0,
                          keep_lines=True, page_break_before=True)
                # 添加书签，供目录超链接跳转
                anchor = _part_anchor(s)
                if anchor:
                    _bm_id[0] += 1
                    bm_start = OxmlElement("w:bookmarkStart")
                    bm_start.set(qn("w:id"), str(_bm_id[0]))
                    bm_start.set(qn("w:name"), anchor)
                    bm_end = OxmlElement("w:bookmarkEnd")
                    bm_end.set(qn("w:id"), str(_bm_id[0]))
                    p._p.append(bm_start)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         sz=30, sz_cs=30, bold=True, bcs=False, color="auto")
                if anchor:
                    p._p.append(bm_end)
                continue

            # ════ 一级子标题（一、二、…） ════
            if phase == "body" and RE_UPPER_NUM.match(s):
                # 第一部分（前言）子标题间不加空行，其余部分各子标题前空一行
                if current_part != 1:
                    ep = doc.add_paragraph()
                    _set_para(ep, line=360, before=0, after=0)
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 二级子标题（（一）（二）…） ════
            if phase == "body" and RE_PAREN_CN.match(s):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 数字条款（1、2、… 或（1）（2）…） ════
            if phase == "body" and (RE_ARAB_DOT.match(s) or RE_PAREN_NUM.match(s)):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 普通正文 ════
            if phase in ("body", "toc"):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)

        # ── 6. 序列化 ──────────────────────────────────────────────────────
        body_sections = list(doc.sections[1:]) if len(doc.sections) > 1 else list(doc.sections)
        for section in body_sections:
            _set_section_doc_grid(section, "312")
        _update_contract_header_text("基金合同（草案）", body_sections)
        _finalize_doc_page_numbers(doc, OxmlElement, qn)
        _set_update_fields_on_open(doc, OxmlElement, qn)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

# ═══════════════════════════════════════════════════════════════════════════════
#  ProspectusEngine — 招募说明书 8步处理管线
# ═══════════════════════════════════════════════════════════════════════════════
class ProspectusEngine:
    def __init__(
        self,
        business_text_overrides: dict | None = None,
        contract_engine: ContractEngine | None = None,
    ):
        with open(PROSPECTUS_CLAUSES_JSON, encoding="utf-8") as f:
            self.pro_clauses = json.load(f)["clauses"]
        self.business_text_overrides = (
            _normalize_business_text_overrides(business_text_overrides)
            if business_text_overrides is not None
            else _load_business_text_overrides()
        )
        self.contract_engine = contract_engine or ContractEngine(business_text_overrides=business_text_overrides)
        self._reference_fixed_cache = {}
        self._reference_fixed_txt_cache = {}
        self._reference_section7_table_cache = {}

    _is_prospectus_toc_placeholder_line = staticmethod(ContractEngine._is_prospectus_toc_placeholder_line)
    _parse_prospectus_chapter_heading = staticmethod(ContractEngine._parse_prospectus_chapter_heading)
    _format_prospectus_reference_heading = staticmethod(ContractEngine._format_prospectus_reference_heading)
    _parse_prospectus_toc_entry = staticmethod(ContractEngine._parse_prospectus_toc_entry)
    _extract_prospectus_toc_entries = ContractEngine._extract_prospectus_toc_entries
    _format_reference_style_prospectus_with_plain_titles = ContractEngine._format_reference_style_prospectus_with_plain_titles
    _format_reference_style_prospectus = ContractEngine._format_reference_style_prospectus

    # ── Step 1: 派生变量（委托给 ContractEngine）────────────────────────────
    def _derive_variables(self, v: dict) -> dict:
        v = self.contract_engine._derive_variables(v)
        market_type = str(v.get("MARKET_TYPE", "") or "").strip().upper()
        market_scope = str(v.get("MARKET_SCOPE", "") or "").strip().upper()
        if market_type in {"KECHUANG", "CHUANGYE"}:
            market_scope = "SINGLE_MARKET"
        elif market_scope not in {"SINGLE_MARKET", "CROSS_MARKET"}:
            market_scope = "CROSS_MARKET"
        v["MARKET_SCOPE"] = market_scope
        product_type = str(v.get("PRODUCT_TYPE", "") or "").strip().upper()
        if not product_type:
            product_type = "ETF"
        v["PRODUCT_TYPE"] = product_type
        return v

    def _prospectus_business_text_replacements(self, v: dict) -> dict[str, str]:
        return {
            "{CSRC_APPROVAL_NO}": str(v.get("CSRC_APPROVAL_NO") or "").strip(),
            "{INDEX_NAME}": str(v.get("INDEX_NAME") or "").strip(),
            "{INDEX_DESCRIPTION}": str(v.get("INDEX_DESCRIPTION") or "").strip(),
            "{INDEX_COMPILER}": self._normalize_index_compiler_name(str(v.get("INDEX_COMPILER") or "").strip()),
            "{INDEX_WEBSITE}": str(v.get("INDEX_WEBSITE") or "").strip(),
            "{MIN_SUB_UNIT}": self._get_prospectus_min_sub_unit(v),
        }

    def _resolve_prospectus_business_text(self, key: str, v: dict, default_text: str = "") -> str:
        spec = PROSPECTUS_BUSINESS_TEXT_SPECS.get(key, {})
        scene = _prospectus_business_text_variant(v, spec.get("variant_mode", "DEFAULT"))
        text = _get_business_text_override(
            self.business_text_overrides,
            "prospectus",
            key,
            scene,
            _business_text_product_type(v),
            _business_text_market_type(v),
            _business_text_exchange(v),
        )
        if text is None:
            text = default_text
        if spec.get("render_placeholders"):
            text = self._replace_prospectus_placeholders(text, self._prospectus_business_text_replacements(v))
        return str(text or "").strip()

    def _apply_business_text_overrides(self, v: dict) -> dict:
        result = dict(v)
        for key, spec in PROSPECTUS_BUSINESS_TEXT_SPECS.items():
            if key not in result:
                continue
            result[key] = self._resolve_prospectus_business_text(key, result, str(result.get(key, "") or ""))
        return result

    @staticmethod
    def _default_chapter_six_product_type_sentence() -> str:
        return "本基金为交易型开放式基金，股票型基金，基金存续期限为不定期。"

    @staticmethod
    def _default_chapter_six_sec8_body() -> str:
        return """八、认购费用
认购费用由投资人承担，不高于0.30%，认购费率如下表所示：

|   |   |
|---|---|
|认购份额（S）|认购费率|
|S＜100万份|0.30%|
|S≥100万份|每笔500元|

基金管理人办理网下现金认购和网下股票认购不收取认购费。发售代理机构办理网上现金认购、网下现金认购、网下股票认购时可参照上述费率结构，按照不高于0.3%的标准收取一定的佣金。投资人申请重复现金认购的，须按每次认购所对应的费率档次分别计费。"""

    def _default_chapter_ten_sec5_body(self, v: dict) -> str:
        min_sub_unit = self._get_prospectus_min_sub_unit(v)
        canonical_item_one = (
            f"1、投资人申购、赎回的基金份额需为最小申购赎回单位的整数倍。"
            f"目前，本基金最小申购赎回单位为{min_sub_unit}，基金管理人有权对其进行调整，"
            "并在调整实施前依照《信息披露办法》的有关规定在规定媒介上公告。"
        )
        return (
            "五、申购和赎回的数额限制\n"
            f"{canonical_item_one}\n"
            "2、基金管理人可以规定本基金当日申购份额及当日赎回份额上限，具体规定请参见申购赎回清单或相关公告。\n"
            "3、基金管理人可以根据市场情况，在法律法规允许的情况下，合理调整上述申购和赎回的数量或比例限制，并在实施前依照《信息披露办法》的有关规定在规定媒介上公告。\n"
            "4、当接受申购申请对存量基金份额持有人利益构成潜在重大不利影响时，基金管理人应当采取设定单一投资者申购份额上限或基金单日净申购比例上限、拒绝大额申购、暂停基金申购等措施。"
        )

    def _get_prospectus_variant_key(self, v: dict) -> str:
        v = self._derive_variables(v)
        return _prospectus_variant_key(v)

    def _get_variant_clause_bundle(self, v: dict) -> dict:
        variants = self.pro_clauses.get("PROSPECTUS_VARIANTS", {}).get("variants", {})
        key = self._get_prospectus_variant_key(v)
        return variants.get(key, variants.get("SSE_CROSS", {}))

    @staticmethod
    def _get_market_clause_key(v: dict) -> str:
        exchange = str((v or {}).get("EXCHANGE", "") or "").strip().upper()
        has_hk = bool((v or {}).get("HAS_HK_CONNECT"))
        if exchange == "SSE":
            return "MARKET_SH_HK_CONNECT" if has_hk else "MARKET_SH_STANDARD"
        return "MARKET_SZ_HK_CONNECT" if has_hk else "MARKET_SZ_STANDARD"

    @staticmethod
    def _get_market_exchange_key(v: dict) -> str:
        exchange = str((v or {}).get("EXCHANGE", "") or "").strip().upper()
        return "MARKET_SH" if exchange == "SSE" else "MARKET_SZ"

    def _resolve_prospectus_clause_text(self, clause_key: str, variant_key: str) -> str:
        clause = self.pro_clauses.get(clause_key, {})
        variant = clause.get("variants", {}).get(variant_key, {})
        text = str(variant.get("text") or "").strip()
        if text:
            return self._finalize_prospectus_clause_text(clause_key, variant_key, text)

        legacy = variant.get("legacy_source") or {}
        legacy_clause_key = legacy.get("clause")
        legacy_variant_key = legacy.get("variant")
        if not legacy_clause_key or not legacy_variant_key:
            return ""

        legacy_clause = self.pro_clauses.get(legacy_clause_key, {})
        legacy_variant = legacy_clause.get("variants", {}).get(legacy_variant_key, {})
        if not legacy_variant:
            legacy_variant = legacy_clause.get(legacy_variant_key, {})
        return self._finalize_prospectus_clause_text(
            clause_key,
            variant_key,
            str(legacy_variant.get("text") or "").strip(),
        )

    def _finalize_prospectus_clause_text(self, clause_key: str, variant_key: str, text: str) -> str:
        body = str(text or "").strip()
        if not body:
            return ""
        if clause_key == "PURCHASE_REDEMPTION_LIST_CLAUSE" and variant_key == "MARKET_SH_STANDARD":
            return self._ensure_market_sh_standard_purchase_redemption_list_clause(body)
        return body

    def _ensure_market_sh_standard_purchase_redemption_list_clause(self, text: str) -> str:
        body = str(text or "").strip()
        if not body or "1）禁止现金替代" not in body:
            return body

        formula_label = self._cash_substitution_ratio_formula_label()
        formula_text = self._cash_substitution_ratio_formula_text()
        formula_block = f"{formula_label}\n{formula_text}"

        if formula_label in body and formula_text in body:
            return body
        if formula_label in body:
            return body.replace(formula_label, formula_block, 1)
        return body.replace("1）禁止现金替代", f"{formula_block}\n1）禁止现金替代", 1)

    @staticmethod
    def _replace_top_level_section(body_text: str, section_cn: str, new_section: str) -> str:
        body = (body_text or "").strip()
        section = (new_section or "").strip()
        if not body or not section:
            return body

        sec_re = re.compile(rf"^{section_cn}、[^\n]*", re.MULTILINE)
        next_sec_re = re.compile(r"^[一二三四五六七八九十百]+、[^\n]*", re.MULTILINE)
        match = sec_re.search(body)
        if not match:
            return body
        next_match = next_sec_re.search(body, match.end())
        sec_end = next_match.start() if next_match else len(body)
        return f"{body[:match.start()]}{section}\n{body[sec_end:].lstrip()}".strip()


    @staticmethod
    def _get_product_type(v: dict) -> str:
        product_type = str((v or {}).get("PRODUCT_TYPE", "") or "").strip().upper()
        return product_type or "ETF"

    @staticmethod
    def _strip_signing_page_from_contract_summary(summary_text: str) -> str:
        return _strip_contract_signing_page_text(summary_text)

    @staticmethod
    def _basket_value_from_min_sub_unit(v: dict) -> str:
        min_sub_unit = ProspectusEngine._get_prospectus_min_sub_unit(v)
        if min_sub_unit.endswith("万份"):
            return min_sub_unit[:-2]
        return min_sub_unit.replace("份", "")

    def _build_index_info_source_clause(self, v: dict) -> str:
        index_compiler = self._normalize_index_compiler_name(str(v.get("INDEX_COMPILER") or "").strip())
        index_website = str(v.get("INDEX_WEBSITE") or "").strip()
        if not index_compiler or not index_website:
            return ""
        return f"有关标的指数具体编制方案及成份股信息详见{index_compiler}网站，网址：{index_website}。"

    def _build_subscription_account_clause(self, v: dict) -> str:
        variant_key = self._get_prospectus_variant_key(v)

        if variant_key == "SSE_CROSS":
            return (
                "投资人投资本基金时需具有上海证券交易所A股账户或基金账户。"
                "其中，上海证券交易所基金账户只能进行基金的现金认购和二级市场交易，"
                "如投资人需要使用标的指数成份股中的上海证券交易所上市股票参与网下股票认购或基金的申购、赎回，"
                "则应开立上海证券交易所A股账户；如投资人需要使用标的指数成份股中的深圳证券交易所上市股票参与网下股票认购，"
                "则还应开立深圳证券交易所A股账户。"
            )

        if variant_key == "SSE_HK":
            return (
                "投资人投资本基金时需具有上海证券交易所A股账户或基金账户。"
                "其中，上海证券交易所基金账户只能进行基金的现金认购和二级市场交易，"
                "如投资人需要参与基金的申购、赎回，则应开立并使用上海证券交易所A股账户。"
            )

        if variant_key == "SSE_SINGLE":
            return (
                "投资人投资本基金时需具有上海证券交易所A股账户或基金账户。"
                "其中，上海证券交易所基金账户只能进行基金的现金认购和二级市场交易，"
                "如投资人需要参与网下股票认购或基金的申购、赎回，则应开立上海证券交易所A股账户。"
            )

        if variant_key == "SZSE_CROSS":
            return (
                "投资人投资本基金时需具有深圳证券交易所A股账户或基金账户。"
                "其中，深圳证券交易所基金账户只能进行基金的现金认购和二级市场交易，"
                "如投资人需要使用标的指数成份股中的深圳证券交易所上市股票参与基金的申购、赎回，则应开立深圳证券交易所A股账户。"
            )

        return (
            "投资人投资本基金时需具有深圳证券交易所A股账户或基金账户。"
            "其中，深圳证券交易所基金账户只能进行基金的现金认购和二级市场交易，"
            "如投资人需要参与基金的申购、赎回，则应开立深圳证券交易所A股账户。"
        )

    @staticmethod
    def _replace_prospectus_placeholders(text: str, replacements: dict[str, str]) -> str:
        body = str(text or "")
        if not body:
            return ""
        for placeholder, value in replacements.items():
            body = body.replace(placeholder, value or "")
        return body

    def _replace_manual_reference_text(self, text: str, v: dict) -> str:
        body = self._replace_reference_fund_name(text, v)
        return self._replace_prospectus_placeholders(
            body,
            {
                "{SUBSCRIPTION_ACCOUNT_CLAUSE}": str(v.get("SUBSCRIPTION_ACCOUNT_CLAUSE") or "").strip(),
                "{SUB_ACCOUNT_OPENING_CLAUSE}": str(v.get("SUB_ACCOUNT_OPENING_CLAUSE") or "").strip(),
                "{LISTING_IOPV_CLAUSE}": str(v.get("LISTING_IOPV_CLAUSE") or "").strip(),
                "{PURCHASE_REDEMPTION_PROCEDURE_CLAUSE}": str(v.get("PURCHASE_REDEMPTION_PROCEDURE_CLAUSE") or "").strip(),
            },
        )

    # ── Step 2: 注入差异条款（合同条款 + 招募说明书专有条款）──────────────
    def _inject_clause_texts(self, v: dict) -> dict:
        v = self._derive_variables(v)
        v = self.contract_engine._inject_clause_texts(v)
        has_hk = v.get("HAS_HK_CONNECT", False)
        exch_cn = v.get("EXCHANGE_NAME_CN", "证券交易所")
        variant_key = self._get_prospectus_variant_key(v)
        variant_bundle = self._get_variant_clause_bundle(v)
        market_clause_key = self._get_market_clause_key(v)
        market_exchange_key = self._get_market_exchange_key(v)

        v["MARKET_SH"] = market_exchange_key == "MARKET_SH"
        v["MARKET_SZ"] = market_exchange_key == "MARKET_SZ"

        vt_key = "HK_CONNECT" if has_hk else "STANDARD"
        vt_variants = self.pro_clauses["VALUATION_TIMING"]["variants"]
        valuation_text = vt_variants.get(vt_key, vt_variants["STANDARD"])["text"]
        valuation_text = valuation_text.replace("{EXCHANGE_NAME_CN}", exch_cn)
        v["VALUATION_TIMING_CLAUSE"] = valuation_text

        v["LISTING_IOPV_CLAUSE"] = self._resolve_prospectus_clause_text(
            "LISTING_IOPV_CLAUSE",
            market_clause_key,
        )
        v["PURCHASE_REDEMPTION_PROCEDURE_CLAUSE"] = self._resolve_prospectus_clause_text(
            "PURCHASE_REDEMPTION_PROCEDURE_CLAUSE",
            market_clause_key,
        )
        v["PURCHASE_REDEMPTION_LIST_CLAUSE"] = self._resolve_prospectus_clause_text(
            "PURCHASE_REDEMPTION_LIST_CLAUSE",
            market_clause_key,
        )
        v["SUB_ACCOUNT_OPENING_CLAUSE"] = self._resolve_prospectus_clause_text(
            "SUB_ACCOUNT_OPENING_CLAUSE",
            market_exchange_key,
        )
        v["SUBSCRIPTION_ACCOUNT_CLAUSE"] = self._build_subscription_account_clause(v)
        v["ONLINE_CASH_SUBSCRIPTION_CLAUSE"] = self._resolve_prospectus_clause_text(
            "ONLINE_CASH_SUBSCRIPTION_CLAUSE",
            market_exchange_key,
        )
        v["OFFLINE_CASH_SUBSCRIPTION_CLAUSE"] = self._resolve_prospectus_clause_text(
            "OFFLINE_CASH_SUBSCRIPTION_CLAUSE",
            market_exchange_key,
        )
        v["OFFLINE_STOCK_SUBSCRIPTION_CLAUSE"] = (
            self._resolve_prospectus_clause_text("OFFLINE_STOCK_SUBSCRIPTION_CLAUSE", "NON_HK_CONNECT")
            if not has_hk
            else ""
        )
        v["RISK_DISCLOSURE_CHUANGYE"] = self.pro_clauses["RISK_DISCLOSURE_CHUANGYE"]["variants"]["DEFAULT"]["text"]
        v["RISK_DISCLOSURE_KECHUANG"] = self.pro_clauses["RISK_DISCLOSURE_KECHUANG"]["variants"]["DEFAULT"]["text"]
        v["RISK_DISCLOSURE_HK_CONNECT"] = self.pro_clauses["RISK_DISCLOSURE_HK_CONNECT"]["variants"]["DEFAULT"]["text"]

        custodian_name = v.get("CUSTODIAN_NAME", "")
        custodian_contacts = self.pro_clauses.get("CUSTODIAN_INFO_PROSPECTUS", {}).get("custodians", {})
        info = custodian_contacts.get(custodian_name, {})
        v.setdefault("CUSTODIAN_DEPT", info.get("dept", "[待填写：托管部门名称]"))
        v.setdefault("CUSTODIAN_PHONE", info.get("phone", "[待填写：服务电话]"))
        v.setdefault("CUSTODIAN_WEBSITE", info.get("website", "[待填写：网址]"))

        v.setdefault("FUND_MANAGER_NAME", "[待填写：基金经理姓名]")
        v.setdefault("FUND_MANAGER_BIO", "[待填写：基金经理简介（学历、从业经历、任职日期等）]")
        v.setdefault("FUND_MANAGER_SEX", "")
        v.setdefault("CSRC_APPROVAL_NO", "202X年X月X日证监许可〔202X〕XXX号")
        v.setdefault("INDEX_DESCRIPTION", "")
        v.setdefault("INDEX_WEBSITE", "")
        v.setdefault("MIN_SUB_UNIT", "100万份")
        v.setdefault("CUSTODIAN_INTRO", "【托管人情况待填写】")
        v.setdefault("ACCOUNTANT", "【待填写】")
        v["INDEX_INFO_SOURCE_CLAUSE"] = self._build_index_info_source_clause(v)
        v["IMPORTANT_NOTICE_APPROVAL_SENTENCE"] = f"本基金经中国证监会{v.get('CSRC_APPROVAL_NO', '202X年X月X日证监许可〔202X〕XXX号')}文注册募集。"
        v["IMPORTANT_NOTICE_INDEX_SOURCE_SENTENCE"] = v["INDEX_INFO_SOURCE_CLAUSE"]
        v["CHAPTER6_INTRO_REGISTRATION_SENTENCE"] = f"本基金由基金管理人依照《基金法》、《运作办法》、《销售办法》、基金合同及其他有关规定，并经中国证监会{v.get('CSRC_APPROVAL_NO', '202X年X月X日证监许可〔202X〕XXX号')}文注册募集。"
        v["CHAPTER6_PRODUCT_TYPE_SENTENCE"] = self._default_chapter_six_product_type_sentence()
        v["CHAPTER6_SEC8_DEFAULT_BODY"] = self._default_chapter_six_sec8_body()
        v["CHAPTER10_SEC5_DEFAULT_BODY"] = self._default_chapter_ten_sec5_body(v)
        v["CHAPTER21_PLACEHOLDER_TEXT"] = "【待填写】"
        v["BASKET"] = self._basket_value_from_min_sub_unit(v)

        chapter6 = variant_bundle.get("chapter_6", {})
        chapter10 = variant_bundle.get("chapter_10", {})
        v["PROSPECTUS_VARIANT_KEY"] = variant_key
        v["METHOD_SUBSCRIBE_DERIVED_FROM_CONTRACT"] = str(v.get("SUBSCRIPTION_METHOD_TEXT") or chapter6.get("section_4", "")).strip()
        v["PROSPECTUS_CH6_SEC4"] = v["METHOD_SUBSCRIBE_DERIVED_FROM_CONTRACT"]
        v["PROSPECTUS_CH6_SEC7"] = v["SUB_ACCOUNT_OPENING_CLAUSE"]
        v["PROSPECTUS_CH6_SEC11"] = (
            v["OFFLINE_STOCK_SUBSCRIPTION_CLAUSE"] or (chapter6.get("section_11", "") if not has_hk else "")
        )
        v["PROSPECTUS_CH6_SEC12"] = chapter6.get("section_12", "")
        v["PROSPECTUS_CH6_SEC13"] = chapter6.get("section_13", "")
        v["PROSPECTUS_CH7_BODY"] = variant_bundle.get("chapter_7", "")
        v["PROSPECTUS_CH9_BODY"] = self._replace_top_level_section(
            variant_bundle.get("chapter_9", ""),
            "四",
            v["LISTING_IOPV_CLAUSE"],
        )
        v["PROSPECTUS_CH10_PRELUDE"] = chapter10.get("prelude", "")
        v["PROSPECTUS_CH10_SEC4"] = v["PURCHASE_REDEMPTION_PROCEDURE_CLAUSE"] or chapter10.get("section_4", "")
        v["PROSPECTUS_CH10_SEC7"] = v["PURCHASE_REDEMPTION_LIST_CLAUSE"] or chapter10.get("section_7", "")
        v["PROSPECTUS_CH10_SEC7_BODY"] = v["PROSPECTUS_CH10_SEC7"]
        v["PROSPECTUS_CH10_SEC7_FORMULA_LABEL"] = ""
        v["PROSPECTUS_CH10_SEC7_FORMULA_TEXT"] = ""
        v["PROSPECTUS_CH10_SEC7_FORMAT_BLOCK"] = ""
        ref = self._load_reference_fixed_content(v)
        reference_sec7 = self._reference_chapter_ten_sec7(ref, v)
        if reference_sec7:
            v["PURCHASE_REDEMPTION_LIST_CLAUSE"] = reference_sec7
            v["PROSPECTUS_CH10_SEC7"] = reference_sec7
            v["PROSPECTUS_CH10_SEC7_BODY"] = reference_sec7
        v["PROSPECTUS_CH21_TITLES"] = self.pro_clauses.get("CHAPTER21_TITLES", {}).get("text", "")

        risk_bodies = self.pro_clauses.get("RISK_CHAPTER_BODIES", {}).get("variants", {})
        if has_hk:
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("HK_CONNECT", "")
        elif v.get("IS_KECHUANG"):
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("KECHUANG", "")
        elif v.get("IS_CHUANGYE"):
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("CHUANGYE", "")
        else:
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("STANDARD_A", "")

        return self._apply_business_text_overrides(v)

    # ── Step 3: 从合同全文提取各关键部分（内容摘要 + 各章节来源段落）───────
    def _extract_contract_sections(self, v: dict) -> dict:
        """
        生成完整合同文本，按以下逻辑提取各部分供招募说明书使用：

        #二章 释义 → 《基金合同》第二部分
        #十章 申购赎回 子条款 → 《基金合同》第八部分 对应子条款：
          二、开放日 → CONTRACT_PART8_SEC2
          三、原则   → CONTRACT_PART8_SEC3
          六、对价   → CONTRACT_PART8_SEC6
          七、拒绝申购 → CONTRACT_PART8_SEC7
          八、暂停赎回 → CONTRACT_PART8_SEC8
          九、其他方式 → CONTRACT_PART8_SEC9
          十、非交易过户 → CONTRACT_PART8_SEC10
          十一、冻结 → CONTRACT_PART8_SEC11
          十二、转让 → CONTRACT_PART8_SEC12
          十三、其他业务 → CONTRACT_PART8_SEC13
          十四、清算交收模式 → CONTRACT_PART8_SEC14
        #十一章 投资 → 《基金合同》第十四部分
        #十二章 财产 → 《基金合同》第十五部分
        #十三章 估值 → 《基金合同》第十六部分
        #十四章 收益分配 → 三、基金收益分配原则 同 《基金合同》三
        #十五章 费用 → 《基金合同》第十七部分
        #十六章 会计审计 → 《基金合同》第十九部分
        #十七章 信息披露 → 《基金合同》第二十部分
        #十九章 变更终止 → 《基金合同》第二十一部分
        #二十章 合同摘要 → 《基金合同》第二十六部分
        """
        try:
            contract_text = engine.generate(v)
        except Exception as _exc:
            import logging
            logging.warning("ProspectusEngine: 合同生成失败，部分章节将使用占位符。原因：%s", _exc)
            contract_text = ""

        # 按"第X部分"标题切分合同全文
        RE_PART = re.compile(r'^第[一二三四五六七八九十百]+部分\s*\S[^\n]*', re.MULTILINE)
        part_iter = list(RE_PART.finditer(contract_text))

        # 中文数字顺序表，用于精确匹配部分编号（避免"第二部分"匹配"第二十部分"）
        _CN_ORDER = [
            "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
            "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九",
            "二十", "二十一", "二十二", "二十三", "二十四", "二十五", "二十六",
        ]

        def _get_part(cn_num: str) -> str:
            """
            按准确中文部分编号提取该部分全文（标题行之后的内容）。
            cn_num 示例：'第二部分'、'第八部分'、'第二十六部分'
            使用精确匹配避免"第二部分"匹配到"第二十部分"。
            """
            # 构造精确匹配模式：编号后紧跟"部分"且后面为空白
            pattern = re.compile(
                r'^' + re.escape(cn_num) + r'\s+\S[^\n]*',
                re.MULTILINE
            )
            matched_idx = []
            for i, m in enumerate(part_iter):
                if pattern.match(m.group()):
                    matched_idx.append(i)
            if matched_idx:
                # Contract text includes TOC + body with duplicate "第X部分" headings.
                # Use the last match to target body section, not TOC line.
                i = matched_idx[-1]
                m = part_iter[i]
                start = m.end()
                end = part_iter[i + 1].start() if i + 1 < len(part_iter) else len(contract_text)
                return contract_text[start:end].strip()
            return ""

        def _get_part_subsection(part_cn: str, cn_num: str) -> str:
            """
            从指定部分提取指定中文序号的子条款全文。
            cn_num 示例：'九', '十', '十一', ...
            使用精确匹配：标题为 "序号、" 格式，用 word boundary 避免"十"匹配"十一"。
            返回该子条款标题行（含序号）及其正文，直到下一个同级序号或部分结束。
            """
            part_text = _get_part(part_cn)
            if not part_text:
                return ""
            # 精确匹配行首的中文序号：序号后紧跟"、"，序号本身是完整单词
            # 构造列表后按行首定位，避免多字符序号前缀混淆
            RE_CN = re.compile(
                r'^([一二三四五六七八九十百]+)、',
                re.MULTILINE
            )
            markers = list(RE_CN.finditer(part_text))
            for i, m in enumerate(markers):
                heading = m.group(1)          # 捕获组1 = 纯序号，不含"、"
                if heading == cn_num:
                    start = m.start()
                    end = markers[i + 1].start() if i + 1 < len(markers) else len(part_text)
                    return part_text[start:end].strip()
            return ""

        def _get_part8_subsection(cn_num: str) -> str:
            return _get_part_subsection("第八部分", cn_num)

        # #二章 释义（完整第二部分，不含标题行）
        v.setdefault("CONTRACT_DEFS_TEXT", _get_part("第二部分") or
                     "[待填写：释义内容，请从基金合同第二部分复制]")

        # #六章 募集 / 七章 生效 / 八章 折算 / 九章 上市
        v.setdefault("CONTRACT_PART4_SEC1", _get_part_subsection("第四部分", "一") or
                     "[待填写：基金份额发售时间、发售方式、发售对象，来自基金合同第四部分第一项]")
        v.setdefault("CONTRACT_PART4_SEC2", _get_part_subsection("第四部分", "二") or
                     "[待填写：基金份额的认购，来自基金合同第四部分第二项]")
        v.setdefault("CONTRACT_PART4_SEC4", _get_part_subsection("第四部分", "四") or
                     "[待填写：基金认购的其他具体规定，来自基金合同第四部分第四项]")
        v.setdefault("CONTRACT_EFFECTIVENESS_TEXT", _get_part("第五部分") or
                     "[待填写：基金合同的生效，来自基金合同第五部分]")
        v.setdefault("CONTRACT_CONVERSION_TEXT", _get_part("第六部分") or
                     "[待填写：基金份额折算与变更登记，来自基金合同第六部分]")
        v.setdefault("CONTRACT_PART7_SEC1", _get_part_subsection("第七部分", "一") or
                     "[待填写：基金份额的上市，来自基金合同第七部分第一项]")
        v.setdefault("CONTRACT_PART7_SEC2", _get_part_subsection("第七部分", "二") or
                     "[待填写：基金份额的上市交易，来自基金合同第七部分第二项]")
        v.setdefault("CONTRACT_PART7_SEC3", _get_part_subsection("第七部分", "三") or
                     "[待填写：终止上市交易，来自基金合同第七部分第三项]")
        v.setdefault("CONTRACT_PART7_SEC5", _get_part_subsection("第七部分", "五") or
                     "[待填写：其他证券交易所上市安排，来自基金合同第七部分第五项]")
        v.setdefault("CONTRACT_PART7_SEC6", _get_part_subsection("第七部分", "六") or
                     "[待填写：上市规则调整安排，来自基金合同第七部分第六项]")
        v.setdefault("CONTRACT_PART7_SEC7", _get_part_subsection("第七部分", "七") or
                     "[待填写：上市交易新增功能安排，来自基金合同第七部分第七项]")

        # #十章 申购赎回 各子条款（来自第八部分）
        v.setdefault("CONTRACT_PART8_SEC1", _get_part8_subsection("一") or
                     "[待填写：申购和赎回场所，来自基金合同第八部分]")
        # 二/三/六 保持固定（模板中已有固定文本），此处仍提取供模板引用
        v.setdefault("CONTRACT_PART8_SEC2", _get_part8_subsection("二") or
                     "[待填写：申购和赎回的开放日及时间，来自基金合同第八部分]")
        v.setdefault("CONTRACT_PART8_SEC3", _get_part8_subsection("三") or
                     "[待填写：申购和赎回的原则，来自基金合同第八部分]")
        v.setdefault("CONTRACT_PART8_SEC6", _get_part8_subsection("六") or
                     "[待填写：申购和赎回的对价、费用及其用途，来自基金合同第八部分]")
        # 七 = 拒绝暂停申购（合同 七，招募 八）
        v.setdefault("CONTRACT_PART8_SEC7", _get_part8_subsection("七") or
                     "[待填写：拒绝或暂停申购的情形，来自基金合同第八部分]")
        # 八 = 暂停赎回（合同 八，招募 九）
        v.setdefault("CONTRACT_PART8_SEC8", _get_part8_subsection("八") or
                     "[待填写：暂停赎回或延缓支付赎回对价的情形，来自基金合同第八部分]")
        # 招募 十 → 合同八.九
        v.setdefault("CONTRACT_PART8_SEC9", _get_part8_subsection("九") or
                     "[待填写：其他申购赎回方式，来自基金合同第八部分]")
        # 招募 十一 → 合同八.十
        v.setdefault("CONTRACT_PART8_SEC10", _get_part8_subsection("十") or
                     "[待填写：基金的非交易过户，来自基金合同第八部分]")
        # 招募 十二 → 合同八.十一
        v.setdefault("CONTRACT_PART8_SEC11", _get_part8_subsection("十一") or
                     "[待填写：基金份额的冻结和解冻，来自基金合同第八部分]")
        # 招募 十三 → 合同八.十二
        v.setdefault("CONTRACT_PART8_SEC12", _get_part8_subsection("十二") or
                     "[待填写：基金份额的转让，来自基金合同第八部分]")
        # 招募 十四 → 合同八.十三
        v.setdefault("CONTRACT_PART8_SEC13", _get_part8_subsection("十三") or
                     "[待填写：其他业务，来自基金合同第八部分]")
        # 招募 十五 → 合同八.十四（交易所相关，文字含交易所名称）
        v.setdefault("CONTRACT_PART8_SEC14", _get_part8_subsection("十四") or
                     "[待填写：清算交收模式，来自基金合同第八部分]")

        # #十一-十七、十九章 各对应合同部分
        v.setdefault("CONTRACT_INVEST_TEXT", _get_part("第十四部分") or
                     "[待填写：基金的投资，来自基金合同第十四部分]")
        v.setdefault("CONTRACT_ASSET_TEXT", _get_part("第十五部分") or
                     "[待填写：基金的财产，来自基金合同第十五部分]")
        v.setdefault("CONTRACT_VALUATION_TEXT", _get_part("第十六部分") or
                     "[待填写：基金资产估值，来自基金合同第十六部分]")
        v.setdefault("CONTRACT_FEE_TEXT", _get_part("第十七部分") or
                     "[待填写：基金的费用与税收，来自基金合同第十七部分]")
        v.setdefault("CONTRACT_AUDIT_TEXT", _get_part("第十九部分") or
                     "[待填写：基金的会计与审计，来自基金合同第十九部分]")
        v.setdefault("CONTRACT_DISCLOSURE_TEXT", _get_part("第二十部分") or
                     "[待填写：基金的信息披露，来自基金合同第二十部分]")
        v.setdefault("CONTRACT_TERMINATION_TEXT", _get_part("第二十一部分") or
                     "[待填写：基金合同的变更、终止与基金财产的清算，来自基金合同第二十一部分]")

        # #二十章 合同内容摘要 → 第二十六部分
        summary = self._strip_signing_page_from_contract_summary(_get_part("第二十六部分"))
        v.setdefault("CONTRACT_SUMMARY_TEXT", summary or
                     "[待填写：基金合同内容摘要，请先生成基金合同，从第二十六部分复制此处]")

        # #??? ???????? ? ???????????
        v.setdefault("CONTRACT_DISTRIBUTION_TEXT", _get_part("\u7b2c\u5341\u516b\u90e8\u5206") or
                     "[\u5f85\u586b\u5199\uff1a\u57fa\u91d1\u7684\u6536\u76ca\u4e0e\u5206\u914d\uff0c\u6765\u81ea\u57fa\u91d1\u5408\u540c\u7b2c\u5341\u516b\u90e8\u5206]")
        # #??? ?????????? ? ??????????? ?
        v.setdefault("CONTRACT_PART18_SEC1", _get_part_subsection("第十八部分", "一") or
                     "[待填写：基金利润的构成，来自基金合同第十八部分第一项]")
        v.setdefault("CONTRACT_PART18_SEC2", _get_part_subsection("第十八部分", "二") or
                     "[待填写：基金可供分配利润，来自基金合同第十八部分第二项]")
        v.setdefault("CONTRACT_PART18_SEC3", _get_part_subsection("第十八部分", "三") or
                     "[待填写：基金收益分配原则，来自基金合同第十八部分第三项]")
        v.setdefault("CONTRACT_PART18_SEC4", _get_part_subsection("第十八部分", "四") or
                     "[待填写：收益分配方案，来自基金合同第十八部分第四项]")
        v.setdefault("CONTRACT_PART18_SEC5", _get_part_subsection("第十八部分", "五") or
                     "[待填写：收益分配方案的确定、公告与实施，来自基金合同第十八部分第五项]")
        v.setdefault("CONTRACT_PART18_SEC6", _get_part_subsection("第十八部分", "六") or
                     "[待填写：基金收益分配中发生的费用，来自基金合同第十八部分第六项]")

        # 新模板兼容映射：填充模板占位符，但不再按章覆盖模板正文。
        v["INTERPRETATION_DERIVED_FROM_CONTRACT"] = str(v.get("CONTRACT_DEFS_TEXT") or "").strip()
        v["OBJECT_SUBSCRIBE_DERIVED_FROM_CONTRACT"] = self._extract_numbered_item_body(v.get("CONTRACT_PART4_SEC1", ""), "3")
        v["METHOD_SUBSCRIBE_DERIVED_FROM_CONTRACT"] = (
            self._extract_numbered_item_body(v.get("CONTRACT_PART4_SEC1", ""), "2")
            or str(v.get("SUBSCRIPTION_METHOD_TEXT") or "").strip()
            or self._subsection_body_or_title(v.get("CONTRACT_PART4_SEC4", ""))
        )
        v["EFFECTIVENESS_DERIVED_FROM_CONTRACT"] = str(v.get("CONTRACT_EFFECTIVENESS_TEXT") or "").strip()
        v["CONVERSION_DERIVED_FROM_CONTRACT"] = str(v.get("CONTRACT_CONVERSION_TEXT") or "").strip()
        v["LISTING_DERIVED_FROM_CONTRACT_SEC7_1"] = self._subsection_body_or_title(v.get("CONTRACT_PART7_SEC1", ""))
        v["LISTING_DERIVED_FROM_CONTRACT_SEC7_2"] = self._subsection_body_or_title(v.get("CONTRACT_PART7_SEC2", ""))
        v["LISTING_DERIVED_FROM_CONTRACT_SEC7_3"] = self._subsection_body_or_title(v.get("CONTRACT_PART7_SEC3", ""))
        v["LISTING_DERIVED_FROM_CONTRACT_SEC7_5"] = self._subsection_body_or_title(v.get("CONTRACT_PART7_SEC5", ""))
        v["LISTING_DERIVED_FROM_CONTRACT_SEC7_6"] = self._subsection_body_or_title(v.get("CONTRACT_PART7_SEC6", ""))
        v["LISTING_DERIVED_FROM_CONTRACT_SEC7_7"] = self._subsection_body_or_title(v.get("CONTRACT_PART7_SEC7", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_1"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC1", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_2"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC2", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_3"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC3", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_6"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC6", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_7"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC7", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_8"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC8", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_9"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC9", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_10"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC10", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_11"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC11", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_12"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC12", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_13"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC13", ""))
        v["PURCHASE_DERIVED_FROM_CONTRACT_SEC8_14"] = self._subsection_body_or_title(v.get("CONTRACT_PART8_SEC14", ""))
        v["INVESTMENT_DERIVED_FROM_CONTRACT_SEC14"] = str(v.get("CONTRACT_INVEST_TEXT") or "").strip()
        v["BELONGINGS_DERIVED_FROM_CONTRACT_SEC15"] = str(v.get("CONTRACT_ASSET_TEXT") or "").strip()
        v["EVALUATION_DERIVED_FROM_CONTRACT_SEC16"] = str(v.get("CONTRACT_VALUATION_TEXT") or "").strip()
        v["DISTRIBUTION_DERIVED_FROM_CONTRACT_SEC18_1"] = self._subsection_body_or_title(v.get("CONTRACT_PART18_SEC1", ""))
        v["DISTRIBUTION_DERIVED_FROM_CONTRACT_SEC18_2"] = self._subsection_body_or_title(v.get("CONTRACT_PART18_SEC2", ""))
        v["DISTRIBUTION_DERIVED_FROM_CONTRACT_SEC18_"] = self._subsection_body_or_title(v.get("CONTRACT_PART18_SEC3", ""))
        v["DISTRIBUTION_DERIVED_FROM_CONTRACT_SEC18_4"] = self._subsection_body_or_title(v.get("CONTRACT_PART18_SEC4", ""))
        v["DISTRIBUTION_DERIVED_FROM_CONTRACT_SEC18_5"] = self._subsection_body_or_title(v.get("CONTRACT_PART18_SEC5", ""))
        v["DISTRIBUTION_DERIVED_FROM_CONTRACT_SEC18_6"] = self._subsection_body_or_title(v.get("CONTRACT_PART18_SEC6", ""))
        v["EXPENSE_DERIVED_FROM_CONTRACT_SEC17"] = str(v.get("CONTRACT_FEE_TEXT") or "").strip()
        v["ACCOUNTING_DERIVED_FROM_CONTRACT_SEC19"] = str(v.get("CONTRACT_AUDIT_TEXT") or "").strip()
        v["DISCLOSURE_DERIVED_FROM_CONTRACT_SEC20"] = str(v.get("CONTRACT_DISCLOSURE_TEXT") or "").strip()
        v["CHANGE_DERIVED_FROM_CONTRACT_SEC21"] = str(v.get("CONTRACT_TERMINATION_TEXT") or "").strip()
        v["SUMMARY_DERIVED_FROM_CONTRACT_SEC22"] = str(v.get("CONTRACT_SUMMARY_TEXT") or "").strip()

        return v

    # ── Step 3 (兼容旧接口) ────────────────────────────────────────────────
    def _extract_contract_summary(self, v: dict) -> dict:
        """向后兼容旧方法名，实际调用 _extract_contract_sections。"""
        return self._extract_contract_sections(v)

    @staticmethod
    def _chapter_num_to_cn(num: int) -> str:
        cn_map = {
            1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
            11: "十一", 12: "十二", 13: "十三", 14: "十四", 15: "十五", 16: "十六", 17: "十七", 18: "十八",
            19: "十九", 20: "二十", 21: "二十一", 22: "二十二", 23: "二十三", 24: "二十四", 25: "二十五",
        }
        return cn_map.get(num, "")

    @staticmethod
    def _split_top_sections(chapter_body: str) -> dict:
        section_re = re.compile(r"^([一二三四五六七八九十百]+)、[^\n]*", re.MULTILINE)
        matches = list(section_re.finditer(chapter_body or ""))
        sections = {}
        for i, m in enumerate(matches):
            sec_cn = m.group(1)
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(chapter_body)
            sections[sec_cn] = chapter_body[start:end].strip()
        return sections

    def _load_template_fixed_content(self) -> dict:
        cache = getattr(self, "_template_fixed_cache", None)
        if cache is not None:
            return cache

        template_text = PROSPECTUS_TEMPLATE_MD.read_text(encoding="utf-8")
        data = {}
        for chapter_num in range(1, 26):
            chapter_cn = self._chapter_num_to_cn(chapter_num)
            span = self._find_chapter_span(template_text, chapter_cn)
            if not span:
                continue
            start, heading_end, end = span
            body = template_text[heading_end:end].strip()
            data[chapter_cn] = {
                "title": template_text[start:heading_end].strip(),
                "body": body,
                "sections": self._split_top_sections(body),
            }

        self._template_fixed_cache = data
        return data

    def _template_chapter_body(self, chapter_cn: str) -> str:
        return self._load_template_fixed_content().get(chapter_cn, {}).get("body", "")

    def _template_section(self, chapter_cn: str, section_cn: str) -> str:
        return self._load_template_fixed_content().get(chapter_cn, {}).get("sections", {}).get(section_cn, "")

    def _get_reference_prospectus_docx(self, v: dict | None = None) -> Path:
        variant_key = self._get_prospectus_variant_key(v or {}) if v is not None else "SSE_CROSS"
        return _resolve_reference_prospectus_docx(variant_key)

    def _load_reference_fixed_content_from_txt(self, v: dict | None = None) -> dict:
        variant_key = self._get_prospectus_variant_key(v or {}) if v is not None else "SSE_CROSS"
        if variant_key in self._reference_fixed_txt_cache:
            return deepcopy(self._reference_fixed_txt_cache[variant_key])

        data = {"important_notice": ""}
        txt_path = self._get_reference_prospectus_docx(v).with_suffix(".txt")
        if not txt_path.exists():
            self._reference_fixed_txt_cache[variant_key] = data
            return deepcopy(data)

        try:
            reference_text = txt_path.read_text(encoding="utf-8").replace("\r\n", "\n").replace("\r", "\n")
        except Exception:
            self._reference_fixed_txt_cache[variant_key] = data
            return deepcopy(data)

        important_notice = ""
        notice_start = reference_text.find("重要提示")
        if notice_start != -1:
            toc_start = reference_text.find("目录", notice_start)
            notice_block = reference_text[notice_start:toc_start if toc_start != -1 else len(reference_text)].strip()
            if notice_block:
                important_notice = notice_block
        if important_notice:
            data["important_notice"] = important_notice

        body_start = reference_text.rfind("目录")
        body_search_start = body_start if body_start != -1 else 0

        def chapter_display_title(chapter_cn: str) -> str:
            title = str(self._load_template_fixed_content().get(chapter_cn, {}).get("title", "") or "").strip()
            if not title:
                return ""
            return re.sub(r"^第[一二三四五六七八九十百]+章", "", title).strip()

        def find_txt_chapter_span(chapter_cn: str):
            display_title = chapter_display_title(chapter_cn)
            if not display_title:
                return None
            heading_re = re.compile(rf"(?m)^{re.escape(display_title)}\s*$")
            matches = list(heading_re.finditer(reference_text, body_search_start))
            if not matches:
                return None
            current_match = matches[-1]

            next_match = None
            chapter_order = [self._chapter_num_to_cn(i) for i in range(1, 26)]
            current_idx = chapter_order.index(chapter_cn) if chapter_cn in chapter_order else -1
            for next_cn in chapter_order[current_idx + 1:]:
                next_title = chapter_display_title(next_cn)
                if not next_title:
                    continue
                candidate = re.search(rf"(?m)^{re.escape(next_title)}\s*$", reference_text[current_match.end():])
                if candidate:
                    candidate_start = current_match.end() + candidate.start()
                    if next_match is None or candidate_start < next_match:
                        next_match = candidate_start
            return current_match.start(), current_match.end(), next_match if next_match is not None else len(reference_text)

        for chapter_num in range(1, 26):
            chapter_cn = self._chapter_num_to_cn(chapter_num)
            span = self._find_chapter_span(reference_text, chapter_cn)
            if not span:
                span = find_txt_chapter_span(chapter_cn)
            if not span:
                continue
            start, heading_end, end = span
            body = reference_text[heading_end:end].strip()
            if not body:
                continue
            data[chapter_cn] = {
                "title": reference_text[start:heading_end].strip(),
                "body": body,
                "sections": self._split_top_sections(body),
            }

        self._reference_fixed_txt_cache[variant_key] = deepcopy(data)
        return deepcopy(data)

    @staticmethod
    def _merge_reference_fixed_content(primary: dict, fallback: dict) -> dict:
        merged = deepcopy(primary or {"important_notice": ""})
        if not str(merged.get("important_notice") or "").strip():
            merged["important_notice"] = str((fallback or {}).get("important_notice") or "").strip()

        chapter_six_fallback_override_sections = {"八", "九", "十", "十一", "十二", "十三"}

        for chapter_cn, fallback_chapter in (fallback or {}).items():
            if chapter_cn == "important_notice" or not isinstance(fallback_chapter, dict):
                continue

            primary_chapter = merged.get(chapter_cn)
            if not isinstance(primary_chapter, dict):
                merged[chapter_cn] = deepcopy(fallback_chapter)
                continue

            if not str(primary_chapter.get("title") or "").strip():
                primary_chapter["title"] = str(fallback_chapter.get("title") or "").strip()
            if not str(primary_chapter.get("body") or "").strip():
                primary_chapter["body"] = str(fallback_chapter.get("body") or "").strip()

            merged_sections = dict(primary_chapter.get("sections") or {})
            for sec_cn, sec_text in dict(fallback_chapter.get("sections") or {}).items():
                should_override_from_txt = chapter_cn == "六" and sec_cn in chapter_six_fallback_override_sections
                if should_override_from_txt or not str(merged_sections.get(sec_cn) or "").strip():
                    merged_sections[sec_cn] = sec_text
            primary_chapter["sections"] = merged_sections
            merged[chapter_cn] = primary_chapter

        return merged

    def _required_docx_reference_assets(self, v: dict | None = None) -> list[tuple[str, Path]]:
        assets = [
            ("reference", self._get_reference_prospectus_docx(v)),
            ("template", self._get_format_template_prospectus_docx()),
        ]
        deduped = []
        seen = set()
        for asset_type, path in assets:
            key = str(path)
            if key in seen:
                continue
            seen.add(key)
            deduped.append((asset_type, path))
        return deduped

    def _ensure_docx_reference_assets(self, v: dict | None = None) -> None:
        missing = []
        for asset_type, path in self._required_docx_reference_assets(v):
            if path.exists():
                continue
            label = "参考招募说明书" if asset_type == "reference" else "格式模板"
            missing.append(f"{label}: {path.name}")
        if missing:
            raise MissingProspectusReferenceAssetsError(missing)

    def _load_reference_fixed_content(self, v: dict | None = None) -> dict:
        """
        Load fixed chapter/section text from the reference prospectus DOCX for the current variant.
        """
        variant_key = self._get_prospectus_variant_key(v or {}) if v is not None else "SSE_CROSS"
        if variant_key in self._reference_fixed_cache:
            return self._reference_fixed_cache[variant_key]

        data = {"important_notice": ""}
        reference_docx = self._get_reference_prospectus_docx(v)
        if reference_docx.exists():
            try:
                from docx import Document
                from docx.document import Document as DocumentObject
                from docx.oxml.table import CT_Tbl
                from docx.oxml.text.paragraph import CT_P
                from docx.table import Table
                from docx.text.paragraph import Paragraph
                doc = Document(str(reference_docx))

                chapter_order = (
                    "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
                    "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
                    "二十一", "二十二", "二十三", "二十四", "二十五",
                )

                def table_to_lines(table) -> list[str]:
                    rows = []
                    for row in table.rows:
                        values = [cell.text.replace("\r", "\n").strip() for cell in row.cells]
                        if any(values):
                            rows.append(values)
                    if not rows:
                        return []
                    return [" | ".join(value for value in row if value) for row in rows]

                blocks = []
                for block_type, block in _iter_docx_blocks(doc):
                    if block_type == "paragraph":
                        text_value = (block.text or "").strip()
                        if text_value:
                            style_name = ""
                            try:
                                style_name = block.style.name or ""
                            except Exception:
                                style_name = ""
                            blocks.append(("paragraph", text_value, style_name))
                    else:
                        lines = table_to_lines(block)
                        if lines:
                            blocks.append(("table", "\n".join(lines), ""))

                notice_started = False
                notice_lines = []
                toc_seen = False
                for _, block_text, _style_name in blocks:
                    stripped = block_text.strip()
                    if not stripped:
                        continue
                    if stripped == "\u91cd\u8981\u63d0\u793a":
                        notice_started = True
                        notice_lines = [stripped]
                        continue
                    if notice_started and stripped == "\u76ee\u5f55":
                        toc_seen = True
                        break
                    if notice_started:
                        notice_lines.extend(line for line in stripped.splitlines() if line.strip())
                if notice_lines:
                    data["important_notice"] = "\n".join(notice_lines).strip()

                chapter_heading_re = re.compile(r'^\u7b2c([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e]+)\u7ae0')
                current_chapter = None
                current_lines = []
                current_title = ""
                chapter_index = 0
                body_started = False

                def flush_chapter():
                    nonlocal current_chapter, current_lines, current_title
                    if not current_chapter:
                        return
                    body = "\n".join(line for line in current_lines if line.strip()).strip()
                    if body:
                        data[current_chapter] = {
                            "title": current_title,
                            "body": body,
                            "sections": self._split_top_sections(body),
                        }
                    current_chapter = None
                    current_lines = []
                    current_title = ""

                for block_type, block_text, style_name in blocks:
                    stripped = block_text.strip()
                    if not stripped:
                        continue
                    if stripped == "\u76ee\u5f55":
                        toc_seen = True
                        body_started = False
                        continue
                    if not toc_seen:
                        continue

                    style_lower = (style_name or "").strip().lower()
                    is_toc_style = "toc" in style_lower
                    is_heading_style = style_lower in {"heading 1", "heading 2", "标题 1", "标题1", "标题 2", "标题2"}
                    match = chapter_heading_re.match(stripped)
                    reference_heading = self._parse_prospectus_toc_entry(stripped)
                    expected_chapter = chapter_order[chapter_index] if chapter_index < len(chapter_order) else None
                    entering_body = False
                    if not body_started:
                        if is_toc_style or self._is_prospectus_toc_placeholder_line(stripped):
                            continue
                        if not (is_heading_style or match or reference_heading):
                            continue
                        body_started = True
                        entering_body = True
                    promote_reference_heading = bool(reference_heading) and (is_heading_style or entering_body)
                    if is_heading_style or match or promote_reference_heading:
                        flush_chapter()
                        if match:
                            current_chapter = match.group(1)
                            if current_chapter in chapter_order:
                                chapter_index = chapter_order.index(current_chapter) + 1
                        elif promote_reference_heading:
                            current_chapter = reference_heading[0]
                            if current_chapter in chapter_order:
                                chapter_index = chapter_order.index(current_chapter) + 1
                        elif expected_chapter:
                            current_chapter = expected_chapter
                            chapter_index += 1
                        else:
                            current_chapter = None
                        current_title = stripped
                        continue
                    if current_chapter:
                        current_lines.extend(line for line in stripped.splitlines() if line.strip())
                flush_chapter()
            except Exception:
                data = {"important_notice": ""}

        data = self._merge_reference_fixed_content(data, self._load_reference_fixed_content_from_txt(v))
        self._reference_fixed_cache[variant_key] = data
        return data

    def _load_reference_section7_table_xmls(self, v: dict | None = None) -> list[dict[str, str | None]]:
        variant_key = self._get_prospectus_variant_key(v or {}) if v is not None else "SSE_CROSS"
        if variant_key in self._reference_section7_table_cache:
            return [
                {
                    **dict(item),
                    "rows": [list(row) for row in (item.get("rows") or [])],
                }
                for item in self._reference_section7_table_cache[variant_key]
            ]

        reference_docx = self._get_reference_prospectus_docx(v)
        if not reference_docx.exists():
            self._reference_section7_table_cache[variant_key] = []
            return []

        try:
            from docx import Document
            doc = Document(str(reference_docx))
        except Exception:
            self._reference_section7_table_cache[variant_key] = []
            return []

        in_format_section = False
        pending_title = None
        table_entries = []
        intro_markers = {
            "申购赎回清单的格式举例如下：",
            "T日申购赎回清单的格式举例如下：",
        }
        for block_type, block in _iter_docx_blocks(doc):
            if block_type == "paragraph":
                text_value = (block.text or "").strip()
                if any(marker in text_value for marker in intro_markers):
                    in_format_section = True
                    pending_title = None
                    continue
                if in_format_section and text_value.startswith("说明："):
                    break
                if in_format_section and text_value:
                    pending_title = text_value
                continue

            if in_format_section:
                table_entries.append({
                    "title": pending_title,
                    "table_xml": block._tbl.xml,
                    "rows": self._docx_table_to_rows(block),
                })
                pending_title = None

        self._reference_section7_table_cache[variant_key] = table_entries
        return [
            {
                **dict(item),
                "rows": [list(row) for row in (item.get("rows") or [])],
            }
            for item in table_entries
        ]

    def _load_canonical_reference_fixed_content(self) -> dict:
        return self._load_reference_fixed_content({"EXCHANGE": "SSE", "MARKET_TYPE": "A_SHARE"})

    @staticmethod
    def _normalize_index_compiler_name(name: str) -> str:
        normalized = (name or "").strip()
        return {
            "深证信息有限公司": "深圳证券信息有限公司",
        }.get(normalized, normalized)

    @staticmethod
    def _section7_format_titles() -> tuple[str, ...]:
        return ("基本信息", "T-1日信息内容", "T日信息内容", "成份股信息内容")

    def _replace_reference_fund_name(self, text: str, v: dict) -> str:
        body = (text or "").strip()
        if not body:
            return ""
        target_name = str(v.get("FUND_NAME") or "").strip()
        if not target_name:
            return body

        reference_names = set()
        for context in ({"EXCHANGE": "SSE", "MARKET_TYPE": "A_SHARE"}, v or {}):
            reference_docx = self._get_reference_prospectus_docx(context)
            reference_name = reference_docx.stem.replace("招募说明书", "").strip()
            if reference_name:
                reference_names.add(reference_name)

        for reference_name in sorted(reference_names, key=len, reverse=True):
            body = body.replace(reference_name, target_name)
        return body

    def _replace_reference_fund_name_in_xml(self, xml_text: str, v: dict) -> str:
        body = xml_text or ""
        target_name = str(v.get("FUND_NAME") or "").strip()
        if not body or not target_name:
            return body

        reference_names = set()
        for context in ({"EXCHANGE": "SSE", "MARKET_TYPE": "A_SHARE"}, v or {}):
            reference_docx = self._get_reference_prospectus_docx(context)
            reference_name = reference_docx.stem.replace("招募说明书", "").strip()
            if reference_name:
                reference_names.add(reference_name)

        for reference_name in sorted(reference_names, key=len, reverse=True):
            body = body.replace(reference_name, target_name)
        return body

    @staticmethod
    def _docx_table_to_rows(table) -> list[list[str]]:
        rows = []
        for row in getattr(table, "rows", []):
            cells = [str(getattr(cell, "text", "") or "").strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        return rows

    @staticmethod
    def _markdown_table_lines_to_rows(table_lines: list[str]) -> list[list[str]]:
        rows = []
        for line in table_lines:
            stripped = str(line or "").strip()
            if not (stripped.startswith("|") and stripped.endswith("|")):
                continue
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if not any(cells):
                continue
            if all(not cell.strip("-: ") for cell in cells):
                continue
            rows.append(cells)
        return rows

    @staticmethod
    def _subscription_fee_table_rows() -> list[list[str]]:
        return [
            ["认购份额（S）", "认购费率"],
            ["S＜100万份", "0.30%"],
            ["S≥100万份", "每笔500元"],
        ]

    def _table_block_from_rows(
        self,
        rows: list[list[str]],
        *,
        table_kind: str,
        title: str | None = None,
        source: str | None = None,
        table_xml: str | None = None,
        bordered: bool = False,
        merge_spans: list[dict] | None = None,
        column_widths: list[int] | None = None,
        header_align: str | None = None,
    ) -> dict:
        block = {
            "type": "table",
            "table_kind": table_kind,
            "rows": [list(row) for row in rows],
        }
        if title:
            block["title"] = title
        if source:
            block["source"] = source
        if table_xml:
            block["table_xml"] = table_xml
        if bordered:
            block["bordered"] = True
        if merge_spans:
            block["merge_spans"] = [dict(span) for span in merge_spans]
        if column_widths:
            block["column_widths"] = [int(width) for width in column_widths]
        if header_align:
            block["header_align"] = str(header_align)
        return block

    def _table_kind_from_rows(self, rows: list[list[str]]) -> str:
        if rows == self._subscription_fee_table_rows():
            return "subscription_fee"
        return "markdown"

    @classmethod
    def _cash_substitution_ratio_formula_render_model(cls) -> dict:
        return {
            "type": "formula",
            "formula_kind": "cash_substitution",
            "plain_text": cls._cash_substitution_ratio_formula_text(),
            "prefix": "现金替代比例（%）＝",
            "sigma": "Σ",
            "upper": "n",
            "lower": "i=1",
            "body_segments": [
                {"text": "（第i只替代证券的数量×该证券参考价格）×100%／（申购基金份额×参考基金份额净值）"},
            ],
        }

    def _section7_reference_table_block(self, table_entry: dict, v: dict | None = None) -> dict:
        context = v or {}
        title_text = self._replace_reference_fund_name((table_entry.get("title") or "").strip(), context)
        target_name = str(context.get("FUND_NAME") or "").strip()
        raw_rows = [list(row) for row in (table_entry.get("rows") or [])]
        rows = []
        for row in raw_rows:
            rendered_row = [self._replace_reference_fund_name(str(cell or ""), context) for cell in row]
            if target_name and len(rendered_row) >= 2 and str(rendered_row[0]).strip().rstrip("：:") == "基金名称":
                rendered_row[1] = target_name
            rows.append(rendered_row)
        table_xml = self._replace_reference_fund_name_in_xml(table_entry.get("table_xml") or "", context)
        if target_name:
            for row in raw_rows:
                if len(row) < 2:
                    continue
                row_label = str(row[0] or "").strip().rstrip("：:")
                sample_name = str(row[1] or "").strip()
                if row_label == "基金名称" and sample_name:
                    table_xml = table_xml.replace(sample_name, target_name)
            if any(len(row) >= 2 and str(row[0] or "").strip().rstrip("：:") == "基金名称" for row in raw_rows):
                normalized_xml = re.sub(r"\s+", "", table_xml or "")
                normalized_target = re.sub(r"\s+", "", target_name)
                if normalized_target and normalized_target not in normalized_xml:
                    table_xml = ""
        return self._table_block_from_rows(
            rows,
            table_kind="purchase_redemption_list_format",
            title=title_text or None,
            source="reference_docx",
            table_xml=table_xml or None,
        )

    def _enhance_prospectus_render_model(self, model: dict, render_context: dict | None = None) -> dict:
        context = render_context or {"EXCHANGE": "SSE", "MARKET_TYPE": "A_SHARE"}
        intro_markers = {
            "申购赎回清单的格式举例如下：",
            "T日申购赎回清单的格式举例如下：",
        }
        section7_reference_tables = self._load_reference_section7_table_xmls(context)

        for chapter in model.get("chapters", []):
            enhanced_blocks = []
            blocks = chapter.get("blocks") or []
            block_index = 0

            while block_index < len(blocks):
                block = blocks[block_index]
                if not isinstance(block, dict):
                    enhanced_blocks.append(block)
                    block_index += 1
                    continue

                if block.get("type") != "paragraph":
                    enhanced_blocks.append(block)
                    block_index += 1
                    continue

                block_text = str(block.get("text") or "").strip()
                if not block_text:
                    block_index += 1
                    continue

                if block_text.startswith("|") and block_text.endswith("|"):
                    table_lines = []
                    while block_index < len(blocks):
                        candidate = blocks[block_index]
                        if not isinstance(candidate, dict) or candidate.get("type") != "paragraph":
                            break
                        candidate_text = str(candidate.get("text") or "").strip()
                        if not (candidate_text.startswith("|") and candidate_text.endswith("|")):
                            break
                        table_lines.append(candidate_text)
                        block_index += 1

                    rows = self._markdown_table_lines_to_rows(table_lines)
                    if rows:
                        table_kind = self._table_kind_from_rows(rows)
                        enhanced_blocks.append(
                            self._table_block_from_rows(
                                rows,
                                table_kind=table_kind,
                                bordered=(table_kind == "subscription_fee"),
                            )
                        )
                    else:
                        enhanced_blocks.extend({"type": "paragraph", "text": line} for line in table_lines)
                    continue

                enhanced_blocks.append(block)
                block_index += 1

                if (
                    chapter.get("chapter_cn") == "十"
                    and section7_reference_tables
                    and any(marker in block_text for marker in intro_markers)
                ):
                    enhanced_blocks.extend(
                        self._section7_reference_table_block(table_entry, context)
                        for table_entry in section7_reference_tables
                    )
                    while block_index < len(blocks):
                        candidate = blocks[block_index]
                        if not isinstance(candidate, dict) or candidate.get("type") != "paragraph":
                            break
                        candidate_text = str(candidate.get("text") or "").strip()
                        if candidate_text.rstrip("。") not in self._section7_format_titles():
                            break
                        block_index += 1

            chapter["blocks"] = enhanced_blocks

        return model

    def _normalize_prospectus_chapter_ten_text(self, text: str) -> str:
        body = (text or "").strip()
        if not body:
            return ""

        body = body.replace("具体办理时间为指", "具体办理时间为")
        body = body.replace("申购赎回清单的内容与格式示例参见招募说明书。", "")
        body = body.replace("依据本基金合同开展其他服务", "依据基金合同开展其他服务")
        body = re.sub(
            r"申购赎回代理券商可按照一定的标准收取佣金(?:，其中包含证券交易所、登记机构等收取的相关费用)?(?:，具体规定请参见招募说明书及基金产品资料概要)?。",
            "申购赎回代理券商可按不超过0.2%的标准收取佣金，其中包含证券交易所、登记机构等收取的相关费用。",
            body,
        )
        body = re.sub(r"\n{3,}", "\n\n", body)
        return body.strip()

    def _indent_section7_format_titles(self, text: str) -> str:
        body = (text or "").strip()
        if not body:
            return ""

        lines = []
        for line in body.splitlines():
            stripped = line.strip()
            if stripped in self._section7_format_titles():
                lines.append(stripped)
            else:
                lines.append(line)
        return "\n".join(lines).strip()

    @staticmethod
    def _replace_numbered_item_in_section(chapter_body: str, section_cn: str, item_no: str, new_item: str) -> str:
        body = (chapter_body or "").strip()
        if not body:
            return chapter_body

        sec_re = re.compile(rf"^{section_cn}、[^\n]*", re.MULTILINE)
        next_sec_re = re.compile(r"^[一二三四五六七八九十百]+、[^\n]*", re.MULTILINE)
        m = sec_re.search(body)
        if not m:
            return chapter_body
        n = next_sec_re.search(body, m.end())
        sec_end = n.start() if n else len(body)

        sec_block = body[m.start():sec_end]
        item_text = (new_item or "").strip()
        if not item_text:
            return chapter_body
        if not re.match(rf"^{re.escape(item_no)}、", item_text):
            item_text = f"{item_no}、{item_text}"

        item_re = re.compile(rf"^{re.escape(item_no)}、[^\n]*", re.MULTILINE)
        next_item_re = re.compile(r"^\d+、[^\n]*", re.MULTILINE)
        im = item_re.search(sec_block)
        if not im:
            if not sec_block.endswith("\n"):
                sec_block += "\n"
            sec_block = f"{sec_block}{item_text}\n"
        else:
            inn = next_item_re.search(sec_block, im.end())
            item_end = inn.start() if inn else len(sec_block)
            sec_block = f"{sec_block[:im.start()]}{item_text}\n{sec_block[item_end:].lstrip()}"

        new_body = f"{body[:m.start()]}{sec_block}{body[sec_end:]}"
        return new_body

    def _apply_reference_fixed_content(self, text: str, v: dict) -> str:
        """
        Apply canonical fixed text from red-dividend prospectus docx.
        """
        canonical_ref = self._load_canonical_reference_fixed_content()
        ref = self._load_reference_fixed_content(v)
        if not canonical_ref and not ref:
            return text

        def ref_chapter(chap_cn: str) -> str:
            return ref.get(chap_cn, {}).get("body", "")

        def ref_section(chap_cn: str, sec_cn: str) -> str:
            return ref.get(chap_cn, {}).get("sections", {}).get(sec_cn, "")

        def canonical_chapter(chap_cn: str) -> str:
            return canonical_ref.get(chap_cn, {}).get("body", "")

        def canonical_section(chap_cn: str, sec_cn: str) -> str:
            return canonical_ref.get(chap_cn, {}).get("sections", {}).get(sec_cn, "")

        chapter_one = self._replace_reference_fund_name(canonical_chapter("一"), v)
        if chapter_one:
            text = self._replace_chapter_body(text, "一", chapter_one)

        ch3 = canonical_chapter("三")
        if ch3:
            manager_name = str(v.get("FUND_MANAGER_NAME") or "[待填写：基金经理姓名]").strip()
            manager_bio = str(v.get("FUND_MANAGER_BIO") or "[待填写：基金经理简介]").strip()
            sec_item3 = f"3、本基金的基金经理为{manager_name}。"
            if manager_bio:
                sec_item3 = f"{sec_item3}\n{manager_bio}"
            ch3 = self._replace_numbered_item_in_section(ch3, "二", "3", sec_item3)
            text = self._replace_chapter_body(text, "三", ch3)

        text = self._replace_chapter_body(text, "四", "【托管人情况待填写】")

        text = self._replace_subsection_in_chapter(text, "五", "一", canonical_section("五", "一"))
        text = self._replace_subsection_in_chapter(text, "五", "二", canonical_section("五", "二"))
        sec3 = canonical_section("五", "三")
        if sec3:
            sec3 = re.sub(r"^经办律师[：:].*$", "经办律师：丁媛、李晓露", sec3, flags=re.MULTILINE)
        text = self._replace_subsection_in_chapter(text, "五", "三", sec3)
        text = self._replace_subsection_in_chapter(text, "五", "四", "四、审计基金财产的会计师事务所\n【待填写】")

        text = self._replace_chapter_body(text, "八", ref_chapter("八"))
        text = self._replace_chapter_body(text, "十四", ref_chapter("十四"))

        for chap_cn in ("二十二", "二十三", "二十四"):
            text = self._replace_chapter_body(text, chap_cn, ref_chapter(chap_cn))

        return text

    def _find_chapter_span(self, text: str, chapter_cn: str):
        """Return (start, heading_end, end) for a chapter, or None."""
        chap_re = re.compile(rf"^第{chapter_cn}章[^\n]*", re.MULTILINE)
        any_chap_re = re.compile(r"^第[一二三四五六七八九十百]+章[^\n]*", re.MULTILINE)
        matches = list(chap_re.finditer(text))
        if matches:
            m = matches[-1]
            n = any_chap_re.search(text, m.end())
            end = n.start() if n else len(text)
            return m.start(), m.end(), end

        toc_entries = self._extract_prospectus_toc_entries(text)
        if not toc_entries:
            return None

        entry_index = next((idx for idx, entry in enumerate(toc_entries) if entry["chapter_cn"] == chapter_cn), None)
        if entry_index is None:
            return None

        lines = text.splitlines(keepends=True)
        line_offsets = []
        offset = 0
        for line in lines:
            line_offsets.append(offset)
            offset += len(line)
        body_start = line_offsets[toc_entries[-1]["line_index"] + 1] if toc_entries[-1]["line_index"] + 1 < len(line_offsets) else 0

        def find_heading(entry: dict, start_pos: int):
            candidates = [entry["title"], entry["display_title"]]
            found = None
            for candidate in candidates:
                pattern = re.compile(rf"^{re.escape(candidate)}\s*$", re.MULTILINE)
                match = pattern.search(text, start_pos)
                if match and (found is None or match.start() < found.start()):
                    found = match
            return found

        current_match = find_heading(toc_entries[entry_index], body_start)
        if not current_match:
            return None

        next_match = None
        for next_entry in toc_entries[entry_index + 1:]:
            candidate = find_heading(next_entry, current_match.end())
            if candidate and (next_match is None or candidate.start() < next_match.start()):
                next_match = candidate

        end = next_match.start() if next_match else len(text)
        return current_match.start(), current_match.end(), end

    def _replace_chapter_body(self, text: str, chapter_cn: str, new_body: str) -> str:
        new_body = (new_body or "").strip()
        if not new_body:
            return text
        span = self._find_chapter_span(text, chapter_cn)
        if not span:
            return text
        start, heading_end, end = span
        prefix = text[:heading_end].rstrip("\n")
        suffix = text[end:].lstrip("\n")
        if suffix:
            return f"{prefix}\n{new_body}\n{suffix}"
        return f"{prefix}\n{new_body}\n"

    def _replace_subsection_in_chapter(self, text: str, chapter_cn: str, subsection_cn: str, new_subsection: str) -> str:
        new_subsection = (new_subsection or "").strip()
        if not new_subsection:
            return text
        span = self._find_chapter_span(text, chapter_cn)
        if not span:
            return text
        start, _heading_end, end = span
        block = text[start:end]

        sec_re = re.compile(rf"^{subsection_cn}、[^\n]*", re.MULTILINE)
        next_sec_re = re.compile(r"^[一二三四五六七八九十百]+、[^\n]*", re.MULTILINE)
        m = sec_re.search(block)
        if not m:
            # If subsection does not exist in template, append it at chapter tail.
            if not block.endswith("\n"):
                block += "\n"
            block = f"{block}{new_subsection}\n"
        else:
            n = next_sec_re.search(block, m.end())
            sec_end = n.start() if n else len(block)
            block = f"{block[:m.start()]}{new_subsection}\n{block[sec_end:].lstrip()}"

        return f"{text[:start]}{block}{text[end:]}"

    @staticmethod
    def _retag_subsection_number(subsection_text: str, new_cn: str) -> str:
        txt = (subsection_text or "").strip()
        return re.sub(r"^[一二三四五六七八九十百]+、", f"{new_cn}、", txt, count=1)

    @staticmethod
    def _ensure_subsection_heading(subsection_text: str, cn_num: str) -> str:
        txt = (subsection_text or "").strip()
        if not txt:
            return ""
        if re.match(r"^[一二三四五六七八九十百]+、", txt):
            return txt
        return f"{cn_num}、\n{txt}"

    @staticmethod
    def _first_nonempty(*values) -> str:
        for value in values:
            if (value or "").strip():
                return value.strip()
        return ""

    @staticmethod
    def _subsection_body_or_title(subsection_text: str) -> str:
        txt = (subsection_text or "").strip()
        if not txt:
            return ""

        lines = txt.splitlines()
        first_line = lines[0].strip()
        match = re.match(r"^[一二三四五六七八九十百]+、(.*)$", first_line)
        if not match:
            return txt

        remainder = "\n".join(lines[1:]).strip()
        return remainder or match.group(1).strip()

    @staticmethod
    def _extract_numbered_item_body(section_text: str, item_no: str) -> str:
        txt = (section_text or "").strip()
        if not txt:
            return ""

        pattern = re.compile(r"^(\d+)、([^\n]*)", re.MULTILINE)
        matches = list(pattern.finditer(txt))
        for idx, match in enumerate(matches):
            if match.group(1) != str(item_no):
                continue
            start = match.end()
            end = matches[idx + 1].start() if idx + 1 < len(matches) else len(txt)
            body = txt[start:end].strip()
            return body or match.group(2).strip()
        return ""

    def _replace_standalone_placeholder_tokens(self, text: str, v: dict) -> str:
        updated = text
        keys = sorted(
            (
                key
                for key in v.keys()
                if isinstance(key, str) and re.fullmatch(r"[A-Z][A-Z0-9_]*", key)
            ),
            key=len,
            reverse=True,
        )
        for key in keys:
            value = v.get(key)
            if value is None:
                continue
            replacement = "是" if isinstance(value, bool) and value else "否" if isinstance(value, bool) else str(value)
            updated = re.sub(
                rf"(?m)^(?P<indent>\s*){re.escape(key)}(?P<trail>\s*)$",
                lambda match: f"{match.group('indent')}{replacement}{match.group('trail')}",
                updated,
            )
        return updated

    @staticmethod
    def _normalize_top_level_section_spacing(text: str) -> str:
        body = (text or "").strip()
        if not body:
            return ""

        lines = [line.rstrip() for line in body.splitlines()]
        result = []
        seen_top_level_section = False
        top_level_re = re.compile(r"^[一二三四五六七八九十百]+、")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if result and result[-1] != "":
                    result.append("")
                continue
            if top_level_re.match(stripped):
                if seen_top_level_section and result and result[-1] != "":
                    result.append("")
                seen_top_level_section = True
            result.append(line)

        return "\n".join(result).strip()

    def _chapter_body_text(self, text: str, chapter_cn: str) -> str:
        span = self._find_chapter_span(text, chapter_cn)
        if not span:
            return ""
        _, heading_end, end = span
        return text[heading_end:end].strip()

    @staticmethod
    def _ensure_chapter_eight_intro(chapter_body: str) -> str:
        intro = "基金合同生效后，为提高交易便利，本基金可以进行基金份额折算。"
        body = (chapter_body or "").strip()
        if not body:
            return intro
        if body.startswith(intro):
            return body
        return f"{intro}\n\n{body}"

    @staticmethod
    def _ensure_chapter_nine_body(chapter_body: str) -> str:
        body = (chapter_body or "").strip()
        if not body:
            return "一、基金份额的上市"
        body = re.sub(r"^一、基金份额的上市交易(?=\n|$)", "一、基金份额的上市", body, count=1)
        if re.match(r"^一、基金份额的上市(?=\n|$)", body):
            return body
        return f"一、基金份额的上市\n\n{body}"

    def _normalize_all_chapter_body_spacing(self, text: str) -> str:
        normalized_text = text
        for chapter_num in range(1, 26):
            chapter_cn = self._chapter_num_to_cn(chapter_num)
            chapter_body = self._chapter_body_text(normalized_text, chapter_cn)
            if not chapter_body:
                continue
            normalized_text = self._replace_chapter_body(
                normalized_text,
                chapter_cn,
                self._normalize_top_level_section_spacing(chapter_body),
            )
        return normalized_text

    @staticmethod
    def _join_nonempty_blocks(blocks) -> str:
        return "\n\n".join(block.strip() for block in blocks if (block or "").strip())

    @staticmethod
    def _normalize_reused_prospectus_chapter(chapter_text: str) -> str:
        text = (chapter_text or "").strip()
        if not text:
            return ""
        text = text.replace("详见招募说明书的规定", "详见招募说明书“侧袋机制”部分的规定")
        return ProspectusEngine._normalize_prospectus_contract_references(text)

    @staticmethod
    def _normalize_prospectus_contract_references(text: str) -> str:
        normalized = (text or "").strip()
        if not normalized:
            return ""
        normalized = re.sub(r"(?m)^4、基金合同或本基金合同：", "4、基金合同：", normalized)
        normalized = re.sub(r"(?m)^6、招募说明书：", "6、招募说明书或本招募说明书：", normalized)
        replacements = [
            ("在本基金合同中", "在本招募说明书中"),
            ("基金合同或本基金合同", "基金合同"),
            ("本基金合同当事人", "基金合同当事人"),
            ("对本基金合同的任何有效修订和补充", "对基金合同的任何有效修订和补充"),
            ("依据本基金合同开展其他服务", "依据基金合同开展其他服务"),
            ("本基金合同", "基金合同"),
        ]
        for old, new in replacements:
            normalized = normalized.replace(old, new)
        return normalized

    def _normalize_prospectus_contract_references_outside_contract_summary(self, text: str) -> str:
        normalized_text = str(text or "")
        if not normalized_text.strip():
            return ""
        for chapter_num in range(1, 26):
            if chapter_num == 20:
                continue
            chapter_cn = self._chapter_num_to_cn(chapter_num)
            chapter_body = self._chapter_body_text(normalized_text, chapter_cn)
            if not chapter_body or "本基金合同" not in chapter_body:
                continue
            normalized_body = self._normalize_prospectus_contract_references(chapter_body)
            if normalized_body != chapter_body:
                normalized_text = self._replace_chapter_body(normalized_text, chapter_cn, normalized_body)
        return normalized_text.strip()

    @staticmethod
    def _get_prospectus_min_sub_unit(v: dict) -> str:
        raw = str(v.get("MIN_SUB_UNIT") or "").strip()
        if not raw:
            return "100万份"

        normalized = raw.replace("，", ",")

        wan_match = re.search(r"(\d+(?:\.\d+)?)\s*(?:万|萬)(?:份)?", normalized)
        if wan_match:
            value = wan_match.group(1)
            if "." in value:
                value = value.rstrip("0").rstrip(".")
            return f"{value}万份"

        count_match = re.search(r"(\d[\d,]*)\s*份", normalized)
        if count_match:
            count = int(count_match.group(1).replace(",", ""))
            if count % 10000 == 0:
                return f"{count // 10000}万份"
            return f"{count / 10000:.4f}".rstrip("0").rstrip(".") + "万份"

        bare_number = normalized.replace(",", "")
        if re.fullmatch(r"\d+(?:\.\d+)?", bare_number):
            if "." in bare_number:
                return bare_number.rstrip("0").rstrip(".") + "万份"
            count = int(bare_number)
            if count >= 10000:
                if count % 10000 == 0:
                    return f"{count // 10000}万份"
                return f"{count / 10000:.4f}".rstrip("0").rstrip(".") + "万份"
            return f"{count}万份"

        return raw

    def _normalize_prospectus_risk_chapter(self, chapter_text: str, v: dict) -> str:
        text = (chapter_text or "").strip()
        if not text:
            return ""
        min_sub_unit = self._get_prospectus_min_sub_unit(v)
        for old in ("1,000,000份（即100万份）", "1,000,000份", "100万份"):
            text = text.replace(old, min_sub_unit)
        text = re.sub(r"\d[\d,]*份（即\d+(?:\.\d+)?万份）", min_sub_unit, text)
        text = re.sub(r"(?<=按原)\d[\d,]*份(?:（即\d+万份）)?", min_sub_unit, text)
        text = re.sub(r"(?<=新的)\d[\d,]*份(?:（即\d+万份）)?", min_sub_unit, text)
        text = re.sub(r"(?<=最小申购赎回单位设置较高（目前为)[^）]+", min_sub_unit, text)
        text = re.sub(r"（目前为\d+(?:\.\d+)?万份）", f"（目前为{min_sub_unit}）", text)
        return text

    def _build_chapter_nine_body(self, v: dict, ref: dict) -> str:
        reference_body = str(ref.get("九", {}).get("body", "") or "").strip()
        if reference_body:
            body = self._ensure_chapter_nine_body(self._replace_manual_reference_text(reference_body, v))
            return self._normalize_top_level_section_spacing(body)

        fallback = self._replace_prospectus_placeholders(
            v.get("PROSPECTUS_CH9_BODY", ""),
            {
                "{LISTING_IOPV_CLAUSE}": str(v.get("LISTING_IOPV_CLAUSE") or "").strip(),
            },
        )
        return self._normalize_top_level_section_spacing(self._ensure_chapter_nine_body(fallback))

    def _build_chapter_eighteen_body(self, v: dict, ref: dict) -> str:
        reference_body = str(ref.get("十八", {}).get("body", "") or "").strip()
        chapter_body = reference_body or v.get("PROSPECTUS_CH18_BODY", "")
        if reference_body:
            chapter_body = self._replace_manual_reference_text(chapter_body, v)
        chapter_body = self._normalize_prospectus_risk_chapter(chapter_body, v)
        return self._normalize_top_level_section_spacing(chapter_body)

    def _build_chapter_ten_sec7(self, v: dict, ref: dict | None = None) -> str:
        reference_section7 = self._reference_chapter_ten_sec7(ref, v) if ref else ""
        section7 = (reference_section7 or v.get("PROSPECTUS_CH10_SEC7_BODY") or v.get("PROSPECTUS_CH10_SEC7") or "").strip()
        if not section7:
            return ""

        section7 = self._ensure_subsection_heading(section7, "七")
        if reference_section7:
            section7 = self._replace_reference_fund_name(section7, v)
            section7 = self._indent_section7_format_titles(section7)
            return self._normalize_prospectus_chapter_ten_text(section7)
        if "申购赎回清单的格式举例如下：" in section7 and "T日申购赎回清单的格式举例如下：" not in section7:
            section7 = section7.replace("申购赎回清单的格式举例如下：", "T日申购赎回清单的格式举例如下：", 1)

        formula_label = (v.get("PROSPECTUS_CH10_SEC7_FORMULA_LABEL") or "").strip()
        formula_text = (v.get("PROSPECTUS_CH10_SEC7_FORMULA_TEXT") or "").strip()
        if formula_label and formula_text:
            formula_block = f"{formula_label}\n{formula_text}"
            if formula_label in section7:
                formula_pattern = re.escape(formula_label) + r"\n[^\n]+"
                if re.search(formula_pattern, section7):
                    section7 = re.sub(formula_pattern, formula_block, section7, count=1)
                elif formula_text not in section7:
                    section7 = section7.replace(formula_label, formula_block, 1)
            elif formula_label not in section7 and "3）替代金额的处理程序" in section7:
                section7 = section7.replace("3）替代金额的处理程序", f"{formula_block}\n3）替代金额的处理程序", 1)

        format_block = (v.get("PROSPECTUS_CH10_SEC7_FORMAT_BLOCK") or "").strip()
        if format_block:
            format_intro = "T日申购赎回清单的格式举例如下："
            if format_intro in section7 and format_block not in section7:
                prefix, suffix = section7.split(format_intro, 1)
                suffix_lines = suffix.lstrip("\n").splitlines()
                removable_prefix_count = 0
                for line in suffix_lines:
                    stripped = line.strip()
                    if not stripped:
                        removable_prefix_count += 1
                        continue
                    if stripped in self._section7_format_titles() or stripped.startswith("|"):
                        removable_prefix_count += 1
                        continue
                    break
                preserved_suffix = "\n".join(suffix_lines[removable_prefix_count:]).lstrip("\n")
                section7 = f"{prefix}{format_intro}\n{format_block}"
                if preserved_suffix:
                    section7 = f"{section7}\n{preserved_suffix}"

        section7 = self._replace_reference_fund_name(section7, v)
        section7 = self._indent_section7_format_titles(section7)
        return self._normalize_prospectus_chapter_ten_text(section7)

    @staticmethod
    def _stock_subscription_formula_render_model() -> dict:
        return {
            "type": "formula",
            "formula_kind": "stock_subscription",
            "plain_text": "投资者的认购份额＝Σ（第i只股票在网下股票认购期最后一日的均价×有效认购数量）／1.00",
            "prefix": "投资者的认购份额＝",
            "sigma": "Σ",
            "upper": "n",
            "lower": "i=1",
            "body_segments": [
                {"text": "（第i只股票在网下股票认购期最后一日的均价×有效认购数量）／1.00"},
            ],
        }

    @classmethod
    def _stock_subscription_formula_text(cls) -> str:
        return cls._stock_subscription_formula_render_model()["plain_text"]

    @staticmethod
    def _subscription_fee_table_markdown() -> str:
        return "\n".join([
            "|认购份额（S）|认购费率|",
            "|---|---|",
            "|S＜100万份|0.30%|",
            "|S≥100万份|每笔500元|",
        ])

    @staticmethod
    def _normalize_standard_subscription_fee_table(text: str) -> str:
        normalized = str(text or "")
        normalized = normalized.replace("|S＜100万份|0.3%|", "|S＜100万份|0.30%|")
        normalized = normalized.replace("S＜100万份 | 0.3%", "S＜100万份 | 0.30%")
        return normalized

    @staticmethod
    def _cash_substitution_ratio_formula_label() -> str:
        return "现金替代比例的计算公式为："

    @classmethod
    def _cash_substitution_ratio_formula_text(cls) -> str:
        return "现金替代比例（%）＝Σ（第i只替代证券的数量×该证券参考价格）×100%／（申购基金份额×参考基金份额净值）"

    def _section7_reference_format_block(self, v: dict | None = None) -> str:
        titles = []
        for entry in self._load_reference_section7_table_xmls(v):
            title = str(entry.get("title") or "").strip()
            if title:
                titles.append(title)
        if not titles:
            titles = ["基本信息", "T-1日信息内容", "T日信息内容", "成份股信息内容"]
        if titles:
            titles[-1] = titles[-1].rstrip("。") + "。"
        return "\n".join(titles).strip()

    def _compact_reference_section7_format_block(self, text: str, v: dict | None = None) -> str:
        body = str(text or "").strip()
        if not body:
            return ""

        format_intro = None
        for candidate in ("T日申购赎回清单的格式举例如下：", "申购赎回清单的格式举例如下："):
            if candidate in body:
                format_intro = candidate
                break
        if not format_intro:
            return body

        prefix, suffix = body.split(format_intro, 1)
        suffix_lines = suffix.lstrip("\n").splitlines()
        note_index = next(
            (idx for idx, line in enumerate(suffix_lines) if str(line or "").strip().startswith("说明")),
            len(suffix_lines),
        )
        preserved_suffix = "\n".join(suffix_lines[note_index:]).strip()
        compacted = f"{prefix}T日申购赎回清单的格式举例如下：\n{self._section7_reference_format_block(v)}"
        if preserved_suffix:
            compacted = f"{compacted}\n{preserved_suffix}"
        return compacted.strip()

    def _reference_chapter_ten_sec7(self, ref: dict | None, v: dict) -> str:
        ref_section = str((ref or {}).get("十", {}).get("sections", {}).get("七", "") or "").strip()
        if not ref_section:
            return ""
        ref_section = self._replace_reference_fund_name(ref_section, v)
        ref_section = self._compact_reference_section7_format_block(ref_section, v)
        return self._ensure_subsection_heading(ref_section, "七")

    def _ensure_stock_subscription_formula(self, section_text: str) -> str:
        section_text = (section_text or "").strip()
        if not section_text:
            return ""

        formula_text = self._stock_subscription_formula_text()
        if formula_text in section_text:
            return section_text

        formula_label = "6、认购份额的计算公式："
        formula_start = section_text.find(formula_label)
        if formula_start == -1:
            return section_text

        between_start = formula_start + len(formula_label)
        explanation_start = section_text.find("\n其中", between_start)
        if explanation_start == -1:
            return "".join([
                section_text[:between_start],
                f"\n{formula_text}",
                section_text[between_start:],
            ]).strip()

        between_lines = [line.strip() for line in section_text[between_start:explanation_start].splitlines() if line.strip()]
        if between_lines == [formula_text]:
            return section_text

        return "".join([
            section_text[:between_start],
            f"\n{formula_text}",
            section_text[explanation_start:],
        ])

    def _restore_missing_prospectus_blocks(self, text: str) -> str:
        restored = str(text or "")

        fee_anchor = "认购费用由投资人承担，不高于0.30%，认购费率如下表所示："
        fee_table = self._subscription_fee_table_markdown()
        if fee_anchor in restored and fee_table not in restored:
            restored = restored.replace(fee_anchor, f"{fee_anchor}\n{fee_table}", 1)

        restored = self._ensure_stock_subscription_formula(restored)

        note_anchor = "成份股信息内容。说明："
        if note_anchor in restored:
            restored = restored.replace(note_anchor, "成份股信息内容。\n说明：", 1)

        formula_label = self._cash_substitution_ratio_formula_label()
        formula_text = self._cash_substitution_ratio_formula_text()
        formula_block = f"{formula_label}\n{formula_text}"
        if formula_label in restored and formula_text not in restored:
            restored = restored.replace(formula_label, formula_block, 1)
        elif formula_label not in restored and "1）禁止现金替代" in restored:
            restored = restored.replace("1）禁止现金替代", f"{formula_block}\n1）禁止现金替代", 1)

        return restored

    def _build_chapter_six_body(self, v: dict, ref: dict) -> str:
        def ref_section(sec_cn: str) -> str:
            return ref.get("六", {}).get("sections", {}).get(sec_cn, "")

        def template_section(sec_cn: str) -> str:
            return self._template_section("六", sec_cn)

        def ensure_named_section(sec_cn: str, section_text: str) -> str:
            text = (section_text or "").strip()
            if not text:
                return ""
            if re.match(r"^[一二三四五六七八九十百]+、", text):
                return text
            heading_source = self._first_nonempty(ref_section(sec_cn), template_section(sec_cn))
            heading_line = str(heading_source or "").strip().splitlines()[0].strip() if str(heading_source or "").strip() else f"{sec_cn}、"
            return f"{heading_line}\n{text}"

        has_stock_subscription = bool(v.get("HAS_STOCK_SUBSCRIPTION"))
        default_sec8 = v.get("CHAPTER6_SEC8_DEFAULT_BODY", "") or self._default_chapter_six_sec8_body()
        sec8 = self._restore_missing_prospectus_blocks(
            self._first_nonempty(ref_section("八"), default_sec8)
        )
        sec8 = self._normalize_standard_subscription_fee_table(sec8)
        sec9 = self._first_nonempty(ref_section("九"), v.get("ONLINE_CASH_SUBSCRIPTION_CLAUSE", ""), template_section("九"))
        sec10 = self._first_nonempty(ref_section("十"), v.get("OFFLINE_CASH_SUBSCRIPTION_CLAUSE", ""), template_section("十"))

        sec11_fallbacks = []
        if has_stock_subscription:
            sec11_fallbacks.extend([
                v.get("OFFLINE_STOCK_SUBSCRIPTION_CLAUSE", ""),
                v.get("PROSPECTUS_CH6_SEC11", ""),
            ])
        else:
            sec11_fallbacks.append(v.get("PROSPECTUS_CH6_SEC11", ""))
        sec11 = self._restore_missing_prospectus_blocks(
            self._first_nonempty(ref_section("十一"), *sec11_fallbacks, template_section("十一"))
        )
        if "认购份额的计算公式" in sec11 or self._stock_subscription_formula_text() in sec11:
            sec11 = self._ensure_stock_subscription_formula(sec11)

        sec12 = self._first_nonempty(ref_section("十二"), v.get("PROSPECTUS_CH6_SEC12", ""), template_section("十二"))
        sec13 = ""
        ref_sec13 = ref_section("十三")
        if has_stock_subscription and ref_sec13:
            sec13 = self._first_nonempty(ref_sec13, v.get("PROSPECTUS_CH6_SEC13", ""), template_section("十三"))

        sec1 = ensure_named_section("一", self._first_nonempty(template_section("一"), ref_section("一")))
        sec2 = ensure_named_section(
            "二",
            self._first_nonempty(v.get("OBJECT_SUBSCRIBE_DERIVED_FROM_CONTRACT", ""), ref_section("二"), template_section("二")),
        )
        sec3 = ensure_named_section("三", self._first_nonempty(template_section("三"), ref_section("三")))
        sec4 = ensure_named_section(
            "四",
            self._first_nonempty(ref_section("四"), v.get("METHOD_SUBSCRIBE_DERIVED_FROM_CONTRACT", ""), v.get("PROSPECTUS_CH6_SEC4", ""), template_section("四")),
        )
        sec5 = ensure_named_section("五", self._first_nonempty(template_section("五"), ref_section("五")))
        sec6 = ensure_named_section("六", self._first_nonempty(template_section("六"), ref_section("六")))
        sec7 = ensure_named_section(
            "七",
            self._first_nonempty(ref_section("七"), v.get("SUB_ACCOUNT_OPENING_CLAUSE", ""), v.get("PROSPECTUS_CH6_SEC7", ""), template_section("七")),
        )
        sec8 = ensure_named_section("八", sec8)
        sec9 = ensure_named_section("九", sec9)
        sec10 = ensure_named_section("十", sec10)
        sec11 = ensure_named_section("十一", sec11)
        sec12 = ensure_named_section("十二", sec12)
        sec13 = ensure_named_section("十三", sec13)

        blocks = [
            str(v.get("CHAPTER6_INTRO_REGISTRATION_SENTENCE") or "").strip(),
            str(v.get("CHAPTER6_PRODUCT_TYPE_SENTENCE") or self._default_chapter_six_product_type_sentence()).strip(),
            sec1,
            sec2,
            sec3,
            sec4,
            sec5,
            sec6,
            sec7,
            sec8,
            sec9,
            sec10,
            sec11,
            sec12,
            sec13,
        ]
        return self._normalize_top_level_section_spacing(self._join_nonempty_blocks(blocks))

    @staticmethod
    def _build_chapter_ten_limits_table(min_sub_unit: str) -> str:
        return "\n".join([
            "|项目|内容|",
            "|---|---|",
            f"|最小申购赎回单位|{min_sub_unit}|",
            "|申购/赎回份额上限|以申购赎回清单或相关公告为准|",
        ])

    def _build_chapter_ten_sec5(self, ref: dict, v: dict) -> str:
        sec5 = ref.get("\u5341", {}).get("sections", {}).get("\u4e94", "")
        min_sub_unit = self._get_prospectus_min_sub_unit(v)
        canonical_item_one = (
            f"1\u3001\u6295\u8d44\u4eba\u7533\u8d2d\u3001\u8d4e\u56de\u7684\u57fa\u91d1\u4efd\u989d\u9700\u4e3a\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d\u7684\u6574\u6570\u500d\u3002"
            f"\u76ee\u524d\uff0c\u672c\u57fa\u91d1\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d\u4e3a{min_sub_unit}\uff0c\u57fa\u91d1\u7ba1\u7406\u4eba\u6709\u6743\u5bf9\u5176\u8fdb\u884c\u8c03\u6574\uff0c"
            "\u5e76\u5728\u8c03\u6574\u5b9e\u65bd\u524d\u4f9d\u7167\u300a\u4fe1\u606f\u62ab\u9732\u529e\u6cd5\u300b\u7684\u6709\u5173\u89c4\u5b9a\u5728\u89c4\u5b9a\u5a92\u4ecb\u4e0a\u516c\u544a\u3002"
        )
        if sec5:
            sec5 = re.sub(
                r"(?ms)^1\u3001\u6295\u8d44\u4eba\u7533\u8d2d\u3001\u8d4e\u56de\u7684\u57fa\u91d1\u4efd\u989d\u9700[\s\S]*?\u89c4\u5b9a\u5a92\u4ecb\u4e0a\u516c\u544a\u3002\s*",
                f"{canonical_item_one}\n",
                sec5.strip(),
                count=1,
            )
            if canonical_item_one not in sec5:
                sec5 = re.sub(
                    r"(?m)^(\u4e94\u3001[^\n]*\n)",
                    lambda m: f"{m.group(1)}{canonical_item_one}\n",
                    sec5.lstrip(),
                    count=1,
                )
            if canonical_item_one not in sec5:
                sec5 = f"{canonical_item_one}\n{sec5.lstrip()}"
            cleaned_lines = []
            canonical_seen = False
            duplicate_prefixes = (
                "1、投资人申购、赎回的基金份额需为最小申购赎回单位的整数倍。",
                "投资人申购、赎回的基金份额需为最小申购赎回单位的整数倍。",
            )
            for raw_line in sec5.splitlines():
                stripped = raw_line.strip()
                if stripped == canonical_item_one:
                    if canonical_seen:
                        continue
                    canonical_seen = True
                    cleaned_lines.append(raw_line)
                    continue
                if stripped.startswith(duplicate_prefixes):
                        continue
                cleaned_lines.append(raw_line)
            sec5 = "\n".join(cleaned_lines).strip()
            return self._normalize_prospectus_chapter_ten_text(sec5)
        return str(v.get("CHAPTER10_SEC5_DEFAULT_BODY") or self._default_chapter_ten_sec5_body(v)).strip()

    def _build_contract_section(self, v: dict, var_name: str, sec_cn: str) -> str:
        sec_text = self._retag_subsection_number(v.get(var_name, ""), sec_cn)
        sec_text = self._normalize_prospectus_chapter_ten_text(sec_text)
        return self._ensure_subsection_heading(sec_text, sec_cn)

    def _build_manual_chapter_ten_section(self, ref: dict, v: dict, sec_cn: str, fallback_text: str = "") -> str:
        ref_sections = (ref or {}).get("十", {}).get("sections", {}) or {}
        if ref_sections:
            section_text = str(ref_sections.get(sec_cn) or "").strip()
            if not section_text:
                return ""
            section_text = self._replace_manual_reference_text(section_text, v)
        else:
            section_text = str(fallback_text or "").strip()
            if not section_text:
                return ""
        section_text = self._normalize_prospectus_chapter_ten_text(section_text)
        return self._ensure_subsection_heading(section_text, sec_cn)

    def _build_chapter_twenty_one_placeholder_body(self, v: dict) -> str:
        titles = [line.strip() for line in (v.get("PROSPECTUS_CH21_TITLES") or "").splitlines() if line.strip()]
        if not titles:
            return ""

        remaining_sections = list((_get_custodian_summary_state().get("sections") or []))
        blocks = []
        for title in titles:
            content = ""
            if remaining_sections:
                match = _match_custodian_summary_section(remaining_sections, title)
                matched_section = match.get("section")
                if matched_section:
                    content = str(matched_section.get("content") or "").strip()
                    remaining_sections = [section for section in remaining_sections if section is not matched_section]
            blocks.extend([title, content or str(v.get("CHAPTER21_PLACEHOLDER_TEXT") or "【待填写】").strip()])
        return "\n".join(blocks)

    def _display_chapter_heading(self, chapter_cn: str) -> str:
        title = self._load_template_fixed_content().get(chapter_cn, {}).get("title", "").strip()
        parsed = self._parse_prospectus_chapter_heading(title)
        if parsed:
            return self._format_prospectus_reference_heading(*parsed)
        if title:
            return self._format_prospectus_reference_heading(chapter_cn, title)
        return f"{chapter_cn}、"

    def _find_display_chapter_span(self, text: str, chapter_cn: str):
        chapter_heading = self._display_chapter_heading(chapter_cn)
        chapter_re = re.compile(rf"^{re.escape(chapter_heading)}$", re.MULTILINE)
        matches = list(chapter_re.finditer(text))
        if not matches:
            return None
        match = matches[-1]
        next_match = None
        chapter_order = [self._chapter_num_to_cn(i) for i in range(1, 26)]
        current_num = chapter_order.index(chapter_cn) + 1 if chapter_cn in chapter_order else 0
        for num in range(current_num + 1, 26):
            next_heading = self._display_chapter_heading(self._chapter_num_to_cn(num))
            candidate_re = re.compile(rf"^{re.escape(next_heading)}$", re.MULTILINE)
            candidate = candidate_re.search(text, match.end())
            if candidate and (next_match is None or candidate.start() < next_match.start()):
                next_match = candidate
        end = next_match.start() if next_match else len(text)
        return match.start(), match.end(), end

    def _replace_display_chapter_body(self, text: str, chapter_cn: str, new_body: str) -> str:
        new_body = (new_body or "").strip()
        if not new_body:
            return text
        span = self._find_display_chapter_span(text, chapter_cn)
        if not span:
            return text
        start, heading_end, end = span
        prefix = text[:heading_end].rstrip("\n")
        suffix = text[end:].lstrip("\n")
        if suffix:
            return f"{prefix}\n{new_body}\n{suffix}"
        return f"{prefix}\n{new_body}\n"

    def _replace_display_subsection_in_chapter(self, text: str, chapter_cn: str, subsection_cn: str, new_subsection: str) -> str:
        new_subsection = (new_subsection or "").strip()
        if not new_subsection:
            return text
        span = self._find_display_chapter_span(text, chapter_cn)
        if not span:
            return text
        start, _heading_end, end = span
        block = text[start:end]

        sec_re = re.compile(rf"^{subsection_cn}、[^\n]*", re.MULTILINE)
        next_sec_re = re.compile(r"^[一二三四五六七八九十百]+、[^\n]*", re.MULTILINE)
        match = sec_re.search(block)
        if not match:
            if not block.endswith("\n"):
                block += "\n"
            block = f"{block}{new_subsection}\n"
        else:
            next_match = next_sec_re.search(block, match.end())
            sec_end = next_match.start() if next_match else len(block)
            block = f"{block[:match.start()]}{new_subsection}\n{block[sec_end:].lstrip()}"

        return f"{text[:start]}{block}{text[end:]}"

    def _apply_canonical_reference_overrides(self, text: str, v: dict) -> str:
        canonical_ref = self._load_reference_fixed_content()

        chapter_fourteen = canonical_ref.get("十四", {})
        section_four = chapter_fourteen.get("sections", {}).get("四", "").strip()
        if not section_four:
            section_four = self._split_top_sections(chapter_fourteen.get("body", "")).get("四", "").strip()
        if section_four:
            text = self._replace_display_subsection_in_chapter(text, "十四", "四", section_four)

        chapter_twenty_one = self._build_chapter_twenty_one_placeholder_body(v)
        if chapter_twenty_one:
            text = self._replace_display_chapter_body(text, "二十一", chapter_twenty_one)

        for chapter_cn in ("二十二", "二十三", "二十四", "二十五"):
            chapter_body = canonical_ref.get(chapter_cn, {}).get("body", "")
            chapter_body = self._replace_reference_fund_name(chapter_body, v)
            if chapter_body:
                text = self._replace_display_chapter_body(text, chapter_cn, chapter_body)

        return text

    def _prospectus_clause_text(self, clause_key: str, variant_key: str = "DEFAULT") -> str:
        return str(
            self.pro_clauses.get(clause_key, {})
            .get("variants", {})
            .get(variant_key, {})
            .get("text", "")
            or ""
        )

    def _prospectus_product_summary_item_four_text(self) -> str:
        return self._prospectus_clause_text("PRODUCT_SUMMARY_DISCLOSURE_ITEM")

    def _rewrite_disclosure_product_summary_item_four(self, text: str) -> str:
        span = self._find_display_chapter_span(text, "十七")
        if not span:
            return text

        start, _heading_end, end = span
        chapter_block = text[start:end]

        section_match = re.search(r"^五、公开披露的基金信息[^\n]*", chapter_block, flags=re.MULTILINE)
        if not section_match:
            return text

        next_section_match = re.search(
            r"^[一二三四五六七八九十百]+、[^\n]*",
            chapter_block[section_match.end():],
            flags=re.MULTILINE,
        )
        section_end = section_match.end() + next_section_match.start() if next_section_match else len(chapter_block)
        section_block = chapter_block[section_match.start():section_end]
        new_item = self._prospectus_product_summary_item_four_text()

        item_match = re.search(r"^4、[^\n]*", section_block, flags=re.MULTILINE)
        next_item_pattern = re.compile(r"^\d+、[^\n]*", re.MULTILINE)
        if item_match:
            next_item_match = next_item_pattern.search(section_block, item_match.end())
            next_disclosure_item_match = re.search(
                r"^（[一二三四五六七八九十百]+）[^\n]*",
                section_block[item_match.end():],
                flags=re.MULTILINE,
            )
            boundary_candidates = []
            if next_item_match:
                boundary_candidates.append(next_item_match.start())
            if next_disclosure_item_match:
                boundary_candidates.append(item_match.end() + next_disclosure_item_match.start())
            item_end = min(boundary_candidates) if boundary_candidates else len(section_block)
            updated_section = f"{section_block[:item_match.start()]}{new_item}\n{section_block[item_end:].lstrip()}"
        else:
            insert_at = len(section_block)
            for candidate in next_item_pattern.finditer(section_block):
                item_no = int(candidate.group(0).split("、", 1)[0])
                if item_no > 4:
                    insert_at = candidate.start()
                    break

            if insert_at == len(section_block):
                suffix = "" if section_block.endswith("\n") else "\n"
                updated_section = f"{section_block}{suffix}{new_item}\n"
            else:
                before = section_block[:insert_at].rstrip("\n")
                after = section_block[insert_at:].lstrip("\n")
                updated_section = f"{before}\n{new_item}\n{after}"

        updated_chapter = f"{chapter_block[:section_match.start()]}{updated_section}{chapter_block[section_end:]}"
        return f"{text[:start]}{updated_chapter}{text[end:]}"

    @staticmethod
    def _normalize_prospectus_contract_term(text: str) -> str:
        return text.replace("本基金合同", "基金合同")

    def _ensure_distribution_conditions_section(self, chapter_body: str) -> str:
        if not chapter_body or re.search(r"(?m)^\u56db\u3001\u6536\u76ca\u5206\u914d\u6761\u4ef6(?:\n|$)", chapter_body):
            return chapter_body
        if not re.search(r"(?m)^\u56db\u3001\u6536\u76ca\u5206\u914d\u65b9\u6848(?:\n|$)", chapter_body):
            return chapter_body
        conditions = self._prospectus_clause_text("DISTRIBUTION_CONDITIONS_SECTION")
        return re.sub(r"(?m)^\u56db\u3001\u6536\u76ca\u5206\u914d\u65b9\u6848", f"{conditions}\n\u4e94\u3001\u6536\u76ca\u5206\u914d\u65b9\u6848", chapter_body, count=1)

    def _build_chapter_ten_body(self, v: dict, ref: dict) -> str:
        blocks = []
        prelude = v.get("PROSPECTUS_CH10_PRELUDE", "")
        if prelude:
            blocks.append(prelude)
        sec1 = (
            self._build_contract_section(v, "CONTRACT_PART8_SEC1", "一")
            if str(v.get("CONTRACT_PART8_SEC1") or "").strip()
            else self._build_manual_chapter_ten_section(ref, v, "一")
        )
        blocks.extend([
            sec1,
            self._build_contract_section(v, "CONTRACT_PART8_SEC2", "二"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC3", "三"),
            self._build_manual_chapter_ten_section(ref, v, "四", v.get("PROSPECTUS_CH10_SEC4", "")),
            self._build_chapter_ten_sec5(ref, v),
            self._build_contract_section(v, "CONTRACT_PART8_SEC6", "六"),
            self._build_chapter_ten_sec7(v, ref),
            self._build_manual_chapter_ten_section(ref, v, "八", v.get("CONTRACT_PART8_SEC7", "")),
            self._build_manual_chapter_ten_section(ref, v, "九", v.get("CONTRACT_PART8_SEC8", "")),
            self._build_manual_chapter_ten_section(ref, v, "十", v.get("CONTRACT_PART8_SEC9", "")),
            self._build_manual_chapter_ten_section(ref, v, "十一", v.get("CONTRACT_PART8_SEC10", "")),
            self._build_manual_chapter_ten_section(ref, v, "十二", v.get("CONTRACT_PART8_SEC11", "")),
            self._build_manual_chapter_ten_section(ref, v, "十三", v.get("CONTRACT_PART8_SEC12", "")),
            self._build_contract_section(v, "CONTRACT_PART8_SEC13", "十四"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC14", "十五"),
        ])
        return self._join_nonempty_blocks(blocks)

    def _get_product_type_chapter_builders(self, v: dict, ref: dict) -> dict:
        product_type = self._get_product_type(v)
        if product_type == "ETF":
            return {
                "十": lambda: self._build_chapter_ten_body(v, ref),
            }
        return {}

    def _apply_prospectus_chapter_logic(self, text: str, v: dict) -> str:
        """
        Apply chapter-level composition rules for prospectus generation.
        """
        text = self._apply_reference_fixed_content(text, v)
        ref = self._load_reference_fixed_content(v)

        chapter_builders = {
            "二": lambda: self._normalize_prospectus_definitions(v.get("CONTRACT_DEFS_TEXT", "")),
            "六": lambda: self._build_chapter_six_body(v, ref),
            "七": lambda: v.get("PROSPECTUS_CH7_BODY", ""),
            "九": lambda: self._build_chapter_nine_body(v, ref),
            "十": lambda: self._build_chapter_ten_body(v, ref),
            "十一": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_INVEST_TEXT", "")),
            "十二": lambda: v.get("CONTRACT_ASSET_TEXT", ""),
            "十三": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_VALUATION_TEXT", "")),
            "十四": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_DISTRIBUTION_TEXT", "")) or ref.get("十四", {}).get("body", ""),
            "十五": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_FEE_TEXT", "")),
            "十六": lambda: v.get("CONTRACT_AUDIT_TEXT", ""),
            "十七": lambda: v.get("CONTRACT_DISCLOSURE_TEXT", ""),
            "十八": lambda: self._build_chapter_eighteen_body(v, ref),
            "十九": lambda: v.get("CONTRACT_TERMINATION_TEXT", ""),
            "二十": lambda: v.get("CONTRACT_SUMMARY_TEXT", ""),
            "二十一": lambda: self._build_chapter_twenty_one_placeholder_body(v),
        }
        for chap_cn, builder in chapter_builders.items():
            text = self._replace_chapter_body(text, chap_cn, builder())

        chapter_eight_body = self._chapter_body_text(text, "八")
        if chapter_eight_body:
            text = self._replace_chapter_body(text, "八", self._ensure_chapter_eight_intro(chapter_eight_body))

        chapter_nine_body = self._chapter_body_text(text, "九")
        if chapter_nine_body:
            text = self._replace_chapter_body(text, "九", self._ensure_chapter_nine_body(chapter_nine_body))

        sec3 = self._ensure_subsection_heading(v.get("CONTRACT_PART18_SEC3", ""), "三")
        text = self._replace_subsection_in_chapter(text, "十四", "三", sec3)
        text = self._normalize_prospectus_contract_references_outside_contract_summary(text)
        return self._normalize_all_chapter_body_spacing(text)

    def _process_conditionals(self, text: str, v: dict) -> str:
        return engine._process_conditionals(text, v)

    def _replace_placeholders(self, text: str, v: dict) -> str:
        return engine._replace_placeholders(text, v)

    # ── Step 6: 重排序号（使用"章"作为重置标志）─────────────────────────
    def _renumber_sequences(self, text: str) -> str:
        lines = text.split("\n")
        RE_NUM = re.compile(r"^(\d+)(、)")
        RE_RESET = re.compile(
            r"^(?:[一二三四五六七八九十百]+、|第[一二三四五六七八九十百]+章)"
        )
        last_num = None
        result = []
        for line in lines:
            if RE_RESET.match(line.strip()):
                last_num = None
                result.append(line)
                continue
            m = RE_NUM.match(line)
            if m:
                num = int(m.group(1))
                if num == 1:
                    last_num = 1
                elif last_num is not None and num > last_num + 1:
                    expected = last_num + 1
                    line = re.sub(r"^\d+、", f"{expected}、", line, count=1)
                    last_num = expected
                else:
                    last_num = num
            result.append(line)
        return "\n".join(result)

    # ── Step 7: 清理（检测招募说明书封面而非基金合同）────────────────────
    def _looks_like_cover_title(self, line: str) -> bool:
        s = line.strip()
        if not s:
            return False
        if s.startswith(("#", ">", "-", "*", "**", "`")):
            return False
        if not s.endswith("\u62db\u52df\u8bf4\u660e\u4e66"):
            return False
        forbidden = ("\u6a21\u677f\u8bf4\u660e", "\u6761\u4ef6\u53d8\u91cf\u5f15\u7528\u8bf4\u660e", "\u5dee\u5f02\u6761\u6b3e", ".json", "_CLAUSE", "_DEF")
        return not any(token in s for token in forbidden)

    def _find_body_start_index(self, lines: list[str]) -> int:
        nonempty = [(idx, line.strip()) for idx, line in enumerate(lines) if line.strip()]
        for pos, (idx, stripped) in enumerate(nonempty):
            if not self._looks_like_cover_title(stripped):
                continue
            window = [item[1] for item in nonempty[pos + 1:pos + 6]]
            has_manager = any(item.startswith("\u57fa\u91d1\u7ba1\u7406\u4eba") for item in window)
            has_custodian = any(item.startswith("\u57fa\u91d1\u6258\u7ba1\u4eba") for item in window)
            if has_manager and has_custodian:
                return idx
        for idx, stripped in nonempty:
            if self._looks_like_cover_title(stripped):
                return idx
        return 0

    def _is_internal_metadata_line(self, line: str) -> bool:
        s = line.strip()
        if not s:
            return False
        if s in {"---", "***"}:
            return True
        if s.startswith((">", "# ", "## ")):
            return True
        keywords = (
            "\u6a21\u677f\u8bf4\u660e",
            "\u6761\u4ef6\u53d8\u91cf\u5f15\u7528\u8bf4\u660e",
            "\u5dee\u5f02\u6761\u6b3e\u5f15\u7528\u8bf4\u660e",
            "VALUATION_TIMING_CLAUSE",
            "WORKING_DAY_DEF",
            "BUSINESS_RULES_DEF",
            "NON_COMPONENT_SCOPE",
            "DISTRIBUTION_FREQ_CLAUSE",
            "MGMT_FEE_PAYMENT_METHOD",
            "CUSTODY_FEE_PAYMENT_METHOD",
        )
        if any(keyword in s for keyword in keywords):
            return True
        if ".json" in s and "\u89c1" in s and "`" in s:
            return True
        return False

    def validate_exportable_text(self, text: str) -> dict:
        metadata_matches = []

        def _append_unique(target: list[str], value: str):
            if value not in target:
                target.append(value)

        for raw_line in text.splitlines():
            stripped = raw_line.strip()
            if not stripped:
                continue
            if self._is_internal_metadata_line(stripped):
                _append_unique(metadata_matches, stripped)

        if metadata_matches:
            return {
                "ok": False,
                "error_type": "template_metadata_leaked",
                "error": "招募说明书正文中仍包含模板说明或内部标记，请先清理后再导出。",
                "matches": metadata_matches[:5],
            }
        # Placeholder content is allowed to export; only internal template metadata blocks delivery.
        return {"ok": True, "matches": []}

    def _cleanup(self, text: str) -> str:
        lines = text.split("\n")
        start_idx = self._find_body_start_index(lines)
        clean = []
        for idx, line in enumerate(lines):
            if idx < start_idx:
                continue
            if self._is_internal_metadata_line(line):
                continue
            if line.strip() in {"{SUBSCRIPTION_ACCOUNT_CLAUSE}", "{SUB_ACCOUNT_OPENING_CLAUSE}"}:
                continue
            clean.append(line)
        text = "\n".join(clean)
        text = text.replace("{SUBSCRIPTION_ACCOUNT_CLAUSE}", "")
        text = text.replace("{SUB_ACCOUNT_OPENING_CLAUSE}", "")
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        return text

    @staticmethod
    def _normalize_prospectus_definitions(definitions_text: str) -> str:
        lines = (definitions_text or "").splitlines()
        if not lines:
            return ""

        normalized = []
        for line in lines:
            stripped = line.strip()
            updated = line
            if stripped.startswith("在本基金合同中"):
                updated = line.replace("在本基金合同中", "在本招募说明书中", 1)
            elif stripped.startswith("4、基金合同或本基金合同："):
                updated = line.replace("4、基金合同或本基金合同：", "4、基金合同：", 1)
                updated = updated.replace("本基金合同", "基金合同")
            elif stripped.startswith("6、招募说明书："):
                updated = line.replace("6、招募说明书：", "6、招募说明书或本招募说明书：", 1)
            elif re.match(r"^\d+、不可抗力：", stripped):
                updated = line.replace("本基金合同当事人", "基金合同当事人")
                updated = updated.replace("本基金合同", "基金合同")
            normalized.append(updated)

        return ProspectusEngine._normalize_prospectus_contract_references("\n".join(normalized).strip())

    @staticmethod
    def _split_important_notice_lines(notice_text: str) -> list[str]:
        return [line.strip() for line in (notice_text or "").splitlines() if line.strip()]

    @staticmethod
    def _normalize_important_notice_style(value: str | None) -> str:
        style = str(value or "AUTO").strip().upper()
        if style in {"DIVIDEND_QUALITY", "CHUANGYE_200", "AUTO"}:
            return style
        return "AUTO"

    def _resolve_important_notice_style(self, v: dict) -> str:
        explicit_style = self._normalize_important_notice_style(v.get("IMPORTANT_NOTICE_STYLE"))
        if explicit_style != "AUTO":
            return explicit_style

        index_name = str(v.get("INDEX_NAME") or "").strip()
        dividend_hit = any(keyword in index_name for keyword in ("中证", "科创"))
        chuangye_hit = any(keyword in index_name for keyword in ("国证", "创业板"))
        if dividend_hit and not chuangye_hit:
            return "DIVIDEND_QUALITY"
        if chuangye_hit and not dividend_hit:
            return "CHUANGYE_200"
        raise ValueError("重要提示参考口径无法自动判断，请先选择重要提示参考口径。")

    def _important_notice_reference_context(self, v: dict) -> dict:
        style = self._resolve_important_notice_style(v)
        if style == "CHUANGYE_200":
            return {"EXCHANGE": "SZSE", "MARKET_TYPE": "CHUANGYE"}
        return {"EXCHANGE": "SSE", "MARKET_TYPE": "A_SHARE"}

    def _validate_prospectus_inputs(self, form_data: dict) -> None:
        index_description = str((form_data or {}).get("INDEX_DESCRIPTION") or "").strip()
        if not index_description:
            raise ValueError("INDEX_DESCRIPTION 不能为空，生成或导出招募说明书前请先填写指数简介。")
        self._resolve_important_notice_style(form_data or {})

    @staticmethod
    def _extract_index_notice_fallbacks(notice_blocks: list[str]) -> dict:
        fallback = {
            "index_name": "",
            "index_description": "",
            "index_compiler": "",
            "index_website": "",
        }
        if not notice_blocks:
            return fallback

        for line in reversed(notice_blocks):
            stripped = line.strip()
            if not fallback["index_website"]:
                website_match = re.match(
                    r"^有关标的指数具体编制方案及成份股信息详见(.+?)网站，网址[:：]([^。]+)。?$",
                    stripped,
                )
                if website_match:
                    fallback["index_compiler"] = website_match.group(1).strip()
                    fallback["index_website"] = website_match.group(2).strip()
                    continue

            if stripped.startswith("本基金标的指数为"):
                name_match = re.match(r"^本基金标的指数为([^。：]+)", stripped)
                if name_match:
                    fallback["index_name"] = name_match.group(1).strip()
                description_match = re.match(r"^本基金标的指数为[^：。]+[：。](.+)$", stripped)
                if description_match:
                    fallback["index_description"] = description_match.group(1).strip()
                break

        return fallback

    def _build_important_notice_blocks(self, v: dict) -> list[str]:
        ref = self._load_reference_fixed_content(v)
        lines = self._split_important_notice_lines(ref.get("important_notice", ""))
        if not lines:
            return []

        if lines[0] == "重要提示":
            lines = lines[1:]
        if not lines:
            return []

        fallback = self._extract_index_notice_fallbacks(lines)
        approval_no = str(v.get("CSRC_APPROVAL_NO") or "202X年X月X日证监许可〔202X〕XXX号").strip()
        index_name = str(v.get("INDEX_NAME") or fallback["index_name"] or "【指数名称】").strip()
        index_description = str(v.get("INDEX_DESCRIPTION") or "").strip()
        index_compiler = self._normalize_index_compiler_name(str(v.get("INDEX_COMPILER") or fallback["index_compiler"] or "").strip())
        index_website = str(v.get("INDEX_WEBSITE") or fallback["index_website"] or "").strip()

        normalized = [self._replace_manual_reference_text(block, v) for block in lines]
        fallback_index_name = fallback["index_name"]
        if fallback_index_name and fallback_index_name != index_name:
            normalized = [block.replace(fallback_index_name, index_name) for block in normalized]

        normalized[0] = str(
            v.get("IMPORTANT_NOTICE_APPROVAL_SENTENCE")
            or f"本基金经中国证监会{approval_no}文注册募集。"
        ).strip()

        target_sentence = f"本基金是跟踪标的指数的交易型开放式基金，本基金标的指数为{index_name}。"
        tracking_sentence = f"本基金被动跟踪标的指数“{index_name}”，因此，本基金的业绩表现与标的指数的表现密切相关。"
        risk_sentence_re = re.compile(
            r"本基金是跟踪标的指数的交易型开放式基金，本基金标的指数为[^。]*(?:。|，及其未来可能发生的变更。)"
        )
        normalized = [risk_sentence_re.sub(target_sentence, block, count=1) for block in normalized]

        tracking_sentence_re = re.compile(
            r"本基金被动跟踪标的指数“[^”]+”(?:及其未来可能发生的变更)?，因此，本基金的业绩表现与标的指数的表现密切相关。"
        )
        normalized = [tracking_sentence_re.sub(tracking_sentence, block, count=1) for block in normalized]

        while normalized and normalized[-1].startswith("有关标的指数具体编制方案及成份股信息详见"):
            normalized.pop()
        while normalized and normalized[-1].startswith("本基金标的指数为"):
            normalized.pop()

        index_intro = f"本基金标的指数为{index_name}。"
        if index_description:
            index_intro = f"{index_intro}{index_description}"
        normalized.append(index_intro)

        if index_compiler and index_website:
            normalized.append(
                str(
                    v.get("IMPORTANT_NOTICE_INDEX_SOURCE_SENTENCE")
                    or f"有关标的指数具体编制方案及成份股信息详见{index_compiler}网站，网址：{index_website}。"
                ).strip()
            )

        return [block for block in normalized if str(block or "").strip()]

    def _normalize_template_important_notice(self, text: str, v: dict) -> str:
        lines = text.splitlines()
        toc_idx = next((idx for idx, line in enumerate(lines) if line.strip() == "目录"), None)
        if toc_idx is None:
            return text

        notice_idx = next((idx for idx, line in enumerate(lines[:toc_idx]) if line.strip() == "重要提示"), None)
        if notice_idx is None:
            return text

        cover = "\n".join(lines[:notice_idx]).strip()
        notice_lines = self._split_important_notice_lines("\n".join(lines[notice_idx:toc_idx]))
        if notice_lines and notice_lines[0] == "重要提示":
            notice_lines = notice_lines[1:]
        if not notice_lines:
            return text

        fallback = self._extract_index_notice_fallbacks(notice_lines)
        approval_no = str(v.get("CSRC_APPROVAL_NO") or "202X年X月X日证监许可〔202X〕XXX号").strip()
        index_name = str(v.get("INDEX_NAME") or fallback["index_name"] or "【指数名称】").strip()
        index_description = str(v.get("INDEX_DESCRIPTION") or "").strip()
        index_compiler = self._normalize_index_compiler_name(str(v.get("INDEX_COMPILER") or fallback["index_compiler"] or "").strip())
        index_website = str(v.get("INDEX_WEBSITE") or fallback["index_website"] or "").strip()

        normalized = [self._replace_manual_reference_text(block, v) for block in notice_lines]
        fallback_index_name = fallback["index_name"]
        if fallback_index_name and fallback_index_name != index_name:
            normalized = [block.replace(fallback_index_name, index_name) for block in normalized]

        normalized[0] = str(
            v.get("IMPORTANT_NOTICE_APPROVAL_SENTENCE")
            or f"本基金经中国证监会{approval_no}文注册募集。"
        ).strip()

        target_sentence = f"本基金是跟踪标的指数的交易型开放式基金，本基金标的指数为{index_name}。"
        tracking_sentence = f"本基金被动跟踪标的指数“{index_name}”，因此，本基金的业绩表现与标的指数的表现密切相关。"
        risk_sentence_re = re.compile(
            r"本基金是跟踪标的指数的交易型开放式基金，本基金标的指数为[^。]*(?:。|，及其未来可能发生的变更。)"
        )
        normalized = [risk_sentence_re.sub(target_sentence, block, count=1) for block in normalized]

        tracking_sentence_re = re.compile(
            r"本基金被动跟踪标的指数“[^”]+”(?:及其未来可能发生的变更)?，因此，本基金的业绩表现与标的指数的表现密切相关。"
        )
        normalized = [tracking_sentence_re.sub(tracking_sentence, block, count=1) for block in normalized]

        while normalized and normalized[-1].startswith("有关标的指数具体编制方案及成份股信息详见"):
            normalized.pop()
        while normalized and normalized[-1].startswith("本基金标的指数为"):
            normalized.pop()

        index_intro = f"本基金标的指数为{index_name}。"
        if index_description:
            index_intro = f"{index_intro}{index_description}"
        normalized.append(index_intro)

        if index_compiler and index_website:
            normalized.append(
                str(
                    v.get("IMPORTANT_NOTICE_INDEX_SOURCE_SENTENCE")
                    or f"有关标的指数具体编制方案及成份股信息详见{index_compiler}网站，网址：{index_website}。"
                ).strip()
            )

        tail = "\n".join(lines[toc_idx:]).strip()
        notice = "\n".join(["重要提示", *[block for block in normalized if str(block or "").strip()]]).strip()
        return "\n\n".join(part for part in (cover, notice, tail) if part).strip()

    def _inject_important_notice_before_toc(self, text: str, v: dict) -> str:
        notice_blocks = self._build_important_notice_blocks(v)
        if not notice_blocks:
            return text

        lines = text.splitlines()
        toc_idx = next((idx for idx, line in enumerate(lines) if line.strip() == "\u76ee\u5f55"), None)
        if toc_idx is None:
            return text

        notice_idx = next(
            (idx for idx, line in enumerate(lines[:toc_idx]) if line.strip() == "\u91cd\u8981\u63d0\u793a"),
            None,
        )
        cover_end = notice_idx if notice_idx is not None else toc_idx
        cover = "\n".join(lines[:cover_end]).strip()
        tail = "\n".join(lines[toc_idx:]).strip()
        notice = "\n".join(["重要提示", *notice_blocks]).strip()
        return "\n\n".join(part for part in (cover, notice, tail) if part).strip()

    def _get_format_template_prospectus_docx(self) -> Path:
        return _resolve_reference_prospectus_docx("SSE_CROSS")

    def _override_prospectus_toc_styles(self, document) -> None:
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        def ensure_style(style_name: str, base_style_name: str | None = None):
            styles = document.styles
            try:
                style = styles[style_name]
            except KeyError:
                style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            if base_style_name:
                try:
                    style.base_style = styles[base_style_name]
                except KeyError:
                    pass
            return style

        def set_style_font(style, *, eastasia: str, ascii_font: str, size_half_pt: int, bold: bool = False):
            style.font.name = ascii_font
            style.font.size = Pt(size_half_pt / 2)
            style.font.bold = bold
            rpr = style._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), ascii_font)
            rfonts.set(qn("w:hAnsi"), ascii_font)
            rfonts.set(qn("w:eastAsia"), eastasia)
            for tag in ("w:sz", "w:szCs"):
                node = rpr.find(qn(tag))
                if node is None:
                    node = OxmlElement(tag)
                    rpr.append(node)
                node.set(qn("w:val"), str(size_half_pt))
            if bold:
                for tag in ("w:b", "w:bCs"):
                    if rpr.find(qn(tag)) is None:
                        rpr.append(OxmlElement(tag))

        toc_heading = ensure_style("TOC Heading", "Normal")
        set_style_font(toc_heading, eastasia="宋体", ascii_font="宋体", size_half_pt=28, bold=False)
        toc_heading.paragraph_format.first_line_indent = Pt(0)

        for style_name in ("TOC 1", "TOC 2", "TOC 3"):
            toc_style = ensure_style(style_name, "Normal")
            set_style_font(toc_style, eastasia="宋体", ascii_font="宋体", size_half_pt=21, bold=False)
            toc_style.paragraph_format.first_line_indent = Pt(0)

        for style in document.styles:
            try:
                style_name = (style.name or "").strip().lower()
                style_id = (style.style_id or "").strip()
            except Exception:
                continue
            if style.type != WD_STYLE_TYPE.PARAGRAPH:
                continue
            if style_id != "3" and style_name not in {"heading 2", "标题 2", "标题2"}:
                continue

            set_style_font(style, eastasia="黑体", ascii_font="Arial", size_half_pt=32, bold=True)
            style.paragraph_format.first_line_indent = Pt(0)
            ppr = style._element.get_or_add_pPr()
            num_pr = ppr.find(qn("w:numPr"))
            if num_pr is not None:
                ppr.remove(num_pr)

    def _build_display_prospectus_text(self, form_data: dict) -> str:
        v = self._derive_variables(form_data)
        v = self._inject_clause_texts(v)
        v = self._extract_contract_sections(v)
        ref = self._load_reference_fixed_content(v)
        template_text = PROSPECTUS_TEMPLATE_MD.read_text(encoding="utf-8").lstrip("\ufeff")
        text = self._process_conditionals(template_text, v)
        text = self._replace_placeholders(text, v)
        text = self._replace_standalone_placeholder_tokens(text, v)
        text = self._normalize_template_important_notice(text, v)
        text = self._inject_important_notice_before_toc(text, v)
        text = self._replace_chapter_body(text, "六", self._build_chapter_six_body(v, ref))
        text = self._replace_chapter_body(text, "九", self._build_chapter_nine_body(v, ref))
        text = self._replace_chapter_body(text, "十", self._build_chapter_ten_body(v, ref))
        text = self._replace_chapter_body(text, "十八", self._build_chapter_eighteen_body(v, ref))
        text = self._renumber_sequences(text)
        text = self._cleanup(text)
        text = self._normalize_prospectus_contract_references_outside_contract_summary(text)
        text = self._restore_missing_prospectus_blocks(text)
        text = self._format_reference_style_prospectus(text)
        chapter_twenty_one = self._build_chapter_twenty_one_placeholder_body(v)
        if chapter_twenty_one:
            text = self._replace_display_chapter_body(text, "二十一", chapter_twenty_one)
        text = self._rewrite_disclosure_product_summary_item_four(text)
        text = self._normalize_all_chapter_body_spacing(text)
        text = self._normalize_prospectus_contract_references_outside_contract_summary(text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _build_render_model_from_display_text(self, text: str, render_context: dict | None = None) -> dict:
        stock_formula_text = self._stock_subscription_formula_text()
        cash_formula_text = self._cash_substitution_ratio_formula_text()
        model = {
            "cover_lines": [],
            "important_notice_title": "\u91cd\u8981\u63d0\u793a",
            "important_notice_blocks": [],
            "toc_title": "\u76ee\u5f55",
            "toc_titles": [],
            "toc_entries": [],
            "chapters": [],
        }
        phase = "cover"
        in_notice = False
        current_chapter = None
        next_toc_chapter_index = 0
        top_level_re = re.compile(r"^([^\u3001]+)\u3001(.+)$")

        for raw in text.splitlines():
            stripped = raw.strip()
            if not stripped:
                continue

            if phase == "cover":
                if stripped == "\u91cd\u8981\u63d0\u793a":
                    in_notice = True
                    continue
                if stripped == "\u76ee\u5f55":
                    phase = "toc"
                    in_notice = False
                    continue
                if in_notice:
                    model["important_notice_blocks"].append(stripped)
                else:
                    model["cover_lines"].append(stripped)
                continue

            match = top_level_re.match(stripped)
            if phase == "toc":
                if not match or self._is_prospectus_toc_placeholder_line(stripped):
                    continue
                chapter_cn, title = match.groups()
                if model["toc_titles"] and title == model["toc_titles"][0]:
                    phase = "body"
                    next_toc_chapter_index = 1
                    current_chapter = {
                        "chapter_cn": chapter_cn,
                        "title": title,
                        "display_title": stripped,
                        "blocks": [],
                    }
                    model["chapters"].append(current_chapter)
                else:
                    model["toc_titles"].append(title)
                    model["toc_entries"].append({
                        "chapter_cn": chapter_cn,
                        "title": title,
                        "display_title": stripped,
                    })
                continue

            if (
                phase == "body"
                and match
                and next_toc_chapter_index < len(model["toc_entries"])
                and stripped == model["toc_entries"][next_toc_chapter_index]["display_title"]
            ):
                chapter_cn, title = match.groups()
                current_chapter = {
                    "chapter_cn": chapter_cn,
                    "title": title,
                    "display_title": stripped,
                    "blocks": [],
                }
                model["chapters"].append(current_chapter)
                next_toc_chapter_index += 1
                continue

            if current_chapter is None:
                continue
            if stripped == stock_formula_text:
                current_chapter["blocks"].append(self._stock_subscription_formula_render_model())
            elif stripped == cash_formula_text:
                current_chapter["blocks"].append(self._cash_substitution_ratio_formula_render_model())
            else:
                current_chapter["blocks"].append({"type": "paragraph", "text": stripped})

        if not model["toc_entries"]:
            model["toc_entries"] = [
                {
                    "chapter_cn": chapter["chapter_cn"],
                    "title": chapter["title"],
                    "display_title": chapter["display_title"],
                }
                for chapter in model["chapters"]
            ]
        if not model["toc_titles"]:
            model["toc_titles"] = [entry["title"] for entry in model["toc_entries"]]
        return self._enhance_prospectus_render_model(model, render_context)

    def generate_bundle(self, form_data: dict) -> dict:
        self._validate_prospectus_inputs(form_data)
        text = self._build_display_prospectus_text(form_data)
        return {
            "text": text,
            "render_model": self._build_render_model_from_display_text(text, form_data),
        }

    def generate_render_model(self, form_data: dict) -> dict:
        return self.generate_bundle(form_data)["render_model"]

    def generate(self, form_data: dict) -> str:
        return self.generate_bundle(form_data)["text"]

    def build_docx_prospectus(self, prospectus_text: str, form_data: dict | None = None) -> bytes:
        return _build_docx_prospectus_with_form_data(self, prospectus_text, form_data)

def _build_docx_prospectus_with_form_data(self, prospectus_text: str, form_data: dict | None = None) -> bytes:
    reference_context = form_data or {"EXCHANGE": "SSE", "MARKET_TYPE": "A_SHARE"}
    self._ensure_docx_reference_assets(reference_context)
    prospectus_text = str(prospectus_text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    model = self._build_render_model_from_display_text(prospectus_text, reference_context)
    total_blocks = sum(len(chapter.get("blocks") or []) for chapter in model.get("chapters", []))
    if not model.get("chapters") or total_blocks == 0:
        formatted_text = self._format_reference_style_prospectus(prospectus_text)
        formatted_text = self._normalize_all_chapter_body_spacing(formatted_text)
        formatted_text = re.sub(r"\n{3,}", "\n\n", formatted_text).strip()
        formatted_model = self._build_render_model_from_display_text(formatted_text, reference_context)
        formatted_blocks = sum(len(chapter.get("blocks") or []) for chapter in formatted_model.get("chapters", []))
        if formatted_model.get("chapters") and formatted_blocks > 0:
            prospectus_text = formatted_text
            model = formatted_model

    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Pt

    template_docx = self._get_format_template_prospectus_docx()
    doc = Document(str(template_docx))

    def set_default_tab_stop(document, value="420"):
        settings = document.settings.element
        default_tab_stop = settings.find(qn("w:defaultTabStop"))
        if default_tab_stop is None:
            default_tab_stop = OxmlElement("w:defaultTabStop")
            settings.insert(0, default_tab_stop)
        default_tab_stop.set(qn("w:val"), str(value))

    def normalize_normal_style(document):
        try:
            normal_style = document.styles["Normal"]
        except KeyError:
            return
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(10.5)
        normal_style.paragraph_format.line_spacing = 1.5
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.space_after = Pt(0)
        rpr = normal_style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")
        for tag in ("w:sz", "w:szCs"):
            node = rpr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                rpr.append(node)
            node.set(qn("w:val"), "21")

    def clear_body_keep_section(document):
        body = document._element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)

    def clear_paragraph_runs(paragraph):
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)

    def remove_heading2_auto_numbering(document):
        try:
            style = document.styles["Heading 2"]
        except KeyError:
            return
        style_element = style.element
        ppr = style_element.find(qn("w:pPr"))
        if ppr is None:
            return
        numpr = ppr.find(qn("w:numPr"))
        if numpr is not None:
            ppr.remove(numpr)

    def set_run_font(run, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21, bold=False, vertical_align: str | None = None):
        run.font.name = ascii_font
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), ascii_font)
        rfonts.set(qn("w:hAnsi"), ascii_font)
        rfonts.set(qn("w:eastAsia"), eastasia)
        run.font.size = Pt(size_half_pt / 2)
        run.bold = bold
        for tag in ("w:sz", "w:szCs"):
            node = rpr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                rpr.append(node)
            node.set(qn("w:val"), str(size_half_pt))
        if bold:
            for tag in ("w:b", "w:bCs"):
                if rpr.find(qn(tag)) is None:
                    rpr.append(OxmlElement(tag))
        if vertical_align:
            run.font.subscript = vertical_align == "subscript"
            run.font.superscript = vertical_align == "superscript"
            vert_align = rpr.find(qn("w:vertAlign"))
            if vert_align is None:
                vert_align = OxmlElement("w:vertAlign")
                rpr.append(vert_align)
            vert_align.set(qn("w:val"), "subscript" if vertical_align == "subscript" else "superscript")

    def set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False):
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.first_line_indent = Pt(21 if first_line else 0)
        ppr = paragraph._p.get_or_add_pPr()
        snap_to_grid = ppr.find(qn("w:snapToGrid"))
        if snap_to_grid is None:
            snap_to_grid = OxmlElement("w:snapToGrid")
            spacing = ppr.find(qn("w:spacing"))
            children = list(ppr)
            insert_at = children.index(spacing) if spacing is not None else len(children)
            ppr.insert(insert_at, snap_to_grid)
        snap_to_grid.set(qn("w:val"), "0")
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        for attr in ("w:firstLine", "w:firstLineChars", "w:hanging", "w:hangingChars"):
            attr_name = qn(attr)
            if ind.get(attr_name) is not None:
                del ind.attrib[attr_name]
        if first_line:
            ind.set(qn("w:firstLine"), "420")
            ind.set(qn("w:firstLineChars"), "200")
        else:
            ind.set(qn("w:firstLine"), "0")

    section7_format_titles = set(self._section7_format_titles())

    def is_section7_format_title(text_value: str) -> bool:
        return (text_value or "").strip() in section7_format_titles

    def set_left_tab_stop(paragraph, position_twips=840):
        ppr = paragraph._p.get_or_add_pPr()
        tabs = ppr.find(qn("w:tabs"))
        if tabs is not None:
            ppr.remove(tabs)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), str(position_twips))
        tabs.append(tab)
        ppr.append(tabs)

    def set_right_dot_tab_stop(paragraph, position_twips=8296):
        ppr = paragraph._p.get_or_add_pPr()
        tabs = ppr.find(qn("w:tabs"))
        if tabs is not None:
            ppr.remove(tabs)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), str(position_twips))
        tabs.append(tab)
        ppr.append(tabs)

    def add_paragraph(text_value, *, style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21, bold=False):
        paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        set_paragraph_format(paragraph, align=align, first_line=first_line)
        if text_value:
            run = paragraph.add_run(text_value)
            set_run_font(run, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt, bold=bold)
        return paragraph

    def add_formula_run(paragraph, text_value: str, *, size_half_pt=21, vertical_align: str | None = None):
        run = paragraph.add_run(text_value)
        set_run_font(
            run,
            eastasia="\u5b8b\u4f53",
            ascii_font="Times New Roman",
            size_half_pt=size_half_pt,
            bold=False,
            vertical_align=vertical_align,
        )
        return run

    def add_native_formula_paragraph(formula_model: dict):
        def math_run_xml(text_value: str) -> str:
            escaped = xml.sax.saxutils.escape(str(text_value or ""))
            return (
                '<m:r>'
                '<w:rPr>'
                '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体"/>'
                '<w:sz w:val="21"/>'
                '<w:szCs w:val="21"/>'
                "</w:rPr>"
                f"<m:t>{escaped}</m:t>"
                "</m:r>"
            )

        paragraph = doc.add_paragraph()
        set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        prefix = str(formula_model.get("prefix") or "")
        upper = str(formula_model.get("upper") or "")
        lower = str(formula_model.get("lower") or "")
        sigma = xml.sax.saxutils.escape(str(formula_model.get("sigma") or "Σ"))
        body_segments = formula_model.get("body_segments") or []
        body_xml = "".join(math_run_xml(segment.get("text") or "") for segment in body_segments if segment.get("text"))

        if prefix and upper and lower and body_xml:
            formula_xml = f"""
            <m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
                         xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
              <m:oMath>
                {math_run_xml(prefix)}
                <m:nary>
                  <m:naryPr>
                    <m:chr m:val="{sigma}"/>
                    <m:limLoc m:val="undOvr"/>
                  </m:naryPr>
                  <m:sub>{math_run_xml(lower)}</m:sub>
                  <m:sup>{math_run_xml(upper)}</m:sup>
                  <m:e>{body_xml}</m:e>
                </m:nary>
              </m:oMath>
            </m:oMathPara>
            """.strip()
        else:
            formula_text = str(formula_model.get("plain_text") or "")
            formula_xml = f"""
            <m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
                         xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
              <m:oMath>
                {math_run_xml(formula_text)}
              </m:oMath>
            </m:oMathPara>
            """.strip()

        paragraph._p.append(parse_xml(formula_xml))
        return paragraph

    def add_cover_line(text_value):
        paragraph = doc.add_paragraph()
        set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        if not text_value:
            return paragraph
        if "\uff1a" in text_value and text_value.startswith(("\u57fa\u91d1\u7ba1\u7406\u4eba", "\u57fa\u91d1\u6258\u7ba1\u4eba")):
            label, value = text_value.split("\uff1a", 1)
            label_run = paragraph.add_run(f"{label}\uff1a")
            set_run_font(label_run, eastasia="\u5b8b\u4f53", ascii_font="\u5b8b\u4f53", size_half_pt=30, bold=True)
            if value:
                value_run = paragraph.add_run(value)
                set_run_font(value_run, eastasia="\u5b8b\u4f53", ascii_font="\u5b8b\u4f53", size_half_pt=30, bold=True)
        else:
            run = paragraph.add_run(text_value)
            set_run_font(run, eastasia="\u5b8b\u4f53", ascii_font="\u5b8b\u4f53", size_half_pt=30, bold=True)
        return paragraph

    bookmark_id = [0]

    def chapter_anchor_name(index: int) -> str:
        return f"toc_part_{index + 1}"

    def append_run_properties(run_element, *, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21, bold=False):
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), ascii_font)
        rfonts.set(qn("w:hAnsi"), ascii_font)
        rfonts.set(qn("w:eastAsia"), eastasia)
        rpr.append(rfonts)
        for tag in ("w:sz", "w:szCs"):
            node = OxmlElement(tag)
            node.set(qn("w:val"), str(size_half_pt))
            rpr.append(node)
        if bold:
            rpr.append(OxmlElement("w:b"))
            rpr.append(OxmlElement("w:bCs"))
        run_element.append(rpr)

    def append_tab_run(paragraph, *, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21):
        run_element = OxmlElement("w:r")
        append_run_properties(run_element, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt)
        run_element.append(OxmlElement("w:tab"))
        paragraph._p.append(run_element)

    def append_field(paragraph, instruction: str, *, placeholder="1", eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21):
        begin = OxmlElement("w:r")
        append_run_properties(begin, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")
        begin.append(fld_begin)
        paragraph._p.append(begin)

        instr_run = OxmlElement("w:r")
        append_run_properties(instr_run, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt)
        instr = OxmlElement("w:instrText")
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instr.text = instruction
        instr_run.append(instr)
        paragraph._p.append(instr_run)

        sep = OxmlElement("w:r")
        append_run_properties(sep, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt)
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        sep.append(fld_sep)
        paragraph._p.append(sep)

        text_run = OxmlElement("w:r")
        append_run_properties(text_run, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt)
        text_node = OxmlElement("w:t")
        text_node.text = placeholder
        text_run.append(text_node)
        paragraph._p.append(text_run)

        end = OxmlElement("w:r")
        append_run_properties(end, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        end.append(fld_end)
        paragraph._p.append(end)

    def append_hyperlink(paragraph, text_value: str, anchor: str, *, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21):
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)
        hyperlink.set(qn("w:history"), "1")
        run_element = OxmlElement("w:r")
        append_run_properties(run_element, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt)
        text_node = OxmlElement("w:t")
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text_node.text = text_value
        run_element.append(text_node)
        hyperlink.append(run_element)
        paragraph._p.append(hyperlink)

    def insert_bookmark(paragraph, name: str):
        bookmark_id[0] += 1
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(bookmark_id[0]))
        bookmark_start.set(qn("w:name"), name)
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(bookmark_id[0]))
        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)

    def add_chapter_heading(chapter, chapter_index: int):
        heading = doc.add_paragraph(style="Heading 2")
        set_paragraph_format(heading, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
        set_left_tab_stop(heading)
        insert_bookmark(heading, chapter_anchor_name(chapter_index))

        chapter_cn = (chapter.get("chapter_cn") or "").strip()
        title = (chapter.get("title") or "").strip()
        if chapter_cn:
            prefix_run = heading.add_run(f"{chapter_cn}\u3001")
            set_run_font(prefix_run, eastasia="\u9ed1\u4f53", ascii_font="Arial", size_half_pt=32, bold=True)
            tab_run = heading.add_run()
            tab_run._r.append(OxmlElement("w:tab"))

        title_run = heading.add_run(title)
        set_run_font(title_run, eastasia="\u9ed1\u4f53", ascii_font="Arial", size_half_pt=32, bold=True)
        return heading

    def add_manual_toc_item(entry, chapter_index: int):
        paragraph = doc.add_paragraph()
        set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
        set_right_dot_tab_stop(paragraph)
        chapter_cn = (entry.get("chapter_cn") or "").strip()
        title = (entry.get("title") or "").strip()
        title_text = f"{chapter_cn}\u3001{title}" if chapter_cn else title
        append_hyperlink(paragraph, title_text, chapter_anchor_name(chapter_index), eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21)
        append_tab_run(paragraph, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21)
        append_field(
            paragraph,
            f" PAGEREF {chapter_anchor_name(chapter_index)} \\\\h ",
            placeholder="1",
            eastasia="\u5b8b\u4f53",
            ascii_font="Times New Roman",
            size_half_pt=21,
        )

    def update_header_text(header_text, sections):
        if not sections:
            return
        for section in sections:
            for header in (section.header, section.first_page_header):
                paragraph = _reset_header_footer_part(header)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                ppr = paragraph._p.get_or_add_pPr()
                existing = ppr.find(qn("w:pBdr"))
                if existing is not None:
                    ppr.remove(existing)
                pbdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:color"), "auto")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                pbdr.append(bottom)
                ppr.append(pbdr)
                run = paragraph.add_run(header_text)
                set_run_font(run, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=18)

    table_border_size = "4"  # Word border size is in eighths of a point; 4 => 0.5 pt

    def ensure_cell_borders(cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is not None:
            tc_pr.remove(borders)
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), table_border_size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            borders.append(border)
        tc_pr.append(borders)

    def ensure_table_borders(table):
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is not None:
            tbl_pr.remove(borders)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), table_border_size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            borders.append(border)
        tbl_pr.append(borders)
        for row in table.rows:
            for cell in row.cells:
                ensure_cell_borders(cell)

    def render_table_rows(rows, *, bordered=False):
        if not rows:
            return
        col_count = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(rows):
            for col_idx in range(col_count):
                cell = table.cell(row_idx, col_idx)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                clear_paragraph_runs(paragraph)
                set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
                value = row[col_idx] if col_idx < len(row) else ""
                if value:
                    run = paragraph.add_run(value)
                    set_run_font(run, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21, bold=(row_idx == 0))
        if bordered:
            ensure_table_borders(table)

    def render_table_block(table_lines, *, bordered=False):
        render_table_rows(self._markdown_table_lines_to_rows(table_lines), bordered=bordered)

    def format_prospectus_draft_header_text(title_text: str) -> str:
        header = str(title_text or "招募说明书").strip().strip("《》")
        if "招募说明书" not in header:
            header = f"{header}招募说明书"
        header = re.sub(r"招募说明书(?:（草案）)?$", "招募说明书（草案）", header)
        if not header.endswith("（草案）"):
            header = header.replace("招募说明书", "招募说明书（草案）", 1)
        return f"《{header}》"

    def insert_reference_table_block(table_entry: dict):
        body = doc._element.body
        title_text = self._replace_reference_fund_name((table_entry.get("title") or "").strip(), reference_context)
        if title_text:
            add_paragraph(
                title_text,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line=is_section7_format_title(title_text),
                eastasia="宋体",
                ascii_font="Times New Roman",
                size_half_pt=21,
                bold=False,
            )
        insert_at = len(body)
        if insert_at and body[-1].tag == qn("w:sectPr"):
            insert_at -= 1
        table_xml = self._replace_reference_fund_name_in_xml(table_entry.get("table_xml") or "", reference_context)
        if table_xml:
            body.insert(insert_at, parse_xml(table_xml))

    clear_body_keep_section(doc)
    set_default_tab_stop(doc)
    normalize_normal_style(doc)
    _set_update_fields_on_open(doc, OxmlElement, qn)
    remove_heading2_auto_numbering(doc)
    _ensure_word_toc_styles(
        doc,
        heading_eastasia="宋体",
        heading_ascii="宋体",
        heading_size_half_pt=28,
        heading_bold=False,
        entry_eastasia="宋体",
        entry_ascii="Times New Roman",
        entry_size_half_pt=21,
        max_level=3,
    )

    header_text = format_prospectus_draft_header_text((model.get("cover_lines") or ["招募说明书"])[0])

    cover_lines = model.get("cover_lines") or []
    if cover_lines:
        for _ in range(4):
            add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        add_paragraph(cover_lines[0], align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, eastasia="\u5b8b\u4f53", ascii_font="\u5b8b\u4f53", size_half_pt=44, bold=True)
        for _ in range(6):
            add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        for line in cover_lines[1:]:
            add_cover_line(line)

    if cover_lines:
        doc.add_section(WD_SECTION.NEW_PAGE)

    if model.get("important_notice_blocks"):
        add_paragraph(model.get("important_notice_title") or "\u91cd\u8981\u63d0\u793a", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, eastasia="\u5b8b\u4f53", ascii_font="\u5b8b\u4f53", size_half_pt=21, bold=True)
        for block in model["important_notice_blocks"]:
            add_paragraph(block, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True)
        doc.add_page_break()

    add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_paragraph(
        model.get("toc_title") or "\u76ee\u5f55",
        style="TOC Heading",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
        eastasia="\u5b8b\u4f53",
        ascii_font="\u5b8b\u4f53",
        size_half_pt=28,
        bold=False,
    )
    toc_field_paragraph = doc.add_paragraph()
    set_paragraph_format(toc_field_paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    _append_word_toc_field(
        toc_field_paragraph,
        OxmlElement,
        qn,
        levels="2-2",
        placeholder="右键更新目录",
        ascii_font="Times New Roman",
        eastasia_font="\u5b8b\u4f53",
        size=21,
    )

    chapters = model.get("chapters", [])
    if chapters:
        doc.add_page_break()

        for chapter_index, chapter in enumerate(chapters):
            add_chapter_heading(chapter, chapter_index)
            add_paragraph("", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)

            blocks = chapter.get("blocks", [])
            table_buffer = []
            seen_top_level_section = False
            for block in blocks:
                if isinstance(block, dict) and block.get("type") == "formula":
                    if table_buffer:
                        render_table_block(table_buffer)
                        table_buffer = []
                    add_native_formula_paragraph(block)
                    continue

                if isinstance(block, dict) and block.get("type") == "table":
                    if table_buffer:
                        render_table_block(table_buffer)
                        table_buffer = []
                    if block.get("table_kind") == "purchase_redemption_list_format" and block.get("table_xml"):
                        insert_reference_table_block(block)
                    else:
                        title_text = str(block.get("title") or "").strip()
                        if title_text:
                            add_paragraph(
                                title_text,
                                align=WD_ALIGN_PARAGRAPH.LEFT,
                                first_line=is_section7_format_title(title_text),
                                bold=False,
                            )
                        render_table_rows(
                            [list(row) for row in (block.get("rows") or [])],
                            bordered=bool(block.get("bordered")),
                        )
                    continue

                block_text = (block.get("text", "") if isinstance(block, dict) else str(block)).strip()
                if not block_text:
                    continue
                if block_text.startswith("|") and block_text.endswith("|"):
                    table_buffer.append(block_text)
                    continue
                if table_buffer:
                    render_table_block(
                        table_buffer,
                        bordered=(self._table_kind_from_rows(self._markdown_table_lines_to_rows(table_buffer)) == "subscription_fee"),
                    )
                    table_buffer = []
                if re.match(r"^[一二三四五六七八九十百]+\u3001", block_text):
                    if seen_top_level_section:
                        add_paragraph("", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
                    seen_top_level_section = True

                add_paragraph(
                    block_text,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line=True,
                    bold=False,
                )
            if table_buffer:
                render_table_block(
                    table_buffer,
                    bordered=(self._table_kind_from_rows(self._markdown_table_lines_to_rows(table_buffer)) == "subscription_fee"),
                )
            if chapter_index < len(chapters) - 1:
                doc.add_page_break()

    body_sections = list(doc.sections[1:]) if len(doc.sections) > 1 else list(doc.sections)
    update_header_text(header_text, body_sections)

    _finalize_doc_page_numbers(doc, OxmlElement, qn, body_start_index=1)

    for table in doc.tables:
        ensure_table_borders(table)

    def finalize_footer_parts(docx_bytes: bytes) -> bytes:
        source = io.BytesIO(docx_bytes)
        target = io.BytesIO()
        with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zout:
            canonical_footer = None
            for info in zin.infolist():
                if info.filename.startswith("word/footer"):
                    footer_xml = zin.read(info.filename)
                    if b"PAGE" in footer_xml:
                        canonical_footer = footer_xml
                        break

            for info in zin.infolist():
                payload = zin.read(info.filename)
                if info.filename == "word/document.xml":
                    document_xml = payload.decode("utf-8")

                    def strip_cover_refs(match):
                        sect_xml = match.group(0)
                        return re.sub(r"<w:(?:headerReference|footerReference)\b[^>]*/>", "", sect_xml)

                    document_xml = re.sub(
                        r"<w:sectPr\b[\s\S]*?</w:sectPr>",
                        strip_cover_refs,
                        document_xml,
                        count=1,
                    )
                    payload = document_xml.encode("utf-8")
                elif info.filename.startswith("word/footer") and canonical_footer is not None and b"PAGE" not in payload:
                    payload = canonical_footer
                zout.writestr(info, payload)
        return target.getvalue()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return finalize_footer_parts(buf.read())


class ProductSummaryEngine(ProspectusEngine):
    DEFAULT_MANAGER_NAME = "南方基金管理股份有限公司"
    DEFAULT_SERVICE_HOTLINE = "400-889-8899"
    DEFAULT_WEBSITE = "www.nffund.com"
    DEFAULT_PLACEHOLDER_CN_DATE = "202X年X月X日"
    DEFAULT_ON_EXCHANGE_SHORT_NAME_TEXT = "场内简称：XXXETF南方"

    @staticmethod
    def _format_cn_date(date_value: datetime | None = None) -> str:
        value = date_value or datetime.now()
        return f"{value.year}年{value.month}月{value.day}日"

    @staticmethod
    def _clean_text(value, default: str = "-") -> str:
        text = str(value or "").strip()
        return text or default

    @classmethod
    def _pick_first_value(cls, data: dict, keys: list[str], default: str = "-") -> str:
        for key in keys:
            if key in data:
                text = cls._clean_text(data.get(key), default="")
                if text:
                    return text
        return default

    @staticmethod
    def _split_text_values(value) -> list[str]:
        text = str(value or "").strip()
        if not text:
            return []
        parts = re.split(r"[\n；;、,，]+", text)
        return [part.strip() for part in parts if str(part).strip()]

    def _require_generated_prospectus(self, form_data: dict) -> tuple[str, dict]:
        prospectus_text = str(form_data.get("PROSPECTUS_TEXT") or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not prospectus_text:
            raise ValueError("请先生成招募说明书，再生成产品资料概要。")
        render_model = form_data.get("PROSPECTUS_RENDER_MODEL")
        if not isinstance(render_model, dict) or not render_model.get("chapters"):
            render_model = self._build_render_model_from_display_text(prospectus_text, form_data)
            total_blocks = sum(len(chapter.get("blocks") or []) for chapter in render_model.get("chapters", []))
            if not render_model.get("chapters") or total_blocks == 0:
                formatted_text = self._format_reference_style_prospectus(prospectus_text)
                formatted_text = self._normalize_all_chapter_body_spacing(formatted_text)
                formatted_text = re.sub(r"\n{3,}", "\n\n", formatted_text).strip()
                formatted_model = self._build_render_model_from_display_text(formatted_text, form_data)
                formatted_blocks = sum(len(chapter.get("blocks") or []) for chapter in formatted_model.get("chapters", []))
                if formatted_model.get("chapters") and formatted_blocks > 0:
                    prospectus_text = formatted_text
                    render_model = formatted_model
        return prospectus_text, render_model

    @staticmethod
    def _chapter_by_title(render_model: dict, title: str) -> dict | None:
        for chapter in render_model.get("chapters") or []:
            if str(chapter.get("title") or "").strip() == title:
                return chapter
        return None

    @staticmethod
    def _chapter_blocks_to_lines(chapter: dict | None) -> list[str]:
        if not isinstance(chapter, dict):
            return []
        lines = []
        for block in chapter.get("blocks") or []:
            if not isinstance(block, dict):
                continue
            block_type = str(block.get("type") or "")
            if block_type == "paragraph":
                text = str(block.get("text") or "").strip()
                if text:
                    lines.append(text)
            elif block_type == "formula":
                text = str(block.get("plain_text") or "").strip()
                if text:
                    lines.append(text)
        return lines

    @classmethod
    def _extract_chapter_sections(cls, chapter: dict | None) -> dict[str, list[str]]:
        section_re = re.compile(r"^[一二三四五六七八九十百]+、")
        sections: dict[str, list[str]] = {}
        current_heading = ""
        for line in cls._chapter_blocks_to_lines(chapter):
            if section_re.match(line):
                current_heading = line
                sections.setdefault(current_heading, [])
                continue
            if current_heading:
                sections.setdefault(current_heading, []).append(line)
        return sections

    @staticmethod
    def _match_section_text(sections: dict[str, list[str]], prefix: str, default: str = "暂无。") -> str:
        for heading, lines in sections.items():
            if str(heading).startswith(prefix):
                text = "\n".join(str(line).strip() for line in lines if str(line).strip()).strip()
                return text or default
        return default

    @staticmethod
    def _match_section_lines(sections: dict[str, list[str]], prefix: str) -> list[str]:
        for heading, lines in sections.items():
            if str(heading).startswith(prefix):
                return [str(line).strip() for line in lines if str(line).strip()]
        return []

    @classmethod
    def _find_heading_lines(
        cls,
        render_model: dict | None,
        heading_keyword: str,
        *,
        chapter_title: str | None = None,
    ) -> list[str]:
        section_re = re.compile(r"^[一二三四五六七八九十百]+、")
        for chapter in (render_model or {}).get("chapters") or []:
            if chapter_title and str(chapter.get("title") or "").strip() != chapter_title:
                continue
            lines = cls._chapter_blocks_to_lines(chapter)
            for index, line in enumerate(lines):
                heading = str(line).strip()
                if not heading:
                    continue
                if heading == heading_keyword or heading.endswith(heading_keyword) or heading.endswith(f"、{heading_keyword}"):
                    collected: list[str] = []
                    for next_line in lines[index + 1:]:
                        text = str(next_line).strip()
                        if not text:
                            continue
                        if section_re.match(text):
                            break
                        collected.append(text)
                    if collected:
                        return collected
        return []

    @staticmethod
    def _normalize_rate_text(value, default: str) -> str:
        text = str(value or "").strip()
        if not text:
            return default
        return text if text.endswith("%") else f"{text}%"

    def _extract_risk_return_text(self, prospectus_render_model: dict | None) -> str:
        lines = self._find_heading_lines(prospectus_render_model, "风险收益特征", chapter_title="基金的投资")
        return "\n".join(lines).strip()

    def _extract_investment_range_text(self, prospectus_render_model: dict | None) -> str:
        lines = self._find_heading_lines(prospectus_render_model, "投资范围", chapter_title="基金的投资")
        return "\n".join(lines).strip()

    def _extract_other_fees_text(self, prospectus_render_model: dict | None) -> str:
        chapter = self._chapter_by_title(prospectus_render_model or {}, "基金的费用与税收")
        sections = self._extract_chapter_sections(chapter)
        lines = self._match_section_lines(sections, "一、")
        cleaned_parts = []
        for line in lines:
            text = re.sub(r"^\s*\d+\s*[、.．)]\s*", "", str(line).strip())
            if not text:
                continue
            if any(keyword in text for keyword in ("管理费", "托管费", "销售服务费")):
                continue
            cleaned_parts.append(text.rstrip("；;。"))
        if not cleaned_parts:
            return ""
        return "；".join(cleaned_parts) + "。"

    def _extract_special_risk_section(self, prospectus_render_model: dict | None) -> tuple[str, list[str]]:
        chapter = self._chapter_by_title(prospectus_render_model or {}, "风险揭示")
        sections = self._extract_chapter_sections(chapter)
        for heading, lines in sections.items():
            if not str(heading).startswith("四、"):
                continue
            normalized_heading = str(heading).split("、", 1)[1].strip() if "、" in str(heading) else str(heading).strip()
            normalized_heading = normalized_heading or "本基金特有的风险"
            normalized_lines = [str(line).strip() for line in lines if str(line).strip()]
            return f"一）{normalized_heading}", normalized_lines
        return "一）本基金特有的风险", []

    @staticmethod
    def _extract_risk_title_before_punctuation(text: str) -> tuple[str, str]:
        raw = str(text or "").strip()
        if not raw:
            return "", ""
        parts = re.split(r"[：:。；;]\s*", raw, maxsplit=1)
        title = str(parts[0] or "").strip()
        remainder = str(parts[1] or "").strip() if len(parts) > 1 else ""
        return title, remainder

    @classmethod
    def _extract_inline_nested_risk_titles(cls, text: str) -> list[str]:
        titles = []
        for match in re.finditer(r"[（(](\d+)[)）]\s*([^（(\n。；;：:]+?风险)", str(text or "")):
            titles.append(f"（{match.group(1)}）{match.group(2).strip()}")
        return titles

    @staticmethod
    def _fixed_market_risk_intro() -> str:
        return "证券市场价格受到经济因素、政治因素、投资心理和交易制度等各种因素的影响，导致基金收益水平变化，产生风险，主要包括："

    def _extract_market_risk_section(self, prospectus_render_model: dict | None) -> tuple[str, list[str]]:
        chapter = self._chapter_by_title(prospectus_render_model or {}, "风险揭示")
        sections = self._extract_chapter_sections(chapter)
        market_lines = self._match_section_lines(sections, "一、")
        intro = self._fixed_market_risk_intro()
        if not market_lines:
            return "二）市场风险", [intro]

        items: list[str] = []
        nested_parent_titles = {"投资股指期货的风险", "本基金参与转融通证券出借业务的风险"}
        current_item: dict | None = None

        def flush_current():
            nonlocal current_item
            if not current_item:
                return
            prefix = str(current_item.get("prefix") or "").strip()
            nested_titles = list(current_item.get("nested") or [])
            if prefix:
                if nested_titles:
                    items.append(prefix + "：" + "".join(nested_titles) + "。")
                else:
                    items.append(prefix + "。")
            current_item = None

        for line in market_lines:
            text = str(line or "").strip()
            if not text or text == intro:
                continue

            top_match = re.match(r"^\s*(\d+)\s*[、.．)]\s*(.+)$", text)
            if top_match:
                flush_current()
                number = top_match.group(1)
                raw_title, remainder = self._extract_risk_title_before_punctuation(top_match.group(2))
                if raw_title and "风险" in raw_title:
                    current_item = {
                        "prefix": f"{number}、{raw_title}",
                        "nested": [],
                        "capture_nested": raw_title in nested_parent_titles,
                    }
                    if current_item["capture_nested"]:
                        current_item["nested"].extend(self._extract_inline_nested_risk_titles(remainder))
                continue

            nested_match = re.match(r"^\s*[（(](\d+)[)）]\s*(.+)$", text)
            if nested_match and current_item and current_item.get("capture_nested"):
                raw_title, _remainder = self._extract_risk_title_before_punctuation(nested_match.group(2))
                if raw_title and "风险" in raw_title:
                    current_item["nested"].append(f"（{nested_match.group(1)}）{raw_title}")

        flush_current()
        return "二）市场风险", [intro + "".join(items)]

    def _template_risk_lines(self) -> list[str]:
        cache = getattr(self, "_product_summary_template_risk_lines_cache", None)
        if isinstance(cache, list):
            return list(cache)

        from docx import Document

        template_doc = Document(str(_resolve_product_summary_template_docx()))
        lines: list[str] = []
        capture = False
        for paragraph in template_doc.paragraphs:
            text = str(paragraph.text or "").strip()
            if not text:
                continue
            if text == "（一） 风险揭示":
                capture = True
            if capture:
                if text == "（二） 重要提示":
                    break
                lines.append(text)

        self._product_summary_template_risk_lines_cache = list(lines)
        return list(lines)

    def _template_important_notice_lines(self) -> list[str]:
        cache = getattr(self, "_product_summary_template_important_notice_lines_cache", None)
        if isinstance(cache, list):
            return list(cache)

        from docx import Document

        template_doc = Document(str(_resolve_product_summary_template_docx()))
        lines: list[str] = []
        capture = False
        for paragraph in template_doc.paragraphs:
            text = str(paragraph.text or "").strip()
            if not text:
                continue
            if text == "（二） 重要提示":
                capture = True
                continue
            if capture:
                if text == "五、其他资料查询方式":
                    break
                lines.append(text)

        self._product_summary_template_important_notice_lines_cache = list(lines)
        return list(lines)

    @staticmethod
    def _product_summary_dispute_resolution_sentence(form_data: dict) -> str:
        _, dispute_config = _resolve_dispute_resolution_venue(
            (form_data or {}).get("DISPUTE_RESOLUTION_VENUE") or (form_data or {}).get("DISPUTE_RESOLUTION_PLACE")
        )
        return dispute_config["product_summary_sentence"]

    @classmethod
    def _apply_dispute_resolution_to_notice_lines(cls, lines: list[str], form_data: dict) -> list[str]:
        dispute_sentence = cls._product_summary_dispute_resolution_sentence(form_data)
        result: list[str] = []
        for line in lines:
            text = str(line or "")
            if (
                "争议解决方式为仲裁" in text
                or "深圳国际仲裁院" in text
                or "中国国际经济贸易仲裁委员会" in text
            ):
                result.append(dispute_sentence)
            else:
                result.append(text)
        return result

    @staticmethod
    def _rows_to_pipe_lines(rows: list[list[str]]) -> list[str]:
        return [f"|{'|'.join(str(cell or '').strip() for cell in row)}|" for row in rows if row]

    def _manager_entries(self, form_data: dict) -> list[dict[str, str]]:
        names = self._split_text_values(form_data.get("PRODUCT_SUMMARY_MANAGER_NAMES") or form_data.get("FUND_MANAGER_NAME"))
        start_dates = self._split_text_values(form_data.get("PRODUCT_SUMMARY_MANAGER_START_DATES") or form_data.get("FUND_MANAGER_START_DATE"))
        if not names:
            names = [self._pick_first_value(form_data, ["FUND_MANAGER_NAME"], default="-")]
        entries = []
        for index, name in enumerate(names):
            entries.append(
                {
                    "name": name or "-",
                    "start_date": start_dates[index] if index < len(start_dates) else "",
                    "security_date": self.DEFAULT_PLACEHOLDER_CN_DATE,
                }
            )
        return entries

    def _build_overview_table_block(self, v: dict, form_data: dict) -> dict:
        fund_short_name = self._pick_first_value(form_data, ["FUND_SHORT_NAME", "FUND_ABBR", "FUND_NAME"], default="-")
        fund_code = self._pick_first_value(form_data, ["FUND_CODE", "ETF_CODE", "TRADE_CODE"], default="-")
        effective_date = self._pick_first_value(form_data, ["FUND_EFFECTIVE_DATE", "CONTRACT_EFFECTIVE_DATE"], default="-")
        listing_date = self._pick_first_value(form_data, ["LISTING_DATE"], default="-")
        fund_type = self._pick_first_value(form_data, ["FUND_TYPE"], default="股票型")
        trading_currency = self._pick_first_value(form_data, ["TRADING_CURRENCY", "CURRENCY"], default="人民币")
        operation_mode = self._pick_first_value(form_data, ["OPERATION_MODE"], default="普通开放式")
        frequency = self._pick_first_value(form_data, ["OPEN_FREQUENCY"], default="每个开放日")
        other_text = self.DEFAULT_ON_EXCHANGE_SHORT_NAME_TEXT
        rows = [
            ["基金简称", fund_short_name, "基金代码", fund_code],
            [
                "基金管理人",
                self._pick_first_value(v, ["FUND_MANAGER_COMPANY", "MANAGER_NAME"], default=self.DEFAULT_MANAGER_NAME),
                "基金托管人",
                self._pick_first_value(v, ["CUSTODIAN_NAME"], default="-"),
            ],
            ["基金合同生效日", effective_date, "上市交易所", self._pick_first_value(v, ["EXCHANGE_NAME_CN"], default="-")],
            ["基金合同生效日", effective_date, "上市日期", listing_date],
            ["基金类型", fund_type, "交易币种", trading_currency],
            ["运作方式", operation_mode, "", ""],
            ["开放频率", frequency, "", ""],
        ]
        for entry in self._manager_entries(form_data):
            rows.append(["基金经理", entry["name"], "开始担任本基金基金经理的日期", entry["start_date"]])
            rows.append(["基金经理", entry["name"], "证券从业日期", entry["security_date"]])
        rows.append(["其他", other_text, "", ""])
        merge_spans = [
            {"row": 5, "start": 1, "end": 3},
            {"row": 6, "start": 1, "end": 3},
            {"row": len(rows) - 1, "start": 1, "end": 3},
        ]
        return self._table_block_from_rows(
            rows,
            table_kind="product_summary_overview",
            bordered=True,
            merge_spans=merge_spans,
            column_widths=[1600, 2600, 1600, 2600],
        )

    @staticmethod
    def _product_summary_variant_key(form_data: dict) -> str:
        market_type = str(form_data.get("MARKET_TYPE") or "").strip().upper()
        exchange = str(form_data.get("EXCHANGE") or "").strip().upper()
        if market_type == "HK_CONNECT":
            return f"{exchange}_HK_CONNECT" if exchange in {"SSE", "SZSE"} else "HK_CONNECT"
        if market_type == "CHUANGYE":
            return "SZSE_CHUANGYE"
        if market_type == "KECHUANG":
            return "SSE_KECHUANG"
        if exchange == "SZSE":
            return "SZSE_A_SHARE"
        return "SSE_A_SHARE"

    def _build_investment_table_rows(self, form_data: dict, prospectus_render_model: dict | None = None) -> list[list[str]]:
        index_name = self._pick_first_value(form_data, ["INDEX_NAME"], default="标的指数")
        variant = self._product_summary_variant_key(form_data)
        objective = "紧密跟踪标的指数，追求跟踪偏离度和跟踪误差最小化。"
        configs = {
            "SSE_A_SHARE": {
                "range": (
                    "本基金主要投资于标的指数成份股、备选成份股（含存托凭证）。为更好地实现投资目标，本基金可少量投资于非成份股（包含主板、创业板、科创板及其他经中国证监会核准或注册发行的股票、存托凭证）、金融衍生品（股指期货、股票期权等）、债券（包括国债、金融债、企业债、公司债、政府机构债券、地方政府债券、次级债、可转换债券、可交换债券、央行票据、中期票据、短期融资券、超短期融资券等）、资产支持证券、债券回购、银行存款（包括协议存款、定期存款及其他银行存款）、同业存单、货币市场工具以及中国证监会允许基金投资的其他金融工具（但须符合中国证监会的相关规定）。\n"
                    "本基金根据相关规定可参与融资、转融通证券出借业务。\n"
                    "如法律法规或监管机构以后允许基金投资其他品种，基金管理人在履行适当程序后，可以将其纳入投资范围。\n"
                    "在建仓完成后，本基金投资于标的指数成份股、备选成份股的资产比例不低于基金资产净值的90%，且不低于非现金基金资产的80%，因法律法规的规定而受限制的情形除外。\n"
                    "如法律法规或监管机构变更投资品种的投资比例限制，基金管理人在履行适当程序后，可以调整上述投资品种的投资比例。\n"
                    f"本基金标的指数为{index_name}。"
                ),
                "strategy": (
                    "本基金主要采用完全复制策略、替代策略及其他适当的策略以更好地跟踪标的指数，实现基金投资目标。"
                    "本基金力争日均跟踪偏离度的绝对值不超过0.2%，年跟踪误差不超过2%。"
                    "主要投资策略包括：完全复制策略、替代策略、金融衍生品投资策略、债券投资策略、可转换债券及可交换债券投资策略、资产支持证券投资策略、融资及转融通证券出借业务投资策略、存托凭证投资策略等。"
                ),
                "benchmark": f"{index_name}收益率",
                "risk_return": "本基金属于股票型基金，一般而言，其长期平均风险和预期收益水平高于混合型基金、债券型基金与货币市场基金。本基金采用完全复制法跟踪标的指数的表现，具有与标的指数、以及标的指数所代表的股票市场相似的风险收益特征。",
            },
            "SZSE_A_SHARE": {
                "range": (
                    "本基金主要投资于标的指数成份股、备选成份股（含存托凭证）。为更好地实现投资目标，本基金可少量投资于非成份股（包含主板、创业板、科创板及其他经中国证监会核准或注册发行的股票、存托凭证）、金融衍生品（股指期货、股票期权等）、债券（包括国债、金融债、企业债、公司债、政府机构债券、地方政府债券、次级债、可转换债券、可交换债券、央行票据、中期票据、短期融资券、超短期融资券等）、资产支持证券、债券回购、银行存款（包括协议存款、定期存款及其他银行存款）、同业存单、货币市场工具以及中国证监会允许基金投资的其他金融工具（但须符合中国证监会的相关规定）。\n"
                    "本基金根据相关规定可参与融资、转融通证券出借业务。\n"
                    "如法律法规或监管机构以后允许基金投资其他品种，基金管理人在履行适当程序后，可以将其纳入投资范围。\n"
                    "在建仓完成后，本基金投资于标的指数成份股、备选成份股的资产比例不低于基金资产净值的90%，且不低于非现金基金资产的80%，因法律法规的规定而受限制的情形除外。\n"
                    "如法律法规或监管机构变更投资品种的投资比例限制，基金管理人在履行适当程序后，可以调整上述投资品种的投资比例。\n"
                    f"本基金标的指数为{index_name}。"
                ),
                "strategy": (
                    "本基金主要采用完全复制策略、替代策略及其他适当的策略以更好地跟踪标的指数，实现基金投资目标。"
                    "本基金力争日均跟踪偏离度的绝对值不超过0.2%，年跟踪误差不超过2%。"
                    "主要投资策略包括：完全复制策略、替代策略、金融衍生品投资策略、债券投资策略、可转换债券及可交换债券投资策略、资产支持证券投资策略、融资及转融通证券出借业务投资策略、存托凭证投资策略等。"
                ),
                "benchmark": f"本基金的业绩比较基准为标的指数收益率。本基金标的指数为{index_name}。",
                "risk_return": "本基金属于股票型基金，一般而言，其长期平均风险和预期收益水平高于混合型基金、债券型基金与货币市场基金。本基金采用完全复制法跟踪标的指数的表现，具有与标的指数、以及标的指数所代表的股票市场相似的风险收益特征。",
            },
            "SZSE_CHUANGYE": {
                "range": (
                    "本基金主要投资于标的指数成份股、备选成份股（含存托凭证）。为更好地实现投资目标，本基金可少量投资于非成份股（包含主板、创业板及其他经中国证监会核准或注册发行的股票、存托凭证）、金融衍生品（股指期货、股票期权等）、债券（包括国债、金融债、企业债、公司债、政府机构债券、地方政府债券、次级债、可转换债券、可交换债券、央行票据、中期票据、短期融资券、超短期融资券等）、资产支持证券、债券回购、银行存款（包括协议存款、定期存款及其他银行存款）、同业存单、货币市场工具以及中国证监会允许基金投资的其他金融工具（但须符合中国证监会的相关规定）。\n"
                    "本基金根据相关规定可参与融资、转融通证券出借业务。\n"
                    "如法律法规或监管机构以后允许基金投资其他品种，基金管理人在履行适当程序后，可以将其纳入投资范围。\n"
                    "在建仓完成后，本基金投资于标的指数成份股、备选成份股的资产比例不低于基金资产净值的90%，且不低于非现金基金资产的80%，因法律法规的规定而受限制的情形除外。\n"
                    "如法律法规或监管机构变更投资品种的投资比例限制，基金管理人在履行适当程序后，可以调整上述投资品种的投资比例。"
                ),
                "strategy": (
                    "本基金主要采用完全复制策略、替代策略及其他适当的策略以更好地跟踪标的指数，实现基金投资目标。"
                    "本基金力争日均跟踪偏离度的绝对值不超过0.2%，年跟踪误差不超过2%。"
                    "主要投资策略包括：完全复制策略、替代策略、金融衍生品投资策略、债券投资策略、可转换债券及可交换债券投资策略、资产支持证券投资策略、融资及转融通证券出借业务投资策略、存托凭证投资策略等。"
                ),
                "benchmark": f"本基金的业绩比较基准为标的指数收益率。本基金标的指数为{index_name}，及其未来可能发生的变更。",
                "risk_return": "本基金属于股票型基金，一般而言，其长期平均风险和预期收益水平高于混合型基金、债券型基金与货币市场基金。本基金采用完全复制法跟踪标的指数的表现，具有与标的指数、以及标的指数所代表的股票市场相似的风险收益特征。本基金投资创业板股票，会面临创业板机制下因投资标的、市场制度以及交易规则等差异带来的特有风险，包括流动性风险、退市风险和投资集中风险等。",
            },
            "SSE_KECHUANG": {
                "range": (
                    "本基金主要投资于标的指数成份股、备选成份股（含存托凭证）。为更好地实现投资目标，本基金可少量投资于非成份股（包含主板、科创板、创业板及其他经中国证监会核准或注册发行的股票、存托凭证）、金融衍生品（股指期货、股票期权等）、债券（包括国债、金融债、企业债、公司债、政府机构债券、地方政府债券、次级债、可转换债券、可交换债券、央行票据、中期票据、短期融资券、超短期融资券等）、资产支持证券、债券回购、银行存款（包括协议存款、定期存款及其他银行存款）、同业存单、货币市场工具以及中国证监会允许基金投资的其他金融工具（但须符合中国证监会的相关规定）。\n"
                    "本基金根据相关规定可参与融资、转融通证券出借业务。\n"
                    "如法律法规或监管机构以后允许基金投资其他品种，基金管理人在履行适当程序后，可以将其纳入投资范围。\n"
                    "在建仓完成后，本基金投资于标的指数成份股、备选成份股的资产比例不低于基金资产净值的90%，且不低于非现金基金资产的80%，因法律法规的规定而受限制的情形除外。\n"
                    "如法律法规或监管机构变更投资品种的投资比例限制，基金管理人在履行适当程序后，可以调整上述投资品种的投资比例。\n"
                    f"本基金标的指数为{index_name}。"
                ),
                "strategy": (
                    "本基金主要采用完全复制策略、替代策略及其他适当的策略以更好的跟踪标的指数，实现基金投资目标。"
                    "本基金力争日均跟踪偏离度的绝对值不超过0.2%，年跟踪误差不超过2%。"
                    "主要投资策略包括：完全复制策略、替代策略、金融衍生品投资策略、债券投资策略、可转换债券及可交换债券投资策略、资产支持证券投资策略、融资及转融通证券出借业务投资策略、存托凭证投资策略等。"
                ),
                "benchmark": f"本基金的业绩比较基准为标的指数收益率。本基金标的指数为{index_name}。",
                "risk_return": "本基金属于股票型基金，一般而言，其长期平均风险和预期收益水平高于混合型基金、债券型基金与货币市场基金。本基金采用完全复制法跟踪标的指数的表现，具有与标的指数以及标的指数所代表的股票市场相似的风险收益特征。本基金投资于科创板股票，会面临科创板机制下因投资标的、市场制度以及交易规则等差异带来的特有风险，包括股价波动风险、流动性风险、退市风险和投资集中风险等。",
            },
            "SSE_HK_CONNECT": {
                "range": (
                    "本基金主要投资于标的指数成份股、备选成份股（包括内地与香港股票市场交易互联互通机制允许买卖的规定范围内的香港联合交易所上市的股票（简称“港股通股票”）、存托凭证，下同）。为更好地实现投资目标，本基金可少量投资于非成份股（包含主板、创业板、科创板及其他经中国证监会核准或注册发行的股票、港股通股票、存托凭证）、金融衍生品（股指期货、股票期权等）、债券（包括国债、金融债、企业债、公司债、政府机构债券、地方政府债券、次级债、可转换债券、可交换债券、央行票据、中期票据、短期融资券、超短期融资券等）、资产支持证券、债券回购、银行存款（包括协议存款、定期存款及其他银行存款）、同业存单、货币市场工具以及中国证监会允许基金投资的其他金融工具（但须符合中国证监会的相关规定）。\n"
                    "本基金根据相关规定可参与融资、转融通证券出借业务。\n"
                    "如法律法规或监管机构以后允许基金投资其他品种，基金管理人在履行适当程序后，可以将其纳入投资范围。\n"
                    "在建仓完成后，本基金投资于标的指数成份股、备选成份股的资产比例不低于基金资产净值的90%，且不低于非现金基金资产的80%，因法律法规的规定而受限制的情形除外。\n"
                    "如法律法规或监管机构变更投资品种的投资比例限制，基金管理人在履行适当程序后，可以调整上述投资品种的投资比例。\n"
                    f"本基金标的指数为{index_name}。"
                ),
                "strategy": (
                    "本基金主要采用完全复制策略、替代策略及其他适当的策略以更好的跟踪标的指数，实现基金投资目标。"
                    "本基金力争日均跟踪偏离度的绝对值不超过0.35%，年跟踪误差不超过4%。"
                    "主要投资策略包括：完全复制策略、替代策略、金融衍生品投资策略、债券投资策略、可转换债券及可交换债券投资策略、资产支持证券投资策略、融资及转融通证券出借业务投资策略、存托凭证投资策略等。"
                ),
                "benchmark": (
                    f"本基金的业绩比较基准为{index_name}（经估值汇率调整后）收益率。本基金标的指数为{index_name}。"
                    if "50" in index_name
                    else f"本基金的业绩比较基准为{index_name}（人民币）收益率。本基金标的指数为{index_name}。"
                ),
                "risk_return": "本基金属于股票型基金，一般而言，其长期平均风险和预期收益水平高于混合型基金、债券型基金与货币市场基金。本基金主要采用完全复制法跟踪标的指数的表现，具有与标的指数以及标的指数所代表的市场相似的风险收益特征。\n本基金投资港股通股票，除了需要承担与境内证券投资基金类似的市场波动风险等一般投资风险之外，本基金还面临汇率风险和港股通机制下因投资环境、投资标的、市场制度以及交易规则等差异带来的特有风险。",
            },
            "SZSE_HK_CONNECT": {
                "range": (
                    "本基金主要投资于标的指数成份股、备选成份股（包括内地与香港股票市场交易互联互通机制允许买卖的规定范围内的香港联合交易所上市的股票（简称“港股通股票”）、存托凭证，下同）。为更好地实现投资目标，本基金可少量投资于非成份股（包含主板、创业板、科创板及其他经中国证监会核准或注册发行的股票、港股通股票、存托凭证）、金融衍生品（股指期货、股票期权等）、债券（包括国债、金融债、企业债、公司债、政府机构债券、地方政府债券、次级债、可转换债券、可交换债券、央行票据、中期票据、短期融资券、超短期融资券等）、资产支持证券、债券回购、银行存款（包括协议存款、定期存款及其他银行存款）、同业存单、货币市场工具以及中国证监会允许基金投资的其他金融工具（但须符合中国证监会的相关规定）。\n"
                    "本基金根据相关规定可参与融资、转融通证券出借业务。\n"
                    "如法律法规或监管机构以后允许基金投资其他品种，基金管理人在履行适当程序后，可以将其纳入投资范围。\n"
                    "在建仓完成后，本基金投资于标的指数成份股、备选成份股的资产比例不低于基金资产净值的90%，且不低于非现金基金资产的80%，因法律法规的规定而受限制的情形除外。\n"
                    "如法律法规或监管机构变更投资品种的投资比例限制，基金管理人在履行适当程序后，可以调整上述投资品种的投资比例。\n"
                    f"本基金标的指数为{index_name}。"
                ),
                "strategy": (
                    "本基金主要采用完全复制策略、替代策略及其他适当的策略以更好的跟踪标的指数，实现基金投资目标。"
                    "本基金力争日均跟踪偏离度的绝对值不超过0.35%，年跟踪误差不超过4%。"
                    "主要投资策略包括：完全复制策略、替代策略、金融衍生品投资策略、债券投资策略、可转换债券及可交换债券投资策略、资产支持证券投资策略、融资及转融通证券出借业务投资策略、存托凭证投资策略等。"
                ),
                "benchmark": (
                    f"本基金的业绩比较基准为{index_name}（经估值汇率调整后）收益率。本基金标的指数为{index_name}。"
                    if "50" in index_name
                    else f"本基金的业绩比较基准为{index_name}（人民币）收益率。本基金标的指数为{index_name}。"
                ),
                "risk_return": "本基金属于股票型基金，一般而言，其长期平均风险和预期收益水平高于混合型基金、债券型基金与货币市场基金。本基金主要采用完全复制法跟踪标的指数的表现，具有与标的指数以及标的指数所代表的市场相似的风险收益特征。\n本基金投资港股通股票，除了需要承担与境内证券投资基金类似的市场波动风险等一般投资风险之外，本基金还面临汇率风险和港股通机制下因投资环境、投资标的、市场制度以及交易规则等差异带来的特有风险。",
            },
        }
        chosen = configs.get(variant, configs["SSE_A_SHARE"])
        investment_range_text = self._extract_investment_range_text(prospectus_render_model) or chosen["range"]
        risk_return_text = self._extract_risk_return_text(prospectus_render_model) or chosen["risk_return"]
        return [
            ["投资目标", objective],
            ["投资范围", investment_range_text],
            ["主要投资策略", chosen["strategy"]],
            ["业绩比较基准", chosen["benchmark"]],
            ["风险收益特征", risk_return_text],
        ]

    def _build_sales_fee_table_rows(self) -> list[list[str]]:
        return [
            ["费用类型", "份额（S）或金额（M）/持有期限（N）", "收费方式/费率", "备注"],
            ["认购费", "S< 100万份", "0.3%", "-"],
            ["认购费", "100万份≤S", "每笔500元", "-"],
        ]

    def _build_operation_fee_table_rows(self, form_data: dict, prospectus_render_model: dict | None = None) -> list[list[str]]:
        mgmt_fee = self._normalize_rate_text(self._pick_first_value(form_data, ["MGMT_FEE_RATE"], default="0.15"), "0.15%")
        custody_fee = self._normalize_rate_text(self._pick_first_value(form_data, ["CUSTODY_FEE_RATE"], default="0.05"), "0.05%")
        other_fees = self._extract_other_fees_text(prospectus_render_model) or (
            "《基金合同》生效后与基金相关的信息披露费用；《基金合同》生效后与基金相关的会计师费、律师费、公证费、诉讼费和仲裁费；"
            "基金份额持有人大会费用；基金的证券/期货/期权交易费用；基金的银行汇划费用；基金上市费及年费；"
            "基金相关账户的开户及维护费用；因参与融资及转融通证券出借业务而产生的各项合理费用；按照国家有关规定和《基金合同》约定，可以在基金财产中列支的其他费用。"
        )
        return [
            ["费用类别", "收费方式/年费率或金额", "收取方"],
            ["管理费", mgmt_fee, "基金管理人和销售机构"],
            ["托管费", custody_fee, "基金托管人"],
            ["销售服务费", "-", "销售机构"],
            ["其他费用", other_fees, other_fees],
        ]

    @staticmethod
    def _normalize_note_text(text: str) -> str:
        return re.sub(r"^\s*注[:：]\s*", "", str(text or "").strip())

    def _build_operation_fee_note_lines(self, form_data: dict) -> list[str]:
        exchange = str(form_data.get("EXCHANGE") or "").strip().upper()
        if exchange == "SSE":
            return [
                "注：标的指数许可使用费应当由基金管理人承担，不得从基金财产中列支。",
                "本基金交易证券等产生的费用和税负，按实际发生额从基金资产扣除。",
            ]
        return ["注：本基金交易证券、基金等产生的费用和税负，按实际发生额从基金资产扣除。"]

    def _build_risk_blocks(self, form_data: dict, prospectus_render_model: dict | None = None) -> list[dict]:
        special_risk_heading, special_risk_lines = self._extract_special_risk_section(prospectus_render_model)
        market_risk_heading, market_risk_lines = self._extract_market_risk_section(prospectus_render_model)
        template_risk_lines = self._template_risk_lines()

        rendered_risk_lines: list[str] = []
        in_special_risk_body = False
        special_risk_body_inserted = False
        in_market_risk_body = False
        market_risk_body_inserted = False
        for line in template_risk_lines:
            if line.startswith("一）"):
                rendered_risk_lines.append(special_risk_heading)
                in_special_risk_body = True
                if special_risk_lines:
                    rendered_risk_lines.extend(special_risk_lines)
                    special_risk_body_inserted = True
                continue
            if line.startswith("二）市场风险"):
                in_special_risk_body = False
                rendered_risk_lines.append(market_risk_heading)
                in_market_risk_body = True
                if market_risk_lines:
                    rendered_risk_lines.extend(market_risk_lines)
                    market_risk_body_inserted = True
                continue
            if re.match(r"^[三四五六七八九十]）", line):
                in_special_risk_body = False
                in_market_risk_body = False
                rendered_risk_lines.append(line)
                continue
            if in_special_risk_body and special_risk_body_inserted:
                continue
            if in_market_risk_body and market_risk_body_inserted:
                continue
            rendered_risk_lines.append(line)

        blocks = [{"type": "paragraph", "text": line} for line in rendered_risk_lines]
        blocks.append({"type": "paragraph", "text": "（二） 重要提示"})
        blocks.extend({"type": "paragraph", "text": line} for line in self._build_important_notice_lines(form_data))
        return blocks

    def _build_important_notice_lines(self, form_data: dict) -> list[str]:
        fund_name = self._pick_first_value(form_data, ["FUND_NAME"], default="本基金")
        approval_no = self._pick_first_value(form_data, ["CSRC_APPROVAL_NO"], default="202X年X月X日证监许可〔202X〕XXX号")
        template_lines = self._template_important_notice_lines()
        template_lines = self._apply_dispute_resolution_to_notice_lines(template_lines, form_data)
        if not template_lines:
            return [
                f"{fund_name}（以下简称“本基金”）经中国证监会{approval_no}文注册募集。中国证监会对本基金募集的注册，并不表明其对本基金的价值和收益作出实质性判断或保证，也不表明投资于本基金没有风险。",
                "基金管理人依照恪尽职守、诚实信用、谨慎勤勉的原则管理和运用基金财产，但不保证基金一定盈利，也不保证最低收益。",
                "基金投资者自依基金合同取得基金份额，即成为基金份额持有人和基金合同的当事人。",
                self._product_summary_dispute_resolution_sentence(form_data),
                "基金产品资料概要信息发生重大变更的，基金管理人将在三个工作日内更新，其他信息发生变更的，基金管理人每年更新一次。因此，本文件内容相比基金的实际情况可能存在一定的滞后，如需及时、准确获取基金的相关信息，敬请同时关注基金管理人发布的相关临时公告等。",
                "投资人知悉并同意基金管理人可为投资人提供营销信息、资讯与增值服务，并可自主选择退订，具体的服务说明详见招募说明书“基金份额持有人服务”章节。",
            ]

        first_line = template_lines[0]
        first_line = re.sub(
            r"^.+?（以下简称“本基金”）经中国证监会.+?文注册募集。",
            f"{fund_name}（以下简称“本基金”）经中国证监会{approval_no}文注册募集。",
            first_line,
        )
        return [first_line, *template_lines[1:]]

    def _build_query_blocks(self, v: dict, form_data: dict) -> list[dict]:
        fund_name = self._pick_first_value(form_data, ["FUND_NAME"], default="本基金")
        website = self._pick_first_value(form_data, ["COMPANY_WEBSITE", "FUND_MANAGER_WEBSITE"], default=self.DEFAULT_WEBSITE)
        hotline = self._pick_first_value(form_data, ["SERVICE_HOTLINE", "FUND_MANAGER_HOTLINE"], default=self.DEFAULT_SERVICE_HOTLINE)
        lines = [
            f"以下资料详见基金管理人网站[{website}][客服电话：{hotline}]",
            f"●《{fund_name}基金合同》、",
            f"《{fund_name}托管协议》、",
            f"《{fund_name}招募说明书》",
            "●定期报告，包括基金季度报告、中期报告和年度报告",
            "●基金份额净值",
            "●基金销售机构及联系方式",
            "●其他重要资料",
        ]
        return [{"type": "paragraph", "text": line} for line in lines]

    def _build_sections(self, form_data: dict, prospectus_render_model: dict | None = None) -> list[dict]:
        v = self._derive_variables(dict(form_data))
        sections = [
            {
                "title": "一、产品概况",
                "blocks": [
                    self._build_overview_table_block(v, form_data),
                ],
            },
            {
                "title": "二、基金投资与净值表现",
                "blocks": [
                    {"type": "paragraph", "text": "（一） 投资目标与投资策略"},
                    {
                        "type": "paragraph",
                        "text": f"注:详见《{self._pick_first_value(form_data, ['FUND_NAME'], default='本基金')}招募说明书》“基金的投资”部分。",
                        "style_id": "-7",
                    },
                    self._table_block_from_rows(
                        self._build_investment_table_rows(form_data, prospectus_render_model),
                        table_kind="product_summary_investment",
                        bordered=True,
                    ),
                    {"type": "paragraph", "text": "（二） 投资组合资产配置图表/区域配置图表"},
                    {"type": "paragraph", "text": "无"},
                    {"type": "paragraph", "text": "（三） 自基金合同生效以来基金每年的净值增长率及与同期业绩比较基准的比较图"},
                    {"type": "paragraph", "text": "无"},
                ],
            },
            {
                "title": "三、投资本基金涉及的费用",
                "blocks": [
                    {"type": "paragraph", "text": "（一） 基金销售相关费用"},
                    self._table_block_from_rows(
                        self._build_sales_fee_table_rows(),
                        table_kind="product_summary_sales_fee",
                        bordered=True,
                        column_widths=[1800, 2800, 1800, 1400],
                        header_align="center",
                    ),
                    {"type": "paragraph", "text": "基金管理人办理网下现金认购和网下股票认购不收取认购费。发售代理机构办理网上现金认购、网下现金认购、网下股票认购时可参照上述费率结构，按照不高于0.3%的标准收取一定的佣金。"},
                    {"type": "paragraph", "text": "投资人申请重复现金认购的，须按每次认购所对应的费率档次分别计费。"},
                    {"type": "paragraph", "text": "场内交易费用以证券公司实际收取为准。"},
                    {"type": "paragraph", "text": "投资人在申购基金份额时，申购代理券商可按照不超过0.2%的标准收取佣金。"},
                    {"type": "paragraph", "text": "投资人在赎回基金份额时，赎回代理券商可按照不超过0.2%的标准收取佣金。"},
                    {"type": "paragraph", "text": "（二） 基金运作相关费用"},
                    {"type": "paragraph", "text": "以下费用将从基金资产中扣除："},
                    self._table_block_from_rows(
                        self._build_operation_fee_table_rows(form_data, prospectus_render_model),
                        table_kind="product_summary_operation_fee",
                        bordered=True,
                        column_widths=[1800, 3600, 1600],
                    ),
                    *[
                        {"type": "paragraph", "text": line}
                        for line in self._build_operation_fee_note_lines(form_data)
                    ],
                    {"type": "paragraph", "text": "（三） 基金运作综合费用测算"},
                    {"type": "paragraph", "text": "-"},
                ],
            },
            {
                "title": "四、风险揭示与重要提示",
                "blocks": self._build_risk_blocks(form_data, prospectus_render_model),
            },
            {
                "title": "五、其他资料查询方式",
                "blocks": self._build_query_blocks(v, form_data),
            },
            {
                "title": "六、其他情况说明",
                "blocks": [{"type": "paragraph", "text": "暂无。"}],
            },
        ]
        return sections

    def _build_render_model(self, form_data: dict, prospectus_render_model: dict | None = None) -> dict:
        compile_date = self.DEFAULT_PLACEHOLDER_CN_DATE
        fund_name = self._pick_first_value(form_data, ["FUND_NAME"], default="本基金")
        return {
            "cover_lines": [
                f"{fund_name}基金产品资料概要",
                f"编制日期：{compile_date}",
                f"送出日期：{compile_date}",
                "本概要提供本基金的重要信息，是招募说明书的一部分。",
                "作出投资决定前，请阅读完整的招募说明书等销售文件。",
            ],
            "sections": self._build_sections(form_data, prospectus_render_model),
        }

    def _render_model_to_text(self, render_model: dict) -> str:
        lines = [str(line).strip() for line in (render_model.get("cover_lines") or []) if str(line).strip()]
        for section in render_model.get("sections") or []:
            title = str(section.get("title") or "").strip()
            if title:
                lines.append(title)
            for block in section.get("blocks") or []:
                if not isinstance(block, dict):
                    continue
                block_type = str(block.get("type") or "")
                if block_type == "paragraph":
                    text = str(block.get("text") or "").strip()
                    if text:
                        lines.append(text)
                elif block_type == "table":
                    title_text = str(block.get("title") or "").strip()
                    if title_text:
                        lines.append(title_text)
                    lines.extend(self._rows_to_pipe_lines(block.get("rows") or []))
        return "\n".join(lines).strip()

    def generate_bundle(self, form_data: dict) -> dict:
        _, prospectus_render_model = self._require_generated_prospectus(form_data)
        render_model = self._build_render_model(form_data, prospectus_render_model)
        return {
            "text": self._render_model_to_text(render_model),
            "render_model": render_model,
        }

    def generate(self, form_data: dict) -> str:
        return self.generate_bundle(form_data)["text"]

    def build_docx_product_summary(self, render_model: dict, form_data: dict | None = None) -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        form_data = dict(form_data or {})
        _, prospectus_render_model = self._require_generated_prospectus(form_data)
        v = self._derive_variables(dict(form_data))
        template_docx = _resolve_product_summary_template_docx()
        doc = Document(str(template_docx))

        def replace_paragraph_text(paragraph, text: str):
            runs = list(paragraph.runs)
            if runs:
                runs[0].text = str(text or "")
                for run in runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(str(text or ""))

        def replace_cell_lines(cell, lines: list[str]):
            normalized_lines = [str(line or "") for line in lines] or [""]
            paragraphs = list(cell.paragraphs)
            base_style = paragraphs[0].style if paragraphs else None
            base_alignment = paragraphs[0].alignment if paragraphs else WD_ALIGN_PARAGRAPH.LEFT
            for index, paragraph in enumerate(paragraphs):
                replace_paragraph_text(paragraph, normalized_lines[index] if index < len(normalized_lines) else "")
            if len(normalized_lines) > len(paragraphs):
                for line in normalized_lines[len(paragraphs):]:
                    paragraph = cell.add_paragraph()
                    if base_style is not None:
                        paragraph.style = base_style
                    paragraph.alignment = base_alignment
                    replace_paragraph_text(paragraph, line)

        def remove_paragraph(paragraph):
            element = paragraph._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)

        def replace_document_range(start_text: str, stop_text: str, lines: list[str]):
            paragraphs = list(doc.paragraphs)
            start_index = next(
                index for index, paragraph in enumerate(paragraphs)
                if paragraph.text.strip() == start_text
            )
            stop_index = next(
                index for index, paragraph in enumerate(paragraphs[start_index + 1:], start_index + 1)
                if paragraph.text.strip() == stop_text
            )
            existing = paragraphs[start_index + 1:stop_index]
            base_style = existing[0].style if existing else None
            base_alignment = existing[0].alignment if existing else WD_ALIGN_PARAGRAPH.LEFT
            for paragraph in existing:
                remove_paragraph(paragraph)
            anchor = doc.paragraphs[next(
                index for index, paragraph in enumerate(doc.paragraphs)
                if paragraph.text.strip() == stop_text
            )]
            for line in lines:
                paragraph = anchor.insert_paragraph_before(str(line or ""))
                if base_style is not None:
                    paragraph.style = base_style
                paragraph.alignment = base_alignment

        cover_lines = [str(line or "").strip() for line in (render_model.get("cover_lines") or [])]
        if len(cover_lines) >= 5:
            replace_paragraph_text(doc.paragraphs[0], cover_lines[0])
            replace_paragraph_text(doc.paragraphs[2], cover_lines[1])
            replace_paragraph_text(doc.paragraphs[3], cover_lines[2])
            replace_paragraph_text(doc.paragraphs[4], cover_lines[3])
            replace_paragraph_text(doc.paragraphs[5], cover_lines[4])
        if cover_lines:
            for section in doc.sections:
                for header in (section.header, section.first_page_header, section.even_page_header):
                    if header.paragraphs:
                        if header.paragraphs[0].text.strip():
                            replace_paragraph_text(header.paragraphs[0], cover_lines[0])
                    else:
                        header.add_paragraph(cover_lines[0])

        replace_paragraph_text(
            next(paragraph for paragraph in doc.paragraphs if "基金的投资" in paragraph.text and "注:" in paragraph.text),
            f"注:详见《{self._pick_first_value(form_data, ['FUND_NAME'], default='本基金')}招募说明书》“基金的投资”部分。",
        )

        overview_block = self._build_overview_table_block(v, form_data)
        investment_rows = self._build_investment_table_rows(form_data, prospectus_render_model)
        operation_fee_rows = self._build_operation_fee_table_rows(form_data, prospectus_render_model)

        merged_aliases = {
            (int(span["row"]), col_index)
            for span in (overview_block.get("merge_spans") or [])
            for col_index in range(int(span["start"]) + 1, int(span["end"]) + 1)
        }
        for row_index, row in enumerate(overview_block["rows"]):
            for col_index, value in enumerate(row):
                if (row_index, col_index) in merged_aliases:
                    continue
                replace_cell_lines(doc.tables[0].cell(row_index, col_index), [value])

        for row_index, row in enumerate(investment_rows):
            replace_cell_lines(doc.tables[1].cell(row_index, 0), [row[0]])
            replace_cell_lines(doc.tables[1].cell(row_index, 1), str(row[1]).splitlines())

        for row_index, row in enumerate(operation_fee_rows):
            for col_index, value in enumerate(row):
                replace_cell_lines(doc.tables[3].cell(row_index, col_index), [value])

        special_risk_heading, special_risk_lines = self._extract_special_risk_section(prospectus_render_model)
        _market_risk_heading, market_risk_lines = self._extract_market_risk_section(prospectus_render_model)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == "一）本基金特有的风险":
                replace_paragraph_text(paragraph, special_risk_heading)
                break
        if special_risk_lines:
            replace_document_range("一）本基金特有的风险", "二）市场风险", special_risk_lines)
        if market_risk_lines:
            replace_document_range("二）市场风险", "三）管理风险", market_risk_lines)

        replace_document_range(
            "（二） 重要提示",
            "五、其他资料查询方式",
            self._build_important_notice_lines(form_data),
        )
        replace_document_range(
            "五、其他资料查询方式",
            "六、其他情况说明",
            [block["text"] for block in self._build_query_blocks(v, form_data)],
        )

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()


def _sample_contract_form_for_business_text_variant(variant: str) -> dict:
    form = {
        "FUND_NAME": "测试交易型开放式指数证券投资基金",
        "FUND_SHORT_NAME": "测试ETF",
        "INDEX_NAME": "中证测试指数",
        "INDEX_CODE": "000001",
        "INDEX_COMPILER": "中证指数有限公司",
        "CONTRACT_DATE": "2026年4月24日",
        "EXCHANGE": "SSE",
        "MARKET_TYPE": "KECHUANG",
        "CUSTODIAN_NAME": "招商银行股份有限公司",
        "DISTRIBUTION_FREQ": "QUARTERLY",
        "MGMT_FEE_PAYMENT_METHOD": "CONSULT",
        "CUSTODY_FEE_PAYMENT_METHOD": "CONSULT",
        "MGMT_FEE_RATE": "0.15",
        "CUSTODY_FEE_RATE": "0.05",
    }
    if variant == "SSE_HK":
        form.update({"EXCHANGE": "SSE", "MARKET_TYPE": "HK_CONNECT"})
    elif variant == "SZSE_HK":
        form.update({"EXCHANGE": "SZSE", "MARKET_TYPE": "HK_CONNECT"})
    elif variant == "HK_CONNECT":
        form.update({"EXCHANGE": "SSE", "MARKET_TYPE": "HK_CONNECT"})
    return form


def _sample_prospectus_form_for_business_text_variant(variant: str) -> dict:
    form = _sample_contract_form_for_business_text_variant("DEFAULT")
    form.update({
        "PRODUCT_TYPE": "ETF",
        "MARKET_TYPE": "A_SHARE",
        "MARKET_SCOPE": "CROSS_MARKET",
        "INDEX_COMPILER": "中证指数有限公司",
        "INDEX_WEBSITE": "https://www.csindex.com.cn",
    })
    if variant == "SSE_SINGLE":
        form.update({"EXCHANGE": "SSE", "MARKET_TYPE": "KECHUANG"})
    elif variant == "SSE_HK":
        form.update({"EXCHANGE": "SSE", "MARKET_TYPE": "HK_CONNECT"})
    elif variant == "SZSE_CROSS":
        form.update({"EXCHANGE": "SZSE", "MARKET_TYPE": "A_SHARE", "MARKET_SCOPE": "CROSS_MARKET"})
    elif variant == "SZSE_SINGLE":
        form.update({"EXCHANGE": "SZSE", "MARKET_TYPE": "CHUANGYE"})
    elif variant == "SZSE_HK":
        form.update({"EXCHANGE": "SZSE", "MARKET_TYPE": "HK_CONNECT"})
    else:
        form.update({"EXCHANGE": "SSE", "MARKET_TYPE": "A_SHARE", "MARKET_SCOPE": "CROSS_MARKET"})
    return form


def _business_text_matrix(
    default_text: str,
    overrides: dict,
    group: str,
    key: str,
    scene: str,
    *,
    supports_conditions: bool,
) -> dict:
    product_types = _business_text_product_type_values() if supports_conditions else ["ALL"]
    market_types = _business_text_market_type_values() if supports_conditions else ["ALL"]
    exchanges = _business_text_exchange_values() if supports_conditions else ["ALL"]
    matrix = {}
    for product_type in product_types:
        matrix[product_type] = {}
        for market_type in market_types:
            matrix[product_type][market_type] = {}
            for exchange in exchanges:
                matrix[product_type][market_type][exchange] = {
                    "default": str(default_text or ""),
                    "override": str(
                        _business_text_exact_override(
                            overrides,
                            group,
                            key,
                            scene,
                            product_type,
                            market_type,
                            exchange,
                        ) or ""
                    ),
                }
    return matrix


def _legacy_business_text_variants_from_scenes(scenes: dict) -> dict:
    variants = {}
    for scene, scene_entry in scenes.items():
        matrix = scene_entry.get("matrix", {})
        variants[scene] = {
            "default": str(matrix.get("ALL", {}).get("ALL", {}).get("ALL", {}).get("default", "") or ""),
            "override": str(matrix.get("ALL", {}).get("ALL", {}).get("ALL", {}).get("override", "") or ""),
        }
    return variants


def _count_business_text_overrides(scenes: dict) -> int:
    count = 0
    for scene_entry in scenes.values():
        matrix = scene_entry.get("matrix", {})
        for product_map in matrix.values():
            for market_map in product_map.values():
                for cell in market_map.values():
                    if str(cell.get("override") or "").strip():
                        count += 1
    return count


def _build_contract_business_text_catalog() -> dict:
    default_engine = ContractEngine(business_text_overrides=_empty_business_text_overrides())
    current_overrides = _load_business_text_overrides()
    catalog = {}
    for key, spec in CONTRACT_BUSINESS_TEXT_SPECS.items():
        variant_mode = spec.get("variant_mode", "DEFAULT")
        scenes = {}
        for scene in _contract_business_text_variant_values(variant_mode):
            form = _sample_contract_form_for_business_text_variant(scene)
            values = default_engine._inject_clause_texts(default_engine._derive_variables(form))
            scenes[scene] = {
                "label": _business_text_scene_label(scene),
                "matrix": _business_text_matrix(
                    str(values.get(key, "") or ""),
                    current_overrides,
                    "contract",
                    key,
                    scene,
                    supports_conditions=True,
                ),
            }
        catalog[key] = {
            "key": key,
            "variant_mode": variant_mode,
            "supports_conditions": True,
            "processing_tags": _business_text_processing_tags(spec),
            "processing_description": _business_text_processing_description(spec),
            "scene_label_map": {scene: data.get("label", scene) for scene, data in scenes.items()},
            "scenes": scenes,
            "variants": _legacy_business_text_variants_from_scenes(scenes),
            "configured_override_count": _count_business_text_overrides(scenes),
        }
    return catalog


def _build_prospectus_business_text_catalog() -> dict:
    empty_overrides = _empty_business_text_overrides()
    default_contract_engine = ContractEngine(business_text_overrides=empty_overrides)
    default_engine = ProspectusEngine(
        business_text_overrides=empty_overrides,
        contract_engine=default_contract_engine,
    )
    current_overrides = _load_business_text_overrides()
    catalog = {}
    for key, spec in PROSPECTUS_BUSINESS_TEXT_SPECS.items():
        variant_mode = spec.get("variant_mode", "DEFAULT")
        scenes = {}
        for scene in _business_text_variant_values("prospectus", variant_mode):
            form = _sample_prospectus_form_for_business_text_variant(scene)
            values = default_engine._inject_clause_texts(default_engine._derive_variables(form))
            scenes[scene] = {
                "label": _business_text_scene_label(scene),
                "matrix": _business_text_matrix(
                    str(values.get(key, "") or ""),
                    current_overrides,
                    "prospectus",
                    key,
                    scene,
                    supports_conditions=True,
                ),
            }
        catalog[key] = {
            "key": key,
            "variant_mode": variant_mode,
            "supports_conditions": True,
            "processing_tags": _business_text_processing_tags(spec),
            "processing_description": _business_text_processing_description(spec),
            "scene_label_map": {scene: data.get("label", scene) for scene, data in scenes.items()},
            "scenes": scenes,
            "variants": _legacy_business_text_variants_from_scenes(scenes),
            "configured_override_count": _count_business_text_overrides(scenes),
        }
    return catalog


def _reload_generation_engines() -> None:
    global engine, prospectus_engine, product_summary_engine
    engine = ContractEngine()
    prospectus_engine = ProspectusEngine(contract_engine=engine)
    product_summary_engine = ProductSummaryEngine()


engine = ContractEngine()
prospectus_engine = ProspectusEngine(contract_engine=engine)
product_summary_engine = ProductSummaryEngine()


_VARIABLE_CONDITION_TYPES = {"boolean", "condition", "condition_boolean"}
_VARIABLE_CONDITION_PREFIXES = ("IS_", "HAS_", "USE_", "NOT_", "INCLUDE_", "SHOW_")
_VARIABLE_OPTION_MISSING = object()


def _schema_json_error(message: str, status: int = 400):
    return jsonify({"ok": False, "success": False, "error": message}), status


def _load_schema_json() -> dict:
    with open(SCHEMA_JSON, encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        return {"groups": {}}
    data.setdefault("groups", {})
    return data


def _save_schema_json(data: dict) -> None:
    data["last_updated"] = datetime.now(timezone.utc).date().isoformat()
    with open(SCHEMA_JSON, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
        f.write("\n")


def _schema_variable_key(item: dict) -> str:
    return str((item or {}).get("key") or (item or {}).get("name") or "").strip()


def _schema_variable_label(item: dict) -> str:
    return str(
        (item or {}).get("label")
        or (item or {}).get("ui_entry")
        or (item or {}).get("applies_to")
        or _schema_variable_key(item)
    ).strip()


def _schema_variable_group_label(group: dict, group_key: str) -> str:
    return str((group or {}).get("label") or (group or {}).get("description") or group_key).strip()


def _schema_variable_is_condition(item: dict) -> bool:
    key = _schema_variable_key(item)
    field_type = str((item or {}).get("type") or "").strip()
    return field_type in _VARIABLE_CONDITION_TYPES or key.startswith(_VARIABLE_CONDITION_PREFIXES)


def _schema_variable_default(item: dict):
    if "default" in (item or {}):
        return item.get("default")
    return False if _schema_variable_is_condition(item) else ""


def _schema_variable_item_payload(item: dict, group_key: str, group: dict) -> dict:
    key = _schema_variable_key(item)
    options = (item or {}).get("options")
    return {
        "key": key,
        "label": _schema_variable_label(item),
        "type": str((item or {}).get("type") or "string"),
        "default": _schema_variable_default(item),
        "group": group_key,
        "group_label": _schema_variable_group_label(group, group_key),
        "required": bool((item or {}).get("required", False)),
        "usage": str((item or {}).get("usage") or (item or {}).get("applies_to") or ""),
        "source_field": str((item or {}).get("source_field") or ""),
        "selection_rule": str((item or {}).get("selection_rule") or ""),
        "options": options,
        "options_json": "" if options in (None, "") else json.dumps(options, ensure_ascii=False, indent=2),
        "is_condition": _schema_variable_is_condition(item),
    }


def _schema_variable_group_options(data: dict) -> list[dict]:
    return [
        {"key": group_key, "label": _schema_variable_group_label(group, group_key)}
        for group_key, group in (data.get("groups") or {}).items()
    ]


def _flatten_schema_variables(data: dict) -> list[dict]:
    items = []
    for group_key, group in (data.get("groups") or {}).items():
        for item in group.get("variables") or []:
            if isinstance(item, dict) and _schema_variable_key(item):
                items.append(_schema_variable_item_payload(item, group_key, group))
    return items


def _find_schema_variable(data: dict, key: str):
    for group_key, group in (data.get("groups") or {}).items():
        variables = group.get("variables") or []
        for index, item in enumerate(variables):
            if isinstance(item, dict) and _schema_variable_key(item) == key:
                return group_key, group, index, item
    return None, None, None, None


def _parse_schema_variable_boolean(value) -> bool:
    if isinstance(value, bool):
        return value
    return str(value or "").strip().lower() in {"1", "true", "yes", "y", "on", "是", "启用", "开启"}


def _schema_variable_generation_default(item: dict):
    if "default" not in (item or {}):
        return _VARIABLE_OPTION_MISSING
    default = item.get("default")
    if _schema_variable_is_condition(item):
        return _parse_schema_variable_boolean(default)
    return default


def _merge_schema_variable_defaults(values: dict | None) -> dict:
    merged = dict(values or {})
    try:
        data = _load_schema_json()
    except Exception:
        return merged
    for group in (data.get("groups") or {}).values():
        for item in group.get("variables") or []:
            if not isinstance(item, dict):
                continue
            key = _schema_variable_key(item)
            if not key or key in merged:
                continue
            default = _schema_variable_generation_default(item)
            if default is _VARIABLE_OPTION_MISSING:
                continue
            merged[key] = default
    return merged


def _parse_schema_variable_options(payload: dict):
    if "options_json" not in payload and "options" not in payload:
        return _VARIABLE_OPTION_MISSING
    raw = payload.get("options_json", payload.get("options"))
    if raw in (None, ""):
        return None
    if isinstance(raw, (dict, list)):
        return raw
    return json.loads(str(raw))


def _normalize_schema_variable_payload(payload: dict, existing_key: str | None = None) -> dict:
    payload = payload or {}
    key = str(payload.get("key") or payload.get("name") or existing_key or "").strip()
    if not key:
        raise ValueError("变量名不能为空")
    if not key.replace("_", "").isalnum():
        raise ValueError("变量名只能包含字母、数字和下划线")
    field_type = str(payload.get("type") or "string").strip() or "string"
    if not field_type.replace("_", "").isalnum():
        raise ValueError("变量类型只能包含字母、数字和下划线")
    try:
        options = _parse_schema_variable_options(payload)
    except json.JSONDecodeError as exc:
        raise ValueError(f"选项 JSON 格式错误：{exc}") from exc
    default = payload.get("default")
    if field_type in _VARIABLE_CONDITION_TYPES:
        default = _parse_schema_variable_boolean(default)
    elif default is None:
        default = ""
    return {
        "key": key,
        "label": str(payload.get("label") or payload.get("ui_entry") or key).strip(),
        "type": field_type,
        "default": default,
        "group": str(payload.get("group") or "custom").strip() or "custom",
        "group_label": str(payload.get("group_label") or payload.get("group") or "自定义变量").strip(),
        "required": _parse_schema_variable_boolean(payload.get("required")),
        "usage": str(payload.get("usage") or payload.get("applies_to") or "").strip(),
        "source_field": str(payload.get("source_field") or "").strip(),
        "selection_rule": str(payload.get("selection_rule") or "").strip(),
        "options": options,
    }


def _schema_prefers_name_key(data: dict) -> bool:
    key_count = 0
    name_count = 0
    for group in (data.get("groups") or {}).values():
        for item in group.get("variables") or []:
            key_count += int("key" in item)
            name_count += int("name" in item)
    return name_count > key_count


def _ensure_schema_variable_group(data: dict, group_key: str, group_label: str) -> dict:
    groups = data.setdefault("groups", {})
    group = groups.setdefault(group_key, {"label": group_label or group_key, "variables": []})
    group.setdefault("variables", [])
    if group_label and not (group.get("label") or group.get("description")):
        group["label"] = group_label
    return group


def _apply_schema_variable_update(item: dict, normalized: dict, *, key_field: str) -> dict:
    item[key_field] = normalized["key"]
    if key_field == "key" and "name" in item:
        item.pop("name", None)
    if key_field == "name" and "key" in item:
        item.pop("key", None)
    if "ui_entry" in item and "label" not in item:
        item["ui_entry"] = normalized["label"]
    else:
        item["label"] = normalized["label"]
    item["type"] = normalized["type"]
    item["default"] = normalized["default"]
    item["required"] = normalized["required"]
    if "applies_to" in item and "usage" not in item:
        item["applies_to"] = normalized["usage"]
    elif normalized["usage"] or "usage" in item:
        item["usage"] = normalized["usage"]
    if normalized["source_field"] or "source_field" in item:
        item["source_field"] = normalized["source_field"]
    if normalized["selection_rule"] or "selection_rule" in item:
        item["selection_rule"] = normalized["selection_rule"]
    if normalized["options"] is not _VARIABLE_OPTION_MISSING:
        if normalized["options"] is None:
            item.pop("options", None)
        else:
            item["options"] = normalized["options"]
    return item


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/schema")
def api_schema():
    with open(SCHEMA_JSON, encoding="utf-8") as f:
        return jsonify(json.load(f))


@app.get("/api/variables")
def api_variables():
    data = _load_schema_json()
    items = _flatten_schema_variables(data)
    return jsonify(
        {
            "source_template": str(SCHEMA_JSON),
            "count": len(items),
            "condition_count": sum(1 for item in items if item.get("is_condition")),
            "group_options": _schema_variable_group_options(data),
            "items": items,
        }
    )


@app.post("/api/variables")
def api_create_variable():
    try:
        payload = request.get_json(silent=True) or {}
        normalized = _normalize_schema_variable_payload(payload)
        data = _load_schema_json()
        if _find_schema_variable(data, normalized["key"])[3] is not None:
            return _schema_json_error("变量名已存在")
        group = _ensure_schema_variable_group(data, normalized["group"], normalized["group_label"])
        key_field = "name" if _schema_prefers_name_key(data) else "key"
        item = _apply_schema_variable_update({}, normalized, key_field=key_field)
        group["variables"].append(item)
        _save_schema_json(data)
        return jsonify({"ok": True, "success": True, "id": normalized["key"], "item": _schema_variable_item_payload(item, normalized["group"], group)})
    except ValueError as exc:
        return _schema_json_error(str(exc))


@app.put("/api/variables/<path:key>")
def api_update_variable(key: str):
    try:
        payload = request.get_json(silent=True) or {}
        normalized = _normalize_schema_variable_payload(payload, existing_key=key)
        data = _load_schema_json()
        old_group_key, old_group, index, item = _find_schema_variable(data, key)
        if item is None:
            return _schema_json_error("未找到该变量", 404)
        duplicate_group, _duplicate_group_obj, _duplicate_index, duplicate_item = _find_schema_variable(data, normalized["key"])
        if normalized["key"] != key and duplicate_item is not None:
            return _schema_json_error("新变量名已存在")
        key_field = "name" if "name" in item and "key" not in item else "key"
        updated = _apply_schema_variable_update(dict(item), normalized, key_field=key_field)
        target_group = _ensure_schema_variable_group(data, normalized["group"], normalized["group_label"])
        if normalized["group"] != old_group_key:
            old_group["variables"].pop(index)
            target_group["variables"].append(updated)
        else:
            old_group["variables"][index] = updated
        _save_schema_json(data)
        return jsonify({"ok": True, "success": True, "id": normalized["key"], "item": _schema_variable_item_payload(updated, normalized["group"], target_group)})
    except ValueError as exc:
        return _schema_json_error(str(exc))


@app.delete("/api/variables/<path:key>")
def api_delete_variable(key: str):
    data = _load_schema_json()
    group_key, group, index, item = _find_schema_variable(data, key)
    if item is None:
        return _schema_json_error("未找到该变量", 404)
    removed = group["variables"].pop(index)
    _save_schema_json(data)
    return jsonify({"ok": True, "success": True, "id": key, "item": _schema_variable_item_payload(removed, group_key, group)})


@app.route("/api/custodians")
def api_custodians():
    """返回已知托管人的自动填充数据"""
    with open(CLAUSES_JSON, encoding="utf-8") as f:
        data = json.load(f)
    mgmt_variants = data["clauses"]["MGMT_FEE_PAYMENT"]["variants"]
    custody_variants = data["clauses"]["CUSTODY_FEE_PAYMENT"]["variants"]
    custodians = {}
    for name, info in data["clauses"]["CUSTODIAN_INFO_KNOWN_VALUES"]["custodians"].items():
        expanded = dict(info)
        legacy_payment = expanded.get("payment_method") or ""
        mgmt_value = (
            expanded.get("mgmt_fee_payment_text")
            or expanded.get("mgmt_fee_payment_method")
            or legacy_payment
        )
        custody_value = (
            expanded.get("custody_fee_payment_text")
            or expanded.get("custody_fee_payment_method")
            or legacy_payment
        )
        mgmt_enum, mgmt_text = _resolve_payment_clause_value(mgmt_value, mgmt_variants)
        custody_enum, custody_text = _resolve_payment_clause_value(custody_value, custody_variants)
        expanded["mgmt_fee_payment_text"] = mgmt_text
        expanded["custody_fee_payment_text"] = custody_text
        expanded["mgmt_fee_payment_enum"] = mgmt_enum or ""
        expanded["custody_fee_payment_enum"] = custody_enum or ""
        custodians[name] = expanded
    ordered_custodians = dict(
        sorted(
            custodians.items(),
            key=lambda item: (item[1].get("directory_order", float("inf")), item[0]),
        )
    )
    return app.response_class(
        json.dumps(ordered_custodians, ensure_ascii=False),
        mimetype="application/json",
    )


@app.route("/api/diff_table")
def api_diff_table():
    """返回差异条款匹配表（Markdown原文）"""
    content = DIFF_TABLE_MD.read_text(encoding="utf-8")
    return jsonify({"content": content})


@app.route("/api/extract_text", methods=["POST"])
def api_extract_text():
    """从上传的 .docx / .txt / .md 文件提取纯文本，用于合同对比。"""
    if "file" not in request.files:
        return jsonify({"error": "未收到文件"}), 400
    f = request.files["file"]
    name = f.filename or ""
    ext  = name.rsplit(".", 1)[-1].lower() if "." in name else ""

    if ext == "docx":
        try:
            import tempfile, os as _os
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            f.save(tmp.name)
            tmp.close()
            text = _extract_docx_text_with_tables(tmp.name)
            _os.unlink(tmp.name)
        except Exception as e:
            return jsonify({"error": f"无法读取 .docx：{e}"}), 400
    elif ext in ("txt", "md"):
        raw = f.read()
        for enc in ("utf-8", "gbk", "utf-16"):
            try:
                text = raw.decode(enc)
                break
            except Exception:
                pass
        else:
            return jsonify({"error": "文件编码无法识别，请另存为 UTF-8 后再上传"}), 400
    else:
        return jsonify({"error": "仅支持 .docx、.txt、.md 格式"}), 400

    return jsonify({"filename": name, "text": text,
                    "lines": len(text.splitlines()),
                    "size":  len(text)})


@app.route("/api/compare", methods=["POST"])
def api_compare():
    """
    合同对比分析：接收两段合同文本，返回逐行差异（含字符级高亮HTML）。
    响应格式：
      summary: { total_a, total_b, changed_lines, similarity, diff_groups }
      hunks:   每个 hunk 为 { type: "lines"|"skip", lines?, count?, preview? }
               每行: { tag, a, b, a_html, b_html, num_a, num_b }
    """
    import difflib
    import html as html_mod

    d = request.get_json(force=True)
    text1  = d.get("text1",  "")
    text2  = d.get("text2",  "")
    label1 = d.get("label1", "合同A")
    label2 = d.get("label2", "合同B")
    CTX    = 3   # 每个差异块上下各保留 3 行上下文

    rows1 = _collect_nonblank_compare_lines(text1)
    rows2 = _collect_nonblank_compare_lines(text2)
    lines1 = [row["line"] for row in rows1]
    lines2 = [row["line"] for row in rows2]
    compare1 = [row["compare"] for row in rows1]
    compare2 = [row["compare"] for row in rows2]
    nums1 = [row["lineno"] for row in rows1]
    nums2 = [row["lineno"] for row in rows2]

    if not lines1 or not lines2:
        return jsonify({"error": "两份合同文本均不能为空"}), 400

    # 字符级差异 HTML（对单行 a, b 比较）
    def char_diff(a: str, b: str):
        m = difflib.SequenceMatcher(None, a, b, autojunk=False)
        ra, rb = [], []
        for tag, i1, i2, j1, j2 in m.get_opcodes():
            ea = html_mod.escape(a[i1:i2])
            eb = html_mod.escape(b[j1:j2])
            if tag == "equal":
                ra.append(ea); rb.append(eb)
            elif tag == "replace":
                ra.append(f'<del>{ea}</del>'); rb.append(f'<ins>{eb}</ins>')
            elif tag == "delete":
                ra.append(f'<del>{ea}</del>')
            elif tag == "insert":
                rb.append(f'<ins>{eb}</ins>')
        return "".join(ra), "".join(rb)

    # 逐行对比，构建 result_lines 列表
    sm = difflib.SequenceMatcher(None, compare1, compare2, autojunk=False)
    result_lines = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for k in range(i2 - i1):
                t = html_mod.escape(lines1[i1 + k])
                result_lines.append({"tag": "equal",
                                     "a": t, "b": t,
                                     "a_html": t, "b_html": t,
                                     "num_a": nums1[i1 + k], "num_b": nums2[j1 + k]})
        elif tag == "replace":
            len_a, len_b = i2 - i1, j2 - j1
            for k in range(max(len_a, len_b)):
                la = lines1[i1+k] if k < len_a else None
                lb = lines2[j1+k] if k < len_b else None
                na = nums1[i1 + k] if la is not None else None
                nb = nums2[j1 + k] if lb is not None else None
                if la is not None and lb is not None:
                    ah, bh = char_diff(la, lb)
                    result_lines.append({"tag": "replace",
                                         "a": html_mod.escape(la), "b": html_mod.escape(lb),
                                         "a_html": ah, "b_html": bh,
                                         "num_a": na, "num_b": nb})
                elif la is not None:
                    ea = html_mod.escape(la)
                    result_lines.append({"tag": "delete",
                                         "a": ea, "b": "",
                                         "a_html": f"<del>{ea}</del>", "b_html": "",
                                         "num_a": na, "num_b": None})
                else:
                    eb = html_mod.escape(lb)
                    result_lines.append({"tag": "insert",
                                         "a": "", "b": eb,
                                         "a_html": "", "b_html": f"<ins>{eb}</ins>",
                                         "num_a": None, "num_b": nb})
        elif tag == "delete":
            for k in range(i2 - i1):
                ea = html_mod.escape(lines1[i1+k])
                result_lines.append({"tag": "delete",
                                     "a": ea, "b": "",
                                     "a_html": f"<del>{ea}</del>", "b_html": "",
                                     "num_a": nums1[i1 + k], "num_b": None})
        elif tag == "insert":
            for k in range(j2 - j1):
                eb = html_mod.escape(lines2[j1+k])
                result_lines.append({"tag": "insert",
                                     "a": "", "b": eb,
                                     "a_html": "", "b_html": f"<ins>{eb}</ins>",
                                     "num_a": None, "num_b": nums2[j1 + k]})

    # 压缩连续的 equal 区域（超过 CTX*2+1 行则折叠中间部分）
    hunks = []
    idx, n = 0, len(result_lines)
    while idx < n:
        r = result_lines[idx]
        if r["tag"] == "equal":
            j = idx
            while j < n and result_lines[j]["tag"] == "equal":
                j += 1
            count = j - idx
            if count > CTX * 2 + 1:
                hunks.append({"type": "lines", "lines": result_lines[idx: idx+CTX]})
                hunks.append({"type": "skip",
                              "count": count - CTX*2,
                              "preview": result_lines[idx+CTX]["a"][:60],
                              "lines": result_lines[idx+CTX: j-CTX]})
                hunks.append({"type": "lines", "lines": result_lines[j-CTX: j]})
            else:
                hunks.append({"type": "lines", "lines": result_lines[idx:j]})
            idx = j
        else:
            j = idx
            while j < n and result_lines[j]["tag"] != "equal":
                j += 1
            hunks.append({"type": "lines", "lines": result_lines[idx:j], "is_diff": True})
            idx = j

    # 统计
    changed = sum(1 for r in result_lines if r["tag"] != "equal")
    total   = len(result_lines)
    sim     = round((total - changed) / total * 100, 1) if total > 0 else 100.0
    diff_groups = sum(1 for h in hunks if h.get("is_diff"))

    return jsonify({
        "label1": label1, "label2": label2,
        "summary": {
            "total_a": len(lines1), "total_b": len(lines2),
            "changed_lines": changed, "similarity": sim,
            "diff_groups": diff_groups
        },
        "hunks": hunks
    })


@app.route("/api/summary_check", methods=["POST"])
def api_summary_check():
    """
    合同摘要一致性检验：
    1. 解析第二十六部分摘要的各子项（一、二、…九、）
    2. 按内容相似度自动匹配正文对应章节
    3. 逐项文字比对，返回字符级高亮差异报告
    """
    import difflib
    import html as html_mod

    d = request.get_json(force=True) or {}
    store_data = _review_store.get("data") or {}
    text = d.get("text") or store_data.get("contract_text") or ""
    if not text:
        return jsonify({"error": "合同文本不能为空"}), 400

    CTX = 2  # 差异块上下文行数

    # ── 0. 序号剥离（用于比对，不影响显示） ──────────────────────────────────
    def strip_num(line: str) -> str:
        """剥除行首序号，仅供相似度匹配/diff 对齐使用，不修改原文显示。"""
        return _strip_review_number_prefix(line)

    def normalize_summary_compare_line(line: str) -> str:
        return _normalize_summary_compare_text(strip_num(line))

    # ── 1. 分割合同各部分（复用统一拆分逻辑，避免签署页混入正文）───────────────
    sections = _split_contract_sections(text)
    summary_context = _prepare_summary_review_context(sections)
    if summary_context.get("error"):
        return jsonify({"error": summary_context["error"]}), 400
    summary_sec = summary_context["summary_section"]
    body_sections_all = summary_context["body_sections_all"]
    candidate_sections = summary_context["candidate_sections"]

    fund_type = _infer_review_fund_type_from_text(text)
    review_rules = _load_review_rules(fund_type)
    summary_rules = review_rules.get("summary_rules", [])

    # ── 2. 解析摘要子项（再次清洗摘要尾部签署页，避免摘要与签署页串连）────────────
    summary_content = _strip_contract_signing_page_text(summary_sec["content"])
    subsections = summary_context["summary_subsections"]
    if not summary_rules:
        summary_rules = _build_builtin_etf_summary_rules(subsections)

    # ── 3. 字符级差异 HTML ──────────────────────────────────────────────────
    def char_diff(a, b):
        sm = difflib.SequenceMatcher(None, a, b, autojunk=False)
        ra, rb = [], []
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            ea = html_mod.escape(a[i1:i2])
            eb = html_mod.escape(b[j1:j2])
            if tag == "equal":
                ra.append(ea); rb.append(eb)
            elif tag == "replace":
                ra.append(f'<del>{ea}</del>'); rb.append(f'<ins>{eb}</ins>')
            elif tag == "delete":
                ra.append(f'<del>{ea}</del>')
            elif tag == "insert":
                rb.append(f'<ins>{eb}</ins>')
        return "".join(ra), "".join(rb)

    # ── 4. 逐行差异 + 折叠 ─────────────────────────────────────────────────
    def build_diff(text_a, text_b):
        rows_a = _collect_nonblank_compare_lines(text_a, normalize_line=normalize_summary_compare_line)
        rows_b = _collect_nonblank_compare_lines(text_b, normalize_line=normalize_summary_compare_line)
        lines_a = [row["line"] for row in rows_a]
        lines_b = [row["line"] for row in rows_b]
        norm_a = [row["compare"] for row in rows_a]
        norm_b = [row["compare"] for row in rows_b]
        nums_a = [row["lineno"] for row in rows_a]
        nums_b = [row["lineno"] for row in rows_b]
        sm = difflib.SequenceMatcher(None, norm_a, norm_b, autojunk=False)
        result_lines = []

        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == "equal":
                for k in range(i2 - i1):
                    t = html_mod.escape(lines_a[i1 + k])
                    result_lines.append({"tag": "equal", "a": t, "b": t,
                                         "a_html": t, "b_html": t,
                                         "num_a": nums_a[i1 + k], "num_b": nums_b[j1 + k]})
            elif tag == "replace":
                la, lb = i2 - i1, j2 - j1
                for k in range(max(la, lb)):
                    al = lines_a[i1 + k] if k < la else None
                    bl = lines_b[j1 + k] if k < lb else None
                    if al is not None and bl is not None:
                        ah, bh = char_diff(al, bl)
                        result_lines.append({"tag": "replace",
                                             "a": html_mod.escape(al), "b": html_mod.escape(bl),
                                             "a_html": ah, "b_html": bh,
                                             "num_a": nums_a[i1 + k], "num_b": nums_b[j1 + k]})
                    elif al is not None:
                        ea = html_mod.escape(al)
                        result_lines.append({"tag": "delete", "a": ea, "b": "",
                                             "a_html": f'<del>{ea}</del>', "b_html": "",
                                             "num_a": nums_a[i1 + k], "num_b": None})
                    else:
                        eb = html_mod.escape(bl)
                        result_lines.append({"tag": "insert", "a": "", "b": eb,
                                             "a_html": "", "b_html": f'<ins>{eb}</ins>',
                                             "num_a": None, "num_b": nums_b[j1 + k]})
            elif tag == "delete":
                for k in range(i2 - i1):
                    ea = html_mod.escape(lines_a[i1 + k])
                    result_lines.append({"tag": "delete", "a": ea, "b": "",
                                         "a_html": f'<del>{ea}</del>', "b_html": "",
                                         "num_a": nums_a[i1 + k], "num_b": None})
            elif tag == "insert":
                for k in range(j2 - j1):
                    eb = html_mod.escape(lines_b[j1 + k])
                    result_lines.append({"tag": "insert", "a": "", "b": eb,
                                         "a_html": "", "b_html": f'<ins>{eb}</ins>',
                                         "num_a": None, "num_b": nums_b[j1 + k]})

        # 折叠连续 equal 区
        hunks = []
        idx, n = 0, len(result_lines)
        while idx < n:
            r = result_lines[idx]
            if r["tag"] == "equal":
                j = idx
                while j < n and result_lines[j]["tag"] == "equal":
                    j += 1
                count = j - idx
                if count > CTX * 2 + 1:
                    hunks.append({"type": "lines", "lines": result_lines[idx: idx + CTX]})
                    hunks.append({"type": "skip", "count": count - CTX * 2,
                                  "preview": result_lines[idx + CTX]["a"][:60]})
                    hunks.append({"type": "lines", "lines": result_lines[j - CTX: j]})
                else:
                    hunks.append({"type": "lines", "lines": result_lines[idx:j]})
                idx = j
            else:
                j = idx
                while j < n and result_lines[j]["tag"] != "equal":
                    j += 1
                hunks.append({"type": "lines", "lines": result_lines[idx:j], "is_diff": True})
                idx = j

        changed = sum(1 for r in result_lines if r["tag"] != "equal")
        total = len(result_lines)
        sim = round((total - changed) / total * 100, 1) if total > 0 else 100.0
        diff_groups = sum(1 for h in hunks if h.get("is_diff"))
        return hunks, changed, total, sim, diff_groups

    # ── 5. 按内容相似度匹配正文章节 ─────────────────────────────────────────
    def make_summary_result(sub, matched_rule, target, *, score=0.0, severity="warning",
                            similarity=0.0, changed=0, total=0, diff_groups=0, hunks=None,
                            rule_source="similarity_fallback", locator_matched=False, rule_locator="",
                            rule_level="", summary_group_path="", summary_group_heading="",
                            summary_path="", summary_level="chapter", detail_results=None,
                            excerpt_fields=None, strict_text_match=False):
        section_heading = _normalize_review_text(target.get("section_heading")) if target and locator_matched else ""
        target_heading = _normalize_review_text(target.get("target_heading")) if target and locator_matched else ""
        matched_locator = _format_review_target_locator(target) if locator_matched else ""
        result = {
            "summary_heading": sub["heading"],
            "summary_text": sub["content"],
            "summary_group_path": summary_group_path,
            "summary_group_heading": summary_group_heading or sub["heading"],
            "summary_path": summary_path or summary_group_path or sub["heading"],
            "summary_level": summary_level,
            "matched_section": section_heading or None,
            "matched_target_heading": target_heading,
            "matched_locator": matched_locator,
            "matched_section_text": (target.get("text") or "").strip() if target and locator_matched else "",
            "content_match_score": round(max(score, 0.0) * 100, 1),
            "similarity": similarity,
            "severity": severity,
            "changed_lines": changed,
            "total_lines": total,
            "diff_groups": diff_groups,
            "hunks": hunks or [],
            "rule_source": rule_source,
            "rule_locator": rule_locator,
            "rule_level": rule_level,
            "locator_matched": bool(locator_matched),
            "strict_text_match": bool(strict_text_match),
            "sheet_name": _normalize_review_text(matched_rule.get("sheet_name")) if matched_rule else "",
            "expected_status": _normalize_review_text(matched_rule.get("status")) if matched_rule else "",
            "expected_method": _normalize_review_text(matched_rule.get("method")) if matched_rule else "",
            "expected_detail": _normalize_review_text(matched_rule.get("detail")) if matched_rule else "",
            "detail_results": detail_results or [],
            "detail_count": len(detail_results or []),
            "problem_count": 0,
            "pass_count": 0,
            "contract_anchor_text": "",
            "summary_anchor_text": "",
            "contract_context_excerpt": "",
            "summary_context_excerpt": "",
        }
        if excerpt_fields:
            result.update(excerpt_fields)
        result["is_problem"] = _summary_check_result_is_problem(result)
        if detail_results is not None:
            result["problem_count"] = sum(1 for item in detail_results if item.get("is_problem"))
            result["pass_count"] = sum(1 for item in detail_results if not item.get("is_problem"))
            if detail_results:
                result["is_problem"] = any(item.get("is_problem") for item in detail_results)
        return result

    def build_single_summary_item(sub, matched_rule, *, section_hint="", prefer_subheading=True,
                                  summary_group_path="", summary_group_heading="", summary_path="",
                                  summary_level="detail"):
        strict_text_match = _summary_rule_requires_strict_text_match(matched_rule)
        target, match_score, locator_matched, rule_source, rule_locator, rule_level = _resolve_summary_rule_target(
            sub,
            matched_rule,
            body_sections_all,
            candidate_sections,
            section_hint=section_hint,
            prefer_subheading=prefer_subheading,
        )

        if not locator_matched:
            return make_summary_result(
                sub,
                matched_rule,
                target,
                score=0.0,
                severity="warning",
                similarity=0.0,
                changed=0,
                total=0,
                diff_groups=0,
                hunks=[],
                rule_source=rule_source,
                locator_matched=False,
                rule_locator=rule_locator,
                rule_level=rule_level,
                summary_group_path=summary_group_path,
                summary_group_heading=summary_group_heading,
                summary_path=summary_path,
                summary_level=summary_level,
                strict_text_match=strict_text_match,
            )

        source_text_for_compare = (target.get("body_text") or target.get("text") or "").strip()
        hunks, changed, total, sim, diff_group_count = build_diff(source_text_for_compare, sub["content"])
        severity = "error" if strict_text_match and changed > 0 else ("info" if sim >= 95 else ("warning" if sim >= 80 else "error"))
        excerpt_fields = _build_review_excerpt_pair(
            source_text_for_compare,
            sub["content"],
            normalize_line=normalize_summary_compare_line,
            left_prefix="contract",
            right_prefix="summary",
        )
        return make_summary_result(
            sub,
            matched_rule,
            target,
            score=match_score,
            severity=severity,
            similarity=sim,
            changed=changed,
            total=total,
            diff_groups=diff_group_count,
            hunks=hunks,
            rule_source=rule_source,
            locator_matched=True,
            rule_locator=rule_locator,
            rule_level=rule_level,
            summary_group_path=summary_group_path,
            summary_group_heading=summary_group_heading,
            summary_path=summary_path,
            summary_level=summary_level,
            excerpt_fields=excerpt_fields,
            strict_text_match=strict_text_match,
        )

    results = []
    for sub in subsections:
        summary_group_path = sub.get("summary_group_path") or ""
        summary_group_heading = sub.get("summary_group_heading") or sub["heading"]
        has_detail_items = bool(sub.get("detail_items"))
        force_detail_compare = _summary_group_forces_detail_compare(summary_group_heading)
        chapter_rule = _find_matching_summary_rule(
            summary_rules,
            sub["heading"],
            summary_group_path=summary_group_path,
            summary_group_heading=summary_group_heading,
            prefer_detail=False,
        )
        if not chapter_rule and not has_detail_items:
            chapter_rule = _find_matching_summary_rule(
                summary_rules,
                sub["heading"],
                summary_group_path=summary_group_path,
                summary_group_heading=summary_group_heading,
                prefer_detail=True,
            )
        chapter_result = build_single_summary_item(
            sub,
            chapter_rule,
            prefer_subheading=not has_detail_items,
            summary_group_path=summary_group_path,
            summary_group_heading=summary_group_heading,
            summary_path=sub.get("summary_path") or summary_group_path or sub["heading"],
            summary_level="group" if has_detail_items else _normalize_review_text((chapter_rule or {}).get("rule_level")) or "chapter",
        )

        comparable_detail_items = [
            detail for detail in (sub.get("detail_items") or [])
            if str(detail.get("content") or "").strip()
        ]
        detail_results = []
        for detail in sub.get("detail_items") or []:
            if force_detail_compare and not str(detail.get("content") or "").strip():
                continue
            detail_rule = _find_matching_summary_rule(
                summary_rules,
                detail["heading"],
                summary_group_path=summary_group_path,
                summary_group_heading=summary_group_heading,
                prefer_detail=True,
            )
            matched_rule = detail_rule if _normalize_review_text((detail_rule or {}).get("rule_level")) == "detail" else None
            if matched_rule is None and force_detail_compare:
                matched_rule = chapter_rule
            if matched_rule is None:
                continue
            detail_section_hint = (
                chapter_result.get("matched_section")
                or _normalize_review_text((chapter_rule or {}).get("contract_pos"))
            )
            detail_results.append(
                build_single_summary_item(
                    detail,
                    matched_rule,
                    section_hint=detail_section_hint,
                    prefer_subheading=True,
                    summary_group_path=summary_group_path,
                    summary_group_heading=summary_group_heading,
                    summary_path=detail.get("summary_path") or detail["heading"],
                    summary_level="detail",
                )
            )

        if detail_results:
            representative = next((item for item in detail_results if item.get("is_problem")), detail_results[0])
            chapter_result["detail_results"] = detail_results
            chapter_result["detail_count"] = len(detail_results)
            chapter_result["problem_count"] = sum(1 for item in detail_results if item.get("is_problem"))
            chapter_result["pass_count"] = sum(1 for item in detail_results if not item.get("is_problem"))
            chapter_result["is_problem"] = any(item.get("is_problem") for item in detail_results)
            chapter_result["severity"] = (
                "error" if any(item.get("severity") == "error" for item in detail_results)
                else "warning" if any(item.get("is_problem") for item in detail_results)
                else "info"
            )
            for field_name in (
                "contract_anchor_text",
                "summary_anchor_text",
                "contract_context_excerpt",
                "summary_context_excerpt",
            ):
                chapter_result[field_name] = representative.get(field_name, "")
        elif chapter_result.get("locator_matched") and _summary_rule_expects_partial(chapter_rule):
            chapter_result["is_problem"] = False
            chapter_result["severity"] = "info"
        elif force_detail_compare and not comparable_detail_items:
            chapter_result["is_problem"] = False
            chapter_result["severity"] = "info"

        results.append(chapter_result)

    return jsonify({"results": results, "total_subsections": len(results)})


@app.route("/api/generate", methods=["POST"])
def api_generate():
    """接收表单数据，返回合同全文"""
    form_data = request.get_json(force=True)
    try:
        contract_text = engine.generate(form_data)
        return jsonify({"success": True, "text": contract_text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export", methods=["POST"])
def api_export():
    """返回 .txt 文件下载"""
    form_data = request.get_json(force=True)
    try:
        contract_text = engine.generate(form_data)
        fund_name = form_data.get("FUND_NAME", "ETF基金合同")
        # Write to a temp file
        import tempfile
        tmp = tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", encoding="utf-8",
            delete=False, prefix="contract_"
        )
        tmp.write(contract_text)
        tmp.close()
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}基金合同.txt",
            mimetype="text/plain; charset=utf-8",
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export_docx", methods=["POST"])
def api_export_docx():
    """返回格式化 .docx 文件下载"""
    form_data = request.get_json(force=True)
    try:
        contract_text = engine.generate(form_data)
        docx_bytes = engine.build_docx(contract_text)
        fund_name = form_data.get("FUND_NAME", "ETF基金合同")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile, os as _os
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="contract_")
        tmp.write(docx_bytes)
        tmp.close()
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}基金合同.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e), "trace": traceback.format_exc()}), 400


@app.route("/api/files")
def api_files():
    """列出所有可编辑文件"""
    return jsonify(_iter_editable_knowledge_files())


@app.route("/api/files/<path:filename>", methods=["GET", "POST"])
def api_file(filename):
    """读取或保存指定文件（防路径遍历）"""
    target, relative_path = _resolve_editable_knowledge_file(filename)

    if request.method == "GET":
        content = target.read_text(encoding="utf-8")
        return jsonify({"name": target.name, "path": relative_path, "content": content})

    # POST: save
    data = request.get_json(force=True)
    content = data.get("content", "")
    target.write_text(content, encoding="utf-8")
    # Reload engine after template changes
    global engine
    engine = ContractEngine()
    return jsonify({"success": True, "message": f"已保存 {relative_path}", "path": relative_path})


@app.route("/api/clause_library")
def api_clause_library():
    """返回面向业务维护的条款库扁平条目。"""
    return jsonify(_build_clause_library_catalog())


@app.route("/api/clause_library", methods=["POST"])
def api_clause_library_save():
    """保存单个条款正文字段，不接受整文件覆盖。"""
    data = request.get_json(silent=True)
    if not isinstance(data, dict):
        return jsonify({"success": False, "error": "请求体必须是 JSON 对象"}), 400
    try:
        result = _save_clause_library_field(
            str(data.get("path_id") or ""),
            str(data.get("content") or ""),
        )
    except FileNotFoundError as exc:
        return jsonify({"success": False, "error": str(exc)}), 404
    except (ValueError, json.JSONDecodeError) as exc:
        return jsonify({"success": False, "error": str(exc)}), 400
    except Exception as exc:
        logger.exception("Save clause library field failed")
        return jsonify({"success": False, "error": f"保存条款正文失败：{exc}"}), 500
    return jsonify(result)


@app.route("/api/business_texts")
def api_business_texts():
    return jsonify({
        "success": True,
        "path": _business_text_override_path().name,
        "condition_options": _business_text_condition_options(),
        "groups": {
            "contract": _build_contract_business_text_catalog(),
            "prospectus": _build_prospectus_business_text_catalog(),
        },
    })


@app.route("/api/business_texts", methods=["POST"])
def api_business_texts_save():
    payload = request.get_json(force=True) or {}
    group = str(payload.get("group") or "").strip()
    key = str(payload.get("key") or "").strip()
    scene = str(payload.get("scene") or payload.get("variant") or "").strip()
    product_type = str(payload.get("product_type") or "ALL").strip().upper() or "ALL"
    market_type = str(payload.get("market_type") or "ALL").strip().upper() or "ALL"
    exchange = str(payload.get("exchange") or "ALL").strip().upper() or "ALL"
    content = str(payload.get("content") or "")

    if group not in BUSINESS_TEXT_GROUP_SPECS:
        return jsonify({"success": False, "error": "未知业务正文分组"}), 400
    if key not in BUSINESS_TEXT_GROUP_SPECS[group]:
        return jsonify({"success": False, "error": "未知业务正文键名"}), 400

    allowed_variants = _business_text_variant_values(
        group,
        BUSINESS_TEXT_GROUP_SPECS[group][key].get("variant_mode", "DEFAULT")
    )
    if scene not in allowed_variants:
        return jsonify({"success": False, "error": "未知业务正文场景"}), 400
    if product_type not in _business_text_product_type_values():
        return jsonify({"success": False, "error": "未知产品类型条件"}), 400
    if market_type not in _business_text_market_type_values():
        return jsonify({"success": False, "error": "未知市场类型条件"}), 400
    if exchange not in _business_text_exchange_values():
        return jsonify({"success": False, "error": "未知上市交易所条件"}), 400

    overrides = _load_business_text_overrides()
    overrides = _set_business_text_override(
        overrides,
        group,
        key,
        scene,
        content,
        product_type,
        market_type,
        exchange,
    )
    _save_business_text_overrides(overrides)
    _reload_generation_engines()
    return jsonify({
        "success": True,
        "path": _business_text_override_path().name,
        "group": group,
        "key": key,
        "scene": scene,
        "variant": scene,
        "product_type": product_type,
        "market_type": market_type,
        "exchange": exchange,
    })


@app.route("/api/business_texts/reset", methods=["POST"])
def api_business_texts_reset():
    payload = request.get_json(force=True) or {}
    group = str(payload.get("group") or "").strip()
    key = str(payload.get("key") or "").strip()
    scene = str(payload.get("scene") or payload.get("variant") or "").strip()
    product_type = str(payload.get("product_type") or "ALL").strip().upper() or "ALL"
    market_type = str(payload.get("market_type") or "ALL").strip().upper() or "ALL"
    exchange = str(payload.get("exchange") or "ALL").strip().upper() or "ALL"

    if group not in BUSINESS_TEXT_GROUP_SPECS:
        return jsonify({"success": False, "error": "未知业务正文分组"}), 400
    if key not in BUSINESS_TEXT_GROUP_SPECS[group]:
        return jsonify({"success": False, "error": "未知业务正文键名"}), 400

    allowed_variants = _business_text_variant_values(
        group,
        BUSINESS_TEXT_GROUP_SPECS[group][key].get("variant_mode", "DEFAULT")
    )
    if scene not in allowed_variants:
        return jsonify({"success": False, "error": "未知业务正文场景"}), 400
    if product_type not in _business_text_product_type_values():
        return jsonify({"success": False, "error": "未知产品类型条件"}), 400
    if market_type not in _business_text_market_type_values():
        return jsonify({"success": False, "error": "未知市场类型条件"}), 400
    if exchange not in _business_text_exchange_values():
        return jsonify({"success": False, "error": "未知上市交易所条件"}), 400

    overrides = _load_business_text_overrides()
    overrides = _delete_business_text_override(
        overrides,
        group,
        key,
        scene,
        product_type,
        market_type,
        exchange,
    )
    _save_business_text_overrides(overrides)
    _reload_generation_engines()
    return jsonify({
        "success": True,
        "path": _business_text_override_path().name,
        "group": group,
        "key": key,
        "scene": scene,
        "variant": scene,
        "product_type": product_type,
        "market_type": market_type,
        "exchange": exchange,
    })


@app.route("/api/preview_clause", methods=["POST"])
def api_preview_clause():
    """根据当前表单值，实时返回关键差异条款预览文本"""
    form_data = request.get_json(force=True)
    try:
        v = engine._derive_variables(form_data)
        v = engine._inject_clause_texts(v)
        preview = {
            "WORKING_DAY_DEF": v.get("WORKING_DAY_DEF", ""),
            "BUSINESS_RULES_DEF": v.get("BUSINESS_RULES_DEF", ""),
            "DISPUTE_RESOLUTION_LABEL": v.get("DISPUTE_RESOLUTION_LABEL", ""),
            "DISPUTE_RESOLUTION_CLAUSE": v.get("DISPUTE_RESOLUTION_CLAUSE", ""),
            "NON_COMPONENT_SCOPE": v.get("NON_COMPONENT_SCOPE", ""),
            "DISTRIBUTION_FREQ_CLAUSE": v.get("DISTRIBUTION_FREQ_CLAUSE", ""),
            "MGMT_FEE_PAYMENT_METHOD": v.get("MGMT_FEE_PAYMENT_METHOD", ""),
            "TRACKING_ERROR_DAILY": v.get("TRACKING_ERROR_DAILY", 0.2),
            "TRACKING_ERROR_ANNUAL": v.get("TRACKING_ERROR_ANNUAL", 2),
        }
        return jsonify(preview)
    except Exception as e:
        return jsonify({"error": str(e)}), 400


# ── 招募说明书路由 ────────────────────────────────────────────────────────────


@app.route("/api/prospectus/custodian_summary/status")
def api_prospectus_custodian_summary_status():
    return jsonify({"success": True, **_serialize_custodian_summary_state()})


@app.route("/api/prospectus/custodian_summary/clear", methods=["POST"])
def api_prospectus_custodian_summary_clear():
    state = _clear_custodian_summary_state()
    return jsonify({"success": True, **_serialize_custodian_summary_state(state)})


@app.route("/api/prospectus/custodian_summary/upload", methods=["POST"])
def api_prospectus_custodian_summary_upload():
    if "file" not in request.files:
        return jsonify({"success": False, "error": "未收到托管协议文件"}), 400

    file = request.files["file"]
    filename = str(file.filename or "").strip()
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext != "docx":
        return jsonify({"success": False, "error": "仅支持上传 .docx 托管协议"}), 400

    try:
        import tempfile

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = tmp.name
        try:
            file.save(tmp_path)
        finally:
            tmp.close()

        try:
            text, sections = _extract_custodian_summary_sections_from_docx(tmp_path)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except Exception as e:
        return jsonify({"success": False, "error": f"托管协议解析失败：{e}"}), 400

    if not sections:
        return jsonify({"success": False, "error": "未识别到可用于托管协议摘要的章节标题"}), 400

    state = _set_custodian_summary_state(filename, text, sections)
    return jsonify({"success": True, **_serialize_custodian_summary_state(state)})


@app.route("/api/generate_prospectus", methods=["POST"])
def api_generate_prospectus():
    """接收表单数据，生成招募说明书全文，返回 JSON"""
    form_data = request.get_json(force=True)
    try:
        bundle = prospectus_engine.generate_bundle(form_data)
        return jsonify({"success": True, **bundle})
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e), "trace": traceback.format_exc()}), 400


@app.route("/api/export_prospectus_docx", methods=["POST"])
def api_export_prospectus_docx():
    """\u8fd4\u56de\u683c\u5f0f\u5316\u62db\u52df\u8bf4\u660e\u4e66 .docx \u6587\u4ef6\u4e0b\u8f7d"""
    form_data = request.get_json(force=True)
    try:
        prospectus_text = str(form_data.get("PROSPECTUS_TEXT") or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not prospectus_text:
            prospectus_text = prospectus_engine.generate(form_data)
        report = prospectus_engine.validate_exportable_text(prospectus_text)
        if not report["ok"]:
            report = {k: v for k, v in report.items() if k != "ok"}
            return jsonify({"success": False, **report}), 400
        docx_bytes = prospectus_engine.build_docx_prospectus(prospectus_text, form_data)
        fund_name = form_data.get("FUND_NAME", "ETF\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="prospectus_")
        tmp.write(docx_bytes)
        tmp.close()
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}\u62db\u52df\u8bf4\u660e\u4e66.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except MissingProspectusReferenceAssetsError as e:
        return jsonify({
            "success": False,
            "error": "缺少招募说明书 Word 导出所需的参考 DOCX 资产，请先补齐参考文件后再导出。",
            "error_type": "missing_reference_assets",
            "matches": e.missing_assets,
        }), 400
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e), "trace": traceback.format_exc()}), 400


@app.route("/api/export_prospectus_txt", methods=["POST"])
def api_export_prospectus_txt():
    """\u8fd4\u56de\u62db\u52df\u8bf4\u660e\u4e66 .txt \u6587\u4ef6\u4e0b\u8f7d"""
    form_data = request.get_json(force=True)
    try:
        prospectus_text = str(form_data.get("PROSPECTUS_TEXT") or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not prospectus_text:
            prospectus_text = prospectus_engine.generate(form_data)
        report = prospectus_engine.validate_exportable_text(prospectus_text)
        if not report["ok"]:
            report = {k: v for k, v in report.items() if k != "ok"}
            return jsonify({"success": False, **report}), 400
        fund_name = form_data.get("FUND_NAME", "ETF\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66")
        import tempfile
        tmp = tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", encoding="utf-8",
            delete=False, prefix="prospectus_"
        )
        tmp.write(prospectus_text)
        tmp.close()
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}\u62db\u52df\u8bf4\u660e\u4e66.txt",
            mimetype="text/plain; charset=utf-8",
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/generate_product_summary", methods=["POST"])
def api_generate_product_summary():
    """接收表单数据，生成产品资料概要全文，返回 JSON"""
    form_data = request.get_json(force=True)
    try:
        bundle = product_summary_engine.generate_bundle(form_data)
        return jsonify({"success": True, **bundle})
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e), "trace": traceback.format_exc()}), 400


@app.route("/api/export_product_summary_docx", methods=["POST"])
def api_export_product_summary_docx():
    """返回格式化基金产品资料概要 .docx 文件下载"""
    form_data = request.get_json(force=True)
    try:
        bundle = product_summary_engine.generate_bundle(form_data)
        docx_bytes = product_summary_engine.build_docx_product_summary(bundle["render_model"], form_data)
        fund_name = form_data.get("FUND_NAME", "ETF基金产品资料概要")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="product_summary_")
        tmp.write(docx_bytes)
        tmp.close()
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}基金产品资料概要.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e), "trace": traceback.format_exc()}), 400


@app.route("/api/export_product_summary_txt", methods=["POST"])
def api_export_product_summary_txt():
    """返回基金产品资料概要 .txt 文件下载"""
    form_data = request.get_json(force=True)
    try:
        bundle = product_summary_engine.generate_bundle(form_data)
        fund_name = form_data.get("FUND_NAME", "ETF基金产品资料概要")
        import tempfile
        tmp = tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", encoding="utf-8",
            delete=False, prefix="product_summary_"
        )
        tmp.write(bundle["text"])
        tmp.close()
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}基金产品资料概要.txt",
            mimetype="text/plain; charset=utf-8",
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


# ═══════════════════════════════════════════════════════════════════════════════
# 复核系统 API 路由
# ═══════════════════════════════════════════════════════════════════════════════

# 全局存储上传解析结果（单用户Flask够用，多用户需改session/DB）
_review_store = {}
_revision_workbench_store = {}

RULES_XLSX = PACKAGED_ASSETS_DIR / "review_rules" / RULES_XLSX_FILENAME
LEGACY_RULES_XLSX = LEGACY_RULES_XLSX_CANDIDATES[0]


def _rules_xlsx_candidates() -> list[Path]:
    candidates = []
    env_file = str(os.getenv("ETF_REVIEW_RULES_XLSX") or "").strip()
    if env_file:
        candidates.append(Path(env_file))
    candidates.append(RULES_XLSX)
    for legacy_candidate in LEGACY_RULES_XLSX_CANDIDATES:
        if legacy_candidate not in candidates:
            candidates.append(legacy_candidate)
    return candidates


def _resolve_rules_xlsx() -> Path:
    candidates = _rules_xlsx_candidates()
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def _split_contract_sections(text):
    """按 '第X部分' 正则拆分合同章节，自动剥离末尾签署页"""
    # 匹配 "第X部分" 后跟任意字符（可有可无空格），兼容 "第一部分前言" 和 "第一部分  前言"
    RE_PART = re.compile(r'^(第[一二三四五六七八九十百]+部分\s*\S[^\n]*)', re.MULTILINE)
    RE_TOC_LINE = re.compile(r'\t\d+\s*$')  # 目录行：末尾有tab+页码

    # 过滤掉目录行，只保留正文标题
    part_iter = [m for m in RE_PART.finditer(text) if not RE_TOC_LINE.search(m.group(1))]

    sections = []
    for i, m in enumerate(part_iter):
        heading = m.group(1).split('\t')[0].strip()
        start = m.end()
        end = part_iter[i + 1].start() if i + 1 < len(part_iter) else len(text)
        content = _strip_contract_signing_page_text(text[start:end])
        sections.append({"heading": heading, "content": content})

    return sections


def _split_prospectus_sections(text, doc=None):
    """按 Heading 样式 或 顶级标题 拆分招募说明书章节。
    招募说明书通常用 Heading 2 (不是 Heading 1) 作为顶级章节标题。
    """
    def heading_style_level(style_name: str):
        style_name = str(style_name or "").strip()
        if not style_name:
            return None
        match = re.match(r"^(?:Heading|标题)\s*(\d+)$", style_name, flags=re.IGNORECASE)
        return int(match.group(1)) if match else None

    def is_toc_line(line: str) -> bool:
        stripped = (line or "").strip()
        return bool(stripped and '\t' in stripped and re.search(r'\t+\d+\s*$', stripped))

    def toc_line_heading(line: str) -> str:
        stripped = (line or "").strip()
        if not stripped:
            return ""
        if is_toc_line(stripped):
            return re.sub(r'\t+\d+\s*$', '', stripped).strip()
        return ""

    def line_is_plain_heading(line: str) -> bool:
        stripped = (line or "").strip()
        if not stripped or is_toc_line(stripped):
            return False
        if len(stripped) > 80:
            return False
        if re.search(r"[。！？；]", stripped):
            return False
        return True

    def unique_ordered(values):
        result = []
        for value in values:
            if value and value not in result:
                result.append(value)
        return result

    def labels_equivalent(left: str, right: str) -> bool:
        left_variants = _review_label_variants(left)
        right_variants = _review_label_variants(right)
        for left_variant in left_variants:
            left_key = _review_heading_key(left_variant)
            left_ordinal = _review_label_ordinal(left_variant)
            if not left_key:
                continue
            for right_variant in right_variants:
                right_key = _review_heading_key(right_variant)
                right_ordinal = _review_label_ordinal(right_variant)
                if not right_key:
                    continue
                if left_ordinal is not None and right_ordinal is not None and left_ordinal != right_ordinal:
                    continue
                if left_key == right_key:
                    return True
        return False

    def build_sections_from_line_matches(matches):
        deduped_matches = []
        seen = set()
        for line_index, pos, heading in matches:
            key = (line_index, heading)
            if key in seen:
                continue
            seen.add(key)
            deduped_matches.append((line_index, pos, heading))

        sections = []
        for i, (line_index, pos, heading) in enumerate(deduped_matches):
            line_text = lines_list[line_index]
            start = pos + len(line_text)
            end = deduped_matches[i + 1][1] if i + 1 < len(deduped_matches) else len(text)
            sections.append({"heading": heading, "content": text[start:end].strip()})
        return sections

    def find_heading_line_matches(headings):
        matches = []
        search_start = 0
        for heading in headings:
            found = None
            for line_index in range(search_start, len(line_offsets)):
                pos, line = line_offsets[line_index]
                stripped = line.strip()
                if not line_is_plain_heading(stripped):
                    continue
                if labels_equivalent(stripped, heading):
                    found = (line_index, pos, stripped)
                    break
            if not found:
                continue
            matches.append(found)
            search_start = found[0] + 1
        return matches

    lines_list = text.split('\n')
    line_offsets = []
    offset = 0
    for line in lines_list:
        line_offsets.append((offset, line))
        offset += len(line) + 1

    styled_headings = []
    if doc:
        paragraphs = list(getattr(doc, "paragraphs", None) or [])
        if not paragraphs:
            try:
                for block_type, block in _iter_docx_blocks(doc):
                    if block_type == "paragraph":
                        paragraphs.append(block)
            except Exception:
                paragraphs = []

        styled_entries = []
        for paragraph in paragraphs:
            ptext = (getattr(paragraph, "text", "") or "").strip()
            if not ptext:
                continue
            style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
            level = heading_style_level(style_name)
            if level is None:
                continue
            styled_entries.append({"text": ptext, "heading_level": level})

        if styled_entries:
            min_level = min(entry["heading_level"] for entry in styled_entries)
            styled_headings = [entry["text"] for entry in styled_entries if entry["heading_level"] == min_level]

    toc_headings = [toc_line_heading(line) for _, line in line_offsets]
    candidate_headings = unique_ordered([*styled_headings, *toc_headings])
    heading_matches = find_heading_line_matches(candidate_headings)

    if heading_matches:
        return build_sections_from_line_matches(heading_matches)

    if not heading_matches:
        RE_TOP = re.compile(r'^([一二三四五六七八九十百]+、[^\n]+)', re.MULTILINE)
        regex_matches = []
        for m in RE_TOP.finditer(text):
            heading_text = m.group(1).strip()
            if '\t' in heading_text and re.search(r'\d+$', heading_text):
                continue
            line_index = text[:m.start()].count("\n")
            regex_matches.append((line_index, m.start(), heading_text))
        heading_matches = regex_matches

    return build_sections_from_line_matches(heading_matches)


def _find_contract_summary_section(sections):
    """优先定位“基金合同内容摘要”，避免误命中“基金托管协议的内容摘要”一类章节。"""
    summary_candidates = []
    for section in sections or []:
        heading = (section.get("heading") or "").strip()
        content = _strip_contract_signing_page_text(section.get("content") or "")
        if "摘要" not in heading or len(content) <= 10:
            continue
        cleaned_section = dict(section)
        cleaned_section["content"] = content
        summary_candidates.append(cleaned_section)

    for section in summary_candidates:
        heading = re.sub(r"\s+", "", section.get("heading") or "")
        if "基金合同" in heading and "托管协议" not in heading:
            return section

    for section in summary_candidates:
        heading = re.sub(r"\s+", "", section.get("heading") or "")
        if "合同" in heading and "托管" not in heading:
            return section

    return summary_candidates[-1] if summary_candidates else None


SUMMARY_PROBLEM_SIMILARITY_THRESHOLD = 95.0
SUMMARY_WARNING_SIMILARITY_THRESHOLD = 80.0
AI_REVIEW_TEXT_LIMIT = 12000
AI_REVIEW_TEXT_EDGE = 6000

_REVIEW_NUM_PREFIX_RE = re.compile(
    r"^(?:"
    r"第[一二三四五六七八九十百千]+条"
    r"|[（(][一二三四五六七八九十百千]+[）)]"
    r"|[一二三四五六七八九十百千]+[、．]"
    r"|[（(]\d+[）)]"
    r"|\d+[.、．]"
    r")\s*"
)
_SUMMARY_SUBSECTION_RE = re.compile(r"^([一二三四五六七八九十百]+、[^\n]+)", re.MULTILINE)
_SUMMARY_DETAIL_HEADING_RE = re.compile(
    r"^((?:"
    r"[（(][一二三四五六七八九十百千\d]+[）)]"
    r"|[一二三四五六七八九十百千]+[、．]"
    r"|\d+[.、．]"
    r")[^\n]+)",
    re.MULTILINE,
)
_REVIEW_CHAPTER_PREFIX_RE = re.compile(r"^(第([一二三四五六七八九十百千\d]+)(?:部分|章|节|条))\s*(.*)$")
_REVIEW_LOCATOR_MISSING_KEYWORDS = (
    "无直接对应",
    "无同名独立条款",
    "无对应独立章节",
    "未纳入摘要",
    "未进入摘要",
    "未收录",
    "无对应正文",
)


def _review_similarity_severity(similarity: float | int | None) -> str:
    try:
        similarity_value = float(similarity)
    except (TypeError, ValueError):
        similarity_value = 0.0
    if similarity_value >= SUMMARY_PROBLEM_SIMILARITY_THRESHOLD:
        return "info"
    if similarity_value >= SUMMARY_WARNING_SIMILARITY_THRESHOLD:
        return "warning"
    return "error"


def _summary_cross_severity(similarity: float | int | None, changed_lines: int | None = None) -> str:
    try:
        changed = int(changed_lines or 0)
    except (TypeError, ValueError):
        changed = 0
    if changed <= 0:
        return "info"
    try:
        similarity_value = float(similarity)
    except (TypeError, ValueError):
        similarity_value = 0.0
    if similarity_value >= SUMMARY_PROBLEM_SIMILARITY_THRESHOLD:
        return "warning"
    return "error"


def _normalize_summary_compare_text(text: str) -> str:
    normalized = _normalize_contract_prospectus_compare_text(text)
    normalized = re.sub(
        r"第[\(（]([一二三四五六七八九十百千\d]+)[\)）](?=(?:条|款|项|目|节|部分|章))",
        r"第\1",
        normalized,
    )
    return normalized


def _summary_check_result_is_problem(result) -> bool:
    if not result:
        return False
    if not result.get("locator_matched"):
        return True
    if result.get("strict_text_match"):
        try:
            return int(result.get("changed_lines") or 0) > 0
        except (TypeError, ValueError):
            return True
    try:
        return float(result.get("similarity") or 0.0) < SUMMARY_PROBLEM_SIMILARITY_THRESHOLD
    except (TypeError, ValueError):
        return True


def _cross_check_result_is_problem(result) -> bool:
    status = str((result or {}).get("status") or "").strip().lower()
    return status in {"fail", "missing", "warning"}


def _summary_cross_result_is_problem(similarity, changed_lines: int | None = None) -> bool:
    try:
        changed = int(changed_lines or 0)
    except (TypeError, ValueError):
        changed = 0
    if changed_lines is not None:
        return changed > 0
    try:
        return float(similarity) < SUMMARY_PROBLEM_SIMILARITY_THRESHOLD
    except (TypeError, ValueError):
        return True


def _strip_review_number_prefix(line: str) -> str:
    return _REVIEW_NUM_PREFIX_RE.sub("", (line or "").strip())


def _truncate_ai_review_text(text: str, *, limit: int = AI_REVIEW_TEXT_LIMIT, edge: int = AI_REVIEW_TEXT_EDGE) -> str:
    text = (text or "").strip()
    if len(text) <= limit:
        return text
    head = text[:edge]
    tail = text[-edge:]
    omitted = max(len(text) - edge * 2, 0)
    return f"{head}\n...[已省略中间 {omitted} 字]...\n{tail}"


def _truncate_review_excerpt(text: str, *, limit: int = 220) -> str:
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(text) <= limit:
        return text
    return text[: max(limit - 1, 1)].rstrip() + "…"


def _collect_nonblank_compare_lines(text: str, *, normalize_line=None):
    normalize_line = normalize_line or (lambda value: value)
    rows = []
    for lineno, raw_line in enumerate(str(text or "").splitlines(), 1):
        if _is_layout_blank_line(raw_line) or _is_markdown_table_separator_line(raw_line):
            continue
        normalized = normalize_line(raw_line)
        if normalized is None:
            continue
        normalized = str(normalized)
        if _is_layout_blank_line(normalized) or _is_markdown_table_separator_line(normalized):
            continue
        rows.append({
            "line": raw_line,
            "compare": normalized,
            "lineno": lineno,
        })
    return rows


def _summary_group_path_from_heading(heading: str) -> str:
    text = _normalize_review_text(heading)
    if not text:
        return ""
    match = re.match(r"^(?:摘要)?([一二三四五六七八九十百千\d]+)", text)
    if not match:
        return ""
    return f"摘要{match.group(1)}"


_SUMMARY_SELECTED_DETAIL_HEADINGS = {
    "基金合同当事人的权利、义务": {
        "基金管理人的权利与义务",
        "基金托管人的权利与义务",
    },
    "基金费用与税收": {
        "基金费用的种类",
        "基金费用计提方法、计提标准和支付方式",
    },
    "基金财产的投资范围和投资限制": {
        "投资范围",
        "投资限制",
    },
    "基金资产净值的计算方法和公告方式": {
        "五、估值程序",
    },
}


def _summary_heading_key(heading: str) -> str:
    return _normalize_review_text(_strip_review_number_prefix(heading))


def _summary_heading_ordinal(heading: str):
    text = _normalize_review_text(heading)
    match = re.match(r"^([一二三四五六七八九十百千\d]+)[、．]", text)
    if not match:
        return None
    token = match.group(1)
    return int(token) if token.isdigit() else ContractEngine._cn_numeral_to_int(token)


def _summary_group_selected_detail_headings(group_heading: str) -> set[str]:
    return {
        _summary_heading_key(heading)
        for heading in _SUMMARY_SELECTED_DETAIL_HEADINGS.get(_summary_heading_key(group_heading), set())
    }


def _summary_group_selected_detail_labels(group_heading: str) -> list[str]:
    labels = list(_SUMMARY_SELECTED_DETAIL_HEADINGS.get(_summary_heading_key(group_heading), set()))
    return labels


def _summary_group_forces_detail_compare(group_heading: str) -> bool:
    return bool(_summary_group_selected_detail_headings(group_heading))


def _normalize_review_locator(locator: str) -> str:
    text = _normalize_review_text(locator)
    if not text:
        return ""
    text = re.sub(r"[（(]\s*整体\s*[）)]", "", text)
    text = text.replace("／", "/").replace("\\", "/").replace("＼", "/")
    text = re.sub(
        r"(第[一二三四五六七八九十百千\d]+(?:部分|章|节|条)[^/—－–-]*)\s*[—－–-]\s*",
        r"\1 / ",
        text,
    )
    text = re.sub(r"\s*/\s*", " / ", text)
    return re.sub(r"\s+", " ", text).strip(" /")


def _is_missing_review_locator(locator: str) -> bool:
    text = _normalize_review_text(locator)
    if not text:
        return False
    return any(keyword in text for keyword in _REVIEW_LOCATOR_MISSING_KEYWORDS)


def _extract_review_chapter_meta(value: str) -> dict:
    text = _normalize_review_locator(value)
    if not text:
        return {"text": "", "prefix": "", "ordinal": None, "title": ""}

    chapter_part = re.split(r"\s*/\s*", text)[0]
    match = _REVIEW_CHAPTER_PREFIX_RE.match(chapter_part)
    if not match:
        return {"text": chapter_part, "prefix": "", "ordinal": None, "title": chapter_part}

    ordinal_token = match.group(2)
    ordinal = ContractEngine._cn_numeral_to_int(ordinal_token) if not ordinal_token.isdigit() else int(ordinal_token)
    title = re.sub(r"^[—－–-]\s*", "", (match.group(3) or "").strip())
    return {
        "text": chapter_part,
        "prefix": match.group(1),
        "ordinal": ordinal,
        "title": title,
    }


def _parse_summary_rule_path(summary_pos: str) -> dict:
    text = _normalize_review_locator(summary_pos)
    if not text:
        return {
            "normalized_summary_path": "",
            "summary_group_path": "",
            "summary_group_heading": "",
            "summary_detail_heading": "",
        }

    parts = [part.strip() for part in re.split(r"\s*/\s*", text) if part and part.strip()]
    first = parts[0] if parts else text
    group_path = _summary_group_path_from_heading(first)
    group_heading = ""
    detail_heading = ""

    if first.startswith("摘要"):
        group_heading = re.sub(r"^摘要([一二三四五六七八九十百千\d]+)", r"\1", first).strip()
        if group_heading and not re.match(r"^[一二三四五六七八九十百千\d]+[、．]", group_heading):
            group_heading = ""
    else:
        group_heading = first if re.match(r"^[一二三四五六七八九十百千\d]+[、．]", first) else ""
        if not group_path:
            group_path = _summary_group_path_from_heading(first)

    if len(parts) > 1:
        detail_heading = parts[-1]

    normalized_path = group_path
    if detail_heading:
        normalized_path = f"{group_path} / {detail_heading}" if group_path else detail_heading
    elif not normalized_path:
        normalized_path = text

    return {
        "normalized_summary_path": normalized_path,
        "summary_group_path": group_path,
        "summary_group_heading": group_heading,
        "summary_detail_heading": detail_heading,
    }


def _summary_locator_group_matches(rule: dict, *, summary_group_path: str = "", summary_group_heading: str = "") -> bool:
    rule_group_path = _normalize_review_text(rule.get("summary_group_path"))
    if summary_group_path and rule_group_path and summary_group_path == rule_group_path:
        return True
    rule_group_heading = _normalize_review_text(rule.get("summary_group_heading"))
    if summary_group_heading and rule_group_heading and _review_labels_match(summary_group_heading, rule_group_heading):
        return True
    if summary_group_path and not rule_group_path and summary_group_heading:
        summary_pos = _normalize_review_text(rule.get("summary_pos"))
        if summary_pos and (_review_labels_match(summary_group_heading, summary_pos) or summary_group_path in summary_pos):
            return True
    return False


def _build_review_excerpt_pair(left_text: str, right_text: str, *, normalize_line=None, left_prefix: str, right_prefix: str) -> dict:
    normalize_line = normalize_line or (lambda value: value)
    left_rows = _collect_nonblank_compare_lines(left_text, normalize_line=normalize_line)
    right_rows = _collect_nonblank_compare_lines(right_text, normalize_line=normalize_line)
    left_lines = [row["line"] for row in left_rows]
    right_lines = [row["line"] for row in right_rows]
    left_compare = [row["compare"] for row in left_rows]
    right_compare = [row["compare"] for row in right_rows]

    matcher = difflib.SequenceMatcher(None, left_compare, right_compare, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue

        left_anchor_source = [line for line in left_lines[i1:i2] if line.strip()]
        right_anchor_source = [line for line in right_lines[j1:j2] if line.strip()]
        left_anchor = "\n".join(left_anchor_source[:2]).strip()
        right_anchor = "\n".join(right_anchor_source[:2]).strip()

        left_lo = max(0, (i1 if i1 < len(left_lines) else max(len(left_lines) - 1, 0)) - 2)
        left_hi = min(len(left_lines), max(i2, i1 + 1) + 2)
        right_lo = max(0, (j1 if j1 < len(right_lines) else max(len(right_lines) - 1, 0)) - 2)
        right_hi = min(len(right_lines), max(j2, j1 + 1) + 2)

        return {
            f"{left_prefix}_anchor_text": _truncate_review_excerpt(left_anchor or (left_lines[i1] if i1 < len(left_lines) else "")),
            f"{right_prefix}_anchor_text": _truncate_review_excerpt(right_anchor or (right_lines[j1] if j1 < len(right_lines) else "")),
            f"{left_prefix}_context_excerpt": _truncate_review_excerpt("\n".join(left_lines[left_lo:left_hi]), limit=520),
            f"{right_prefix}_context_excerpt": _truncate_review_excerpt("\n".join(right_lines[right_lo:right_hi]), limit=520),
        }

    return {
        f"{left_prefix}_anchor_text": "",
        f"{right_prefix}_anchor_text": "",
        f"{left_prefix}_context_excerpt": "",
        f"{right_prefix}_context_excerpt": "",
    }


def _split_contract_summary_subsections(summary_text: str) -> list[dict]:
    summary_content = _strip_contract_signing_page_text(summary_text or "")
    markers = []
    max_group_ordinal = 0
    for marker in _SUMMARY_SUBSECTION_RE.finditer(summary_content):
        ordinal = _summary_heading_ordinal(marker.group(1).strip())
        if ordinal is None:
            continue
        if not markers or ordinal > max_group_ordinal:
            markers.append(marker)
            max_group_ordinal = ordinal
    subsections = []
    for index, marker in enumerate(markers):
        start = marker.end()
        end = markers[index + 1].start() if index + 1 < len(markers) else len(summary_content)
        group_heading = marker.group(1).strip()
        group_content = summary_content[start:end].strip()
        group_path = _summary_group_path_from_heading(group_heading)
        detail_markers = list(_SUMMARY_DETAIL_HEADING_RE.finditer(group_content))
        detail_items = []
        for detail_index, detail_marker in enumerate(detail_markers):
            detail_start = detail_marker.end()
            detail_end = detail_markers[detail_index + 1].start() if detail_index + 1 < len(detail_markers) else len(group_content)
            detail_heading = detail_marker.group(1).strip()
            detail_items.append({
                "heading": detail_heading,
                "content": group_content[detail_start:detail_end].strip(),
                "summary_group_heading": group_heading,
                "summary_group_path": group_path,
                "summary_path": f"{group_path} / {detail_heading}" if group_path else detail_heading,
            })
        selected_detail_headings = _summary_group_selected_detail_headings(group_heading)
        if selected_detail_headings:
            detail_items = [
                item for item in detail_items
                if _summary_heading_key(item.get("heading")) in selected_detail_headings
            ]
            if not detail_items:
                selected_labels = _summary_group_selected_detail_labels(group_heading)
                if len(selected_labels) == 1 and group_content:
                    synthetic_heading = selected_labels[0]
                    detail_items = [{
                        "heading": synthetic_heading,
                        "content": group_content,
                        "summary_group_heading": group_heading,
                        "summary_group_path": group_path,
                        "summary_path": f"{group_path} / {synthetic_heading}" if group_path else synthetic_heading,
                    }]
            comparable_detail_items = [item for item in detail_items if str(item.get("content") or "").strip()]
            if comparable_detail_items:
                group_content = "\n".join(
                    part
                    for item in comparable_detail_items
                    for part in (item.get("heading"), item.get("content"))
                    if str(part or "").strip()
                ).strip()
        subsections.append({
            "heading": group_heading,
            "content": group_content,
            "summary_group_heading": group_heading,
            "summary_group_path": group_path,
            "summary_path": group_path or group_heading,
            "detail_items": detail_items,
        })
    return subsections


def _summary_content_similarity(summary_content: str, section_content: str) -> float:
    normalized_summary = "\n".join(
        row["compare"]
        for row in _collect_nonblank_compare_lines(
            summary_content,
            normalize_line=lambda line: _normalize_summary_compare_text(_strip_review_number_prefix(line)),
        )
    )
    normalized_section = "\n".join(
        row["compare"]
        for row in _collect_nonblank_compare_lines(
            section_content,
            normalize_line=lambda line: _normalize_summary_compare_text(_strip_review_number_prefix(line)),
        )
    )
    if not normalized_summary:
        return 0.0
    matcher = difflib.SequenceMatcher(None, normalized_summary, normalized_section, autojunk=False)
    matched = sum(block.size for block in matcher.get_matching_blocks())
    return matched / len(normalized_summary)


def _empty_review_target(missing_reason: str = "section_missing") -> dict:
    return {
        "section": None,
        "section_heading": "",
        "target_heading": "",
        "text": "",
        "body_text": "",
        "matched": False,
        "match_method": "",
        "missing_reason": missing_reason,
        "locator_parts": [],
    }


def _resolve_summary_rule_target(subsection, matched_rule, body_sections_all, candidate_sections, *, section_hint: str = "", prefer_subheading: bool = True):
    rule_source = "similarity_fallback"
    rule_level = _normalize_review_text((matched_rule or {}).get("rule_level"))
    rule_locator = _normalize_review_locator((matched_rule or {}).get("contract_pos"))

    if matched_rule:
        rule_source = "summary_rules"
        if rule_locator and not _is_missing_review_locator(rule_locator):
            if rule_level == "detail":
                direct_target = _locate_review_rule_target(body_sections_all, rule_locator, section_hint)
                if direct_target.get("matched"):
                    return (
                        direct_target,
                        _summary_content_similarity(subsection["content"], direct_target["text"]),
                        True,
                        rule_source,
                        rule_locator,
                        rule_level,
                    )
                return direct_target, 0.0, False, rule_source, rule_locator, rule_level

            chapter_target = _locate_review_rule_target(body_sections_all, rule_locator, section_hint)
            if chapter_target.get("matched") and not prefer_subheading:
                return (
                    chapter_target,
                    _summary_content_similarity(subsection["content"], chapter_target["text"]),
                    True,
                    rule_source,
                    rule_locator,
                    rule_level or "chapter",
                )
            if chapter_target.get("matched") and chapter_target.get("section"):
                sub_target = _locate_review_subheading(
                    chapter_target["section"],
                    subsection["heading"],
                    subsection["content"],
                    score_fn=_summary_content_similarity,
                )
                if sub_target.get("matched"):
                    return (
                        sub_target,
                        _summary_content_similarity(subsection["content"], sub_target["text"]),
                        True,
                        rule_source,
                        rule_locator,
                        rule_level or "chapter",
                    )
                sub_target["section"] = chapter_target.get("section")
                sub_target["section_heading"] = chapter_target.get("section_heading")
                sub_target["locator_parts"] = chapter_target.get("locator_parts") or []
                return sub_target, 0.0, False, rule_source, rule_locator, rule_level or "chapter"
            return chapter_target, 0.0, False, rule_source, rule_locator, rule_level or "chapter"

        return _empty_review_target("section_missing"), 0.0, False, rule_source, rule_locator, rule_level

    best_section = None
    best_score = -1.0
    for section in candidate_sections:
        score = _summary_content_similarity(subsection["content"], section.get("content", ""))
        if score > best_score:
            best_score = score
            best_section = section

    if best_section is None or best_score < SUMMARY_FALLBACK_SECTION_SCORE_THRESHOLD:
        return _empty_review_target("section_missing"), 0.0, False, rule_source, "", ""

    if not prefer_subheading:
        return (
            {
                "section": best_section,
                "section_heading": _normalize_review_text(best_section.get("heading")),
                "target_heading": _normalize_review_text(best_section.get("heading")),
                "text": (best_section.get("content") or "").strip(),
                "body_text": (best_section.get("content") or "").strip(),
                "matched": True,
                "match_method": "fallback",
                "missing_reason": "",
                "locator_parts": [_normalize_review_text(best_section.get("heading"))] if _normalize_review_text(best_section.get("heading")) else [],
            },
            best_score,
            True,
            rule_source,
            "",
            "",
        )

    sub_target = _locate_review_subheading(
        best_section,
        subsection["heading"],
        subsection["content"],
        score_fn=_summary_content_similarity,
    )
    if sub_target.get("matched"):
        fallback_score = _summary_content_similarity(subsection["content"], sub_target["text"])
        if fallback_score >= SUMMARY_FALLBACK_SUBHEADING_SCORE_THRESHOLD and sub_target.get("match_method") != "fallback":
            return sub_target, fallback_score, True, rule_source, "", ""

    sub_target["section"] = best_section
    sub_target["section_heading"] = _normalize_review_text(best_section.get("heading"))
    sub_target["target_heading"] = _normalize_review_text(subsection.get("heading"))
    sub_target["locator_parts"] = [_normalize_review_text(best_section.get("heading"))] if _normalize_review_text(best_section.get("heading")) else []
    return sub_target, 0.0, False, rule_source, "", ""


def _prepare_summary_review_context(contract_sections):
    summary_section = _find_contract_summary_section(contract_sections or [])
    if not summary_section:
        return {"error": "未找到基金合同内容摘要章节"}

    summary_heading = _normalize_review_text(summary_section.get("heading"))
    body_sections_all = [
        section
        for section in (contract_sections or [])
        if not _review_labels_match(section.get("heading"), summary_heading)
    ]
    body_sections = [section for section in body_sections_all if len((section.get("content") or "").strip()) > 100]
    subsections = _split_contract_summary_subsections(summary_section.get("content") or "")
    if not subsections:
        return {"error": "第二十六部分未找到子项标题（一、二、…）"}

    return {
        "summary_section": summary_section,
        "summary_subsections": subsections,
        "body_sections_all": body_sections_all,
        "candidate_sections": body_sections or body_sections_all,
    }


def _build_review_cross_rules(review_rules, fund_type="ETF"):
    def load_general_rules():
        rules = []
        rules_xlsx = _resolve_rules_xlsx()
        if not rules_xlsx.exists():
            return rules

        import openpyxl

        wb = openpyxl.load_workbook(str(rules_xlsx), read_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        wb.close()
        for row in rows:
            if not row or not row[0]:
                continue
            row_fund_type = str(row[0]).strip()
            if row_fund_type not in ("通用", fund_type):
                continue
            rules.append({
                "source": "general",
                "fund_type": row_fund_type,
                "contract_chapter": str(row[1] or "").strip(),
                "prospectus_chapter": str(row[2] or "").strip(),
                "relation": str(row[3] or "").strip(),
                "consistency": str(row[3] or "").strip(),
                "prompt": str(row[4] or "").strip(),
                "identical_items": str(row[5] or "").strip(),
                "diff_items": str(row[6] or "").strip(),
                "detail": str(row[6] or "").strip(),
            })
        return rules

    detailed = []
    chapter_rules = review_rules.get("chapter_rules") or review_rules.get("chapter_level", [])
    detail_rules = review_rules.get("detail_rules", [])
    key_diff_rules = review_rules.get("key_diff_rules") or review_rules.get("key_diffs", [])
    for row in chapter_rules:
        contract_chapter = str(row.get("contract") or "").strip()
        prospectus_chapter = str(row.get("prospectus") or "").strip()
        contract_section_name = (_split_review_locator(contract_chapter) or [contract_chapter])[0]
        prospectus_section_name = (_split_review_locator(prospectus_chapter) or [prospectus_chapter])[0]
        if not contract_chapter and not prospectus_chapter:
            continue
        matched_detail_rules = _find_matching_detail_rules(detail_rules, row)
        matched_key_diffs = _find_matching_key_diffs(key_diff_rules, row)
        if _review_chapter_prefers_detail_results(row, matched_detail_rules):
            for detail_rule in matched_detail_rules:
                detail_contract = str(detail_rule.get("contract") or "").strip()
                detail_prospectus_section = str(detail_rule.get("prospectus") or prospectus_section_name or prospectus_chapter).strip()
                detail_prospectus_point = str(detail_rule.get("prospectus_point") or "").strip()
                if detail_prospectus_point and not any(sep in detail_prospectus_point for sep in ("/", "／")):
                    detail_prospectus = f"{detail_prospectus_section} / {detail_prospectus_point}" if detail_prospectus_section else detail_prospectus_point
                else:
                    detail_prospectus = detail_prospectus_point or detail_prospectus_section
                detail_contract_section_name = (_split_review_locator(detail_contract) or [contract_section_name or detail_contract])[0]
                detail_prospectus_section_name = (_split_review_locator(detail_prospectus_section or detail_prospectus) or [prospectus_section_name or detail_prospectus])[0]
                detail_key_diffs = _find_matching_key_diffs(
                    key_diff_rules,
                    {
                        "contract": detail_contract or contract_chapter,
                        "prospectus": detail_prospectus or prospectus_chapter,
                    },
                )
                expected_diff_parts = []
                for text_part in (
                    str(detail_rule.get("detail") or "").strip(),
                    *(str(key_diff.get("description") or "").strip() for key_diff in detail_key_diffs),
                ):
                    if text_part and text_part not in expected_diff_parts:
                        expected_diff_parts.append(text_part)
                detailed.append({
                    "source": "detail_rules",
                    "contract_chapter": detail_contract or contract_chapter,
                    "prospectus_chapter": detail_prospectus or prospectus_chapter,
                    "contract_locator": detail_contract or contract_chapter,
                    "prospectus_locator": detail_prospectus or prospectus_chapter,
                    "contract_section_name": detail_contract_section_name or contract_section_name,
                    "prospectus_section_name": detail_prospectus_section_name or prospectus_section_name,
                    "relation": str(detail_rule.get("relation") or "").strip(),
                    "consistency": str(detail_rule.get("consistency") or "").strip(),
                    "detail": str(detail_rule.get("detail") or "").strip(),
                    "similarity": str(detail_rule.get("similarity") or "").strip(),
                    "sheet_name": str(detail_rule.get("sheet_name") or row.get("sheet_name") or "").strip(),
                    "expected_diff": "；".join(expected_diff_parts),
                    "suggestion": next(
                        (str(key_diff.get("suggestion") or "").strip() for key_diff in detail_key_diffs if str(key_diff.get("suggestion") or "").strip()),
                        "",
                    ),
                })
            continue

        expected_diff_parts = []
        if str(row.get("detail") or "").strip():
            expected_diff_parts.append(str(row.get("detail") or "").strip())
        for detail_rule in matched_detail_rules:
            detail_text = str(detail_rule.get("detail") or "").strip()
            if detail_text and detail_text not in expected_diff_parts:
                expected_diff_parts.append(detail_text)
        for key_diff in matched_key_diffs:
            description = str(key_diff.get("description") or "").strip()
            if description and description not in expected_diff_parts:
                expected_diff_parts.append(description)
        detailed.append({
            "source": "chapter_rules",
            "contract_chapter": contract_chapter,
            "prospectus_chapter": prospectus_chapter,
            "contract_locator": contract_chapter,
            "prospectus_locator": prospectus_chapter,
            "contract_section_name": contract_section_name,
            "prospectus_section_name": prospectus_section_name,
            "relation": str(row.get("relation") or "").strip(),
            "consistency": str(row.get("consistency") or "").strip(),
            "detail": str(row.get("detail") or "").strip(),
            "similarity": str(row.get("similarity") or "").strip(),
            "sheet_name": str(row.get("sheet_name") or "").strip(),
            "expected_diff": "；".join(expected_diff_parts),
            "suggestion": next(
                (str(key_diff.get("suggestion") or "").strip() for key_diff in matched_key_diffs if str(key_diff.get("suggestion") or "").strip()),
                "",
            ),
        })
    return detailed or load_general_rules()


def _build_ai_cross_review_items(contract_sections, prospectus_sections, review_rules, fund_type="ETF"):
    items = []
    seen = set()
    cross_rules = _build_review_cross_rules(review_rules, fund_type)
    if not cross_rules and fund_type == "ETF":
        cross_rules = _build_builtin_etf_cross_rules(contract_sections, prospectus_sections)
    for rule in cross_rules:
        contract_target = _locate_review_rule_target(
            contract_sections,
            rule.get("contract_locator") or rule.get("contract_chapter") or "",
            rule.get("contract_section_name") or rule.get("contract_chapter") or "",
        )
        if not contract_target.get("matched"):
            contract_target = _fallback_review_rule_target_from_text(
                contract_text,
                rule.get("contract_locator") or rule.get("contract_chapter") or "",
                rule.get("contract_section_name") or rule.get("contract_chapter") or "",
            )
        prospectus_target = _locate_review_rule_target(
            prospectus_sections,
            rule.get("prospectus_locator") or rule.get("prospectus_chapter") or "",
            rule.get("prospectus_section_name") or rule.get("prospectus_chapter") or "",
        )
        if not prospectus_target.get("matched"):
            prospectus_target = _fallback_review_rule_target_from_text(
                prospectus_text,
                rule.get("prospectus_locator") or rule.get("prospectus_chapter") or "",
                rule.get("prospectus_section_name") or rule.get("prospectus_chapter") or "",
            )
        if not contract_target.get("matched") or not prospectus_target.get("matched"):
            continue

        contract_locator = _format_review_target_locator(contract_target)
        prospectus_locator = _format_review_target_locator(prospectus_target)
        dedupe_key = (contract_locator, prospectus_locator)
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)

        contract_text = (contract_target.get("body_text") or contract_target.get("text") or "").strip()
        prospectus_text = (prospectus_target.get("body_text") or prospectus_target.get("text") or "").strip()
        if not contract_text or not prospectus_text:
            continue

        items.append({
            "label": f"合同↔招募: {contract_locator} ↔ {prospectus_locator}",
            "contract_locator": contract_locator,
            "prospectus_locator": prospectus_locator,
            "contract_text": _truncate_ai_review_text(contract_text),
            "prospectus_text": _truncate_ai_review_text(prospectus_text),
            "rule": rule,
        })
    return items


def _build_ai_summary_review_items(contract_sections, review_rules):
    context = _prepare_summary_review_context(contract_sections)
    if context.get("error"):
        return []

    items = []
    summary_rules = review_rules.get("summary_rules") or review_rules.get("summary") or []
    for subsection in context["summary_subsections"]:
        summary_group_path = subsection.get("summary_group_path") or ""
        summary_group_heading = subsection.get("summary_group_heading") or subsection["heading"]
        force_detail_compare = _summary_group_forces_detail_compare(summary_group_heading)
        chapter_rule = _find_matching_summary_rule(
            summary_rules,
            subsection["heading"],
            summary_group_path=summary_group_path,
            summary_group_heading=summary_group_heading,
            prefer_detail=False,
        )
        chapter_target, _, chapter_matched, _, _, _ = _resolve_summary_rule_target(
            subsection,
            chapter_rule,
            context["body_sections_all"],
            context["candidate_sections"],
            prefer_subheading=False,
        )

        detail_items = subsection.get("detail_items") or [subsection]
        for detail_item in detail_items:
            prefer_detail = True
            if force_detail_compare and not str(detail_item.get("content") or "").strip():
                continue
            detail_rule = _find_matching_summary_rule(
                summary_rules,
                detail_item["heading"],
                summary_group_path=summary_group_path,
                summary_group_heading=summary_group_heading,
                prefer_detail=prefer_detail,
            )
            matched_rule = detail_rule if _normalize_review_text((detail_rule or {}).get("rule_level")) == "detail" else None
            if matched_rule is None and force_detail_compare:
                matched_rule = chapter_rule
            method_text = _normalize_review_text((matched_rule or {}).get("method"))
            status_text = _normalize_review_text((matched_rule or {}).get("status"))
            if "未纳入摘要" in method_text or "未进入摘要" in method_text or "未纳入摘要" in status_text:
                continue

            target, _, locator_matched, _, rule_locator, rule_level = _resolve_summary_rule_target(
                detail_item,
                matched_rule,
                context["body_sections_all"],
                context["candidate_sections"],
                section_hint=chapter_target.get("section_heading") if chapter_matched else "",
                prefer_subheading=prefer_detail,
            )
            if not locator_matched:
                continue

            contract_locator = _format_review_target_locator(target)
            contract_text = (target.get("body_text") or target.get("text") or "").strip()
            summary_text = (detail_item.get("content") or "").strip()
            if not contract_text or not summary_text:
                continue

            items.append({
                "label": f"正文↔摘要: {contract_locator}",
                "summary_heading": detail_item["heading"],
                "summary_text": _truncate_ai_review_text(summary_text),
                "contract_locator": contract_locator,
                "contract_text": _truncate_ai_review_text(contract_text),
                "rule_locator": rule_locator,
                "rule_level": rule_level,
                "rule": matched_rule or {},
            })
    return items


def _infer_review_fund_type_from_text(text: str) -> str:
    rules_path = _review_store.get("rules_xlsx_path")
    if rules_path:
        return "联接基金" if "联接" in Path(rules_path).name else "ETF"

    text = text or ""
    linked_keywords = ("联接基金", "目标ETF", "发起式联接", "A类基金份额", "C类基金份额")
    return "联接基金" if any(keyword in text for keyword in linked_keywords) else "ETF"


REVIEW_SOFT_HEADING_MATCH_MIN_KEY_LENGTH = 16
REVIEW_SOFT_HEADING_MATCH_RATIO = 0.88
SUMMARY_FALLBACK_SECTION_SCORE_THRESHOLD = 0.55
SUMMARY_FALLBACK_SUBHEADING_SCORE_THRESHOLD = 0.6


def _fallback_review_filename(filename: str) -> str:
    stem = Path(str(filename or "")).stem.strip()
    if not stem:
        return ""
    stem = re.sub(r"(基金合同|招募说明书|托管协议|法律意见书|合同|说明书)$", "", stem).strip()
    return stem or Path(str(filename or "")).stem.strip()


def _infer_review_fund_name(text: str, filename: str = "") -> str:
    text = str(text or "")
    patterns = (
        r"([^\s，。；：“”\"《》]{4,80}证券投资基金)",
        r"([^\s，。；：“”\"《》]{4,80}交易型开放式指数证券投资基金)",
        r"([^\s，。；：“”\"《》]{4,80}联接基金)",
    )
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    return _fallback_review_filename(filename)


def _review_heading_key(value: str) -> str:
    text = unicodedata.normalize("NFKC", str(value or ""))
    text = re.sub(r'\([^)]*\)', '', text)
    text = re.sub(r'[（(][^）)]*[）)]', '', text)
    return re.sub(r'[\s的与及和、，,：:；;/-]', '', text)


def _review_label_variants(value: str) -> list[str]:
    text = unicodedata.normalize("NFKC", str(value or "")).strip()
    if not text:
        return []

    parts = [part.strip() for part in re.split(r'[+/／]', text) if part and part.strip()]
    variants = []
    for part in parts or [text]:
        normalized = part.strip()
        candidates = [normalized]
        candidates.append(re.sub(r'^摘要[一二三四五六七八九十百]+[、．]?\s*', '', normalized).strip())
        candidates.append(re.sub(r'^(?:[（(]?[一二三四五六七八九十百]+[）)]|[一二三四五六七八九十百]+[、．])\s*', '', normalized).strip())
        for candidate in candidates:
            candidate = candidate.strip(" /")
            if candidate and candidate not in variants:
                variants.append(candidate)
    return variants


def _review_label_ordinal(value: str):
    text = unicodedata.normalize("NFKC", str(value or "")).strip()
    if not text:
        return None

    patterns = (
        r'^(?:摘要)?([一二三四五六七八九十百千]+)[、．]',
        r'^(?:摘要)?[（(]([一二三四五六七八九十百千]+)[）)]',
        r'^(?:摘要)?(\d+)[.、．]',
        r'^(?:摘要)?[（(](\d+)[）)]',
    )
    for pattern in patterns:
        match = re.match(pattern, text)
        if not match:
            continue
        ordinal = match.group(1)
        return ContractEngine._cn_numeral_to_int(ordinal) if not ordinal.isdigit() else int(ordinal)
    return None


def _review_labels_match(left: str, right: str) -> bool:
    left_original_ordinal = _review_label_ordinal(left)
    right_original_ordinal = _review_label_ordinal(right)
    if left_original_ordinal is not None and right_original_ordinal is not None and left_original_ordinal != right_original_ordinal:
        return False

    left_variants = _review_label_variants(left)
    right_variants = _review_label_variants(right)
    for left_variant in left_variants:
        left_key = _review_heading_key(left_variant)
        left_ordinal = _review_label_ordinal(left_variant)
        if not left_key:
            continue
        for right_variant in right_variants:
            right_key = _review_heading_key(right_variant)
            right_ordinal = _review_label_ordinal(right_variant)
            if not right_key:
                continue
            if left_ordinal is not None and right_ordinal is not None and left_ordinal != right_ordinal:
                continue
            if left_key in right_key or right_key in left_key:
                return True
    return False


def _review_soft_heading_key(value: str) -> str:
    text = _review_heading_key(value)
    replacements = (
        ("上海证券交易所", "证券交易所"),
        ("深圳证券交易所", "证券交易所"),
        ("上海交易所", "交易所"),
        ("深圳交易所", "交易所"),
        ("本基金合同", "基金合同"),
        ("本基金", "基金"),
    )
    for source, target in replacements:
        text = text.replace(source, target)
    return text


def _review_soft_heading_match(left: str, right: str) -> bool:
    if _review_labels_match(left, right):
        return True

    left_ordinal = _review_label_ordinal(left)
    right_ordinal = _review_label_ordinal(right)
    if left_ordinal is not None and right_ordinal is not None and left_ordinal != right_ordinal:
        return False

    left_key = _review_soft_heading_key(left)
    right_key = _review_soft_heading_key(right)
    if not left_key or not right_key:
        return False
    if min(len(left_key), len(right_key)) < REVIEW_SOFT_HEADING_MATCH_MIN_KEY_LENGTH:
        return False
    ratio = difflib.SequenceMatcher(None, left_key, right_key, autojunk=False).ratio()
    return ratio >= REVIEW_SOFT_HEADING_MATCH_RATIO


def _match_review_section(sections, chapter_name):
    candidates = _review_label_variants(chapter_name)
    if not candidates:
        return {"section": None, "match_method": "", "score": 0.0}

    best = {"section": None, "match_method": "", "score": 0.0}
    method_rank = {"exact": 4, "normalized": 3, "chapter_title": 2, "soft_heading": 1, "": 0}

    def update(section, match_method, score):
        current_rank = method_rank.get(best["match_method"], 0)
        new_rank = method_rank.get(match_method, 0)
        if new_rank > current_rank or (new_rank == current_rank and score > best["score"]):
            best["section"] = section
            best["match_method"] = match_method
            best["score"] = score

    for candidate in candidates:
        candidate_text = _normalize_review_text(candidate)
        candidate_key = _review_heading_key(candidate)
        candidate_meta = _extract_review_chapter_meta(candidate)
        candidate_title_key = _review_soft_heading_key(candidate_meta.get("title"))
        for section in sections:
            heading = section.get("heading", "")
            heading_text = _normalize_review_text(heading)
            heading_key = _review_heading_key(heading)
            section_meta = _extract_review_chapter_meta(heading)
            section_title_key = _review_soft_heading_key(section_meta.get("title"))

            if candidate_text and heading_text and candidate_text == heading_text:
                update(section, "exact", float(len(candidate_text)))
                continue

            if candidate_key and heading_key and candidate_key == heading_key:
                update(section, "normalized", float(len(candidate_key)))
                continue

            if candidate_key and heading_key and (candidate_key in heading_key or heading_key in candidate_key):
                update(section, "normalized", float(min(len(candidate_key), len(heading_key))))
                continue

            if candidate and heading and (candidate in heading or heading in candidate):
                update(section, "normalized", float(min(len(candidate), len(heading))))
                continue

            if candidate_title_key and section_title_key:
                if candidate_title_key == section_title_key or candidate_title_key in section_title_key or section_title_key in candidate_title_key:
                    title_score = float(min(len(candidate_title_key), len(section_title_key)))
                    if candidate_meta.get("ordinal") is not None and candidate_meta.get("ordinal") == section_meta.get("ordinal"):
                        title_score += 0.5
                    update(section, "chapter_title", title_score)
                    continue

                if min(len(candidate_title_key), len(section_title_key)) >= 6:
                    title_ratio = difflib.SequenceMatcher(None, candidate_title_key, section_title_key, autojunk=False).ratio()
                    if title_ratio >= REVIEW_SOFT_HEADING_MATCH_RATIO:
                        if candidate_meta.get("ordinal") is not None and candidate_meta.get("ordinal") == section_meta.get("ordinal"):
                            title_ratio += 0.25
                        update(section, "chapter_title", title_ratio)
                        continue

            if _review_soft_heading_match(heading, candidate):
                ratio = difflib.SequenceMatcher(
                    None,
                    _review_soft_heading_key(heading),
                    _review_soft_heading_key(candidate),
                    autojunk=False,
                ).ratio()
                update(section, "soft_heading", ratio)

    return best


def _find_review_section(sections, chapter_name):
    return _match_review_section(sections, chapter_name)["section"]


def _custodian_summary_heading_direction(value: str) -> str:
    key = _review_heading_key(value)
    if "基金托管人对基金管理人" in key or "托管人对管理人" in key:
        return "custodian_to_manager"
    if "基金管理人对基金托管人" in key or "管理人对托管人" in key:
        return "manager_to_custodian"
    return ""


def _custodian_summary_heading_key(value: str) -> str:
    text = _review_soft_heading_key(value)
    replacements = (
        ("基金托管协议", "托管协议"),
        ("基金托管人", "托管人"),
        ("基金管理人", "管理人"),
        ("业务监督和核查", "监督核查"),
        ("业务核查", "核查"),
        ("基金财产的保管", "基金财产保管"),
        ("基金资产净值计算和会计核算", "基金资产净值计算会计核算"),
        ("基金份额持有人名册的保管", "基金份额持有人名册保管"),
        ("适用法律与争议解决方式", "适用法律争议解决方式"),
        ("基金托管协议的变更终止与基金财产的清算", "托管协议变更终止基金财产清算"),
    )
    for source, target in replacements:
        text = text.replace(source, target)
    return text


def _match_custodian_summary_section(sections, chapter_name):
    target_direction = _custodian_summary_heading_direction(chapter_name)
    candidates = _review_label_variants(chapter_name) or [str(chapter_name or "").strip()]
    best = {"section": None, "match_method": "", "score": 0.0}
    method_rank = {"exact": 4, "normalized": 3, "soft_heading": 2, "": 0}

    def update(section, match_method, score):
        current_rank = method_rank.get(best["match_method"], 0)
        new_rank = method_rank.get(match_method, 0)
        if new_rank > current_rank or (new_rank == current_rank and score > best["score"]):
            best["section"] = section
            best["match_method"] = match_method
            best["score"] = score

    for candidate in candidates:
        candidate_key = _custodian_summary_heading_key(candidate)
        candidate_ordinal = _review_label_ordinal(candidate)
        if not candidate_key:
            continue
        for section in sections or []:
            heading = str((section or {}).get("heading") or "").strip()
            heading_key = _custodian_summary_heading_key(heading)
            heading_ordinal = _review_label_ordinal(heading)
            section_direction = _custodian_summary_heading_direction(heading)
            if not heading_key:
                continue
            if target_direction and section_direction and section_direction != target_direction:
                continue
            if candidate_ordinal is not None and heading_ordinal is not None and candidate_ordinal != heading_ordinal:
                continue
            if candidate_key == heading_key:
                update(section, "exact", float(len(candidate_key)))
                continue
            if candidate_key in heading_key or heading_key in candidate_key:
                update(section, "normalized", float(min(len(candidate_key), len(heading_key))))
                continue
            if min(len(candidate_key), len(heading_key)) < 4:
                continue
            ratio = difflib.SequenceMatcher(None, candidate_key, heading_key, autojunk=False).ratio()
            if ratio >= 0.55:
                update(section, "soft_heading", ratio)
    return best


def _review_rule_chapter_labels(value: str) -> list[str]:
    text = _normalize_review_text(value)
    if not text:
        return []
    locator_parts = _split_review_locator(text)
    chapter_part = locator_parts[0] if locator_parts else text
    return [part.strip() for part in re.split(r"\s*[+＋]\s*", chapter_part) if part and part.strip()]


def _find_matching_summary_rule(summary_rules, summary_heading, *, summary_group_path: str = "", summary_group_heading: str = "", prefer_detail: bool = False):
    matched_rules = []
    for rule in summary_rules:
        summary_pos = _normalize_review_text(rule.get("summary_pos"))
        if not summary_pos:
            continue

        rule_level = _normalize_review_text(rule.get("rule_level"))
        detail_heading = _normalize_review_text(rule.get("summary_detail_heading"))
        group_heading = _normalize_review_text(rule.get("summary_group_heading"))
        group_path = _normalize_review_text(rule.get("summary_group_path"))
        normalized_path = _normalize_review_text(rule.get("normalized_summary_path"))
        group_matched = _summary_locator_group_matches(
            rule,
            summary_group_path=summary_group_path,
            summary_group_heading=summary_group_heading,
        )

        if prefer_detail:
            if rule_level == "detail":
                if summary_group_path and group_path and summary_group_path != group_path:
                    continue
                if not group_matched and summary_group_path:
                    continue
                if not detail_heading and _review_labels_match(summary_heading, group_heading or summary_pos):
                    matched_rules.append((3, rule))
                    continue
                if detail_heading and _review_labels_match(summary_heading, detail_heading):
                    matched_rules.append((4, rule))
                    continue
                if normalized_path and summary_group_path and normalized_path == f"{summary_group_path} / {summary_heading}":
                    matched_rules.append((5, rule))
                    continue
                continue

            if group_matched and _review_labels_match(summary_group_heading or summary_heading, group_heading or summary_pos):
                matched_rules.append((1, rule))
            continue

        if rule_level == "detail":
            continue
        if _review_labels_match(summary_heading, group_heading or summary_pos) or (
            group_matched and _review_labels_match(summary_group_heading or summary_heading, group_heading or summary_pos)
        ):
            matched_rules.append((3, rule))

    if not matched_rules:
        return None

    matched_rules.sort(
        key=lambda pair: (
            -pair[0],
            -len(_normalize_review_text(pair[1].get("contract_pos"))),
            -len(_normalize_review_text(pair[1].get("summary_pos"))),
        )
    )
    return matched_rules[0][1]


def _find_matching_detail_rules(detail_rules, chapter_rule):
    matched = []
    chapter_prospectus_labels = _review_rule_chapter_labels(chapter_rule.get("prospectus"))
    chapter_contract_labels = _review_rule_chapter_labels(chapter_rule.get("contract"))
    for detail_rule in detail_rules:
        detail_prospectus_labels = _review_rule_chapter_labels(detail_rule.get("prospectus"))
        detail_contract_labels = _review_rule_chapter_labels(detail_rule.get("contract"))
        if (
            len(chapter_prospectus_labels) == 1
            and len(detail_prospectus_labels) == 1
            and _review_labels_match(chapter_prospectus_labels[0], detail_prospectus_labels[0])
        ):
            matched.append(detail_rule)
            continue
        if (
            len(chapter_contract_labels) == 1
            and len(detail_contract_labels) == 1
            and _review_labels_match(chapter_contract_labels[0], detail_contract_labels[0])
        ):
            matched.append(detail_rule)
    return matched


def _find_matching_key_diffs(key_diff_rules, chapter_rule):
    matched = []
    for key_diff in key_diff_rules:
        location = _normalize_review_text(key_diff.get("location"))
        description = _normalize_review_text(key_diff.get("description"))
        haystack = " ".join(filter(None, [location, description]))
        if not haystack:
            continue
        if _review_labels_match(chapter_rule.get("contract"), haystack) or _review_labels_match(chapter_rule.get("prospectus"), haystack):
            matched.append(key_diff)
    return matched


def _review_chapter_prefers_detail_results(chapter_rule, matched_detail_rules):
    if not matched_detail_rules:
        return False
    contract_label = _normalize_review_text(chapter_rule.get("contract"))
    prospectus_label = _normalize_review_text(chapter_rule.get("prospectus"))
    if any(separator in contract_label for separator in ("+", "＋")):
        return False
    if any(separator in prospectus_label for separator in ("+", "＋")):
        return False
    priority_labels = (
        "基金份额的申购赎回",
        "基金份额的申购与赎回",
        "基金份额的上市交易",
    )
    for label in priority_labels:
        if _review_labels_match(contract_label, label) or _review_labels_match(prospectus_label, label):
            return True
    return False


def _split_review_locator(locator: str) -> list[str]:
    return [part.strip() for part in re.split(r"\s*/\s*", _normalize_review_locator(locator)) if part and part.strip()]


def _format_review_locator_path(section_heading: str = "", target_heading: str = "") -> str:
    section_heading = _normalize_review_text(section_heading)
    target_heading = _normalize_review_text(target_heading)
    if section_heading and target_heading and not _review_labels_match(section_heading, target_heading):
        return f"{section_heading} / {target_heading}"
    return section_heading or target_heading


def _format_review_target_locator(target: dict | None) -> str:
    if not target:
        return ""
    locator_parts = [_normalize_review_text(part) for part in (target.get("locator_parts") or []) if _normalize_review_text(part)]
    if locator_parts:
        deduped_parts = []
        for part in locator_parts:
            if not deduped_parts:
                deduped_parts.append(part)
                continue
            if _review_labels_match(deduped_parts[-1], part):
                continue
            deduped_parts.append(part)
        return " / ".join(deduped_parts)
    return _format_review_locator_path(target.get("section_heading"), target.get("target_heading"))


_DOCX_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_DOCX_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_DOCX_NS = {"w": _DOCX_W_NS}


_REVIEW_REVISION_SCOPE_LABELS = {
    "contract_body": "基金合同正文",
    "contract_summary": "基金合同摘要",
    "prospectus_body": "招募说明书",
    "prospectus_contract_summary": "招募说明书内基金合同摘要",
    "custody_agreement": "托管协议",
    "product_summary": "产品资料概要",
}


_REVIEW_REVISION_IMPACT_TARGETS = {
    "contract_body": ["contract_summary", "prospectus_body", "prospectus_contract_summary", "product_summary"],
    "contract_summary": ["prospectus_contract_summary", "product_summary"],
    "prospectus_body": ["product_summary"],
    "prospectus_contract_summary": ["product_summary"],
    "custody_agreement": ["prospectus_body"],
    "product_summary": [],
}


def _docx_local_name(tag: str) -> str:
    return str(tag or "").rsplit("}", 1)[-1]


def _docx_attr(node, name: str) -> str:
    return str(node.attrib.get(f"{{{_DOCX_W_NS}}}{name}") or node.attrib.get(name) or "").strip()


def _docx_xml_text(node) -> str:
    parts = []
    for child in node.iter():
        name = _docx_local_name(child.tag)
        if name in {"t", "delText"} and child.text:
            parts.append(str(child.text))
        elif name == "tab":
            parts.append("\t")
        elif name in {"br", "cr"}:
            parts.append("\n")
    return "".join(parts).strip()


def _read_docx_xml_part(docx_bytes: bytes, part_name: str) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
            with zf.open(part_name) as part:
                return part.read().decode("utf-8")
    except (KeyError, zipfile.BadZipFile, UnicodeDecodeError):
        return ""


def _load_docx_comments(docx_bytes: bytes) -> dict[str, dict]:
    xml_text = _read_docx_xml_part(docx_bytes, "word/comments.xml")
    if not xml_text:
        return {}
    import xml.etree.ElementTree as ET

    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return {}

    comments = {}
    for comment in root.findall(".//w:comment", _DOCX_NS):
        comment_id = _docx_attr(comment, "id")
        if not comment_id:
            continue
        comments[comment_id] = {
            "author": _docx_attr(comment, "author"),
            "date": _docx_attr(comment, "date"),
            "comment_text": _docx_xml_text(comment),
        }
    return comments


def _review_revision_base_document_kind(document_kind: str) -> str:
    kind = str(document_kind or "").strip()
    if kind.startswith("product_summary"):
        return "product_summary"
    return kind


def _review_revision_default_scope(document_kind: str) -> str:
    kind = _review_revision_base_document_kind(document_kind)
    if kind == "contract":
        return "contract_body"
    if kind == "prospectus":
        return "prospectus_body"
    if kind == "product_summary":
        return "product_summary"
    if kind == "custody_agreement":
        return "custody_agreement"
    return kind or "unknown"


def _review_revision_scope_for_paragraph(document_kind: str, paragraph_text: str, current_scope: str) -> str:
    document_kind = _review_revision_base_document_kind(document_kind)
    text = _normalize_review_text(paragraph_text)
    compact = re.sub(r"\s+", "", text)
    if document_kind == "contract" and re.fullmatch(r"第二十六部分基金合同内容摘要", compact):
        return "contract_summary"
    if document_kind == "prospectus":
        appendix_prefix = r"(?:附件[一二三四五六七八九十]+[、:：]?)?"
        numbered_prefix = r"(?:[一二三四五六七八九十]+[、:：])?"
        if re.fullmatch(rf"{appendix_prefix}{numbered_prefix}基金合同内容摘要", compact):
            return "prospectus_contract_summary"
        if re.fullmatch(rf"{appendix_prefix}{numbered_prefix}(?:基金)?托管协议内容摘要", compact):
            return "prospectus_body"
    return current_scope


def _review_revision_impact_targets(scope: str) -> list[dict]:
    return [
        {"key": target, "label": _REVIEW_REVISION_SCOPE_LABELS.get(target, target)}
        for target in _REVIEW_REVISION_IMPACT_TARGETS.get(scope, [])
    ]


def _build_review_revision_item(
    *,
    document_kind: str,
    scope: str,
    filename: str,
    revision_type: str,
    text: str,
    line: int,
    context: str,
    author: str = "",
    date: str = "",
    comment_text: str = "",
) -> dict:
    label = _REVIEW_REVISION_SCOPE_LABELS.get(scope, scope or document_kind)
    impact_targets = _review_revision_impact_targets(scope)
    path_labels = [label, *[target["label"] for target in impact_targets]]
    return {
        "id": "",
        "document_kind": document_kind,
        "document_scope": scope,
        "document_label": label,
        "source_file": filename,
        "revision_type": revision_type,
        "text": text,
        "comment_text": comment_text,
        "line": line,
        "context": context,
        "author": author,
        "date": date,
        "decision": "pending",
        "impact_targets": impact_targets,
        "impact_path": " → ".join(path_labels),
    }


def _review_revision_label_path(scope: str) -> str:
    label = _REVIEW_REVISION_SCOPE_LABELS.get(scope, scope)
    targets = _review_revision_impact_targets(scope)
    return " → ".join([label, *[target["label"] for target in targets]])


def _build_review_reader_revision(
    *,
    revision_id: str,
    document_kind: str,
    scope: str,
    filename: str,
    revision_type: str,
    block_id: str,
    context: str,
    text: str = "",
    old_text: str = "",
    new_text: str = "",
    author: str = "",
    date: str = "",
    comment_text: str = "",
) -> dict:
    label = _REVIEW_REVISION_SCOPE_LABELS.get(scope, scope or document_kind)
    return {
        "id": revision_id,
        "document_kind": document_kind,
        "document_scope": scope,
        "document_label": label,
        "source_file": filename,
        "revision_type": revision_type,
        "text": text,
        "old_text": old_text,
        "new_text": new_text,
        "comment_text": comment_text,
        "block_id": block_id,
        "context": context,
        "author": author,
        "date": date,
        "decision": "pending",
        "impact_targets": _review_revision_impact_targets(scope),
        "impact_path": _review_revision_label_path(scope),
    }


def _append_review_reader_text_run(
    runs: list[dict],
    *,
    text: str,
    mark_type: str = "normal",
    revision_type: str = "",
    author: str = "",
    date: str = "",
    comment_id: str = "",
    comment_text: str = "",
) -> None:
    if not text:
        return
    runs.append(
        {
            "text": text,
            "mark_type": mark_type,
            "revision_type": revision_type,
            "author": author,
            "date": date,
            "comment_id": comment_id,
            "comment_text": comment_text,
        }
    )


def _append_review_reader_plain_text_run(
    runs: list[dict],
    text: str,
    active_comments: list[str],
    comments: dict[str, dict],
) -> None:
    if not text:
        return
    comment_id = active_comments[-1] if active_comments else ""
    if comment_id and comment_id in comments:
        comment = comments[comment_id]
        _append_review_reader_text_run(
            runs,
            text=text,
            mark_type="comment",
            revision_type="comment",
            author=comment.get("author", ""),
            date=comment.get("date", ""),
            comment_id=comment_id,
            comment_text=comment.get("comment_text", ""),
        )
    else:
        _append_review_reader_text_run(runs, text=text)


def _append_review_reader_inline_runs(node, runs: list[dict], active_comments: list[str], comments: dict[str, dict]) -> None:
    name = _docx_local_name(node.tag)
    if name in {"ins", "del", "moveTo", "moveFrom"}:
        revision_text = _docx_xml_text(node)
        if revision_text:
            revision_type = "insert" if name in {"ins", "moveTo"} else "delete"
            _append_review_reader_text_run(
                runs,
                text=revision_text,
                mark_type=revision_type,
                revision_type=revision_type,
                author=_docx_attr(node, "author"),
                date=_docx_attr(node, "date"),
            )
        return
    if name == "commentRangeStart":
        comment_id = _docx_attr(node, "id")
        if comment_id:
            active_comments.append(comment_id)
        return
    if name == "commentRangeEnd":
        comment_id = _docx_attr(node, "id")
        if comment_id in active_comments:
            active_comments[:] = [cid for cid in active_comments if cid != comment_id]
        elif active_comments:
            active_comments.pop()
        return
    if name in {"t", "delText"}:
        _append_review_reader_plain_text_run(runs, str(node.text or ""), active_comments, comments)
        return
    if name == "tab":
        _append_review_reader_plain_text_run(runs, "\t", active_comments, comments)
        return
    if name in {"br", "cr"}:
        _append_review_reader_plain_text_run(runs, "\n", active_comments, comments)
        return

    for child in list(node):
        _append_review_reader_inline_runs(child, runs, active_comments, comments)


def _parse_review_reader_paragraph_runs(paragraph, comments: dict[str, dict]) -> list[dict]:
    runs: list[dict] = []
    active_comments: list[str] = []

    for child in list(paragraph):
        name = _docx_local_name(child.tag)
        if name == "commentRangeStart":
            comment_id = _docx_attr(child, "id")
            if comment_id:
                active_comments.append(comment_id)
            continue
        if name == "commentRangeEnd":
            comment_id = _docx_attr(child, "id")
            if comment_id in active_comments:
                active_comments = [cid for cid in active_comments if cid != comment_id]
            elif active_comments:
                active_comments.pop()
            continue
        if name in {"ins", "del", "moveTo", "moveFrom"}:
            revision_text = _docx_xml_text(child)
            if not revision_text:
                continue
            revision_type = "insert" if name in {"ins", "moveTo"} else "delete"
            _append_review_reader_text_run(
                runs,
                text=revision_text,
                mark_type=revision_type,
                revision_type=revision_type,
                author=_docx_attr(child, "author"),
                date=_docx_attr(child, "date"),
            )
            continue
        if name in {"r", "hyperlink", "sdt"}:
            _append_review_reader_inline_runs(child, runs, active_comments, comments)

    return runs


def _finalize_review_reader_runs(
    *,
    runs: list[dict],
    document_kind: str,
    scope: str,
    filename: str,
    block_id: str,
    context: str,
    revision_counter: int,
) -> tuple[list[dict], list[dict], int]:
    revisions: list[dict] = []
    finalized = [dict(run) for run in runs]
    index = 0
    comment_revision_ids: dict[str, str] = {}

    while index < len(finalized):
        run = finalized[index]
        mark_type = run.get("mark_type")
        next_run = finalized[index + 1] if index + 1 < len(finalized) else None
        next_type = (next_run or {}).get("mark_type")
        is_modify_pair = {mark_type, next_type} == {"delete", "insert"}

        if is_modify_pair:
            revision_counter += 1
            revision_id = f"{document_kind}-reader-{revision_counter}"
            from_run = run if mark_type == "delete" else next_run
            to_run = run if mark_type == "insert" else next_run
            run["mark_type"] = "modify_from" if mark_type == "delete" else "modify_to"
            next_run["mark_type"] = "modify_from" if next_type == "delete" else "modify_to"
            run["revision_id"] = revision_id
            next_run["revision_id"] = revision_id
            revisions.append(
                _build_review_reader_revision(
                    revision_id=revision_id,
                    document_kind=document_kind,
                    scope=scope,
                    filename=filename,
                    revision_type="modify",
                    block_id=block_id,
                    context=context,
                    text=f"{from_run.get('text', '')} → {to_run.get('text', '')}",
                    old_text=from_run.get("text", ""),
                    new_text=to_run.get("text", ""),
                    author=to_run.get("author") or from_run.get("author", ""),
                    date=to_run.get("date") or from_run.get("date", ""),
                )
            )
            index += 2
            continue

        if mark_type in {"insert", "delete"}:
            revision_counter += 1
            revision_id = f"{document_kind}-reader-{revision_counter}"
            run["revision_id"] = revision_id
            revisions.append(
                _build_review_reader_revision(
                    revision_id=revision_id,
                    document_kind=document_kind,
                    scope=scope,
                    filename=filename,
                    revision_type=mark_type,
                    block_id=block_id,
                    context=context,
                    text=run.get("text", ""),
                    author=run.get("author", ""),
                    date=run.get("date", ""),
                )
            )
        elif mark_type == "comment":
            comment_id = run.get("comment_id") or f"{block_id}-comment-{index}"
            revision_id = comment_revision_ids.get(comment_id)
            if not revision_id:
                revision_counter += 1
                revision_id = f"{document_kind}-reader-{revision_counter}"
                comment_revision_ids[comment_id] = revision_id
                revisions.append(
                    _build_review_reader_revision(
                        revision_id=revision_id,
                        document_kind=document_kind,
                        scope=scope,
                        filename=filename,
                        revision_type="comment",
                        block_id=block_id,
                        context=context,
                        text=run.get("text", ""),
                        author=run.get("author", ""),
                        date=run.get("date", ""),
                        comment_text=run.get("comment_text", ""),
                    )
                )
            run["revision_id"] = revision_id
        index += 1

    for run in finalized:
        for key in ("revision_type", "comment_id"):
            run.pop(key, None)
    return finalized, revisions, revision_counter


def _build_review_revision_paragraph_block(
    *,
    paragraph,
    document_kind: str,
    filename: str,
    block_id: str,
    scope: str,
    comments: dict[str, dict],
    revision_counter: int,
) -> tuple[dict, list[dict], int, str]:
    runs = _parse_review_reader_paragraph_runs(paragraph, comments)
    text = "".join(run.get("text", "") for run in runs).strip()
    next_scope = _review_revision_scope_for_paragraph(document_kind, text, scope) if text else scope
    finalized_runs, revisions, revision_counter = _finalize_review_reader_runs(
        runs=runs,
        document_kind=document_kind,
        scope=next_scope,
        filename=filename,
        block_id=block_id,
        context=text,
        revision_counter=revision_counter,
    )
    block = {
        "id": block_id,
        "type": "paragraph",
        "document_scope": next_scope,
        "document_label": _REVIEW_REVISION_SCOPE_LABELS.get(next_scope, next_scope),
        "text": text,
        "runs": finalized_runs,
        "revision_ids": [revision["id"] for revision in revisions],
    }
    return block, revisions, revision_counter, next_scope


def _build_review_revision_table_block(
    *,
    table,
    document_kind: str,
    filename: str,
    block_id: str,
    scope: str,
    comments: dict[str, dict],
    revision_counter: int,
) -> tuple[dict, list[dict], int, str]:
    rows = []
    revisions: list[dict] = []
    current_scope = scope
    cell_texts: list[str] = []
    revision_ids: list[str] = []

    for row_index, row in enumerate(table.findall(".//w:tr", _DOCX_NS), start=1):
        cells = []
        for cell_index, cell in enumerate(row.findall("./w:tc", _DOCX_NS), start=1):
            cell_blocks = []
            for paragraph_index, paragraph in enumerate(cell.findall("./w:p", _DOCX_NS), start=1):
                child_block_id = f"{block_id}-r{row_index}c{cell_index}p{paragraph_index}"
                paragraph_block, paragraph_revisions, revision_counter, current_scope = _build_review_revision_paragraph_block(
                    paragraph=paragraph,
                    document_kind=document_kind,
                    filename=filename,
                    block_id=child_block_id,
                    scope=current_scope,
                    comments=comments,
                    revision_counter=revision_counter,
                )
                cell_blocks.append(paragraph_block)
                revisions.extend(paragraph_revisions)
                revision_ids.extend(paragraph_block["revision_ids"])
                if paragraph_block.get("text"):
                    cell_texts.append(paragraph_block["text"])
            cells.append({"blocks": cell_blocks, "text": "\n".join(block.get("text", "") for block in cell_blocks if block.get("text"))})
        rows.append(cells)

    block = {
        "id": block_id,
        "type": "table",
        "document_scope": current_scope,
        "document_label": _REVIEW_REVISION_SCOPE_LABELS.get(current_scope, current_scope),
        "text": "\n".join(cell_texts),
        "rows": rows,
        "revision_ids": revision_ids,
    }
    return block, revisions, revision_counter, current_scope


def _build_review_revision_document_view(docx_source, *, document_kind: str, filename: str = "") -> dict:
    if hasattr(docx_source, "read"):
        current_pos = None
        try:
            current_pos = docx_source.tell()
        except Exception:
            current_pos = None
        docx_bytes = docx_source.read()
        if current_pos is not None:
            try:
                docx_source.seek(current_pos)
            except Exception:
                pass
    else:
        docx_bytes = bytes(docx_source or b"")

    import xml.etree.ElementTree as ET

    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        return {
            "document_kind": document_kind,
            "filename": filename,
            "blocks": [],
            "revisions": [],
            "revision_counts": _count_review_revision_items([]),
        }

    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return {
            "document_kind": document_kind,
            "filename": filename,
            "blocks": [],
            "revisions": [],
            "revision_counts": _count_review_revision_items([]),
        }

    body = root.find("w:body", _DOCX_NS)
    comments = _load_docx_comments(docx_bytes)
    blocks = []
    revisions: list[dict] = []
    scope = _review_revision_default_scope(document_kind)
    revision_counter = 0
    block_counter = 0

    for child in list(body) if body is not None else []:
        name = _docx_local_name(child.tag)
        if name not in {"p", "tbl"}:
            continue
        block_counter += 1
        block_id = f"{document_kind}-block-{block_counter}"
        if name == "p":
            block, block_revisions, revision_counter, scope = _build_review_revision_paragraph_block(
                paragraph=child,
                document_kind=document_kind,
                filename=filename,
                block_id=block_id,
                scope=scope,
                comments=comments,
                revision_counter=revision_counter,
            )
        else:
            block, block_revisions, revision_counter, scope = _build_review_revision_table_block(
                table=child,
                document_kind=document_kind,
                filename=filename,
                block_id=block_id,
                scope=scope,
                comments=comments,
                revision_counter=revision_counter,
            )
        if block.get("text") or block.get("type") == "table":
            blocks.append(block)
        revisions.extend(block_revisions)

    return {
        "document_kind": document_kind,
        "filename": filename,
        "blocks": blocks,
        "revisions": revisions,
        "revision_counts": _count_review_revision_items(revisions),
    }


def _normalize_review_export_decision(value: str) -> str:
    raw = str(value or "").strip().lower()
    if raw in {"accept", "accepted", "接受", "已接受"}:
        return "accept"
    if raw in {"reject", "rejected", "refuse", "拒绝", "已拒绝"}:
        return "reject"
    return "pending"


def _register_docx_xml_namespaces(ET) -> None:
    namespaces = {
        "w": _DOCX_W_NS,
        "r": _DOCX_R_NS,
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
        "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
        "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    }
    for prefix, uri in namespaces.items():
        try:
            ET.register_namespace(prefix, uri)
        except ValueError:
            pass


def _copy_docx_element(ET, node):
    return ET.fromstring(ET.tostring(node, encoding="utf-8"))


def _docx_root_start_tag(xml_text: str) -> str:
    match = re.search(r"<(?![!?])[^>]+>", str(xml_text or ""))
    return match.group(0) if match else ""


def _serialize_docx_document_xml(ET, root, original_xml_text: str | bytes = "") -> bytes:
    rendered = ET.tostring(root, encoding="utf-8", xml_declaration=True).decode("utf-8")
    original = bytes(original_xml_text).decode("utf-8", errors="ignore") if isinstance(original_xml_text, (bytes, bytearray)) else str(original_xml_text or "")
    original_start = _docx_root_start_tag(original)
    rendered_start = _docx_root_start_tag(rendered)
    if not original_start or not rendered_start:
        return rendered.encode("utf-8")

    missing_declarations = []
    declaration_names = set()
    for declaration in re.findall(r'xmlns(?::[\w.\-]+)?="[^"]+"', original_start):
        declaration_name = declaration.split("=", 1)[0]
        if f"{declaration_name}=" not in rendered_start:
            missing_declarations.append(declaration)
            declaration_names.add(declaration_name)

    ignorable_match = re.search(r'\bmc:Ignorable="([^"]*)"', rendered_start)
    if ignorable_match:
        known_ignorable_namespaces = {
            "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
            "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
            "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
            "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
        }
        for prefix in ignorable_match.group(1).split():
            declaration_name = f"xmlns:{prefix}"
            if f"{declaration_name}=" in rendered_start or declaration_name in declaration_names:
                continue
            uri = known_ignorable_namespaces.get(prefix)
            if uri:
                missing_declarations.append(f'{declaration_name}="{uri}"')
                declaration_names.add(declaration_name)
    if not missing_declarations:
        return rendered.encode("utf-8")

    insert_at = rendered.find(">", rendered.find(rendered_start))
    if insert_at < 0:
        return rendered.encode("utf-8")
    rendered = f"{rendered[:insert_at]} {' '.join(missing_declarations)}{rendered[insert_at:]}"
    return rendered.encode("utf-8")


def _convert_deleted_run_to_plain_text(run) -> None:
    for child in run.iter():
        if _docx_local_name(child.tag) == "delText":
            child.tag = f"{{{_DOCX_W_NS}}}t"


def _deleted_revision_to_plain_runs(ET, revision_node) -> list:
    runs = []
    for run in revision_node.findall(".//w:r", _DOCX_NS):
        copied = _copy_docx_element(ET, run)
        _convert_deleted_run_to_plain_text(copied)
        runs.append(copied)
    if runs:
        return runs

    text = _docx_xml_text(revision_node)
    if not text:
        return []
    run = ET.Element(f"{{{_DOCX_W_NS}}}r")
    text_node = ET.SubElement(run, f"{{{_DOCX_W_NS}}}t")
    text_node.text = text
    return [run]


def _deleted_revision_to_plain_text_nodes(ET, revision_node) -> list:
    text = _docx_xml_text(revision_node)
    if not text:
        return []
    text_node = ET.Element(f"{{{_DOCX_W_NS}}}t")
    text_node.text = text
    return [text_node]


def _replace_docx_child(parent, child, replacements: list) -> None:
    children = list(parent)
    try:
        index = children.index(child)
    except ValueError:
        return
    parent.remove(child)
    for offset, replacement in enumerate(replacements or []):
        parent.insert(index + offset, replacement)


def _remove_docx_child(parent, child) -> None:
    try:
        parent.remove(child)
    except ValueError:
        pass


def _docx_revision_token_from_child(child, active_comments: list[str], comments: dict[str, dict]) -> dict | None:
    name = _docx_local_name(child.tag)
    if name in {"ins", "del", "moveTo", "moveFrom"}:
        revision_text = _docx_xml_text(child)
        if not revision_text:
            return None
        revision_type = "insert" if name in {"ins", "moveTo"} else "delete"
        return {
            "mark_type": revision_type,
            "revision_type": revision_type,
            "text": revision_text,
            "element": child,
            "author": _docx_attr(child, "author"),
            "date": _docx_attr(child, "date"),
        }
    if name in {"r", "hyperlink", "sdt"}:
        text = _docx_xml_text(child)
        if not text:
            return None
        comment_id = active_comments[-1] if active_comments else ""
        if comment_id and comment_id in comments:
            return {
                "mark_type": "comment",
                "revision_type": "comment",
                "text": text,
                "element": child,
                "comment_id": comment_id,
                "comment_text": comments[comment_id].get("comment_text", ""),
                "author": comments[comment_id].get("author", ""),
                "date": comments[comment_id].get("date", ""),
            }
        return {"mark_type": "normal", "text": text, "element": child}
    return None


def _append_review_export_text_token(
    tokens: list[dict],
    *,
    text: str,
    element,
    parent,
    active_comments: list[str],
    comments: dict[str, dict],
) -> None:
    if not text:
        return
    comment_id = active_comments[-1] if active_comments else ""
    if comment_id and comment_id in comments:
        tokens.append({
            "mark_type": "comment",
            "revision_type": "comment",
            "text": text,
            "element": element,
            "parent": parent,
            "comment_id": comment_id,
            "comment_text": comments[comment_id].get("comment_text", ""),
            "author": comments[comment_id].get("author", ""),
            "date": comments[comment_id].get("date", ""),
        })
        return
    tokens.append({"mark_type": "normal", "text": text, "element": element, "parent": parent})


def _append_review_export_tokens_from_node(
    node,
    parent,
    tokens: list[dict],
    active_comments: list[str],
    comments: dict[str, dict],
) -> None:
    name = _docx_local_name(node.tag)
    if name in {"ins", "del", "moveTo", "moveFrom"}:
        revision_text = _docx_xml_text(node)
        if revision_text:
            revision_type = "insert" if name in {"ins", "moveTo"} else "delete"
            tokens.append({
                "mark_type": revision_type,
                "revision_type": revision_type,
                "text": revision_text,
                "element": node,
                "parent": parent,
                "author": _docx_attr(node, "author"),
                "date": _docx_attr(node, "date"),
            })
        return
    if name == "commentRangeStart":
        comment_id = _docx_attr(node, "id")
        if comment_id:
            active_comments.append(comment_id)
        return
    if name == "commentRangeEnd":
        comment_id = _docx_attr(node, "id")
        if comment_id in active_comments:
            active_comments[:] = [cid for cid in active_comments if cid != comment_id]
        elif active_comments:
            active_comments.pop()
        return
    if name in {"t", "delText"}:
        _append_review_export_text_token(
            tokens,
            text=str(node.text or ""),
            element=node,
            parent=parent,
            active_comments=active_comments,
            comments=comments,
        )
        return
    if name == "tab":
        _append_review_export_text_token(
            tokens,
            text="\t",
            element=node,
            parent=parent,
            active_comments=active_comments,
            comments=comments,
        )
        return
    if name in {"br", "cr"}:
        _append_review_export_text_token(
            tokens,
            text="\n",
            element=node,
            parent=parent,
            active_comments=active_comments,
            comments=comments,
        )
        return

    for child in list(node):
        _append_review_export_tokens_from_node(child, node, tokens, active_comments, comments)


def _review_export_paragraph_tokens(paragraph, comments: dict[str, dict]) -> list[dict]:
    tokens: list[dict] = []
    active_comments: list[str] = []
    for child in list(paragraph):
        _append_review_export_tokens_from_node(child, paragraph, tokens, active_comments, comments)
    return tokens


def _apply_rejected_docx_revision(parent, token: dict, ET) -> None:
    mark_type = token.get("mark_type")
    element = token.get("element")
    token_parent = token.get("parent")
    if token_parent is not None:
        parent = token_parent
    if element is None:
        return
    if mark_type == "insert":
        _remove_docx_child(parent, element)
    elif mark_type == "delete":
        if _docx_local_name(getattr(parent, "tag", "")) == "r":
            replacements = _deleted_revision_to_plain_text_nodes(ET, element)
        else:
            replacements = _deleted_revision_to_plain_runs(ET, element)
        _replace_docx_child(parent, element, replacements)


def _strip_docx_comment_markers(root, comment_ids: set[str]) -> None:
    if not comment_ids:
        return
    for parent in root.iter():
        for child in list(parent):
            name = _docx_local_name(child.tag)
            if name in {"commentRangeStart", "commentRangeEnd"} and _docx_attr(child, "id") in comment_ids:
                parent.remove(child)
                continue
            if name == "r":
                references = child.findall(".//w:commentReference", _DOCX_NS)
                if any(_docx_attr(ref, "id") in comment_ids for ref in references):
                    parent.remove(child)


def _copy_docx_run_properties(ET, run):
    if run is None:
        return None
    r_pr = run.find("w:rPr", _DOCX_NS)
    return _copy_docx_element(ET, r_pr) if r_pr is not None else None


DEFAULT_REVIEW_AUTHOR = "审核工作台"


def _normalize_review_author(value) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", str(value or ""))
    text = re.sub(r"\s+", " ", text).strip()
    return (text[:80].strip() or DEFAULT_REVIEW_AUTHOR)


def _normalize_review_revision_edits(revision_edits: dict | None) -> dict[str, dict[str, str]]:
    if not isinstance(revision_edits, dict):
        return {}
    normalized: dict[str, dict[str, str]] = {}
    for revision_id, payload in revision_edits.items():
        clean_id = str(revision_id or "").strip()
        if not clean_id:
            continue
        text = payload.get("text", "") if isinstance(payload, dict) else payload
        clean_text = str(text or "").strip()
        if clean_text:
            normalized[clean_id] = {"text": clean_text}
            if isinstance(payload, dict) and str(payload.get("mode") or "").strip().lower() == "paragraph":
                normalized[clean_id]["mode"] = "paragraph"
    return normalized


def _review_revision_edit_text(revision_edits: dict[str, dict[str, str]], revision_id: str) -> str:
    payload = revision_edits.get(revision_id) or {}
    return str(payload.get("text") or "").strip()


def _review_revision_edit_mode(revision_edits: dict[str, dict[str, str]], revision_id: str) -> str:
    payload = revision_edits.get(revision_id) or {}
    return str(payload.get("mode") or "").strip().lower()


def _docx_original_text_from_node(node) -> str:
    def iter_text(current):
        name = _docx_local_name(current.tag)
        if name in {"ins", "moveTo"}:
            return
        if name in {"del", "moveFrom"}:
            text = _docx_xml_text(current)
            if text:
                yield text
            return
        if name in {"t", "delText"} and current.text:
            yield str(current.text)
        elif name == "tab":
            yield "\t"
        elif name in {"br", "cr"}:
            yield "\n"
        for child in list(current):
            yield from iter_text(child)

    return "".join(iter_text(node)).strip()


def _copy_docx_revision_run_properties(ET, revision_node):
    if revision_node is None:
        return None
    for run in revision_node.findall(".//w:r", _DOCX_NS):
        run_properties = _copy_docx_run_properties(ET, run)
        if run_properties is not None:
            return run_properties
    return None


def _append_docx_revision_text_runs(ET, container, text: str, *, tag_name: str, run_properties=None) -> None:
    lines = str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    for index, line in enumerate(lines):
        if index:
            break_run = ET.SubElement(container, f"{{{_DOCX_W_NS}}}r")
            if run_properties is not None:
                break_run.append(_copy_docx_element(ET, run_properties))
            ET.SubElement(break_run, f"{{{_DOCX_W_NS}}}br")
        if line or len(lines) == 1:
            run = ET.SubElement(container, f"{{{_DOCX_W_NS}}}r")
            if run_properties is not None:
                run.append(_copy_docx_element(ET, run_properties))
            run.append(_docx_revision_text_element(ET, tag_name, line))


def _replace_docx_revision_text(ET, revision_node, text: str, *, review_author: str | None = None) -> bool:
    if revision_node is None:
        return False
    if review_author:
        revision_node.set(f"{{{_DOCX_W_NS}}}author", _normalize_review_author(review_author))
        revision_node.set(f"{{{_DOCX_W_NS}}}date", datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"))
    tag_name = "delText" if _docx_local_name(revision_node.tag) in {"del", "moveFrom"} else "t"
    run_properties = _copy_docx_revision_run_properties(ET, revision_node)
    for child in list(revision_node):
        revision_node.remove(child)
    _append_docx_revision_text_runs(ET, revision_node, text, tag_name=tag_name, run_properties=run_properties)
    return True


def _insert_docx_revision_after_token(ET, root, token: dict, text: str, *, review_author: str = DEFAULT_REVIEW_AUTHOR) -> bool:
    parent = token.get("parent")
    element = token.get("element")
    if parent is None or element is None:
        return False
    children = list(parent)
    try:
        index = children.index(element)
    except ValueError:
        return False

    change = ET.Element(f"{{{_DOCX_W_NS}}}ins")
    change.set(f"{{{_DOCX_W_NS}}}id", str(_next_docx_change_id(root)))
    change.set(f"{{{_DOCX_W_NS}}}author", _normalize_review_author(review_author))
    change.set(
        f"{{{_DOCX_W_NS}}}date",
        datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
    )
    _append_docx_revision_text_runs(
        ET,
        change,
        text,
        tag_name="t",
        run_properties=_copy_docx_revision_run_properties(ET, element),
    )
    parent.insert(index + 1, change)
    return True


def _docx_revision_change_with_text_runs(ET, *, text: str, change_type: str, change_id: int, author: str, date: str, run_properties=None):
    change = ET.Element(f"{{{_DOCX_W_NS}}}{change_type}")
    change.set(f"{{{_DOCX_W_NS}}}id", str(change_id))
    change.set(f"{{{_DOCX_W_NS}}}author", author)
    change.set(f"{{{_DOCX_W_NS}}}date", date)
    _append_docx_revision_text_runs(
        ET,
        change,
        text,
        tag_name="delText" if change_type == "del" else "t",
        run_properties=run_properties,
    )
    return change


def _replace_paragraph_with_full_redline(ET, paragraph, old_text: str, new_text: str, *, start_change_id: int, author: str, date: str) -> int:
    p_pr = _copy_docx_paragraph_properties(ET, paragraph)
    run_properties = _copy_docx_first_run_properties(ET, paragraph)
    for child in list(paragraph):
        paragraph.remove(child)
    if p_pr is not None:
        paragraph.append(p_pr)

    next_id = start_change_id
    if old_text:
        paragraph.append(_docx_revision_change_with_text_runs(
            ET,
            text=old_text,
            change_type="del",
            change_id=next_id,
            author=author,
            date=date,
            run_properties=run_properties,
        ))
        next_id += 1
    if new_text:
        paragraph.append(_docx_revision_change_with_text_runs(
            ET,
            text=new_text,
            change_type="ins",
            change_id=next_id,
            author=author,
            date=date,
            run_properties=run_properties,
        ))
        next_id += 1
    return next_id


_REVIEW_NUMBER_PREFIX_PATTERN = re.compile(
    r"^(\s*(?:[（(]\s*(?:[0-9０-９]+|[一二三四五六七八九十百千]+)\s*[）)]|(?:[0-9０-９]+|[一二三四五六七八九十百千]+)[、.．]))(.*)$",
    re.S,
)


def _split_review_number_prefix(text: str) -> tuple[str, str]:
    match = _REVIEW_NUMBER_PREFIX_PATTERN.match(str(text or ""))
    if not match:
        return "", str(text or "")
    return match.group(1), match.group(2)


def _review_lines_need_number_prefix_redline(old_line: str, new_line: str) -> bool:
    old_prefix, old_body = _split_review_number_prefix(old_line)
    new_prefix, new_body = _split_review_number_prefix(new_line)
    if not old_prefix or not new_prefix or old_prefix == new_prefix:
        return False
    return _normalize_contract_prospectus_compare_text(old_body) == _normalize_contract_prospectus_compare_text(new_body)


def _review_lines_have_number_prefix_redlines(old_lines: list[str], new_lines: list[str]) -> bool:
    return any(
        _review_lines_need_number_prefix_redline(old_line, new_line)
        for old_line, new_line in zip(old_lines, new_lines)
    )


def _docx_first_text_node(run, tag_name: str = "t"):
    if run is None or _docx_local_name(getattr(run, "tag", "")) != "r":
        return None
    return run.find(f".//w:{tag_name}", _DOCX_NS)


def _normalize_docx_leading_number_revision_marks(root, ET, *, review_author: str = DEFAULT_REVIEW_AUTHOR) -> None:
    number_text_pattern = re.compile(r"^\s*(?:[0-9０-９]+|[一二三四五六七八九十百千]+)\s*$")
    opening_pattern = re.compile(r"^\s*[（(]\s*$")
    for paragraph in root.findall(".//w:p", _DOCX_NS):
        changed = True
        while changed:
            changed = False
            children = list(paragraph)
            for index in range(0, max(len(children) - 3, 0)):
                open_run, first_change, second_change, close_run = children[index:index + 4]
                if _docx_local_name(open_run.tag) != "r" or _docx_local_name(close_run.tag) != "r":
                    continue
                first_name = _docx_local_name(first_change.tag)
                second_name = _docx_local_name(second_change.tag)
                if {first_name, second_name} != {"del", "ins"}:
                    continue
                delete_node = first_change if first_name == "del" else second_change
                insert_node = first_change if first_name == "ins" else second_change
                old_number = _docx_xml_text(delete_node).strip()
                new_number = _docx_xml_text(insert_node).strip()
                if not number_text_pattern.fullmatch(old_number) or not number_text_pattern.fullmatch(new_number):
                    continue
                open_text_node = _docx_first_text_node(open_run)
                close_text_node = _docx_first_text_node(close_run)
                open_text = str(open_text_node.text or "") if open_text_node is not None else ""
                close_text = str(close_text_node.text or "") if close_text_node is not None else ""
                if not opening_pattern.fullmatch(open_text) or close_text[:1] not in {"）", ")"}:
                    continue

                open_char = "（" if "（" in open_text else "("
                close_char = close_text[0]
                old_prefix = f"{open_char}{old_number}{close_char}"
                new_prefix = f"{open_char}{new_number}{close_char}"
                run_properties = _copy_docx_run_properties(ET, open_run) or _copy_docx_run_properties(ET, close_run)
                delete_id = _docx_attr(delete_node, "id") or str(_next_docx_change_id(root))
                if _docx_attr(insert_node, "id"):
                    insert_id = _docx_attr(insert_node, "id")
                elif str(delete_id).isdigit():
                    insert_id = str(int(str(delete_id)) + 1)
                else:
                    insert_id = str(_next_docx_change_id(root))
                delete_change = _docx_inline_change(
                    ET,
                    text=old_prefix,
                    change_type="del",
                    change_id=delete_id,
                    author=_docx_attr(delete_node, "author") or _normalize_review_author(review_author),
                    date=_docx_attr(delete_node, "date") or datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
                    run_properties=run_properties,
                )
                insert_change = _docx_inline_change(
                    ET,
                    text=new_prefix,
                    change_type="ins",
                    change_id=insert_id,
                    author=_docx_attr(insert_node, "author") or _normalize_review_author(review_author),
                    date=_docx_attr(insert_node, "date") or datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
                    run_properties=run_properties,
                )

                close_text_node.text = close_text[1:]
                insert_at = index
                for node in (open_run, first_change, second_change):
                    try:
                        paragraph.remove(node)
                    except ValueError:
                        pass
                if not close_text_node.text:
                    try:
                        paragraph.remove(close_run)
                    except ValueError:
                        pass
                paragraph.insert(insert_at, insert_change)
                paragraph.insert(insert_at, delete_change)
                changed = True
                break


def _remove_comments_from_comments_xml(comments_xml: bytes | str, comment_ids: set[str]) -> bytes | str:
    if not comment_ids:
        return comments_xml
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    is_bytes = isinstance(comments_xml, (bytes, bytearray))
    raw = bytes(comments_xml).decode("utf-8", errors="ignore") if is_bytes else str(comments_xml or "")
    try:
        root = ET.fromstring(raw)
    except ET.ParseError:
        for comment_id in comment_ids:
            raw = re.sub(
                rf"<w:comment\b[^>]*\bw:id=\"{re.escape(comment_id)}\"[\s\S]*?</w:comment>",
                "",
                raw,
            )
        return raw.encode("utf-8") if is_bytes else raw

    for parent in root.iter():
        for child in list(parent):
            if _docx_local_name(child.tag) == "comment" and _docx_attr(child, "id") in comment_ids:
                parent.remove(child)
    rendered = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    return rendered if is_bytes else rendered.decode("utf-8")


def _repack_docx_parts(docx_bytes: bytes, replacements: dict[str, bytes | str], *, removals: set[str] | None = None) -> bytes:
    output = io.BytesIO()
    removals = removals or set()
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as src, zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        written = set()
        for item in src.infolist():
            if item.filename in removals:
                continue
            data = replacements.get(item.filename)
            if data is None:
                data = src.read(item.filename)
            if isinstance(data, str):
                data = data.encode("utf-8")
            dst.writestr(item, data)
            written.add(item.filename)
        for name, data in replacements.items():
            if name in written or name in removals:
                continue
            if isinstance(data, str):
                data = data.encode("utf-8")
            dst.writestr(name, data)
    return output.getvalue()


def _review_export_revision_actions(tokens: list[dict], document_kind: str, revision_counter: int, decisions: dict, revision_edits: dict) -> tuple[list[dict], int]:
    actions: list[dict] = []
    index = 0
    while index < len(tokens):
        token = tokens[index]
        mark_type = token.get("mark_type")
        next_token = tokens[index + 1] if index + 1 < len(tokens) else None
        next_type = (next_token or {}).get("mark_type")
        is_modify_pair = {mark_type, next_type} == {"delete", "insert"}

        if is_modify_pair:
            revision_counter += 1
            revision_id = f"{document_kind}-reader-{revision_counter}"
            actions.append({
                "revision_id": revision_id,
                "kind": "modify",
                "decision": _normalize_review_export_decision(decisions.get(revision_id)),
                "token": token,
                "next_token": next_token,
                "edit_text": _review_revision_edit_text(revision_edits, revision_id),
                "edit_mode": _review_revision_edit_mode(revision_edits, revision_id),
            })
            index += 2
            continue

        if mark_type in {"insert", "delete"}:
            revision_counter += 1
            revision_id = f"{document_kind}-reader-{revision_counter}"
            actions.append({
                "revision_id": revision_id,
                "kind": mark_type,
                "decision": _normalize_review_export_decision(decisions.get(revision_id)),
                "token": token,
                "edit_text": _review_revision_edit_text(revision_edits, revision_id),
                "edit_mode": _review_revision_edit_mode(revision_edits, revision_id),
            })
        elif mark_type == "comment":
            comment_id = token.get("comment_id") or ""
            revision_id = token.get("_revision_id")
            if not revision_id:
                revision_counter += 1
                revision_id = f"{document_kind}-reader-{revision_counter}"
                for peer in tokens:
                    if peer.get("mark_type") == "comment" and peer.get("comment_id") == comment_id:
                        peer["_revision_id"] = revision_id
                actions.append({
                    "revision_id": revision_id,
                    "kind": "comment",
                    "decision": _normalize_review_export_decision(decisions.get(revision_id)),
                    "token": token,
                    "comment_id": comment_id,
                })
        index += 1
    return actions, revision_counter


def _apply_docx_paragraph_revision_edit(ET, root, paragraph, action: dict, *, review_author: str = DEFAULT_REVIEW_AUTHOR) -> bool:
    edit_text = str(action.get("edit_text") or "").strip()
    if not edit_text:
        return False
    token = action.get("token") or {}
    old_text = _docx_original_text_from_node(paragraph)
    start_change_id = _next_docx_change_id(root)
    _replace_paragraph_with_full_redline(
        ET,
        paragraph,
        old_text,
        edit_text,
        start_change_id=start_change_id,
        author=_normalize_review_author(review_author),
        date=datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
    )
    return True


def _apply_review_revision_decisions_to_docx(
    docx_bytes: bytes,
    *,
    document_kind: str,
    filename: str = "",
    decisions: dict | None = None,
    revision_edits: dict | None = None,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    """Apply review decisions while keeping accepted changes as Word redlines."""
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    decisions = decisions or {}
    revision_edits = _normalize_review_revision_edits(revision_edits)
    review_author = _normalize_review_author(review_author)
    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        raise ValueError("DOCX 缺少 word/document.xml，无法生成红线终稿")

    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError as exc:
        raise ValueError("DOCX 正文 XML 无法解析，无法生成红线终稿") from exc

    comments = _load_docx_comments(docx_bytes)
    revision_counter = 0
    accepted_count = 0
    rejected_count = 0
    pending_count = 0
    rejected_comment_ids: set[str] = set()
    applied_ids: list[str] = []
    edited_revision_ids: list[str] = []

    for paragraph in root.findall(".//w:p", _DOCX_NS):
        tokens = _review_export_paragraph_tokens(paragraph, comments)
        actions, revision_counter = _review_export_revision_actions(tokens, document_kind, revision_counter, decisions, revision_edits)
        if not actions:
            continue

        for action in actions:
            revision_id = action["revision_id"]
            decision = action["decision"]
            applied_ids.append(revision_id)
            if decision == "reject":
                rejected_count += 1
                if action["kind"] == "comment" and action.get("comment_id"):
                    rejected_comment_ids.add(action["comment_id"])
            elif decision == "pending":
                pending_count += 1
            else:
                accepted_count += 1

        paragraph_edit = next(
            (
                action
                for action in actions
                if action["kind"] in {"insert", "delete", "modify"}
                and action["decision"] == "accept"
                and action.get("edit_text")
                and action.get("edit_mode") == "paragraph"
            ),
            None,
        )
        if paragraph_edit:
            if _apply_docx_paragraph_revision_edit(ET, root, paragraph, paragraph_edit, review_author=review_author):
                edited_revision_ids.append(paragraph_edit["revision_id"])
            continue

        for action in actions:
            decision = action["decision"]
            kind = action["kind"]
            token = action.get("token")
            next_token = action.get("next_token")
            revision_id = action["revision_id"]

            if decision == "reject":
                if kind == "modify":
                    for part in (token, next_token):
                        if part and part.get("mark_type") == "delete":
                            _apply_rejected_docx_revision(paragraph, part, ET)
                    for part in (token, next_token):
                        if part and part.get("mark_type") == "insert":
                            _apply_rejected_docx_revision(paragraph, part, ET)
                elif kind in {"insert", "delete"}:
                    _apply_rejected_docx_revision(paragraph, token, ET)
                continue
            if decision != "accept":
                continue

            edit_text = str(action.get("edit_text") or "").strip()
            if not edit_text:
                continue
            if kind == "modify":
                insert_part = token if token and token.get("mark_type") == "insert" else next_token if next_token and next_token.get("mark_type") == "insert" else None
                if insert_part and _replace_docx_revision_text(ET, insert_part.get("element"), edit_text, review_author=review_author):
                    edited_revision_ids.append(revision_id)
            elif kind == "insert" and _replace_docx_revision_text(ET, token.get("element"), edit_text, review_author=review_author):
                edited_revision_ids.append(revision_id)
            elif kind == "delete" and _insert_docx_revision_after_token(ET, root, token, edit_text, review_author=review_author):
                edited_revision_ids.append(revision_id)

    _strip_docx_comment_markers(root, rejected_comment_ids)
    _normalize_docx_leading_number_revision_marks(root, ET, review_author=review_author)
    replacements: dict[str, bytes | str] = {
        "word/document.xml": _serialize_docx_document_xml(ET, root, xml_text)
    }
    comments_xml = _read_docx_xml_part(docx_bytes, "word/comments.xml")
    if comments_xml and rejected_comment_ids:
        replacements["word/comments.xml"] = _remove_comments_from_comments_xml(comments_xml, rejected_comment_ids)

    final_bytes = _repack_docx_parts(docx_bytes, replacements)
    return final_bytes, {
        "document_kind": document_kind,
        "filename": filename,
        "accepted_count": accepted_count,
        "rejected_count": rejected_count,
        "pending_count": pending_count,
        "edited_count": len(edited_revision_ids),
        "applied_revision_ids": applied_ids,
        "edited_revision_ids": edited_revision_ids,
        "rejected_comment_ids": sorted(rejected_comment_ids),
    }


def _docx_effective_text_from_bytes(docx_bytes: bytes) -> str:
    import xml.etree.ElementTree as ET

    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        return ""
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return ""

    def iter_effective_text(node, excluded: bool = False):
        name = _docx_local_name(node.tag)
        if name in {"del", "moveFrom"}:
            excluded = True
        if not excluded:
            if name == "t" and node.text:
                yield str(node.text)
            elif name == "tab":
                yield "\t"
            elif name in {"br", "cr"}:
                yield "\n"
        for child in list(node):
            yield from iter_effective_text(child, excluded)

    lines = []
    for paragraph in root.findall(".//w:p", _DOCX_NS):
        text = "".join(iter_effective_text(paragraph)).strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def _build_review_decision_csv(
    decisions: dict,
    documents: list[dict],
    revision_edits: dict | None = None,
    applied_edit_ids: set[str] | None = None,
) -> str:
    import csv

    revision_edits = _normalize_review_revision_edits(revision_edits)
    applied_edit_ids = applied_edit_ids or set()
    revisions_by_id = {
        revision.get("id"): revision
        for document in documents or []
        for revision in (document or {}).get("revisions", [])
        if revision.get("id")
    }
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["revision_id", "decision", "document", "scope", "type", "text", "edited_text", "edit_applied"])
    for revision_id in sorted(revisions_by_id):
        revision = revisions_by_id[revision_id]
        edited_text = _review_revision_edit_text(revision_edits, revision_id)
        writer.writerow([
            revision_id,
            _normalize_review_export_decision(decisions.get(revision_id)),
            revision.get("source_file", ""),
            revision.get("document_label", ""),
            revision.get("revision_type", ""),
            revision.get("text") or revision.get("new_text") or revision.get("old_text") or "",
            edited_text,
            "yes" if revision_id in applied_edit_ids else "no",
        ])
    return output.getvalue()


def _render_review_contract_summary_from_text(contract_text: str) -> tuple[str, dict]:
    template_text = TEMPLATE_MD.read_text(encoding="utf-8")
    template_summary = _find_contract_summary_section(_split_contract_sections(template_text))
    if not template_summary:
        raise ValueError("未在基金合同模板中找到第二十六部分摘要结构")
    scaffold = "\n".join([
        str(template_summary.get("heading") or "第二十六部分  基金合同内容摘要"),
        str(template_summary.get("content") or "").strip(),
    ]).strip()
    merged = f"{contract_text.rstrip()}\n\n{scaffold}"
    rendered = engine._replace_contract_summary_placeholders(merged)
    rendered_sections = _split_contract_sections(rendered)
    rendered_summary_candidates = [
        section
        for section in rendered_sections
        if _review_docx_is_fund_contract_summary_heading(section.get("heading") or "")
        and len(_strip_contract_signing_page_text(section.get("content") or "")) > 10
    ]
    rendered_summary = rendered_summary_candidates[-1] if rendered_summary_candidates else _find_contract_summary_section(rendered_sections)
    if not rendered_summary:
        raise ValueError("基金合同摘要重算后无法定位摘要章节")
    summary_text = "\n".join([
        str(rendered_summary.get("heading") or "").strip(),
        str(rendered_summary.get("content") or "").strip(),
    ]).strip()
    return summary_text, {
        "status": "generated",
        "message": "已根据应用审核决定后的基金合同正文重算基金合同摘要",
        "characters": len(summary_text),
    }


def _effective_docx_text_from_node(node) -> str:
    def iter_text(current, excluded: bool = False):
        name = _docx_local_name(current.tag)
        if name in {"del", "moveFrom"}:
            excluded = True
        if not excluded:
            if name == "t" and current.text:
                yield str(current.text)
            elif name == "tab":
                yield "\t"
            elif name in {"br", "cr"}:
                yield "\n"
        for child in list(current):
            yield from iter_text(child, excluded)

    return "".join(iter_text(node)).strip()


def _docx_revision_text_element(ET, tag_name: str, text: str):
    text_node = ET.Element(f"{{{_DOCX_W_NS}}}{tag_name}")
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    return text_node


def _next_docx_change_id(root) -> int:
    max_id = 0
    for node in root.iter():
        node_id = _docx_attr(node, "id")
        if node_id.isdigit():
            max_id = max(max_id, int(node_id))
    return max_id + 1


def _tracked_review_paragraph(ET, *, text: str, change_type: str, change_id: int, author: str, date: str):
    paragraph = ET.Element(f"{{{_DOCX_W_NS}}}p")
    change = ET.SubElement(paragraph, f"{{{_DOCX_W_NS}}}{change_type}")
    change.set(f"{{{_DOCX_W_NS}}}id", str(change_id))
    change.set(f"{{{_DOCX_W_NS}}}author", author)
    change.set(f"{{{_DOCX_W_NS}}}date", date)
    run = ET.SubElement(change, f"{{{_DOCX_W_NS}}}r")
    text_tag = "delText" if change_type == "del" else "t"
    run.append(_docx_revision_text_element(ET, text_tag, text))
    return paragraph


def _review_docx_summary_heading_compact(text: str) -> str:
    compact = re.sub(r"\s+", "", text or "")
    replacements = {
        "基金合同的内容摘要": "基金合同内容摘要",
        "基金托管协议的内容摘要": "基金托管协议内容摘要",
        "托管协议的内容摘要": "托管协议内容摘要",
    }
    for old, new in replacements.items():
        compact = compact.replace(old, new)
    return compact


def _review_docx_heading_has_toc_page_number(compact_heading: str) -> bool:
    return bool(re.search(r"[0-9０-９]+$", compact_heading or ""))


def _review_docx_is_fund_contract_summary_heading(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    if _review_docx_heading_has_toc_page_number(compact):
        return False
    return "基金合同内容摘要" in compact and "托管协议" not in compact


def _strip_matching_heading_from_summary(summary_text: str, heading_text: str) -> str:
    lines = [line for line in str(summary_text or "").splitlines()]
    if not lines:
        return ""
    heading_compact = _review_docx_summary_heading_compact(heading_text)
    first_compact = _review_docx_summary_heading_compact(lines[0])
    if heading_compact and first_compact == heading_compact:
        return "\n".join(lines[1:]).strip()
    if "基金合同内容摘要" in first_compact:
        return "\n".join(lines[1:]).strip()
    return "\n".join(lines).strip()


def _docx_text_is_contract_signing_page_start(text: str) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if not compact:
        return False
    if "签署页" in compact and ("基金合同" in compact or "无正文" in compact):
        return True
    if "本页" in compact and "基金合同" in compact and ("无正文" in compact or "签署页" in compact):
        return True
    return False


def _replace_docx_section_content_with_redline(
    docx_bytes: bytes,
    *,
    is_heading,
    is_stop_heading,
    new_text: str,
    report_label: str,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：DOCX 缺少正文 XML"}
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：正文 XML 无法解析"}
    _normalize_docx_leading_number_revision_marks(root, ET, review_author=review_author)

    body = root.find("w:body", _DOCX_NS)
    if body is None:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：未找到正文 body"}

    children = list(body)
    heading_index = None
    heading_text = ""
    for index, child in enumerate(children):
        if _docx_local_name(child.tag) != "p":
            continue
        text = _effective_docx_text_from_node(child)
        if is_heading(text):
            heading_index = index
            heading_text = text
            break
    if heading_index is None:
        return docx_bytes, {"status": "not_found", "message": f"{report_label}：未定位到目标摘要标题，未自动替换"}

    end_index = len(children)
    for index in range(heading_index + 1, len(children)):
        child = children[index]
        if _docx_local_name(child.tag) == "sectPr":
            end_index = index
            break
        if _docx_local_name(child.tag) == "p":
            if child.find("w:pPr/w:sectPr", _DOCX_NS) is not None:
                end_index = index
                break
            paragraph_text = _effective_docx_text_from_node(child)
            if is_stop_heading(paragraph_text) or _docx_text_is_contract_signing_page_start(paragraph_text):
                end_index = index
                break

    old_blocks = children[heading_index + 1:end_index]
    old_units = []
    for child in old_blocks:
        if _docx_local_name(child.tag) != "p":
            continue
        line = _effective_docx_text_from_node(child).strip()
        if not line:
            continue
        old_units.append({
            "node": child,
            "line": line,
            "compare": _normalize_contract_prospectus_compare_text(_strip_review_number_prefix(line)),
        })
    old_text = "\n".join(unit["line"] for unit in old_units).strip()
    replacement_text = _strip_matching_heading_from_summary(new_text, heading_text)
    if not replacement_text:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：新摘要为空，未替换"}
    if re.sub(r"\s+", "", old_text) == re.sub(r"\s+", "", replacement_text):
        return docx_bytes, {"status": "unchanged", "message": f"{report_label}：内容已一致，未新增系统红线"}

    for child in old_blocks:
        try:
            body.remove(child)
        except ValueError:
            pass

    next_id = _next_docx_change_id(root)
    now = datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
    author = _normalize_review_author(review_author)
    replacements = []
    new_rows = _review_sync_line_rows(replacement_text)
    new_lines = [row["line"].strip() for row in new_rows]
    old_compare = [unit["compare"] for unit in old_units]
    new_compare = [row["compare"] for row in new_rows]
    old_nodes = [unit["node"] for unit in old_units]
    old_lines = [unit["line"] for unit in old_units]
    matcher = difflib.SequenceMatcher(None, old_compare, new_compare, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(i2 - i1):
                paragraph = old_nodes[i1 + offset]
                old_line = old_lines[i1 + offset]
                new_line = new_lines[j1 + offset]
                if _review_lines_need_number_prefix_redline(old_line, new_line):
                    next_id = _replace_paragraph_with_minimal_redline(
                        ET,
                        paragraph,
                        old_line,
                        new_line,
                        start_change_id=next_id,
                        author=author,
                        date=now,
                    )
                replacements.append(paragraph)
            continue
        if tag == "replace" and (i2 - i1) == (j2 - j1):
            for offset in range(i2 - i1):
                paragraph = old_nodes[i1 + offset]
                next_id = _replace_paragraph_with_minimal_redline(
                    ET,
                    paragraph,
                    old_lines[i1 + offset],
                    new_lines[j1 + offset],
                    start_change_id=next_id,
                    author=author,
                    date=now,
                )
                replacements.append(paragraph)
            continue
        template = old_nodes[i1] if i1 < len(old_nodes) else (old_nodes[i1 - 1] if i1 > 0 else None)
        for old_line in old_lines[i1:i2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=old_line, change_type="del", change_id=next_id, author=author, date=now))
            next_id += 1
        for new_line in new_lines[j1:j2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=new_line, change_type="ins", change_id=next_id, author=author, date=now))
            next_id += 1

    insert_index = heading_index + 1
    for offset, replacement in enumerate(replacements):
        body.insert(insert_index + offset, replacement)

    final_bytes = _repack_docx_parts(
        docx_bytes,
        {"word/document.xml": _serialize_docx_document_xml(ET, root, xml_text)},
    )
    return final_bytes, {
        "status": "redlined",
        "message": f"{report_label}：已以红线方式替换摘要内容",
        "old_characters": len(old_text),
        "new_characters": len(replacement_text),
    }


def _contract_summary_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    return compact == "第二十六部分基金合同内容摘要"


def _contract_summary_stop_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    return bool(re.match(r"^第[一二三四五六七八九十百]+部分", compact or "") and compact != "第二十六部分基金合同内容摘要")


def _prospectus_contract_summary_heading_match(text: str) -> bool:
    return _review_docx_is_fund_contract_summary_heading(text)


def _prospectus_contract_summary_stop_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    if _review_docx_heading_has_toc_page_number(compact):
        return False
    return "基金托管协议内容摘要" in compact or "托管协议内容摘要" in compact


def _prospectus_custody_summary_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    if _review_docx_heading_has_toc_page_number(compact):
        return False
    return "基金托管协议内容摘要" in compact or "托管协议内容摘要" in compact


def _prospectus_custody_summary_stop_heading_match(text: str) -> bool:
    compact = _review_docx_summary_heading_compact(text)
    if not compact or _review_docx_heading_has_toc_page_number(compact):
        return False
    if _prospectus_custody_summary_heading_match(text):
        return False
    stop_keywords = (
        "基金份额持有人服务",
        "其他应披露事项",
        "其它应披露事项",
        "招募说明书存放",
        "备查文件",
    )
    return any(keyword in compact for keyword in stop_keywords)


def _sync_review_contract_summary_redlines(contract_docx: bytes, prospectus_docx: bytes, *, review_author: str = DEFAULT_REVIEW_AUTHOR) -> tuple[bytes, bytes, dict]:
    contract_text = _docx_effective_text_from_bytes(contract_docx)
    try:
        summary_text, summary_report = _render_review_contract_summary_from_text(contract_text)
    except Exception as exc:
        return contract_docx, prospectus_docx, {
            "contract_summary": {"status": "skipped", "message": "未能自动重算基金合同摘要", "error": str(exc)},
            "prospectus_contract_summary": {"status": "skipped", "message": "因合同摘要未重算，未同步招募说明书内合同摘要"},
        }

    contract_docx, contract_apply_report = _replace_docx_section_content_with_redline(
        contract_docx,
        is_heading=_contract_summary_heading_match,
        is_stop_heading=_contract_summary_stop_heading_match,
        new_text=summary_text,
        report_label="基金合同第 26 部分摘要",
        review_author=review_author,
    )
    prospectus_docx, prospectus_apply_report = _replace_docx_section_content_with_redline(
        prospectus_docx,
        is_heading=_prospectus_contract_summary_heading_match,
        is_stop_heading=_prospectus_contract_summary_stop_heading_match,
        new_text=summary_text,
        report_label="招募说明书内基金合同摘要",
        review_author=review_author,
    )
    return contract_docx, prospectus_docx, {
        "contract_summary": {**summary_report, "apply": contract_apply_report},
        "prospectus_contract_summary": prospectus_apply_report,
    }


def _build_custodian_summary_text_from_docx_bytes(docx_bytes: bytes) -> dict:
    source_text = _docx_effective_text_from_bytes(docx_bytes)
    sections = _clean_custodian_summary_sections(_split_prospectus_sections(source_text))
    titles = [
        line.strip()
        for line in str(prospectus_engine.pro_clauses.get("CHAPTER21_TITLES", {}).get("text", "")).splitlines()
        if line.strip()
    ]
    remaining_sections = list(sections)
    blocks = []
    missing = []
    for title in titles:
        content = ""
        if remaining_sections:
            match = _match_custodian_summary_section(remaining_sections, title)
            matched_section = match.get("section")
            if matched_section:
                content = str(matched_section.get("content") or "").strip()
                remaining_sections = [section for section in remaining_sections if section is not matched_section]
        if not content:
            missing.append(title)
            content = "【待填写】"
        blocks.extend([title, content])
    return {
        "source_type": "custody_agreement",
        "summary_text": "\n".join(blocks).strip(),
        "matched_sections": [section.get("heading", "") for section in sections],
        "missing_sections": missing,
    }


def _sync_review_custody_summary_redlines(
    custody_docx: bytes,
    prospectus_docx: bytes,
    *,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    summary_report = _build_custodian_summary_text_from_docx_bytes(custody_docx)
    summary_text = str(summary_report.get("summary_text") or "").strip()
    if not summary_text:
        return prospectus_docx, {
            "status": "skipped",
            "message": "未能从托管协议提取可同步的内容摘要",
            "summary": summary_report,
            "apply": {"status": "skipped", "message": "未同步招募说明书托管协议摘要"},
        }
    prospectus_docx, apply_report = _replace_docx_section_content_with_redline(
        prospectus_docx,
        is_heading=_prospectus_custody_summary_heading_match,
        is_stop_heading=_prospectus_custody_summary_stop_heading_match,
        new_text=summary_text,
        report_label="招募说明书内托管协议摘要",
        review_author=review_author,
    )
    return prospectus_docx, {
        "status": apply_report.get("status") or "skipped",
        "message": apply_report.get("message") or "已尝试同步招募说明书托管协议摘要",
        "summary": summary_report,
        "apply": apply_report,
    }


def _review_cross_rule_is_direct_sync_candidate(rule: dict) -> tuple[bool, str]:
    relation = _normalize_review_text((rule or {}).get("relation"))
    consistency = _normalize_review_text((rule or {}).get("consistency"))
    combined = " ".join([relation, consistency, _normalize_review_text((rule or {}).get("detail"))])
    contract_locator = _normalize_review_text((rule or {}).get("contract_locator") or (rule or {}).get("contract_chapter"))
    prospectus_locator = _normalize_review_text((rule or {}).get("prospectus_locator") or (rule or {}).get("prospectus_chapter"))
    locator_pair = f"{contract_locator} {prospectus_locator}"
    is_whole_chapter_rule = (
        not any(separator in contract_locator for separator in ("/", "／"))
        and not any(separator in prospectus_locator for separator in ("/", "／"))
    )
    is_whole_distribution_chapter = (
        "基金的收益与分配" in locator_pair
        and is_whole_chapter_rule
    )
    if is_whole_distribution_chapter:
        return False, "收益分配章节含招募细化条款，不做整章自动替换"
    if any(keyword in combined for keyword in ("无直接对应", "招募独有", "合同独有", "无同名独立条款", "无对应独立章节")):
        return False, "无直接对应或单方独有"
    if any(keyword in relation for keyword in ("招募细化", "细化", "部分对应", "说明性对应")):
        return False, "招募细化或部分对应规则不自动替换"
    if is_whole_chapter_rule:
        return False, "整章对应需经过招募说明书后处理，不做整章自动替换"
    if "直接对应" not in relation and "完全一致" not in consistency:
        return False, "非直接对应规则"
    if consistency and not any(keyword in consistency for keyword in ("完全一致", "基本一致")):
        return False, "一致性口径不是完全一致或基本一致"
    return True, ""


def _review_sync_minimal_change_guard(source_text: str, target_text: str) -> tuple[bool, str]:
    source_rows = _review_sync_line_rows(source_text)
    target_rows = _review_sync_line_rows(target_text)
    source_compare = [row["compare"] for row in source_rows if row.get("compare")]
    target_compare = [row["compare"] for row in target_rows if row.get("compare")]
    if not source_compare or not target_compare:
        return False, "同步源或目标为空"
    if len(source_compare) != len(target_compare):
        if len(target_compare) > len(source_compare):
            return False, "招募说明书存在额外内容，未自动删除"
        return False, "合同来源内容存在新增段落，需人工确认后同步"
    source_joined = "\n".join(source_compare)
    target_joined = "\n".join(target_compare)
    similarity = difflib.SequenceMatcher(None, source_joined, target_joined, autojunk=False).ratio()
    if similarity < 0.72:
        return False, f"文本差异较大（相似度 {similarity:.0%}），未自动替换"
    return True, ""


def _review_sync_line_rows(text: str) -> list[dict]:
    return _collect_nonblank_compare_lines(
        text,
        normalize_line=lambda line: _normalize_contract_prospectus_compare_text(_strip_review_number_prefix(line)),
    )


def _fallback_review_rule_target_from_text(full_text: str, locator: str, section_hint: str = "") -> dict:
    parts = _split_review_locator(locator)
    labels = [part for part in [*(reversed(parts)), section_hint, locator] if _normalize_review_text(part)]
    seen = set()
    for label in labels:
        normalized = _normalize_review_text(label)
        if normalized in seen:
            continue
        seen.add(normalized)
        block = _find_review_block_by_label(full_text, label)
        if not block:
            continue
        heading = _normalize_review_text(block.get("heading"))
        body_text = (block.get("body") or block.get("text") or "").strip()
        return {
            "section": None,
            "section_heading": "",
            "target_heading": heading,
            "text": (block.get("text") or "").strip(),
            "body_text": body_text,
            "matched": bool(body_text),
            "match_method": "full_text_block",
            "missing_reason": "" if body_text else "body_empty",
            "locator_parts": [heading] if heading else [],
        }
    return {
        "section": None,
        "section_heading": "",
        "target_heading": locator,
        "text": "",
        "body_text": "",
        "matched": False,
        "match_method": "",
        "missing_reason": "fallback_block_missing",
        "locator_parts": [],
    }


def _docx_body_paragraphs(root) -> list[dict]:
    body = root.find("w:body", _DOCX_NS)
    if body is None:
        return []
    paragraphs = []
    for child in list(body):
        if _docx_local_name(child.tag) != "p":
            continue
        text = _effective_docx_text_from_node(child)
        if text:
            paragraphs.append({"node": child, "text": text, "compare": _normalize_contract_prospectus_compare_text(_strip_review_number_prefix(text))})
    return paragraphs


def _find_docx_paragraph_span_by_text(root, target_text: str) -> tuple[list, str]:
    target_rows = _review_sync_line_rows(target_text)
    target_compare = [row["compare"] for row in target_rows if row.get("compare")]
    if not target_compare:
        return [], "target_empty"

    paragraphs = _docx_body_paragraphs(root)
    paragraph_compare = [row["compare"] for row in paragraphs]
    target_len = len(target_compare)
    for start in range(0, max(len(paragraph_compare) - target_len + 1, 0)):
        if paragraph_compare[start:start + target_len] == target_compare:
            return [row["node"] for row in paragraphs[start:start + target_len]], ""

    if target_len == 1:
        target = target_compare[0]
        for row in paragraphs:
            if row["compare"] == target:
                return [row["node"]], ""
    return [], "target_not_found_in_docx"


def _copy_docx_paragraph_properties(ET, paragraph):
    p_pr = paragraph.find("w:pPr", _DOCX_NS)
    return _copy_docx_element(ET, p_pr) if p_pr is not None else None


def _copy_docx_first_run_properties(ET, paragraph):
    if paragraph is None:
        return None
    for run in paragraph.findall(".//w:r", _DOCX_NS):
        r_pr = run.find("w:rPr", _DOCX_NS)
        if r_pr is not None:
            return _copy_docx_element(ET, r_pr)
    return None


def _docx_plain_text_run(ET, text: str, run_properties=None):
    run = ET.Element(f"{{{_DOCX_W_NS}}}r")
    if run_properties is not None:
        run.append(_copy_docx_element(ET, run_properties))
    run.append(_docx_revision_text_element(ET, "t", text))
    return run


def _docx_inline_change(ET, *, text: str, change_type: str, change_id: int, author: str, date: str, run_properties=None):
    change = ET.Element(f"{{{_DOCX_W_NS}}}{change_type}")
    change.set(f"{{{_DOCX_W_NS}}}id", str(change_id))
    change.set(f"{{{_DOCX_W_NS}}}author", author)
    change.set(f"{{{_DOCX_W_NS}}}date", date)
    run = ET.SubElement(change, f"{{{_DOCX_W_NS}}}r")
    if run_properties is not None:
        run.append(_copy_docx_element(ET, run_properties))
    run.append(_docx_revision_text_element(ET, "delText" if change_type == "del" else "t", text))
    return change


def _review_sync_diff_tokens(text: str) -> list[str]:
    return re.findall(r"\d+(?:\.\d+)?%?|[A-Za-z]+(?:[-_][A-Za-z0-9]+)*|[\u4e00-\u9fff]+|[^\u4e00-\u9fffA-Za-z\d]+", str(text or ""))


def _append_minimal_redline_text_runs(ET, paragraph, old_text: str, new_text: str, *, next_id: int, author: str, date: str, run_properties=None) -> int:
    old_tokens = _review_sync_diff_tokens(old_text)
    new_tokens = _review_sync_diff_tokens(new_text)
    matcher = difflib.SequenceMatcher(None, old_tokens, new_tokens, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        old_part = "".join(old_tokens[i1:i2])
        new_part = "".join(new_tokens[j1:j2])
        if tag == "equal":
            if old_part:
                paragraph.append(_docx_plain_text_run(ET, old_part, run_properties=run_properties))
            continue
        if tag in {"replace", "delete"} and old_part:
            paragraph.append(_docx_inline_change(ET, text=old_part, change_type="del", change_id=next_id, author=author, date=date, run_properties=run_properties))
            next_id += 1
        if tag in {"replace", "insert"} and new_part:
            paragraph.append(_docx_inline_change(ET, text=new_part, change_type="ins", change_id=next_id, author=author, date=date, run_properties=run_properties))
            next_id += 1
    return next_id


def _replace_paragraph_with_minimal_redline(ET, paragraph, old_text: str, new_text: str, *, start_change_id: int, author: str, date: str) -> int:
    p_pr = _copy_docx_paragraph_properties(ET, paragraph)
    run_properties = _copy_docx_first_run_properties(ET, paragraph)
    for child in list(paragraph):
        paragraph.remove(child)
    if p_pr is not None:
        paragraph.append(p_pr)

    next_id = start_change_id
    old_prefix, old_body = _split_review_number_prefix(old_text)
    new_prefix, new_body = _split_review_number_prefix(new_text)
    if old_prefix and new_prefix and old_prefix != new_prefix:
        paragraph.append(_docx_inline_change(ET, text=old_prefix, change_type="del", change_id=next_id, author=author, date=date, run_properties=run_properties))
        next_id += 1
        paragraph.append(_docx_inline_change(ET, text=new_prefix, change_type="ins", change_id=next_id, author=author, date=date, run_properties=run_properties))
        next_id += 1
        return _append_minimal_redline_text_runs(
            ET,
            paragraph,
            old_body,
            new_body,
            next_id=next_id,
            author=author,
            date=date,
            run_properties=run_properties,
        )
    return _append_minimal_redline_text_runs(
        ET,
        paragraph,
        old_text,
        new_text,
        next_id=next_id,
        author=author,
        date=date,
        run_properties=run_properties,
    )


def _tracked_review_paragraph_from_template(ET, *, template=None, text: str, change_type: str, change_id: int, author: str, date: str):
    paragraph = ET.Element(f"{{{_DOCX_W_NS}}}p")
    p_pr = _copy_docx_paragraph_properties(ET, template) if template is not None else None
    run_properties = _copy_docx_first_run_properties(ET, template)
    if p_pr is not None:
        paragraph.append(p_pr)
    paragraph.append(_docx_inline_change(ET, text=text, change_type=change_type, change_id=change_id, author=author, date=date, run_properties=run_properties))
    return paragraph


def _redline_docx_paragraph_span_to_text(
    docx_bytes: bytes,
    target_text: str,
    replacement_text: str,
    *,
    report_label: str,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    import xml.etree.ElementTree as ET

    _register_docx_xml_namespaces(ET)
    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：DOCX 缺少正文 XML"}
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：正文 XML 无法解析"}
    body = root.find("w:body", _DOCX_NS)
    if body is None:
        return docx_bytes, {"status": "skipped", "message": f"{report_label}：未找到正文 body"}

    span_nodes, missing_reason = _find_docx_paragraph_span_by_text(root, target_text)
    if not span_nodes:
        return docx_bytes, {"status": "not_found", "message": f"{report_label}：未定位到招募正文段落，未自动替换", "reason": missing_reason}

    old_rows = _review_sync_line_rows(target_text)
    new_rows = _review_sync_line_rows(replacement_text)
    old_lines = [row["line"].strip() for row in old_rows]
    new_lines = [row["line"].strip() for row in new_rows]
    old_compare = [row["compare"] for row in old_rows]
    new_compare = [row["compare"] for row in new_rows]
    if old_compare == new_compare and not _review_lines_have_number_prefix_redlines(old_lines, new_lines):
        return docx_bytes, {"status": "unchanged", "message": f"{report_label}：内容已一致，未新增系统红线"}

    children = list(body)
    try:
        start_index = children.index(span_nodes[0])
    except ValueError:
        return docx_bytes, {"status": "not_found", "message": f"{report_label}：目标段落位置失效，未自动替换"}
    span_set = set(span_nodes)
    for node in span_nodes:
        try:
            body.remove(node)
        except ValueError:
            pass

    now = datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
    author = _normalize_review_author(review_author)
    next_id = _next_docx_change_id(root)
    replacements = []
    matcher = difflib.SequenceMatcher(None, old_compare, new_compare, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(i2 - i1):
                paragraph = span_nodes[i1 + offset]
                old_line = old_lines[i1 + offset]
                new_line = new_lines[j1 + offset]
                if _review_lines_need_number_prefix_redline(old_line, new_line):
                    next_id = _replace_paragraph_with_minimal_redline(
                        ET,
                        paragraph,
                        old_line,
                        new_line,
                        start_change_id=next_id,
                        author=author,
                        date=now,
                    )
                replacements.append(paragraph)
            continue
        if tag == "replace" and (i2 - i1) == (j2 - j1):
            for offset in range(i2 - i1):
                paragraph = span_nodes[i1 + offset]
                next_id = _replace_paragraph_with_minimal_redline(
                    ET,
                    paragraph,
                    old_lines[i1 + offset],
                    new_lines[j1 + offset],
                    start_change_id=next_id,
                    author=author,
                    date=now,
                )
                replacements.append(paragraph)
            continue
        template = span_nodes[i1] if i1 < len(span_nodes) else (span_nodes[i1 - 1] if i1 > 0 else None)
        for old_line in old_lines[i1:i2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=old_line, change_type="del", change_id=next_id, author=author, date=now))
            next_id += 1
        for new_line in new_lines[j1:j2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=new_line, change_type="ins", change_id=next_id, author=author, date=now))
            next_id += 1

    for offset, node in enumerate(replacements):
        body.insert(start_index + offset, node)

    final_bytes = _repack_docx_parts(
        docx_bytes,
        {"word/document.xml": _serialize_docx_document_xml(ET, root, xml_text)},
    )
    return final_bytes, {
        "status": "redlined",
        "message": f"{report_label}：已按最小文本差异写入红线",
        "old_lines": len(old_lines),
        "new_lines": len(new_lines),
        "changed_paragraphs": max(len(old_lines), len(new_lines)),
    }


def _redline_docx_paragraph_nodes_in_parent(
    ET,
    root,
    parent,
    span_nodes: list,
    replacement_text: str,
    *,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bool, dict]:
    span_nodes = [node for node in span_nodes or [] if _docx_local_name(getattr(node, "tag", "")) == "p"]
    new_rows = _review_sync_line_rows(replacement_text)
    new_lines = [row["line"].strip() for row in new_rows]
    new_compare = [row["compare"] for row in new_rows]
    if not span_nodes or not new_compare:
        return False, {"status": "skipped", "reason": "empty_paragraphs_or_replacement"}

    old_lines = [_effective_docx_text_from_node(node).strip() for node in span_nodes]
    old_lines = [line for line in old_lines if line]
    old_compare = [
        _normalize_contract_prospectus_compare_text(_strip_review_number_prefix(line))
        for line in old_lines
    ]
    if old_compare == new_compare and not _review_lines_have_number_prefix_redlines(old_lines, new_lines):
        return False, {"status": "unchanged"}

    children = list(parent)
    try:
        start_index = children.index(span_nodes[0])
    except ValueError:
        return False, {"status": "not_found", "reason": "paragraph_position_lost"}

    for node in span_nodes:
        _remove_docx_child(parent, node)

    now = datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
    author = _normalize_review_author(review_author)
    next_id = _next_docx_change_id(root)
    replacements = []
    matcher = difflib.SequenceMatcher(None, old_compare, new_compare, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(i2 - i1):
                paragraph = span_nodes[i1 + offset]
                old_line = old_lines[i1 + offset]
                new_line = new_lines[j1 + offset]
                if _review_lines_need_number_prefix_redline(old_line, new_line):
                    next_id = _replace_paragraph_with_minimal_redline(
                        ET,
                        paragraph,
                        old_line,
                        new_line,
                        start_change_id=next_id,
                        author=author,
                        date=now,
                    )
                replacements.append(paragraph)
            continue
        if tag == "replace" and (i2 - i1) == (j2 - j1):
            for offset in range(i2 - i1):
                paragraph = span_nodes[i1 + offset]
                next_id = _replace_paragraph_with_minimal_redline(
                    ET,
                    paragraph,
                    old_lines[i1 + offset],
                    new_lines[j1 + offset],
                    start_change_id=next_id,
                    author=author,
                    date=now,
                )
                replacements.append(paragraph)
            continue
        template = span_nodes[i1] if i1 < len(span_nodes) else (span_nodes[i1 - 1] if i1 > 0 else None)
        for old_line in old_lines[i1:i2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=old_line, change_type="del", change_id=next_id, author=author, date=now))
            next_id += 1
        for new_line in new_lines[j1:j2]:
            replacements.append(_tracked_review_paragraph_from_template(ET, template=template, text=new_line, change_type="ins", change_id=next_id, author=author, date=now))
            next_id += 1

    for offset, node in enumerate(replacements):
        parent.insert(start_index + offset, node)
    return True, {
        "status": "redlined",
        "old_lines": len(old_lines),
        "new_lines": len(new_lines),
        "changed_paragraphs": max(len(old_lines), len(new_lines)),
    }


def _product_summary_prospectus_driven_table_values(render_model: dict) -> dict[str, list[str]]:
    target_labels = {"投资范围", "风险收益特征", "其他费用"}
    rows_by_label: dict[str, list[str]] = {}
    for section in (render_model or {}).get("sections", []):
        for block in section.get("blocks", []):
            if not isinstance(block, dict) or block.get("type") != "table":
                continue
            for row in block.get("rows") or []:
                if not row:
                    continue
                label = str(row[0] or "").strip()
                if label in target_labels:
                    rows_by_label[label] = [str(cell or "") for cell in row[1:]]
    return rows_by_label


def _product_summary_model_risk_disclosure_text(render_model: dict) -> str:
    for section in (render_model or {}).get("sections", []):
        if section.get("title") != "四、风险揭示与重要提示":
            continue
        lines = []
        capturing = False
        for block in section.get("blocks") or []:
            if not isinstance(block, dict) or block.get("type") != "paragraph":
                continue
            text = str(block.get("text") or "").strip()
            normalized = _normalize_review_text(text)
            if normalized == _normalize_review_text("（一）风险揭示"):
                capturing = True
                continue
            if normalized == _normalize_review_text("（二）重要提示"):
                break
            if capturing and text:
                lines.append(text)
        return "\n".join(lines).strip()
    return ""


def _review_product_summary_prospectus_render_model_from_text(text: str) -> dict:
    target_aliases = {
        "基金的投资": ("基金的投资", "基金投资"),
        "基金的费用与税收": ("基金的费用与税收", "基金费用与税收"),
        "风险揭示": ("风险揭示",),
    }

    def normalize_heading(value: str) -> str:
        normalized = _normalize_review_text(value)
        normalized = re.sub(r"\t+\d+\s*$", "", normalized).strip()
        normalized = re.sub(r"^第[一二三四五六七八九十百千]+(?:部分|章|节)\s*", "", normalized).strip()
        normalized = re.sub(r"^[一二三四五六七八九十百千]+、", "", normalized).strip()
        return normalized

    normalized_aliases = {
        canonical: {_normalize_review_text(alias) for alias in aliases}
        for canonical, aliases in target_aliases.items()
    }

    def match_chapter_title(line: str) -> str:
        normalized = normalize_heading(line)
        if not normalized:
            return ""
        for canonical, aliases in normalized_aliases.items():
            if normalized in aliases:
                return canonical
        return ""

    current_chapter: dict | None = None
    chapters_by_title: dict[str, dict] = {}
    for raw_line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").splitlines():
        stripped = str(raw_line or "").strip()
        if not stripped:
            continue
        chapter_title = match_chapter_title(stripped)
        if chapter_title:
            current_chapter = {
                "chapter_cn": "",
                "title": chapter_title,
                "display_title": stripped,
                "blocks": [],
            }
            chapters_by_title[chapter_title] = current_chapter
            continue
        if current_chapter is not None:
            current_chapter.setdefault("blocks", []).append({"type": "paragraph", "text": stripped})

    chapters = [
        chapter
        for canonical in target_aliases
        if (chapter := chapters_by_title.get(canonical)) and chapter.get("blocks")
    ]
    return {
        "cover_lines": [],
        "important_notice_title": "重要提示",
        "important_notice_blocks": [],
        "toc_title": "目录",
        "toc_titles": [chapter["title"] for chapter in chapters],
        "toc_entries": [
            {
                "chapter_cn": chapter.get("chapter_cn", ""),
                "title": chapter.get("title", ""),
                "display_title": chapter.get("display_title", ""),
            }
            for chapter in chapters
        ],
        "chapters": chapters,
    }


def _redline_product_summary_table_rows_from_model(
    docx_bytes: bytes,
    render_model: dict,
    *,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    import xml.etree.ElementTree as ET

    expected_rows = _product_summary_prospectus_driven_table_values(render_model)
    report = {
        "status": "completed",
        "checked_fields": sorted(expected_rows),
        "updated_fields": [],
        "unchanged_fields": [],
        "skipped": [],
    }
    if not expected_rows:
        report["status"] = "skipped"
        report["message"] = "未生成可同步的产品资料概要字段"
        return docx_bytes, report

    _register_docx_xml_namespaces(ET)
    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        report["status"] = "skipped"
        report["message"] = "产品资料概要 DOCX 缺少正文 XML"
        return docx_bytes, report
    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        report["status"] = "skipped"
        report["message"] = "产品资料概要正文 XML 无法解析"
        return docx_bytes, report

    matched_labels: set[str] = set()
    changed = False
    for row in root.findall(".//w:tr", _DOCX_NS):
        cells = row.findall("./w:tc", _DOCX_NS)
        if len(cells) < 2:
            continue
        label = _effective_docx_text_from_node(cells[0]).strip()
        if label not in expected_rows:
            continue
        matched_labels.add(label)
        expected_values = expected_rows[label]
        row_changed = False
        for offset, expected_value in enumerate(expected_values, start=1):
            if offset >= len(cells):
                report["skipped"].append({"field": label, "reason": "target_cell_missing", "cell_index": offset + 1})
                continue
            paragraphs = cells[offset].findall("./w:p", _DOCX_NS)
            if not paragraphs:
                paragraph = ET.Element(f"{{{_DOCX_W_NS}}}p")
                cells[offset].append(paragraph)
                paragraphs = [paragraph]
            cell_changed, cell_report = _redline_docx_paragraph_nodes_in_parent(
                ET,
                root,
                cells[offset],
                paragraphs,
                expected_value,
                review_author=review_author,
            )
            row_changed = row_changed or cell_changed
            changed = changed or cell_changed
            if cell_report.get("status") not in {"redlined", "unchanged"}:
                report["skipped"].append({"field": label, **cell_report})
        if row_changed:
            report["updated_fields"].append(label)
        else:
            report["unchanged_fields"].append(label)

    for label in sorted(set(expected_rows) - matched_labels):
        report["skipped"].append({"field": label, "reason": "row_not_found"})

    if changed:
        report["status"] = "redlined"
        final_bytes = _repack_docx_parts(
            docx_bytes,
            {"word/document.xml": _serialize_docx_document_xml(ET, root, xml_text)},
        )
        return final_bytes, report
    report["status"] = "unchanged"
    return docx_bytes, report


def _sync_review_product_summary_from_prospectus_redlines(
    product_summary_docx: bytes,
    prospectus_docx: bytes,
    *,
    form_data: dict | None = None,
    filename: str = "",
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    prospectus_text = _docx_effective_text_from_bytes(prospectus_docx)
    if not prospectus_text:
        return product_summary_docx, {
            "status": "skipped",
            "message": "未能从最终招募说明书提取文本，产品资料概要未同步",
        }
    try:
        sync_form_data = {**(form_data or {}), "PROSPECTUS_TEXT": prospectus_text}
        fallback_model = _review_product_summary_prospectus_render_model_from_text(prospectus_text)
        if fallback_model.get("chapters"):
            sync_form_data["PROSPECTUS_RENDER_MODEL"] = fallback_model
        else:
            sync_form_data.pop("PROSPECTUS_RENDER_MODEL", None)
        bundle = product_summary_engine.generate_bundle(sync_form_data)
        render_model = bundle.get("render_model") or {}
        current_docx, table_report = _redline_product_summary_table_rows_from_model(
            product_summary_docx,
            render_model,
            review_author=review_author,
        )

        risk_text = _product_summary_model_risk_disclosure_text(render_model)
        risk_report = {"status": "skipped", "message": "未生成风险揭示同步文本"}
        if risk_text:
            current_docx, risk_report = _replace_docx_section_content_with_redline(
                current_docx,
                is_heading=lambda value: _normalize_review_text(value) == _normalize_review_text("（一）风险揭示"),
                is_stop_heading=lambda value: _normalize_review_text(value) in {
                    _normalize_review_text("（二）重要提示"),
                    _normalize_review_text("五、其他资料查询方式"),
                    _normalize_review_text("六、其他情况说明"),
                },
                new_text=risk_text,
                report_label="产品资料概要风险揭示",
                review_author=review_author,
            )

        updated_fields = list(table_report.get("updated_fields") or [])
        if risk_report.get("status") == "redlined":
            updated_fields.append("风险揭示")
        status = "synced_from_prospectus" if updated_fields else "unchanged"
        return current_docx, {
            "status": status,
            "message": "已根据应用审核决定后的招募说明书同步产品资料概要红线" if updated_fields else "产品资料概要与最终招募说明书派生内容已一致",
            "updated_fields": updated_fields,
            "table_sync": table_report,
            "risk_sync": risk_report,
        }
    except Exception as exc:
        logger.exception("Failed to sync product summary from final prospectus")
        return product_summary_docx, {
            "status": "skipped",
            "message": "产品资料概要根据最终招募说明书同步失败",
            "error": str(exc),
        }


_ETF_PROSPECTUS_TEMPLATE_SYNC_RULES = (
    ("第二部分 释义", "释义", "normalize_definitions"),
    ("第五部分 基金备案", "基金合同的生效", "normalize_contract_references"),
    ("第六部分 基金份额的折算", "基金份额折算与变更登记", "normalize_contract_references"),
    ("第八部分 基金份额的申购与赎回", "基金份额的申购赎回", "normalize_contract_references"),
    ("第十四部分 基金的投资", "基金的投资", "normalize_reused_chapter"),
    ("第十五部分 基金的财产", "基金的财产", "normalize_contract_references"),
    ("第十六部分 基金资产估值", "基金资产估值", "normalize_reused_chapter"),
    ("第十七部分 基金费用与税收", "基金的费用与税收", "normalize_reused_chapter"),
    ("第十八部分 基金的收益与分配", "基金的收益与分配", "normalize_reused_chapter"),
    ("第十九部分 基金的会计与审计", "基金的会计与审计", "normalize_contract_references"),
    ("第二十部分 基金的信息披露", "基金的信息披露", "normalize_contract_references"),
    ("第二十一部分 基金合同的变更、终止与基金财产的清算", "基金合同的变更、终止和基金财产的清算", "normalize_contract_references"),
)


def _build_review_prospectus_template_sync_rules() -> list[dict]:
    rules = []
    for contract_locator, prospectus_locator, postprocess in _ETF_PROSPECTUS_TEMPLATE_SYNC_RULES:
        rules.append({
            "source": "prospectus_template",
            "sheet_name": "招募说明书模板",
            "contract_chapter": contract_locator,
            "prospectus_chapter": prospectus_locator,
            "contract_locator": contract_locator,
            "prospectus_locator": prospectus_locator,
            "contract_section_name": contract_locator,
            "prospectus_section_name": prospectus_locator,
            "relation": "模板取自基金合同",
            "consistency": "模板来源同步",
            "detail": "招募说明书模板变量由基金合同正文派生，导出时先执行招募说明书后处理再写入红线。",
            "postprocess": postprocess,
        })
    return rules


def _review_rule_is_prospectus_template_source(rule: dict) -> bool:
    return (rule or {}).get("source") == "prospectus_template"


def _postprocess_review_prospectus_template_source(text: str, rule: dict) -> str:
    mode = (rule or {}).get("postprocess") or ""
    source = (text or "").strip()
    if not source:
        return ""
    if mode == "normalize_definitions":
        return ProspectusEngine._normalize_prospectus_definitions(source)
    if mode == "normalize_reused_chapter":
        return prospectus_engine._normalize_reused_prospectus_chapter(source)
    if mode == "normalize_contract_references":
        return ProspectusEngine._normalize_prospectus_contract_references(source)
    return source


def _expand_empty_review_target_from_following_sections(target: dict, sections: list[dict]) -> dict:
    if not target.get("matched") or (target.get("body_text") or target.get("text") or "").strip():
        return target
    section = target.get("section")
    if section not in sections:
        return target
    section_index = sections.index(section)
    base_ordinal = _review_label_ordinal(section.get("heading"))
    if base_ordinal is None:
        return target

    parts = []
    for following in sections[section_index + 1:]:
        heading = (following.get("heading") or "").strip()
        next_ordinal = _review_label_ordinal(heading)
        if next_ordinal is not None and next_ordinal > base_ordinal:
            break
        content = (following.get("content") or "").strip()
        block = "\n".join(part for part in (heading, content) if part).strip()
        if block:
            parts.append(block)

    if not parts:
        return target
    expanded = deepcopy(target)
    expanded_text = "\n".join(parts).strip()
    expanded["text"] = expanded_text
    expanded["body_text"] = expanded_text
    expanded["match_method"] = f"{target.get('match_method') or 'matched'}+following_sections"
    return expanded


def _sync_review_direct_contract_sources_redlines(
    contract_docx: bytes,
    prospectus_docx: bytes,
    *,
    rules: list[dict] | None = None,
    review_author: str = DEFAULT_REVIEW_AUTHOR,
) -> tuple[bytes, dict]:
    contract_text = _docx_effective_text_from_bytes(contract_docx)
    prospectus_text = _docx_effective_text_from_bytes(prospectus_docx)
    contract_sections = _split_contract_sections(contract_text)
    prospectus_sections = _split_prospectus_sections(prospectus_text)
    candidate_rules = list(rules) if rules is not None else (
        _build_builtin_etf_cross_rules(contract_sections, prospectus_sections)
        + _build_review_prospectus_template_sync_rules()
    )

    report = {
        "status": "completed",
        "message": "已检查招募说明书正文中可确定为直接对应的合同来源内容",
        "checked_count": 0,
        "updated_count": 0,
        "unchanged_count": 0,
        "template_checked_count": 0,
        "template_updated_count": 0,
        "template_unchanged_count": 0,
        "skipped": [],
        "updated": [],
    }
    current_prospectus = prospectus_docx

    for rule in candidate_rules:
        is_template_source = _review_rule_is_prospectus_template_source(rule)
        ok, reason = (True, "") if is_template_source else _review_cross_rule_is_direct_sync_candidate(rule)
        label = f"{rule.get('contract_locator') or rule.get('contract_chapter') or ''} → {rule.get('prospectus_locator') or rule.get('prospectus_chapter') or ''}".strip()
        if not ok:
            report["skipped"].append({"label": label, "reason": reason})
            continue
        report["checked_count"] += 1
        if is_template_source:
            report["template_checked_count"] += 1

        prospectus_text = _docx_effective_text_from_bytes(current_prospectus)
        prospectus_sections = _split_prospectus_sections(prospectus_text)
        contract_target = _locate_review_rule_target(
            contract_sections,
            rule.get("contract_locator") or rule.get("contract_chapter") or "",
            rule.get("contract_section_name") or rule.get("contract_chapter") or "",
        )
        if not contract_target.get("matched"):
            contract_target = _fallback_review_rule_target_from_text(
                contract_text,
                rule.get("contract_locator") or rule.get("contract_chapter") or "",
                rule.get("contract_section_name") or rule.get("contract_chapter") or "",
            )
        prospectus_target = _locate_review_rule_target(
            prospectus_sections,
            rule.get("prospectus_locator") or rule.get("prospectus_chapter") or "",
            rule.get("prospectus_section_name") or rule.get("prospectus_chapter") or "",
        )
        if not prospectus_target.get("matched"):
            prospectus_target = _fallback_review_rule_target_from_text(
                prospectus_text,
                rule.get("prospectus_locator") or rule.get("prospectus_chapter") or "",
                rule.get("prospectus_section_name") or rule.get("prospectus_chapter") or "",
            )
        if is_template_source:
            prospectus_target = _expand_empty_review_target_from_following_sections(prospectus_target, prospectus_sections)
        if not contract_target.get("matched") or not prospectus_target.get("matched"):
            report["skipped"].append({
                "label": label,
                "reason": "locator_not_matched",
                "contract_matched": bool(contract_target.get("matched")),
                "prospectus_matched": bool(prospectus_target.get("matched")),
            })
            continue

        contract_source = (contract_target.get("body_text") or contract_target.get("text") or "").strip()
        prospectus_target_text = (prospectus_target.get("body_text") or prospectus_target.get("text") or "").strip()
        if not contract_source or not prospectus_target_text:
            report["skipped"].append({"label": label, "reason": "empty_source_or_target", "source": rule.get("source", "")})
            continue
        replacement_text = (
            _postprocess_review_prospectus_template_source(contract_source, rule)
            if is_template_source
            else contract_source
        )
        if not replacement_text:
            report["skipped"].append({"label": label, "reason": "empty_postprocessed_source", "source": rule.get("source", "")})
            continue
        minimal_ok, minimal_reason = _review_sync_minimal_change_guard(replacement_text, prospectus_target_text)
        if not minimal_ok:
            report["skipped"].append({"label": label, "reason": minimal_reason, "source": rule.get("source", "")})
            continue

        current_prospectus, apply_report = _redline_docx_paragraph_span_to_text(
            current_prospectus,
            prospectus_target_text,
            replacement_text,
            report_label=label or "招募说明书合同来源正文",
            review_author=review_author,
        )
        if apply_report.get("status") == "redlined":
            report["updated_count"] += 1
            if is_template_source:
                report["template_updated_count"] += 1
            report["updated"].append({"label": label, "source": rule.get("source", ""), "postprocess": rule.get("postprocess", ""), **apply_report})
        elif apply_report.get("status") == "unchanged":
            report["unchanged_count"] += 1
            if is_template_source:
                report["template_unchanged_count"] += 1
        else:
            report["skipped"].append({"label": label, "reason": apply_report.get("reason") or apply_report.get("message", ""), "source": rule.get("source", "")})

    return current_prospectus, report


def _build_revision_workbench_sync_report(
    *,
    contract_bytes: bytes,
    prospectus_bytes: bytes,
    product_summary_bytes: bytes | None,
    form_data: dict,
    document_reports: dict,
    summary_sync_report: dict | None = None,
    direct_sync_report: dict | None = None,
    product_summary_sync_report: dict | None = None,
    custody_sync_report: dict | None = None,
) -> dict:
    summary_report = (summary_sync_report or {}).get("contract_summary") or {
        "status": "skipped",
        "message": "未能自动重算基金合同摘要",
    }

    prospectus_report = {
        "status": "report_only",
        "message": "已保留应用审核决定后的招募说明书红线；无法确定为直接对应的合同来源片段未做猜测替换",
    }
    if summary_sync_report and summary_sync_report.get("prospectus_contract_summary"):
        prospectus_report["contract_summary_apply"] = summary_sync_report["prospectus_contract_summary"]
    if summary_report.get("status") == "generated":
        prospectus_report["contract_summary_source"] = "已生成新的基金合同摘要，可用于招募说明书内合同摘要同步"
    if direct_sync_report:
        prospectus_report["direct_contract_sources"] = direct_sync_report
    if custody_sync_report:
        prospectus_report["custody_summary_apply"] = custody_sync_report

    product_summary_report = {
        "status": "skipped",
        "message": "产品资料概要未上传，未生成该终稿",
    }
    if product_summary_bytes:
        product_summary_report = {
            "status": "kept_redline",
            "message": "已输出应用审核决定后的产品资料概要红线；本版未在缺少完整表单数据时强制重算",
        }
        if product_summary_sync_report:
            product_summary_report = product_summary_sync_report
        if form_data:
            product_summary_report["form_data_keys"] = sorted(str(key) for key in form_data.keys())[:30]

    return {
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
        "contract_summary": summary_report,
        "prospectus": prospectus_report,
        "custody_agreement": custody_sync_report or {"status": "skipped", "message": "托管协议未上传，未同步招募说明书托管协议摘要"},
        "product_summary": product_summary_report,
        "documents": document_reports,
    }


def _extract_review_docx_revision_items(docx_source, *, document_kind: str, filename: str = "") -> list[dict]:
    if hasattr(docx_source, "read"):
        current_pos = None
        try:
            current_pos = docx_source.tell()
        except Exception:
            current_pos = None
        docx_bytes = docx_source.read()
        if current_pos is not None:
            try:
                docx_source.seek(current_pos)
            except Exception:
                pass
    else:
        docx_bytes = bytes(docx_source or b"")

    xml_text = _read_docx_xml_part(docx_bytes, "word/document.xml")
    if not xml_text:
        return []

    import xml.etree.ElementTree as ET

    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return []

    comments = _load_docx_comments(docx_bytes)
    items = []
    scope = _review_revision_default_scope(document_kind)
    seen_comment_ids = set()

    for line_no, paragraph in enumerate(root.findall(".//w:p", _DOCX_NS), start=1):
        paragraph_text = _docx_xml_text(paragraph)
        if paragraph_text:
            scope = _review_revision_scope_for_paragraph(document_kind, paragraph_text, scope)
        context = paragraph_text

        for node in paragraph.iter():
            name = _docx_local_name(node.tag)
            if name not in {"ins", "del", "moveTo", "moveFrom", "commentRangeStart"}:
                continue

            if name == "commentRangeStart":
                comment_id = _docx_attr(node, "id")
                if not comment_id or comment_id in seen_comment_ids or comment_id not in comments:
                    continue
                seen_comment_ids.add(comment_id)
                comment = comments[comment_id]
                comment_text = comment.get("comment_text", "")
                if not comment_text:
                    continue
                items.append(
                    _build_review_revision_item(
                        document_kind=document_kind,
                        scope=scope,
                        filename=filename,
                        revision_type="comment",
                        text=comment_text,
                        line=line_no,
                        context=context,
                        author=comment.get("author", ""),
                        date=comment.get("date", ""),
                        comment_text=comment_text,
                    )
                )
                continue

            revision_text = _docx_xml_text(node)
            if not revision_text:
                continue
            revision_type = "insert" if name in {"ins", "moveTo"} else "delete"
            items.append(
                _build_review_revision_item(
                    document_kind=document_kind,
                    scope=scope,
                    filename=filename,
                    revision_type=revision_type,
                    text=revision_text,
                    line=line_no,
                    context=context,
                    author=_docx_attr(node, "author"),
                    date=_docx_attr(node, "date"),
                )
            )

    for index, item in enumerate(items, start=1):
        item["id"] = f"{item['document_scope']}-{index}"
    return items


def _count_review_revision_items(items: list[dict]) -> dict:
    counts = {scope: 0 for scope in _REVIEW_REVISION_SCOPE_LABELS}
    for item in items or []:
        scope = str((item or {}).get("document_scope") or "").strip()
        if scope:
            counts[scope] = counts.get(scope, 0) + 1
    return counts


def _split_review_blocks(text: str) -> list[dict]:
    text = (text or "").strip()
    if not text:
        return []

    marker_re = re.compile(
        r'^((?:'
        r'第[一二三四五六七八九十百千]+条'
        r'|[（(][一二三四五六七八九十百千]+[）)]'
        r'|[一二三四五六七八九十百千]+[、．]'
        r'|[（(]\d+[）)]'
        r'|\d+[.、．]'
        r')[^\n]*)',
        re.MULTILINE,
    )
    markers = list(marker_re.finditer(text))
    blocks = []
    for index, marker in enumerate(markers):
        start = marker.start()
        end = markers[index + 1].start() if index + 1 < len(markers) else len(text)
        heading = marker.group(1).strip()
        heading_end = marker.end()
        blocks.append({
            "heading": heading,
            "text": text[start:end].strip(),
            "body": text[heading_end:end].strip(),
        })
    return blocks


def _review_block_heading_style(heading: str) -> str:
    text = _normalize_review_text(heading)
    if re.match(r"^[一二三四五六七八九十百千]+[、．]", text):
        return "cn_enum"
    if re.match(r"^\d+[.、．]", text):
        return "digit_enum"
    if re.match(r"^[（(][一二三四五六七八九十百千]+[）)]", text):
        return "paren_cn"
    if re.match(r"^[（(]\d+[）)]", text):
        return "paren_digit"
    return "other"


def _expand_review_block_with_descendants(blocks: list[dict], start_index: int) -> dict:
    if start_index < 0 or start_index >= len(blocks):
        return {"text": "", "body": ""}

    block = blocks[start_index]
    base_style = _review_block_heading_style(block.get("heading"))
    descendant_styles = {
        "cn_enum": {"digit_enum", "paren_cn", "paren_digit"},
        "paren_cn": {"digit_enum", "paren_digit"},
    }.get(base_style, set())
    if not descendant_styles:
        return {
            "text": block.get("text") or "",
            "body": block.get("body") or "",
        }

    text_parts = [block.get("text") or ""]
    body_parts = [part for part in [block.get("body") or ""] if str(part).strip()]
    for idx in range(start_index + 1, len(blocks)):
        current = blocks[idx]
        current_style = _review_block_heading_style(current.get("heading"))
        if current_style == base_style:
            break
        if current_style not in descendant_styles:
            break
        text_parts.append(current.get("text") or "")
        body_parts.append(current.get("text") or "")

    return {
        "text": "\n".join(part for part in text_parts if str(part).strip()).strip(),
        "body": "\n".join(part for part in body_parts if str(part).strip()).strip(),
    }


def _match_review_block_by_label(text: str, label: str):
    text = (text or "").strip()
    label = _normalize_review_text(label)
    if not text or not label:
        return {"block": None, "match_method": "", "score": 0.0}

    best_block = None
    best_method = ""
    best_score = 0.0
    for block in _split_review_blocks(text):
        heading = block.get("heading", "")
        heading_text = _normalize_review_text(heading)
        heading_key = _review_heading_key(heading)
        label_key = _review_heading_key(label)
        if heading_text == label:
            return {"block": block, "match_method": "exact", "score": float(len(label))}
        if _review_labels_match(heading, label):
            score = float(min(len(heading_key or heading), len(label_key or label)))
            if heading_key and label_key and heading_key == label_key:
                score += 0.5
            if score > best_score or best_method != "normalized":
                best_block = block
                best_method = "normalized"
                best_score = score
            continue
        if _review_soft_heading_match(heading, label):
            heading_key = _review_soft_heading_key(heading)
            label_key = _review_soft_heading_key(label)
            if not heading_key or not label_key:
                continue
            score = difflib.SequenceMatcher(None, heading_key, label_key, autojunk=False).ratio()
            if best_method in ("", "soft_heading") and score > best_score:
                best_block = block
                best_method = "soft_heading"
                best_score = score
    return {"block": best_block, "match_method": best_method, "score": best_score}


def _find_review_block_by_label(text: str, label: str):
    return _match_review_block_by_label(text, label)["block"]


def _extract_review_block_by_label(text: str, label: str) -> str:
    text = (text or "").strip()
    label = _normalize_review_text(label)
    if not text or not label:
        return ""

    block = _find_review_block_by_label(text, label)
    return (block or {}).get("text", "")


def _locate_review_subheading(section, label_hint: str = "", content_hint: str = "", score_fn=None):
    result = {
        "section": section,
        "section_heading": _normalize_review_text(section.get("heading")) if section else "",
        "target_heading": "",
        "text": "",
        "body_text": "",
        "matched": False,
        "match_method": "",
        "missing_reason": "subheading_missing" if section else "section_missing",
        "locator_parts": [_normalize_review_text(section.get("heading"))] if section and _normalize_review_text(section.get("heading")) else [],
    }
    if not section:
        return result

    section_text = (section.get("content") or "").strip()
    if not section_text:
        return result

    blocks = _split_review_blocks(section_text)
    if not blocks:
        result["target_heading"] = result["section_heading"]
        result["text"] = section_text
        result["body_text"] = section_text
        result["matched"] = True
        result["missing_reason"] = ""
        result["locator_parts"] = [part for part in [result["section_heading"]] if part]
        return result

    label_hint = _normalize_review_text(label_hint)
    matched_info = _match_review_block_by_label(section_text, label_hint) if label_hint else {"block": None, "match_method": "", "score": 0.0}
    matched_block = matched_info["block"]
    if matched_block:
        matched_index = next((index for index, block in enumerate(blocks) if block == matched_block), -1)
        expanded_block = _expand_review_block_with_descendants(blocks, matched_index)
        result["target_heading"] = matched_block["heading"]
        result["text"] = expanded_block["text"] or matched_block["text"]
        result["body_text"] = expanded_block["body"] or matched_block.get("body") or matched_block["text"]
        result["matched"] = True
        result["match_method"] = matched_info["match_method"]
        result["missing_reason"] = ""
        result["locator_parts"] = [
            part for part in [result["section_heading"], _normalize_review_text(matched_block["heading"])] if part
        ]
        return result

    if content_hint and score_fn:
        best_block = None
        best_score = 0.0
        for block in blocks:
            score = score_fn(content_hint, block["body"] or block["text"])
            if score > best_score:
                best_score = score
                best_block = block
        if best_block is not None and best_score >= SUMMARY_FALLBACK_SUBHEADING_SCORE_THRESHOLD:
            result["target_heading"] = best_block["heading"]
            result["text"] = best_block["text"]
            result["body_text"] = best_block.get("body") or best_block["text"]
            result["matched"] = True
            result["match_method"] = "fallback"
            result["missing_reason"] = ""
            result["locator_parts"] = [
                part for part in [result["section_heading"], _normalize_review_text(best_block["heading"])] if part
            ]
            return result

    return result


def _locate_review_rule_target(sections, locator: str, section_hint: str = ""):
    locator = _normalize_review_locator(locator)
    result = {
        "section": None,
        "section_heading": None,
        "target_heading": locator,
        "text": "",
        "body_text": "",
        "matched": False,
        "match_method": "",
        "missing_reason": "section_missing" if locator else "",
        "locator_parts": [],
    }
    if not locator:
        return result
    if _is_missing_review_locator(locator):
        result["missing_reason"] = ""
        return result

    segments = _split_review_locator(locator)
    section_candidates = []
    sub_labels = []
    if segments and ("第" in segments[0] or "部分" in segments[0] or "章" in segments[0]):
        section_candidates.append(segments[0])
        sub_labels = segments[1:]
    else:
        sub_labels = segments

    section_hint = _normalize_review_text(section_hint)
    if section_hint:
        section_candidates.insert(0, section_hint)
    if locator not in section_candidates:
        section_candidates.append(locator)

    for candidate in section_candidates:
        section_match = _match_review_section(sections, candidate)
        section = section_match["section"]
        if not section:
            continue
        result["section"] = section
        result["section_heading"] = section.get("heading")
        result["target_heading"] = section.get("heading")
        result["match_method"] = section_match["match_method"]
        result["locator_parts"] = [_normalize_review_text(section.get("heading"))] if _normalize_review_text(section.get("heading")) else []
        section_text = section.get("content", "")
        effective_sub_labels = list(sub_labels)
        if effective_sub_labels and _review_labels_match(section.get("heading"), effective_sub_labels[0]):
            effective_sub_labels = effective_sub_labels[1:]
        if not effective_sub_labels or (len(effective_sub_labels) == 1 and _review_labels_match(section.get("heading"), effective_sub_labels[0])):
            result["text"] = section_text
            result["body_text"] = section_text
            result["matched"] = True
            result["missing_reason"] = ""
            return result

        current_text = section_text
        current_body = section_text
        locator_parts = list(result["locator_parts"])
        for label in effective_sub_labels:
            block_match = _match_review_block_by_label(current_text, label)
            section_block_match = _match_review_block_by_label(section_text, label)
            method_rank = {"exact": 3, "normalized": 2, "soft_heading": 1, "": 0}
            if (
                section_block_match["block"]
                and (
                    not block_match["block"]
                    or method_rank.get(section_block_match.get("match_method") or "", 0) > method_rank.get(block_match.get("match_method") or "", 0)
                    or (
                        method_rank.get(section_block_match.get("match_method") or "", 0) == method_rank.get(block_match.get("match_method") or "", 0)
                        and section_block_match.get("score", 0.0) > block_match.get("score", 0.0)
                    )
                )
            ):
                block_match = section_block_match
            block = block_match["block"]
            result["target_heading"] = block.get("heading") if block else label
            if not block:
                result["text"] = ""
                result["matched"] = False
                result["missing_reason"] = "subheading_missing"
                return result
            current_text = block.get("text", "")
            current_body = block.get("body") or current_text
            if block_match["match_method"]:
                result["match_method"] = block_match["match_method"]
            normalized_heading = _normalize_review_text(block.get("heading"))
            if normalized_heading and (not locator_parts or not _review_labels_match(locator_parts[-1], normalized_heading)):
                locator_parts.append(normalized_heading)
        result["text"] = current_text.strip()
        result["body_text"] = current_body.strip()
        result["matched"] = bool(result["text"])
        result["missing_reason"] = ""
        result["locator_parts"] = locator_parts
        return result

    return result


@app.route("/api/review/upload", methods=["POST"])
def api_review_upload():
    """上传基金合同DOCX + 可选招募说明书DOCX，解析并存储"""
    from docx import Document

    if "contract" not in request.files:
        return jsonify({"error": "请上传基金合同DOCX文件"}), 400

    contract_file = request.files["contract"]
    contract_bytes = contract_file.read()
    contract_doc = Document(io.BytesIO(contract_bytes))
    contract_text_lines = []
    for block_type, block in _iter_docx_blocks(contract_doc):
        if block_type == "paragraph":
            t = (block.text or "").strip()
            if not _is_layout_blank_line(t):
                contract_text_lines.append(t)
        else:
            tl = _docx_table_to_markdown_lines(block)
            if tl:
                contract_text_lines.extend(tl)
    contract_text = "\n".join(contract_text_lines)
    contract_sections = _split_contract_sections(contract_text)
    inferred_fund_name = _infer_review_fund_name(contract_text, contract_file.filename)

    result = {
        "contract_text": contract_text,
        "contract_sections": contract_sections,
        "contract_filename": contract_file.filename,
        "fund_name": inferred_fund_name,
        "contract_doc_meta": _extract_review_doc_metadata(
            contract_doc,
            filename=contract_file.filename,
            title_reference=inferred_fund_name or contract_file.filename,
        ),
        "prospectus_text": None,
        "prospectus_sections": [],
        "prospectus_filename": None,
        "prospectus_doc_meta": {},
    }

    if "prospectus" in request.files and request.files["prospectus"].filename:
        pros_file = request.files["prospectus"]
        pros_bytes = pros_file.read()
        pros_doc = Document(io.BytesIO(pros_bytes))
        pros_lines = []
        for block_type, block in _iter_docx_blocks(pros_doc):
            if block_type == "paragraph":
                t = (block.text or "").strip()
                if not _is_layout_blank_line(t):
                    pros_lines.append(t)
            else:
                tl = _docx_table_to_markdown_lines(block)
                if tl:
                    pros_lines.extend(tl)
        pros_text = "\n".join(pros_lines)
        pros_doc2 = Document(io.BytesIO(pros_bytes))
        pros_sections = _split_prospectus_sections(pros_text, pros_doc2)
        prospectus_fund_name = _infer_review_fund_name(pros_text, pros_file.filename) or result["fund_name"]
        result["prospectus_text"] = pros_text
        result["prospectus_sections"] = pros_sections
        result["prospectus_filename"] = pros_file.filename
        result["prospectus_doc_meta"] = _extract_review_doc_metadata(
            pros_doc,
            filename=pros_file.filename,
            title_reference=prospectus_fund_name or pros_file.filename,
        )

    _review_store["data"] = result

    # 复核规则表改为自动按基金名称/基金类型选择，清理旧的手动上传痕迹
    _review_store.pop("rules_xlsx_path", None)

    return jsonify({
        "success": True,
        "fund_name": result.get("fund_name", ""),
        "contract_sections": len(contract_sections),
        "contract_section_titles": [s["heading"] for s in contract_sections],
        "prospectus_sections": len(result["prospectus_sections"]),
        "prospectus_section_titles": [s["heading"] for s in result["prospectus_sections"]],
    })


@app.route("/api/revision_workbench/upload", methods=["POST"])
def api_revision_workbench_upload():
    """上传修订版 DOCX，抽取修订痕迹并按四件套影响链分类。"""
    from docx import Document

    file_specs = (
        ("contract", "contract", "基金合同"),
        ("prospectus", "prospectus", "招募说明书"),
        ("custody_agreement", "custody_agreement", "托管协议"),
        ("custody", "custody_agreement", "托管协议"),
        ("product_summary", "product_summary", "产品资料概要"),
    )
    uploaded = []
    revision_items = []
    documents = []
    texts = {}
    filenames = {}
    file_bytes = {}

    for field_name, document_kind, label in file_specs:
        file = request.files.get(field_name)
        if not file or not file.filename:
            continue
        docx_bytes = file.read()
        uploaded.append(label)
        filenames[document_kind] = file.filename
        file_bytes[document_kind] = docx_bytes
        document_view = _build_review_revision_document_view(
            docx_bytes,
            document_kind=document_kind,
            filename=file.filename,
        )
        documents.append(document_view)
        revision_items.extend(
            _extract_review_docx_revision_items(
                docx_bytes,
                document_kind=document_kind,
                filename=file.filename,
            )
        )

        try:
            doc = Document(io.BytesIO(docx_bytes))
            lines = []
            for block_type, block in _iter_docx_blocks(doc):
                if block_type == "paragraph":
                    text = (block.text or "").strip()
                    if not _is_layout_blank_line(text):
                        lines.append(text)
                else:
                    table_lines = _docx_table_to_markdown_lines(block)
                    if table_lines:
                        lines.extend(table_lines)
            texts[f"{field_name}_text"] = "\n".join(lines)
        except Exception:
            texts[f"{field_name}_text"] = ""

    if not uploaded:
        return jsonify({"error": "请至少上传一份带修订痕迹的 DOCX 文件"}), 400

    revision_counts = _count_review_revision_items(revision_items)
    data = {
        "uploaded": uploaded,
        "filenames": filenames,
        "file_bytes": file_bytes,
        "documents": documents,
        "revision_items": revision_items,
        "revision_counts": revision_counts,
        **texts,
    }
    _revision_workbench_store["data"] = data

    return jsonify({
        "success": True,
        "uploaded": uploaded,
        "filenames": filenames,
        "documents": documents,
        "revision_items": revision_items,
        "revision_total": len(revision_items),
        "revision_counts": revision_counts,
        "scope_labels": _REVIEW_REVISION_SCOPE_LABELS,
        "impact_targets": _REVIEW_REVISION_IMPACT_TARGETS,
    })


@app.route("/api/revision_workbench/export_final", methods=["POST"])
def api_revision_workbench_export_final():
    """根据审核工作台决策导出红线终稿 ZIP，不覆盖原文件。"""
    data = _revision_workbench_store.get("data") or {}
    file_bytes = data.get("file_bytes") or {}
    if not file_bytes:
        return jsonify({"error": "请先上传修订版 DOCX 文件"}), 400
    if not file_bytes.get("contract") or not file_bytes.get("prospectus"):
        return jsonify({"error": "生成红线终稿至少需要上传基金合同和招募说明书"}), 400

    payload = request.get_json(force=True, silent=True) or {}
    decisions = payload.get("decisions") or {}
    revision_edits = _normalize_review_revision_edits(payload.get("revision_edits") or {})
    review_author = _normalize_review_author(payload.get("review_author"))
    form_data = payload.get("form_data") or {}
    session_operations = payload.get("session_operations") or []
    documents = data.get("documents") or []
    revisions = [
        revision
        for document in documents
        for revision in (document or {}).get("revisions", [])
        if revision.get("id")
    ]
    pending_ids = [
        revision["id"]
        for revision in revisions
        if _normalize_review_export_decision(decisions.get(revision["id"])) == "pending"
    ]
    if pending_ids:
        return jsonify({
            "error": f"仍有 {len(pending_ids)} 条待定修订，请先全部标记为接受或拒绝后再生成红线终稿",
            "pending_count": len(pending_ids),
            "pending_ids": pending_ids[:50],
        }), 400

    filenames = data.get("filenames") or {}
    document_reports = {}
    try:
        contract_docx, contract_report = _apply_review_revision_decisions_to_docx(
            file_bytes["contract"],
            document_kind="contract",
            filename=filenames.get("contract", ""),
            decisions=decisions,
            revision_edits=revision_edits,
            review_author=review_author,
        )
        prospectus_docx, prospectus_report = _apply_review_revision_decisions_to_docx(
            file_bytes["prospectus"],
            document_kind="prospectus",
            filename=filenames.get("prospectus", ""),
            decisions=decisions,
            revision_edits=revision_edits,
            review_author=review_author,
        )
        document_reports["contract"] = contract_report
        document_reports["prospectus"] = prospectus_report
        contract_docx, prospectus_docx, summary_sync_report = _sync_review_contract_summary_redlines(
            contract_docx,
            prospectus_docx,
            review_author=review_author,
        )
        prospectus_docx, direct_sync_report = _sync_review_direct_contract_sources_redlines(
            contract_docx,
            prospectus_docx,
            review_author=review_author,
        )

        custody_docx = None
        custody_sync_report = None
        if file_bytes.get("custody_agreement"):
            custody_docx, custody_report = _apply_review_revision_decisions_to_docx(
                file_bytes["custody_agreement"],
                document_kind="custody_agreement",
                filename=filenames.get("custody_agreement", ""),
                decisions=decisions,
                revision_edits=revision_edits,
                review_author=review_author,
            )
            document_reports["custody_agreement"] = custody_report
            prospectus_docx, custody_sync_report = _sync_review_custody_summary_redlines(
                custody_docx,
                prospectus_docx,
                review_author=review_author,
            )

        product_summary_docx = None
        product_summary_sync_report = None
        if file_bytes.get("product_summary"):
            product_summary_docx, product_summary_report = _apply_review_revision_decisions_to_docx(
                file_bytes["product_summary"],
                document_kind="product_summary",
                filename=filenames.get("product_summary", ""),
                decisions=decisions,
                revision_edits=revision_edits,
                review_author=review_author,
            )
            document_reports["product_summary"] = product_summary_report
            product_summary_docx, product_summary_sync_report = _sync_review_product_summary_from_prospectus_redlines(
                product_summary_docx,
                prospectus_docx,
                form_data=form_data,
                filename=filenames.get("product_summary", ""),
                review_author=review_author,
            )

        sync_report = _build_revision_workbench_sync_report(
            contract_bytes=contract_docx,
            prospectus_bytes=prospectus_docx,
            product_summary_bytes=product_summary_docx,
            form_data=form_data,
            document_reports=document_reports,
            summary_sync_report=summary_sync_report,
            direct_sync_report=direct_sync_report,
            product_summary_sync_report=product_summary_sync_report,
            custody_sync_report=custody_sync_report,
        )
        sync_report["review_author"] = review_author
        applied_edit_ids = {
            revision_id
            for report in document_reports.values()
            for revision_id in (report.get("edited_revision_ids") or [])
        }

        decision_record = {
            "generated_at": sync_report["generated_at"],
            "review_author": review_author,
            "decisions": decisions,
            "revision_edits": revision_edits,
            "edited_count": len(applied_edit_ids),
            "applied_edit_ids": sorted(applied_edit_ids),
            "session_operations": session_operations,
            "documents": [
                {
                    "document_kind": (document or {}).get("document_kind", ""),
                    "filename": (document or {}).get("filename", ""),
                    "revision_count": len((document or {}).get("revisions", []) or []),
                }
                for document in documents
            ],
        }
        csv_text = _build_review_decision_csv(decisions, documents, revision_edits, applied_edit_ids)
        contract_docx = _finalize_review_docx_for_export(contract_docx)
        prospectus_docx = _finalize_review_docx_for_export(prospectus_docx)
        if custody_docx:
            custody_docx = _finalize_review_docx_for_export(custody_docx)
        if product_summary_docx:
            product_summary_docx = _finalize_review_docx_for_export(product_summary_docx)

        output = io.BytesIO()
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("基金合同.docx", contract_docx)
            zf.writestr("招募说明书.docx", prospectus_docx)
            if custody_docx:
                zf.writestr("托管协议.docx", custody_docx)
            if product_summary_docx:
                zf.writestr("产品资料概要.docx", product_summary_docx)
            zf.writestr("审核决策记录.json", json.dumps(decision_record, ensure_ascii=False, indent=2).encode("utf-8"))
            zf.writestr("审核决策记录.csv", "\ufeff" + csv_text)
            zf.writestr("同步报告.json", json.dumps(sync_report, ensure_ascii=False, indent=2).encode("utf-8"))
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name="审核工作台红线终稿.zip",
            mimetype="application/zip",
        )
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": f"生成红线终稿失败：{exc}"}), 500


@app.route("/api/review/get_text", methods=["GET"])
def api_review_get_text():
    """返回已上传的合同全文（供摘要核对等使用）"""
    data = _review_store.get("data")
    if not data:
        return jsonify({"error": "请先上传文件"}), 400
    return jsonify({"contract_text": data["contract_text"]})


@app.route("/api/review/cross_check", methods=["POST"])
def api_review_cross_check():
    """合同与招募说明书章节交叉比对"""
    import difflib
    import html as html_mod

    data = _review_store.get("data")
    if not data or not data.get("prospectus_text"):
        return jsonify({"error": "请先上传合同和招募说明书"}), 400

    req = request.get_json(force=True) or {}
    fund_type = req.get("fund_type", "ETF")
    review_rules = _load_review_rules(fund_type)

    contract_sections = data["contract_sections"]
    pros_sections = data["prospectus_sections"]

    rules = _build_review_cross_rules(review_rules, fund_type)
    if not rules and fund_type == "ETF":
        rules = _build_builtin_etf_cross_rules(contract_sections, pros_sections)

    re_num = re.compile(
        r'^(?:第[一二三四五六七八九十百千]+条'
        r'|[（(][一二三四五六七八九十百千]+[）)]'
        r'|[一二三四五六七八九十百千]+[、．]'
        r'|[（(]\d+[）)]'
        r'|\d+[.、．])\s*'
    )
    heading_cleanup_re = re.compile(r'[()（）]')

    def strip_num(line):
        return re_num.sub('', (line or '').strip())

    def normalize_compare_line(line):
        return strip_num(_normalize_contract_prospectus_compare_text(line))

    def normalized_heading_tokens(name):
        text = unicodedata.normalize("NFKC", str(name or "")).strip()
        if not text:
            return []
        parts = [part.strip() for part in re.split(r'[+＋]', text) if part and str(part).strip()]
        tokens = []
        for part in parts or [text]:
            tokens.append(part)
            cleaned = re.sub(r'\([^)]*\)', '', part).strip()
            cleaned = re.sub(r'[（(][^）)]*[）)]', '', cleaned).strip()
            if cleaned:
                tokens.append(cleaned)
        deduped = []
        for token in tokens:
            if token and token not in deduped:
                deduped.append(token)
        return deduped

    def heading_key(value):
        text = unicodedata.normalize("NFKC", str(value or ""))
        text = re.sub(r'\([^)]*\)', '', text)
        text = re.sub(r'[（(][^）)]*[）)]', '', text)
        text = heading_cleanup_re.sub('', text)
        return re.sub(r'[\s的与及和、，,：:；;/-]', '', text)

    def find_section(sections, chapter_name):
        return _find_review_section(sections, chapter_name)

    def char_diff(a, b):
        matcher = difflib.SequenceMatcher(None, a, b, autojunk=False)
        ah, bh = [], []
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            ea = html_mod.escape(a[i1:i2])
            eb = html_mod.escape(b[j1:j2])
            if tag == "equal":
                ah.append(ea)
                bh.append(eb)
            elif tag == "replace":
                ah.append(f'<del>{ea}</del>')
                bh.append(f'<ins>{eb}</ins>')
            elif tag == "delete":
                ah.append(f'<del>{ea}</del>')
            elif tag == "insert":
                bh.append(f'<ins>{eb}</ins>')
        return "".join(ah), "".join(bh)

    def build_diff_hunks(lines_c, lines_p, compare_c, compare_p, nums_c, nums_p):
        matcher = difflib.SequenceMatcher(None, compare_c, compare_p, autojunk=False)
        diff_lines = []
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == "equal":
                continue
            if tag == "replace":
                for k in range(max(i2 - i1, j2 - j1)):
                    al = lines_c[i1 + k] if i1 + k < i2 else ""
                    bl = lines_p[j1 + k] if j1 + k < j2 else ""
                    ah, bh = char_diff(al, bl)
                    diff_lines.append({
                        "contract_ln": nums_c[i1 + k] if i1 + k < i2 else None,
                        "prospectus_ln": nums_p[j1 + k] if j1 + k < j2 else None,
                        "contract_text": al,
                        "prospectus_text": bl,
                        "contract_html": ah,
                        "prospectus_html": bh,
                        "tag": "replace",
                    })
            elif tag == "delete":
                for k in range(i1, i2):
                    diff_lines.append({
                        "contract_ln": nums_c[k],
                        "prospectus_ln": None,
                        "contract_text": lines_c[k],
                        "prospectus_text": "",
                        "contract_html": '<del>' + html_mod.escape(lines_c[k]) + '</del>',
                        "prospectus_html": "",
                        "tag": "delete",
                    })
            elif tag == "insert":
                for k in range(j1, j2):
                    diff_lines.append({
                        "contract_ln": None,
                        "prospectus_ln": nums_p[k],
                        "contract_text": "",
                        "prospectus_text": lines_p[k],
                        "contract_html": "",
                        "prospectus_html": '<ins>' + html_mod.escape(lines_p[k]) + '</ins>',
                        "tag": "insert",
                    })
        return matcher.ratio(), diff_lines[:80]

    def rule_mode(rule):
        relation = str(rule.get("relation") or "").strip()
        consistency = str(rule.get("consistency") or "").strip()
        detail = str(rule.get("detail") or "").strip()
        summary = " ".join(x for x in [relation, consistency, detail] if x)

        allowed_diff_hints = (
            "招募细化",
            "招募补充",
            "新增",
            "补充",
            "细化",
            "略有差异",
            "自指",
            "口径",
            "说明书口径重述",
            "详见招募说明书",
            "侧袋机制",
            "本基金合同",
        )
        allowed_diff_text = " ".join(
            str(rule.get(key) or "").strip()
            for key in ("relation", "consistency", "detail", "expected_diff")
        )
        has_allowed_differences = any(hint in allowed_diff_text for hint in allowed_diff_hints)

        if ("完全一致" in relation or "完全一致" in consistency) and not has_allowed_differences:
            return "strict"
        if has_allowed_differences or "个别表述差异" in summary or consistency in ("基本一致", "基本对应") or "基本一致" in summary:
            return "normalized"
        return "descriptive"

    def descriptive_message(rule):
        consistency = str(rule.get("consistency") or "").strip()
        relation = str(rule.get("relation") or "").strip()
        detail = str(rule.get("detail") or "").strip()
        headline = "；".join([part for part in [consistency, relation] if part])
        if detail:
            return f"{headline}：{detail}" if headline else detail
        return headline or "该章节按勾稽规则属于说明性对应关系"

    def missing_message(side, reason):
        side_label = "合同" if side == "contract" else "招募"
        if reason == "subheading_missing":
            return f"章节已找到，但{side_label}子条款未找到"
        if reason == "section_missing":
            return f"{side_label}一级章节未找到"
        return f"{side_label}章节未找到"

    results = []
    for rule in rules:
        mode = rule_mode(rule)
        c_target = _locate_review_rule_target(
            contract_sections,
            rule.get("contract_locator") or rule["contract_chapter"],
            rule.get("contract_section_name") or rule["contract_chapter"],
        )
        p_target = _locate_review_rule_target(
            pros_sections,
            rule.get("prospectus_locator") or rule["prospectus_chapter"],
            rule.get("prospectus_section_name") or rule["prospectus_chapter"],
        )
        c_sec = c_target["section"]
        p_sec = p_target["section"]

        item = {
            "rule": rule,
            "contract_found": bool(c_target["matched"]),
            "prospectus_found": bool(p_target["matched"]),
            "contract_heading": c_sec["heading"] if c_sec else None,
            "prospectus_heading": p_sec["heading"] if p_sec else None,
            "contract_target": c_target["target_heading"] if c_target["matched"] else (c_target["target_heading"] or None),
            "prospectus_target": p_target["target_heading"] if p_target["matched"] else (p_target["target_heading"] or None),
            "contract_locator": (
                _format_review_target_locator(c_target)
                if c_target["matched"]
                else str(rule.get("contract_locator") or rule.get("contract_chapter") or "")
            ),
            "prospectus_locator": (
                _format_review_target_locator(p_target)
                if p_target["matched"]
                else str(rule.get("prospectus_locator") or rule.get("prospectus_chapter") or "")
            ),
            "status": "skip",
            "severity": "info",
            "message": "",
            "hunks": [],
            "rule_source": str(rule.get("source") or ""),
            "sheet_name": str(rule.get("sheet_name") or ""),
            "expected_relation": str(rule.get("relation") or ""),
            "expected_diff": str(rule.get("expected_diff") or rule.get("detail") or ""),
            "suggestion": str(rule.get("suggestion") or ""),
            "similarity": None,
            "contract_missing_reason": c_target.get("missing_reason") or "",
            "prospectus_missing_reason": p_target.get("missing_reason") or "",
            "contract_match_method": c_target.get("match_method") or "",
            "prospectus_match_method": p_target.get("match_method") or "",
            "contract_anchor_text": "",
            "prospectus_anchor_text": "",
            "contract_context_excerpt": "",
            "prospectus_context_excerpt": "",
        }

        missing_parts = []
        if not c_target["matched"] and not _review_cross_section_optional(rule, "contract"):
            missing_parts.append(missing_message("contract", c_target.get("missing_reason")))
        if not p_target["matched"] and not _review_cross_section_optional(rule, "prospectus"):
            missing_parts.append(missing_message("prospectus", p_target.get("missing_reason")))

        if missing_parts:
            item["status"] = "missing"
            item["severity"] = "warning"
            item["message"] = "；".join(part for part in missing_parts if part)
            item["is_problem"] = _cross_check_result_is_problem(item)
            results.append(item)
            continue

        if mode == "descriptive":
            item["status"] = "pass"
            item["severity"] = "info"
            item["message"] = descriptive_message(rule)
            item["is_problem"] = _cross_check_result_is_problem(item)
            results.append(item)
            continue

        c_text = (c_target.get("body_text") or c_target["text"]) if c_target["matched"] else ""
        p_text = (p_target.get("body_text") or p_target["text"]) if p_target["matched"] else ""
        rows_c = _collect_nonblank_compare_lines(c_text, normalize_line=normalize_compare_line)
        rows_p = _collect_nonblank_compare_lines(p_text, normalize_line=normalize_compare_line)
        lines_c = [row["line"] for row in rows_c]
        lines_p = [row["line"] for row in rows_p]
        compare_c = [row["compare"] for row in rows_c]
        compare_p = [row["compare"] for row in rows_p]
        nums_c = [row["lineno"] for row in rows_c]
        nums_p = [row["lineno"] for row in rows_p]

        if mode == "strict":
            ratio, hunks = build_diff_hunks(lines_c, lines_p, compare_c, compare_p, nums_c, nums_p)
            item["similarity"] = round(ratio * 100, 1)
            if ratio > 0.98:
                item["status"] = "pass"
                item["severity"] = "info"
                item["message"] = f"内容一致 (相似度 {ratio*100:.1f}%)"
            else:
                item["status"] = "fail"
                item["severity"] = "error"
                item["message"] = f"应完全一致但存在差异 (相似度 {ratio*100:.1f}%)"
                item["hunks"] = hunks
        elif mode == "normalized":
            ratio, hunks = build_diff_hunks(lines_c, lines_p, compare_c, compare_p, nums_c, nums_p)
            item["similarity"] = round(ratio * 100, 1)
            if ratio > 0.95:
                item["status"] = "pass"
                item["severity"] = "info"
                item["message"] = f"标准化后一致 (相似度 {ratio*100:.1f}%)"
            elif str(rule.get("detail") or rule.get("expected_diff") or "").strip():
                item["status"] = "pass"
                item["severity"] = "info"
                item["message"] = descriptive_message(rule)
            else:
                item["status"] = "warning"
                item["severity"] = "warning"
                item["message"] = f"存在表述差异 (标准化后相似度 {ratio*100:.1f}%)"
                item["hunks"] = hunks

        if item["status"] in {"fail", "warning"} and (c_text or p_text):
            item.update(
                _build_review_excerpt_pair(
                    c_text,
                    p_text,
                    normalize_line=normalize_compare_line,
                    left_prefix="contract",
                    right_prefix="prospectus",
                )
            )

        item["is_problem"] = _cross_check_result_is_problem(item)
        results.append(item)

    deduped_results = []
    seen_result_keys = set()
    for item in results:
        dedupe_key = (
            item.get("contract_locator"),
            item.get("prospectus_locator"),
            item.get("status"),
            item.get("message"),
            item.get("rule_source"),
            item.get("expected_relation"),
        )
        if dedupe_key in seen_result_keys:
            continue
        seen_result_keys.add(dedupe_key)
        deduped_results.append(item)

    return jsonify({"results": deduped_results, "total_rules": len(deduped_results)})


@app.route("/api/review/format_check", methods=["POST"])
def api_review_format_check():
    """格式与标点检查"""
    import html as html_mod
    data = _review_store.get("data")
    if not data:
        return jsonify({"error": "请先上传文件"}), 400

    text = data["contract_text"]
    lines = text.splitlines()
    issues = []

    def _ctx(line_idx, highlight_start=None, highlight_end=None):
        """构建上下文：返回当前行全文 + 前后各2行，问题片段高亮标记"""
        ctx_lines = []
        for offset in (-2, -1, 0, 1, 2):
            li = line_idx + offset
            if 0 <= li < len(lines):
                raw = lines[li]
                if offset == 0 and highlight_start is not None:
                    # 在原文中用 «» 标记问题位置
                    hl_end = highlight_end if highlight_end else highlight_start + 2
                    marked = (html_mod.escape(raw[:highlight_start])
                              + '<mark>' + html_mod.escape(raw[highlight_start:hl_end]) + '</mark>'
                              + html_mod.escape(raw[hl_end:]))
                    ctx_lines.append({"ln": li + 1, "text": raw, "html": marked, "current": True})
                else:
                    ctx_lines.append({"ln": li + 1, "text": raw, "html": html_mod.escape(raw), "current": offset == 0})
        return ctx_lines

    def _ctx_full_line(line_idx):
        """仅返回整行原文 + 前后各2行的纯文本上下文"""
        return _ctx(line_idx)

    # ── 标题连续性检查 ──
    # 先找到摘要章节的起始行，摘要内容会重复 "第X部分" 标题，不应纳入检查
    # 使用最后一个匹配（跳过目录中的同名行）
    RE_PART_HEADING = re.compile(r'^第([一二三四五六七八九十百]+)部分')
    # 收集正文中的 "第X部分" 标题（过滤目录行和摘要内部重复标题）
    # 目录行特征：含 tab+页码（如 "第一部分  前言\t1"）
    # 摘要内部特征：在最后一个含"摘要"的 "第X部分" 之后
    RE_TOC_LINE = re.compile(r'\t\d+\s*$')

    # 先找摘要起始行（取最后一个不含tab的摘要行）
    summary_start_line = len(lines)
    for i, line in enumerate(lines):
        stripped = line.strip()
        if '摘要' in stripped and RE_PART_HEADING.match(stripped) and not RE_TOC_LINE.search(line):
            summary_start_line = i

    part_nums = []
    for i, line in enumerate(lines):
        if i >= summary_start_line:
            break  # 摘要之后的 "第X部分" 不参与连续性检查
        stripped = line.strip()
        if RE_TOC_LINE.search(line):
            continue  # 跳过目录行
        m = RE_PART_HEADING.match(stripped)
        if m:
            n = ContractEngine._cn_numeral_to_int(m.group(1))
            if n is not None:
                part_nums.append((i + 1, n, stripped, i))

    for idx in range(1, len(part_nums)):
        ln, num, txt, li = part_nums[idx]
        _, prev_num, _, _ = part_nums[idx - 1]
        if num != prev_num + 1:
            issues.append({
                "line": ln, "type": "heading_sequence",
                "severity": "error",
                "description": f"标题序号跳号: 第{prev_num}部分 → 第{num}部分 (期望第{prev_num+1}部分)",
                "text": txt,
                "context": _ctx_full_line(li),
                "suggestion": f"检查是否遗漏第{prev_num+1}部分",
            })

    # 章节内: 一、→二、→三、
    RE_CN_SEQ = re.compile(r'^([一二三四五六七八九十]+)、')
    current_list = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if RE_PART_HEADING.match(stripped):
            current_list = []
            continue
        m = RE_CN_SEQ.match(stripped)
        if m:
            cn = m.group(1)
            n = ContractEngine._cn_numeral_to_int(cn)
            if n is not None:
                if current_list and n != current_list[-1][1] + 1 and n != 1:
                    issues.append({
                        "line": i + 1, "type": "heading_sequence",
                        "severity": "warning",
                        "description": f"章节序号可能跳号: {current_list[-1][2]}、 → {cn}、",
                        "text": stripped[:80],
                        "context": _ctx_full_line(i),
                        "suggestion": "检查序号连续性",
                    })
                if n == 1:
                    current_list = []
                current_list.append((i + 1, n, cn))

    # 子项: (一)→(二)→(三)
    RE_PAREN_SEQ = re.compile(r'^[（(]([一二三四五六七八九十]+)[）)]')
    paren_list = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if RE_PART_HEADING.match(stripped) or RE_CN_SEQ.match(stripped):
            paren_list = []
            continue
        m = RE_PAREN_SEQ.match(stripped)
        if m:
            cn = m.group(1)
            n = ContractEngine._cn_numeral_to_int(cn)
            if n is not None:
                if paren_list and n != paren_list[-1][1] + 1 and n != 1:
                    issues.append({
                        "line": i + 1, "type": "heading_sequence",
                        "severity": "warning",
                        "description": f"子项序号可能跳号: ({paren_list[-1][2]}) → ({cn})",
                        "text": stripped[:80],
                        "context": _ctx_full_line(i),
                        "suggestion": "检查序号连续性",
                    })
                if n == 1:
                    paren_list = []
                paren_list.append((i + 1, n, cn))

    # 数字序号按样式分层检查，避免把 "1、" / "1）" / "（1）" 混为同一级别
    digit_patterns = {
        "plain_comma": re.compile(r'^(\d+)、(?=\s|$|[\u4e00-\u9fffA-Za-z]|[《〈“"\'【\[])'),
        "plain_dot": re.compile(r'^(\d+)[.．](?=\s|$|[\u4e00-\u9fffA-Za-z]|[《〈“"\'【\[])'),
        "plain_rparen": re.compile(r'^(\d+)[）)](?=\s|$|[\u4e00-\u9fffA-Za-z]|[《〈“"\'【\[])'),
        "wrapped_paren": re.compile(r'^[（(](\d+)[）)](?=\s|$|[\u4e00-\u9fffA-Za-z]|[《〈“"\'【\[])'),
    }
    digit_lists = {name: [] for name in digit_patterns}
    for i, line in enumerate(lines):
        stripped = line.strip()
        if RE_PART_HEADING.match(stripped) or RE_CN_SEQ.match(stripped):
            digit_lists = {name: [] for name in digit_patterns}
            continue
        for style_name, pattern in digit_patterns.items():
            m = pattern.match(stripped)
            if not m:
                continue
            n = int(m.group(1))
            digit_list = digit_lists[style_name]
            if digit_list and n != digit_list[-1] + 1 and n != 1:
                issues.append({
                    "line": i + 1, "type": "heading_sequence",
                    "severity": "warning",
                    "description": f"数字序号跳号: {digit_list[-1]} → {n}",
                    "text": stripped[:80],
                    "context": _ctx_full_line(i),
                    "suggestion": "检查序号连续性",
                })
            if n == 1:
                digit_list = []
            digit_list.append(n)
            digit_lists[style_name] = digit_list
            break

    # ── 标点符号混用 ──
    RE_EN_PUNCT = re.compile(r'[\u4e00-\u9fff][,.:;!?]|[,.:;!?][\u4e00-\u9fff]')
    EN_PAREN_IN_CN = re.compile(r'[\u4e00-\u9fff]\(|\)[\u4e00-\u9fff]')
    SINGLE_RIGHT_PAREN_ENUM_RE = re.compile(r'^\s*(?:\d+|[A-Za-z]|[一二三四五六七八九十百千]+)\uFF09$')

    def _is_allowed_single_right_paren_enumerator(line: str, char_index: int, opener: str, closer: str) -> bool:
        if opener != "\uFF08" or closer != "\uFF09":
            return False
        prefix = (line or "")[: char_index + 1]
        if not SINGLE_RIGHT_PAREN_ENUM_RE.match(prefix):
            return False
        return "\uFF08" not in prefix[:-1]

    for i, line in enumerate(lines):
        for m in RE_EN_PUNCT.finditer(line):
            issues.append({
                "line": i + 1, "type": "punctuation",
                "severity": "warning",
                "description": f"中文语境使用英文标点: '{m.group()}'",
                "text": line.strip()[:120],
                "context": _ctx(i, m.start(), m.end()),
                "suggestion": "建议使用中文标点",
            })
        for m in EN_PAREN_IN_CN.finditer(line):
            issues.append({
                "line": i + 1, "type": "punctuation",
                "severity": "warning",
                "description": f"中文语境使用英文括号: '{m.group()}'",
                "text": line.strip()[:120],
                "context": _ctx(i, m.start(), m.end()),
                "suggestion": "建议使用中文括号（）",
            })

    # ── 引号/括号配对 ──
    PAIRS = [("\u201c", "\u201d"), ("\u300a", "\u300b"), ("\uff08", "\uff09")]
    for opener, closer in PAIRS:
        # 用栈匹配找出所有未配对的位置；对“1）/2）/3）”这类合法编号样式，不参与括号配对统计。
        stack = []  # 存 (line_idx, char_idx, char) 的开符号
        unmatched = []  # 未配对的 (line_idx, char_idx, char)
        effective_open = 0
        effective_close = 0
        for li, ln in enumerate(lines):
            for ci, ch in enumerate(ln):
                if ch == opener:
                    stack.append((li, ci, ch))
                    effective_open += 1
                elif ch == closer:
                    if _is_allowed_single_right_paren_enumerator(ln, ci, opener, closer):
                        continue
                    effective_close += 1
                    if stack:
                        stack.pop()
                    else:
                        unmatched.append((li, ci, ch))
        # 栈中剩余的是未配对的开符号
        unmatched.extend(stack)
        if unmatched:
            # 按行号排序，取前5个
            unmatched.sort(key=lambda x: (x[0], x[1]))
            unmatched_ctxs = []
            for li, ci, ch in unmatched[:5]:
                unmatched_ctxs.append(_ctx(li, ci, ci + 1))

            issues.append({
                "line": unmatched[0][0] + 1 if unmatched else 0,
                "type": "bracket_mismatch",
                "severity": "warning",
                "description": f"符号不配对: '{opener}' 出现{effective_open}次, '{closer}' 出现{effective_close}次 (差{abs(effective_open-effective_close)}个)",
                "text": "",
                "context": unmatched_ctxs[0] if unmatched_ctxs else [],
                "extra_locations": unmatched_ctxs[1:] if len(unmatched_ctxs) > 1 else [],
                "suggestion": f"以下{len(unmatched_ctxs)}处为疑似未配对位置，请逐个检查",
            })

    # ── 多余空格 ──
    RE_EXTRA_SPACE = re.compile(r'(?<=[\u4e00-\u9fff])\s{1,}(?=[\u4e00-\u9fff])')
    # 白名单：这些模式中的多余空格是正常的
    # - "第X部分  标题" 标题格式
    # - "年    月    日" 待填写日期
    # - "签订日期：" 后的空格
    RE_SPACE_OK = re.compile(
        r'第[一二三四五六七八九十百]+部分\s'  # 标题格式
        r'|[年月]\s+[年月日]'                  # 日期占位
        r'|签订日'                              # 签订日期行
        r'|^目\s+录$'                           # 目录标题
    )
    for i, line in enumerate(lines):
        if RE_SPACE_OK.search(line):
            continue
        for m in RE_EXTRA_SPACE.finditer(line):
            issues.append({
                "line": i + 1, "type": "extra_space",
                "severity": "info",
                "description": "中文之间存在多余空格",
                "text": line.strip()[:120],
                "context": _ctx(i, m.start(), m.end()),
                "suggestion": "删除多余空格",
            })

    # ── DOCX 空白页 / 页眉与标题一致性 ──
    review_docs = [
        ("contract_doc_meta", "contract_filename", "基金合同"),
        ("prospectus_doc_meta", "prospectus_filename", "招募说明书"),
    ]
    for meta_key, filename_key, doc_label in review_docs:
        doc_meta = data.get(meta_key) or {}
        if not doc_meta:
            continue

        blank_pages = [int(page) for page in (doc_meta.get("blank_pages") or []) if str(page).isdigit()]
        if blank_pages:
            locations = []
            for page in blank_pages[:5]:
                location_text = f"{doc_label}第{page}页疑似空白页"
                locations.append({"ln": page, "text": location_text, "html": html_mod.escape(location_text)})
            issues.append({
                "type": "blank_page",
                "severity": "warning",
                "description": f"{doc_label}检测到 {len(blank_pages)} 处疑似空白页",
                "locations": locations,
                "suggestion": "检查分页符、分节符或多余空段落是否造成空白页",
            })

        header_entries = doc_meta.get("header_entries") or []
        title_reference = (
            str(doc_meta.get("title_reference") or "").strip()
            or str(data.get("fund_name") or "").strip()
            or _fallback_review_filename(data.get(filename_key, ""))
        )
        title_key = _normalize_review_doc_title_key(title_reference)
        mismatched_headers = []
        for entry in header_entries:
            header_text = str(entry.get("text") or "").strip()
            header_key = _normalize_review_doc_title_key(header_text)
            if not header_key or not title_key:
                continue
            if title_key in header_key or header_key in title_key:
                continue
            similarity = difflib.SequenceMatcher(None, header_key, title_key, autojunk=False).ratio()
            if similarity >= 0.86 and min(len(header_key), len(title_key)) >= 6:
                continue
            mismatched_headers.append(entry)

        if mismatched_headers:
            locations = []
            for entry in mismatched_headers[:5]:
                section_no = entry.get("section") or ""
                header_kind = "首页页眉" if entry.get("kind") == "first_page" else "页眉"
                location_text = f"第{section_no}节{header_kind}: {entry.get('text', '')}"
                locations.append({
                    "ln": section_no,
                    "text": location_text,
                    "html": html_mod.escape(location_text),
                })
            issues.append({
                "type": "header_title_mismatch",
                "severity": "warning",
                "description": f"{doc_label}页眉与文件标题不一致",
                "locations": locations,
                "suggestion": "核对页眉中的基金名称和文种是否与文件标题保持一致",
            })

    # ── 交叉引用验证 ──
    RE_XREF = re.compile(r'详见第([一二三四五六七八九十百]+)部分')
    existing_parts = set()
    for _, num, _, _ in part_nums:
        existing_parts.add(num)
    for i, line in enumerate(lines):
        for m in RE_XREF.finditer(line):
            ref_num = ContractEngine._cn_numeral_to_int(m.group(1))
            if ref_num and ref_num not in existing_parts:
                issues.append({
                    "line": i + 1, "type": "xref_broken",
                    "severity": "error",
                    "description": f"交叉引用指向不存在的章节: {m.group()}",
                    "text": line.strip()[:120],
                    "context": _ctx(i, m.start(), m.end()),
                    "suggestion": f"第{ref_num}部分不存在，请核查引用",
                })

    return jsonify({"issues": issues, "total": len(issues)})


@app.route("/api/review/consistency_check", methods=["POST"])
def api_review_consistency_check():
    """前后文一致性检查"""
    import html as html_mod
    data = _review_store.get("data")
    if not data:
        return jsonify({"error": "请先上传文件"}), 400

    text = data["contract_text"]
    lines = text.splitlines()
    issues = []

    def _find_occurrences(pattern_re, group_idx=0, max_hits=5):
        """在 lines 中找到匹配正则的所有行，返回 [{ln, text, html}]"""
        hits = []
        for i, line in enumerate(lines):
            for m in pattern_re.finditer(line):
                hl = (html_mod.escape(line[:m.start()])
                      + '<mark>' + html_mod.escape(m.group(group_idx)) + '</mark>'
                      + html_mod.escape(line[m.end():]))
                hits.append({"ln": i + 1, "text": line.strip()[:150], "html": hl})
                if len(hits) >= max_hits:
                    return hits
        return hits

    # ── 基金名称一致性 ──
    # 基金全称必须以"证券投资基金"结尾（不匹配截断的"基金"）
    # 排除断句干扰：不允许中间出现句号/分号/换行
    RE_FUND_NAME = re.compile(
        r'(南方[^，。；：\n""]{4,50}证券投资基金)'
    )
    fund_names_found = []
    for i, line in enumerate(lines):
        for m in RE_FUND_NAME.finditer(line):
            fund_names_found.append((m.group(1), i, m.start(), m.end()))

    if fund_names_found:
        name_variants = {}
        for fn, li, s, e in fund_names_found:
            name_variants.setdefault(fn, []).append((li, s, e))
        if len(name_variants) > 1:
            sorted_names = sorted(name_variants.items(), key=lambda x: -len(x[1]))
            most_common = sorted_names[0][0]
            for name, occurrences in sorted_names[1:]:
                # 跳过是主名称子串的变体（如简称包含在全称中）
                if most_common.startswith(name) or name.startswith(most_common):
                    continue
                if name != most_common and len(name) > 10:
                    locs = []
                    for li, s, e in occurrences[:5]:
                        raw = lines[li]
                        hl = (html_mod.escape(raw[:s])
                              + '<mark>' + html_mod.escape(raw[s:e]) + '</mark>'
                              + html_mod.escape(raw[e:]))
                        locs.append({"ln": li + 1, "text": raw.strip()[:150], "html": hl})
                    issues.append({
                        "type": "fund_name",
                        "severity": "warning",
                        "description": f"基金名称不一致: '{name}' (出现{len(occurrences)}次) vs 主要名称 '{most_common}' (出现{len(sorted_names[0][1])}次)",
                        "suggestion": "请确认基金全称是否统一",
                        "locations": locs,
                    })

    # ── 当事人名称一致性 ──
    # 管理人/托管人名称通常以"公司"结尾，用更精确的模式避免断句
    RE_MANAGER = re.compile(r'基金管理人[：:]\s*(.+?(?:公司|有限公司))')
    mgr_hits = {}
    for i, line in enumerate(lines):
        for m in RE_MANAGER.finditer(line):
            val = m.group(1).strip().rstrip('（(')
            # 剥离释义中的"指"前缀，如 "指南方基金管理股份有限公司"
            if val.startswith('指'):
                val = val[1:]
            if len(val) < 3 or len(val) > 30:
                continue
            mgr_hits.setdefault(val, []).append((i, m.start(), m.end()))
    if len(mgr_hits) > 1:
        locs = []
        for val, occ in mgr_hits.items():
            for li, s, e in occ[:3]:
                raw = lines[li]
                hl = (html_mod.escape(raw[:s])
                      + '<mark>' + html_mod.escape(raw[s:e]) + '</mark>'
                      + html_mod.escape(raw[e:]))
                locs.append({"ln": li + 1, "text": raw.strip()[:150], "html": hl, "value": val})
        issues.append({
            "type": "party_name",
            "severity": "error",
            "description": f"基金管理人名称不一致: {', '.join(mgr_hits.keys())}",
            "suggestion": "统一基金管理人名称",
            "locations": locs,
        })

    RE_CUSTODIAN = re.compile(r'基金托管人[：:]\s*(.+?(?:公司|有限公司|银行))')
    cust_hits = {}
    for i, line in enumerate(lines):
        for m in RE_CUSTODIAN.finditer(line):
            val = m.group(1).strip().rstrip('（(')
            if val.startswith('指'):
                val = val[1:]
            if len(val) < 3 or len(val) > 30:
                continue
            cust_hits.setdefault(val, []).append((i, m.start(), m.end()))
    if len(cust_hits) > 1:
        locs = []
        for val, occ in cust_hits.items():
            for li, s, e in occ[:3]:
                raw = lines[li]
                hl = (html_mod.escape(raw[:s])
                      + '<mark>' + html_mod.escape(raw[s:e]) + '</mark>'
                      + html_mod.escape(raw[e:]))
                locs.append({"ln": li + 1, "text": raw.strip()[:150], "html": hl, "value": val})
        issues.append({
            "type": "party_name",
            "severity": "error",
            "description": f"基金托管人名称不一致: {', '.join(cust_hits.keys())}",
            "suggestion": "统一基金托管人名称",
            "locations": locs,
        })

    # ── 数值一致性 ──
    RE_RATE = re.compile(r'(管理费|托管费|销售服务费|申购费|赎回费)([^\n]{0,20}?)(\d+\.?\d*)\s*%')
    rate_map = {}
    for i, line in enumerate(lines):
        for m in RE_RATE.finditer(line):
            key = m.group(1)
            val = m.group(3)
            rate_map.setdefault(key, {}).setdefault(val, []).append((i, m.start(), m.end()))

    for key, val_dict in rate_map.items():
        if len(val_dict) > 1:
            locs = []
            for val, occ in val_dict.items():
                for li, s, e in occ[:3]:
                    raw = lines[li]
                    hl = (html_mod.escape(raw[:s])
                          + '<mark>' + html_mod.escape(raw[s:e]) + '</mark>'
                          + html_mod.escape(raw[e:]))
                    locs.append({"ln": li + 1, "text": raw.strip()[:150], "html": hl, "value": f"{val}%"})
            issues.append({
                "type": "numeric",
                "severity": "warning",
                "description": f"{key}费率出现多个不同数值: {', '.join(v + '%' for v in val_dict.keys())}",
                "suggestion": "请核查费率数值是否正确",
                "locations": locs,
            })

    return jsonify({"issues": issues, "total": len(issues)})


REVIEW_XLSX_CANDIDATES = [
    PACKAGED_ASSETS_DIR / "review_workbooks" / filename for filename in REVIEW_WORKBOOK_FILENAMES
]


def _normalize_review_text(value):
    text = str(value or "").strip()
    if not text:
        return ""
    text = re.sub(r"[（(]\s*[PpＰｐ]\s*\d+\s*[）)]", "", text)
    return re.sub(r"\s+", " ", text).strip()


def _normalize_review_workbook_key(value: str) -> str:
    text = unicodedata.normalize("NFKC", str(value or "")).strip()
    if not text:
        return ""
    text = Path(text).stem
    for marker in ("勾稽关系整理", "勾稽规则", "规则整理", "规则表"):
        text = text.replace(marker, "")
    text = re.sub(r"[《》【】\[\]（）()_\-\s]+", "", text)
    return text


def _enrich_summary_rule_entry(entry: dict) -> dict:
    summary_meta = _parse_summary_rule_path(entry.get("summary_pos"))
    contract_locator = _normalize_review_locator(entry.get("contract_pos"))
    chapter_meta = _extract_review_chapter_meta(contract_locator)
    enriched = dict(entry)
    enriched["summary_pos"] = _normalize_review_text(entry.get("summary_pos"))
    enriched["contract_pos"] = contract_locator
    enriched["normalized_summary_path"] = summary_meta["normalized_summary_path"]
    enriched["summary_group_path"] = summary_meta["summary_group_path"]
    enriched["summary_group_heading"] = summary_meta["summary_group_heading"]
    enriched["summary_detail_heading"] = summary_meta["summary_detail_heading"]
    enriched["normalized_contract_locator"] = contract_locator
    enriched["contract_locator_parts"] = _split_review_locator(contract_locator)
    enriched["contract_chapter_title"] = chapter_meta["title"]
    enriched["contract_chapter_ordinal"] = chapter_meta["ordinal"]
    return enriched


def _find_first_existing_review_workbook(fund_type="ETF"):
    review_path = _review_store.get("rules_xlsx_path")
    if review_path and Path(review_path).exists():
        return str(review_path)

    review_candidates = []
    for candidate in REVIEW_XLSX_CANDIDATES:
        if candidate not in review_candidates:
            review_candidates.append(candidate)
    for candidate in LEGACY_REVIEW_XLSX_CANDIDATES:
        if candidate not in review_candidates:
            review_candidates.append(candidate)

    existing_candidates = [p for p in review_candidates if p.exists()]
    if not existing_candidates:
        return None

    stored_data = _review_store.get("data") or {}
    fund_name = str(stored_data.get("fund_name") or "").strip()
    fund_key = _normalize_review_workbook_key(fund_name)
    if fund_key:
        for candidate in existing_candidates:
            candidate_key = _normalize_review_workbook_key(candidate.name)
            if candidate_key and fund_key == candidate_key:
                return str(candidate)

    return None


def _build_builtin_etf_summary_rules(subsections) -> list[dict]:
    if not subsections:
        return []

    heading_matches = 0
    for subsection in subsections:
        heading = _normalize_review_text(subsection.get("heading"))
        if any(_review_labels_match(heading, rule["summary_pos"]) for rule in BUILTIN_ETF_SUMMARY_RULES):
            heading_matches += 1

    if heading_matches < min(7, len(BUILTIN_ETF_SUMMARY_RULES)):
        return []

    return [
        _enrich_summary_rule_entry({
            **rule,
            "sheet_name": "内置ETF摘要规则",
            "rule_level": "chapter",
            "detail": rule.get("method", ""),
            "similarity": "",
        })
        for rule in BUILTIN_ETF_SUMMARY_RULES
    ]


def _review_cross_section_optional(rule, side):
    chapter_name = str(
        rule.get(f"{side}_chapter")
        or rule.get(f"{side}_locator")
        or rule.get(f"{side}_section_name")
        or ""
    ).strip()
    combined = " ".join(
        str(rule.get(key) or "").strip()
        for key in (
            "relation",
            "consistency",
            "detail",
            "expected_diff",
            f"{side}_chapter",
            f"{side}_locator",
            f"{side}_section_name",
        )
    )
    if any(keyword in chapter_name for keyword in ("无直接对应", "无同名独立条款", "无对应独立章节")):
        return True
    if side == "contract" and any(keyword in combined for keyword in ("招募独有", "无直接对应", "无同名独立条款", "无对应独立章节")):
        return True
    if side == "prospectus" and "合同独有" in combined:
        return True
    return False


def _build_builtin_etf_cross_rules(contract_sections, prospectus_sections) -> list[dict]:
    built_rules = []
    for rule in BUILTIN_ETF_CROSS_RULES:
        contract_optional = _review_cross_section_optional(rule, "contract")
        prospectus_optional = _review_cross_section_optional(rule, "prospectus")
        contract_locator = str(rule.get("contract_locator") or rule.get("contract_chapter") or "").strip()
        prospectus_locator = str(rule.get("prospectus_locator") or rule.get("prospectus_chapter") or "").strip()
        contract_hint = str(rule.get("contract_section_name") or rule.get("contract_chapter") or "").strip()
        prospectus_hint = str(rule.get("prospectus_section_name") or rule.get("prospectus_chapter") or "").strip()

        contract_matched = contract_optional or bool(
            _locate_review_rule_target(contract_sections, contract_locator, contract_hint).get("matched")
        )
        prospectus_matched = prospectus_optional or bool(
            _locate_review_rule_target(prospectus_sections, prospectus_locator, prospectus_hint).get("matched")
        )
        if contract_matched and prospectus_matched:
            built_rules.append(dict(rule))
    return built_rules


def _summary_rule_expects_partial(rule: dict | None) -> bool:
    if not rule:
        return False
    text = " ".join(
        _normalize_review_text((rule or {}).get(field))
        for field in ("status", "method", "detail")
    )
    return any(
        keyword in text
        for keyword in ("部分", "选择性摘录", "选择性", "节录", "仅节录", "仅纳入", "零散吸收")
    )


def _summary_rule_requires_strict_text_match(rule: dict | None) -> bool:
    if not rule or _summary_rule_expects_partial(rule):
        return False
    text = " ".join(
        _normalize_review_text((rule or {}).get(field))
        for field in ("status", "method", "detail")
    )
    return any(
        keyword in text
        for keyword in ("已收录", "全文摘录", "近全文摘录", "全文收录", "完全摘录", "完全收录")
    )


def _find_sheet_by_keywords(sheetnames, include_keywords, optional_keywords=()):
    for sheet_name in sheetnames:
        if all(keyword in sheet_name for keyword in include_keywords):
            if optional_keywords and not any(keyword in sheet_name for keyword in optional_keywords):
                continue
            return sheet_name
    return None


def _find_header_row(ws, required_keywords, max_row=12):
    required_keywords = tuple(required_keywords)
    for row_index, row in enumerate(ws.iter_rows(min_row=1, max_row=max_row, values_only=True), 1):
        values = [_normalize_review_text(cell) for cell in row]
        joined = "|".join(values)
        if all(keyword in joined for keyword in required_keywords):
            return row_index, values
    return None, []


def _parse_review_workbook_meta(ws):
    meta = {
        "sheet_name": ws.title,
        "fund_type": "",
        "special_notes": [],
        "key_points": [],
        "raw_pairs": [],
    }
    current_bucket = None

    for row in ws.iter_rows(values_only=True):
        values = [_normalize_review_text(cell) for cell in row]
        if not any(values):
            continue

        key = next((value for value in values if value), "")
        key_index = values.index(key)
        tail = " ".join(value for value in values[key_index + 1:] if value).strip()
        if tail:
            meta["raw_pairs"].append({"key": key, "value": tail})

        if key == "基金类型" and tail:
            meta["fund_type"] = tail
            current_bucket = None
            continue
        if key == "特别说明":
            current_bucket = "special_notes"
            if tail:
                meta["special_notes"].append(tail)
            continue
        if key in {"本基金的几个关键判断", "重点结论"}:
            current_bucket = "key_points"
            if tail:
                meta["key_points"].append(tail)
            continue
        if re.match(r"^\d+[.、]", key):
            if current_bucket == "special_notes":
                meta["special_notes"].append(" ".join(part for part in (key, tail) if part).strip())
            elif current_bucket == "key_points":
                meta["key_points"].append(" ".join(part for part in (key, tail) if part).strip())
            continue
        if current_bucket == "special_notes" and tail:
            meta["special_notes"].append(tail)
        elif current_bucket == "key_points" and tail:
            meta["key_points"].append(tail)

    return meta


def _load_review_rules(fund_type="ETF"):
    """从参考Excel加载勾稽规则，供AI检查使用。
    自动检测 ETF 格式（sheet名含括号）和联接基金格式（sheet名含下划线）。
    """
    import openpyxl

    rules = {
        "cross": [],
        "summary": [],
        "chapter_level": [],
        "key_diffs": [],
        "workbook_meta": {
            "sheet_name": "",
            "fund_type": "",
            "special_notes": [],
            "key_points": [],
            "raw_pairs": [],
        },
        "chapter_rules": [],
        "detail_rules": [],
        "summary_rules": [],
        "key_diff_rules": [],
    }

    # 1. 基金合同与招募说明书规则.xlsx (通用规则)
    rules_xlsx = _resolve_rules_xlsx()
    if rules_xlsx.exists():
        wb = openpyxl.load_workbook(str(rules_xlsx), read_only=True)
        ws = wb.active
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row or not row[0]:
                continue
            rules["cross"].append({
                "fund_type": str(row[0] or "").strip(),
                "contract": str(row[1] or "").strip(),
                "prospectus": str(row[2] or "").strip(),
                "relation": str(row[3] or "").strip(),
                "prompt": str(row[4] or "").strip(),
                "identical": str(row[5] or "").strip(),
                "diff": str(row[6] or "").strip(),
            })
        wb.close()

    # 2. 详细勾稽关系整理表 — 优先用上传的，否则按fund_type选候选文件
    review_path = _find_first_existing_review_workbook(fund_type)

    if review_path and Path(review_path).exists():
        wb = openpyxl.load_workbook(review_path, read_only=True)
        sheets = wb.sheetnames

        meta_sheet = _find_sheet_by_keywords(sheets, ("总览",)) or _find_sheet_by_keywords(sheets, ("说明",))
        if meta_sheet:
            rules["workbook_meta"] = _parse_review_workbook_meta(wb[meta_sheet])
        if not rules["workbook_meta"]["fund_type"]:
            rules["workbook_meta"]["fund_type"] = fund_type

        # ── 招募-合同 章节级 ──
        # ETF格式: "招募-合同(章节级)" cols: 序号|招募|合同|关系|一致性|相似度|说明
        # 联接格式: "招募-合同_总表"    cols: 序号|招募|合同|关系判定|勾稽结论
        cross_sheet = _find_sheet_by_keywords(sheets, ("招募", "合同"), ("章节", "总表"))
        if cross_sheet:
            ws = wb[cross_sheet]
            header_row, _ = _find_header_row(ws, ("序号",))
            if header_row:
                for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
                    if not row or not row[1]:
                        continue
                    vals = [_normalize_review_text(cell) for cell in row]
                    is_linked_format = "总表" in cross_sheet
                    entry = {
                        "prospectus": vals[1] if len(vals) > 1 else "",
                        "contract": vals[2] if len(vals) > 2 else "",
                        "relation": vals[3] if len(vals) > 3 else "",
                        "consistency": vals[4] if len(vals) > 4 else "",
                        "similarity": vals[5] if len(vals) > 5 else "",
                        "detail": (vals[4] if is_linked_format else (vals[6] if len(vals) > 6 else "")),
                        "sheet_name": cross_sheet,
                        "rule_level": "chapter",
                    }
                    rules["chapter_rules"].append(entry)
                    rules["chapter_level"].append(dict(entry))

        detail_sheet = _find_sheet_by_keywords(sheets, ("招募", "合同"), ("详细", "明细"))
        if detail_sheet:
            ws = wb[detail_sheet]
            header_row, _ = _find_header_row(ws, ("序号",))
            if header_row:
                is_linked_format = "明细" in detail_sheet and "总表" not in detail_sheet and "_" in detail_sheet
                for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
                    if not row or not row[1]:
                        continue
                    vals = [_normalize_review_text(cell) for cell in row]
                    entry = {
                        "prospectus": vals[1] if len(vals) > 1 else "",
                        "prospectus_point": vals[2] if len(vals) > 2 else "",
                        "contract": vals[3] if len(vals) > 3 else "",
                        "relation": vals[4] if len(vals) > 4 else "",
                        "consistency": vals[5] if len(vals) > 5 else "",
                        "similarity": vals[6] if len(vals) > 6 else "",
                        "detail": vals[6] if is_linked_format else (vals[7] if len(vals) > 7 else ""),
                        "sheet_name": detail_sheet,
                        "rule_level": "detail",
                    }
                    rules["detail_rules"].append(entry)

        # ── 正文-摘要 章节级 / 明细 ──
        chapter_summary_entries = []
        summary_total_sheet = _find_sheet_by_keywords(sheets, ("正文", "摘要"), ("章节", "总表"))
        if summary_total_sheet:
            ws = wb[summary_total_sheet]
            header_row, _ = _find_header_row(ws, ("序号",))
            if header_row:
                is_etf_format = "章节级" in summary_total_sheet
                for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
                    if not row or not row[1]:
                        continue
                    vals = [_normalize_review_text(cell) for cell in row]
                    entry = {
                        "contract_pos": vals[1] if len(vals) > 1 else "",
                        "summary_pos": vals[2] if len(vals) > 2 else "",
                        "status": vals[3] if len(vals) > 3 else "",
                        "method": vals[4] if len(vals) > 4 else "",
                        "similarity": vals[5] if len(vals) > 5 else "",
                        "detail": (vals[6] if is_etf_format and len(vals) > 6 else (vals[4] if len(vals) > 4 else "")),
                        "sheet_name": summary_total_sheet,
                        "rule_level": "chapter",
                    }
                    chapter_summary_entries.append(_enrich_summary_rule_entry(entry))

        # ETF格式: "正文-摘要(详细)" cols: 序号|摘要位置|正文位置|收录情况|收录方式|相似度|说明
        # 联接格式: "正文-摘要_明细"  cols: 序号|正文位置|摘要位置|纳入方式|一致性|差异点|备注
        summary_sheet = _find_sheet_by_keywords(sheets, ("正文", "摘要"), ("详细", "明细"))
        if summary_sheet:
            ws = wb[summary_sheet]
            header_row, header_values = _find_header_row(ws, ("序号",))
            if header_row:
                col2_name = header_values[1] if len(header_values) > 1 else ""
                is_etf_format = "摘要" in col2_name
                for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
                    if not row or not row[1]:
                        continue
                    vals = [_normalize_review_text(cell) for cell in row]
                    if is_etf_format:
                        entry = {
                            "summary_pos": vals[1] if len(vals) > 1 else "",
                            "contract_pos": vals[2] if len(vals) > 2 else "",
                            "status": vals[3] if len(vals) > 3 else "",
                            "method": vals[4] if len(vals) > 4 else "",
                            "similarity": vals[5] if len(vals) > 5 else "",
                            "detail": vals[6] if len(vals) > 6 else "",
                            "sheet_name": summary_sheet,
                            "rule_level": "detail",
                        }
                    else:
                        entry = {
                            "contract_pos": vals[1] if len(vals) > 1 else "",
                            "summary_pos": vals[2] if len(vals) > 2 else "",
                            "method": vals[3] if len(vals) > 3 else "",
                            "status": vals[4] if len(vals) > 4 else "",
                            "detail": vals[5] if len(vals) > 5 else "",
                            "similarity": "",
                            "sheet_name": summary_sheet,
                            "rule_level": "detail",
                        }
                    enriched_entry = _enrich_summary_rule_entry(entry)
                    rules["summary_rules"].append(enriched_entry)
                    rules["summary"].append(dict(enriched_entry))

        for entry in chapter_summary_entries:
            rules["summary_rules"].append(entry)
            rules["summary"].append(dict(entry))

        # ── 重点差异 / 参考规则 ──
        for sn in sheets:
            if '重点差异' in sn or '参考规则' in sn:
                ws = wb[sn]
                header_row = None
                header_values = []
                if '重点差异' in sn:
                    header_row, header_values = _find_header_row(ws, ("序号",))
                else:
                    header_row, header_values = _find_header_row(ws, ("基金类型", "合同", "招募"))
                if header_row:
                    for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
                        if not row or not row[1]:
                            continue
                        vals = [_normalize_review_text(cell) for cell in row]
                        if '重点差异' in sn:
                            entry = {
                                "category": vals[1] if len(vals) > 1 else "",
                                "location": vals[2] if len(vals) > 2 else "",
                                "description": vals[3] if len(vals) > 3 else "",
                                "suggestion": vals[4] if len(vals) > 4 else "",
                                "sheet_name": sn,
                            }
                        else:
                            entry = {
                                "category": vals[0] if len(vals) > 0 else "",
                                "location": " ↔ ".join(part for part in (vals[1] if len(vals) > 1 else "", vals[2] if len(vals) > 2 else "") if part),
                                "description": vals[6] if len(vals) > 6 else (vals[3] if len(vals) > 3 else ""),
                                "suggestion": vals[4] if len(vals) > 4 else "",
                                "sheet_name": sn,
                                "identical": vals[5] if len(vals) > 5 else "",
                            }
                        rules["key_diff_rules"].append(entry)
                        rules["key_diffs"].append(dict(entry))

        wb.close()
    return rules


def _build_cross_rules_prompt(rules, fund_type="ETF"):
    """将规则整理成AI可参考的文本"""
    lines = []
    lines.append("【勾稽规则参考】以下是基金合同与招募说明书各章节的预期对应关系：\n")

    # 参考规则
    for r in rules.get("cross", []):
        if r["fund_type"] in ("通用", fund_type):
            line = f"- 合同「{r['contract']}」↔ 招募「{r['prospectus']}」: {r['relation']}"
            if r["diff"]:
                line += f"\n  已知差异: {r['diff'][:200]}"
            if r["prompt"]:
                line += f"\n  检查要点: {r['prompt'][:200]}"
            lines.append(line)

    # 章节级详细规则
    for r in rules.get("chapter_level", []):
        detail = r.get("detail", "") or r.get("consistency", "")
        relation = r.get("relation", "")
        if detail and relation not in ("完全对应", "无直接对应"):
            lines.append(
                f"- 招募「{r['prospectus']}」↔ 合同「{r['contract']}」: {relation}"
                f"\n  说明: {detail[:250]}"
            )

    # 重点差异提示
    if rules.get("key_diffs"):
        lines.append("\n【重点差异提示】以下是本基金勾稽中值得特别关注的差异点：\n")
        for kd in rules["key_diffs"]:
            lines.append(
                f"- [{kd['category']}] {kd['location']}: {kd['description'][:250]}"
                + (f"\n  建议: {kd['suggestion'][:150]}" if kd.get("suggestion") else "")
            )

    return "\n".join(lines)


def _build_summary_rules_prompt(rules):
    """将正文-摘要规则整理成AI参考文本"""
    lines = ["【正文-摘要勾稽规则】以下是基金合同正文与摘要的预期对应关系：\n"]
    for r in rules.get("summary", []):
        line = f"- {r['summary_pos']} ↔ {r['contract_pos']}: {r['status']}({r['method']})"
        if r["detail"]:
            line += f" — {r['detail'][:150]}"
        lines.append(line)
    return "\n".join(lines)


@app.route("/api/review/ai_check", methods=["POST"])
def api_review_ai_check():
    """AI智能检查（使用DeepSeek API），基于勾稽规则进行深度审查。
    使用SSE流式推送，每完成一项检查即发送结果，避免浏览器超时。"""
    from flask import Response, stream_with_context

    data = _review_store.get("data")
    if not data:
        return jsonify({"error": "请先上传文件"}), 400

    api_key = os.environ.get("DEEPSEEK_API_KEY", "")
    if not api_key:
        return jsonify({"error": "未配置DEEPSEEK_API_KEY环境变量", "issues": []}), 200

    req = request.get_json(force=True) or {}
    fund_type = req.get("fund_type", "ETF")

    # 将整个AI检查逻辑放在生成器里，通过SSE逐条推送
    def generate():
        rules = _load_review_rules(fund_type)
        has_prospectus = bool(data.get("prospectus_text"))
        contract_sections = data["contract_sections"]
        pros_sections = data.get("prospectus_sections", [])

        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com", timeout=60)
        except ImportError:
            yield f"data: {json.dumps({'type':'error','message':'未安装openai库'})}\n\n"
            yield "data: {\"type\":\"done\"}\n\n"
            return

        def _call_ai(prompt, chunk_label):
            """调用AI并yield SSE事件"""
            issues = []
            try:
                response = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.1,
                    max_tokens=3000,
                )
                resp_text = response.choices[0].message.content.strip()
                if resp_text.startswith("```"):
                    resp_text = re.sub(r'^```\w*\n?', '', resp_text)
                    resp_text = re.sub(r'\n?```$', '', resp_text)
                result = json.loads(resp_text)
                for issue in result.get("issues", []):
                    issue["chapter"] = chunk_label
                    issue["source"] = "AI"
                    issues.append(issue)
            except Exception as e:
                issues.append({
                    "chapter": chunk_label,
                    "type": "error",
                    "description": f"AI检查出错: {str(e)}",
                    "source": "AI",
                })
            return issues

        # ── 1. 合同↔招募 规则驱动的章节级对比 ──
        if has_prospectus:
            checked = 0
            for item in _build_ai_cross_review_items(contract_sections, pros_sections, rules, fund_type):
                if checked >= 8:
                    break
                rule = item["rule"]
                specific_prompt = (
                    str(rule.get("prompt") or "").strip()
                    or str(rule.get("detail") or "").strip()
                    or str(rule.get("consistency") or "").strip()
                )
                known_diff = str(rule.get("expected_diff") or rule.get("diff") or "").strip()
                relation = str(rule.get("relation") or "").strip() or "按勾稽规则对应"

                prompt = f"""你是基金法律文件复核专家。请根据勾稽规则审查以下基金合同与招募说明书的对应章节。

【勾稽规则】
章节: 合同「{item['contract_locator']}」↔ 招募「{item['prospectus_locator']}」
预期对应关系: {relation}
{"检查要点: " + specific_prompt if specific_prompt else ""}
{"已知合理差异（不需要标记为问题）: " + known_diff if known_diff else ""}

【基金合同原文】
{item['contract_text']}

【招募说明书原文】
{item['prospectus_text']}

请逐字逐句对比两段文本，重点检查：
1. 是否存在不应有的实质性内容差异（排除已知的合理差异如"本基金合同"→"本招募说明书"等文件自指口径）
2. 数值、费率、比例、日期等关键数据是否一致
3. 权利义务条款是否遗漏或矛盾
4. 招募说明书是否按预期补充了合同未覆盖的内容

以JSON格式返回：
{{"issues": [{{"location": "问题所在原文", "type": "内容差异/数据不一致/条款遗漏/表述矛盾", "description": "具体问题描述", "suggestion": "修改建议"}}]}}
没有问题则返回 {{"issues": []}}
只返回JSON。"""

                label = item["label"]
                yield f"data: {json.dumps({'type':'progress','message':f'正在检查: {label}'}, ensure_ascii=False)}\n\n"
                issues = _call_ai(prompt, label)
                for iss in issues:
                    yield f"data: {json.dumps({'type':'issue','data':iss}, ensure_ascii=False)}\n\n"
                checked += 1

        # ── 2. 正文↔摘要 规则驱动审查 ──
        summary_review_items = _build_ai_summary_review_items(contract_sections, rules)
        if summary_review_items:
            checked = 0
            for item in summary_review_items:
                if checked >= 5:
                    break
                rule = item["rule"]
                method = str(rule.get("method") or "").strip() or "摘录/转录"
                detail = str(rule.get("detail") or "").strip()

                prompt = f"""你是基金法律文件复核专家。请审查基金合同正文与合同内容摘要的对应关系。

【规则】
摘要位置: {item['summary_heading']}
对应正文: {item['contract_locator']}
预期收录方式: {method}
参考说明: {detail[:200]}

【合同正文原文】
{item['contract_text']}

【合同摘要原文】
{item['summary_text']}

请检查：
1. 摘要是否准确摘录了正文的核心条款
2. 摘要中是否存在与正文不一致的表述（数值、条件、比例等）
3. 关键条款是否遗漏

以JSON格式返回：
{{"issues": [{{"location": "问题所在原文", "type": "摘要不一致/关键遗漏/数据错误", "description": "具体问题描述", "suggestion": "修改建议"}}]}}
没有问题则返回 {{"issues": []}}
只返回JSON。"""

                label = item["label"]
                yield f"data: {json.dumps({'type':'progress','message':f'正在检查: {label}'}, ensure_ascii=False)}\n\n"
                issues = _call_ai(prompt, label)
                for iss in issues:
                    yield f"data: {json.dumps({'type':'issue','data':iss}, ensure_ascii=False)}\n\n"
                checked += 1

        # ── 3. 通用文本质量检查 ──
        key_chapters = ["前言", "基金的基本情况", "基金份额的发售", "违约责任", "争议", "基金合同的效力"]
        checked = 0
        for kw in key_chapters:
            if checked >= 3:
                break
            sec = _find_review_section(contract_sections, kw)
            if not sec or len(sec["content"]) < 50:
                continue

            prompt = f"""请审查以下基金法律文件章节，检查是否存在以下问题：
1. 错别字或语病
2. 金融/法律术语使用不当
3. 数值或日期明显错误
4. 条款内容前后矛盾

章节: {sec['heading']}
文本:
{sec['content'][:4000]}

以JSON格式返回：
{{"issues": [{{"location": "问题所在原文", "type": "错别字/语病/术语误用/数据错误/逻辑矛盾", "description": "问题描述", "suggestion": "修改建议"}}]}}
没有问题则返回 {{"issues": []}}
只返回JSON。"""

            label = sec["heading"]
            yield f"data: {json.dumps({'type':'progress','message':f'正在检查: {label}'}, ensure_ascii=False)}\n\n"
            issues = _call_ai(prompt, label)
            for iss in issues:
                yield f"data: {json.dumps({'type':'issue','data':iss}, ensure_ascii=False)}\n\n"
            checked += 1

        yield f"data: {json.dumps({'type':'done'})}\n\n"

    return Response(stream_with_context(generate()),
                    mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


@app.route("/api/review/summary_cross_check", methods=["POST"])
def api_review_summary_cross_check():
    """合同摘要 ↔ 招募摘要 交叉比对:
    比较合同中的"基金合同内容摘要"与招募说明书中的"基金合同的内容摘要" """
    import difflib
    import html as html_mod
    context_lines = 2

    data = _review_store.get("data")
    if not data:
        return jsonify({"error": "请先上传文件"}), 400
    if not data.get("prospectus_text"):
        return jsonify({"error": "请上传招募说明书"}), 400

    # 优先定位“基金合同内容摘要”，避免误选“基金托管协议的内容摘要”
    contract_summary = _find_contract_summary_section(data["contract_sections"])
    pros_summary = _find_contract_summary_section(data["prospectus_sections"])

    if not contract_summary:
        return jsonify({"error": "合同中未找到摘要章节"}), 400
    if not pros_summary:
        return jsonify({"error": "招募说明书中未找到合同摘要章节"}), 400

    # 整段逐行逐字对比（不拆子项，直接比较两个摘要的全文）
    contract_summary_text = _strip_contract_signing_page_text(contract_summary.get("content") or "")
    prospectus_summary_text = _strip_contract_signing_page_text(pros_summary.get("content") or "")
    c_rows = _collect_nonblank_compare_lines(contract_summary_text)
    p_rows = _collect_nonblank_compare_lines(prospectus_summary_text)
    c_lines = [row["line"] for row in c_rows]
    p_lines = [row["line"] for row in p_rows]
    c_nums = [row["lineno"] for row in c_rows]
    p_nums = [row["lineno"] for row in p_rows]

    # 严格按字面内容对齐，摘要之间默认应完全一致，仅做签署页裁剪，不做口径归一化。
    sm = difflib.SequenceMatcher(None, c_lines, p_lines, autojunk=False)
    ratio = sm.ratio()

    diff_lines = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            # 仅保留少量 equal 行作为上下文
            if i2 - i1 <= context_lines * 2:
                for k in range(i1, i2):
                    diff_lines.append({
                        "contract_ln": c_nums[k], "prospectus_ln": p_nums[j1 + (k - i1)],
                        "contract_html": html_mod.escape(c_lines[k]),
                        "prospectus_html": html_mod.escape(p_lines[j1 + (k - i1)]),
                        "tag": "equal",
                    })
            else:
                # 首尾各保留两行上下文，其余折叠
                for k in range(context_lines):
                    diff_lines.append({
                        "contract_ln": c_nums[i1 + k],
                        "prospectus_ln": p_nums[j1 + k],
                        "contract_html": html_mod.escape(c_lines[i1 + k]),
                        "prospectus_html": html_mod.escape(p_lines[j1 + k]),
                        "tag": "equal",
                    })
                diff_lines.append({
                    "tag": "skip", "count": i2 - i1 - context_lines * 2,
                    "contract_ln": None, "prospectus_ln": None,
                    "contract_html": "", "prospectus_html": "",
                })
                for k in range(context_lines):
                    diff_lines.append({
                        "contract_ln": c_nums[i2 - context_lines + k],
                        "prospectus_ln": p_nums[j2 - context_lines + k],
                        "contract_html": html_mod.escape(c_lines[i2 - context_lines + k]),
                        "prospectus_html": html_mod.escape(p_lines[j2 - context_lines + k]),
                        "tag": "equal",
                    })
        elif tag == "replace":
            for k in range(max(i2 - i1, j2 - j1)):
                al = c_lines[i1 + k] if i1 + k < i2 else ""
                bl = p_lines[j1 + k] if j1 + k < j2 else ""
                csm = difflib.SequenceMatcher(None, al, bl, autojunk=False)
                ah, bh = [], []
                for ct, ci1, ci2, cj1, cj2 in csm.get_opcodes():
                    ea = html_mod.escape(al[ci1:ci2])
                    eb = html_mod.escape(bl[cj1:cj2])
                    if ct == "equal":
                        ah.append(ea); bh.append(eb)
                    elif ct == "replace":
                        ah.append(f'<del>{ea}</del>'); bh.append(f'<ins>{eb}</ins>')
                    elif ct == "delete":
                        ah.append(f'<del>{ea}</del>')
                    elif ct == "insert":
                        bh.append(f'<ins>{eb}</ins>')
                diff_lines.append({
                    "contract_ln": c_nums[i1 + k] if i1 + k < i2 else None,
                    "prospectus_ln": p_nums[j1 + k] if j1 + k < j2 else None,
                    "contract_html": "".join(ah), "prospectus_html": "".join(bh),
                    "tag": "replace",
                })
        elif tag == "delete":
            for k in range(i1, i2):
                diff_lines.append({
                    "contract_ln": c_nums[k], "prospectus_ln": None,
                    "contract_html": '<del>' + html_mod.escape(c_lines[k]) + '</del>',
                    "prospectus_html": "", "tag": "delete",
                })
        elif tag == "insert":
            for k in range(j1, j2):
                diff_lines.append({
                    "contract_ln": None, "prospectus_ln": p_nums[k],
                    "contract_html": "",
                    "prospectus_html": '<ins>' + html_mod.escape(p_lines[k]) + '</ins>',
                    "tag": "insert",
                })

    changed = sum(1 for d in diff_lines if d["tag"] not in ("equal", "skip"))
    total = len(c_lines)
    similarity = round(ratio * 100, 1)
    excerpt_fields = _build_review_excerpt_pair(
        contract_summary_text,
        prospectus_summary_text,
        normalize_line=lambda line: line,
        left_prefix="contract",
        right_prefix="prospectus",
    )

    return jsonify({
        "contract_summary_heading": contract_summary["heading"],
        "prospectus_summary_heading": pros_summary["heading"],
        "contract_summary_locator": _normalize_review_text(contract_summary.get("heading")),
        "prospectus_summary_locator": _normalize_review_text(pros_summary.get("heading")),
        "contract_summary_text": contract_summary_text,
        "prospectus_summary_text": prospectus_summary_text,
        "similarity": similarity,
        "severity": _summary_cross_severity(similarity, changed),
        "is_problem": _summary_cross_result_is_problem(similarity, changed),
        "changed_lines": changed,
        "total_lines": total,
        "hunks": diff_lines,
        **excerpt_fields,
    })


@app.route("/api/review/report", methods=["POST"])
def api_review_report():
    """汇总所有检查结果，生成自包含HTML报告"""
    d = request.get_json(force=True) or {}
    data = _review_store.get("data", {})
    summary_cross_check = d.get("summary_cross_check")
    if not isinstance(summary_cross_check, dict):
        summary_cross_check = {}
    fund_name = (
        str(d.get("fund_name") or "").strip()
        or str(data.get("fund_name") or "").strip()
        or _fallback_review_filename(data.get("contract_filename", ""))
        or "基金"
    )

    cross_check = [dict(item) for item in d.get("cross_check", [])]
    for item in cross_check:
        item.setdefault("is_problem", _cross_check_result_is_problem(item))

    summary_check = [dict(item) for item in d.get("summary_check", [])]
    for item in summary_check:
        item.setdefault("is_problem", _summary_check_result_is_problem(item))

    summary_cross_check = dict(summary_cross_check or {})
    if summary_cross_check.get("similarity") is not None:
        changed_lines = summary_cross_check.get("changed_lines")
        summary_cross_check.setdefault("is_problem", _summary_cross_result_is_problem(summary_cross_check.get("similarity"), changed_lines))
        summary_cross_check.setdefault("severity", _summary_cross_severity(summary_cross_check.get("similarity"), changed_lines))

    format_issues = []
    for raw_issue in d.get("format_issues", []):
        issue = dict(raw_issue or {})
        if issue.get("locations") and not issue.get("context"):
            issue["context"] = issue.get("locations")
        if issue.get("locations") and not issue.get("line"):
            first_loc = (issue.get("locations") or [{}])[0]
            issue["line"] = first_loc.get("ln")
        format_issues.append(issue)

    findings = {
        "fund_name": fund_name,
        "contract_filename": data.get("contract_filename", ""),
        "prospectus_filename": data.get("prospectus_filename", ""),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "cross_check": cross_check,
        "summary_check": summary_check,
        "summary_cross_check": summary_cross_check,
        "format_issues": format_issues,
        "consistency_issues": d.get("consistency_issues", []),
        "ai_issues": d.get("ai_issues", []),
    }

    stats = {"error": 0, "warning": 0, "info": 0}
    for item in findings["cross_check"]:
        sev = item.get("severity", "info")
        stats[sev] = stats.get(sev, 0) + 1
    scc = findings.get("summary_cross_check", {})
    if scc.get("similarity") is not None:
        sim = scc["similarity"]
        sev = "info" if sim >= 95 else ("warning" if sim >= 80 else "error")
        stats[sev] = stats.get(sev, 0) + 1
    for item in findings["format_issues"]:
        sev = item.get("severity", "info")
        stats[sev] = stats.get(sev, 0) + 1
    for item in findings["consistency_issues"]:
        sev = item.get("severity", "info")
        stats[sev] = stats.get(sev, 0) + 1
    for item in findings["ai_issues"]:
        stats["info"] += 1
    findings["stats"] = stats

    cross_problem_count = sum(1 for item in cross_check if item.get("is_problem"))
    summary_problem_count = sum(1 for item in summary_check if item.get("is_problem"))
    has_summary_cross = bool(summary_cross_check.get("hunks"))
    summary_cross_problem_count = 1 if has_summary_cross and summary_cross_check.get("is_problem") else 0
    hidden_pass_count = (
        max(len(cross_check) - cross_problem_count, 0)
        + max(len(summary_check) - summary_problem_count, 0)
        + (1 if has_summary_cross and not summary_cross_check.get("is_problem") else 0)
    )
    findings["problem_summary"] = {
        "total_problems": cross_problem_count + summary_problem_count + summary_cross_problem_count
        + len(findings["format_issues"]) + len(findings["consistency_issues"]) + len(findings["ai_issues"]),
        "cross_problem_count": cross_problem_count,
        "summary_problem_count": summary_problem_count,
        "summary_cross_problem_count": summary_cross_problem_count,
        "format_problem_count": len(findings["format_issues"]),
        "consistency_problem_count": len(findings["consistency_issues"]),
        "ai_problem_count": len(findings["ai_issues"]),
        "hidden_pass_count": hidden_pass_count,
    }

    html = render_template("review_report.html", findings=findings)
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


# \u2500\u2500 \u542f\u52a8
def open_browser():
    time.sleep(1.2)
    webbrowser.open("http://127.0.0.1:5000")


if __name__ == "__main__":
    t = threading.Thread(target=open_browser, daemon=True)
    t.start()
    app.run(debug=False, port=5000, use_reloader=False)
