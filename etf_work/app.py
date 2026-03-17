"""
ETF基金合同知识库 Web 应用
Flask 后端服务器 + 合同生成引擎
"""

import io
import json
import os
import re
import threading
import time
import webbrowser
from datetime import datetime
from pathlib import Path

from flask import Flask, abort, jsonify, render_template, request, send_file

# ── 路径常量 ────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
TEMPLATE_MD = BASE_DIR / "01_通用模板.md"
SCHEMA_JSON = BASE_DIR / "02_变量定义表.json"
DIFF_TABLE_MD = BASE_DIR / "03_差异条款匹配表.md"
CLAUSES_JSON = BASE_DIR / "04_差异条款原文库.json"
ENTRY_TABLE_MD = BASE_DIR / "05_要素录入表.md"
PROSPECTUS_TEMPLATE_MD = BASE_DIR / "01_招募说明书模板.md"
PROSPECTUS_CLAUSES_JSON = BASE_DIR / "06_招募说明书差异条款库.json"
REFERENCE_PROSPECTUS_DOCX_MAP = {
    "SSE_CROSS": Path("C:/Users/12534/Desktop/\u57fa\u91d1\u5408\u540c/\u5357\u65b9\u4e2d\u8bc1\u5168\u6307\u7ea2\u5229\u8d28\u91cf\u4ea4\u6613\u578b\u5f00\u653e\u5f0f\u6307\u6570\u8bc1\u5238\u6295\u8d44\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66.docx"),
    "SSE_SINGLE": Path("C:/Users/12534/Desktop/\u57fa\u91d1\u5408\u540c/\u5357\u65b9\u4e0a\u8bc1\u79d1\u521b\u677f\u4eba\u5de5\u667a\u80fd\u4ea4\u6613\u578b\u5f00\u653e\u5f0f\u6307\u6570\u8bc1\u5238\u6295\u8d44\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66.docx"),
    "SSE_HK": Path("C:/Users/12534/Desktop/\u57fa\u91d1\u5408\u540c/\u5357\u65b9\u4e2d\u8bc1\u6e2f\u80a1\u901a\u4e92\u8054\u7f51\u4ea4\u6613\u578b\u5f00\u653e\u5f0f\u6307\u6570\u8bc1\u5238\u6295\u8d44\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66.docx"),
    "SZSE_CROSS": Path("C:/Users/12534/Desktop/\u57fa\u91d1\u5408\u540c/\u5357\u65b9\u4e2d\u8bc1\u7535\u6c60\u4e3b\u9898\u4ea4\u6613\u578b\u5f00\u653e\u5f0f\u6307\u6570\u8bc1\u5238\u6295\u8d44\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66.docx"),
    "SZSE_SINGLE": Path("C:/Users/12534/Desktop/\u57fa\u91d1\u5408\u540c/\u5357\u65b9\u521b\u4e1a\u677f\u4e2d\u76d8200\u4ea4\u6613\u578b\u5f00\u653e\u5f0f\u6307\u6570\u8bc1\u5238\u6295\u8d44\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66.docx"),
    "SZSE_HK": Path("C:/Users/12534/Desktop/\u57fa\u91d1\u5408\u540c/\u5357\u65b9\u4e2d\u8bc1\u6e2f\u80a1\u901a50\u4ea4\u6613\u578b\u5f00\u653e\u5f0f\u6307\u6570\u8bc1\u5238\u6295\u8d44\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66.docx"),
}

ALLOWED_SUFFIXES = {".md", ".json"}


def _clear_xml_children(node):
    for child in list(node):
        node.remove(child)


def _append_run_properties(run, OxmlElement, qn, ascii_font=None, eastasia_font=None, size=None, bold=False):
    rPr = OxmlElement("w:rPr")
    if ascii_font or eastasia_font:
        rFonts = OxmlElement("w:rFonts")
        if ascii_font:
            rFonts.set(qn("w:ascii"), ascii_font)
            rFonts.set(qn("w:hAnsi"), ascii_font)
        if eastasia_font:
            rFonts.set(qn("w:eastAsia"), eastasia_font)
        rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size))
        rPr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(size))
        rPr.append(sz_cs)
    run.append(rPr)


def _append_page_field(paragraph, OxmlElement, qn, ascii_font="Times New Roman", eastasia_font="Times New Roman", size=18):
    fld_begin = OxmlElement("w:r")
    _append_run_properties(fld_begin, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.append(fld_char_begin)

    instr_run = OxmlElement("w:r")
    _append_run_properties(instr_run, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = "PAGE   \\* MERGEFORMAT"
    instr_run.append(instr_text)

    fld_sep = OxmlElement("w:r")
    _append_run_properties(fld_sep, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_sep.append(fld_char_sep)

    fld_text = OxmlElement("w:r")
    _append_run_properties(fld_text, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_text_run = OxmlElement("w:t")
    fld_text_run.text = "1"
    fld_text.append(fld_text_run)

    fld_end = OxmlElement("w:r")
    _append_run_properties(fld_end, OxmlElement, qn, ascii_font=ascii_font, eastasia_font=eastasia_font, size=size)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    fld_end.append(fld_char_end)

    paragraph._p.append(fld_begin)
    paragraph._p.append(instr_run)
    paragraph._p.append(fld_sep)
    paragraph._p.append(fld_text)
    paragraph._p.append(fld_end)


def _set_section_page_numbers(section, OxmlElement, qn, start=None):
    try:
        section.footer.is_linked_to_previous = False
    except Exception:
        pass

    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    _clear_xml_children(paragraph._p)
    pPr = paragraph._p.get_or_add_pPr()
    _clear_xml_children(pPr)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    _append_page_field(paragraph, OxmlElement, qn, ascii_font="Times New Roman", eastasia_font="Times New Roman", size=18)

    sectPr = section._sectPr
    for node in list(sectPr.findall(qn("w:pgNumType"))):
        sectPr.remove(node)
    if start is not None:
        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:start"), str(start))
        sectPr.append(pg_num_type)


def _clear_header_footer_part(part):
    try:
        part.is_linked_to_previous = False
    except Exception:
        pass
    for paragraph in part.paragraphs:
        _clear_xml_children(paragraph._p)


def _set_cover_section_no_page_number(section):
    section.different_first_page_header_footer = True
    for part in (section.header, section.first_page_header, section.footer, section.first_page_footer):
        _clear_header_footer_part(part)


def _finalize_doc_page_numbers(doc, OxmlElement, qn, body_start_index=1):
    sections = list(doc.sections)
    if not sections:
        return

    body_start_index = max(1, min(body_start_index, len(sections)))
    for section in sections[:body_start_index]:
        _set_cover_section_no_page_number(section)

    body_sections = sections[body_start_index:] or sections[-1:]
    for idx, section in enumerate(body_sections):
        _set_section_page_numbers(section, OxmlElement, qn, start=1 if idx == 0 else None)


app = Flask(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
#  ContractEngine — 6步处理管线
# ═══════════════════════════════════════════════════════════════════════════════
class ContractEngine:
    def __init__(self):
        with open(CLAUSES_JSON, encoding="utf-8") as f:
            self.clauses = json.load(f)["clauses"]

    @staticmethod
    def _cn_numeral_to_int(cn: str):
        """
        Convert common Chinese numerals to int.
        Supports values used in headings like: 一, 十, 十一, 二十六, 一百零二.
        Returns None when conversion fails.
        """
        if not cn:
            return None
        if cn.isdigit():
            return int(cn)

        digits = {
            "零": 0, "一": 1, "二": 2, "三": 3, "四": 4,
            "五": 5, "六": 6, "七": 7, "八": 8, "九": 9,
        }
        units = {"十": 10, "百": 100, "千": 1000}

        total = 0
        current = 0
        for ch in cn:
            if ch in digits:
                current = digits[ch]
                continue
            if ch in units:
                if current == 0:
                    current = 1
                total += current * units[ch]
                current = 0
                continue
            return None
        total += current
        return total if total > 0 else None

    # ── Step 1: 派生变量 ─────────────────────────────────────────────────────
    def _derive_variables(self, v: dict) -> dict:
        v = dict(v)  # 浅拷贝，不改原始输入

        # 交易所 → 中文名 / 业务规则类型
        exchange = v.get("EXCHANGE", "")
        if exchange == "SZSE":
            v.setdefault("EXCHANGE_NAME_CN", "深圳证券交易所")
            v.setdefault("BUSINESS_RULES_TYPE", "SZSE")
        elif exchange == "SSE":
            v.setdefault("EXCHANGE_NAME_CN", "上海证券交易所")
            v.setdefault("BUSINESS_RULES_TYPE", "SSE")

        # 市场类型 → HAS_HK_CONNECT / 跟踪误差默认值 / 市场标志
        market = v.get("MARKET_TYPE", "")
        if market == "HK_CONNECT":
            v["HAS_HK_CONNECT"] = True
            v.setdefault("TRACKING_ERROR_DAILY", 0.35)
            v.setdefault("TRACKING_ERROR_ANNUAL", 4)
        else:
            v.setdefault("HAS_HK_CONNECT", False)
            v.setdefault("TRACKING_ERROR_DAILY", 0.2)
            v.setdefault("TRACKING_ERROR_ANNUAL", 2)
        v["IS_CHUANGYE"] = (market == "CHUANGYE")
        v["IS_KECHUANG"] = (market == "KECHUANG")

        # 中文布尔 → Python bool
        for key in ("HAS_HK_CONNECT", "HAS_AML", "HAS_CUSTODIAN_TRANSFER_SPECIAL",
                    "HAS_STOCK_SUBSCRIPTION", "HAS_CDR", "CUSTODIAN_HAS_OFFICE_ADDRESS"):
            if key in v:
                raw = v[key]
                if isinstance(raw, str):
                    v[key] = raw.lower() in ("true", "1", "yes", "是")

        # 托管人有无办公地址
        custodian = v.get("CUSTODIAN_NAME", "")
        custodians_with_office = {"交通银行股份有限公司", "中信证券股份有限公司"}
        v.setdefault(
            "CUSTODIAN_HAS_OFFICE_ADDRESS",
            custodian in custodians_with_office,
        )

        return v

    # ── Step 2: 注入差异条款原文 ─────────────────────────────────────────────
    def _inject_clause_texts(self, v: dict) -> dict:
        v = dict(v)

        # WORKING_DAY_DEF
        wdt = v.get("WORKING_DAY_TYPE", "SZSE_ONLY")
        wday_variants = self.clauses["WORKING_DAY_DEF"]["variants"]
        v["WORKING_DAY_DEF"] = wday_variants.get(wdt, wday_variants["SZSE_ONLY"])["text"]

        # BUSINESS_RULES_DEF
        brt = v.get("BUSINESS_RULES_TYPE", "SZSE")
        br_variants = self.clauses["BUSINESS_RULES_DEF"]["variants"]
        v["BUSINESS_RULES_DEF"] = br_variants.get(brt, br_variants["SZSE"])["text"]

        # NON_COMPONENT_SCOPE
        mt = v.get("MARKET_TYPE", "CHUANGYE")
        nc_variants = self.clauses["NON_COMPONENT_SCOPE"]["variants"]
        # Map MARKET_TYPE value to variant key
        mt_to_key = {
            "CHUANGYE": "CHUANGYE",
            "KECHUANG": "KECHUANG",
            "A_SHARE": "A_SHARE",
            "HK_CONNECT": "HK_CONNECT",
        }
        nc_key = mt_to_key.get(mt, "CHUANGYE")
        v["NON_COMPONENT_SCOPE"] = nc_variants.get(nc_key, nc_variants["CHUANGYE"])["text"]

        # NON_COMPONENT_SCOPE_INTRO
        if mt == "HK_CONNECT":
            v["NON_COMPONENT_SCOPE_INTRO"] = (
                "包括内地与香港股票市场交易互联互通机制允许买卖的规定范围内的"
                "香港联合交易所上市的股票（简称\u201c港股通股票\u201d）、存托凭证，下同"
            )
        else:
            v["NON_COMPONENT_SCOPE_INTRO"] = "含存托凭证"

        # DISTRIBUTION_FREQ_CLAUSE
        df = v.get("DISTRIBUTION_FREQ", "MONTHLY")
        df_variants = self.clauses["DISTRIBUTION_FREQ_CLAUSE"]["variants"]
        v["DISTRIBUTION_FREQ_CLAUSE"] = df_variants.get(df, df_variants["MONTHLY"])["text"]

        # MGMT_FEE_PAYMENT_METHOD text (replaces the enum value with actual text)
        mfpm = v.get("MGMT_FEE_PAYMENT_METHOD", "CONSULT")
        mgmt_variants = self.clauses["MGMT_FEE_PAYMENT"]["variants"]
        # Store original enum for custody lookup
        v["_MGMT_FEE_PAYMENT_ENUM"] = mfpm
        v["MGMT_FEE_PAYMENT_METHOD"] = mgmt_variants.get(mfpm, mgmt_variants["CONSULT"])["text"]

        # CUSTODY_FEE_PAYMENT_METHOD text
        cfpm = v.get("CUSTODY_FEE_PAYMENT_METHOD", v.get("_MGMT_FEE_PAYMENT_ENUM", "CONSULT"))
        cust_variants = self.clauses["CUSTODY_FEE_PAYMENT"]["variants"]
        v["CUSTODY_FEE_PAYMENT_METHOD"] = cust_variants.get(cfpm, cust_variants["CONSULT"])["text"]

        # FUND_PROFIT_DEF (for HK_CONNECT)
        fp_variants = self.clauses["FUND_PROFIT_DEF"]["variants"]
        if v.get("HAS_HK_CONNECT"):
            v["FUND_PROFIT_DEF"] = fp_variants["HK_CONNECT"]["text"]
        else:
            v["FUND_PROFIT_DEF"] = fp_variants["STANDARD"]["text"]

        # SUBSCRIPTION_METHOD_TEXT
        sub_variants = self.clauses["SUBSCRIPTION_METHOD"]["variants"]
        if v.get("HAS_STOCK_SUBSCRIPTION"):
            sub_text = sub_variants["THREE_TYPES"]["text"]
        else:
            sub_text = sub_variants["TWO_TYPES"]["text"]
        # Replace {EXCHANGE_NAME_CN} inside subscription text
        sub_text = sub_text.replace("{EXCHANGE_NAME_CN}", v.get("EXCHANGE_NAME_CN", "深圳证券交易所"))
        v["SUBSCRIPTION_METHOD_TEXT"] = sub_text

        return v

    # ── Step 4: 处理条件块（支持嵌套） ──────────────────────────────────────
    def _process_conditionals(self, text: str, v: dict) -> str:
        # We process from innermost outward using a stack-based approach
        max_passes = 10
        for _ in range(max_passes):
            new_text = self._single_pass_conditionals(text, v)
            if new_text == text:
                break
            text = new_text
        return text

    def _single_pass_conditionals(self, text: str, v: dict) -> str:
        # Match innermost {{IF ...}}...{{ENDIF}} (no nested IF inside)
        pattern_if = re.compile(
            r'\{\{IF\s+(\w+)\}\}((?:(?!\{\{IF)[\s\S])*?)\{\{ENDIF\}\}',
            re.DOTALL
        )
        pattern_if_not = re.compile(
            r'\{\{IF_NOT\s+(\w+)\}\}((?:(?!\{\{IF)[\s\S])*?)\{\{ENDIF\}\}',
            re.DOTALL
        )

        def replace_if(m):
            var_name = m.group(1)
            content = m.group(2)
            val = v.get(var_name, False)
            if isinstance(val, str):
                val = val.lower() in ("true", "1", "yes", "是")
            return content if val else ""

        def replace_if_not(m):
            var_name = m.group(1)
            content = m.group(2)
            val = v.get(var_name, False)
            if isinstance(val, str):
                val = val.lower() in ("true", "1", "yes", "是")
            return "" if val else content

        text = pattern_if.sub(replace_if, text)
        text = pattern_if_not.sub(replace_if_not, text)
        return text

    # ── Step 5: 替换占位符 ───────────────────────────────────────────────────
    def _replace_placeholders(self, text: str, v: dict) -> str:
        def replacer(m):
            key = m.group(1)
            val = v.get(key)
            if val is None:
                return m.group(0)  # 保留未定义的占位符
            if isinstance(val, bool):
                return "是" if val else "否"
            return str(val)

        return re.sub(r"\{([A-Z_][A-Z0-9_]*)\}", replacer, text)

    # ── Step 6: 清理 ─────────────────────────────────────────────────────────
    @staticmethod
    def _is_prospectus_toc_placeholder_line(line: str) -> bool:
        return bool(re.search(r'[\u3010\[](?:\u5f85\u586b\u5199|\u5f85\u8865\u5145)', line or ""))

    @staticmethod
    def _parse_prospectus_chapter_heading(line: str):
        match = re.match("^\u7b2c([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e]+)\u7ae0\\s+(.+?)(?:[\\t ]+\\d+)?$", (line or "").strip())
        if not match:
            return None
        return match.group(1), match.group(2).strip()

    @staticmethod
    def _format_prospectus_reference_heading(chapter_cn: str, title: str) -> str:
        return f"{chapter_cn}\u3001{(title or '').strip()}"

    def _format_reference_style_prospectus(self, text: str) -> str:
        lines = text.split("\n")
        formatted = []
        phase = "cover"
        toc_headings = []

        for raw in lines:
            stripped = raw.strip()
            if not stripped:
                if phase == "body" and formatted and formatted[-1] != "":
                    formatted.append("")
                continue

            if phase == "cover":
                if re.match(r"^\u76ee\s*\u5f55$", stripped):
                    formatted.append("\u76ee\u5f55")
                    phase = "toc"
                    continue
                formatted.append(stripped)
                continue

            parsed = self._parse_prospectus_chapter_heading(stripped)
            if phase == "toc":
                if self._is_prospectus_toc_placeholder_line(stripped):
                    continue
                if parsed:
                    heading = self._format_prospectus_reference_heading(*parsed)
                    if toc_headings and heading == toc_headings[0]:
                        phase = "body"
                        formatted.append(heading)
                    elif heading not in toc_headings:
                        toc_headings.append(heading)
                        formatted.append(heading)
                    continue
                if re.match(r"^[???????????]+?", stripped):
                    continue
                continue

            if parsed:
                formatted.append(self._format_prospectus_reference_heading(*parsed))
            else:
                formatted.append(stripped)

        text = "\n".join(formatted)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _inject_important_notice_before_toc(self, text: str, v: dict) -> str:
        ref = self._load_reference_fixed_content(v)
        notice = (ref.get("important_notice") or "").strip()
        if not notice or "重要提示" in text:
            return text

        lines = text.splitlines()
        toc_idx = next((idx for idx, line in enumerate(lines) if line.strip() == "目录"), None)
        if toc_idx is None:
            return text

        cover = "\n".join(lines[:toc_idx]).strip()
        tail = "\n".join(lines[toc_idx:]).strip()
        return "\n\n".join(part for part in (cover, notice, tail) if part).strip()

    def _cleanup(self, text: str) -> str:
        # Remove markdown blockquote header lines (> ...)
        lines = text.split("\n")
        clean = []
        skip_header = True
        for line in lines:
            # Skip the preamble before the actual contract starts
            if skip_header:
                if line.strip().startswith("{FUND_NAME}") or (
                    "基金合同" in line and "模板说明" not in line and not line.startswith(">")
                    and not line.startswith("#") and not line.startswith("**") and not line.startswith("---")
                ):
                    skip_header = False
                    clean.append(line)
                # Skip preamble lines
                continue
            # Remove comment/annotation lines
            if line.strip().startswith(">") or line.strip().startswith("**条件变量") or line.strip().startswith("**差异条款"):
                continue
            clean.append(line)

        text = "\n".join(clean)

        # Collapse 3+ consecutive blank lines into 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        return text

    # ── 主方法：生成合同 ─────────────────────────────────────────────────────
    def generate(self, form_data: dict) -> str:
        # Step 1
        v = self._derive_variables(form_data)
        # Step 2
        v = self._inject_clause_texts(v)
        # Step 3: read template
        template_text = TEMPLATE_MD.read_text(encoding="utf-8")
        # Step 4
        text = self._process_conditionals(template_text, v)
        # Step 5
        text = self._replace_placeholders(text, v)
        # Step 5b: 修复条件删除导致的序号跳跃
        text = self._renumber_sequences(text)
        # Step 6
        text = self._cleanup(text)
        return text

    # ── Step 5b: 重排阿拉伯序号（修复条件块删除项目后的跳跃）──────────────
    def _renumber_sequences(self, text: str) -> str:
        """
        修复条件块删除项目后阿拉伯序号列表的跳跃。
        规则：
          - 只处理行首 `数字、` 格式
          - num == 1：新列表开始，重置计数器，不修改
          - num > expected_next（且 num > 1）：检测到跳跃，将该行及后续连续项重排
          - 遇到中文序号标题（一、二、…）或章节标题（第X部分）时重置计数器
        """
        lines = text.split("\n")
        RE_NUM = re.compile(r"^(\d+)(、)")
        RE_RESET = re.compile(
            r"^(?:[一二三四五六七八九十百]+、|第[一二三四五六七八九十百]+部分)"
        )
        last_num = None
        result = []
        for line in lines:
            if RE_RESET.match(line.strip()):
                last_num = None
                result.append(line)
                continue
            m = RE_NUM.match(line)
            if m:
                num = int(m.group(1))
                if num == 1:
                    last_num = 1
                elif last_num is not None and num > last_num + 1:
                    expected = last_num + 1
                    line = re.sub(r"^\d+、", f"{expected}、", line, count=1)
                    last_num = expected
                else:
                    last_num = num
            result.append(line)
        return "\n".join(result)

    # ── 中文标点规范化 ───────────────────────────────────────────────────────
    @staticmethod
    def _to_chinese_punct(text: str) -> str:
        """将文本中的半角标点转为全角中文标点（保留数字中的逗号与小数点）。"""
        text = text.replace('(', '（').replace(')', '）')
        text = text.replace('[', '【').replace(']', '】')
        text = text.replace(';', '；')
        # 逗号仅在两侧均非数字时替换，保留千位分隔符如"1,000"
        text = re.sub(r'(?<!\d),(?!\d)', '，', text)
        return text

    # ── Word (.docx) 导出 ────────────────────────────────────────────────────
    def build_docx(self, contract_text: str) -> bytes:
        """
        将合同纯文本转换为格式化 Word 文档，与参考合同格式完全对齐。
        - 封面：独立 section + vAlign=center，彻底消除内容溢出第二页的问题
        - 签署页：左对齐 + 精确空行间距（来自参考文档XML）
        - 章节标题：Times New Roman 15pt 加粗，pageBreakBefore
        - 正文：宋体 12pt，两端对齐，首行缩进2字，1.5倍行距
        - 数字/英文：Times New Roman；中文：宋体
        """
        # 标准化中文标点
        contract_text = self._to_chinese_punct(contract_text)
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Twips

        # ── 1. 文档与页面设置 ──────────────────────────────────────────────
        doc = Document()
        sec = doc.sections[0]
        sec.page_width        = Twips(11906)
        sec.page_height       = Twips(16838)
        sec.top_margin        = Twips(1440)
        sec.bottom_margin     = Twips(1440)
        sec.left_margin       = Twips(1800)
        sec.right_margin      = Twips(1800)
        sec.header_distance   = Twips(851)
        sec.footer_distance   = Twips(992)
        sectPr = sec._sectPr
        docGrid = OxmlElement("w:docGrid")
        docGrid.set(qn("w:type"), "lines")
        docGrid.set(qn("w:linePitch"), "312")
        sectPr.append(docGrid)

        # ── 2. Normal 样式：中文宋体/英数 Times New Roman 12pt，两端对齐，孤行控制关闭 ──
        normal_style = doc.styles["Normal"]
        normal_rpr = normal_style.element.get_or_add_rPr()
        for old in normal_rpr.findall(qn("w:rFonts")):
            normal_rpr.remove(old)
        rFonts_n = OxmlElement("w:rFonts")
        rFonts_n.set(qn("w:ascii"),    "Times New Roman")
        rFonts_n.set(qn("w:hAnsi"),    "Times New Roman")
        rFonts_n.set(qn("w:eastAsia"), "宋体")
        normal_rpr.insert(0, rFonts_n)
        sz_n = OxmlElement("w:sz"); sz_n.set(qn("w:val"), "24"); normal_rpr.append(sz_n)
        normal_ppr = normal_style.element.get_or_add_pPr()
        wc = OxmlElement("w:widowControl"); wc.set(qn("w:val"), "0"); normal_ppr.insert(0, wc)
        jc_n = OxmlElement("w:jc"); jc_n.set(qn("w:val"), "both"); normal_ppr.append(jc_n)
        sp_n = OxmlElement("w:spacing")
        sp_n.set(qn("w:line"), "360"); sp_n.set(qn("w:lineRule"), "auto")
        sp_n.set(qn("w:before"), "0"); sp_n.set(qn("w:after"), "0")
        normal_ppr.append(sp_n)

        # ── 3. XML 辅助函数 ───────────────────────────────────────────────
        def _set_para(p, jc=None, line=None, line_rule="auto",
                      before=None, after=None,
                      first_line=None, first_line_chars=None, left_ind=None,
                      keep_lines=False, page_break_before=False,
                      snap_to_grid=None):
            pPr = p._p.get_or_add_pPr()
            if keep_lines:
                pPr.append(OxmlElement("w:keepLines"))
            if page_break_before:
                pPr.append(OxmlElement("w:pageBreakBefore"))
            if snap_to_grid is not None:
                sg = OxmlElement("w:snapToGrid")
                sg.set(qn("w:val"), "1" if snap_to_grid else "0")
                pPr.append(sg)
            if before is not None or after is not None or line is not None:
                sp = OxmlElement("w:spacing")
                if before is not None: sp.set(qn("w:before"), str(before))
                if after  is not None: sp.set(qn("w:after"),  str(after))
                if line   is not None:
                    sp.set(qn("w:line"),     str(line))
                    sp.set(qn("w:lineRule"), line_rule)
                pPr.append(sp)
            if first_line is not None or first_line_chars is not None or left_ind is not None:
                ind = OxmlElement("w:ind")
                if left_ind          is not None: ind.set(qn("w:left"),           str(left_ind))
                if first_line_chars  is not None: ind.set(qn("w:firstLineChars"), str(first_line_chars))
                if first_line        is not None: ind.set(qn("w:firstLine"),      str(first_line))
                pPr.append(ind)
            if jc is not None:
                jc_el = OxmlElement("w:jc"); jc_el.set(qn("w:val"), jc); pPr.append(jc_el)

        def _set_run(r, ascii_font=None, eastasia_font=None, hint=None,
                     sz=None, sz_cs=None, bold=False, bcs=False, color=None):
            rPr = r._r.get_or_add_rPr()
            if ascii_font or eastasia_font or hint:
                rF = OxmlElement("w:rFonts")
                if ascii_font:    rF.set(qn("w:ascii"), ascii_font); rF.set(qn("w:hAnsi"), ascii_font)
                if eastasia_font: rF.set(qn("w:eastAsia"), eastasia_font)
                if hint:          rF.set(qn("w:hint"), hint)
                rPr.insert(0, rF)
            if bold: rPr.append(OxmlElement("w:b"))
            if bcs:  rPr.append(OxmlElement("w:bCs"))
            if color:
                col = OxmlElement("w:color"); col.set(qn("w:val"), color); rPr.append(col)
            if sz is not None:
                s = OxmlElement("w:sz"); s.set(qn("w:val"), str(sz)); rPr.append(s)
            if sz_cs is not None:
                sc = OxmlElement("w:szCs"); sc.set(qn("w:val"), str(sz_cs)); rPr.append(sc)

        def _signing_empty(n=1):
            """签署页专用空行：左对齐，snapToGrid=0，宋体，bCs，与参考文档一致。"""
            for _ in range(n):
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                sg = OxmlElement("w:snapToGrid"); sg.set(qn("w:val"), "0"); pPr.append(sg)
                sp = OxmlElement("w:spacing")
                sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
                sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
                pPr.append(sp)
                jc_el = OxmlElement("w:jc"); jc_el.set(qn("w:val"), "left"); pPr.append(jc_el)
                mrPr = OxmlElement("w:rPr")
                rF2 = OxmlElement("w:rFonts")
                rF2.set(qn("w:ascii"), "Times New Roman"); rF2.set(qn("w:hAnsi"), "Times New Roman")
                mrPr.append(rF2)
                mrPr.append(OxmlElement("w:bCs"))
                szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "21"); mrPr.append(szCs)
                pPr.append(mrPr)

        def _cover_section_break():
            """
            在封面最后一段的 pPr 中嵌入 sectPr（封面 section 定义）。
            使用 vAlign=center 确保内容永远不溢出，并通过 nextPage 分隔目录。
            """
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            # 封面 section 的属性
            cover_sp = OxmlElement("w:sectPr")
            # 分节类型：下一页
            t = OxmlElement("w:type"); t.set(qn("w:val"), "nextPage"); cover_sp.append(t)
            # 页面尺寸（A4）
            pgSz = OxmlElement("w:pgSz")
            pgSz.set(qn("w:w"), "11906"); pgSz.set(qn("w:h"), "16838")
            cover_sp.append(pgSz)
            # 页边距（与正文一致）
            pgMar = OxmlElement("w:pgMar")
            pgMar.set(qn("w:top"),    "1440"); pgMar.set(qn("w:right"),  "1800")
            pgMar.set(qn("w:bottom"), "1440"); pgMar.set(qn("w:left"),   "1800")
            pgMar.set(qn("w:header"), "851");  pgMar.set(qn("w:footer"), "992")
            pgMar.set(qn("w:gutter"), "0")
            cover_sp.append(pgMar)
            # 文档网格
            dg = OxmlElement("w:docGrid")
            dg.set(qn("w:type"), "lines"); dg.set(qn("w:linePitch"), "312")
            cover_sp.append(dg)
            # 垂直居中——确保内容不管多少行都不溢出第二页
            vAlign = OxmlElement("w:vAlign"); vAlign.set(qn("w:val"), "center")
            cover_sp.append(vAlign)
            pPr.append(cover_sp)

        # ── 4. 正则分类 ──────────────────────────────────────────────────
        RE_TOC_ENTRY = re.compile(r"^(第[一二三四五六七八九十百]+部分\s+.+?)[\t ]+(\d+)$")
        RE_PART_HEAD = re.compile(r"^第[一二三四五六七八九十百]+部分\s+\S")
        RE_UPPER_NUM = re.compile(r"^[一二三四五六七八九十百]+、")
        RE_PAREN_CN  = re.compile(r"^（[一二三四五六七八九十百]+）")
        RE_ARAB_DOT  = re.compile(r"^\d+、")
        RE_PAREN_NUM = re.compile(r"^（\d+）")
        RE_MD_TABLE_ROW = re.compile(r"^\|.*\|\s*$")
        RE_MD_TABLE_DIVIDER = re.compile(r"^\|\s*:?-{3,}:?(?:\|\s*:?-{3,}:?)+\|\s*$")

        # ── 书签计数器与锚点辅助 ────────────────────────────────────────────────
        _bm_id = [0]  # 可变单元格，允许在内层代码中递增

        def _part_anchor(heading_s: str) -> str:
            """从"第X部分..."中提取 ASCII 书签锚名，例如 'part_26'。"""
            m = re.match(r'^第([一二三四五六七八九十百]+)部分', heading_s)
            if not m:
                return ""
            part_no = self._cn_numeral_to_int(m.group(1))
            return f"part_{part_no}" if part_no is not None else ""

        # ── 5. 主处理循环 ─────────────────────────────────────────────────
        lines = contract_text.split("\n")
        phase          = "cover"
        cover_idx      = 0      # 非空封面行计数
        signing_started = False
        signing_idx    = 0      # 签署页内容行计数
        current_part   = 0      # 当前所在部分编号（用于判断前言是否加空行）

        for raw in lines:
            s = raw.strip()

            # 所有空行跳过（签署页内的空行由代码主动添加）
            if not s:
                continue

            # ════ 封面 → 目录 ════
            if phase == "cover" and re.match(r"^目\s*录$", s):
                phase = "toc"
                # 封面 section break（vAlign=center, nextPage）
                _cover_section_break()
                p = doc.add_paragraph()
                _set_para(p, jc="center", line=360, before=0, after=0)
                r = p.add_run("目    录")
                _set_run(r, sz=28, sz_cs=28, bcs=True)
                continue

            # ════ 封面内容：仅4行，vAlign=center 负责垂直定位 ════
            if phase == "cover":
                p = doc.add_paragraph()
                if cover_idx == 0:
                    # 合同标题：sz=48(24pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=48, sz_cs=48, bold=True, bcs=True)
                    # 标题与管理人之间加3个空行，制造视觉距离
                    for _ in range(3):
                        ep = doc.add_paragraph()
                        _set_para(ep, jc="center", line=360, before=0, after=0)
                elif s.startswith("基金管理人") or s.startswith("基金托管人"):
                    # 管理人/托管人：sz=36(18pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=36, sz_cs=36, bold=True, bcs=True)
                else:
                    # 日期：sz=36(18pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=36, sz_cs=36, bold=True, bcs=True)
                cover_idx += 1
                continue

            # ════ 目录条目 ════
            if phase == "toc":
                m_toc = RE_TOC_ENTRY.match(s)
                if m_toc:
                    toc_text = m_toc.group(1).strip()
                    toc_page = m_toc.group(2).strip()
                    p = doc.add_paragraph()
                    _set_para(p, line=360, before=0, after=0)
                    pPr = p._p.get_or_add_pPr()
                    tabs_el = OxmlElement("w:tabs")
                    tab_el  = OxmlElement("w:tab")
                    tab_el.set(qn("w:val"),    "right")
                    tab_el.set(qn("w:leader"), "dot")
                    tab_el.set(qn("w:pos"),    "8296")
                    tabs_el.append(tab_el)
                    pPr.append(tabs_el)
                    anchor = _part_anchor(toc_text)
                    if anchor:
                        hlink = OxmlElement("w:hyperlink")
                        hlink.set(qn("w:anchor"), anchor)
                        hlink.set(qn("w:history"), "1")
                        for txt_val in ([toc_text] + (["\t" + toc_page] if toc_page else [])):
                            rel = OxmlElement("w:r")
                            rPr_h = OxmlElement("w:rPr")
                            # Explicit hyperlink style improves compatibility (e.g. WPS).
                            rStyle = OxmlElement("w:rStyle")
                            rStyle.set(qn("w:val"), "Hyperlink")
                            rPr_h.append(rStyle)
                            rF_h = OxmlElement("w:rFonts")
                            rF_h.set(qn("w:ascii"), "Times New Roman")
                            rF_h.set(qn("w:hAnsi"), "Times New Roman")
                            rF_h.set(qn("w:eastAsia"), "宋体")
                            rPr_h.append(rF_h)
                            sz_h = OxmlElement("w:sz")
                            sz_h.set(qn("w:val"), "24")
                            rPr_h.append(sz_h)
                            rel.append(rPr_h)
                            t_h = OxmlElement("w:t")
                            t_h.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                            t_h.text = txt_val
                            rel.append(t_h)
                            hlink.append(rel)
                        p._p.append(hlink)
                    else:
                        r1 = p.add_run(toc_text)
                        _set_run(r1, ascii_font="Times New Roman", eastasia_font="宋体", sz=24)
                        if toc_page:
                            r2 = p.add_run("\t" + toc_page)
                            _set_run(r2, ascii_font="Times New Roman", eastasia_font="宋体", sz=24)
                    continue
                else:
                    phase = "body"

            # ════ 签署页检测 ════
            if phase == "body" and not signing_started and "签署页" in s and "无正文" in s:
                signing_started = True
                p = doc.add_paragraph()
                _set_para(p, snap_to_grid=False, jc="left", line=360, before=0, after=0,
                          page_break_before=True)   # 签署页强制分页
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                _signing_empty(2)   # 签署页标题后 2 个空行（参考文档）
                signing_idx = 1
                continue

            # ════ 签署页内容行 ════
            if phase == "body" and signing_started:
                p = doc.add_paragraph()
                is_seal = "（盖章）" in s
                _set_para(p, snap_to_grid=False, jc="left", line=360, before=0, after=0)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                # 每一内容行后补充空行（对应参考文档签名留白）
                if "（盖章）" in s:
                    _signing_empty(4)   # 盖章行后 4 个空行（管理人/托管人签名区）
                elif "（签字或盖章）：" in s or "（签名）" in s:
                    _signing_empty(4)   # 法定代表人签名行后 4 个空行
                # 签订地点/日期行后不加空行
                signing_idx += 1
                continue

            # ════ 部分标题（第X部分…） ════
            if phase == "body" and RE_PART_HEAD.match(s):
                current_part += 1
                p = doc.add_paragraph(style="Heading 1")
                _set_para(p, jc="center", line=360, before=0, after=0,
                          keep_lines=True, page_break_before=True)
                # 添加书签，供目录超链接跳转
                anchor = _part_anchor(s)
                if anchor:
                    _bm_id[0] += 1
                    bm_start = OxmlElement("w:bookmarkStart")
                    bm_start.set(qn("w:id"), str(_bm_id[0]))
                    bm_start.set(qn("w:name"), anchor)
                    bm_end = OxmlElement("w:bookmarkEnd")
                    bm_end.set(qn("w:id"), str(_bm_id[0]))
                    p._p.append(bm_start)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         sz=30, sz_cs=30, bold=True, bcs=False, color="auto")
                if anchor:
                    p._p.append(bm_end)
                continue

            # ════ 一级子标题（一、二、…） ════
            if phase == "body" and RE_UPPER_NUM.match(s):
                # 第一部分（前言）子标题间不加空行，其余部分各子标题前空一行
                if current_part != 1:
                    ep = doc.add_paragraph()
                    _set_para(ep, line=360, before=0, after=0)
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480, first_line_chars=200)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 二级子标题（（一）（二）…） ════
            if phase == "body" and RE_PAREN_CN.match(s):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480, first_line_chars=200)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 数字条款（1、2、… 或（1）（2）…） ════
            if phase == "body" and (RE_ARAB_DOT.match(s) or RE_PAREN_NUM.match(s)):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480, first_line_chars=200)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 普通正文 ════
            if phase in ("body", "toc"):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480, first_line_chars=200)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)

        # ── 6. 序列化 ──────────────────────────────────────────────────────
        _finalize_doc_page_numbers(doc, OxmlElement, qn)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

# ═══════════════════════════════════════════════════════════════════════════════
#  ProspectusEngine — 招募说明书 8步处理管线
# ═══════════════════════════════════════════════════════════════════════════════
class ProspectusEngine:
    def __init__(self):
        with open(PROSPECTUS_CLAUSES_JSON, encoding="utf-8") as f:
            self.pro_clauses = json.load(f)["clauses"]
        self._reference_fixed_cache = {}

    # ── Step 1: 派生变量（委托给 ContractEngine）────────────────────────────
    def _derive_variables(self, v: dict) -> dict:
        v = engine._derive_variables(v)
        market_scope = str(v.get("MARKET_SCOPE", "") or "").strip().upper()
        if market_scope not in {"SINGLE_MARKET", "CROSS_MARKET"}:
            market_type = str(v.get("MARKET_TYPE", "") or "").strip().upper()
            if market_type in {"KECHUANG", "CHUANGYE"}:
                market_scope = "SINGLE_MARKET"
            else:
                market_scope = "CROSS_MARKET"
        v["MARKET_SCOPE"] = market_scope
        product_type = str(v.get("PRODUCT_TYPE", "") or "").strip().upper()
        if not product_type:
            product_type = "ETF"
        v["PRODUCT_TYPE"] = product_type
        return v

    def _get_prospectus_variant_key(self, v: dict) -> str:
        v = self._derive_variables(v)
        exchange = v.get("EXCHANGE", "")
        has_hk = bool(v.get("HAS_HK_CONNECT"))
        market_scope = v.get("MARKET_SCOPE", "CROSS_MARKET")

        if has_hk:
            return "SSE_HK" if exchange == "SSE" else "SZSE_HK"
        if exchange == "SSE":
            return "SSE_SINGLE" if market_scope == "SINGLE_MARKET" else "SSE_CROSS"
        return "SZSE_SINGLE" if market_scope == "SINGLE_MARKET" else "SZSE_CROSS"

    def _get_variant_clause_bundle(self, v: dict) -> dict:
        variants = self.pro_clauses.get("PROSPECTUS_VARIANTS", {}).get("variants", {})
        key = self._get_prospectus_variant_key(v)
        return variants.get(key, variants.get("SSE_CROSS", {}))


    @staticmethod
    def _get_product_type(v: dict) -> str:
        product_type = str((v or {}).get("PRODUCT_TYPE", "") or "").strip().upper()
        return product_type or "ETF"

    @staticmethod
    @staticmethod
    def _strip_signing_page_from_contract_summary(summary_text: str) -> str:
        text = (summary_text or "").strip()
        if not text:
            return ""
        markers = [
            "签署页",
            "本页无正文",
            "（盖章）",
            "（签字或盖章）",
            "（签名）",
        ]
        positions = [pos for pos in (text.find(marker) for marker in markers) if pos >= 0]
        signing_patterns = [
            r"（?本页为《[^》]+基金合同》签署页[^）]*）?",
            r"本页为《[^》]+基金合同》",
        ]
        for pattern in signing_patterns:
            m = re.search(pattern, text)
            if m:
                positions.append(m.start())
        if positions:
            text = text[:min(positions)].rstrip()
        return text

    # ── Step 2: 注入差异条款（合同条款 + 招募说明书专有条款）──────────────
    def _inject_clause_texts(self, v: dict) -> dict:
        v = self._derive_variables(v)
        v = engine._inject_clause_texts(v)
        has_hk = v.get("HAS_HK_CONNECT", False)
        exch_cn = v.get("EXCHANGE_NAME_CN", "证券交易所")
        variant_key = self._get_prospectus_variant_key(v)
        variant_bundle = self._get_variant_clause_bundle(v)

        vt_key = "HK_CONNECT" if has_hk else "STANDARD"
        vt_variants = self.pro_clauses["VALUATION_TIMING"]["variants"]
        valuation_text = vt_variants.get(vt_key, vt_variants["STANDARD"])["text"]
        valuation_text = valuation_text.replace("{EXCHANGE_NAME_CN}", exch_cn)
        v["VALUATION_TIMING_CLAUSE"] = valuation_text

        v["RISK_DISCLOSURE_CHUANGYE"] = self.pro_clauses["RISK_DISCLOSURE_CHUANGYE"]["variants"]["DEFAULT"]["text"]
        v["RISK_DISCLOSURE_KECHUANG"] = self.pro_clauses["RISK_DISCLOSURE_KECHUANG"]["variants"]["DEFAULT"]["text"]
        v["RISK_DISCLOSURE_HK_CONNECT"] = self.pro_clauses["RISK_DISCLOSURE_HK_CONNECT"]["variants"]["DEFAULT"]["text"]

        custodian_name = v.get("CUSTODIAN_NAME", "")
        custodian_contacts = self.pro_clauses.get("CUSTODIAN_INFO_PROSPECTUS", {}).get("custodians", {})
        info = custodian_contacts.get(custodian_name, {})
        v.setdefault("CUSTODIAN_DEPT", info.get("dept", "[待填写：托管部门名称]"))
        v.setdefault("CUSTODIAN_PHONE", info.get("phone", "[待填写：服务电话]"))
        v.setdefault("CUSTODIAN_WEBSITE", info.get("website", "[待填写：网址]"))

        v.setdefault("FUND_MANAGER_NAME", "[待填写：基金经理姓名]")
        v.setdefault("FUND_MANAGER_BIO", "[待填写：基金经理简介（学历、从业经历、任职日期等）]")
        v.setdefault("CSRC_APPROVAL_NO", "202X年X月X日证监许可〔202X〕XXXX号")
        v.setdefault("MIN_SUB_UNIT", "1,000,000份（即100万份）")

        chapter6 = variant_bundle.get("chapter_6", {})
        chapter10 = variant_bundle.get("chapter_10", {})
        v["PROSPECTUS_VARIANT_KEY"] = variant_key
        v["PROSPECTUS_CH6_SEC4"] = chapter6.get("section_4", "")
        v["PROSPECTUS_CH6_SEC7"] = chapter6.get("section_7", "")
        v["PROSPECTUS_CH6_SEC11"] = chapter6.get("section_11", "")
        v["PROSPECTUS_CH6_SEC12"] = chapter6.get("section_12", "")
        v["PROSPECTUS_CH6_SEC13"] = chapter6.get("section_13", "")
        v["PROSPECTUS_CH7_BODY"] = variant_bundle.get("chapter_7", "")
        v["PROSPECTUS_CH9_BODY"] = variant_bundle.get("chapter_9", "")
        v["PROSPECTUS_CH10_PRELUDE"] = chapter10.get("prelude", "")
        v["PROSPECTUS_CH10_SEC4"] = chapter10.get("section_4", "")
        v["PROSPECTUS_CH10_SEC7"] = chapter10.get("section_7", "")
        v["PROSPECTUS_CH21_TITLES"] = self.pro_clauses.get("CHAPTER21_TITLES", {}).get("text", "")

        risk_bodies = self.pro_clauses.get("RISK_CHAPTER_BODIES", {}).get("variants", {})
        if has_hk:
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("HK_CONNECT", "")
        elif v.get("IS_KECHUANG"):
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("KECHUANG", "")
        elif v.get("IS_CHUANGYE"):
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("CHUANGYE", "")
        else:
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("STANDARD_A", "")

        return v

    # ── Step 3: 从合同全文提取各关键部分（内容摘要 + 各章节来源段落）───────
    def _extract_contract_sections(self, v: dict) -> dict:
        """
        生成完整合同文本，按以下逻辑提取各部分供招募说明书使用：

        #二章 释义 → 《基金合同》第二部分
        #十章 申购赎回 子条款 → 《基金合同》第八部分 对应子条款：
          二、开放日 → CONTRACT_PART8_SEC2
          三、原则   → CONTRACT_PART8_SEC3
          六、对价   → CONTRACT_PART8_SEC6
          七、拒绝申购 → CONTRACT_PART8_SEC7
          八、暂停赎回 → CONTRACT_PART8_SEC8
          九、其他方式 → CONTRACT_PART8_SEC9
          十、非交易过户 → CONTRACT_PART8_SEC10
          十一、冻结 → CONTRACT_PART8_SEC11
          十二、转让 → CONTRACT_PART8_SEC12
          十三、其他业务 → CONTRACT_PART8_SEC13
          十四、清算交收模式 → CONTRACT_PART8_SEC14
        #十一章 投资 → 《基金合同》第十四部分
        #十二章 财产 → 《基金合同》第十五部分
        #十三章 估值 → 《基金合同》第十六部分
        #十四章 收益分配 → 三、基金收益分配原则 同 《基金合同》三
        #十五章 费用 → 《基金合同》第十七部分
        #十六章 会计审计 → 《基金合同》第十九部分
        #十七章 信息披露 → 《基金合同》第二十部分
        #十九章 变更终止 → 《基金合同》第二十一部分
        #二十章 合同摘要 → 《基金合同》第二十六部分
        """
        try:
            contract_text = engine.generate(v)
        except Exception as _exc:
            import logging
            logging.warning("ProspectusEngine: 合同生成失败，部分章节将使用占位符。原因：%s", _exc)
            contract_text = ""

        # 按"第X部分"标题切分合同全文
        RE_PART = re.compile(r'^第[一二三四五六七八九十百]+部分\s+\S[^\n]*', re.MULTILINE)
        part_iter = list(RE_PART.finditer(contract_text))

        # 中文数字顺序表，用于精确匹配部分编号（避免"第二部分"匹配"第二十部分"）
        _CN_ORDER = [
            "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
            "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九",
            "二十", "二十一", "二十二", "二十三", "二十四", "二十五", "二十六",
        ]

        def _get_part(cn_num: str) -> str:
            """
            按准确中文部分编号提取该部分全文（标题行之后的内容）。
            cn_num 示例：'第二部分'、'第八部分'、'第二十六部分'
            使用精确匹配避免"第二部分"匹配到"第二十部分"。
            """
            # 构造精确匹配模式：编号后紧跟"部分"且后面为空白
            pattern = re.compile(
                r'^' + re.escape(cn_num) + r'\s+\S[^\n]*',
                re.MULTILINE
            )
            matched_idx = []
            for i, m in enumerate(part_iter):
                if pattern.match(m.group()):
                    matched_idx.append(i)
            if matched_idx:
                # Contract text includes TOC + body with duplicate "第X部分" headings.
                # Use the last match to target body section, not TOC line.
                i = matched_idx[-1]
                m = part_iter[i]
                start = m.end()
                end = part_iter[i + 1].start() if i + 1 < len(part_iter) else len(contract_text)
                return contract_text[start:end].strip()
            return ""

        def _get_part_subsection(part_cn: str, cn_num: str) -> str:
            """
            从指定部分提取指定中文序号的子条款全文。
            cn_num 示例：'九', '十', '十一', ...
            使用精确匹配：标题为 "序号、" 格式，用 word boundary 避免"十"匹配"十一"。
            返回该子条款标题行（含序号）及其正文，直到下一个同级序号或部分结束。
            """
            part_text = _get_part(part_cn)
            if not part_text:
                return ""
            # 精确匹配行首的中文序号：序号后紧跟"、"，序号本身是完整单词
            # 构造列表后按行首定位，避免多字符序号前缀混淆
            RE_CN = re.compile(
                r'^([一二三四五六七八九十百]+)、',
                re.MULTILINE
            )
            markers = list(RE_CN.finditer(part_text))
            for i, m in enumerate(markers):
                heading = m.group(1)          # 捕获组1 = 纯序号，不含"、"
                if heading == cn_num:
                    start = m.start()
                    end = markers[i + 1].start() if i + 1 < len(markers) else len(part_text)
                    return part_text[start:end].strip()
            return ""

        def _get_part8_subsection(cn_num: str) -> str:
            return _get_part_subsection("第八部分", cn_num)

        # #二章 释义（完整第二部分，不含标题行）
        v.setdefault("CONTRACT_DEFS_TEXT", _get_part("第二部分") or
                     "[待填写：释义内容，请从基金合同第二部分复制]")

        # #十章 申购赎回 各子条款（来自第八部分）
        # 二/三/六 保持固定（模板中已有固定文本），此处仍提取供模板引用
        v.setdefault("CONTRACT_PART8_SEC2", _get_part8_subsection("二") or
                     "[待填写：申购和赎回的开放日及时间，来自基金合同第八部分]")
        v.setdefault("CONTRACT_PART8_SEC3", _get_part8_subsection("三") or
                     "[待填写：申购和赎回的原则，来自基金合同第八部分]")
        v.setdefault("CONTRACT_PART8_SEC6", _get_part8_subsection("六") or
                     "[待填写：申购和赎回的对价、费用及其用途，来自基金合同第八部分]")
        # 七 = 拒绝暂停申购（合同 七，招募 八）
        v.setdefault("CONTRACT_PART8_SEC7", _get_part8_subsection("七") or
                     "[待填写：拒绝或暂停申购的情形，来自基金合同第八部分]")
        # 八 = 暂停赎回（合同 八，招募 九）
        v.setdefault("CONTRACT_PART8_SEC8", _get_part8_subsection("八") or
                     "[待填写：暂停赎回或延缓支付赎回对价的情形，来自基金合同第八部分]")
        # 招募 十 → 合同八.九
        v.setdefault("CONTRACT_PART8_SEC9", _get_part8_subsection("九") or
                     "[待填写：其他申购赎回方式，来自基金合同第八部分]")
        # 招募 十一 → 合同八.十
        v.setdefault("CONTRACT_PART8_SEC10", _get_part8_subsection("十") or
                     "[待填写：基金的非交易过户，来自基金合同第八部分]")
        # 招募 十二 → 合同八.十一
        v.setdefault("CONTRACT_PART8_SEC11", _get_part8_subsection("十一") or
                     "[待填写：基金份额的冻结和解冻，来自基金合同第八部分]")
        # 招募 十三 → 合同八.十二
        v.setdefault("CONTRACT_PART8_SEC12", _get_part8_subsection("十二") or
                     "[待填写：基金份额的转让，来自基金合同第八部分]")
        # 招募 十四 → 合同八.十三
        v.setdefault("CONTRACT_PART8_SEC13", _get_part8_subsection("十三") or
                     "[待填写：其他业务，来自基金合同第八部分]")
        # 招募 十五 → 合同八.十四（交易所相关，文字含交易所名称）
        v.setdefault("CONTRACT_PART8_SEC14", _get_part8_subsection("十四") or
                     "[待填写：清算交收模式，来自基金合同第八部分]")

        # #十一-十七、十九章 各对应合同部分
        v.setdefault("CONTRACT_INVEST_TEXT", _get_part("第十四部分") or
                     "[待填写：基金的投资，来自基金合同第十四部分]")
        v.setdefault("CONTRACT_ASSET_TEXT", _get_part("第十五部分") or
                     "[待填写：基金的财产，来自基金合同第十五部分]")
        v.setdefault("CONTRACT_VALUATION_TEXT", _get_part("第十六部分") or
                     "[待填写：基金资产估值，来自基金合同第十六部分]")
        v.setdefault("CONTRACT_FEE_TEXT", _get_part("第十七部分") or
                     "[待填写：基金的费用与税收，来自基金合同第十七部分]")
        v.setdefault("CONTRACT_AUDIT_TEXT", _get_part("第十九部分") or
                     "[待填写：基金的会计与审计，来自基金合同第十九部分]")
        v.setdefault("CONTRACT_DISCLOSURE_TEXT", _get_part("第二十部分") or
                     "[待填写：基金的信息披露，来自基金合同第二十部分]")
        v.setdefault("CONTRACT_TERMINATION_TEXT", _get_part("第二十一部分") or
                     "[待填写：基金合同的变更、终止与基金财产的清算，来自基金合同第二十一部分]")

        # #二十章 合同内容摘要 → 第二十六部分
        summary = self._strip_signing_page_from_contract_summary(_get_part("第二十六部分"))
        v.setdefault("CONTRACT_SUMMARY_TEXT", summary or
                     "[待填写：基金合同内容摘要，请先生成基金合同，从第二十六部分复制此处]")

        # #十四章 三、基金收益分配原则 → 《基金合同》第十八部分 三
        v.setdefault("CONTRACT_PART18_SEC3", _get_part_subsection("第十八部分", "三") or
                     "[待填写：基金收益分配原则，来自基金合同第十八部分第三项]")

        return v

    # ── Step 3 (兼容旧接口) ────────────────────────────────────────────────
    def _extract_contract_summary(self, v: dict) -> dict:
        """向后兼容旧方法名，实际调用 _extract_contract_sections。"""
        return self._extract_contract_sections(v)

    @staticmethod
    def _chapter_num_to_cn(num: int) -> str:
        cn_map = {
            1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
            11: "十一", 12: "十二", 13: "十三", 14: "十四", 15: "十五", 16: "十六", 17: "十七", 18: "十八",
            19: "十九", 20: "二十", 21: "二十一", 22: "二十二", 23: "二十三", 24: "二十四", 25: "二十五",
        }
        return cn_map.get(num, "")

    @staticmethod
    def _split_top_sections(chapter_body: str) -> dict:
        section_re = re.compile(r"^([一二三四五六七八九十百]+)、[^\n]*", re.MULTILINE)
        matches = list(section_re.finditer(chapter_body or ""))
        sections = {}
        for i, m in enumerate(matches):
            sec_cn = m.group(1)
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(chapter_body)
            sections[sec_cn] = chapter_body[start:end].strip()
        return sections

    def _load_reference_fixed_content(self) -> dict:
        """
        Load fixed chapter/section text from the canonical prospectus DOCX once.
        """
        if self._reference_fixed_cache is not None:
            return self._reference_fixed_cache

        data = {}
        if not REFERENCE_PROSPECTUS_DOCX.exists():
            self._reference_fixed_cache = data
            return data

        try:
            from docx import Document
        except Exception:
            self._reference_fixed_cache = data
            return data

        try:
            doc = Document(str(REFERENCE_PROSPECTUS_DOCX))
        except Exception:
            self._reference_fixed_cache = data
            return data

        paras = list(doc.paragraphs)
        chapter_starts = []
        for i, p in enumerate(paras):
            txt = (p.text or "").strip()
            if not txt:
                continue
            style_name = ""
            try:
                style_name = (p.style.name or "")
            except Exception:
                style_name = ""
            style_lower = style_name.lower()
            if ("heading 2" in style_lower) or ("标题 2" in style_name) or ("标题2" in style_name):
                chapter_starts.append(i)

        if not chapter_starts:
            chap_re = re.compile(r"^第[一二三四五六七八九十百]+章")
            chapter_starts = [
                i for i, p in enumerate(paras) if chap_re.match((p.text or "").strip())
            ]

        for idx, start_i in enumerate(chapter_starts):
            chap_cn = self._chapter_num_to_cn(idx + 1)
            if not chap_cn:
                continue
            end_i = chapter_starts[idx + 1] if idx + 1 < len(chapter_starts) else len(paras)
            chapter_title = (paras[start_i].text or "").strip()

            body_lines = []
            for p in paras[start_i + 1:end_i]:
                line = (p.text or "").strip()
                if line:
                    body_lines.append(line)
            body = "\n".join(body_lines).strip()
            if not body:
                continue

            data[chap_cn] = {
                "title": chapter_title,
                "body": body,
                "sections": self._split_top_sections(body),
            }

        self._reference_fixed_cache = data
        return data

    @staticmethod
    def _replace_numbered_item_in_section(chapter_body: str, section_cn: str, item_no: str, new_item: str) -> str:
        body = (chapter_body or "").strip()
        if not body:
            return chapter_body

        sec_re = re.compile(rf"^{section_cn}、[^\n]*", re.MULTILINE)
        next_sec_re = re.compile(r"^[一二三四五六七八九十百]+、[^\n]*", re.MULTILINE)
        m = sec_re.search(body)
        if not m:
            return chapter_body
        n = next_sec_re.search(body, m.end())
        sec_end = n.start() if n else len(body)

        sec_block = body[m.start():sec_end]
        item_text = (new_item or "").strip()
        if not item_text:
            return chapter_body
        if not re.match(rf"^{re.escape(item_no)}、", item_text):
            item_text = f"{item_no}、{item_text}"

        item_re = re.compile(rf"^{re.escape(item_no)}、[^\n]*", re.MULTILINE)
        next_item_re = re.compile(r"^\d+、[^\n]*", re.MULTILINE)
        im = item_re.search(sec_block)
        if not im:
            if not sec_block.endswith("\n"):
                sec_block += "\n"
            sec_block = f"{sec_block}{item_text}\n"
        else:
            inn = next_item_re.search(sec_block, im.end())
            item_end = inn.start() if inn else len(sec_block)
            sec_block = f"{sec_block[:im.start()]}{item_text}\n{sec_block[item_end:].lstrip()}"

        new_body = f"{body[:m.start()]}{sec_block}{body[sec_end:]}"
        return new_body

    def _apply_reference_fixed_content(self, text: str, v: dict) -> str:
        """
        Apply canonical fixed text from red-dividend prospectus docx.
        """
        ref = self._load_reference_fixed_content()
        if not ref:
            return text

        def ref_chapter(chap_cn: str) -> str:
            return ref.get(chap_cn, {}).get("body", "")

        def ref_section(chap_cn: str, sec_cn: str) -> str:
            return ref.get(chap_cn, {}).get("sections", {}).get(sec_cn, "")

        ch3 = ref_chapter("三")
        if ch3:
            manager_name = str(v.get("FUND_MANAGER_NAME") or "[待填写：基金经理姓名]").strip()
            manager_bio = str(v.get("FUND_MANAGER_BIO") or "[待填写：基金经理简介]").strip()
            sec_item3 = f"3、本基金的基金经理为{manager_name}。"
            if manager_bio:
                sec_item3 = f"{sec_item3}\n{manager_bio}"
            ch3 = self._replace_numbered_item_in_section(ch3, "二", "3", sec_item3)
            text = self._replace_chapter_body(text, "三", ch3)

        text = self._replace_chapter_body(text, "四", "【托管人情况待填写】")

        text = self._replace_subsection_in_chapter(text, "五", "一", ref_section("五", "一"))
        text = self._replace_subsection_in_chapter(text, "五", "二", ref_section("五", "二"))
        sec3 = ref_section("五", "三")
        if sec3:
            sec3 = re.sub(r"^经办律师[：:].*$", "经办律师：丁媛、李晓露", sec3, flags=re.MULTILINE)
        text = self._replace_subsection_in_chapter(text, "五", "三", sec3)
        text = self._replace_subsection_in_chapter(text, "五", "四", "四、审计基金财产的会计师事务所\n【待填写】")

        text = self._replace_chapter_body(text, "八", ref_chapter("八"))
        text = self._replace_chapter_body(text, "十四", ref_chapter("十四"))

        for chap_cn in ("二十二", "二十三", "二十四"):
            text = self._replace_chapter_body(text, chap_cn, ref_chapter(chap_cn))

        return text

    @staticmethod
    def _find_chapter_span(text: str, chapter_cn: str):
        """Return (start, heading_end, end) for a chapter, or None."""
        chap_re = re.compile(rf"^第{chapter_cn}章[^\n]*", re.MULTILINE)
        any_chap_re = re.compile(r"^第[一二三四五六七八九十百]+章[^\n]*", re.MULTILINE)
        matches = list(chap_re.finditer(text))
        if not matches:
            return None
        # Template contains a TOC copy and a body copy; use the last one (body section).
        m = matches[-1]
        n = any_chap_re.search(text, m.end())
        end = n.start() if n else len(text)
        return m.start(), m.end(), end

    def _replace_chapter_body(self, text: str, chapter_cn: str, new_body: str) -> str:
        new_body = (new_body or "").strip()
        if not new_body:
            return text
        span = self._find_chapter_span(text, chapter_cn)
        if not span:
            return text
        start, heading_end, end = span
        prefix = text[:heading_end].rstrip("\n")
        suffix = text[end:].lstrip("\n")
        if suffix:
            return f"{prefix}\n{new_body}\n{suffix}"
        return f"{prefix}\n{new_body}\n"

    def _replace_subsection_in_chapter(self, text: str, chapter_cn: str, subsection_cn: str, new_subsection: str) -> str:
        new_subsection = (new_subsection or "").strip()
        if not new_subsection:
            return text
        span = self._find_chapter_span(text, chapter_cn)
        if not span:
            return text
        start, _heading_end, end = span
        block = text[start:end]

        sec_re = re.compile(rf"^{subsection_cn}、[^\n]*", re.MULTILINE)
        next_sec_re = re.compile(r"^[一二三四五六七八九十百]+、[^\n]*", re.MULTILINE)
        m = sec_re.search(block)
        if not m:
            # If subsection does not exist in template, append it at chapter tail.
            if not block.endswith("\n"):
                block += "\n"
            block = f"{block}{new_subsection}\n"
        else:
            n = next_sec_re.search(block, m.end())
            sec_end = n.start() if n else len(block)
            block = f"{block[:m.start()]}{new_subsection}\n{block[sec_end:].lstrip()}"

        return f"{text[:start]}{block}{text[end:]}"

    @staticmethod
    def _retag_subsection_number(subsection_text: str, new_cn: str) -> str:
        txt = (subsection_text or "").strip()
        return re.sub(r"^[一二三四五六七八九十百]+、", f"{new_cn}、", txt, count=1)

    @staticmethod
    def _ensure_subsection_heading(subsection_text: str, cn_num: str) -> str:
        txt = (subsection_text or "").strip()
        if not txt:
            return ""
        if re.match(r"^[一二三四五六七八九十百]+、", txt):
            return txt
        return f"{cn_num}、\n{txt}"

    @staticmethod
    def _join_nonempty_blocks(blocks) -> str:
        return "\n\n".join(block.strip() for block in blocks if (block or "").strip())

    @staticmethod
    def _normalize_reused_prospectus_chapter(chapter_text: str) -> str:
        text = (chapter_text or "").strip()
        if not text:
            return ""
        text = text.replace("详见招募说明书的规定", "详见招募说明书“侧袋机制”部分的规定")
        text = text.replace("本基金合同", "基金合同")
        return text

    @staticmethod
    def _get_prospectus_min_sub_unit(v: dict) -> str:
        return str(v.get("MIN_SUB_UNIT") or "1,000,000份（即100万份）").strip()

    def _normalize_prospectus_risk_chapter(self, chapter_text: str, v: dict) -> str:
        text = (chapter_text or "").strip()
        if not text:
            return ""
        min_sub_unit = self._get_prospectus_min_sub_unit(v)
        for old in ("1,000,000份（即100万份）", "1,000,000份"):
            text = text.replace(old, min_sub_unit)
        text = re.sub(r"(?<=按原)\d[\d,]*份(?:（即\d+万份）)?", min_sub_unit, text)
        text = re.sub(r"(?<=新的)\d[\d,]*份(?:（即\d+万份）)?", min_sub_unit, text)
        return text

    def _build_chapter_six_body(self, v: dict, ref: dict) -> str:
        def ref_section(sec_cn: str) -> str:
            return ref.get("六", {}).get("sections", {}).get(sec_cn, "")

        sec8 = """八、认购费用
认购费用由投资人承担，不高于0.30%，认购费率如下表所示：

|   |   |
|---|---|
|认购份额（S）|认购费率|
|S＜100万份|0.30%|
|S≥100万份|每笔500元|

基金管理人办理网下现金认购和网下股票认购不收取认购费。发售代理机构办理网上现金认购、网下现金认购、网下股票认购时可参照上述费率结构，按照不高于0.3%的标准收取一定的佣金。投资人申请重复现金认购的，须按每次认购所对应的费率档次分别计费。"""

        blocks = [
            f"本基金由基金管理人依照《基金法》、《运作办法》、《销售办法》、基金合同及其他有关规定，并经中国证监会{v.get('CSRC_APPROVAL_NO', '202X年X月X日证监许可〔202X〕XXXX号')}文注册募集。",
            "本基金为交易型开放式基金，股票型基金，基金存续期限为不定期。",
            ref_section("一"),
            ref_section("二"),
            ref_section("三"),
            v.get("PROSPECTUS_CH6_SEC4", ""),
            ref_section("五"),
            ref_section("六"),
            v.get("PROSPECTUS_CH6_SEC7", ""),
            sec8,
            ref_section("九"),
            ref_section("十"),
            v.get("PROSPECTUS_CH6_SEC11", ""),
            v.get("PROSPECTUS_CH6_SEC12", ""),
            v.get("PROSPECTUS_CH6_SEC13", ""),
        ]
        return self._join_nonempty_blocks(blocks)

    @staticmethod
    def _build_chapter_ten_limits_table(min_sub_unit: str) -> str:
        return "\n".join([
            "|项目|内容|",
            "|---|---|",
            f"|最小申购赎回单位|{min_sub_unit}|",
            "|申购/赎回份额上限|以申购赎回清单或相关公告为准|",
        ])

    def _build_chapter_ten_sec5(self, ref: dict, v: dict) -> str:
        sec5 = ref.get("十", {}).get("sections", {}).get("五", "")
        min_sub_unit = self._get_prospectus_min_sub_unit(v)
        if sec5:
            sec5 = re.sub(r"\u76ee\u524d\uff0c\u672c\u57fa\u91d1\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d\u4e3a[^\uff0c\u3002\uff1b]+", f"\u76ee\u524d\uff0c\u672c\u57fa\u91d1\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d\u4e3a{min_sub_unit}", sec5, count=1)
            if min_sub_unit not in sec5:
                sec5 = re.sub(r"\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d[^\u3002]*", f"\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d\u4e3a{min_sub_unit}", sec5, count=1)
            return sec5
            if "|项目|内容|" not in sec5:
                sec5 = self._join_nonempty_blocks([sec5, "申购赎回要点如下表所示：", table_block])
            return sec5
        return f"""五、申购和赎回的数额限制
1、投资人申购、赎回的基金份额需为最小申购赎回单位的整数倍。目前，本基金最小申购赎回单位为{min_sub_unit}，基金管理人有权对其进行调整，并在调整实施前依照《信息披露办法》的有关规定在规定媒介上公告。
2、基金管理人可以规定本基金当日申购份额及当日赎回份额上限，具体规定请参见申购赎回清单或相关公告。
3、基金管理人可根据市场情况，在法律法规允许的情况下，合理调整上述申购和赎回的数量或比例限制，并在实施前依照《信息披露办法》的有关规定在规定媒介上公告。

申购赎回要点如下表所示：

{table_block}"""

    def _build_contract_section(self, v: dict, var_name: str, sec_cn: str) -> str:
        sec_text = self._retag_subsection_number(v.get(var_name, ""), sec_cn)
        return self._ensure_subsection_heading(sec_text, sec_cn)

    def _ensure_distribution_conditions_section(self, chapter_body: str) -> str:
        if not chapter_body or "\u56db\u3001\u6536\u76ca\u5206\u914d\u6761\u4ef6" in chapter_body:
            return chapter_body
        if "\u56db\u3001\u6536\u76ca\u5206\u914d\u65b9\u6848" not in chapter_body:
            return chapter_body
        conditions = (
            "\u56db\u3001\u6536\u76ca\u5206\u914d\u6761\u4ef6\n"
            "1\u3001\u57fa\u91d1\u7ba1\u7406\u4eba\u53ef\u6bcf\u6708\u5bf9\u57fa\u91d1\u76f8\u5bf9\u4e1a\u7ee9\u6bd4\u8f83\u57fa\u51c6\u7684\u8d85\u989d\u6536\u76ca\u7387\u4ee5\u53ca\u57fa\u91d1\u7684\u53ef\u4f9b\u5206\u914d\u5229\u6da6\u8fdb\u884c\u8bc4\u4f30\uff0c\u5728\u7b26\u5408\u57fa\u91d1\u6536\u76ca\u5206\u914d\u6761\u4ef6\u4e0b\uff0c\u53ef\u5b89\u6392\u6536\u76ca\u5206\u914d\u3002\n"
            "2\u3001\u57fa\u91d1\u6536\u76ca\u5206\u914d\u6761\u4ef6\u3001\u8bc4\u4f30\u65f6\u95f4\u3001\u5206\u914d\u65f6\u95f4\u3001\u5206\u914d\u65b9\u6848\u53ca\u6bcf\u6b21\u57fa\u91d1\u6536\u76ca\u5206\u914d\u6570\u989d\u7b49\u5185\u5bb9\uff0c\u57fa\u91d1\u7ba1\u7406\u4eba\u53ef\u4ee5\u6839\u636e\u5b9e\u9645\u60c5\u51b5\u786e\u5b9a\u5e76\u6309\u7167\u6709\u5173\u89c4\u5b9a\u516c\u544a\u3002"
        )
        return chapter_body.replace("\u56db\u3001\u6536\u76ca\u5206\u914d\u65b9\u6848", f"{conditions}\n\u4e94\u3001\u6536\u76ca\u5206\u914d\u65b9\u6848", 1)

    def _build_chapter_ten_body(self, v: dict, ref: dict) -> str:
        blocks = []
        prelude = v.get("PROSPECTUS_CH10_PRELUDE", "")
        if prelude:
            blocks.append(prelude)
        blocks.extend([
            ref.get("十", {}).get("sections", {}).get("一", ""),
            self._build_contract_section(v, "CONTRACT_PART8_SEC2", "二"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC3", "三"),
            v.get("PROSPECTUS_CH10_SEC4", ""),
            self._build_chapter_ten_sec5(ref, v),
            self._build_contract_section(v, "CONTRACT_PART8_SEC6", "六"),
            v.get("PROSPECTUS_CH10_SEC7", ""),
            self._build_contract_section(v, "CONTRACT_PART8_SEC7", "八"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC8", "九"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC9", "十"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC10", "十一"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC11", "十二"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC12", "十三"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC13", "十四"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC14", "十五"),
        ])
        return self._join_nonempty_blocks(blocks)

    def _get_product_type_chapter_builders(self, v: dict, ref: dict) -> dict:
        product_type = self._get_product_type(v)
        if product_type == "ETF":
            return {
                "十": lambda: self._build_chapter_ten_body(v, ref),
            }
        return {}

    def _apply_prospectus_chapter_logic(self, text: str, v: dict) -> str:
        """
        Apply chapter-level composition rules for prospectus generation.
        """
        text = self._apply_reference_fixed_content(text, v)
        ref = self._load_reference_fixed_content()

        chapter_builders = {
            "二": lambda: v.get("CONTRACT_DEFS_TEXT", ""),
            "六": lambda: self._build_chapter_six_body(v, ref),
            "七": lambda: v.get("PROSPECTUS_CH7_BODY", ""),
            "九": lambda: v.get("PROSPECTUS_CH9_BODY", ""),
            "十": lambda: self._build_chapter_ten_body(v, ref),
            "十一": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_INVEST_TEXT", "")),
            "十二": lambda: v.get("CONTRACT_ASSET_TEXT", ""),
            "十三": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_VALUATION_TEXT", "")),
            "十五": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_FEE_TEXT", "")),
            "十六": lambda: v.get("CONTRACT_AUDIT_TEXT", ""),
            "十七": lambda: v.get("CONTRACT_DISCLOSURE_TEXT", ""),
            "十八": lambda: self._normalize_prospectus_risk_chapter(v.get("PROSPECTUS_CH18_BODY", ""), v),
            "十九": lambda: v.get("CONTRACT_TERMINATION_TEXT", ""),
            "二十": lambda: v.get("CONTRACT_SUMMARY_TEXT", ""),
            "二十一": lambda: v.get("PROSPECTUS_CH21_TITLES", ""),
        }
        for chap_cn, builder in chapter_builders.items():
            text = self._replace_chapter_body(text, chap_cn, builder())

        sec3 = self._ensure_subsection_heading(v.get("CONTRACT_PART18_SEC3", ""), "三")
        text = self._replace_subsection_in_chapter(text, "十四", "三", sec3)
        return text

    def _process_conditionals(self, text: str, v: dict) -> str:
        return engine._process_conditionals(text, v)

    def _replace_placeholders(self, text: str, v: dict) -> str:
        return engine._replace_placeholders(text, v)

    # ── Step 6: 重排序号（使用"章"作为重置标志）─────────────────────────
    def _renumber_sequences(self, text: str) -> str:
        lines = text.split("\n")
        RE_NUM = re.compile(r"^(\d+)(、)")
        RE_RESET = re.compile(
            r"^(?:[一二三四五六七八九十百]+、|第[一二三四五六七八九十百]+章)"
        )
        last_num = None
        result = []
        for line in lines:
            if RE_RESET.match(line.strip()):
                last_num = None
                result.append(line)
                continue
            m = RE_NUM.match(line)
            if m:
                num = int(m.group(1))
                if num == 1:
                    last_num = 1
                elif last_num is not None and num > last_num + 1:
                    expected = last_num + 1
                    line = re.sub(r"^\d+、", f"{expected}、", line, count=1)
                    last_num = expected
                else:
                    last_num = num
            result.append(line)
        return "\n".join(result)

    # ── Step 7: 清理（检测招募说明书封面而非基金合同）────────────────────
    def _looks_like_cover_title(self, line: str) -> bool:
        s = line.strip()
        if not s:
            return False
        if s.startswith(("#", ">", "-", "*", "**", "`")):
            return False
        if not s.endswith("\u62db\u52df\u8bf4\u660e\u4e66"):
            return False
        forbidden = ("\u6a21\u677f\u8bf4\u660e", "\u6761\u4ef6\u53d8\u91cf\u5f15\u7528\u8bf4\u660e", "\u5dee\u5f02\u6761\u6b3e", ".json", "_CLAUSE", "_DEF")
        return not any(token in s for token in forbidden)

    def _find_body_start_index(self, lines: list[str]) -> int:
        nonempty = [(idx, line.strip()) for idx, line in enumerate(lines) if line.strip()]
        for pos, (idx, stripped) in enumerate(nonempty):
            if not self._looks_like_cover_title(stripped):
                continue
            window = [item[1] for item in nonempty[pos + 1:pos + 6]]
            has_manager = any(item.startswith("\u57fa\u91d1\u7ba1\u7406\u4eba") for item in window)
            has_custodian = any(item.startswith("\u57fa\u91d1\u6258\u7ba1\u4eba") for item in window)
            if has_manager and has_custodian:
                return idx
        for idx, stripped in nonempty:
            if self._looks_like_cover_title(stripped):
                return idx
        return 0

    def _is_internal_metadata_line(self, line: str) -> bool:
        s = line.strip()
        if not s:
            return False
        if s in {"---", "***"}:
            return True
        if s.startswith((">", "# ", "## ")):
            return True
        keywords = (
            "\u6a21\u677f\u8bf4\u660e",
            "\u6761\u4ef6\u53d8\u91cf\u5f15\u7528\u8bf4\u660e",
            "\u5dee\u5f02\u6761\u6b3e\u5f15\u7528\u8bf4\u660e",
            "VALUATION_TIMING_CLAUSE",
            "WORKING_DAY_DEF",
            "BUSINESS_RULES_DEF",
            "NON_COMPONENT_SCOPE",
            "DISTRIBUTION_FREQ_CLAUSE",
            "MGMT_FEE_PAYMENT_METHOD",
            "CUSTODY_FEE_PAYMENT_METHOD",
        )
        if any(keyword in s for keyword in keywords):
            return True
        if ".json" in s and "\u89c1" in s and "`" in s:
            return True
        return False

    def validate_exportable_text(self, text: str) -> dict:
        metadata_matches = []

        def _append_unique(target: list[str], value: str):
            if value not in target:
                target.append(value)

        for raw_line in text.splitlines():
            stripped = raw_line.strip()
            if not stripped:
                continue
            if self._is_internal_metadata_line(stripped):
                _append_unique(metadata_matches, stripped)

        if metadata_matches:
            return {
                "ok": False,
                "error_type": "template_metadata_leaked",
                "error": "招募说明书正文中仍包含模板说明或内部标记，请先清理后再导出。",
                "matches": metadata_matches[:5],
            }
        # Placeholder content is allowed to export; only internal template metadata blocks delivery.
        return {"ok": True, "matches": []}

    def _cleanup(self, text: str) -> str:
        lines = text.split("\n")
        start_idx = self._find_body_start_index(lines)
        clean = []
        for idx, line in enumerate(lines):
            if idx < start_idx:
                continue
            if self._is_internal_metadata_line(line):
                continue
            clean.append(line)
        text = "\n".join(clean)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        return text

    def generate(self, form_data: dict) -> str:
        # Step 1
        v = self._derive_variables(form_data)
        # Step 2
        v = self._inject_clause_texts(v)
        # Step 3: read template
        template_text = TEMPLATE_MD.read_text(encoding="utf-8")
        # Step 4
        text = self._process_conditionals(template_text, v)
        # Step 5
        text = self._replace_placeholders(text, v)
        # Step 5b: 修复条件删除导致的序号跳跃
        text = self._renumber_sequences(text)
        # Step 6
        text = self._cleanup(text)
        return text

    # ── Step 5b: 重排阿拉伯序号（修复条件块删除项目后的跳跃）──────────────
    def _renumber_sequences(self, text: str) -> str:
        """
        修复条件块删除项目后阿拉伯序号列表的跳跃。
        规则：
          - 只处理行首 `数字、` 格式
          - num == 1：新列表开始，重置计数器，不修改
          - num > expected_next（且 num > 1）：检测到跳跃，将该行及后续连续项重排
          - 遇到中文序号标题（一、二、…）或章节标题（第X部分）时重置计数器
        """
        lines = text.split("\n")
        RE_NUM = re.compile(r"^(\d+)(、)")
        RE_RESET = re.compile(
            r"^(?:[一二三四五六七八九十百]+、|第[一二三四五六七八九十百]+部分)"
        )
        last_num = None
        result = []
        for line in lines:
            if RE_RESET.match(line.strip()):
                last_num = None
                result.append(line)
                continue
            m = RE_NUM.match(line)
            if m:
                num = int(m.group(1))
                if num == 1:
                    last_num = 1
                elif last_num is not None and num > last_num + 1:
                    expected = last_num + 1
                    line = re.sub(r"^\d+、", f"{expected}、", line, count=1)
                    last_num = expected
                else:
                    last_num = num
            result.append(line)
        return "\n".join(result)

    # ── 中文标点规范化 ───────────────────────────────────────────────────────
    @staticmethod
    def _to_chinese_punct(text: str) -> str:
        """将文本中的半角标点转为全角中文标点（保留数字中的逗号与小数点）。"""
        text = text.replace('(', '（').replace(')', '）')
        text = text.replace('[', '【').replace(']', '】')
        text = text.replace(';', '；')
        # 逗号仅在两侧均非数字时替换，保留千位分隔符如"1,000"
        text = re.sub(r'(?<!\d),(?!\d)', '，', text)
        return text

    # ── Word (.docx) 导出 ────────────────────────────────────────────────────
    def build_docx(self, contract_text: str) -> bytes:
        """
        将合同纯文本转换为格式化 Word 文档，与参考合同格式完全对齐。
        - 封面：独立 section + vAlign=center，彻底消除内容溢出第二页的问题
        - 签署页：左对齐 + 精确空行间距（来自参考文档XML）
        - 章节标题：Times New Roman 15pt 加粗，pageBreakBefore
        - 正文：宋体 12pt，两端对齐，首行缩进2字，1.5倍行距
        - 数字/英文：Times New Roman；中文：宋体
        """
        # 标准化中文标点
        contract_text = self._to_chinese_punct(contract_text)
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Twips

        # ── 1. 文档与页面设置 ──────────────────────────────────────────────
        doc = Document()
        sec = doc.sections[0]
        sec.page_width        = Twips(11906)
        sec.page_height       = Twips(16838)
        sec.top_margin        = Twips(1440)
        sec.bottom_margin     = Twips(1440)
        sec.left_margin       = Twips(1800)
        sec.right_margin      = Twips(1800)
        sec.header_distance   = Twips(851)
        sec.footer_distance   = Twips(992)
        sectPr = sec._sectPr
        docGrid = OxmlElement("w:docGrid")
        docGrid.set(qn("w:type"), "lines")
        docGrid.set(qn("w:linePitch"), "312")
        sectPr.append(docGrid)

        # ── 2. Normal 样式：中文宋体/英数 Times New Roman 12pt，两端对齐，孤行控制关闭 ──
        normal_style = doc.styles["Normal"]
        normal_rpr = normal_style.element.get_or_add_rPr()
        for old in normal_rpr.findall(qn("w:rFonts")):
            normal_rpr.remove(old)
        rFonts_n = OxmlElement("w:rFonts")
        rFonts_n.set(qn("w:ascii"),    "Times New Roman")
        rFonts_n.set(qn("w:hAnsi"),    "Times New Roman")
        rFonts_n.set(qn("w:eastAsia"), "宋体")
        normal_rpr.insert(0, rFonts_n)
        sz_n = OxmlElement("w:sz"); sz_n.set(qn("w:val"), "24"); normal_rpr.append(sz_n)
        normal_ppr = normal_style.element.get_or_add_pPr()
        wc = OxmlElement("w:widowControl"); wc.set(qn("w:val"), "0"); normal_ppr.insert(0, wc)
        jc_n = OxmlElement("w:jc"); jc_n.set(qn("w:val"), "both"); normal_ppr.append(jc_n)
        sp_n = OxmlElement("w:spacing")
        sp_n.set(qn("w:line"), "360"); sp_n.set(qn("w:lineRule"), "auto")
        sp_n.set(qn("w:before"), "0"); sp_n.set(qn("w:after"), "0")
        normal_ppr.append(sp_n)

        # ── 3. XML 辅助函数 ───────────────────────────────────────────────
        def _set_para(p, jc=None, line=None, line_rule="auto",
                      before=None, after=None,
                      first_line=None, first_line_chars=None, left_ind=None,
                      keep_lines=False, page_break_before=False,
                      snap_to_grid=None):
            pPr = p._p.get_or_add_pPr()
            if keep_lines:
                pPr.append(OxmlElement("w:keepLines"))
            if page_break_before:
                pPr.append(OxmlElement("w:pageBreakBefore"))
            if snap_to_grid is not None:
                sg = OxmlElement("w:snapToGrid")
                sg.set(qn("w:val"), "1" if snap_to_grid else "0")
                pPr.append(sg)
            if before is not None or after is not None or line is not None:
                sp = OxmlElement("w:spacing")
                if before is not None: sp.set(qn("w:before"), str(before))
                if after  is not None: sp.set(qn("w:after"),  str(after))
                if line   is not None:
                    sp.set(qn("w:line"),     str(line))
                    sp.set(qn("w:lineRule"), line_rule)
                pPr.append(sp)
            if first_line is not None or first_line_chars is not None or left_ind is not None:
                ind = OxmlElement("w:ind")
                if left_ind          is not None: ind.set(qn("w:left"),           str(left_ind))
                if first_line_chars  is not None: ind.set(qn("w:firstLineChars"), str(first_line_chars))
                if first_line        is not None: ind.set(qn("w:firstLine"),      str(first_line))
                pPr.append(ind)
            if jc is not None:
                jc_el = OxmlElement("w:jc"); jc_el.set(qn("w:val"), jc); pPr.append(jc_el)

        def _set_run(r, ascii_font=None, eastasia_font=None, hint=None,
                     sz=None, sz_cs=None, bold=False, bcs=False, color=None):
            rPr = r._r.get_or_add_rPr()
            if ascii_font or eastasia_font or hint:
                rF = OxmlElement("w:rFonts")
                if ascii_font:    rF.set(qn("w:ascii"), ascii_font); rF.set(qn("w:hAnsi"), ascii_font)
                if eastasia_font: rF.set(qn("w:eastAsia"), eastasia_font)
                if hint:          rF.set(qn("w:hint"), hint)
                rPr.insert(0, rF)
            if bold: rPr.append(OxmlElement("w:b"))
            if bcs:  rPr.append(OxmlElement("w:bCs"))
            if color:
                col = OxmlElement("w:color"); col.set(qn("w:val"), color); rPr.append(col)
            if sz is not None:
                s = OxmlElement("w:sz"); s.set(qn("w:val"), str(sz)); rPr.append(s)
            if sz_cs is not None:
                sc = OxmlElement("w:szCs"); sc.set(qn("w:val"), str(sz_cs)); rPr.append(sc)

        def _signing_empty(n=1):
            """签署页专用空行：左对齐，snapToGrid=0，宋体，bCs，与参考文档一致。"""
            for _ in range(n):
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                sg = OxmlElement("w:snapToGrid"); sg.set(qn("w:val"), "0"); pPr.append(sg)
                sp = OxmlElement("w:spacing")
                sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
                sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "0")
                pPr.append(sp)
                jc_el = OxmlElement("w:jc"); jc_el.set(qn("w:val"), "left"); pPr.append(jc_el)
                mrPr = OxmlElement("w:rPr")
                rF2 = OxmlElement("w:rFonts")
                rF2.set(qn("w:ascii"), "Times New Roman"); rF2.set(qn("w:hAnsi"), "Times New Roman")
                mrPr.append(rF2)
                mrPr.append(OxmlElement("w:bCs"))
                szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "21"); mrPr.append(szCs)
                pPr.append(mrPr)

        def _cover_section_break():
            """
            在封面最后一段的 pPr 中嵌入 sectPr（封面 section 定义）。
            使用 vAlign=center 确保内容永远不溢出，并通过 nextPage 分隔目录。
            """
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            # 封面 section 的属性
            cover_sp = OxmlElement("w:sectPr")
            # 分节类型：下一页
            t = OxmlElement("w:type"); t.set(qn("w:val"), "nextPage"); cover_sp.append(t)
            # 页面尺寸（A4）
            pgSz = OxmlElement("w:pgSz")
            pgSz.set(qn("w:w"), "11906"); pgSz.set(qn("w:h"), "16838")
            cover_sp.append(pgSz)
            # 页边距（与正文一致）
            pgMar = OxmlElement("w:pgMar")
            pgMar.set(qn("w:top"),    "1440"); pgMar.set(qn("w:right"),  "1800")
            pgMar.set(qn("w:bottom"), "1440"); pgMar.set(qn("w:left"),   "1800")
            pgMar.set(qn("w:header"), "851");  pgMar.set(qn("w:footer"), "992")
            pgMar.set(qn("w:gutter"), "0")
            cover_sp.append(pgMar)
            # 文档网格
            dg = OxmlElement("w:docGrid")
            dg.set(qn("w:type"), "lines"); dg.set(qn("w:linePitch"), "312")
            cover_sp.append(dg)
            # 垂直居中——确保内容不管多少行都不溢出第二页
            vAlign = OxmlElement("w:vAlign"); vAlign.set(qn("w:val"), "center")
            cover_sp.append(vAlign)
            pPr.append(cover_sp)

        # ── 4. 正则分类 ──────────────────────────────────────────────────
        RE_TOC_ENTRY = re.compile(r"^(第[一二三四五六七八九十百]+部分\s+.+?)[\t ]+(\d+)$")
        RE_PART_HEAD = re.compile(r"^第[一二三四五六七八九十百]+部分\s+\S")
        RE_UPPER_NUM = re.compile(r"^[一二三四五六七八九十百]+、")
        RE_PAREN_CN  = re.compile(r"^（[一二三四五六七八九十百]+）")
        RE_ARAB_DOT  = re.compile(r"^\d+、")
        RE_PAREN_NUM = re.compile(r"^（\d+）")
        RE_MD_TABLE_ROW = re.compile(r"^\|.*\|\s*$")
        RE_MD_TABLE_DIVIDER = re.compile(r"^\|\s*:?-{3,}:?(?:\|\s*:?-{3,}:?)+\|\s*$")

        # ── 书签计数器与锚点辅助 ────────────────────────────────────────────────
        _bm_id = [0]  # 可变单元格，允许在内层代码中递增

        def _part_anchor(heading_s: str) -> str:
            """从"第X部分..."中提取 ASCII 书签锚名，例如 'part_26'。"""
            m = re.match(r'^第([一二三四五六七八九十百]+)部分', heading_s)
            if not m:
                return ""
            part_no = self._cn_numeral_to_int(m.group(1))
            return f"part_{part_no}" if part_no is not None else ""

        # ── 5. 主处理循环 ─────────────────────────────────────────────────
        lines = contract_text.split("\n")
        phase          = "cover"
        cover_idx      = 0      # 非空封面行计数
        signing_started = False
        signing_idx    = 0      # 签署页内容行计数
        current_part   = 0      # 当前所在部分编号（用于判断前言是否加空行）

        for raw in lines:
            s = raw.strip()

            # 所有空行跳过（签署页内的空行由代码主动添加）
            if not s:
                continue

            # ════ 封面 → 目录 ════
            if phase == "cover" and re.match(r"^目\s*录$", s):
                phase = "toc"
                # 封面 section break（vAlign=center, nextPage）
                _cover_section_break()
                p = doc.add_paragraph()
                _set_para(p, jc="center", line=360, before=0, after=0)
                r = p.add_run("目    录")
                _set_run(r, sz=28, sz_cs=28, bcs=True)
                continue

            # ════ 封面内容：仅4行，vAlign=center 负责垂直定位 ════
            if phase == "cover":
                p = doc.add_paragraph()
                if cover_idx == 0:
                    # 合同标题：sz=48(24pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=48, sz_cs=48, bold=True, bcs=True)
                    # 标题与管理人之间加3个空行，制造视觉距离
                    for _ in range(3):
                        ep = doc.add_paragraph()
                        _set_para(ep, jc="center", line=360, before=0, after=0)
                elif s.startswith("基金管理人") or s.startswith("基金托管人"):
                    # 管理人/托管人：sz=36(18pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=36, sz_cs=36, bold=True, bcs=True)
                else:
                    # 日期：sz=36(18pt), bold, center
                    _set_para(p, jc="center", line=360, before=0, after=0)
                    r = p.add_run(s)
                    _set_run(r, sz=36, sz_cs=36, bold=True, bcs=True)
                cover_idx += 1
                continue

            # ════ 目录条目 ════
            if phase == "toc":
                m_toc = RE_TOC_ENTRY.match(s)
                if m_toc:
                    toc_text = m_toc.group(1).strip()
                    toc_page = m_toc.group(2).strip()
                    p = doc.add_paragraph()
                    _set_para(p, line=360, before=0, after=0)
                    pPr = p._p.get_or_add_pPr()
                    tabs_el = OxmlElement("w:tabs")
                    tab_el  = OxmlElement("w:tab")
                    tab_el.set(qn("w:val"),    "right")
                    tab_el.set(qn("w:leader"), "dot")
                    tab_el.set(qn("w:pos"),    "8296")
                    tabs_el.append(tab_el)
                    pPr.append(tabs_el)
                    anchor = _part_anchor(toc_text)
                    if anchor:
                        hlink = OxmlElement("w:hyperlink")
                        hlink.set(qn("w:anchor"), anchor)
                        hlink.set(qn("w:history"), "1")
                        for txt_val in ([toc_text] + (["\t" + toc_page] if toc_page else [])):
                            rel = OxmlElement("w:r")
                            rPr_h = OxmlElement("w:rPr")
                            # Explicit hyperlink style improves compatibility (e.g. WPS).
                            rStyle = OxmlElement("w:rStyle")
                            rStyle.set(qn("w:val"), "Hyperlink")
                            rPr_h.append(rStyle)
                            rF_h = OxmlElement("w:rFonts")
                            rF_h.set(qn("w:ascii"), "Times New Roman")
                            rF_h.set(qn("w:hAnsi"), "Times New Roman")
                            rF_h.set(qn("w:eastAsia"), "宋体")
                            rPr_h.append(rF_h)
                            sz_h = OxmlElement("w:sz")
                            sz_h.set(qn("w:val"), "24")
                            rPr_h.append(sz_h)
                            rel.append(rPr_h)
                            t_h = OxmlElement("w:t")
                            t_h.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                            t_h.text = txt_val
                            rel.append(t_h)
                            hlink.append(rel)
                        p._p.append(hlink)
                    else:
                        r1 = p.add_run(toc_text)
                        _set_run(r1, ascii_font="Times New Roman", eastasia_font="宋体", sz=24)
                        if toc_page:
                            r2 = p.add_run("\t" + toc_page)
                            _set_run(r2, ascii_font="Times New Roman", eastasia_font="宋体", sz=24)
                    continue
                else:
                    phase = "body"

            # ════ 签署页检测 ════
            if phase == "body" and not signing_started and "签署页" in s and "无正文" in s:
                signing_started = True
                p = doc.add_paragraph()
                _set_para(p, snap_to_grid=False, jc="left", line=360, before=0, after=0,
                          page_break_before=True)   # 签署页强制分页
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                _signing_empty(2)   # 签署页标题后 2 个空行（参考文档）
                signing_idx = 1
                continue

            # ════ 签署页内容行 ════
            if phase == "body" and signing_started:
                p = doc.add_paragraph()
                is_seal = "（盖章）" in s
                _set_para(p, snap_to_grid=False, jc="left", line=360, before=0, after=0)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         hint="eastAsia", sz=24, bcs=True)
                # 每一内容行后补充空行（对应参考文档签名留白）
                if "（盖章）" in s:
                    _signing_empty(4)   # 盖章行后 4 个空行（管理人/托管人签名区）
                elif "（签字或盖章）：" in s or "（签名）" in s:
                    _signing_empty(4)   # 法定代表人签名行后 4 个空行
                # 签订地点/日期行后不加空行
                signing_idx += 1
                continue

            # ════ 部分标题（第X部分…） ════
            if phase == "body" and RE_PART_HEAD.match(s):
                current_part += 1
                p = doc.add_paragraph(style="Heading 1")
                _set_para(p, jc="center", line=360, before=0, after=0,
                          keep_lines=True, page_break_before=True)
                # 添加书签，供目录超链接跳转
                anchor = _part_anchor(s)
                if anchor:
                    _bm_id[0] += 1
                    bm_start = OxmlElement("w:bookmarkStart")
                    bm_start.set(qn("w:id"), str(_bm_id[0]))
                    bm_start.set(qn("w:name"), anchor)
                    bm_end = OxmlElement("w:bookmarkEnd")
                    bm_end.set(qn("w:id"), str(_bm_id[0]))
                    p._p.append(bm_start)
                r = p.add_run(s)
                _set_run(r, ascii_font="Times New Roman", eastasia_font="宋体",
                         sz=30, sz_cs=30, bold=True, bcs=False, color="auto")
                if anchor:
                    p._p.append(bm_end)
                continue

            # ════ 一级子标题（一、二、…） ════
            if phase == "body" and RE_UPPER_NUM.match(s):
                # 第一部分（前言）子标题间不加空行，其余部分各子标题前空一行
                if current_part != 1:
                    ep = doc.add_paragraph()
                    _set_para(ep, line=360, before=0, after=0)
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480, first_line_chars=200)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 二级子标题（（一）（二）…） ════
            if phase == "body" and RE_PAREN_CN.match(s):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480, first_line_chars=200)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 数字条款（1、2、… 或（1）（2）…） ════
            if phase == "body" and (RE_ARAB_DOT.match(s) or RE_PAREN_NUM.match(s)):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480, first_line_chars=200)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)
                continue

            # ════ 普通正文 ════
            if phase in ("body", "toc"):
                p = doc.add_paragraph()
                _set_para(p, line=360, before=0, after=0,
                          first_line=480, first_line_chars=200)
                r = p.add_run(s)
                _set_run(r, hint="eastAsia", sz=24, bcs=True)

        # ── 6. 序列化 ──────────────────────────────────────────────────────
        _finalize_doc_page_numbers(doc, OxmlElement, qn)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

# ═══════════════════════════════════════════════════════════════════════════════
#  ProspectusEngine — 招募说明书 8步处理管线
# ═══════════════════════════════════════════════════════════════════════════════
class ProspectusEngine:
    def __init__(self):
        with open(PROSPECTUS_CLAUSES_JSON, encoding="utf-8") as f:
            self.pro_clauses = json.load(f)["clauses"]
        self._reference_fixed_cache = {}

    # ── Step 1: 派生变量（委托给 ContractEngine）────────────────────────────
    def _derive_variables(self, v: dict) -> dict:
        v = engine._derive_variables(v)
        market_scope = str(v.get("MARKET_SCOPE", "") or "").strip().upper()
        if market_scope not in {"SINGLE_MARKET", "CROSS_MARKET"}:
            market_type = str(v.get("MARKET_TYPE", "") or "").strip().upper()
            if market_type in {"KECHUANG", "CHUANGYE"}:
                market_scope = "SINGLE_MARKET"
            else:
                market_scope = "CROSS_MARKET"
        v["MARKET_SCOPE"] = market_scope
        product_type = str(v.get("PRODUCT_TYPE", "") or "").strip().upper()
        if not product_type:
            product_type = "ETF"
        v["PRODUCT_TYPE"] = product_type
        return v

    def _get_prospectus_variant_key(self, v: dict) -> str:
        v = self._derive_variables(v)
        exchange = v.get("EXCHANGE", "")
        has_hk = bool(v.get("HAS_HK_CONNECT"))
        market_scope = v.get("MARKET_SCOPE", "CROSS_MARKET")

        if has_hk:
            return "SSE_HK" if exchange == "SSE" else "SZSE_HK"
        if exchange == "SSE":
            return "SSE_SINGLE" if market_scope == "SINGLE_MARKET" else "SSE_CROSS"
        return "SZSE_SINGLE" if market_scope == "SINGLE_MARKET" else "SZSE_CROSS"

    def _get_variant_clause_bundle(self, v: dict) -> dict:
        variants = self.pro_clauses.get("PROSPECTUS_VARIANTS", {}).get("variants", {})
        key = self._get_prospectus_variant_key(v)
        return variants.get(key, variants.get("SSE_CROSS", {}))


    @staticmethod
    def _get_product_type(v: dict) -> str:
        product_type = str((v or {}).get("PRODUCT_TYPE", "") or "").strip().upper()
        return product_type or "ETF"

    @staticmethod
    @staticmethod
    def _strip_signing_page_from_contract_summary(summary_text: str) -> str:
        text = (summary_text or "").strip()
        if not text:
            return ""
        markers = [
            "签署页",
            "本页无正文",
            "（盖章）",
            "（签字或盖章）",
            "（签名）",
        ]
        positions = [pos for pos in (text.find(marker) for marker in markers) if pos >= 0]
        signing_patterns = [
            r"（?本页为《[^》]+基金合同》签署页[^）]*）?",
            r"本页为《[^》]+基金合同》",
        ]
        for pattern in signing_patterns:
            m = re.search(pattern, text)
            if m:
                positions.append(m.start())
        if positions:
            text = text[:min(positions)].rstrip()
        return text

    # ── Step 2: 注入差异条款（合同条款 + 招募说明书专有条款）──────────────
    def _inject_clause_texts(self, v: dict) -> dict:
        v = self._derive_variables(v)
        v = engine._inject_clause_texts(v)
        has_hk = v.get("HAS_HK_CONNECT", False)
        exch_cn = v.get("EXCHANGE_NAME_CN", "证券交易所")
        variant_key = self._get_prospectus_variant_key(v)
        variant_bundle = self._get_variant_clause_bundle(v)

        vt_key = "HK_CONNECT" if has_hk else "STANDARD"
        vt_variants = self.pro_clauses["VALUATION_TIMING"]["variants"]
        valuation_text = vt_variants.get(vt_key, vt_variants["STANDARD"])["text"]
        valuation_text = valuation_text.replace("{EXCHANGE_NAME_CN}", exch_cn)
        v["VALUATION_TIMING_CLAUSE"] = valuation_text

        v["RISK_DISCLOSURE_CHUANGYE"] = self.pro_clauses["RISK_DISCLOSURE_CHUANGYE"]["variants"]["DEFAULT"]["text"]
        v["RISK_DISCLOSURE_KECHUANG"] = self.pro_clauses["RISK_DISCLOSURE_KECHUANG"]["variants"]["DEFAULT"]["text"]
        v["RISK_DISCLOSURE_HK_CONNECT"] = self.pro_clauses["RISK_DISCLOSURE_HK_CONNECT"]["variants"]["DEFAULT"]["text"]

        custodian_name = v.get("CUSTODIAN_NAME", "")
        custodian_contacts = self.pro_clauses.get("CUSTODIAN_INFO_PROSPECTUS", {}).get("custodians", {})
        info = custodian_contacts.get(custodian_name, {})
        v.setdefault("CUSTODIAN_DEPT", info.get("dept", "[待填写：托管部门名称]"))
        v.setdefault("CUSTODIAN_PHONE", info.get("phone", "[待填写：服务电话]"))
        v.setdefault("CUSTODIAN_WEBSITE", info.get("website", "[待填写：网址]"))

        v.setdefault("FUND_MANAGER_NAME", "[待填写：基金经理姓名]")
        v.setdefault("FUND_MANAGER_BIO", "[待填写：基金经理简介（学历、从业经历、任职日期等）]")
        v.setdefault("CSRC_APPROVAL_NO", "202X年X月X日证监许可〔202X〕XXXX号")
        v.setdefault("MIN_SUB_UNIT", "1,000,000份（即100万份）")

        chapter6 = variant_bundle.get("chapter_6", {})
        chapter10 = variant_bundle.get("chapter_10", {})
        v["PROSPECTUS_VARIANT_KEY"] = variant_key
        v["PROSPECTUS_CH6_SEC4"] = chapter6.get("section_4", "")
        v["PROSPECTUS_CH6_SEC7"] = chapter6.get("section_7", "")
        v["PROSPECTUS_CH6_SEC11"] = chapter6.get("section_11", "")
        v["PROSPECTUS_CH6_SEC12"] = chapter6.get("section_12", "")
        v["PROSPECTUS_CH6_SEC13"] = chapter6.get("section_13", "")
        v["PROSPECTUS_CH7_BODY"] = variant_bundle.get("chapter_7", "")
        v["PROSPECTUS_CH9_BODY"] = variant_bundle.get("chapter_9", "")
        v["PROSPECTUS_CH10_PRELUDE"] = chapter10.get("prelude", "")
        v["PROSPECTUS_CH10_SEC4"] = chapter10.get("section_4", "")
        v["PROSPECTUS_CH10_SEC7"] = chapter10.get("section_7", "")
        v["PROSPECTUS_CH21_TITLES"] = self.pro_clauses.get("CHAPTER21_TITLES", {}).get("text", "")

        risk_bodies = self.pro_clauses.get("RISK_CHAPTER_BODIES", {}).get("variants", {})
        if has_hk:
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("HK_CONNECT", "")
        elif v.get("IS_KECHUANG"):
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("KECHUANG", "")
        elif v.get("IS_CHUANGYE"):
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("CHUANGYE", "")
        else:
            v["PROSPECTUS_CH18_BODY"] = risk_bodies.get("STANDARD_A", "")

        return v

    # ── Step 3: 从合同全文提取各关键部分（内容摘要 + 各章节来源段落）───────
    def _extract_contract_sections(self, v: dict) -> dict:
        """
        生成完整合同文本，按以下逻辑提取各部分供招募说明书使用：

        #二章 释义 → 《基金合同》第二部分
        #十章 申购赎回 子条款 → 《基金合同》第八部分 对应子条款：
          二、开放日 → CONTRACT_PART8_SEC2
          三、原则   → CONTRACT_PART8_SEC3
          六、对价   → CONTRACT_PART8_SEC6
          七、拒绝申购 → CONTRACT_PART8_SEC7
          八、暂停赎回 → CONTRACT_PART8_SEC8
          九、其他方式 → CONTRACT_PART8_SEC9
          十、非交易过户 → CONTRACT_PART8_SEC10
          十一、冻结 → CONTRACT_PART8_SEC11
          十二、转让 → CONTRACT_PART8_SEC12
          十三、其他业务 → CONTRACT_PART8_SEC13
          十四、清算交收模式 → CONTRACT_PART8_SEC14
        #十一章 投资 → 《基金合同》第十四部分
        #十二章 财产 → 《基金合同》第十五部分
        #十三章 估值 → 《基金合同》第十六部分
        #十四章 收益分配 → 三、基金收益分配原则 同 《基金合同》三
        #十五章 费用 → 《基金合同》第十七部分
        #十六章 会计审计 → 《基金合同》第十九部分
        #十七章 信息披露 → 《基金合同》第二十部分
        #十九章 变更终止 → 《基金合同》第二十一部分
        #二十章 合同摘要 → 《基金合同》第二十六部分
        """
        try:
            contract_text = engine.generate(v)
        except Exception as _exc:
            import logging
            logging.warning("ProspectusEngine: 合同生成失败，部分章节将使用占位符。原因：%s", _exc)
            contract_text = ""

        # 按"第X部分"标题切分合同全文
        RE_PART = re.compile(r'^第[一二三四五六七八九十百]+部分\s+\S[^\n]*', re.MULTILINE)
        part_iter = list(RE_PART.finditer(contract_text))

        # 中文数字顺序表，用于精确匹配部分编号（避免"第二部分"匹配"第二十部分"）
        _CN_ORDER = [
            "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
            "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九",
            "二十", "二十一", "二十二", "二十三", "二十四", "二十五", "二十六",
        ]

        def _get_part(cn_num: str) -> str:
            """
            按准确中文部分编号提取该部分全文（标题行之后的内容）。
            cn_num 示例：'第二部分'、'第八部分'、'第二十六部分'
            使用精确匹配避免"第二部分"匹配到"第二十部分"。
            """
            # 构造精确匹配模式：编号后紧跟"部分"且后面为空白
            pattern = re.compile(
                r'^' + re.escape(cn_num) + r'\s+\S[^\n]*',
                re.MULTILINE
            )
            matched_idx = []
            for i, m in enumerate(part_iter):
                if pattern.match(m.group()):
                    matched_idx.append(i)
            if matched_idx:
                # Contract text includes TOC + body with duplicate "第X部分" headings.
                # Use the last match to target body section, not TOC line.
                i = matched_idx[-1]
                m = part_iter[i]
                start = m.end()
                end = part_iter[i + 1].start() if i + 1 < len(part_iter) else len(contract_text)
                return contract_text[start:end].strip()
            return ""

        def _get_part_subsection(part_cn: str, cn_num: str) -> str:
            """
            从指定部分提取指定中文序号的子条款全文。
            cn_num 示例：'九', '十', '十一', ...
            使用精确匹配：标题为 "序号、" 格式，用 word boundary 避免"十"匹配"十一"。
            返回该子条款标题行（含序号）及其正文，直到下一个同级序号或部分结束。
            """
            part_text = _get_part(part_cn)
            if not part_text:
                return ""
            # 精确匹配行首的中文序号：序号后紧跟"、"，序号本身是完整单词
            # 构造列表后按行首定位，避免多字符序号前缀混淆
            RE_CN = re.compile(
                r'^([一二三四五六七八九十百]+)、',
                re.MULTILINE
            )
            markers = list(RE_CN.finditer(part_text))
            for i, m in enumerate(markers):
                heading = m.group(1)          # 捕获组1 = 纯序号，不含"、"
                if heading == cn_num:
                    start = m.start()
                    end = markers[i + 1].start() if i + 1 < len(markers) else len(part_text)
                    return part_text[start:end].strip()
            return ""

        def _get_part8_subsection(cn_num: str) -> str:
            return _get_part_subsection("第八部分", cn_num)

        # #二章 释义（完整第二部分，不含标题行）
        v.setdefault("CONTRACT_DEFS_TEXT", _get_part("第二部分") or
                     "[待填写：释义内容，请从基金合同第二部分复制]")

        # #十章 申购赎回 各子条款（来自第八部分）
        # 二/三/六 保持固定（模板中已有固定文本），此处仍提取供模板引用
        v.setdefault("CONTRACT_PART8_SEC2", _get_part8_subsection("二") or
                     "[待填写：申购和赎回的开放日及时间，来自基金合同第八部分]")
        v.setdefault("CONTRACT_PART8_SEC3", _get_part8_subsection("三") or
                     "[待填写：申购和赎回的原则，来自基金合同第八部分]")
        v.setdefault("CONTRACT_PART8_SEC6", _get_part8_subsection("六") or
                     "[待填写：申购和赎回的对价、费用及其用途，来自基金合同第八部分]")
        # 七 = 拒绝暂停申购（合同 七，招募 八）
        v.setdefault("CONTRACT_PART8_SEC7", _get_part8_subsection("七") or
                     "[待填写：拒绝或暂停申购的情形，来自基金合同第八部分]")
        # 八 = 暂停赎回（合同 八，招募 九）
        v.setdefault("CONTRACT_PART8_SEC8", _get_part8_subsection("八") or
                     "[待填写：暂停赎回或延缓支付赎回对价的情形，来自基金合同第八部分]")
        # 招募 十 → 合同八.九
        v.setdefault("CONTRACT_PART8_SEC9", _get_part8_subsection("九") or
                     "[待填写：其他申购赎回方式，来自基金合同第八部分]")
        # 招募 十一 → 合同八.十
        v.setdefault("CONTRACT_PART8_SEC10", _get_part8_subsection("十") or
                     "[待填写：基金的非交易过户，来自基金合同第八部分]")
        # 招募 十二 → 合同八.十一
        v.setdefault("CONTRACT_PART8_SEC11", _get_part8_subsection("十一") or
                     "[待填写：基金份额的冻结和解冻，来自基金合同第八部分]")
        # 招募 十三 → 合同八.十二
        v.setdefault("CONTRACT_PART8_SEC12", _get_part8_subsection("十二") or
                     "[待填写：基金份额的转让，来自基金合同第八部分]")
        # 招募 十四 → 合同八.十三
        v.setdefault("CONTRACT_PART8_SEC13", _get_part8_subsection("十三") or
                     "[待填写：其他业务，来自基金合同第八部分]")
        # 招募 十五 → 合同八.十四（交易所相关，文字含交易所名称）
        v.setdefault("CONTRACT_PART8_SEC14", _get_part8_subsection("十四") or
                     "[待填写：清算交收模式，来自基金合同第八部分]")

        # #十一-十七、十九章 各对应合同部分
        v.setdefault("CONTRACT_INVEST_TEXT", _get_part("第十四部分") or
                     "[待填写：基金的投资，来自基金合同第十四部分]")
        v.setdefault("CONTRACT_ASSET_TEXT", _get_part("第十五部分") or
                     "[待填写：基金的财产，来自基金合同第十五部分]")
        v.setdefault("CONTRACT_VALUATION_TEXT", _get_part("第十六部分") or
                     "[待填写：基金资产估值，来自基金合同第十六部分]")
        v.setdefault("CONTRACT_FEE_TEXT", _get_part("第十七部分") or
                     "[待填写：基金的费用与税收，来自基金合同第十七部分]")
        v.setdefault("CONTRACT_AUDIT_TEXT", _get_part("第十九部分") or
                     "[待填写：基金的会计与审计，来自基金合同第十九部分]")
        v.setdefault("CONTRACT_DISCLOSURE_TEXT", _get_part("第二十部分") or
                     "[待填写：基金的信息披露，来自基金合同第二十部分]")
        v.setdefault("CONTRACT_TERMINATION_TEXT", _get_part("第二十一部分") or
                     "[待填写：基金合同的变更、终止与基金财产的清算，来自基金合同第二十一部分]")

        # #二十章 合同内容摘要 → 第二十六部分
        summary = self._strip_signing_page_from_contract_summary(_get_part("第二十六部分"))
        v.setdefault("CONTRACT_SUMMARY_TEXT", summary or
                     "[待填写：基金合同内容摘要，请先生成基金合同，从第二十六部分复制此处]")

        # #??? ???????? ? ???????????
        v.setdefault("CONTRACT_DISTRIBUTION_TEXT", _get_part("\u7b2c\u5341\u516b\u90e8\u5206") or
                     "[\u5f85\u586b\u5199\uff1a\u57fa\u91d1\u7684\u6536\u76ca\u4e0e\u5206\u914d\uff0c\u6765\u81ea\u57fa\u91d1\u5408\u540c\u7b2c\u5341\u516b\u90e8\u5206]")
        # #??? ?????????? ? ??????????? ?
        v.setdefault("CONTRACT_PART18_SEC3", _get_part_subsection("第十八部分", "三") or
                     "[待填写：基金收益分配原则，来自基金合同第十八部分第三项]")

        return v

    # ── Step 3 (兼容旧接口) ────────────────────────────────────────────────
    def _extract_contract_summary(self, v: dict) -> dict:
        """向后兼容旧方法名，实际调用 _extract_contract_sections。"""
        return self._extract_contract_sections(v)

    @staticmethod
    def _chapter_num_to_cn(num: int) -> str:
        cn_map = {
            1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
            11: "十一", 12: "十二", 13: "十三", 14: "十四", 15: "十五", 16: "十六", 17: "十七", 18: "十八",
            19: "十九", 20: "二十", 21: "二十一", 22: "二十二", 23: "二十三", 24: "二十四", 25: "二十五",
        }
        return cn_map.get(num, "")

    @staticmethod
    def _split_top_sections(chapter_body: str) -> dict:
        section_re = re.compile(r"^([一二三四五六七八九十百]+)、[^\n]*", re.MULTILINE)
        matches = list(section_re.finditer(chapter_body or ""))
        sections = {}
        for i, m in enumerate(matches):
            sec_cn = m.group(1)
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(chapter_body)
            sections[sec_cn] = chapter_body[start:end].strip()
        return sections

    def _get_reference_prospectus_docx(self, v: dict | None = None) -> Path:
        variant_key = self._get_prospectus_variant_key(v or {}) if v is not None else "SSE_CROSS"
        return REFERENCE_PROSPECTUS_DOCX_MAP.get(variant_key, REFERENCE_PROSPECTUS_DOCX_MAP["SSE_CROSS"])

    def _load_reference_fixed_content(self, v: dict | None = None) -> dict:
        """
        Load fixed chapter/section text from the reference prospectus DOCX for the current variant.
        """
        variant_key = self._get_prospectus_variant_key(v or {}) if v is not None else "SSE_CROSS"
        if variant_key in self._reference_fixed_cache:
            return self._reference_fixed_cache[variant_key]

        data = {"important_notice": ""}
        reference_docx = self._get_reference_prospectus_docx(v)
        if not reference_docx.exists():
            self._reference_fixed_cache[variant_key] = data
            return data

        try:
            from docx import Document
            from docx.document import Document as DocumentObject
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table
            from docx.text.paragraph import Paragraph
            doc = Document(str(reference_docx))
        except Exception:
            self._reference_fixed_cache[variant_key] = data
            return data

        def iter_blocks(parent):
            if isinstance(parent, DocumentObject):
                parent_elm = parent.element.body
                parent_obj = parent
            else:
                parent_elm = parent._tc
                parent_obj = parent
            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield ("paragraph", Paragraph(child, parent_obj))
                elif isinstance(child, CT_Tbl):
                    yield ("table", Table(child, parent_obj))

        def table_to_lines(table) -> list[str]:
            rows = []
            for row in table.rows:
                values = [cell.text.replace("\r", "\n").strip() for cell in row.cells]
                if any(values):
                    rows.append(values)
            if not rows:
                return []
            return [" | ".join(value for value in row if value) for row in rows]

        blocks = []
        for block_type, block in iter_blocks(doc):
            if block_type == "paragraph":
                text_value = (block.text or "").strip()
                if text_value:
                    blocks.append(("paragraph", text_value))
            else:
                lines = table_to_lines(block)
                if lines:
                    blocks.append(("table", "\n".join(lines)))

        notice_started = False
        notice_lines = []
        for _, block_text in blocks:
            stripped = block_text.strip()
            if not stripped:
                continue
            if stripped == "\u91cd\u8981\u63d0\u793a":
                notice_started = True
                notice_lines = [stripped]
                continue
            if notice_started and stripped == "\u76ee\u5f55":
                break
            if notice_started:
                notice_lines.extend(line for line in stripped.splitlines() if line.strip())
        if notice_lines:
            data["important_notice"] = "\n".join(notice_lines).strip()

        chapter_heading_re = re.compile(r'^\u7b2c([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e]+)\u7ae0')
        current_chapter = None
        current_lines = []
        current_title = ""

        def flush_chapter():
            nonlocal current_chapter, current_lines, current_title
            if not current_chapter:
                return
            body = "\n".join(line for line in current_lines if line.strip()).strip()
            if body:
                data[current_chapter] = {
                    "title": current_title,
                    "body": body,
                    "sections": self._split_top_sections(body),
                }
            current_chapter = None
            current_lines = []
            current_title = ""

        for _, block_text in blocks:
            stripped = block_text.strip()
            if not stripped:
                continue
            match = chapter_heading_re.match(stripped)
            if match:
                flush_chapter()
                current_chapter = match.group(1)
                current_title = stripped
                continue
            if current_chapter:
                current_lines.extend(line for line in stripped.splitlines() if line.strip())
        flush_chapter()

        self._reference_fixed_cache[variant_key] = data
        return data

    @staticmethod
    def _replace_numbered_item_in_section(chapter_body: str, section_cn: str, item_no: str, new_item: str) -> str:
        body = (chapter_body or "").strip()
        if not body:
            return chapter_body

        sec_re = re.compile(rf"^{section_cn}、[^\n]*", re.MULTILINE)
        next_sec_re = re.compile(r"^[一二三四五六七八九十百]+、[^\n]*", re.MULTILINE)
        m = sec_re.search(body)
        if not m:
            return chapter_body
        n = next_sec_re.search(body, m.end())
        sec_end = n.start() if n else len(body)

        sec_block = body[m.start():sec_end]
        item_text = (new_item or "").strip()
        if not item_text:
            return chapter_body
        if not re.match(rf"^{re.escape(item_no)}、", item_text):
            item_text = f"{item_no}、{item_text}"

        item_re = re.compile(rf"^{re.escape(item_no)}、[^\n]*", re.MULTILINE)
        next_item_re = re.compile(r"^\d+、[^\n]*", re.MULTILINE)
        im = item_re.search(sec_block)
        if not im:
            if not sec_block.endswith("\n"):
                sec_block += "\n"
            sec_block = f"{sec_block}{item_text}\n"
        else:
            inn = next_item_re.search(sec_block, im.end())
            item_end = inn.start() if inn else len(sec_block)
            sec_block = f"{sec_block[:im.start()]}{item_text}\n{sec_block[item_end:].lstrip()}"

        new_body = f"{body[:m.start()]}{sec_block}{body[sec_end:]}"
        return new_body

    def _apply_reference_fixed_content(self, text: str, v: dict) -> str:
        """
        Apply canonical fixed text from red-dividend prospectus docx.
        """
        ref = self._load_reference_fixed_content(v)
        if not ref:
            return text

        def ref_chapter(chap_cn: str) -> str:
            return ref.get(chap_cn, {}).get("body", "")

        def ref_section(chap_cn: str, sec_cn: str) -> str:
            return ref.get(chap_cn, {}).get("sections", {}).get(sec_cn, "")

        ch3 = ref_chapter("三")
        if ch3:
            manager_name = str(v.get("FUND_MANAGER_NAME") or "[待填写：基金经理姓名]").strip()
            manager_bio = str(v.get("FUND_MANAGER_BIO") or "[待填写：基金经理简介]").strip()
            sec_item3 = f"3、本基金的基金经理为{manager_name}。"
            if manager_bio:
                sec_item3 = f"{sec_item3}\n{manager_bio}"
            ch3 = self._replace_numbered_item_in_section(ch3, "二", "3", sec_item3)
            text = self._replace_chapter_body(text, "三", ch3)

        text = self._replace_chapter_body(text, "四", "【托管人情况待填写】")

        text = self._replace_subsection_in_chapter(text, "五", "一", ref_section("五", "一"))
        text = self._replace_subsection_in_chapter(text, "五", "二", ref_section("五", "二"))
        sec3 = ref_section("五", "三")
        if sec3:
            sec3 = re.sub(r"^经办律师[：:].*$", "经办律师：丁媛、李晓露", sec3, flags=re.MULTILINE)
        text = self._replace_subsection_in_chapter(text, "五", "三", sec3)
        text = self._replace_subsection_in_chapter(text, "五", "四", "四、审计基金财产的会计师事务所\n【待填写】")

        text = self._replace_chapter_body(text, "八", ref_chapter("八"))
        text = self._replace_chapter_body(text, "十四", ref_chapter("十四"))

        for chap_cn in ("二十二", "二十三", "二十四"):
            text = self._replace_chapter_body(text, chap_cn, ref_chapter(chap_cn))

        return text

    @staticmethod
    def _find_chapter_span(text: str, chapter_cn: str):
        """Return (start, heading_end, end) for a chapter, or None."""
        chap_re = re.compile(rf"^第{chapter_cn}章[^\n]*", re.MULTILINE)
        any_chap_re = re.compile(r"^第[一二三四五六七八九十百]+章[^\n]*", re.MULTILINE)
        matches = list(chap_re.finditer(text))
        if not matches:
            return None
        # Template contains a TOC copy and a body copy; use the last one (body section).
        m = matches[-1]
        n = any_chap_re.search(text, m.end())
        end = n.start() if n else len(text)
        return m.start(), m.end(), end

    def _replace_chapter_body(self, text: str, chapter_cn: str, new_body: str) -> str:
        new_body = (new_body or "").strip()
        if not new_body:
            return text
        span = self._find_chapter_span(text, chapter_cn)
        if not span:
            return text
        start, heading_end, end = span
        prefix = text[:heading_end].rstrip("\n")
        suffix = text[end:].lstrip("\n")
        if suffix:
            return f"{prefix}\n{new_body}\n{suffix}"
        return f"{prefix}\n{new_body}\n"

    def _replace_subsection_in_chapter(self, text: str, chapter_cn: str, subsection_cn: str, new_subsection: str) -> str:
        new_subsection = (new_subsection or "").strip()
        if not new_subsection:
            return text
        span = self._find_chapter_span(text, chapter_cn)
        if not span:
            return text
        start, _heading_end, end = span
        block = text[start:end]

        sec_re = re.compile(rf"^{subsection_cn}、[^\n]*", re.MULTILINE)
        next_sec_re = re.compile(r"^[一二三四五六七八九十百]+、[^\n]*", re.MULTILINE)
        m = sec_re.search(block)
        if not m:
            # If subsection does not exist in template, append it at chapter tail.
            if not block.endswith("\n"):
                block += "\n"
            block = f"{block}{new_subsection}\n"
        else:
            n = next_sec_re.search(block, m.end())
            sec_end = n.start() if n else len(block)
            block = f"{block[:m.start()]}{new_subsection}\n{block[sec_end:].lstrip()}"

        return f"{text[:start]}{block}{text[end:]}"

    @staticmethod
    def _retag_subsection_number(subsection_text: str, new_cn: str) -> str:
        txt = (subsection_text or "").strip()
        return re.sub(r"^[一二三四五六七八九十百]+、", f"{new_cn}、", txt, count=1)

    @staticmethod
    def _ensure_subsection_heading(subsection_text: str, cn_num: str) -> str:
        txt = (subsection_text or "").strip()
        if not txt:
            return ""
        if re.match(r"^[一二三四五六七八九十百]+、", txt):
            return txt
        return f"{cn_num}、\n{txt}"

    @staticmethod
    def _join_nonempty_blocks(blocks) -> str:
        return "\n\n".join(block.strip() for block in blocks if (block or "").strip())

    @staticmethod
    def _normalize_reused_prospectus_chapter(chapter_text: str) -> str:
        text = (chapter_text or "").strip()
        if not text:
            return ""
        text = text.replace("详见招募说明书的规定", "详见招募说明书“侧袋机制”部分的规定")
        text = text.replace("本基金合同", "基金合同")
        return text

    @staticmethod
    def _get_prospectus_min_sub_unit(v: dict) -> str:
        return str(v.get("MIN_SUB_UNIT") or "1,000,000份（即100万份）").strip()

    def _normalize_prospectus_risk_chapter(self, chapter_text: str, v: dict) -> str:
        text = (chapter_text or "").strip()
        if not text:
            return ""
        min_sub_unit = self._get_prospectus_min_sub_unit(v)
        for old in ("1,000,000份（即100万份）", "1,000,000份"):
            text = text.replace(old, min_sub_unit)
        text = re.sub(r"(?<=按原)\d[\d,]*份(?:（即\d+万份）)?", min_sub_unit, text)
        text = re.sub(r"(?<=新的)\d[\d,]*份(?:（即\d+万份）)?", min_sub_unit, text)
        return text

    def _build_chapter_six_body(self, v: dict, ref: dict) -> str:
        def ref_section(sec_cn: str) -> str:
            return ref.get("六", {}).get("sections", {}).get(sec_cn, "")

        sec8 = """八、认购费用
认购费用由投资人承担，不高于0.30%，认购费率如下表所示：

|   |   |
|---|---|
|认购份额（S）|认购费率|
|S＜100万份|0.30%|
|S≥100万份|每笔500元|

基金管理人办理网下现金认购和网下股票认购不收取认购费。发售代理机构办理网上现金认购、网下现金认购、网下股票认购时可参照上述费率结构，按照不高于0.3%的标准收取一定的佣金。投资人申请重复现金认购的，须按每次认购所对应的费率档次分别计费。"""

        blocks = [
            f"本基金由基金管理人依照《基金法》、《运作办法》、《销售办法》、基金合同及其他有关规定，并经中国证监会{v.get('CSRC_APPROVAL_NO', '202X年X月X日证监许可〔202X〕XXXX号')}文注册募集。",
            "本基金为交易型开放式基金，股票型基金，基金存续期限为不定期。",
            ref_section("一"),
            ref_section("二"),
            ref_section("三"),
            v.get("PROSPECTUS_CH6_SEC4", ""),
            ref_section("五"),
            ref_section("六"),
            v.get("PROSPECTUS_CH6_SEC7", ""),
            sec8,
            ref_section("九"),
            ref_section("十"),
            ref_section("??") or v.get("PROSPECTUS_CH6_SEC11", ""),
            ref_section("??") or v.get("PROSPECTUS_CH6_SEC12", ""),
            ref_section("??") or v.get("PROSPECTUS_CH6_SEC13", ""),
        ]
        return self._join_nonempty_blocks(blocks)

    @staticmethod
    def _build_chapter_ten_limits_table(min_sub_unit: str) -> str:
        return "\n".join([
            "|项目|内容|",
            "|---|---|",
            f"|最小申购赎回单位|{min_sub_unit}|",
            "|申购/赎回份额上限|以申购赎回清单或相关公告为准|",
        ])

    def _build_chapter_ten_sec5(self, ref: dict, v: dict) -> str:
        sec5 = ref.get("\u5341", {}).get("sections", {}).get("\u4e94", "")
        min_sub_unit = self._get_prospectus_min_sub_unit(v)
        if sec5:
            sec5 = re.sub(r"\u76ee\u524d\uff0c\u672c\u57fa\u91d1\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d\u4e3a[^\uff0c\u3002\uff1b]+", f"\u76ee\u524d\uff0c\u672c\u57fa\u91d1\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d\u4e3a{min_sub_unit}", sec5, count=1)
            if min_sub_unit not in sec5:
                sec5 = re.sub(r"\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d[^\u3002]*", f"\u6700\u5c0f\u7533\u8d2d\u8d4e\u56de\u5355\u4f4d\u4e3a{min_sub_unit}", sec5, count=1)
            return sec5
        return f"""五、申购和赎回的数额限制
1、投资人申购、赎回的基金份额需为最小申购赎回单位的整数倍。目前，本基金最小申购赎回单位为{min_sub_unit}，基金管理人有权对其进行调整，并在调整实施前依照《信息披露办法》的有关规定在规定媒介上公告。
2、基金管理人可以规定本基金当日申购份额及当日赎回份额上限，具体规定请参见申购赎回清单或相关公告。
3、基金管理人可以根据市场情况，在法律法规允许的情况下，合理调整上述申购和赎回的数量或比例限制，并在实施前依照《信息披露办法》的有关规定在规定媒介上公告。
4、当接受申购申请对存量基金份额持有人利益构成潜在重大不利影响时，基金管理人应当采取设定单一投资者申购份额上限或基金单日净申购比例上限、拒绝大额申购、暂停基金申购等措施。"""

    def _build_contract_section(self, v: dict, var_name: str, sec_cn: str) -> str:
        sec_text = self._retag_subsection_number(v.get(var_name, ""), sec_cn)
        return self._ensure_subsection_heading(sec_text, sec_cn)

    def _ensure_distribution_conditions_section(self, chapter_body: str) -> str:
        if not chapter_body or re.search(r"(?m)^\u56db\u3001\u6536\u76ca\u5206\u914d\u6761\u4ef6(?:\n|$)", chapter_body):
            return chapter_body
        if not re.search(r"(?m)^\u56db\u3001\u6536\u76ca\u5206\u914d\u65b9\u6848(?:\n|$)", chapter_body):
            return chapter_body
        conditions = (
            "\u56db\u3001\u6536\u76ca\u5206\u914d\u6761\u4ef6\n"
            "1\u3001\u57fa\u91d1\u7ba1\u7406\u4eba\u53ef\u6bcf\u6708\u5bf9\u57fa\u91d1\u76f8\u5bf9\u4e1a\u7ee9\u6bd4\u8f83\u57fa\u51c6\u7684\u8d85\u989d\u6536\u76ca\u7387\u4ee5\u53ca\u57fa\u91d1\u7684\u53ef\u4f9b\u5206\u914d\u5229\u6da6\u8fdb\u884c\u8bc4\u4f30\uff0c\u5728\u7b26\u5408\u57fa\u91d1\u6536\u76ca\u5206\u914d\u6761\u4ef6\u4e0b\uff0c\u53ef\u5b89\u6392\u6536\u76ca\u5206\u914d\u3002\n"
            "2\u3001\u57fa\u91d1\u6536\u76ca\u5206\u914d\u6761\u4ef6\u3001\u8bc4\u4f30\u65f6\u95f4\u3001\u5206\u914d\u65f6\u95f4\u3001\u5206\u914d\u65b9\u6848\u53ca\u6bcf\u6b21\u57fa\u91d1\u6536\u76ca\u5206\u914d\u6570\u989d\u7b49\u5185\u5bb9\uff0c\u57fa\u91d1\u7ba1\u7406\u4eba\u53ef\u4ee5\u6839\u636e\u5b9e\u9645\u60c5\u51b5\u786e\u5b9a\u5e76\u6309\u7167\u6709\u5173\u89c4\u5b9a\u516c\u544a\u3002"
        )
        return re.sub(r"(?m)^\u56db\u3001\u6536\u76ca\u5206\u914d\u65b9\u6848", f"{conditions}\n\u4e94\u3001\u6536\u76ca\u5206\u914d\u65b9\u6848", chapter_body, count=1)

    def _build_chapter_ten_body(self, v: dict, ref: dict) -> str:
        blocks = []
        prelude = v.get("PROSPECTUS_CH10_PRELUDE", "")
        if prelude:
            blocks.append(prelude)
        section7 = ref.get("\u5341", {}).get("sections", {}).get("\u4e03", "") or v.get("PROSPECTUS_CH10_SEC7", "")
        if section7 and '\u7533\u8d2d\u8d4e\u56de\u6e05\u5355\u7684\u683c\u5f0f\u4e3e\u4f8b\u5982\u4e0b\uff1a' in section7 and 'T\u65e5\u7533\u8d2d\u8d4e\u56de\u6e05\u5355\u7684\u683c\u5f0f\u4e3e\u4f8b\u5982\u4e0b\uff1a' not in section7:
            section7 = section7.replace('\u7533\u8d2d\u8d4e\u56de\u6e05\u5355\u7684\u683c\u5f0f\u4e3e\u4f8b\u5982\u4e0b\uff1a', 'T\u65e5\u7533\u8d2d\u8d4e\u56de\u6e05\u5355\u7684\u683c\u5f0f\u4e3e\u4f8b\u5982\u4e0b\uff1a', 1)
        blocks.extend([
            ref.get("十", {}).get("sections", {}).get("一", ""),
            self._build_contract_section(v, "CONTRACT_PART8_SEC2", "二"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC3", "三"),
            v.get("PROSPECTUS_CH10_SEC4", ""),
            self._build_chapter_ten_sec5(ref, v),
            self._build_contract_section(v, "CONTRACT_PART8_SEC6", "六"),
            section7,
            self._build_contract_section(v, "CONTRACT_PART8_SEC7", "八"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC8", "九"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC9", "十"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC10", "十一"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC11", "十二"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC12", "十三"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC13", "十四"),
            self._build_contract_section(v, "CONTRACT_PART8_SEC14", "十五"),
        ])
        return self._join_nonempty_blocks(blocks)

    def _get_product_type_chapter_builders(self, v: dict, ref: dict) -> dict:
        product_type = self._get_product_type(v)
        if product_type == "ETF":
            return {
                "十": lambda: self._build_chapter_ten_body(v, ref),
            }
        return {}

    def _apply_prospectus_chapter_logic(self, text: str, v: dict) -> str:
        """
        Apply chapter-level composition rules for prospectus generation.
        """
        text = self._apply_reference_fixed_content(text, v)
        ref = self._load_reference_fixed_content(v)

        chapter_builders = {
            "二": lambda: v.get("CONTRACT_DEFS_TEXT", ""),
            "六": lambda: self._build_chapter_six_body(v, ref),
            "七": lambda: v.get("PROSPECTUS_CH7_BODY", ""),
            "九": lambda: v.get("PROSPECTUS_CH9_BODY", ""),
            "十": lambda: self._build_chapter_ten_body(v, ref),
            "十一": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_INVEST_TEXT", "")),
            "十二": lambda: v.get("CONTRACT_ASSET_TEXT", ""),
            "十三": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_VALUATION_TEXT", "")),
            "\u5341\u56db": lambda: self._ensure_distribution_conditions_section(self._normalize_reused_prospectus_chapter(v.get("CONTRACT_DISTRIBUTION_TEXT", "")) or ref.get("\u5341\u56db", {}).get("body", "")),
            "十五": lambda: self._normalize_reused_prospectus_chapter(v.get("CONTRACT_FEE_TEXT", "")),
            "十六": lambda: v.get("CONTRACT_AUDIT_TEXT", ""),
            "十七": lambda: v.get("CONTRACT_DISCLOSURE_TEXT", ""),
            "十八": lambda: self._normalize_prospectus_risk_chapter(v.get("PROSPECTUS_CH18_BODY", ""), v),
            "十九": lambda: v.get("CONTRACT_TERMINATION_TEXT", ""),
            "二十": lambda: v.get("CONTRACT_SUMMARY_TEXT", ""),
            "二十一": lambda: v.get("PROSPECTUS_CH21_TITLES", ""),
        }
        for chap_cn, builder in chapter_builders.items():
            text = self._replace_chapter_body(text, chap_cn, builder())

        sec3 = self._ensure_subsection_heading(v.get("CONTRACT_PART18_SEC3", ""), "三")
        text = self._replace_subsection_in_chapter(text, "十四", "三", sec3)
        return text

    def _process_conditionals(self, text: str, v: dict) -> str:
        return engine._process_conditionals(text, v)

    def _replace_placeholders(self, text: str, v: dict) -> str:
        return engine._replace_placeholders(text, v)

    # ── Step 6: 重排序号（使用"章"作为重置标志）─────────────────────────
    def _renumber_sequences(self, text: str) -> str:
        lines = text.split("\n")
        RE_NUM = re.compile(r"^(\d+)(、)")
        RE_RESET = re.compile(
            r"^(?:[一二三四五六七八九十百]+、|第[一二三四五六七八九十百]+章)"
        )
        last_num = None
        result = []
        for line in lines:
            if RE_RESET.match(line.strip()):
                last_num = None
                result.append(line)
                continue
            m = RE_NUM.match(line)
            if m:
                num = int(m.group(1))
                if num == 1:
                    last_num = 1
                elif last_num is not None and num > last_num + 1:
                    expected = last_num + 1
                    line = re.sub(r"^\d+、", f"{expected}、", line, count=1)
                    last_num = expected
                else:
                    last_num = num
            result.append(line)
        return "\n".join(result)

    # ── Step 7: 清理（检测招募说明书封面而非基金合同）────────────────────
    def _looks_like_cover_title(self, line: str) -> bool:
        s = line.strip()
        if not s:
            return False
        if s.startswith(("#", ">", "-", "*", "**", "`")):
            return False
        if not s.endswith("\u62db\u52df\u8bf4\u660e\u4e66"):
            return False
        forbidden = ("\u6a21\u677f\u8bf4\u660e", "\u6761\u4ef6\u53d8\u91cf\u5f15\u7528\u8bf4\u660e", "\u5dee\u5f02\u6761\u6b3e", ".json", "_CLAUSE", "_DEF")
        return not any(token in s for token in forbidden)

    def _find_body_start_index(self, lines: list[str]) -> int:
        nonempty = [(idx, line.strip()) for idx, line in enumerate(lines) if line.strip()]
        for pos, (idx, stripped) in enumerate(nonempty):
            if not self._looks_like_cover_title(stripped):
                continue
            window = [item[1] for item in nonempty[pos + 1:pos + 6]]
            has_manager = any(item.startswith("\u57fa\u91d1\u7ba1\u7406\u4eba") for item in window)
            has_custodian = any(item.startswith("\u57fa\u91d1\u6258\u7ba1\u4eba") for item in window)
            if has_manager and has_custodian:
                return idx
        for idx, stripped in nonempty:
            if self._looks_like_cover_title(stripped):
                return idx
        return 0

    def _is_internal_metadata_line(self, line: str) -> bool:
        s = line.strip()
        if not s:
            return False
        if s in {"---", "***"}:
            return True
        if s.startswith((">", "# ", "## ")):
            return True
        keywords = (
            "\u6a21\u677f\u8bf4\u660e",
            "\u6761\u4ef6\u53d8\u91cf\u5f15\u7528\u8bf4\u660e",
            "\u5dee\u5f02\u6761\u6b3e\u5f15\u7528\u8bf4\u660e",
            "VALUATION_TIMING_CLAUSE",
            "WORKING_DAY_DEF",
            "BUSINESS_RULES_DEF",
            "NON_COMPONENT_SCOPE",
            "DISTRIBUTION_FREQ_CLAUSE",
            "MGMT_FEE_PAYMENT_METHOD",
            "CUSTODY_FEE_PAYMENT_METHOD",
        )
        if any(keyword in s for keyword in keywords):
            return True
        if ".json" in s and "\u89c1" in s and "`" in s:
            return True
        return False

    def validate_exportable_text(self, text: str) -> dict:
        metadata_matches = []

        def _append_unique(target: list[str], value: str):
            if value not in target:
                target.append(value)

        for raw_line in text.splitlines():
            stripped = raw_line.strip()
            if not stripped:
                continue
            if self._is_internal_metadata_line(stripped):
                _append_unique(metadata_matches, stripped)

        if metadata_matches:
            return {
                "ok": False,
                "error_type": "template_metadata_leaked",
                "error": "招募说明书正文中仍包含模板说明或内部标记，请先清理后再导出。",
                "matches": metadata_matches[:5],
            }
        # Placeholder content is allowed to export; only internal template metadata blocks delivery.
        return {"ok": True, "matches": []}

    def _cleanup(self, text: str) -> str:
        lines = text.split("\n")
        start_idx = self._find_body_start_index(lines)
        clean = []
        for idx, line in enumerate(lines):
            if idx < start_idx:
                continue
            if self._is_internal_metadata_line(line):
                continue
            clean.append(line)
        text = "\n".join(clean)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        return text

    def _inject_important_notice_before_toc(self, text: str, v: dict) -> str:
        ref = self._load_reference_fixed_content(v)
        notice = (ref.get("important_notice") or "").strip()
        if not notice or "\u91cd\u8981\u63d0\u793a" in text:
            return text

        lines = text.splitlines()
        toc_idx = next((idx for idx, line in enumerate(lines) if line.strip() == "\u76ee\u5f55"), None)
        if toc_idx is None:
            return text

        cover = "\n".join(lines[:toc_idx]).strip()
        tail = "\n".join(lines[toc_idx:]).strip()
        return "\n\n".join(part for part in (cover, notice, tail) if part).strip()

    def _get_format_template_prospectus_docx(self) -> Path:
        return REFERENCE_PROSPECTUS_DOCX_MAP["SSE_CROSS"]

    def _build_display_prospectus_text(self, form_data: dict) -> str:
        v = self._derive_variables(form_data)
        v = self._inject_clause_texts(v)
        v = self._extract_contract_sections(v)
        template_text = PROSPECTUS_TEMPLATE_MD.read_text(encoding="utf-8")
        text = self._process_conditionals(template_text, v)
        text = self._replace_placeholders(text, v)
        text = self._apply_prospectus_chapter_logic(text, v)
        text = self._renumber_sequences(text)
        text = self._cleanup(text)
        text = self._format_reference_style_prospectus(text)
        text = self._inject_important_notice_before_toc(text, v)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _build_render_model_from_display_text(self, text: str) -> dict:
        model = {
            "cover_lines": [],
            "important_notice_title": "\u91cd\u8981\u63d0\u793a",
            "important_notice_blocks": [],
            "toc_title": "\u76ee\u5f55",
            "toc_titles": [],
            "toc_entries": [],
            "chapters": [],
        }
        phase = "cover"
        in_notice = False
        current_chapter = None
        top_level_re = re.compile(r"^([^\u3001]+)\u3001(.+)$")

        for raw in text.splitlines():
            stripped = raw.strip()
            if not stripped:
                continue

            if phase == "cover":
                if stripped == "\u91cd\u8981\u63d0\u793a":
                    in_notice = True
                    continue
                if stripped == "\u76ee\u5f55":
                    phase = "toc"
                    in_notice = False
                    continue
                if in_notice:
                    model["important_notice_blocks"].append(stripped)
                else:
                    model["cover_lines"].append(stripped)
                continue

            match = top_level_re.match(stripped)
            if phase == "toc":
                if not match or self._is_prospectus_toc_placeholder_line(stripped):
                    continue
                chapter_cn, title = match.groups()
                if model["toc_titles"] and title == model["toc_titles"][0]:
                    phase = "body"
                    current_chapter = {
                        "chapter_cn": chapter_cn,
                        "title": title,
                        "display_title": stripped,
                        "blocks": [],
                    }
                    model["chapters"].append(current_chapter)
                else:
                    model["toc_titles"].append(title)
                    model["toc_entries"].append({
                        "chapter_cn": chapter_cn,
                        "title": title,
                        "display_title": stripped,
                    })
                continue

            if match and any(match.group(2) == item["title"] for item in model["toc_entries"]):
                chapter_cn, title = match.groups()
                current_chapter = {
                    "chapter_cn": chapter_cn,
                    "title": title,
                    "display_title": stripped,
                    "blocks": [],
                }
                model["chapters"].append(current_chapter)
                continue

            if current_chapter is None:
                continue
            current_chapter["blocks"].append({"type": "paragraph", "text": stripped})

        if not model["toc_entries"]:
            model["toc_entries"] = [
                {
                    "chapter_cn": chapter["chapter_cn"],
                    "title": chapter["title"],
                    "display_title": chapter["display_title"],
                }
                for chapter in model["chapters"]
            ]
        if not model["toc_titles"]:
            model["toc_titles"] = [entry["title"] for entry in model["toc_entries"]]
        return model

    def generate_bundle(self, form_data: dict) -> dict:
        text = self._build_display_prospectus_text(form_data)
        return {
            "text": text,
            "render_model": self._build_render_model_from_display_text(text),
        }

    def generate_render_model(self, form_data: dict) -> dict:
        return self.generate_bundle(form_data)["render_model"]

    def generate(self, form_data: dict) -> str:
        return self.generate_bundle(form_data)["text"]

    def build_docx_prospectus(self, prospectus_text: str) -> bytes:
        prospectus_text = engine._to_chinese_punct(prospectus_text)
        model = self._build_render_model_from_display_text(prospectus_text)

        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        template_docx = self._get_format_template_prospectus_docx()
        doc = Document(str(template_docx)) if template_docx.exists() else Document()

        def clear_body_keep_section(document):
            body = document._element.body
            for child in list(body):
                if child.tag != qn("w:sectPr"):
                    body.remove(child)

        def clear_paragraph_runs(paragraph):
            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)

        def ensure_update_fields(document):
            settings = document.settings.element
            for node in list(settings.findall(qn("w:updateFields"))):
                settings.remove(node)
            update_fields = OxmlElement("w:updateFields")
            update_fields.set(qn("w:val"), "true")
            settings.append(update_fields)

        def set_run_font(run, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21, bold=False):
            run.font.name = ascii_font
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), ascii_font)
            rfonts.set(qn("w:hAnsi"), ascii_font)
            rfonts.set(qn("w:eastAsia"), eastasia)
            run.font.size = Pt(size_half_pt / 2)
            run.bold = bold
            for tag in ("w:sz", "w:szCs"):
                node = rpr.find(qn(tag))
                if node is None:
                    node = OxmlElement(tag)
                    rpr.append(node)
                node.set(qn("w:val"), str(size_half_pt))
            if bold:
                for tag in ("w:b", "w:bCs"):
                    if rpr.find(qn(tag)) is None:
                        rpr.append(OxmlElement(tag))

        def set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False):
            paragraph.alignment = align
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.first_line_indent = Pt(21) if first_line else Pt(0)

        def add_paragraph(text_value, *, style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21, bold=False):
            paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            set_paragraph_format(paragraph, align=align, first_line=first_line)
            if text_value:
                run = paragraph.add_run(text_value)
                set_run_font(run, eastasia=eastasia, ascii_font=ascii_font, size_half_pt=size_half_pt, bold=bold)
            return paragraph

        def add_toc_field():
            field_para = add_paragraph("", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
            begin = OxmlElement("w:r")
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            fld_begin.set(qn("w:dirty"), "true")
            begin.append(fld_begin)
            field_para._p.append(begin)

            instr_run = OxmlElement("w:r")
            instr = OxmlElement("w:instrText")
            instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            instr.text = ' TOC \\o "1-3" \\h \\z \\u '
            instr_run.append(instr)
            field_para._p.append(instr_run)

            sep = OxmlElement("w:r")
            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            sep.append(fld_sep)
            field_para._p.append(sep)

            add_paragraph("\u66f4\u65b0\u76ee\u5f55\u540e\u663e\u793a\u9875\u7801", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)

            end = OxmlElement("w:r")
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            end.append(fld_end)
            field_para._p.append(end)

        def update_header_text(header_text):
            if not doc.sections:
                return
            header = doc.sections[0].header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            clear_paragraph_runs(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ppr = paragraph._p.get_or_add_pPr()
            existing = ppr.find(qn("w:pBdr"))
            if existing is not None:
                ppr.remove(existing)
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:color"), "auto")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            pbdr.append(bottom)
            ppr.append(pbdr)
            run = paragraph.add_run(header_text)
            set_run_font(run, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=18)

        def render_table_block(table_lines):
            rows = []
            header_row = [cell.strip() for cell in table_lines[0].strip().strip('|').split('|')]
            if any(header_row):
                rows.append(header_row)
            for row_text in table_lines[2:]:
                row = [cell.strip() for cell in row_text.strip().strip('|').split('|')]
                if any(row):
                    rows.append(row)
            if not rows:
                return
            col_count = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row_idx, row in enumerate(rows):
                for col_idx in range(col_count):
                    cell = table.cell(row_idx, col_idx)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    clear_paragraph_runs(paragraph)
                    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
                    value = row[col_idx] if col_idx < len(row) else ""
                    if value:
                        run = paragraph.add_run(value)
                        set_run_font(run, eastasia="\u5b8b\u4f53", ascii_font="Times New Roman", size_half_pt=21, bold=(row_idx == 0))

        clear_body_keep_section(doc)
        ensure_update_fields(doc)
        if doc.sections:
            _set_section_page_numbers(doc.sections[0], OxmlElement, qn, start=0)

        header_text = (model.get("cover_lines") or ["\u62db\u52df\u8bf4\u660e\u4e66"])[0]
        if not header_text.endswith("\u62db\u52df\u8bf4\u660e\u4e66"):
            header_text = f"{header_text}\u62db\u52df\u8bf4\u660e\u4e66"
        update_header_text(header_text)

        cover_lines = model.get("cover_lines") or []
        if cover_lines:
            add_paragraph(cover_lines[0], align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, eastasia="\u5b8b\u4f53", ascii_font="\u5b8b\u4f53", size_half_pt=44, bold=True)
            for _ in range(3):
                add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
            for line in cover_lines[1:]:
                add_paragraph(line, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, eastasia="\u5b8b\u4f53", ascii_font="\u5b8b\u4f53", size_half_pt=30, bold=True)

        if model.get("important_notice_blocks"):
            add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
            add_paragraph(model.get("important_notice_title") or "\u91cd\u8981\u63d0\u793a", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, eastasia="\u5b8b\u4f53", ascii_font="\u5b8b\u4f53", size_half_pt=21, bold=True)
            for block in model["important_notice_blocks"]:
                add_paragraph(block, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False)

        add_paragraph("", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        add_paragraph(model.get("toc_title") or "\u76ee\u5f55", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        add_toc_field()

        for chapter in model.get("chapters", []):
            heading = doc.add_paragraph(style="Heading 2")
            set_paragraph_format(heading, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
            run = heading.add_run(chapter.get("title", ""))
            set_run_font(run, eastasia="\u9ed1\u4f53", ascii_font="Arial", size_half_pt=32, bold=True)

            blocks = chapter.get("blocks", [])
            table_buffer = []
            for block in blocks:
                block_text = (block.get("text", "") if isinstance(block, dict) else str(block)).strip()
                if not block_text:
                    continue
                if block_text.startswith("|") and block_text.endswith("|"):
                    table_buffer.append(block_text)
                    continue
                if table_buffer:
                    render_table_block(table_buffer)
                    table_buffer = []
                if re.match(r"^[^\u3001]+\u3001", block_text) or re.match(r"^\uff08[^\uff09]+\uff09", block_text) or re.match(r"^\d+\u3001", block_text) or re.match(r"^\uff08\d+\uff09", block_text):
                    add_paragraph(block_text, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
                else:
                    add_paragraph(block_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True)
            if table_buffer:
                render_table_block(table_buffer)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
ProspectusEngine._is_prospectus_toc_placeholder_line = staticmethod(ContractEngine._is_prospectus_toc_placeholder_line)
ProspectusEngine._parse_prospectus_chapter_heading = staticmethod(ContractEngine._parse_prospectus_chapter_heading)
ProspectusEngine._format_prospectus_reference_heading = staticmethod(ContractEngine._format_prospectus_reference_heading)
ProspectusEngine._format_reference_style_prospectus = ContractEngine._format_reference_style_prospectus

engine = ContractEngine()
prospectus_engine = ProspectusEngine()


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/schema")
def api_schema():
    with open(SCHEMA_JSON, encoding="utf-8") as f:
        return jsonify(json.load(f))


@app.route("/api/custodians")
def api_custodians():
    """返回已知托管人的自动填充数据"""
    with open(CLAUSES_JSON, encoding="utf-8") as f:
        data = json.load(f)
    custodians = data["clauses"]["CUSTODIAN_INFO_KNOWN_VALUES"]["custodians"]
    return jsonify(custodians)


@app.route("/api/diff_table")
def api_diff_table():
    """返回差异条款匹配表（Markdown原文）"""
    content = DIFF_TABLE_MD.read_text(encoding="utf-8")
    return jsonify({"content": content})


@app.route("/api/extract_text", methods=["POST"])
def api_extract_text():
    """从上传的 .docx / .txt / .md 文件提取纯文本，用于合同对比。"""
    if "file" not in request.files:
        return jsonify({"error": "未收到文件"}), 400
    f = request.files["file"]
    name = f.filename or ""
    ext  = name.rsplit(".", 1)[-1].lower() if "." in name else ""

    if ext == "docx":
        try:
            from docx import Document as DocxDocument
            import tempfile, os as _os
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            f.save(tmp.name)
            tmp.close()
            doc = DocxDocument(tmp.name)
            _os.unlink(tmp.name)
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(lines)
        except Exception as e:
            return jsonify({"error": f"无法读取 .docx：{e}"}), 400
    elif ext in ("txt", "md"):
        raw = f.read()
        for enc in ("utf-8", "gbk", "utf-16"):
            try:
                text = raw.decode(enc)
                break
            except Exception:
                pass
        else:
            return jsonify({"error": "文件编码无法识别，请另存为 UTF-8 后再上传"}), 400
    else:
        return jsonify({"error": "仅支持 .docx、.txt、.md 格式"}), 400

    return jsonify({"filename": name, "text": text,
                    "lines": len(text.splitlines()),
                    "size":  len(text)})


@app.route("/api/compare", methods=["POST"])
def api_compare():
    """
    合同对比分析：接收两段合同文本，返回逐行差异（含字符级高亮HTML）。
    响应格式：
      summary: { total_a, total_b, changed_lines, similarity, diff_groups }
      hunks:   每个 hunk 为 { type: "lines"|"skip", lines?, count?, preview? }
               每行: { tag, a, b, a_html, b_html, num_a, num_b }
    """
    import difflib
    import html as html_mod

    d = request.get_json(force=True)
    text1  = d.get("text1",  "")
    text2  = d.get("text2",  "")
    label1 = d.get("label1", "合同A")
    label2 = d.get("label2", "合同B")
    CTX    = 3   # 每个差异块上下各保留 3 行上下文

    lines1 = text1.splitlines()
    lines2 = text2.splitlines()

    if not lines1 or not lines2:
        return jsonify({"error": "两份合同文本均不能为空"}), 400

    # 字符级差异 HTML（对单行 a, b 比较）
    def char_diff(a: str, b: str):
        m = difflib.SequenceMatcher(None, a, b, autojunk=False)
        ra, rb = [], []
        for tag, i1, i2, j1, j2 in m.get_opcodes():
            ea = html_mod.escape(a[i1:i2])
            eb = html_mod.escape(b[j1:j2])
            if tag == "equal":
                ra.append(ea); rb.append(eb)
            elif tag == "replace":
                ra.append(f'<del>{ea}</del>'); rb.append(f'<ins>{eb}</ins>')
            elif tag == "delete":
                ra.append(f'<del>{ea}</del>')
            elif tag == "insert":
                rb.append(f'<ins>{eb}</ins>')
        return "".join(ra), "".join(rb)

    # 逐行对比，构建 result_lines 列表
    sm = difflib.SequenceMatcher(None, lines1, lines2, autojunk=False)
    result_lines = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for k in range(i2 - i1):
                t = html_mod.escape(lines1[i1 + k])
                result_lines.append({"tag": "equal",
                                     "a": t, "b": t,
                                     "a_html": t, "b_html": t,
                                     "num_a": i1+k+1, "num_b": j1+k+1})
        elif tag == "replace":
            len_a, len_b = i2 - i1, j2 - j1
            for k in range(max(len_a, len_b)):
                la = lines1[i1+k] if k < len_a else None
                lb = lines2[j1+k] if k < len_b else None
                na = i1+k+1 if la is not None else None
                nb = j1+k+1 if lb is not None else None
                if la is not None and lb is not None:
                    ah, bh = char_diff(la, lb)
                    result_lines.append({"tag": "replace",
                                         "a": html_mod.escape(la), "b": html_mod.escape(lb),
                                         "a_html": ah, "b_html": bh,
                                         "num_a": na, "num_b": nb})
                elif la is not None:
                    ea = html_mod.escape(la)
                    result_lines.append({"tag": "delete",
                                         "a": ea, "b": "",
                                         "a_html": f"<del>{ea}</del>", "b_html": "",
                                         "num_a": na, "num_b": None})
                else:
                    eb = html_mod.escape(lb)
                    result_lines.append({"tag": "insert",
                                         "a": "", "b": eb,
                                         "a_html": "", "b_html": f"<ins>{eb}</ins>",
                                         "num_a": None, "num_b": nb})
        elif tag == "delete":
            for k in range(i2 - i1):
                ea = html_mod.escape(lines1[i1+k])
                result_lines.append({"tag": "delete",
                                     "a": ea, "b": "",
                                     "a_html": f"<del>{ea}</del>", "b_html": "",
                                     "num_a": i1+k+1, "num_b": None})
        elif tag == "insert":
            for k in range(j2 - j1):
                eb = html_mod.escape(lines2[j1+k])
                result_lines.append({"tag": "insert",
                                     "a": "", "b": eb,
                                     "a_html": "", "b_html": f"<ins>{eb}</ins>",
                                     "num_a": None, "num_b": j1+k+1})

    # 压缩连续的 equal 区域（超过 CTX*2+1 行则折叠中间部分）
    hunks = []
    idx, n = 0, len(result_lines)
    while idx < n:
        r = result_lines[idx]
        if r["tag"] == "equal":
            j = idx
            while j < n and result_lines[j]["tag"] == "equal":
                j += 1
            count = j - idx
            if count > CTX * 2 + 1:
                hunks.append({"type": "lines", "lines": result_lines[idx: idx+CTX]})
                hunks.append({"type": "skip",
                              "count": count - CTX*2,
                              "preview": result_lines[idx+CTX]["a"][:60]})
                hunks.append({"type": "lines", "lines": result_lines[j-CTX: j]})
            else:
                hunks.append({"type": "lines", "lines": result_lines[idx:j]})
            idx = j
        else:
            j = idx
            while j < n and result_lines[j]["tag"] != "equal":
                j += 1
            hunks.append({"type": "lines", "lines": result_lines[idx:j], "is_diff": True})
            idx = j

    # 统计
    changed = sum(1 for r in result_lines if r["tag"] != "equal")
    total   = len(result_lines)
    sim     = round((total - changed) / total * 100, 1) if total > 0 else 100.0
    diff_groups = sum(1 for h in hunks if h.get("is_diff"))

    return jsonify({
        "label1": label1, "label2": label2,
        "summary": {
            "total_a": len(lines1), "total_b": len(lines2),
            "changed_lines": changed, "similarity": sim,
            "diff_groups": diff_groups
        },
        "hunks": hunks
    })


@app.route("/api/summary_check", methods=["POST"])
def api_summary_check():
    """
    合同摘要一致性检验：
    1. 解析第二十六部分摘要的各子项（一、二、…九、）
    2. 按内容相似度自动匹配正文对应章节
    3. 逐项文字比对，返回字符级高亮差异报告
    """
    import difflib
    import html as html_mod

    d = request.get_json(force=True)
    text = d.get("text", "")
    if not text:
        return jsonify({"error": "合同文本不能为空"}), 400

    CTX = 2  # 差异块上下文行数

    # ── 0. 序号剥离（用于比对，不影响显示） ──────────────────────────────────
    _RE_NUM = re.compile(
        r'^(?:'
        r'第[一二三四五六七八九十百千]+条'          # 第X条
        r'|[（(][一二三四五六七八九十百千]+[）)]'   # （一）
        r'|[一二三四五六七八九十百千]+[、．]'        # 一、
        r'|[（(]\d+[）)]'                           # （1）
        r'|\d+[.、．]'                              # 1. 或 1、
        r')\s*'
    )

    def strip_num(line: str) -> str:
        """剥除行首序号，仅供相似度匹配/diff 对齐使用，不修改原文显示。"""
        return _RE_NUM.sub('', line.strip())

    # ── 1. 分割合同各部分 ────────────────────────────────────────────────────
    RE_PART = re.compile(r'^(第[一二三四五六七八九十百]+部分\s+\S[^\n]*)', re.MULTILINE)
    part_iter = list(RE_PART.finditer(text))
    sections = []
    for i, m in enumerate(part_iter):
        heading = m.group(1).split('\t')[0].strip()
        start = m.end()
        end = part_iter[i + 1].start() if i + 1 < len(part_iter) else len(text)
        sections.append({"heading": heading, "content": text[start:end].strip()})

    # 目录条目的"内容"近乎空白（相邻两行之间无实质文本），用长度过滤掉
    body_sections = [s for s in sections
                     if '第二十六部分' not in s['heading'] and len(s['content']) > 100]
    # 取最后一次出现的第二十六部分（跳过目录行，使用正文实体）
    summary_sec = next((s for s in reversed(sections)
                        if '第二十六部分' in s['heading'] and len(s['content']) > 100), None)

    if not summary_sec:
        return jsonify({"error": "未找到第二十六部分（基金合同内容摘要）"}), 400

    # ── 2. 解析摘要子项 ──────────────────────────────────────────────────────
    RE_SUM = re.compile(r'^([一二三四五六七八九十百]+、[^\n]+)', re.MULTILINE)
    sum_iter = list(RE_SUM.finditer(summary_sec['content']))
    subsections = []
    for i, m in enumerate(sum_iter):
        start = m.end()
        end = sum_iter[i + 1].start() if i + 1 < len(sum_iter) else len(summary_sec['content'])
        subsections.append({
            "heading": m.group(1).strip(),
            "content": summary_sec['content'][start:end].strip()
        })

    if not subsections:
        return jsonify({"error": "第二十六部分未找到子项标题（一、二、…）"}), 400

    # ── 3. 字符级差异 HTML ──────────────────────────────────────────────────
    def char_diff(a, b):
        sm = difflib.SequenceMatcher(None, a, b, autojunk=False)
        ra, rb = [], []
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            ea = html_mod.escape(a[i1:i2])
            eb = html_mod.escape(b[j1:j2])
            if tag == "equal":
                ra.append(ea); rb.append(eb)
            elif tag == "replace":
                ra.append(f'<del>{ea}</del>'); rb.append(f'<ins>{eb}</ins>')
            elif tag == "delete":
                ra.append(f'<del>{ea}</del>')
            elif tag == "insert":
                rb.append(f'<ins>{eb}</ins>')
        return "".join(ra), "".join(rb)

    # ── 4. 逐行差异 + 折叠 ─────────────────────────────────────────────────
    def build_diff(text_a, text_b):
        lines_a = text_a.splitlines()
        lines_b = text_b.splitlines()
        # 用剥离序号后的文本做对齐，消除"一、"vs"（一）"等纯序号差异；原文仍用于显示
        norm_a = [strip_num(l) for l in lines_a]
        norm_b = [strip_num(l) for l in lines_b]
        sm = difflib.SequenceMatcher(None, norm_a, norm_b, autojunk=False)
        result_lines = []

        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == "equal":
                for k in range(i2 - i1):
                    t = html_mod.escape(lines_a[i1 + k])
                    result_lines.append({"tag": "equal", "a": t, "b": t,
                                         "a_html": t, "b_html": t,
                                         "num_a": i1 + k + 1, "num_b": j1 + k + 1})
            elif tag == "replace":
                la, lb = i2 - i1, j2 - j1
                for k in range(max(la, lb)):
                    al = lines_a[i1 + k] if k < la else None
                    bl = lines_b[j1 + k] if k < lb else None
                    if al is not None and bl is not None:
                        ah, bh = char_diff(al, bl)
                        result_lines.append({"tag": "replace",
                                             "a": html_mod.escape(al), "b": html_mod.escape(bl),
                                             "a_html": ah, "b_html": bh,
                                             "num_a": i1 + k + 1, "num_b": j1 + k + 1})
                    elif al is not None:
                        ea = html_mod.escape(al)
                        result_lines.append({"tag": "delete", "a": ea, "b": "",
                                             "a_html": f'<del>{ea}</del>', "b_html": "",
                                             "num_a": i1 + k + 1, "num_b": None})
                    else:
                        eb = html_mod.escape(bl)
                        result_lines.append({"tag": "insert", "a": "", "b": eb,
                                             "a_html": "", "b_html": f'<ins>{eb}</ins>',
                                             "num_a": None, "num_b": j1 + k + 1})
            elif tag == "delete":
                for k in range(i2 - i1):
                    ea = html_mod.escape(lines_a[i1 + k])
                    result_lines.append({"tag": "delete", "a": ea, "b": "",
                                         "a_html": f'<del>{ea}</del>', "b_html": "",
                                         "num_a": i1 + k + 1, "num_b": None})
            elif tag == "insert":
                for k in range(j2 - j1):
                    eb = html_mod.escape(lines_b[j1 + k])
                    result_lines.append({"tag": "insert", "a": "", "b": eb,
                                         "a_html": "", "b_html": f'<ins>{eb}</ins>',
                                         "num_a": None, "num_b": j1 + k + 1})

        # 折叠连续 equal 区
        hunks = []
        idx, n = 0, len(result_lines)
        while idx < n:
            r = result_lines[idx]
            if r["tag"] == "equal":
                j = idx
                while j < n and result_lines[j]["tag"] == "equal":
                    j += 1
                count = j - idx
                if count > CTX * 2 + 1:
                    hunks.append({"type": "lines", "lines": result_lines[idx: idx + CTX]})
                    hunks.append({"type": "skip", "count": count - CTX * 2,
                                  "preview": result_lines[idx + CTX]["a"][:60]})
                    hunks.append({"type": "lines", "lines": result_lines[j - CTX: j]})
                else:
                    hunks.append({"type": "lines", "lines": result_lines[idx:j]})
                idx = j
            else:
                j = idx
                while j < n and result_lines[j]["tag"] != "equal":
                    j += 1
                hunks.append({"type": "lines", "lines": result_lines[idx:j], "is_diff": True})
                idx = j

        changed = sum(1 for r in result_lines if r["tag"] != "equal")
        total = len(result_lines)
        sim = round((total - changed) / total * 100, 1) if total > 0 else 100.0
        diff_groups = sum(1 for h in hunks if h.get("is_diff"))
        return hunks, changed, total, sim, diff_groups

    # ── 5. 按内容相似度匹配正文章节 ─────────────────────────────────────────
    def content_sim(summary_content, section_content):
        """
        计算摘要在正文章节中的召回率：摘要中有多少比例的字符能在章节中找到。
        比双向 ratio() 更适合：摘要是正文的精简版，长度远短于原章节。
        序号已剥离，避免"一、"vs"（一）"干扰匹配得分。
        """
        norm_s   = '\n'.join(strip_num(l) for l in summary_content.splitlines())
        norm_sec = '\n'.join(strip_num(l) for l in section_content.splitlines())
        if not norm_s:
            return 0.0
        sm = difflib.SequenceMatcher(None, norm_s, norm_sec, autojunk=False)
        matched = sum(blk.size for blk in sm.get_matching_blocks())
        return matched / len(norm_s)

    results = []
    for sub in subsections:
        best_score = -1.0
        best_sec = None
        for sec in body_sections:
            score = content_sim(sub['content'], sec['content'])
            if score > best_score:
                best_score = score
                best_sec = sec

        if best_sec is None:
            results.append({
                "summary_heading": sub['heading'],
                "matched_section": None,
                "content_match_score": 0,
                "similarity": 0,
                "changed_lines": 0,
                "total_lines": 0,
                "diff_groups": 0,
                "hunks": []
            })
            continue

        hunks, changed, total, sim, diff_groups = build_diff(best_sec['content'], sub['content'])
        results.append({
            "summary_heading": sub['heading'],
            "matched_section": best_sec['heading'],
            "content_match_score": round(best_score * 100, 1),
            "similarity": sim,
            "changed_lines": changed,
            "total_lines": total,
            "diff_groups": diff_groups,
            "hunks": hunks
        })

    return jsonify({"results": results, "total_subsections": len(results)})


@app.route("/api/generate", methods=["POST"])
def api_generate():
    """接收表单数据，返回合同全文"""
    form_data = request.get_json(force=True)
    try:
        contract_text = engine.generate(form_data)
        return jsonify({"success": True, "text": contract_text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export", methods=["POST"])
def api_export():
    """返回 .txt 文件下载"""
    form_data = request.get_json(force=True)
    try:
        contract_text = engine.generate(form_data)
        fund_name = form_data.get("FUND_NAME", "ETF基金合同")
        # Write to a temp file
        import tempfile
        tmp = tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", encoding="utf-8",
            delete=False, prefix="contract_"
        )
        tmp.write(contract_text)
        tmp.close()
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}基金合同.txt",
            mimetype="text/plain; charset=utf-8",
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/export_docx", methods=["POST"])
def api_export_docx():
    """返回格式化 .docx 文件下载"""
    form_data = request.get_json(force=True)
    try:
        contract_text = engine.generate(form_data)
        docx_bytes = engine.build_docx(contract_text)
        fund_name = form_data.get("FUND_NAME", "ETF基金合同")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile, os as _os
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="contract_")
        tmp.write(docx_bytes)
        tmp.close()
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}基金合同.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e), "trace": traceback.format_exc()}), 400


@app.route("/api/files")
def api_files():
    """列出所有可编辑文件"""
    files = []
    for p in sorted(BASE_DIR.iterdir()):
        if p.suffix in ALLOWED_SUFFIXES and p.name != "app.py":
            stat = p.stat()
            files.append({
                "name": p.name,
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
            })
    return jsonify(files)


@app.route("/api/files/<filename>", methods=["GET", "POST"])
def api_file(filename):
    """读取或保存指定文件（防路径遍历）"""
    # Security: only allow .md and .json files in BASE_DIR
    target = (BASE_DIR / filename).resolve()
    if BASE_DIR.resolve() not in target.parents or target.suffix not in ALLOWED_SUFFIXES:
        abort(403)
    if not target.is_file() or not target.name == filename:
        abort(404)

    if request.method == "GET":
        content = target.read_text(encoding="utf-8")
        return jsonify({"name": filename, "content": content})

    # POST: save
    data = request.get_json(force=True)
    content = data.get("content", "")
    target.write_text(content, encoding="utf-8")
    # Reload engine after template changes
    global engine
    engine = ContractEngine()
    return jsonify({"success": True, "message": f"已保存 {filename}"})


@app.route("/api/preview_clause", methods=["POST"])
def api_preview_clause():
    """根据当前表单值，实时返回关键差异条款预览文本"""
    form_data = request.get_json(force=True)
    try:
        v = engine._derive_variables(form_data)
        v = engine._inject_clause_texts(v)
        preview = {
            "WORKING_DAY_DEF": v.get("WORKING_DAY_DEF", ""),
            "BUSINESS_RULES_DEF": v.get("BUSINESS_RULES_DEF", ""),
            "NON_COMPONENT_SCOPE": v.get("NON_COMPONENT_SCOPE", ""),
            "DISTRIBUTION_FREQ_CLAUSE": v.get("DISTRIBUTION_FREQ_CLAUSE", ""),
            "MGMT_FEE_PAYMENT_METHOD": v.get("MGMT_FEE_PAYMENT_METHOD", ""),
            "TRACKING_ERROR_DAILY": v.get("TRACKING_ERROR_DAILY", 0.2),
            "TRACKING_ERROR_ANNUAL": v.get("TRACKING_ERROR_ANNUAL", 2),
        }
        return jsonify(preview)
    except Exception as e:
        return jsonify({"error": str(e)}), 400


# ── 招募说明书路由 ────────────────────────────────────────────────────────────


@app.route("/api/generate_prospectus", methods=["POST"])
def api_generate_prospectus():
    """接收表单数据，生成招募说明书全文，返回 JSON"""
    form_data = request.get_json(force=True)
    try:
        bundle = prospectus_engine.generate_bundle(form_data)
        return jsonify({"success": True, **bundle})
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e), "trace": traceback.format_exc()}), 400


@app.route("/api/export_prospectus_docx", methods=["POST"])
def api_export_prospectus_docx():
    """\u8fd4\u56de\u683c\u5f0f\u5316\u62db\u52df\u8bf4\u660e\u4e66 .docx \u6587\u4ef6\u4e0b\u8f7d"""
    form_data = request.get_json(force=True)
    try:
        prospectus_text = prospectus_engine.generate(form_data)
        report = prospectus_engine.validate_exportable_text(prospectus_text)
        if not report["ok"]:
            report = {k: v for k, v in report.items() if k != "ok"}
            return jsonify({"success": False, **report}), 400
        docx_bytes = prospectus_engine.build_docx_prospectus(prospectus_text)
        fund_name = form_data.get("FUND_NAME", "ETF\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="prospectus_")
        tmp.write(docx_bytes)
        tmp.close()
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}\u62db\u52df\u8bf4\u660e\u4e66.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e), "trace": traceback.format_exc()}), 400


@app.route("/api/export_prospectus_txt", methods=["POST"])
def api_export_prospectus_txt():
    """\u8fd4\u56de\u62db\u52df\u8bf4\u660e\u4e66 .txt \u6587\u4ef6\u4e0b\u8f7d"""
    form_data = request.get_json(force=True)
    try:
        prospectus_text = prospectus_engine.generate(form_data)
        report = prospectus_engine.validate_exportable_text(prospectus_text)
        if not report["ok"]:
            report = {k: v for k, v in report.items() if k != "ok"}
            return jsonify({"success": False, **report}), 400
        fund_name = form_data.get("FUND_NAME", "ETF\u57fa\u91d1\u62db\u52df\u8bf4\u660e\u4e66")
        import tempfile
        tmp = tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", encoding="utf-8",
            delete=False, prefix="prospectus_"
        )
        tmp.write(prospectus_text)
        tmp.close()
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", fund_name)
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name=f"{safe_name}\u62db\u52df\u8bf4\u660e\u4e66.txt",
            mimetype="text/plain; charset=utf-8",
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


# \u2500\u2500 \u542f\u52a8
def open_browser():
    time.sleep(1.2)
    webbrowser.open("http://127.0.0.1:5000")


if __name__ == "__main__":
    t = threading.Thread(target=open_browser, daemon=True)
    t.start()
    app.run(debug=False, port=5000, use_reloader=False)
